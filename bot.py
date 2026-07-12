# =====================================================
# BOT DICOR FINAL RESOLVIDO - MESAS + PROCURADOS + CATALOGO HTML
# Feito para GTA RP / personagens ficticios
# =====================================================

import os
import re
import json
import html
import base64
import hmac
import secrets
import string
import time
import asyncio
import datetime
import unicodedata
from pathlib import Path
from typing import Optional, Dict, Any, List

import discord
from discord import app_commands
from discord.ext import commands, tasks
from discord.ui import View, Button, Modal, TextInput
from dotenv import load_dotenv
from aiohttp import web, ClientSession, ClientTimeout

# Dependências do Dossiê Operacional Automático DICOR
# Railway/Render: adicione no requirements.txt:
# reportlab
# python-docx
# qrcode
# pillow
import io
import math
from collections import defaultdict

try:
    import qrcode
    from PIL import Image as PILImage
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm, mm
    from reportlab.pdfgen import canvas
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, Image as RLImage, KeepTogether
    )
except Exception:
    qrcode = None
    PILImage = None
    colors = None
    TA_CENTER = TA_JUSTIFY = TA_LEFT = TA_RIGHT = 0
    A4 = None
    cm = mm = 1
    canvas = None
    SimpleDocTemplate = Paragraph = Spacer = Table = TableStyle = PageBreak = RLImage = KeepTogether = None
    getSampleStyleSheet = ParagraphStyle = None

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except Exception:
    Document = None
    WD_ALIGN_PARAGRAPH = WD_TABLE_ALIGNMENT = WD_CELL_VERTICAL_ALIGNMENT = None
    OxmlElement = qn = Inches = Pt = RGBColor = None

# =====================================================
# CONFIGURACAO
# =====================================================

load_dotenv()


def env_int(nome: str, padrao: int = 0) -> int:
    """Lê ID/contador do Railway sem crashar se ficar placeholder tipo ID_DO_CARGO_INSPETOR."""
    valor = str(os.getenv(nome, str(padrao)) or "").strip()
    try:
        return int(valor)
    except (TypeError, ValueError):
        print(f"⚠️ Variável {nome} inválida: {valor!r}. Usando {padrao}.")
        return int(padrao)


def env_float(nome: str, padrao: float = 0.0) -> float:
    valor = str(os.getenv(nome, str(padrao)) or "").strip().replace(",", ".")
    try:
        return float(valor)
    except (TypeError, ValueError):
        print(f"⚠️ Variável {nome} inválida: {valor!r}. Usando {padrao}.")
        return float(padrao)


DISCORD_TOKEN = os.getenv("DISCORD_TOKEN", "").strip()
GUILD_ID = env_int("GUILD_ID", 0)

# Procurados
PROCURADOS_CHANNEL_ID = env_int("PROCURADOS_CHANNEL_ID", 0)
HISTORICO_PROCURADOS_ID = env_int("HISTORICO_PROCURADOS_ID", 1490200536207855857)
LOGS_CHANNEL_ID = env_int("LOGS_CHANNEL_ID", 1490205503228477610)
PROCURADOS_TEMP_CATEGORY_ID = env_int("PROCURADOS_TEMP_CATEGORY_ID", 0)

# Boletins
BOLETINS_CHANNEL_ID = env_int("BOLETINS_CHANNEL_ID", 1490200514837745754)
BOLETIM_TEMP_CATEGORY_ID = env_int("BOLETIM_TEMP_CATEGORY_ID", 0)

# Atendimento automático de boletins
BOLETIM_ATENDIMENTO_CHANNEL_ID = env_int("BOLETIM_ATENDIMENTO_CHANNEL_ID", 1525762770253910136)
BOLETINS_ARQUIVADOS_CHANNEL_ID = env_int("BOLETINS_ARQUIVADOS_CHANNEL_ID", 1525762226269720696)
# Rodízio do atendimento: somente Estagiário e Investigador.
# Mesmo que a variável do Railway esteja errada ou tenha cargos da diretoria,
# o bot filtra e mantém apenas estes dois cargos.
BOLETIM_RODIZIO_CARGOS_FIXOS_IDS = {1490200391239864352, 1490200390426165290}
_boletim_rodizio_env_ids = [
    int(x)
    for x in os.getenv(
        "BOLETIM_RODIZIO_CARGOS_IDS",
        "1490200391239864352,1490200390426165290",
    ).replace(";", ",").split(",")
    if x.strip().isdigit()
]
BOLETIM_RODIZIO_CARGOS_IDS = [
    cid for cid in _boletim_rodizio_env_ids
    if cid in BOLETIM_RODIZIO_CARGOS_FIXOS_IDS
] or sorted(BOLETIM_RODIZIO_CARGOS_FIXOS_IDS)
# Cargos superiores que NÃO devem entrar no rodízio, mesmo se também tiverem cargo de Investigador/Estagiário.
BOLETIM_RODIZIO_EXCLUIR_CARGOS_IDS = {
    int(x)
    for x in os.getenv(
        "BOLETIM_RODIZIO_EXCLUIR_CARGOS_IDS",
        "1490200388912156692,1490200383614615725,1490200382776021132,1490200384818647051",
    ).replace(";", ",").split(",")
    if x.strip().isdigit()
}

# Mesas
CATEGORIA_MESAS_ABERTAS_ID = env_int("CATEGORIA_MESAS_ABERTAS_ID", 1490200456855552192)
CATEGORIA_MESAS_FECHADAS_ID = env_int("CATEGORIA_MESAS_FECHADAS_ID", 1515165416815722586)
BACKUP_CHANNEL_ID = env_int("BACKUP_CHANNEL_ID", 1515165673276440677)

# Dossiê Operacional Automático DICOR
# Coloque no .env/Railway para escolher o canal que receberá PDF + DOCX:
# DOSSIE_CHANNEL_ID=ID_DO_CANAL_DOS_DOSSIES
DOSSIE_CHANNEL_ID = env_int("DOSSIE_CHANNEL_ID", BACKUP_CHANNEL_ID)
DOSSIE_HISTORY_LIMIT = env_int("DOSSIE_HISTORY_LIMIT", 0)  # 0 = varrer todo o histórico disponível
DOSSIE_ENVIAR_NA_MESA = os.getenv("DOSSIE_ENVIAR_NA_MESA", "1").strip().lower() not in {"0", "false", "nao", "não", "off"}
DOSSIE_SKIP_BOT_BOILERPLATE = os.getenv("DOSSIE_SKIP_BOT_BOILERPLATE", "1").strip().lower() not in {"0", "false", "nao", "não", "off"}
DOSSIE_PROGRESS_INTERVAL = env_float("DOSSIE_PROGRESS_INTERVAL", 2)

# Controle de quem pode fechar mesas
# Opção 1: coloque o ID do cargo Inspetor. Quem tiver esse cargo ou cargo acima poderá fechar.
# DOSSIE_CARGO_MINIMO_FECHAR_ID=ID_DO_CARGO_INSPETOR
# Opção 2: coloque uma lista de cargos autorizados, separados por vírgula.
# CARGOS_FECHAR_MESA_IDS=ID_INSPETOR,ID_DELEGADO_DICOR,ID_DELEGADO_GERAL
DOSSIE_CARGO_MINIMO_FECHAR_ID = env_int("DOSSIE_CARGO_MINIMO_FECHAR_ID", 0)

# Assinaturas do Dossiê
ASSINATURA_DELEGADO_GERAL_NOME = os.getenv("ASSINATURA_DELEGADO_GERAL_NOME", "Delegado Geral").strip() or "Delegado Geral"
ASSINATURA_DELEGADO_DICOR_NOME = os.getenv("ASSINATURA_DELEGADO_DICOR_NOME", "Delegado DICOR").strip() or "Delegado DICOR"
ASSINATURA_DELEGADO_GERAL_IMAGEM = os.getenv("ASSINATURA_DELEGADO_GERAL_IMAGEM", "").strip()
ASSINATURA_DELEGADO_DICOR_IMAGEM = os.getenv("ASSINATURA_DELEGADO_DICOR_IMAGEM", "").strip()
ASSINATURA_AGENTE_RESPONSAVEL_IMAGEM = os.getenv("ASSINATURA_AGENTE_RESPONSAVEL_IMAGEM", "").strip()
ASSINATURA_DELEGADO_GERAL_TEXTO = os.getenv("ASSINATURA_DELEGADO_GERAL_TEXTO", "").strip()
ASSINATURA_DELEGADO_DICOR_TEXTO = os.getenv("ASSINATURA_DELEGADO_DICOR_TEXTO", "").strip()
ASSINATURA_AGENTE_RESPONSAVEL_TEXTO = os.getenv("ASSINATURA_AGENTE_RESPONSAVEL_TEXTO", "").strip()

# Catalogo
CATALOG_PUBLIC_URL = os.getenv("CATALOG_PUBLIC_URL", "http://127.0.0.1:8000/").strip()
CATALOG_ADMIN_PASSWORD = os.getenv("CATALOG_ADMIN_PASSWORD", "").strip()
PORT = env_int("PORT", 8000)


def _ids_env(nome: str) -> List[int]:
    ids: List[int] = []
    for parte in os.getenv(nome, "").split(","):
        parte = parte.strip()
        if parte.isdigit():
            ids.append(int(parte))
    return ids

CARGOS_ADMIN_IDS = _ids_env("CARGOS_ADMIN_IDS")
CARGOS_EQUIPE_IDS = _ids_env("CARGOS_EQUIPE_IDS") or CARGOS_ADMIN_IDS
CARGOS_FECHAR_MESA_IDS = _ids_env("CARGOS_FECHAR_MESA_IDS")

BASE_DIR = Path(__file__).parent
DATA_DIR = BASE_DIR / "data"
PUBLIC_DIR = BASE_DIR / "public"
UPLOADS_DIR = PUBLIC_DIR / "uploads"
BACKUP_DIR = BASE_DIR / "backups"

CATALOGO_JSON = DATA_DIR / "procurados.json"
CATALOGO_HTML = PUBLIC_DIR / "index.html"
MESAS_JSON = DATA_DIR / "mesas.json"
BOLETINS_JSON = DATA_DIR / "boletins.json"
BOLETINS_PENDENTES_JSON = DATA_DIR / "boletins_pendentes.json"
USUARIOS_RG_JSON = DATA_DIR / "usuarios_rg.json"
BOLETINS_CONTADOR_JSON = DATA_DIR / "boletins_contador.json"
CATALOG_ADMIN_JSON = DATA_DIR / "catalog_admin.json"
BOLETIM_ATENDIMENTOS_JSON = DATA_DIR / "boletins_atendimentos.json"
BOLETIM_RODIZIO_JSON = DATA_DIR / "boletins_rodizio.json"
BOLETIM_ARQUIVOS_DIR = DATA_DIR / "boletins_arquivos"
ORGANIZACOES_JSON = DATA_DIR / "organizacoes.json"
HISTORICO_ORGANIZACOES_JSON = DATA_DIR / "historico_organizacoes.json"
DOSSIES_JSON = DATA_DIR / "dossies_operacionais.json"
DOSSIES_DIR = DATA_DIR / "dossies"
DOSSIE_ASSETS_DIR = DATA_DIR / "assets_dossie"
ASSINATURAS_DOSSIE_JSON = DATA_DIR / "assinaturas_dossie.json"
ASSINATURAS_DOSSIE_DIR = DOSSIE_ASSETS_DIR / "assinaturas"
RELATORIOS_OPERACIONAIS_JSON = DATA_DIR / "relatorios_operacionais.json"
ADMIN_REPORTS_DIR = DATA_DIR / "relatorios_administrativos"

for pasta in [DATA_DIR, PUBLIC_DIR, UPLOADS_DIR, BACKUP_DIR, DOSSIES_DIR, DOSSIE_ASSETS_DIR, ASSINATURAS_DOSSIE_DIR, ADMIN_REPORTS_DIR, globals().get("BOLETIM_ARQUIVOS_DIR", DATA_DIR / "boletins_arquivos")]:
    pasta.mkdir(exist_ok=True)

# =====================================================
# BOT
# =====================================================

intents = discord.Intents.default()
intents.guilds = True
intents.members = True
intents.messages = True
intents.message_content = True

bot = commands.Bot(command_prefix="!", intents=intents)

cadastros_pendentes: Dict[int, Dict[str, Any]] = {}
boletins_pendentes: Dict[int, Dict[str, Any]] = {}
catalogo_tentativas_senha: Dict[str, List[float]] = {}

# =====================================================
# FUNCOES GERAIS
# =====================================================

def agora_br() -> str:
    tz = datetime.timezone(datetime.timedelta(hours=-3))
    return datetime.datetime.now(tz).strftime("%d/%m/%Y %H:%M")


def data_caso() -> str:
    tz = datetime.timezone(datetime.timedelta(hours=-3))
    return datetime.datetime.now(tz).strftime("%Y%m%d-%H%M%S")


def slugify(texto: str) -> str:
    texto = str(texto or "").strip().lower()
    texto = unicodedata.normalize("NFKD", texto).encode("ascii", "ignore").decode("ascii")
    texto = re.sub(r"[^a-z0-9]+", "-", texto)
    texto = texto.strip("-")
    return texto[:70] or "item"


def limpar_rg(rg: str) -> str:
    return re.sub(r"\s+", "", str(rg or "").lower())


def escape(texto: Any) -> str:
    return html.escape(str(texto or ""))


def cortar_discord(texto: str, limite: int = 1900) -> str:
    texto = str(texto or "")
    if len(texto) <= limite:
        return texto
    return texto[:limite - 40].rstrip() + "\n\n... texto cortado pelo limite do Discord."


def carregar_json(caminho: Path, padrao):
    if not caminho.exists():
        return padrao
    try:
        with open(caminho, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return padrao


def salvar_json(caminho: Path, dados) -> None:
    with open(caminho, "w", encoding="utf-8") as f:
        json.dump(dados, f, ensure_ascii=False, indent=2)


def carregar_assinaturas_dossie() -> Dict[str, Any]:
    dados = carregar_json(ASSINATURAS_DOSSIE_JSON, {})
    return dados if isinstance(dados, dict) else {}


def salvar_assinaturas_dossie(dados: Dict[str, Any]) -> None:
    salvar_json(ASSINATURAS_DOSSIE_JSON, dados)


def caminho_relativo_base(caminho: Path) -> str:
    try:
        return str(caminho.resolve().relative_to(BASE_DIR.resolve())).replace("\\", "/")
    except Exception:
        return str(caminho)


def caminho_assinatura_registrada(registro: Any) -> Optional[Path]:
    if not isinstance(registro, dict):
        return None
    valor = str(registro.get("arquivo", "") or "").strip()
    if not valor:
        return None
    caminho = Path(valor)
    if not caminho.is_absolute():
        caminho = BASE_DIR / caminho
    return caminho if caminho.exists() else None


async def salvar_imagem_assinatura(arquivo: discord.Attachment, destino: Path) -> Path:
    extensao = Path(arquivo.filename or "").suffix.lower()
    if extensao not in [".png", ".jpg", ".jpeg", ".webp"]:
        extensao = ".png"
    destino = destino.with_suffix(extensao)
    await arquivo.save(str(destino))
    return destino


# -----------------------------------------------------
# Tratamento visual das imagens do dossiê
# Remove fundo preto/branco de borda, corta sobra vazia e deixa brasões/assinaturas maiores e limpos.
# -----------------------------------------------------
def _is_dark_pixel(px) -> bool:
    r, g, b, a = px
    return a > 0 and r < 38 and g < 38 and b < 38


def _flood_transparent_from_edges(img, predicate) -> None:
    """Deixa transparente apenas o fundo conectado às bordas, preservando detalhes internos."""
    if PILImage is None:
        return
    w, h = img.size
    pix = img.load()
    visitados = set()
    pilha = []

    for x in range(w):
        pilha.append((x, 0))
        pilha.append((x, h - 1))
    for y in range(h):
        pilha.append((0, y))
        pilha.append((w - 1, y))

    while pilha:
        x, y = pilha.pop()
        if x < 0 or y < 0 or x >= w or y >= h or (x, y) in visitados:
            continue
        visitados.add((x, y))
        r, g, b, a = pix[x, y]
        if not predicate((r, g, b, a)):
            continue
        pix[x, y] = (r, g, b, 0)
        pilha.append((x + 1, y))
        pilha.append((x - 1, y))
        pilha.append((x, y + 1))
        pilha.append((x, y - 1))


def limpar_imagem_brasao_dossie(caminho: Optional[Path]) -> Optional[Path]:
    """Remove o quadrado preto externo dos brasões e gera PNG com transparência."""
    if not caminho or not Path(caminho).exists() or PILImage is None:
        return caminho
    try:
        caminho = Path(caminho)
        cache = DOSSIE_ASSETS_DIR / f"limpo_{caminho.stem}_{int(caminho.stat().st_mtime)}.png"
        if cache.exists():
            return cache

        img = PILImage.open(caminho).convert("RGBA")
        _flood_transparent_from_edges(img, _is_dark_pixel)

        # Corta a sobra transparente ao redor, mantendo pequena margem para não encostar no texto.
        alpha = img.getchannel("A")
        bbox = alpha.getbbox()
        if bbox:
            margem = 18
            left = max(0, bbox[0] - margem)
            top = max(0, bbox[1] - margem)
            right = min(img.size[0], bbox[2] + margem)
            bottom = min(img.size[1], bbox[3] + margem)
            img = img.crop((left, top, right, bottom))

        img.save(cache, "PNG")
        return cache
    except Exception:
        return caminho


def limpar_imagem_assinatura_dossie(caminho: Optional[Path]) -> Optional[Path]:
    """Corta espaços brancos da assinatura e remove fundo branco quando houver."""
    if not caminho or not Path(caminho).exists() or PILImage is None:
        return caminho
    try:
        caminho = Path(caminho)
        cache = DOSSIE_ASSETS_DIR / "assinaturas" / f"limpa_{caminho.stem}_{int(caminho.stat().st_mtime)}.png"
        if cache.exists():
            return cache

        img = PILImage.open(caminho).convert("RGBA")
        pix = img.load()
        w, h = img.size

        # Remove fundo branco/claro em imagens de assinatura escaneada/printada.
        for y in range(h):
            for x in range(w):
                r, g, b, a = pix[x, y]
                if a == 0:
                    continue
                if r > 232 and g > 232 and b > 232:
                    pix[x, y] = (r, g, b, 0)

        alpha = img.getchannel("A")
        bbox = alpha.getbbox()
        if bbox:
            margem = 12
            left = max(0, bbox[0] - margem)
            top = max(0, bbox[1] - margem)
            right = min(img.size[0], bbox[2] + margem)
            bottom = min(img.size[1], bbox[3] + margem)
            img = img.crop((left, top, right, bottom))

        img.save(cache, "PNG")
        return cache
    except Exception:
        return caminho


def garantir_senha_catalogo() -> str:
    """Carrega a senha do Railway ou gera uma senha local persistente."""
    global CATALOG_ADMIN_PASSWORD

    if CATALOG_ADMIN_PASSWORD:
        return CATALOG_ADMIN_PASSWORD

    configuracao = carregar_json(CATALOG_ADMIN_JSON, {})
    if isinstance(configuracao, dict):
        senha_salva = str(configuracao.get("senha", "") or "").strip()
        if senha_salva:
            CATALOG_ADMIN_PASSWORD = senha_salva
            return CATALOG_ADMIN_PASSWORD

    alfabeto = string.ascii_letters + string.digits
    CATALOG_ADMIN_PASSWORD = "DICOR-" + "".join(
        secrets.choice(alfabeto) for _ in range(14)
    )
    salvar_json(CATALOG_ADMIN_JSON, {"senha": CATALOG_ADMIN_PASSWORD})
    print("ATENÇÃO: CATALOG_ADMIN_PASSWORD não foi configurada no Railway.")
    print(f"SENHA TEMPORÁRIA DO CATÁLOGO: {CATALOG_ADMIN_PASSWORD}")
    print(
        "Cadastre essa senha na variável CATALOG_ADMIN_PASSWORD "
        "para ela não mudar em novos deploys."
    )
    return CATALOG_ADMIN_PASSWORD


def senha_catalogo_valida(senha: str) -> bool:
    senha_ativa = garantir_senha_catalogo()
    return bool(senha_ativa) and hmac.compare_digest(
        str(senha or ""),
        senha_ativa,
    )


def catalogo_limite_tentativas(ip: str) -> bool:
    """Permite até 8 tentativas erradas por IP a cada 10 minutos."""
    agora = time.monotonic()
    janela = 600.0
    tentativas = [
        instante
        for instante in catalogo_tentativas_senha.get(ip, [])
        if agora - instante < janela
    ]
    catalogo_tentativas_senha[ip] = tentativas
    return len(tentativas) < 8


def registrar_falha_senha_catalogo(ip: str) -> None:
    catalogo_tentativas_senha.setdefault(ip, []).append(time.monotonic())


def carregar_procurados() -> List[Dict[str, Any]]:
    dados = carregar_json(CATALOGO_JSON, [])
    return dados if isinstance(dados, list) else []


def salvar_procurados(lista: List[Dict[str, Any]]) -> None:
    salvar_json(CATALOGO_JSON, lista)


def carregar_mesas() -> List[Dict[str, Any]]:
    dados = carregar_json(MESAS_JSON, [])
    return dados if isinstance(dados, list) else []


def salvar_mesas(lista: List[Dict[str, Any]]) -> None:
    salvar_json(MESAS_JSON, lista)


def procurar_por_rg(rg: str) -> Optional[Dict[str, Any]]:
    alvo = limpar_rg(rg)
    for p in carregar_procurados():
        if limpar_rg(p.get("rg", "")) == alvo:
            return p
    return None


def remover_duplicados(lista: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    vistos = set()
    saida = []
    for item in lista:
        rg = limpar_rg(item.get("rg", ""))
        chave = rg or item.get("id") or item.get("caso")
        if chave in vistos:
            continue
        vistos.add(chave)
        saida.append(item)
    return saida


async def enviar_log(texto: str) -> None:
    """Envia logs operacionais no canal fixo da DICOR.

    Canal padrão: 1490205503228477610.
    Usa fetch_channel como fallback para não depender apenas do cache do bot.
    """
    if not LOGS_CHANNEL_ID:
        return

    try:
        canal = bot.get_channel(LOGS_CHANNEL_ID)
        if canal is None:
            canal = await bot.fetch_channel(LOGS_CHANNEL_ID)

        if not canal or not hasattr(canal, "send"):
            return

        texto = str(texto or "")
        if not texto.strip():
            texto = "Ação registrada sem descrição."

        for parte in [texto[i:i + 1900] for i in range(0, len(texto), 1900)]:
            await canal.send(parte)
    except Exception:
        pass


def usuario_tem_admin(member: discord.Member) -> bool:
    if not CARGOS_ADMIN_IDS:
        return True
    cargos = {role.id for role in member.roles}
    return any(cargo in cargos for cargo in CARGOS_ADMIN_IDS)


def usuario_pode_fechar_mesa(member: discord.Member) -> bool:
    """Permite fechar mesa apenas para ADM DICOR / Inspetor para cima."""
    if not isinstance(member, discord.Member):
        return False

    if member.guild_permissions.administrator:
        return True

    cargos_usuario = {role.id for role in member.roles}

    # Lista direta de cargos autorizados.
    cargos_autorizados = set(CARGOS_FECHAR_MESA_IDS or CARGOS_ADMIN_IDS)
    if cargos_autorizados and cargos_usuario.intersection(cargos_autorizados):
        return True

    # Cargo mínimo: quem estiver no cargo informado ou acima dele também pode.
    if DOSSIE_CARGO_MINIMO_FECHAR_ID:
        cargo_minimo = member.guild.get_role(DOSSIE_CARGO_MINIMO_FECHAR_ID)
        if cargo_minimo and member.top_role.position >= cargo_minimo.position:
            return True

    return False


def mensagem_sem_permissao_fechar_mesa() -> str:
    return (
        "❌ Apenas a administração da DICOR pode fechar mesas.\n"
        "Cargo mínimo recomendado: **Inspetor DICOR para cima**."
    )


def cargos_equipe_permissoes(guild: discord.Guild) -> Dict[Any, discord.PermissionOverwrite]:
    overwrites: Dict[Any, discord.PermissionOverwrite] = {
        guild.default_role: discord.PermissionOverwrite(view_channel=False)
    }
    for cargo_id in set(CARGOS_EQUIPE_IDS + CARGOS_ADMIN_IDS):
        cargo = guild.get_role(cargo_id)
        if cargo:
            overwrites[cargo] = discord.PermissionOverwrite(
                view_channel=True,
                send_messages=True,
                attach_files=True,
                read_message_history=True,
                create_public_threads=True,
                send_messages_in_threads=True,
                manage_threads=True,
            )
    if guild.me:
        overwrites[guild.me] = discord.PermissionOverwrite(
            view_channel=True,
            send_messages=True,
            manage_channels=True,
            attach_files=True,
            read_message_history=True,
            create_public_threads=True,
            send_messages_in_threads=True,
            manage_threads=True,
        )
    return overwrites

# =====================================================
# CATALOGO HTML COM FOTO EM TELA CHEIA
# =====================================================

def valor_crimes_registro(registro: Dict[str, Any]) -> str:
    """Aceita o nome atual e também chaves usadas por versões antigas."""
    for chave in ("crimes", "crimes_cometidos", "crimes_imputados", "crime"):
        valor = str(registro.get(chave, "") or "").strip()
        if valor and valor.lower() not in {
            "não informado",
            "nao informado",
            "não identificado na mensagem antiga",
            "nao identificado na mensagem antiga",
        }:
            return valor
    return "Não informado"


def registro_tem_valor_util(valor: Any) -> bool:
    texto = str(valor or "").strip().lower()
    return bool(texto) and texto not in {
        "não informado",
        "nao informado",
        "não identificado na mensagem antiga",
        "nao identificado na mensagem antiga",
        "importado automaticamente do histórico do canal.",
    }


def gerar_catalogo_html() -> None:
    procurados = carregar_procurados()

    # Corrige registros criados por versões antigas.
    alterou = False
    for registro in procurados:
        crimes = valor_crimes_registro(registro)
        if registro.get("crimes") != crimes:
            registro["crimes"] = crimes
            alterou = True
    if alterou:
        salvar_procurados(procurados)

    visiveis = [
        p for p in procurados
        if str(p.get("status", "A PROCURAR") or "A PROCURAR").upper() != "APAGADO"
    ]
    ativos = [
        p for p in visiveis
        if str(p.get("status", "A PROCURAR") or "A PROCURAR").upper() != "RETIRADO"
    ]
    retirados = [
        p for p in visiveis
        if str(p.get("status", "") or "").upper() == "RETIRADO"
    ]

    def img_html(src: str, alt: str, vazio: str) -> str:
        if not src:
            return f'<div class="placeholder">{vazio}</div>'
        src_e = escape(src)
        return (
            f'<img class="zoom-img" src="{src_e}" alt="{escape(alt)}" '
            f'onclick="abrirFoto(this.src, this.alt)">'
        )

    def card_html(registro: Dict[str, Any]) -> str:
        status = str(registro.get("status", "A PROCURAR") or "A PROCURAR").upper()
        classe_status = "retirado" if status == "RETIRADO" else "ativo"
        foto_ind = registro.get("foto_individuo", "") or ""
        foto_rg = registro.get("foto_rg", "") or ""
        link_msg = registro.get("mensagem_url", "") or ""
        crimes = valor_crimes_registro(registro)
        registro_id_js = json.dumps(
            str(
                registro.get("id")
                or registro.get("caso")
                or registro.get("rg")
                or ""
            ),
            ensure_ascii=False,
        )
        nome_js = json.dumps(
            str(registro.get("nome") or "Sem nome"),
            ensure_ascii=False,
        )

        return f"""
        <article class="card {classe_status}">
            <div class="card-head">
                <div><span>Nº DO CASO</span><b>{escape(registro.get('caso'))}</b></div>
                <div><span>DATA</span><b>{escape(registro.get('data'))}</b></div>
                <div class="status">{escape(status)}</div>
            </div>

            <div class="card-grid">
                <div class="photos">
                    <div class="photo big">
                        <div class="photo-title">Foto do Indivíduo</div>
                        {img_html(foto_ind, 'Foto do indivíduo', 'Sem foto')}
                    </div>
                    <div class="photo small">
                        <div class="photo-title">Foto do RG</div>
                        {img_html(foto_rg, 'Foto do RG', 'Sem RG')}
                    </div>
                </div>

                <div class="info">
                    <div class="linha"><b>Nome</b><p>{escape(registro.get('nome'))}</p></div>
                    <div class="linha"><b>RG</b><p>{escape(registro.get('rg'))}</p></div>
                    <div class="box crimes"><b>Crimes Cometidos</b><p>{escape(crimes)}</p></div>
                    <div class="box destaque"><b>Último Avistamento</b><p>{escape(registro.get('ultimo_avistamento'))}</p></div>
                    <div class="box"><b>Número do Boletim</b><p>{escape(registro.get('numero_boletim') or registro.get('informacoes') or 'Não informado')}</p></div>
                    {f'<a class="msg" href="{escape(link_msg)}" target="_blank" rel="noopener noreferrer">Abrir mensagem no Discord</a>' if link_msg else ''}
                    {f'<div class="motivo"><b>Motivo da retirada:</b> {escape(registro.get("motivo_retirada"))}</div>' if registro.get('motivo_retirada') else ''}
                    <button class="apagar-registro" type="button" onclick='abrirExclusao({registro_id_js}, {nome_js})'>
                        🗑️ Apagar permanentemente
                    </button>
                </div>
            </div>
        </article>
        """

    cards_ativos = "".join(card_html(p) for p in ativos)
    cards_retirados = "".join(card_html(p) for p in retirados)

    if not cards_ativos:
        cards_ativos = """
        <div class="vazio">
            <h2>Nenhum procurado ativo.</h2>
            <p>Os novos procurados aparecerão nesta aba automaticamente.</p>
        </div>
        """

    if not cards_retirados:
        cards_retirados = """
        <div class="vazio">
            <h2>Nenhum procurado retirado.</h2>
            <p>Os registros retirados aparecerão nesta aba.</p>
        </div>
        """

    html_final = f"""
<!DOCTYPE html>
<html lang="pt-BR">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Catálogo de Procurados - DICOR</title>
    <style>
        * {{ box-sizing: border-box; }}
        body {{
            margin: 0;
            font-family: Arial, Helvetica, sans-serif;
            background:
                radial-gradient(circle at top, rgba(230,169,55,.12), transparent 35%),
                linear-gradient(135deg, #07111f, #02060d 70%);
            color: #f7f7f7;
        }}
        header {{
            border-bottom: 2px solid #d99a2b;
            padding: 28px 18px 22px;
            text-align: center;
            background: rgba(0,0,0,.38);
            position: sticky;
            top: 0;
            z-index: 10;
            backdrop-filter: blur(7px);
        }}
        header h1 {{ margin:0; color:#f0b64d; font-size:clamp(30px,5vw,62px); letter-spacing:2px; text-transform:uppercase; }}
        header p {{ margin:8px 0 0; color:#ddd; }}
        .stats {{ display:flex; gap:14px; justify-content:center; flex-wrap:wrap; margin-top:18px; }}
        .stat {{ border:1px solid rgba(240,182,77,.55); background:rgba(8,19,35,.9); padding:10px 18px; border-radius:12px; min-width:150px; }}
        .stat span {{ display:block; font-size:12px; color:#bfc6d1; text-transform:uppercase; }}
        .stat b {{ font-size:24px; color:#fff; }}
        .abas {{ display:flex; justify-content:center; gap:10px; flex-wrap:wrap; margin-top:18px; }}
        .aba-btn {{
            border:1px solid rgba(240,182,77,.65);
            background:#081323;
            color:#f7f7f7;
            padding:11px 18px;
            border-radius:12px;
            font-weight:800;
            cursor:pointer;
            transition:.15s ease;
        }}
        .aba-btn:hover {{ transform:translateY(-1px); border-color:#f0b64d; }}
        .aba-btn.ativa {{ background:#f0b64d; color:#07111f; }}
        .aba-conteudo {{ display:none; }}
        .aba-conteudo.ativa {{ display:block; }}
        .grade {{ width:min(1380px,96%); margin:24px auto 50px; display:grid; grid-template-columns:repeat(auto-fit,minmax(520px,1fr)); gap:22px; }}
        .card {{ border:1px solid rgba(240,182,77,.55); background:linear-gradient(180deg,rgba(8,19,35,.97),rgba(3,8,16,.97)); box-shadow:0 0 0 1px rgba(255,255,255,.05),0 18px 50px rgba(0,0,0,.32); border-radius:18px; overflow:hidden; }}
        .card.retirado {{ opacity:.82; filter:grayscale(.25); }}
        .card-head {{ display:grid; grid-template-columns:1fr 1fr auto; gap:12px; padding:14px; border-bottom:1px solid rgba(240,182,77,.32); align-items:center; }}
        .card-head span {{ display:block; font-size:11px; color:#f0b64d; text-transform:uppercase; font-weight:bold; }}
        .card-head b {{ display:block; margin-top:4px; }}
        .status {{ color:#101820; background:#f0b64d; font-weight:900; padding:8px 12px; border-radius:10px; white-space:nowrap; }}
        .retirado .status {{ background:#c94a4a; color:#fff; }}
        .card-grid {{ display:grid; grid-template-columns:210px 1fr; gap:16px; padding:16px; }}
        .photos {{ display:grid; gap:14px; }}
        .photo {{ border:1px solid rgba(240,182,77,.55); background:#081323; border-radius:14px; padding:10px; }}
        .photo-title {{ text-align:center; color:#f0b64d; font-weight:800; margin-bottom:8px; text-transform:uppercase; font-size:13px; }}
        .photo img {{ width:100%; height:240px; object-fit:cover; border-radius:10px; background:#ddd; display:block; }}
        .photo.small img {{ height:125px; object-fit:cover; }}
        .zoom-img {{ cursor:zoom-in; transition:transform .15s ease,filter .15s ease; }}
        .zoom-img:hover {{ transform:scale(1.02); filter:brightness(1.08); }}
        .placeholder {{ height:240px; border:1px dashed #98a2b3; display:grid; place-items:center; color:#98a2b3; border-radius:10px; }}
        .photo.small .placeholder {{ height:125px; }}
        .info {{ display:grid; gap:10px; }}
        .linha,.box {{ background:#f6f7f9; color:#07111f; border-radius:12px; padding:10px 12px; border-left:5px solid #f0b64d; }}
        .linha b,.box b {{ display:block; color:#07111f; margin-bottom:3px; }}
        .linha p,.box p {{ margin:0; white-space:pre-wrap; line-height:1.4; overflow-wrap:anywhere; }}
        .box.crimes {{ min-height:90px; }}
        .destaque {{ background:#fff4dc; }}
        .msg {{ display:inline-block; color:#f0b64d; text-decoration:none; font-weight:bold; margin-top:4px; }}
        .motivo {{ background:rgba(201,74,74,.18); color:#ffd7d7; border:1px solid rgba(201,74,74,.4); padding:10px; border-radius:10px; }}
        .apagar-registro {{
            border:1px solid rgba(222,75,75,.72);
            background:rgba(138,29,29,.25);
            color:#ffdede;
            padding:10px 12px;
            border-radius:10px;
            font-weight:800;
            cursor:pointer;
            transition:.15s ease;
            justify-self:start;
        }}
        .apagar-registro:hover {{
            background:#a92f2f;
            color:#fff;
            transform:translateY(-1px);
        }}
        .vazio {{ grid-column:1/-1; text-align:center; padding:60px 20px; border:1px dashed rgba(240,182,77,.6); border-radius:18px; background:rgba(8,19,35,.7); }}
        #modalExclusao {{
            display:none;
            position:fixed;
            inset:0;
            z-index:10000;
            background:rgba(0,0,0,.86);
            align-items:center;
            justify-content:center;
            padding:20px;
        }}
        #modalExclusao.ativo {{ display:flex; }}
        .modal-exclusao-caixa {{
            width:min(460px,96vw);
            background:#081323;
            border:1px solid rgba(240,182,77,.65);
            border-radius:16px;
            padding:22px;
            box-shadow:0 20px 80px rgba(0,0,0,.55);
        }}
        .modal-exclusao-caixa h2 {{
            margin:0 0 8px;
            color:#f0b64d;
        }}
        .modal-exclusao-caixa p {{
            color:#d8dde6;
            line-height:1.45;
        }}
        .modal-exclusao-caixa input {{
            width:100%;
            background:#02060d;
            color:#fff;
            border:1px solid #586274;
            border-radius:10px;
            padding:12px;
            font-size:16px;
            margin:8px 0 12px;
        }}
        .modal-acoes {{
            display:flex;
            gap:10px;
            flex-wrap:wrap;
        }}
        .modal-acoes button {{
            border:0;
            padding:11px 14px;
            border-radius:10px;
            font-weight:800;
            cursor:pointer;
        }}
        .confirmar-exclusao {{ background:#c83d3d; color:#fff; }}
        .cancelar-exclusao {{ background:#2a3546; color:#fff; }}
        #mensagemExclusao {{
            min-height:22px;
            margin-top:10px;
            color:#ffd078;
        }}
        footer {{ text-align:center; color:#bfc6d1; border-top:1px solid rgba(240,182,77,.25); padding:18px; }}
        #lightbox {{ display:none; position:fixed; inset:0; background:rgba(0,0,0,.94); z-index:9999; align-items:center; justify-content:center; padding:22px; }}
        #lightbox.ativo {{ display:flex; }}
        #lightbox img {{ max-width:96vw; max-height:92vh; object-fit:contain; border-radius:12px; box-shadow:0 0 30px rgba(0,0,0,.7); }}
        #fecharLightbox {{ position:fixed; top:18px; right:22px; background:#f0b64d; color:#06111f; border:0; padding:10px 15px; border-radius:10px; font-weight:900; cursor:pointer; }}
        #legendaLightbox {{ position:fixed; bottom:16px; left:50%; transform:translateX(-50%); color:#fff; background:rgba(0,0,0,.65); padding:8px 12px; border-radius:10px; }}
        @media (max-width:720px) {{
            .grade {{ grid-template-columns:1fr; }}
            .card-grid {{ grid-template-columns:1fr; }}
            .card-head {{ grid-template-columns:1fr; }}
            .photo img {{ height:330px; }}
        }}
    </style>
</head>
<body>
    <header>
        <h1>Catálogo de Procurados</h1>
        <p>DICOR • Atualizado automaticamente pelo bot • Clique nas fotos para abrir em tela cheia</p>
        <div class="stats">
            <div class="stat"><span>Registros totais</span><b>{len(visiveis)}</b></div>
            <div class="stat"><span>Ativos</span><b>{len(ativos)}</b></div>
            <div class="stat"><span>Retirados</span><b>{len(retirados)}</b></div>
            <div class="stat"><span>Última atualização</span><b style="font-size:16px">{agora_br()}</b></div>
        </div>
        <div class="abas">
            <button id="btn-ativos" class="aba-btn ativa" onclick="mostrarAba('ativos')">🚨 Procurados Ativos ({len(ativos)})</button>
            <button id="btn-retirados" class="aba-btn" onclick="mostrarAba('retirados')">✅ Retirados ({len(retirados)})</button>
        </div>
    </header>

    <section id="aba-ativos" class="aba-conteudo ativa">
        <main class="grade">{cards_ativos}</main>
    </section>

    <section id="aba-retirados" class="aba-conteudo">
        <main class="grade">{cards_retirados}</main>
    </section>

    <footer>Uso exclusivo para GTA RP / sistema interno autorizado.</footer>

    <div id="modalExclusao" onclick="fecharExclusao(event)">
        <div class="modal-exclusao-caixa" onclick="event.stopPropagation()">
            <h2>🗑️ Apagar registro</h2>
            <p id="textoExclusao">
                Essa ação remove o procurado de forma permanente das abas de ativos e retirados.
            </p>
            <input
                id="senhaExclusao"
                type="password"
                autocomplete="current-password"
                placeholder="Senha administrativa do catálogo"
            >
            <div class="modal-acoes">
                <button class="confirmar-exclusao" type="button" onclick="confirmarExclusao()">
                    Apagar permanentemente
                </button>
                <button class="cancelar-exclusao" type="button" onclick="fecharExclusao()">
                    Cancelar
                </button>
            </div>
            <div id="mensagemExclusao"></div>
        </div>
    </div>

    <div id="lightbox" onclick="fecharFoto()">
        <button id="fecharLightbox" onclick="fecharFoto()">Fechar ✕</button>
        <img id="lightboxImg" src="" alt="Foto em tela cheia">
        <div id="legendaLightbox"></div>
    </div>

    <script>
        let registroExclusaoId = '';
        let registroExclusaoNome = '';

        function abrirExclusao(id, nome) {{
            registroExclusaoId = id || '';
            registroExclusaoNome = nome || 'Sem nome';
            document.getElementById('textoExclusao').textContent =
                'Apagar permanentemente "' + registroExclusaoNome +
                '"? Ele não aparecerá mais em nenhuma aba do catálogo.';
            document.getElementById('senhaExclusao').value = '';
            document.getElementById('mensagemExclusao').textContent = '';
            document.getElementById('modalExclusao').classList.add('ativo');
            setTimeout(
                () => document.getElementById('senhaExclusao').focus(),
                50
            );
        }}

        function fecharExclusao(event) {{
            if (
                event
                && event.target
                && event.target.id !== 'modalExclusao'
            ) {{
                return;
            }}
            document.getElementById('modalExclusao').classList.remove('ativo');
            registroExclusaoId = '';
        }}

        async function confirmarExclusao() {{
            const senha = document.getElementById('senhaExclusao').value;
            const mensagem = document.getElementById('mensagemExclusao');

            if (!senha) {{
                mensagem.textContent = 'Digite a senha administrativa.';
                return;
            }}

            mensagem.textContent = 'Apagando registro...';
            try {{
                const resposta = await fetch('/api/catalogo/apagar', {{
                    method: 'POST',
                    headers: {{'Content-Type': 'application/json'}},
                    body: JSON.stringify({{
                        id: registroExclusaoId,
                        senha: senha
                    }})
                }});
                const dados = await resposta.json();

                if (!resposta.ok || !dados.ok) {{
                    mensagem.textContent =
                        dados.erro || 'Não foi possível apagar o registro.';
                    return;
                }}

                mensagem.textContent =
                    'Registro apagado. Atualizando o catálogo...';
                window.location.reload();
            }} catch (erro) {{
                mensagem.textContent =
                    'Falha ao conectar ao servidor do catálogo.';
            }}
        }}

        function mostrarAba(nome) {{
            document.querySelectorAll('.aba-conteudo').forEach(el => el.classList.remove('ativa'));
            document.querySelectorAll('.aba-btn').forEach(el => el.classList.remove('ativa'));
            document.getElementById('aba-' + nome).classList.add('ativa');
            document.getElementById('btn-' + nome).classList.add('ativa');
            window.scrollTo({{ top: 0, behavior: 'smooth' }});
        }}
        function abrirFoto(src, alt) {{
            document.getElementById('lightboxImg').src = src;
            document.getElementById('legendaLightbox').textContent = alt || '';
            document.getElementById('lightbox').classList.add('ativo');
            document.body.style.overflow = 'hidden';
        }}
        function fecharFoto() {{
            document.getElementById('lightbox').classList.remove('ativo');
            document.getElementById('lightboxImg').src = '';
            document.getElementById('legendaLightbox').textContent = '';
            document.body.style.overflow = '';
        }}
        document.addEventListener('keydown', function(e) {{
            if (e.key === 'Escape') {{
                fecharFoto();
                document
                    .getElementById('modalExclusao')
                    .classList.remove('ativo');
            }}
        }});
        document.getElementById('lightboxImg').addEventListener('click', function(e) {{ e.stopPropagation(); }});
        document.getElementById('fecharLightbox').addEventListener('click', function(e) {{ e.stopPropagation(); }});
    </script>
</body>
</html>
    """
    with open(CATALOGO_HTML, "w", encoding="utf-8") as f:
        f.write(html_final)

async def salvar_anexo_publico(anexo: discord.Attachment, prefixo: str) -> str:
    ext = Path(anexo.filename).suffix.lower()
    if ext not in [".png", ".jpg", ".jpeg", ".webp", ".gif"]:
        ext = ".png"
    nome = f"{data_caso()}-{slugify(prefixo)}-{slugify(Path(anexo.filename).stem)}{ext}"
    caminho = UPLOADS_DIR / nome
    await anexo.save(str(caminho))
    return f"uploads/{nome}"


async def salvar_procurado_catalogo(registro: Dict[str, Any]) -> None:
    registro["crimes"] = valor_crimes_registro(registro)
    lista = carregar_procurados()
    lista.append(registro)
    lista = remover_duplicados(lista)
    salvar_procurados(lista)
    gerar_catalogo_html()

# =====================================================
# SERVIDOR HTML
# =====================================================

async def pagina_inicial(request):
    gerar_catalogo_html()
    return web.FileResponse(CATALOGO_HTML)


async def obter_canal_por_id(canal_id: int):
    if not canal_id:
        return None

    canal = bot.get_channel(canal_id)
    if canal is not None:
        return canal

    try:
        return await bot.fetch_channel(canal_id)
    except Exception:
        return None


async def apagar_mensagens_discord_do_registro(
    registro: Dict[str, Any],
) -> None:
    """Apaga o post ativo e/ou o post arquivado relacionado ao registro."""
    alvos = [
        (PROCURADOS_CHANNEL_ID, registro.get("mensagem_id")),
        (
            HISTORICO_PROCURADOS_ID,
            registro.get("mensagem_arquivada_id"),
        ),
    ]

    for canal_id, mensagem_id in alvos:
        if not canal_id or not mensagem_id:
            continue

        canal = await obter_canal_por_id(int(canal_id))
        if canal is None or not hasattr(canal, "fetch_message"):
            continue

        try:
            mensagem = await canal.fetch_message(int(mensagem_id))
            await mensagem.delete()
        except discord.NotFound:
            pass
        except Exception as erro:
            await enviar_log(
                f"⚠️ Não foi possível apagar a mensagem `{mensagem_id}` "
                f"do registro `{registro.get('id')}`: {erro}"
            )


def apagar_arquivos_locais_do_registro(
    registro: Dict[str, Any],
) -> None:
    """Remove as fotos locais pertencentes ao registro apagado."""
    for chave in ("foto_individuo", "foto_rg"):
        caminho_relativo = str(registro.get(chave, "") or "").strip()

        if (
            not caminho_relativo
            or not caminho_relativo.startswith("uploads/")
        ):
            continue

        caminho = PUBLIC_DIR / caminho_relativo
        try:
            caminho_resolvido = caminho.resolve()
            uploads_resolvido = UPLOADS_DIR.resolve()

            if (
                uploads_resolvido in caminho_resolvido.parents
                and caminho_resolvido.exists()
            ):
                caminho_resolvido.unlink()
        except Exception:
            pass


async def api_apagar_procurado(
    request: web.Request,
) -> web.Response:
    ip = request.remote or "desconhecido"

    if not catalogo_limite_tentativas(ip):
        return web.json_response(
            {
                "ok": False,
                "erro": "Muitas tentativas. Aguarde 10 minutos.",
            },
            status=429,
        )

    try:
        dados = await request.json()
    except Exception:
        return web.json_response(
            {"ok": False, "erro": "Requisição inválida."},
            status=400,
        )

    senha = str(dados.get("senha", "") or "")
    registro_id = str(dados.get("id", "") or "").strip()

    if not senha_catalogo_valida(senha):
        registrar_falha_senha_catalogo(ip)
        return web.json_response(
            {"ok": False, "erro": "Senha incorreta."},
            status=401,
        )

    if not registro_id:
        return web.json_response(
            {"ok": False, "erro": "Registro não informado."},
            status=400,
        )

    lista = carregar_procurados()
    encontrado = None
    nova_lista: List[Dict[str, Any]] = []

    for registro in lista:
        identificador = str(
            registro.get("id")
            or registro.get("caso")
            or registro.get("rg")
            or ""
        )

        if encontrado is None and identificador == registro_id:
            encontrado = registro
            continue

        nova_lista.append(registro)

    if encontrado is None:
        return web.json_response(
            {
                "ok": False,
                "erro": "O procurado já foi apagado ou não existe.",
            },
            status=404,
        )

    # Remove fisicamente do JSON. Assim não aparece em nenhuma aba.
    salvar_procurados(nova_lista)
    gerar_catalogo_html()

    # Limpa, quando possível, os posts e as fotos relacionados.
    await apagar_mensagens_discord_do_registro(encontrado)
    apagar_arquivos_locais_do_registro(encontrado)

    catalogo_tentativas_senha.pop(ip, None)

    await enviar_log(
        f"🗑️ **Procurado apagado permanentemente pelo catálogo**\n"
        f"Nome: {encontrado.get('nome')}\n"
        f"RG: {encontrado.get('rg')}\n"
        f"Registro: {registro_id}"
    )

    return web.json_response(
        {
            "ok": True,
            "mensagem": "Registro apagado permanentemente.",
        }
    )


async def start_web_server():
    gerar_catalogo_html()
    app = web.Application()
    app.router.add_get("/", pagina_inicial)
    app.router.add_get("/index.html", pagina_inicial)
    app.router.add_post(
        "/api/catalogo/apagar",
        api_apagar_procurado,
    )
    app.router.add_static(
        "/uploads/",
        path=str(UPLOADS_DIR),
        show_index=False,
    )
    runner = web.AppRunner(app)
    await runner.setup()
    site = web.TCPSite(runner, "0.0.0.0", PORT)
    await site.start()
    print(f"Catalogo rodando na porta {PORT}")

# =====================================================
# PROCURADOS
# =====================================================

def mencao_auditoria_criador(registro: Dict[str, Any]) -> str:
    """Linha discreta usada para auditoria/UP.
    A menção fica em formato de citação cinza no Discord, sem poluir o texto oficial.
    O bot usa essa linha para contar a atividade para o agente correto.
    """
    usuario_id = (
        registro.get("criado_por_id")
        or registro.get("autor_id")
        or registro.get("comunicante_id")
        or registro.get("publicado_por_id")
    )
    if usuario_id and str(usuario_id).strip().isdigit() and str(usuario_id).strip() != "0":
        return f"\n\n> 👮 Criado por: <@{int(usuario_id)}>"

    nome = (
        registro.get("criado_por_nome")
        or registro.get("autor_nome")
        or registro.get("comunicante_nome")
        or registro.get("publicado_por_nome")
        or ""
    )
    if str(nome).strip():
        return f"\n\n> 👮 Criado por: {nome}"
    return ""


def criar_texto_procurado(registro: Dict[str, Any]) -> str:
    return f"""
🚨 **MANDADO DE PRISÃO E PROCURAÇÃO INVESTIGATIVA** 🚨

A Polícia Polícia Federal de Capital Morada, por intermédio da **Divisão de Investigações Criminais (DICOR)**, informa que o indivíduo abaixo encontra-se oficialmente procurado pelas autoridades competentes.

As investigações apontam seu envolvimento em atividades criminosas, havendo mandado ativo para sua localização, abordagem e condução para os procedimentos cabíveis.

📍 **ÚLTIMO AVISTAMENTO:** {registro.get('ultimo_avistamento', 'Não informado')}

⚠️ **CRIMES IMPUTADOS:**
{registro.get('crimes', 'Não informado')}

📋 **NÚMERO DO BOLETIM:** {registro.get('numero_boletim') or registro.get('informacoes') or 'Não informado'}

━━━━━━━━━━━━━━━━━━━━━━━

🆔 **IDENTIFICAÇÃO DO PROCURADO**

👤 **Nome:** {registro.get('nome', 'Não informado')}
🆔 **RG:** {registro.get('rg', 'Não informado')}

━━━━━━━━━━━━━━━━━━━━━━━

📞 Qualquer informação sobre o paradeiro deste indivíduo deverá ser repassada imediatamente a um agente da Polícia Federal ou da DICOR.

🔒 O sigilo do denunciante será integralmente preservado.

🔹 Polícia Polícia Federal de Capital Morada
🔹 Divisão de Investigações Criminais (DICOR)
""".strip() + mencao_auditoria_criador(registro)


def criar_embed_procurado(registro: Dict[str, Any]) -> discord.Embed:
    cor = discord.Color.red() if registro.get("status") == "A PROCURAR" else discord.Color.dark_gray()
    embed = discord.Embed(description=criar_texto_procurado(registro), color=cor)
    embed.set_footer(text=f"Caso: {registro.get('caso')} • Cadastro: {registro.get('data')}")
    return embed


async def postar_procurado_oficial(registro: Dict[str, Any]) -> Optional[discord.Message]:
    canal = bot.get_channel(PROCURADOS_CHANNEL_ID)
    if not canal:
        return None

    arquivos = []
    for chave, nome in [("foto_individuo", "foto_individuo"), ("foto_rg", "foto_rg")]:
        rel = registro.get(chave, "")
        if rel:
            caminho = PUBLIC_DIR / rel
            if caminho.exists():
                arquivos.append(discord.File(str(caminho), filename=f"{nome}_{registro.get('rg','')}.png"))

    # Mensagem normal, sem embed, para ficar igual ao modelo antigo.
    # As imagens seguem anexadas no mesmo envio e aparecem embaixo da mensagem.
    texto = cortar_discord(criar_texto_procurado(registro), 1900)
    msg = await canal.send(
        content=texto,
        files=arquivos,
        allowed_mentions=discord.AllowedMentions(users=False, roles=False, everyone=False),
    )
    return msg


async def atualizar_post_procurado_discord(registro: Dict[str, Any]) -> bool:
    """Atualiza o texto do procurado no canal oficial ou no histórico."""
    status = str(registro.get("status", "A PROCURAR") or "A PROCURAR").upper()

    alvos = []
    if status == "RETIRADO":
        alvos.append((HISTORICO_PROCURADOS_ID, registro.get("mensagem_arquivada_id")))
        alvos.append((PROCURADOS_CHANNEL_ID, registro.get("mensagem_id")))
    else:
        alvos.append((PROCURADOS_CHANNEL_ID, registro.get("mensagem_id")))
        alvos.append((HISTORICO_PROCURADOS_ID, registro.get("mensagem_arquivada_id")))

    texto = cortar_discord(criar_texto_procurado(registro), 1900)

    for canal_id, mensagem_id in alvos:
        if not canal_id or not mensagem_id:
            continue

        canal = await obter_canal_por_id(int(canal_id))
        if canal is None or not hasattr(canal, "fetch_message"):
            continue

        try:
            mensagem = await canal.fetch_message(int(mensagem_id))
            await mensagem.edit(content=texto)
            registro["mensagem_url"] = mensagem.jump_url
            return True
        except discord.NotFound:
            continue
        except Exception as erro:
            await enviar_log(
                f"⚠️ Não foi possível editar a mensagem do procurado "
                f"`{registro.get('nome')}` / RG `{registro.get('rg')}`: {erro}"
            )
            continue

    return False


def juntar_crimes_procurado(crimes_atuais: str, novos_crimes: str) -> str:
    atual = str(crimes_atuais or "").strip()
    novos = str(novos_crimes or "").strip()

    if not atual:
        return novos or "Não informado"
    if not novos:
        return atual

    atual_norm = atual.casefold()
    novos_norm = novos.casefold()
    if novos_norm in atual_norm:
        return atual

    separador = "\n" if atual.endswith(("\n", ";", ",")) else "\n"
    return (atual + separador + novos).strip()


def arquivos_locais_procurado(
    registro: Dict[str, Any],
) -> List[discord.File]:
    arquivos: List[discord.File] = []

    for chave, nome_base in (
        ("foto_individuo", "foto_individuo"),
        ("foto_rg", "foto_rg"),
    ):
        caminho_relativo = str(registro.get(chave, "") or "").strip()
        if not caminho_relativo:
            continue

        caminho = PUBLIC_DIR / caminho_relativo
        if not caminho.exists():
            continue

        extensao = caminho.suffix.lower() or ".png"
        arquivos.append(
            discord.File(
                str(caminho),
                filename=(
                    f"{nome_base}_{registro.get('rg', '')}{extensao}"
                ),
            )
        )

    return arquivos


async def arquivar_procurado_discord(
    registro: Dict[str, Any],
    motivo: str,
    retirado_por: str,
) -> discord.Message:
    """
    Publica o procurado completo no histórico antes de apagar a publicação ativa.
    Se o envio ao histórico falhar, a função gera erro e nada é apagado.
    """
    canal_historico = await obter_canal_por_id(HISTORICO_PROCURADOS_ID)
    if canal_historico is None or not hasattr(canal_historico, "send"):
        raise RuntimeError(f"Canal de histórico `{HISTORICO_PROCURADOS_ID}` não encontrado.")

    mensagem_original = None
    if registro.get("mensagem_id") and PROCURADOS_CHANNEL_ID:
        canal_ativo = await obter_canal_por_id(PROCURADOS_CHANNEL_ID)
        if canal_ativo is not None and hasattr(canal_ativo, "fetch_message"):
            try:
                mensagem_original = await canal_ativo.fetch_message(int(registro["mensagem_id"]))
            except Exception:
                mensagem_original = None

    arquivos: List[discord.File] = []
    links_extras: List[str] = []
    if mensagem_original is not None:
        for anexo in mensagem_original.attachments[:10]:
            try:
                arquivos.append(await anexo.to_file(use_cached=True))
            except Exception:
                links_extras.append(anexo.url)
    if not arquivos:
        arquivos = arquivos_locais_procurado(registro)

    texto_original = ""
    if mensagem_original:
        try:
            texto_original = coletar_texto_embed(mensagem_original) or mensagem_original.content or ""
        except Exception:
            texto_original = mensagem_original.content or ""

    texto_arquivado = cortar_discord(
        "📁 **PROCURADO ARQUIVADO**\n\n"
        f"**Nome:** {registro.get('nome', 'Não informado')}\n"
        f"**RG:** {registro.get('rg', 'Não informado')}\n"
        f"**Características:** {registro.get('caracteristicas', 'Não informado')}\n"
        f"**Crimes:**\n{valor_crimes_registro(registro)}\n\n"
        f"**Outras informações:** {registro.get('informacoes') or registro.get('ultimo_avistamento') or 'Não informado'}\n"
        f"**Motivo da retirada:** {motivo}\n"
        f"**Retirado por:** {retirado_por}\n"
        f"**Data da retirada:** {agora_br()}\n"
        f"**Autor original:** {registro.get('autor_nome', 'Não informado')}\n"
        f"**Data original do cadastro:** {registro.get('data', 'Não informado')}\n"
        f"**Publicação original:** {registro.get('mensagem_url') or 'Não informado'}\n"
        + ("\n**Links/anexos adicionais:**\n" + "\n".join(links_extras) if links_extras else "")
        + ("\n\n**Texto original preservado:**\n" + texto_original[:900] if texto_original else ""),
        1900,
    )

    mensagem_arquivada = await canal_historico.send(content=texto_arquivado, files=arquivos)

    if mensagem_original is not None:
        try:
            await mensagem_original.delete(reason="Procurado retirado e arquivado no histórico")
        except Exception as erro:
            await enviar_log(
                "⚠️ Procurado arquivado, mas não foi possível apagar a mensagem ativa "
                f"`{registro.get('mensagem_id')}`: {erro}"
            )

    return mensagem_arquivada


class NovoProcuradoModal(Modal, title="Cadastrar Novo Procurado"):
    nome = TextInput(label="Nome", placeholder="Nome do procurado", max_length=100)
    rg = TextInput(label="RG", placeholder="RG do procurado", max_length=50)
    crimes = TextInput(label="Crimes Cometidos", placeholder="Ex: Art. 8.3\nFormação de Quadrilha", style=discord.TextStyle.paragraph, max_length=1000)
    ultimo = TextInput(label="Último Avistamento", placeholder="Ex: Caixa d'água", style=discord.TextStyle.paragraph, max_length=600)
    numero_boletim = TextInput(label="Número do boletim", placeholder="Ex: 1, 01 ou 001", max_length=20, required=True)

    async def on_submit(self, interaction: discord.Interaction):
        guild = interaction.guild
        if guild is None:
            await interaction.response.send_message("❌ Use isso dentro de um servidor.", ephemeral=True)
            return

        numero_boletim = normalizar_boletim_procurado(str(self.numero_boletim.value))
        if not numero_boletim:
            await interaction.response.send_message(
                "❌ O número do boletim é obrigatório. Use apenas número, exemplo: `1`, `01` ou `001`.",
                ephemeral=True,
            )
            return

        categoria = guild.get_channel(PROCURADOS_TEMP_CATEGORY_ID) if PROCURADOS_TEMP_CATEGORY_ID else None
        overwrites = {
            guild.default_role: discord.PermissionOverwrite(view_channel=False),
            interaction.user: discord.PermissionOverwrite(view_channel=True, send_messages=True, attach_files=True, read_message_history=True),
        }
        if guild.me:
            overwrites[guild.me] = discord.PermissionOverwrite(view_channel=True, send_messages=True, manage_channels=True, read_message_history=True, attach_files=True)
        for cargo_id in set(CARGOS_ADMIN_IDS + CARGOS_EQUIPE_IDS):
            cargo = guild.get_role(cargo_id)
            if cargo:
                overwrites[cargo] = discord.PermissionOverwrite(view_channel=True, send_messages=True, read_message_history=True, attach_files=True)

        nome_canal = f"📸・procurado-{slugify(str(self.nome.value))}"
        canal = await guild.create_text_channel(name=nome_canal, category=categoria, overwrites=overwrites)

        cadastros_pendentes[canal.id] = {
            "nome": str(self.nome.value),
            "rg": str(self.rg.value),
            "crimes": str(self.crimes.value),
            "ultimo_avistamento": str(self.ultimo.value),
            "numero_boletim": numero_boletim,
            "autor_id": interaction.user.id,
            "autor_nome": str(interaction.user),
        }

        await canal.send(
            f"🚨 **Cadastro de Procurado — DICOR**\n\n"
            f"👤 **Nome:** {self.nome.value}\n"
            f"🪪 **RG:** {self.rg.value}\n"
            f"📋 **Boletim vinculado:** {numero_boletim}\n\n"
            f"📸 Envie aqui **2 imagens obrigatórias**:\n"
            f"**1. Foto do indivíduo**\n"
            f"**2. Foto do RG**\n\n"
            f"Depois clique em **✅ Finalizar Cadastro**.",
            view=FinalizarProcuradoView()
        )

        await enviar_log(f"📁 Canal provisório de procurado criado por {interaction.user.mention}: {canal.mention}")
        await interaction.response.send_message(f"✅ Canal provisório criado: {canal.mention}", ephemeral=True)


class FinalizarProcuradoView(View):
    def __init__(self):
        super().__init__(timeout=None)

    @discord.ui.button(label="Finalizar Cadastro", emoji="✅", style=discord.ButtonStyle.green, custom_id="dic_finalizar_procurado")
    async def finalizar(self, interaction: discord.Interaction, button: Button):
        canal = interaction.channel
        if not isinstance(canal, discord.TextChannel):
            await interaction.response.send_message("❌ Canal inválido.", ephemeral=True)
            return

        dados = cadastros_pendentes.get(canal.id)
        if not dados:
            await interaction.response.send_message("❌ Não encontrei os dados desse cadastro. Crie novamente pelo painel.", ephemeral=True)
            return

        if not normalizar_boletim_procurado(str(dados.get("numero_boletim", ""))):
            await interaction.response.send_message(
                "❌ Esse cadastro está sem número de boletim válido. Cancele e cadastre novamente usando `1`, `01` ou `001`.",
                ephemeral=True,
            )
            return

        if interaction.user.id != dados.get("autor_id") and isinstance(interaction.user, discord.Member) and not usuario_tem_admin(interaction.user):
            await interaction.response.send_message("❌ Apenas quem criou ou a equipe pode finalizar.", ephemeral=True)
            return

        anexos: List[discord.Attachment] = []
        async for msg in canal.history(limit=100, oldest_first=True):
            for a in msg.attachments:
                if (a.content_type and a.content_type.startswith("image/")) or Path(a.filename).suffix.lower() in [".png", ".jpg", ".jpeg", ".webp", ".gif"]:
                    anexos.append(a)

        if len(anexos) < 2:
            await interaction.response.send_message("❌ Envie as 2 imagens: foto do indivíduo e foto do RG.", ephemeral=True)
            return

        await interaction.response.defer(ephemeral=True)
        foto_ind = await salvar_anexo_publico(anexos[0], f"individuo-{dados['rg']}")
        foto_rg = await salvar_anexo_publico(anexos[1], f"rg-{dados['rg']}")

        registro = {
            "id": data_caso(),
            "caso": f"DICOR-{data_caso()}",
            "data": agora_br(),
            "status": "A PROCURAR",
            "nome": dados["nome"],
            "rg": dados["rg"],
            "crimes": dados["crimes"],
            "ultimo_avistamento": dados["ultimo_avistamento"],
            "numero_boletim": dados.get("numero_boletim", ""),
            "informacoes": dados.get("numero_boletim", ""),
            "foto_individuo": foto_ind,
            "foto_rg": foto_rg,
            "autor_id": dados["autor_id"],
            "autor_nome": dados["autor_nome"],
            "criado_por_id": dados["autor_id"],
            "criado_por_nome": dados["autor_nome"],
            "finalizado_por_id": interaction.user.id,
            "finalizado_por_nome": str(interaction.user),
            "mensagem_id": None,
            "mensagem_url": None,
        }

        msg = await postar_procurado_oficial(registro)
        if msg:
            registro["mensagem_id"] = msg.id
            registro["mensagem_url"] = msg.jump_url

        await salvar_procurado_catalogo(registro)
        cadastros_pendentes.pop(canal.id, None)

        await enviar_log(f"✅ Procurado cadastrado: {registro['nome']} | RG {registro['rg']} | Boletim {registro.get('numero_boletim')} | Catálogo: {CATALOG_PUBLIC_URL}")
        await interaction.followup.send(f"✅ Procurado cadastrado e enviado ao catálogo: {CATALOG_PUBLIC_URL}", ephemeral=True)
        await canal.send("✅ Cadastro finalizado. Este canal será apagado em 5 segundos.")
        await asyncio.sleep(5)
        try:
            await canal.delete(reason="Cadastro de procurado finalizado")
        except Exception:
            pass

    @discord.ui.button(label="Cancelar", emoji="❌", style=discord.ButtonStyle.red, custom_id="dic_cancelar_procurado")
    async def cancelar(self, interaction: discord.Interaction, button: Button):
        canal = interaction.channel
        if isinstance(canal, discord.TextChannel):
            cadastros_pendentes.pop(canal.id, None)
            await interaction.response.send_message("Cadastro cancelado. Canal será apagado.", ephemeral=True)
            await asyncio.sleep(3)
            try:
                await canal.delete(reason="Cadastro de procurado cancelado")
            except Exception:
                pass


class RetirarProcuradoModal(Modal, title="Retirar Procurado"):
    rg = TextInput(label="RG", placeholder="RG do procurado", max_length=50)
    motivo = TextInput(label="Motivo", placeholder="Motivo da retirada", style=discord.TextStyle.paragraph, max_length=800, required=False)

    async def on_submit(self, interaction: discord.Interaction):
        await retirar_procurado(interaction, str(self.rg.value), str(self.motivo.value or "Não informado"))


class BuscarModificarProcuradoModal(Modal, title="Modificar Procurado"):
    rg = TextInput(
        label="RG do procurado",
        placeholder="Digite o RG para localizar o cadastro",
        max_length=50,
    )

    async def on_submit(self, interaction: discord.Interaction):
        if not isinstance(interaction.user, discord.Member) or not usuario_tem_equipe(interaction.user):
            await interaction.response.send_message(
                "❌ Apenas a equipe DICOR pode modificar procurados.",
                ephemeral=True,
            )
            return

        alvo = limpar_rg(str(self.rg.value))
        encontrado = None
        for registro in carregar_procurados():
            if limpar_rg(registro.get("rg", "")) == alvo:
                encontrado = registro
                break

        if not encontrado:
            await interaction.response.send_message(
                "❌ Não encontrei procurado com esse RG no catálogo.",
                ephemeral=True,
            )
            return

        await interaction.response.send_modal(EditarCrimesProcuradoModal(encontrado))


class EditarCrimesProcuradoModal(Modal):
    def __init__(self, registro: Dict[str, Any]):
        nome = str(registro.get("nome", "Procurado") or "Procurado")
        rg = str(registro.get("rg", "") or "")
        super().__init__(title="Editar Crimes do Procurado")

        self.rg_registro = rg
        self.nome_registro = nome

        crimes_atuais = valor_crimes_registro(registro)
        if len(crimes_atuais) > 3300:
            crimes_atuais = crimes_atuais[:3300].rstrip() + "\n..."

        self.crimes_cadastrados = TextInput(
            label=f"Crimes já cadastrados - {nome[:35]}",
            placeholder="Crimes atuais do procurado",
            style=discord.TextStyle.paragraph,
            default=crimes_atuais,
            required=False,
            max_length=3500,
        )
        self.novos_crimes = TextInput(
            label="Adicionar mais crimes",
            placeholder="Digite aqui somente os novos crimes que serão adicionados",
            style=discord.TextStyle.paragraph,
            required=False,
            max_length=1000,
        )

        self.add_item(self.crimes_cadastrados)
        self.add_item(self.novos_crimes)

    async def on_submit(self, interaction: discord.Interaction):
        if not isinstance(interaction.user, discord.Member) or not usuario_tem_equipe(interaction.user):
            await interaction.response.send_message(
                "❌ Apenas a equipe DICOR pode modificar procurados.",
                ephemeral=True,
            )
            return

        await interaction.response.defer(ephemeral=True, thinking=True)

        lista = carregar_procurados()
        alvo = limpar_rg(self.rg_registro)
        encontrado = None

        for registro in lista:
            if limpar_rg(registro.get("rg", "")) == alvo:
                encontrado = registro
                break

        if not encontrado:
            await interaction.followup.send(
                "❌ Esse procurado não existe mais no catálogo.",
                ephemeral=True,
            )
            return

        crimes_antigos = valor_crimes_registro(encontrado)
        crimes_editados = str(self.crimes_cadastrados.value or "").strip()
        crimes_adicionados = str(self.novos_crimes.value or "").strip()
        crimes_finais = juntar_crimes_procurado(crimes_editados, crimes_adicionados)

        if crimes_finais.strip() == crimes_antigos.strip():
            await interaction.followup.send(
                "⚠️ Nenhum crime novo foi adicionado ou alterado.",
                ephemeral=True,
            )
            return

        encontrado["crimes"] = crimes_finais
        encontrado.setdefault("historico_edicoes", []).append({
            "data": agora_br(),
            "tipo": "ALTERAÇÃO DE CRIMES",
            "campo": "crimes",
            "usuario": str(interaction.user),
            "usuario_id": interaction.user.id,
            "valor_anterior": crimes_antigos,
            "valor_novo": crimes_finais,
            "acrescimo": crimes_adicionados,
        })

        salvar_procurados(lista)
        gerar_catalogo_html()
        post_atualizado = await atualizar_post_procurado_discord(encontrado)

        await enviar_log(
            "✏️ **Procurado modificado**\n"
            f"Nome: {encontrado.get('nome')}\n"
            f"RG: {encontrado.get('rg')}\n"
            f"Alterado por: {interaction.user.mention}\n"
            f"Post Discord atualizado: {'sim' if post_atualizado else 'não'}"
        )

        aviso_post = (
            "✅ Post oficial atualizado."
            if post_atualizado
            else "⚠️ Catálogo atualizado, mas não encontrei o post antigo no Discord para editar."
        )

        await interaction.followup.send(
            f"✅ **Procurado atualizado com sucesso.**\n"
            f"👤 **Nome:** {encontrado.get('nome')}\n"
            f"🪪 **RG:** `{encontrado.get('rg')}`\n"
            f"{aviso_post}\n"
            f"🔗 {CATALOG_PUBLIC_URL}",
            ephemeral=True,
        )


class PainelProcuradosView(View):
    def __init__(self):
        super().__init__(timeout=None)

    @discord.ui.button(label="Novo Procurado", emoji="➕", style=discord.ButtonStyle.danger, custom_id="dic_novo_procurado")
    async def novo(self, interaction: discord.Interaction, button: Button):
        await interaction.response.send_modal(NovoProcuradoModal())

    @discord.ui.button(label="Lista de Procurados", emoji="📋", style=discord.ButtonStyle.blurple, custom_id="dic_lista_procurados")
    async def lista(self, interaction: discord.Interaction, button: Button):
        procurados = carregar_procurados()
        ativos = [p for p in procurados if p.get("status", "A PROCURAR") == "A PROCURAR"]
        if not ativos:
            texto = "Nenhum procurado ativo cadastrado."
        else:
            linhas = [f"• **{p.get('nome','Sem nome')}** — RG: `{p.get('rg','')}`" for p in ativos[:20]]
            texto = "\n".join(linhas)
            if len(ativos) > 20:
                texto += f"\n... e mais {len(ativos)-20} no catálogo."
        await interaction.response.send_message(f"📋 **Procurados ativos:**\n{texto}\n\n🔗 {CATALOG_PUBLIC_URL}", ephemeral=True)

    @discord.ui.button(label="Retirar Procurado", emoji="❌", style=discord.ButtonStyle.gray, custom_id="dic_retirar_procurado")
    async def retirar(self, interaction: discord.Interaction, button: Button):
        await interaction.response.send_modal(RetirarProcuradoModal())

    @discord.ui.button(label="Modificar Procurado", emoji="✏️", style=discord.ButtonStyle.primary, custom_id="dic_modificar_procurado", row=1)
    async def modificar(self, interaction: discord.Interaction, button: Button):
        if not isinstance(interaction.user, discord.Member) or not usuario_tem_equipe(interaction.user):
            await interaction.response.send_message(
                "❌ Apenas a equipe DICOR pode modificar procurados.",
                ephemeral=True,
            )
            return
        await interaction.response.send_modal(BuscarModificarProcuradoModal())

    @discord.ui.button(label="Abrir Catálogo", emoji="📄", style=discord.ButtonStyle.green, custom_id="dic_abrir_catalogo")
    async def abrir_catalogo(self, interaction: discord.Interaction, button: Button):
        await interaction.response.send_message(f"📄 **Catálogo de Procurados:**\n{CATALOG_PUBLIC_URL}", ephemeral=True)

    @discord.ui.button(label="Sincronizar Antigos", emoji="🔄", style=discord.ButtonStyle.secondary, custom_id="dic_sync_catalogo")
    async def sync_old(self, interaction: discord.Interaction, button: Button):
        await sincronizar_catalogo_core(interaction)


async def retirar_procurado(
    interaction: discord.Interaction,
    rg: str,
    motivo: str,
):
    if not interaction.response.is_done():
        await interaction.response.defer(ephemeral=True, thinking=True)

    if not isinstance(interaction.user, discord.Member) or not usuario_pode_fechar_mesa(interaction.user):
        await enviar_log(
            "🚫 **Tentativa sem permissão para retirar procurado**\n"
            f"Usuário: {interaction.user.mention} (`{interaction.user.id}`)\n"
            f"RG informado: `{rg}`\n"
            f"Canal: {getattr(interaction.channel, 'mention', 'Sem canal')}\n"
            f"Data: {agora_br()}"
        )
        await interaction.followup.send(
            "❌ Você não possui permissão para retirar procurados. Essa ação está disponível somente para Inspetor ou superior.",
            ephemeral=True,
        )
        return

    lista = carregar_procurados()
    alvo = limpar_rg(rg)
    encontrado = None

    for registro in lista:
        if limpar_rg(registro.get("rg", "")) == alvo:
            encontrado = registro
            break

    if not encontrado:
        await interaction.followup.send("❌ Não encontrei procurado com esse RG no catálogo.", ephemeral=True)
        return

    if str(encontrado.get("status", "") or "").upper() == "RETIRADO" and encontrado.get("mensagem_arquivada_id"):
        await interaction.followup.send("⚠️ Esse procurado já foi retirado e está arquivado.", ephemeral=True)
        return

    try:
        mensagem_arquivada = await arquivar_procurado_discord(encontrado, motivo, interaction.user.mention)
    except Exception as erro:
        await enviar_log(f"❌ Falha ao arquivar procurado {encontrado.get('nome')} (RG {encontrado.get('rg')}): {erro}")
        await interaction.followup.send(
            "❌ Não foi possível enviar o procurado para o canal de arquivados. Nada foi removido. Confira o ID e as permissões do canal.",
            ephemeral=True,
        )
        return

    encontrado["status"] = "RETIRADO"
    encontrado["motivo_retirada"] = motivo
    encontrado["data_retirada"] = agora_br()
    encontrado["retirado_por"] = str(interaction.user)
    encontrado["retirado_por_id"] = interaction.user.id
    encontrado["mensagem_original_id"] = encontrado.get("mensagem_id")
    encontrado["mensagem_original_url"] = encontrado.get("mensagem_url")
    encontrado["mensagem_arquivada_id"] = mensagem_arquivada.id
    encontrado["mensagem_arquivada_url"] = mensagem_arquivada.jump_url
    encontrado["mensagem_url"] = mensagem_arquivada.jump_url

    salvar_procurados(lista)
    gerar_catalogo_html()

    await enviar_log(
        "📁 **Procurado retirado e arquivado**\n"
        f"Nome: {encontrado.get('nome')}\n"
        f"RG: {encontrado.get('rg')}\n"
        f"Motivo: {motivo}\n"
        f"Retirado por: {interaction.user.mention} (`{interaction.user.id}`)\n"
        f"Arquivo: {mensagem_arquivada.jump_url}"
    )

    await interaction.followup.send(
        "✅ Procurado retirado, movido para o canal de arquivados e catálogo atualizado.\n"
        f"📁 {mensagem_arquivada.jump_url}\n"
        f"🔗 {CATALOG_PUBLIC_URL}",
        ephemeral=True,
    )

def limpar_markdown_extracao(texto: str) -> str:
    texto = str(texto or "").replace("**", "").replace("__", "")
    return texto.replace("`", "")


def extrair_bloco_rotulado(texto: str, rotulo: str, proximos: List[str]) -> str:
    fim = "|".join(proximos)
    padrao = rf"(?:{rotulo})\s*:\s*(.*?)(?=\n\s*(?:{fim})\s*:|\n\s*━+|\Z)"
    m = re.search(padrao, texto, flags=re.I | re.S)
    if not m:
        return ""
    valor = re.sub(r"\n{3,}", "\n\n", m.group(1).strip())
    return valor[:1800]


def extrair_por_labels(texto: str) -> Dict[str, str]:
    texto_limpo = limpar_markdown_extracao(texto)
    dados: Dict[str, str] = {}

    m_nome = re.search(r"(?:nome|indiv[ií]duo)\s*[:\-]\s*([^\n]+)", texto_limpo, flags=re.I)
    m_rg = re.search(r"(?:rg|passaporte)\s*[:\-]\s*([^\n]+)", texto_limpo, flags=re.I)
    if m_nome:
        dados["nome"] = m_nome.group(1).strip()[:200]
    if m_rg:
        dados["rg"] = m_rg.group(1).strip()[:100]

    proximos_comuns = [
        r"(?:📍\s*)?(?:último avistamento|ultimo avistamento|avistamento)",
        r"(?:⚠️\s*)?(?:crimes cometidos|crimes imputados|crimes|crime)",
        r"(?:ℹ️\s*)?(?:informações|informacoes|observações|observacoes|detalhes)",
        r"(?:🆔\s*)?(?:identificação do procurado|identificacao do procurado)",
        r"(?:👤\s*)?nome",
        r"(?:🆔\s*)?rg",
    ]

    crimes = extrair_bloco_rotulado(
        texto_limpo,
        r"(?:⚠️\s*)?(?:crimes cometidos|crimes imputados|crimes|crime)",
        proximos_comuns,
    )
    ultimo = extrair_bloco_rotulado(
        texto_limpo,
        r"(?:📍\s*)?(?:último avistamento|ultimo avistamento|avistamento)",
        proximos_comuns,
    )
    informacoes = extrair_bloco_rotulado(
        texto_limpo,
        r"(?:ℹ️\s*)?(?:informações|informacoes|observações|observacoes|detalhes)",
        proximos_comuns,
    )

    if crimes:
        dados["crimes"] = crimes
    if ultimo:
        dados["ultimo_avistamento"] = ultimo
    if informacoes:
        dados["informacoes"] = informacoes
    return dados


async def importar_mensagem_antiga(msg: discord.Message) -> Optional[Dict[str, Any]]:
    texto_total = msg.content or ""
    for embed in msg.embeds:
        if embed.title:
            texto_total += "\n" + embed.title
        if embed.description:
            texto_total += "\n" + embed.description
        for field in embed.fields:
            texto_total += f"\n{field.name}: {field.value}"

    dados = extrair_por_labels(texto_total)
    if not dados.get("rg") or not dados.get("nome"):
        return None

    existente = procurar_por_rg(dados.get("rg", ""))

    imagens = []
    for anexo in msg.attachments:
        if (
            anexo.content_type and anexo.content_type.startswith("image/")
        ) or Path(anexo.filename).suffix.lower() in [".png", ".jpg", ".jpeg", ".webp", ".gif"]:
            imagens.append(anexo)

    foto_ind = str(existente.get("foto_individuo", "") if existente else "")
    foto_rg = str(existente.get("foto_rg", "") if existente else "")
    if not foto_ind and len(imagens) >= 1:
        foto_ind = await salvar_anexo_publico(imagens[0], f"sync-ind-{dados.get('rg', '')}")
    if not foto_rg and len(imagens) >= 2:
        foto_rg = await salvar_anexo_publico(imagens[1], f"sync-rg-{dados.get('rg', '')}")

    return {
        "id": f"SYNC-{msg.id}",
        "caso": f"SYNC-{msg.id}",
        "data": msg.created_at.strftime("%d/%m/%Y %H:%M"),
        "status": str(existente.get("status", "A PROCURAR") if existente else "A PROCURAR"),
        "nome": dados.get("nome", ""),
        "rg": dados.get("rg", ""),
        "crimes": dados.get("crimes", "Não identificado na mensagem antiga"),
        "ultimo_avistamento": dados.get("ultimo_avistamento", "Não identificado na mensagem antiga"),
        "informacoes": dados.get("informacoes", "Importado automaticamente do histórico do canal."),
        "foto_individuo": foto_ind,
        "foto_rg": foto_rg,
        "autor_id": None,
        "autor_nome": "Sincronização antiga",
        "mensagem_id": msg.id,
        "mensagem_url": msg.jump_url,
    }

async def sincronizar_catalogo_core(interaction: discord.Interaction):
    if not PROCURADOS_CHANNEL_ID:
        await interaction.response.send_message("❌ Configure PROCURADOS_CHANNEL_ID.", ephemeral=True)
        return

    canal_ativos = bot.get_channel(PROCURADOS_CHANNEL_ID)
    canal_arquivados = bot.get_channel(HISTORICO_PROCURADOS_ID) if 'HISTORICO_PROCURADOS_ID' in globals() else None
    
    if not isinstance(canal_ativos, discord.TextChannel):
        await interaction.response.send_message("❌ Não encontrei o canal de procurados ativos.", ephemeral=True)
        return

    await interaction.response.defer(ephemeral=True)
    importados = 0
    atualizados = 0
    analisados = 0
    lista = carregar_procurados()

    # 1. PROCESSA PRIMEIRO OS PROCURADOS ATIVOS
    async for msg in canal_ativos.history(limit=1000, oldest_first=True):
        analisados += 1
        registro = await importar_mensagem_antiga(msg)
        if not registro:
            continue

        # Garante que mensagens do canal ativo fiquem com status correto
        registro["status"] = "A PROCURAR"

        alvo = limpar_rg(registro.get("rg", ""))
        existente = next((p for p in lista if limpar_rg(p.get("rg", "")) == alvo), None)

        if existente is None:
            registro["crimes"] = valor_crimes_registro(registro) if 'valor_crimes_registro' in globals() else registro.get("crimes")
            lista.append(registro)
            importados += 1
            continue

        # Se já existe na lista mas veio do canal ativo, mantém ativo
        existente["status"] = "A PROCURAR"
        mudou = False
        for campo in ("nome", "crimes", "ultimo_avistamento", "informacoes", "foto_individuo", "foto_rg", "mensagem_id", "mensagem_url"):
            novo_valor = registro.get(campo)
            valor_atual = existente.get(campo)
            if 'registro_tem_valor_util' in globals() and registro_tem_valor_util(novo_valor) and not registro_tem_valor_util(valor_atual):
                existente[campo] = novo_valor
                mudou = True
            elif novo_valor and not valor_atual:
                existente[campo] = novo_valor
                mudou = True

        if mudou:
            atualizados += 1

    # 2. PROCESSA OS PROCURADOS PEGOs/ARQUIVADOS (SE O CANAL EXISTIR)
    if isinstance(canal_arquivados, discord.TextChannel):
        async for msg in canal_arquivados.history(limit=1000, oldest_first=True):
            analisados += 1
            registro = await importar_mensagem_antiga(msg)
            if not registro:
                continue

            # Registros vindos deste canal estão marcados como RETIRADO / PEGO
            registro["status"] = "RETIRADO"

            alvo = limpar_rg(registro.get("rg", ""))
            existente = next((p for p in lista if limpar_rg(p.get("rg", "")) == alvo), None)

            if existente is None:
                registro["crimes"] = valor_crimes_registro(registro) if 'valor_crimes_registro' in globals() else registro.get("crimes")
                lista.append(registro)
                importados += 1
                continue
            
            # Se já existia como ativo, o histórico do canal de arquivados tem prioridade (já foi pego)
            existente["status"] = "RETIRADO"
            existente["mensagem_arquivada_id"] = msg.id
            existente["mensagem_url"] = msg.jump_url
            atualizados += 1

    if 'remover_duplicados' in globals():
        lista = remover_duplicados(lista)
        
    salvar_procurados(lista)
    gerar_catalogo_html()

    await interaction.followup.send(
        f"✅ Sincronização e separação finalizadas.\n"
        f"Mensagens analisadas: `{analisados}`\n"
        f"Procurados importados: `{importados}`\n"
        f"Registros atualizados/separados: `{atualizados}`\n"
        f"Catálogo: {CATALOG_PUBLIC_URL}",
        ephemeral=True,
    )

# =====================================================
# MESAS
# =====================================================

TOPICOS_MESA = [
    "📌 Painel",
    "📑・informações-gerais",
    "🖼️ Fotos líderes",
    "👥 Fotos dos membros",
    "📻 Rádio",
    "📍 Localização",
    "⚖️ Crimes da comunidade",
    "📦 Baú de líder",
    "📦 Baú de membros",
    "🌾 Rota de farm",
    "⚙️ Rota de produção",
    "🧪 Ingredientes base e produtos",
    "🕵️ Informante",
    "💬 Chat",
]



async def criar_topicos_reais_mesa(canal: discord.TextChannel) -> List[int]:
    """Cria cada item da mesa como tópico/thread público e já aberto."""
    topicos_ids: List[int] = []

    await canal.send(
        "🧵 **Tópicos da investigação**\n"
        "Abra o tópico desejado para adicionar textos, imagens e provas."
    )

    for topico in TOPICOS_MESA:
        mensagem = await canal.send(f"🧵 **{topico}**")
        try:
            try:
                thread = await mensagem.create_thread(
                    name=topico[:100],
                    auto_archive_duration=10080,
                    reason="Tópico automático da mesa de investigação",
                )
            except discord.HTTPException:
                thread = await mensagem.create_thread(
                    name=topico[:100],
                    auto_archive_duration=1440,
                    reason="Tópico automático da mesa de investigação",
                )

            await thread.send(
                "**📑 INFORMAÇÕES GERAIS DA INVESTIGAÇÃO**\n\n"
                "Utilize este espaço para registrar informações complementares, observações importantes, identificação de envolvidos, veículos, telefones, horários, rotinas, locais e qualquer outro dado relevante para a investigação."
                if "informações-gerais" in topico.lower() or "informacoes-gerais" in topico.lower()
                else f"## {topico}\n> Envie aqui todas as informações, imagens e provas relacionadas a este tópico."
            )
            topicos_ids.append(thread.id)
            await asyncio.sleep(0.25)
        except Exception as erro:
            await mensagem.reply(
                f"⚠️ Não foi possível abrir este tópico automaticamente: `{erro}`"
            )
            await enviar_log(
                f"⚠️ Erro ao criar tópico `{topico}` no canal {canal.id}: {erro}"
            )

    return topicos_ids

def registrar_mesa(dados: Dict[str, Any]) -> None:
    mesas = carregar_mesas()
    mesas = [m for m in mesas if m.get("canal_id") != dados.get("canal_id")]
    mesas.append(dados)
    salvar_mesas(mesas)


def buscar_mesa_por_canal(canal_id: int) -> Optional[Dict[str, Any]]:
    for mesa in carregar_mesas():
        if int(mesa.get("canal_id", 0)) == canal_id:
            return mesa
    return None


class CriarMesaModal(Modal, title="Criar Mesa de Investigação"):
    apelido = TextInput(label="Apelido do agente", placeholder="Ex: Baiano", max_length=50)
    familia = TextInput(label="Nome da família/organização", placeholder="Ex: Olimpo", max_length=80)
    observacao = TextInput(label="Observação inicial", placeholder="Opcional", style=discord.TextStyle.paragraph, required=False, max_length=600)

    async def on_submit(self, interaction: discord.Interaction):
        # Defere rápido para o Discord não mostrar "Esta interação falhou"
        # enquanto o bot cria o canal e envia todos os tópicos.
        await criar_mesa_core(interaction, str(self.apelido.value), str(self.familia.value), str(self.observacao.value or ""))

    async def on_error(self, interaction: discord.Interaction, error: Exception) -> None:
        await enviar_log(f"❌ Erro no modal de criar mesa: {error}")
        try:
            if interaction.response.is_done():
                await interaction.followup.send("❌ Deu erro ao criar a mesa. Veja os logs do bot/Railway.", ephemeral=True)
            else:
                await interaction.response.send_message("❌ Deu erro ao criar a mesa. Veja os logs do bot/Railway.", ephemeral=True)
        except Exception:
            pass


class FecharMesaView(View):
    def __init__(self, dados_mesa=None):
        super().__init__(timeout=None)
        self.dados_mesa = dados_mesa or {}

    @discord.ui.button(
        label="Fechar Mesa",
        emoji="🔒",
        style=discord.ButtonStyle.red,
        custom_id="dic_fechar_mesa_botao",
    )
    async def fechar(self, interaction: discord.Interaction, button: Button):
        # Responde/defera IMEDIATAMENTE para o Discord não mostrar "Esta interação falhou".
        try:
            if not interaction.response.is_done():
                await interaction.response.defer(ephemeral=True, thinking=True)
        except Exception:
            pass

        try:
            canal = interaction.channel
            if not isinstance(canal, discord.TextChannel):
                await interaction.followup.send("❌ Este botão só funciona dentro de um canal de mesa.", ephemeral=True)
                return

            if not usuario_pode_fechar_mesa(interaction.user):
                await interaction.followup.send(mensagem_sem_permissao_fechar_mesa(), ephemeral=True)
                return

            dados = dict(self.dados_mesa or {})

            # Fallback para recuperar os dados se a memória do bot limpar/reiniciar.
            if not dados:
                mesa_banco = buscar_mesa_por_canal(canal.id)
                if mesa_banco:
                    dados = {
                        "nome": f"OPERAÇÃO {mesa_banco.get('familia', canal.name).upper()}",
                        "comunidade": mesa_banco.get("familia", "Mapeada em logs"),
                        "delegado": mesa_banco.get("autor_nome", "Superintendência DICOR"),
                        "data_abertura": mesa_banco.get("criada_em", agora_br()),
                        "processo": f"PF-DICOR-{canal.id}",
                        "numero": f"INV-{str(canal.id)[-6:]}",
                    }

            if not dados:
                dados = {
                    "nome": canal.name.upper(),
                    "comunidade": "Setor Mapeado em Campo",
                    "delegado": getattr(interaction.user, "display_name", str(interaction.user)),
                    "data_abertura": agora_br(),
                    "processo": f"PF-DICOR-{canal.id}",
                    "numero": f"INV-{str(canal.id)[-6:]}",
                }

            self.dados_mesa = dados

            await interaction.followup.send(
                content=(
                    "⚠️ **Confirmação DICOR:** Você tem certeza que deseja encerrar esta mesa "
                    "tática e consolidar o Dossiê Operacional?"
                ),
                view=ConfirmacaoFecharMesaView(dados),
                ephemeral=True,
            )

        except Exception as erro:
            await enviar_log(f"❌ Erro ao abrir confirmação de fechamento da mesa: {erro}")
            try:
                await interaction.followup.send(
                    "❌ O botão de fechar mesa deu erro. Use `/fecharmesa` neste canal ou veja os logs do Railway.",
                    ephemeral=True,
                )
            except Exception:
                pass


class ConfirmacaoFecharMesaView(View):
    def __init__(self, dados_mesa=None):
        super().__init__(timeout=120)
        self.dados_mesa = dados_mesa or {}

    @discord.ui.button(
        label="Confirmar Encerramento",
        emoji="✅",
        style=discord.ButtonStyle.danger,
        custom_id="dic_confirmar_fechamento_mesa",
    )
    async def confirmar(self, interaction: discord.Interaction, button: Button):
        # Defer imediato para impedir "Esta interação falhou" durante a geração do dossiê.
        try:
            if not interaction.response.is_done():
                await interaction.response.defer(ephemeral=True, thinking=True)
        except Exception:
            pass

        try:
            for item in self.children:
                item.disabled = True
            if interaction.message:
                await interaction.message.edit(
                    content="⏳ Encerramento confirmado. Coletando dados e gerando o Dossiê Operacional...",
                    view=self,
                )
        except Exception:
            pass

        try:
            await fechar_mesa_core(
                interaction,
                motivo="Fechada por botão",
                dados_confirmacao=self.dados_mesa,
            )
        except Exception as erro:
            await enviar_log(f"❌ Erro crítico ao fechar mesa pelo botão: {erro}")
            try:
                await interaction.followup.send(
                    "❌ Deu erro ao fechar a mesa. Use `/fecharmesa` neste canal e veja os logs do Railway.",
                    ephemeral=True,
                )
            except Exception:
                pass

    @discord.ui.button(
        label="Cancelar",
        emoji="❌",
        style=discord.ButtonStyle.secondary,
        custom_id="dic_cancelar_fechamento_mesa",
    )
    async def cancelar(self, interaction: discord.Interaction, button: Button):
        try:
            await interaction.response.edit_message(
                content="❌ Encerramento cancelado. A mesa continua aberta.",
                view=None,
            )
        except Exception:
            try:
                await interaction.response.send_message("❌ Encerramento cancelado.", ephemeral=True)
            except Exception:
                pass

    async def on_error(self, interaction: discord.Interaction, error: Exception, item) -> None:
        await enviar_log(f"❌ Erro em botão de confirmação de fechamento: {error}")
        try:
            if interaction.response.is_done():
                await interaction.followup.send("❌ O botão deu erro. Use `/fecharmesa` neste canal.", ephemeral=True)
            else:
                await interaction.response.send_message("❌ O botão deu erro. Use `/fecharmesa` neste canal.", ephemeral=True)
        except Exception:
            pass


class PainelMesasView(View):
    def __init__(self):
        super().__init__(timeout=None)

    @discord.ui.button(label="Criar Mesa", emoji="➕", style=discord.ButtonStyle.green, custom_id="dic_criar_mesa")
    async def criar(self, interaction: discord.Interaction, button: Button):
        await interaction.response.send_modal(CriarMesaModal())

    @discord.ui.button(label="🔒 Fechar Mesa", emoji="🔒", style=discord.ButtonStyle.red, custom_id="dic_fechar_mesa")
    async def fechar_mesa(self, interaction: discord.Interaction, button: discord.ui.Button):
        mesa = buscar_mesa_por_canal(interaction.channel.id) if interaction.channel else None
        if not mesa:
            await interaction.response.send_message(
                "⚠️ Para fechar uma mesa, use este botão dentro do canal da mesa de investigação ou use `/fecharmesa` no canal correto.",
                ephemeral=True,
            )
            return
        if not usuario_pode_fechar_mesa(interaction.user):
            await interaction.response.send_message(mensagem_sem_permissao_fechar_mesa(), ephemeral=True)
            return
        dados = {
            "nome": f"OPERAÇÃO {mesa.get('familia', 'PADRÃO').upper()}",
            "comunidade": mesa.get('familia', 'Mapeada em logs'),
            "delegado": mesa.get('autor_nome', 'Superintendência'),
            "data_abertura": mesa.get('criada_em', agora_br()),
            "processo": f"PF-DICOR-{interaction.channel.id}",
            "numero": f"INV-{str(interaction.channel.id)[-6:]}",
        }
        await interaction.response.send_message(
            "⚠️ **Confirmação DICOR:** deseja encerrar esta mesa e gerar o Dossiê Operacional?",
            view=ConfirmacaoFecharMesaView(dados),
            ephemeral=True,
        )

    @discord.ui.button(label="Reabrir Mesa", emoji="🔓", style=discord.ButtonStyle.blurple, custom_id="dic_reabrir_mesa")
    async def reabrir(self, interaction: discord.Interaction, button: Button):
        await interaction.response.send_message("Use `/reabrirmesa canal:#canal` para escolher a mesa que será reaberta.", ephemeral=True)

    @discord.ui.button(label="Histórico", emoji="📂", style=discord.ButtonStyle.gray, custom_id="dic_historico_mesas")
    async def historico(self, interaction: discord.Interaction, button: Button):
        await historico_mesa_core(interaction)

    @discord.ui.button(label="Backup", emoji="📦", style=discord.ButtonStyle.gray, custom_id="dic_backup_mesas")
    async def backup(self, interaction: discord.Interaction, button: Button):
        await backup_core(interaction)

    @discord.ui.button(label="Estatísticas", emoji="📊", style=discord.ButtonStyle.gray, custom_id="dic_stats_mesas")
    async def stats(self, interaction: discord.Interaction, button: Button):
        await estatisticas_core(interaction)


    async def on_error(self, interaction: discord.Interaction, error: Exception, item) -> None:
        await enviar_log(f"❌ Erro em botão do painel de mesas: {error}")
        try:
            if interaction.response.is_done():
                await interaction.followup.send("❌ O botão deu erro. Tente usar o comando direto ou veja os logs.", ephemeral=True)
            else:
                await interaction.response.send_message("❌ O botão deu erro. Tente usar o comando direto ou veja os logs.", ephemeral=True)
        except Exception:
            pass


async def responder_interacao(interaction: discord.Interaction, mensagem: str, ephemeral: bool = True):
    if interaction.response.is_done():
        await interaction.followup.send(mensagem, ephemeral=ephemeral)
    else:
        await interaction.response.send_message(mensagem, ephemeral=ephemeral)


async def criar_mesa_core(interaction: discord.Interaction, apelido: str, familia: str, observacao: str = ""):
    guild = interaction.guild
    if guild is None:
        await responder_interacao(interaction, "❌ Use dentro de um servidor.", ephemeral=True)
        return

    if not interaction.response.is_done():
        await interaction.response.defer(ephemeral=True, thinking=True)

    categoria = guild.get_channel(CATEGORIA_MESAS_ABERTAS_ID) if CATEGORIA_MESAS_ABERTAS_ID else None
    overwrites = cargos_equipe_permissoes(guild)
    overwrites[interaction.user] = discord.PermissionOverwrite(
        view_channel=True,
        send_messages=True,
        attach_files=True,
        read_message_history=True,
        send_messages_in_threads=True,
    )

    nome_canal = f"🕵️┃{slugify(apelido)}-{slugify(familia)}"
    canal = await guild.create_text_channel(
        name=nome_canal,
        category=categoria,
        overwrites=overwrites,
    )

    await canal.send(
        f"🕵️ **Mesa de Investigação — {familia}**\n"
        f"👮 **Agente:** {interaction.user.mention}\n"
        f"📛 **Apelido:** {apelido}\n"
        f"🏷️ **Organização/Família:** {familia}\n"
        f"🕒 **Criada em:** {agora_br()}\n"
        f"📝 **Observação:** {observacao or 'Nenhuma'}\n\n"
        f"━━━━━━━━━━━━━━━━━━━━━━━\n"
        f"🧵 **Use os tópicos abertos abaixo para organizar a investigação.**"
    )

    topicos_ids = await criar_topicos_reais_mesa(canal)
    await canal.send(
        "🔒 Para encerrar esta mesa, clique no botão abaixo.",
        view=FecharMesaView()
    )

    registrar_mesa({
        "canal_id": canal.id,
        "nome_canal": canal.name,
        "apelido": apelido,
        "familia": familia,
        "autor_id": interaction.user.id,
        "autor_nome": str(interaction.user),
        "status": "ABERTA",
        "criada_em": agora_br(),
        "fechada_em": None,
        "topicos_ids": topicos_ids,
    })

    await enviar_log(
        f"➕ Mesa criada: {canal.mention} | Família: {familia} | Por: {interaction.user.mention}"
    )
    await responder_interacao(interaction, f"✅ Mesa criada: {canal.mention}", ephemeral=True)

# =====================================================
# DOSSIÊ OPERACIONAL AUTOMÁTICO DICOR
# =====================================================

URL_REGEX = re.compile(r"https?://[^\s<>()]+", flags=re.I)


def carregar_dossies() -> List[Dict[str, Any]]:
    dados = carregar_json(DOSSIES_JSON, [])
    return dados if isinstance(dados, list) else []


def salvar_dossies(lista: List[Dict[str, Any]]) -> None:
    salvar_json(DOSSIES_JSON, lista)


def registrar_dossie_operacional(registro: Dict[str, Any]) -> None:
    lista = carregar_dossies()
    lista.append(registro)
    salvar_dossies(lista[-500:])


def formatar_data_discord(dt: Optional[datetime.datetime]) -> str:
    try:
        if dt is None:
            return agora_br()
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=datetime.timezone.utc)
        tz = datetime.timezone(datetime.timedelta(hours=-3))
        return dt.astimezone(tz).strftime("%d/%m/%Y %H:%M")
    except Exception:
        return agora_br()


def texto_limpo_dossie(texto: Any, limite: int = 2500) -> str:
    valor = str(texto or "").replace("\x00", " ").strip()
    valor = re.sub(r"[ \t]+", " ", valor)
    valor = re.sub(r"\n{4,}", "\n\n", valor)
    if limite and len(valor) > limite:
        valor = valor[:limite - 20].rstrip() + "..."
    return valor


def normalizar_busca(texto: Any) -> str:
    valor = str(texto or "").lower()
    valor = unicodedata.normalize("NFKD", valor).encode("ascii", "ignore").decode("ascii")
    return valor


DOSSIE_CIDADE_OPERACIONAL = os.getenv("DOSSIE_CIDADE_OPERACIONAL", "Capital Morada do Valley").strip() or "Capital Morada do Valley"
DOSSIE_INSTITUICAO_CABECALHO = os.getenv(
    "DOSSIE_INSTITUICAO_CABECALHO",
    "POLÍCIA FEDERAL DE CAPITAL MORADA DO VALLEY\nDIRETORIA DE INVESTIGAÇÃO E COMBATE AO CRIME ORGANIZADO\nCENTRO DE INTELIGÊNCIA OPERACIONAL - DICOR",
).strip()


def nome_operacional_dossie(usuario: Any) -> str:
    """Extrai o nome operacional do Discord, ex: '[INSP.DIC] Baiano | 6027' -> 'Baiano'."""
    bruto = ""
    try:
        bruto = str(getattr(usuario, "display_name", "") or getattr(usuario, "name", "") or usuario or "")
    except Exception:
        bruto = str(usuario or "")

    texto = re.sub(r"<@!?\d+>", "", bruto).strip()
    texto = texto.replace("＠", "@").strip()

    # Remove prefixos de cargo em colchetes/parênteses, ficando só o que vem depois.
    for sep in ("]", ")"):
        if sep in texto:
            texto = texto.split(sep, 1)[1].strip()

    # Mantém somente o que está antes da barra vertical.
    if "|" in texto:
        texto = texto.split("|", 1)[0].strip()

    # Limpeza final de marcações comuns.
    texto = re.sub(r"^[\s@•\-–—:]+", "", texto).strip()
    texto = re.sub(r"\s+", " ", texto).strip()
    return texto or bruto or "Não informado"


def integrantes_mesa_dossie(mesa: Optional[Dict[str, Any]], autores: Dict[int, str], interaction: discord.Interaction) -> List[str]:
    """No dossiê, o campo Integrantes da Investigação usa o apelido cadastrado na criação da mesa."""
    apelido = ""
    if mesa and isinstance(mesa, dict):
        apelido = str(mesa.get("apelido") or "").strip()
    if apelido:
        return [apelido]
    if autores:
        return [nome_operacional_dossie(nome) for nome in autores.values()]
    return [nome_operacional_dossie(interaction.user)]


def nome_arquivo_seguro(nome: str, limite: int = 90) -> str:
    base = slugify(Path(nome or "arquivo").stem)[:limite] or "arquivo"
    ext = Path(nome or "").suffix.lower()
    if not ext or len(ext) > 10:
        ext = ".bin"
    return base + ext


def tipo_anexo_dossie(anexo: discord.Attachment) -> str:
    content_type = (anexo.content_type or "").lower()
    ext = Path(anexo.filename or "").suffix.lower()
    if content_type.startswith("image/") or ext in [".png", ".jpg", ".jpeg", ".webp", ".gif"]:
        return "imagem"
    if content_type.startswith("video/") or ext in [".mp4", ".mov", ".avi", ".mkv", ".webm"]:
        return "video"
    return "arquivo"


def topico_dossie_por_nome(nome: str) -> str:
    n = normalizar_busca(nome)
    if "lider" in n:
        return "liderancas"
    if "membro" in n or "integrante" in n:
        return "integrantes"
    if "painel" in n:
        return "painel"
    if "local" in n or "base" in n or "mapa" in n:
        return "localizacao"
    if "farm" in n or "produc" in n or "fabric" in n or "ingrediente" in n:
        return "producao"
    if "bau" in n or "armazen" in n:
        return "baus"
    if "informante" in n:
        return "informantes"
    if "crime" in n:
        return "crimes"
    if "radio" in n:
        return "radio"
    if "chat" in n:
        return "chat"
    return "geral"


def coletar_texto_embed(msg: discord.Message) -> str:
    partes = [msg.content or ""]
    for embed in msg.embeds:
        if embed.title:
            partes.append(embed.title)
        if embed.description:
            partes.append(embed.description)
        for field in embed.fields:
            partes.append(f"{field.name}: {field.value}")
    return texto_limpo_dossie("\n".join([p for p in partes if p]), 5000)


def texto_dossie_eh_boilerplate_bot(texto: str) -> bool:
    """Ignora mensagens automáticas da própria mesa para mesa vazia não virar dossiê pesado."""
    if not DOSSIE_SKIP_BOT_BOILERPLATE:
        return False

    n = normalizar_busca(texto)
    padroes = [
        "topicos da investigacao",
        "abra o topico desejado",
        "envie aqui todas as informacoes",
        "para encerrar esta mesa",
        "mesa de investigacao",
        "use os topicos abertos",
        "encerramento confirmado",
        "coletando dados",
        "gerando o dossie operacional",
        "mesa encerrada com sucesso",
        "dossie operacional gerado",
        "investigacao arquivada",
    ]
    return any(p in n for p in padroes)


def texto_dossie_tem_conteudo_util(texto: str) -> bool:
    """Retorna True quando o texto parece ser dado real da investigação."""
    if not texto:
        return False
    if texto_dossie_eh_boilerplate_bot(texto):
        return False
    n = normalizar_busca(texto)
    if len(n.strip()) < 4:
        return False
    return True


async def editar_progresso_dossie(mensagem: Optional[discord.Message], texto: str) -> None:
    if mensagem is None:
        return
    try:
        await mensagem.edit(content=texto[:1900])
    except Exception:
        pass


async def enviar_arquivos_dossie_destino(
    destino,
    dados_dossie: Dict[str, Any],
    arquivos: Dict[str, str],
    nome_pdf: str,
    nome_docx: str,
    canal_mesa: discord.TextChannel,
    usuario: discord.abc.User,
    titulo: str = "🏛️ DOSSIÊ OPERACIONAL AUTOMÁTICO DICOR",
) -> Optional[discord.Message]:
    if not destino or not arquivos or not hasattr(destino, "send"):
        return None

    embed_oficial = discord.Embed(
        title=titulo,
        description=(
            f"**Processo:** `{dados_dossie.get('processo')}`\n"
            f"**Investigação:** `{dados_dossie.get('numero_investigacao')}`\n"
            f"**Operação:** {dados_dossie.get('nome_operacao')}\n"
            f"**Mesa:** {canal_mesa.mention}\n"
            f"**Encerrada por:** {usuario.mention if hasattr(usuario, 'mention') else usuario}"
        ),
        color=discord.Color.from_rgb(0, 43, 91),
    )
    embed_oficial.add_field(
        name="Conteúdo",
        value=(
            f"Mensagens: `{dados_dossie.get('estatisticas', {}).get('mensagens_analisadas', 0)}` • "
            f"Evidências: `{dados_dossie.get('estatisticas', {}).get('evidencias', 0)}` • "
            f"Links: `{dados_dossie.get('estatisticas', {}).get('links', 0)}`"
        ),
        inline=False,
    )
    embed_oficial.set_footer(text="Polícia Federal - DICOR • Dossiê gerado automaticamente")

    itens: List[tuple[str, str]] = []
    if "pdf" in arquivos and Path(arquivos["pdf"]).exists():
        itens.append((arquivos["pdf"], nome_pdf))
    if "docx" in arquivos and Path(arquivos["docx"]).exists():
        itens.append((arquivos["docx"], nome_docx))

    if not itens:
        return None

    # Primeiro tenta enviar tudo junto. Se o Discord retornar 413/Payload Too Large,
    # envia um arquivo por mensagem. Isso evita perder o dossiê quando PDF+DOCX ficam grandes.
    try:
        arquivos_discord = [discord.File(caminho, filename=nome) for caminho, nome in itens]
        return await destino.send(embed=embed_oficial, files=arquivos_discord)
    except discord.HTTPException as erro:
        await enviar_log(f"⚠️ Envio conjunto do dossiê falhou. Tentando arquivos separados: {erro}")

    primeira_msg: Optional[discord.Message] = None
    for indice, (caminho, nome) in enumerate(itens, start=1):
        try:
            tamanho = Path(caminho).stat().st_size
            if tamanho > 24 * 1024 * 1024:
                aviso = (
                    f"⚠️ `{nome}` ficou grande demais para envio direto pelo Discord "
                    f"({tamanho / (1024 * 1024):.1f} MB). O arquivo foi preservado no armazenamento interno do bot."
                )
                msg = await destino.send(embed=embed_oficial if primeira_msg is None else None, content=aviso)
                primeira_msg = primeira_msg or msg
                continue
            msg = await destino.send(
                embed=embed_oficial if primeira_msg is None else None,
                content=f"📄 **Dossiê Operacional — arquivo {indice}/{len(itens)}**",
                file=discord.File(caminho, filename=nome),
            )
            primeira_msg = primeira_msg or msg
        except Exception as erro_individual:
            await enviar_log(f"❌ Falha ao enviar `{nome}` individualmente: {erro_individual}")
    return primeira_msg


async def bloquear_mesa_para_novas_mensagens(canal: discord.TextChannel) -> None:
    guild = canal.guild
    try:
        overwrite_default = canal.overwrites_for(guild.default_role)
        overwrite_default.send_messages = False
        overwrite_default.create_public_threads = False
        overwrite_default.create_private_threads = False
        overwrite_default.send_messages_in_threads = False
        await canal.set_permissions(guild.default_role, overwrite=overwrite_default)
    except Exception as erro:
        await enviar_log(f"⚠️ Não consegui bloquear @everyone na mesa {canal.id}: {erro}")

    # Remove envio dos usuários/cargos que já tinham overwrite no canal, mantendo o bot liberado.
    try:
        for alvo, overwrite in list(canal.overwrites.items()):
            if guild.me and alvo == guild.me:
                continue
            overwrite.send_messages = False
            overwrite.create_public_threads = False
            overwrite.create_private_threads = False
            overwrite.send_messages_in_threads = False
            await canal.set_permissions(alvo, overwrite=overwrite)
            await asyncio.sleep(0.05)
    except Exception as erro:
        await enviar_log(f"⚠️ Bloqueio parcial de permissões na mesa {canal.id}: {erro}")

    if guild.me:
        try:
            await canal.set_permissions(
                guild.me,
                view_channel=True,
                send_messages=True,
                manage_channels=True,
                manage_threads=True,
                attach_files=True,
                read_message_history=True,
            )
        except Exception:
            pass


async def listar_threads_da_mesa(canal: discord.TextChannel, mesa: Optional[Dict[str, Any]]) -> List[discord.Thread]:
    threads: Dict[int, discord.Thread] = {}

    for thread in getattr(canal, "threads", []):
        threads[thread.id] = thread

    if mesa and isinstance(mesa, dict):
        for tid in mesa.get("topicos_ids", []) or []:
            try:
                tid_int = int(tid)
            except Exception:
                continue
            thread = canal.guild.get_thread(tid_int) or bot.get_channel(tid_int)
            if thread is None:
                try:
                    thread = await bot.fetch_channel(tid_int)
                except Exception:
                    thread = None
            if isinstance(thread, discord.Thread):
                threads[thread.id] = thread

    try:
        async for thread in canal.archived_threads(limit=100, private=False):
            threads[thread.id] = thread
    except Exception:
        pass

    try:
        async for thread in canal.archived_threads(limit=100, private=True):
            threads[thread.id] = thread
    except Exception:
        pass

    return list(threads.values())


async def travar_threads_mesa(threads: List[discord.Thread]) -> None:
    for thread in threads:
        try:
            if thread.archived:
                await thread.edit(archived=False)
            await thread.edit(locked=True, reason="Mesa encerrada e dossiê operacional gerado")
        except Exception:
            pass


async def salvar_anexo_dossie(anexo: discord.Attachment, pasta_evidencias: Path, indice: int) -> Optional[Path]:
    try:
        pasta_evidencias.mkdir(parents=True, exist_ok=True)
        nome = f"{indice:03d}-{nome_arquivo_seguro(anexo.filename)}"
        caminho = pasta_evidencias / nome
        contador = 2
        while caminho.exists():
            caminho = pasta_evidencias / f"{indice:03d}-{contador}-{nome_arquivo_seguro(anexo.filename)}"
            contador += 1
        await anexo.save(str(caminho))
        return caminho
    except Exception as erro:
        await enviar_log(f"⚠️ Falha ao salvar anexo `{getattr(anexo, 'filename', 'arquivo')}` no dossiê: {erro}")
        return None


def extrair_valor_por_rotulos(textos: List[str], rotulos: List[str], limite: int = 600) -> str:
    for texto in textos:
        if not texto:
            continue
        for rotulo in rotulos:
            padrao = rf"(?:^|\n)\s*(?:[*_`>\-•\s]*)(?:{re.escape(rotulo)})\s*[:\-–]\s*(.+)"
            m = re.search(padrao, texto, flags=re.I)
            if m:
                valor = texto_limpo_dossie(m.group(1), limite)
                if valor:
                    return valor
    return "Não informado"


def extrair_blocos_pessoas(textos: List[str], funcao_padrao: str = "Integrante") -> List[Dict[str, str]]:
    pessoas: List[Dict[str, str]] = []
    vistos = set()
    for texto in textos:
        linhas = [linha.strip(" •-*_") for linha in str(texto or "").splitlines() if linha.strip()]
        for i, linha in enumerate(linhas):
            m_nome = re.search(r"(?:nome|indiv[ií]duo)\s*[:\-–]\s*(.+)", linha, flags=re.I)
            if not m_nome:
                continue
            janela = "\n".join(linhas[i:i + 8])
            nome = texto_limpo_dossie(m_nome.group(1), 120)
            rg = extrair_valor_por_rotulos([janela], ["RG", "Passaporte"], 80)
            cargo = extrair_valor_por_rotulos([janela], ["Função", "Funcao", "Cargo", "Posição", "Posicao"], 120)
            periculosidade = extrair_valor_por_rotulos([janela], ["Grau de periculosidade", "Periculosidade"], 80)
            observacoes = extrair_valor_por_rotulos([janela], ["Observações", "Observacoes", "Obs"], 250)
            chave = (normalizar_busca(nome), normalizar_busca(rg))
            if not nome or chave in vistos:
                continue
            vistos.add(chave)
            pessoas.append({
                "nome": nome,
                "rg": rg,
                "cargo": cargo if cargo != "Não informado" else funcao_padrao,
                "funcao": cargo if cargo != "Não informado" else funcao_padrao,
                "periculosidade": periculosidade,
                "observacoes": observacoes,
                "foto": "",
            })
    return pessoas[:80]


def filtrar_evidencias_por_topico(evidencias: List[Dict[str, Any]], topicos: List[str], limite: int = 12) -> List[Dict[str, Any]]:
    topicos_norm = {normalizar_busca(t) for t in topicos}
    saida = []
    for ev in evidencias:
        chave = normalizar_busca(ev.get("topico", ""))
        titulo = normalizar_busca(ev.get("origem", ""))
        if chave in topicos_norm or any(t in titulo for t in topicos_norm):
            saida.append(ev)
    return saida[:limite]


def resumir_textos_topico(textos: List[str], limite: int = 1200) -> str:
    relevantes = []
    for texto in textos:
        t = texto_limpo_dossie(texto, 500)
        if not t:
            continue
        if "Tópicos da investigação" in t or "Envie aqui" in t or "Para encerrar esta mesa" in t:
            continue
        relevantes.append(t)
    if not relevantes:
        return "Nenhum registro textual específico foi encontrado neste tópico."
    return texto_limpo_dossie("\n\n".join(relevantes[:8]), limite)


def gerar_qr_code_dossie(valor: str, pasta: Path, nome: str = "qr_reabertura.png") -> Optional[Path]:
    if qrcode is None:
        return None
    try:
        qr = qrcode.QRCode(version=1, box_size=8, border=2)
        qr.add_data(valor)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")
        caminho = pasta / nome
        img.save(caminho)
        return caminho
    except Exception:
        return None


async def coletar_dados_operacionais_mesa(
    canal: discord.TextChannel,
    mesa: Optional[Dict[str, Any]],
    interaction: discord.Interaction,
    pasta_dossie: Path,
    dados_confirmacao: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    pasta_evidencias = pasta_dossie / "evidencias"
    pasta_evidencias.mkdir(parents=True, exist_ok=True)

    limite = DOSSIE_HISTORY_LIMIT if DOSSIE_HISTORY_LIMIT > 0 else None
    threads = await listar_threads_da_mesa(canal, mesa)

    mensagens: List[Dict[str, Any]] = []
    textos_por_topico: Dict[str, List[str]] = defaultdict(list)
    evidencias: List[Dict[str, Any]] = []
    links: List[Dict[str, str]] = []
    autores: Dict[int, str] = {}
    contador_anexo = 1

    async def processar_historico(entidade, origem: str):
        nonlocal contador_anexo
        topico = topico_dossie_por_nome(origem)
        try:
            async for msg in entidade.history(limit=limite, oldest_first=True):
                if msg.author and not getattr(msg.author, "bot", False):
                    autores[msg.author.id] = str(msg.author)

                texto_msg = coletar_texto_embed(msg)
                eh_bot = bool(getattr(msg.author, "bot", False))
                texto_util = texto_dossie_tem_conteudo_util(texto_msg)
                links_msg = [link.rstrip(".,)];") for link in URL_REGEX.findall(texto_msg)] if texto_msg else []

                # Mensagens automáticas do próprio bot não entram no relatório,
                # mas mensagens de usuários, links e anexos continuam sendo coletados.
                if texto_msg and (texto_util or (not eh_bot and not texto_dossie_eh_boilerplate_bot(texto_msg))):
                    textos_por_topico[topico].append(texto_msg)
                    mensagens.append({
                        "id": msg.id,
                        "autor": str(msg.author),
                        "autor_id": getattr(msg.author, "id", None),
                        "data": formatar_data_discord(msg.created_at),
                        "origem": origem,
                        "topico": topico,
                        "conteudo": texto_msg,
                        "url": getattr(msg, "jump_url", ""),
                    })

                for link in links_msg:
                    if texto_util or not eh_bot:
                        links.append({
                            "url": link,
                            "data": formatar_data_discord(msg.created_at),
                            "autor": str(msg.author),
                            "origem": origem,
                        })

                for anexo in msg.attachments:
                    tipo = tipo_anexo_dossie(anexo)
                    caminho_local = await salvar_anexo_dossie(anexo, pasta_evidencias, contador_anexo)
                    contador_anexo += 1
                    evidencias.append({
                        "tipo": tipo,
                        "arquivo": anexo.filename,
                        "content_type": anexo.content_type or "",
                        "url": anexo.url,
                        "local": str(caminho_local) if caminho_local else "",
                        "data": formatar_data_discord(msg.created_at),
                        "autor": str(msg.author),
                        "origem": origem,
                        "topico": topico,
                        "mensagem_url": getattr(msg, "jump_url", ""),
                    })
        except Exception as erro:
            await enviar_log(f"⚠️ Erro ao coletar histórico de `{origem}` na mesa {canal.id}: {erro}")

    await processar_historico(canal, "Canal principal")
    for thread in threads:
        await processar_historico(thread, thread.name)

    todos_textos = [m["conteudo"] for m in mensagens]
    dados_base = dados_confirmacao or {}

    processo = str(dados_base.get("processo") or f"PF-DICOR-{canal.id}").replace(" ", "-")
    investigacao = str(dados_base.get("numero") or f"INV-{str(canal.id)[-6:]}")
    nome_operacao = texto_limpo_dossie(
        dados_base.get("nome")
        or (f"OPERAÇÃO {str(mesa.get('familia', canal.name)).upper()}" if mesa else f"OPERAÇÃO {canal.name.upper()}"),
        150,
    )

    comunidade = extrair_valor_por_rotulos(
        todos_textos,
        ["Comunidade investigada", "Comunidade", "Local", "Localização", "Localizacao"],
        150,
    )
    if comunidade == "Não informado":
        comunidade = dados_base.get("comunidade") or (mesa.get("familia") if mesa else "Não informado")

    faccao = extrair_valor_por_rotulos(
        todos_textos,
        ["Facção investigada", "Faccao investigada", "Facção", "Faccao", "Organização", "Organizacao", "Família", "Familia"],
        150,
    )
    if faccao == "Não informado":
        faccao = mesa.get("familia", "Não informado") if mesa else "Não informado"

    objetivo = extrair_valor_por_rotulos(
        todos_textos,
        ["Objetivo da investigação", "Objetivo", "Finalidade"],
        700,
    )
    if objetivo == "Não informado":
        objetivo = "Consolidar informações de inteligência, identificar a estrutura criminosa investigada e registrar evidências coletadas na mesa operacional."

    prioridade = extrair_valor_por_rotulos(todos_textos, ["Prioridade", "Prioridade da investigação"], 80)
    if prioridade == "Não informado":
        prioridade = "ALTA"

    status = extrair_valor_por_rotulos(todos_textos, ["Status", "Status da investigação"], 80)
    if status == "Não informado":
        status = "ENCERRADA / ARQUIVADA"

    liderancas = extrair_blocos_pessoas(textos_por_topico.get("liderancas", []) + textos_por_topico.get("painel", []), "Liderança")
    integrantes = extrair_blocos_pessoas(textos_por_topico.get("integrantes", []) + textos_por_topico.get("chat", []), "Integrante")
    informantes = extrair_blocos_pessoas(textos_por_topico.get("informantes", []), "Informante")

    # Tenta vincular imagens às pessoas pela ordem em cada tópico.
    imgs_lideres = [ev for ev in evidencias if ev.get("tipo") == "imagem" and ev.get("topico") == "liderancas"]
    imgs_integrantes = [ev for ev in evidencias if ev.get("tipo") == "imagem" and ev.get("topico") == "integrantes"]
    imgs_informantes = [ev for ev in evidencias if ev.get("tipo") == "imagem" and ev.get("topico") == "informantes"]
    for pessoa, ev in zip(liderancas, imgs_lideres):
        pessoa["foto"] = ev.get("local", "")
    for pessoa, ev in zip(integrantes, imgs_integrantes):
        pessoa["foto"] = ev.get("local", "")
    for pessoa, ev in zip(informantes, imgs_informantes):
        pessoa["foto"] = ev.get("local", "")

    def extrair_resultado(rotulos: List[str]) -> str:
        return extrair_valor_por_rotulos(todos_textos, rotulos, 300)

    resultados = {
        "prisoes": extrair_resultado(["Prisões efetuadas", "Prisoes efetuadas", "Presos", "Prisões", "Prisoes"]),
        "procurados_capturados": extrair_resultado(["Procurados capturados", "Capturados"]),
        "veiculos": extrair_resultado(["Veículos apreendidos", "Veiculos apreendidos", "Veículos", "Veiculos"]),
        "materiais": extrair_resultado(["Materiais apreendidos", "Materiais"]),
        "dinheiro": extrair_resultado(["Dinheiro apreendido", "Dinheiro", "Valores"]),
        "armas": extrair_resultado(["Armas apreendidas", "Armas", "Armamentos"]),
        "drogas": extrair_resultado(["Drogas apreendidas", "Drogas", "Entorpecentes"]),
        "municoes": extrair_resultado(["Munições", "Municoes", "Munição", "Municao"]),
        "outros": extrair_resultado(["Outros itens", "Outros"]),
    }

    reabrir_url = f"https://discord.com/channels/{canal.guild.id}/{canal.id}"
    qr_path = gerar_qr_code_dossie(reabrir_url, pasta_dossie)

    return {
        "processo": processo,
        "numero_investigacao": investigacao,
        "nome_operacao": nome_operacao,
        "nome": nome_operacao,
        "canal_id": canal.id,
        "canal_nome": canal.name,
        "guild_id": canal.guild.id,
        "guild_nome": canal.guild.name,
        "data_abertura": (mesa or {}).get("criada_em") or dados_base.get("data_abertura") or "Não informado",
        "data_encerramento": agora_br(),
        "delegado_responsavel": nome_operacional_dossie(dados_base.get("delegado") or (mesa or {}).get("autor_nome") or "Superintendência DICOR"),
        "agente_encerramento": nome_operacional_dossie(interaction.user),
        "agente_encerramento_id": interaction.user.id,
        "integrantes_investigacao": integrantes_mesa_dossie(mesa, autores, interaction),
        "cidade_operacional": DOSSIE_CIDADE_OPERACIONAL,
        "cabecalho_institucional": DOSSIE_INSTITUICAO_CABECALHO,
        "comunidade": comunidade,
        "faccao": faccao,
        "objetivo": objetivo,
        "status": status,
        "prioridade": prioridade,
        "mensagens": mensagens,
        "textos_por_topico": dict(textos_por_topico),
        "evidencias": evidencias,
        "links": links,
        "liderancas": liderancas,
        "integrantes": integrantes,
        "informantes": informantes,
        "resultados": resultados,
        "resumos": {
            "painel": resumir_textos_topico(textos_por_topico.get("painel", [])),
            "localizacao": resumir_textos_topico(textos_por_topico.get("localizacao", [])),
            "producao": resumir_textos_topico(textos_por_topico.get("producao", [])),
            "baus": resumir_textos_topico(textos_por_topico.get("baus", [])),
            "informantes": resumir_textos_topico(textos_por_topico.get("informantes", [])),
            "crimes": resumir_textos_topico(textos_por_topico.get("crimes", [])),
            "radio": resumir_textos_topico(textos_por_topico.get("radio", [])),
            "chat": resumir_textos_topico(textos_por_topico.get("chat", []), 900),
        },
        "qr_reabertura": str(qr_path) if qr_path else "",
        "reabrir_url": reabrir_url,
        "estatisticas": {
            "mensagens_analisadas": len(mensagens),
            "evidencias": len(evidencias),
            "imagens": len([e for e in evidencias if e.get("tipo") == "imagem"]),
            "videos": len([e for e in evidencias if e.get("tipo") == "video"]),
            "links": len(links),
            "threads": len(threads),
        },
    }


async def atualizar_status_mesa_fechada(canal_id: int, dados_dossie: Dict[str, Any], arquivos: Dict[str, str]) -> None:
    try:
        mesas = carregar_mesas()
        for m in mesas:
            if int(m.get("canal_id", 0)) == int(canal_id):
                m["status"] = "FECHADA"
                m["fechada_em"] = dados_dossie.get("data_encerramento", agora_br())
                m["fechada_por_id"] = dados_dossie.get("agente_encerramento_id")
                m["fechada_por_nome"] = dados_dossie.get("agente_encerramento")
                m["dossie"] = {
                    "processo": dados_dossie.get("processo"),
                    "numero_investigacao": dados_dossie.get("numero_investigacao"),
                    "pdf": arquivos.get("pdf"),
                    "docx": arquivos.get("docx"),
                    "gerado_em": agora_br(),
                }
        salvar_mesas(mesas)
    except Exception as erro:
        await enviar_log(f"⚠️ Erro ao atualizar status da mesa {canal_id}: {erro}")


async def fechar_mesa_core(
    interaction: discord.Interaction,
    motivo: str = "Fechada",
    dados_confirmacao: Optional[Dict[str, Any]] = None,
):
    inicio_total = time.monotonic()
    canal = interaction.channel
    if not isinstance(canal, discord.TextChannel):
        return

    try:
        if not interaction.response.is_done():
            await interaction.response.defer(ephemeral=True, thinking=True)
    except Exception:
        pass

    guild = interaction.guild
    if guild is None:
        await responder_interacao(interaction, "❌ Use dentro de um servidor.", ephemeral=True)
        return

    if not usuario_pode_fechar_mesa(interaction.user):
        await responder_interacao(interaction, mensagem_sem_permissao_fechar_mesa(), ephemeral=True)
        return

    msg_aviso = None
    try:
        msg_aviso = await canal.send(
            "⏳ **[DICOR] Encerramento confirmado.**\n"
            "`Etapa 1/6` Bloqueando a mesa e preparando o Dossiê Operacional..."
        )
    except Exception:
        pass

    mesa = None
    try:
        mesa = buscar_mesa_por_canal(canal.id)
    except Exception as erro:
        await enviar_log(f"⚠️ Mesa não localizada no banco local durante fechamento: {erro}")

    await editar_progresso_dossie(
        msg_aviso,
        "🔒 **[DICOR] Etapa 1/6**\nBloqueando novas mensagens na mesa e protegendo os tópicos..."
    )
    try:
        await bloquear_mesa_para_novas_mensagens(canal)
    except Exception as erro:
        await enviar_log(f"⚠️ Falha parcial ao bloquear mesa {canal.id}: {erro}")

    processo_base = str((dados_confirmacao or {}).get("processo") or f"PF-DICOR-{canal.id}").replace(" ", "-")
    processo_limpo = slugify(processo_base).upper().replace("-", "_") or str(canal.id)
    pasta_dossie = DOSSIES_DIR / processo_limpo
    pasta_dossie.mkdir(parents=True, exist_ok=True)

    await editar_progresso_dossie(
        msg_aviso,
        "🔎 **[DICOR] Etapa 2/6**\nColetando mensagens, tópicos, imagens, vídeos, links e anexos.\n"
        "Mesas completas podem demorar mais, pois o relatório preserva todas as evidências."
    )

    t_coleta = time.monotonic()
    dados_dossie = await coletar_dados_operacionais_mesa(
        canal,
        mesa,
        interaction,
        pasta_dossie,
        dados_confirmacao=dados_confirmacao,
    )
    tempo_coleta = time.monotonic() - t_coleta

    try:
        threads = await listar_threads_da_mesa(canal, mesa)
        await travar_threads_mesa(threads)
    except Exception:
        pass

    nome_pdf = f"DOSSIE_OPERACIONAL_{processo_limpo}.pdf"
    nome_docx = f"DOSSIE_OPERACIONAL_{processo_limpo}.docx"
    caminho_pdf = pasta_dossie / nome_pdf
    caminho_docx = pasta_dossie / nome_docx
    arquivos: Dict[str, str] = {}

    stats = dados_dossie.get("estatisticas", {})
    mesa_vazia = (
        int(stats.get("evidencias", 0) or 0) == 0
        and int(stats.get("links", 0) or 0) == 0
        and int(stats.get("mensagens_analisadas", 0) or 0) == 0
    )

    await editar_progresso_dossie(
        msg_aviso,
        "📄 **[DICOR] Etapa 3/6**\nGerando o PDF profissional com capa, índice, páginas oficiais e conclusão.\n"
        f"Conteúdo coletado: `{stats.get('mensagens_analisadas', 0)}` mensagens • `{stats.get('evidencias', 0)}` evidências • `{stats.get('links', 0)}` links."
    )

    erros_geracao: List[str] = []
    t_pdf = time.monotonic()
    try:
        await asyncio.to_thread(gerar_pdf_dossie, dados_dossie, caminho_pdf)
        if caminho_pdf.exists():
            arquivos["pdf"] = str(caminho_pdf)
    except Exception as erro:
        erros_geracao.append(f"PDF: {erro}")
        await enviar_log(f"❌ Erro ao gerar PDF do dossiê da mesa {canal.id}: {erro}")
    tempo_pdf = time.monotonic() - t_pdf

    await editar_progresso_dossie(
        msg_aviso,
        "📝 **[DICOR] Etapa 4/6**\nGerando o DOCX editável do dossiê, com tabelas e seções organizadas..."
    )

    t_docx = time.monotonic()
    try:
        await asyncio.to_thread(gerar_docx_dossie, dados_dossie, caminho_docx)
        if caminho_docx.exists():
            arquivos["docx"] = str(caminho_docx)
    except Exception as erro:
        erros_geracao.append(f"DOCX: {erro}")
        await enviar_log(f"❌ Erro ao gerar DOCX do dossiê da mesa {canal.id}: {erro}")
    tempo_docx = time.monotonic() - t_docx

    await editar_progresso_dossie(
        msg_aviso,
        "📤 **[DICOR] Etapa 5/6**\nEnviando PDF e DOCX para o canal oficial de dossiês e deixando uma cópia na própria mesa..."
    )

    canal_dest = guild.get_channel(DOSSIE_CHANNEL_ID) or guild.get_channel(BACKUP_CHANNEL_ID)
    mensagem_dossie_url = ""
    erros_envio: List[str] = []

    try:
        if canal_dest and arquivos:
            msg_dossie = await enviar_arquivos_dossie_destino(
                canal_dest,
                dados_dossie,
                arquivos,
                nome_pdf,
                nome_docx,
                canal,
                interaction.user,
            )
            if msg_dossie:
                mensagem_dossie_url = msg_dossie.jump_url
        elif not canal_dest:
            erros_envio.append("Canal de dossiês não encontrado. Configure DOSSIE_CHANNEL_ID ou BACKUP_CHANNEL_ID.")
        elif not arquivos:
            erros_envio.append("Nenhum arquivo foi gerado para envio.")
    except Exception as erro:
        erros_envio.append(f"Envio canal oficial: {erro}")
        await enviar_log(f"❌ Não foi possível enviar dossiê no canal configurado `{DOSSIE_CHANNEL_ID}`: {erro}")

    # Cópia na própria mesa para não ficar parecendo que sumiu em outro canal.
    if DOSSIE_ENVIAR_NA_MESA and arquivos:
        try:
            if not canal_dest or getattr(canal_dest, "id", None) != canal.id:
                msg_mesa = await enviar_arquivos_dossie_destino(
                    canal,
                    dados_dossie,
                    arquivos,
                    nome_pdf,
                    nome_docx,
                    canal,
                    interaction.user,
                    titulo="📎 CÓPIA DO DOSSIÊ OPERACIONAL DICOR",
                )
                if not mensagem_dossie_url and msg_mesa:
                    mensagem_dossie_url = msg_mesa.jump_url
        except Exception as erro:
            erros_envio.append(f"Envio na mesa: {erro}")
            await enviar_log(f"⚠️ Dossiê gerado, mas não consegui enviar cópia na mesa {canal.id}: {erro}")

    await atualizar_status_mesa_fechada(canal.id, dados_dossie, arquivos)

    registrar_dossie_operacional({
        "processo": dados_dossie.get("processo"),
        "numero_investigacao": dados_dossie.get("numero_investigacao"),
        "nome_operacao": dados_dossie.get("nome_operacao"),
        "canal_id": canal.id,
        "canal_nome": canal.name,
        "guild_id": guild.id,
        "gerado_em": agora_br(),
        "encerrado_por": str(interaction.user),
        "encerrado_por_id": interaction.user.id,
        "pdf": arquivos.get("pdf"),
        "docx": arquivos.get("docx"),
        "mensagem_dossie_url": mensagem_dossie_url,
        "estatisticas": dados_dossie.get("estatisticas", {}),
        "tempos_segundos": {
            "coleta": round(tempo_coleta, 2),
            "pdf": round(tempo_pdf, 2),
            "docx": round(tempo_docx, 2),
            "total": round(time.monotonic() - inicio_total, 2),
        },
    })

    await editar_progresso_dossie(
        msg_aviso,
        "📚 **[DICOR] Etapa 6/6**\nArquivando a investigação, movendo a mesa e finalizando o registro oficial..."
    )

    try:
        categoria_fechada = guild.get_channel(CATEGORIA_MESAS_FECHADAS_ID) if CATEGORIA_MESAS_FECHADAS_ID else None
        novo_nome = canal.name if canal.name.startswith("🔒") else f"🔒-{canal.name}"
        if categoria_fechada:
            await canal.edit(category=categoria_fechada, name=novo_nome, reason="Mesa encerrada e arquivada pela DICOR")
        else:
            await canal.edit(name=novo_nome, reason="Mesa encerrada e arquivada pela DICOR")
    except Exception as erro:
        await enviar_log(f"⚠️ Não foi possível mover/renomear a mesa arquivada {canal.id}: {erro}")

    try:
        if msg_aviso:
            await msg_aviso.delete()
    except Exception:
        pass

    embed_sucesso = discord.Embed(
        title="🏛️ Polícia Federal - DICOR",
        description=(
            "✅ **Mesa encerrada com sucesso.**\n\n"
            "📄 **Dossiê Operacional gerado.**\n"
            "📁 **Arquivos salvos com sucesso.**\n"
            "📎 **Arquivos enviados na mesa e/ou no canal oficial.**\n"
            "📚 **Investigação arquivada.**\n"
            "🏛️ **Polícia Federal - DICOR.**"
        ),
        color=discord.Color.from_rgb(0, 43, 91),
    )
    embed_sucesso.add_field(name="Processo", value=f"`{dados_dossie.get('processo')}`", inline=True)
    embed_sucesso.add_field(name="Investigação", value=f"`{dados_dossie.get('numero_investigacao')}`", inline=True)
    embed_sucesso.add_field(name="Evidências coletadas", value=f"`{stats.get('evidencias', 0)}`", inline=True)
    embed_sucesso.add_field(
        name="Tempo de geração",
        value=(
            f"Coleta: `{tempo_coleta:.1f}s` • PDF: `{tempo_pdf:.1f}s` • "
            f"DOCX: `{tempo_docx:.1f}s` • Total: `{time.monotonic() - inicio_total:.1f}s`"
        ),
        inline=False,
    )
    if mesa_vazia:
        embed_sucesso.add_field(
            name="Observação",
            value="A mesa estava sem registros úteis de investigação. O dossiê foi gerado com estrutura oficial e campos não informados.",
            inline=False,
        )
    if mensagem_dossie_url:
        embed_sucesso.add_field(name="Arquivo oficial", value=f"[Abrir envio do dossiê]({mensagem_dossie_url})", inline=False)
    avisos = erros_geracao + erros_envio
    if avisos:
        embed_sucesso.add_field(name="Avisos", value="\n".join(avisos)[:900], inline=False)

    try:
        await canal.send(embed=embed_sucesso, view=ReabrirMesaView())
        if interaction.message:
            try:
                await interaction.message.delete()
            except Exception:
                pass
    except Exception:
        pass

    try:
        await enviar_log(
            f"✅ **Mesa encerrada e dossiê gerado**\n"
            f"Mesa: {canal.mention}\n"
            f"Processo: `{dados_dossie.get('processo')}`\n"
            f"Arquivos: PDF={'sim' if 'pdf' in arquivos else 'não'} | DOCX={'sim' if 'docx' in arquivos else 'não'}\n"
            f"Evidências: `{stats.get('evidencias', 0)}` | Mensagens: `{stats.get('mensagens_analisadas', 0)}`\n"
            f"Tempo total: `{time.monotonic() - inicio_total:.1f}s`\n"
            f"Canal dos dossiês: <#{DOSSIE_CHANNEL_ID}>"
        )
    except Exception:
        pass

    try:
        await interaction.followup.send(
            "✅ Mesa encerrada e Dossiê Operacional gerado com sucesso. Os arquivos foram enviados na mesa/canal oficial.",
            ephemeral=True,
        )
    except Exception:
        pass


async def reabrir_mesa_core(interaction: discord.Interaction, canal: Optional[discord.TextChannel] = None):
    """Reabre uma mesa fechada pelo botão ou pelo comando /reabrirmesa."""
    try:
        if not interaction.response.is_done():
            await interaction.response.defer(ephemeral=True, thinking=True)
    except Exception:
        pass

    guild = interaction.guild
    if guild is None:
        await responder_interacao(interaction, "❌ Use dentro de um servidor.", ephemeral=True)
        return

    if canal is None:
        canal = interaction.channel

    if not isinstance(canal, discord.TextChannel):
        await responder_interacao(interaction, "❌ Canal inválido para reabrir mesa.", ephemeral=True)
        return

    mesa = buscar_mesa_por_canal(canal.id)

    try:
        categoria_aberta = guild.get_channel(CATEGORIA_MESAS_ABERTAS_ID) if CATEGORIA_MESAS_ABERTAS_ID else None
        novo_nome = canal.name
        while novo_nome.startswith("🔒") or novo_nome.startswith("-"):
            novo_nome = novo_nome.replace("🔒", "", 1).lstrip("-").strip()
        novo_nome = novo_nome or canal.name.replace("🔒-", "")

        if categoria_aberta:
            await canal.edit(
                category=categoria_aberta,
                name=novo_nome,
                reason="Mesa reaberta pela DICOR",
            )
        else:
            await canal.edit(name=novo_nome, reason="Mesa reaberta pela DICOR")
    except Exception as erro:
        await enviar_log(f"⚠️ Não consegui mover/renomear a mesa reaberta {canal.id}: {erro}")

    try:
        # Restaura permissão de envio para equipe/admin, autor da mesa, usuário que reabriu e o bot.
        for cargo_id in set(CARGOS_EQUIPE_IDS + CARGOS_ADMIN_IDS):
            cargo = guild.get_role(cargo_id)
            if cargo:
                overwrite = canal.overwrites_for(cargo)
                overwrite.view_channel = True
                overwrite.send_messages = True
                overwrite.attach_files = True
                overwrite.read_message_history = True
                overwrite.create_public_threads = True
                overwrite.send_messages_in_threads = True
                await canal.set_permissions(cargo, overwrite=overwrite)
                await asyncio.sleep(0.05)

        if mesa and mesa.get("autor_id"):
            autor = guild.get_member(int(mesa.get("autor_id")))
            if autor:
                overwrite = canal.overwrites_for(autor)
                overwrite.view_channel = True
                overwrite.send_messages = True
                overwrite.attach_files = True
                overwrite.read_message_history = True
                overwrite.create_public_threads = True
                overwrite.send_messages_in_threads = True
                await canal.set_permissions(autor, overwrite=overwrite)

        if isinstance(interaction.user, discord.Member):
            overwrite = canal.overwrites_for(interaction.user)
            overwrite.view_channel = True
            overwrite.send_messages = True
            overwrite.attach_files = True
            overwrite.read_message_history = True
            overwrite.create_public_threads = True
            overwrite.send_messages_in_threads = True
            await canal.set_permissions(interaction.user, overwrite=overwrite)

        if guild.me:
            overwrite = canal.overwrites_for(guild.me)
            overwrite.view_channel = True
            overwrite.send_messages = True
            overwrite.manage_channels = True
            overwrite.manage_threads = True
            overwrite.attach_files = True
            overwrite.read_message_history = True
            overwrite.create_public_threads = True
            overwrite.send_messages_in_threads = True
            await canal.set_permissions(guild.me, overwrite=overwrite)
    except Exception as erro:
        await enviar_log(f"⚠️ Falha parcial ao restaurar permissões da mesa {canal.id}: {erro}")

    try:
        threads = await listar_threads_da_mesa(canal, mesa)
        for thread in threads:
            try:
                if thread.archived:
                    await thread.edit(archived=False)
                if thread.locked:
                    await thread.edit(locked=False, reason="Mesa reaberta pela DICOR")
            except Exception:
                pass
    except Exception:
        pass

    try:
        mesas = carregar_mesas()
        for item in mesas:
            if int(item.get("canal_id", 0)) == canal.id:
                item["status"] = "ABERTA"
                item["reaberta_em"] = agora_br()
                item["reaberta_por"] = str(interaction.user)
                item["nome_canal"] = canal.name
                break
        salvar_mesas(mesas)
    except Exception as erro:
        await enviar_log(f"⚠️ Não consegui atualizar banco da mesa reaberta {canal.id}: {erro}")

    try:
        await canal.send(
            f"🔓 **Mesa reaberta com sucesso.**\n"
            f"👮 **Reaberta por:** {interaction.user.mention}\n"
            f"🕒 **Data:** {agora_br()}"
        )
    except Exception:
        pass

    await enviar_log(f"🔓 Mesa reaberta: {canal.mention} | Por: {interaction.user.mention}")
    await responder_interacao(interaction, f"✅ Mesa reaberta com sucesso: {canal.mention}", ephemeral=True)


async def backup_core(interaction: discord.Interaction):
    guild = interaction.guild
    if guild is None:
        await interaction.response.send_message("❌ Use dentro de um servidor.", ephemeral=True)
        return
    await interaction.response.defer(ephemeral=True)
    arquivo = BACKUP_DIR / f"backup-{data_caso()}.json"
    dados = {
        "data": agora_br(),
        "servidor": guild.name,
        "procurados": carregar_procurados(),
        "mesas": carregar_mesas(),
        "organizacoes": carregar_organizacoes(),
        "historico_organizacoes": carregar_historico_organizacoes(),
    }
    salvar_json(arquivo, dados)
    canal_backup = bot.get_channel(BACKUP_CHANNEL_ID) if BACKUP_CHANNEL_ID else None
    if canal_backup:
        try:
            await canal_backup.send(f"📦 **Backup Executado**\nServidor: {guild.name}\nData: {agora_br()}", file=discord.File(str(arquivo)))
        except Exception:
            pass
    await enviar_log(f"📦 Backup executado: {arquivo.name}")
    await interaction.followup.send(f"✅ Backup criado: `{arquivo.name}`", ephemeral=True)


async def estatisticas_core(interaction: discord.Interaction):
    procurados = carregar_procurados()
    mesas = carregar_mesas()
    organizacoes = carregar_organizacoes()
    ativos = len([p for p in procurados if p.get("status", "A PROCURAR") == "A PROCURAR"])
    retirados = len([p for p in procurados if p.get("status") == "RETIRADO"])
    abertas = len([m for m in mesas if m.get("status") == "ABERTA"])
    fechadas = len([m for m in mesas if m.get("status") == "FECHADA"])
    organizacoes_editadas = len([
        org for org in organizacoes
        if int(org.get("versao", 1) or 1) > 1
    ])
    await interaction.response.send_message(
        f"📊 **Estatísticas DICOR**\n\n"
        f"🚨 Procurados totais: `{len(procurados)}`\n"
        f"🟢 Procurados ativos: `{ativos}`\n"
        f"🔴 Procurados retirados: `{retirados}`\n\n"
        f"🕵️ Mesas totais: `{len(mesas)}`\n"
        f"🟢 Mesas abertas: `{abertas}`\n"
        f"🔒 Mesas fechadas: `{fechadas}`\n\n"
        f"🏴 Organizações fixas: `{len(organizacoes)}`\n"
        f"✏️ Organizações já editadas: `{organizacoes_editadas}`\n\n"
        f"📄 Catálogo: {CATALOG_PUBLIC_URL}",
        ephemeral=True,
    )




# =====================================================
# BOLETINS DE OCORRENCIA
# =====================================================

def agora_datetime_br() -> datetime.datetime:
    tz = datetime.timezone(datetime.timedelta(hours=-3))
    return datetime.datetime.now(tz)


def data_atual_br() -> str:
    return agora_datetime_br().strftime("%d/%m/%Y")


def horario_atual_br() -> str:
    return agora_datetime_br().strftime("%Hh%M")


def carregar_boletins() -> List[Dict[str, Any]]:
    dados = carregar_json(BOLETINS_JSON, [])
    return dados if isinstance(dados, list) else []


def salvar_boletins(lista: List[Dict[str, Any]]) -> None:
    salvar_json(BOLETINS_JSON, lista)


def carregar_usuarios_rg() -> Dict[str, str]:
    dados = carregar_json(USUARIOS_RG_JSON, {})
    return dados if isinstance(dados, dict) else {}


def salvar_usuarios_rg(dados: Dict[str, str]) -> None:
    salvar_json(USUARIOS_RG_JSON, dados)


def salvar_boletins_pendentes() -> None:
    salvar_json(
        BOLETINS_PENDENTES_JSON,
        {str(canal_id): dados for canal_id, dados in boletins_pendentes.items()},
    )


def carregar_boletins_pendentes_memoria() -> None:
    boletins_pendentes.clear()
    dados = carregar_json(BOLETINS_PENDENTES_JSON, {})
    if not isinstance(dados, dict):
        return
    for canal_id, rascunho in dados.items():
        try:
            boletins_pendentes[int(canal_id)] = rascunho
        except (TypeError, ValueError):
            continue


def usuario_tem_equipe(member: discord.Member) -> bool:
    if member.guild_permissions.administrator:
        return True
    permitidos = set(CARGOS_ADMIN_IDS + CARGOS_EQUIPE_IDS)
    if not permitidos:
        return False
    return any(cargo.id in permitidos for cargo in member.roles)


def obter_rg_usuario(usuario_id: int) -> str:
    return str(carregar_usuarios_rg().get(str(usuario_id), "") or "").strip()


def vincular_rg_usuario(usuario_id: int, rg: str) -> None:
    dados = carregar_usuarios_rg()
    dados[str(usuario_id)] = str(rg).strip()
    salvar_usuarios_rg(dados)


def remover_rg_usuario(usuario_id: int) -> bool:
    dados = carregar_usuarios_rg()
    chave = str(usuario_id)
    if chave not in dados:
        return False
    dados.pop(chave, None)
    salvar_usuarios_rg(dados)
    return True


def gerar_numero_boletim() -> str:
    """Gera numeração global e crescente para todos os boletins.
    Exemplo: BO-DICOR-001, BO-DICOR-002, BO-DICOR-003.
    Não reinicia por dia nem por usuário.

    Também olha boletins já publicados para evitar repetir número caso o bot
    tenha usado o modelo antigo por data.
    """
    contador = carregar_json(BOLETINS_CONTADOR_JSON, {})
    if not isinstance(contador, dict):
        contador = {}

    ultimo_contador = int(contador.get("ultimo", 0) or 0)
    maior_existente = 0
    boletins = carregar_boletins()
    for boletim in boletins:
        numero_salvo = str(boletim.get("numero", "") or "")
        m = re.search(r"(\d+)$", numero_salvo)
        if m:
            try:
                maior_existente = max(maior_existente, int(m.group(1)))
            except Exception:
                pass

    # len(boletins) ajuda a não voltar para 001 quando existiam boletins antigos
    # com numeração diária repetida.
    ultimo = max(ultimo_contador, maior_existente, len(boletins)) + 1

    salvar_json(
        BOLETINS_CONTADOR_JSON,
        {
            "ultimo": ultimo,
            "modelo": "global_crescente",
            "atualizado_em": agora_br(),
        },
    )
    return f"BO-DICOR-{ultimo:03d}"

def numero_curto_boletim(numero: str) -> str:
    parte = str(numero or "").rsplit("-", 1)[-1]
    return parte.zfill(3) if parte.isdigit() else parte


def normalizar_numero_boletim_consulta(numero: str) -> List[str]:
    texto = str(numero or "").strip().upper()
    candidatos = {texto}
    m = re.search(r"(\d+)$", texto)
    if m:
        curto = m.group(1).zfill(3)
        candidatos.add(curto)
        candidatos.add(f"BO-DICOR-{curto}")
    return [c for c in candidatos if c]


def buscar_boletim_numero(numero: str) -> Optional[Dict[str, Any]]:
    candidatos = normalizar_numero_boletim_consulta(numero)
    for boletim in carregar_boletins():
        numero_salvo = str(boletim.get("numero", "")).strip().upper()
        if numero_salvo in candidatos or numero_curto_boletim(numero_salvo) in candidatos:
            return boletim
    return None

def obter_rascunho_boletim(
    interaction: discord.Interaction,
    exigir_dono: bool = True,
) -> Optional[Dict[str, Any]]:
    canal = interaction.channel
    if not isinstance(canal, discord.TextChannel):
        return None

    dados = boletins_pendentes.get(canal.id)
    if not dados:
        return None

    if exigir_dono and interaction.user.id != int(dados.get("comunicante_id", 0)):
        if not isinstance(interaction.user, discord.Member) or not usuario_tem_equipe(interaction.user):
            return None
    return dados


def dividir_texto_discord(texto: str, limite: int = 1900) -> List[str]:
    texto = str(texto or "").strip()
    if not texto:
        return [""]
    if len(texto) <= limite:
        return [texto]

    partes: List[str] = []
    atual = ""
    for bloco in texto.split("\n"):
        linha = bloco + "\n"
        if len(linha) > limite:
            if atual.strip():
                partes.append(atual.rstrip())
                atual = ""
            restante = linha
            while len(restante) > limite:
                corte = restante.rfind(" ", 0, limite)
                if corte < limite // 2:
                    corte = limite
                partes.append(restante[:corte].rstrip())
                restante = restante[corte:].lstrip()
            atual = restante
            continue

        if len(atual) + len(linha) > limite:
            partes.append(atual.rstrip())
            atual = linha
        else:
            atual += linha

    if atual.strip():
        partes.append(atual.rstrip())
    return partes


def resumir_relato_sem_inventar(relato: str, limite: int = 430) -> str:
    texto = re.sub(r"\s+", " ", str(relato or "")).strip()
    if not texto:
        return ""

    sentencas = re.split(r"(?<=[.!?])\s+", texto)
    selecionadas: List[str] = []
    total = 0
    for sentenca in sentencas:
        sentenca = sentenca.strip()
        if not sentenca:
            continue
        if selecionadas and total + len(sentenca) + 1 > limite:
            break
        if not selecionadas and len(sentenca) > limite:
            selecionadas.append(sentenca[:limite].rstrip(" ,;:") + "...")
            break
        selecionadas.append(sentenca)
        total += len(sentenca) + 1
        if len(selecionadas) >= 2:
            break

    return " ".join(selecionadas).strip()


def gerar_conclusao_automatica(
    relato: str,
    tipo_identificacao: str,
    dados_individuo: Optional[Dict[str, Any]] = None,
    dados_veiculo: Optional[Dict[str, Any]] = None,
) -> str:
    resumo = resumir_relato_sem_inventar(relato)
    tipo = str(tipo_identificacao or "").lower()

    if tipo == "veiculo":
        encerramento = (
            "Os dados do veículo e as provas coletadas foram registrados para "
            "análise e continuidade das diligências investigativas."
        )
    elif tipo == "individuo":
        encerramento = (
            "As informações do indivíduo e as provas coletadas foram registradas "
            "para continuidade das diligências investigativas."
        )
    else:
        encerramento = (
            "Os dados do indivíduo, do veículo e as provas coletadas foram "
            "registrados para análise e continuidade das investigações."
        )

    if resumo:
        return f"Conforme o relato, {resumo[0].lower() + resumo[1:] if len(resumo) > 1 else resumo.lower()} {encerramento}"
    return encerramento


def formatar_secao_individuo(dados: Dict[str, Any]) -> str:
    if not dados:
        return ""
    return (
        "👤 **INDIVÍDUO IDENTIFICADO**\n\n"
        f"**Nome:** {dados.get('nome') or 'Não identificado'}\n"
        f"**RG:** {dados.get('rg') or 'Não identificado'}\n"
        f"**Características:** {dados.get('caracteristicas') or 'Não identificado'}\n"
        f"**Participação no fato:** {dados.get('participacao') or 'Não identificado'}\n"
        f"**Outras informações:** {dados.get('outras_informacoes') or 'Não identificado'}"
    )


def formatar_secao_veiculo(dados: Dict[str, Any]) -> str:
    if not dados:
        return ""
    return (
        "🚘 **VEÍCULO IDENTIFICADO**\n\n"
        f"**Modelo:** {dados.get('modelo') or 'Não identificado'}\n"
        f"**Placa:** {dados.get('placa') or 'Não identificado'}\n"
        f"**Telefone:** {dados.get('telefone') or 'Não identificado'}\n"
        f"**Proprietário:** {dados.get('proprietario') or 'Não identificado'}\n"
        f"**RG do proprietário:** {dados.get('rg_proprietario') or 'Não identificado'}"
    )


def formatar_boletim(dados: Dict[str, Any], previa: bool = False) -> str:
    secoes: List[str] = []
    tipo = str(dados.get("tipo_identificacao", "") or "").lower()

    if tipo in {"individuo", "ambos"}:
        secao = formatar_secao_individuo(dados.get("dados_individuo", {}))
        if secao:
            secoes.append(secao)

    if tipo in {"veiculo", "ambos"}:
        secao = formatar_secao_veiculo(dados.get("dados_veiculo", {}))
        if secao:
            secoes.append(secao)

    identificacoes = "\n\n━━━━━━━━━━━━━━━━━━━━━━━\n\n".join(secoes)
    if identificacoes:
        identificacoes = (
            "\n\n━━━━━━━━━━━━━━━━━━━━━━━\n\n"
            + identificacoes
            + "\n\n━━━━━━━━━━━━━━━━━━━━━━━\n"
        )
    else:
        identificacoes = "\n\n━━━━━━━━━━━━━━━━━━━━━━━\n"

    provas_texto = (
        "As provas encontram-se anexadas abaixo."
        if not previa
        else "As provas enviadas neste canal serão anexadas na publicação."
    )

    numero = str(dados.get("numero", ""))
    comunicante = str(
        dados.get("comunicante_mention")
        or f"<@{dados.get('comunicante_id', 0)}>"
    )

    return (
        f"📋 **BOLETIM DE OCORRÊNCIA — {numero_curto_boletim(numero)}**\n\n"
        f"**Comunicante:** {comunicante}\n"
        f"**RG:** {dados.get('rg') or 'Não informado'}\n\n"
        f"**Data do fato:** {dados.get('data_fato') or data_atual_br()}\n"
        f"**Horário aproximado:** {dados.get('horario_fato') or horario_atual_br()}\n"
        f"**Local:** {dados.get('local') or 'Não informado'}\n\n"
        f"**RELATO:**\n\n"
        f"{dados.get('relato') or 'Não informado'}"
        f"{identificacoes}\n"
        f"📎 **PROVAS:**\n\n"
        f"{provas_texto}\n\n"
        f"**CONCLUSÃO:**\n\n"
        f"{dados.get('conclusao') or 'A conclusão ainda não foi gerada.'}\n\n"
        f"━━━━━━━━━━━━━━━━━━━━━━━\n\n"
        f"**Número do boletim:** {numero}\n"
        f"**Registrado por:** {comunicante}\n"
        f"**Data do registro:** {data_atual_br()} — {horario_atual_br()}\n"
        f"> 👮 Criado por: {comunicante}"
    )


async def enviar_previa_boletim(
    canal: discord.TextChannel,
    dados: Dict[str, Any],
) -> None:
    texto = "🔎 **PRÉVIA DO BOLETIM**\n\n" + formatar_boletim(dados, previa=True)
    partes = dividir_texto_discord(texto)
    mensagens_ids: List[int] = []
    for indice, parte in enumerate(partes):
        view = PreviaBoletimView() if indice == len(partes) - 1 else None
        mensagem = await canal.send(parte, view=view)
        mensagens_ids.append(mensagem.id)

    dados["preview_message_ids"] = mensagens_ids
    dados["etapa"] = "PREVIA"
    boletins_pendentes[canal.id] = dados
    salvar_boletins_pendentes()


async def enviar_etapa_provas(
    interaction: discord.Interaction,
    dados: Dict[str, Any],
) -> None:
    canal = interaction.channel
    if not isinstance(canal, discord.TextChannel):
        await responder_interacao(interaction, "❌ Canal inválido.", ephemeral=True)
        return

    dados["etapa"] = "PROVAS"
    boletins_pendentes[canal.id] = dados
    salvar_boletins_pendentes()

    mensagem = (
        "📎 **ETAPA DE PROVAS**\n\n"
        "Envie neste canal todas as fotos, vídeos, prints, documentos ou outros "
        "arquivos relacionados ao boletim.\n\n"
        "Quando terminar, clique em **Finalizar Provas**."
    )
    if interaction.response.is_done():
        await canal.send(mensagem, view=ProvasBoletimView())
    else:
        await interaction.response.send_message(
            mensagem,
            view=ProvasBoletimView(),
        )


async def depois_identificacao(
    interaction: discord.Interaction,
    dados: Dict[str, Any],
) -> None:
    canal = interaction.channel
    if not isinstance(canal, discord.TextChannel):
        await responder_interacao(interaction, "❌ Canal inválido.", ephemeral=True)
        return

    retorno = dados.pop("retorno_identificacao", "")
    dados["conclusao"] = gerar_conclusao_automatica(
        dados.get("relato", ""),
        dados.get("tipo_identificacao", ""),
        dados.get("dados_individuo", {}),
        dados.get("dados_veiculo", {}),
    )
    boletins_pendentes[canal.id] = dados
    salvar_boletins_pendentes()

    if retorno == "PREVIA":
        if not interaction.response.is_done():
            await interaction.response.defer(ephemeral=True)
        await enviar_previa_boletim(canal, dados)
        await interaction.followup.send("✅ Identificação atualizada.", ephemeral=True)
        return

    await enviar_etapa_provas(interaction, dados)


async def cancelar_boletim_core(interaction: discord.Interaction) -> None:
    canal = interaction.channel
    if not isinstance(canal, discord.TextChannel):
        await responder_interacao(interaction, "❌ Canal inválido.", ephemeral=True)
        return

    dados = obter_rascunho_boletim(interaction)
    if not dados:
        await responder_interacao(
            interaction,
            "❌ Você não pode cancelar este boletim ou ele não foi encontrado.",
            ephemeral=True,
        )
        return

    boletins_pendentes.pop(canal.id, None)
    salvar_boletins_pendentes()
    await responder_interacao(
        interaction,
        "❌ Boletim cancelado. O canal será apagado em 3 segundos.",
        ephemeral=True,
    )
    await asyncio.sleep(3)
    try:
        await canal.delete(reason="Boletim cancelado")
    except Exception:
        pass


async def coletar_anexos_boletim(
    canal: discord.TextChannel,
) -> List[discord.Attachment]:
    anexos: List[discord.Attachment] = []
    async for mensagem in canal.history(limit=None, oldest_first=True):
        if mensagem.author.bot:
            continue
        anexos.extend(mensagem.attachments)
    return anexos


async def enviar_anexos_boletim(
    canal_oficial: discord.abc.Messageable,
    anexos: List[discord.Attachment],
    numero: str,
) -> List[Dict[str, Any]]:
    resultados: List[Dict[str, Any]] = []

    for indice in range(0, len(anexos), 5):
        lote = anexos[indice:indice + 5]
        arquivos: List[discord.File] = []
        originais: List[discord.Attachment] = []
        for anexo in lote:
            try:
                arquivos.append(await anexo.to_file())
                originais.append(anexo)
            except Exception as erro:
                await enviar_log(
                    f"⚠️ Não foi possível preparar o anexo `{anexo.filename}` "
                    f"do boletim {numero}: {erro}"
                )

        if not arquivos:
            continue

        try:
            mensagem = await canal_oficial.send(
                content=f"📎 **Provas — {numero}**",
                files=arquivos,
            )
            for anexo_publicado in mensagem.attachments:
                resultados.append({
                    "id": anexo_publicado.id,
                    "nome": anexo_publicado.filename,
                    "url": anexo_publicado.url,
                    "mensagem_id": mensagem.id,
                    "mensagem_url": mensagem.jump_url,
                })
        except Exception as erro:
            await enviar_log(
                f"⚠️ Falha ao publicar um lote de provas do boletim {numero}: {erro}"
            )
            # Tenta novamente um arquivo por mensagem.
            for original in originais:
                try:
                    arquivo = await original.to_file()
                    mensagem = await canal_oficial.send(
                        content=f"📎 **Prova — {numero}**",
                        file=arquivo,
                    )
                    for anexo_publicado in mensagem.attachments:
                        resultados.append({
                            "id": anexo_publicado.id,
                            "nome": anexo_publicado.filename,
                            "url": anexo_publicado.url,
                            "mensagem_id": mensagem.id,
                            "mensagem_url": mensagem.jump_url,
                        })
                except Exception as erro_individual:
                    await enviar_log(
                        f"⚠️ Falha ao publicar `{original.filename}` no boletim "
                        f"{numero}: {erro_individual}"
                    )

    return resultados


async def publicar_boletim_core(interaction: discord.Interaction) -> None:
    canal_temp = interaction.channel
    if not isinstance(canal_temp, discord.TextChannel):
        await responder_interacao(interaction, "❌ Canal inválido.", ephemeral=True)
        return

    dados = obter_rascunho_boletim(interaction)
    if not dados:
        await responder_interacao(
            interaction,
            "❌ Boletim não encontrado ou você não tem permissão.",
            ephemeral=True,
        )
        return

    if dados.get("publicando"):
        await responder_interacao(
            interaction,
            "⏳ Este boletim já está sendo publicado.",
            ephemeral=True,
        )
        return

    if buscar_boletim_numero(str(dados.get("numero", ""))):
        await responder_interacao(
            interaction,
            "⚠️ Este boletim já foi publicado.",
            ephemeral=True,
        )
        return

    dados["publicando"] = True
    boletins_pendentes[canal_temp.id] = dados
    salvar_boletins_pendentes()
    await interaction.response.defer(ephemeral=True, thinking=True)

    try:
        canal_oficial = bot.get_channel(BOLETINS_CHANNEL_ID)
        if canal_oficial is None:
            canal_oficial = await bot.fetch_channel(BOLETINS_CHANNEL_ID)
        if not isinstance(canal_oficial, discord.TextChannel):
            raise RuntimeError("O canal oficial de boletins não foi encontrado.")

        anexos_origem = await coletar_anexos_boletim(canal_temp)
        partes = dividir_texto_discord(formatar_boletim(dados, previa=False))
        mensagens_publicadas: List[discord.Message] = []
        anexos_publicados: List[Dict[str, Any]] = []

        anexos_anexados_na_principal: set[int] = set()
        arquivos_principais: List[discord.File] = []
        tamanho_principal = 0
        for anexo in anexos_origem:
            if len(arquivos_principais) >= 10:
                break
            try:
                tamanho = int(getattr(anexo, "size", 0) or 0)
                if tamanho_principal + tamanho > 24 * 1024 * 1024:
                    continue
                arquivos_principais.append(await anexo.to_file())
                anexos_anexados_na_principal.add(anexo.id)
                tamanho_principal += tamanho
            except Exception as erro:
                await enviar_log(f"⚠️ Não consegui preparar anexo para a mensagem principal do boletim {dados.get('numero')}: {erro}")

        for indice_parte, parte in enumerate(partes):
            try:
                if indice_parte == 0 and arquivos_principais:
                    msg_publicada = await canal_oficial.send(
                        parte,
                        files=arquivos_principais,
                        allowed_mentions=discord.AllowedMentions(users=False, roles=False, everyone=False),
                    )
                    for anexo_publicado in msg_publicada.attachments:
                        anexos_publicados.append({
                            "id": anexo_publicado.id,
                            "nome": anexo_publicado.filename,
                            "url": anexo_publicado.url,
                            "mensagem_id": msg_publicada.id,
                            "mensagem_url": msg_publicada.jump_url,
                        })
                else:
                    msg_publicada = await canal_oficial.send(
                        parte,
                        allowed_mentions=discord.AllowedMentions(users=False, roles=False, everyone=False),
                    )
                mensagens_publicadas.append(msg_publicada)
            except discord.HTTPException as erro:
                await enviar_log(f"⚠️ Falha ao enviar boletim com anexos juntos. Tentando sem anexos: {erro}")
                msg_publicada = await canal_oficial.send(
                    parte,
                    allowed_mentions=discord.AllowedMentions(users=False, roles=False, everyone=False),
                )
                mensagens_publicadas.append(msg_publicada)
                anexos_anexados_na_principal.clear()

        mensagem_principal = mensagens_publicadas[0]
        anexos_restantes = [a for a in anexos_origem if a.id not in anexos_anexados_na_principal]
        if anexos_restantes:
            anexos_publicados.extend(await enviar_anexos_boletim(
                canal_oficial,
                anexos_restantes,
                str(dados.get("numero", "")),
            ))

        registro = dict(dados)
        registro.pop("preview_message_ids", None)
        registro.pop("publicando", None)
        registro["status"] = "PUBLICADO"
        registro["mensagem_id"] = mensagem_principal.id
        registro["mensagem_url"] = mensagem_principal.jump_url
        registro["mensagens_oficiais"] = [
            {"id": mensagem.id, "url": mensagem.jump_url}
            for mensagem in mensagens_publicadas
        ]
        registro["anexos"] = anexos_publicados
        registro["data_criacao"] = agora_br()
        registro["criado_por_id"] = registro.get("comunicante_id")
        registro["criado_por_nome"] = registro.get("comunicante_nome")
        registro["publicado_por_id"] = interaction.user.id
        registro["publicado_por_nome"] = str(interaction.user)
        registro["canal_provisorio_id"] = canal_temp.id

        boletins = carregar_boletins()
        if not any(
            str(item.get("numero", "")).upper()
            == str(registro.get("numero", "")).upper()
            for item in boletins
        ):
            boletins.append(registro)
            salvar_boletins(boletins)

        boletins_pendentes.pop(canal_temp.id, None)
        salvar_boletins_pendentes()

        await enviar_log(
            f"📋 Boletim publicado: {registro.get('numero')} | "
            f"Comunicante: <@{registro.get('comunicante_id')}> | "
            f"Mensagem: {mensagem_principal.jump_url}"
        )
        await interaction.followup.send(
            f"✅ Boletim publicado com sucesso: {mensagem_principal.jump_url}",
            ephemeral=True,
        )
        await canal_temp.send("✅ Boletim publicado com sucesso. Este canal será apagado em 5 segundos.")
        await asyncio.sleep(5)
        await canal_temp.delete(reason="Boletim publicado com sucesso")
    except Exception as erro:
        dados["publicando"] = False
        boletins_pendentes[canal_temp.id] = dados
        salvar_boletins_pendentes()
        await enviar_log(f"❌ Erro ao publicar boletim {dados.get('numero')}: {erro}")
        await interaction.followup.send(
            f"❌ Não foi possível publicar o boletim: `{erro}`\n"
            "O canal e o rascunho foram mantidos.",
            ephemeral=True,
        )


async def abrir_boletim_core(interaction: discord.Interaction) -> None:
    guild = interaction.guild
    if guild is None:
        await responder_interacao(interaction, "❌ Use dentro de um servidor.", ephemeral=True)
        return

    if not interaction.response.is_done():
        await interaction.response.defer(ephemeral=True, thinking=True)

    numero = gerar_numero_boletim()
    rg = obter_rg_usuario(interaction.user.id)

    categoria: Optional[discord.CategoryChannel] = None
    if BOLETIM_TEMP_CATEGORY_ID:
        categoria_encontrada = guild.get_channel(BOLETIM_TEMP_CATEGORY_ID)
        if isinstance(categoria_encontrada, discord.CategoryChannel):
            categoria = categoria_encontrada
    if categoria is None and isinstance(interaction.channel, discord.TextChannel):
        categoria = interaction.channel.category

    overwrites: Dict[Any, discord.PermissionOverwrite] = {
        guild.default_role: discord.PermissionOverwrite(view_channel=False),
        interaction.user: discord.PermissionOverwrite(
            view_channel=True,
            send_messages=True,
            attach_files=True,
            read_message_history=True,
        ),
    }
    if guild.me:
        overwrites[guild.me] = discord.PermissionOverwrite(
            view_channel=True,
            send_messages=True,
            attach_files=True,
            read_message_history=True,
            manage_channels=True,
        )
    for cargo_id in set(CARGOS_ADMIN_IDS + CARGOS_EQUIPE_IDS):
        cargo = guild.get_role(cargo_id)
        if cargo:
            overwrites[cargo] = discord.PermissionOverwrite(
                view_channel=True,
                send_messages=True,
                attach_files=True,
                read_message_history=True,
            )

    nome_usuario = slugify(
        getattr(interaction.user, "display_name", interaction.user.name)
    )[:35]
    canal = await guild.create_text_channel(
        name=f"📋・boletim-{nome_usuario}-{numero_curto_boletim(numero)}",
        category=categoria,
        overwrites=overwrites,
        reason=f"Canal provisório do boletim {numero}",
    )

    dados = {
        "numero": numero,
        "comunicante_id": interaction.user.id,
        "comunicante_mention": interaction.user.mention,
        "comunicante_nome": str(interaction.user),
        "rg": rg,
        "data_fato": data_atual_br(),
        "horario_fato": horario_atual_br(),
        "local": "",
        "relato": "",
        "tipo_identificacao": "",
        "dados_individuo": {},
        "dados_veiculo": {},
        "conclusao": "",
        "etapa": "RG" if not rg else "DADOS_INICIAIS",
        "status": "RASCUNHO",
        "criado_em": agora_br(),
        "canal_id": canal.id,
    }
    boletins_pendentes[canal.id] = dados
    salvar_boletins_pendentes()

    await canal.send(
        f"📋 **Boletim de Ocorrência — {numero}**\n\n"
        f"**Comunicante:** {interaction.user.mention}\n"
        f"**RG:** {rg or 'Ainda não vinculado'}\n"
        f"**Data:** {dados['data_fato']}\n"
        f"**Horário:** {dados['horario_fato']}\n\n"
        "O bot vai ajudar no preenchimento do boletim por etapas."
    )

    if rg:
        await canal.send(
            "Clique abaixo para informar o local e escrever o relato.",
            view=IniciarBoletimView(),
        )
    else:
        await canal.send(
            "Seu Discord ainda não possui RG vinculado. Informe o RG para continuar.",
            view=InformarRGBoletimView(),
        )

    await enviar_log(
        f"📁 Canal provisório do boletim {numero} criado por "
        f"{interaction.user.mention}: {canal.mention}"
    )
    await interaction.followup.send(
        f"✅ Canal provisório criado: {canal.mention}",
        ephemeral=True,
    )


def texto_consulta_boletim(boletim: Dict[str, Any]) -> str:
    comunicante = boletim.get("comunicante_mention") or f"<@{boletim.get('comunicante_id', 0)}>"
    return (
        f"📋 **{boletim.get('numero', 'Boletim')}**\n"
        f"**Comunicante:** {comunicante}\n"
        f"**Data:** {boletim.get('data_fato', 'Não informado')}\n"
        f"**Local:** {boletim.get('local', 'Não informado')}\n"
        f"**Status:** {boletim.get('status', 'Não informado')}\n"
        f"**Mensagem:** {boletim.get('mensagem_url', 'Sem link')}"
    )


class InformarRGBoletimModal(Modal, title="Informar RG"):
    rg = TextInput(
        label="RG do comunicante",
        placeholder="Ex: 6027",
        max_length=50,
    )

    async def on_submit(self, interaction: discord.Interaction):
        dados = obter_rascunho_boletim(interaction)
        if not dados:
            await interaction.response.send_message(
                "❌ Rascunho não encontrado ou sem permissão.",
                ephemeral=True,
            )
            return

        valor = str(self.rg.value).strip()
        dados["rg"] = valor
        dados["etapa"] = "DADOS_INICIAIS"
        vincular_rg_usuario(interaction.user.id, valor)
        boletins_pendentes[interaction.channel.id] = dados
        salvar_boletins_pendentes()

        await interaction.response.send_message(
            f"✅ RG `{valor}` vinculado. Agora preencha o local e o relato.",
            view=IniciarBoletimView(),
        )


class DadosIniciaisBoletimModal(Modal, title="Local e Relato"):
    local = TextInput(
        label="Local do fato",
        placeholder="Ex: Porto, Caixa d'água, Ilha...",
        max_length=200,
    )
    relato = TextInput(
        label="Relato completo",
        placeholder="Descreva toda a ocorrência...",
        style=discord.TextStyle.paragraph,
        max_length=4000,
    )

    async def on_submit(self, interaction: discord.Interaction):
        dados = obter_rascunho_boletim(interaction)
        if not dados:
            await interaction.response.send_message(
                "❌ Rascunho não encontrado ou sem permissão.",
                ephemeral=True,
            )
            return

        dados["local"] = str(self.local.value).strip()
        dados["relato"] = str(self.relato.value).strip()
        dados["etapa"] = "TIPO_IDENTIFICACAO"
        boletins_pendentes[interaction.channel.id] = dados
        salvar_boletins_pendentes()

        await interaction.response.send_message(
            "✅ Local e relato registrados.\n\n"
            "Agora escolha o que foi identificado:",
            view=TipoIdentificacaoBoletimView(),
        )


class VeiculoBoletimModal(Modal, title="Veículo Identificado"):
    modelo = TextInput(label="Modelo", placeholder="Não identificado", max_length=100)
    placa = TextInput(label="Placa", placeholder="Não identificado", max_length=50)
    telefone = TextInput(label="Telefone", placeholder="Não identificado", max_length=100)
    proprietario = TextInput(label="Proprietário", placeholder="Não identificado", max_length=150)
    rg_proprietario = TextInput(label="RG do proprietário", placeholder="Não identificado", max_length=80)

    async def on_submit(self, interaction: discord.Interaction):
        dados = obter_rascunho_boletim(interaction)
        if not dados:
            await interaction.response.send_message(
                "❌ Rascunho não encontrado ou sem permissão.",
                ephemeral=True,
            )
            return

        dados["dados_veiculo"] = {
            "modelo": str(self.modelo.value).strip(),
            "placa": str(self.placa.value).strip(),
            "telefone": str(self.telefone.value).strip(),
            "proprietario": str(self.proprietario.value).strip(),
            "rg_proprietario": str(self.rg_proprietario.value).strip(),
        }
        boletins_pendentes[interaction.channel.id] = dados
        salvar_boletins_pendentes()
        await depois_identificacao(interaction, dados)


class IndividuoBoletimModal(Modal, title="Indivíduo Identificado"):
    nome = TextInput(label="Nome", placeholder="Nome do indivíduo", max_length=150)
    rg = TextInput(label="RG", placeholder="Não identificado", max_length=80, required=False)
    caracteristicas = TextInput(
        label="Características",
        placeholder="Não identificado",
        style=discord.TextStyle.paragraph,
        max_length=800,
        required=False,
    )
    participacao = TextInput(
        label="Participação no fato",
        placeholder="Não identificado",
        style=discord.TextStyle.paragraph,
        max_length=800,
        required=False,
    )
    outras_informacoes = TextInput(
        label="Outras informações",
        placeholder="Não identificado",
        style=discord.TextStyle.paragraph,
        max_length=800,
        required=False,
    )

    def __init__(self, continuar_veiculo: bool = False):
        super().__init__()
        self.continuar_veiculo = continuar_veiculo

    async def on_submit(self, interaction: discord.Interaction):
        dados = obter_rascunho_boletim(interaction)
        if not dados:
            await interaction.response.send_message(
                "❌ Rascunho não encontrado ou sem permissão.",
                ephemeral=True,
            )
            return

        dados["dados_individuo"] = {
            "nome": str(self.nome.value).strip(),
            "rg": str(self.rg.value or "").strip(),
            "caracteristicas": str(self.caracteristicas.value or "").strip(),
            "participacao": str(self.participacao.value or "").strip(),
            "outras_informacoes": str(self.outras_informacoes.value or "").strip(),
        }
        boletins_pendentes[interaction.channel.id] = dados
        salvar_boletins_pendentes()

        if self.continuar_veiculo:
            await interaction.response.send_message(
                "✅ Indivíduo registrado. Agora clique abaixo para preencher o veículo.",
                view=ContinuarVeiculoBoletimView(),
            )
            return

        await depois_identificacao(interaction, dados)


class EditarLocalBoletimModal(Modal, title="Editar Local"):
    local = TextInput(label="Novo local", max_length=200)

    async def on_submit(self, interaction: discord.Interaction):
        dados = obter_rascunho_boletim(interaction)
        if not dados:
            await interaction.response.send_message("❌ Rascunho não encontrado.", ephemeral=True)
            return
        dados["local"] = str(self.local.value).strip()
        boletins_pendentes[interaction.channel.id] = dados
        salvar_boletins_pendentes()
        await interaction.response.defer(ephemeral=True)
        await enviar_previa_boletim(interaction.channel, dados)
        await interaction.followup.send("✅ Local atualizado.", ephemeral=True)


class EditarRelatoBoletimModal(Modal, title="Editar Relato"):
    relato = TextInput(
        label="Novo relato",
        style=discord.TextStyle.paragraph,
        max_length=4000,
    )

    async def on_submit(self, interaction: discord.Interaction):
        dados = obter_rascunho_boletim(interaction)
        if not dados:
            await interaction.response.send_message("❌ Rascunho não encontrado.", ephemeral=True)
            return
        dados["relato"] = str(self.relato.value).strip()
        dados["conclusao"] = gerar_conclusao_automatica(
            dados["relato"],
            dados.get("tipo_identificacao", ""),
            dados.get("dados_individuo", {}),
            dados.get("dados_veiculo", {}),
        )
        boletins_pendentes[interaction.channel.id] = dados
        salvar_boletins_pendentes()
        await interaction.response.defer(ephemeral=True)
        await enviar_previa_boletim(interaction.channel, dados)
        await interaction.followup.send(
            "✅ Relato atualizado e conclusão refeita.",
            ephemeral=True,
        )


class EditarConclusaoBoletimModal(Modal, title="Editar Conclusão"):
    conclusao = TextInput(
        label="Conclusão",
        style=discord.TextStyle.paragraph,
        max_length=2000,
    )

    async def on_submit(self, interaction: discord.Interaction):
        dados = obter_rascunho_boletim(interaction)
        if not dados:
            await interaction.response.send_message("❌ Rascunho não encontrado.", ephemeral=True)
            return
        dados["conclusao"] = str(self.conclusao.value).strip()
        boletins_pendentes[interaction.channel.id] = dados
        salvar_boletins_pendentes()
        await interaction.response.defer(ephemeral=True)
        await enviar_previa_boletim(interaction.channel, dados)
        await interaction.followup.send("✅ Conclusão atualizada.", ephemeral=True)


class ConsultarBoletimModal(Modal, title="Consultar Boletim"):
    numero = TextInput(
        label="Número do boletim",
        placeholder="BO-DICOR-20260702-001",
        max_length=40,
    )

    async def on_submit(self, interaction: discord.Interaction):
        boletim = buscar_boletim_numero(str(self.numero.value))
        if not boletim:
            await interaction.response.send_message(
                "❌ Boletim não encontrado.",
                ephemeral=True,
            )
            return
        await interaction.response.send_message(
            texto_consulta_boletim(boletim),
            ephemeral=True,
        )


class PainelBoletimView(View):
    def __init__(self):
        super().__init__(timeout=None)

    @discord.ui.button(
        label="Abrir Boletim",
        emoji="📝",
        style=discord.ButtonStyle.green,
        custom_id="dic_boletim_abrir",
    )
    async def abrir(self, interaction: discord.Interaction, button: Button):
        await abrir_boletim_core(interaction)

    @discord.ui.button(
        label="Meus Boletins",
        emoji="📂",
        style=discord.ButtonStyle.blurple,
        custom_id="dic_boletim_meus",
    )
    async def meus(self, interaction: discord.Interaction, button: Button):
        boletins = [
            boletim
            for boletim in carregar_boletins()
            if int(boletim.get("comunicante_id", 0)) == interaction.user.id
        ][-20:]
        if not boletins:
            await interaction.response.send_message(
                "📂 Você ainda não possui boletins publicados.",
                ephemeral=True,
            )
            return
        linhas = [
            f"• **{b.get('numero')}** | {b.get('data_fato')} | "
            f"{b.get('local')} | {b.get('mensagem_url', 'Sem link')}"
            for b in reversed(boletins)
        ]
        await interaction.response.send_message(
            "📂 **Meus últimos boletins:**\n\n" + "\n".join(linhas),
            ephemeral=True,
        )

    @discord.ui.button(
        label="Consultar Boletim",
        emoji="🔎",
        style=discord.ButtonStyle.gray,
        custom_id="dic_boletim_consultar",
    )
    async def consultar(self, interaction: discord.Interaction, button: Button):
        await interaction.response.send_modal(ConsultarBoletimModal())

    @discord.ui.button(
        label="Histórico",
        emoji="🗃️",
        style=discord.ButtonStyle.gray,
        custom_id="dic_boletim_historico",
    )
    async def historico(self, interaction: discord.Interaction, button: Button):
        if not isinstance(interaction.user, discord.Member) or not usuario_tem_equipe(interaction.user):
            await interaction.response.send_message(
                "❌ Apenas a equipe autorizada pode consultar o histórico.",
                ephemeral=True,
            )
            return
        boletins = carregar_boletins()[-25:]
        if not boletins:
            await interaction.response.send_message(
                "🗃️ Nenhum boletim publicado.",
                ephemeral=True,
            )
            return
        linhas = []
        for b in reversed(boletins):
            comunicante = b.get("comunicante_mention") or f"<@{b.get('comunicante_id', 0)}>"
            linhas.append(
                f"• **{b.get('numero')}** | {comunicante} | "
                f"{b.get('data_fato')} | {b.get('local')} | "
                f"{b.get('mensagem_url', 'Sem link')}"
            )
        await interaction.response.send_message(
            "🗃️ **Últimos boletins:**\n\n" + "\n".join(linhas),
            ephemeral=True,
        )

    async def on_error(self, interaction: discord.Interaction, error: Exception, item) -> None:
        await enviar_log(f"❌ Erro no painel de boletins: {error}")
        try:
            await responder_interacao(
                interaction,
                "❌ Ocorreu um erro no painel de boletins.",
                ephemeral=True,
            )
        except Exception:
            pass


class InformarRGBoletimView(View):
    def __init__(self):
        super().__init__(timeout=None)

    @discord.ui.button(
        label="Informar RG",
        emoji="🪪",
        style=discord.ButtonStyle.blurple,
        custom_id="dic_boletim_informar_rg",
    )
    async def informar(self, interaction: discord.Interaction, button: Button):
        await interaction.response.send_modal(InformarRGBoletimModal())


class IniciarBoletimView(View):
    def __init__(self):
        super().__init__(timeout=None)

    @discord.ui.button(
        label="Preencher Local e Relato",
        emoji="📝",
        style=discord.ButtonStyle.green,
        custom_id="dic_boletim_dados_iniciais",
    )
    async def iniciar(self, interaction: discord.Interaction, button: Button):
        if not obter_rascunho_boletim(interaction):
            await interaction.response.send_message(
                "❌ Rascunho não encontrado ou sem permissão.",
                ephemeral=True,
            )
            return
        await interaction.response.send_modal(DadosIniciaisBoletimModal())


class TipoIdentificacaoBoletimView(View):
    def __init__(self):
        super().__init__(timeout=None)

    @discord.ui.button(
        label="Veículo",
        emoji="🚘",
        style=discord.ButtonStyle.blurple,
        custom_id="dic_boletim_tipo_veiculo",
    )
    async def veiculo(self, interaction: discord.Interaction, button: Button):
        dados = obter_rascunho_boletim(interaction)
        if not dados:
            await interaction.response.send_message("❌ Rascunho não encontrado.", ephemeral=True)
            return
        dados["tipo_identificacao"] = "veiculo"
        dados["dados_individuo"] = {}
        boletins_pendentes[interaction.channel.id] = dados
        salvar_boletins_pendentes()
        await interaction.response.send_modal(VeiculoBoletimModal())

    @discord.ui.button(
        label="Indivíduo",
        emoji="👤",
        style=discord.ButtonStyle.blurple,
        custom_id="dic_boletim_tipo_individuo",
    )
    async def individuo(self, interaction: discord.Interaction, button: Button):
        dados = obter_rascunho_boletim(interaction)
        if not dados:
            await interaction.response.send_message("❌ Rascunho não encontrado.", ephemeral=True)
            return
        dados["tipo_identificacao"] = "individuo"
        dados["dados_veiculo"] = {}
        boletins_pendentes[interaction.channel.id] = dados
        salvar_boletins_pendentes()
        await interaction.response.send_modal(IndividuoBoletimModal())

    @discord.ui.button(
        label="Ambos",
        emoji="🔄",
        style=discord.ButtonStyle.green,
        custom_id="dic_boletim_tipo_ambos",
    )
    async def ambos(self, interaction: discord.Interaction, button: Button):
        dados = obter_rascunho_boletim(interaction)
        if not dados:
            await interaction.response.send_message("❌ Rascunho não encontrado.", ephemeral=True)
            return
        dados["tipo_identificacao"] = "ambos"
        boletins_pendentes[interaction.channel.id] = dados
        salvar_boletins_pendentes()
        await interaction.response.send_modal(
            IndividuoBoletimModal(continuar_veiculo=True)
        )


class ContinuarVeiculoBoletimView(View):
    def __init__(self):
        super().__init__(timeout=None)

    @discord.ui.button(
        label="Preencher Veículo",
        emoji="🚘",
        style=discord.ButtonStyle.green,
        custom_id="dic_boletim_continuar_veiculo",
    )
    async def continuar(self, interaction: discord.Interaction, button: Button):
        if not obter_rascunho_boletim(interaction):
            await interaction.response.send_message("❌ Rascunho não encontrado.", ephemeral=True)
            return
        await interaction.response.send_modal(VeiculoBoletimModal())


class ProvasBoletimView(View):
    def __init__(self):
        super().__init__(timeout=None)

    @discord.ui.button(
        label="Finalizar Provas",
        emoji="✅",
        style=discord.ButtonStyle.green,
        custom_id="dic_boletim_finalizar_provas",
    )
    async def finalizar(self, interaction: discord.Interaction, button: Button):
        dados = obter_rascunho_boletim(interaction)
        if not dados:
            await interaction.response.send_message("❌ Rascunho não encontrado.", ephemeral=True)
            return

        dados["conclusao"] = gerar_conclusao_automatica(
            dados.get("relato", ""),
            dados.get("tipo_identificacao", ""),
            dados.get("dados_individuo", {}),
            dados.get("dados_veiculo", {}),
        )
        boletins_pendentes[interaction.channel.id] = dados
        salvar_boletins_pendentes()

        await interaction.response.defer(ephemeral=True)
        await enviar_previa_boletim(interaction.channel, dados)
        anexos = await coletar_anexos_boletim(interaction.channel)
        await interaction.followup.send(
            f"✅ Prévia gerada. Provas encontradas: `{len(anexos)}`.",
            ephemeral=True,
        )

    @discord.ui.button(
        label="Voltar",
        emoji="↩️",
        style=discord.ButtonStyle.gray,
        custom_id="dic_boletim_voltar_identificacao",
    )
    async def voltar(self, interaction: discord.Interaction, button: Button):
        if not obter_rascunho_boletim(interaction):
            await interaction.response.send_message("❌ Rascunho não encontrado.", ephemeral=True)
            return
        await interaction.response.send_message(
            "Escolha novamente o tipo de identificação:",
            view=TipoIdentificacaoBoletimView(),
        )

    @discord.ui.button(
        label="Cancelar Boletim",
        emoji="❌",
        style=discord.ButtonStyle.red,
        custom_id="dic_boletim_cancelar_provas",
    )
    async def cancelar(self, interaction: discord.Interaction, button: Button):
        await cancelar_boletim_core(interaction)


class PreviaBoletimView(View):
    def __init__(self):
        super().__init__(timeout=None)

    @discord.ui.button(
        label="Publicar Boletim",
        emoji="✅",
        style=discord.ButtonStyle.green,
        custom_id="dic_boletim_publicar",
        row=0,
    )
    async def publicar(self, interaction: discord.Interaction, button: Button):
        await publicar_boletim_core(interaction)

    @discord.ui.button(
        label="Editar Local",
        emoji="✏️",
        style=discord.ButtonStyle.blurple,
        custom_id="dic_boletim_editar_local",
        row=0,
    )
    async def editar_local(self, interaction: discord.Interaction, button: Button):
        if not obter_rascunho_boletim(interaction):
            await interaction.response.send_message("❌ Rascunho não encontrado.", ephemeral=True)
            return
        await interaction.response.send_modal(EditarLocalBoletimModal())

    @discord.ui.button(
        label="Editar Relato",
        emoji="✏️",
        style=discord.ButtonStyle.blurple,
        custom_id="dic_boletim_editar_relato",
        row=0,
    )
    async def editar_relato(self, interaction: discord.Interaction, button: Button):
        if not obter_rascunho_boletim(interaction):
            await interaction.response.send_message("❌ Rascunho não encontrado.", ephemeral=True)
            return
        await interaction.response.send_modal(EditarRelatoBoletimModal())

    @discord.ui.button(
        label="Editar Identificação",
        emoji="✏️",
        style=discord.ButtonStyle.gray,
        custom_id="dic_boletim_editar_identificacao",
        row=1,
    )
    async def editar_identificacao(self, interaction: discord.Interaction, button: Button):
        dados = obter_rascunho_boletim(interaction)
        if not dados:
            await interaction.response.send_message("❌ Rascunho não encontrado.", ephemeral=True)
            return
        dados["retorno_identificacao"] = "PREVIA"
        boletins_pendentes[interaction.channel.id] = dados
        salvar_boletins_pendentes()
        await interaction.response.send_message(
            "Escolha novamente o tipo de identificação:",
            view=TipoIdentificacaoBoletimView(),
        )

    @discord.ui.button(
        label="Editar Conclusão",
        emoji="✏️",
        style=discord.ButtonStyle.gray,
        custom_id="dic_boletim_editar_conclusao",
        row=1,
    )
    async def editar_conclusao(self, interaction: discord.Interaction, button: Button):
        if not obter_rascunho_boletim(interaction):
            await interaction.response.send_message("❌ Rascunho não encontrado.", ephemeral=True)
            return
        await interaction.response.send_modal(EditarConclusaoBoletimModal())

    @discord.ui.button(
        label="Cancelar",
        emoji="❌",
        style=discord.ButtonStyle.red,
        custom_id="dic_boletim_cancelar_previa",
        row=1,
    )
    async def cancelar(self, interaction: discord.Interaction, button: Button):
        await cancelar_boletim_core(interaction)


def embed_painel_boletim_padrao() -> discord.Embed:
    return discord.Embed(
        title="📋 Sistema de Boletins - DICOR",
        description=(
            "Utilize os botões abaixo para registrar e consultar boletins.\n\n"
            "📝 **Abrir Boletim** — Cria um canal provisório privado.\n"
            "📂 **Meus Boletins** — Mostra seus últimos registros.\n"
            "🔎 **Consultar Boletim** — Consulta pelo número.\n"
            "🗃️ **Histórico** — Exibe os boletins recentes para a equipe."
        ),
        color=discord.Color.blue(),
    )



# =====================================================
# TABELA COMPARTILHADA DE 56 ORGANIZACOES
# =====================================================

TOTAL_ORGANIZACOES = 56
ORGANIZACOES_POR_PAGINA = 8

# Dados importados da planilha compartilhada. Apenas os campos solicitados
# são mantidos no sistema: Nome, Zona de risco, Status, Produto,
# Possui informante, Líder, Características e Histórico operacional.
ORGANIZACOES_INICIAIS: List[Dict[str, Any]] = [{'id': 1,
  'nome': '4 IRMÃOS / 4M',
  'zona_risco': '1',
  'status': 'Já Pacificada',
  'produto': 'H.',
  'informante': 'Não',
  'lider': 'Bruno Diniz 18130',
  'caracteristicas': 'Não informado',
  'historico': 'Invadimos e n tinha ngm'},
 {'id': 2,
  'nome': 'NOVA HOLANDA',
  'zona_risco': '2',
  'status': 'Investigar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Lucas Leal 5766',
  'caracteristicas': 'Os de Vermelho [milan]/Caixa dagua',
  'historico': 'ja incursamos varias vezes, chorão e faz feio.'},
 {'id': 3,
  'nome': 'ZR03',
  'zona_risco': '3',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'Não informado',
  'historico': 'Sem registros.'},
 {'id': 4,
  'nome': 'Não informado',
  'zona_risco': '4',
  'status': 'Conhecendo',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'FAVELA DA PRAIA',
  'historico': 'Sem registros.'},
 {'id': 5,
  'nome': 'LOS MALDITOS',
  'zona_risco': '5',
  'status': 'Investigando',
  'produto': 'Lança',
  'informante': 'Não',
  'lider': 'Vinicius Thugnine 28280',
  'caracteristicas': 'Não informado',
  'historico': 'Tel Vini: 305-133'},
 {'id': 6,
  'nome': 'ZR06',
  'zona_risco': '6',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'PLAYBOY',
  'historico': 'Sem registros.'},
 {'id': 7,
  'nome': 'ELIPA',
  'zona_risco': '7',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'Não informado',
  'historico': 'Sem registros.'},
 {'id': 8,
  'nome': 'ITALIA',
  'zona_risco': '8',
  'status': 'Conhecendo',
  'produto': 'Ticket | Placa | Nitro',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'Não informado',
  'historico': 'Sem registros.'},
 {'id': 9,
  'nome': 'ZR09',
  'zona_risco': '9',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'VINHEDOS',
  'historico': 'Sem registros.'},
 {'id': 10,
  'nome': 'ZR10',
  'zona_risco': '10',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'BAHAMAS',
  'historico': 'Sem registros.'},
 {'id': 11,
  'nome': 'BOLA MAIS UM',
  'zona_risco': '11',
  'status': 'Investigando',
  'produto': 'Balinha',
  'informante': 'Não',
  'lider': 'Caio Henrique 27365',
  'caracteristicas': 'FAZENDINHA',
  'historico': 'Sem registros.'},
 {'id': 12,
  'nome': 'ZR12',
  'zona_risco': '12',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'VANILLA',
  'historico': 'Sem registros.'},
 {'id': 13,
  'nome': 'ZR13',
  'zona_risco': '13',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'ANTIGA BROOCLIN',
  'historico': 'Sem registros.'},
 {'id': 14,
  'nome': 'ROYALT',
  'zona_risco': '14',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Vice Ragnar',
  'caracteristicas': 'Identidade Preto com Rosa. Carros Rosa',
  'historico': 'Sem registros.'},
 {'id': 15,
  'nome': 'MORRO DA SERPENTE',
  'zona_risco': '15',
  'status': 'Conhecendo',
  'produto': 'Armas',
  'informante': 'Não',
  'lider': 'Vulgo Portuga',
  'caracteristicas': 'Identidade "escama" preto com verde.',
  'historico': 'Sem registros.'},
 {'id': 16,
  'nome': 'PORTUGAL',
  'zona_risco': '16',
  'status': 'Conhecendo',
  'produto': 'Armas',
  'informante': 'Não',
  'lider': '02 Vulgo Bibi Perigosa',
  'caracteristicas': 'Identidade banca com dourado.',
  'historico': 'Tel 02: 675-403'},
 {'id': 17,
  'nome': 'ZR17',
  'zona_risco': '17',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'BOATE',
  'historico': 'Sem registros.'},
 {'id': 18,
  'nome': 'ZR18',
  'zona_risco': '18',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'PEGAR INFORMES HISTÓRICO COM CIVIL - PIKACHU DO BOPE OU CONSTANTINO DA CIVIL',
  'historico': 'Sem registros.'},
 {'id': 19,
  'nome': 'BALLAS',
  'zona_risco': '19',
  'status': 'Fora do Radar',
  'produto': 'Desmanche | CAPOFOL',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'Não informado',
  'historico': 'Informação proviniente de interrogatório durante prisional'},
 {'id': 20,
  'nome': 'FILHOS DA ANARQUIA',
  'zona_risco': '20',
  'status': 'Observando',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'Identidade "motociclista" preto com vermelho.',
  'historico': 'Local com suspeita de venda de droga, seguido de fuga para o interior do "bar".'},
 {'id': 21,
  'nome': 'ZR21',
  'zona_risco': '21',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'MANSÃO SCARFACE/PODEROSO CHEFÃO',
  'historico': 'Sem registros.'},
 {'id': 22,
  'nome': 'ZR22',
  'zona_risco': '22',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'MUSEU DE ARTE',
  'historico': 'Sem registros.'},
 {'id': 23,
  'nome': 'ZR23',
  'zona_risco': '23',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'GALAX',
  'historico': 'Sem registros.'},
 {'id': 24,
  'nome': 'DINASTIA',
  'zona_risco': '24',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'HOTEL CALISTO',
  'historico': 'Sem registros.'},
 {'id': 25,
  'nome': 'BLACK DRAGONS',
  'zona_risco': '25',
  'status': 'Conhecendo',
  'produto': 'Skunk',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'Identidade Cammo Cinza com Preto',
  'historico': 'Sem registros.'},
 {'id': 26,
  'nome': 'ZR26',
  'zona_risco': '26',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'PETROLÍFERA',
  'historico': 'Sem registros.'},
 {'id': 27,
  'nome': 'ZR27',
  'zona_risco': '27',
  'status': 'Investigar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'LIFE INVADER',
  'historico': 'Sem registros.'},
 {'id': 28,
  'nome': 'MORRO DO MINEIRO',
  'zona_risco': '28',
  'status': 'Observando',
  'produto': 'Desmanche | Lança',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'OBSERVATÓRIO',
  'historico': 'Sem registros.'},
 {'id': 29,
  'nome': 'ZR29',
  'zona_risco': '29',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'SUCATINGA',
  'historico': 'Sem registros.'},
 {'id': 30,
  'nome': 'MEDELIN',
  'zona_risco': '30',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'CEMITÉRIO',
  'historico': 'Sem registros.'},
 {'id': 31,
  'nome': 'ZR31',
  'zona_risco': '31',
  'status': 'Investigar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'Não informado',
  'historico': 'Sem registros.'},
 {'id': 32,
  'nome': 'ZR32',
  'zona_risco': '32',
  'status': 'Investigar',
  'produto': 'Desconhecido',
  'informante': 'Sim',
  'lider': 'Não informado',
  'caracteristicas': 'BUNKER DO FAROL',
  'historico': 'INFORMANTE M.B. - DO CORVO'},
 {'id': 33,
  'nome': 'ZR33',
  'zona_risco': '33',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'BOLA DA BENNYS',
  'historico': 'Sem registros.'},
 {'id': 34,
  'nome': 'VERA CRUZ',
  'zona_risco': '34',
  'status': 'Conhecendo',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'JOSE PARAÍBA 29063',
  'caracteristicas': 'ANTIGO CAMPINHO',
  'historico': 'Sem registros.'},
 {'id': 35,
  'nome': 'ZR35',
  'zona_risco': '35',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'Não informado',
  'historico': 'Sem registros.'},
 {'id': 36,
  'nome': 'ANONIMOUS',
  'zona_risco': '36',
  'status': 'Observando',
  'produto': 'Attachs',
  'informante': 'Não',
  'lider': 'Evon Zayon',
  'caracteristicas': 'MORRO DO DINO',
  'historico': 'Sem registros.'},
 {'id': 37,
  'nome': 'ZR37',
  'zona_risco': '37',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'Não informado',
  'historico': 'Sem registros.'},
 {'id': 38,
  'nome': 'ZR38',
  'zona_risco': '38',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'Não informado',
  'historico': 'Sem registros.'},
 {'id': 39,
  'nome': 'ZR39',
  'zona_risco': '39',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'YAKUZA',
  'historico': 'Sem registros.'},
 {'id': 40,
  'nome': 'ZR40',
  'zona_risco': '40',
  'status': 'Investigando',
  'produto': 'Rapé',
  'informante': 'Sim',
  'lider': 'Não informado',
  'caracteristicas': 'Não informado',
  'historico': '[INFO JOAQUIM] Elemento entrando na família. Lá tem Rapé.'},
 {'id': 41,
  'nome': 'ZR41',
  'zona_risco': '41',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'Não informado',
  'historico': 'Sem registros.'},
 {'id': 42,
  'nome': 'ZR42',
  'zona_risco': '42',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'PROXIMO A RED LINE',
  'historico': 'Sem registros.'},
 {'id': 43,
  'nome': 'ZR43',
  'zona_risco': '43',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'Não informado',
  'historico': 'Sem registros.'},
 {'id': 44,
  'nome': 'ZR44',
  'zona_risco': '44',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'Não informado',
  'historico': 'Sem registros.'},
 {'id': 45,
  'nome': 'ZR45',
  'zona_risco': '45',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'Não informado',
  'historico': 'Sem registros.'},
 {'id': 46,
  'nome': 'ZR46',
  'zona_risco': '46',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'DE FRENTE PRA LIFE',
  'historico': 'Sem registros.'},
 {'id': 47,
  'nome': 'ZR47',
  'zona_risco': '47',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'Não informado',
  'historico': 'Sem registros.'},
 {'id': 48,
  'nome': '????',
  'zona_risco': '48',
  'status': 'Observando',
  'produto': 'Desmanche',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'MANICÔMIO',
  'historico': 'Sem registros.'},
 {'id': 49,
  'nome': 'ELEMENTS',
  'zona_risco': '49',
  'status': 'Investigando',
  'produto': 'Master Pick | Flipper mk2',
  'informante': 'Não',
  'lider': 'Tomas Biel [01/gerente]',
  'caracteristicas': 'Não informado',
  'historico': 'Tel do Gerente Tomas: 764-013 // ofereceu masterpick'},
 {'id': 50,
  'nome': 'ZR50',
  'zona_risco': '50',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'Não informado',
  'historico': 'Sem registros.'},
 {'id': 51,
  'nome': 'ZR51',
  'zona_risco': '51',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'Não informado',
  'historico': 'Sem registros.'},
 {'id': 52,
  'nome': 'ZR52',
  'zona_risco': '52',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'Não informado',
  'historico': 'Sem registros.'},
 {'id': 53,
  'nome': 'ZR53',
  'zona_risco': '53',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'Não informado',
  'historico': 'Sem registros.'},
 {'id': 54,
  'nome': 'ZR54',
  'zona_risco': '54',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'Não informado',
  'historico': 'Sem registros.'},
 {'id': 55,
  'nome': 'ZR55',
  'zona_risco': '55',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'Não informado',
  'historico': 'Sem registros.'},
 {'id': 56,
  'nome': 'ZR56',
  'zona_risco': '56',
  'status': 'Fora do Radar',
  'produto': 'Desconhecido',
  'informante': 'Não',
  'lider': 'Não informado',
  'caracteristicas': 'Não informado',
  'historico': 'Sem registros.'}]
ORGANIZACOES_INICIAIS_POR_ID = {int(item["id"]): item for item in ORGANIZACOES_INICIAIS}


def modelo_organizacao(numero: int) -> Dict[str, Any]:
    inicial = dict(ORGANIZACOES_INICIAIS_POR_ID.get(numero, {}))
    return {
        "id": numero,
        "nome": str(inicial.get("nome", f"Organização {numero:02d}")),
        "zona_risco": str(inicial.get("zona_risco", numero)),
        "status": str(inicial.get("status", "Não informado")),
        "produto": str(inicial.get("produto", "Desconhecido")),
        "informante": str(inicial.get("informante", "Não")),
        "lider": str(inicial.get("lider", "Não informado")),
        "caracteristicas": str(inicial.get("caracteristicas", "Não informado")),
        "historico": str(inicial.get("historico", "Sem registros.")),
        "versao": 1,
        "ultima_edicao": None,
        "editado_por": None,
        "editado_por_id": None,
        "origem": "Planilha Google",
    }


def registro_organizacao_foi_editado(registro: Dict[str, Any]) -> bool:
    return bool(
        registro.get("editado_por_id")
        or registro.get("ultima_edicao")
        or int(registro.get("versao", 1) or 1) > 1
    )


def normalizar_organizacao(registro: Dict[str, Any], numero: int) -> Dict[str, Any]:
    base = modelo_organizacao(numero)

    # Mantém edições realizadas pelo Discord. Registros antigos que nunca
    # foram editados são atualizados com os dados reais da planilha.
    if isinstance(registro, dict) and registro_organizacao_foi_editado(registro):
        for chave in base:
            if chave in registro:
                base[chave] = registro[chave]

    base["id"] = numero
    try:
        base["versao"] = max(1, int(base.get("versao", 1) or 1))
    except (TypeError, ValueError):
        base["versao"] = 1
    return base


def garantir_56_organizacoes() -> List[Dict[str, Any]]:
    dados = carregar_json(ORGANIZACOES_JSON, [])
    existentes: Dict[int, Dict[str, Any]] = {}

    if isinstance(dados, list):
        for registro in dados:
            if not isinstance(registro, dict):
                continue
            try:
                numero = int(registro.get("id", 0) or 0)
            except (TypeError, ValueError):
                continue
            if 1 <= numero <= TOTAL_ORGANIZACOES and numero not in existentes:
                existentes[numero] = registro

    organizacoes = [
        normalizar_organizacao(existentes.get(numero, {}), numero)
        for numero in range(1, TOTAL_ORGANIZACOES + 1)
    ]
    salvar_json(ORGANIZACOES_JSON, organizacoes)
    return organizacoes


def carregar_organizacoes() -> List[Dict[str, Any]]:
    return garantir_56_organizacoes()


def salvar_organizacoes(lista: List[Dict[str, Any]]) -> None:
    por_id: Dict[int, Dict[str, Any]] = {}
    for registro in lista:
        if not isinstance(registro, dict):
            continue
        try:
            numero = int(registro.get("id", 0) or 0)
        except (TypeError, ValueError):
            continue
        if 1 <= numero <= TOTAL_ORGANIZACOES:
            por_id[numero] = registro

    organizacoes = [
        normalizar_organizacao(por_id.get(numero, {}), numero)
        for numero in range(1, TOTAL_ORGANIZACOES + 1)
    ]
    salvar_json(ORGANIZACOES_JSON, organizacoes)


def carregar_historico_organizacoes() -> List[Dict[str, Any]]:
    dados = carregar_json(HISTORICO_ORGANIZACOES_JSON, [])
    return dados if isinstance(dados, list) else []


def obter_organizacao_por_id(numero: int) -> Optional[Dict[str, Any]]:
    for organizacao in carregar_organizacoes():
        if int(organizacao.get("id", 0) or 0) == numero:
            return organizacao
    return None


def texto_normalizado_busca(texto: str) -> str:
    valor = unicodedata.normalize("NFKD", str(texto or ""))
    valor = valor.encode("ascii", "ignore").decode("ascii").lower().strip()
    return re.sub(r"\s+", " ", valor)


def pesquisar_organizacoes(termo: str) -> List[Dict[str, Any]]:
    termo_limpo = texto_normalizado_busca(termo)
    if not termo_limpo:
        return []

    if termo_limpo.isdigit():
        numero = int(termo_limpo)
        organizacao = obter_organizacao_por_id(numero)
        return [organizacao] if organizacao else []

    organizacoes = carregar_organizacoes()
    exatas = [
        org for org in organizacoes
        if texto_normalizado_busca(org.get("nome", "")) == termo_limpo
    ]
    if exatas:
        return exatas

    return [
        org for org in organizacoes
        if termo_limpo in texto_normalizado_busca(org.get("nome", ""))
    ]


def valor_org(registro: Dict[str, Any], chave: str, padrao: str = "Não informado") -> str:
    valor = str(registro.get(chave, "") or "").strip()
    return valor or padrao


def cortar_campo_org(texto: str, limite: int) -> str:
    valor = str(texto or "").strip()
    if len(valor) <= limite:
        return valor
    return valor[: limite - 3].rstrip() + "..."


def formatar_ficha_organizacao(organizacao: Dict[str, Any]) -> str:
    numero = int(organizacao.get("id", 0) or 0)
    ultima = valor_org(organizacao, "ultima_edicao", "Nunca editada")
    editor = valor_org(organizacao, "editado_por", "Nenhum")

    texto = (
        f"🏴 **ORGANIZAÇÃO {numero:02d} — {valor_org(organizacao, 'nome')}**\n\n"
        f"⚠️ **Zona de risco:** {valor_org(organizacao, 'zona_risco')}\n"
        f"📊 **Status:** {valor_org(organizacao, 'status')}\n"
        f"📦 **Produto:** {valor_org(organizacao, 'produto')}\n"
        f"🕵️ **Possui informante:** {valor_org(organizacao, 'informante')}\n"
        f"👑 **Líder:** {valor_org(organizacao, 'lider')}\n\n"
        f"🔎 **Características:**\n"
        f"{cortar_campo_org(valor_org(organizacao, 'caracteristicas'), 620)}\n\n"
        f"📂 **Histórico operacional:**\n"
        f"{cortar_campo_org(valor_org(organizacao, 'historico', 'Sem registros.'), 620)}\n\n"
        f"━━━━━━━━━━━━━━━━━━━━━━━\n"
        f"✏️ **Última edição:** {ultima}\n"
        f"👤 **Editado por:** {editor}\n"
        f"🔢 **Versão:** {int(organizacao.get('versao', 1) or 1)}"
    )
    return cortar_discord(texto, 1900)


def formatar_lista_organizacoes(pagina: int = 0) -> str:
    organizacoes = carregar_organizacoes()
    total_paginas = max(
        1,
        (len(organizacoes) + ORGANIZACOES_POR_PAGINA - 1)
        // ORGANIZACOES_POR_PAGINA,
    )
    pagina = max(0, min(pagina, total_paginas - 1))
    inicio = pagina * ORGANIZACOES_POR_PAGINA
    fim = inicio + ORGANIZACOES_POR_PAGINA

    linhas = [
        "📋 **TABELA DE ORGANIZAÇÕES — DICOR**",
        f"Página `{pagina + 1}/{total_paginas}` • Total: `{TOTAL_ORGANIZACOES}`",
        "",
    ]

    for org in organizacoes[inicio:fim]:
        linhas.append(
            f"**{int(org.get('id', 0)):02d}. {cortar_campo_org(valor_org(org, 'nome'), 35)}**\n"
            f"> ⚠️ {cortar_campo_org(valor_org(org, 'zona_risco'), 18)} | "
            f"📊 {cortar_campo_org(valor_org(org, 'status'), 24)} | "
            f"📦 {cortar_campo_org(valor_org(org, 'produto'), 22)}\n"
            f"> 🕵️ Informante: {cortar_campo_org(valor_org(org, 'informante'), 12)} | "
            f"👑 {cortar_campo_org(valor_org(org, 'lider'), 32)}"
        )

    return cortar_discord("\n".join(linhas), 1900)


def formatar_resultados_organizacoes(resultados: List[Dict[str, Any]]) -> str:
    linhas = ["🔎 **Resultados encontrados:**", ""]
    for org in resultados[:15]:
        linhas.append(
            f"• `{int(org.get('id', 0)):02d}` — **{valor_org(org, 'nome')}** "
            f"| {valor_org(org, 'status')}"
        )
    if len(resultados) > 15:
        linhas.append(f"\n... e mais `{len(resultados) - 15}` resultado(s).")
    linhas.append("\nUse o número exato no botão **Editar Organização**.")
    return cortar_discord("\n".join(linhas), 1900)


async def registrar_edicao_organizacao(
    organizacao_antes: Dict[str, Any],
    organizacao_depois: Dict[str, Any],
    usuario: Any,
    alteracoes: Dict[str, Dict[str, str]],
) -> None:
    historico = carregar_historico_organizacoes()
    historico.append({
        "organizacao_id": int(organizacao_depois.get("id", 0) or 0),
        "nome": valor_org(organizacao_depois, "nome"),
        "data": agora_br(),
        "usuario_id": usuario.id,
        "usuario": str(usuario),
        "versao_anterior": int(organizacao_antes.get("versao", 1) or 1),
        "versao_nova": int(organizacao_depois.get("versao", 1) or 1),
        "alteracoes": alteracoes,
    })
    salvar_json(HISTORICO_ORGANIZACOES_JSON, historico)

    campos = ", ".join(alteracoes.keys())
    await enviar_log(
        f"✏️ **Organização editada**\n"
        f"Organização: `{int(organizacao_depois.get('id', 0)):02d}` — "
        f"{valor_org(organizacao_depois, 'nome')}\n"
        f"Editado por: <@{usuario.id}>\n"
        f"Campos alterados: {campos}\n"
        f"Versão: {int(organizacao_depois.get('versao', 1) or 1)}"
    )


async def aplicar_edicao_organizacao(
    numero: int,
    versao_esperada: int,
    novos_valores: Dict[str, str],
    usuario: Any,
) -> tuple[bool, str, Optional[Dict[str, Any]]]:
    organizacoes = carregar_organizacoes()
    indice = next(
        (
            i for i, org in enumerate(organizacoes)
            if int(org.get("id", 0) or 0) == numero
        ),
        None,
    )
    if indice is None:
        return False, "❌ Organização não encontrada.", None

    atual = organizacoes[indice]
    versao_atual = int(atual.get("versao", 1) or 1)
    if versao_atual != versao_esperada:
        return (
            False,
            "⚠️ Essa organização foi editada por outra pessoa enquanto você preenchia. "
            "Abra a edição novamente para não sobrescrever dados mais recentes.",
            atual,
        )

    antes = dict(atual)
    alteracoes: Dict[str, Dict[str, str]] = {}

    for chave, valor in novos_valores.items():
        if chave not in {
            "nome",
            "zona_risco",
            "status",
            "produto",
            "informante",
            "lider",
            "caracteristicas",
            "historico",
        }:
            continue

        novo = str(valor or "").strip() or "Não informado"
        antigo = str(atual.get(chave, "") or "").strip() or "Não informado"
        if novo != antigo:
            alteracoes[chave] = {"antes": antigo, "depois": novo}
            atual[chave] = novo

    if not alteracoes:
        return False, "ℹ️ Nenhuma informação foi alterada.", atual

    atual["versao"] = versao_atual + 1
    atual["ultima_edicao"] = agora_br()
    atual["editado_por"] = str(usuario)
    atual["editado_por_id"] = usuario.id
    organizacoes[indice] = atual
    salvar_organizacoes(organizacoes)
    await registrar_edicao_organizacao(antes, atual, usuario, alteracoes)
    return True, "✅ Organização atualizada e alteração salva.", atual


class PaginacaoOrganizacoesView(View):
    def __init__(self, pagina: int = 0):
        super().__init__(timeout=300)
        total_paginas = max(
            1,
            (TOTAL_ORGANIZACOES + ORGANIZACOES_POR_PAGINA - 1)
            // ORGANIZACOES_POR_PAGINA,
        )
        self.pagina = max(0, min(pagina, total_paginas - 1))
        self.total_paginas = total_paginas
        for item in self.children:
            if isinstance(item, Button) and item.label == "Anterior":
                item.disabled = self.pagina <= 0
            elif isinstance(item, Button) and item.label == "Próxima":
                item.disabled = self.pagina >= self.total_paginas - 1

    @discord.ui.button(label="Anterior", emoji="⬅️", style=discord.ButtonStyle.gray)
    async def anterior(self, interaction: discord.Interaction, button: Button):
        nova_pagina = max(0, self.pagina - 1)
        await interaction.response.edit_message(
            content=formatar_lista_organizacoes(nova_pagina),
            view=PaginacaoOrganizacoesView(nova_pagina),
        )

    @discord.ui.button(label="Próxima", emoji="➡️", style=discord.ButtonStyle.gray)
    async def proxima(self, interaction: discord.Interaction, button: Button):
        nova_pagina = min(self.total_paginas - 1, self.pagina + 1)
        await interaction.response.edit_message(
            content=formatar_lista_organizacoes(nova_pagina),
            view=PaginacaoOrganizacoesView(nova_pagina),
        )


class EditarOrganizacaoBasicoModal(Modal):
    def __init__(self, organizacao: Dict[str, Any]):
        numero = int(organizacao.get("id", 0) or 0)
        super().__init__(title=f"Editar organização {numero:02d}")
        self.numero = numero
        self.versao = int(organizacao.get("versao", 1) or 1)

        self.nome = TextInput(
            label="Nome",
            default=valor_org(organizacao, "nome"),
            max_length=100,
        )
        self.zona_risco = TextInput(
            label="Zona de risco",
            default=valor_org(organizacao, "zona_risco"),
            placeholder="Ex: Baixa, Média, Alta ou Crítica",
            max_length=80,
        )
        self.status = TextInput(
            label="Status",
            default=valor_org(organizacao, "status"),
            placeholder="Ex: Ativa, Monitoramento, Em investigação",
            max_length=100,
        )
        self.produto = TextInput(
            label="Produto",
            default=valor_org(organizacao, "produto"),
            max_length=150,
        )
        self.informante = TextInput(
            label="Possui informante?",
            default=valor_org(organizacao, "informante"),
            placeholder="Sim, Não ou Não identificado",
            max_length=60,
        )

        for item in (
            self.nome,
            self.zona_risco,
            self.status,
            self.produto,
            self.informante,
        ):
            self.add_item(item)

    async def on_submit(self, interaction: discord.Interaction):
        sucesso, mensagem, organizacao = await aplicar_edicao_organizacao(
            self.numero,
            self.versao,
            {
                "nome": str(self.nome.value),
                "zona_risco": str(self.zona_risco.value),
                "status": str(self.status.value),
                "produto": str(self.produto.value),
                "informante": str(self.informante.value),
            },
            interaction.user,
        )
        conteudo = mensagem
        view = None
        if organizacao:
            conteudo += "\n\n" + formatar_ficha_organizacao(organizacao)
            view = OrganizacaoAcoesView(
                int(organizacao.get("id", 0)),
                int(organizacao.get("versao", 1) or 1),
            )
        await interaction.response.send_message(
            cortar_discord(conteudo, 1900),
            view=view,
            ephemeral=True,
        )


class EditarOrganizacaoDetalhesModal(Modal):
    def __init__(self, organizacao: Dict[str, Any]):
        numero = int(organizacao.get("id", 0) or 0)
        super().__init__(title=f"Detalhes da organização {numero:02d}")
        self.numero = numero
        self.versao = int(organizacao.get("versao", 1) or 1)

        self.lider = TextInput(
            label="Líder",
            default=valor_org(organizacao, "lider"),
            max_length=150,
        )
        self.caracteristicas = TextInput(
            label="Características",
            default=valor_org(organizacao, "caracteristicas"),
            style=discord.TextStyle.paragraph,
            max_length=1800,
        )
        self.historico = TextInput(
            label="Histórico operacional",
            default=valor_org(organizacao, "historico", "Sem registros."),
            placeholder="Registre ocorrências, datas e informações relevantes.",
            style=discord.TextStyle.paragraph,
            max_length=3000,
        )

        self.add_item(self.lider)
        self.add_item(self.caracteristicas)
        self.add_item(self.historico)

    async def on_submit(self, interaction: discord.Interaction):
        sucesso, mensagem, organizacao = await aplicar_edicao_organizacao(
            self.numero,
            self.versao,
            {
                "lider": str(self.lider.value),
                "caracteristicas": str(self.caracteristicas.value),
                "historico": str(self.historico.value),
            },
            interaction.user,
        )
        conteudo = mensagem
        view = None
        if organizacao:
            conteudo += "\n\n" + formatar_ficha_organizacao(organizacao)
            view = OrganizacaoAcoesView(
                int(organizacao.get("id", 0)),
                int(organizacao.get("versao", 1) or 1),
            )
        await interaction.response.send_message(
            cortar_discord(conteudo, 1900),
            view=view,
            ephemeral=True,
        )


class OrganizacaoAcoesView(View):
    def __init__(self, numero: int, versao: int):
        super().__init__(timeout=300)
        self.numero = numero
        self.versao = versao

    @discord.ui.button(label="Editar Dados", emoji="✏️", style=discord.ButtonStyle.blurple)
    async def editar_dados(self, interaction: discord.Interaction, button: Button):
        organizacao = obter_organizacao_por_id(self.numero)
        if not organizacao:
            await interaction.response.send_message("❌ Organização não encontrada.", ephemeral=True)
            return
        await interaction.response.send_modal(EditarOrganizacaoBasicoModal(organizacao))

    @discord.ui.button(label="Editar Detalhes", emoji="📝", style=discord.ButtonStyle.green)
    async def editar_detalhes(self, interaction: discord.Interaction, button: Button):
        organizacao = obter_organizacao_por_id(self.numero)
        if not organizacao:
            await interaction.response.send_message("❌ Organização não encontrada.", ephemeral=True)
            return
        await interaction.response.send_modal(EditarOrganizacaoDetalhesModal(organizacao))

    @discord.ui.button(label="Atualizar Ficha", emoji="🔄", style=discord.ButtonStyle.gray)
    async def atualizar(self, interaction: discord.Interaction, button: Button):
        organizacao = obter_organizacao_por_id(self.numero)
        if not organizacao:
            await interaction.response.send_message("❌ Organização não encontrada.", ephemeral=True)
            return
        await interaction.response.edit_message(
            content=formatar_ficha_organizacao(organizacao),
            view=OrganizacaoAcoesView(
                self.numero,
                int(organizacao.get("versao", 1) or 1),
            ),
        )


class PesquisarOrganizacaoModal(Modal, title="Pesquisar Organização"):
    termo = TextInput(
        label="Nome ou número",
        placeholder="Ex: 12 ou Medelim",
        max_length=100,
    )

    async def on_submit(self, interaction: discord.Interaction):
        resultados = pesquisar_organizacoes(str(self.termo.value))
        if not resultados:
            await interaction.response.send_message(
                "❌ Nenhuma organização encontrada.",
                ephemeral=True,
            )
            return

        if len(resultados) == 1:
            organizacao = resultados[0]
            await interaction.response.send_message(
                formatar_ficha_organizacao(organizacao),
                view=OrganizacaoAcoesView(
                    int(organizacao.get("id", 0)),
                    int(organizacao.get("versao", 1) or 1),
                ),
                ephemeral=True,
            )
            return

        await interaction.response.send_message(
            formatar_resultados_organizacoes(resultados),
            ephemeral=True,
        )


class SelecionarOrganizacaoEditarModal(Modal, title="Editar Organização"):
    termo = TextInput(
        label="Nome ou número",
        placeholder="Ex: 12 ou Medelim",
        max_length=100,
    )

    async def on_submit(self, interaction: discord.Interaction):
        resultados = pesquisar_organizacoes(str(self.termo.value))
        if not resultados:
            await interaction.response.send_message(
                "❌ Nenhuma organização encontrada.",
                ephemeral=True,
            )
            return

        if len(resultados) > 1:
            await interaction.response.send_message(
                formatar_resultados_organizacoes(resultados),
                ephemeral=True,
            )
            return

        organizacao = resultados[0]
        await interaction.response.send_message(
            formatar_ficha_organizacao(organizacao),
            view=OrganizacaoAcoesView(
                int(organizacao.get("id", 0)),
                int(organizacao.get("versao", 1) or 1),
            ),
            ephemeral=True,
        )


class PainelOrganizacoesView(View):
    def __init__(self):
        super().__init__(timeout=None)

    @discord.ui.button(
        label="Ver Organizações",
        emoji="📋",
        style=discord.ButtonStyle.blurple,
        custom_id="dic_org_ver_organizacoes",
    )
    async def ver_organizacoes(self, interaction: discord.Interaction, button: Button):
        garantir_56_organizacoes()
        await interaction.response.send_message(
            formatar_lista_organizacoes(0),
            view=PaginacaoOrganizacoesView(0),
            ephemeral=True,
        )

    @discord.ui.button(
        label="Pesquisar Organização",
        emoji="🔎",
        style=discord.ButtonStyle.gray,
        custom_id="dic_org_pesquisar",
    )
    async def pesquisar(self, interaction: discord.Interaction, button: Button):
        await interaction.response.send_modal(PesquisarOrganizacaoModal())

    @discord.ui.button(
        label="Editar Organização",
        emoji="✏️",
        style=discord.ButtonStyle.green,
        custom_id="dic_org_editar",
    )
    async def editar(self, interaction: discord.Interaction, button: Button):
        await interaction.response.send_modal(SelecionarOrganizacaoEditarModal())

    async def on_error(self, interaction: discord.Interaction, error: Exception, item) -> None:
        await enviar_log(f"❌ Erro no painel de organizações: {error}")
        try:
            if interaction.response.is_done():
                await interaction.followup.send(
                    "❌ Ocorreu um erro no sistema de organizações.",
                    ephemeral=True,
                )
            else:
                await interaction.response.send_message(
                    "❌ Ocorreu um erro no sistema de organizações.",
                    ephemeral=True,
                )
        except Exception:
            pass


def embed_painel_organizacoes_padrao() -> discord.Embed:
    return discord.Embed(
        title="🏴 Tabela de Organizações - DICOR",
        description=(
            "As **56 organizações** são fixas e compartilhadas. "
            "Todos podem consultar e editar qualquer ficha.\n\n"
            "📋 **Ver Organizações** — Abre a tabela completa por páginas.\n"
            "🔎 **Pesquisar Organização** — Busca pelo nome ou número.\n"
            "✏️ **Editar Organização** — Atualiza os dados da ficha.\n\n"
            "🔒 Todas as alterações ficam salvas com usuário, data, "
            "valores anteriores e valores novos. Nenhuma organização pode ser apagada."
        ),
        color=discord.Color.dark_blue(),
    )

# =====================================================
# COMANDOS DE TEXTO (!) - PLANO B PARA QUANDO O DISCORD NÃO MOSTRAR SLASH
# =====================================================

def embed_painel_procurados_padrao() -> discord.Embed:
    return discord.Embed(
        title="🚨 Sistema de Procurados - DICOR",
        description=(
            "Utilize os botões abaixo para gerenciar procurados.\n\n"
            "➕ **Novo Procurado** — Cadastrar um novo procurado.\n"
            "📋 **Lista de Procurados** — Ver procurados ativos.\n"
            "❌ **Retirar Procurado** — Retirar um procurado pelo RG.\n"
            "✏️ **Modificar Procurado** — Buscar pelo RG e adicionar novos crimes.\n"
            "📄 **Abrir Catálogo** — Abrir o catálogo HTML.\n"
            "🔄 **Sincronizar Antigos** — Importar procurados antigos do canal oficial."
        ),
        color=discord.Color.red(),
    )


def embed_painel_mesas_padrao() -> discord.Embed:
    return discord.Embed(
        title="🕵️ Sistema de Mesas - DICOR",
        description=(
            "Utilize os botões abaixo para gerenciar as mesas.\n\n"
            "➕ **Criar Mesa** — Abre uma nova mesa de investigação.\n"
            "🔓 **Reabrir Mesa** — Reabre uma mesa fechada.\n"
            "📂 **Histórico** — Mostra mesas registradas.\n"
            "📦 **Backup** — Gera backup dos dados.\n"
            "📊 **Estatísticas** — Mostra números do sistema."
        ),
        color=discord.Color.gold(),
    )


async def criar_mesa_por_texto(ctx: commands.Context, apelido: str, familia: str, observacao: str = ""):
    guild = ctx.guild
    if guild is None:
        await ctx.reply("❌ Use dentro de um servidor.")
        return

    categoria = guild.get_channel(CATEGORIA_MESAS_ABERTAS_ID) if CATEGORIA_MESAS_ABERTAS_ID else None
    overwrites = cargos_equipe_permissoes(guild)
    overwrites[ctx.author] = discord.PermissionOverwrite(
        view_channel=True,
        send_messages=True,
        attach_files=True,
        read_message_history=True,
        send_messages_in_threads=True,
    )

    nome_canal = f"🕵️┃{slugify(apelido)}-{slugify(familia)}"
    canal = await guild.create_text_channel(
        name=nome_canal,
        category=categoria,
        overwrites=overwrites,
    )

    await canal.send(
        f"🕵️ **Mesa de Investigação — {familia}**\n"
        f"👮 **Agente:** {ctx.author.mention}\n"
        f"📛 **Apelido:** {apelido}\n"
        f"🏷️ **Organização/Família:** {familia}\n"
        f"🕒 **Criada em:** {agora_br()}\n"
        f"📝 **Observação:** {observacao or 'Nenhuma'}\n\n"
        f"━━━━━━━━━━━━━━━━━━━━━━━\n"
        f"🧵 **Use os tópicos abertos abaixo para organizar a investigação.**"
    )

    topicos_ids = await criar_topicos_reais_mesa(canal)
    
    # Coleta os dados em tempo de execução para alimentar a capa do PDF
    dados_iniciais_mesa = {
        "nome": f"OPERAÇÃO {familia.upper()}",
        "comunidade": familia,
        "delegado": str(interaction.user.display_name if hasattr(interaction, 'user') else ctx.author.display_name),
        "data_abertura": agora_br()
    }

    await canal.send(
        "🔒 Para encerrar esta mesa, clique no botão abaixo.",
        view=FecharMesaView()
    )

    registrar_mesa({
        "canal_id": canal.id,
        "nome_canal": canal.name,
        "apelido": apelido,
        "familia": familia,
        "autor_id": ctx.author.id,
        "autor_nome": str(ctx.author),
        "status": "ABERTA",
        "criada_em": agora_br(),
        "fechada_em": None,
        "topicos_ids": topicos_ids,
    })

    await enviar_log(
        f"➕ Mesa criada por comando de texto: {canal.mention} | Família: {familia} | Por: {ctx.author.mention}"
    )
    await ctx.reply(f"✅ Mesa criada: {canal.mention}")

@bot.command(name="painelorganizacoes")
async def cmd_painelorganizacoes(ctx: commands.Context):
    """Envia o painel compartilhado das 56 organizações."""
    garantir_56_organizacoes()
    await ctx.send(
        embed=embed_painel_organizacoes_padrao(),
        view=PainelOrganizacoesView(),
    )


@bot.command(name="painelboletim")
async def cmd_painelboletim(ctx: commands.Context):
    """Abre o painel de boletins pelo comando !painelboletim."""
    if not isinstance(ctx.author, discord.Member) or not usuario_tem_equipe(ctx.author):
        await ctx.reply("❌ Apenas a equipe autorizada pode enviar este painel.")
        return
    await ctx.send(embed=embed_painel_boletim_padrao(), view=PainelBoletimView())


@bot.command(name="consultarboletim")
async def cmd_consultarboletim(ctx: commands.Context, *, numero: str = ""):
    """Consulta um boletim pelo número."""
    if not numero.strip():
        await ctx.reply("❌ Use assim: `!consultarboletim BO-DICOR-20260702-001`")
        return
    boletim = buscar_boletim_numero(numero)
    if not boletim:
        await ctx.reply("❌ Boletim não encontrado.")
        return
    await ctx.reply(texto_consulta_boletim(boletim))


@bot.command(name="painelmesas")
@commands.has_permissions(administrator=True)
async def cmd_painelmesas(ctx: commands.Context):
    """Abre o painel de mesas pelo comando !painelmesas."""
    await ctx.send(embed=embed_painel_mesas_padrao(), view=PainelMesasView())


@bot.command(name="painelprocurados")
@commands.has_permissions(administrator=True)
async def cmd_painelprocurados(ctx: commands.Context):
    """Abre o painel de procurados pelo comando !painelprocurados."""
    await ctx.send(embed=embed_painel_procurados_padrao(), view=PainelProcuradosView())


@bot.command(name="criarmesa")
@commands.has_permissions(administrator=True)
async def cmd_criarmesa(ctx: commands.Context, *, texto: str = ""):
    """Cria mesa direto por texto. Ex: !criarmesa Baiano | Olimpo | Observação"""
    partes = [p.strip() for p in texto.split("|")]
    if len(partes) < 2 or not partes[0] or not partes[1]:
        await ctx.reply("❌ Use assim: `!criarmesa Apelido | Família | Observação opcional`\nExemplo: `!criarmesa Baiano | Olimpo | Mesa inicial`")
        return
    apelido = partes[0]
    familia = partes[1]
    observacao = partes[2] if len(partes) >= 3 else ""
    await criar_mesa_por_texto(ctx, apelido, familia, observacao)


@bot.command(name="fecharmesa")
async def cmd_fecharmesa(ctx: commands.Context, *, motivo: str = "Fechada por comando de texto"):
    if isinstance(ctx.author, discord.Member) and not usuario_pode_fechar_mesa(ctx.author):
        await ctx.reply(mensagem_sem_permissao_fechar_mesa())
        return
    await fechar_mesa_por_texto(ctx, motivo)


@bot.command(name="catalogo")
async def cmd_catalogo(ctx: commands.Context):
    await ctx.reply(f"📄 **Catálogo de Procurados:**\n{CATALOG_PUBLIC_URL}")


@bot.command(name="comandos")
async def cmd_comandos(ctx: commands.Context):
    await ctx.reply(
        "📌 **Comandos disponíveis:**\n\n"
        "**Slash:** `/painelprocurados`, `/painelmesas`, `/painelboletim`, `/painelorganizacoes`, `/criarmesa`, `/fecharmesa`, `/backup`, `/estatisticas`, `/consultarboletim`\n"
        "**Plano B por texto:** `!painelprocurados`, `!painelmesas`, `!painelboletim`, `!painelorganizacoes`, `!criarmesa Apelido | Família | Observação`, `!fecharmesa`, `!catalogo`, `!consultarboletim NUMERO`"
    )


@bot.command(name="synccomandos")
@commands.has_permissions(administrator=True)
async def cmd_synccomandos(ctx: commands.Context):
    """Força sincronização dos slash commands no servidor atual."""
    if ctx.guild is None:
        await ctx.reply("Use dentro do servidor.")
        return
    guild_obj = discord.Object(id=ctx.guild.id)
    try:
        bot.tree.copy_global_to(guild=guild_obj)
        comandos = await bot.tree.sync(guild=guild_obj)
        await ctx.reply(f"✅ Slash commands sincronizados neste servidor: `{len(comandos)}`. Agora aperte `Ctrl + R` no Discord.")
    except Exception as e:
        await ctx.reply(f"❌ Erro ao sincronizar: `{e}`")


# =====================================================
# SLASH COMMANDS - TODOS
# =====================================================


@bot.tree.command(name="painelorganizacoes", description="Envia a tabela compartilhada das 56 organizações.")
async def painelorganizacoes(interaction: discord.Interaction):
    garantir_56_organizacoes()
    await interaction.response.send_message(
        embed=embed_painel_organizacoes_padrao(),
        view=PainelOrganizacoesView(),
    )


@bot.tree.command(name="painelboletim", description="Envia o painel de boletins de ocorrência.")
async def painelboletim(interaction: discord.Interaction):
    if not isinstance(interaction.user, discord.Member) or not usuario_tem_equipe(interaction.user):
        await interaction.response.send_message(
            "❌ Apenas a equipe autorizada pode enviar este painel.",
            ephemeral=True,
        )
        return
    await interaction.response.send_message(
        embed=embed_painel_boletim_padrao(),
        view=PainelBoletimView(),
    )


@bot.tree.command(name="vincularrg", description="Vincula um RG a um usuário do Discord.")
@app_commands.describe(usuario="Usuário que receberá o RG", rg="RG do usuário")
@app_commands.checks.has_permissions(administrator=True)
async def vincularrg(
    interaction: discord.Interaction,
    usuario: discord.Member,
    rg: str,
):
    vincular_rg_usuario(usuario.id, rg)
    await interaction.response.send_message(
        f"✅ RG `{rg}` vinculado a {usuario.mention}.",
        ephemeral=True,
    )


@bot.tree.command(name="removervinculorg", description="Remove o RG vinculado de um usuário.")
@app_commands.describe(usuario="Usuário que terá o RG removido")
@app_commands.checks.has_permissions(administrator=True)
async def removervinculorg(
    interaction: discord.Interaction,
    usuario: discord.Member,
):
    removido = remover_rg_usuario(usuario.id)
    mensagem = (
        f"✅ RG de {usuario.mention} removido."
        if removido
        else f"⚠️ {usuario.mention} não possuía RG vinculado."
    )
    await interaction.response.send_message(mensagem, ephemeral=True)


@bot.tree.command(name="consultarboletim", description="Consulta um boletim pelo número.")
@app_commands.describe(numero="Ex: BO-DICOR-20260702-001")
async def consultarboletim(
    interaction: discord.Interaction,
    numero: str,
):
    boletim = buscar_boletim_numero(numero)
    if not boletim:
        await interaction.response.send_message(
            "❌ Boletim não encontrado.",
            ephemeral=True,
        )
        return
    await interaction.response.send_message(
        texto_consulta_boletim(boletim),
        ephemeral=True,
    )


@bot.tree.command(name="painelprocurados", description="Envia o painel de gerenciamento de procurados.")
@app_commands.checks.has_permissions(administrator=True)
async def painelprocurados(interaction: discord.Interaction):
    await interaction.response.send_message(embed=embed_painel_procurados_padrao(), view=PainelProcuradosView())


@bot.tree.command(name="painelmesas", description="Envia o painel de gerenciamento de mesas.")
@app_commands.checks.has_permissions(administrator=True)
async def painelmesas(interaction: discord.Interaction):
    await interaction.response.send_message(embed=embed_painel_mesas_padrao(), view=PainelMesasView())


@bot.tree.command(name="catalogo", description="Mostra o link do catálogo de procurados.")
async def catalogo(interaction: discord.Interaction):
    await interaction.response.send_message(f"📄 **Catálogo de Procurados:**\n{CATALOG_PUBLIC_URL}", ephemeral=True)


@bot.tree.command(name="retirarprocurado", description="Retira um procurado pelo RG.")
@app_commands.describe(rg="RG do procurado", motivo="Motivo da retirada")
async def retirarprocurado(interaction: discord.Interaction, rg: str, motivo: str = "Não informado"):
    await retirar_procurado(interaction, rg, motivo)


@bot.tree.command(name="sincronizarcatalogo", description="Importa procurados antigos do canal oficial para o catálogo.")
@app_commands.checks.has_permissions(administrator=True)
async def sincronizarcatalogo(interaction: discord.Interaction):
    await sincronizar_catalogo_core(interaction)


@bot.tree.command(name="regerarcatalogo", description="Recria o HTML do catálogo.")
@app_commands.checks.has_permissions(administrator=True)
async def regerarcatalogo(interaction: discord.Interaction):
    gerar_catalogo_html()
    await interaction.response.send_message(f"✅ Catálogo recriado: {CATALOG_PUBLIC_URL}", ephemeral=True)


@bot.tree.command(name="criarmesa", description="Cria uma mesa de investigação.")
@app_commands.describe(apelido="Apelido do agente", familia="Nome da família/organização", observacao="Observação opcional")
async def criarmesa(interaction: discord.Interaction, apelido: str, familia: str, observacao: str = ""):
    await criar_mesa_core(interaction, apelido, familia, observacao)


@bot.tree.command(name="fecharmesa", description="Fecha a mesa do canal atual.")
@app_commands.describe(motivo="Motivo do fechamento")
async def fecharmesa(interaction: discord.Interaction, motivo: str = "Fechada por comando"):
    await fechar_mesa_core(interaction, motivo)


@bot.tree.command(name="reabrirmesa", description="Reabre uma mesa fechada.")
@app_commands.describe(canal="Canal da mesa")
async def reabrirmesa(interaction: discord.Interaction, canal: discord.TextChannel):
    await reabrir_mesa_core(interaction, canal)


@bot.tree.command(name="minhaassinatura", description="Salva sua assinatura por arquivo ou texto para aparecer no dossiê.")
@app_commands.describe(
    arquivo="Opcional: imagem da assinatura em PNG/JPG/WebP",
    texto="Opcional: assinatura por texto, caso não queira usar imagem"
)
async def minhaassinatura(
    interaction: discord.Interaction,
    arquivo: Optional[discord.Attachment] = None,
    texto: Optional[str] = None,
):
    texto_limpo = str(texto or "").strip()

    if arquivo is None and not texto_limpo:
        await interaction.response.send_message(
            "❌ Envie um **arquivo de imagem** ou escreva um **texto de assinatura**.",
            ephemeral=True,
        )
        return

    if arquivo is not None and (not arquivo.content_type or not arquivo.content_type.startswith("image/")):
        await interaction.response.send_message("❌ O arquivo precisa ser uma imagem PNG, JPG ou WebP.", ephemeral=True)
        return

    await interaction.response.defer(ephemeral=True, thinking=True)
    try:
        dados = carregar_assinaturas_dossie()
        chave = f"agente_{interaction.user.id}"

        # Se trocar de imagem para texto, remove a imagem antiga para o texto aparecer no dossiê.
        registro_antigo = dados.get(chave, {}) if isinstance(dados, dict) else {}
        if arquivo is None:
            caminho_antigo = caminho_assinatura_registrada(registro_antigo)
            if caminho_antigo:
                try:
                    caminho_antigo.unlink()
                except Exception:
                    pass

        registro = {
            "nome": "",
            "texto": texto_limpo[:400],
            "atualizado_por": str(interaction.user),
            "atualizado_em": agora_br(),
        }

        if arquivo is not None:
            destino = ASSINATURAS_DOSSIE_DIR / f"agente_{interaction.user.id}"
            caminho = await salvar_imagem_assinatura(arquivo, destino)
            registro["arquivo"] = caminho_relativo_base(caminho)

        dados[chave] = registro
        salvar_assinaturas_dossie(dados)

        if arquivo is not None and texto_limpo:
            msg = "✅ Sua assinatura foi salva com **arquivo** e **texto reserva**. No dossiê, a imagem terá prioridade."
        elif arquivo is not None:
            msg = "✅ Sua assinatura por **arquivo** foi salva. Ela aparecerá como **Agente Responsável** nos próximos dossiês."
        else:
            msg = "✅ Sua assinatura por **texto** foi salva. Ela aparecerá como **Agente Responsável** nos próximos dossiês."

        await interaction.followup.send(msg, ephemeral=True)
        await enviar_log(f"✍️ Assinatura do agente atualizada: {interaction.user} | Tipo: {'arquivo' if arquivo else 'texto'}")
    except Exception as erro:
        await enviar_log(f"❌ Erro ao salvar assinatura do agente {interaction.user}: {erro}")
        await interaction.followup.send("❌ Não consegui salvar a assinatura. Veja os logs do Railway.", ephemeral=True)


@bot.tree.command(name="assinaturadicor", description="Configura assinatura oficial por arquivo ou texto.")
@app_commands.describe(
    tipo="Tipo de assinatura",
    arquivo="Opcional: imagem da assinatura em PNG/JPG/WebP",
    texto="Opcional: assinatura por texto, caso não queira usar imagem"
)
@app_commands.choices(tipo=[
    app_commands.Choice(name="Delegado Geral", value="delegado_geral"),
    app_commands.Choice(name="Delegado DICOR", value="delegado_dicor"),
])
async def assinaturadicor(
    interaction: discord.Interaction,
    tipo: app_commands.Choice[str],
    arquivo: Optional[discord.Attachment] = None,
    texto: Optional[str] = None,
):
    if not isinstance(interaction.user, discord.Member) or not usuario_pode_fechar_mesa(interaction.user):
        await interaction.response.send_message("❌ Apenas a ADM da DICOR pode configurar essas assinaturas.", ephemeral=True)
        return

    texto_limpo = str(texto or "").strip()

    if arquivo is None and not texto_limpo:
        await interaction.response.send_message(
            "❌ Envie um **arquivo de imagem** ou escreva um **texto de assinatura**.",
            ephemeral=True,
        )
        return

    if arquivo is not None and (not arquivo.content_type or not arquivo.content_type.startswith("image/")):
        await interaction.response.send_message("❌ O arquivo precisa ser uma imagem PNG, JPG ou WebP.", ephemeral=True)
        return

    await interaction.response.defer(ephemeral=True, thinking=True)
    chave = tipo.value
    try:
        dados = carregar_assinaturas_dossie()

        # Se trocar de imagem para texto, remove a imagem antiga para o texto aparecer no dossiê.
        registro_antigo = dados.get(chave, {}) if isinstance(dados, dict) else {}
        if arquivo is None:
            caminho_antigo = caminho_assinatura_registrada(registro_antigo)
            if caminho_antigo:
                try:
                    caminho_antigo.unlink()
                except Exception:
                    pass

        registro = {
            "nome": "",
            "texto": texto_limpo[:400],
            "atualizado_por": str(interaction.user),
            "atualizado_em": agora_br(),
        }

        if arquivo is not None:
            destino = ASSINATURAS_DOSSIE_DIR / chave
            caminho = await salvar_imagem_assinatura(arquivo, destino)
            registro["arquivo"] = caminho_relativo_base(caminho)

        dados[chave] = registro
        salvar_assinaturas_dossie(dados)

        if arquivo is not None and texto_limpo:
            msg = f"✅ Assinatura de **{tipo.name}** salva com **arquivo** e **texto reserva**. No dossiê, a imagem terá prioridade."
        elif arquivo is not None:
            msg = f"✅ Assinatura de **{tipo.name}** salva por **arquivo**."
        else:
            msg = f"✅ Assinatura de **{tipo.name}** salva por **texto**."

        await interaction.followup.send(msg, ephemeral=True)
        await enviar_log(f"✍️ Assinatura oficial atualizada: {tipo.name} | Por: {interaction.user} | Tipo: {'arquivo' if arquivo else 'texto'}")
    except Exception as erro:
        await enviar_log(f"❌ Erro ao salvar assinatura {chave}: {erro}")
        await interaction.followup.send("❌ Não consegui salvar a assinatura. Veja os logs do Railway.", ephemeral=True)


@bot.tree.command(name="verassinaturas", description="Mostra quais assinaturas do dossiê estão configuradas.")
async def verassinaturas(interaction: discord.Interaction):
    registros = carregar_assinaturas_dossie()
    agente_chave = f"agente_{interaction.user.id}"

    def status(chave: str, titulo: str) -> str:
        reg = registros.get(chave, {}) if isinstance(registros, dict) else {}
        tem_texto = bool(str(reg.get("texto", "") or "").strip())
        tem_img = bool(caminho_assinatura_registrada(reg))
        nome = str(reg.get("nome", titulo) or titulo)
        return f"• **{titulo}:** {nome} | Texto: {'✅' if tem_texto else '❌'} | Imagem: {'✅' if tem_img else '❌'}"

    texto = "\n".join([
        status("delegado_geral", "Delegado Geral"),
        status("delegado_dicor", "Delegado DICOR"),
        status(agente_chave, "Minha assinatura"),
    ])
    await interaction.response.send_message(f"📋 **Assinaturas configuradas:**\n{texto}", ephemeral=True)


@bot.tree.command(name="apagarminhaassinatura", description="Apaga sua assinatura de agente responsável.")
async def apagarminhaassinatura(interaction: discord.Interaction):
    dados = carregar_assinaturas_dossie()
    chave = f"agente_{interaction.user.id}"
    reg = dados.pop(chave, None)
    if reg:
        caminho = caminho_assinatura_registrada(reg)
        if caminho:
            try:
                caminho.unlink()
            except Exception:
                pass
        salvar_assinaturas_dossie(dados)
        await interaction.response.send_message("✅ Sua assinatura foi apagada.", ephemeral=True)
    else:
        await interaction.response.send_message("⚠️ Você ainda não tinha assinatura salva.", ephemeral=True)


@bot.tree.command(name="historicomesa", description="Mostra histórico das mesas.")
async def historicomesa(interaction: discord.Interaction):
    await historico_mesa_core(interaction)


@bot.tree.command(name="backup", description="Gera backup dos dados do bot.")
@app_commands.checks.has_permissions(administrator=True)
async def backup(interaction: discord.Interaction):
    await backup_core(interaction)


@bot.tree.command(name="estatisticas", description="Mostra estatísticas do sistema.")
async def estatisticas(interaction: discord.Interaction):
    await estatisticas_core(interaction)

# =====================================================
# PAINEL ADMINISTRATIVO / ESTATÍSTICAS POR AGENTE
# =====================================================

ADMIN_STATS_HISTORY_LIMIT = env_int("ADMIN_STATS_HISTORY_LIMIT", 0)
ADMIN_STATS_SCAN_ALL_CHANNELS = str(os.getenv("ADMIN_STATS_SCAN_ALL_CHANNELS", "1") or "1").strip().lower() not in {"0", "false", "nao", "não", "no"}


def limite_historico_admin():
    """0 ou negativo = varrer tudo que o Discord permitir no canal."""
    try:
        limite = int(ADMIN_STATS_HISTORY_LIMIT or 0)
    except Exception:
        limite = 0
    return None if limite <= 0 else limite


def usuario_pode_painel_adm(member: discord.Member) -> bool:
    """Painel restrito para Inspetor DICOR para cima."""
    return isinstance(member, discord.Member) and usuario_pode_fechar_mesa(member)


def texto_sem_markdown_admin(texto: Any) -> str:
    texto = str(texto or "")
    texto = texto.replace("**", "").replace("__", "").replace("`", "")
    texto = re.sub(r"<@!?(\d+)>", r"ID:\1", texto)
    return re.sub(r"\s+", " ", texto).strip()


def conteudo_mensagem_admin(msg: discord.Message) -> str:
    partes = [msg.content or ""]
    for embed in msg.embeds:
        if embed.title:
            partes.append(str(embed.title))
        if embed.description:
            partes.append(str(embed.description))
        for field in embed.fields:
            partes.append(f"{field.name}: {field.value}")
    return "\n".join([p for p in partes if p]).strip()


def data_msg_br(msg: discord.Message) -> str:
    try:
        tz = datetime.timezone(datetime.timedelta(hours=-3))
        return msg.created_at.astimezone(tz).strftime("%d/%m/%Y %H:%M")
    except Exception:
        return agora_br()


def nome_usuario_relatorio(guild: Optional[discord.Guild], usuario_id: str, fallback: str = "") -> str:
    if guild and str(usuario_id).isdigit():
        membro = guild.get_member(int(usuario_id))
        if membro:
            return str(membro.display_name or membro.name)
    return str(fallback or f"Usuário {usuario_id}").strip()


def limpar_nome_admin_stats(nome: Any) -> str:
    texto = nome_operacional_dossie(nome)
    texto = texto.lower()
    texto = unicodedata.normalize("NFKD", texto).encode("ascii", "ignore").decode("ascii")
    texto = re.sub(r"[^a-z0-9]+", " ", texto)
    return re.sub(r"\s+", " ", texto).strip()


def localizar_membro_por_nome_admin(guild: Optional[discord.Guild], nome: Any) -> Optional[discord.Member]:
    if guild is None or not nome:
        return None

    texto_original = str(nome or "")
    m_mention = re.search(r"<@!?(\d+)>", texto_original)
    if m_mention:
        membro = guild.get_member(int(m_mention.group(1)))
        if membro:
            return membro

    alvo = limpar_nome_admin_stats(nome)
    if not alvo:
        return None

    for membro in guild.members:
        candidatos = {
            limpar_nome_admin_stats(membro.display_name),
            limpar_nome_admin_stats(membro.name),
            limpar_nome_admin_stats(str(membro)),
            limpar_nome_admin_stats(nome_operacional_dossie(membro)),
        }
        if alvo in candidatos:
            return membro

    for membro in guild.members:
        nome_membro = limpar_nome_admin_stats(nome_operacional_dossie(membro))
        if alvo and nome_membro and (alvo in nome_membro or nome_membro in alvo):
            return membro

    return None


def resolver_autor_admin(guild: Optional[discord.Guild], usuario_id: Any = None, nome: Any = "", *extras: Any) -> tuple[str, str]:
    """
    Resolve quem acionou o bot.
    Prioridade: ID salvo > menção em texto > nome salvo > Sem autor.
    """
    candidatos_nome = [nome, *extras]

    if usuario_id is not None and str(usuario_id).strip().isdigit() and str(usuario_id).strip() not in {"0", ""}:
        uid = str(usuario_id).strip()
        return uid, nome_usuario_relatorio(guild, uid, str(nome or ""))

    for candidato in candidatos_nome:
        m_mention = re.search(r"<@!?(\d+)>", str(candidato or ""))
        if m_mention:
            uid = m_mention.group(1)
            return uid, nome_usuario_relatorio(guild, uid, str(candidato or ""))

    for candidato in candidatos_nome:
        membro = localizar_membro_por_nome_admin(guild, candidato)
        if membro:
            return str(membro.id), str(membro.display_name or membro.name)

    nome_limpo = ""
    for candidato in candidatos_nome:
        nome_limpo = nome_operacional_dossie(candidato)
        if nome_limpo and nome_limpo != "Não informado":
            break

    if nome_limpo:
        return f"nome:{slugify(nome_limpo)}", nome_limpo

    return "sem_autor", "Sem autor identificado"


def categoria_relatorio_admin(tipo: Any, texto: str = "") -> str:
    tipo_txt = str(tipo or "").lower().strip()
    texto_upper = str(texto or "").upper()
    if "TOCAIA" in texto_upper or tipo_txt == "tocaia":
        return "Tocaias"
    if "OLB" in texto_upper or tipo_txt == "olb":
        return "OLBs"
    if "PERÍCIA EXTERNA" in texto_upper or "PERICIA EXTERNA" in texto_upper or tipo_txt in {"pericia_externa", "pericia", "pericia_ext"}:
        return "Perícias externas"
    return "Relatórios operacionais"


def add_item_estatistica(agentes: Dict[str, Dict[str, Any]], usuario_id: Any, nome: str, categoria: str, item: Dict[str, Any]) -> None:
    if usuario_id is None or str(usuario_id).strip() in {"", "0", "None"}:
        return
    uid = str(usuario_id)
    agentes.setdefault(uid, {"id": uid, "nome": nome or f"Usuário {uid}", "itens": {}, "_vistos": set()})
    if nome and (not agentes[uid].get("nome") or agentes[uid]["nome"].startswith("Usuário ")):
        agentes[uid]["nome"] = nome

    chave = str(item.get("key") or item.get("titulo") or item.get("data") or "")
    chave = f"{categoria}:{chave}"
    if chave in agentes[uid].setdefault("_vistos", set()):
        return
    agentes[uid]["_vistos"].add(chave)
    agentes[uid]["itens"].setdefault(categoria, []).append(item)


def normalizar_texto_admin_bruto(texto: Any) -> str:
    """Limpa markdown sem destruir menções, mantendo quebras de linha para extração."""
    texto = str(texto or "")
    texto = texto.replace("**", "").replace("__", "").replace("`", "")
    texto = texto.replace("｜", "|").replace("–", "-").replace("—", "-")
    return texto


def extrair_label_admin(texto: str, rotulos: List[str]) -> str:
    """Extrator reforçado: pega labels com emoji, markdown e variações de pontuação."""
    texto_limpo = normalizar_texto_admin_bruto(texto)

    for rotulo in rotulos:
        # Ex.: 🆔 RG: 123 | **RG:** 123 | RG - 123
        padrao = rf"(?:^|\n)\s*(?:[\W_]*\s*)?(?:{rotulo})\s*(?:[:\-])\s*(.+?)(?=\n|$)"
        m = re.search(padrao, texto_limpo, flags=re.I)
        if m:
            valor = re.sub(r"[*_`]+", "", m.group(1).strip())
            return valor[:220]

    # Segunda tentativa: procura o rótulo no meio da linha e corta depois dos dois pontos.
    for linha in texto_limpo.splitlines():
        linha_s = linha.strip()
        if not linha_s:
            continue
        for rotulo in rotulos:
            if re.search(rotulo, linha_s, flags=re.I):
                partes = re.split(r"[:\-]", linha_s, maxsplit=1)
                if len(partes) == 2:
                    valor = re.sub(r"[*_`]+", "", partes[1].strip())
                    return valor[:220]
    return ""


def identificar_autor_admin(guild: Optional[discord.Guild], texto: str, *fallbacks: Any) -> tuple[str, str]:
    """Tenta descobrir o agente responsável por menção, label textual ou nome."""
    texto_limpo = normalizar_texto_admin_bruto(texto)

    # Prioriza labels de responsabilidade para não pegar menções aleatórias.
    labels = [
        r"Respons[áa]vel", r"Perito respons[áa]vel", r"Perito", r"Comunicante",
        r"Cadastrado por", r"Publicado por", r"Autor", r"Agente", r"Criado por",
    ]
    for label in labels:
        padrao = rf"{label}\s*[:\-]\s*(<@!?(\d+)>)"
        m = re.search(padrao, texto_limpo, flags=re.I)
        if m:
            uid = m.group(2)
            return uid, nome_usuario_relatorio(guild, uid, m.group(1))

    for label in labels:
        valor = extrair_label_admin(texto_limpo, [label])
        if valor:
            uid, nome = resolver_autor_admin(guild, None, valor)
            if uid != "sem_autor":
                return uid, nome

    # Se não encontrou label, usa a primeira menção do texto.
    m_mention = re.search(r"<@!?(\d+)>", texto_limpo)
    if m_mention:
        uid = m_mention.group(1)
        return uid, nome_usuario_relatorio(guild, uid, m_mention.group(0))

    # Fallbacks do banco interno.
    for fb in fallbacks:
        uid, nome = resolver_autor_admin(guild, None, fb)
        if uid != "sem_autor":
            return uid, nome

    return "sem_autor", "Sem autor identificado"


def extrair_numero_relatorio_admin(texto: str) -> str:
    texto_limpo = normalizar_texto_admin_bruto(texto)
    m = re.search(r"N[ºO]?\s*([0-9]{1,6})", texto_limpo, flags=re.I)
    if m:
        return m.group(1).zfill(3)
    m = re.search(r"RELAT[ÓO]RIO.*?(\d{1,6})", texto_limpo, flags=re.I | re.S)
    if m:
        return m.group(1).zfill(3)
    return ""


def extrair_numero_boletim_texto(texto: str) -> str:
    m = re.search(r"\bBO\s*-?\s*DICOR\s*-?\s*(\d{1,6})\b", texto, flags=re.I)
    if m:
        return f"BO-DICOR-{int(m.group(1)):03d}"
    m = re.search(r"BOLETIM(?:\s+DE\s+OCORR[ÊE]NCIA)?\s*[—\-–:]\s*(\d{1,6})", texto, flags=re.I)
    if m:
        return f"BO-DICOR-{int(m.group(1)):03d}"
    return ""


async def obter_canal_texto_admin(guild: discord.Guild, canal_id: int) -> Optional[discord.TextChannel]:
    if not canal_id:
        return None
    canal = guild.get_channel(int(canal_id))
    if canal is None:
        try:
            canal = await bot.fetch_channel(int(canal_id))
        except Exception:
            canal = None
    return canal if isinstance(canal, discord.TextChannel) else None


def canal_estatistica_permitido(canal: Any) -> bool:
    return isinstance(canal, discord.TextChannel)


async def canais_varredura_admin(guild: discord.Guild, ids_preferidos: List[int]) -> List[discord.TextChannel]:
    """Canais usados pela auditoria.
    Primeiro usa os canais oficiais. Depois, se ADMIN_STATS_SCAN_ALL_CHANNELS=1,
    varre todos os canais de texto acessíveis do servidor para contar registros antigos.
    """
    canais: List[discord.TextChannel] = []
    vistos = set()

    for canal_id in ids_preferidos:
        canal = await obter_canal_texto_admin(guild, int(canal_id or 0))
        if canal and canal.id not in vistos:
            canais.append(canal)
            vistos.add(canal.id)

    if ADMIN_STATS_SCAN_ALL_CHANNELS:
        for canal in guild.text_channels:
            if canal.id not in vistos and canal_estatistica_permitido(canal):
                canais.append(canal)
                vistos.add(canal.id)

    return canais


async def coletar_procurados_discord(guild: discord.Guild) -> List[Dict[str, Any]]:
    """Varre canais de procurados para contar registros antigos e atuais.
    Quando o autor não existir no texto/banco, entra como Sem autor identificado.
    """
    registros: List[Dict[str, Any]] = []
    vistos = set()

    for canal in await canais_varredura_admin(guild, [PROCURADOS_CHANNEL_ID, HISTORICO_PROCURADOS_ID]):
        canal_id = canal.id
        try:
            async for msg in canal.history(limit=limite_historico_admin(), oldest_first=False):
                texto = conteudo_mensagem_admin(msg)
                if not texto:
                    continue
                upper = texto.upper()
                if not any(p in upper for p in ["MANDADO", "PROCURADO", "PROCURAÇÃO INVESTIGATIVA", "PROCURACAO INVESTIGATIVA"]):
                    continue

                rg = extrair_label_admin(texto, [r"RG", r"🆔\s*RG"])
                nome = extrair_label_admin(texto, [r"Nome", r"👤\s*Nome", r"Identifica[çc][ãa]o do procurado"])
                numero_boletim = extrair_label_admin(texto, [r"N[úu]mero do boletim", r"Boletim vinculado", r"Boletim"])
                if numero_boletim:
                    try:
                        numero_boletim = normalizar_boletim_procurado(numero_boletim) or numero_boletim
                    except Exception:
                        pass

                key_base = limpar_rg(rg) or str(msg.id)
                key = f"procurado-rg-{key_base}"
                if key in vistos:
                    continue
                vistos.add(key)

                uid, autor_nome = identificar_autor_admin(guild, texto)
                registros.append({
                    "key": key,
                    "autor_id": uid,
                    "autor_nome": autor_nome,
                    "data": data_msg_br(msg),
                    "titulo": f"{nome or 'Procurado sem nome'} | RG {rg or 'Não informado'} | Boletim: {numero_boletim or 'Não informado'}",
                    "rg": rg,
                    "nome": nome,
                })
        except Exception as erro:
            await enviar_log(f"⚠️ Estatística ADM: não consegui varrer procurados do canal `{canal_id}`: {erro}")

    return registros


async def coletar_boletins_discord(guild: discord.Guild) -> List[Dict[str, Any]]:
    """Varre boletins no canal oficial e, quando ativado, em todos os canais do servidor."""
    registros: List[Dict[str, Any]] = []
    vistos = set()

    for canal in await canais_varredura_admin(guild, [BOLETINS_CHANNEL_ID]):
        try:
            async for msg in canal.history(limit=limite_historico_admin(), oldest_first=False):
                texto = conteudo_mensagem_admin(msg)
                if not texto:
                    continue
                upper = texto.upper()
                if "BOLETIM" not in upper and "BO-DICOR" not in upper:
                    continue

                numero = extrair_numero_boletim_texto(texto) or extrair_label_admin(texto, [r"Boletim", r"N[úu]mero", r"Nº", r"NO"])
                if not numero:
                    numero = f"MSG-{msg.id}"
                key = f"boletim-{str(numero).upper()}"
                if key in vistos:
                    continue
                vistos.add(key)

                comunicante = extrair_label_admin(texto, [r"Comunicante", r"Respons[áa]vel", r"Autor", r"Criado por", r"Registrado por"])
                uid, nome = identificar_autor_admin(guild, texto, comunicante)
                local = extrair_label_admin(texto, [r"Local dos fatos", r"Local", r"Localiza[çc][ãa]o"])

                registros.append({
                    "key": key,
                    "autor_id": uid,
                    "autor_nome": nome,
                    "data": data_msg_br(msg),
                    "titulo": f"{numero} | {local or 'Local não informado'}",
                })
        except Exception as erro:
            await enviar_log(f"⚠️ Estatística ADM: não consegui varrer boletins do canal `{canal.id}`: {erro}")

    return registros


async def coletar_relatorios_operacionais_discord(guild: discord.Guild) -> List[Dict[str, Any]]:
    """Lê os canais oficiais e, se ativado, todos os canais do servidor para pegar relatórios antigos."""
    registros: List[Dict[str, Any]] = []
    vistos = set()
    canais_map = globals().get("CANAIS_RELATORIOS", {}) or {}
    ids_oficiais = [int(canal_id or 0) for canal_id in canais_map.values()]
    tipo_por_canal = {int(v): k for k, v in canais_map.items() if v}

    for canal in await canais_varredura_admin(guild, ids_oficiais):
        tipo = tipo_por_canal.get(canal.id, "")
        try:
            async for msg in canal.history(limit=limite_historico_admin(), oldest_first=False):
                conteudo = conteudo_mensagem_admin(msg)
                if not conteudo:
                    continue
                upper = conteudo.upper()
                if not any(p in upper for p in ["RELATÓRIO", "RELATORIO", "PERÍCIA", "PERICIA", "OLB", "TOCAIA"]):
                    continue

                key = f"relatorio-discord-{msg.id}"
                if key in vistos:
                    continue
                vistos.add(key)

                uid, autor_nome = identificar_autor_admin(guild, conteudo)
                numero = extrair_numero_relatorio_admin(conteudo)
                categoria = categoria_relatorio_admin(tipo, conteudo)
                registros.append({
                    "key": key,
                    "id": key,
                    "tipo": tipo,
                    "categoria": categoria,
                    "tipo_nome": nome_tipo_relatorio(tipo) if tipo else categoria,
                    "numero": numero,
                    "autor_id": uid,
                    "autor_nome": autor_nome,
                    "data": data_msg_br(msg),
                    "mensagem_url": msg.jump_url,
                    "titulo": f"{categoria[:-1] if categoria.endswith('s') else categoria} Nº {numero or 'sem número'}",
                })
        except Exception as erro:
            await enviar_log(f"⚠️ Estatística ADM: não consegui varrer relatórios do canal `{canal.id}`: {erro}")

    return registros


async def coletar_mesas_discord(guild: discord.Guild) -> List[Dict[str, Any]]:
    """Varre categorias de mesas e, se ativado, todos os canais do servidor.
    Só conta canais que tenham mensagem com indicação de mesa/investigação.
    """
    registros: List[Dict[str, Any]] = []
    vistos = set()
    canais: List[discord.TextChannel] = []

    for categoria_id in [CATEGORIA_MESAS_ABERTAS_ID, CATEGORIA_MESAS_FECHADAS_ID]:
        categoria = guild.get_channel(int(categoria_id or 0))
        if isinstance(categoria, discord.CategoryChannel):
            for canal in categoria.text_channels:
                if canal.id not in vistos:
                    canais.append(canal)
                    vistos.add(canal.id)

    if ADMIN_STATS_SCAN_ALL_CHANNELS:
        for canal in guild.text_channels:
            if canal.id not in vistos:
                nome = str(canal.name or "").lower()
                if any(p in nome for p in ["mesa", "investig", "dicor", "🕵"]):
                    canais.append(canal)
                    vistos.add(canal.id)

    for canal in canais:
        try:
            achou_mensagem = False
            async for msg in canal.history(limit=80, oldest_first=True):
                texto = conteudo_mensagem_admin(msg)
                if "Mesa de Investigação" not in texto and "Mesa de Investigacao" not in texto and "Investigação" not in texto and "Investigacao" not in texto:
                    continue
                achou_mensagem = True
                uid, nome = identificar_autor_admin(guild, texto)
                familia = extrair_label_admin(texto, [r"Organiza[çc][ãa]o/Fam[íi]lia", r"Fam[íi]lia", r"Organiza[çc][ãa]o", r"Comunidade"])
                apelido = extrair_label_admin(texto, [r"Apelido", r"Agente"])
                registros.append({
                    "key": f"mesa-canal-{canal.id}",
                    "autor_id": uid,
                    "autor_nome": nome,
                    "data": data_msg_br(msg),
                    "titulo": f"{familia or canal.name} | Apelido: {apelido or 'Não informado'}",
                })
                break
            if not achou_mensagem and canal.category_id in {CATEGORIA_MESAS_ABERTAS_ID, CATEGORIA_MESAS_FECHADAS_ID}:
                registros.append({
                    "key": f"mesa-canal-{canal.id}",
                    "autor_id": "sem_autor",
                    "autor_nome": "Sem autor identificado",
                    "data": "Não informado",
                    "titulo": f"{canal.name} | Apelido: Não informado",
                })
        except Exception:
            continue

    return registros


async def montar_estatisticas_administrativas(guild: discord.Guild, alvo_id: Optional[int] = None) -> str:
    agentes: Dict[str, Dict[str, Any]] = {}
    vistos_globais = set()

    def marcar(chave: str) -> bool:
        if chave in vistos_globais:
            return False
        vistos_globais.add(chave)
        return True

    # Procurados cadastrados pelo bot.
    for p in carregar_procurados():
        rg_limpo = limpar_rg(p.get("rg", ""))
        key = f"procurado-rg-{rg_limpo or p.get('id') or p.get('caso') or p.get('mensagem_id')}"
        if not marcar(key):
            continue
        uid, nome = resolver_autor_admin(
            guild,
            p.get("criado_por_id") or p.get("autor_id"),
            p.get("criado_por_nome") or p.get("autor_nome", ""),
        )
        add_item_estatistica(agentes, uid, nome, "Procurados cadastrados", {
            "key": key,
            "data": p.get("data", "Não informado"),
            "titulo": (
                f"{p.get('nome', 'Sem nome')} | RG {p.get('rg', 'Não informado')} | "
                f"Boletim: {p.get('numero_boletim') or p.get('boletim') or 'Não informado'}"
            ),
        })

    # Procurados antigos varridos no Discord. Quando não der para descobrir autor, cai em Sem autor identificado.
    for p in await coletar_procurados_discord(guild):
        if not marcar(str(p.get("key"))):
            continue
        add_item_estatistica(agentes, p.get("autor_id"), p.get("autor_nome", ""), "Procurados cadastrados", p)

    # Boletins publicados pelo bot.
    for b in carregar_boletins():
        numero = str(b.get("numero") or b.get("id") or b.get("mensagem_id") or "").upper()
        key = f"boletim-{numero or data_caso()}"
        if not marcar(key):
            continue
        uid, nome = resolver_autor_admin(
            guild,
            b.get("criado_por_id") or b.get("comunicante_id") or b.get("autor_id") or b.get("publicado_por_id"),
            b.get("criado_por_nome") or b.get("comunicante_nome", "") or b.get("autor_nome", "") or b.get("publicado_por_nome", ""),
            b.get("comunicante_mention", ""),
        )
        add_item_estatistica(agentes, uid, nome, "Boletins publicados", {
            "key": key,
            "data": b.get("data_criacao") or b.get("data_registro") or b.get("criado_em") or "Não informado",
            "titulo": f"{b.get('numero', 'BO sem número')} | {b.get('local', 'Local não informado')}",
        })

    # Boletins antigos já enviados no canal oficial.
    for b in await coletar_boletins_discord(guild):
        if not marcar(str(b.get("key"))):
            continue
        add_item_estatistica(agentes, b.get("autor_id"), b.get("autor_nome", ""), "Boletins publicados", b)

    # Mesas registradas no JSON.
    for m in carregar_mesas():
        key = f"mesa-canal-{m.get('canal_id') or m.get('nome_canal')}"
        if marcar(key):
            uid, nome = resolver_autor_admin(
                guild,
                m.get("autor_id"),
                m.get("autor_nome", ""),
                m.get("apelido", ""),
            )
            add_item_estatistica(agentes, uid, nome, "Mesas criadas", {
                "key": key,
                "data": m.get("criada_em", "Não informado"),
                "titulo": f"{m.get('familia', 'Sem organização')} | Apelido: {m.get('apelido', 'Não informado')}",
            })

        uid_fechou_val = m.get("fechada_por_id")
        if uid_fechou_val:
            key_fechou = f"mesa-fechada-{m.get('canal_id') or m.get('nome_canal')}"
            if marcar(key_fechou):
                uid_fechou, nome_fechou = resolver_autor_admin(
                    guild,
                    uid_fechou_val,
                    m.get("fechada_por_nome", ""),
                )
                add_item_estatistica(agentes, uid_fechou, nome_fechou, "Mesas encerradas", {
                    "key": key_fechou,
                    "data": m.get("fechada_em", "Não informado"),
                    "titulo": f"{m.get('familia', 'Sem organização')} | Processo: {(m.get('dossie') or {}).get('processo', 'Não informado')}",
                })

    # Mesas antigas lidas nas categorias.
    for m in await coletar_mesas_discord(guild):
        if not marcar(str(m.get("key"))):
            continue
        add_item_estatistica(agentes, m.get("autor_id"), m.get("autor_nome", ""), "Mesas criadas", m)

    # Dossiês gerados.
    for d in carregar_dossies():
        key = f"dossie-{d.get('processo') or d.get('arquivo_pdf') or d.get('gerado_em')}"
        if not marcar(key):
            continue
        uid, nome = resolver_autor_admin(
            guild,
            d.get("encerrado_por_id") or d.get("agente_encerramento_id"),
            d.get("encerrado_por", "") or d.get("agente_encerramento", ""),
        )
        add_item_estatistica(agentes, uid, nome, "Dossiês gerados", {
            "key": key,
            "data": d.get("gerado_em", "Não informado"),
            "titulo": f"{d.get('processo', 'Sem processo')} | {d.get('nome_operacao', 'Operação não informada')}",
        })

    # Relatórios operacionais do JSON, separados por tipo.
    for r in carregar_relatorios_operacionais():
        key = f"relatorio-{r.get('id') or r.get('mensagem_url') or r.get('tipo') or r.get('numero') or r.get('data')}"
        if not marcar(key):
            continue
        uid, nome = resolver_autor_admin(
            guild,
            r.get("autor_id") or r.get("criado_por_id"),
            r.get("autor_nome", "") or r.get("criado_por_nome", ""),
            r.get("autor_mention", ""),
        )
        categoria = categoria_relatorio_admin(r.get("tipo"), r.get("resumo", ""))
        add_item_estatistica(agentes, uid, nome, categoria, {
            "key": key,
            "data": r.get("data", "Não informado"),
            "titulo": f"{r.get('tipo_nome') or nome_tipo_relatorio(r.get('tipo'))} Nº {r.get('numero', '')}".strip(),
        })

    # Relatórios antigos varridos nos canais oficiais.
    for r in await coletar_relatorios_operacionais_discord(guild):
        key = str(r.get("key") or r.get("mensagem_url"))
        if not marcar(key):
            continue
        add_item_estatistica(agentes, r.get("autor_id"), r.get("autor_nome", ""), r.get("categoria") or categoria_relatorio_admin(r.get("tipo")), {
            "key": key,
            "data": r.get("data", "Não informado"),
            "titulo": r.get("titulo") or f"{r.get('tipo_nome') or nome_tipo_relatorio(r.get('tipo'))} Nº {r.get('numero', '')}".strip(),
        })

    if alvo_id is not None:
        chave = str(alvo_id)
        agentes = {
            chave: agentes.get(
                chave,
                {"id": chave, "nome": nome_usuario_relatorio(guild, chave), "itens": {}, "_vistos": set()},
            )
        }

    linhas: List[str] = []
    linhas.append("RELATÓRIO ADMINISTRATIVO DICOR")
    linhas.append("Polícia Federal de Capital Morada do Valley")
    linhas.append(f"Gerado em: {agora_br()}")
    linhas.append("Modelo: auditoria por agente, com banco interno e varredura geral dos canais do servidor.")
    linhas.append("=" * 78)
    linhas.append("")

    ordem_categorias = [
        "Procurados cadastrados",
        "Boletins publicados",
        "Tocaias",
        "OLBs",
        "Perícias externas",
        "Relatórios operacionais",
        "Mesas criadas",
        "Mesas encerradas",
        "Dossiês gerados",
    ]

    total_agentes = len([a for a in agentes.values() if any(a.get("itens", {}).values())])
    total_itens = sum(len(v) for a in agentes.values() for v in a.get("itens", {}).values())
    linhas.append(f"Agentes com registro: {total_agentes}")
    linhas.append(f"Atividades encontradas: {total_itens}")
    linhas.append("")

    agentes_ordenados = sorted(agentes.values(), key=lambda a: (a.get("nome") or "").lower())
    for agente in agentes_ordenados:
        itens = agente.get("itens", {})
        total = sum(len(itens.get(cat, [])) for cat in ordem_categorias)
        if alvo_id is None and total == 0:
            continue

        linhas.append("-" * 78)
        linhas.append(f"AGENTE: {agente.get('nome', 'Sem nome')}")
        linhas.append(f"ID: {agente.get('id')}")
        linhas.append(f"TOTAL DE ATIVIDADES: {total}")
        linhas.append("")

        for categoria in ordem_categorias:
            registros = itens.get(categoria, [])
            linhas.append(f"{categoria.upper()}: {len(registros)}")
            for idx, item in enumerate(registros, 1):
                linhas.append(f"  {idx}. {item.get('data', 'Sem data')} — {item.get('titulo', 'Sem título')}")
            linhas.append("")

    if total_itens == 0:
        linhas.append("Nenhum registro encontrado para esta consulta.")

    linhas.append("=" * 78)
    linhas.append("Observação: registros antigos sem autor salvo, sem menção ou sem nome reconhecível entram como 'Sem autor identificado'. Daqui para frente, procurados e boletins recebem uma linha discreta de auditoria para contar no agente correto.")
    return "\n".join(linhas)


async def gerar_arquivo_estatistica_admin(guild: discord.Guild, membro: Optional[discord.Member] = None) -> Path:
    texto = await montar_estatisticas_administrativas(guild, membro.id if membro else None)
    nome_base = f"relatorio_agente_{membro.id}" if membro else "relatorio_geral_dicor"
    caminho = ADMIN_REPORTS_DIR / f"{nome_base}_{data_caso()}.txt"
    caminho.write_text(texto, encoding="utf-8")
    return caminho


async def enviar_relatorio_estatistica(interaction: discord.Interaction, membro: Optional[discord.Member] = None) -> None:
    if not isinstance(interaction.user, discord.Member) or not usuario_pode_painel_adm(interaction.user):
        await responder_interacao(interaction, "❌ Apenas Inspetor DICOR para cima pode usar este comando.", ephemeral=True)
        return

    guild = interaction.guild
    if guild is None:
        await responder_interacao(interaction, "❌ Use dentro de um servidor.", ephemeral=True)
        return

    if not interaction.response.is_done():
        await interaction.response.defer(thinking=True, ephemeral=True)

    caminho = await gerar_arquivo_estatistica_admin(guild, membro)

    resumo = "📊 **Relatório administrativo gerado.**\n"
    if membro:
        resumo += f"Agente analisado: {membro.mention}\n"
    else:
        resumo += "Tipo: varredura geral, nome por nome.\n"
    resumo += "O TXT mostra apenas quantidades e o que foi criado."

    await interaction.followup.send(
        resumo,
        file=discord.File(str(caminho), filename=caminho.name),
        ephemeral=True,
    )
    await enviar_log(
        "📊 **Relatório administrativo gerado**\n"
        f"Solicitante: {interaction.user.mention} (`{interaction.user.id}`)\n"
        f"Tipo: {'Individual' if membro else 'Varredura Geral'}\n"
        f"Alvo: {membro.mention if membro else 'Todos os agentes'}\n"
        f"Arquivo: `{caminho.name}`\n"
        f"Canal de origem: {getattr(interaction.channel, 'mention', 'Sem canal')}"
    )


class AdminAgenteSelect(discord.ui.UserSelect):
    def __init__(self):
        super().__init__(
            placeholder="Selecionar agente para auditoria individual",
            min_values=1,
            max_values=1,
            custom_id="dic_adm_select_agente_painel",
        )

    async def callback(self, interaction: discord.Interaction):
        if not isinstance(interaction.user, discord.Member) or not usuario_pode_painel_adm(interaction.user):
            await interaction.response.send_message("❌ Apenas Inspetor DICOR para cima pode usar este painel.", ephemeral=True)
            return

        membro = self.values[0]
        if interaction.guild and not isinstance(membro, discord.Member):
            try:
                membro = await interaction.guild.fetch_member(membro.id)
            except Exception:
                pass

        await enviar_relatorio_estatistica(interaction, membro)


class PainelAdministrativoView(View):
    def __init__(self):
        super().__init__(timeout=None)
        self.add_item(AdminAgenteSelect())

    @discord.ui.button(label="Varredura Geral", emoji="📊", style=discord.ButtonStyle.blurple, custom_id="dic_adm_varredura_geral")
    async def varredura_geral(self, interaction: discord.Interaction, button: Button):
        await enviar_relatorio_estatistica(interaction, None)


@bot.tree.command(name="paineladministrativo", description="Envia o painel administrativo da DICOR.")
async def paineladministrativo(interaction: discord.Interaction):
    if not isinstance(interaction.user, discord.Member) or not usuario_pode_painel_adm(interaction.user):
        await interaction.response.send_message("❌ Apenas Inspetor DICOR para cima pode usar este comando.", ephemeral=True)
        return

    embed = discord.Embed(
        title="🏛️ Painel Administrativo DICOR",
        description=(
            "Área restrita para Inspetor DICOR para cima.\n\n"
            "📊 **Varredura Geral:** gera um relatório com todos os agentes, nome por nome.\n"
            "👤 **Selecionar agente:** escolha um membro no menu abaixo para gerar o relatório individual.\n\n"
            "O relatório conta procurados, boletins, tocaias, OLBs, perícias externas, relatórios diários, mesas e dossiês.\n"
            "Ele mostra apenas **quantidades** e **o que foi criado**, sem exibir o texto completo das mensagens."
        ),
        color=discord.Color.dark_blue(),
    )
    embed.set_footer(text="DICOR • Estatística operacional e auditoria interna")
    await interaction.response.send_message(embed=embed, view=PainelAdministrativoView())


@bot.tree.command(name="estatisticaagente", description="Gera relatório completo de atividade de um agente.")
@app_commands.describe(agente="Agente que será analisado. Se vazio, gera varredura geral.")
async def estatisticaagente(interaction: discord.Interaction, agente: Optional[discord.Member] = None):
    await enviar_relatorio_estatistica(interaction, agente)
# =====================================================
# EVENTOS
# =====================================================

comandos_ja_sincronizados = False

class ReabrirMesaView(discord.ui.View):
    def __init__(self):
        super().__init__(timeout=None)

    @discord.ui.button(
        label="Reabrir Mesa",
        emoji="🔓",
        style=discord.ButtonStyle.green,
        custom_id="dic_reabrir_mesa_botao"
    )
    async def reabrir(self, interaction: discord.Interaction, button: discord.ui.Button):
        # Garante que a função interna recebe a interação e o canal de texto corretamente
        await reabrir_mesa_core(interaction, interaction.channel)
@bot.event
async def on_ready():
    global comandos_ja_sincronizados
    
    # 1. ATIVA A PERSISTÊNCIA DOS BOTÕES DE TODOS OS PAINÉIS
    try:
        # Painel de Relatórios e Sub-botões dos canais temporários
        bot.add_view(RelatoriosPainelView())
        bot.add_view(IniciarFormularioRelatorioView(tipo="tocaia"))
        bot.add_view(IniciarFormularioRelatorioView(tipo="olb"))
        bot.add_view(IniciarFormularioRelatorioView(tipo="pericia_externa"))
        
        # Painel de Boletins
        bot.add_view(PainelBoletimView())
        bot.add_view(IniciarBoletimView())
        bot.add_view(InformarRGBoletimView())
        bot.add_view(TipoIdentificacaoBoletimView())
        bot.add_view(ContinuarVeiculoBoletimView())
        bot.add_view(ProvasBoletimView())
        bot.add_view(PreviaBoletimView())
        
        # Outros Painéis do Sistema
        bot.add_view(PainelProcuradosView())
        bot.add_view(FinalizarProcuradoView())
        bot.add_view(PainelMesasView())
        bot.add_view(FecharMesaView())
        bot.add_view(ReabrirMesaView())
        bot.add_view(PainelOrganizacoesView())
        bot.add_view(PainelAdministrativoView())
        
        print("✅ Todas as persistências de botões (Relatórios com Tickets, Boletins e Sistemas) foram carregadas!")
    except Exception as e:
        print(f"Aviso ao carregar persistência dos painéis: {e}")
        
        print("✅ Todas as persistências de botões (Relatórios, Boletins e Sistemas) foram carregadas!")
    except Exception as e:
        print(f"Aviso ao carregar persistência dos painéis: {e}")

    # Lógica de sincronização do bot
    try:
        if GUILD_ID > 0:
            guild = discord.Object(id=GUILD_ID)
            bot.tree.copy_global_to(guild=guild)
            comandos_servidor = await bot.tree.sync(guild=guild)
            
            bot.tree.clear_commands(guild=None)
            comandos_globais = await bot.tree.sync()

            print(f"Comandos sincronizados no servidor: {len(comandos_servidor)}")
            print(f"Comandos globais antigos limpos: {len(comandos_globais)}")
            print("Comandos ativos:", ", ".join(f"/{cmd.name}" for cmd in comandos_servidor))
        else:
            comandos_globais = await bot.tree.sync()
            print(f"Comandos sincronizados globalmente: {len(comandos_globais)}")

        comandos_ja_sincronizados = True
    except Exception as e:
        print(f"Erro ao sincronizar comandos: {e}")

    print('VERSAO COMPLETA: mesas | procurados | catalogo | boletins | organizacoes | relatorios')
    print(f"Bot online como {bot.user}")

# =====================================================
# REGISTRO LOCAL DE RELATÓRIOS OPERACIONAIS
# =====================================================

def carregar_relatorios_operacionais() -> List[Dict[str, Any]]:
    dados = carregar_json(RELATORIOS_OPERACIONAIS_JSON, [])
    return dados if isinstance(dados, list) else []


def salvar_relatorios_operacionais(lista: List[Dict[str, Any]]) -> None:
    salvar_json(RELATORIOS_OPERACIONAIS_JSON, lista[-2000:])


def registrar_relatorio_operacional(registro: Dict[str, Any]) -> None:
    lista = carregar_relatorios_operacionais()
    chave = str(registro.get("id") or registro.get("mensagem_url") or "")
    if chave and any(str(item.get("id") or item.get("mensagem_url") or "") == chave for item in lista):
        return
    lista.append(registro)
    salvar_relatorios_operacionais(lista)


def nome_tipo_relatorio(tipo: str) -> str:
    return {
        "tocaia": "Tocaia",
        "olb": "OLB",
        "pericia_externa": "Perícia Externa",
    }.get(str(tipo or ""), str(tipo or "Relatório"))


# =====================================================
# SISTEMA DE PAINEL DE RELATÓRIOS OPERACIONAIS (ATUALIZADO)
# =====================================================

RELATORIOS_CONTADOR_JSON = DATA_DIR / "relatorios_contador.json"
CATEGORIA_RELATORIOS_TEMPORARIOS = 1490200421661147187

CANAIS_RELATORIOS = {
    "tocaia": 1490200477248520333,
    "olb": 1490200479995789374,
    "pericia_externa": 1490200524367200297
}

def maior_numero_relatorio_local(tipo_relatorio: str) -> int:
    maior = 0
    contadores = carregar_json(RELATORIOS_CONTADOR_JSON, {})
    if isinstance(contadores, dict):
        try:
            maior = max(maior, int(contadores.get(tipo_relatorio, 0) or 0))
        except Exception:
            pass

    for item in carregar_relatorios_operacionais():
        if str(item.get("tipo")) != str(tipo_relatorio):
            continue
        try:
            maior = max(maior, int(str(item.get("numero", "0")).strip() or 0))
        except Exception:
            pass
    return maior


def extrair_numero_relatorio_tipo(texto: str, tipo_relatorio: str) -> int:
    texto = str(texto or "")
    texto_norm = normalizar_busca(texto) if 'normalizar_busca' in globals() else texto.lower()
    chaves = {
        "tocaia": ["relatorio de tocaia", "relatório de tocaia"],
        "olb": ["relatorio de olb", "relatório de olb"],
        "pericia_externa": ["relatorio de pericia externa", "relatório de perícia externa", "pericia externa", "perícia externa"],
    }.get(tipo_relatorio, [tipo_relatorio])

    if not any((normalizar_busca(c) if 'normalizar_busca' in globals() else c.lower()) in texto_norm for c in chaves):
        return 0

    m = re.search(r"N[º°O]?\s*[:#-]?\s*(\d{1,6})", texto, flags=re.I)
    if not m:
        return 0
    try:
        return int(m.group(1))
    except Exception:
        return 0


async def maior_numero_relatorio_canal(tipo_relatorio: str, guild: Optional[discord.Guild]) -> int:
    if guild is None:
        return 0
    canal_id = CANAIS_RELATORIOS.get(tipo_relatorio)
    if not canal_id:
        return 0
    canal = guild.get_channel(canal_id)
    if canal is None:
        try:
            canal = await bot.fetch_channel(canal_id)
        except Exception:
            return 0
    if not hasattr(canal, "history"):
        return 0

    limite_env = os.getenv("RELATORIOS_NUMERO_SCAN_LIMIT", "0").strip()
    limite = None if limite_env in {"", "0"} else int(limite_env) if limite_env.isdigit() else 1000
    maior = 0
    try:
        async for msg in canal.history(limit=limite, oldest_first=False):
            conteudo = coletar_texto_embed(msg) if 'coletar_texto_embed' in globals() else (msg.content or "")
            maior = max(maior, extrair_numero_relatorio_tipo(conteudo, tipo_relatorio))
    except Exception as erro:
        await enviar_log(f"⚠️ Não consegui varrer numeração antiga de `{tipo_relatorio}`: {erro}")
    return maior


async def obter_proximo_numero_relatorio_async(tipo_relatorio: str, guild: Optional[discord.Guild]) -> str:
    # Não reinicia a contagem quando o código muda: usa o contador salvo,
    # os registros locais e também o maior número já publicado no canal oficial.
    maior = maior_numero_relatorio_local(tipo_relatorio)
    maior = max(maior, await maior_numero_relatorio_canal(tipo_relatorio, guild))
    proximo = maior + 1
    contadores = carregar_json(RELATORIOS_CONTADOR_JSON, {})
    if not isinstance(contadores, dict):
        contadores = {}
    contadores[tipo_relatorio] = proximo
    salvar_json(RELATORIOS_CONTADOR_JSON, contadores)
    return f"{proximo:03d}"


def obter_proximo_numero_relatorio(tipo_relatorio: str) -> str:
    maior = maior_numero_relatorio_local(tipo_relatorio)
    proximo = maior + 1
    contadores = carregar_json(RELATORIOS_CONTADOR_JSON, {})
    if not isinstance(contadores, dict):
        contadores = {}
    contadores[tipo_relatorio] = proximo
    salvar_json(RELATORIOS_CONTADOR_JSON, contadores)
    return f"{proximo:03d}"

class RelatoriosPainelView(View):
    def __init__(self):
        super().__init__(timeout=None)

    async def criar_canal_temporario_operacional(self, interaction: discord.Interaction, tipo_nome: str) -> Optional[discord.TextChannel]:
        guild = interaction.guild
        if not guild:
            return None
        
        overwrites = {
            guild.default_role: discord.PermissionOverwrite(view_channel=False),
            interaction.user: discord.PermissionOverwrite(
                view_channel=True, send_messages=True, attach_files=True, embed_links=True, read_message_history=True
            )
        }
        try:
            if 'cargos_equipe_permissoes' in globals():
                overwrites = cargos_equipe_permissoes(guild)
                overwrites[interaction.user] = discord.PermissionOverwrite(
                    view_channel=True, send_messages=True, attach_files=True, embed_links=True, read_message_history=True
                )
        except Exception:
            pass

        categoria = guild.get_channel(CATEGORIA_RELATORIOS_TEMPORARIOS)
        nome_canal = f"relatorio-{tipo_nome}-{interaction.user.name}".lower().replace(" ", "-")
        
        try:
            canal = await guild.create_text_channel(name=nome_canal, category=categoria, overwrites=overwrites, reason="Abertura de canal temporário de relatório.")
            return canal
        except Exception as e:
            print(f"Erro crítico ao criar canal de relatório: {e}")
            return None

    async def gerenciar_abertura_ticket(self, interaction: discord.Interaction, tipo: str, titulo_bonito: str):
        await interaction.response.defer(ephemeral=True)
        canal = await self.criar_canal_temporario_operacional(interaction, tipo)
        if not canal:
            return await interaction.followup.send("❌ Não foi possível criar o canal temporário. Verifique as permissões da categoria.", ephemeral=True)
        
        await interaction.followup.send(f"✅ Canal de relatório criado: {canal.mention}", ephemeral=True)
        
        embed = discord.Embed(
            title=f"📝 {titulo_bonito.upper()} — INSTRUÇÕES",
            description=(
                f"Olá {interaction.user.mention}, o seu canal de relatório provisório foi gerado com sucesso!\n\n"
                "**Siga os passos abaixo para concluir o procedimento:**\n\n"
                "📸 **1. ENVIE AS PROVAS NESTE CANAL**\n"
                "Antes de preencher o formulário, faça o upload de todas as imagens, prints, links ou registros fotográficos necessários diretamente aqui no chat.\n\n"
                "✍️ **2. PREENCHA O FORMULÁRIO A SEGUIR**\n"
                "Clique no botão azul abaixo para abrir o formulário oficial e preencher os dados requeridos. Ao enviar, o relatório será gerado automaticamente."
            ),
            color=discord.Color.blue()
        )
        embed.set_footer(text="DICOR • Procedimento Operacional Padrão")
        await canal.send(content=interaction.user.mention, embed=embed, view=IniciarFormularioRelatorioView(tipo))

    @discord.ui.button(label="👀 TOCAIA", style=discord.ButtonStyle.secondary, custom_id="rel_btn_tocaia")
    async def btn_tocaia(self, interaction: discord.Interaction, button: Button):
        await self.gerenciar_abertura_ticket(interaction, "tocaia", "Relatório de Tocaia")

    @discord.ui.button(label="🚔 OLB", style=discord.ButtonStyle.secondary, custom_id="rel_btn_olb")
    async def btn_olb(self, interaction: discord.Interaction, button: Button):
        await self.gerenciar_abertura_ticket(interaction, "olb", "Relatório de OLB")

    @discord.ui.button(label="🧪 PERÍCIA EXTERNA", style=discord.ButtonStyle.secondary, custom_id="rel_btn_pericia_ext")
    async def btn_pericia_ext(self, interaction: discord.Interaction, button: Button):
        await self.gerenciar_abertura_ticket(interaction, "pericia_externa", "Perícia Externa")


class IniciarFormularioRelatorioView(View):
    def __init__(self, tipo: str = "tocaia"):
        super().__init__(timeout=None)
        self.tipo = tipo

    @discord.ui.button(label="✍️ Preencher Formulário", style=discord.ButtonStyle.primary, custom_id="rel_btn_preencher")
    async def preencher(self, interaction: discord.Interaction, button: Button):
        tipo_atual = self.tipo
        nome_canal = interaction.channel.name if interaction.channel else ""
        if "olb" in nome_canal:
            tipo_atual = "olb"
        elif "pericia" in nome_canal or "externa" in nome_canal:
            tipo_atual = "pericia_externa"
        elif "tocaia" in nome_canal:
            tipo_atual = "tocaia"

        if tipo_atual == "tocaia":
            await interaction.response.send_modal(TocaiaModal())
        elif tipo_atual == "olb":
            await interaction.response.send_modal(OlbModal())
        elif tipo_atual == "pericia_externa":
            await interaction.response.send_modal(PericiaExternaModal())

async def finalizar_e_postar_relatorio(interaction: discord.Interaction, tipo: str, texto_conteudo: str):
    canal_id = CANAIS_RELATORIOS.get(tipo)
    canal_destino = interaction.guild.get_channel(canal_id) if interaction.guild and canal_id else None
    
    anexos_encontrados = []
    canal_atual = interaction.channel
    if isinstance(canal_atual, discord.TextChannel):
        async for msg in canal_atual.history(limit=50, oldest_first=True):
            for anexo in msg.attachments:
                anexos_encontrados.append(anexo.url)
    
    if anexos_encontrados:
        texto_conteudo += "\n\n**🔗 REGISTROS FOTOGRÁFICOS / PROVAS EM ANEXO:**\n" + "\n".join([f"• {url}" for url in anexos_encontrados])
    
    mensagem_publicada = None
    if canal_destino:
        mensagem_publicada = await canal_destino.send(
            texto_conteudo,
            allowed_mentions=discord.AllowedMentions(users=False, roles=False, everyone=False),
        )

    try:
        numero_match = re.search(r"Nº\s+([0-9]+)", texto_conteudo, flags=re.I)
        registrar_relatorio_operacional({
            "id": str(mensagem_publicada.id if mensagem_publicada else f"TEMP-{data_caso()}"),
            "tipo": tipo,
            "tipo_nome": nome_tipo_relatorio(tipo),
            "numero": numero_match.group(1) if numero_match else "",
            "autor_id": interaction.user.id,
            "autor_nome": str(interaction.user),
            "autor_mention": interaction.user.mention,
            "data": agora_br(),
            "canal_destino_id": canal_id,
            "mensagem_url": mensagem_publicada.jump_url if mensagem_publicada else "",
            "resumo": cortar_discord(texto_conteudo.replace("*", ""), 800),
            "provas": anexos_encontrados,
        })
    except Exception as erro:
        await enviar_log(f"⚠️ Relatório publicado, mas não consegui registrar estatística: {erro}")
    
    if isinstance(canal_atual, discord.TextChannel):
        await canal_atual.send("✅ Relatório enviado com sucesso para o canal oficial! Este canal provisório será apagado em 5 segundos...")
        await asyncio.sleep(5)
        try:
            await canal_atual.delete(reason="Relatório operacional concluído e armazenado.")
        except Exception:
            pass

class TocaiaModal(Modal, title="Relatório de Tocaia"):
    local = TextInput(label="Local", placeholder="Local de interesse observado", max_length=150)
    tempo = TextInput(label="Tempo de tocaia", placeholder="Duração da vigilância", max_length=100)
    infos = TextInput(label="Informações obtidas", style=discord.TextStyle.paragraph, placeholder="Detalhes colhidos durante a tocaia", max_length=1000)

    async def on_submit(self, interaction: discord.Interaction):
        await interaction.response.defer(ephemeral=True)
        num = await obter_proximo_numero_relatorio_async("tocaia", interaction.guild)
        data_hora = agora_br()
        texto = (
            f"━" * 15 + f"\n**👀 RELATÓRIO DE TOCAIA Nº {num}**\n" + f"━" * 15 + "\n"
            f"**👤 RESPONSÁVEL:** {interaction.user.mention}\n"
            f"**📍 LOCAL OBSERVADO:** {self.local.value}\n"
            f"**⏱️ TEMPO DE VIGILÂNCIA:** {self.tempo.value}\n\n"
            f"**📝 INFORMAÇÕES OBTIDAS:**\n{self.infos.value}\n\n"
            f"📅 *Enviado em {data_hora.split()[0]} às {data_hora.split()[1]}*"
        )
        await finalizar_e_postar_relatorio(interaction, "tocaia", texto)

class OlbModal(Modal, title="Relatório de OLB"):
    dicors = TextInput(label="DICORs envolvidos", placeholder="Agentes participantes da operação", max_length=200)
    relato = TextInput(label="Relatório da emboscada", style=discord.TextStyle.paragraph, placeholder="Como ocorreu a operação", max_length=1000)
    itens = TextInput(label="Itens ilegais apreendidos", style=discord.TextStyle.paragraph, placeholder="Lista de itens retidos", max_length=500)
    pericia = TextInput(label="Houve perícia?", placeholder="Sim / Não", max_length=150)
    prejuizo = TextInput(label="Prejuízo estimado", placeholder="Valor aproximado do impacto", max_length=100)

    async def on_submit(self, interaction: discord.Interaction):
        await interaction.response.defer(ephemeral=True)
        num = await obter_proximo_numero_relatorio_async("olb", interaction.guild)
        data_hora = agora_br()
        texto = (
            f"━" * 15 + f"\n**🚔 RELATÓRIO DE OLB Nº {num}**\n" + f"━" * 15 + "\n"
            f"**👤 RESPONSÁVEL:** {interaction.user.mention}\n"
            f"**👥 AGENTES ENVOLVIDOS:** {self.dicors.value}\n\n"
            f"**💥 RELATO DA OPERAÇÃO:**\n{self.relato.value}\n\n"
            f"**📦 APREENSÕES:**\n{self.itens.value}\n\n"
            f"**🔬 REALIZADA PERÍCIA:** {self.pericia.value}\n"
            f"**💰 PREJUÍZO ESTIMADO:** ≈ {self.prejuizo.value}\n\n"
            f"📅 *Enviado em {data_hora.split()[0]} às {data_hora.split()[1]}*"
        )
        await finalizar_e_postar_relatorio(interaction, "olb", texto)

class PericiaExternaModal(Modal, title="Perícia Externa"):
    codigo = TextInput(label="Código da ocorrência", placeholder="Número de referência", max_length=100)
    local = TextInput(label="Local", placeholder="Local da perícia", max_length=150)
    suspeito = TextInput(label="Suspeito", placeholder="Nome ou descrição", max_length=150, required=False)
    conclusao = TextInput(label="Conclusão", style=discord.TextStyle.paragraph, placeholder="Resultados das análises", max_length=1000)

    async def on_submit(self, interaction: discord.Interaction):
        await interaction.response.defer(ephemeral=True)
        num = await obter_proximo_numero_relatorio_async("pericia_externa", interaction.guild)
        data_hora = agora_br()
        texto = (
            f"━" * 15 + f"\n**🔬 RELATÓRIO DE PERÍCIA EXTERNA Nº {num}**\n" + f"━" * 15 + "\n"
            f"**👤 PERITO:** {interaction.user.mention}\n"
            f"**🔢 CÓD. OCORRÊNCIA:** {self.codigo.value}\n"
            f"**📍 LOCALIZAÇÃO:** {self.local.value}\n"
            f"**👤 SUSPEITO ALVO:** {self.suspeito.value or 'Não identificado'}\n\n"
            f"**📋 CONCLUSÃO DO LAUDO:**\n{self.conclusao.value}\n\n"
            f"📅 *Enviado em {data_hora.split()[0]} às {data_hora.split()[1]}*"
        )
        await finalizar_e_postar_relatorio(interaction, "pericia_externa", texto)


@bot.tree.command(name="painel-relatorios", description="Envia o painel com os Relatórios Operacionais.")
async def painel_relatorios(interaction: discord.Interaction):
    if not usuario_tem_admin(interaction.user):
        return await interaction.response.send_message("❌ Apenas membros autorizados podem usar este comando.", ephemeral=True)
        
    embed = discord.Embed(
        title="📑 PAINEL DE RELATÓRIOS OPERACIONAIS",
        description=(
            "Selecione uma das opções abaixo para iniciar um procedimento operacional.\n\n"
            "👀 **TOCAIA:** Registrar vigilâncias em locais de interesse.\n"
            "🚔 **OLB:** Registro de emboscadas e operações planejadas.\n"
            "🧪 **PERÍCIA EXTERNA:** Análises e levantamentos periciais externos."
        ),
        color=discord.Color.blurple()
    )
    embed.set_footer(text="DICOR • Sistema Operacional Seguro")
    await interaction.response.send_message(embed=embed, view=RelatoriosPainelView())

# =====================================================
# GERADORES DO DOSSIÊ: PDF PROFISSIONAL + DOCX EDITÁVEL
# =====================================================


# Brasões padrão embutidos no próprio bot.
# Caso queira trocar depois, basta colocar arquivos brasao_pf.png e brasao_dicor.png na pasta do bot.
DOSSIE_BRASAO_PF_PADRAO_B64 = """/9j/4AAQSkZJRgABAQAAAQABAAD/2wBDAAMCAgICAgMCAgIDAwMDBAYEBAQEBAgGBgUGCQgKCgkICQkKDA8MCgsOCwkJDRENDg8QEBEQCgwSExIQEw8QEBD/2wBDAQMDAwQDBAgEBAgQCwkLEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBD/wAARCAJtAmwDASIAAhEBAxEB/8QAHgAAAQQDAQEBAAAAAAAAAAAAAAUGBwgDBAkCAQr/xABiEAABAwMCAwUEBQcIAwoMAgsBAgMEAAURBgcSITEIE0FRYRQicYEJIzKRoRVCUmKCscEWM0NykqLR8CRzshclNFNjg7O0wuEYJig1RFSTo6TD0vE2RmR0J0VVVoSUZtPi/8QAHAEAAQUBAQEAAAAAAAAAAAAAAAMEBQYHAgEI/8QARBEAAQMDAQUFBgQFAwMEAQUAAQACAwQFESEGEjFBURMiYXGBBxQykaGxI0LB0RVSYuHwJDNyFoLxJTRDstI1U2Nzov/aAAwDAQACEQMRAD8A5VUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFfeFQ6pIzXzB8qEIooooQiiiihCKKKKEIooooQiiiihCKKKKEIoor20y684lpptSlq5AAZNCF5GRSpG0/KfiGSs92SMoQRzPxp+bWbRXXWFwUsobYjRMKlTZB4Y0RP6Sj+crxCRnofAEiQZW7+2+g5idBaWsKL3pmUA1qCbIT9bcFDHNHMABDmVpHhyAPF75QfNk7rEq2POrlXF5h5hwsvIKFJ6gisfPHSp41/tPaLrala327mKu+mnB3nfJIVJgeaHxgEhP6WOQxnGQTClytUu1vFiQ30JCVDmFY64NdxyCQLl7C1aVFGCOoopRcIooooQiiiihCKKKKEIooooQiiiihCKKKKEIoowfI0YNCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRQAfKhCMHyp0aE0TdtYXyHardb3ZT0twNstJH2z4lRPIISASpR5AA5IHOtbSOlLpqW5x4cC3vyVvLCGmWkFSnVc/dA8ehz6A1PWr9TWbYHTD+jNLym5Gubkz3d0ntFJRbWv/AFdkjPPrk5yeSjj3MIySY7o4pRjeZ4Lzqbs4RZ0GS3oXUlv1NcbUhLdzgR+FuSy/jKgyCohwAHkDhR5gDiBAgK6aYudtWtC4zuWyQtCklLiMeCkkZB/zyqXdP9mztO23QEPfjTui74LRLdWpmVDVmXwD+nLCT3vdE5wvh4Tg88YJzxd+dN6ujt2PezSAlPt4ZF7tyAzOZA5cSh0URzynknl9nPOm8c2pDHB2OOOSULWu4jCgMhSeRBFeRnrViZ2xVk1jD/KO12q7VqhKwpQjd4Ik5KR4cKsJzjqF8BPgDUU37bPUFjUEz4cmCtRIQ3MYWyVEeAJAB+X305bO12nBJujLUzKK3ZVnuUMnvobgA8QOIfeK0iCDggg0qk0UUUUIRRRRg+VCEUUYJ6VtR7ZcJPNiI4oHx4cD8a8yAjC1RX0DJ5A057Ht/fr5JMSFCkynh1ZiMqeXz6ckjlUx2/YO0aNt4vG6mpbbpZlKEOBhxwSZ7iT+i2gkJPkMqUOWUiknzsYu2xudwUG2jTU65uoQGXAFkAJSgqWrPQJSOZqbIe0Wm9uLXG1JuxcFWplzm1aGSFXGXgjKVD+jB8fEdDwHBrQu2+Gm9Ktqs+yuk1RHiC0b5ckh6e9k9Uj7KP6oHD5JzzqRNoexRu1vGVbh7z317RumVnvHLjeifbJac5w0wvB4cE4WvhT+iFHAplWVsdJH2tS8Mb480vDEXu3YhvFRrdNT7h79X+DtftHpJ6Nb5LgYh2S1pP1vLm4+vlnAGVKUeEAEk8iqrT6b7BmyGldHuaH3c1+67uFf0I7idbVKMKxLTkoSrkEuBZPCpSyARzTj7Zm3QGktMbfaXc0RsDppWnLW6nu52pZqOK5XQ8sqBVzAPh9lAHRAPOnTZ9HWe121y2CMJCZIPtK5A7xclR6qcUrmonPj09PDK797QQ14gtxwBz5lWmg2eLhv1WueXRc3tabZb5di/WyVz08dtme9HnsBblsubQ6pUcAhQBwUnhWnOehBOdqHtdv28pzTciNpPVLyR3lomcoMo5JKmilOUqHkkeWEZClV0NuVtkWmyytN3OxM6y0bLGJNguKe9U0nHPuVHny8BnPIcJBxVUN2Po+oOoIknWvZj1AuchBLz2lrm+Gp0XHMoacWQFFPglZBOBhSzyqxWLbWjugDap3ZycM8io+vs09J3oxvN+yqLq/a3VGk7k7bLtanoUhv3u5ewQpPgptYylxJH5ySR6+TNeYeYcLbzS0KHIhQxU+2/erXW35XtzvXoxd+hRVBLlvvcdTU2JywOFak8Y5dM8/0VJrfRtttfu20ZG12tY8acQCLDf1d2+hSj9ht4AlQ58iOL9ZQ5mr02oIGXajqoN0QJ7vHoq3Uc6k7Wexer9GPPM360TrYWVlKlyG+NhR8OF5BKFA+ijTBl2K6QyUuRVLH6SPeH4U4bI1/ApIscOKT6K+qQpBIWkpI8CK+YPlXa5RRRRg+VCEUUUYPlQhFFZmIkqQeFmO4s/qpzSzbdH3ea6hpTfdlw4CccSyfIJHMmvC4DivQCUgYPSty32ubcHAiO0SDyKiMAfOpY0xsReHmhdb62xZ7alXCqXd1hhPL9FvHEr4kBJ8xSkrWG0+3ThVZIi9ZXdoEJefR3MBhWeXAgfbIx1OR5Y60i6fkwZSgj5lJu32yE+7D8s3qVHtVmZ5u3OdhLYwRlLSFHLh59eSR0JGRnV3s2lmaHviUxFpmQZjXtNunMpHdzI5+yRgqAWOik5OOR5gglx2XbztM9qyRLuGlNJXW7W63sOv4Yb7iA0ltOS22VEIW4egQCVqPh1x52o3EgG3ydld2++Ytq3FNQZjyeF60TBlIB4uiOLkQQMAqBIBJCAlJcSHZI4gcl3hp0woFIPPlivlSLuvtnedE6glQ5sbCk/XIdQPq5LR6Ot4J90+I8CCDzBqO1CnjHh4yEg4FpwV8ooorpeIooooQiiiihCKKKKEIooooQiiiihCKKKKEIooooQiiiihCKKKKEIooooQiiiihCKKKKEIooooQiiiihCKKKKEIooooQiiigAk4FCEY5UrafsEy+S0x4zTi+NYbCW0lSlqJACUgA5JJFYbNaX7tKDLYIbBAWrHTP8TU+Q2rV2ftNt3u6Rm3dZz2VC2W2QgEQG1DhU88j9M5V7p6A468WEZJdzut4rtjM6ngtl+92/s56Z9mhhl/X92jcJ5JWm0R1Dp4grV55ycHlw/bkbsZ9kxncZxW/W+jEtOjYj/eQIjvF32oJYJ91OfeLQV9oge8Twg4CiEXsjdlu47/AOoJm8m8r8yPoK3SFPS5Tyyly8Shj6hpXUpzjjUAenAOZyOmFltirjNh3WXbWLZbbW0I9hs7TQQ1AjgcKVFI5cfDywPsjHmcUHa7amOyRGmgdmZ30CnbXbTXP7SQdwfVZmbJrO4Pt6ma1E5Y5gR3US2MNpVCixgMIYU3gBRA6qGMHknkBUS7zdmbZzeLie3U26TYb7wd2jU+mz3YUf0nkgYX/WcQpXUcWKsI28ke8Pu/hnxrOlxLg5/H5VjVNfa6ilM9PIQ4lWmWkgmbuSMyPsuWm5X0a262lSvUux2sYGs4CDxNtR3xDuTQ/qlXdqxyxwuZ9Kieduz2iNonG9Jbo6bflsRlYEPUtrKyodCO8UOJwf1ioeWK7GSdC2CXJ9uiNO26ZxcftUBwsrB8zjAJ+RpOv2kNR3C1yLJNcsuqbVJ/nLfe4KVocB6g8ihf7SBV9t/tGnADK+IP8RoVCT2Jmcwvx5rjwzufsXqN15eqNubrpyU+cmRYZfEyk+YjuEIQPRAFKSNA7AX6MJFi3pjRn3Oka72spx6FRKR8xmr/AOuexf2eNXK47zsRO00+QQuRpiRhvPn3aTwpHoG/lUOan+jP2XuLik6V3kvunXT0Yu1sRKz8VcTBH3GrZS7dWWc957mHxGijpbTVs4tDvJVmT2Y2rvHK7Br/AG+uOTkJZvTjbnyC28D4cWKTpHZP1+kHuLLEkpT+dGvUQg/2nB+6plvP0W+4scKVpXd7Qt0RklPfvSIilDwynu1gH9o/Gm279Gf2lmiRHm6RkgdCzezj5cSBUzHtJbJBllW35pm6hqB8UTlHUPsqbgOrAd0y20P0nr1DCR8cO5pUHZWl25tTl+1foa1JxkiZe1rWn4d2hQ/Gng19Gh2mXFf6RK0owPFTt6yB9yD+6l6z/RfbpPEfyl3U0Fak5ye7lyX1AefCGUgn5iupNo7Wwd6qb816KCoPCJyjNvbfYrTzK16j3utq1tjIatNuU8o+gWCon5pFJkrcPs+6bS23p/Q9/wBUymlcQdvEv2aOT/q2T749Fj4irLWH6NTam1vpOsd+Z92/Sj2mzpZz6d4px1X9ypW0r2O+zHo15L1v2mvmrXhjCr9JWlkq9QOAEfFGKhKrbuyUxx2xef6QnkVmrZODA3zVGEby74bmSl6X2t0n+TG5gDKYGmLSEuBPkFNpy38U4z41JegPo793tUqGp98tVwtEWviCnjcJQlXJ4dcpaCuEHOAe8cSfQ9K6A2Zm9WOyjTul7RYNHWdJ9yDZoKEAfHhSlGfXhUfWsKNJwpEv8oXlT10l8QUl2YvvSkjpgHkMeeKqFx9p7gCy3xBvidSpen2Z3tal+fAKLNpNitk9nSl3anb/APlJfkpLS9T6g+sQ2eWVMoUOFs5/ObQF/rHqJIcsEq73BF51fcV3mY2eJpDiAmMyfNLX2c+qgT49TTnQzhACgMDpWJxA54HSs3uF8rrq/tKiQnzVkpqKnpRuxNwtINDxHIdP8/486+KSkHIrMvCTzNYlkH7JFRQdnVOsI4Uq5EZPn1pOuGn4kp5FxiuuwJ7We7lxFlt1J9SPtD4jwpQSSDjlWRByDXrZHR6gr3A4FMrcDSul90bT/JzfDb2FquFz7m7Qmu5uEQgdUqQUqSPMJVwK55Saq3uH9Gmqe07qDs87iw7yG8uizXZfsssHqEtu4CFEdPrA38Tyq8UdgKOVJyMdMV6f01AuT3frYLcnAw+2SlfLp7w6/MVabRttdLN3Y37zeh1CiKyy0tV3iMHqFytuep+1X2dyvTuurPemrctQJjXyMZMZYHg28ScJ9EL4T5Vhb3s2c1bOVN3D2iNrluJAXL008IwB/SMfKGlH1IJrrUzadWoYfjpujF3hyUd27EurPeJdR+ipXPiHopKqizWPZB7P2uGSrUWwBs0pZVxzdLr7niJ/O7pspbHx4K0i2+0qgqRirjLXdW6qt1NgnhP4Tw4dFz1/k12dNVsKXZ92nbY+r7Me82sKKf6y8oT80qNa6ezJDuzantO7iaCuiF8wlu6usu/cpspHw4jVpdVfRk7IXJ7g0rvFfdLOkZEe8W5Mkfflkp/vVHd5+ip3JRlWjN4tBXlJ+z7RJfiKUPVPdrA/tVbKXam01AyyoA8DoVFyUFQzR0fyUDzuy9rqO53bNhakJH50a7xVJPrzWD+Fa3/g1a45JOlZCSfFVziAf9JUty/ove1XFViE1peaB0VGvicf3kpNYE/Rjdrx4gPWmyIH6Tl+bI/fUi280ThkVLfmE3NPINOzKjyP2Yr8hsSLhM09awnmfb702PvDfFX2JtdtvYnODUu6ulo3ASSILa5iv7xSc/BNS1A+iu7QKnUDUOs9v7Q0ftqevTi1J+SGTn76eFp+jC0TbFoVrztM2viz78Sy2ovqV6Bxbowf+bPwpvLtDbY879QD5apRlHO7QRFVok6i2D04XG4y9SamWB7nd8MFg/tYDg+VYmd8tR5Ys212grZZXVe60YsUypayfNagVLPqriNX90b2AezFY1Mvs6I3A1u+OY9rJixHPioBrA+GetWI2+2fkaMS7H27220VtrCcSErXEiokznPVSkhPEf6yz86g6nba3RZELXSH5BOWWydw/EIaFy+0L2Ie1jvvKTeNWxZGn7W6C8u6aqlKjoQOvuM+858AlAHLwq1G0HYK2F0DJYlXCJct3dRx1IUGUt+z2iO6P0+fCtOfBSlA+KfCrpJ21tkhxqRqe5T78410bmO4YR8GUYQB6YNORhpiIyiPGZbZZbTwoQhISlAHQAAYAqpXHbW4VQ3IyI2+HH5qQgtkEeru8foo8b281su1M+yalh6acgEO221WWOlmDHcH2UuAD6xHgUgBAzyBqk3bT7Jn+61bLvvjt3YFW3XtmTxax0yynjMpKUjM2OBzUSAFHGQoc/tA8XR0yPexn4etNLWenp8x9jVOlH2oupLYnEdxYw3KbBJMd0eKFZOP0SSfE1EWfaCW11PbtOQeOvEfunNRSipZuOAHTwXGva/X1n3M0+xs1uTMEaTGJGn7y5zXHcwQGF56pPTBxnABOUoxE+5G3160XepduucFUd+Mr65vqkjqHEHxSocwR88HlV1O2/2SYl2gXHtE7LWRyKphxbms9ONp+stkge8uS2lP9GeZVywMcSeXFwwZozWVt3800xt3ree3E1fbGS3Ybs+rlLRzPszxP2iTy5Ec1cQBUFBe20FfFWQippjlp4+BVZmicxxik0cFWkg8+WMV8p0620Xd9K3STAuEB6M9GcLb7C0YLSuv3EEEHoQQRyIprEVLNcHDITMgtOCiiiiul4iiiihCKKKKEIooooQiiiihCKKKKEIooooQiiiihCKKKKEIooooQiiiihCKKKKEIooooQiiiihCKKKKEIooowT0oQvoFbdrtsm5yQxHSemVK8Ep8SaxQoj8yQmOy2VKV15dB1JqdtEabsG22mU7ia3jFxlZKrVbl8nLg8M8LigefdDB8MHkckDDiUkm55rtjN4rc05arDsnpljXOq4qJN4lN95YLQ9yUtRAKZTo6gDkRyHIg8yU8K/2ZOzrqrtW69ma/wBwZ0iFoi0yA9fbuR3Yc4RxCIwcYCynA90EISQcZKQUbZDZjcXtl7rSJVwmLhWKDiVfLwtBEe3xeLIbRk8PGRnhT1OFKPJJNdMtKaZ0zCsNq0RoS1i16B0+gIt8Ug95PdH2n3lYBVkjJJA4iSSOgFK2q2mjsEBDSDO4aeAUzbba64PxwjHPqlqw2a2SIlrgWe0NWjSVgaQzYbO0jgSlCRydWM81Hwz5knnkl4oWACM9OX+H+fjSYlzpywE8hwjAFZ0OkDBz51871NZLWyumlOXE65V8ZCyJoawYA4JSS6TjJPLyrbadGE8zjFJjLhwPeHM+dZWn+EkE8sfKuA7Gi8c1LbT2AQOlbrK85JIpFbdyQRgefP4Vutv8uRB8OVKtekHMylNCwASSPjWR1iNJTwSGUOp8lpBH3Gk9LwwSSfkaYO4XaW2K2ojqd11uXZYLyTwmC2/7TMB//V2uJwfNIHrUpSUtRWENhYXHwGUzmkZCMuOFIL2ldOPj37JC5+KWAn91aatEaYycWpr4AkfxqomuPpVdmbE8tjRuk9QalUBydc4ILRPnlfErHPxT8qr7rD6UXtAapeciaC0pYdPx3Mpa4Izk+WP23Dwfc38qtNNsJdKzvSRhg/q0TEXYZ3IsuPQDK6au6L00VcIt7eemMn/GkjUEHb/ScQztSSLFaYyRxd9cn2mEfHLpArkndd2+2PuaXBeNe6njsO/bSuWm2tY+CeAY+VNB7Z/VNwk+0am1jEW8s8S3C89KXk/rYwT+1Ui32e0rdKuqA8GjKmqWi2grhmmo3kdToPquol+7VnZh00FoO7mm3FNjmi3OGTn0HcJUn8ajO+fSG9nC2rU3Gm325cHMKiWvkr/2ik1RZjZKwMkLn6jmyEnmUsxktg/NRP7qWI+2W3ENIzZrhLWP/WZ3L7kJH76ds2M2ch0e57/LRTkGxm1dSMuaxg8TkqyeoPpMds2iRp/bjUlwJHuqlSmIgz8AHaZcn6Tm7+8LXs1b0eRl3px38ENIqNIOktDxknu9HWoeXGlxw/3lmlVFjszKQYmmLc36ogNn/sk0+isezcA3RSl3mT+6kW+zi9OGZ6xrPIJYuX0lO7UxPDb9AaQieR7qW6R972Pwpvv9v7tDTeTESws58GrWs/vWaVI7LzWC3ELafDgi8A/ACtr8qPpGEyXBj9c06bR2SDRlE31TqD2XVU4JdcM9d0f3TUd7bHabk8234iR+pZkn94NYv/DF7Uv2hISR6WRH/wBNO4XKSo4Mx3r/AMYa8O3R1CSozFoABUpS3SAkDxJJwB5k8sUoG2zg2jZ8ks72U9m0ufXvAHgE1R2z+1AyfrJDWPHisiP/AKazx+3b2i4qgXkWh0A/0lqI/cRWxD1ou7RjNstr1RcIoWpAkw7ZIdZUQfBSQR69fEetfJuvZVniLmXS16thRG8d49ItT6G056ZUoYHMjrTv+GUjx3re0+irT9mbK0kC88EtWv6SDea3KH5S0bpOeB144slsn+y8P3U7rP8ASk32OtKr1sza3hnn7HeHWeXoFtufvqKrZrW3ao41wnH320gHikx0AL/q5JJ+6vT0CyPKK5NktjpPP3obR/HhzTCa02EuxPQgHzwpKH2dVlVEJ6K4hzDwODhWj019KxtYeH+UO1epYJGMmLMjygPhkNmpJ0/9KJ2Yp7rbcxvU9tCjhSpVrSoJ9T3a1VQGTp7SUlX1uk7WU9SUoWjl+wsVoyNAbfPjiXY5kXxzFmFP4LSr99JNsOz2e7C5vkcpnUezXaHP4c7H+eQuslm7cHZC1Qw2h3eDT7Hef0d0SuMB8e+QE/jUk6bOxm4kUTtJyNHX5lYz3lqejv4z590TiuHM3avTD5Jt99uMUeCZDKHfvUkp/dSWdrLpCkd9Y9XRO9T0Wouxlgj9bGB/apz/ANP2eUDs348wq7U7FbUUTj+Dvf8AE5Xep3aTb1xX/mBlGfBLi0fgDWE7P6AB/wDMQUB5vuH96q4pae3C7Xm3TCTpTcjU6o7XNLcO8GYgD/V8S/xFSdpb6TbtbaFUmLqhNn1C22RxJvNoLTpHlxslog+pCqbSbHRya072n6KAqIrpQnFVE9voV1jZ2q2+injRpK3rI/41rvP9o0rQbHZLZj8mWiFD8vZ46G+X7IFUE0Z9MFoyd3EXX+1F0tayQHZNtnIlNp8yG1pQr5cRqxm3Xbp7MG6EkQbNulb7ZMUB/o98zblFR8At7Daj4YCzUJV7O1tHxh06gJtHWiTRzj81PDmMnPXHzNYVKCcnIrXi3WBco6JlvmMyY7oCkOMuBaFg+KSMgj4GvLruc48ardUXw6PGCnsQD9Qcr6t3IPOtJbuM8zXtxRA6/GtR1fPAqGllT+OML6p7ByCaxl08zk48eWawuOcP314LppkZSCnTWAhN/UtnnxJ38r9LstKuiG+6lxHP5q5Rx1aX4Zx9kkenSuZPbS7J0XRSl9oPY6O//I+VIBu1vaH1unp3EAUkDmloqKcZxwqUB9kpNdVFOE/a6fwpmaosZhPTb3AtTNygXBhUe+2dxILdwjqGFEJPLvAkkfrDlzzVn2Y2olstRh2sZ4j9kyr7a2sZp8Q59Vyfs9+tnaO0umy3h9iLuJaY3BFecKUovMZPPu1nlhYGSFZ6jpwqUUV91LpubYpbrMiM60pp1TTrbqSFtOJOFIUDzBBBGDzqz3a+7LcvYm+Qt5Nn5UuToG9Se/gSW/eds8o8/ZnSPs4IUEkgHlwnJBJRGza+0jpZdxgIYY3AtcYJuEIYAuzKRhLjfLPeADlgnrw4GUit+o6yOoibUU5yx30VNmjcHlkgw4KsJGKKU73Zn7VIKShRaUo8CiDkY/NPkRSZg4zipMHIyE1IwiiiivV4iiiihCKKKKEIooooQiiiihCKKKKEIooooQiiiihCKKKKEIooooQiiiihCKKKKEIooooQiiiihCKyMMuPupZaQVLUcADxrwhKlKCUgkkgADxqXNrduY8huTqPU0owLNbhxzZJPCSRz7lsnqrB5kZx0wSQDw94YMldNaXHAShtroiw2KzP691ust2iAoFLauS574PJlHjwA44iMknIBGCobmhND7m9sbd1jT1hjqRFQnjddKcRLTARgFxZ6DkAB4qUQkDoBotRNedqHcm07d7d2NRYChGt8JlHCzEYTyU+6eiUgcyo8gAB8elu1O1GlNlNE/7jW3b4eCiHdV6gSMO3CURzZbI6NgcgPzUqIzxKVVU2i2hhsFOZpNZHfCOnipS30ElfJ2bPhHEpY0JoXR+itIQNo9s4vs+krMoflC4H+fvk0ABbrihjKSoHOOXIJTyFSAwGYzaUNoCUAYCQMYHh6VqRWGYzLceO2ltDYCUpHQAeny61sJUkED0x0r5wudxnuc7p6h2XErQ6enjpoxFGMALYChxZwSD4VmbUBzB5n1rWCgrwHxzWG4XK32eC9dLrPjw4kUBT0iQ6lppsealqOB86YsY6QhrBknolHFrBlxSslzHL3qyd8G08RPIc+LiH35qpW8H0hG0+hS7Z9BtvayuqeJKnY5LMFlQ5YLqhxOc/BCSnr71VN112ke0z2gFLgRrlJtFlWChUK05iRSn/AJV0kKc5eC1keQFXi17CXOuAlqMQs6u4/JRhuImk7GkYZH9GjK6N7mdqnY3aBp5vWGuoi7iykK/JtvPtMxR8BwI5IP8AXUmqrbh/SmXN4vQdpdvWWsgpanXtZdWkn84MtKAz6FahVYbbsxAZw9qi/LlvHmtqF0B8i6sYJ/qjx60+7HYbNaGeDTmn48Yp5KkBsuvE+riskfs4q6UWy9htmC5pneOujfQKx0mw1/uQEla9tPGevxJH1Duv2s971GTqDWl7bgODHD3qbbCCT191HAlQ5+Sj6mm5C2RZQ5m/6oDq1nicEJsq4v214z8cGpGcRLQDIeS6tB6uEk5+JrA5J4hkEc8YNWH+IyRDs6ZgYOgGFcbX7NLHAO0qSZndSdPkEm2nQegrI0O50y3OeByHrg8p0j14ElKPvBpdbkMW+KuUlcO1wmcd48Eoistg+fCAD6DmT4A0lSrpHhsh1xDkhxxxLEaMyCXpT6jhDTYHMqJP3A/A397KX0edvSxbd1e0zao92vy0Jft2lVHjt1oQeaUut9Hnh1PEVJBPMqIyO6aknuH4kzzu+PPySW020to2JYKaggaZjywO74k8fRUk0VpDcrdWUtO0m0+rNYcDnAZrEMRYYPge/eHCAf1sGpdidiDtiTGRJXtHpaMcZ7mTqlJdHoS2rhz+FdUNQap0LtrZPa9QXa2WG2RUAAuKS022nwAT0A5eGBTI012ruz1qu7osli3SsUmc84Gm2kSUqK15wAMH8KlmW6jj0DM+ZWS1ftD2jq3bwn3R0aMBcmNzdqt79nmlTN0tmNQWG3BQSbpCcauUJJPQqcZ+wOv2jn0NMVNxbcZbkRX2n47oy28yriQsehxnw6HmPHGK76PR4N3hKakNMy4slvhUlYC23EK6g56giuT30iPZZtnZ6vUbeHbC0+y6K1LNEW82xofUW2aoZQ60P6NtzChge6FcgACkBGotEMjN6nG6RyznKtGy/tRrWVLae7neYSBvcx59R1VVtT3+XCscx6I+tp0NKCVoVgjIxyq+3Zb+jv2X3c2S0ZubrzVWvp87UFqZnSYrd8DUZLi/tBIS3xBP7ZrnTqlwyLFKUFAgtnmK7cdg4j/wRtrj4jT0YH7jXdrjbHEdATnokfanXzOr4uycQ3d5E4Oqox9In2TtkOzfttpPUO1mn7hCusy/tsPTJV2kyVqa7tauHhcWUjmAc4zyqtNqvi5sVtxwnjKQFfKr8fTFhK9m9G4HP+UrWP8A2Ltc5LQ93cVtOegpG8sD2sdjqpj2R1cmJw9xIOOfgnoJpA6/jSLq5L13sT9rYdUgSOEKCVlJVz6E+WccjkVg9qJOQfxr0JAP2v31BxgxvDxyWxVjWVlO+CT4XAgpwdlnthbpdmGUm226e9fdFpfV7XZVq9+IskFa2ufuk45jPCrxAIBCd2ke1Tuv2irkubrm+ybVpRKyuFp5h092oA/bdSMd4vkCc+6nGE8PjLXZP7FNn7UWx2qdSaevg09rmxanfYgT3eJcV9kssq7mQ2AcpyVEKSCRxHIUOQm7e7sN7c9mbsXbgahdc/lLrqXBjNzL7Lbz3STJay3GbJw0jPiMqPLJwMC7NBd3t7918lSzUNNK4NjJeCQB+XQ4z19FQvTBiRoDTsQn30JVk8iAR0pdVOYZjuTrhKbixWMcbrnPn4JSBzUo45Ac+vQAkNbThxbI4B4fq0/j1q6n0cPZltG9mpHt89w7QiXpbSU1UPTtufGWJs5HNcpxJ+2EEoA8OID9DFV6OhFVO7ePdHFbld9rG7NWSF0TR2jgA0cuGp9EwNnex52kd8Lci/6a0fB0lYJOFxbvqpS0PSUH89mKjJIPUFYIPgTU1OfRZbzuwgt3fuwpkY5tI022GwfLI8PlXSiZKt9ogOTpjzcaJGaK3FnAQhAHXl4VUfXn0ovZe0RqF/Tv5el3ZyMstuuwIy3Wkkcj76Rwkf1STUwymp4u4xg+5WJVW1l9uDzI+pd6HA9FUfXH0e3ar0NaX7nbrTpjXrLI4lN2eQqJcAnx4WnAGl8vzRlR8M1XF9LjFwlWadBnWu7wFlEu1XCOqPKYI6koUBkcj06eIFdpNi+1Vsx2hYri9udVR5UlrJciLyh5I5cyhQCwOY6pHXlmm12uuyHortJaPefaix7Tri1oLtkvrTfC8h0DKWXlDmtlRGMHPDnKcHOUJaCCcaDdd4fsp+xe0e72qVrKp/axcw7U48CuPIXwL7wK97wIyPx61uNvXKU2QpbjzQGeBeHEkf1TkH7qRnjdoFyuGntRwFQL3ZpTkC5RFJALT7asK6csE+XrjlUhdknbPRW93achbebkW1+fZZtimq7tmU4w424k5QtC0KBBSc8jkHyNRUNCXymJ53SOa2C97Y01HaWXSFnaMcRp59UxJdh0xcONu6aWhlzoXWAY6xnxwnCc/s035u1enH0k2q9TYhWMBMlpLif7ScH+6a6Wax+ies8ZBf2p3qv9nbSVFMK+R27jGSPLonlj9WueRnSvyzeLLcExParPcpFuccipUhp4tL4ONKVEkZIJx05inUkVXQt345Mt/wA6qtWq6bL7bT+6T0m5KQTwHLjqMJN0vJ392qUJu2mtrtHZZVx4s1wVw5/WYyM/ApqftuPpSN7tHYte6Wl7bqtlCglUhTXsE0AeBKB3Svm2D61DoWTzCsGs7jjU1r2a6QY9xZIwG5DfeED0P2h8iKZTVFNWtLK6FrvoUpcfZTTE9pbZyzwOo+fFdENqPpEOznue6i23S/P6OuagkBm+IDTK1nwTIBU3/bKKsfFmwrjHRNt8tiVHdQlbbrLgWhQPiCkkEfA1w4uO22lbqouW1+RanVeBPfMj7/eA++t7RmuO0R2f5xum3msbkzDSoLcbiOmRDeSOf1sdeUEYH5yMjwIqu1mxltrsmhlLHdHcFR7hs5f7GN+ph7Rg/MzUfuu27iMniAOTWFYOM+fIVQjZ36U2z3F+PZt8tHKtq3FBtd4syFOMo5faXHUSsDzKSo+ST4XU0TuNoPcu0i/aC1XbL7CKApTkN8LLYPTjScLQfRYB9BWe3jZq5Wg/js7vUaj5qOpq6GfQOweh0KXCSnqD8a+BzBzxY686FqSnKVD7/CtdRPMAAYNVondUm3BTA1xpq3wod2RPsjV50ff2SxqKyLGUqQcf6QyBghY6kDrgEcxkcuu0t2dNW9lPXlu1roi6yLjo66umRYL40kKA8THeI90LCT0OApPgCCB18WorJSrGMYwelRtrvQ2mJ+nbtonWFpFx2/1GFNz4oA47U6fsymDzKUhRCsJ+yRkDqKvWyG1r7NL2FRrGdD4eKirpaxVt7SPR4+q5h6mi2PfXSsnX2mI6Y9+itcV/tLRB4gOftDQPNRxkqzkkIJ+1zXXm621+3P8AdOYUlXNCh0Iqf96tm9xOxpuxHlQZaptnkqMqyXhAzHuMU4yhfCftpBAUkHxBBwQa0tbab0/uVppzcXRTSQMcd4tiACuE7g8TyAOrZxlQxyySCQVBO90tRHIwSRnLHag/oqVKw5IcMOCgDBorYmw3oTymH04IPLyPrWvUgmxGEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEV9SOdfBmn1ttt3c9XXZplpLbaeEvKcd5NstAjidX+qMgY8SQBzIrlzgwZK9A3jhKO2G3a7/AC1zrjJEGFDR30yWocorWCeWeRWrBAGR4noCaU9R3e/bs6ktW1e2VmfkQvaEw7Zb4iCpcpwnktXmScnJxj3iepNedw9YN3Ixdrdumn3rYiQGyWQVPXGWtSU54U8iSQgADOcJHQJq/nZb7N6ezfpxu4XFhmTuzqOKfa1rAWjT0FzGWx/ypGQvzOE/ZBKq/erzBZaY1dSdfyjqVI0lHJVyCCP1PgnN2f8AYiz9nDR50XpyQzM1zd2mzqy+pHEIicZ9jjnwSOg8SrKzjCAmYIFtj2+M3FiNcCEDkABzPiT655+uSa17Tao9rjCOyeJayXHXSAVOuHqpXry8/wCFKSOJI65PTwFfN96vFReqp1ROck/boFodJSR0UQjjH9/FZUkdSB18qyJ58uufCsbfp0IGOfjWw02RzzkjyqHcnOVEnaa3z/8AB823OsYlmbudwlzG4EKO6soZDi0qPGsjmQAg8gRzxzxXOvVutd+u0vNauetNROM2YECO0tRjW5gAn+bZT9teSQVYUs+KjVrPpPZyW9q9IwEkhUm+ur5eTbB//wBn41W+yIMfT9hR4ptUc8jjmcq/7Qratj6eC32WOtjjBleSN4jUDwXWzlhj2qu8tPWOd2cYBIGmcnqk2zbW6WsD7QZiuX2UMEOyW8NcQ5kpazzHqrPwpZd1RAkXBmyRlv3K4uL7pi2WiN7U8pXglDTY4U/DI+FYlssak1do3Q81+Si2ahv0SFcUR3i0XmFuISUkjmOpPjzwcchXa3aPs97ObGwDC2v0DarIpxKUPymWMypAHTvH1ZcX54KsZ8BVugpTWtE1U8nwU1tLtTFsVUOtVkp2sIAy4jJ1+65m7V9iLtT7qOxpsrTdq23sylDMvUCjKmlvzREAwevRzhHPrVntJ/RS7NRwZG5+uNXa1knr3k0wow9Ettc0p9Arl4VJe+/by7PvZ+uErTmptRuTNQRMh21QGlOPIUBkBXLhSTkEcXUEHpVftC/TB6G1RrmBp+9be3Cy2e4Skx0XB9xBDfEcJU4ArCRkjJBVjqQBzErBDHE0iOPHjjVZXcb9dru4uqp3O8M4A9ApB179FbsFcLDI/wBy24ai0Vf0oJizGLm9JaDgBwHG3VHiSTjIBHKuZOpYmrdv9Z37bTX8NELUenpaocxDYw27yBQ63yHurSQochkKSQPAfoFZeDzDb6D7riQoeYBH7+lcaPpYoUS0dreLd7c2G37lpWE/LKfz3EOvICj68DaB+zXtRTsq491+A7kn2ye0tXZa5rg8uYdHAnI/8p6fRq7IxN2d57lu5qFkSbJt7wsWtCh9W9dnBlTpHiGkg4z1UWz4GuoO5uvbLtboK96/v7xZt1khuS31BPEQhCc4AyMk9MZGTgVVT6JixR7T2T4t6SP9Jv8AerhMfOPFDpZT+DX40ofSpX+baOybeo0VxaU3CbEiuFJx7qnkgg+hBIrtrGMAhby0/dQd0rpbrXvqpTkvdn9lyh327Qe5naj1xL1Dqq6Svyb3qvydaUukMRWskjKR7pcORlWMnGBgACo8e0tOtKUy46lMutkLQ60ohaVDmCCOnOnVpW0NxIwWEe8cAnp/n/vpckREPNltfQ8udMZbiY5d1g7q1u07A08tuEk/+44ZC6o/Rhdoa/b1bKSrJq6Q5Kvej5Qtz8pauIyGlJCm3T6kEpPmUE+NTV2vtG2zXnZo3G0/dY6HmlWCTKQFJzwvMJ75tY9QttJ5eVU2+hztrsO07rXBIIjflSHGTz5caUrUfngj8Kuf2ntQM2Ts/bgTn3AlAsE1OSfNpX8M097QNdlY7UQGKZ0Y/KSPUFcDlXD2nR5KjlQYKD8Ry/hXc7sGK/8AJH2ySfCwx/3GuEMVtf8AJFw56NK+fM/413R7CL/B2TdtQFdLFG6/1aSaBHvAdSrXtVLJVR0rpOPZhQT9MUoDZ7RvP/8AMrX/AETtc2oDh7hv4V0Z+mDc7zZzRpzz/lK3y/5lyudVpgTX4ySzEeXy/NbJplcTmNp81bfZaS3t2+X2W2hw56mvangeQNeHIFxaOFQZCfi2oVjU080nidaUnJ5ZBFRIZrla+akEFoIyuj30O+Bs3rok9dWOf9XZqYfpI1/+Rzr8A/0Eb/rTVQl9EJIKdm9bgH/82O/9XZqXPpHn1q7Huu05PNuID/8A3LdWYPw7C+TaphdVPP8AUfuVxrbmmJpV2UgkKREwPioBP8c/Ku6XYp0pF0V2Wdt7EwwhvhsbMl7hAHE89l1xR9SpxXX0rhBJSVaQcT5x0kD4cJ/hXdnsf6yj6q7NW39zjqGU2aOyvByApKE5H3H8KQh3WNdnm5W7bZ8kzqZh+ERj+6hD6V7ey8bdbJ2/ROnpyo9w1rNMF5xCilSYqUFTpBHQn3EfBaq5GW7SCpTAcU6W0qGcDqQfPzNdKPpiNNXKbZNs9XJQo2+LdZNufUTyQ64gLQPmGl/dVDorSW2EoSDgDzrmqqnRMG5xOVJbAbP0t27SSpGQ3A+a0dq9d3/s87pWHcqwPvp/JstsymmFFHtMUnDjSuEjqnOPI4PUCus+g/pS+zDqwIavGopVmeOEn2uMtHPxJykJHyUceZrk/PtzM1spkLQhHipSwkD5kikmbotttlLimCUq5ha2woK+CiK4jrmSNHa6Hqnl82Ae6pc6jIDenP8AdWC7dN/2o1J2j07i7Qartt4tmtbQ3IuCYTgPcXBklpfEkHkVNoaVk9SpVbP0dqS52zLIoHpZJp/Ef41Wy36QbiTEywgAozgDkAT41ZT6Ok8HbLs4z0sM3n+0mlWvjmmLmfy6pjc7fW2nZw01XyeMeRXa6UAIrv8AUV+6vzvzl43B1mM8zqS5f9YVX6HppxDeOejav3V+dq5rP+6FrIjx1Jcv+sKrqrANO70UV7PHbl7YfApZbcI9aVtF6F1PvBuPpXZ/Rrzjdw1TcW233mjhUeGhXG88T5IQhS/XgwMnlSChZUAkDJIIFXe+iX2mfvep9XdoW6xiYscHTtjccTz/ADVSHEHoORSnl4lQ8DUZboA+XtHcB/gWse0HaB1stRp4T35dB4DmUk7rfRqb7aFT+UdqNXWvcO2MoUv2K6NiBc0jmcJcBKHjgdVKSeWAnFVZvtr1hoe/o0zrvSt90he1JUpES4sKaDgHUtOgcLg6cxjr1Nd7XpMRl1iPKkNodmLLbDalAFxQSVFKR44CSflVG/padv5N32Y01uha2gZGh76hcpQ6phyh3az6gOhj+0fWpKeigqQTugHwWXbObe3a2Txwyy78JIBDtdDxwVzkvOnNK6k9682lLconnLhBLLg9SnHAs/EJJ86QrPYd0Nq721rDaDVlwEiKQ4h62uqZlhKTkpcZzhaeXNPvpI60tw5PeoCsgg9CDyPjmlW2OlmdGfyQW3215HooVCRVs1P+G7vN5g6rbr1sXZ9oo/eGx7kh1Dm6H15FW87FPbP1bvvfJe3W4lniC9QYC5rdziI7pMhKXEoUlxoe6k++k5Twjr7o5Vbxa0gkcQHPNctuwRKFs7XtytwIAlQbuwB/UBcH/R5rp64+vJOVZxzGRWV7e0NPQXMe7NDWuaHYHDJWIW3fbvwyOzuOLdeOhWRxwkE+ArE4hDyC2tsKQpJCkqHIjyOaxlQUDjln1zQ47ywFcvOqMTjVSoCifc/bHRusdIS9pNxkLc0bdV8VsuA/n7BN58C0qV0Rk8skjBUknB58vtdaK3L7H+7kjTl9YKkj32HUgiNdYSiQlxB8QQD6pVkc+eexc6LHuUZ2HNYS8y+gocbV0Uk+FQzvBsvpvePR/wDuO7gye5SCtej9RLwpy3y1JwGHj+chXCgEZHEE4ByE40vYra40EnuVYcxH6eKgLvau3b28Pxjl1XNPcLR1g1HaG9c6Jc7y0Sf55pKQFW94n7BA6IJIAHLhPIZBSTDMqM9FdUw+gpWg4INTC/H172Y9y7noHXdocZ4FGNcoToJblRnBgONnopCkE4V0UCR5itXc/Q1tMSJqnTEoSrNcgTDeByppeCSw5+sMKwepA54IIG6Qy8MHLTqCqW9pPEYIURUV6cbW2soWkpUnkRXmnaQRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFfQORPlXzB8qWdN6dmX+aliO0teVBICBlS1nolI8zXhIbqULe0VpC4aousaJDjl1Ty+BtGcBR6lRP6KQCSfIGn/rrWsfSdjXtno2SFOOlIvE5rkZDgGA2Mfmp4iADnAJ6EqrNqG+NbTWN7SNmdZc1FcWgidKaPKI0QMMoPxHM4wevMcJFlOxF2VIdsgQe0Xu5ajLL6uPSNgeR78x8n3ZjgV/RpIJSCCkn3jyCeKHudygt9OaupOGt4eJT6np3zSCCLVx59E7+xx2X29k7Vb94Nf2pMncG9MBzTVmkJ4fySypP/C3geYcIIIBA4R098+5aW0Wv2JLj8l3vpstXey31fadWfHp0HMAdAB0517h2+WqU/e70/7RdZq+N93BIQM8m0eSQPv5mlEY54+QFfOO0e0M9+qnSyHDeQ6BaDb7eygiDG8TxK+cIV5Y6VlTjPDjGPOvGRjCsjPmK+oUAOgPqeVVzU6p+thoJJzy5eVZSQkJ5ciR++tUKASME8Ofga8OTAkEk4HpQNVyRgKkX0oFxJtOgLcFABcy5vAegbiAf7VQpFUG7XbWiRlu3xU8/wDUp/xp/wD0lV8g3HU+iLTGuEd9+ExOcfZbdSpbJccZA40g5TnuzjPUCo5JUhtls5wmOwn7m0/4VvFohMez9I0jBOT9SrN7MRm61z+gaPqvWmFF3ebbVIOf/GiB0/1zdd/Ejh8OYrgDorLu9u2jeef8qIOP/bN1+gDHu5B6DrVtoW4p2+qoPtLdvbQSeQ+y4LdvS3uXXtk69bVn/hEY+eAY7dMPbPZS9bsbk6a2w0jHW9dLrNb9qWhPH7DFBBcfcx9kJTzwevIDqK6Wbl/RmSd6e0LqjdXWe5j1lsN5ebKIFoZBmOIQ0hJCnVjhb5p8EqqzWynZr2T7Nlpej7d6YjWxx9CUTbpLc72XJ4fF15eSeZOQOFPkBTx5wQ7e06fuoCK508NvdSxxZkdxceQ8FKURLUC3ts8XC1GaAyo9EpHU/dXCjt87mQ92u1XqSXZXS9BsUdFjZczzUpsqW79zjjiPggedXi7eX0gOmtBafn7V7Q3dq46tuLSmHX2FcbdvQrKeNfgFgZKUE9eEqHDyXy801ZX3EuXKepbjz6i4pazlSyTlSufPJNJulEDTI/08U72Zss13qw2MaDUldZfomdUR7x2YGtPocT7RYLrMjOozzw48p0H7nMfKpO7fW1t03e7MGrrBp6K5KusKOm6Q2G0lS3SwoOlCQOZUQggDxrnT2BO0hG7Ou7EzROqHwxprV7iC28sgIZldMFRICQrkAc8ihGcAkjsZZ7xa7/b27lbJiJMZ1IWlaD5gciDzSroCCMg8sDnXheA4OHA6qJr6OWhqXwSjDmkj6r8+OlSuZZ2pCUFQB7tzlyS54pPkfHB9KUrgzcWnYlut8J6Xdrm4hi2wGkFciU6s8KAlscyCrl059K65bs/R3dmjde/yNXybBdtN3iYtTsyRp24GGmSsnJK21JU3knqUpSSTk07Nm+yN2f8AYJ78r6D0Syi7toUF3q5PqlzVJPU944cI5deAIHL1po6jjEvbF2nTC0FvtJnZbBRxQ4lxjezp0zjqknsSbBOdnjYe16TvLSE6kuTrl3v7iTxf6Y7j6sHybbS2jyylR8agf6VvfSNpPauPtBZ5p/LGrXgiQltYyzFTwqcJxzyr3UfBxXlU09pntn7W9nrTspUq8sT78pJah2+MQ4646Rxe6OhAyMrPuJPIkq908ctb641tvruHN3F13KU7IlK+pZBJbiMZJQyjPlnmTzJySSSacl7dZX8B/mFSbVbai7VTYYhkk6+vEpClRzG0s6gJA+pOPhjl+6u2nYUV/wCShtt6WKP/ALNcZNRQT+QZSWxkpZPLH7qt/sn9JrpHZzZrR+2sTQ93nzbFa2Ykp5bKeBTqRhXDhwEj1NIQb0sZLeOdVbtvKJ1JUwRAaBmPqr+dofsxaA7TtgtumtxLpf4cG1TBOaNokNMuLWEqSAVONL5YUTyGenPziKP9GB2XouEyBri5BPL/AErUa8f+7bTUBz/phpZSW7VtHKUT+cpaUH8Sv+NNyZ9LXubKB9g2vdR5F25x0/h7J/GnH4rWhoA9cKlQQ1jT+DvDyyPsrYtfRp9kNA43Nv7i6RzPe36YrPxw4K5Wbi2W1aI3x3E0VpyIYtms+o5cKFG7xSw00h1aUJ4lEk4HLJNWEl/Sk9oOUSYeibUgeHfTSs/ehKPwqtU2/XzXuudTbhalgx4czUdyduLzUcq7tC3FFSgnJJwCeWcmkp5D2Lu0I5cx+ituydDcBdY3EPA1ySDjh4rox9EM0Bsxrh0Hrqtf/V2x/CpW+kccCeyDrkK/OTEA/wD7luuZuxPat3m7N2nbppHRNis70O4XNdxcekOuocUpSEoxlDiQUgIBGR4nzpb3e7dW9W9G3V52z1VpiD7FeG0NqcZlE92UuJWFAHJP2cYJ8a6ALpA5pGPNV+qtVY2Z5dE7ieR11UHQ4pl6fbZAzlhIx+z/AIV0E+io36jIs1y2D1VckszLQ6qTae+WBxxVrB4Ug9eF1av/AGqfAVROxQVxbZHadbwpLaQU8uuBmschd/0ze4OtNFXJ+23q1uB6NJZOFJPkR4gjIIPIgkEEGmkM7C90TjoStR2m2Wnr7VDVxD8SNo06jA0XcrtDbL6f392jvm19/UlkXBsLiSwgKVDmNniZeT6pVkEZGUqUPGuOe5ux+9WxM5y17n7f3hMVhwtMX23RVyrdNSDkKS8nklRHVKsK8wKt52bPpRNN3KPE0fvoyLLdEcLKZ68+zOknH2yCEJ8RxkADIKyMYvPpPdLbvWEJUjTWsLTNZGAQmUjhV/VJPCschzQSPWnUgDx2c7dOR/usvtN5uGz8plpXYJ0II0PmFx07PnZq192n9Z2ez2zSt3iaJjzmnr9f5MRTMdEdCvfZaUrAcdUMgJBJBIJAAzXQLU30YPZkvbbn8nm9U6QknJS5ZryopzzwFNvhwEDlkJxnHWrTC+WNpBccvVvabRz4lymwkevM4qs/aZ7fe1GyWnpMWxXpi+6llNuNw4MRwKWFjlxKGctpGchSwM4ykOdK6Y9rQIoRp80lcLxX3SqNXO/Djjhpp0GFzh7S+2No7Pm8re1dh3IkawZat4kTXJMBth2I6srCGlKQpXGeEJWTy5KHKnb9HovPbHtK08ibJMyPXjRVfZN71FuDq677i6sk+0XS8vqkvqySEk9EJzk8KU4A59AOZqffo9jwdsW0qH/8Glj/AN4iuO72zt3Gd3XCttw98bsyz31xJLwRk5OOS7Ty3z7G/wC9/Rq/dX55p/8A+PdYBQ//ADFcf+nVX6EpXuxX+XINqz/ZNfntuKkL15rBaSDxahuJ5H/l1GvKjWB3oorYTS9MPgfsi6t3KYiNZbJDel3K8SG7dDYZSVOOOunhSlKRzJOcAeZFd4ezntTA2S2U0ntlCQkKs1vbRJWnH1klzK3l59XFLNcu/o3dp17qdpT+Wdyhh6ybaxxL4lfYVcXeJLI/WI4XFendj0rrdrfU9v0XpC8asuTvdxbTCemOL/RCElRJ9BjOfAZPhXdPGYog0cTqUntxd/4rdC1h7jO6P1+qpB2pe1cdIduPaPbqz3b/AECzPrReWgeFvvJ7ZZQF+qQUuD0I86th2i9v4+7Gw+tdDKY75d4sshEdJGSXgkraI+Cwk1wJ1/uDf9xNx77u7NmPLuM+7quza1K4lNoS59WnI6BKQkAeSPKv0F7Math672n0rq2DJD7N0tMd9DmOa0qbBSpQ8ykpJHrTos3SHfNVWWF0QbnnqFwT0hLW/amEvhSXGk90tJ5YKTw4+PLnTnZcwQQfs8/upwdorQidqO0zuNodlsMRPyyu629oDCUxpX1yUj0Tx8HypsNK90kE5qtV8W5M7H+ZX1VsTcv4lZ4JM6gYPmNEvdkef+Te2tbPex7ZIurHx44j9dV1EFOcg865EbDXiBp3tf6Uut1nR4cVN6CXX5DqW20BxpacqUogAZWBz8661tyUOMNusqCm1oSpKknIKSORz61n3tJjPvNPLjQsGvqsOhw2uqmdJHfdZlOjPPkB5GvPeHPMAj0rHlWOLBIJ8q9JUOvPGfEVmZ1UjjC9FXMZ549a0bva4N3hOW+eyHWXU4UPEHwIPgQcEVtqKfA/hX3mfE+HPFc5LTkLrlqq99oTs/W/tG6UTobUUtiLr2ysrc0pf3B7txZGVKhvnz68+eD7wHJXFzh0vebztBqi8bZbnWV9qCX1Qrvb5CClyM6k440gjqMA8uXQjwNdkb5YYl9gqivKU26lXeR30clsOg5StJ8wedVq7VHZmc7SOn3rtbIbEPdvTEXJCU8KNRwkjkAeQ70DGDgnOUnAI4df2F2uwBbq52nI9CqrerTn/UwDXmFz33Q24dscpE23vpmQpaS9BlN80yWeRAP66QRkc/tAg4IqL1DHIg5qX9vNZJt6ZW1W4cVxFufeUy2p9JS7bpScpzhRHCQonOcY6HkVAtjcjQs3TF1daXwucg6HGwQh9tX2XEAjOD4g8wQQeYOdkieQdx3oqnIN4b4CY1FBBGMiinKQRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUCigUIW7a7c/c5SGGUE5+0cdB/jUyNqj7L6eamPMNr1Rc2+GHHOFewtE/zivJxWAccikHh68aUtDavVdk05docm4WZuWGFqUsrweFZ+w8AftcBwQCCMgcqV93dOPWTWUTWVxku6gsF6cRKYeUvgUtAKS5GWRngUnPCCORSUrAGcBtKd5+6eCXYMDeHFTd2M+y+zuXNd343kac/kZapXHFiuD6y/Tgc90OhLQUBxHxPu4+0U9DYSJ1wmqv15joZkLbDUSIjk3Bjj7LSAMJBAxnAAGMAAUkbeXrQW6e3elNe7dPNJ0zb7e1EhWdpIQ3aVI90oU2Oi0/YJOeY4skKBLoQCEAdRjHUnNfPG220FVcqx1O8FjGaBv8AnNX2y0MVNB2jTlx5r0XCkYBI8POvuUjJKhWIgDJGDX1SiAfDn5VRHYAyVN+CyDA4eefnWJxSkcXMj50j6u1ppbQdjd1JrG/QbRbmPtSJTvCCcZ4UjqpXokE+Qqje9P0gF/1DMf0vsTanojKz3IvElkKlunP22GhkNDyKuJXiOE4AsNj2XuF8d/p24ZzcdAP7+SY1VwhpTuauceAGpVvt0t+ts9nra5O1xqVmK+G+JmAwQ5LkE9A22Dnn+krhSPFQqi+73bf3T3Ucf01tZbpGmbW/xIUuOouT3k58XR/NDl+Zg8z7xHWK2NtrxeZDuo9zr9MVNlK75bC3S/NkFRyVOLVngz+sSfSlxYjWuCbbZ4TMGJy4ktDKnPVaz7yj8TjyArVbRsxabGQ4jtpepHdB8Apa37JXa8s7euPYQHl+c/sokvtuucG6BV5uAmTZR715ZdLi+Pi58Sz9o+oJqdJoCJCkZwEobH3IAqF9V5VfYwHihPz941NFywJjw8lcP3DFWa7PL4oSQrJ7N6OKhuNfBADutIAzqdF4289/ffbNPX/xng/9Miu/6SOEZ8sfjXALbVJVv5tkP/8AJoX/AEqK76lZKAAfCndI7ELTjqs19og3r9L6fZUl7Uv0m+kNh9WXXbTT+irnedT2vhbd7zhaitFSQpJKsknKSDySR59KoBvH28u0rvqH7YdQjTdnkKGYVnBQ4pI+yFPfaHXmElAPiDXztuQkXbtl6+S6MpEiN/1dumNAtMWMEhphCceSedFTVR0p3QMlPtktiXX+H3l7sMzjRNew6MX3nttxUpTizxqKlcSlq8So+f8AjTxaipbbCAnhAAwK20ISkAYr7w8+VQs9U+odlxyt3s2z1JZIRDTjzPVIF+sLF4jlh5GeRAJA5fCpI2e7XHaO2BaatlqvDeorPGVlqNcVrLjScAcKXUkKKQEjCF8aB4Jprd2VYCU8RPQAUnXCfabeP9NuMVpXilSxkfIc/wAKVpquRrezA3goLabZC13P/U1REbv5sgfPPFWhd+lw3nKOBvbhgPYwCqU3w8XnyYz+IqKNfdvjtW7kIlQol6i6biTUltxNvaU48UnqErcUoNn1b4DUH3DXelIS8Nh2YsdS0jhT96uf4Ukvbky33gzZ7A0VL5JSviWs/wBnH7qlWNmeN5sWPElZRLZtmLe/E9aX+DRk/RLCdM3a73Jy/aoucq4z31AuyJTynnVn1Uokn506YlvbZaShlnhSkeHIU0oUPerUigzbNOT2ArkFNwwzgf1yAfxpbt/Z03z1FIHt5QwFYBclTS5j5N8R+VMap8Q1q6hjccsqft+2FosrOytdG9/idM/qt+TLgxMtyp0ZrPIhx5Kf3kUjuS9GoKlu3a3JUfDJX/spNPy19hvWU9Q/KepUjJGe4hqWPkVqQfwp8WzsA2woH5QvN0dV5pU22Pu4T++oiS9WOl41JP8AxC9qtubjWHu0LABw3iSQoAd1PoWGf+Ed9jxajnH4kVg/3QNGoPutTSP1Y6B/2qtVC7AuhElPtMy5LA5kGWOf3Ip127sQ7PxkpEjTokEeLkyRk/2XAPwpq/a7Z9n5nu/zzTH/AKn2gP8AtsiaP+OVSpe5umEDDUCafiAP4msS9z7Nj3LTK/tir6xex3sog/WaKgKT15uvqOPm5XlXZl7ODWp4+k3dMWpu6y4pmR46kO4dbSopVwkq+1yJ4euAo9Aabt2zsjstjieca/5qvHbR7TE4EkY/7VQtG5lmPNdslDPiFCsqNxtLpXxriT0nxwhJ/jXQGT2N9kXSQnRVvR547xJ/urrRldiPZKQ2QNIx0HzRJlD9zteN22sLviY8fL90DaPaZpz2kZ/7VRlG5Ok3AMLmoJ5ZUwP4KrdY1bpWYMm9sN+jqFj9yTVsrn2BNrZBJiNS4oPTupaz/tcVN2X9HtpEk9zdbokAeEpP8W6XbtTs7Jwe8eYT1m220jDiRkbh5YVZpzWk7unhFwgLJ5A96kZ+RIP7q0I2mL1ZlKkaU1BNglz7RhSlIyP2Tn8ani8/R8TGV5s+qZiUjnh2Kh38QtH7qZGoOxTu1Y2y9Z7vAnJAzwZeYc+QKSj+/UtS360uAZDVjycFE118bcTvXC2tJ6sdgpluO7qPNGPK3F1E4yoYUhdzfwoeIIJ5ik6BoZhD4ckqK1Z4lHJJV8SaU5+1vaC0k0e9tNycaR+Y1IRIR/ZSVfupvydV65sCwi/6cLJJ6yIimCfhjFSzZJKgf6eVjh/SQlrfc9lKeQOmppGEfzd4fRPZERtlpLbKAEpGBS5stu872ct64u6S9Mu3plq3OxO6Q5w8JWsHi9ccPTl1qPIO6VnXwpudtksnxW0oLA+SsfvpZGotM3NIQxdYygscku/VkfEKwPuJptGyppH70jMjgrnc57DtbQe6U9S1p0IHAjHLBwunGh/pUNgdbf72X56TpiQ+2UH8psFtIWRg8K0caMeq1orldd7qwxfdXXmA8h+M5e5zzDiDxIcCnjwKB8Qcg1vydIWyejvTGSUq6KRgpPwNa8jSaWbemPa5LsZxp4PpLayDxj7JB6gggHr4U794pz3TnXHFVGh2KutqldV05DwGu3SOpXZr6PzZg7Ndm2xsXWAlm/6nKtQXhSk4cLr6QUIWTz9xoIGPA58c1OeodW6MtsljTWqLtb2nLs2sNRZigBIR0UMK90jnzHlXGHb7t9dqLacpizL43qeIVha1T0YfVjqVKT7q1HnlawtR8xTN7UXaX1Z2q9e2fVgtU6ws2S2IjtRESSrhkFXE66lQAxxYQOn5vOnPekfvggDr/ZZ5Jaa0TmGWN2+TwI/VdUN4fo7OzDvWl67s6UTpS7voKU3LTixFSpR8VsAFlzmeZ4Qo+dSh2a9nrhsNtdb9r5upRfo1mW43Cmra7pxcfjJQhxOSAUpITkHBCRyrjptX27+09sulMJnVTl+tyAlAj3cFxYSnoA4DxK5chx8QGeQFdIOxl2+rZ2oLvL0dP0m9ZL/AhplPo71K2VI5JUpBHNQ4ykEEDHGn7XUduLy3vajqEznpJqZxZKCD4qtn0sOhG9Ob36F3LixVpa1Ra37XKcA90yI60lOfUpdH9g+VVNaX7oJNdRfpRNCSNX9mCZfoMVLkzR12iXxC+HK0NpKm3MHy4HiT4e6PKuWMN5t5ht9s+64ApPwPOou5gPa2T0+S2n2TV2aeajz8JyPX+6jjXMczdbPRUupaMlbSeNWcJ4kp5nAzUn7U9pDfrs9yGIse4O3bTxKf977gTIiLQD/Qrzlo/wBUgchxJOKjjWqeDXrJ8zFOf2U0/bDOdZiqjraaeYdGHWHkBaHB6pPL55BpatMMlKyOojD2EcCFA0mzEV9udcBIY5GyHBH2IV/tkO2btJvM0zbH5403qRXCFWu5OpSH1HxYeOErGce6cL/VI51O5X05qwrBznmQa473ja61XhJm6UlC3zs5ECQ59U4f+ScP2T05Lx/WNSRtJ2yd5Nin2dIa/hy9R2KKpKExLisplxmhn3WH1ZJT5JVxJGPd4RnObXfYCCsBnsz+9zY4/Y/omFdS3TZyTs7pHlnJ41BH6LqEVZ/q+GTXtIwDnB+dRxs/vztlvbakT9DagbekobDkm3SMNTI5x7wW1k5A/STxJ9fCpGCcq6j19Ky+qopqCUw1DS1w5FLRTRzsEkRyFlbJJwSQPHFJd/sDl1ajy4Er2K7W9YfgS0HBacHgf0kHmCPXl0pVbUD4D7qyJIxzUByxnmKSjLonBzThdHXiqRds7svDeOz3HefbuwJha8saM6t0+ykAzmk/+mM4xxrwOI4+0kE5Uoe9ULbe/wAHXNpG22r5QYmtEmzXB4Z7peMFlZ5nhVnB+RwSlIrsNMssdq8RtbC7/kVyyAuSZpICFxuqmXM8iCE/aP2ccXhXILfGVpPeLtFXyTsdYm7ZbLjPUppSD3LC3E835YT/AELSiFOcPgD0B90b9sLe57tSmnnae5wd+hVHvNJHSzb8fPkou1ppadpy6SIkuIqO6y4W3mSQrul+WU5Ch4gjkQeRIwabdThuvqGxOSmbSVflV60R/ZpF2dSUOzsAAJUjoMKClAn3vf4TnhGIRWQSSkYGeQ8hWiROLm5KgHjB0XmiiilVwiiiihCKKKKEIooooQiiiihCKKKKEIooooQiiiihCKKKKELNEkuwpCJDJ95B+/0qbNA6ms+rbG7tvqUg2u6e9CdUCVQZY4glaCOYxxK5dFAqT4giDaVtO3H2Ke2laylC1DmPzVZ5GkpWbwyOK7Y7BwrS9jXfa9dm/duRtjr9fd6avkgRJrb6j3cR9QAalIPTgUlSQVDkULCsnhFdMbhE9hkltPGW1DibUpOFFOeQV+sOh9Qa5B6/itbg6FY1lDwL5ptAYuKQpI76HkBDiehJQpQ5czwL5ABs1fjsO72Sd4dmTpu9SfaNQ6L7uG4tZJdcZ4T3KzzJPEhCkn9ZoH87nkntFsLZYv4rAMEaO/f9Fa7DXFjzTuOnL9lPQUniBKvH7qrv2m+2Dp7Y2U7o6w2py86wU024G3krbixG3EZS44rkXFYIIQjl1yoYwbBgJXjn4gZ+H+PPNQ52k+z7p3ezTH+mRu6vMBpQgzmkguM554UD9tvi5lGefPBCudZps4+3R3Bhujcx/QHx6hWO4sqXwO91OHf5wXPu+T91e0Ndm9Zbl6pcRbgVJYW9kNMt8uJEWOnkOnPGMnqrOaclktVi0i0pnSsAsuFPAqe/70lweBB6N58k+XU1H8lvVmymppGlNX254RSvjAKSEOtk4DrKj1HIg46EFJ5ggP6FLjXCI1NgvpeYeGUKT+IPkfSt2rnSBjRDgQnG7u8Meivns1o7DNDvsbvVQ+Lf+IeI8EO8/eOcnJOfPxNJk7mhWaVHcUnTU5Qr4VHRcVp9Y38MhRfqhONRQx+qj/bVUx3NJE+SfALNQ/qkf+MsMeiP9s1Mt0A9vlf61Q/GpW5f7UXks72Ebi43L/mPstDRl9tOmt5tvr/fJ7MG32+/R5MmS8rhQ02haCpRPoBV097fpboTaHLNshphy6SlKcQqdJy3GbH5qkkjjdJ644UAdOJVUWvNhgXoNomI4g2oqBz519/JVkscQSJa48VpIyFOKAJ+A6n5V1DVRNja0tJPRQ+0Gw012usldNK1kRxqT0HjwXifqDW25evrpuVr19p26XUoU8ptHCDwpCU8iSfspHMk+pyaW2mVKUlCEFRPLAGaYt53NtMNKm7HEVKcHR10cLY/Z6n7x8KRIUzXW4brsWFI+rb4S6hCw2hPEeWQOauh8+ldS0k9T+NNhjR1XdLtbYtkKUW+3l07xnhwJ81IFy1Vp6zEpmXNorTyU2z76x8gcZ+JFNW6brhwqZsVpV1wlyQfePwSOh+ZqQdC9lKZe+F+8PyZKlYIQ0numx8VK5q+5PxqxGguzHYNN8D4gR2nABlTbY4/ktWVfjUDV3yy2vQkyOHoFC1u220d20psQMPTV3zKpvF07vJrQtpRDmxY7493vP8AR21DzxyJ+40/dJdkLVN8He6gubyFk+63EaKwR5lxfDj+yqr12bb3T9sTxMQW+LOSSMkn99OSPbYzHCEspSBVWrfaJUYLaNgYPAaquvtclW/frZXSO8SVVfSnYn0lGDS7pCW+sdS+8peT644R+GKmbS2wujdNtBEG2R4vIBXszKWcj14McXzqUm2kD9HFewEhRPLl1qn1m0lyrv8AelPzTuK300Awxibtt0VYLacs21oHl7xTkmltNvhtYLcdAx4hArYyni5fv5VsBsOwHHkAERlp4sDBCFHhz/aCR+0POoh0j5TknKdANZoAsCGW/IAfCsiQMc+maxgjBGVcz18/WtmEhLinEudAy4ofFKSr+FJHXiuivISkeFbTLPe8kjn6VpNL4zjI5Dx/z61sXG4qsmnrjfEFtLkGK7JSVjKQptJUCoeWQCf8mk3Ekho5nC94N3ltKaRFT3r5CEYKgo8gMc+tMbSmgZm8uktZbm2t72CS/cGGtCy15AV+TQ7/AKUCefduuuPJPmggU3zrmZvsnTe0kWM7YLtcGEzdYBC/ehWgstOFbC/zkyA80htXVIcUFAEE1PGpdS23SFnsls05HjRrbFZbj2yKyQhHcNoAQEjoAQOWB9nHyeVVQ/Z9m4Rmd/AdG54/92Pl5pvEHV7mth9SmFt9rA6701Hvq7c5AlpW5EuEJY96HNaUUPseoS4lWCeZSQTzNOc5ayhSSk+IUMGq8sal1dftX7wae22gSmGXLUNWm5sHAZCoSglpsj/0h9aGCD4BLp5qwRMO38yHcNutL3G3O8cV6yQeAk5IIjoByTzznr6g5pxX0Hu7BM3QOwd3oHDe/wDCI5t6QxHiE4FrGeR8M18SQskAdOta/e5OD4DnW4lKWoCX1dXHSnp4pAJ/2hUSc40Tg6Fee7B5YFfPZ2XMgoSfDpXwOBXQjPhW3FbWtt6SU+4w3kknkSTgD49T8jQCeS8ccJFn6dtU9BRJhNOAjoUDnTVvG0umblGXHMFKG3BhTSR7h9Cnofup/ko6CjhTz6gjrml4qqaE5Y4hcuY13EKsOsexpt5fOJ9qxQ2Hf0o7Xs/390Ug/dUIas7Cio7jz1muc2KMZbSpCZCSfLqhQH3muhKkIzzFaz0dlxJykHNWSi2yu9CcNlJHQ6/dMZrTSzauaM+C5SXjYreXQii7C4n055iDIUOIeGULAUfhg02Xtb6osE1UDVFkWl5OONDzKmHfu/xBrrJdtJ225JLb8VpQPmkfxFRnrDs+6X1DHXGkW9pbas4acbStA+APT5VcKL2hRT4bcYQfEaLyBlytZzbqlzR0JyPkVQGBrLTF1CUOPqiuK+0l9OE59FJ6/MCl+NFjqjh6IG1sr+ytrBSfmKlvX3YyswafkWtqTBdIJQuMeNAPhxNLPP8AZUKr1qnbDX+2Yk3ONcO8hxQkvOtKKCElQT7yFdeZT0J61bqOutl2w2km3SfylT9Jt3W0JzdKYPHNzdD8uCckm3x5Qw40hY9Ug0qbDbtXPsyb3Wzc232qRPgNsOw50VlYSp2OvBPUEHhWlC8HGSgDI61G9m3Qdz3V+iJeSP6ZkcKwc+XQ/LHxp2wrtY9RIKYEpp1ZH82vk4Pkf4ZqRDKihJErct8FNVclh23pQymlDJc5AOjvLXiuqlj7Y3Z+7Wu3WodpxqJqz3LVMGTaWok9KWXcutlKF92on3gog4QpY5farlXpxEqJCNqntqblW55yFIbUMFDjSikpPqKTbjouJKJW0nu3M5BBwc/EdDSnpmyPWmO+mTKW+t94vcSzlRJHMk+JP+etc1M0UsBaDrngU22P2YuWz95y8Zic05I+iaOvCRrmOD4pin+6mntaQO6TTK3C5a8jJT4txf8AYTT3tQIQAa9rf/bReSktlB/65cP/AOz9Ess/DPpW5Lah3aCm1X2C1cYSP5tt3IU1y/o1j3kffj0NajIOK3UI4k4Tg/Hw9T8s1CZc12W8fBalJTw1MRimaHNI1z0TNm7fX/Rtxj6x2m1DORKhqDrKWXixPirHXgKMd4PVPM+KetWc7PP0ha5DsbR++sThdUpLMe/xGvz8gcMhlIOPH30DyBT+dVUtY68kznf5K6SQ7JefWGFusJKluqJwG2wnnzyBy+1Vu+x52Qo+nnIm5O5ENMi7NrS9AhrALcUjmFnzcBHI+BHnzHO0b6BlsL720Ocfg5P9F837Q09uivHZbOEgD4+bAfBXhQ6lTaVJWcKAUMgjIPMHnz+R5jxxWVgOuupaayVLPCPD8a1uPkFFIHmPD4ef+fGmXvPunb9mNrr5uFPW2XIcctRW1qwp55fuNoTzGVKWpIOPzAs+FYZR0r66dsEQyXHACVmkEMe+88Aqu/SO9pZyz29vYDREwtyJLYe1BIaJC+5VhSGBj7IcGFKHXgCB0Uc1WixP9xnR67XIZLGsb42PblLThyBGyCGTz91QKQpQwDxYSccCsp+3LkzU+q75vhr2V7Z7DKcm5cIKpVxcypHukEFCPtnwGG08uMGmNr7VUu+XGTNlOKMqa4XHDxZ4U55Jz4nzPj18a+nLLaI7TSMo4+I1cepWc1dSaiR0rvTyTfvN0VPe4Gz9Sjp5qPio+ppMr7nlXyrAAGjATAnOqKKKK9XiKKKKEIooooQiiiihCKKKKEIooooQiiiihCKKKKEIooooQivqevyr5QKEKVNodTNQ7y0i5IEiJMBhTWl/ntrSU8s9CQSnP6xp+9nbWr3Zz7TkGNJuZRZJr6bZJlOgtIcgSClTUggfZxlpznzGFDqKgCwylx56G+MhL2EHnjx5GpM3QhOXvR+ntfstuKdYJtdycz7qXea2lYHTiw/05Dux6VHV1JHVxPppfheCE6gmdG4SN4t1XYKc2Y0xaFDgB94AdAT1x5DPLHpWMuJWnCuZ8jUedn3cFe6ewmk9YPyUvT2YiYM9YOSXmT3alK/WVwhR9V+tPjiI5nnzPXwr5TrqR9BUvpZOLCQVp9NKKiJsg5jKh/fvYDTG6thksyYKfagFONPtDDjbnDycT69Mg5BAwfAjnVf7HrDYnVK7FfGVPwHVZSoA90+jwUknoseIwCOXUGuvC8uNkKTxA8qh7ejZHTm5VklRLlCQtxxPElQACwvBwoHwUPAg/HIyKuWyu1rrWfc6zvQHlzHiE3lglp5xXUDiyZvA9fAqjMG5wrvBRPgPB1lfj+ck+KVDwIrzK+weXPBplar01qHZLXi9OuviQ26ErS2RgONKJACk+CgQememc0+50dxha47wKXG/dWjOeFQ6p+RyM+OK0+aFkW7LEd5jhlp8Fruy+1A2mpHtmbuyx6OA4Z6jz6KLNWJxqmCAeZDf+2al6/T4NrVImXaUiKz3iykrPvL5n7Keqqh/Xy1Rr/HeRgFtlKhnoSFKpvXC7y9QzVyLrcSVqVkqWVED4AVNm3++xRuccABZvFth/wBJ3Cvijj3pXv0zwGBxKed+3SKypjTsfuznHfvDKj/VHQUzn1Xi8ul64SHXFKPVaiSKU7Va7SVDN1iZPUlwpJ+agKdlvtFvGO6kxXCf0XkK/caU34KEbsTNfqmPu912sf2txqBun8oIDR6ZTTa08W7e/IW3zS0tWVdeQqcOxNYIt5vN+MphDiWjDA4hnHF33+FMa7wi3aJ3C0eERnDnHL7JqWOwQ2gyNVKJ5pVb8efP2ioW/wBXI+yVMnMY+4TO82WC03GlhiGhaT5q61ssUKA0lLTCBwgdBilNIbSOQSDWu04tQGcYAr0XgPHPwr58e5zzlxypMNAGAFtJCcZz4V9UsoPCTgkZGfGtUuEgczj0pQhNC4NGIEf6SAVMhP5/iW/MnH2cHrkc8jHG4XIJwsSXDwgnn8PEVsRmw60+pPEpbSQsJAzlA+0fxB+APlWgVKPifux05VlhTVwZbMlvHE2oKA8FDxSfQjIPmCRQG64K8dwyFlQSoEgH0pS0+FSLkmAtIUmclUQg8gSsYRnywvgPyrXvEVNtl92yFGO+hD8ZeOS2ljKcfDmk+oI8BWi3JfZcDrKglaDxJV+iRzBo1Y7ReHvtWwvjZWpDySlbZKVDHQ9CPlit2wvtC8wVSSO5ElsOg/oFQCh/ZzW1ryO0q6s3qIkJj3qO3PQE+C1Ad4n4hfET8aSoTROcnHLOfH0/GuHObGdUMJkZkc17MNUOU7CUr3461MrJHik4P7qjnWeuIOulDZ3R1wxedQXd3T1wLiOBUOE2CZr2D1BZSttChyKlED3kkU6t7Nct7cr1LqT2YyDEYE1pgHHeuupQW0Z/WcWkZ9aiHXtnO2V8261GpkyLlL0/KtNzuCRzacU8Zsl8n9IBUxQ9QkdDipS2UjHOEzhl53uz6FzRva+A0SDi6UiNhw3TePQKWty9JJibt2u6aNS1b7c9plWnrlMSsIcajMPx1tpSfzlqHGlIH2QSvomoi3D3kn3Pca3T0oQjTFlf/JDjqeiJUmOt2OkgDAIbjk46jvwKWN7db3jTGktLwbGZFw1Bf2sWyEkEly4S1lLbeOfutIynPk0CeZNb2gNhoV47M980bMnSJMmVd1y13vktcqe2VJXMa8S0l4EI8Voa4vz68ojDS00dyuuHDPZN64JIc/04D6cCnYkjt7m09OMvPePgOQ8ylPstagtdj2vu13cc4pV6vT8J6Qs8Sw0grQ2keGEIQ2APDHxpB2y1vJ0hfZO0iYYW1E1FIabSXOUC3vtB1hHPmfr3g0j0Sfk2NpF3BvaC+7XSnG42rNLyXFuqUkFAkF8qS8PBTSytlXF4ocOKWezZoa375aN3J1HqaVIiXPVBjwmENK4JENy3pRh1J8Fe0LIz491TyeGCkfX1Va7MQc3xxkjcI8N0kjySFS0dlG+Md85Pp09FL1n1lp/U9wuMHTs1cw2p4Rn30NKEcvYyptCyAFlIKeIpyn3hgk5p33Br2W32thWStyOqUserjigP7qEVCeyd3aMzVOm0xRFcsV7kwSxj7Lee8ZI8SktLQAT5elTpq5vub07EbPuRGmYw9OFtOfxzUPXxMppnRgaaY8QRnPrxXkbjJu58SkpGc8vGlWQtUW0xIxThUpSpSyT9pKfcb/EOn5ik+2xHbhMYhMc1PuJQk+pOM1u6qlMrvr8WKoezQgmIwBz91scJPzIJpm06ZXTu88N9VqgpwMnn0ya9oRxx1PqUQkEJScfaWeePu/eK1Gw686hloKWtZCUpSMkk8gK37oBFkJtiF5TCBQs55KdP21fHI4fglPrQBoui7XAWr1JOQTnwr6j3sAdetY1LAyRj/vrdZbDEUy3hkuEoaT4Ej7SiPEDkAB+d8DRgnVdE40Wpwkq9D518Wyg8sDI86+rVg5zk56+f/fX0EqXyPiRXoOEcUnTrWxISpDiAQR1x1qtfbD0pb7fs3qO5Mx0JcSywSQPOWwP41adQATnx6dar120GyvYrVisH3Y0b/r0arFsvK9t3p2g6F7fumFzANK8+BXNiz2hMy1qeU2FfWlPr0Fa7lrkRHeOM4pCk8xnkR86dWhYynbA6oIKiJJTyGfzRSnMtccpPerabJH56wn99fQ8lY6OZzDqF1S7Mw1dsgqmndcW544SBaNw7zbCI91R7WyOXvnDifgrx+dSDY9Q2i/tJNulgPcuJhfJz5DxqN7jbrcgq4rhE4T4d8kn8KQXEs29wPQrgkqB5cGQR88VzJQwVurBuu8E5oNrrtsw/sqh4miHIkFw8jxT13ETwbgQ0nI+qif7CafNub9xJzUPNXWbeb5Bl3B8vOhbTfH4kJIAz8qmq3NEICh4GmF0jMEUcZ5BW32fVbLtX1tZEDuueCOuoSggttoKnVBCUglRUcAADmTny6/Ko71LrO5anmp0po9h54SF92VMoJXIJPIAeCc9enmaya+mX+5ahj6Et3C37SplCML4Q+XcFHEfL3hV1uyr2UrfoSIjUOpoqJV0dwS6tHNAzyCQeYHLPgTnJ8hDV9wpNnaUVlV3nn4G/umu2G11Vcqh9ltmWMacPdwPkOg+6T+yb2RIukmmtb63ityrsscTTZHE2xy6JGOZPPKj8seNyIzCIyEtNoCUoASkY6Ach+6vEdlqK0GkAJSkYAHpWYuk8+VYdeLzU3qqdUVLsnkOQHQKr01JHSR9lGNFmCyDgkjPL/urnb9JHus/qnW1m2bs76n2bMlEuWw2CSJbqcNN+qg2rix5vHxzV/rteIlltU69T3A3Gt8Z2W8v9FttBUo/IAn5VyJ0veJ25O9V73Qv+JLUR+RfpRdPInjwwjn1AWtvl+ihXgDV59mVrFTWvr38Ixp/yP9sqD2hqtyJsA/N9gtvXzsXRmm7dtvDeym0tqkXNfL66Yr+c4SPzQrKB5pQg+JqEJchcp9b6+qznHkKdGub49cpb0l1ay9NdLi8nklH5qfgBgD4U0Tjwrd4m6ZVJkdnRFFFFLJNFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCF7Q4ppaHE9UEKHxFTBp9bOodA6ks7sgp/0QT2EDxUjCycenAR+2fOodzUrbLXOLHvML29Hex3A5Ffb/TaVhRH3CkZxplKRnXBVu/oxdbquOmtabYPkH2V1q7Qs/8AKjunR6ALbYP7Rq3mUkZCsjNc3+wZeJmie09/JjvMJnsTrW6M9S0FOJ+PNmujspSUzXm08khxQA9PCvnj2i0Ypb257eD2h3rwK0DZ2Yy0m4eLSsxdAHpjFackF1tSUJypXIDzr2Xc4Gfjimnujq06F271Lq9DqG3LTbX5DalDo7w8LY+JWpAHqoVS6aF1RMyFvFxA+ZUzNIImF55Bc6t1dQQtye1Beriwe8tdsnqjseS2IY4Ar9stlR9VVtPrU9xOOKK3HPeUSepP/fmmjtTA4oN51RMBU7IWiE0rrkrPG4r4gJSP26di0nA5V9D1jGwllKzhG0N+QV/9m1A6nsxqpB3pXF3pwCZ+prAi7OJcdZSsoHCknIIHXH76b6tCNOD+bUk+hqSHGgeZTXlMYuHhQjKvLFLRXCWJoDTol7jsZQV87qiVmS7io0/3PFn7Lq0j5f4VhVoO5J4lx3+FKOZWoYSB5k08b9rCy2JC2QoTJQ5BttQKUn9ZQ6/KmtFj633EeLUNhSYoV1GUMp+J/OPP1NSsM9SW9pM4Nb1KzS/0mzNpd2EAL5ejD9zyTekyLhbH3IrV271KkFB7pZUlQIwRVquwVbZYd1PLWw4lpxcHu1qSQF49ozg9DjI++mlt72bFOrTIubJlrHUuow3n0T1PxP3Dobd7T6GTpOCtqO0ouPqQMBPNWBhKQAPXAGPHpVN2u2kpJKJ9DTd5ztM8OarNvo531AqJiQ0cATnA6aqSC9wpCB5c8V7jIadWUOqLfHzQs4KQfAHPgTyz4Zz4U2J2u9FQJkm2P3tcufERl+Haoy577Sjy4Vhv6to5/wCNcbx402LFvdC1FuEjb6Pom7W8KtirmJ82dHUpSe9CEgMsFxKScK/pScJ6c6y6OyVhidOWYaBk50OPJTZutJ2opw8F55BSY8p2O4tl1KkLQeFQUOYPr+HOvrUhTbiXkOKQtBBStKsFKs8iCOhpStjDN+bbtch9KJyE8MF5ahhzyZUonl19xXgfdPIjhSH2JMR92LKZWy6wSlba0lJQoHHME+np0HWop7N05CkGuDtCnZeGEXm1DVlvCCvjDVyZQMd28ejmP0XPwVkeIFN3GeZUPhW3pS/IsVy/0qP7RAkpMeZHxkOMq6jHiR1HrjmM5rZ1Lp92xzkoZUXrfKT30OQDkOtHGOfiRnB/wIrl2CN4Lhp3Xbh9Eq29s6j0lIiJHHNsGZDKce8qKo/WIB/VV7/wJFNtKSvHDjzHl1/d6UpaUuzmnr3FuSUqUlCsOozyW2eSknzyD9+D4Uqap08LDfXGWSFQn8Pxlg5yyoZA+XMfLPjScj8jeXrcsfu9eH6pSVATd9t2nA2C/YZZTxZ59w6ef97B9MGkS3JCgrI+yrBHlkcv3GnboNvinSbFLCu5ucdUYjw4iMoX/h6EUhRYMqLPfgyGCyspyUfoLQrmPl74z5VGzS7zSOf7LuHuPcPX0Kinti6euFx0C1crMw49JftcZ0NITkvLhyErWkeZKEpAHicDxqtvaO3PjbjbXWu56YnreXCLUpTaDhYSlKkuDlzxhw58CkEHmlVXC7TuoVaW2psd0TbFPtN3dFvmSSwp9EGK8cuPqQkp4uHCcAqSM4ycVRCNo/ZK5Xybdhu3qKzquT5eeZjRYAYS4eZIT3xA8/PqK0bYqGKqpYqqp07Jxc0jXjgOaQOCYS1ZhL4GtzvDB6joQpBTupar5uDs3ruVcWl21luXAW4VYSxNdjvJjlWfEKdAz1CkK8ufQraWzWS57N2CfYVJTb02mKljHLuwhpKeE5/OBSQc8wQc1zv0x2UuzheokiP/AOFbKt8WceN+M8xCCOLORgd7gEHGCAMefLFTlprbLQOi9IztIab+kjuUO23EqW/HSYhVxq+0pKy6VIUo8yUKGep5k0tfdl7ZtHCyCmqN0Q50LXDQkkHhqRk8OKh6q4zsn7R7CCcfT91E141lY9G7y7p6plTmm7FFgNWghCwSt0v96lCR+cpDDaEY6gqQk4JpC7Om5I2/03eNWXqU6xP1hcO/YZBGGYnfrdWGhnPFhbxJPLKmRnrW3q7sp9nePBa7ztbyrqzGcU+3HKYa0rcUriUspL2FKJySVZJ5Amo+vmgNmGu9Yc3k1HcgpQLnEmEA8lOCEK+s+x4cOQPHFS7KC0VdI6kbIX7wY1x3HahgwOI5nU/JSNNcp4wC+PQZxqOJ6qzHZXcf3I3L1Rrx1lKIGpLxHXHSk8nGo7RStweHCo8YHnw56YJsNqRYk3CbMHRbri/x5fjgVAnYh1uu46+umm7XE9t0xabF7W3cRBRGEd/vODuMtLU24O7HECOEjmMYwasFf4boSloNlSn3EI+I5qI+8JrNNpxJFdOyOgAGB0aAAM+OBqnlFKJsu58PXiVl0RH7l2dqN1vLdpjrdbz0U8r3UD8SfupsKSULIUpXETk8XU+Z/Gnu4ybToqNDWktvXN4y3fVpP2fv90/OmjIhLduLUFhta33OBHd9CVq6D9wPwNRgeHP3RyS8btXSHnw8gljTTXsMObqhzhHsSQ3F4hnMhQ5HHjwjmflSG2CVkqJUTzJzkk+NOLV7htrMLSUVwqbgNhb6k9HJC+Z+Qzy+J9KbkdmTJkNxo7RW68tKEJTzJUTgAfGlcg6BcxZ3TI7/AAJRtVsVcppC3AzGjp76S8ejbY8v1j0A8SaxXGeZ8guISW2UJDbLYPJttIwlI+A6nxJJNKOoHY9oit6YgKC1JUHZ7qTnje/Qz4pSPxz45pASoj3lADIz5ChxIGF0w7/4h9PJZwSBzIx45P763Vw2orQLue/cGUoIxwJ5YUT5nPTwHxFbUWEi0w27rcUZkvYVDjq8M/0q/QeA8fhUUbz7wP7UW+03g6YkaiN4vDNtcbTPEZY71LhSvjUhYJ40pThWPtjn1y6oaSWumbBD8R4ZSU9S2Fhkfo0cVIjhwCSc8qgftjxJE3YnVjMFhbzqosbhbQkqUcTY6jgDySkn4Anwp42jfrb+fETJvq7pporUEcV3jJMUFR5D2thbrA8vfU2fSnDqu0QdRWR2K4W5USW0FIcbWFtuI6pUlSSQU8s5B8Kl6aCpsVbFVTsOGuB8ND1Caipp7pA5kEgOQuNTcue017A1NWw0XMlJJSArpkgUuxNDXa4t+0IuDL7WftNLK/v8qt9uz2XbLeXHZ0eD3chXETIjpCHFf1ugXz8VDJxzPjVYtQ7aa/20luzbeXZMZokl1hJyE/rtnpjx6j1reLftHSXdn+meGP6O5qCpqdlDKG3NjpIv6XEYHlwSMnbl44JeUc888qyp2+DfVKlH40s6f3EgT+GPfWxEfPIPIGW1/wBYdR8qeiGkuNodQpK21jKFpIIUPiOtdz1tZTnEui1Gy7K7NXuPtaAh3UH4h5hMWy6NajTG5CmBltQUMk9akm3s8CEhXPHLOPGtdmOnP2cYpUYQkJwojHpURVVL6gjeK0ixWGmsjCynbjPFR7u3GXb5Vk1JEWUPpyyVp6hba+JB+OFAfs11Y2u1dB1zt7pzWcJIS1ebZHmlAAw24tA7xH7LgWn9muaGv7Qi86JuLQRl2DwT2vP3TwrH9lef2atf9HxrZy/7KPabkyuN/TVycjJRxDLbDv1qPkVd7/ZNVnbam98sjKlvGJ2D5FZDtXSG2bTPIHdmbvevNWpU5x5IHL4V8LmRnPjWAr548RyOa+E4VnBx41jSa4zooS7beupGiOzpqFUJfdy7+4zZGiDg926rL2Pi0hxPpxVzo0Y3Htm2VxuJbWmVeJ4iNLBwChCASB8nHPuFWu+k4vpj2LQelm5HOU7LuLrYPUJCEIJ/tLx86qrqiS5a9FaZsIaDYjQDMXgYKnHvrCT6jvCn4AV9D+zyjFNY2PxrI4u9OA+yz++z9rWOHJuijG9yPaJ7nCrKEHgT8q0K+qJUoqPUnNfK0MDAwoA6lFFFFerxFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEU9NtZSGLqwHThKJTKifIE8J/CmWKW9LrKJLykucKg0eEZ8RzB+WPxpOQZaV0w4KmDQV7Oke1tpe/NrS2n+UtvU8o8h3by0Je+9K15+ddS7gFie4V5SoYB4vEgAH91citx5BtW6EO7Me6U+yy0EdQUnkfvTXX7Ua23bu880RwO4cTgcsEZB/GsV9qcI36ao6gj5YKuey78dqzxWkMkApqsH0gWuTp7aKDo+IVJk6ouKEvEHH+iRh3i0fNxTB/ZNWa70pSTjp5jl5Vz97c2ontYb82bb6PLS6zZYkaOoDmEyJOHV59QhbST5FJqtbB0PvV5jkcO7GC8+nD6qWvJc6AQM4vIb80x9LwE2vR9nh44VutKmup9XT7p/sJR99bx6jnW1KfRLmr9mbCWkfVsgDAS2nCUj4cIFMnVO4lvs6lwrKUS5Y93vftNNn0/SP4Vq7I5K6YlgzlbpJW2/ZO0xirkDWsaBjmSBwA5pw3W522xxfbbrIDDaslAAytzHgkePx6VH1x1fqLV0s2bS8N5lp0YDbRJcWnxK1DkB59B51v6T2v1puVPRcLqqShiQoELWMuLH6qT0HkTgeWat5tP2d7Fp+GhT0JKcpBKQPeWf1ldSTzpK4Xa37PtzIQ+XpyCyC9bXXTaV5hpAYYD0+IjxP6KvW2HZquF7dbm35svnIPACe6T8T+cfQYHxq0+m9qdH6O04/c7o5EhQ7ZGVIkyX8IajtAgFRwOQyUjkMkqA5kgVKFt0/EtrAQy2hpCRgqPJKU+ZPgAPH0qCNa7jRNWyWbnGMr+T7Lw/IUVYKEyXErKTc3kHmon3hHSeSE5cPvKFUk3at2kmMkri2JvHH0A8SqZc6ik2bpu0cMvPDqVJmh7tpLUljY1BpYyHLXKC0sLkspacUELUhXEgE8PNJ90k8sZppbj7kSVXSbo60vSrbaLYtyHdpsd5TD8+ZwjvILSk+800ylQ75SSFLK+7BSAVFH2s3Os+mbbuRcn1FS7B3+pIsdfNJdkRu+CRnPu+1IWT/AKznmq7RdTh6DFbmOyFobaQ5Nd4z3sh5eXFgLPPjWtSlKV1AOeZwC6t1l3KmacDQEBueQIBz540UFf7/ADOtkIptDINT06qTjPW/DaihpMTTkRZXCtrDfcxXnASkqSgcikFKgtzmpZSUZyFFLm7OLkrV+vNTaxcSnuGS3bY6x+clkYJ6dOJasfd4VC7uqrzqxxGnNOtNJmvIDfCyn6uEzyHER4Y/NT1Jx1q3Wyegou3ujotsZaLbvDlfEBnqSc+ZJJJ8yT4UX+QUNE9kn+5JoBxOOJJ80z2HtM0lUa6YHA0GeJ8VIeC2kHJ5eR64p5QX2NwoHsUt5LepIjf+jvuKATOaSPsqP6YHRR6+Z8GqYi5rCpkAqcSwOJ9r+kbHisfpJzj4dD5nFBElh5MqOpxDrP1iVoJCkkfnAjpjryrNs44rV3N3+HFZVxX2XnI0hhTTrSihxtYwUqBwUnyOQR99PjRrqNUWpzRNxdQH2yqRanVD7DnMqaJ/RI+7KiM8hWw2xH3HtRmRkpb1FDbHfthIQJjfTiH63QeHM9MFOECHFdQUvRy6xMiq41cAIVwpIwvz4kn7Xpg9QSUJHhq5yJG7vBwQxalolO259styUE8KOAc3ByKPQ4Ax5kAePJ9xG06l0a22r3pVmwppWf5yMeqfkAceXCB416u8CNqO2I1OwlLc+Phq5IbTgpUAOF8eePdPqnl+bWXSik264B1IDQWpRcbIHAFDHeo9U/ZcT6AetMJZMLwu325HEJMsgkRXW3eZfguJSVZ+0E+8j70cv2D5Gl3WVtZTc2LpHcy0+6guOeBae9xSs+hIV+1W1KsaYN072MkKjPfVhKvzRnKP7Jyj+qRSk9AbuOnBCWTwttrjheOfBjAP+zULU1XEt/zkV46QAtcPL/PVJMixs3fT8623BKgpKkv+6cKSoclY8c+7+NUT7ROwVz25ujmrdKO3AWF9+Xd7v3aWliLxJ4i4kFBUEcSCFfa8DVxtyN9tudorU9L1beWjNdircFuYeR36vdzlQJCWkZyAtZAJ6ZOAefe9G5e9PaUtx1U3p52Lt5Z5iHWoqyuOxcUoWCptsngXIVjJU57qE8gAg4KrtsDa7wypFWD2dNwcXaB3TAPE9Co2trYjmMDednQD/OCkvZXX11SlhmBd9fBORhdpsaJHLlzx7GrqKtW1qWS7YONzWG+4dKchKtCoKAfj+TuYqg+zu72i3pzFvOxVzuDjzaltItSZKFqShQQrBEr81WB+FTpG342ll7fPa6jdmXdZ+xxwvvbs1dpQiJ4SUn/0rHIpIPPwPlWgNoKuCqmDQWAjmRgk8MY59FHVjmyhpbqfBaO8WsLyG3Ur1RuOEDP/AJx06mOk+HhDT/D41BmgNE6k3n1V+RoV3vSrSpUiBc5Cm0N926pKeBoFTY5njVkY5AZ8Kc26+9Gi3oqLZA2L1HAlyYpnMquch93/AEZIClOYW+RjCk9T+cPMUwtrYG6Td7c3l2csyIbzDLa37PBUr2iY2AVl9DbhWl5RHVnPEoJ91OTy7tdFUwUL8HckOjHO3SC48OH08QlpJoxu93eA4gcV0x2h26haG0m9bIoORHj25ClqJIbQnhxz9DSw9bRMu4i8ZKQAAQMlK3Tgfc2hKvmahDs+9tzb7cluNpjWSUaYvzSu6eD+Woynwk4QrjOY5JGAHDw55Bajyqy1uisJeVOSoL4iXkqByD7oSPuSMZ59TWM3eguVqqiLkw7558jzJB5qRgrY5AXRHT/Am7rAIduTUVIKokNvBSDyDTSQVDPxUlPxFJ2k7cuKuXrC6gumOlTjSMc3H19Mj9pPzWPKl6da1PueyoQVB4jvz4lsK41p/acUB8vSst5juMwWLVHcSFJ+sU4UhIW8c+/y6JQOJX9gUlTTtcN7qli7DWxt/wA6qO5xeW5IuEpSXJDzqkBQ/Pdzlwj0TnHxV6Us2Zo6YsR1NIbT7dOSWrehYOEJI950+eQenw86z2KxxLxOS5ISW7TETxKLiuEdyM4Cj4FZClHHQE/q1q6lni53ZydJbUGIyUhDXDwgZyUNjyJAyryBx4Cn8TwlHntD2Y4c/wBk1ypbzqnXlFalq4ipZ4iT5nzOacFrtkS1Qk6kvrRKDzgxVfakKxnjI/QHI+AOR16HNZbNGSyvVGpE8EFo/VMkc5S88kgeKc/I8/JWEe+XWdqCaq4zBhPNLaB9htPgkUvgYXW92hw3gOP7LFcLhKu0tyXKeLjrh5nJyPQeg8qhTtTWC4Xbae5S7VwmXaS3cGgeYJYWHB88pqaY8dxbTkl893GZIS46scuI8+FI8VenXqfCka/xkXW3yISgSl1Kh0Gcefx/z8HdvqHUdTHO38pBXlRCyphdAeBBHzVJEalcu0VjUmk3X48qcwCptkkqkNLAy2cfaVn3SnGFjwJHvOHbbdedtyluTBYXJ08Cp652VhR7pDJ95yXEbwUtONjLi20YQ4kHISoBVMLW2lLts1qeXb5TfBp6XIUuE+lOEQ3F8yyv9FBPNB6dU9QQG/ctUwC4JEXCMlJdQfsgk/aA/RJPTwJJ6HA1ptLHVMxEN+J+vp0PiPosG3K3ZuvLGZG6fQhdA4xtV6YYkxHkPx5raH2HRnhcbcHElY5cwQQRnqCKimzK2v3dt027aOm+1Nw5Sokht1tLcmM4OLhC0gkAKCSUKBwoDwIUAy9ldynGOzDq6Q5MWLhoZFwtsRSjlXdLB9k6/orfSgDwShI8BTOtF5c0tqeBI0etbNwskKPBiFaiGp8ZlKWlxZH6SHeDiCuqHFBQ8artPYWxOniLyHNdhpHlnX5haTXbUx0TqbtBlsgyfBJO8HZUhTVv3CyRxDkH3g4y1hCzz+2gfvGD161XeWxr3aecmJdIqlRHFkI4gVR3inGeA+BGR5HnzArqDpq9af3D0vC1RY0uGHPRnu5COB2O4DwuMup6JWhQUkj0yORFMrcTZmxartr7DkNpXejBStAUhXlxA8z5Z68qf2vbSe3vNDc27zQcHPEeRU/HSujc2utbyx/EEH7ql+mtWWXUyUIgudzKP2ozh97P6p/OH4/GnKlCsA9Kbm53Zn1JpOa9ctJpeWhBUv2Un6wYBJ7tWff6fZJ4j0HEaa+md0ZkJYtWq2nOJtXd98UkLQc8wsHr+Bq6Nhgr4veLe7eb05habs17Rw94or43s38A/wDKfPofopSS00+lcWQAW30KZcHmlaSk/gaUewHqt7R++V30BPCksX+G7GCCeSJUZXeIVjz4A8n9qkiJOjTY6X4ryHG1p4krQQRg9KZLWpndtN+9O68ZV3LTc2NNeKRzLZVwPj4kBf303bTe/UlTb3j42HTxHBHtRp2SUtNdYjncdjPLDvFdZyU44k4I6/OsT7nAysg8OEknl6VnCEp4kh5t1I5BxJ91f6wPiMY++sMsDujwg5WtDfTP2lAfxr54LS1265U0OBG8ucf0jN4eu2/1r03xHgstggw+HPRbylvk/Eh5A+QqHd25zSrg9H7sAR4rTISOXDkcQ/ApHyp99re4fyj7YuoUZ4kIu0CCAfANtMt4/umow3YebXfrqpaTlchKEehSlI/cDX1Ts7TimtlLCOTB9RlZhXSb9RI7+oqOjRQTRViUciiiihCKKKKEIooooQiiiihCKKKKEIooooQiiiihCKKKKEIooooQiiiihCKKKKEIpY0uOK4rT5srFI9LGllcNyPq0ofDpXL/AISvRxT13iSUaptTpHNdqjrPzKq60W+YZ1is81auJUi3RXFH1LSDXJrejJ1ZbGj1btbCPnlddXrRHMWwWaIerNsiNnPgQwisf9qI/wBNS+bvsFb9l/8Aelz0CzS5TMaOt9/PdtpKlfAda5Vp1dD1Ju1qfcu/ye7aemS5yOJWSVuLUG20+ZAUCB5Jrob2jtZfyE2P1dfmpCW5S7eq3xATzU9JIZ5eqUOLWP6ma5xbZ7Man3JlodZaXEtmcGSpBJdOQChsdVHmefJI8SDgFD2e00NLQVFfUu3Wk7ufAanCkq+vkpblA+Jge5neweGeWViuOpdRa9njT2kre+hl08Pdt/bcGccTiuiR09BUz7R9l91LjV0v7SZchWMEpy0jz4Qoe8fU8hjkOlTltP2frFouC0gQ0oIAUor95bqv0lq8T6AADwHPJm6DbI0ZpKGmkpAAHIeVJ3zbcNBpLYN1vXmU4kgqrvUe+3V5e8/IeACaOkNvLVZoyENxUhZAClkZPL16/wCfCnwxbPZ+H6rgSrBSAQfI49D9k4/WHnSNrbVFu0DpC5atmRzJMFCUxIaVYVNmOKCGGB6KcUnPjwhRHMVCu0+tpeh9aPWrUt7fmw9WTu6ucl9fE1Hv6hlLqMcksSM90Bj3VNtkkJNVOmtdRdaaSrLtdd0H82NSmtXeKW21cdG84LlIPaI1K5pzbVy0xWlibq2cxp+O+lzgMdt0Kcku5HP+YZdb9O+z4VW6beEPzQstoQ1GZK0pRyQ2BhtttI8AElWB5IGOlS32xL97JZtFM8JQGr69xcuilxHEJ+B5H7/Sq0tXrj9qK1KdWpTaG2E9VuAHhSPiV4q3bP0W9bYyBxLifPOPss126Ms9xDT8IGg81sRr7bGFbhmY85/vnbYdpaQ0ApziJ4irhzzwMJ+Kkg4BzWppPau533uomrNUsWpAV/wKO27KnOK5E/VR0Oupzy5qSgEdCa09O6ftypZut3u3fxWpRWmPEcW2u6SjnjVxpwoMI5oBBClZPBgFRD/Q7qRtpPeyY+nba0kqjxkgs8QJ5huOjB4ueSpwp4upUo5q11EopMshOumfDAxp4pkbpDBEyN7N4sGB0HmpK2xsW2G1TyXnYmpoq8BTkmXpC4d0P0jlLS3OniQSPPynuxX7ROrYTsjQ2rrZe1xUpVKiRnVolRknpxsOpS6Bjx4cfxpxP1Va4CURdNx5F61DLKERA8vi4VdCruweBKQepVxY8CDip52W0sNmrE5fdeXptu8354y5Up90B5xeOfDn3ilII5AYGR51Rr9RxSQmolDjM74RnJd1OOgV22VvNZcAS+NrIhz4a9FL1vdehPtTY7imnWzxIWk4II5Z+6nhFiW/V6O9tgbt97bTxuQ0kNtST4qaPRCj1KemeYwM1G8/fPZ6U0PbdZRY8w5V3rcZ5SFnJOXEttnnn85OTnqCcmkFjfLazuxLe1eWJDWFoWzb5KOec5Ci2Ak/MfI1Szaa+Ud2F3yV197ps/GMqY7KqSxOVMiF63XCEoqda7vBSR9pQT4j9JPkTgYyKe8m2x9Tk3yEj2e4Ix7UhlWQT+a6jPUEE5xniAI+0nnWPU3bq2nt0RtEy2XCfeGSAmexJgMJdSPFQMnvAryPBnxyOYrTj9vW0GOXdBaOfvdwPP2WImVKWlRHvZCY6EDJwcBwjIyMZrs7I32ZoLICB44/wJo+60gO8HahW0spXCmpRwjBy06yfskEZ4B4KH5yD5caT4UoybfBtaO8ckIRGPCpLjqgAgA+4SfHhzw+qSM/ZqlDu+3bY3NYdRoPZWbp9C1YM6TFjweFOQftTVqR1weQyDkik669nrtDbhPQJu7G/ESI50cgWtUi5yWx4++7wR21YzzY5c/GuhsQ+DL7rVxxDpnePyCQNydMf9NE5x8laXcPtVbI7e2wsXXVrFznIUW/ZbapL3d46pddJDLfCMfbXxHHJKulV2mdqrfvtAyHtObDaJkQbK4pSJt6dWqLGHLGVzXAjgHjwMpS6eiVVpM7YdkvYnvLrrmVDvN1ZPetS9WzTNdbUOf1cNGG8kjPvoWR51H26n0iFni95atq9OKnFr3G5k9Abjox0LbKcYA8AcD0qwWaw0DXj+DUj6l4/PJ3YwfXims3aMGa2URjoNSpJ032eNv9Iqf3E7QGtouq7i0fae6mqMezRXgPtlpZ4pKwMAF04IHvJVk4gHtD9siJdkyNKbVMcUZLYjC6vMhKUtjlwsNEcKU45AkcgeQ6Gq6bh7rbh7n3Zd01lqSVOUoktsFfAy0DzwhsYSkfKkCBZp11JMGM7IUnqGmlOH58INaha9jppJm1d6lEjm6tYNI2+Q5+qham9xwMMNA3APFx1cfXkrt9hRh2bqjb29PFTi3pF2guqUc5IW294/EnnVhtHLZifRuataUgcUR28NK/rJnyBn91Uo7Om9d12Wf043c9vdQXAWO9S7iUsR1ILiHowaCE5T14khXwqSm+2PZtN9nDVeyGoNstV2+ZfrhcJcSW60hLLaJEpx5KF8ZSchLnCSAeeagLpZbjLXTSRxZaXNI4cAXZ+hTtldB2UXfGca+afHa9iNwtS3OQlrh9i2vUU8sAKVIiND4chVSNku0pqjal9u13JlV3sJPOOo4dY81Nr6+XunI8sdalztCdqrQG8CdU3PTFuvlvEvTUOxRkzoaRxOJkd44Cppa0j3UpxkjOD6VVCLEDjQKFJV0zzyfn41N7I7OGotbqK6xYzyPqcg+vJM7ndXU07ZqV3yXR6HaOz72qdPC+W66eyamQ0Qm6W9SY13iKPVEhHSQ2fELCsjklafHWs9z7V3ZeWmfbnjuDodopRIEJpb6mmQMZdhlQcjnB/nGVd15hXMHnpbZF8sFwautjuUuDMjqC23o7qm1pIPLBBBq0O1X0gGtNLNxrZuXZvy9HYBH5QhkMTceahgIcI5+CSfE+NR9z2Uu9sYWRNFXT/wAkmN9o/pdz9cJeC7UNwOZPwZOo+E+Y5K7O1Xb22a1rKYh6hcd09MUk8S1BT0dCwMELUEh1rmVHLjaUZ/PPjPka+6c1ZHRN0/doN0jy0YS7CkpdbUg+SkEjCuXjnGPKqVHVHZE7UkNBujNkdvrrIS2+VG03qKrpgOpKQ8RnHvd8mtJ/scay07LYuO0O+MuAGEBUdjUaHUqQrHIJmw8PpHP81ITWY19gsFQ/shI6klz8MgIGfAqXZ77D+IwCVvVp/RXruDfscREBooSOa3FlPuk468I6pTy5eOEppCbsEabI9puanEQox4le9gqJ5qHqtXLiI8wkHI92nUDcTt9bXMyRqTQcjW9ta/m5ER+NdAOHopJbKJWP9aSR6VuW76QG726GqJuntfK04Whwp9ojTIRJxz4T3T6ckcuahjJOQabO2HukeXUjmSt5Frgc+i7Zc42DdkBaeeQrUamnC6PoeljuIEX3GY7Y4Rn9FPhkcuJQ5JxgZI56LMCM5F/Ll+UYkAe6wynk49z5JbHXHmr59edV3sP0hOyFxkoZu2lLxlOOAolwHGuXQFKpCF8A8AEZ8wTmnFI7UW0OqHzdblrOQ25jDbK7bJ4Gh4J+rQpP3H502m2dvFLgS07hn1+ydxV1LJ3WPHmpJu9ydu76VqYRHjM+7HjtjCGk+nmTgZPiaSblcLRZbc5edRXq3Wi1x1BLk64ykR2EqV0HGsgE58BknypsWzf3ZB5XeStcMSHCeFLCYcptGAeXEpTacDryHrkjx+at1Vt3vRY52gGr7bLkbpFcaTBStAJQBn6tvr7vI8hxDGevOuIbdOyQe9xODM64HLmlnVMZaWwOG9yTB3Q3G2V1jEXZolxut/BSUKXbNMzZTZHlxqZCFg46AkVVjVm0ul2pDkjR2sV2mQoEi2XqHJt6Vp5jDZloQ2Mjlwh4k9APClm9Q73t5qV7Qe4qpBSx9XZ7t3nCp5oH3Wysci4OY4Vgg9QAedY5b98WTJ07fnJbSVhPsjjgbcKSOfJR7txPgUkgnOOFQzjUbdCy2gNoydw4IOctPzGP1Cyq8bQ1BqTT19ONPQ+hSPpeS9prafczQd6kqbvk5MS4JZWlSUutR5DLiwknqe7QVY6ngXnkMlZdu0dT0R9tHEVuraWCeYCwTn4cSEj9qo/1hbROV7XDZ9jkRmVNy7YkFtqQ0cBwtoV/MrHJXCkBGR7oBGFZINyRDs8dpyWHjD7rgd8HWQQUODyPCOFQ8FJIPOpiopWzx9qDkudkjpkAfoFB3ipiuLIzENG5HjhWd7MetXE601BoOTH7xq8xE36GsK5NvtrSxLTjzdS4w4T5tE/nGrIyIJcwlKSSeXl6Zz4Y65qlPZsvqXN+bT7OOIs2ScXceAcU0lI/tYqfO0RrJLNtTtzb7k9FfvcNyVen4oPfQrOCEr4SOQckqzHbB6fWKV7uKoN7tTq27RwjTLAXHoNQT8lo+zd09ysIqKk/DoMp8X7TMO5RlMSYyXAvqMA//fpmq0bzdmK26pbenxY6mJqCVIlNJHH8Fj88Y8zxeRxyqYOz3rVV8sknb+6uqXc9IpZajvOLClTLW4D7K6SAMrQAWVkAAlCT1JqU5cFpbZBQDkY+INRAq63ZiuMcLzkcOhH91ZoH017pGzAZDguU1ztOvdl7smJd4ynILqstrSSph4ePCeqV8uaThQ8Rivmu73ZtWacg3eAsiVDe7p9pYHGltfMHl1SFA8/1h0rolr/amx6vhSIsu3R3kPJIcbdRlLg8iB48vtAgjqOYzVIN3uy/qPRDz100yy9cLekkuRieJ9lOOahyw4nrzA4vMY5nUtn9qqG7yNNRiOYfJyTqKy5W+3S20O7SndyOpaeRC6Bdl/Witd7FaQu761rks29FvkLJyVux/qST64Qk/OpUebXxx0HqZLOP7Yqn/wBG9rNM7RWqdBvuoD1luDVwjN/nKZkJKV49EraR83KuEtZQ7GWeeJLOfmsVj21VB/DrxPABgZyPI6pzaqj3ijY49MeoXKDetRf7XupFLJOdXudfIPcv3VH+6h47pcF4/wD3k6P7yx/Cn/vUruu15qXiGOHV7v4v0wd0wE3O4oH/APEnFf3l19JWv/20H/Bv2Czup+N/mUwKKD1oqYTNFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRS1plHE/JWAStEdXD8+X+FIo88U99prexcdVQY0rHcvTIzbuT/AEfeAq/AGuH6NK9aMnCVt2iuduSYTSSVIS1HSPiSR/tV10uUYQZ7lvB96IlDGPVLaU/wrlJbLX/LDtLafsKElYuOo7bDUOvJTzSVfcCc11Sflrn3SdcF/aflOK9OtY37UZQRTRH+o/YK5bLtO9K/yTF3Z2ysm6VgiWS+pcdiwpiJ6WAspaecSlSUhzHMgcZOAfKsOm9B2vT7aBGjJCwkJThASEpHIJCQAAAOgAAHhyqQHUAjAGT0rXUxwj7OSKzEXCoEApi7uDgPPirR7vGH7+NeqTm4oA+HpWdLJCCDxcuoFbSWAcDBAHrWVTYA4seHiaZ74J4pUKCe0TqRVuv+kdNtBJU0zL1M6CnPvA+yRDg8spWqS4M+KU+VQY73V4auNrfwhiYVR3DjmglDa0rHqhRStJ80inf2m7s8zvTbi4v6pWkWWUDPQonyCofeQajMXJbNzkhSyEOIaeTz8SCn/wCWK1a2QdnRwuZ/ICPPOSsE2vllN5e7+Xgps3PZl7x9m+yaj7lo3F6MPaVJHEWp8dwtOKHiCXmScfoOVUuxwJN4tcuRLvdttFydjSJUdMx1SeFoDhwgAH61zBSk9AAo/nCp30pu3H0vsBqPTK21KnXHUtzQy6lXKFEHdKedx5qUtLSBy+sfBzyILA0/oSI8GtTz2rc87IUVruEsKdjMBGEhiKwhQVI7sYSFcSG8pI4ldRP2gOtrJmyDDd8lvrx9FbKiroagwT1gyGsG9k415DrlNnTWnNz+5akWuwsttFCeF5ExKHCkDoFY4kpPkMZ8cnFOuybVbm6nkiLLZj21s81Oo4nnD8OIJR96j8DT+hbsXm1vt2XTl/1RwpSn6qFcjbAsjqtTcANcI8+NxzA/ONauqNyNx2IyI2ndc6vi327uez29iFq66L4XCfeccKpBBCU4OMEZxz8D7JVzTP03Wk8Duk+up+uFAsmsctU1kcTnEnQZUlbP9nu16BkC8XVPtkx3hU466e8Usjpk4+z+qAB5jlTN7Zd2N1l6bsiJLwZStbK0pWUnhW9H4wCDnBAHpyqyuiJmpI+jLXA1pepd/uUeMEPTLm4X5DiuZJU4ffPMnqfIdAKqJ2uSWtW2FyO0UBL3EfeyP55jpyz++qdYqmSvv2/K/fI3sH05LTquCKnoAxjMDTRTVL7O/ZK0hGauGtdMxLcyp4sJcueoJbaHVgZwkh1OTgZ9KyNsfR+abxIRA26WUc8Py5k7+6tawr7qZfb1WJGymnZBGFDUoSfnHe/wqhCGFqHOp7ZiyXDaih9+nr5W5c4brSANDjplRlzrqa1zdi2BpwBxzzXSyX2l+xXp1fDarLopHAMD8l6FjKVj0WpgH55zSNf/AKR3a20wBb9HWLVU9KBwJZUlmDHA9AFrwP2RXPFETxxWdENP6NWlns1opHB9VPLJ5yHHyGFEu2nkYCIo2N9P3Vtb/wDSQa1kNLb0tt/arepXR6XIckrz8BwCoe1v2sN+9fJVHuOuJFujKyCxaUCGlQ8lFvC1/tE1F3suE+6BmnLtvtpuFutqyNonbrSUu+XeUCQxGSVBtPitxX2UIHipRA/Cpuj2KsVqd2sdO3PU94/N2UxlvldVab58homhKduE19UiY86+64okrcXxEn1J61L+xelrHrOeLXbez5qvce6ICVOMwLq40wn+sGmCUg+qxnzq/wBsR9Fntloi1M6x7Tmpol2uDQD7tmiySzbYYx9l973VukY80JHmqrmbZ6m2nh25GkNkdPwXbPbQW2/yFDSm3IIxkB5pPcqUcjOFEnxp/UVUTBuwjQdEjEx79ZOKqtsh2ctzIFviiF2Jdi9IsOYd7/U1weulwQSOpJQ8sHkDwlSQMnAGTVvdHaI1lYWG2Zw0XEQEjiYtOn1sNoV4hKi9zH7I+FPeI9clJBkJZj4H2EHiIHl5VsPzCw2VpUjIHVauH+FR5nL+SW7PHBY0WdhSQp9qOs+P1AGfxP768z9MWK7MKi3WzQZjKgQtp9hLjagfNKgUmmRrvea27f21V2v07TkCMnJD10vogIXjwStxrhJ9M1V2T9LdsRaNRGxXux30MpJC59tVHuEVOCcniQ4lRHLwQTXsYLtW5XjgVYa9dkLYG5XBV5tmgYGnrms5ck2NpMMSPHhfZSO4kJOBkOtrB8RUT7ifRmdn3cBabqLDDs1/bWSJdtiCNFkJ8n4baktHkSCpru1eIIIBqT9r+2b2dt4VoiaG3NtUmesJKYLzns8pRPglp3hWs/1QallGorepOVyWwPM8sfEeFKGpdEcOOF4It7gFzH3S+ie1N+QpN02kuzDN4iuFLmn7nILkSW2Oi4svHGjiTj6p0e6eJJXgBauf2utu9W7d6jl6R13pmfYr1D5vQprRQsD9JJ6LQfBQyOR51+jd7UFtbDanJCQH3EsoVnIK1EBI+ZIA9SPOoH7YW0myO823i4m5c6Ba7nGSs2O9cQEmA+MgFOASpriThaTlOM9DipCluxDg2XUJrNQhwy3Qrg2q1gZUklOehBxT80Jv5vdtgtKNI7iXZiKgcoUlz2mL8O6c4kfcBTdu7Itl2nWR96KuVb5TsR0x18Tbim1FPEg+KTjI59DSavCj051K1drobnFuzxte08iAfumEFZU0j8scQfAqz2l/pHt1bU0GNWaMsN6HIFxhS4az5/ZJT/dFSnpX6TDb6QkxtW6O1PaEOfaMGQ3Mb59coV3ZxVBnWQR9kYPStRcZJPNNUmq9mVgly6KIxE82Oc37HH0U9BtRXMGHkOHiAV0lR2quxRq5eb7ZNLuOOfaVfNCx1rOf0nEtrPz4vnWF4fR9apJfctu3DfHzxFky4GPgltxsD4Yrm0uKg+FYvZik5/jUc/2dugbu0tdM0f8AIH7hOm7Rtf8A7sDD6Y+y6YWrYjsV63kuWzQ9st9wlIaL6o9r1LKWsNggFWO+UQMqSM+oqDtPWO2bTdrGRpjTKpcS0wZjCozDspbxbLkBwq95XM5LhGTzxSR9HEytO9WoXRz7rSkg/DMuKP40ta1ccm9s+7pVnhbnxkkA4PK3+eDVW7CuttyrLVUVDpoxDvd7Gck45KYjkhqoI6mOMNO/jRWm3W2rse6dhMCbFaXxt4Clozjly59cg8woEEVU/V+w2utCOFuxzhOaCylLEtKipKPAJebClHx+0gfE1fK1ShHgMOxGUIWhAUlX21ZHj72R+FVA1XqTc7Se5V10xq3crWc2DdFql2BcnU1yTHDQP1kZSWnke8nmUn3vdPQ4wKzsrWVDi+nbIN1uu6RnI548vDkutq4KNlN29VBvgcxxCg68aU3U9nIkaejFKBhIVOCyk9coyOJJ9Rg+fLNNGFYbomPIau1yt0KR3riUNl3C2yoA8JGBxIUo4GPsqAPTNWOmbqawti24yNR6qjNO8SOA6klyWHs8ingkrdZUceC0E4IyBTY1Lp2za2SZd8gt3AFKgl+NGRHnxhjOS23hqSkAH3mw2v8AUVjB0GluZjw14GD0GvpkkKj2+osMmYwzDiMDeOmVIPYX0M5Pduetrk0lMme8iKySnCQlBKcA+WSon4+lNTWWuHNTO6j1s0jhTf5r8llROViBHCmYbXoAhK1+qnifAV92L3ic0xofWu3sqWlbzVtmy7ROYUO7U6GFqUnP2suJAdT5KQ8MA8qZdsfZf0TYLajAD8SJH+HeBAP4rP41ES00wrqieYfE5oH/ABxn9EbTVbW2ynpYRgEnIUu7YX86a3c0fIKk8F1lvaakHHVMsEsZ9EyUNn041+Bq3DrRHTn58ulc/HLnJVr7SIhlQUrVtrcbwehTJSo/xroSyO+QXOfPBAz0B51S9rYt0QScyHfIHT7q0ezyR7qB0buAOi0lxgoc05PwpKuumoNyb7qUwF+RxzFOcsKx08fOvSI4JAJzz+NU1kzmHLVoBaDoo52+2N0vo/cKZuLaEuxLhOhrhSUMnhakhSkKK3EdOIFAPEOaiSTk86lS4rLMdDqT/NuMr6cuS0n9wojNcHunixyourZXbpHAD/NqPXyGa6qqyeteH1DskADJ6DgEjHFHEMRjHguWPalg/kHtganSrl/v9FmHPL+dQ07/ANuo83fjLa1BcEpHupkBxX7SEq/eqpm+kHhC39pdy+oRwJvFqttzbPnwt9yT/aYVUV7vQlF9ycrP+lR2JPLxx9X/ANivqKxSia3U0g5sb9AFmdY3dmkb4lROaKDRU+mCKKKKEIooooQiiiihCKKKKEIooooQiiiihCKKKKEIooooQiiiihCKKKKEIooooQvQPKpM2dtSpVxYeQQkoU7IUT4JQ2cfjioxFS3t5FXadH6h1MqQGlQ4SYzKMYKnHiVBXyDXCR+uKSlPdwlIxk5T27HEJOqe1bZLs82HEQpMy64I6FDThaPyWW/wrpJDKJDIkISMOrU5j4nIqiH0d9k4NQa51sWCr8jWRENheOSX5TmEY9cNLq+9vjBiO22M/VpSg5HLkKwL2l1Jluwi/kaB89VetmY92mL+pQG8cwkc6+lkZ5jFKCGClIJyeR8OlfFsEK4gPLqazoFWMlJFwl26y2+Xe7q8lEK2sOTZRJAwy2grXjPLPCk4qJ9p9/2NzGrfE1NYk6Z1BdoxnQYvHxRrixjiCo6/FaUEFTR94DBGQeSl2pbmIGx9+gofU0/e3oVjYI6hcqQlKv8A3SHarGYkYXRcVZktWp/gVFRHd7t6L3XutOsLx9W82koIVjqADlJNXmx2ujqraZKod55ODzAb/c/RUHaXac2Wuij4txqn32t9Fzu7t+uLdA71dqK1vhAPEuKsDvsAdeHgS4PIBdVjuWomUONvNvZbUyXErB6oGFfxUauzpDW/+6npG56M1RJQ/qaxobEl4N8AnR1p+pmNjHLiwpK0jklYV4KAqlu4WjtIWi6zrBx3iLfPyucsAJENELhy6sDHECcKPI8OB06irnsu52Tbqwd6PgRzadR6c/VMr1ZYb6BdqV4A3cnOmfAeKQdPlE6LHhX+Qowu8VeJzXeFJcQpRLLHpxqPGr9Ug/mipRt1wfuLyJWqp4hRUISlqG0kB4MpwENISRwsoCeScg4HRJyKjq06V1Vfbs5NsD8WWuZiTwobU4WAeSUe6MAhOBjlj0p/2LYDdK/qQ9cHZMOIVZceajhpHqVPOr4UdPtK8ulWC5PppDuOkA8OJ+WFXK7Zu4yO3HtwNMa9Qsl81ba7dGTB0/CIcmLS23DiJUp2SoH3QckqWQVKwScDJxjNSj2ddsrk1eF6/wBYR0OXCQjuozSlBbcVrqEI5YJ5+8rODyx0JVj0ht/s7tWsTtY7k6RTLQngdLV3RdHlKz0cVFDpJzgcICR+rnBqz7FhWGe7tUZLyGzwpEfmSB+p9oH0IB9Kom0FzfRwGCBhAfoXkanwHQdVa9ltmqeheZ5Xh0g+iH0MLY4gOA8J6e8M1THtWOkaktLbq0q97kQCP6djwq5r1tvcZol+yXMAePsbhH34ql/a4WXNQ2p0MuNlsqJCkFP9Mx51E7FYN0a3OdD9lbLsQKc46hSH2+GUo2RsCkjkdToH/wAM/VEmmxwj4VfTt7JP+4Hp5Rzkanj5+cSTVD2vsj4CtY9lI3rFj+t//wBlR9q3f63ToPsvSUisiUpr4MCvi+Dg4VjI6mtOAxwVWzqnHo626MnyRO1trFNntEdwJkJiMGVPeGMlLDOQnPQca1pQCeXFgirebO9oLc3VUtvZD6P3ZCLpKK4jiul7m4nXKQgZHtMySpIZZSMnCeE8zhPMgGAuzH2XdadqXXSdJ6TZTBtcDD16vDicMwWcj3Qfz3FdEpHPqeQBI687R7c6F2f0qdttkYLNh05BPdXnU7iOKTcX204dU0pXJxYIILpy0hWUpSSCEwdfO1h3XHe/RStJGcZaMJobSdjG2RHod67Sutrlu7q9r/SUQ7nIUu0QM8j3URR4VYOfrFJwfBKfG2LXs0SM1DbbaYbZbShqOygBLaQMBIA6D0xUb2LWNhuTz9h2tgG5MMOKEy7JeK4vfHmQZSioyXBn3g3x8JBClIOBT9gsusNJS6UB0/aKUnr6AkkVX5HuDt08FJhrceKzSpLzDJUUtsgjCeNXM/LrmmBqrSd/1M29Gdvt9isvHIEB5EABOMcikF3r48dSSiKftcHvnxWck/5+VYZFuak8QkvvOZ6obOE/cP8AvpMh3EL1rmt4rnx2jewFtnqO2XLV9yu2rol1CFrXeZdzkXFMcAZ43w6HFKaGOfvoCRzKgBUD9kDaLSNw2P3qitaLs+qN1LIXoVviTWGpnBhohlTIPEklSku8KkfaUlAz0z1vl2GKUFMe0DOOSlLA+fnVSdw+yG47rI7k6A0zcdHavDjhXddG32PDE1CiOUiNJa7lQOBkdSRzUrkacQSuDezcUnI0HUBcuNObMX7e/eb+T2h9AP6KgB5pM+OX3n0WVpsBMh95x0JU2ElKlkL4cEhIpR0V2wO0xt1PTp/SO6F0vcZl0xIbE5sT23BngSG0uhSgDgYSD4jFXn3b7O/bf3bsStBR9aXlq03EJbnSLpdLYw0tscihbVvZK3cj9cJx1BzT37Pf0cG0+zzLEjcqda73eC0pyQ9MWnCVlCkgNJWQltCOLjCiFLUoJ95IBSXZqARmTXwSW7j4NFD2ye43bh3f1BZXddTI2nNGxprMyZcbbZ0PLdQg8YUFIUpPuENucyeEFDikKbBB6MM6o0xtbopcy93Fq2WOxxFPPzJb+UtNIGVOOOKOSSSTknKirpzApT0zZdNQYqDp2G23GCQELbQQlYGcK4jjj6nBBPxpsbk2a0a2bG3cjR0C+MOFmZMNzZLkCMltwKaK0f0q+JHEGgfzQVEAjLHfLpAWjAS+O7g6qCbzc+zD23dPXWXZNsYuobm9LTZoN3uNleilx/hJLiJKQhxbbSQVqwoYASDgrTUHbs/Q9KjwPylspui69IbZTx2zUbKVJdWBhRRJaAKB4hKkKxn7VdGbXFt1ohx4cSDFYREZDSEMMJaQhPiEpTySDyOE8unlW+iahZ+1j/Ppy8qcx1z4D+GSkH07ZB3gvzobp7Wa/wBmtZSdDblaYlWO7MJ40tu+82+1zAcacHuuIOD7wJ8R4U0Ft+HjX6Bu032bdE9pfbSdozUTDMe5ttqdst27vLtvlY91YIwSgn7Sc4UPXBrghq3TF80Rqm7aN1NEMW62Sa7Als+CXW1EEjzB6g+Rqy0NcKtuHcVEVVMYTkcEhrQAOlYVpGDyrZWMmsKxyzTt4TcFWk+jiQf92LVi/BOlHB986LSjq4BfbPvobUgEXhCfe6coIHhWt9HE2TutrNwfmaVI/wDjotbN+WGu2xf1utqWlN9WkpHjiKkViV5fnaeu8KcfdX62j/QQ/wDP9FdyzRmU22MXZDr47sBSUthIHpxKJP8AdqMO0JthA3L02qN7MUS4qu/iPx8B5h4fYWlRHMg4ynooZGU54hLkGNcJdvZVDscspCAB3MV1Q+/n++vD9ouaUrU/bJTQTzJdZLY+9WKx6jq5KSoFTCcOB0VwmhjnjMUmoOi59w7jctLSP5N66iojzXUgIWUgxbihJGFoPQkEAlB95Ks8h4a17uUVCEv2i4JjOjCi0tR4CSchSVkkpPoeXkR0qw28kDaG73ybo6+a303EuzXcvSYFxK2UZcbS4hXerbS2CUke+lwYwRxAg1CN57NGpyyqdo7UQuME8RQYcxi7xADzwH2FkpHgApRV++tXoZ4qlgqKhpjJ1OQd3XmOmf8ACsfuWxrWVLnUMgIz8J4hRJflOuXly7sumC9cgYs9WeHhfUD3b5HLrnClHqFZ55zXzT2opUOFaoLyil+3OllaVfmONZIBHxCaXNS7Sa+tsNbM5cJtlZ4SpTDqO7Hnk55AAcv+6kCLZdKyhdX9VXS8N3gsezQ0QFJ7p2c3lDhcUQSQR3R93GQsnwq2QSU1XBgPDgOmv+afZdRbNV9RuQTgNycAk6ZA4Z64U19nrTB3G3St0gR1O2vTC1POPH7LksjCUjzKASfQqTV0twNfae2utcOTdkvTJs50x7XaoaQuXOeSMlDYPJKQMFTivdQDknoDDu1kPTvZ+2f/AJb3iC9GbTFbdbhJH17zzuA3HQOpcW4oDnnxJ5Com1Dq6936+SLhqCb7XqC9EMTXUqPdxWOLlBjD8xhtOeL85xXGVciM5xVUzL5WGolH4EfdaP5iOP14/JT09dFsZbW0sWDIdT+6tXsdug5u9olWpLhbIdtucO4y7XPgxXu8RHcZdUEgKOeIFote9nmeIjlipF7ogAevKqv9kGfHt+vdeaZIKBOiwL6ynORnKo0hXxUvuSfhVrS0kcwTyyenWqdtJSR0dwc2BuGOAcB0BHBW7Z64fxKgjqM5yP8AysTaU8OelZXGQ4ytr3eFSSk8/MVlaRyweXrXtz3U5USfgKr2uVMlc+vpKdOxw/txq9pIDz1sk2h9XmWXA4n8XnPuquOuVv3jSdhu5TlEi29yTnqpHDn+9xVdXt/6dTeNivyn3BXJ0rqNAzj7MaSggE+Q41oHxxVLbAh687SrVxcQsU8hY8Qw5kf7bg+6vpDYSq95scJ/kJb9Vnl5j7OteOuqiGisshosvrZVnKFFJ+XKsR5Gr4oNFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUCigDNCFmiNGTJaYAOVqCfxqWtWuRLRtTDjtngk3eapwgcuJtJCcD0HdJ/tGo50nCemXdsMtd4tPJCB1Ws8kgeuTT03kktnUVv0rEcS+zZIqIaVt9HlAAFwf1ynj/apCX4gOSVZndKu92HNHO2LYSHLkx0od1hf3JylY5uRY44EA+gUh0j/AFg86tK1Gwj3QMHn8jTI2r0f/Iuwad0M37ydKWGNEfJ5D211IW/j9ofjUjIbxjHnivlnaOu/iFzmnzxcceXALTrdD7tSsj8AsIQMBPLOOdeFtE5OB5VtFASPIgc6xLI4MAnnz9KhRqE8yq59sGQY2kdMqWcNtawtbj3kPqpYBP7RTVeLxdGURo7yUK+rdSnOeXCscJ+9XBVmu1jpe4ak26uMa3xu+lMtplxk+KXmVpdSQPEkJWjH69U9jX+Pe7Rwx3kNCeyFMKUeTThAKM/1FhP41p+zUTZ7dH/Q5wPrrn/Oixz2g0jjXxzngRhPzQGq16Z3K0teJLjiWJ0s2GY4DyLMv3UcQ8UofDa/2l+dZ+0fCgav3Vh7faeiI9qwzAmSGx7ynXeF+QnP6jKWEkfpOKHjUU33VjMeyJuiAUqivxJYCeoUh5C/vGK2hq25as3JuVxiOvQ1SWHp8+Xkccf2xzvHUt8jh5SFMsJOPd4SrljItFNb3iQVuMOa1wz6jGfLKa2i5dnbfdpz3A7PoB+6m2FK01tfBfs23zVtdeiq7mffJLXfx25I+0xHaHKU6gK97iIYR0UVkhBZs3Uj1+uCW1MLvk8t4VOvyxPW2nxU20sezR0cuQQ0CBgBR8WzIuNxv81q0W1tthiK0hoKAIYhxxySBjn0z7o95aiSeZUqlCZerNpC1OJbeKW0jvXnF4LjqvAnHLPTCRyHh1JU07B0OkeS53zJ6k/oom73+suLgxhIbwACyLam3/W2nNAMTH5TAlJu1wJPCgtsn6pHCOSU8ZyU+OBnyFvreURGW2kIwEJAGefh+/1zUH9mnQjDSpuv9YQ5Iul0cStLSXAksNAe40cpJBAJJGep58xVg/aNI5BXCu/LxQ+0f3pqlbT1onnbTDURjBPVx4n9PRazspbn22gb2ur3ala8i6T0sFuFcJDXL+ieUn+NU27VL9wVcoq7hKW8QlZSVqCj/PM+NXQM/QmD37F9aAznPckfvqmva7m2Fy5xRZnJCgG1lffDBH1zOOnKldidbqwNbyPJSl3INMT5KVO3y3ns92BzHXU0T8Ykn/CqCtnlXQTt+N/+ThYVgcv5SQD98WXXPxsDArWPZOP/AEIj/wDkf/8AYqi7VHNZ6D7L0Tk08dptqtXb17hWbbPREL2i63l7ugs/zcZoc3H3D4IQnKifTlk4FM/3UklRAA6/Cro7Tbj6X7BW0bmtn48S771bhQ0uQbY5hSbBale80t/HvIWsYXwDmfcHu4UTodZOYmYbxKgaeMSOy7grg6q1VsB9HfsrbNujO4pC2uOVDgpH5T1BJIPGpSsju2yrkpxRASgcCcq6IG02mN5u185A3I3oZkaU24SQuw6Dt6lR2ZjCSO7enKGFOtchwtnhSrGeEJPvVf7FOwGqe2Hu7c9/N7JU692S1TQ44uYcpuU4YUGcdEstgglKeX2EdCa64OzbTpq1F5SmmGWEAg8kpQkcs+g8vLOKqlQ/s/w85dzKnY2l3ePBbVmsltsFujW+JGaYZjtpaaZZQEpSkcglKQAEgdMAYFK7YHJzg4OWcZzSfbnG1x0zSFFbnMcScHHgAD06ePMdOtR/P1nrHca8PaY2rcRb7RDUpu56uebQ80Hc8KmILSgQ86Me84oFpPT6xXEkIxs7Q4+q6eccFJqZsZyUIAkt+0FBX3IWOMpzjOOuM+NbqUobHupGelNzRmjbJoiAqFbA8/JeV3s2fLdL0qW6eq3nVHKj5D7IHJIAAA17zuBaod2OmbRDkX29goK4EIJJYSogcchwkIaSM5wpXER9lKjgFRrcnDUkTjVORXE8vhSjjSPEnCR8B40xt29T6k0bp2NcdL2MTHZNwZhvylR3HWbcwsK4pTjLP1rqEqCElKOfvgkhIUoPhUpcaMXZpQlSU5WGyVBJ8cHGT8gD6VDm63aEtu29nm3rVj9n0jZ2XO7auWpZakrlnx9nhsBbzv7RaPoRXjGZdjivSSBlMe6agfvU9tm7bxwbuw0lKi2i6Sra0Vkcx7NAaS6tIIGA5IcJGemKSZ+4cOyPuTIOvdKWJbaSXJLO2d8kOAD/APSFKHF06n8ardu39MDBt61WzZfbszFoUpBu18AYQ5j85uM2oqCemOJwHHUDpVbr59J72sbu84+1q+1wFuE+6xamlobB6BDbvEgY6cWCo+JNPI4pC3GMJMvbnJK6k6NuW62vI/tult/rZKYKA4lx7QUqM2pJ6YU88gH5HNO6LYd/7elK1a80RduE5Ul2zzIiljPMcSZToTy8eA1xTu3b27XN4UFSd9dQtFIwBFLUYfABpCRisNt7cfa6iOd4zvtqdwePtD6Hk/c4lVHucmcjCO2aOJXZHUG5W/OjEGbf9lI1/gIVlbmldQNPvhIPNQjykMKVyHRJUfI140V2p9p9W3ZGnJV4naZ1A6QlFm1Lb3LbMcP6gdHA7n/k1q9RXJFn6QPtYttd1I3HiTB0Jk2C3uE/MsZNNjXva2303Ksb+ntYamtkqC9gqbb0/AaKSDkKSpDQUhQP5yVA1x/DJnHLmhemsY0aFd+4FyDo4eIHiyOR5GuN30o2jrbpHtWzLhbEJSjVFihXh5KeQ77icYWrHme4yfUmrE/Rm9qLUW5VmuW0+urtInXrTcZuRAlvqLjkmDkNkLWTkrbJbTnxSpOeaSTWH6SrXUXW/aluESE6lxrS1pi2RSh/xqVOPOD5KeKfik0ra45IK0xOXFcWvg3gqsqJzWMjIxXtRGcCvBPWrW5QIVt/o3WgNwNeyiOadONN5/rTGVf9isl07+R21NVJiyUsvN6hmJStSkpA4WMdVcqUfo2oIXP3HuOP5qJbo+f67jqv/l0lxmocztuayVcVP+zJ1NdgruT73IYrB7xj/qS5P6QNWg23/wBlTj+oq7Vnk35u3Npl3tx5HABgSir8AcV5ku957qkhWemTz+PiKVYTuj0W9pI/LBwgdO6H768k6Xc5oYu4/rOMj+FYyx+TktKumMHgqhdoi0O6R3Ls2r0B1q3agjmyTXW18ASvJcYUefJXH3gz5Kx4Ypmx3I2npCFvRY72VDurhGzCmtq5jhEmPwOBQB5BRUg8vdODVod+dBaf1/oe4WRceYh1bRLKy4gltxJ4kLGE9UqAVy8seNVBsdyeVHkWDVbCGbzbD7NOZdGUufouJPilYwoEefKtQsla6st7CCd6PukeH5T+h9FkG3VBPSVYuFOSGu445FSbF3Rck252z65ad1PYh7rkpxhJu8NrJJcX3aQiYhORxqAS9jJwsDlFt20xYtqd17RqOIyxc9KTy3dI7zKg406yghLy0HzLDhV8Uc+Y5bImrsquNpxbkdJHCoq+saV4An9yv49WZrjWT8CHEWUh+1ouftIaQAPZXHARISkYwGnkKUSgcgtJIA4uc/b6cvmc2MYDwQcaA58ORH+aplbto5qmMU1WclurT4hWr7RGqI7+ttP6JjPpeZsNvXf1ls5bVIcUY8NXkeD650eGQjyqE2ksC8OqKllUVo8JKskLVyGf2QsfOmpD1tOvV+uUyY4tT1sgW2yhaznvER2loCx6EDPzzWeHc5PcuuPvpJnuqcbA6hpJKE5+Kg4r5imrLa+hjEA4NAHqdSVD7R1j7hXPmJ7umFM3Ziujiu0BMLeS0NJutuEdOdwjYz9xq8TKS4kE9McsVSTsW238pau1JrERVezulq1xXj0WhkqLhA8i4pPP9SrssOEKwEjHMkEdKz7bAgXARj8jWg+fH9cLXtj6d1LaImu8T81t8CU5IJ/yaxqbJHFwZx5msiQSPc5g17UhZygAVVcAqyknkon3s0YzrHQ2ttFrjh5zUWnn1x0eUyOO8ZP9sNfIGuWO0zwlLvulJKygT4SlJQT1cTyHzBIP7NdhdUpEUwr0AlQhSUh8f8i59WvPw4gflXJTc/TzmzfaVvljTyatt6c7k4wFRn/fbPzadQa2T2X1u/HPRE9HD7H9FU9pIMPZMPJRVqWN7NdFHGA6kL6eOOf4ikg8+dP7da0ew3mQppCUtJeLjfCOXdue+kD4ZKT6imDitijdvNyqi4YKKKKK7XKKKKKEIooooQiiiihCKKKKEIooooQiiiihCKKKKEIooooQiiiihCK+p618r02kqWEpGSTgDzoQpH2ftra7s3cH461tRguY6ocghDY90k+A4ygfEgU9OzLptW6/adsDhhIkw49yN4ktOfzfs0b6wJXn833EJP8AWpAti2dK7YXic1I4H5xRbkJHIqSOa+fjxFQVj/khVjPo8tCRmNM643OlMOmfMLGmLQvmEgukLkqHqP8ARxny4/Wq1tJcBb7ZPU51xgeZ0UlQQe8VEcY4EglXb0spUqNIuzifeuMlyQDjqnPCkf2Up++nCAQPd5E/hWrCitxIrMZkAIaQltOPJIAraCTnKuXPzr5ded5xytNwBoOS8qKUniLmMDqajeVv/tFE3CuO3EzVrVtvNqdaYfE5IYaW8tpDnC26TwKI4+EgkHIPIgA1I74T3RChhJBBx41z4us72vcDdaQUof8AaNZXJDjD7QeaebbWEpSptQKTyHIkZBAweVWPZ61U9zbMajIDQMY6k4Vb2kvTrHCycNyCcK8esbU3dLY6wVEKIC2yOoV1BBPTwNc797tu5u3l/m3a3xB+QZbqn5LDKOEwHVKJUpKAMdwokHI5IJI5DGZL283Uv2i34qtH99LtYPDJ00/IKmHmzj3oTrhKmHepDaiptRzjhJAMj7m64221fpzS99h3L2aFqee5EcmOR/rbehsJ9pDjRPJ5PeITwHI5nqMZs1opquw1Qaz8SJ5weWPPoRxB4fVM6eSi23ibTRaScgeKo1Pms3NEexMKSszZrYyk5BaJ4ir4DB+6t3TN5aW3cpTKu6Xcpa5DpxkNsgkNgDxIBVgdDkdAMhU3L0zoTQ+5ctvQN6cvVuuCJFvt0p9KWAw64AjvSlPIBIcVgAYzg+lOIdni0WS56f01B17B1RIvsQSg3b3loTGdKSpbDiRhRWAk5wevLFafNJTspu8TunB4anr+ir8+wdxjMsAAwzGTy11+yQGdYFtabbbGF9464e5iMkuvvuHqTjmpRAxnkAAMYA5PjbjajUWqLvDv2rAWksuB+PGDmUMkYIcWRyU55Ack9eZp9WDS22e20edCh2o365Qvcnxra83Hix3SBhqVPc4kJP8AySO8XkEcI54xXHd3XNsm6ZnWi7We0MTNQRoZtllgEtqjp4lOZkvkvLV7qEqKUoTgkc/Ct1FRNM1wpG7mQe874jpnQcshNaCgtdqqWtndvyE4A6FT/bLXGslrajMBLbSEgJyceH/2+NK9vs91uaP97rbLlkdAwypf7hWCPd70xFQGJimPcScsIS2c+PNIH+fOm1uLvE9oKyMS7vMutzlz3VR7XbGZKy9PeAHuJ4shKRxJKl4PCCDjJAOUxU09bN2UIy5x/wAz4LT5Z2U0ZllO6AlbWhlaMsL951ZHZssBPuqkXIJaTxkHhSOPmVHBwACT4CqNb7a2Gs3WLtZ7U7Hs382zPmNiOqcorbJ7hscylJCfe8jzA5U/NQ3G4arSvXe696TcW7UPaRDR70GAgqASzHaUSHnlqIQFryPE5SniqvmvNTXnV16N7vTvvlSW2WQT3cZoH3W0eg8zzJyTnNaxsnYIqF3avO88cSOAPQdfMrOaval10k7KnGGZxnqr5dvxCXOy/Y5GD/8AiO2n74sz/CudiD7o/wA+FdG+3Uwt7siWqSeYZvtodPwMeUP3qFc40fZGTzFT/ss0sjgP/wBx/wD9l5tP/wC6BPQfZbFsuk2z3eJdILEVx2G6l5CZLQdaUtPNPEhXJYzg8JBBxggjIK7pLTWuN8d07bpaI/JvOpNX3JLa5MhRcccccVlTrijnkBxKUfJJ8qbSjgV0N+io2bW2nUvaRvduSpERK7Jp1x1WAp7hzIcz+ahKShJV5Kd/RNXqveyFheTqomlaZDu8ld/TEbbbs7bVw9FJukW1aQ27tiHL1OJCEuPqIXwrJ58Ti3O8Kc8R7xtPReKbOwWr7z2qbpJ3ivVrkW7b+FOUzo+0yCQq4qZUUm5SU/ZVhYw03kpQpKlc1gLFK97twtS9uLfCzdnza2X3e31nuqETrit3hbuUoqUX5ivE8g73aBk8KVLxjPB1O0RpGz6I01atLafiIi2+2Q2ocVlHRpltISkD5AE+fM5POq3MCxuXDvH6KXjIc7T4Vm1NY5epYxsftjsW2PDE11lwodeb5fUII+yCM8Ss5A5DmcjetLVvs0SLYLJBjw4sVtLEeNHbDbTLSRhKUJTgJSBgAADA5fBqbobtaP2p01K1Jq66JhQIpS2ogFbj7yjhthpA5rcWrACRknnyxmt/b9vUMy2fyr1nCVbrjcMOMWglKjbmT9lpZSSlb2MFahyBylPIZLduXM3m8F2cNOCni80tyKuM3JWypaSO9QAVI/WGfH76SrNZLRpiELLp6EhkLWp10A5U66o5W64o81rUeZUoknzrU1TrC2aSs8q83mazFjRWlPOuuqwltA8SfjyGOZOAMnlWexSZMqzpufs7rMqajiSl8AKbB6cSegPiU/LqCaO1JGAjcxqUyt4Ny9TaeYXpHazTB1PrqayDDiKX3USGFFSRJmvdGmhwrwACtZTwpHVSaCdpTb/bHZ+ON2+1xrWdutuFckEWqwofXDtyXE5JYZabIWiM2T7yyU56BBUav5ra+6X2U0ZqHXF3c7mJDYcut3mq9558gYHEoglSj7jaf2UjAAA5sTdqbvvTE1H22u10mZbtIRoiptg0s0tTDkmIDiKzxn3mmVKWhIUBxuFZVyHMr0x3j0H38FxIN0df0VH9xNxLluLckTJVttVqiRQUQrXaoaIsOE2rBKW20jxwMqUSpWMlRNM0jB504dW3tWpb/Mva4EGAH1e5EhMJZYjtAAIaQkfmpSAATlRxlRJ5lAWDnl8qm245JkcoZaLjgHh40ooaAGByA6VrwUZWonwpQCedPqdmW5TOd+DhY+69aC0R61mCPSvQSCeeKdCMJuXqUOzfv9P7Omrb3ray24TbvIsEm2W1DgBZRIdcbw44D1SkJKuHxIA6E1HFyu11v92m369znZlwuUhcqU+6riW66tRUpRPmSSfnWtwJPPxNegOHpXDKdrJDIeK6fMXMDF5IOa8mvZPpWNRNLHwSTVeX6NNkGw7oSeX/AAyxtA+hbnH+Apu6Ut1yu3bM1uLQGVvjVF9WA6kKBCXADyIIPXypf+jlkex6K3CWU/8ACL1aUA+fAxL/APrqBNS3kyN9de3NToQFanvTqTnGOKUACD4cvEVhlXEa7aG6MYcfhsb145WgUhEFDTPd/MV06tendUKiIUu0wnBgZLLTOD/Zx+6tt7TF+4VLTZpeMcRKWDgY6nI8Kovtxuhvnqe5HTu2cy+3p+IEreCj7SzFbHVTrrp4Gk4/TWnHhmlHWOop8xxX+7HuzOv6Qo405pGSlLHHn7MicpPdkDoe6Qs9RxDrVFi2RqI5+zqZWjnoDvfLl6qRrNpaalZk8fNWyVJi3FoqiTI0poOrZUtl5LiQ4nHEglORxDlkdR5VW7tF7Nm5S29QWHMO4x0q7uQjJ4Qo5Lbice82Tk+aSriGeYOr2QdZJtm4evbBYYf5Ltk0W+exDRIcdQ2ktKBBLhJVzPMn5dOTj7ROutw7PuZpK12W/tQrbd4s0OsPW9uUy4+0jvBlCsHnyBKVBWCetSFHbJrXefdqaTg3eyem7vYIHyXE9wpq+1e9VLPwzx+fFVecuty025+TdX272CQ5xJbUvJjSAOR4FfZUOYGOo6HrTa1HIjSoUuORiPKRhJKieFXVJyeo4sc/iPEE2ZuN/wBK6pssmFrbT8eCCB7ZMtIXMtwGCO8eiuD2iPgjBXwuJSfzkeMa7gbD6XsVlg3u2TV/kae4gmXbHy60Y6hnvAOIpWCMjkR168sVd6K4QslAnaWOJ4jVp8jy8lSKfZenuMoktcwLSeB5KHGL89ES3MQ2pX5SgNJx5vsEtFPqT7p/ap0aUt181ncW9M6bB71ZQ1LmdURUYxy/SWQk4A59T0zWtrC16D1DqexWHR1tkWC1tJK5a1vl8KcHJx1vi5jKUo93OAQelW12ZjbTaK1zbNC6e1HGv8GXbXZbM9URMZUZ9ttTjyVkEhQLbaz3hAOUjPLo/vVYKamM8Ee9IRkA/c9cdFYqb2ez9o+eocDHGTkg6Ej/AMqb9jtuLbtvpCDbmmEMFlhKCAMcIAyQT+lnJUfEkmlXc3fzaTaFDidaaqaZmxx3irbFT384jGRhlJ93OOqykc+uOsC7lb7XHWSe40TdZVk0ggrbFzjgIn3kp4kH2VSgfZ44UCnv8FayDwAY5QnqNcBGhdRx7ZBYgQZUSQVsxwrifUEE8TzqiXHTkcuNSufF6VmlHs2KuoE10cS551aOOp5nr4fNJXPbOloHilpBk8M8l0otlyhXOBHuVudUuLMabfZWoYK21pC0nGeXuqH7s8q2gvOQk888xTA2PuKrns9oiW5zcc07bCpR659lbz+I9KfJVgeRz99UeriEFRJE38pI+RwrzA4yRNeeYWG6w03CBIhKT7r7amyfEEg4PxBrnL9InpJf5e0dunGgIaRe7YbbNdb5BUuKrHP14VJSD5IFdI1D3fjnA9fCqwdsTbeNrPY/WUJtLpnaXlM6qtiWxniaPuSEegCS8r4oTVn2GuQt95jLj3XHdProo+80/vFE7HEaqg2qm0al0ZaNQNoUpciMqHIVjOH2+Y+asLI+B8qiRQ8flUr6Alpumir5pp54pXE4Z7CcdeHmSPhjH/OGo1vUb2S4vNgYSVcafgedfScOAS1Z7JrgrRooopdJIooooQiiiihCKKKKEIooooQiiiihCKKKKEIooooQiiiihCKKKKEIpR0/GEi6NBacoRlavl0/HFJ1Pnauz/lO+MIW0VIU4FLwOjbaStf4Jrl53Wkr1oycJZ3adTbIFh0c0WwYbBkyijoX15Ks+fCStIPkBXR/s2aDc0TtTt7o15Cg4m2HU1zbxgpkSfeZQr1SheD6g+Vc79Aabd3o7RNl0wyz3rFzvTbCkE/+itqy4fhwJUfnXWywSEXW63e+juw3JlKYaCAAjumcoTw46J4uM/MVkntMuBhpoaEcXd4+nD6q2bNQb8r5zwGgS0nGQOEYHlXpSQFYHQ+dfVAFPLAPpXttHeHjUeQ86xTGuquBWpIQFoWhXEBw+VUB1vahpvercCyrWgpnXCPemiOhRKZStX9lYKT6g+RroHICeAjHUY+VUr7WGlnNNawgbnRYwMRhBgXZKSSpMZa+Jt7HThbdPPx+sHhVt2RmDap9Of8A5G4HmCCB64wqntpb3V1sducWnIUbGIzZLiXFoxHmguNn80OD7Y/EK/a9KTtWNm93JFrS48pnU8hLTiErAKLpgpZfSo+LqeJpzI5ngKvDG7c58a/Wxy3tPtB3IdaVjkh0c0q+B6E/oqNMTVE+Q9p6alzvWJkZpTqAThxp5o8Q5jxBSRmtDoI5HShx48D+hWP2Oslt9YyoacYOqQrMWdvtxFTtaWlm8N2x5CGIk1kBpaglKyhSRgA4UP7WetPKzzrVqLcJzVEJarLGWw67IiW1wNrCHDgMtrSfqAoEpWpPvBHIDKjWtedT2bWeqntVP2pi4OX1ic/BiOgpS5KfDEVskeSFJWvHkis1vhxdPtt2fT8VpyQ4lLTa0o4eIJGFOr9OqlfHH6NTVXM8sG8D2mMeHjotFue200NofbYxl0jjh35gAeJPPwToefN1mtWqIhEW3QcnuY6QhmMknmhtI5cSiMZzn3SpROOenOSb9uro/T0MIbatyXZakDIS2DhKR49Dj/Csr9wtVgsiiH0cLKCt59Z95xQ+0s/uA8AEpHSlLs02SZqrU0/Xd1ZU2JCgGEqSMoZSMIHnzGVH1INQD5DBTy1buDQQM83OGPsqXsrQyXC5skOSGHUnmVaCddNPWLTcq9X+9JixLfGXJkOpZUrhQhJUccXDkkDAHXJGOdVLk3yVrm/ytdXZtbcu6NBuIw4sLEC3nJaYTgD3lBXEtWBkq5Y55k3tR35i32DT2h1PgK1VeWIzyAeZiNKQp74ZUpoffUbR0LW8uStCQpZ4yEjABPPA8gOXKoexUvulF71+aXOP+I6eZ+ysPtDu0kbmUDDgYyfFNjeKY/bLLYtP8KR+UlPXZxWOa20KLLXPxHGJBPqE1CV5USjJ/NIP41Me9KhMuGl5GFBDNret5JHLvESVvHH7MhP3Goh1EyUsLUlPgP31o9jwIIx1GfqqxaThkf8AnNdG+2BGXcuxGZQAIYbsM4/BXAnP/va5oo6Zrp/vW81fewDdHRgqd0fYJKc88FL0Ek/dmuYCefQePL76a+ytx/hk0R4tlf8AdXXacf6hp8B9luWiyXLU93t+mrLDcl3G7ymoENhsZU684sIQkepUQPnXSPtYaxkbM7U6C7B+xajJ1feIUaFdlW5JC+F84WniHRUhwrJzzDeSeSxVb+wfpeLb9dXnfe+Wt6bA27iBdtjoUlBnXqQFIiR0qPIKGFuFR5ICOJWAKsFsLfNHaNh7gdtzdGZEuaoM2XDttxQCXLnclpSHkxUnowniRHZ5Dl3y1cscNyuEu/JgcB91HUrN1nmp67I3Zit+yU1qKw9HU5Z7WluS6F8Ts25SCovyyD9htKUezteHCHiOalZs7edX2rR+lrnq3Udxbg22BHdlPvLOEsx0AkqPieXPHPmceIqunYqvl61lsq9uvqeWV3/c2/zbkeeQywlwxWmUZ6IbajLI+JPUmqnfSmdpybcb412eNJXUpg2/u5Wou4XkOvEBTMZWPBAIWpPTiKAeaagmxvnqNzOccVJFzY4s4wpJ7Pmpr324+05N3hvsOTG2z2vcH5AtbmCl64LJDbrg6KcwlTh5HgJbSCRzN/Y2oUv3aXBWwoM21pBdc5BtTqxxBtIPThRgnw+sR61Wf6OPQETQPZp0uZMRMa4akSu/zVqzlSXlEMcR647pLR++oi7dPaLl6c0JA2j2xlynNXbuy1yltNkpejW6S4G2EeYW8gNoHjwpcz1TSr2drLuR8AuQd1uTxKsBt5rZfad3KnaitqQduNF3BUW3P5V/v5dmSC5JT4FhgFKWz+c4pS85QjFkOJLKOLkEoHh0GB4enhUWdnjbK3bP7YWHQNtSjhtEBqO6tJz3jvVxzP6znGr504d1tUtaX0Y88FKEy5yotlgBAyfapj6I7Rx5JU6FH9VKqavIee4lMFo3nJr7ibesbz3ezae1Oyt3S0CUm+XGIVfV3F5CsRo7o/OaQQXVI/OU2z4ZqqX0rOrbsvRWiNotOR3JMvVt9CxFjj6x1MdsJZaA5ciuQ2cYxlA8ud6LKWm4yltJ90qKU56lI90Z9cCq4bh7XK3J7Ymir1eWG5Fp0PZJl57twZCpjz6GGAB/VQpz4sjzruCUBwJ4NyUSM0IHNcx+1jsVC2ARoHbNHcyLy5Y1Xu/S2+jst91SQ0g+LbaGUpB8SVK5cWKrzIhKZSXFAjAyBV+/paLSLVvTpW+8i1O0yphOPBxmS4FD+8k+nFUEdtXZ9/Z/c63WwMqTbLrpu1zIjnAUpWUxkMOjPQkOMqJ/rA+NTUMpcxpPNMHN3SVXuCsd/wAJ6LHL40pJTmkr2SUwUqW0tCkpDiQpJBKCMgj0IOfhSs24HEhY5A/v8amaR4c3cUbVMIOV9wR5UHJNBPryoFPeCaL5XwnnX0kZ5c68gKWvhSniJOAnzPlQThetGSvmc16ZZfkupYitLddWcJQhJUonyAHjTw0ZtZq7W18i2Cz2G43G4SlJ4LdBZ7ySUH89Yzwso/WcKQOuMZNWLsexe3O2C2HtyZbWo7owhK16asMsiEHQOaZ88e88rP2m4/CgFOMjxr122hpLW3D3Zd0HEpZwZTt353Bo8Uj9kLWWutLad1Do7RW3cm/3i5XJuauQ86lm3wG20FPHLfOEoAKlHBUnPMZr47t/tXpHUFw1Nrm6DcHUs2U9NdgWhxcaytSXVFayt/8AnZCc8uFvgSQMBRFOvWOvLxqmG3Z22YVmscfCYtktLAiwY6ccgGk44zjqpZKs5yfCmG/GHPIwRn7qyd9R29XNVRDszL8WD3jjhl3IeDceZTCs2vkkibT03wt4E/slPUG42pr7bPyGl2NabKn7FmtLAiQEeWGkfaPqviV5mmgtRcVlQ5cq3HG1BRQEnn6VqRnHblcm7Hp+2Tr3dnVBCLfa2VSHsnzCeSB55xTiCEMadwYHP9yf1Kr7TU3GTTL3Hon72Y7gi3b2XRji4RLsbCjn/k30IH7zUsdrS0XBiHpbXUYIcY0/dmlyCDzRHey04ojy4lND51F20GhdZ6F7QzNr11aY1tlztMe1IitykPKaZM0ABxSPdC8pVkAnHLPPNWv3p0pbdXbcXOwKRxNyo6m8kElGRjiGOpHUeoFVy8VkVLe4KgOBaWtyRqCDlp156LabRb3zWA0UwwcEeqppdfbGrgm6wXXGS0viWW1EKQfBYI6ev31hkarvml7e/FscttiJOV3E2K4gKjjvCAZLSDyaeGDkj3Vg8xkZOlp+8XCEw9YdRI4LtaleyTEK94LI+y4D+chaQFAjkefUGtO4uRXA5BksByLJz3YUc8uvDnzHL1KceuLFFEY39m4ZA9cjkVlVsq6iw17XAEdmRlvUDj80m7wXHQ0iyWSdouxqt1zgsobfe9pS+qYfe4l8gPd4eHkc45jODX3T+mLpb78zarvIfYakRlSLj3SA2VQirgLSAMYLyxwE9eFKuo5FqXKPAabiafbhttzRKkuMSBkKejPsKHdn0Q42cei6kiNqe1all3rVzXJuU9HjRwrKQlmPGbR0PMcThWT65qcqHOgpQGajXXzIx+pWk7VbZ+/RPkoxuNe1vdB5niT9k6NQXB6+uMQIxAkPpShtKByZZQjkEj81CEJAA6ABI8qQteFETRVxgw0njeYESOlPX3yEJA8SSVfM8/OvlinO2tty5TE4dmpBTxAZQxkFPw4iQsjy4K39A2mTuhuraLTGbDlvsshu4zlE+6XgfqGeXU8RLhT5IPlVdYwUxMknwR94nqRr99PNZfaaCa4V0cLeZySugO2Nk/kvoey6bKgpVqgR4OQcglltLZ+R4c/OnUOLjGMY8q0rHBTAgMx0jHC2Egcv8k0oJJCQQAOXxrFp5DNK6V3FxJ+q+jmNDGho4DRfeFRTjBAPl5fGmvqa3xHrpATcBmBemn7FcSU5SGX0e7nw5KH96naEpCEqPL1pH1dBcn2GayzzdSgOtHHMOIIWnH7SRXsDzDKHt4hD2h7S081x0g2ubtZvZcNI3JoKdtd0k2aShX2XOFwo+4lKT93nTb3LsRs13Wwk8SGVltK/02z7zaj6lBBPrVhfpBtKC07x2vcm1MpYjaztEW4pU3yAkNIDa/nhKCT4nJqIdyFNX+zQNRMM8KbnCDhPgHW8cQ+QIHyr6rtFcLhSQVbfzNGfP/ysxqoTBI+I8QVE1FfTnHTxr5U0mSKKKKEIooooQiiiihCKKKKEIooooQiiiihCKKKKEIooooQiiiihC++VS7ottnT+hr1fu97t5mII6R48ThClHPnlLY+BI8aiq3R1TJzEdI+2sA+g8T92ak/XThsm3lnsiWu5duzyp73mW84QD6cIbX+3SMuuGpWPQFymz6PDRzMnVWr9050dbh0tZjFtxHRM+YVNpUT44aS8Pi6g+VdEtN2wWm0RLaFZ9nZS2o4zxKA95XzIJ+dVl7EejZmm+z/p32hlLa9b32ReVpTjiMSMe7TxeOFONDHp8atQ3xJKUJPMDr5185bfV5rrxIM91ndHpx+q0CwwdhRDPF2q2SOFBUa8N8WeEeNesFQwoYBHM5r6ClBynOAfHwql8lMrG4jhAUR7uM4BqON19IM6rsj0RyOh7KFAocGUqScgpPoc4NNXtC7y7gbeay0dYdJsWD2XUAnNvLuzLykBxltDqRxs5WARxgDCvh4hBsXaekSy5C1dt2+pTaglUzTc5u5IIPIq9nV3b4A8gFEeVWKmsde2JlbT4ORkYIzoccPRQ0t6omTGlqHhruhVRNwtB6g2vmKTbWJUyzNqURg8ci3jmShaRzW0B0WB05HGBTPcvSLs0t0vodRIR3aiFcgeHAV8wADnn7o9avrqIbbbgaTvOstMXeLdjZor0hxuOtTchpbaFK4FJPC42okY94DB6Zqle6tl2sumkLdctLSrkrW8oKN1ZYaQ1EUslIQhlDYBJOeZOSTg5Oa1LZ65G5Ds6phbI3AJ6nxH6qu1WwZr9+roHABupHXmmGxP/J0/Tj4cKfYbI4oEeBced5/H3/3U5bfrOLBiOrmvobdfSniUvwSOYbQPjnOOpI8AKbEfR1/uV+slvfbfh8dpjuj6vKltqUeEpB5cyfHyqYbXsNY9ORWdQ36Q4XHne6YbwZEuU+oZS0wjA4ln5AdSUgZqeuVTRx7rJjlx5Dnqqo7Zqokz7z3N3jniOabGmtHak3LusZLkSQ1aS8CmNwkOSSOeVJHRAHh1q3u2WnrLabTwWeXDkcBUy8qM8h0IdT1SeEnBGRy6jNQKuTJMQaelOrYbeJ7+022SpLYaGRwzJSMLfJyMtNlLQPD7ylZAlDssNQYuntWwYbTbDEfUsxLLSOSG2+L3UpHkAKo+0znS0TpC7AaRho4YJxknmVZtl6u3sqjQUQyQMkppb42KVqrc6ZIjMKku6P0aq+ttDnhaZqCsgefdhXyTTZgKiPFCorwfYWAWlgfbbP2T8xUq6Zuyx2ntVKQhDgZ0xEQlKwFJUn2hJUhQ8UnOCPEEimPfdv3NuNXp0z3LqLBdeOTpqS4pSkOsZyuEVn+lYJA81IKVYxXFPUtdG2jPxMja4DqCMn1B18lB7d2eWaEXKPXBcD+iRtW6Dc1tply1w5AYuDSxKt61/YD6RjgUfALBKM9AeEnABNVr1IHIbL8OdGXHlRlFl9hxJC2nEnBSoHx5VcRBEZkkdSkHoKgbdy3xtSSHHbuktzo6Cw1ObQC4pCeSUOjl3gA5BRPEkYAJCQmrFs3cSH9hL8Gcg8wf2VAsdd3uyk4Z0/ZW4RcP5U9gN3uzx8WiUxiPWIEgn744rnHZYVyvVxh2Wz296dcJzqI8aMynjW86ogJQB4kmuhHZlalat7I69PNHjItt5gFI8VEvqSPmHR99U/291NoLa7TzuovYzqLWtwaVHiNLJbg2iMrKVlwjCnX1jlwo4UpQr7aieFLv2ePdT1Vyo28WzHHkdVrO0TGyNp5jplg+ikbVzGrYekLD2e9t5Uf2KRLRBcdjkF3Ud9lOJRJWwBhSo7fAlgPYKSG1cJw4oVrdsXcO1mZY+zdoGchejdro/wCTlvsjhFyuo/4VJVjkocfEE58Ss/nYEdaU3gu9l11N3Rud2WvUsOE+qyrDQCGJak900ptCQENIZSsrbQkBKS2kACoyWqQ9xKcUpa1klRUSSonmST4551o7ab8TL+X3VeMw3e6uyW0G49h2L7FukdYXRADemdDG5hKSB3kqQo90gZ5FTjj/AAj+sT4GuQGob7eteaouGo7tLVJut8nrkvvLJJW667xFRP8AWV08KtP2lNzX5XZL2T0Cw8pP5bs0adJKF4C2oZcjpCh4jvCv9puqsaaDMW/Wt5w4Q3NjqVnoE94nOflTOgpy1j5SOqcVEwLmsC722mVA282zjMTZCIVusNpTHW4eSWo0OOEFQHgAEZ+Vcv8As+aql9qTt72/X17acZjsy3rnBhuK40w48VoiKz+wQ1nwJyatx9IRugjRnZdmsW98JmanuC7EzwqwoIU+6p5Y/wCbZKT/AKwedVT+igsjc3fe/wB6eKgq16fKkc/znJTKf3cVRNKwsgkndzTyZwc9sa6+2zuWm3Wm0nk4UDzwAAP3VEW8+og9u/s7oRxQ7m5Xi63ZxJ/OVCtMt1rl44cKD6EA1IumbrGlm7oSvhUxcShXPkB3LKh+8moJ7QEuyW7dDSG6961pBsVu21TcJlwD8dbynWrhEMFpKEpUkZUsuYGeqOQOcVFGripW9rMdP80Tssc87rVYeEDFtjKc/wBHkgeJA8/iKR4bEdeu7hdkt8Tr8GG0hfXLSS8pP4rUfkKrfd+3HtRG0x3Fv1sLvKxiNDskBw3F8FQTwlpRKWOHnlRUTjmAnoYq2Y7ZOpJ+s27JqfSLdtcusV+TbzEusiUFGOkumLIRIUUhSmwSFt9FfE1FfxcDJjYS0DU4x54B1OOeEb0BkbE6Qbx4DPFO76VDZx3XGy8Pce0wnHbloict9a0cyYEghL4x6LSwoeQQrzre7QexznbF7J+mb/p1lDurbTaI1/siCcd930ZtT8Xix1WU4GcDjQjOASRbS6QrPrbSsm0XGEi4Wy+Q1x347ycpdZeQQUqB8wflkkeFRDsY2rY3R0jbHUEl6UnR63WIRUttDjlqUsriyVqcUhtI4VLaJKgONheKnxWiKBshd8J+hSRiL3nA4rlHbtr7jups5c5VhtSxrvajiYuttDXDLm2ZTiyHS2eanIzpcQsAZ7paB0Ryr8iSlp48uFC+qf0D/h5V2O3O2ked3PZ7UvZpjso1kykL1FphbraWNSQFHhcKFIUprvsAEKCiFFKTyX9qAu1D2CLfr+xq387L1vLybgtx68aSUnu3mHs/XBhs4KVpVxBbCsFP5vLlU1R3KKYCaF2WpjNTO+FwXP5KwRkEEHofSjHrzrXkRZ1rnO2yTEkMyWXCy9FebKHW3AcFJSeYOfDr0BqVdn9itX7nXJ1iBakqTCAcnPy1mPAtreRlyW/jDY58m0ZcUQQADjMzJcYIIu1mdgBRZpiDpwTDttknXRIeRwsxyvuu/d4uDix9kBIKlK5jCUgqORyqy23vZfZsNvZ1Vu7OnaSt7oColvbQhWobkgHI7tlR4IDZ/wCMcPGQfdBBwZC07D292ZltObepRqXVcdnularnxQliGrAym2w1ZS0P+WXlwnJ5EklMnzJdzkO3CfMflSZCu8defcK3FqPUqUeZPqazq9bZvmJipNG/zKBuN+hoO5TDed15BLcvV0e2WBOjdurBE0fpoY72HAUpUicSOa5kpX1slXM/aISPBIwMNKQQtPMYPTBHh4VsHPCRwZwen8KTlXRMq8s6XsdsuF/1BKwI9ntTJflOHplQTybSM81Kxgc+fSqZG2eslLsFzuf91Uy+su02dXuK1pDKlqCUAknkABzNIrBm326/yb0ZY7hqe+L+zbrSyXlI544nVjKWwDjJPPpy6kWQ267F2qdVRRdd9tQrsMJ0hSdLafeCpDiRnlLm8wkH/i2wR4kpIqzWltD6R250+nS2g9NW+wWlvqxDawXlfpPOHK3lc/tLUT64qEu+19o2eBje8TS/yt4A/wBTv0HzWlbPezmqrSJK87renNVH0V2L71eY7c3evUv5NbWrjVpzTzwK+Age7JmnIz5paBA8wasBpbQmjtvLUiw6H0xb7LASnhKYjWFvY6KdcP1jqvVSj6Yp8TI+VHB54A+6kp8AkNhOVE4AHifKssum2Fyv3cldux8mN0H9/VbRadm7fZmBsDBkc+aq/rEOvdsS0s8am0K0SBxAZx/vi7zPL0qx11TCs9iduV+lsRrewwXn5Dy8NJbAyok8wBy55qseorlHldscsB0PJh6QRFdSw6CWlmcV8J64PvA4PnUndpibJh9nzVfDJWUvW5xskjBwpBHMc/3mtAmpjKbfTO03o4x5ZKj2Shkc0rdcOcVCe+21cfU81jWOh5UVT7iFCLLZdSY0xvJ+qdWnkRnosH3T15dK7XK5y4bqrTfYz9tmoIUY8jCVZzkLQfsrT5EZz65xU0WCc1pJTZ05cPyGmbGaLraEqdtshwIQcyImOEHH9K1wuDKiOLmCrfkbTm7bDunr7bolvvKG1Sza5C++YdQOXtER3kSgnqU4WkghYIGTfqSf+HM7OTMkTeB4OaPLmPt4LOX/AMN2sd2kJ7OcaY6qt0m6NzdR6becSlDrVxQyvl+atQGfhgE/M+RrUtMt2JbY0TvsMIcdce545d4Rw/PGM+Az5Us7k7XzNC6rti4aJCGGbgyhLTjpWkArGOEnnjOc5J+VGzmltubnqCS/u7KvLFm43ShuEEpWpWFcCgF+6pPGMHzANXCKSllpGyxuyz+5/dc27ZKruNaLbERvYzrw0SjYJGoNwrim2WJjjLq+Fcwo/wBHjjPhj7a+fJCc8z5ZxePs1bHRtuLa1MfbWp9ZLilOnLi1q5qcX4cRIHIdAAOeMlqbAWnb2DoOVru8Ow7fHtEt2Mqa+pLLPdoSlaCBnCDwrSngGBkHGMmnXdu1bp22Q1PaF0TdNRMcCQ1cH1t2y3Kz0KXX8Kd/YbI/WNZjfqutuznUNDGRGDg+J8ScD0VuorPQ7Jl3bvG+M6/fCsWgK7sJSFcvXnyNehySMBRAAxj51VDSXaa3d1Luvo/RkiFpBq139ySuQ1bjIkvNtNN5IU44EoySoAKQCOvMdKtYo5VnOSMfndfI1Rbpap7SWNmxlwyMHPPH6Kdt9wguTDJTnLRzWZHTkrJ8OvSvqyFDBTgqHimsHGrzIOcgV4CionClZxk58KiSn+FTbt37fM3XY83ZpLvtu31+9wJ+z+TZuOR/quFkDwxx1TXTC0X7bKVCddy7ZZgcSD17teRj4HiX/ZFdT90dLL1hbdSaCYaS5/LPTkyGyHMEGYyCpn7wo/2R5Vyc2peKNSTtOS0kNXWK5FWknHCscwr4j3q372c15qrW+B3GN2nkdf3VHv8AB2VWJBwcEwJscw5bsZXPu1EA+Y8DWvS7rCI9FupS+yWncFDqP0XEkpUPkRSFWmtORlVsjBRRRRXq8RRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFAooAzQhL+jYSpl1CU8lEBCPRSiAP306t01Pz9axtORnO/9jQ1DbCTyKjgAD9ngHyo2atcaRfGJc5JDLC1yicdSykqQPgXeAfOnT2ddNr3J7UGloBYMlhy+Ge+PtYjx+J5efMcDRphVzinjknPBrSfkE4iZv7rBzK6gaF0xbtLNWfSltwYuk9OwLMhWPtOhtKnV48CopSo+qzT0LZ4gQAkqHLl0pB0opcpqZeXR71ymvSEjyQFFCB8koGKcHeAYUsFWPEcq+T6yc1E7pXnUn7rUooxEwRjkva1cBwnGFYNBV0WORyevhWM5JwASPxA/ya8OOqCcI+ApslcBVn7b9quC9H2bWEDhJ0ldo9wex9r2dau5cPwy4jI8s1Ahg+4qcyA64v6xaAnPepPPjTjkVcsnz6/aHvXY3U05G1XpS42icx3zEqO4y6gpzltSSlQx8FHHkcHwqhVplXjSs+ZobVBUidZFBpLhTykRsEsPp8SFI4Tn7+daVs3P75buyHGMn1a79j91kvtDtchkZXMGhGCl6beZOoZzN/VdZTN2DPct31hXFMS2U4CFqJ/0hrGMoczkcgU+DZ0ltVI1JqER2L0mDd7WtiX7Ow8ruijj+rlRkkDLYV+aQChXukDllckNIbUu728BSE5XLjgfbGclxGOh8SPTiGeYrG/q5WlnLfrywoQ9M028JiWyrHfxsj2mMo/orTk+ik5HPObNTTyt/CiOp0B6HkD4f+Qq/s5tFNbZRDI4mNxHovPZlgRLlrW53HUkzvE2axtPKee5pabSp3jUSeiUpQo+gGacOotXy7m85qdmSpsz2XW7bHSOEW61uJwlHInD0hHvvHxQpCDyBTUKWXWDJ9stNpLrLWt4aLYlR5FEcXFa3MkdD3XI+hPnT2FyRdJqYigG43847jOEspx7gx4nIQPLOegpxX0T46p1Q7i4AY6ADXHmdPRPtor5UTUzKYOO8S4uJ4kZ/Zb0BK2WjOUAn2nCk56hocxnPTJPF5YIPPwlHsnpel6Pvd4wQ3crvLkoz+iXFAfuNQnrDUiottVEhe9NuB9jiITkguLHCMeWAT8MCrXbIaVOlduLZZw3wpbYSM8OOI45q+Kjkn1NVzaCQQWw72hkIx5Nz+pCfezuje6olqyNMYCjPTcxu09ryRa55ARqLTZjNc/6RtzvR8eTShVktQ6H0pr3Sb+jtWw3lw3XEPtvRnO7kw5Cc93IYcweF1GTg+RKSKpT2mpd10zvVpzU1slexTRBUqFKxkMvsulaVeuCU5HkTkEHFXB2V15E3Q0fF1VGYTFfKzHnxAeIwpiMd41z54BIKc9UlJ51X9raCuoKGh2lpT3d0NJHJzeGfAhaVa6qmqJ6i01A4nIB5gquurrJqva67s6Z3IcQtiW4pFo1I2CiFdU9QlZz9RIAxxNq65yOWCY13LsxalOlxBTx4KgRjCh1z8xXSi5aX0vq3TknTGrbHEutqmJCX4klsKQrHQg9UqHgtJCh4VTbfXst640VBl3HbZ2Zq7SrKC4m3PEG7WtABJQ2ektoeAwHMcsciS72Y2sobxM3LhFNzadGu8WngPI+nRZrtH7P57bUe9W4ZZzHRbXYK1GmLoy/6Xf5ogXtfED07taE5Hwwg1SXcHTMjRe4GptGTP52x3eXb18uvdOqSD8wM1YvsU6lQzurfNLiRhu8QEvJbWCFB9lXCpODzB4Vnl6elM7ts6OOld/bnP6p1FEj3UnHV0oDb3xy42s/OtA2Yf8Aw7bCrpjp27Gv9Rx/VTdwDqux083NmWlQNwAq4jzIGK9J93p+7NfOYr0By5itgxlU/JSvfdaXzVVn05p+6FoRtJW9y2QO7SQS05Kekq4sk5PeSF9McsUhuIWU4QrCvAjzrNwpGSOp615BxzpJsQazcSnalzg5Wz7dW5jGvNIbPMW+4B6FKsj1/cQFg5fld13mQP0HEOo9OEjzrL9G1qFOm9xtaPIc7p6TYGGmVFBWEqM5hHEUgjix3gOMjOAMjOaqncdS3O9W60WS4qbVHsEZyNC4c8QaW6t0pOTg4UtXQDANS52TdVRtK7uW9E5BVHvDaratoYw8pa0KQ2ckD3ltpTz5ZUM1AXKmfFaZmxaEAkeY1UnTz71Ywu4HCudq3tr3Sz6s1Hp3RG3Eh5Vruy4b8ubqN2K5IcSEpKlsoQQlJASElPDkAHwFI2r9+tpN3ND3y3a2uFwt9+uttFrd017I7NuAltuKfiSIroUC42haiviVyIODwg4qFt2oVws2+evIU6U7MdkTWLn7Q42ULS3IZCkMLT4d2kcIHlnNIrTr/GHA6oEJ4cjrwnmUjlnGfAfGsmlpop4o3uLgcNdkOOScA65yMZ6AEKPr9p6q21k1M5rXNGQNF5s868OxEu3lIRNfaa9r55K1JSAniUOZx5Hp0yaX9FyrbpbX1m1df0Sk2aLFuEaZIjRRJXBL7HAiQGhzISeLJHQEedJ7DYWAcYx0pQYU61hTLq0LT9lSSUkH0I5ivZJsucTwcCD69OngqVSXB9DWCsYMkHOCpv1d2zNV3xDmm9ttPRbhZVw247d61FIfZYfWAAXUQEkpLPI8IVknrk4OY33Y3x3E3issfSupdN6ctFrX7IbmIsiRJkTTHKlIQFLJS3H43FkNc+EE+JJpsd2vjU84ta3FHmpRJJPiST1rG8y57y2sBfCeEqzjixyzjJxnrSYkja8FrdRwJJJz144+nkpap2uuFQXNaQ1rtOCnTsUbsWjQltvmhZ6pAmablSbomMtsdwi3yVtJS2w4FYQoLKlBopAUpZGRxZF1dc7l7ZbT2S57ka4u0SyNEIRLkZCVzXUo4WkcA5vOcKAlOPe4RjISnly10Zfdc7VXCHfrJfNQzfyjeY0i+xrXb21odjJGFN8JSpXCEDgBUsD3s4GSakGTv1b7rqCFctxdhLdd7LanjLtxTeFSLjBcTj3196ru3VHA91BCeQHveMlT1c1BLI6JvaRv72jgO9jvAb2M5Oo/dXKhutHVUcfayhrx3deeOa397JVt3/11C3r3N0c3oLS7TQcssdiKhOqdTtAcLbrrmPqGClICXFA8KThHH9oNG566kXK0RdKWO1Q9OaVt7hchWG28SYyFYx3jh6vOlOAXV5UefQcqfV22t1Fuvc16+2+1EjVFovjCLgy9PuTQnNIWP5p1JwU8Cst55D3cYGDUfam0DrnRjQk6m0bfLWwpZaEiTb3EMqWPBLhTwKPlg4PhnnUZUXk3V27I7Ufk4bvodcqqX51zcXR9kWx9RrnxyFoocW6vOST1rxeL9b9PxW5N1lIYS6vumBgqcecJA4G0jJWrJHIDxrd250nrvdu5SrTtjbGHmYiu7nagnAptVvUc8i4B9e7y91tAJOM4IBq3Oz+wm3u0zzV/ZD2p9ZJTl/U91bBcbUeqYbP2IzYIODgr81DoI65V1BYo+1ub8dGDV59OQ8Sk9nthK6+vD3Ddj6lQjt52Z9z9w3Y9z19KlbcaYUEuezBtK9QT0k+CSeGElQ6FfvjyV4Wo212z2+2isr1j210vGskeUkCY+hRcmTsEkGQ+ffc69M8PM4SKWBxPLU8tSlLUcrWo5JPmT1PxNZ21EfwrH9oNvrhdmmmpPwYT+VvE+Z4krdrLshbrFGBE3LuZK3gtIB4QPLp/npWBzicWAkEk9ABnr06VDu7nam2y2llqsT817UGp/wCjsVnAfkg9cOqA4WfDPFzAOcEZqmO6nac3c3VekxLzev5NWBWUpsWn5HvOAH7L837Sv1gjAPQBNSOxfsg2i2xeJYo+yhPF78j5Dn9kndtrLfZhuudvP6BW53a7Ve0e20iTYk3FepNRsJcBs9mCZDqHEg5S8sfVsgY94kkp8QapfuZ2rN39xXn4qbuNJWh0qQiBY31B9bShjgdl9TyPPhGDnoKjJ0tsxlxIUZmJHJJLTKOFJyc+8SSV/tEn150mPEjr0z48via+sNlPYtYNlGtlnb28v8z+APg3h+qzG6bbXC6ZZGdxnIDj81MHY5tLK92ru4zHShtiAw0cea1BRJJ5k+7nnVwO0BZnL5sxqa1IbypdvdUjAyeIJJFVb7DSVStW6mmFHIvMNA+PuoUMVeK9QUT7LJguhBDzJRhScisg2/qxDtO6Vg+At08loGz8JdahG7iQfqufen5IvGk7VJZeSpxUKOQpSjycSgJOceoIOOdb0abcZTUcQbo7bJdvfXJt8nPvQpXCEqOP0VhIbcAyFJBznApNXandvtZ3XQ8ppTEbv3Jdt4sgdwpRJQCevArKc+WD40TpgiyOJKk8DoGRk5CgPLxBSOXqlXnU+4DtN+HUHvDyPJYhVwT2qucxp3XNOik7eK72fWu0eltcIhNM3By+woM1hKf+DzG3wHmxzOACnI/VUnrmoqiaERJ2ytWurnfXm0vLchwWGVkPPFUp0IYaQB76lK4iM+ZPQU3Jeslx5i9IPcfsdxvEC/oQMlKVsJcS8cdPeCEE/AeRpf261RI1Bp+0qKQ1B0zDNugBSjxe1OkuSZOOnFhwNJPgnJGFEkPIaGS3UY3SQ0O3v+08G/PRaAdpRTtNfnDjGBp1yle125+0WqDbr84zd3YDrshmE99dAtrrijxFKFe68+T9t5YIyMJGAFFWmOy9Yyn73eJsiXxL+sfedUt2QtOE8IUSTgAYyDgcOB1ynTjd1dnO74i1bGF/XuAEF9Y590jHQ8hlXgCPE4r5qjU8Ow2pcsNpDTKA2yw0kDiJ5IbSkeJ5dOY686YSSS1MgaB3jy6Z/VZ1V1lVcpt+UkuJ0CfnZet83UvaAk31PdiDpS1IhtgABKH31cQSkDkAEIV06cvSr3DKRw4GeVVv7Ie3U7RukfyjfmQi8XZ1c2d7oBS65jKfgkBKfThNWQC+JolIV1Bx5edZttVVsqriWxnusAaPHHE+pyt72doDbbfHC7jjJ8yvoI6E8snmKHEo4Txcx1rx3gUpKTzyTWRKivOQrA5+h/zmqyFOFN/VCe4Rb7wnCfyZOZkqWnqlsngc+XCsn5Vyp7Remo+0nai1LDtzYZgIvP5RjJAwER5JDvAPRIcKf2a61Xm3i6WmXbu8SkyWVtA+WUmub30h1lW/qrRG4YjqbGotPpjyF9AZUZXC6D6gOIrTfZlWGG4ups6PafmNVXNpIt6nEvQ/dQFu3bg1dZD3EFEuh7iHiFjJ/HNRyalvVcVm66Ls17ypb8mAqO6euVt8wfiSV/dUSnlW8wnu4VIkGDlfKKKKVSaKKKKEIooooQiiiihCKKKKEIooooQiiiihCKKKKEIr6nrivle2ULcdQ22jiUtQSE+Z8qEBS3oZTdn0berqpPCWLeEBXhxLJUBnzKkIqavo5rC4jVut9wu7SRY9PKhsuK6ofkuJxj14G3B+161B14cVa9qnY55G53JKAenE22lOPlxJcFW47Amm34Wy16vy08KdTanYtqOXNSIzaXVn1HvEVSts6o0tkncDguwPmVNWiETVsbOQ1VxrQyiJb4sNoDgYZS2nPoAP4fea3gU4KikeYxWi04EhOFYArIXs5TgZPjXzUTk5C0fHMLbUtKRg8OT0rEtYCgknp68qwFRPMKBGBXkJV4gfEmvR4LxeZraHmyhQ5KBSfhiqpdo3ZZF3dTfbUfYp8MkxZiUlSEA5Utp1I5qaUefLmk8xnmDbFSQU5JxgeFJN2tMW6sKjyEtrCvdIKck+fUVJ2q5y2uobPCeHHxTWrpI62EwStyCuZLuqbxpa5Ks15iuW26MKOWXFgodwccbLoOFpPgpJ58utaDl+Q4+qYhkNtkjvWce55cQxy4SDzHLHPHLkm6W4fZwtV6bdQI7D0V3JVGfb4mz5FOeaFeSkkEZ5Gq2X3s1Xc3u82bQlx9hXZ4f5QmImTUqjtt4zwoS4ONSsEHGVHBGM1r1ovNvuRw0brzqeY+fEfXzWbS+z6qklLbeN7wULw4ybbHt0lBAFvu0mMFE/ZbcbSpv8QulxGq48GIp2StQW8scIyVKWAfdQlPU5JJOP0vSmm43PiNXPS0+Qhy4vvsPMKQcgOpcKSMj9VRPljFWA0psDatBfyQvuo9W2G9o1OtDa3IL5dkQCpQQltaVYCEglPER7wBPPAKatNwdFFF2kx3scAOJHFN6LYu4XSR7ZGECM4cfLUhbmwe283U19/llrCM2wzDb4mm31AJiNAcSlLWfdSsj7RP2E8Q5HOJusG/tjvu4Vt0Jp2wufkC4QZS7fe3ctCc+wU5DLZIPc8PeYUscSzwkYGQWlupcbTxHQNjeCdMWVxDNzbYc4U3a4JSFGMspVlUZnKe8A+26QPsp4kxfeNQNWqbp/V81xRftF6iy1OpwOGM4oNOIAHIJ4FIIA5e74VQ5aZl2cZKhvec0hreTNO76nn+6eSXyKxVkNppRgAjeKcPbJtCnGrBqXusC3XENLUB9lp9PCon0Cm0f2qZO0O6t92n1Ki/WKKZbbqAzc7X3hSm4x0/Y4c/ZkIHEUK/O5JwSfesL2gNLuav0DdLKylPfyY5DJI/pkgKQPmpIGfWqVWSYbjb2nXSe84SlzI6LHJXLzyK0X2ZMo9pNnp7DcWBzQToeh+2Eltj7xbLlHcqU4PX911X2z3G01ubp+PqfSNxTMgyBgjkHGHByU26jOULB5FJ9CMgg0+HIrTrKkrGUqRz9RXL7a/cLVm2upm9U6NurEOc4ngmx5XEYV0Qn7LchKfsr5BKXU4Kcgk/aUL5bP9o7RO7sE29oLsmp44xNsU5QEhBGMuNEcnmyTyUnny5gcs/PXtK9kl02Kn96pGmSlOocOLf8Al+6vOze2NJfWCOUhsoGCOR8lAe/ljg7c7hWHde2W1lqbaru2q4SWkJS69FdJac4iPtkFaTlXP7XhSf26ttLpr/QmmdzdM25ydKs8kwZDLCeN52JKHGyoJHXgcSoHHi8BU09obQcbWOmpjD7YKH2VNqPlkEE/jSFsJfIO6GxsvbXVs0xbtbkPaeuMg/WLjuo5NyU46lKkpd9SBin+z18kiio7wDvPgd2byT+V/D0Bz80+u9AJo5adgADhkf8AIf2XMNTDrKTxJI4TwqBGCk8+Sh4HkeVfM46irK6u0XbdU32VpfcFpvTWsbe85AdvYbUtl51BKSJjYzxoJAKXkjj4VJJ4sAiFdwNt9UbcXY2rVFrVDWtPeRn0KDkWY0c4djvDKVoPxyPHxx9KWvaGnr3di87snTr4g81je5vglvLj1HmmpmvnInJFBz1I6da+fHP34qf3gOKTCkfYLZeZvnuNE0YxNXb4gQuVcZyWuP2aM2MqVgkDiJwkZI5kVZjSm6ux2hFRY2yXZ9g3OLaZBMbVGorglqfJWjIEhBKVOZJzgtttpT04AQcNHbODb9vOyZP1LYnZjl+3Mugscic273UezssqXwhahzCnApZ6HOU4x4t1MCDbSI9ujhhhoBLTQz7iPBPMk/PJzzrJr5cTdKiaJxJiadwNBIBI+IuwQTqcDXknNwuLrFHGyJv4rtcnXA5YSrdLlMvl7vOpbsQu4324Oznz3yni2gnDTPGoJyG0jAwkDB9KxNNgq6Vga989cnr8q3WEYyeXKq+87v8AnL+yodRUPqZXTSnLnHJW5GaAwceNbraQR0rUYcHMcvKttpQOR6UzeU0cs3AFV9SlOCCPvr4FDOQa9Zx4ikiVwvhbSSR4H/PStd5kkgtjhAPLyHyrOFEqxmvYRxDnXud3VdAkJBRpjT8dHDDsrUQkKSXIr77DmCcnK23ApWTzwoqTnGAKcsLX+42jV+26c1pf74yptSJtj1FMRcIlxZwMscKm08BKQQFe94eudVaUpBrXXwjx8KWMzn6Sd4dCM/U6j0KkKW71tG4OikOByzkK3+1O7e0mr9O2yw6Fv9ntLicMtaYlLbt8qI6QFKYSwoICsKyOIDCgCSTjNSPDfZfj94kn4EYIPjkeeeXPn0HhXNfX0tSdL3CUkcMmMhEhh8EhbLiFpKXEKHNKh4Gpr7QPaU1xtxfJu3uiYkdqa6wxOVfpYDoQh2M0U90xjCnSQpXve6Cs8gOZo9y9nVTfqyJtqcXPlJ0ceAbjJ3uYGemVuGzW28c9C+asYGCPA0556KyW5W9W3uzloF11zfm4q3MCLBZ+tmSlHkA00n3iM9VHCR4keNOd2O1huXuY0q2WiRI0Lp1xI4okFxK7tLTkj657ATGBH5iRxdMhQxUGuXS43O4SL5dbjMuN1mFSpNymvKdkuknOOI8kJ8kpx6k8sfcpAyeqRn7+tfQOwfsJs2zwbWXUCecdfhB8B+pVPv8At5V3EmGk7jPDiV7bWiIy4xbo7cRhzk6lpSip05yS44olayc5yo9c4xjFazqfTry9T5Vv2m13fUE9u0WC2SrjNeICI8ZlTrhycDkkE/OpcsnZyYssFF93m1nC0zHUpYVa2FpkTxy6uFP1TefLK1jxQK1+4X222WMRPcAeTWjJPkBqqjSW2suLu0Y0kc3Hh6kqB1BxxxLTCFOLWQlKUp4iok4AAqQbFsTqF+Ci9azlMaZgOBRbbl+9OdSB1EYELQD4Kd4AfzeKnhdd7Nu9v4S7TszoxqNJLQQ7eZRD0x5Q8S8oFSf6jYQ3z+zUC6213qa+iXPvV3fWkhag2lRCQVchyHU5I5nn+FVyW4Xa5NMkTOwj6u1cR4N5eqm46K30hDZn9q/o34R68/RWp7EOnYlu05d78ttbgm3uQESEnqygJSCUnwJ4uhzy8fCaN995522b2lrRYLC3fp99lPKdhtuhLpgMMKW6tvmBx8Sk8IP6Kk9SCGv2TdOosG0tjhqSDxRg6snkSt0lauvkVEfIVHO9d7gXHfS9NNjvDp6xxbW2rBw29KUp9xY8ld0hKCR4LNfMszGXfaGomnG+xu8SPoM+q0m5VpslmEsYwQB9Ut7vae09vFoyNuHoqYkOtkux5IaPfRHk8lNuIHvJyPdcRjPJKhnhGar3O73SHOesN7j+wXJhXNviBbWQeS21D3XEHzBPImpqsmupWkZitXRm33m0N41BbknKLlGRyU6Ek4ElocRSr85IUFEeK9uJoHb7dHVtv0pZ9R2OHEusBN0F0uLikRY6HQstFtScrStQSVEchgpCgroLFaSaGRtK9pdEc7p5s6g9R9/DVQEFFDtzGJqZuJxxCq1OdTJvC54GfZbNLdJJ5pUoKbSPiFKFLVmmG22ODaEyO4YjMAyVJ6lSsqKQB1UeLHPoADz5At7UFo/kS/qDTr94ZnyH1R4kaU2TwOt94VKVzGcAoT94p9WfYTWEnRcHXmoL+yjTr0tMEtw5SESGVEDKlIxkpA5n0zjnyq51DYRC3tHjc5eOn91GRbG3WrkkpWR/7RAd0CyStwojTLVvZjr4k4bjw2U8RPiEpSMqJ8zjmeZyTkypsjsjqLXOoYep9YxlNpi4diQgoFEVRwe8cI5F7B6DPB4+8MCVNreyjpuwd0+9AabyElakkrcd8uJxWVEeOAQPSrJad0tb7BDbhwI7bSG+gAxk/wAay+9bUQQMdBbh3joXHj6dM/NWKw7GxW+Tt6jvPCy2KyRLNbmYENpKUtIShIBwEgClZKwnAUDjyr5wgnGRjAxyrypJJHDjlWZuO8d5XwDOvJZU8K8EIwoZxg0cZJwATgZ5q6V8BASSoAdfCviUlXQYHL4GkjjOq9xhZ+NZGT5VT3t7aUTcthjOixuJej9VlxK8c2ostB4h8ONTA+VW+WSj7eMCoP7SNlf1BtZufpyO2Frf02m8Npx4xFlSz8eFKPuqxbK1RorrBLng4Z8jofuo+5wiakezwXN7TryLjti4nOXLdcOh6hKwMfI8Z+6osltKZkONKGChRGKkfbA9/b9TWcc1uwQ+2k88rbKiPnkpFMfUiAm7PKxgOcKx8CK+pIz3yFmr9WglJdFFFLpJFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEVv2FPFd4pI+y4F/dz/hWgOdLGlm0quKnFjk22pXzrxxwF6OKfe503g0rpe1JHDwMLeI/S4iV5PzdI+VdCuyfp2Rp3YvbO1ymi2qVFuF+cTjqX3lhpXzaW3XO7eZTbWpLdbIx9yFb2m8D9Ik5/ACuqegIggWHS1lLfdmzaUtkQJHgoNJSr/oh+NZb7S6gxWyKD+ZxPyH91aNnIt6qe7oPunorH2UH0oTxYCyc5HxrwgYOSDjJNZfdISFc8+p/yawkq8cl7LhwEpzjxzXkqIznPU9K8gJBzjkD516JCzyAGK8C8SPrZGopWj71C0pcV2+8PwXUwJQZQ6WXwOJtXCsEEcQAPI8iR41TXTm9G7UuwsXuVudqBcpIUzOaTa7UoMSG8h1Cm1R8hIIJByRgiruvK4UKKFc8HB9cVSHdywp2x3Vlahajqj2DVssqcU2Tww7tzVknokPDK/D3wvHJNXjZGojkZLSOALvibkAk44tBI6a+ipe2bK2KmFXQuILeIHTqnxZd/tduwe7m65tMqSn3kIvOm0xmlkeCn7e4CPiY6vlUFb76j1Lf9Qu6qYsD9kEzumJLkecJVtcXw8KSmShIKCoADu3UA5GeLORSvdrb+U3Vv2lxLD/2uAABlzmc5A+wT1BAx54yVD7Z9QTLDJWxIbEeUEhD7D6AtK0q/NWhQKXEKHgcoI86t9H2VLJ7xHG0uxqOB+mn0VCtG3V0oJGvc845kcQFGekpGqNtJ8jUTlohvzbs2u0Pe3MJdS2h4JSFJ4gcHPIk9QrwBrQsjmpF2p3Vk91TEO2gyQ2STxOtOcCG+fIZdCAR1xUgawlQpNjnQQtLVomMlpGQf96pIVltSVDmYxXwjB/m1dPdNNiLqCK7tZdrTKW24ufcLbdWV45lDkhCJSM+IDrST8Fg+NWaGd9XEJSwZJA9P3H91eYdsagxFtNN+GQ445l3HVPe/TVItcOIy8tQjRwyD+k6pRU4snqStxalE+oHhSHr5tkaNunCTwojAoz4YUnh/cKJV1bNyiMMo4wZBUfRKAVj8UAfOtTVKhfG7ZpZpZDl8uDMbhRzV3QUC4r4AJ/GoanhMczCeGS70GqyUS1FxuTZZPje7KuLMji76cbJypTqOL1zz/jVLN0dKSNDa+kSVNlNq1E6p9pXIBqZ/SN48M8iPPw6Vd9+6af0hpy3L1NdotvRLfZhMqfXgOSXc8LSeR5nn6ADJ5VGe8+2sDWlmkxZDK1MOjJU2PfaKfsuJ8MjmcfEdDUDsdtFLs3dRVjPZOJHmM6/Jbbe7Wy7URp/ztVWWHuFOM9R54r0qS6FRne/kNuwXQ9ElR3lNSIiwftNODmnzIxw58PEJhYuun7ovTWpGi3PaGWXcYbmNeDjZPXocjrkHOCCBtlZOAD1r68gqKS9UgezD43Dz0PULDZoJ7dUFpy1wPJWM0L2u7m3ATpndziusAthLeooscJdaAAA9sjp9eriMjOCQcnHzT2pbdtfu5C1hZ7wxM0nq4pgypEVzjZDgwGXQUk5/QPIEFKAeoqt6nVsOB5hZQ4MkFJwRy/z8uVId0cmxos1q2SFRGZ6T7UwykBh08sK7r7KVAjkpOCnwrH737I7cx8s1p7jZAQ9n5TnmOhB1HLRX6zbcVDY2w13e3eDufqr0dp3bRmUxH3VsaAtTgah3RtGVBRICY8lJ6cKhhpQ8wyfE1Cdn1gwuxu6G1pZW9R6Vlqy5b5CylcVfMd/FcHNl0ZJynkrGFBQqVOyVvTb9zdFSdstcFidOixFxZUWTzRNiKBCgeeSCD1HNJyQRhJpjbo7T3LbjUjkYLdkWyQS/bpbnP2hjiA5qHIuoKkocSOYPCr7K01kdAJ6Vz7VXaTwHAPDLfyuB4pPa62vge2/W0/hv445HxUK7udnqboq2HX2hLg7qXQjqhmahJEu1OK5Bic0Ae7VnkFglC+RBycCGwCk88Y5Y8iKt3pfWGotGT1XCySUpD7KosyK8gOxpkdXJbLzavdcQodUkeRGCAabO53Z8suqrJL3L2PtzraIjKpWoNIlzvJFtSD78iGo+8/G8x9tvIByMKrTNn9qDpS1zs8g7r5+PioCjr468aaP6dfJaPZn1xYJ+j9a7Danv8eyp1aw09aZsz/g7UxtYV3a/ILKUc+X2TjmQC7NWbZb1aYtVzv+o9uGkwrSwqZMlRr1HcaUyOZW0kErXy97GOgNVHkshSuSs4PIjr/npT52t3r3C2l1IxqOw3VcsNMriLhT1KfjOsqxxtqbUSOE4GRRddmpzPJV0G67e7xac8cYOHAjGR1BHNTTm0dfGyOtBy0YBHIePkpVTfrLFZ9rXdmVtpTxJS24lbqvIJTkZUeQA8/KnfrTSOvNAWO3ag1XBsto9vuUa2tWf2oyZ2HEFaluKbIbYKU8JCFe8SojHmzY3bSk2iQ9dNMbDbZ2W7ue8i4RbG0l1hfgtshI4VA88pwc1obe9o3Tt8tGpdGb9s3OdC1FdEX1q6W1SUSok5GAFpJBwMJAHI4GRjByKzLZ7lGO2NPhjSMjIc4jnugaacddTywm9Ps3bWB7HSb7j8OmAPNO2OXOLKgRk+FKLbxCeeTmnBG2suGpIf8AKLZfXundaWSWG3m4N2uoj3eOopHeNLUUBGArJBOCQcBIxzZNvuMt2XcbdcoKYc61T37bLaRIS+gPMr4VcDiQApOfHFQxAla5zSDjGRwcPMHUKo3OxVlsG9M0bvI8vRLzbmfhWZSgRypPDuBgeHKvSZHQdabFqgyMlbjaxxdcVl77HWtJLgI5jBry4+ccuvxrzdRxWZ58HOK1i8pXJGSSQAB1OT/n/vrEpwkE5/7q3NMWKLrTcHTmh57s5Fvubc+TPTBd7p91thnjS33nVKFHkoj3qVZFkne4AE/IZT2gonXCpbTM4krPozRsrcR9Fzusd207f215My96imJDUZ1lh3K2GCc94pakBJV0GTkHxjHdDXv+6juZqDXDTbiY90mqXGQ59pLCQEN58jwpBx61vX+96v7QNzteittNBv2vS+mGFtwbXHcKmmG1Kyp6S6ohHESOqj1JwSSaftk2o2i2xiou272r03q6dUWK0KPcpwM4cf8AtOHJwUoCQP01dK0/ZCmiss/8Qq8mUjDImjLgDqSePecMZ1AHAK+VFtEVJ7hR6MBy57jgE+GeQUYaW0dqfWM42zSlim3KS2kLdTGb4gygnAWtXIITn85RAqZYOwugtvoAvG9uu2BMHF/vFZnkqWnAyEuPkEcX6rYIGf5zPKmzqTtKX1NqGlNubJB0nYWvdQzCbCCpXiteMFayOXeLKlkcicZFRBPulwucpc25THZT7nNTjqySf8PgOXStBfFert/uu93i6N1f5E8B6aqKD7bbzhg7Z/Xg30HEqb772nGbJaHNKbP6RgaYtakpSpxls96+odFuLUe8dV+s4SR4YFQjetRXvUMlcy9XGRLcWpS/rFkgE9cDoM+fU1oKVknNYFKOTjw608obPRW3LoWd88XHVx8yUjUXGorBuvdoODRoPkguEZyOg6V7tlgd1hquyaSjDiM6Slx/0YQcrP3A49a15D7UdlT8hQQ237yifDy+PPHKp27JG2M+5Xh3XdziK4pGBHQRktsj7J5+JIz8h51Xtt75HZLZI9x77hgKT2dtzq+racaN4q4+gLIbHpyJCba7sIbSAkDkMdB+4fKqlaswrdjcFbiypxdzhKyf0TCbwPgMmrhM6z0qzfm9DN3eN+XfYjcE28K+s9mC+Arx0HPPLrgZ6VTfdyMdOb7XxqX9W1qGGzKZ4uQU4yO7IHrwcP3GvmjZsSmon7UYc9mR494HPyV324Z21oPZa7pGUixpvc3GSlpSkFl1KwfIqGc/eD6HNMyVGn3NM6LZCGV2F1MUtgkJ9mc43GeQPLhJeT8MeVL825Mx7qltQT/pEckkHkS2rw8/5z8KTtP3hMQ6+YYLYnXeDb4EIrGQl51TiSs/1Gy4vPhw1eqKJwDnNHIH6gH6ZVI2Ju1RaKx8kTt0ObqeeRqm68mZu/NtFliWGPGVbIKnuKE2kOOq5JSVqIGc8IVk5JyfE0qWRrWcfUKtLz7XcJvsZ7ydDhLKlIQDj3+LCGQc44lk4ByATilLaSQ5bn58rTNxMZyQ6O+uobC126AjKGQ2k+6qS5wrKAeSc8R6U9bldLNbrMLPaWW4MPvC+WgoqckPK/pXln33nD04lHIBITwjlTqvrTDIYN3LcaDp1J6eHXw0VruPtCraFzpmPzPIO9jhjlnxwpWlb8bmSkuqcv8Ap3SyGUBEeJbISrxJTgc+Na1tRvuK/gKaSd5N05S3n5W7WomWGQpTkhFrtkdtIA58ksqzy8Of8aZ0axXJtXtl7W5HQSCiIDhxQ83Dn3PLh+158OOeC6Id3J1DbNrNPt8Lk9QM91hJCYkFJ+sPLoTkJT6q59ecBFS07XkNjaG8XHdGgHHU5yqZFtBebpVNiilOXHGBwAVs+yRqXX+r9AStVa51HMujd1nLXajKistONw0ZQgnukJyVcyc5H2ceNTlxYUSOeOdNzQlggaY0vbrBAjJjRYMduOyyhWQ0hCQAkeeAAM+PM+NLoUeYT0Hn8ayi7VTKytkmjGGk6AaYHp4LcaOF0EDI3EkganxW0SSgZPTwoDp/R5dBisJWoAZI5k+NCVZbAwOnOoshOlm4sFXX5/CmzqC3IuF5g254ZZvcKfZJBIyC2+yCM/2VU4SrhHLOB1z0P+f8+dI92UpE+zSicdzc2cH+uFN/9unVNJ2MzX9Nfkk5W77C3quPu1SnLPuCm3zkFK0peYdbVy95HMpI/YI+VNnW0MxZyUk82ipj48Kjz/GpN3QhNac7V+p4kVvu46dWSg0gDAS088rhA9OFymZuowhi5SmQBluYs59FDi/jX1rSTdu1kw/M0H5jKyyVu7lvQkJg0UGin6bIooooQiiiihCKKKKEIooooQiiiihCKKKKEIooooQgdadGgoKp9yMcHm8tpgDHipYFNceXnT92fQhWqbetY9xM+MpXwSriP7q4l0YV03it7VjP5a3i/J2eJLs+NFSOvI8Ax95NdcI8dtnUl3bY5NxfZoiPIBDQP71GuTul435Z7RWnoSuZlapgMn1zIbFdYbWsyJ14lLHN24u/cnCR+6sa9qchHu8fgT9lctl25Mr/ACSuFZ5cGeHnQFDOFJ5gYGTzrAFjjwU9Tj8K+nAweLJxnHrWPq3rKpxQHDy6+VZFLAQMnzrAVY94nyry/JixW1yp8kMRY6VOvOno22kZUo+gAJ+Veta55DWcSuHEN1K9uji58yMcz0qJ96dAwdW2eTBnQky4kpotSWR7pcRnIIJ+ytJAUlWORAqR7Hf7TqizQb/Y5iJMC4R25UZ1I5LbWkKSeXjg8x4EVmlRUPtqS4CoHl0zjPKnEEs1BUB7e65p8sELhzGVMW6RkFc6bum+bXTUWTUrzkq3uOqat92LeEPAdGnR0Q6PEE885BIr1L1Bbr201GvDS3WWspQUngcbSSOINrOeHPqCknBKTVvNf7QQdSsyEhhlQkpIcbdbDjbo544kkHiHM8+RwfSqva27PV0sEpcnSzkq3NheTCdSqVESMdEEnvED0Jc69fCtUtl7oLjh0rtyX6H5cPssqvOxD2ymooeB5JiKmP2VDynZbVyt6k924soCSEqGOFxBPug5x4jJxnJAqPJURLkS5aeguOEwlm7W8BeCpkYLyPiAlKvi2ae970dryEQEMW95QBwpqYpnl4jDqUnHh401V6W15bJ8e+osjTfsBU8tKJTa+NAGVgAHPNIPIfKrvbnwtG82RuviOI/zCgoLNXwZLoz+mU4LZe2R/pbzqA0qH3inVdAFkEn0OEkffUk9n/R8/WWska/ntqYgW9siEp4hCUtge88onATnGQTyCQc9ajLRO2c3WOtWrXGkqesr7KLhDZyAHGlOKASsDwQvjSc+OPAirea+0vC252YGm4UxMSRqOXFsPtGPeSl8kvLAHP3WG3sfAelQN9q4qaRtDAfxJcD/AItPH1xxU3s/aI6Vr7pLqGgkeYUVbh320brTTc7ktyVaXUuQLGg8vZoQyHJiUHmHnlpStB5cLbSAc8fuzLs5qReu9Cew3iW3LvlhkKtV0WHStTziRlEjnzw62UrCiOauMdQcQRdXmDdmUxWEsMJBS0wgYDTKWylCB6JSAkfDnTr2Eurtk3lchd4lUXVFmU04k/8ArERwcK//AGbuB86g7pAypoXxsGAwbzPDHEeo1Pimmy20M9TeXCU5bL/gSpvFszA1LBMd+MoFKlOtON8nGF9ONB8DyGR44GeYqsF7tuodCPmHqloriFfCxcWkkoV5BwfmK69evPr1rpPdLOzNbIU34eA/Com13tnCuTTqHIiFocSQpPd8QUPIjBFe7Gbf1mz7gwnLOYP6K+3zZynurcuGHdVTEOtvISttaVIUMgpOU49D41qSEcSSMZ6+HWn/AKz2IuVikrm6PeVCGcrhPEqYJ8eE81Izy8x6jwjidLnWV8QdT2l+2uj88pC2V+qVp5Ef5NfR1n2ytd+jG6/ddzBWVV+ztZbXHLd5vIha9g1DfdvdTw9WaakGPMhOBSDkgLGeaFY6p/wroVo3VWju0ntfzkORnlhR42+b1rnhBSF8J+0jnhaDjjQSk4ykjn062xMbK2nEOoXzCkKCgfmKWtp90r/stq5F9tfE9CeUlufE4vdfazzPkFAdPPpVI9oOxrrowXO26VMerTycP5T58lY9mr4yEGgrxmJ+hB5Hqpm1Rpm/aPvkrT9/iCPMiKSVhPNt1C8lt1pX9I2sAlKh0wpJwpJAx2PUd70xdYl+0/cX4NwgOh2PIZUUrbXjHI/AkEHkQSDyJqyDy9Hdojb6Nc7TcG0yUtFdsn8OfZFnBXHeT1UwvGFp5FJHGn3k4Nbr7p+86XvUzTmoLa9CuEBSUvsrGcBQyhaVjk4hacKStJ4VA8ufEBklFWsrWElu7IzRzObTz9Oirm1OzMthqBUU5zE7VrgtTcjZi1b2W2VrzbC0w7ZreKyuTfNLRGu7ZuiE5UuXARn3V8PNyOMnPEpHLIFWvcUOhyORB6g+tWxtN0n2WfGu1qmPRJcNxLzD7SylxC0nIUkjmDyrDuztbA32tc/cjQ8BmHuBbI6pmobJHaS23e2U83J0VtPIPgc3WUjnzUgDmKvmzu0zqd/ulYcs5Hp4Hw+y4t1wFwb2cmjx9f7qp6kBXhisS2EnOedbBBT9rIIJCs9c9P4V8PPl5VpGA4Z/zCkcuBTq2W1tD233V01rG6JWuFbZgU+EDKghSVIKhy6gKz8qsxdtr7fC09q3dPazdHSmptL259d6VZ47LonRY8iR/Mqwohoo4jzUMHhJIHMCm7rYVyKeVWj2bDlp7HW4lwhu8Dl4v9utshbZ4S1HDjfEXlde6IJScdSvHjWfbX0Jp3RVkTsFzmMIIBDgXc+YIycEKWpuyraaSmnbvAAuHgQFtuyGuMpZdS4nPurCuIKT4EHy8fn86+tukEGkpl1HfONsuNqDaikhs5CQOgOOQwPDwFbaFhQwFc+WKpckZYSFlj2bhxhb5fGOZ5157zi6VpFSueD0619D3BzJ50l2a43QOK21q4BkHoCevpSxtvP/ACFd79u7cXIsey6Dtk23R1OuDikXiUykNNttjn9k8yeXXp1Dal3SLEjqeeDi15DLLLTSnXHnlfYbQhPNSifAeANbO7Lb+2u1Ng2tk26SjUesUx9V6odmhKHIzhyGooaSBwBISOR5jBHiQJC3UT66aOjYO9Kd3x3eLz6N09Va9mIOxe+4SDDYxofHkE0V7v6ylWWPp+1yGbTamUqAiQUd00SoAFZSOqjjmVZJ8xTZdfekOrkyHVuuuYK1rUSVHHjny8B4VrsoCEBIGMDpQV8yOeMZ5V9K0Nuo7XEGUzA37+pTKqrqivfmVxP+dELIGawqX86xzZ8OKgOSZTbafVWfwHOkheo2JbncWqHMnOnkkMtnmf3/AIUVFfTQjL3gL2Gjml0a0pUWeRIPStKZcYkFAckvhJPNKRzUr4ClK16H13qR5DZZatraiBwpHeu4+Pgfuqb9rezIyJjcy4MuOvBWTIkDicJ9B0T+J9TVFvm39utjD2bt5ys1v2WqqpwMvdCjPa/aDUG5F4YmXiCti2NlK2oiwcufrOY6DHPHU1fPS9vs+2WjZN4vCG48K0w3JUpxRSn6ttBUrmeXEQnAHicAcyK39B7cWzTMdKY8dIOBxE4OfnTF7Wc2Q7prTm38NCAnU16Z9tJ/9Si5kPAjxGG0n1xjxr58u1/qNsroyGZ2GZ4dANT9FokdLDs9b3zNGrQq/wBwk3e7apG46pZY1zLf/LjEwuBXsquEBiASQPqRGwytPL3lnqEYVIW81ohb6bZ2fdHRoU3cI7ftDDPFlTMhHuPRl+uUqSfP3FAAVHUxUmXeC64QkuodcVgAAKKkeXlxfu8qkvs4r9ov+sNBOOtqYuMZm/tNYPJwqDMnA/WUWV4HLKzUvXSljG1rNHQ6jpucC3yxr81nuyl4dcKqWgqjvNlyfVVlF7ZnxospWUutv9y6hY4VtLIKVIVnxyU/dTWmSX3H7tfIa197PeFot6UKx3jmAHFAeiVcP/O1NHaW2TuOmL0i96efLT1zltRscQSh5SyeHiPTiABPF+ilRPTJh+LYr9cZrC9KwXJVrtAVFZf40tF1wfbWOLzJPPqAQOoq6WmopqinFVCRunry8PmkpbHJb6t0EY3jx9E+NO3RFnsjOjrI2wj2RHey5LhCUF5Q951avPlwoHMkJGAaWIEqzaffbntXFUu4oypUxwE8Csn+ZT/R4BxxHKsgkFIITTHiaS3BnuBMaxx22wriCDNSUg+fCnJJ9etPzTew24Oq3w1cpqokZWciCyriVnwK3QCPiEn4VG1hpIiXyzNGeOuSfumbdm7hUvyGkZ4krSn62uN6nI05p+Gubc5I+raHQAdVuH81IHVRPTFWl7Lex/8AI6KrUF3V7TdLhwuy5KkY7w9UpSD0QkHAHjzJ64G1tH2arPpdpDy7a2yk8KnEBRUp4jHNxZ95ZJ69B5AVYW2W9qC0hhoAIRy6fD/Cs32h2iifEaKg0YeJ5n+3h81o+zWysVmHav1kPNboAQhI5Dl/nx/dX0kjmonmeuMV5WtBJVxDl6dKTLRqWyX9y5N2iamQbTPctkspIw3JbShS2+XXhDifvqjiN5YXtGgxn1VvLwCGpYSoj3gCM5HLxrGkE5UrAGOpFeVrCsYPIdK88WBjiOevIcqSxldrI5yVkYxnIyOVI2qlkWtMoHnGkx3+XIDgeQf3ClcHOTxH7qR9Uo7zTtxR0IYWoH1Az/CvRo4IXMTtiwDpftTajeAKcTIc5PzZaV/Co/3aaJu89QTnjWy4k+haTk/galr6QpIX2ipU1AA9tstvf5eJLZTn+6KjTdRlMhDcpJAL9tZe+4qH7hX1VYZe1ttNJ/QPssvrG7s8jfEqJj0ooOaKsCj0UUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIQKkHapHDc4ziOvfO8/wCqyoio+HWpE2m53CKMf00n/oDSUwywruP4k7ez/DF47Vmi2lDiT/K+M8oH9Ft8LP4JrqJYXcsSl9eOdJUfj3qv8K5ndkxtL/a00n3nNIvMlfzS28R+IFdL7EkiEriSOch8/e6usP8Aam4msgZ/T+qu2yw/CkPilUA8XFzxk9K+8QCSnJyAMEivgCSDyNe+AE5Ixyx16issOgVqXnBUcgHn4YqJu1Vq9zROwmsLkyvhfmQvyVHV+tKUGlf+6U6flUvZCQTgefWqg/SNaokwNBaX0jHWEm8XR2UtOeZRHQAPlxP/AN01YdkaMV95p4iNM5P/AG6qLu03Z0jy3idF5+j83XXeNMXDaa93BTkmxcU60hxY/wCBqV9a0nxPA4rjx/yqvAVbZ10HCeLn+l/nxrktpqZfdiddaW17bZKlBkNSVpHQgpAfYV8UqIP6qh4iuo9mv9v1HZ7fqC0yA9BucZuXGc5HiaWnKeXTxwR4Yx1FWX2g2dtNXC4U4/Dlzn/kOPz4pG0mopQaGqG69mPkeCyav1tovQ8BifrLVdssrM1xbbC5i3EhxwJSVAcKDzwRyznryphL3s2SuRUiNq1VyWOqLfYrlJUfmmNw/jW/vDomFuNoebpeY+Y6llD8SUE8S4kps5afSRzBSeuOfCpQAJqo1vvV1femac1XPmWzUtsWUTWmZJU09y5PtpzwKQrqCggGmFgs9BXUhmfvF7eIBxpyPDOOR8fNQu01+rrEQ6NmWHn0VkZWv9mpTpS7E1G0Dn35GkLgED4lLKj+FeX0bSa0gKg6V1VYpUxz3BGQ4WpSVEcvqHUod/uY9arvEVqtl0GDNXcEqBP1DpS6jn0U2rBz/VKx6itSbrB6Upy231tu4JUQFxroz3y2yOnApf1jJ/WbUkjzNTxskG8Oxc4Y6Oz9MBVSDb+TO7NCCD0Wv+R77s1ue0yww4hLUxyda21e6hSlAiRCWo/muIyU+Sg161KvaQ1NCv0fbK/WWeJFjuE1cppwYAKzHXwhQ/NWMrSpJ5pUFJ6g1B2udTXO8WBdnjzJdxiw8OQ482Sp6VCWk5SqNJV76gnmO7cJ90nhPk2JW4N2u+3aW3IpfZi3Zi5vtoVwmHNGUOOBPg2+lZKgOQcHhxDNpZan1roauQgvZlpPUEEZI5H9c8cr110p5aeaGnd+HIOHNrv2KkV6eyq7JXkEBlwj8B/Gl3bCfnefR6EZUUt3BSgB+apDSf3kfdUW/lps3GOtp8LakMLLagftJPAQfw/CpU7M8T+UG7cnUaozi49tji3x1+HeA8TuPQEgeucUxrqc0dJLM/gGEep0Vd2VoHuusQH5dVdtpCVNJ4gMkcuXz/iK151tYktlK20knxAqLta9pbSWh9bTdETdI6rmyLdHYlSJltiIkspQ6kEKKeMKCRnBODzBwB4uzQ28e2u4ajG0pquDLk496OpSmn0/q904Erz8qyySz18ETakxHdIzka8fJbmysgkcYw4ZBWpe9EwZjakqa6+OOdRNqzaaFJQ62YaHG1A5QpIKT8jyqybsdteUFPMZHrSROs7bucIzhOTRSXWekd3Cu5IGyjUKg+q+z9HalOSbMl+1vHmFRlHgPxQf4YqO77t5ra0cZcis3VhsdWj3bwH9Xof2c10IvmlGXypQbzjl0qOtQbfMSwoLZ4ck8wK0iz+0GvpQGF+R0KrVbs5Sz94tweqqbstv1f8AYzVPfNxpBt0lQFxtj6SnI/TR5KHw5/HmL5yG9A9pLQsS82K4NIlMoCrdc2wVOQlq5rYeQDlbKvzmyMgjiSUqTVYdfbJxLnGLgTh9rJad4ASgeRH5w/zy61HO2W4euOztqxc21IU7DWo+3WxSuJqS10K2ifT5jx6DDq7UkG0rv4paiI6tvEcn+B8+vzSVPmjhNBXjfgPXUt8lLuodH6m0pdJVm1HalxJsM5dSlYcaUhRwh1txPJxtYwQoY/RICgQNK0Xi5WC6Rb1Z57sKdCcS9HkNHC21joR/geR8eVWisGrNuu0doWNc7W824tocSFlAVKtzqxhbShy4m1ZwpvISr3VDC0pUmAdxtr9QaKujgcZQuJzUFtKK0BHEAHEEgKU0SQASApJyFpSccVdoLj7450E7Ozlbo5h5eXULPNpNkZrO73yidvRHUEclHHaB29tOu7VO3w0FAbiTmCg6uskdsJQytXIT46R/QrOONP5ij5EVXAqyeIf5+FW00rqS4aMvrF8tgaUWwUPxn08bEphQ4XGXUdFIUklJHhkEYIBEOb8bY2vRl4jar0SlbmjdSZet4UriXb3x/OwnDn7SDzSfzkFJ65A1HZS9EEW+d2f5T+h8uXgk6KtbcYt46PaNfHx/dRaSc5HX91SLs3v1fNn13S2/kSDftPX5tLVxtM5tK2XuE+6RnIB8OYPgeoBEcnOeH7/8KkDZvYrWm+N6kWvSyIcaHb0B65XSe8WYUBo5wp1YBPPBwEgqODgciRaLzDRT0j21+Oz558OGD18tVJ0T5o5QYPi/RWDu0TQ+5e1No3p290DHsF0tNzXAvsGwI4kN24oc+tejpTlRSUo+sGCAv3s8iGe3LZcbD0V9DrTg40OIOUqSfEef7+tO/UOr9O7Yaek7I9nqVFRbnYqWtUayQlLsm6Pra4XGYy8kpTgqGBgJ+OVKYlujMRYrEKG2UNMoCEJKuYSBjmfE+v8A9qyWOHs43DJ3N47m8SXBh1G9nPA8MnOOKh9rTTSTsfHjfAAcBwyla2t3y83eNpzS+nZt6vMtlchmLHA4UISrh711aiAhHH7pJ8j505tTaH2n0pcUWPcXtCPx9RpQ0gM2W2IVEt0gpSVGQU8XfI+0jmU9QemaS9C3C3aYte4+5U2QVM2exjTEaOwMrdlz0lSVqV+alCkgHHMdRzqucaM4pRcdK3XHTxKUolRJPjz5kn8TirBY9n5bxLJiTso4wMkAEkkZOrsgADHAa51KeUVLR22jjnli35ZNe9yHD6qxSdS7LbQ3OPqxGt5u5eo7a8h2zxI7AhQ4royQ873ZKVKHFyBVy/RPhFtrj6/3X1PNuDdsuN8vVxkKffSwhTvd8aiQFKJAbQM4BUQAB1wKmTaDskSrpDa1LuiJNmgLAXHtbaQibKQRkLXn/g7ZGcZBcPgEj3qdG6XaD0Js1af9zvauxW9y5MILSY8QD2aLj899ZOXXOfNSypR8SOVIQ7QUtmr3U1hBq6vG6XnAYwemgHXHEq0R2V9fTB1UBBBx3RxJUfXDaLTO3FmN43g19EtznjbbRwypKz+iHT9WVDzSlxI/SNRs7rezXeUqzbWbbR0HiObjfFma8lJ6LUlX1KD44CPgK0bZpfWW7eoHL/qCfImLkLKjKX0weqWUHkAPPGPLPjYjbzY+LamW09zhIAPD4k+ZP5x9Tz/gvc9oZqVpfc6gyy/yg7rG+AA4+qc0Vqgz/pYt1vU6uPz4KH7LtJc79LRJ1DIkXV0HIS5lDKR4YbGB8unpUvaX2OZ+qa9nShA/MS2EJ+4f/ept05oCPHUjiZGRjoKkC32CPHKOFlI4fMVm102vqagkB2is1NaoohqFH2kNooNvS2oxm+WMculSba9PxrfwJQhIPEfzaVmGmm0gADAGcjlSLrDcLRGg4Qm6w1HAtTRI4RIeAWv4IzxK+QNVF89VcZN1uXE8hxUkGx07cnQJytoaCDgcgnwHWqr9rO7vM7k6HZTkNtRrr73mVRgP3cX4073e2Dt3IvVttFr05qyfDuc+PbWrom3hqIXHV8KSC6sLXz8OEcgcVHPbNRPbjWHW8ZOUWaYgvcPi0sFDn4K8+nwqybP2qpoLpEKtm72gcBnqRj9VBX9zLja544DkgKKXrwoy+8QBn2ZwnH+sa/hTj7P825TN/oSoa+AJsMvv+eAW++axnPL7XD15VEzl+4bglCVcSFMueP2hxtkfuFe9Ka6et0nVdwhFbEF2Cm1S7ghYQpMUq4nWmjn3nHlJCARnCApXhkaCbW+WCSMD4m4+eix7Z8e5VzamQaNCk7tAa2VuXrRuDpeb7SnD1tsqkjLfc5Al3A9ThRT3bZ/QSo4BIw29M7haT2/26uOiJOmWk32O7xM3lUlaJMJxtR4glIODkjPCQeLPjnkhbf3S+SEXDVigixrn4jtS0oCpTcZIIQxDQRwtJAAy6eY4U8IJGaUrVcrDZn3TbILDD8g5dlun2ic6o9VGQvK0k9fcKQfKnHZR0UQowMsYBoOvPJ4ffXJV5odsoLGZZ3xiWWTOc8GjkB+qs5pHfPa9vT0GfL0xepk5+M25KTa9MyVtofKAVhKi0lJAP6JI8qcMDtJbRwUl+dZ9R2xA6qmadlJCfiUtqA++qzuy9U3ltDZXIjsOJJEic8pIxy6g5Wr5A/GtedNt2k7U5Oul8lvqQCA0hwthxfgkJSeIk9Bz8eYxVSfYaGd+C12848A7+yi49v6zeDWxjXlxKuvo3tAbP61vkfS2mNYR5V1lJKmogiyEuEAEknLY4QADzJA9akBT4B4UHh8qrF2TtBXmIZ+6etGUpvV5ZTFitEf8Chp+y0MdCcAr8c4zzBqxTr5JJ5kEeWM+fT41Q79SUlHVmCjJLW8dc688Y+S1K2zVE9M2WpG648uiae+W6TWz22N512osqmRWgzbWnkcaHprgIaSpPiAcrI8UoVVcPo59ezLqxrnS93mLkSXJjF8Q44rKlrc4kSFH1JSyT65NR723tyrluVupbtnrFJzB024pEoBZCVXBzHeqX4fVJAR5JIcPiaSexTLd0Z2kl6W7/LV0hzIIKj/Od2nvkHHmQ1y/rVpNHs4yn2UmjkH4sjd/xAGo+ihppaiWsbWBv4LXbgPLOF0pLiRn4nlQXFcyOYx06ViQsEJPFzJ8a9qUACfe59axocVaF7KxgjJHSk69IDlonIKuS47w/uGtnvQVDI5fHNa9ySkwZGOhaWD/AGTXQGXBAXOHt/N//td09LUPek6Rtzh9cLeH8KizcdfHarYsHmq0tj8c/wAaln6QE53J0aoDroyCT/7eR/hUR6uCnLXZ1L5/72t9fgmvqDZY5s1K7+kLNLlpVy+ais0UGirSopFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEVI20ykpuEMn/1iR/0FRzT82ufCbtDZ85Kh/abKaTl+AruP4lJHZNcTH7W2kUucs3qQ3j1U26n+NdM7SlIhkE80vPD7nFc65cbGzBZe1do99SilI1fGbJ9HHwk/gqupFoBS1JbXzU1LkIH/ALRR/jWG+1JpNZTu6s/VXbZc5hkHilFsJynJGfjQXcDljlyrEThQGOvn4GsSlFKRlR6+I5VmHAaqzrYLwGTx1zt7cd/d1l2ibXoyM+Fs2ODDtqADkJdeUXnPmO+AP9UeVdBSXHVpQB76lBKeXiTyrlxdL5G1t2gdWaxacL8RM6bKjLJzlAUUMq/2DWjezmmArJqxw/22YHmUxqac19ZTULPzvGfIcUoattP8qbNNt3CO9SoyIeP00j7HwKeXxAqcewZuiLhY7jtFd5YRMtJXOtaFDmtgqy+geqVHjwPBaz0BqH0OKb4XGyQpJ4gR4GmhMus/arcmzblaXR3YRIEoNhWE94nk80SPzVJJ+SxV7qKJt6oJbdJxOrfBw/dXX2iWb3J0V7phozDXj+nkfRdPX2u+Ckk5ByD/ABFQBvnsxB1Epu8RlKgz4hUY1yYby7Hyc8CwP51oknKSSRklPUhU26a1NZtX6cturNOyw/brtFRJjrxzCVDmhXkpBylQ8CkitqbCbnNlt1IIUMHI8/8A7/dWLUVZU2Wqy07rm5BH3B8FT6mmhuEG68bzXKgEu/ah0We51nbFR0hwtN3GPlyI8rHMBzqhWOqVYOMcq2G9axbmwy1IXHuMRgHuWX20vIbz/wAWSCW/2CDVv79tFaZ6XnGmw2p4cLiOEKQ4nnhK0EFKh6EEfdUO3nsvafMx56Ppu2EuEqy2HGMfANLSPwrQKTaK21IzMCx3hw+vD5rN67YLDy+ldjPJV+vEuyEreYbkRVhXEFIcK0I/ZcyT/bFMV3VCYN89sihqWqUCzLZQDwym1ZCgtI/OI8QT55zVqW+ybaJZw5aIqMHmVPPuY+S3CPwp02bs32OxNgMxmGlJHIsMJbPzIHEfvqZj2rtdM3ukvJ8h+64o9iaiP/ccqdqsl3Sq2r0uFyLbdXvZrfJcUMx3XAPqHD4OJIPD04gMgDoLwdm7biPonTjQ7jCwOEuFOFK581H4kk/OoC3T21ue3k6VdLXAM+1z+dytnNKZAB4uNBT0cBycgZ6kE5IqZNhe0XYplni6a1fdUrZdcRFtt/fcSO+UrITHm9O6kDGA6fddAJGFJVlvtLLPerWJKDvNzkjn/wCR/friTskcVmrHQ1Yw48Hdf7pg9pND6Nx9R3KDKdhzHNIe1R32HChxt+LNbcC0qHMHuypOR4GoCG915uzYa3BtMPVhVwj8oSMs3NCQeiZiMOL6cu97wDy8KlTtQa8sbu4upIMSaHZEHTwswCOafa3pKVvcxy91lOD+sceFQVpXbPWerLbIv9utaG7REd7l65TZLMOKh08wjv3lJSV4wSkHIB6VbdlqONlqjNc3HdbjkRpy59E2vMxdVHsTnUq2u0faS1BboKbo/fZ2stGREI/KiJyR+XbC2SEBxeCfao6SUjvM5wQCGzwpq28OXAuECLcrfNZlQpjCZEaQyribeaWApK0nxBBBz6jpXNra3TsXbu+zdZ3nU9int2i1z1O2i2TkTlykLjra4HFtgtIaKnUpJUsnmAASRVw+ynNmvbLaZgTHnFqjRSpPEc8KFrUtKfQBKh/kVQNvLPRQR+90p1DsHTGcjPz8easNgq6h5MUvDClx+M26SClJGcfCm1dLQl1JCEjGOfL0p4OI+rwnI4uXrWo6wVE8j5YNZpHKWahWctyopu+nQpBBSD4dKiTcPam33uGtYjfWoPEgo5LQr9JCh0P4EZB61ZS5W8OAkJI+VNO62kJSrKfxqft10kpnh8bsEJnUUzJRhyo3p/UOuti9Zp1Dpd8tOlXC6zgiNcGhnIWnoF4zkdQckZBzV9Nqd4Nvu0RotTGAxNawZltWtPtEN7hKVLRxAhWUqWnJBSpKlJWkpJBiLXu2lr1BBeWuIlal83EnlxEdDkcwfIjpVZrrbNZ7SaqZ1Tpa4SIUuM6AzLSeEEH+ieA5HyB6H06Vc6impdr42ysd2VW34XdfB3UH5j6KHa99tzHI3fhdoQeXiFYreHZ++aLuBm29gv2t0uLbcabUGylKePKRk8Kkp4iptXNPCSkrR7yWZZmrLfLLdNvNXrUiyX4J4XgecGanPs8pP9VRwvzQpVWD2A7Q2k99LG7pbU8ZiHqJCEpm2xaigPcJBDrKhhSVBQyCghSCAQelR3v1s7P0a85qKyj2nT7jiU+0jkqM4tQCW30gcKCVEJS4MIX7vJCvdMZb7jUMqf4dXDs6hn16Fp5qlXzZs29zbrae9HzHQc1R3Ulguulr/cNM3qP3M+1vrjPoPTiScZHmD1B8cg1K/Z93G0dZ7BrLa3Xl2lWi16wjtpRcmGwsx3W1ZHEnI4gR4EgYykkcWRs76WI3jTts3EUyBPglFmvBxhTiUpPszyh4HhSWyfEoT4moLebS8QlPvHGcjw/+9bLEI9pbWGSHByM44hzTxHrqm9LUiB4njGWkfQ8VaTUG3V/0VEj36KtOpNEyoTMiPqG1W9TUePxLKO7fQFENg8iMEjnnn1HzTGmdUa21WxovQ1sblXBa21y5UhCvZLbHUTh59QHLocJ5k4+VMLYveK1afg37bDc273U6Q1FbjAS4youqtzneBYW2hWQATkkADJA+NTfpPQWkNbaP1HofZPUd5v8AJvUuGu+aqukIxIsWO0SW2yhOQVfzmEpJUsqOeFIqh3Jk1qe9tcOBGH7p3CD+Z3IFuuRzOMDVLN2co6+pZUQ/BjVudc/rlNPXEeLqCDYeztsiwrUqmXzOvdzhwu6NzuAUeJwqJ5ttg/zizgJ4emDmd9p+z1ofZG3Na01pKt9z1BAb9oXNfSDCtKxz+o4v5xxOP55Q5H7CR1LjitbLdlbb95ZkIj8bQRImvJT+ULqtOSE458KM/ZaT7qfHiIzVNN2t6tc7+338lx2n4VhSriiWplXDxIzycfPTy+ePHnUHT1lx2midb7Y50NECTJKdHSk8SPDwGgCt4oae2vFVWND5sANYODQOCfW+/a0uOrHZGkdrX3otrWpTMi7kHvpJOQUtY5468+pz4daYe2mx8y8rRcL5HcQwVcfcqIUpfjlw/nHr7vQepp7bV7FJiuM3S6tpdfSkcGU4CPRA/NAx15k+gOKshpvSrEdttCGkAAYA4R19MdKVqbrQ7O038PtA3RzdzPjn/PBOI6We4SdtVnyHIJv6K28gW5toMR0jAAzj8P8AP/2lOz2IR+D6vy6it2z2VLASooBUB0H/AH04kRUIR7oGeXhWc1lfJUPLnHOVPwwNiaAAscWKhopGQOY5Ac+lbaykYAOR6jw/zj8axrURjiITj99I+o7su0W1+6NLwuMguJP6JHj0P+fTNMY2mWQN6pR53WkqE97+0TcLdLuWm9C3aFaoFjdEPUGqZLfftw5BBPskRr+nkgAkp6JKeZABUKc3ne1Au8u56asqpVzfc7xWoNSFNyuTjni6A5lplR5dEqUnGAo9aVtQRGdc7eWO3q1PbLTc7Dd7vHmtXN7uG5r7ryXe+S8QUhzhKUKDhSMITgnnhiXbajWcK0yr5DiW+7QICOOW/ablGniOnpxuBhxRQnOOZAFfR2z9ltNthEZxnh0yfE889M4Wa3KrrKl5J4f5yUk6MvGqdQp0nqK/32bd7hedYSHnH5j6nXMW6IytocSiSE8UpfLpyHlVv9z7K7q3bxcFxrvOOPwrSR9oFGFj5gn76optbre1WlnTEe6TgyLHqN6WSpJIMeYy208vl+gWEHHjxnyq3u52/tp07YZNh0m/Bn3VEZJlS3F8cS2NrSSFuHotwg5Q0D7x5qGAQaltbQVctygFLHqCSDwA73Xywpm21lPBQudO7kM/JU1OmNTM32XpkkMfkxoiRPdVhuPEUQUuKPnw4AHUkgAZrSuuoLI4/FsNtY9nssA94w3IUU+1uHkp50gcycDkOWAEg8s1J2hdHXfdu6LhxmZLFjW+ZT7jx+vub3QvvHHJJ54SOQHujIJUZ/V2ZLG9FbbeQHUJTwlDraVoHwChgfLyqardpKO2SNiqTl2NQOR6+fhyVepdnZbjvTt7rSe6P1Kq9b7u3JjKXMeekhQASGVJaQkAchyCiQPinw8uS/ZtVW2xr9rt8SPCfSjhMhCT33MYP1pysfAHFS7M7H9jK1KhwI6QTnCVOt/glYFfYPY+tbiwmVAjnBzlbrzn4KUah5doLPKCTIfLH902fsNVOO7vaKHDuQblJ/J+n4Uu6TCOTUVGeHnjKlA4SMnqenrUr7MbE3bVF3i6n1kpuS+yeNlhA4o8LHIFPEMOPePFzCeWCVc0zRors46dsTKQqMylrkO6aZS22ojxKU9T8edS5abJFszCWYzaUJAxgDyqrXXaqBkZhtrd3Om9z8vD0VosmxtNbXiWTvOXy225izQWYbCQG2kJSAM+AwPWmfvfurG2e23u2t3HmhOjpEa2IWebs1YPdhI8eHClkeSDT8WvvFBIyVDwHXnXPXtf7iO7u7vRNtdMzO/tGmHHIZW2vLS5ZI9pdGORSjhCM+SCR1qJ2Ts/8buIMv8Ats7z/Tl6lWW4zvjjEMA77zutHmoz0DEmPJm6zurzz8+5OuIadeUVrcKjl11RPMqJIGfElXjWxGvsrQW9WlNctOcAiToNwUQeqEOhLiT6EJUD6KpyoahxWmoUBBTFiNhpniHNSR0UfUnJPqaZ+6cYLtcGekDjjvKaJx0SRlP4hVbRDM2rqXNd8LgW48MYV22g2ZbbNkRBH8UW68nqc95dcAEJQlLbiXGwfdWOihjkRXokFIA5E0wNitTM6r2f0df23y97VZ44cUcE962nu3AT5hbax8qfvEnhJBxjpyFfN9XTupaiSBw1a4j6qrQydrE2TqAgEJyela09Y9jf8CGnD/dNZ3efME8seFaN0Kk26UrPMMuf7JpAcQlFzw+kAcCtz9KIT/R6PhJI8vrnz/Gok1llFstIJACba2T8OVSl2/1cG+MKDn/gem4DOPLm4r/tVGO44DcOHHSMlNqRy/aV/hX09sw3dtFK3+kLNLmf9ZL5qKKKDRVpUUiiiihCKKKKEIooooQiiiihCKKKKEIooooQigePwooFCECnntm401eYrjnINzY6lH9Xiwf30zB60uaUfW1Kc4FHiCONI9UnNcSfCV004KfVrfXY+0JZJnFw+yakt8jPTH1rS8/jXWR4JhXW5R0nAEoqH7SEq/7Vckdzgqza9h3xoDMliNOSR5pJT/8ALrrBBuSb0mPfEfzd0gxpgx48aevzHCfnWN+1GEuFNL4EfYq47MPwZWeSU21cY97AyepNeXDgcgPSsDi+A8sHnmvpcSoDoVKzisjI0VtTV3Y1azoXbTVOr1kBy2WiU7HOf/SC2UM4/wCdUg/KuZe1kIpttyuawcvPIjoUfHHvK/Hh++rqdufVKrDsbItTYAcvtxjQuuCUJ4nVf9GkH+sPOqmaPiJgaTtkXh4VOpXKX6lZwP7qB99bDsVT+62N83OV+PRqfbI0ortpWvPwwsJ9ToEtAYR08KTr1amr9aJVneSCpz6yOojml4fZPwIykj1B8K3yTjrWNRycg8xzBqfjcY3BzTqFtNdTRVsD6aYZa4EEeak7sKbtdxInbL6gdUl0rcm2dSieSxkvR/TIAcT6pWPzhV0WG+AAK4cgcwB4+PlXLXU7l00Xqe1bl6ZeMaXHkodU4jkESUHOSB4KAzjxPEOldHdttx7Vunoq1a2syEMt3Bol5hCiRGfTydZyeeEq6eaSk+NUfby099t2px3ZNHeDv7r57jpZbLWS2ip4sOWnq08FrbyXTcC1aMfn7YuwfyzEWH0sy4zbyJSACFMZXzSpWRwkHmoAZ51Ceit990rpZ4Wqb5ads7hAmFTS0/l52yPsSEHC2XVS1KZQ6P0epThQykg1Zl6K1KaLL6AtCxzB8v8A7VWferaG5Wa5StVaMaiBNwCRdrfKZLkC5pB5e0ND7K+ZAcRhQ5EEKHHUbs1WUD2+41kbc50cR9CfseX2Z3KCpa7t4HHxGU7X9/rwgBDe08N5fiqDuNZpSflwJrWVv7eHElCtq47KhkZma/tEdP3rFUe1VpG0JkvuWv2iwykjvFWi6ryk5P2Y8oDgcAHMcZQSOhUaZKipCihJzg4GM1p9LsLa6tu9E0eXe/8AyVXnvtXT91+Va/dvtA3xaRDXZtCxSpRPCdRfllXy9iwlJ+JqukrW14Zvky5Q0wOCenu5UZlnhYeA80K6nocnJ5U2QQr3lKHx/wC6nTftrNwdMWazaj1BpSbEtl/hmdbZSkhbb8cEZWCknGOJOQoA+8nzq1UNht9pb2YwC70z8yeChauvnrxh+oCajslTryyUFvJJCVdPWpd3pjSJ+3+31zszmbRZbFGts6IysqREnucUlTix4F9DqV8R6qStOT3fKLkSWSjun2GpCOnCvI+4ggj/AD16U7NFaylWa7iWL4203KZRCmw7sx38CXHSkJS093ac44UgBQRxJIBBSQFB5cKaQGOWI53M6dcjh+yRppI8OadCcBe9CuvI0lOt7BIk6quEWy8iM+ypUH3h581pjfcRXRjaO3t2zT0RiO2lDfdpw0kcgMDkB6VR21aOtNyucTUeiX2LX7K+5KZst2nIVBecXyUiNcE/VZAxhMgoUMAcSj1tDt1v7o7TqhpfcmNcNF3NlKUhF1YIjucvzHk5QU+IUrAI6Gst24gqbgxvYMLtSSOY5cPADiFcbFJHTktecdP/ACrFtAFPvHhR/Wr5ISAkhPPmRg+XhSVatS2a9Rm5llu0SWwtIU25HeS4gj0UCQfkaUW32neLKxxcgCPKsekjfCd2QYI5K3BzXDIK11RwtKlnAJHMkf5xSHc4CXDxJSTnyGcUvvFCklRBI6H41gLSTniIOfIdK6ZJu6heEZCYFxtZHGAnkefKo+1jt3BvcV0rYaUtaCFJUkFKvjnlU3y4QIUAEhI5dOdNy4wBzCE8jUpSVz4XAxnCQlha8YcqG6/0RetB3hF9sEiRDXDcDjElri72Lg8snqpv1PMdDnqbT9mLtTWzcBA2/wB0HI8W/OsqipdeAMe5tKHCpCgfdypJwUnkfHkOSpqbRMG/R1tyWkhXCcKA5jlVWN29oJenn1z7Gwsd2oOhpk4U2oc+Nn7iSj48PkL6JaHa2mbR13dlHwSDiDy1/Tn5qDMU1seZYxvMPFp4FWY3p2EbgR7tN0c2+/pe4xTHu9tSQ69BQSD3rQ5KdabcDawACtsI97iQCpPPyTa5VouEu1zm1CVFfWw4g9QpJIPT4H486uR2bu185avYdFbwSiccLdvvvF1Hgl4+Y6Z+R8zN07sv7H6r141u8rTwdWpHtDtrjKSLTcJBOUynEjmPMtoIQo4JAAUFOLNtNNsVLJSX4aYy14+F+OHqeajaiwQ3FgktugJyW9M8VVTs9dje9brR4uudaIlWrRneZQpBCZNzSkjiTGCs4QehfIKQRhIWelm9zd9dr+zjo6LofTlkhiawlSbbYoK+TZIGXXT9ok8srVlSiOflTR7QHbOh6MRJ2/26fj3C9BruHp6cezW9AHDwpAwnIHIADA+BqqmktudV7kXteobxJlSBNd7x+U8o97J8Tj9FB+R8sdaSngq9r5P4ltATFSA5ZHnBd0LvP5/dOoBDZm+7UXfm5u6eSw3CXuDvhq96936SqSsuEAlJMaIM/YQnxUOX7z5VYra/ZuBY47a34vE4cKUpZytavNR8f3U79vtr7dZYTLCIbaA0kJSEpwE+n/f99Sxa7CEJSkpAxjHKoy9bTh7PdKRu5G3QAaaJ/RW3dd205y480lWewIZS2hTIz0wBjFPG02tDPVIzywPIVt263tN8IWklXPHKlVLTYAwMHHOqDPO6TUlTjWho0XxEdsNkAAHoP85r0fdHl/n/ALq+KwFcYUkDpWGRKJSEJIyetMxqdV1lfHODOcZV1JzypsawPtFsfiHklxtSTg+Y6GlOZd41vSXJMhDSOeVLOB0qH9yd+ttLRFehR9RJu1zyW24FrSqU8pR5Y90FKSPIqHSpa30NTVyjsYy70/VITzxwtJccKoW5drRFVrfSaIqUB4xtRxlHr3rClMvBJ/WQ+tZ/qCm72boUx/dK1Xtx5bFqtDyJV1dIPduRyoIVHV4fXFQaAPXj+NSBq2x3LUuo0X3VqJOnG1trZRaGGxKvMpl5OVKLKfcjJ4c83lIPPICudMfXGpY8Bi12nS8uDYrPbH0SWrXbnxNfckN9JMuQAG3XSfBJ4UZIShI6/QNvM01GKQfE4DJ4gHGPU6fPis9q3RiYSngPrqoyuCRAmSYykcIbeWjhBzjBPL/D/OFZGrbo8zBhkR2oEFwOtxFNBTS1+K1pPJZ9Vc60XpbchTjvcfWOrUtbizxLWpRyfQc/AAV4gQZtwccTCiPPqbTxK4EE8I8zgchVtdTs7NpqCNAoFzg6Qlg5qf8AbPtA6msMtuNAToVQVjiRKckQFqx/yqj3SfwAqxVs7Q2onoqJDm3unpmQOcHcS0u5+CT7w+dU63H2E1NtHpaxah1hfLCiffj3jVmjzkvTGWSCUurCcpKTgj3VKwcA4JIEbFZKuZwPWqfLsnaL0DUQbrhk64PLjqCFONvNXR4jeD810Y/8IW/BzP8AuTIQPN7W9qaT/aVyrQuu+O496StnS2ntubOIzK5Eh2562ZufdNIGVLWmCR3YH6xGSQOZIFUPsdgiS1IkXe5eyxCoJDMdJelvgnGG2xjmOfvLIHLkT0NldqdlbzuL7LbE2o6c0ewtC3beCFPTnE/0ktwAFxXQ8AwhPQJBzmuXLZ2zWVpkeG6dQT6AF2p+nUqWpbhXVxDW5Hqp67Mm5G7+48q437Wj9qTpxtJjwGodv7lMo8Q+uCnCXAnA5cR58WcDlU/OJC04xjI6A+VI+nbBb9PWmPbLZHS0ywgIAwPAY/hW8t5DZPfOIQhIJUpRwlIHMknwA8ax+61bLhVF8DAxvANA4f3Kt0EZpo8PcSeZ8VE/af3cZ2b2tm3SLKSm+Xjit9maQvC0urSeN/0S2nn/AFi34GqCbc2r2K3SL/MQTNuZKGlq6oZz76x6qOR8Aqnfv7uW92id51otjriNMWMKiQRxZT3CFZdkHyU4vmB5FCfCtVS0LX9U2lppAShpsdEIAwlI+AxWy2W2ixWttOR+LJhz/wBApnYi2G63F12mHcjy1g8eZ9FnSAOZAz6fCkXWMQTdNT2+HiWylL6B6pPP8CaWEq9KOBt/jYeHuPoUyr4KBT/GncD+ykDxyK1m50ra+hlpnfmaR9NPqrO/R/6wbvWy8jTLqyZWnbs80nJ6R30pdR/f7+rNOKUDy6eNUG+j91E5Z9wtWaJeOBNgiSlJ6hxh3hPL+q6T8BV9ELChkgY6ZzWU7cUfut7lLeD8O+YWA2dx91Eb+LSQfMFZUrJTkgjPia0butSoLjIGVOkNjn+koD+NbKnQTggCsbiG31MJUAfr21YHiEnix/dqptxlSZ0C5u9vKY3de0zeY8VXF3EaDETjwPcpOPvV+NRrunJAnyWULCQwy1HHr44+5VOfemUb52qtQpeWHUx9SGMtSeYKY6whR+GGjTH3KeC7jPWlSSHJIScnySBn+7X1PZYuwoaePoxv1Cy+sfvzSO8SmGa+UHwoqcTBFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUpaefUzc2wOiwpJ+7P8ACk2s0J9UaWzIQMltaVY88GvCMjC9HFSPuOhyXYtOXpZyoMriqJ8gUkfiXK6WbFalb1hsdt/qOKDzsjUB/PXvYpMdRPxLJV865vXKOq77Zylk5XaJqHUjx7pagP3ufhVzvo99WJvOyd00rI953Tl5U60c8wy+lKuH099Lh+dZp7RqXt7OJhxjcPkdP2Vk2elMdZuj8wU/XjUmnNPtoe1LqC22pDhVwLnTGo6VY64LihnGR0qMtQ9qTYTTzzjMzdO1vLR/RwWJEwkeimm1Nn+3UXfSCaKevWkbRrO2sJ7+wOqafUCQpUZw8/jwrCcDHRaj0FVR0RpbSt4srU6bHlSHSoodSHwlIUPIBOehHjVW2f2StVxtzbhUvccnBaMDBVtpY7ldbibfRBoOM5dzCkztd9obSe+UrTGm9Am4qtdmXJeddmR0sd/Ie7tIKUhSiUpS0McWDlR5UktRDBabhE59mbSzn+qAn+FaEfTGmoDqH4NhjJWhQUlTiluEc/InH4VvEqHJwnmr958KuwZT01NHRUQIYzhnjqtO2O2aqdnXz1NwcHSPx8PAALIVDzrGtQHQ5xTBuW5s5IcRBs0fCSQVrWpSuvXAxisDOptxJKT7Nb+6SsciYqfwUsE/jTxtqqC0OeQ31XM/tGtG+YaYPkI47reafFwQzc4b9plZ7mUnhUQM8KuqVD4HFPDsX7oT9vdxZG02oZHd23UMhLbHeH3WZ4yG1JJ5DvAeAnxPdn82oJct+urio+03ZxIUclK5YCR8gcUnXm0z7A4xIdmNPPO5Ulba1FScfrEA/dTl9qgq6WShleCHj5Hqs62vu813dHcWUr4zGcFxxgtPIhdhFp7tJbKSkgkY6Y+XhWtJtLVwR3LyAoLGMEf91aOj7wu96atc990uPvQo6n1E5KllpJKj8SSfn8aWnwpTRQgnmMZzXzbJG6CYxE/CcZ9Uq0iRgd1VJe0fr/bDTkiVadI2VrUc1K1MSpKkf72Rn0nm2pZH16x4pTkDkeLORVP5AQ6+7J4UguKKihCAlCcnOAB0HgKuJvLoGHovU14Z1Pa3pWhtUyBKnmMkqetEwkgXBhOQMjJDieYUkq8SkprHrjQN729uLMW5KYlQprff225xHO8iz45OA60vxB6FJwpJyFAEEV9IbFyUDaVrYHHLhxJzk8/Ijp6rNr62pEpMnDyTu2y7P2sNY690Hpu/Wt6BbtahM6PJC0qK7elR75wBJJSeFKsBWM5T4Grhbu6xgWrb3W+p9Kx2RL1S+1tto6BHytLNtjqLbqmPPvXVPAY8GmznnVZ7H2vtR2bSkG2xtFWhWpbVYTpm3akK1B+LBwOFIb+zxJwPf5Hl45wbB6buG2txZ2q3Ln7p6ci6O20sCHjZlvH8oovPAONwtcOF4UlsgZyVoPIhWTDbQG4Oqop61m61uQANd7GvLhvd0fNLUDIBG5kLsk8c8v8AxqUztxOyjtRaWUGVqJrStt0HZoTGrLu207Jdl3SSQUt91xlOUpUhRKBz75CcciRFeuuxzurar5PiaBt/8sLTDgxrm1MhKQ24uNI4wj6hZCysFtYIRxdAehqetEwtT3LX87bjddUO7WHf+1PakitwnXC9bC2lS2lLJSAlSUt4BQVDiabzkZBZ3Zv/AJUaX3D3E3C1dq+deo+1tmmW9h6ZMceDikqcDSBxqOUhAewk8gpxOBnFcUd6uVHE9wmDi0AgEZDg7AGCNch+8DnwwupaOCV7RuYzpkcsa/ZVemW3Wm3Fy9iukC76fuCPf7p9tyO5/ZOCR++t+HulriIktm5MzIi/52HLiNPR3v67S0lCifPGauLoW86t3x2Elz9ylw7lfdfXuNpSwypFvYUY8dKyt1xHuYBQgyTxgBWW05PLIZu4eyvZ/n6Y3OtOgrPebPe9qUttPXKTc+/RcXghfeILZHukKQpJKcDODyFT8O01JUPNNXwd9pwd3UDUDOdDjJwmDrdPGO0pn5adRnQ/5zUFWnc3S8SWm4N6cummJ6cFMrTN1cYbSr9LuXePHwSpI8sVKOme0vri1Nd9aN2bfc0nH+92qrUttzP6klgrUfi4tI9KYSuzPrU7G/7vC7laW7RlSjCcccTLLXehpK0jgKFAqPQrBwM+NQ+txRBHU+h5f41Im0Wi8B4hcHbp3TkBwBHLXXTzSTbjW0hBk4nXocK81i7ac6KG4+stsZL6vz5NgntTEkfpJQCc/Dj+NStpDtI7Ma4WmLadcxYsxQyqHc0mG+hX6J7wBtRyRyQtVcw2HHWHApp1aFDoUqIx8xThgavuUNtTF6ix7zDWUlbNxa7xQA/Qd5OtfBChVdr/AGbW+dhdC3dPgcfQ5H1UpTbUTsdh5z5rrM5HcABWFBKkpcBx1SoZB+B5c6TpcVCweEHIODkdKrN2bN1n7DM0pp1mXMXpLW4ntQbbPd752yXCLw8QYePNcdwKSOE9FE+KSpVo5AHCSjBzz6k1jt9sslhqhCTvAjIP7q5W6ubcIu0AwU2JcAhJAByoeApo6h0lEurK0TGApBGOYzz8PxqSlMcSeI8+LlgdSfSoc3J3hYgzJOj9u2Id81EyMTXnXeGBZkZ4S5Kd6FQOMNpJJPI8/dPlqp6qtl3KYZI4ngB5lc19VT0cRfUOwFWLffbiFZZqXWCkSp6ssQkIKlzvewVpSOaVABRK/snhIJzzpvwd6d39G6Ef21OqXo9okqyr3iqS0nHNpC/I8sjoOeKc2rNWtWm6y49omzNR6nvRDcu6vD6+SMD6tlPRhgdAgdUpTnoEp2bb2ar/AKpsbl6uspxNzV9Y0Ug8DZ/RCT9odMk4PLl665HNS09JHHcyHtGMFwySeoHQdT/ZZ9T1dbdKh8lvBZHgjPVMzaXQVq3Bv6YinwiW2e8RAd+04nGe8Of5w8s8unj63b0NoRmxw2o3cc08iTjOcCqbWwfkK4I0jq+E9bLlFd76JNiOFl5tQyA9HeA4gM8ynpnBIyMi0O2++ki2vM2fdl+OYry0tW7VbKEtRZSvBuW2nIjPYwc54CQfD3jB7VwVdcN+ndvNAzu+HVvX7p/s/d6eCX3SsbuS+PNTXbrU0gD3cEY5AUustICRn3SkZ+NeG0NBCHGlIUlYC0qQeIKSehB8QRz8RzznnWRGQCrpyHMHwx/hWRzPdkh4wei0JpBG8OCzDiSSkE5zw9D18uX7qYett+Nr9vFKa1Rra3xJTaihUdKlPvhQ5EFtoKUPHqB6mmH2htx7m1crftlYr69Z0T7fPu99nwgDMbt0VouFmNnGHHOEjIIOOvu8VUXue5DiFyWdD2GLp2M8U5eSO/uCseKpax3iSeRw13Y9K0LZfYN14ibU1DtHagDp1JPD5Kt3K/iikMTBr1V1r92vlON8WjdsNRzoq0nup9xSiCw5nooFRIKfUkfCozvfaU1/I74XvcPR+k2wMhi2Mru8wjyTwBccH+s4k1Um4Xa7Xd9T91ucua6onK5DynFH5knNe7Ra7xqC4M2exWmZcZsjIbjxGFOuqwMnCUDPIDJ8q0ij2AtlE3ee0DHM6/U6fRVqfaKpmOGk/b7KT9T7t2TUjpdvd01pqdaObaZ05EKPz/5NsrUAfIOCkE7z6qgIXD0rFtemoC2+7XHtcbgLo/5V5ZU8754WtSeXIClSy9m7dy9zk29WmFWxblmXqIOXJ1EdCrelYSt8En3sE5IAKgMHFTLpnsWaSg3nS1o1/vBCnytVe2KiQ9NtF1LzTDTqypuQ9jjPG0pGO7HvcOCakJJ7Da8RlwceTRry6DRNMXCrOQMfRVivmqtRamSBdbvIkBOQlvIS2PPCE4SOngBWXSe3+t9a3GFa9NaXnT37iXEQwhHCh8toK1JQtRCVEJSTgHPLlVodM2Xsz2XQ+o9d2Xbm46gsLep4NguLl84lXC2wH4x4pDCWlEBYfQopOAr3kjl0qctOas3E1ppOFqSwRXLpYLbpUTIjHsyWprOobe/hrhZSO8BkNpSVIHu4dV06GNuG2IomZpIMDO73sAZxwx114HCWp7SZ/wDcfk8ccVVGZ2TLnpDS9r19e9V2PUERtmHeb5YrPKWLhHtD6x9cFLQARw5BVghJ/SANWJtmi7rsNddX6G2WmxUwb60dR6ZfdjNvuSXISQqbZ31OJJWVsALQD1SsEEcSyFHcK6adsErTW6VtZcesVtkpgXlh2Px99pS95UkLHiGXjIbxyAKSPEUuXmRobT0HT9jvuqJMDUc24NwbXcMpESJerTxNMuPOnmj2hj2dlwZKVtkZB6inVm0FfXsj94y4OJGANMg5wR4Yznm0qZgt8MDnGPAxjUph7zaT2x17s8vdC4aTuV/NqsRl2D8nzktKiw318Km3cDiUIT5cSE88I4AoH3qoDNt0tlLZlMOsd42HW+NBSFIV9lSQeqT4EcutXO3K3+0Ns6i2wNqbhG1HJevk68z7ZJbQ7CgxZrKRItpUAUuI71AXy5ZScjrxVv15uLuR2iNeRJV0hNz7xIQi32+FboaWw20knhaQlIyQnJwVHkPQcrfsayspoXdu09gSSC44wOm7y65UXeDFI4GM9/mBz8cp2dn/AF9oC0XJi0630qpmOgDjusFkvd0hPMuPtgFZSOpKc454SeldFNKQbD+SYc3TciLJt0llLsZ+MoKbdbUMhSSOoPrzHTwqh+3+3MW4SBtFpVcaXPncB1hqGM5xttICuIW2K4BhSc8JdcGQpQABKU+/fnTNljacsUCxQGw3FgMIYZbSMBKUjAGPgKzz2hSUhnDoCd48RnTHI45Z6dNVZ9ne37L8QadfFK4JSkpOemOXh8ard24N5XNvdBo0NYZRavurWlocI5OMW/JDigfDvVAtj9XvMdRViZEhDDanFdEji6+Vcv8AtIz7jrLtE6rbMwKKpqWWC6s8LTSGkhKR1wMDGB+81G7A2mK43Iyz/DEN7HInl8uKeXqWVsLYYBkvOPFa+kLS9puzpZeSpMy4hLz5V1Sj8xGfgcn4jypb7wHoR/hUcHT+t4X/AAK7KIzy7mWef4j91bCLtuJBb4HIq5A/SMdDp/tYz+Na9PRe8vMokBJV5tW07bLSsopqKVjWDoD9uqkJKjWXnjiJPLnUcRdwr3GcUzOtDHEg++FhaFD7zyp7Wa5i729qaGu775JPDnPQkdflTOooZqYZfwVssm1ttvrzDTE741IIwtbbjXzGzG/0HWk9L67e0657YiOgLWqPIaUhfCkkAkBeQCRzTV3bZ2xuzxd0Nob3DFvcVgd3cbbKbIJ81IbWgfEqxVMJdmslzc7242liS4EhIcUVhWPAclAdPSkifo3RrTDsj2OTHDaFOKKJHIAAk9Umo+62i138sfWBwe0YBafkqJX7G3einnqKKRhjcS7B0IHFdL9Obi6A1WW2dO66sF0fdz3bMS4suOKwM8kBXFn0xSu5PRbHzdZaj7Ja2X7g+T0DTSCVfgTXOfsV6Ik6k3jZ1N7Opy36eBkZJ6vqyGxyxkj3lfsgeNXd7Smojpfs+a3vMZ4NPvW78mtLHh36g2oDx5hz99ZvednYLZeIbdSvL97dzkcMnh8lVaavmnpHzzADGeHPC5s6ClSdVbizdRXd3vJMgyprzh/OdcJ4ifm5TW1lKL76cqP1jjjoHoSade3Eb2XTGptQOApV3TcJhfQcS+Iqx8Pqz+1TF1K+HbmptJ5MgNgeWK+hY2Brt0DACz55yN7rqkqiiinKRRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFfU5zyr5X1JwaEKX9s5DV0hSrNJAcTdbe5G4P18FAV8UlQV+zUt/R56xcsm7V42+nvJZb1JbnEtoVyHtcbKwPj3ffj44quu3l2ct10YdbWQY76Vdf6NXurH3E083bm9s/v9adZQnuNiPcY94aUnkFNLVl1HwBLjZ8+E1BXugFxoZ6Mj42nHmNR9U/o5zDMyYciF0m3M0pE1bo676buDaVolsLRzGcEg8xXMnSqZukNZ3LRN1HCrv1MYIxl1J90j4gn7xXWG5OMTVImxsLjzEJfaUOYKVgKBHmCCD8K59dtfbF/SWsIu4dnZUiNPUA+tPLgdBJSrl58xn4Vk2wFwDZZbVOdH8P8AkFoMtVJbKmC6xamMjOObTxCQcgA5x1OMVpTnghHGDgjBrHZ7ui92mPdAAFPJ+sA/NWPtD78n4EVr3d0COog+FXlkZbLunkVutTXxz2/3qA5a5uR5EJj6G0ozqyZc4qriIUhlouMLWfqyvjwEq8gc9fA4zyrzPj33SNwVaL1BcjutcuBZ5FPgUq6YPXI5GlHa1IVLvBJ/oAP/AHqaktuVbbjERadV2tN1tyU8KErVwvR88stOfmkeRBSfEdCJiqruxnMbxvNwPMLK7Hsn/FLOyuon9nPvO15OG9wIUcwJzcpOUq5jHEDyI+I/jSHr5eUwMHolf76eGodqbjGS/ddDTl3qCwFLW1whM2OjPVTfPI6e8kkZPPmajq7uy5zbLT4CVx+JKgscKySfEHlS9IyJ8gmhdkKE2krq6koH265wljyRh3FrsHkV1O2ffU9o618XX2OP/wBGmpBCkhJHKow2IutrvWg7ZPs1yjTY6I7TC3I7qVhDqW08SF4PurHilWDjnjBqTM8xzJzivm27xuirZQ8Ed4/dd0rg+FpB5BNrV+kbfqe3uQpjCFpWCASnJyenx+H38qqNrva297fsv2FFmb1HoyTJXIe0+8soUw6rkXoLuCWHuH3eXJXLKVj3Rd84JAJ5ZxSPqDTNvvkRcSXGQ4hwFJChnIp/ZNoJrS/A1ZzH7ePikK2hjq2ne4rmZetnRdxKu+0Fwf1JBjpLsm1Otd1ebcnJ4kvRgSXOHmC41lPLJCT7ojNQUVlCwRw8sYxjwPrmr+7j9nREqZ/KCyLlQ7pHIXHnxHSzKaXywoLT1Ix0PwyM5qGtU6fvTsl2PuvogauQ4SlN4tPDBvqP11DCmpJ8woKUf0hW22XbKKsaGvO94EjeH6H6HzVIrrC+ElzdPt/ZMzSva03n0ZpFrSFpu0JxqPHXDiSpMUOSY7SgkFKHM9AEpwDn7KeuBjLa91dMWns13vbyI5MXq7VV/RJu8hxvCFQ0ELQkOBRye8QknIH21dRk0ip2gseppTjO3W4dsmySrgatF7xargpWcd2A6Sy4rPL3XMnypl6v0Xq7QlwNp1dpq5WiYgA93LjKaJHmkkYUP1gSPI1Mx0NmqjmEBj94OI+EkjUZB4jy0Ua+atg7r9Rgj0Kvpt1dtFJ1Ttrtxo7UMK7J0DomZfI7cKQ26mXeHuFDwHCSCrLslYT1SFDwqtsTsq7sapventRaxkKt0vXV9falwZCHW7m21xLcfkqaWkZQEpUeas5UnIwRUX7N6jsOjNxbZqm+TLnFbti++ZdgMB5XfAjCVoLjZUgpKgQlYPMdRkGQrj2pNawtyNQa005MelRpjEy2WYXlRfftkF5eU92pBAS6AlGVYOeHByKgm2O60FVIy3kO3m53nDnlxwDw4nJ8k99+pZo2mfIwcYHTCtH2h7de7BsDuJAm2hy12eE/aLJp9jiQUGBH4CFJ4ScAuOrHPmcDODXO8pBOc9eXKplV2ghcOzpcNnr3Jvc+8PXYTWJkl4PMpY+qw0CpXEnHdEgAY949KhNTgCOflj/P+fCpnY+3VNrp54qoDe3zg40IwNfVMbtOyqfG6M6AfLwUkwNOac01tZC3Iv1hl3+Ve578CDHZlLjw4IYCcmQUDvHFuFR4UJUgcKFHJ6DRi7n6eZZxG2i0gHBzQt8THuH1wp8g/PPwp56p1je9Fbe7YRrYthyNcdLSVTYMtlL0aUk3adwhxtWQSMEgjBGeRHOmszL2M1YOK822/aIuPD7zlrCblb3FeYYdWh5rPj9c56AU6jeakGedhc0k4wTpg41H7Lx7GRO7ONwBH7KQNt73dNS6624vFwfZ79gTfZ2I7CGGIzCM8LbbbYCUJB4zgAZKlE5KiavzEkhyOCVjmARy++uc+kGWNLXi0ai0juVoe4rswdajxLouVCLqVlSjx8YSkHKvBzHLrVlrT2hNyHoAkK2OfvLDYyt/TN5auDYGMk8DSXifPm4KzTa+yVNymY+mALWgjjjmTzwrZZ7hDTRFsuhynH2m9ZXjSGjoCLTeZdqbvF0jwZ0yIvun2Yawvve7cx7ilABPF5KNVY1hqm3W+HC0jtpZlwW5JHBb0K4y4/jhMha+alHhwMqJxz4SAcCU9y+0dozVdne0/f8ASWqLc6U4XFuNqSkeHIguE45eQ5imFtXqzYDT97FwvDrsIp8VQXVHH9bB5c+g8qVsVLU2y39nNA4uaScAAh3TJHRRd7t8V4rWSGXDBxClbs87BtwG/wCUmpEmTPlgFxx0c1Drw8+iemB6Anyqy8WxRY8fuWmUhITjkkfu6enzqLbV2k9gm2kNxdxITCMAcLkKSDyHo0cUuo7R+yJQCnc61Yx4tSR/8qqTdobzcJzNPE/wGD8grVRuoaSIRxOAA8U2t69j7Nrm0rdbjhubHUXWH2hwuNufpIPgeQz+lgZzgYqvadQXfRNzk6F1gwl5pxBbcS4yHGbhHCskFtXIqyAeHIwoAgjkauBM7R2x3MO7kW5ST14Y8kg/+6qB97NZ9nnWSUvI1IzOWlfElLUOQMHPUHhSasWzslxjHulZC8x8jg5aeo/ZQG0FrorqzfY8CQc06Oy3uJcHdZ3fbuyOT3NJQ4CJsCLNdL67W4XfeZbX+a2sKUsI8CfPJNoXpCA2eFWDzzkf58Kpjs3vjsttaiTHgW67ue0fbMWDlSj+2oE/M5qTrh2g9V3tlT+hNjNaXGGRzmTGTCYSMf8AGFDiP71Mr9ZK243B0kMJa0gDLsDPideKkLTUQ0FGyCWTecOajDf68TYW+FvvESWuPMiWe4mG+jkpl5EZxaFD1CgOtV8mbp26RFDNw2y0i88RwqkMxXYylnz4GXEIB8eSAOdSTuZP1HrbU7V/1PfdD6UXBaeZTGVehPfKXElCgpMYrOcE9Qn4VHrETZXTTSJVyuF61hPQf+Aw2hb4Jx07yQvjdWOnupbQf1k1p1ipGU9HHDM3ee0Ad3PHJ5jAx6quXGbtZi9hwNeK27ZY9La0271Rqu36cm2G4aXDElx5Epb9uktvOBsR+FwFbb/NS0kLIUltzkMZPzs86uOh97NHanQ6EIYurDD6jyHcOnunCf2FqPPypd0tr666u0/rDT640O22SFpyZJi2mA0WorToW0S4QSVOOcKcd4tSlYGM4qHIMpUR1uQ2SFsqStKh1yDkfuFT0TZKqGoo5RjTAGc4yFFv3Y3RzN6/Yrq9amFXCRYrahhv8o6cuN5sjzbif5+EVFt+N6n2R+I+kDr3JHjVU7dcL3pzbOzXdwlu47I7iOsSAnktm3vvc/2Q+CnH65rd132rdJoQL3pS7S5V+FwsuomSiOpLAmNxPZJzCyvB4VtJbyQME5qI9Z9oeReNQ7lStPaVjx7TuQ2wmVGnOqeXGcbKVF5BQUDjKwpXMKAJ8cVnlj2duTQ5hjO4ccdDnOCPItcceSn6yupgQS7UdPmritytewtz9UaTtWirTp7RF3tzRsOorbZfZYK7sVpkwFyHccDjvfFbZVknKh6Cm3vFrd7Sejo8/VupbPD1/bEWrUoZanJLxvsTgaltLSj/AI5hbKxnCTl3H2appqLfXdvVGnoulL3rae/Z4kZmK1ESrgb7tpICAoJA4scIPM9R586ZpdkXCQpZLr77h5lRK1LPqcZOf8KmKbYl+8ySqe1obyGucaZJPUcUykvDBkQtJJ6/2Vs9Z9tTS8JcmFt/o/8ALDSlS20flttJiGJIcbkhhxjB7zuZCVlGSAAScZPKt+s9z9c69lzndTXx99mfc3bsuKj3WUynEhKnAkePClIHPkE8utLMPY3WkaAxfNZ+x6MtMhHeMStQPGKp9PUFpjm86D4FCCPWnPp2xaDtzKf5I6Uma6u68JTdLy2qHaI6vHuo6Vd7IV6qUkebdStLS2WxtL6Vu84c85HzOg9NU3e+uru6/RqY+jdr9TawgOahJjWjTsVYbl3u5L7mG2oYJQFY+sXj+jQFL9OdS9onTyLilekNkoM2PCmMmLdtXS44buF0aVycbip5+zR1cwRkqUB7yjkop86V2H1xubJhXLce6vTGIaQmLE7sNQ4iM/YYjpwhA9cDPXFWi0VtzZdGw0x4EZIXwjjXgZOPWqftHtu1oMcbg4jgB8I//I+engp22WDXfk0H1Tf2X2dsm2dkZhwoqEu8I7xWOefHn48/H8ak0ng5HAxXhSQ3yA4fIeleffJGcnNY/VVctbKZpiSTqrlHE2Fga0YSZf3yi3vEEfYP7q5j7jL73tBX9aj9q4qJ/sCul+pifYH0JUBhtRUScBKQOaiegA6knkADXLvcu7Mr3cv9+tcmNNjqmrU0604FNuDhxlJTkFOfEda0n2awuMk5A0LcZ5cVD3WoZBPTyP4NcCfJLEuSiI0pbiglA6qJ5Cm3Ov8AIuDoh25pxSnTwpUkErWc/milKw6J1VrZCrlIUmHbGlBK50s92ykeSBjKzjPIZPI0/LZadPaTbDemmFPTP6a5yUfWk+IbTnCB68yeXTmK0Nz6ejGHd5/0Cu8NPeNqnfgAwU5/MfiI/pCjvUmg5OndL/la8Pd3PddRwxEEFTaFZOXT4KOOSfjmlfRD6vySy1+iDyH9Y1tbiBatJyHFLUtSpLRUokkqOFcyTSToZwiGgHxH/aNdOmfU0Re8817SWym2f2nbS0wIb2YzzJJ4k+JT0AykkimruJdRBtKLcyoB6cfe8+6B/iR+Bp0JI4cqICQCST0AHUn086Qtq9Ky94t4IcQtKXb476XHcnkGUH3U+WVHr+0aY03ZwB9TL8MYyf0U1t7dzRW8UcJ/Em7o8BzP6K5PY520VonbKPNmxkouF2PtTxKcFOfso+SeEH1zTZ+kS1pHs+2eltvmXMS77OVdpKQccMdlPCjI8lKdB+LVWd0/bI8GJEtaMttpCW8pTzwRgkD4furnh21dTy9zO0nJ05BcC02lTFiYSk5Sh3iy4P2XHFD9n0rPNlN6/wC0jq+XUNy/y5BZLdnNoreIGeATEV3lk2stFrfQEG4yXZ6gOXUDJP7KGuXh86iSW6X5Lrx/PWTUpbpT4jMtcC1PLVAtkduDFDh97hA6n1KeHPwqKTW5QjOXdVSJdDhfKKKKXSSKKKKEIooooQiiiihCKKKKEIooooQiiiihCKKKKEIooooQlLT0pMW5tlZIQ4ChXz/78VI2v+PUeg7TqFtpHfWZ32CTj7QSoZSSfEZGc+btRQglKgQTn0qX9sJEW9xpumZzqEsXyKpjiUMgOjx+IICx6oFITaEOSsZyN1Xy7HmvVbj9n+0CYpTlw0otVmkrK8qU0jBYV58mlIR5+5S1vnt5C3F0HcrA+0ONbRLS8A8CxzSoeuRVROwnuQ5t3vHO23v2W4OqUqt7qVqKe4nNFRbVg+J+sbx1y4PKr/S4hQp2NJSnIODgciD0I9DyPwxXz3tVRybP34zw6Bx32+vH6q/2mZtfRCN/EDB/zyXJ3Sbk/Suo5ui7wgsqDymOFXIB1JwCPiPvyKct7JLCvA+NSJ2z9pXbFe0biWRjDTigiYEJwEqyeFf8PuqJIV7/AC7ZESSsKfThD48eIdD8x/GtWpamO7U8Vxi/NjeHQ81aNmLy6GkmslQclmSzxaeXotPa1R9pu/8AqQP/AHgqQm0mo82tBEm7KPi0n/pBUhpPKi6A+8nHQK+ezzWxR/8AJ/8A9lkYkSYT6JMV9xl5ByFtqKSD5gisl4i6W1kwWdXWjhmdE3SCAiQOfVxJ91wf2T5k+GBXPwzXzhHLHPn0zTFsjonb7Dgq3VlDT3CMw1LA5p5EaJvxNMblbYSv5X7WaokvMNcQMi2uFLyUg9HWDzx0OFJKevxE87XdvdKG2LPvFppSHElKF3W1t4KgOqnY6jji8coKR+pk1F8Zu5NrD8ZL7axyC05SR8DyNa95ttp1E0pvU9gQt7/11hAafHh7xA4V/MfOvauGjurNy5RB39QGHD1WXXT2exteZLJUBjj+QnI/sugei9eaL3GtbV50RqiBeIq0FZ9ndy61g4IdaIDjR9FpGRzGRzpwKHup5A5HXNcsm9v9V6bmp1Btrql72lnJQI76o05rHgACM/sk/wAKl3QXbp3K0epNl3X06L801woMrgEOchI68RA4HT6qSFHxUapVx9n73gy2qUPH8p0cP3VMq3XCzP7O6wFg/mGrVe11hpaMKSD5450277oey3tstyoDSs8skdPuptbddpPZvdJhhFh1fHiXJ5QQbXdMRZYPpxEtuendrUfMCpRS0OQI+0OXTn8KotVR1tqfuTscwhLxVENW3MbgVXnXXZp09qBsuGGy6tKTw96jJGenvDmOfkaiO67KbqaViqt2m9W3RNuzzt8vhnwin9ENO5CR6c6vAplKwcp+NYn7dFcSQppBJx+aKl6Hauuo9Cd4dDr902mtsMuuMHquc970aVpW3qvZ23LcaBzN01MVbn1+pYe42VH0QgVE9+07p1lTq2rhebYpP2Gbra1IJ9CtBV9/CB6Cup130RZbkkokQGSCf0RTCvmw+lbiFpEMI4s5A6fdVytntDEOGzAjyOnyOVDVWzrJdW8fJcw3e6Yd7ovNLI/OQr3T8OQ/HFYX+HBOeXp51frUXZS0tKQ4VW2GvPMK7kIUP2kYV+NQFuX2cIWmrdcrtE9ojohRnHyA4VJPCknHMZ8vE1frdt5b7hiPUE6KvVGzs9OS/OiZe7cwrtO3duB5RNHReWeneSZLv/zAfnUegZOcU8t2EOt3+1QnOkLTdmZA8BmC04fxcNM5PhVptDdyiY0eJ+ZKhK3WchZYsR6dLZhRwC7JcS0gHoVKOBUz607JHaA23tUzUFw06xOttuUoSZVouDUpTKUk5cU0g9+lAxkqKAAOZNMHaWzi97oaRtaxlMm9wkK/ql5Ofwq/d41toPbjerWG62q92YHcsWZEFWlwvLq5AShRCm+I8SlYUE4T+fz5Zqr7TbQVFrro6enaH5aXFuCSTkAAY4ceJ0Ura7eyrgfK8kYIGQcclRWFa97bdamr1Ct+t2Lc8137MkMy0sONnnxJURwqSeZyCRSKncPX8tfcrv8ANkkn7DqUunPwUDXRvTFt3SatmyA0LfWLdpmBbG3tTxVSkJ9piqYjq4UtEFbhx3gBGOEqySBzqMF3vTm3u2m8W/W3tvtbt3e1ibXbZb0UKbZZC4wUUI5dVPunlgEhBIIFRNLthDUv3DTtJJAGDwJcWgO00OmfJOn2eSIAiQjmc9MZ0VL3db6ibWQ87DUsdeO3xyfnlFfU7g3zKR7NZF/1rLDV/wDLq3e5msbfofXW3GuoGgNNTLjuFY4UO7okQQporWWFl5KOSQ59YU8RGSEgHOM077pqTTty7Ull2Xf2y0i1Bt6xdPbEWiOHnyq2Or7twd3wqRxPZwcjKQfCnB2nhLQ73XulrnZ3hoGnDvkeHVci2SfD22uQOHEngqNOa31C5z9mtLZ8OCzQ0/uarH/ui6xiLSti5Nx8dFNQ2UH5cKBVu9z5mpdZ7jWzazVXZ/tmmdHStZM24XuJZEw1TI6H18KEvIbSPrG0k8j05jpkJ+6+8u2JuO5WzuutF2a32mLbEMaZNps6Q7HnBoKQriSQEDiKSSPBODnJy7iv0NRuMhpd5zhvEBwOG5AzpxOuo4pJ1DLEHOfLgA44Y1URpt99l6PVqa7doxuI+za1XB+2Q1S3XkrUkFhoLaT3BWolsKHeDu8ji6EUxtG6E3g3keea0hYL/qZcYEvuJK1tN4GeFTiyEBXXCScnwFXS2M0/a5/Zt07tBMQ1HOutN32TxKwMvd6O7WRj3iEraVk8/q+XSool3m67f9j/AEBcNNT5FqfZ1ZJF1MVwtKeWmRJGHOE5OEtoGPJKc5xUbBtCe0miijBeJA1u8BjBB1GBni0hLut2+1j3uO7u5OOKq1erPd7Bc5Nmv9slW+4xHC2/FlMqadaWOoUlXvA+hpPWgYNWJ7c+JW88e/jOLtZYzpJGM92pxlP3IbQKrqo5NX2zVwulDFVFu6XDUDkVAVlOaWd0R1wnhtjLLU3UEMDJmaauzYHmRGWsD+5TIjp5eeRT+2Ttpu+5lps45i5pkW8p8++Ycbx/epd2y2IVraJHuTrslbLxUFIaISEkKKSMkHyphWXGntU8stQcDDT9/wBk8gpZK2FjI+OqjCO0ZLyI7Bb41HGVOJQkfFRIA+JNOK3ab08XFovWsYrTiRyYt8dya4o+QUjDf3LNWUsnZJtCXm1m3tgj/jVKcPzycfhUp6b7Ndht6kqWwhOAMhpCUA/IAVVbh7QqJmkJPpj+6mKXZqU6vCqLYNK6UVhdu281DqJ9GPfu0pNuiKPq02O9UP6roqR9M6T3cnuoGnGrXo5hXuFvT9uDTiUnlj2hwl0n1KzVu7LtLpi1pSWoDZUD4pz++nhB0/boyUoZiNIHogDwqjXDb18hxEzPi7X6HT6Kdp7Axg7xx5KrGl+yr+UZKLnqJ2ROlKUFrlTnlSHic9SVfwFTxpPZnTNgbQRDbccR+ctOfu8B4VILTDbICEpwCOXKsqDxD0qm1t+rq/SR5I6clMQ0UEGrRr4rBEgx4iA2y0lHCOQSMVlUo55HHnz8ab2stxdDbfxDM1pqy2WZoJUoCU79YsD9BpIK1n0Sk1V3cHt/Wxpx237W6UfnuhJQibdctt5/SSw2eIjy4lD1TStu2aul3P4MRA/mOg+qTnuVPTHdJyeg1Vu58mPCYVJmymYzDeC49IdS202D4rWohKU+ZJA9arvuj239sNGOPWvR6F6uubRUgrjrLUFCwcD64jLiR5tpwR0X4iq2pLpvlvfIF23B1O/Hti1BQEtQjREeXdR0AD5pTz8TWS16I0np5QXHiKvUwE8LstHC0FDxS0Ov7R+VXu37FW6gIfcZO0f/ACt0b6nmpCgsV9vo3qePsov5n8fQcV71dr/fPtDynXbxcfYrJx8XszavZLcwnH2eZ+sIGeaipXqa+WHRukdLYfca/L88YIckJUmK2rzSj7S/irFLEmdMlpSiU+VIRyQ2PdQgeSUAAAfACsHXORj+NWr3kxsEVO0RsHIBaHZNg7bayJ5syy/zO5HwHBZp1xnXNSDNkFwIGEI6IQM9AkYA+Va2MZHXB69K9YAHKvhx+Ipt4K7NAY3dCbu4OP5Iv/8A6wz+5VN/RhUIreP0T/GlzcP/APCr5B59+1y+Sqb2l5LcS3e1SFJSy0grWT5Anl8elTVOwuosDqsqvcrYNrBI44AjBJSnrm9O2+1pt0dWJE7koA8w2Ov3n78GrbdjPaNWk9MDUt0YxPuWH8KSMoQR7qc+gyT6lXlVXNktDTt3Ny2pM5pS4ER1LzwPMcjlDfw5ZPoD510ssVuZtFrZiNJSENNDPLl064+X7qo+3F09xpW2uE952r/0CptRXvv9yfcHf7bdGeQ5+qyar1TH2+0XqDX8otFuw296QhpZwHXeH3Qc+f2cfriuW22ipV71XetfXuQpaobb86Q+rmpyS8VYwT+ccuKH9Q1bb6Qfco6f0DYtoYEhImXx0XO6gHC0NIOUMkDzWU59WQfGqqIxpTbSBaXPq5F9cNwfSRglvA4ST1wEcJSP+VVjrUr7O7YaS2uq3jvSnTyH7nKqd/qhPUCNp0aPqmDq65PTJSlPYS6+tT7qU9AVHOB6Cm7W1cpJly3XirOVYHw8K1a0tg3WhVtxyUUUUV0uUUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhApyaOvTltlpDah3ja0vs56cSTnHzpt1kjOqYfQ8jqghQrxw3hheg4OVLG5Ymwb1Y91rC440qcW3hJSfsTGeEhR8QopCVHzUFHzrpttxuHC3c2y0/uTbeBCpsZLdwYQcCPJSOFaQP0QsKAP6PB51za0c7E1ppa47fr7lSpyfa7atz+jlJBKQnxByVJPo4s+AqZPo+t4v5MaoueymqXlM2/UBU7DU4cLjzU4C2wk8srSM8/zmkD841ne3lmNztvbxj8SHXzbz/dWKx1vutRuHg/7q1u5GjLfrDTk20XKOHGZDKm1AjPUH/GuZOrtPXDaXWs7TtxZWYqiUhRBHeNE8lA+Y5fMetdY7iw62t2I/wAPG37pKT7qh1BB8iOYPrVTu1fs+NWW78qwmEJmxSe7XjBVn80+h6ZqkbC3z3GpNJUH8N/0PVWO6wPwKqmOHt4Hr4KqO1rnE/cz+m2k/wB8VIaemTmo52uSUP3JJ6hpIP8AaqREnwB5860u6ge8kDoFrns5fvWCIu5l33WQlHCVKUlKQCVFRwAB1JPhUtdmzsx7q9pSY5c9P3H+R+hIzymJGp3WCp+WpOApuGgkEnnzUOEDxVz4ai/SWibju1uLpXaKzulp/VdzbjSHU9WYiTl5wfBAUflXbnR2lNOaD0ratFaTtyLfZbFGRBgx08+7aQMAk/nKPNSlHmVEmu6SFkbO1eMk8B+qqXtD2sqGT/wuicWgfEQddeWfuqqWf6LHs1xGwu/XDWl9kqGXHpF2DOVeeEI+f2jSNqj6KnaPgXJ243H1npe4AfVOOSETGQfIpw2sjz5mrK75doPbDs76ca1PubfDDYkuFuLGYa76VKUn7QbbBHFjxOQByyRkZgbS/wBKl2V9TXZu0zXdXWAOuBtMu6WtHs/X7SlMuuKSP2T60/a+Z4y1uQskE8rDvB7geuTlU03u7LW/vZ9Lt01lp5jVGlmQf/GXTrRUGkfpSGMJLZAznKQOf2zyqNHJrF5trQuDUS9250EMrdTxkY8Er5KBGehxjPQV3Ksd7sOrrFFvmnrjBu9nujPGxIjrS9HktK5ciMpUk8wQfn5Vzr7cXY0Z2qjzt89j7V3WmXF97qrTUcngigq5TIyOYSlJUriSOSMjhHCThrNSMm1hG68cuqv2zW3tRSPFJdj2sLtMnUjxPUKktw2w0rdgt+yXR20SOqGJQLrCj6OJ94eHVPLzpZ0juX2kdk0pasF/mz7K2viMJ1QuMBXmQhXEGif0kcKvWvMeUxJZRJiuhxl5IUlY5ZB/z08K2osqRDdD8R91lf6SFFJ/A1HOqnOaYqloe3hhwytIrtgLLdQKijzETqCw6HPhwU06B+kLs6iiBufoWXCcKuFU2zOd4lA9WHVA49e8JPrVg9I77bP6/GNK7g2iS4oABh54xn8+XdPBKs8/zQR6mqI3NFqv2TfbBBmq8XQ0GnT6laMEn4g02521mlpiC/aL5LtslfMNS0d61n0cT72P2ar1XsrZLh3mh0LvDUfJVCr2Kv8AbiTRvbO0cj3Xf3XU5bSkFCHEKQVjKQsYKh5jzrXdSE4Bx051zS0prPtH7YDuNDa0uEqCweL2JmWJkQ/GM6FJ/u1IWnu3vulYMxNe6Dtt1UkgFSS5AfT6ZAUgfDgx8KrlT7PK0HNHK2QeeD9VXZqupoe7X072emnzCu7LjhaCkp8cc6i/dvSDV90zcLctjvGpkdxhxKeR4VDHL1HUeopo6b7cmyF/YYTeVXzT01eA83MiJfjoPo8ySoj4tCnyd39pNXpRF07uDp+Y68nCW1TUNLJPgEucKvwqEFmu1omD5IXDBzkDI+i8FdSVjcNcNVQfcnQWrIKUy59uduSYTKI6blGTxuFltPCgSEcyFIQEo4xywlIyajJCC5ktjjwM+7z/AM/OuhWtNAul1MxCH46j7zTwBSD6pUP3ion1btRY7m4ZF203DkvFJBfYSYzqj+kVN4Cz6qBPxrWbLtwI4xFUN1HTT6FVWvsPavMkJVedq9axdudwrFreXbF3BqzS0ylRkOBtSyM4AUQcHnnp4V53m1xE3P3Jv+uYEB6GxdpPfNR3lhS20BKUgEjkfs+FPO67GWxTTi4N4usKRn3WZMZElrH+tSUK/wDd023Nl9WRmi6zMtEvHg28tsj5OJTVljuVoqar30vxJu7uvTj91EGhroYuwaMtzn1Vn5W+u1UvdjZa5WzWjQtGmNPuwrnIU26yll8xyhKFcaQeqU8+Y51qP6k0pvHoPdHZ+xa0sFruEzWj19gPzpXdRpkda2vsLSk8/cV4H83pkkVRf281oFlpFhW8Ry+ofQ6T8Ak5rSk6O1Tbj/pWlbyyfEqhLx9+Khm7MW0FppqjDmYxnHEPL8kc9SU7NxqhkSxEj+2FZvfjVWl525m0uiLNf4N0/kqIMWbLiuBTAXxso5L+DXER4cXPnmnZcb7a/wDw/m57M+MuKLa2nvg6ko/80AEcXTqcfhVJn2JbQLb0SQ2f0VtKH8K10oIWFKC8jqcHJ+7/AD8qdDZWBsDYmzZ/DezP/M5J/skRc5C8vfH+YO58uSuprpO7uitx4O4u4e6Vvu2gLZrBi4s2xu7l9yPHVIKW1BkJwOBpwjAPu88Vua42k2wRqfcPerc262G7adulq73TaI10c7x6cGwnAS0ocQBSE88glQPQKIpH3aniOJLiiT+iTk+fKpNvG592vu1do2qt+hlIjWt8SRL4XZD3ee9xBriH1KVqVlYSSFKAPhTKbZ6aldCaaYA43HFoDe5kHIGeORx55Tj+INlDw+M44jOTqrXXbfnbPbC67P6VFssF3MWzRmJN9TLLi7Kl9IYkABBKQSlKiri54xTFVuN2d9Q2PWezeu9Vy4tib1S5qCy3O3NrdbW07hamhwpUQQXHAcpAOc55YNXGtD6zmNhX8lrnw45KcjlA+9WK3IO1msphx7JEjp8S7LQcfHhJrgbLWmkAc6pO8Nd4O13g4uB544kYGmvVAuNZLo2LTp4YxhO7tHbp2bdbXbVx01GlNWW1Qm7dBVKADzyEFRLigPs5KjgHnjBOCSBFTbLr3NtHu9OInAHz6VJ1k2RckLKb9qMISB7qLfELylHy4nSgJ+I4vnUgaY2QsqVsqTY1ynEKyF3B4ujl+oAlHyKSPPNTLdoLXZaYU1MS4NCbm11ddKZptCVGG2FuvkO4s3nSsZbl1YP+jXBf1cSC7zHeBRGXVpzkDBAODhWKud2dNuzp7SMaE6pTq0FS1OKTgrJUVFWDzHXpSLpfa7vZzHeJLqk8kNIbASgfooQnkkegGPKpt/KulNurcy1qe/2qz+6DwzpjbCsf1VkH8KzDajaCa8DsqdpOeOhJ04K02ygjoBvSEBOGJbWmR9gYHiR0rdS0gHhAGTUS6o7WmwOlQEu64FxXz4kWqKuSv5fZQf7YqG9WfSDW5Ml2Pt/t3ImNkYRJvEoIJPn3DIP4umqxSbJ3m4d5kJA6u7o+qkZLtSRHd3snw1VwsBAycdRzrQ1DqjTekWDM1XfoFob4CoGdIQxxJ9Asji+VUCuXaL7VWvY7ibPLk2OC4Se+tsURClJ8PaVZWkfBYpjI23uV/muXfX2ue+ee951SXFzZKyTz4lE4J/aqx03s/awb1wqAPBupTqmhu90cG0FK7HV2g+quJuD26dm9KIMXS/t+rZ/Cc+xtmNESfAKedAUo/wBVtQ/WqAtQ9rbtEblhdu0LDOnIyyrjetLSjJ4T+lJXlTfXGUFFNSDo/R1mdxatOOXFxv3u/uCu8+ZbT7oHxJpTlX6Zc2EsImtGK0eAR4q0BlBHgEt4SDg/GrVQ2i0WwA0sG8R+Z+v9lYKTYKtqXht2qQ3P5G8fn+yaTO18mc85ddfavLklxRccZZWZclaj14nCeEHzJVn406LLGsOmU/8Ai3ZGGHyAn2uSA8+DnqCocKfkPnWInCQCPDlWCWsJjrOfzSKkH1E1Qdxx06DQfJXug2XtOz8RlpoQXgZydT8yt20vXvW1+Nj0jZL1rS/vHnFtbK5BTzwCtzBAHPwBHqKlTcfsa76aC2B1FvJuVebdpRNsaaLGnoZ7+Y/xuoQPaHU+6gAKJ4QpXMDIT42/+ikbYZ7N0x1hhtDrt9kd64hAStwY5BShgkD1p8fSKpA7IOusY/m4nT/9ZbqWhgp6dwaxuT1KxC9bZ3i6l0bpdyPPBunzPFcorHKXItENa1Eq7lGST190c/ia3+KkTTSibPF/1KP9kUr8VQszQJHea+hbVKZKGJ7jnut+y+lXxoyT0ryTlWK+cQz864wnpcm9uCeHS73PH1zR/fUexnZ12bhactbC3XnnAkpSPtKJOB8AMn/7U/8AcHP8l3jn+ma/jUg9kvaEX26I1bcmwtvmWAeYQnPCVfElOPgPWpQ18NrtjqqXkTgeKwjb8TTX/sojjeY0HyVkOzTtTC0DpSIVspMp1PeOuEc1rPU5/D0AFT3bzHZdVLnvNswoDapMpxascKEAkfiOf6oUfCka3REQ4yGmvdQ2AkADkABUJ9tzdxW2m1De3lqfCb9rAKTK5kLiwgPfCSPFQCUHP6bo8BWIwx1O012DM5c92vlzPoExmdHbKU8t0Kou4OrJHaL3/uWoZCHlWlb6uBtKgks21gHAJ8FFIyT+m4fPFNbdTU6r1d5k7jQQpZYZ4BhKUAnkkeCR0AGBgAY5Uq6abOiNvnLu6lKLnqXLbJUnDiIiVZBTnolS0lR9UNHoajC8zfbJR4FZbb91Pl619GUtOyna2CIYYwAD0WczSF+Xu4krQNFFFP02RRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRQDiiihCc2i747a7gy405wvMOB1rJwDjqn5jNPrclU2De7RvJpF9+Mqa8h92Sznjj3FHvcalYwlS+FSvIqQ761ETTimXEuNnCk8wamLbi6wdQ2uZou+O8NvvSDwLPEruJI5pWEjqcgEePIj86m8zAO8RkHj5JaN3LnyXSDaTc22bz7U2fcKD3aJgbEW6MNq5NyEABaefMYJ4gD+a4jyyUrcuELhZH0lIPTqPWqW9lLd2f2ft3ZGhtbSPZ7DeHhAuKVqKmWXCcNyQc4KCDzUORQrPgKvlry0ezxnmAeJpeHGFZzxIJ+7IIwfUeWKwG/WR1gu7dz/bcctPgTw9FfrdXCupMH4hof3XLzbr3bnd0+SSP79PoqPM5pj6Bwm9XpIPLn/0lPRRJSceX+NarcNZ/QfZat7P3btgjx1d91On0fNrZvPbFtkmSni/I+np0xnP5rhKWwf8A3prrofewkHh54HnmuPn0f+o2LB2xrPHlnAvNlmW5B/XPC6P+irr+57yTwKyroD6/5NOCMMaD0CxnahxfeKgn+Y/ouJnb03FvO53an1PElTFP27SzosttZyShhDY+sIGeqnCsk/DyqD1WR1TZ4lKWk9UqOc/KpF3st0mN2kdz2p7akvI1NOOFDnguqKevpikRUdIR48xStTUOhcGNVw2U2ehraD3iTi7KsV9G/wBpK+7U7uwtmb7Jce0hrV4sR2nncIt1xIy24jIICXCnu1J5Z40nOUnPWy526FeoL9qu0FmXBmsLjyY7yQpDzK0lKm1A/mlJUPXNfn0el3KwXe36hsUkRLjapbU2I9wBXA42eJJwQQeYB5g9KuftJ9KxreyJj2/eTRzN4ZbP1l0sx4HFD9JxhZxn1QptP6tducZ2iSPiqdfLLJbqpzGtJZxUBby7Zp2H3v1js+y867AtkwzLO45niVb3vfaBJ6qSlSUk+JCjypuhwlIqQ+2XvntfvxvPpfcXbObIWZFhEC5syI6mnWnm3XCkKBGFZQpPMEjljPKo2S6CnPn4VHXCItkDsYz91sns7uklZaBFKcujJb6cltJdSn7VO/ajYzfTtBwNQ37aDTVmuFr09JagONzpIYXJeKCpXduEpCj4kFXuhSaji/T1wra4phJXJewzHQBkqdVyTgeh5/KuyvZK2bZ2H2E0toFLHBclRhc7ysnK3J8gBbmf6o4WwPJA8c11RxRtYZZBnkP1UN7Q9pqqhkioqGQtcO8SOPgP1XJ/XW2+6u0b3d7n7P6u02Gscc4xDJgk+aX0e4B+0o03I+sGL+2qMxcWLk0gAlqQ2HMD4OJz8xXeUBSkKbWvLahhSCcpIPgRg5Fcx/pPtprVoTcDQ27OlLBFt8HUDb9kuiYkZDDQlIIW0pQQAOJSFq5+PdGlXUkEmTGCHKCsvtCuRqI6S4br4nEA6ajPPoVUx7TWjbmSbjpploq/Pgulgj5HiT+ApLlbV6UcSV2m/wA+ET0RMYS6B+2gg/3aWml8SR4eNZFK5YzTBtVURcHn9FqlXstZLk0mWnaSeY0P0TTt1j3J0rI7/R+vXYhz/wCiT3Y2fiCUg0uK3H7RtuCVyJz92CRjjdiMyifisJJ/Gs65E1642yw2exy7vdLxJREhxIyglbzqjwpQAQeIlRHLxpU1Jp7XWh1FjW21Gu9POp5ET7G8j+8UpB8elO8S1LQ90LXjyGVn9w2e2TpKl1K+pdC8Y5k8fNIbnaI3Liju7vo20PKHIl2E62o/2XAPwrdi9pGwrGL9ts5xD86Hci2P7LjS/wDapMa1pbM927dJccDwkMuJA+QBrIbvpmd9q42R4nxcS2k/e4kUg6ipP/lpcHwyEw/6Vo6jWiubSP6sJUVvttnJHErTV7ik+BDL4Hz4kVuQt2toXhl6bc4yv1oAGPmlZpGTZbLJa7/8iWuQ1/xjTSVJ+9s4rVc0vpd0+9p+KOfPhKx/Gmxo7dwLHj1/dK/9AXR43oamN48j+ieQ3M2oXzb1vNbHkW5H+FB3F2qPP+Xsj/2Mj/CmZ/IzRq04VYMeqJKh+/NYXdDaNA5WeQPhK/8A+a890t3DekHqEi72fX0cDGfU/snyNytrEA/+Pkk/BmQP4VqSdzNpWjx/ylnSVekd05Pzpot6G0bjP5HeJxnBkn/6aBpXSiVcKNNtq9VPOH9xFdijt55vPqEmdgr0343Rj1KWX94NsWXMss3aR/UhIB/F0Vnb7Q230RsoY0PeJa8clOzW2h9wQv8AfSSzp/TTBJTpyDkeYWr96qySo1ogHDtutETB/pm2UH+/QKO3OOBG53/d+ycf9BV8Ld6pqo2en7rBK7SF6W/jT+hbayj80SHHX1fengH4VmRvdvzd0hFntaIiD0MW1ggH+ssGvjeorHCHCL7b2D5Rzk/Lu0kfjWMaojz5Hs0Bq83RxRwlMWEt0q+Azn8KctpqfP4NKPXJ+65Gy1rpRm4XMeTcf3WO6Xvf7UTJjX3cCeywsYUwbmGUfNtsj7sUhwts0SXS7d9ToKicqSwypwn5qwKl/S+wnaO1uhDuk+zxreSw9jgkzIhgxz68bwSkj9qpd019HX2tb22h67s6I0khZyW5E8S3kD+q33ifxp7HHVMb+G1rB4ABIObsRROzI58xHiSD9gqvx9C6HjrCU2643B/HWRICEH9hCc/3qWoAasDiRBsMC3rHRSYg48einAVfjV3tIfRSpWRK3J7QV6kJVzdi2GAmKkjxSFuFX+xVCRbWdPar1TpuE6+7Etl7lxGFPr4nChtwpSVHxOAPvNITwTdmXPlz81YNm9orNV3BlHQUQjyCckDl804Zlzm3R0Lny3n1jHD3iyeH4Cnpsx2c98e0S1LvO09js0fT8Kcu2vXu7TAlPfJSFKKGxzIAWnok1E+obj+T7TJfQspc4OFPnxE8Ix65P4V2R7GG1j2znZu0XpSfE9nuUmAm63RtQwoSZP1qkq8ilKkJOeY4K5ooY2MMsjc8gFz7Qtp6qjdFRULyxx1cRxxyCr3tv9Fnpktoe303YvWp1jHDa7Ij8nwUHxClK4lr8eaUoqr3a+2EtnZs3/Xp7SMIxdIast7dxsrXEpQYcbHA8yVqJKlBSVK+DqeVdi0SIjrrzMeQ245GUEPISfebUoBQB8sgg/Aiqi/SbbUydc7Afy8s0Rb922/nIug7pOV+xqIQ/wCuEjhcPokmn292oMJAwVl1tu9TRXGOvkeXEHXJJ04ELmyFnAPhWrcFER1j0rzDmImxmZTeMOtpWOfTIzisdwX/AKOv4VDNZuyY8V9KVdQ2ahdJGcgtz6ELpX9FCri7N00H82+v/wAafv0ioB7IGvfRqKf/AIluo9+ifKh2c7ikjl+XXsfcakP6RPn2QdfejEY//EN1Nf8Azeq+VpT3nea5G6aXi0Rcf8Sj9wpXUsjmaRNMq/3pjf6pH7hSvk5qGnH4jvNfUVlfmghH9I+yyAkZyOdBPL1rXkS4UItJmy2mFvEd22snjWCevCOeOvPpWfBBA8jn0POuS0txlPGVUUxLY3AkccHh5pu6/V/4ruj/AJZr/tVbTsfxgja60yUoHEtDmTjr9csVUrcAg6Xe8+/a/wC1VwuyKFK2j0+0w2VuOIcQkAZJJfc6VD7WkssQ/wCf6FZDtZ3to/8AsCsA3OtdmtszU19kpZttojLmSXFdAlCSrJ/VABUfMJIHMiuZms9U3rtM723HU1yD6bS0tS8ZPBAtrawAnI+zxFQGeWXHf1qnrt373MwbdG2N0rPy4sIl31xok9cKbZKhgnPurIx0DYPjVf3I/wDuZaITYVju75fAmTccg8UZrHuNZzzIBOR4LUoH7FKbAWI0VN/EZh+JIO74N6+v2VAv1eJ5ewadG8fPom3uXq5N3mOLYAQw2lMWIhP2UNISEjhHgkBKQkfogVHBNbdzmGbJU70QDhI8hWoceFabGzcGFWHHJRRRRXa5RRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhApW09d3bbLSEuqQhSgoFJwUKHRQPhikmgDJ50HVAOFM2r7WncTSSNWQUYvtiZ4LghA/4RGTzDnCPFAOcjkUZ6cHO1/ZK3xO7O3CtqNTSR/KTTjCVwHnFc5cRJAAP6yQAkq8uBX5qqpTtxreXp25MymFkusjhUgEfWt+I58sj15efI4pwagYvG2GprRuxtxP7iA8+JMZbOQIr45qYWk59w88A5BSSk5KVCq7e7My7U3u79HN1aehH6FSdFWOpHiRvDmP1SdotC2tQXlpwEL97KT9rPec6e2CE5AyKb9/0l+WrSN2NuHVKbUsruNuQoqcguk5WAOqkE5P9VQ6+8E4tP61h3gCNP4IswcsHkhZ9D4fD/IaVMZqPxo9cYDhzBHgtj2D2jpKanFrqTuOyS0ng4E548j4Jd0xrWXtdudpXcuDnNgubUh7GcqY4gFjl48JIru3pnUFs1bYLfqeyzWJkG5x25Ud9o5Q42scQKT+jg8vQ1wYu9vRLjuNuIyCkjB64I/zzq7f0b/a1jWppHZw3Fnqacj8StOTn1e66jJJjklWQoZ93GcjI5cI4lInCSLA4tVb28tT6Wt98A7r+Pn/AHUV/SH7Xzdue03I1ipo/kXcKG3Oiv8ADhCZbQDTzRPTiylK8dcOCq+hJI59Dz++u0naC2B0T2k9vJGgdYlyOS4JVtuUcAvW+UkEJdRnHEkg4Un85JPjg1yt3V7Lm/8AsZc37fqrb+5ajszSj7NqKwMqlMSGj0WtAyptfmF8J5dVfaPVRCakCWPjwITzYvaimt0XuFc7dbnLXcteRUSSojLvJ5xtA6ArWE5PgBnqfSkuXp9SDxBJSVDI5YyKnPs69lnXnaI3Ps0K96Bvls0BbpHtV7uFwaXDDjSRnuWlclFazhICSSOIqyADVmNyPosW2XJFx2I3PetxOVN2TUyDIikH80SEAqSfUoJ9epr1lK+NoO9h3Qp1cdsLW+sdEYy+Mab7frpzC54W6z8M1EhxocSc4Vw8/vp2JOEkY5Dl8gKU9zdvNxtlNaNaB3W0mi0XV9jv47kWW3IjyGslIcQpJPukg9SDy6Ugy57cCK7Lf5JZQVkA4ycch8ScCmlTHK6QNfx5K5bP1ltbRPqKR3cGpOMa4U09izalreztNWKDc2O+0/odP8obmlYyhx5sgsNkernBkHqAoV2NWUo4luch9o5PL7/AcqqJ9Grsw5trsUnXl5YKb/uJI/Kr3GjC2YacpjN5PM5HG5/zgHhU39pPdeFs7shqzX0mS2y7At7iInH1XJX7jSQOpJWpIHkSD4U6IAIjby09Vh14uD7pXSVT/wAx+nIJt7I9o6y7wbwbobf2y4tOs6LlRYsYIOe8HCoPrB8cOhSSenuZ8aSu3ztq7uZ2XNXwIbHeTLG2jUUT3eYcikqVjxyWi6K5tfR3blXHQ/arsjlxc42NcIk2y4LWrHE46e9QoZ/O7xCBz8FHzrtRJYjTI7sCeyh2NIQpp5CxlKkKBSoKHiME0q+PsJBg/wCc1G4ew94YP+ELgTY7gbhbo8rOS4gFX9Ycj+NKBcIGc1ta90NI2k3a1rtXIBSmwXd9uIM5KoqlcTJ/sFPzNJ6le7yqKqYwyTTgfsvpTZu6fxG2RTZ1xr5jRK22cpad+tsVp6taogKHj/6Qiu7jjzjiS2okpOcgkkEfDOK4NbXr/wD29bakjpqeCf8A4hFd4jyHMcyM0/aS2JuD1+6xPbR4lvUrneH2SHd9H6RviFNXvSlkuCV9RKtzLuf7Sc0xLt2WOzZqBalXvYfRD5UCCU2lDKvvbwfnVWvpEe1TvnsDudpG1bTayTZ4lztDkmWw9BYlNOuB9SQoh1CiPdAHLFV9tH0p/artriPys3ou8tJI4gu0JZUefmgpH4U7jE26HNd81WGUz5ASxmQFdXX30bHZY1iwVad0tddEzQDwyrBdHcZxy4mny4g+HIcPxFU27RfY13b7N1sc1a3PGvNDNKCXrjFjludbB4F9vnlv9cEpB8U+Nxex/wBvPTHaWlu6Kv1g/k3rOM134ihZdjzGgR7zSjzBBPNJ8wRxc8WokxYs+I/Bmx2pMeU0tl5l5IU262tJSpCgeSkkEgg+FJvcHHs5xlPqC61tpl7WkkLSOWTjyIXBxuSw+2iRHeS606niQtOMFPh/3+uR4VrXWcqLbpbzK1JcSyvhUk8wcdafnaN2mTsF2hNWbZW/i/ITikXqyFXRMN8BQQP6iuJv/m81Gd/cH5Lk48Wlj+6ajHQCKcNGo/Rb3R7Q/wAXsjq5gw7dII8caq5XZj+j30LvnsdpfdLUe6euYU++NyFvxoTrHctlD7jfu8aScYQDzJr52q/o+9rdh9hr/ujpzX2ubjd7UuMllu4S2CwrvHkIPElDQUeSj+cKtb9Hrz7H233+pm/9dfpO+kiJ/wDBA1iOn1sA/wDxTdS7ZMSbmBjyXz371UE5Mjs/8j+65OxLm45amJClEuBtPMnmVYH/AHf99Xu+j77LPZ93V2Ag673H2pteoL67c5rS5kuVLwttDxSkFtDqW+WMfZ51z/guYtDaRn+aT0PoKsv2XPpBoXZs2dibat7V3LUE5mZKkuyTNTGYAccUpKUjgWpXI8/s/OkKZjmh4j0OVfNuJZKqKkLsu7vjxwF0z092cOzzpfhVpzZHRUFwfZUiztLVn+svJp9221W21JSm2WyFDSB7qY8ZDI+A4AB+FcvdRfS5boygtOmNntOwCR7ipk158pPngKSKs52Au1NuX2mbHrO6bkotDT1lnMx4jFthhlDaCjKskqUpRyepNLPEwG+9yzx0HZ4yzCtm4ou8lcRJ6ZyTTfv2ttHaZaLmpNW2W0tpGSZs9pr16KOTypbW5wpUs4IAJIPQjHQ1+dS7Ln328XF25zJM1wTX08T7yl4ws+dIMj7UFznJelpn1cgiiGpXanXvbx7KuhkLYnbrW+4Sef1NqQuac+X1SVAfPHxrkMm9QtRaw1Tf7epz2W53qVMjlxJSotOOKUnIPTkelNCNYHE+6ltCE5zySM/f1py2qN7Czwg5yc8zmuZzGyItYc5wtB2UsVTQ17aubQAH6qR9iNujvJ2htB7auMF63u3FFzuYAyBFj5Wvi9CElPxUK7hOPhCC86OFKAVK4R9kdcDyxXOP6KTbpq43bXW+NxhK40BGm7Q4roUYDklSc+PutDI6cx41bntd7nJ2l7OuttWtvKRMFuXBhKScK9oeHdoIPhwlXF8qN3dAhHL7lVHaOv8A4lc5qgHTOB5DRVt7Inape3G7XG7uj5cwrtl9m+1WQJVlCfZMR1AeYWhLRB/V9avHfLRA1BZLjp66sh6DdYj0GU0oclNOoKFpP7JNcBdj9wJGzG6ujdzWVyA3ap7SpgYVhxcZfJ0DzPCpRGfzsV39tk9i726Nc4rgcYltIebWnooKGQR9/Wu6hgjcHM4fqFEyQupzuScePoVwn1Toy57Ubj6s2mvCHEv6Xuj0ZouDBXHJy0vHkpBSoHx46TZ6v9GVjyq4f0pu2f8AJrcTSO+VuhqEe/MfkC8OJ+x37eVMKV6lvIyfBoeVU4mlRYUPHxprUMHbNkHA/fmtp2RuxrbI+B570YI9OS6WfRQrP/g63EHwvr2PxqRvpDDnsibgZGcRWP8ArDdRp9FG4f8Awermgj7N8d/7VWU302ph75bWX7a64Xhy1x762hlyW00HVtJDiVEpSSBnCcc/Pxp2527Nr1WKSghzsf5quGtmuFuh2aJ7TKHeKaSG2GxxuuEgYCUjnz8+Q9asbsf2K+0JvWGrtJsx290q4QoXa9MqE2Qnx9njcl8weSiEjlyUa6DbF9jLYPs+8M3Rujm7hfeFKXL1eSJcoFPi0FDu2cn9BIPmo1I+4u6ugdqbMvUe4+rYFkgpBIclvAKc8whOSpZ5HkP+6uS2Bry5o3jn0Voq9s7lUQMpoXdkwDGnE+ZXNLtydk/avs0bR6Gf0ei4XK/3LU4buV9uT3HJkp9ncVwBIwhtAUMhKRnpxFRwarhHWpTKD1yAam7tw9snSvafNi0Nt3pie3Z7BdxPF2mkNqlEIW37rfVKCF5BJCuX2RVfbnf7dYoQ9rVxPKHuR0/bWcdT5D4864rBJKGMx3laNgallvp6ipq37rMg5PPTx4laevVA6aXk8jIbHXryVVodp9zrPsX2WYGv7olqRcZbDsWzwyeb7xedAUT4DPET+qD+mmqo6f0vfdzZjl2vEv8AJlggAuSJi8htpA5nh8zjnnn08TgH7cpl13bv1q0jp9S4unLAwY0FDqj3cOLxFTj7nPmtRJUcZJJSkZwkU2q7RFc42UUjstY7ed5cm+ZVV2j2ibX3B9dA3ALQ1ueJxzxyWzopL2or5dN39duqmIjyVyQXekycfeGfApSSFFI6koT0NNDXGpJd3uD8qU8VyJaipzn9hPgn5c6c+utRwYkSNp6z8SLTaAWYzZPNxWclSgDjJJUSfEk+ATUXOuKdWXFnKlHJPrVohjboWjAGg8lR5HHGD/hXg4ooopykUUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCFkZfcYdS60opWg5B9alHQmr4EmE/p3ULXtNnuQ7mSyo82Vno6k9UkHByOhHMEZBiqtmBNdgyUvtkYH2knooeRrh7N8Lprt0qU7Ldr5sJrnupCRPstwACwBlubEJwFIznhWMEehyDy6yfuLsJYNW6dZ3J26mtrt89IWH2hlDLhOOB0AZSCeQVjKVZSeL3SY107ebJrOyDRmo3eBpwg26djK4T+MAHn7yTgJIPUAdCE4cWy27eo+zxrCRo3W0Jx/TdwPDPiLBUkNOJGJDX6SVIIJA5KTkdQkit3SjqM++UJxK3iOTh08+ilaOoYB2M/wAB4HmCmCq9al0bIRZ9VQlqaTybUrmrhH6K+ih8c49K23m4d5Sm5WiYtLjJDiHGjwvMq8FDnkYP/wB/EXO11sfpbVtib1DposXfTV0SHozqVB5LRUDwgqGMdFBKuRISQcKBFVd152f9Q6Tlrn6WedKU8SlMLXwqT/UXyBB8jg+GTUVb79R1rw1/4cv0J/RXCO5VcEHu1YO2g/8A9N8irgdln6S1dhjRdvu0ZxlDIDUTUrSCrjA8JCR0V0PFggnqE5Kq6E6P1zpPW9uZv2itSwbpCkDLT0OQFhXwwevn61+euTcX40t22apt7jD7J4HMtcK0HHRSDjP4efOnDorWet9AyV3XbXW11sy1KSt78nyVBCyPs9439lQH6wxU7JT6bwOD4cFBuoY6kl1E7eH8p+IL9COXFOfWLUSOgUonA+dIGu9caX28sMzUurbvGt8KCyXnXH3OABPr1IB6ZwevLJwDxra7efa3ixxG/wB0XjwMd4qOEK+Pu4H4VG+vd1d395Fo/wB0bW9zvDCXA6mKpQbjhYzhXdpASVAEjiIzgnnzpIU5Or3aLiG11kr+zZGQfFOrfze2T2gt+bxuMnvha0kQ7Sl4YUmK3kBRT+bxHiVw88cWCTjNetqttJm+e7ukdobeVhq8Tm3Lo4jl7PAbPE8vPgeALx6lNMi02tuG1xrAQAMlXkBzz6D/AAroB9FVs2pMfU/aJvbSVG7OqsVgStPvpYbOZDv7SuBAx+gv0rzea+Qyjg3QK4XDf2dsbaIHvSnXy5n9F0Bt9thWmBGs9rYQxDgMtxYzSQQlDSEhKE49EgfdVEfpU2t0r5pDTGltJaKv9z02iSq43mZb4Dj7aFIGG0LKRyHvKUQeWUg+Aq+Fwnx7fFenyXUtsR2lOuLPRKEgk/gKa2hN3ttdzopm7dbgWS9lrCXUwJ7a3WT4BaEniSfjiuY3hr97GcLPhkaDzXA606hRpi+WnUtond1PsU5ic0jBbcCm3AeHB8fPHlX6BtE6qga50hZdYWqSJEO8wGJrLuPtpcQFBX4g/Om7r7YjZrdBLytx9rdOXyU6MLlyrcgS/T69AS5nn14qXtC6O07t7pa36L0jDch2a0tdxDjqfU8WW+I4QFK5kDoMk8vGlZXsc1pZ9U4qat1ZJ2jgAfBc2PpQNtF6S3t0vuvBjBMPWNvNumLSOXtjBwCfVTam8efAaqshQKEnzrqd9I9t2dwOy/fLhDjlc3SD7eoIxQnJCG/dex4/zalHy92uUtrme2wmZI6uISo/P/vzTSpbvMDumn7LTvZ3cSY5KJ3LUevFLe2Jzvxtrj/+ZYH/AE7dd5AOIZJFcHNr+e/G2vL/APMsH/pkV3eK+XToKWziJmPH7qn7Xa3iX0+y5afS1Jbd3m0Ela0ISbG8OJRwP59XjVO0W22qQMyYp8yH04/fXdTX+yu0e6UpmVuTtvYdSOx0Fpl64xA462g8ylLgIUkZ54Bx6VFt3+j+7IV0eDv+49FjHPSLcZTSfhjvKUD43sAc4jHRN7Re/wCEtcwwh+ddVzr7Aek7xfO19pWdpxK3omn2pM65vtc0Ms90pABUOXNa0AeprtKog+8OXMf9w/fTI2x2i2y2esbmn9stD2rTsR8pVJ9jZAdklIOC66olxzGTjiVy54FK+s9caY2+0zO1Rqu7x7dbrewt9x15wISEpxk8zzGSB6kgdSAeJZBKQ1qh6qY1M7pi0N3jnA5LmH9J5NiOdqKxMxVpU9H0o0iQE+ALzpSDVVL6vNskZ8Wl/uNODeLdSdvlvbqXdB9C241wkdzb2jn3Ibfut9fE44j+spXSm1ejxWx//VL/ANk0nK0NlYOmAtR2Za+GwSh3PeP0XYL6PVX/AJH23wH/ABM3/rr9J/0kJ4ux/rP/AFsD/rTdb/0e4Kex/t9gf0M3/rr9JX0jpUeyHrLrjvYH/WkU4H+96rJwuR1vQXIDSfNsfurWXZEKVxKpQtCf9BYH/Jp/dW6Wh5UyMrmPdjqt8it8dXTRmQZ0H2CQVWhttBISCa6H/RDBLWn90GsfZu0f/YP+FUKfaAbJxV9Poj1gWjdRHld43+wql45S+N2fBUfbehjo44uzHM/ZdB5AHcOgHqhX7q/PVameOfclEZP5RkeH61foSfP1Czn8xX7q/PvY05lXI4//AHlJ/wBoUF27C70UNsbEJLs1h6FK6GEjBGK0NQShBtby0541J4EAdSo8gB6/4Ur8kjA8aefZ821e3l7R2hNvu5W5CTN/K1zITkJixwXFZ9CU8HxWKa0435RngNVqm0tS222yWVuhxgeZ0XV3sg7Sp2X7OmjNEyEf74LgIuly5YPtcrDziD4nh4koB8kCq7fSgM7ja/tmjdkdsNJ3nUcqfMVd7pGtcRbvC037jCXFgFKQpSlEcWMd0TV5GnSgYwAPzQByAHh8OvL5V5ud4hWO1PXK93ONbba0C49JlvIZjpHQqWpRCeg6mnscwEnaEZXz+7OM81yO25+jM7R24DSGNbq0/oK3EAlc58y5h9EsskgK8+JSa6nbWaMm7cbeWDQlx1O7qF2yQmoQuLkUMLkJbSAFFHErnnP53pUIbkfSC9lTbYqSNxEanlJ4gmNpxr2zmPAu5S2Pko02Ozf9IPZO0nvDK25g6HVpqCIC5MF+ZLDsmWtKwMFKUhKOR5JHET58q7kMj2ZLcAJaWeWokEkrt44xqpe7W+1Ebejs/wCrdEmIh24GIZ9qV+c1NY+sbKT+sUlB9FmuLNvnuTrWkvJKX2SWX0+IcScHPr/jXe++aisFkjOyL9d4dvjx21PPOSXQhLbY/OVnoP1jgDzriPv6nbi07/6wjbVait960xeX/wArQlwV8bbDjvvPMZ6YS4XOHBI4OHnzpFv4jN0cW6/urLsncPc6wwk92QEHz5K/P0UawjYG8BwjhTenFEnwHv8Aj4VYbd3tJ7KbKwTK3D15BhOFAU3DaV30l0HOCltJyRyPPpyrjHoztA736D2/f2w0VfvyFaX5jsp92I1iS4pWMfWE+7jwKcEZ60yX4E27TX77qO5vy5MhXeSJk14uOOLP5ylKOVH486WfEwuLnu49OKj4bLV1TyQ3AydTpzV496fpYNXXxx+ybBaObtEVXEgXm8JD0hfgFtsj3EeY4uI+YHSqa6u1Frbcy9Pax3U1jcb7PdI7yVcpKl8I8EpBOEj9UADyFNuTqK125Pd2tj2p/pxqB4B/E0v6Z2q13r+U0u4odjRiRhKke8lPX3UcsfPHzruSRlMztJCI29TxKcNjt9udho7eToPhB8SkVeoksvC26UiKekLJAeLeVE/qJ8OXifwqR9puzxftfXUyr4CsIAfkqcUS2w2T9p1XiTzwkEZPjyOJ22h7MkFhwR40EB0ICn3nMngRkDiWrGcZIwkYycAerN7Tu/Nk05Cd2M2WkByOlRZvF0ZUFuSXiOEtIUnkTzIUQT+gCRxFVc/jMl0nNFaRgni88h1/skq6aZ7RLXu7o4MGgHkFF28WsYl2uDGzu1jftVpjPpZWuMkKVcJOeiSkYKEny5Eji6BONC4OQNuNNr0la32VXOWkOXie2eLmByaQrwSnJAx1IUonBSBigWxray0vGe0FasmtrbdbWnItzJylTZB/PPMKz0+z04sxnfb07cXFAPLcC1la3FHmtR6mrfRUbKaIQR8BxPMnmSqxPO6R5ldx+y1LrcFT5BVzDaDhCc9B5/OtKiipIDAwmZOdUUUUV6hFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCFu2u5OW6QHBzQT76fMenrUsQ7ladybIzp3UEgNTGf/NVzOeJlRP807j7SSc+oOCD1C4bGMVvWm6vWx8LT7zZ+0g/w8jSb2b2o4rtrsaFWM7P/aE1V2dNTu6C17EXK0xJXwSozie89nCsEPM45LbVyUQDhQwpJCgDV1rjp3S+r7DH1TpaSxc9P3FpDrL7SuMNpXnCTn83ORxYyCcKCVcjzpj32x66tLVi1O8ONrIt9zxl2Mo8+7czzWknOQc9cjHPLy2T363B7LerDY76w5dNMSjxPwieJtaD0fjFQx06pIwoZSoAjlQ9ptl/4jmqoxuzDiOAd/f7qftd2NL+HLq37KxWuuz3Y9UsEO21p8JSeEODCkZ/RUCCPkfD5VWvW3Zg1RpmQuZpqctTYUngaeUELTz6d4PdIz4kJrofpq+6K3T0qzrfbOaidAkJKnYqB9dGWACtBQSSCnxQeYGCOJODSTcbPEmhQeaSpJGDy/HFZ7Q7U3KySGCXOBxaf7qzvttNWgSRadCFzBk3PUGmZZt2rrC8hbaj7zjZbWeeMpV0UPX8aXbRe7DcuFLEpDTp6oePCr7zyP31ePVGzen7/HLT0BhSAThtbYUgE+SSP41AWvOyFDBVJ04XYDgCiUNkuNE+iVHiT8iR6CrrRbW2y5ANnBjcefJOKS5Xa0EEYlYOR4/NRXcovtFvfh8Smy8jhyORA8v8+FSbtB21e0HsbZrfpG0SrXedOW5IajwJkfu1NNg5whxsjByTzWF9enSohve2m6+hld4zGduEZsE5j5dSB4gtH3h8SMetIjOtGB9Vd7WppfRRbz/sq8as0DA6MmAiRngU5uN2tO0Lm++gwyAYyeHzV5dyfpQo+u9mNT6Ri6Qutg1VdIHsUVw8L7YKyEOL75PDzCVKUk92PeSnrVB7VCudufbnwpEmJIbPEh6O6W1pV5gjn91OSHI09c8OR5rPEeZQ57h+48j8qWTBbI5IBGOmK7NT7uNxrcdcrq27HQVWZGzCRvLGE79HdtHtV7dd0xZd3LtMiMgBMa7JTNbIHgQ4FHHzq9vYg7emsO0TrWXtjuVYrFAuzNsVPhy7Y242JXAtKVhSFrUAcKB5dcE8sVzbdtTLgOUD7qW9pNwrpsDuvYN17Hb/AG1y0OrD8YLCe+aWkpKcnPgfI9K6bNHUDcIwVHXvZKWgiNRDqAdQu7d1tFrvlrnWO9R0vQbnHchymlD3Vsup4FpOfRRrg5d9KTtu9c6n21uQIk6Yu0iCSfFtKyEn5gZ/aroVpD6VzZW7FlnWmmr/AKceyAtxcXv2U+uUFS8fsn4VVHtp6s2p1tv3H3V2i1jabvatZWxtM9uK4Uux57I4Vd40sJcSFo7oglIyorA6GkxHJuOY9p6hR2y9b/D7tG9xwHHB8io62rCTvztonz1LC/6ZFd1SrhHFz+XUVwl2rUte/e2qk5//ABLBx/7dFd025CVN54gRgePWkj3Y2ZPJd7W4dd5ccNPsqS9u7tobvdm3cbTenduounZUC7Wgy5LN0gqdPeh9afdUhaVAYA8ahCxfS3b0sD/f3ajRE31ivSYx/vOrr59LC0mRvXoZtCSM2Bfh/wDpDtU/asrSkABOcDnyp1vRRxt3gkLLYH3VjngnQ4VyL59LbuzMiKasG0ulrc6eXeSJ70hP9lJbJ+SqrDu32gt59/5pd3E1UuRCLveotkVJZiJUOmUg5XjwKiT602U2Nn9AfdW9FtrbWMIHL0pL3qNmsbdVaaPYkiQGUkhYbXD7lA4gOIcj6f5/h8K+X04t0gf8kv8A2TSoGQkcqTb4B+TpJPg0v9xppG8vlDj1V3q6VtLbnxNGgafsuvf0er4X2QdvwFA8DU0Hn/8Apj3+IpO+keUR2RNYDzdgf9ZRXr6O1wHsk6LTnHB7YMf/ANS5WP6Rzl2RtYFXi9BGP/6lFPGH8b1WBY1AXJuzD/QWP9Wn91KQFaFn/wCBR/8AVJ/dShjHOo6X/cPqvo+3N/0kf/EfZYJIPdqA8qvV9Ei4kQd1msc/ypGVn9lVUWkH3FfCry/RKrQIm6uTjNyjnPphVLwf7bvRUH2hD8OHzP2XQaUs9w5z/MV+6uANjGJNz9LnJ/2hXfK5zY0Vgl+Yw1xZSnvHAnJPLAzXBC0gpn3ZB/Nusof36UJzC4jwVc2I/wD1dvk5K6sKSQTy8avD9E7oXv3Ne70zGQfaXEabtylDmltHC69g+v1IOKoZfpqrfapEpv7YQUoGOqjyH3dflVm9n/pELZsXs1prarbTaqVKdtUVS5s+dIRHEuc8ouPOkBK8p4lcIGUq4UJ5+XVJETG57fL91O+0CtdJJFQsznicfILp3utrSHt3ttqbXc5wJbsVsfmdeZUlB4U+mTgCuCer9fbk7tTjdte6zvV7ecJUBMmLW2jPglJJAHoMVLu7vbH7RO+Ftm6V1BebfatP3PCZFutkY8K0hQUAXXVLcxkDISsA46eFRrAtbUZlKQgDAxjGcClu0FK3e4uKgLBsvNcZSahpDG9dE3I2nlgBHAEj0Hh8TS1pK9az2t1XB1zoK4ewXiAHUtPlCXAlLjakKBSoEH3VK6jlkEc8Upvu2+GjjmS2WARn31YyPQdT8qb8zXNoYUW4sdyWc4zjhSfv51zFNUTHLW5CtFytFjoYOzqZQ3Pjr8glXVOsNy9z5ft+4GtLxelFWQ3Kkq7kH/Vj3R8QPjWjEYs1hBdnSGY68FOMDjPwT1+fT1pOip17rB0NWS2OtMrVw8TSS2gfFxX+JqQdH9l7Ud9cbkXyW4UlXvIZzk/tkcvurirq6ekbmrlDB0HFVqG60lE7/wBKpy938ztB+6j6VrHv3PZ7HbVPPKVhLjieI/JAzz++lrTu0G4OuJCHJ4cjtLwcOAlQB8ODon4HFW10J2dNN6fQ2pMBlpaeeQCVHzyo5UT8al6yaPtFrbSzHiNNgciQPP8AHnVMuG3VPSAx0LNep1XM8dwuzt6tk0/lGgVdtuOzLZbMW5EmFxOpHvOugKcJI+4fKrC6P2wjN5bgRm2W2khT76k+62Dgcz1J8h1+QJDuhacjoiO3e4SGrfbIqVOOyHCEJASMkAqIGABzUSAkdT0Bpz2lu2ZK1eF7UbGd/Gs6nDGkXCOkh6eScFDXLi4Fc8rPvK6YSnka3RMuu1lTuAkt5k8B/nRJzyUlpiwBr05lKXal7V1vtkeTs3shJJytUe63hnmt1YylTTKh1Oc++nlg8KSealV3stnRtkwL1dkIe1Q+nijx1kkW5JH218x9bzHqjChji+x6t1ngbYtqnS1R5+p1IHCE++1bwRzwSObvPHFggYPCeiqYOodRSbi86VPl1bij3jh6k+XoPIVsVptFPaYBT0405nmT1VMq6uSqkMsh8h0RqLUMme+sF9S3HDl104yryHLp8Bj4DoEAgYr4aKnGtDRgJiTniiiiivV4iiiihCKKKKEIooooQiiiihCKKKKEIooooQiiiihCKKKKEIooooQiiiihCKKKKEIooooQiiiihCKKKKEIooooQtmBOdguhxvmnopBPJQ/hUp2DWVi1BZRpnVjBlW0kcCz/PwVcvfQRkkegHPOCDhOIjrNGkOxnQ6yspUPGk3xh66a4tUx6O1buX2btUR9aaDuwl2h9QCiBxxJjeRll9APuq5YyMKHVCvGr87J7+bc9o63I/J8tuxatQkiZbHzguEDPeJxycSc440gKGPfT+ernFofcBy2LciPpYeZkAokQ5KQqNITjoQehP78HwGFaZozjktap2puMmNcIrge/J6XiiTHWCOFTK8hShk9OShyxxEnFUv+zNLfGfijdlHBw+x6qWt9zmonZYct6fsuns+DIgOliawplafBQ6+oPQj1zjypOejtO5K08845iqx7F9vRLpa0R2hYRebQQ01em2OFxpwEAd+hIynP5y0DrzUhWSRbNmHbb3ao+oNF3Ni92qS2HGnoziXFYA5/ZJCgMHmPmEnlWK3jZ+uskm5O3TkRwPqrxRXKnrm9w69Ez7lpC2XFJU/GRxHxxj93Ooz1j2fdNaiSTJtEaTjmA43k/JYwsfIipqSriBIxgkHn5V7wgfaBBzg00pbpU0bt6F5B804kpoZviGVRzWHZKRG716xSJcYklQQvDqR6DPvAfEk1GV02h3U0m4pdtK5jaRkmM7j5FC8E/IGulLtvjSBlbaSCPGke46Mtc0EKjNnNW6h2+rYhuzgPHiFFutAifv07iw+BIXNF/VGq7OQ3f7IUYOCtxhTSvh4D8Kzp1rZLgQ1LZfj55FRAUP4H8Kv1dNn7LNStkx08C8+7gcJ+I6EfGow1X2WdJ3FanfyKw0vH24w7nPyRgfhVkpdsrXUn8ZhYfAp6y432lG62XtG9HDP1VXE/kSZlES6Rlj0XwZ+SsE18i6aYTLRMCUq7tQVxADr8qk/UHZMVHeU5bJ0xhs9ApAdx8vd/fTHuGwO5FqWpVrlNPIT9kpdW0v7iMD+0asFPdbfUN/BqAM8iuf46d8OraMEjXLSsF1iSHFxZsCa/FlQ3O9YfYcKFtqHRQI5gjGacFp357R2mkCPZt39UJZHMNyZPtKP7LvEPwpl/kLdqyqU09a5b+OvG2l8ffzrRe1Hqi3LxdbAEefeR1t/4Cn8LXAbrHNePMJSvuthu8vb1LHxuPHTP1Ce2tdxN094bzZ7puXqVF3eskdUWM4YbTC0tFRVwnukJB5k+tfG2QlIGKaLO4LSMd7Zyk+PA7/iKUI+4NgcT/pDM1lX6jaVj7+IUlUU9VLxbw6KybP3XZ22RGGGcDJz3spwhsA8hX0pFIzes9Nu/ZmqR/rWin92a2G9S2BzpeIw+PEP+zTM08zeLSrdFerZL8E7D6hKHDzpLvrZXbpCUAqUWlAAfCtk3mxqGU3aGfi4B+/FYTdLSvkbpDwf+WT/jRG2Rjg7dKTrqijqoXRCVveBHEc/VWo7Mv0hmktidmrBtlcdB6hnSbUHy6+y2z3alreWv3cuJOMKT1HUVpdpz6Qmx7/bR3na+27eXmA5c3I60SXu74Ult1K+YSok5ANVVkmxnmm4Q/X65P+NeGHtPsELXc4gP9bNSO+3e3ww59Vlv/SUIfuuqW48x+6VbWgtwo6VpIUG05B8OQrdByPnSOvU9hbGBdmDy8OP/AOmsadaacRzdkyF4/wCLZ4h+JFMTTzSEuDCtKhvlspGNidUMwABxHJLLiQQc9D+NedJ7m7wbXi727bvWsqyQ7xLEqSmOw0pxxaQQk94pJUMAnoQKb0rcS0IPBDhy3h5rw3+AJ/fSe5rpbmURrMjjPTiWVfup3TQVMX5cA9cKsbRXfZ66Bsck29unPdBTovmuN29ZOlzVG5Oqbgs8/rrm6cfAZwK8WS0i2w+4RxLUVFxRJ4iSep/AU3oty17NObbp5fCroRDKgPgVA0oxdvd3dRH30SG0KP2XZHCn+yM/uomOG7s0rWDzCg6C92W0v7aige9/DOMfdK8tyHwluXJZaSRzDjgT+8im/NuWlYCQ6mWl5f6DKCoj4k4H3E08bV2XdUTUpcutzIcV+Yy0VD+0rH7qkHTnZNtgCF3CO+8Qfe710gfcnhqLku9qox3ps+AXVXtTX1+sdM1p6nUqAla5QFJatVnLiz9kOEkn5Jx++ssa1bnarPdw7c+y0o4J4Qyj7zgn8auhp7s9adtqEpYt7LIAH8y0lGfuGT8yafdm2ysttKUphoyPMc/vqDqNt6KAkU0WT1Oqj5X3evG7UTkN6N0H0VKdO9mfUl3fSq8XBawrmQwgn+8vH7sVNeieyzYLV3Lki3NuOpOS88O8WfXnyHyAqy0TT8OKgFphKcUqIispSBwAfOqpcNtrhV91rt0dBouYLLAw7zhvHx1Ue2Hayw2lKVJipVw+JH8OlO2LbY0VPCwykBPiABSqEFIBCeXUVvW/Tsy6o75AbYje8FSHlYRy64GDxEDwAPy6iqy1k1S7MhJJUm2KOEYAwktpjicShsKUonASBkk+g8edZNX6l0RtDp1WsN0Ly1CjI4lNwslTr608+7CUkKUpXL3AcjOVKQOdQ9vV2x9uNnm5Vh24U3qfVHNv2hKgY0Yc+alpOAoED3G85GQXEjkqm19l6/3tvKtf7samfj29Y+rdcSAA3z4WorHIJTkdRhIyT7x5G6bP7EVNzxUV2WRfU+XQeKgbjfY4Mxw953Xontvd2mNyO0xdRpHSsFy06XaPuQW1JR3iAeTklYwgJHLCRhA8iedMFxyxbew3IempSJl2cSUzbsRhLfXKGAcEJPLn1PRXIlNat71XCs1r/IOnWBbrWDxKSk5fkr/ScXyKvTPIAnAGTmPbldZE9XCTwtD7KAf3+dbJQW+GhhFPTMDWDl18SqfPUSSvL5Dl3+cFsXW+PTMttuKIKuJSyeaj40kGiipIANGAmpOeKKKKK9XiKKKKEIooooQiiiihCKKKKEIooooQiiiihCKKKKEIooooQiiiihCKKKKEIooooQiiiihCKKKKEIooooQiiiihCKKKKEIo6UUUIQDg5pdsmqZtrkNOKfd+qOUOJV76Ph5jwI8jSFX0AedeFodxQpbXetJa8RjVbAZmFASi7xEjvQQnCe+T0cxgdSDjHvJxSzorWm9fZwmL1BoTUBmWB11Knu5Up+A+QfcLrfLu18+SiEq6hJx1hGNKfiuB2O8pCvMGndpncK52V/jalhjiBQtJTxNOJV9oKT5Hx8D400npWTMMcjQ5p5HVOI5ix2804Pgr/bVdtvaLdVLNv3JaRo/UayEuSSv/AEWQs8uMOnkM8sh3h/1h8Jzf0xcEs+2WpbV0iKQHEOwyFcTZ58YSOZGPFOU/rHx5YzYOgdZNmQhtGnbgokiRHwqG5noFIGODn4pxjOOFXUOfRW6naG7OqG5emb87L0/GcCktFRmW9Clc/s8iyTjw4CefPxrObv7PaepJlt7txx/KeHoVYqLaGWIBs43h1C6IcQQAEqHPPPrkeFe0KPirNQft19IDtHrn2eHu3ph/TtyX7j02MVOx1uHqsKSO8R8FJdx+kasNaoOl9XWROpNvdZ2u+2lRwH2ZDfCFfoKWklKVeiig+g6VmVysNwtLi2pjIHXl81aKa5U1Xqx2vRJK1NgjGKxLbadQQUJPLxFbFztl1tclLFzgSIrjg40B1BSFjzSTyUPUEisTPFjmfl51EjLeKejB1C0nbVGdTgtJyeXQUlv6VgOZC2E/2acw4cHiryUoIOCCenKu2TvZ8JXJYHcQmFO26tEjJXGR05jhFJL+0lldGQwlOf0RipRKWuHChknlWMJQegNOo7hUM+FxSbqeN3EKFbn2eNK3bnLtcJ8nmC7FbWfvKc03ZPZL0S7nFmYST+iCgf3T/CrJNMoxkjp51kLTeeWKfxbR3GIYZKfmU2dbqd57zR8lU+b2OtKP5DERbfqh1f8AjSHJ7GFsJPdLmpx5Pj+KTVzURW8e8BzHlXpMVofmA86eR7ZXaMYEpSLrPRu/IqRHsZsBWPaLiB/rUf8A0Vhe7GXPLdwuA/sH8eGry+yNKzy5ZPrXtuIwftoGB5ilxtzdW8ZEmbHSng1UUb7GXiqbcVc/0kD/ALNbCOxrHOCXbic9cuo/+irymLHT/Rj5YxXlEVo5+r+/Fenbm6u/+ReCx0n8ipbC7F9sPN/2sjrze/wFL0TsfaXYA7y38ZH6bq//AKqtquM0kfYxyrGWGc+8nxps/bG6ycZThKts9I38qrPb+ybo1lwKVZIis8/fb4/9oGnVA2GsFuQGY8RllAA5NNpbH3JAqawhAHDgczyryWEBWQnFMpdobjN8chPqUuy3QM+FoUawdpLBGSD7M2o+ZSM0txNDWeMAW4jYHnwinaWyCE4o7sJAHjUe+vnk1e4pwII28GpGY0/EQoFLKcA8sCt1FvaaThKUjl5VtjCVe6eY8K+cSuL0PQ03Mz3cSlWtHILwhhKOH3Rzr13YOegx5ivbbL8p9uLHaW866QltttBUpZ8gBzPypYY0xMYQ67eJLNrajpK3RIWOJlAGSpaR9gf1yn414N55wBlcuc1upSQkgDOAeXLlW9AsdyuSS/FZIZBCVPrPCgKz0yeRPoMn0qJdw+2N2dtsSuHapy9aXRpCs+wcLjBV+jxn6pI9fraq3uF2w+0LvpMkWnRqXtP2ltB4WLUpSVsRzy+tlHBSnnzI7tB6YFWy17FXS5nec3s2fzO0+Q4lQ9XfKam7rDvHwVxN0u0FsZsch2Nqa9Iv1+aBP5LhAOLCjgBChnCfUuFJx+YfGmW6Pan3z7RkyVp7Tjciz2BxJQ5CgOlI9nB91MiRyHB0HD7qCcDh6VGMTSGmrHIdla0uartNyT+T4bp4eMjq671I80jGc8lV4v2u3UwfyXF7m225JCm4MMBKSrGOJWMcSsZyo8+eOfWtUsuyFvtOHhvaSfzHgPIKrVt3nqyQTut6Dis0Sz6P0cOOepjUV5SMlB/4HHOOhBH1hHjnl4YOc0hap1xPushT8qZ7TIJ5H8xHToBy9PgMeFNmZc3ZPuIw23+ik1omre2Pm4qHLjjAWSRJelOF59wrUT1JrFRRSq4RRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIWeLNlQ1ccd9SPTqD8qdVg3BuVpeQtiU7FwQSG/ebV55QfD0FM6vuciuSwO4roOLeCldczQOqowVe7GmLIWSpU+1K4VZ/XbI4VD04Qr9Y19sFh1vpiei+bT69WJCTxoVEm+xyBjoCCoBR/VyT6VFTbq2jltwpPmk4NKULUc2I4lxSy4U9FZIWB/WFIOgDm7h1HQ6pRsu7rzVptD9vnf3btw2zcSysajhlYStNxjKiyUjx4XEAIUcZ5rbXmp60l27uzdq5LcLVtnuulHj9qSGi4kKPmWkqSR/wAyOnhVC7TujPabbhyZCXoyT/MTW0vN48uYOR/WBHpSlMd23vzK3ZmnFW+SvmH7VI+rz5qaVlPySB8KrFw2PtVecyRbp6t0+nBSVPdqqDRr8jxXUuwSNv8AW8FNx0FuRZrxHfwUdy+l1YHktLZUtB/rJT8BWzM0hqSEz3hgKlI5+9CWJGB8EEqT+0BXKNnRUND7U3Re4CGJIwU+1BcZwK9FoyfmQKfumt5u1xtq+l+0auvVwjM49x5xu7RiB+qrvEgfcaptX7NDkupJ/Rwx9lMw7TOGBMz5LoK6HWSUOoUhfksYP3GvIdAGceGKqfpn6SncO3rEXcvb61XYpwFOMKcjrI8ctO96158koSKlbTXbx7M+rC2nVunLxpeU5hKnWWCphJIHM90VDxPRkdKrFVsTd6Uaxbw/p1UrDfqSTicKXUPHz6+VZErUep/GjSV52i1+2HNAbq2m5uL5oZ75KlgeSkAh0H/mh8KVrjpa920uFUYSGkjPexlh5GPXh5p+CgDVZmpZ6Y7srC0+Kk46iOUZYcpPS6ByUa9BeRzNaiSrxB+FZUrOOvrSGEtxW22Rj3gcVk4s+HXyNajbvj5jFbLOSeoPLwrgr1ZSOLrkH16V5yEqwnPInNfFOFHU/jXnvSvlk/dXPFC+qcB5YOa8qOeWOlfCCrkCedb0KzXO4pSYsJxSFHh71WENj4rVhI++gHC5ccalJo5L588dK+8QGPSt6/OaN0e0ZGt9d2e1JAz9bJQhI9ON1SAf2eKog1T2zuyxozjjxZ9y1VOayMQWXFRyr+sruUkfAr+dStHZq+vOKeFzh4AppNX01P8AG/Ck7IW5w45k8hit6Lp+93BzhiWyQpB/pFo7tv5rVhI+ZqpmpPpN30tLj6C2mhRHByaclSAgD9hhCFH5uH41Emp+1p2tty0n2O7S7ZEOf/NFvTGCR/r8F35ldWmj9nt2qD+LusHidfkFFTbRUrNIwSfBdFZenYdpbU7qPUtqtaEc1KdfHdpHmXVFLQ+a/nUT6v7WPZg25eeZ/lGvVlwZGCzASp9ni/rIKGz/AO1WOfQ1z7vOn9d6leTcNxtwEqcXzPttxVNeA+AKgD6FQrUas+3lpdWZDdxvix9jiX7MwT64977lfOrXQ+zijiOaqUuPRug+fFRM+0U79GN3R4qzeuvpJNTzY71j2h0HEtbbo4UTJqe+eTn9BloIR8OPvKgjVV9343bCZ+5WtZjFvUThFxk9w2B5piowceoRj1pCk7ht2tKWbHGt1nQ2kpAgsDj5+az7yv2io+tNGdqyVJWXCpx5WSeN9RUcnxwTirnb7DQW4f6WENPUjJ+ZUNPXT1B/Fef0TwiWnb7Tag8+09qSYnIPeqLMRB8wEnK8efFw+lat63HlyIqbc3LbjxGkcKIkJIab8euAMnzzTCk3CZLOX5ClemcCtYnI8KmRECclNN/HwpSl32S8OFnLSSckg5JpNUpSuaiSfMnNfKKVAA4LgklFFFFerxFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFekLWk5QopI8jXmgHFCEoM3y4sgJU93iUnIS4OL8etKkHWs+GviaW6x5hl0gfceRptk5orksaeK9DiFJkfdm4uQzDmyo09o8u6uMRD4x5ALBA+IANe++2+vhbM3SohrI+sctUtQSfXgWVBPwSMVGAr0hakHiSopI6YNcGL+XRdB55qS16I08++49pDW5ivNp4m2rg2psqPklxGef8AWQgU/dEdoDtQbLtoUxep91srWFGNcFC4RUpH6LgJU0P6i0moDaus1AIL3eDyc978etOjT+41ysi0LhzX4bjYKUKbWeEA9cDOU/KmtVRRVLdydgePEJWKYxHMZLT4FXj26+kI251sWbVu/pldhmr5C7RuJ5geXGU/Wj04u9Hw61YO2MWjV9sa1BoK/RL5b3SSkMuoWtIxnkRlK/kQvzSK5ZKuWkdZOZvtlbhvuFKfyhbQEEZPNa2wOAnxxgE/pClzS8vdbZCY5rDbTUzsm14433Iv1jK2AcASo56DmASQoAkYVnnVCvHs/o6rMlEezd0Orf7Kfo7/ADQ4bN3h9V0sQhaFFDiFJWklJSoYII6jB5jw5VmS7w88+HWoP2R7aWjN4jG0xuOyzYNTngZakh0d1JOBw8Kl/nZ5d2tXT7Kx9mp1k2yRGdaZSlLy3hltbRyhzzxnGPUHBHjisludqq7TN2FU3B5dCPNW6mrIqxm/GVjcWhQ94gY58z5VuRrUoMKnXSSza4SEqdVIkEJHCBnKR1PxOE+ZFMTdbfrazs/2ku6jmtXrUi2u8iWmMpKiVEHhJyCEIyBlxQI68KVY4qonuVvVvb2mp0n2mcuDp5p7K4zbpZhNKOSkvLPN1eBkcXErAPCBzqdsGxtbesSu7kXU8/Ic/so+43mKk7jdXK3e5nbc2P2wWuHpOIvWN7SnhHdFJjsLT+m4oKbPPwQlzHgoVWfWHbJ7TG7qnYOk3XrHCc9xf5HbIcwenHKWVLT8lJHpUYx7ftzpVpKksr1Pc20pW46/lqGheeaQjqtI/SUTnpwJpH1PuXdrqluObiURmSruYkVIZYZSfBCE4SkegA9a1i1bIWu24LIt9w/M7X6cFUqq7VNTnLsDoFuzNF3KeV3TX2u20y1L95gPqnS1HxJVxcH3uE+lfG4W3NpfSfyJPuwQPeVLld02o+eEcKh+740xl3+d7waUlsK8hz+81pOyn3yVOPKUT1yatTYS0bvAeCiy8HXipMO48a1ocFhtdltKDyT7JESXU/8AOn3z81Gm9cNwLhOQUSZ8x/PM5c4QT646/OmcaK6EQHHVeb55JUc1BNUriZ4G8dOWT95rSkTpUlXG/IcWfVVYKKUDQOC5JJQTmiiivV4iiiihCKKKKEIooooQiiiihCKKKKEIooooQiiiihCKKKKEIooooQiiiihCKKKKEIooooQiiiihCKKKKEIooooQiiiihCKKKKEIooooQiiiihCKKKKEIooooQiiiihCKKKKEIooooQiiiihCKKKKEIooooQiiiihCKKKKEIooooQiiiihCzxJkmG4HYzxbV6Hkadum9xLtZJCJEKUqI+FZKkk8C/Agj1GQfMEg8iaZdfQeuedcuY13xBehxHBSo/p2z7gpMzTzUe06gR764qVBEeXz6o8ELzz5YSf1TzVKOje2tupt/t9d9u73EXcLyykQ4E2crLsPh9xaHkq5ukAYGSCMYVxAAVXXTt5lw5aGG1HCiShQOFIOOoNS3DmW3U8JWrNR2dubdbGkoS9x8KZAQlPD3qMEK4QQE56Y55ThIirjbqauAjrGB4ByPA/5xTynqZYSXwnBTcYs0u6Sl7jbu3KZMXc1CU1GdeUZM8ryQtZzlLZx0GCoY4eEELGlq7X65yUQmUIiwmElMe3xvcaaSfPHLJ6kjmSBnOAaTNa6oul2cTNuDpekSQVFxSiSkHwGcnwHPPgKZhUVcyST55p/HECBpgDkm73rYkzn5KiVL4Un81JwK1yMV8opcDHBJcUUUUV6hFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCEUUUUIRRRRQhFFFFCF//9k="""
DOSSIE_BRASAO_DICOR_PADRAO_B64 = """/9j/4AAQSkZJRgABAQAAAQABAAD/2wBDAAMCAgICAgMCAgIDAwMDBAYEBAQEBAgGBgUGCQgKCgkICQkKDA8MCgsOCwkJDRENDg8QEBEQCgwSExIQEw8QEBD/2wBDAQMDAwQDBAgEBAgQCwkLEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBD/wAARCAPoA0cDASIAAhEBAxEB/8QAHgAAAAYDAQEAAAAAAAAAAAAAAQIDBAYHAAUICQr/xABqEAACAQIFAQYCBgcEBwMEAiMBAgMEEQAFBhIhMQcTIkFRYXGBCBQyQpGhFSNSscHR8BZicuEJJDNDU4LxF2OSJTRzg6KywtImRFSTGDVkdHaEo7PDNkVVVmVmlNPiGXWFhpWktMQn1PL/xAAbAQABBQEBAAAAAAAAAAAAAAAAAQIDBAUGB//EAEURAAEDAQQHBwMBBQcFAQACAwEAAgMRBBIhMQVBUWGRofATInGBscHRFDLh8QYjM0JSFSSSorLC0kNTYnKCNBbiJUTT/9oADAMBAAIRAxEAPwDyyN7n1PnjP6+eMuB5c4FSWva1z6eeBCL58kfPBrbeLc3tjAoJuQLe2DA88n2464ELADYWAJ6WGBG291Nj0+fpjAfCTx5YwegsP68sIhGHS9weh4PHwxjAE7v3dMCbdevnf1/zwB4F7WBPJ9sKhAnB3C+70t1wbcAoawve3Jt88EG6/Ug35t6YU2EL4utvX92BCAbiATzc+Z6/PAC7Hgc3tyP4YEkKLE3b2/fgV+1dgV9eLc4ELADc3Fh6Xv8AuwIAve1vT3wJUixP8hgVReCxIuPQW4wIRAqhuOh56f1+GDqSDzYHi1xx8cHsosx4HTpb+hjJFsArcEG3Tp7e3nhEIAQQAfEb8X8/bAguvi2E2tb3xihT0vyOlrn8cCAOtxc+drfPCoQ7bx9AbAkfDBSroNtgLC54wYP4b7vtAc39D0wQgFubW6c3H/TAhHUWva2089fL43tg7XcbfETbkDr8v5YCJQLMjksD5HofX3wYAruF2Ful+SPn/HAhZuDAEsCL8WBBHywZQd+3cgJ4AY2B9LHpgBdlVQLkC/Pt1Hy64yIqzEBzzcceXofx8/jgQg5vwSbjqfL3/rphRGA+0b8XFlP9W9cEa5Yszcvf7vHP9eWAHN+VuAADu8z0/d+7AhOQ+9AWsvoLEj8QLf15YyRrhbCwJsCBuI/PphIAoCStibA7gQL+mB3brIEPW1urX+OEQhJJHiXm9jY3/DAfa8K+IAkDqL+h/wAsYxG2wYCwuSOh9ievywK+KMhV6C5sOV/mMCEXZvW1iQfK3P8ALGWYKxa/HBJ48vT1PrhTbtZVFiLWO5QRe/UeYHTAEEllBDKOSQQbgenr/XTAhCnDWCMtugPI+f8AlgW5e5ktYE/5W9+mCuSrH9bdU8Js1wRbp8P3YMk118uBYhb2BPnb0PnhUiTJDRiVRYE889Pb1HzwoGsLFDZeWDeH5euCncAGO4gGxZhx09bYElFuATZRcHyA9PgMCVKAjmyICeRxc/I+Y/PBwFSzJta4tblSp8wbfxGG9lBC3I9mYAfHCiErYi5JHLnkH4H19zhEJR1bcrOOWHQ9B5dR1+XTzwibnxGw5HU839Pf4jy64VszeFRYkdQbj4n+vPABdwvsABNh47EH28sCFkbtGOT16g/ZA8iD1v58G1sGanaRlUsSQeAT0waNALEt4b354sPP4EHn4E2xhkuoHHUGwAt+Qv8AO9/3YVCSvezOyAH7R2X2+nA8j64OlipFk3fesLgD+Hx8sAC11Oy/xNvO1hbz98Kb+AA91uQxLGzWPH4H1v7cYTehFcPGnKgG9jcGwN+AT0Nxgp5DC21rC9z5jyt5YPZXK+G118RJtbjp7+nODBpO6CpI4+z0N7HzA87HjjphUIiEAgFRYizWvuH+G/mMKh1UbFkd7dCt/EPUf54CNliXmVrt1C8Wt+8+noRa+E3QK1u/LBQdwsRduvF/P9/I64EJVXDowF+eODa/r/Xn8cY8chGxb2PAuR++1vlgkX2dwYXQ3P3gvPn8/wDPCqsyqHAbi4uSV5I6EjqbXtxyLjAhImKzl1A3EC/HIB6X8ufLn2wvChB7wbrng+Q/E8j4eVupwk0u915BO0ABADfy/H2PX0GDmRSd3FwNpCn068enn54EJTkkyksb8BB4QLed/MfD52wRdgIEi+Ld/Qv5j8xg0r7ACCE53A3G7d7n09COOvngGIdboSRa5awVgfvAn29D1vgQk0YHcViUsRuXagAvcDr6eVup6YNua4usfmwG+49lAHPyxl+78O9Ra6kmwF/NTu+zwOh+RwVyVkAkYqobnd4T7H/PywIQKxCM4dOOoDAEg+x6j38sDFMrKxIANtwYNcED2t5+o/dhR4JXX9arIb3beCo/Hpz7H9+E1UEKB6HmwJB9D6H2488CEWXkhmKXUnpwG5tb0+fAwqrC3AVh9nklR8vM/HpjJIzv4RwGA6nduI+H7umMRRfdvFhZmswJHw97WPyOBCTmKhQNw3E9LXNvUfy/gcCBu2sisihrAHm//wCd7dLH44OynbtRV8ZDA7wWQg8HpYE+h8sEuVchWUgpuIPS3wPp6dRbAhYN8hKFbXO6+659unx4sOcCouCz8nhf2SPba1j+BxkbiQqm1G28lQF/ePP8Rbg3woWG9yXALMTvYhUPqbfwwISMn29pcjnglgov/wA37vXA7gTsIVw/IKeLp8PUX4ueRhR4Iebsw23Pi44Pw6H9/lbCjI1w3JAQLYjgenTyPr1PxwISUU0cUoVRtZrqnhJJNvQEeXywor99IVUxyEq20A7Sx9CD1+XJw0qlDDc8zEA3F5SfkAb/AJHB0WNVK/bUhelxbi/N/T0/PCIQfV5gwdAhvfaEYK7H2vf8xbBWLbOV3HoW6X9bD+PtxhyTE4JMQKkceI8H+PXkYKEaa+25ANi1woU+pPQfvwqEz3KSCqC56DoT7/H2wsoIue9UMPNbsD6/MfgR7jCq0kLSAkylSeVLAk/DdwD8fnzg4QGzqVDC3gIuSfcji3HW+BCakbCd1jwb+K2w35vfr8uOcHjdTIhlUMtyCqDxcexFucKNTuXUR7/E/gBG4t6H0N/P+GEY3KnwsV2m4ZT9keg+BuAcCEo24qLRjd5neRc+/t78YQlY3u1gGAUkkcN6Cx/PDkykXZjZd1mv4fGBcXte4/64R2d7e7hrDxdCLfy8vy64EJsQFBIA3eXFunr6f54VMniuQ5NgV3jaT8vP5HnAvTMbNcEqBwLn534v8fLpbCfc2DAm27o1iAD5n+vXAhLpZorM3JF1tfp+NvxsffCLiURsWBAtfcR0Hw6j4nB9rbTduSLDcAAfbjpxjHRTuAVtx8SMo6ny5HQ/vwiE3IHBsORfm9hjG2A3Z1QdDc8j5YcmMkEKq2HFhY2Pp+Pl+GE9ht4QF43H046gf+6n4DCoSZs6/rE5BsebN87cfPrgGLsSSLkXJtbp8D6euFAhS+1HFiTYX/ofwwk/jUobDkMD0N/jgQsUAAuqEDzuQTY9D7m/FvfAEGxO0A2uLG5+BwZv9mdwHla3QX5v88AzXu5AsSWuRYD1Hvz0wIQId20GwHN73Hwt7/uxniVTvddqdbC359L+x64IXFxtINz4rm+FBblXbbwD1t+PHzwITbh7L1PpyD19fLB227ghJt6Hkj+eDIi7g6m6gk3Cnn5H92MHk6kX9QSOfb0OBCSAPCkHxfs82+Xpg4JZNpJNuRbpgTZRtAJ8iD5fj54KrCxG7rwLLzx5YEITe63AXn9r8zfGHawJtz7nGEi1gAL+9zf0v/DBSxcnkCwt4SePhf8AdhEIrK5AutvPCZBI6eXrhVoyBu2m9vUj88FA6AgWPIF+tv4fvwqEiVPRf88F4sR+Ywoyk8kj44DaAdzDj8flgQkz0AvYng4y27kdfywYqBwAPS/rgbXFuSAL2v0HrgQiW4DW5PAGAtc8+X5/LByAbcEgDr0scBzezXHH5YEJM/H8PLBTZmAAJ8ve+DkE/s+/88YVPA9rYEJKxva/TGYHoLWGMwlULOvReT0wYA24H4+WDGw8hjP8XXqPQYVCxSOBfgHyPlg1iAb8XH9Wxl78m/J9cZ5c3BI8xgQgJt1Nr884MBa52/wvgOOFJHW48ucDZvskj3wIWA3JuefPjBjY2I6Hg++AC9Cx4HOM8JsrEG37/TAhYim4PTr/ANMKNa3Cjjnjz/zwmQLdLEHChBJ3WsLXvbrgQgPNiTx5WxihitrjgccHBiCCG2ix6H9/OBVeeefe17YELApIuvUgcHz+P8MH225V738r/v8ALAqLgKSLXJAPB/o4y6khieLc+Vv8sIhArFTwbgrYE8kg9f44MSGsCDt6c+f9dMFAv9k346K3At6WxliSQL+HkkC/zGFQlLMoDd4A3UckH2Psf39cJm4BPkLDp09PjhQMW22Ul3JWx58Xkb+d/fBXAKB782Nrn5W9rHqMCEQ223IPv05/r4YMD4j4QLWsPT+OC7QSSCeRyR/XX+uuCg7SOhI9/wB/xwiEr3njMhJJJvyOSf3YOAbbwDzcgg2PuDhLcGUFj1J4PH4f18cGBs19wW4v6A+nx+GFQhKb9oHiJI8JN7n+eMQFmCgHk7QD6/8AXCwKgEoDsa4NvT4dcGBNyjEm1r3H2vQE+fseuBCTQBlD7m5NiWHN/bAWe5UqCoHqOPa54wqfCftsQOvNre3xwlsBayeEjyZeOfUemEQjRq7L4jwgJtfgDztfr+/Akq3HI5vc9OfW39efmcClrMhcAhuR159vXm/PTBxGW2qBe1zbgA++FQgJtYhbkGxDgHn4Hj5YKqbTt23ZOPc/j1+WHCq7OVRmJNkW1n3+g6fh8MJEqApvcdCFAA9jxgQg3BgB1ta/FrH09cGa7DdYubkHxDofQ/8AXCVmA3FF4PHPPPp6/wAMC5L8eLw8gOd1v3flhEIzGxBHHIA9uODfzxiLdtrCzDn0Pxv5H+jg3BUr3lrDcfFbr+8/0cZsJlAXbdgLWPBPn16G/r64VCy3Vgo454A48vP92MkO1jFwdrW6XHX8cAfELANYG/qAOh+f8PhgHRVsGcfOxAPwI4wIWKWD90puRxa9h8Ph+7Dnx7N8buSLDcrEbSfu/H8sNVszAtusOCALfAi/A+eFY0ZgCXUMeh3ALb59MCEDBm2uTcAA35vfz+duePLCqEi7Fr3JIsCwtbqPIA/ywMYO3wX2sSVIJtf14Hr+GAO9w0pJ2g3L7LgG/rwPkMIhG3MRYXJVR1sOT68/u/LBOGVgSSOvi8/Xj5/MYP4RILxHxKDtPi68g3HXnp5j44EoIyQX4vyCLgelxxc+h4I88KhAy7v1RXaoa/Ucnpci/mPh5YKVI4YkWPIJ6H19CfQ36YUV9rFpN0ZBubpuIPwPn8cB3kVgqS7VYm5I22N+U4Pz/LCIWKrKg2FWPG3hiQPIcflfkfDAMzRBWDgIxuCBc287jytyCPmMKqo2o+4efLAEH1vfoRwfncHywBVS1l3sxJYkWJIFuelx8ebYVCB4327FUExgkeK4J9Dfy9COtxgyXVwSwAta2ws3r8L+WCbyyN5qo3/Yuo5A4A6fHBrMksqJGVIaxF+fnxwD59fwwISccjyASBSp5BfdYrfqDbnp+IxjBhGL3C34NiPLqfTj52+ODsGAExiZfJCVO1/2l3ccH4HywhKyNzEX2ggrubxE36EdCRz+/CISodtiHhl5Q7l5b2vwQPTz9cGV3QMQt7WJXyt7n+RBOMRzYixNvDcdB529CPOx598OBvjtaR2KEi8beAj1Fut/ceWBCSnJEomIuXJIcGxYfHlSPa2BAdSXEYsIyOBtG0c23X5B9/PAyRlZFQWVntYhbd5cHi3RvYjCKkMFKqg9l5JPW1v4exwIQCY7u+aQEud4a/2r89enXy9rcYDxsQtmJsSLAkHnr6A/O+D96HjKK4V38RsfER63H7+CfhhaNZQwWFC5uAUKE3P94XNwenN/jhUII02qVWKxI3eD7Vx5gXsOPa9vbBVvYd8oBS7eI2svoSDe3mPTy4wcqGH6rhS4CeGxv1AB8mHl6/ljEA3yCyMRdiCoRbdCTa/zA/LAhGETGQxqrO7XG2x3dPSwJt7XwKys5EjtuZluwfkcixLeoPqOfPrg0LNeNO8Tlb7V8LGxuL358PkRfjCpvYoSrEs20Kt/F1Is3PzHn1HngQmrMthydxO25AJ3DryOT8bXwQI4bdtIU9PF6HkHr8cOLO5H62QIVtuDEcevw9bc+eBaNQgIWS32W32YqR8AARb09iPPCISjxR90A3iADPZjdrAdRbz+Hxw3RRtUu9yellBB4vz52v6G9ubHphRtpJTzDD/edPQ9ARgGRLhluLeNgF8Snm4uCPT8+uBCRJ7ogAndYi3PS9rc9T14+fPTA+J2bagAY2JvyR1tfy6c2HlgxS8qpIjHwJ4TY3BFwL+QPt0NrYEd8CpAHeWsLAgm3nfrfyJ88CEiN0iMBNwwK7gFI587Hz8xbkegwZQpCxtLtdgbmQGxt15Fx+VzhUTFEH1hWDNypdrcE2ueL/MYR23mJjmAQsQbKVsP2ueR6g+mBCISXUGNzfqtr828wB5/na18GG0/7aEEqRbffw+pUAgemA2sxBLKWHNm5PA5sAOfxv54GBEDAhjGgIuNm5AvnuPp7++BCVeWVWZHLMyjad0m4g+YF+V+OC8KLKu7kbVW5/5T5X8xa/mMAZCzAszsqsLm+7jp5mx8uvlbCigKVuQoB6KDxbz29bj0BwISNvCzBi/3j5An0N7WPy/EYxEVpFHPNtx8P7gOnTqemMDKLFldrrfbuN+b+fN/l1wDeFg5l3vJze1wx8mA9T5A4ELFj3BmKKGB55Pg9uvT0Pp+GDDep2EspI6MB0+Xkbc4Ed5HbwMpsRbbYr8R90H3wRZY1a6OnhI22NyPW4txY/G+FQhkvYsW8NuQzdD5EevxHXzthFvAhbvNxLBr/eNvUHiw/wAsHRixBUAg+E7AGv8AkOfbBEUyBkUfZBLkNtCgHm4PWxtzz8MCETvJYydoUWBW3hHJ6EHy/iDgFVja7ABQSF8/gB59flhR2lZVMYIIBAIBva9yOOvr7YxVJDOVICkFrgkX8xf1wiEDi5BYJHYbbgk+/UfuAt+/BbBnUDyIBB/Kw9vfn1vhSRCfAVbcFvtBDFvmOLj09MYDdRGAQCADYkh/Tjp/ngQk2TcSO7kC36E+fS9+l/f4/HCTIRGS4J5AJ2g8/vJ/hg42Ky8eEEkm1ifysfwvgxC23PdWUFduzw8+V73De/78CEjJdrsxN2PJ4N74SWy2sGve97C9+lufzwvawIJfpZrN4iD7n+uMJ8hW7y3HX1P8wf8ArgQkGDEEXY28/j/DCkbNsDLJYDi4e35/1+OBZSygNcjyv/P+OFm7xgjuD0tdyAT8v4jrgQkmuwJJ3eKysxPJ9/M8c4KR+sFwbE2sbW98KbZBYubDmwI4v52H8sJSKSQCNpB+1a5+BPngQsZLqbMCxHXqLfxOCrFYlnBAsLAep+OFCV5FuSdxFwMYpI5F1uOSfP8A8WFQiSgKBvUk+dze/sB/XywmVvc/a8gB1+PvhRjcBEbzI+1a/wAB/ngrKFA6+e64vz6e4wIRSBtI5G4W6fvwU82JDe59/TBm/ZUEC3kb/wBDBSCCN1wRyRa1v6/LAhFYFhxtA6kW4H8cFAv4gAN3Xi/Ppg1wwsrg8nw3vz+4YDoNote4YC5/G3rgQkmXgn5HGck3JNr2ucZwAWa5Frjnn+vfArcXYenUDnAhAL9TwfT3wUg7iSb+tzycKMBbnbYDkjn8MEsSegv1tgQidLHr+7AEAXFjg9iSW6HzscFYix5tf0wISfUmw4xmMNrEnz8sZgQjDcGIt7WwcHjkm5N+vngp6m/rgwuLFiefX2wIWAbR14tx6YxTc8Dg364Hk9QDzycZY36dOt/IYEI9yRYH5YKSS3Qm/IsPLAiwWxBHJ6nAMATtI568+WEQhtfk3uPLzxjjxbl446emMXaDZRt6f5XwZyu69vDfj1wqEVQTwOfwFhhRAoU3HlgoUnkoNvkD/A4Od+3qQD0/j88CEVyAAenHn1HtgQwsCoHwPNv5YwL6dB6jy/kMG2qTYsOvPx9zgQsU39bDnp5euDKC19rH9oHj9+A22F9oIJtfqAcCRa26xseC4vb2+GEQjg7h7dSx6C3lx8bYK5AAWx69AehwcAFiHPNiDcjn4euCFNw9No4DN74EI3sQfF144v6D1xmwNwD1N7AfmL9fh54EqQt14P2lBNrYBDtIJD2A4sdp/cfywqEPdkEsDe1re/w/fgjx7jyWN+pY3Pp88OLbWCStyLggi469CPM/nhNw28i/HxJB+J/ngQkxuJIUdbCw56cef78HRTxtDC458/ngFXexXduYkgEdD/n+/CgANjZb9bEce9vfzwIR7tsuOgu1+Ln+v4X8sAoIDAbF228jYXPX4fzxgAPG0hj0bg2+Hv8A9PPGRAgk7SLg8Djy9fT/AKYRCM1i1ieOAbgAgefT06j4YwRK9nVLgnqOeb88fj8sACSPBYWG7jj8Af4YUXaTz9sm4YjqfO5/j+PGBCAKquQXHPQnlfl64U3MLBQLcOFuTcnjgnyPp64IWKtdLjqvJU8/LjBkKJ4QjXIuNynj+YPngQhks6hdquBa3U+H0PoR5fMYTIJbhnC34YW4+Q/dhRxvCq26zrvNyQL3/DBHJJLtuB45Kcnjz9PfjAhEAVm3HzNjxYD2A8sG2oA3Lg9BdAfx549j54EfaYEsL2+z0t6H+eDrZVIALXO29ttuPXz+GBCFSEQbGLEi5sTfrxYD87/hbBWJZu6VTvubrb26W9fyt0wYMg28WsenRr+nt8emCMGBVSDtA4AG0WvzwPx9x0wIQvtDK7EP93gm449un8cAkanwX5PhHI5I9iLH4ce2BJuSXuysP2vzv/Pr+eAZ2YMzF5LcXHIv73/hgQsG2y2AQKbnyAN+vthZGG+5ZtxFzzyWPX7PHw6XHphJHuAxIZj0ZtwNvTjp88HKMTa1yQW5NioHnc+XxwIWF2kci+/c1xYm5PoPj7jB2ELHhQ9wLe/wv0I9B8sF3b1uzA7TuuOtz18Vuelx88HBYPuL253HwCxPmf5288CEe27kBPGdxVVsAxHJA8gRY2HH4YNJH+rAKExuOdvIJHpfn4i/wwR9wVWDOEN1C7bWC+fHpf3PJwMDuF8CrIGYgqTdSLXI46jzBHT8RhULG7w7bKSpQKzLyCvkOevHr6YMYmR1IW3eiwDL4WA6C3ngv2VXwKSWvuYWIHTm32ufO1+uBaMmQ7gFtxZrKRb38yP5HAhAmwFSxU7lvckX2nkE/wBeRxk0SBtjsH2k9LEA+t/w5HJuMCsjj7DsVBJNmCgHrckjw3+PPlgAV7m52k9Cp+HW34g89cIhFlisAHUWvwxF72HT5XHlhSKLeg2pbaQt72APUD16dCPcYBI2szFCQouSyNyPUkfxwK8gBgCAvU82B5uP66jjAhAVjkBlkLszHwlxfjm/ivxY8bbYKsAN9rKeRwCeB6+48uOh9sLnvZJVNzK8p4s92e/F/W58r2PTCTKJLAotwrEMBtJIB5I6X4tx+/AhKASX3BlZvMKoA2j1W1iPbm3njC/efrWisCx527Rxbhbe37x8cJBCwYbGj2J3hBH4Ko8jz1OFoEckqkhXzABt5f8AX8x6YSqVJSIAW8JN2J4Umwve4I4A6efGBCNIVlmBKhtrB7X6eZHJ8ufTnDuOjnuQCAV8YW5DH4ep8/hycBNBZ1dkRWQEC/UA8nj7wvz+7CVRRIRrMWUsxIDEta3W3PHS/Uehvz1wadIo370KVkYtfk3H4enQKeflY4VuUQBhuNxYld4A+fDX/C1/bCbvtWRFuLsNwU2ufQi/P8hbDgUEIhd1QxkBSwCsOfktvTzHW3tgWvJeTvCzuCxLk8cfAfl8ecEZjvUju0t1G6/Px6j2Hl54UgkXxSNZm6qdilju4v4vfjpf8sCRDLPGyspY7d6spcdAVBPsLHkfE+WMVW3WkAXbYgC6kA+V/TgEHywlIyqNu/b92zHn4X9PzwKiwKKgBW3NuT8fIn8/jgQloQO/dAslwSLbl3XAudpHhJt5W5wbu4wAY4wxP3dxtbysOgJHxGApiYXuu3cpVls32COnl+Rt8cHZjazEsHUttACEKT1Hot/I8fDAhEOzvSd7MCeJDfkH1Avcnp8Rxbrgjyd34yhLht3JIAt0Wx5HS+7r5HC1vu3VwPED0BNvF8j/AJjCc4MTXDGOxBHBXYfYAm1vUE3v8cCElH4mEQDqGb9Xstex5288etvXkYUY+Itff0H2uT6Hr069MAOrKXijVbncx+0N3HT71/T1wZ1JFlLleT0v8b+nx6H2OBCBobWMW08kAx32m3UjcByL9enlgEURL+sjYrwNu4rfi/UXtYfw63wYyGxjDPcNu2sob0+x5eXPF8ANkq+G/Nw3B4HkR7dfhzgSpKKIS7iFZxGCeHCkAng88+1xzgyFe8IRm33AAtbx/si/F/7p4PS+D1MRBUSbXO2xLq118xa4Bvbz/fhFQzKSsbluhJIIPxFr2PvgSJSTugXVSikgAcm1vMC/l5XNzbAhXNlELlwOeDu9iQOtvfn1wRRsIDFt1+drBQfMkEA7SPO3GFFjiRDddyFfCS1uL2sTxYg8G/B4wIRViZQY/Em4+Lpu5Hpe1/O3B44IwR2G6Q3QAG569D7enpcYMQEBVH2ArYgD7Sggg+4B/f5jAOQbq8zMbXN+Lelr8/LgEe+BCKUVigIOwXLIlmA9CbcEH18vTCTIxK2sQCRYNdR7Hz/HpwRhUKgH6sRSX56XBB87iwJ9iL+uAKysFTefCCBvuSB6WPIA/HAhJq6F7oWICix3g8W/8VxgCF8JMibRyrdAPe/J+flgrHdwkjHiwRmub+4tcYMrlTcNbaL9bG/n5EH4YEJRY3cAlrMoDXS6sB1BG3i9uthghIjkLWAcggAWIUenJ/f188CZGuWuCtgSAl+vPyPnbrgDIFA3ScL0HG0X6rbpfzv0wqEDygqFC2A4SwBJ9rjz9L+R6+WCMCIr7l63MbAgsenToPiecA6Sqb7QBawYMPzt1+Fh7YFGaIKQrL1tuF7DyN/yv/HCFCGwUgybW56tu6W6H0P4298FkVlb9XuUAWsT5e4PB/d8MAZJEO9WYE8sT1Hvz94fxxllQlV2PsbpbeAfUEc/nbAhFCsSAyrcHoFsCfh6f54DYS9rtYAkgevwJ/jgb7Qd3jA5BZA1+fK/nf8ADBbuHZZywa5J3i3X1P8APAhYybNpW4W3DFePf16flgdpF1MXB4uAN1vl1wC3ZV58Ti20g3HoP+mAZAAGYG1wDZbH4gnj2v5YELG8NwEc7ftcWKn0wix4F9w5PAbj8ucGKuPtRm49G8vmOcCIwysz+LdwOhBH8fTCoSaMzLYW23sVvxe3HB9cHAIAuOLGxsLD4X6Yxo1UAogta1mYXHx9cE+yLuAbG3QAn5fvOBCwX+yQBu4sT5+hvxgoYFbADjzHkPbBr3PRRza5tb2wVnYgjcDY3BYXuPa+BCCwJPgDWF79OPhgrKgttHJHzPxwZSSbBVJBFr+V/UYwguSQFsT0v0t6++BCI53eLk+5HIwm/UsBYXv8MKSm53A7gSfMnBGAB8XoL+H+r4EIpLEkMWAHH+WMC3PQ8cGwwITkAW9hby+HrjDyLtc3PJJ5wIRfIEsTfk88e2Cm3vwR+OFHNiwaxI464SLC/BF/LAhB04BHXCbHp06WwoQbnoSOef4YKy2Xobjr7DAhFCliAOhxmMW9wAT8cZgQj7Rc3Pna/lfGEeh/H+OC7rdef4YPctwVF+lrYEIVsL/iPfGEm58V7c8+fwxlhx8SR7jAHrwRcdPXCIRt1/Mc+Vri38cCF4AUj+v4YKiEtybA+/XAgk25a3Xpx8PbAhFuCCQLceXnhaJQfEQW55GCbVA4N/Q+388HAsADZvY8c+mFQhVBusVPA5PoPng5I8V7gnkgHoPQ/lbBUIDEm5UDqeRb3xm4hfFc28vhhEIQh4LdPX19hb92B8I4PBHn/D+vPBl6+ErdvPyP+XrjCxHpboQTcHCoRdoZ+ALKt7niw9cG5sCNqre3nYH/AKYEptHi2naQCLeK/kfQ/HGG9xbgAWvfz+XXAhCGILAXNza3S/xBxhN7kc7RwSeTzyQPXy/PAfZPiBBHFmB4Hz9D+/AA2UqACW46ci/p8cCEdQVjGxbDrb1HuPQ4xQqkkiQBxcbXAJF/T4/DBDYC5sSR8b+1/wCumDjYrBeSCDdgn8OuBCMW3MSCtjzfki/nfzwZWsL+begt+A/hghJYgvs9blub4M3HHA28MLXKt/G+BCMEDbje7FeAR/X54wobFbWBF7Hm9vM/1xgIt2/wBrgehJHlfCuxfCjEEXJ8PNuMIhJWUXsPB1Ks17j5efwwCICLiM2Hqb3v05wtcKSU4DLcD1sPzPuPfpgoVQTsBuB5C/4+3v8AjgQsC8NusD+J+X8fxwZWNiF8VwCQA12IPnfpbnpgqsNt9vQ+XQk+V8KAtawjMl+ef69vnzhUIm/adyNYfH99/mMYm2Ngzjcqn717Mo8vX+r4OitKwdFNjZrg83639sDtdGLMUvY7mby58z/V74RCKdqqRv4texa5N+h9xgrOdoNrW4W3IX25/HCvX7IZbLYA/tdfkDgojTcoYHk3PmR18/L+WBCSUNdWLNbkAjqf7v8ALC6uFupVhccC/Q/x44/AjCKqVuGJuPfi3sP44VVV6rGngHi3EggeXy9D8jgQjLY2TaW2kg3te3sfTz56fPCZ6b0ZgOhJUKfbpe2DMQB1YA2uC32fP/PByoZ27xuEBYttJt+FjY8YVCRuWO4g+LgsRa7fEYMih7MRzfbfp5X/AOo64USmXxNJtU25u1hbyve/GCspDBXVgOltvIv7+Yv/ANcIhCEF1uzbRdrXO0D19/jz5HpgTuPgVRdjYAoOW+HkfL3weKSW4jSQgMDxew+F+nkeMAFTa292HlbyCefw+HocKhYjsWId28S7tzXFiDcdOh4IuRa974wghzGI1BA8VlH7hyR626eXTCwjcFV3Mrj7INyb+1vU/L88AiNHUrIHUhSGXaTz8PXnj/LCISRJVLkhg/nusDbjcLc/Ow98DGSoVQwXxE+YI9xYfv8AXCzxwKn6p2LG4cWtz67uhB+FxgRGNq2DEgbR149iR0B6i4IwIRCRGlnlXcxuNzcqfX5j8evFsChAV37y32QApux/hbyv8sCrgK7Byse7k2ZQw9Oh5+Ptgdm5mdS5BUKTv8dienle/W3tgQkmN5Nx27hZlHFxf+Hvg6feuLI3VgTuX5Hg4VVSVUxuV7rcw4/Hwnm3S6+98AqBHDxRsoJvtvusR7cc/E/jgQhvZgHkCGM28Fx4vX4/OwxllRbmNBbluOLX546ed/TjpgxRVBIDLdS2zbYi3B9R58nCkcalirpEdwPhsbHi3B6qfTz+POEqlTZVuuzbww6lfTzI/q2MXZuLeFbEcEmyn2NjfnC4jAnAYMst0UHozPa3ht1J44F784sfR/YPrrVubQZJDktfFXVSiSHLIKN6jMpYz98UaHdFH/3k7RRjrc4iknjhpfPhtPgMynsidJ9oVdxQsxXYxuW2rzuLEnpb7xPS3X2xbXZn9HbVeus0/R/6HqpqyECWfK6N0inp4j0lrqiT9RlsXneYmU/ciJx072V/Q0y/TCx5lrjMPqTqpElDlFak1c3HKVGaKO6ph6xUKs/kZh1x0FlmT5RpTLKfTumcqy/LMkpCaimy+hhEUEUlrvKyEkySWsxnlZ3N1uw3WxzOkf2hbDVkWJ3H1OQ8qn/1K2LHop0pvSYDf8a/Og8VAtIfRF0rP2K5zoWbMsievObUmbSJJlExyZH+qyRrEs9xXkW3E1wKvvsRGY/CeSe2H6Jut+zzNY4cvymq3Vd2pMtrJo5pKkDm9BWR2p8xQdQq93UDo0JPOPSPs5z6Kops4nmaSOORsvnglcfbh7ude+v17vfdO9+zu4viXz5FlOc5TV5RmGWUGbZZXMPreU10CTUlZxcgxsNveW+zIu1wbcm98Z2jtOSRUbKM8cPHn6+Ku2nR0bq3Tl8DrZ4LwwrcsmoWljqA8c1MWWeOVTG8Tg8o6EAq1/u2BB8uMa+SAmQrsbaObkC4FufUDm/Hwx6r9r/0J9H69Wer0zUFnhXbT02a1bLPAlrqlNmhVpEUAi0Nas0fkrqORw/2qfRY152e5sMs/Q1dUVDAtHRT0wp6+RB9+OIM0dWtv95SSSe6LjrLPb4p21B6620WLNYnxnrrhVULHEyum5yot0sbj1uB05tc+V7+eFNl0HhG29gCBa3px/R6HGyqaBoWaJ94eFikqEFWjcfdKmzKw6WNjhnJFZSxYBipBJNyPb1/EeR88Xg5Ui2ibmN5CkaFnc327Rcv/dsbX+fJHHNsYsV3XcviKqq+InafIkHjnyGFJITLKAVFj1ZuAF8ifQefnbBUYMjxgK3GxbeIPY9SCOhHp7Xw7NFEfu2KrFsZ7EsARZ7Dqovz18va4wi7Ib8wqJGFrHi5HX1F+h9fTDmOAOSstmUoArg7SCOgsQbkHy6jywUxm5KgbnvJwCp23N7H4i/ra/wwVSIqNwA62IY/rCoPiHQG/Bt5kC5wISN49wDEMQGFuAT5Br8n8MN5AFKMXUNsKXZgwIHkPPz6fnh1TFukscaKVIvIBtYkgAG/W/TjoOfLCpEd41KmRlDMFJC8IWF+AlhYMDfi1iL4IrIjI9mJ3bid3dspAPAI6Hzv06jCxXxt38R37bKWYbmbptI68D73oObnnCageMytvKWS3e82636XA6G46H54RCTQq8W9Etu5G6S1rdeg4Pv06YUDrGo2xD9Zfo5Ive1wCAQeoKnoefbAFzFHs4Qsx+w7Hd8T5keakAedvPBTKW2yFi55te556G4/ony9cKhJysZGMpJKkW3sCd3xv1selzxhPu2B7woVO+/PUkdbkHy87extg4UISYBtnTxckFv+UHhh6g84MzRyXIjXfe5uOp9Q1rsPfqPhhELIRaXlDKFFm2OoNjzx5cGxBA5vY84VjSS25FdZLld0bm4J8unQj14PxGERFC5Q3fcVUhEQNYDjr0Hz6jBTIveqWhWw52KCysOtuouD5j42wIRpfA5VmuOp2tuIPS1z5+3S/wAcIAKXKtY7xx4SQT87bfkScLjY5PdqbqbAxe/lbpb2PTBQIxEAiFXAAP7Li53c36dBa3zwJUMZk7vwx7t1hu6hSOpB8j8f4YCVgzbmO5dwLG56ngeM829T+GAERW6SobEbrFSpNvj1A/DBpWQDfu8ZG31br9kfxBBGBCTk4vHKzAEC4A2gWPHQXK/1zhMptUsV2qbdfIeVvUeVyOMLsoLlgXcyLcNt2eL1ub3+WE+QCfvHgnm5J8r+/wDVsFUIPEkILqxCt+0Ra/pb+PwwgoYvuk+zc3AW5UX6f3vTn1w5WEkhlUeJtt9vP9eoOBEa3IDICW8NjcA/H93PHX0wVQm4iklXegv+0AOb9R8R+74YxVBjVGHmSSp8j0v73+R+NsHZy4ChiQSGtfkG3W3rgWO+LbYWPityAQPOx5J9vL0wqRJiMFSp8YP2SAQCfPpyfnz+GEZSxIXl7WYWe34EYc9S7M7EhetiPQC/HAH/AEwRxGEQsguFtcn8Lj+P44EIFEbnc5AB4AA8j5W9vO/44CRWWHu2L+D7pJsD5XB5H/XBkN223FlHS3T+JvgrNKUCufCguCx6D0B629sCEkW2IIwW5N7hrA+xH8uuA4sxYsptssALD434t/lgJJNpCccjaCQeB6W8vfz+WDKoIu1rjm/PA6eV/wCQwISexfEDe4W4HUgjy9x74BlWwJBHW9rC48yf3cYVYJ3SxhuSdwAsLH5cH2wU/sIwBYjwEi5I9v68sCEk+1gFNyPU+Z9SMAOANjE+Hm4BFr+fmcGVTIQVVbdfbg8k39LjjBzta7OWPiv5c+5Fv3YEJFiQw6jcLEKAL+3v+WEX4+zcAHjof6+A6YcMihiNoFvz/wAsFaJAb7WJB5BPT0+Nv8sCESwK2+163J/A/wBc4w9eVBP4fI4N08XSwtZW4A9vMYzzVVW9+BtFr3+PngQkz0Y7ufM+WAC+G1za126ED3thX7QBuD7gHj5W4wmVt5EH7p4uMKhFYg+QF7GxJ/lgNpF5ORfhbcn3+GBKi5PPkeR+71xjH0YjyBt+VsCEk21TusvTkA8fj5HBL3BWw4PI6f1/DB5ACPLgc+WCWAH3j5XPHzOBCOLC63IF7W6/O/ngjC4tf429PfBr2G02FhbgdcAxPr7W9sCEmABa/OMwUk7gdt79cZhEIxXxHbbw8/DBl4XkeXQ+eMYDkgc+VuhwdL/Hi55/EWwqETxcXP5jnBrEnyt0v0xnFySBx+Y/ngRcNcHkeZ/ccIhFZCCCRb4c2wIsLG/2vbr/ACxlrXFjcdfI3xhuALk8nCoSlwLEWsPIj8cAAAAo+0DwSf34A2Xow4sDze3HtgyDi4FrcX63wIQqo/Z6cC/IOBtwL/Z8rm4Nv4YwC5J5N+eOTf8ArzwdTbqLXuTb1wIQXC2JAIPX0PtfB7oORc2G0bj+7+Xrgqg/8Tm/BJPPHl54Elj683sR5/10wiEKqHU92dzXBA9vP+GDbSOSnC2B3G1vif44IpVifCLWBKg3Fx6XwYKNu4dPI2FvfAhZscLtCsQDx629fh/lgps1tzgC/XqF/jhYIt77TY8iy3PtzggAYFioLC4Fzwfb164EISwe7G3W5JHBv6+/vjAy95+r6g3HNwD7HzwUkW23v0sQbfL0NvTywVTuF9pB9ByG9/b54VCVsw6kAD+6OeehP9cYKwsDz5bSQbix9/69cLIfCwNwbW3X4sRbp1GBZizsT3a7rFrgkDjkcfj8DgQk0DEgWLg+nW/oecKI5cAfaYjw83ubdbH93Q4JtNztvdj4rgH8P5H5YM5tYhVU36A3t68enx+GBCEEAG+0HggC4BPra9vwwMm1wbW+fW38fjhMsO74UIPbpf8Au+mFFJDs5JBI4AG0jjjg4EIoRlUni17dRz6cemHtFSV9fVwUFBTS1NVUypFTwQKXkllYgKqgcsxawHrhofAp/Wr4Rzc8jzvb09cdjfRu7Lsv7Kcg/wC1rW1HPJqesi25PRqQlRQLJHuXu9wstbLH4gxFqeEmVuSitRt9tbYor5xcchtPwNZ1BWbLZjaX3chrPWvYtZl30DNYag079T0Vqiiru0HKvFnmnJjHEp7xdwWhkLD6y0JBjnC8LIQB0IxQetuyjtB7OMylyvX2j8zyaoiJQ9/TuF/MXHxtjranrJ80z+hzaphiNaz/AFNVjiZYoAFJggpwx3LGpVghVt7MxZmMjk4til7ade5pSUulMzGWa2yqrZaePKtT0q16FjwqxzsVnjvaynebG4IBHODZtNuYLsxqdtPjIbM8Fry6LY/FgoPHqvJeYxh2gEeMW+0jCxPxwnskCq5iS6m7C39cH9/GO/M47Cfotdq0S5m+Q6l7K8zqwx+uZW/6VykPcqS8dlnjXcCLlWA8zittUf6OjtoFLPm3ZJqXS/ablcYJvk9en1pR6NAxV1b2642oNKQTjA9evEBZstgliOPx+OBK5KaNNx2NYMbqp4/dx+GDyIkYAiu2yzG0YABP5keVz1xLdXdlfadoOpNNrns71BkkkLFSauhdE9+SB+WIvenLFRKL2sLgrY+RN/nfF5krJBVpqqj43xmjwQkUtu3BiL3LqQG5v79fY9cH5VlaS6up4U3J9iT0IPkflgwTwshNiwBG0g3selx0v19PXBGgl2gsCqkFfEeAfMcdD5+4OHVTUIcJtjXqOOGNh8AenHle2A8AHgkU3ANwD3ZuevI/cLYAROxVU3MXuu0C5t7g/L8DgxQ8nabFQLbw4Jta+70wtUiTK3kVpFY34u5Kgj3Pn8bWPGFZCFW21QAb7iTa49xc/I4TDfrQFsGNlJC9fTCqgbQrbbsNwseCfI/wJwIR1cqoRmYhgdvjIVbdQLcW/dhN0UoWRT1tuAuLjzt6gWuvoARhTu2VLyA/tX8NyLHmxBwVUckttLHgHz6fh08j/wBMIhJwh4m8e82O7xC3Pr6i/r8MKuVtbduKkW6km/xHB9uecLinKxgrYBibLa1z6+gP4Xwb6rJIY9y25G7hjz6E9LfD8cFaIScaTFf1LhOjKd3Jvxf39LfLB5opPC6FQgbbfcGsD+0ep+JwqlIERmkliQEkEu6ixHUG5HP5HqMSHS/Z3rfW1Si6K0bnGdbR/tKKhdotx+0DIQE2245b3wySVkTb7yANpwTmMLzRoqVF4IW7vejRLdj9u9+PP2NrG49cG7qTZGQt2YEFUbcePW3T1B9D8cdI6H+gn2zam7uTPKFcipyAWCRNXzhffuysEf8A6yYYubSv0CtJ5QBUZzmlJXtGQGetkbNGB9qWkaKlB9pamQeoOMyTTNkYKtdXeMuJoD5ElW2WCd5yp68M1w3kWm9Q6lq/qWm8oq8zqEIZ0o4jJ3Y/adh4Ix6l2UAXB4xbHZ39GjWutyzUVFLWxwuBL+hFSpihb/vcwkZKGnt5/rJGHkpx3NlPYz2Y5Ll70FTp183mo5VWmjzru5KJfCGEkdBAsdCg543RzEeZJteXRbpIqdq6q+sRU/6qmgmayL91Y4kFo4iW8IEagABj904wbX+0wBux8vkj0b5rUg0MSLz+fwPnyVEdm30W9J6QiXMM5r4jKFLyw5BM4dkH2jPm8yd+6+q0UMS+knni9Mrh0zpLJZMi05k9BkuVu2+ejy6Du4pnsT3kpuZKl7/endyfmMJam1FluhMobUOrMxyzJaBQX/SecVgoaeST1Qm8sx9FhR2tZQRcnHK/aT9NPSGUxSUfZ7lT6gqANqZpm0D0GWRgdDDSBvrFQB5GV41NhcG2Mtv9oaUBLRRp8QD4k1c7n5K642Wx4HFw86eAFAOS6ZzTWdC2XVWYZjXU2W5bQRWrcwr6haekpgykWeVrIp/uAlzY2TkY5v7WvprZBlEM+WdmVPHmclzfPs2pStFGb3BpaGSzVFjyJamyX5EJ4xyNrrtn132pZ3FUZ3mlbqKsguKKFkCUtGvpT00YEUKj1VQfVj1xD8xFLBN3+p8yOYVg/wDfGkkBVD6SSi4HwW59xjasmhYrOQZO87w9G+7sNwWfPpF8oN3AePqfYKT1vb52hjtGg7S8q13qKm1JEu2TOhmDmqmO69mv4e7tZe6293tFttsdpdh/+kYyfNaWDTnbpQJlMxaMxaqyWmY0rSJ9lquiTmEj/iQcefd4878wr6nMJVmakghhVNkUEUe1EQHoPPr5kknDemqZaaTvaOdonv8AZJ4P8D88bs1gitEQZI3V1jmPRZsdqfE+8w9eGte/Ondb5PnGXUOrNP5rl+b5PmHgGYZZVrU0j+YKyr9lgTwj7WFyCtsb/Ncr05qrJpMmzzLMvzLLqhtz0dXTpNTu/wC13beEP6MhjkHkT1x4UdmPbZr7snzs55oDV+ZaWzFyO+aka9NVD9mena8cq+zKcdydjX+kqyOoSDLO3LR8uVM1kOodLR9/Rv8A3p6BjdB5numt6JjnJ9ET2Q3rM6u44HyORWvHbo5hSYU36vkK+O1r6F/Z32lSPVUbNRVcaFVNY0s3dAD7KVqD61EPRahJ0HrbHIXaZ9ALtf0pGcw0zSnPaAk92GeJZSP+7qEJp5/gTE393HpH2cdpOju03Ko9Q9lmtso1fQxKGkGWVPeVEA/vRNaeEj0ZTb1IxNMtqUikNXQSH/WSxdIhZagfeGzp3i/ejI56jBDpO02U3ZhTxHQPlRLJZYZu8DXrrOq8INRaF1Ho+r/RGqMmr8mrQ3hp8wienkP+FHADD3Ut7Y1dRRKkjLKuxpBdd4KlbdeLHdfgix4OPebWHZb2d9oWVPl+o8ho56WZb7EijeFr+ZglVor/AOEKb45W7R/9GNoTPXkqezvUkmn5HBbuISVgJ/8AlectHb/BInsMbEOmY3fdyx5Z8KrOksBH2Hj1TjReWrLJYXZhsjKqH+6L3tbzBJN1HPxwR3YssjM1xbeS223pYre3tx8sda9oP+ji+kfpHvZsoyHLtTUa8q1BUmnncevczeAn/DIRjnvVvZlrfQtW9PrTRef6eniYhlzDL5VUDyKzAGNvju/LGlHbIJcGuFdmvhnyVR9nkYKlppt1cclCnjfvSwRlLC9mblgffz+WFIoRFJxG8d1O0MlmbjkBTdT6244B88O4aemqbqlTTTHayKkbhrHqCLdOb9epOBmpJKdFcU8kdrFkRNzAftEnqR7WA8vXFmqhomoCyCLwx7iQABe636WbyF/L7PwtgZNsDsFjJsCCDIF5B5N1uODf2PthRV42FNwKk2NiORxyvr1B6e2AMRULs3eEi3PFunw9uP54KpKJIx2YsZQwsFdVdX4HS5HFx5HCW2Zpls0m7cTZTc3HTkW4tzfgDDnudqrG+4d2SttwAWx6dCRxfpxbnCMwVRu5AB+yfKx5HyHr7euFqhIP47oURhfdYOCC37Qvbn52wqVKxyRm8btw6bQj387nk/hb4YOys0zKLuhBuHfcCvQ7r8KPfqOOuDfVnAKgc2HFgLrbg+IE9B+PywVQkHeQRi7IwKg2ViCR+0wHHl7X9MJMit4iCoVCW2kXbnjoeRfi9rjrY4WSAcEMNhv14uR0PsQD16EcHCrRBBvIXcT1JFz+fA+P5YSqKJDahheRwzkDg3ICkEXvbki1xtPnhMrcKwjNmYABmFif2Qfn+7D2Sn2oZXuoUW3SDao4uOTYj2PPzGHuRaQz7VtUtHpjJq/PJeECZbRy1jW9P1S7Rz5kgXw1z2sFXGgSgEmgWrG2NDYKzDctvujyNz/L3wcRrK4AZ3ZlAKsoLW9CRw3HQ/ji0KX6PGv0qEh1I+R6XkmbwxZzmiGrN/2aGl72odv7u0Yvfs3+gXqDM4IKyp0NrjU0fVJcyePR+TcnqZKjvK6Vf8ESk+VsVjboSLzDe8Mv8X2jzKnFnkycKeOfDPkuODFIaiGLZ3k8pEccIW7y+gWNbsT8BcHFm5Z9HHtDaKlr9cw5foCgrbfVZNUzNTVVSG8qfL41esqGP92IX45GPRXsy+hVnOnmT6zrDI9D00gCyUPZvk5p6yUHqsud1++qPu0Sr7WxfGhOyfss7ISuY6O0RR0Ga1zCOTNmD1uaVLOeGmrpy07XvuJ3KPQYyrRpuOMZgeHePH7RxcrUWj5JDSnHDlnyC838z+hHHlnZMuebNZ0ee5rndPlmQ1eosviyqgzJ3p3YRyUzEzUSzyBYaaaoYb5SquqKwOOTMzy2tyuuqcuzCjnpaqklenmp6iLupYJkYq0bo1ijKwIIPIItj3Z1DFkGrUzjK87yKkzOi1DTtR1lHVm6ZlR8osEhJ4NvEjcMjMCDjgj6Vn0Xs2zCvgqMpepznUE8bjIM1mX/AFjVtJAniy+r8hnlLGtgT/57DHcXlSxk0bpWO2VbrG3ZtyHph4JbXYXwAHrr1XCGxdpRrgk/tkg+xBJ/Lj2whILEJdATcuA24j03Hpf4dPnhepj7uxUg7vELCwt/PqLdeOcIFWswCWQG4JXhefu+fy6Y3AVmkIj3ddwfkdFHPB629B8cFFwNpWwPNgu3jyuLD88Kut36NzwAByfgD+7qPhgjXW5sm6xAvcH4jnkj34w5NRrqVCs4ZCLDeQwuPS/Tjyv64QbxM3hFzzx4QT7egt1PwwrJK8Xhl3eFAOoDBfIe3tf4HjGQOEIeRf1YHW/l6j39vlgQkjEy+IRSEgWsklmXjg+4/rjAKAHVNhLAeRBW/naxuMGqLTgKIvFY8dDcet/+mC7Y0PJVVUAgPcAjzuACB8uPTnAhEkCEkEKwILWbz+eExvZSx/2bjbuvwT5j+umFGIY32WCqOfM+5ubX6fIX64A2VjusT05Y3J/ukcH4YEIr3iACELzZrtdRx5k+vTzwKbmttW/Qhib3Pp8P34OrG6Aq248ABrEn1B/q/TBbEKUKhrEghgp+WBCGVUDBdsi887gAR6cX5+PF8A0YVihUAL9qwIHHr1OAZQQF2rtK3HmDz5n+HGDBSXP2TYXsUa1rdCBzgCElzazf4ul7e9sJ8fdYew6fPB3JA8VutiARyfLgYKGAHJXaRe5I5Pv6YVCC6kAhhwTYf154IQLWBBt6ji39e2DOqhSevla3zwAFwb+XPN7n8MCERgWZgLE25uRa3x6YKqjdyDe1h63woSCxY8gg+1/l6YywJG3wgD2HPnhEJIqT0F/Prz+HtgpT1Ysfz+d8LENYbSGJ9PLBSRzbz6A8kfHCoSbKoBFrc/O2E2224vcYU3KeB4QfLqbYBtzHkgg+mBCR5DgqvyxmBud90a2MwIRmv0uLDoAeRgyItgCeD525wkzW4APB+Fzg4It4j064EISCAVBBsfL+eMPmLgm9umAHWxFha/XAkXNh58Ak4EIw9SBYG9vIYDncCvPP9cYKPQX+QwKhgQVPPkL4EJYo203Futvf4fv+Rxl7BlKk2HJPl8+uCC5Xavlz09Pb8cCoCm/AP9dfbAhG4J6XA6/3T/HA3IUXYccEnyPxwHi63Nh7fvwc8DcbWWw3X6enOEQhVjstHyebAte3rjGUECzCwHrfr6ev88CynxAsx+8d3J+eBKngAm3JBHT5X5/o4EIigq4JFl6Hi9/X54XuxK7mJIFreh9PT34wiCpJbuwA1iALm3Hlg4vddl+hv5fH42wISjkGRg6kNfcbKV58yPTn5eeCltxJdRcnoWsfe2AAdgV2kBTb7XBv048vgMCkV9oBBJv1PX5HCoWEA7vsgHnxNxb0v5k4xo2Vd6hu8HHAPiXzuR6evng3kpJuDfhT5+98AqIHBYXUm/s3txhEIgYlQATY2238va/nhQbilwhueLXuwI+OBeMWDDbYr1JvY+1uvtfAGNTZWKNuUEkNfn39DhUI8ZJZSNtiLc+XsQf3fhjNt4w0QJBIsSLc+Y9j8cJwlXfczlT5lgbn+eFXBBBKi55BYC1j7+eBCIFCuLNtLG3A5HxP5euBQBSbrs4IIbyPzwZT4bhzwbdeotjpX6Ov0fsrp6Cn7Ze2OkEWQUqxVmVZXVQNIKxWa0VRND9qSJ3G2CmFmqnB+zCkj4q2m0sssd9/kNZOwKSKJ0zrrVufo1fR5yvIsqg7b+2aBqWgjWOpyLLZYFkkmZ7mGcwPYSSSbSaeBrK21p5dsKeKaag1BVanzFs4q6impo4meOnoxVMzU8bNvYB3Ud9I7eOWa4Z3F7Kioim13r3N9e5hFX5vDPSwRPKaSCWQSrTrIPEHkUbXnlKjvZLBTtCIFiiUHR1bZTklDXag1XU1FBp7KJ1iqZYDeoqqhl3R0FIrXVqqQeLdysSXlk4CK/HSyTW2e87Fx4AbBuGs6zjsC6SJkdki2AddbOKfZxmOS5JpGu1VqJpv0bG01Dl1PGS0ubZqFDx0sIY/YjO2Sol4ESWW/eyLaA5T9Insr1hTCi1JLnujKyRgzTFDmlA7A3szoUqUQnnkSN0uSQDioO1TtUznWGcCpeGGlmEH1HLMroyxpsoodxYU0NySeSWeRrvI7M7Ekk4qc95FIyLKpIPNjcE41rPoaGVnfrXUa8cMqbKhZ8ukZWO7uWzrWvTLR2eLn2Rz5np3Nsr1pBCO/qRkdclVUTKRZpBAwWaOoA6howsoHiAazYkFHm+UjMYpIQJJVQKGPeQTtGegLC0yMD9m/jjYAcg48taTM63LqiOspnkgqIiGjmhdo5FPqGXkfLFuac+lj20ZPFBSV2tGz+khBC02o6SLM0A9BJIDKv8AyuMUrR+zcgxheD44elfQK1FpluUjevNeoWTdruu8kylMvzHXNRmmXxERPHn8EOYQPCbBZHMg3rtJVZPEQFdJR4N4VlnulexbXlVJSdo30cdCVlUwYmfLRLk87W5awTcu4dWF7geKxXnHFmj/AKdNKkSUmsey6Jk6NPkGaMoAsf8A33qhIpHJ8O4CzEdCcXDpf6XfYBnFLDBVaxzLTckRRYIs7yOdEhVeVQzUzyghTyjWDL0BK+HGdJZ9K2Y5E08HfNFcZaLDKNQ4j0oppmf0JfodaqaSOjyjtF0VWqQGWjq4cwhTcLqfNiCOVPRh0viBZ7/o4uyeSSRdJ/SbqKSSNgjR6g03UQ923kruBYH44trL+0rRGsDJPoTtO0dmFWpTu0ps+phIFLAsO5qO6JHXwEWPltJuJ/l9NrClrDUx5NmVRHNHtNVl6GVFHxTcrL5+Y8io6ln9r6QhHfBHje+ackfRWSQ90g8PivNci1P+jA7SJ4JH0124dlucwL4rNWNA3PmQ3S/vxjRyf6Mr6UBDHKoNGZyvS9HqCL8uTjumbNJv1f6Sy+ljmHhJqoJIzf8AaUODtPsHK+gthxTZhkcsoM+U5dFM5sBSRxMtWbXLRADiUC5MfG4coL3XEg/aW0DM/wCn/iozoiE5D1/5LgGp/wBGd9MCJif+y6hqCed0OfU9jfrwel/MYND/AKNj6Y0yCIdlVLH4t93z6nAva3PyGPRWlk03UwLLFJQ73G6PYSI6hf2o7MPgyfaU34OHKVeSRd3voVcOt7By1he1777EX6W5PNwCMTD9p5qVoOP4UZ0PHt9fleetL/ovfpZ1DXqdL6ZyxHG0/WtSxtt6ciy3vwMbmh/0VfbjPMFz7tG7NcmiuCxavaVh8ALfhexx3vVHT8c8ciUtHs6w1HdiVGcC7Ip5s46lWtccqSOMBBPSfW1tQxVKzhliaGlAVZQhZQ21LqrAEbr33AdCcRP/AGnmd3RgfEf8U5uh4hicfI/8lyRlX+ix0jTwhdZfScy24F5I8lydXNvOxLE9fb0xP9N/6O36LOm135hmuvdVyxWGyWZaKF2sTtsNvUA+vF/THQDU2bikjqTkFYYGkDOYaOZWB+4SGCrsBueOhIJvYY0uoNedm+jafvNb680lkWze+3M8+pIXDGw4UOz9AB9ndycVnaT0jajRlSN172oOSk+jssOLiB43ffHmonk3YF9HzRzpHo7sB03TzLYLU18n1uTdzz9gnja33h0xMIHqqCKakoKfLaQrTkRx0tAqGJ5PDCwMhcqftOBwbLipNQ/TL+jDkykN2s/pkxcGLTuT1tY5Iv8AZlYRotyTyG+fOKk1H/pJuzPJQ8GhuyvP8yIdmV83zWny2HcerGOATSsT5ln3H1GIho3SVpN5zaHaaA/KcbVZYhRuI3VP4XUNdBXSrG9T9YzAltqyVcrShAByxLkgewQXLHiwBOGWYVktHSfpBBNHRxNN3mYzuKelg7shSzTyskUfi3H7fO0WBGPP7W/+kc7dM93w6cn05o2mP2RkuVLNUAf/AC1WGR7+4UY571j2r687TswNXqnP881VW3uj5nWTVxT/AAox7tB8FFsaEX7NSO708nD5PwVA/SwGEbOPwPlejusvpZfR80i0/f6yq9a5jHx9R0nB9YhB/Zeum2U6L7Rhx58m2OadefT717PNVw9m2U5NoWGqlMr1UJOcZuzEbf8AzqoHdQeHgCGNQBxjlSpizBwDnmcwUijgRM/eyD2EacD52w2StyenfZQ5ZLmEvlJWN4L+0aH95ONuyaJstm/hNqdufM4cFnT22aXCR1N2XIY8VIM+1rrbtCzyXOa+uzXPs0fmTMcyqnq5wPeWQkIPYWGNLUpllO5lzvNGzKp6/VqN7oD/AH5Tx/4QfjhKqkznMwIqupCQqARDGAkaA8jwLYYXpcqp4LFozK/Ugmx2kfkfT368Y2G2dxNXGnhnx+KeKoGYDIV8cuHzVJvmGZ1sRpKZI8uomHMNMNocf3j9p/mThanyeKNiFjDc2DEEg8+3l8MbP6uIpjGrRs6IhuV2ld67l3D7rW8ulzgxMULlZ43Xadx2HabdbgdA3l8vLFhkbYxRoULnl5q5MVy1JHKm67owxJYHq5HB9OOuNdX5QVNwpYk26W5+Pri1dOdnuV6p7O851nluu8gpc2o6+GKLK81roaM1dGUZ3eN3IUTrIovE1i6G63IsYOVCKWV3ktch2BsB62628/fgnAyRslQ05YJXMcyl4ZqKPTVMCqGjLow3BXHl6j8DgaWrancSUVS0Dj7rNwfn5/PEoGWmoEVNMv6ySNVPNyLrt+Y8wenONTNk8c0UbqAGZAWI8j/1wrmB2CGvLcU909rfOdLZzBn+WV2Y5NmsDbocyyupelqEPqHQi+Ouuy3/AEnXbxpiFMq1q2S9peVi25c2T6lmigfs1UVtzDyLBj744nmpKyiO2NiUPkRcH5YREtM9u9ieFv2ozx/4T/A4qy2VrxdIqOI4fCsMtBBrrXtV2Y/6Rn6L3aIY6DUufZp2d5tUIA0Go4L0vej7y1UV0IPQ7wu4c8HHUOns9ynPMrXN8hzWkzjLpAGWuyypSsp2B/vxE/8AsgD63x84dNW5lGuykrI6qP8A4Tcm3+Fv4YkuiO1DVHZ7mq5ppbUWf6Ur1YMtRklfJRtx+0qnaw+IxkzaHicS6PA7vg+yuMtbqUJqN/yvooWu7g7qad4lcFw0LkKR5t4eOPO63GE6+teaIpVw0VbSOo3x1dMrgC/2ri11Prb8emPHzs//ANJz9JXSxp4M71Pp3XtHA25YtRZYIakX44qINrXtxc3v546M0P8A6VrQNZSRU3aB2I6iyaQMW+sadzGHMYUJ6lY5djhT5ryDig/R1oYC1jxTZUjkcFOJ4zi5uPHmurtVfRm+jhr6okk1b2DaXqKmZNzTU1KkcknrYqFJYfG9jcXGKjz/AP0Yv0WM/eSfI5ddaaeQ3AoM1Z4Yz6BJA4t7Y3Omf9IL9ETVKxUs/afHkExACQagyiqoinqGazRkf8wt5HFx6W7ZOyPWibtJ9qejM+LAbVpNQ0ssvwN2DH2JG7yN8Pjba4D9ppuwH+WiY90MmBPH/wDtVceZ7/olskqHlXR30hnUwtsaHOMhgmeNrDws0bIb2t1Hpiv84/0TvbXRtMuR9o+gs4VrkB1nomJIsbgbhz+/nrj02ozW1lW8sdLVd0sMR72ILJFPcEWBF/EtrG/UFbYc11QuW0s0n6MNQ6NHK6NH4mjJ2sQSPunn4Ys/2hM0FzqgD28QSouwYSAKE9bCF5L1n+i4+lPF46TL9Izm9/1eobA8384x5+vqca6X/Rg/S0IdBorTsgPTu9RoLfivTy+HGPXcTQtJMwp403qC0auCyEeZA6ow6kdMFpqxYWWI0E0imUIGXxGJGHBcX8Sg8XHQWwo0o4kY+n/FIbMKVoOf/JeReV/6Lv6XNROVr9PaXohcWafPlbp0JCKbm3mcSuk/0TP0gahh9c1NoShRuWtJUTgfC1seqdQKYRNG88asb8rLsYehU+TDyP44QiaCS31mtpmkZCkjtKg3HoHsDYX4v73thXaQeTS8Qd133aUNgbndHP8A5Lzdyj/RB6vbY2o+23JaZL+IUWVC4B62MkmJxk/+iR7JaGTdqHtQ1LmTEAlYHhgW3wjjJt88dywyCry6BZqdmqWgcymnTcV2Gxbgc3IFh1wjVQLQQy1dXUClp3BlLTukSKdvQd4QADYcHgXOKclqtElLjnGu+n+kNUzY4m4OAHl8krm/IP8AR7/Rs0lNDLSafq5nRFU1FZTx1czv695UlwOgsoQdOMWpQfR47K6CjVP0NVV8BXwRV2ZzGnIHpDCUj+ItYYaao+kd2BaIt/bXtn0NlUqDxRNn0U8oPnZYtz/x/LFSas/0k30VdOwvHkeotTarkUkoMhyGVYwo+yglqNiheOTY3JvhGRTvN8s8znxdUpXSMaLodwy5YLojT2l9G6QjKaY0zk+SbbLJHlOXRUsjMfsruUbyW8rte1yeMbCrknLJsjDTOCGIcu4PoOCzfH2Hrjzw1z/peHgBh7PexChpnG4x1mpc6M0gZhYsYaYdbcfb6cY5k7Sf9Il9JjXolpZu1apyKhlBU0WmKSPKo9p8jKN0x/8AEMSu0dNaP4sleJ668mC0Mi+xvoF67a/7QdFdl2WSZv2na+yLR9M62D5vVokkqtYHu4QTK7gbwLLyTjj3tv8A9Kl2YZGlTl/Y1o/MNV1i3SDMs4Vsvy5G5CssX+3mA4Nj3YPrjzGzTVWe53Uy5pMKmqqpOZK6qkeeZifWaUlvzxGalpZ5C9TVq8jH9rd+J6Yt2XQ1mZg8V8fgKKW2zHEGnh8r3T7AO3DRX0j+z2XNqChp6POsrjiodT5FI1zl9QVsGHN3pJbExyD7PQkEYnGodKZHqfTtbpPV9FPV5VXbC4eUxVEc0RDQzRzizR1MTBWiqF5BUbvXHiZ2A9vGvOxPtAy3WOkcxSnznLl+qmOqJNNmVGSN9DUr96NgPCeqmxBFgR7G9lvbD2d/SH7OaTXego441jkWmzPKqog1WS1tuaaXzVSb91L9lhx7DF0jYX6Pf9TZ8KZ01bxuPJadltLbSOxkxrlX0O9cG/S0+jJn9NqCr1Fl0X17PKtJ695aamEUeqaeIbp6yGFRaLM4R4qyjUfrADUwggyKvHT04Y+FA+8AqUVSWB6EEdfkce4udafyPWWUVGltTUsstLPLHURyU8xp6qGeI3hq6aXrT1kLWKSD/C11aw89vpTfRczjT+azZ9lNFFV5hViesYZfSdxTaghju01dRwrxDVxgE1mXjlTungDRl1Td0RpVluZT+YZjrV6LNt9idZ311HriuOyAbrcMQSLDm9ubWH5nBWUkWYbeQtzwefK3kP5jDmVNv+8jKtZlKuNtvIg9B7YRkjUXUheV54HB8jxwR6+uNsFZhCDY0JN/A0fPJsyn0t6/icFA2hiU69bAm/vfrhUqqtGt7AgML+K3kebG3PBt7YTvGAi7BYEkXU7r+Yv1t52OHJEmLMSFYsApB2t5W55HX3H44BiykLYg2B4XbusOTx5j1/LCsjRuni2m3hHgA48rcbR6evxwkqtJdeAG+16cevwsbkc/LAhJKCzLdlBNyu0c38rKOCT6/HBthHiWO9x4owhIBHkbenketsGjjD2MasDfixPItcm/r04+eFCWY2+tMoYgbQx3W+dtw9ib4EJvtdkB2Gx4BTj8LdPcHrgWAayd2oVW459fMt5k4WaKNlIKbiTsNzy48t3Nj8PIYQs0dySoC+Hgci/l8OPhgQgFkbkcEndtAUE+VyOfywRHe5Qc2v0549QPbBiLkG5uvJHU/iOo/MYOEA2FmsCObcj87YVCQK2jUlUUX43Xvb09P664SIJu4PFjcnoR74ckx7iI1FmPFm/l+78MIkIA12G4Hi1zcelvPAhJu903AkevHT3v0wYqQAfNvXj8L4LtawZVNwPYnA2Z2uCffg3/AAwiEYBjuRd+77wHn8j+/BL2sptYG1rX/wCuAFum0dbWvxbBmFgQCAbegB59T/LCoRFdioVWsDxa/BwUeJLL5D1HOD7iFK7gt/LoP69sBbebqAb+VufngQkzySbFvMljb8cEK2UEk8eo64U2K27nj7tvI++CyKbeVxz7/wDXAhI2LeY64zAi9ywFifLqLYzAhDuI8/F06eWCngX8J9sGYksePf44GwK8n2scCEHBsQBx8RjLWII+XOABswPB5t1/fhQXNrG5PAsMCEA9xaw/PBrWXlePMX5v64MVA5VfOwseo/njLAHwkG4uObG3r8fzwIWEXUki46nz2jAbbfaZRxZrA3HufW+DnluTu3WIN7X+P9eeAV7bRuHrYnkA+WEQgUX8RBNuOvT+WFALA2AI4Nzxz6/Dy/PAXAsUAaw636G/kcGDAqv3Sb2uo6+5wIQAMv3FUeYIHB9h/LAr4ioBsV6MxsPn/PGdLHdza/S9vKx9R/lg4stkBAFyCb8X+Pl8/PAhEYK12NlW1gPID0wcbSLC5C9fS5PW3ytfANZSF2Ai973vY/L938cY2wNc2sBbde4A9b+YwIQi5BJFvIkjj8+vywqNvBJJBJJF+vt+fX0OEFLs21Sbkcc9Ldf5/PCnjtcrx8LXP/T0wIQBzHZG5BHA+8v+E+Y9sGiJPmtlPI+7z6/zwHdFm62HHN/LywVTu2hlJU8W4/d64VCcPzwbCROLmw3fHyv79DhCZrubnaAOhF9vx/njHZlO02IB2/4R8f6GBjjTvAG2j2PFvj5D44RCyIG20feW9wLm1+o9PjgTazEkWHiseh9TheCnLTKiq7Mx2hQp3FjwBYckk8C179Bjrnsd+jrkXZLltJ2o9u9OjZuz7sn03JEk8kUygPeaJzslqVBVu6f9TTgh6kltsDVbXbI7I28/M5DWetZ1KaGF8zqNWk+jz9HnIsooIe17t0o+5ymNYp8oyOppzI1YzgtDNUQXDSK+0mCkuGnsXcpTqztOtaa4zzX+YvnNdTzU9FQz97R06yd8tOXAUyyyqAr1DooVpCAgUCGJY4ksW2t9ZZ1rjNY9TZ6Gjp3776lSRzGSOJCf1tnexlmkKr3s8lmk8NgkapGNNl1RQ5Vk1frTUmYy5Hp6jmemqa6CMNU1NS43Ghoka3f1LC11bwRraSWwAVuVlllt0tcycNw3D3OtbsUMdkZVy20c+W/obMNS6pzqTKtPZQ0aZjmQiWSRWcXipKWE2E1XKF/VQiyqAZZdkSc84drva7mGtM3gZaD9HUeXxvS5DkcU5mjyqndrsWc8zVMreOadhud/RQqql2s9rOa6xzCjp46JMsoMrV4cjyCnmaaHK45Dd5Hc8z1cpAaWdvE7W+yqqq1kEk71kWTvKiT/AGst77b9QD6+p+WNqxWFsTaZ1zO38eqz7VanSGpwpq2fn0SqOwEkaPunmuJ5up5+4vt6nz6DjDaSjBYgLY2HAw7giENlB8J46gE89Rf0Iws6hZVUjbuV1HHQizD+ONlrA3JZrnErUmGaIkKxsCR7YKQR9uEH3HGNrLCpjbaLEi1gbgnBpKcSAWW282B/h8R/XXC0CQOK0wEBPDuh9xfC8VVVxcQ1pt6FyPyOHUlBE43gADzt5e38sJNltywuUAPUjgfH0whaDmnB1EJq6kjdPRRy3+8UBxsMo1hnWnpVnyTNMzymVTcPQ1ktOQfUbGGNT9SnWwQkE+WM7uuUAiRyP8WGdmE6+rWyz6VHb7lKotB27a/iRBYRtnk8iAfB2YHEoo/pzfSdoYwtL265+Stivf01LPyOQbvETcHzxQDNWDhrN/yg/wAMB3s3V6aJv/Vj+GIX2WNxq5o5fCe2WmAPXFdGp9P/AOlSpYy9rUE5Zt7mbT+WsWb1N6fk+/XB/wD9IP8AStUloO1WlhLdTFp/LlJ/+gY5wE5A5oYD/wAh/ngRU2JH6Op7jrdG/nhv0cP/AGxwCO1d/V6roSp+nz9LWqQpL2855Ah520lPS04/+hwjEczL6Xn0js5JGafSD7QZUPVEz2eNT8kKjFQitl42UFKPT9SD+/CqZlm0Q/Ud3ECfuQoP4YcILv2tA4D2Rfacz1xUkzPtE1dqSRpM31JqDOHf7Rq8wqakt8d7kY1hjzlgO4yU04YX3NAkV/fcQP34ZNmWoJF2HMqkBhfaspUW+WG8lFWTXknkdiBcliTg7F5ONOZ+ECVjRh8fK2ZgrHA+v55l8AHUPUGVh8l3YQZ8liJE2a1lWP2aeERA/Nj/AAw2jysnaXawNgT5Lfofhxh1DlSIVEsZFyPK9/WxHBw4QHW7hT9eaQzYYDr05IEzWgiYDL9OU27pvqWedvwJC/lg5rdQ1Y+rvXSQxHgxQgRrb02pYYd0mXwtHdwQuw2Yi4JtwCByL+RHqLi3OH9LTOxWlWDve8HdgxKe8Nx6Dhj8OT8Thws8YxpXxx9U0yvOFeGHotNQ5ErEd4m7zN7/AND542CZckcrJFAAYQCylejHmxHsAAR/eOH9PQzsv+sIX2cSbRe5Btex63468db84BpkkZqYpABHEhHeKu4KwJHLC/HQDyvidRpjUyx0ifWVhZlAUSIG5ZB6Nbhl9ehU89MbGRKDvt1CaiSnEasv1pFhk2kchgpYcEmzAnyuLcYbzPQxFYpK2mtfa6mZDdCCptz6N526YNSiGeOFmmgmJjXeVkBKsPDyOt+L/nfm2E11Qj1dFSQ5pVVULlzI+4GxTuwQCysAeW3XFr28xcEYLXQF6aWmEjgOD41lZvkbk3HuP4HDmGjnYSMsTuxCMGCMwIMdj5EHlQD8cKS5ZnLyOgyqsfaQCUo5SCbXuGC/C1vfAhaqhy5IYgiTTNSCVXellbbecIbMSOOFJseRz6Y2AE4imIZ5Ngu3FnUHzYDpxezC689fLG2pchzL6pLJJkuaIplS++gmu1kI48FvmLeXAwnPkdcIpXko62NoUaWMmlmQhwOikp97kFTYG9+uDAZJcSmFBDA1TTw1veSUkcqieOGYJIybhvETEHu3I6Hlb9RhCmo4u52Uh/UjxIZ23BVY3CsRY9COQODe4t02Sq0bd2O7WO4I3BkNrg9GXi3Swtzz0wwSMUaRxtW0hjKM0bCZFcLuI2MSQVYenBI5BscJrQmNRTTS7VCrG6zpHLE/2yeSRb9kBSCehPywlPkUAjcmPcEViRe5G3n5cY3U8VOkSyrW0WyFhIpWojb+6QoDG91NyB12+Rw4ekqamhnWKCJlMEjtJDLvJUKSdvHAPqbkjjChIoRW5A8ElkuOvvyMNlqMxph3feCWMfckUOPwbpidy5VPIymVJgGjDbpI25BAsbgdPe3secMXyaKZu7sreEsTG6swIHHh4IF+vHHOEc0OzShxbkoj9boJD/rWVhCfvU8hT8jcYcQTUStupM7qaYjoJob/AJqT+7G4lyDvLlYgQnilCkWjB++P7vv08sNZtNQ/7tw3BK7b2a3uemIzCDgCfX1qpBKRjTryoloM7z9UCQ5lRViX4QygE/J7Yc/peqA35hoyKUftpAG/Ox/fjRNkMpj76Mkp5tbhW81PofP3HOE0y7MKfxQySp15QkXt8MV3WNpyA5j0I9FMLW4Z19fUKZ5X2oZrp2QTZHnGpMjkHQ0Oa1NMR8ke35YnWS/TD7fMkUJlX0h+0mlQdEbPJJ1HykvilzUZ9GgJr52Xp423AfjgDWZvf9fDBKevip42/hg+neBSp/xH3CTtmE1I5fldH0v+kF+ldQkCi+kZqSQAcd/RUcp/FoicbWH/AEk/0wYV47eatwOneZHl5v8A/QscsNV1B+3lNEx/9Bb9xGEhVF7hcppLjrZG/wDdsK2Fwy9khew7Oa6tn/0lv0wZU2/9t7r/AHlyDLwfx7vGqn/0hX0tKos1R9IbOo1PRYMvoYj+Kw8Y5nWpckhcppbj1jb+eFVlrXuY8vpEAFyRAvA+JwGEnoJA9g6KuzP/AKZX0itQqYs4+kDr2pjPVIs6emQ/KELit841vn2p2aTM6rOM4kkPiesrKmrLfHvGI/LEdWXOrFkdI1Av4I1X9w64U7rO518dbVNfgAM1ifT0wdg7L3PwEvasGrl+VsI6nOaYAw5XBRAfedI4bfM84Z1dZUzsTX59B8EZpT+XH54TOn6lnCuSWJsxJvY+l8PqbTCkhplYAkAKyncxPoDb+WHCzgGp9PmqQz7OuFFp5JstHQ1dSf7xWNT8hc4Vp6qqJAy/LqeH0fu97f8Aia+N7FkKxSRutHujQESFeOWF1PyW3xucP4cpCuE2XP2fGqWTi/Hp6m/HqcTdk3Xj1wUZkdqUXloq2qINbVyyuRcC99o9bH9wwk2UlQykXIHUHryen4YlSQw7FYSR2YXZhexPS4JHKn1Hw6c4bCMNUTl1sFSIcta193T158v5YkADRQKMknNRgr3gFNUuFkTiGU8XHkre3ofL4YtbsC7ftddiWvKfVmlqxVzJFFJW0FUT9Uzmj+9SVKjrf7r9VaxBBGK1zSjLF9o4Unk+vp/XwwwhKVFqeocJIvEcjHj/AAsfT38vhivNC1zccuuW0KeKQtPXVdhXud2Wdqui+3nQ9L2gdnddIYA4gzGgqWBrcmrdp/1eoHmeTslttlUetxjeZ1lWRa009PpTV8FTU0s0kTu3eilqIaqI/qaymnj/ANhUxsAY5lsL3RuDjxu7B+3/AF52G65g1VpSuSDM4lFLXUVVc0ecUlxupapfMH7r9VNiCCAcetXZD2w6D+kBopddaAlmp5qeRYM3yipZTWZTVMLdzN5MrWPdy22Srw1nFscLpHR0ui5vq7Jg3WNnyCukstqZa2dhPidu34K4j+kp9EjUmS6ieuyWnSvzPMDNPSyUVIIIdTBAWlaGBfDT5rGoLz0QsswDS0/3o8ciTxd21g6yKRwVBA+A9/UfI849vs4odO6nySp0tq3LhX5bXhe/pWkaIlomvHLFKv6yGaJgGSVSJImA6ryOHPpWfRQzGPMW1NllYtVWZtU7KPO5Ujp4c+nYXFJX7bRUmbn7svhp67qCk1w/RaJ0zHbm3XYOGY9xu9PBZVusDrO6oxHXW9cQqpBukpTuyFW6ji/3Qvp1vb3OG8rc7rqxDdNxNvibf542Vdl9Zlk89DX0k1NU08jwTQSxvHNE6GzxujW2spuCrAEHjDBlJspfcByDY26eV/P2xvNNVlEUSDOyBltIreZIseeo48j5/lhSzlC1pH7teGsX2G9+o4A69PzwkVAs1k8XFm6Nf4cdeh45w5SK97rfZwWZdvI9f2fb91sOSJGzGVjIjuLABCSbnr1Pl58dfLAhZXG50Vw58Ujn73S1z0t0t54V7slIyx3Rxiw9U5vax5+X4HCcyREg2Dg3uxHhv6Hzvbr6e+BCMNwUiSzK62YMxAIBvtJHTn146YCyS3O47WFwoUXPHkRwb/h7YIDGSr2O4gksPYeYuPLzHXzF8KRBZEJZEc2uqi+5vUkLY8fjgQm8sbRNZ0J3c8C+74HzP+eCAsWIQg7lKyEHdZD1v7+mFwkcyOI6dpCGXeF5sOoPHJN/M/DBQjJ+t3Na5tfpbz4t8B/0wqEkIyjDcDe3VfCbYQYBbLt4IsOOG9/8vLDt7hQCCd4sGP3jxe388JTx7ieFHPjPW5+AP54EJJlW9uhHqOfbjy87+eCXAsCeeQt+pHl/1wpLZHtZRfgqL2+V8EIv5DmzAXvf3+OBCJ4SeX6WHn+FsGsCLWPS/Xnr+GBsBcXuQSDc2H4+ZwJQ3J3c2FwR6ed+mBCSZeo28DkkHn5jywAF+eACLXIt8vngzKpuSb3PU9fxxljcg23N+z6e+BCJKem57m1j5kAdPlgkg5I22N+R0thR+QBb4g9P5nBWW4uB19b3JwISB8RJ3WvySOL4zCgTqr/P3xmEQk2BtcrYfv8ATGW224I4HTjjBiD9q/uMYCAQAOvFhhULOQbci9gB5nBvtct09wfwwUC4PhFh1FsDtAvyCfUG9j/LAhGBLKBbk/mcYeTtBFhze1ucZbwC7AjrYE8fj+/BjcnkBS3JB4t6WwIQFTydxN/fz+GDC5VQGJANwLkAXwVrm5uBxb5eg9MH3nbztubliVF8CECm1mYuT03dbeR6+XwwcAAgF1APFmI/H+uowmbs25mBJseT0Hp8PhhS7WLbWNja9jz7evH7sIhAAdwZm8RH2j6/DzGMhvusuwet/EB8CDgTuHDhVPJHrtHp7/wxkZtKrIyoxsQSeB8f688KhGeMsCFF79eOn9euCFHLtIIiwPJJ459sKgkoqG6lbgX6/wDX1HpgVZAFkHJUkG9mtfi/PJHp6YRCSEYVAxU2vY8XHqBfChKAAhb35F16/H4fvwZkCb4wWBUDi/2iD/AG4wnuHG7aPIC/X1HOBCOzEbXEikNwLny9CP4fhhOIKJCWbwtz0ud3l1wLI7ESckni5HoAOR6YPHYObp4SDwLH5i+BCNtI8TC1mANxcC/t7emNvpHSWo9aagodL6UyarzXN8xl7qkoqWPfLM/U2HkAOSxsqgEkgAnEi7LeyHV/azmM1NkUUNHlmXBGzPOK8tHQ5cjmymVwCWkc8RwoGlkYhUVicdV5XBovsByKp0n2fZc1VnNUppc3zPMoE+sTsOsdTGCVRAbFcuViimzVru1qYZ9rt7LPVjcXch4/GZVmCzOmx1dZIvZ/2Zdn/wBGCgpNTakmo9U9o9bD32XRUtR/qtCD4S9NIASEBurV9rsQY6RWO6dYjn2oc/1Znj5tnlcamseN4YwkOyOJFN1p4IluIolLFggv4rsxdyWw0kzCpzHN5c6zmqnqZnl7/MKqpqDI9SALGaRzYqwXbxwFVSqhU2qHeq9Q5H2YUozPW+Vx5hqDMIxU5TpNS1O7xtcpU5iy7XpaTaQVhBWaYHnuo23NzbjLapa5k5nrIDZxxWyxsdmZs69UepzDIdM6Zg1dryeop8jleVcto6QhK/PJ1IEkNETcRxBlAlqiDHGeEEkllXnrtN7X8/7Qs7gqZ6ekpvqMJosnyjLkKUGS0t79zTISTuJuzysS7sWZ2ZiTjV9oPaRqftE1FU55nubfpHNZ41gepCLFT0lOgslPTRqAkMCL4VRQFA6D1icYsjwUZJUi0kvQyew9F/f5+mN2x2FsTaUqT1w9Vm2i0mQ1Orrj6I3ETskc3eTyX72e9wL9VU/vb8MKxRrEuzfsRuQbXA9QbeXvjO5EcdrDw2Kn9ocEX/nh+tBWSwSVVBSyVMVMhqJniUsIYwwBMgHKAbgpJ45HW+NZrQ1UHOLk2QFQI+BuAFri3XzwvlGYZfFntMuYZelbSqxjmp2YoZd3hsjryjD7rdAbHkXBQ3OVBaVrRp3aBgLhRew9+pw2EUc0xLqTZCFI6bxyL+wwrhUUSNNDVbDMjSfpCc0QqEplkZacVDoZEQNwHKAKT6kAA/DAwgSOsUe0s5uFLAG4HqeDxcdefjhKSPwmaJQEYmyqeU4uQR1FueelvPBT3iMrd34lXcDYeXN78g/P/LCpEodix79oTzV77SfVT63Bva2AKqzXQ79se7cvQ2axHuNpxnfQylpi5kdiSwQBbAm58rX9gLe+ABR5LIpbZZVK+Djr09cCEewZVVRuBNjtUk29z6ew64P3BCXCm28re3R/2T6X6j15GAUMUMge5VAeevUjmx9xz74xxZLKgWNluRvIVrHi9/Y3APvgQk3hUKr2UkN9nzI87jr88KJSpuZWQKI/Ff7W6/FvgP3432ktIaq1xWvQaO0xm+oKooFanyvL5q12543dyrbf8RPHwJxYtT9FnthyaOKt1fkmW6Jp3cL3+q9R5flSpfqNjy96x9QEv7YQuAzKACclTj0qm7Ar4bbj+z5An0HlfyuL4ARJDbfDyGCEAeIN0BHr6EY6i0X9DjINVUTV03bplua0SyGGRtH6cr84QuAC8Yqplp6a44udxHTEzpuwT6HehpmOvNX5/WmnA2rnWrMtyxXa3/wPQiaf5Fx6XxRn0nZbObrnVOwYngFZjsc0oqBhtOA5ri5MuUqXENk5PiuoB8xfyB8vQ+2Bgy4zTtDl96pmICRw/rZCT93agNyPhjtSbtl+gX2fKy6Z0HpfMquEXjl/QNXnMhYdP1uYSbOvmFwV/wDSFS90lD2X9n+q4Y/s91lhp8oiK+myhhuB878e+K50q5wrHA8jeA0f5iFN9CBg6RvlU+gXN+m/o3dveqwsuR9ieuqyFhuWoTIaiOME9DvlVEI9QT09MSU/RE7ZqCJZ9V0GkdJIHs7ai1lllDb2MTTs9x7C/tiXat+k325amoYqWo7Ostihiuxm1JmM9ZIxPVmNVMqk/wDLb2xXjds3bFBJBHT9oWj9Nx0zySQRZTQ0oMDSWLshghcqTYXIbmwwjLfaJftY0f8A1X/SCldZI2ZuPCnqQpFR/Rdo6kLBmH0gOzYMj7zBlUWcZxIvHIvS0RRhze2/r0xJIfoz9mOQ0D5nWdpuss3UsaeemynRC0Ti4+2BmlVC20DpLt4vipcz7TteZzI0WovpD63r1C7mjp5KrYV9hJNGLf8ALa2IzmMWkq51lr6jV+azG/jq6uBbkdf+Iw+HXg+mLDZLTJu8j73VAWRN3+Y9qq+27Pvo2ZPC8hy7XGZNDGm2HNtc5BlDs1wCNsEdQVABvfdc8jCkOa/RkyKRw2jtAm0LNH+kde59mriUDwhhRQwJ16kE28hjnymj0jTklNFxTkC4+t5rM4P/AMLVP3+Y9cLrnWVoWSl0TpuHwoylqaae4YXBBklP7vI+mJBFO77nEeAA9ym3425N9fwr3pe1f6PWTzFodK9iqgc3Ojc/zhrnrzWViA4NL9Jzszy6op/0S2j6GnRyZlyTsUyhZHFuArVdQ4HPqMc/jPKtpjJT0OT0vQjuMqplsOh6ob2PB58wcbuj11renYJludVMG1RuSmghSx8yNsfQi1r9OhwroZdTifMD/aUCRhzHXFdB5b9M/JqFN2U1uv6lz0FDorSdAnyCUUpH4420P02NWVFP9Wp9G9seZqVZDLHmlHQvzexX6pli2I4I58sc6f8AaZ2kVMSOnaVniyISvdx5vJFyD1CoRwfTyNxzwcOo8z7Ws/3071+u65ZFBV4ZcynW9+q+oHr58jjg4pSWZzsJHU/+3e11WGSgfaCf/kflXs/0tO3OqiEeW6Y+kPZV2ju9X1CA+57ugGNjSfSy+kVO9Oy9lnblNLBGEuNd5wO8t95kENiT5+XtigqfRfbLVRm2mu0CoA/Yy/NHsfTleVPyI6G/XD1uxntcld2k7ONb1Ctwr/oGvcHjr4owbHr64hNkidnIB/8AT/8AmE/tH6mHgP8Aiujh9MX6SqfqW7Fe2Jiq3YPq7M1Iv0N+6GCUv0yfpF0FVJU/9jPbFM0sLBY5NaZnKqC4BkVO56g+ZuOemOX6jsn7SYZHSbs31xHZ7qx0/W+JbdGHdjofPz46eZf7F6ry10hrtN59Ry7Wt9ZpZ6Vip6jx7b/K1/QnEf8AZkDsGyDH/wAn/wD/AEUgtMzcSw8B/wAV1ND9OLtoiVUzXsW7Y5ZFFmZ86qZLn18dIRhCf6dGeRsr552PdqkYCusvevQShiWurHv8tblV8Nr2PU845SqTqrT7ok2ZZ7RbmHdLJV1MDH+4pZhuUjpa5Vh6HDqn1/rWjdkptfaqgK3UrFnNYg9rqJL/ABA5HUA2sXjRBYag/wCeQf7imnSDnCn+1v8AxXQ+dfTZ7OMyRRmml9YZWwPjWp0RpWvVh6EyUkTfO4xoT9J/6PmZTDvdP6SZSLFs87GMqlPzajrIz+AxUdT2r9qNJEtRB2maoeMOrMXzWSUmMjod5YXtcgjhgOOQQHEHajr+p+tR1mopqyVY3ZXq6ejnRWHmVkhN15B+dybci1FZZ423Qf8AOT6tKgdNG81P+kexCtPLO2j6LmffUqbUnZl2PKzPKa2SmyPUOToiBwI+6NLUPd2UklSqqpFgzdcSzKofoO6vMaVGWaaypUqWRGyvtNzXL2jh+65jzKnkBY/shxbFALrzNKtz+lNMaHrztB3VWj8v3dQOWiRDf4Hny64Sq8w0zUzSRZp2WaJkLRqY3pKaqpQGJ+y3d1AHiH2WHG6ynk8DrPaq1DyPAtPq0JRLDTFoPjUehK6bzn6NH0WM7qaOXSuvdR1EM0vJy7W+m8xaE+TMs4hYg9Ot/XCeZfQh0tNlL5nkvaNqrL4oAZDJmGiFrPgd+V1jg/EIfhjlnMss7MpY7jsleMnq9Bn1Uje42ypKAevF/L14wWhyXs6y5jWZBVdomn6nhg1JmVJPtN7XDDuWIvx6344PUEVuaP4lfFo/2uCC+zuP2cCfcKz8z+jHlFHA9RT/AEiOy1YzUdwYszlzXJqhJbEqGSej2oOvLHbe/i8sauH6KPazXSxPpeDRWrUuUb+zut8prmmU9CiCdXRhwQCvXjoSMafJe0DX+TSOmn/pPa7oUC37rM6WeaLb5htk0y2HndbDzwrnOp+0/NWM+c6t7MNVbgBvzfTdLHK/mP1ktKhv7h/niyH2gDEA8R7e6hLYicCeSb5/9HLtz03Oyah7HNYZXQxAmSq/QFTOikdLyQrInl0U8euK6kyuGmqpqGpVYJUbbtqyKZgT/dl2n8fXF5aA7a+3bQeZxZjp/QtE/d06QBNK6iloxLGhNrpTTtuI9dt8XNQ/6Qaei3U3bV2R9os1MwCsmYtS5tCg9AlfSgkfF7++In2uePOInwIPqQnizxuyeB4g/BXFc+QMkInqKcAMbFk8Se21luLe1/54bCho3kMawhWjIR+8W1iRx4Pl1J+WO8KTt0/0a/aWxj1JoXTmR1M4JaSfTdZkkxf+9PlkrJ/7EDDqp+jb9CHtLWGp7M+1/M6aulQJ9WyvV2X5n3dj4UFNmIimdb8gbzb2w3+1Im/xWub4tPqKjml+kefscD4Ee9FwVLksa+JYRuIFtxsW/j59QDa48uMBFl1OJSv1W7W8CSAhVb++3TnybyPkL3x2Rnv0HIMtinr8m7Ucvgjp45EI1XpbMsk2qesgqoBUU5PpIBtxWeZ/Rk7Xo6Z67SeUU2s4IwrvJpPUGX6iIhJtvEUMgqEB6XMd/Lg84uxTxTisbgVA+J8WDxRUdFkTTyBDQT93ZxI2y2xulib8MBf5kYw5ZJSt3ZpWMsRZJjEAwDKfQdSftW+IHPGJTnOR5lp7M5cv1Llc2TVykD6rmcL5fUhfJTFVIhJBvYktb3tjXTUVdSzTIsMlP3iLJ3jqSqI/JHPDEsm7xXAuSbnjEqjWqhpIJqgRwwbkeL6xHcbT3RJHRuu1gyn22n1wZqWnows8qbEdLWILbm5KstgDbllYWuAVPI6LxsY0fZM1U4kUteUFj1A/WHz4IAAsLHoDg8VVJPcCwichB+sJJA52i/BP4AXHtcQtazQiVpImux7txJtNzZLGxIB29OgsbcYTZy144mMDlQFZHBNyeU2+QK+ZuG6HD6ah7ypjRzIGfcwWGJZNzM1wLsQysbjjkHrYE2xkmVUAp56j65O9RFLGDTLCZ+9VrrIVlA294hCFk+ywYgHctsBNEBa6taVYCZ6qenEbjfP3Rdox5lUaxF+hF/QcDEi1XH2UZflOn8s0rqyuOftlxmzuqqaAmKoeSQtFDEqM5hlSM7Ta6kEXKsCMaXMIkWKOMic7VaZFqftuV4BKXOxQfIm5PXGnoMnjps1lrJYxHHUp3lMz+FQb+NAbEbh5DzHTEbmFzmuvEAatvinNcACKVqjzUCkeEKUYeHnp8R5fEX98R+vylwSxULztFvMn/L94xMJ5Iqcv+rLOAwQL1LWNhbnqbW/McA4e6r0LmuR5BS5vW/VHpJ4qZ4JY6gLIxlTeFMEm2a45DEKyjjmzDDy9rSGk4lIGkgkalWZmE6Cmqjtkj8Mch8h+y3t6Hy+GLI7Fu3TX/YrrKn1Vo7OFoM1gjNLIKhO8pcwpTbdSVcZ4liawseq8EEEAiC5nlpssgZWYqGDJypHx88a4N4RBV3AA8D25X+Y9sRSQhwp11uUrJSD115r2o7Eu3XQ30ktIyag0wP0dnuWIn6e0/US7qnLn6CQNwZacnhKgcgWWSxAJnE0UElFW5JqLL8vr8kzGm+p1tPmEZeCsidiGhqYiNpiNxscHerG4Hr4odnPaXq7sx1Tl2qNL6hqckzvKnD0GZ07XKKeDHIvIlhYXDIwIIJBBFxj1U+jZ9KPSP0ksvh05WU1Jp/tBpor1uQo14MwS3NVl1z+sjIBLQXLp1Xco44XSmhpLHJ9VY8KY0GreN3oukslvbM3sZ9eGPod6rD6Uv0VMp/RE+rMvq6l8mpokWHUNRunrMhUCyUmdbQXq6EDwxZkoaaAALOJEF14M1RpbO9K5xUZDn+XS01fSld8TSBwVYbkkR1JWRGXxJIhKstiCRj22oXrcrmecJUosSkCcRq8bKfCy9SGvezRutm62xzT9IH6HeSazyeXNezDI5pqWBJZ30vlwtWZaeWlnyTvCN8RJ3y5XIQvUwsjEX09D6ebaaRTYP9fDfu4bFTt2jTFV7MR6ePzxXmG5BY72I4uT5/A+uBVhboW8KkWFz1twT04HXEm1hoLNtHSwTVAirMsrGkjoczpw/wBWqSnDx+IB4pU6SQSBZEP2ltZjGZFlZjKxYqtlLk8WA4G7zsPnbHVNcHCoWG5paaFEdm2qdyefJN+B7+Vri49LH1wQv+rLN0vZthDbj5XtwPifzwo7PtDXbbHZgR0T4W6fx98IS7iWBdWYi4If+XANvww8JqEMHVbTqjAHcqryx8jbz/rjzwCMpUEqQw4Nrjn+H8PLGMXNoOWcgNZRckeZPoenTr1xkQ2ON17A8ggng/dt539MCEsGEpcyiRwPEArk89L3/iefLCY7tizSN1sxZlvf3v1/rpjJGLbUJaygAA2Fh5fZ6/Hrgu9S42sCeWsPQ+d/58+hwIQSAoygcg832gXHlfi/42wWwPiKFlv0vwvxsPzGMdvstuAuORby+PS1/mMFUDkMocgEAhj4fkMKhZKo23FxdrEBr2Hy+VvXCbRNtsU4cm3IsT8RwfhhzypBSQpdQtwpvc9Rb0/64RfxAFip8gP8v54EJIhgVUgC35H0J6YE7eV3DpdfIH39vn5YHZuN/tFuBt8+bdBwPb/PGFGYmwDA9dvr6H/LCIRd1lB3bfugi1renPHOCXG07WUcW2ngE+XHpgzKSrAWY252g/n54ID7EA+psCfPCoQHoLtyQOfYdPbAOFCBbkE9Rf8Ao2wYNusCxHHU8AfDqfwwWQ8AKTa3Nh/Af9cCERSR6jnk+ZxmCXUX8/QYzCIWEeIi/AHxPxtjDdlA45Hl5e4xgWzEXAt5nAADi7Dk3Iva3vhUIyhg1gv4fwGDHi/jFj06/v6YC9i1wVPqeOcGIP2hax4v5fPAhAoci5BsL9PI/wBeXngLi5NtvPmb/DGMNhuVsSL8i1vTGJcsLg3tzZiCRhEIxJU2DLfqOR1/jgTcLa/laxtce1/68sByQARfqOen4+eMAU3HIt1JI5/lgQhF73Y3Pr6n0wB63IsDfnbf92BXcwItwD0vYdPLGBbENuW4PPJHHt5YEIYwhUKqLcDcCvUH3Ht64UDb1VDYjzueD8f3YT+yuy42Kbgnp/MDBgXvcsfDcXuDb5+YwISu0iRUUm5Xwg9V56X8782+OCqSLgoNrjkgEWt52Pl7Yzdzu2nbbxAp9n+GMWUgMHNwQDYjcL+XHn6euFQjqdzoJLDm5dzzf+WCbnXkyHcTZlXqtuD7fD1GA3qOQb7gRwen9enzxIdFaD1d2h6hp9LaNyKpzTM50MgggAtHEou0sjkhIolHLSOVRRckjDXvbG0ucaAJQC40C0KxlVDbgCD1B6Afev6Yvfsj+jVW6lgg1b2jzzZLkMlMcyp6FJEp8wzKkU+Ko3S/q6GiuCDWz2QnwwpO5C4nWg+w3RvZfSUmq9T5pQZ/mr3ko6qKnFVl6OvU5fTy2XMpFYAGrmC0ERvb604CY3ea64rM/q5TVEx0r1aVjxPUvUNLUhbCpqpnHeVc4AA71wBGDaGOFBtxhWrS2qLLbrPgNQ3nHYMitCz2Evxd+Pyn2pe0GHIcvyzR2g6JMlyigRpaVsvikp0pw4Kv9URz30ZcXElbOfrkw/8AgWM92Yjl1PUZpVUeVZVRS1NVNKtPSUdNCS8h5/VxxqOflwOSbAFsPqjLKKuoajOszzGmyrJMmTuq/OK5iKeFixZEuoLyzEEhYIwzuQGAVbuKg1/240v6Lq9L9ncNblGS10Zgrcxn2pm2cxecbFCRSUp/4EZO7/ePIbWzIon2x9GjDrieesrQdI2yimtT3V3abknZWZ8t0zWZfnutKdw9RmoZKnKNPSj/AIJ5Svrl/wCJzBCf9mJWHerzVqDUdfn9XV1c9dVztWTNPWVtXK0tRWTMbtJI7Es7E89fjhlXZhUVwjFSBFBGLQ00Y2gD4fxPOCRxO5Vn2qFHhXyA9v6+OOisljbCMB115Dmsia0OeakokMDyKVUER3+z5sffGxpoSOFC2sFBHS9+hHUH44KHSMXSwN72PP5eXwwLupLMv6sk9A+4rfyuOfx5ti/SmSqVrmlalIop5vqsiyQmRzGwJXw3NuSMNoquanqXcFwxURdbLsP2lYeYOFY5xEbQSFAvlfi3rz5+XyBwH6sot1ttv4l633Eg/hYEeeCiKoqsVUBTcncD7f5/z5wKKm9AXAQ8A3ttv0P44wJ3amR5VADEbgRa/UWv5/5jE5yTsY7Tc/yiLPm01JlmSN+sizfPJ48roiCOSs1SUWQW803HASBiUBQ6KYxPHKwZniPF+LjzHwPIwApVKTNAS0MV2u1lO0C4Jv5jp8Riz8l7OOzekZ3znXOaatqKeQLLl+icpeaJWPQPX1SpEg/vLE49CcSI670HowL/AGf7PdBaaniFhV5/PJq7N1YdGWI2o429jGtjiB9qjYbtcdmvhnyUjYnuxAwVYaS7P9d6+l+q6J0hnOoZLgP+jcukqFi5uCzqNi+nLAW5vfE2X6O+ZZHJv7T+0LQeglNt1LmmdCtrwvlajoRNLcHpu2emA1H9I/MtQ0U2XZxWan1nEyqscWbZk9Bl0AU/7uhomSMA+hYgDyxBT2k6nHgyj9Baajtbbk+WRRS8efeBTIT7774i7ed57jKDeafJ4gJ5jjbm7h1TmrgXs07DtPUr1LTdpOtJ1QslQtHTaUylweD+vrGkndLHmyKbYJU9onZZo+SE6W092Y5K6RhnkGU1Oqa5JebgS1rdwbcchQL3txihKuoOaVZr88qsyzeoc3aWqqjdj7k7mP4jGyy3P5cnSL9GZZldKwBUz/UlmmB/xS7rH4AYZJHaXjB3Cg5mv+lOY6JuY9/j1VyZt9JXXefU9MKPMu0DUeV0JZ5cveaPKclkAv8AbpMvRF2gckF+bc8YjNZ2i63zKaGtaXs80m0KladzTQ1FTCh5stxNIt734AN+euK4zzO82zghsyzmsrCrdJZmZAPQL9kW9AMa+CPuiO5WxU7l9fce+GCwF+MhFfM+uH+VO+pDcGZcPTHmpdnGoKnPa5INQ9peotRU53tJHAskcQckEKiysFAbm52C1uhvhqlToyiCmj0THUyHpLmFfLID77Iu7B/E4j1lRrqUu3i8DXUj+Hw8sOJZp0Y1IZmJAPjNyARazfHn49euLAsbKBpJ8jd/00UZtDq1oOFfWq2sud5lFaTLKTKqBb8mky6IED/GwZvz6YQhzzUuZlKeXO8xqCbAwpUOQD7Ipt+WLJ+jh2i6Y032jZHlWouzrRmqcnzfPaPLqj+0uUmpajhmkVGliZXQgqSSAxYEDpzj0vy/VOotHy1GV6eyLQmlqihmkpZv7P6ZpaZe8jcoSryCRwPDfkAi/U3Bxi6Q0pBop4Y6IVORw+PdaNksc1vFQ/rrcvLbJOxHtf100Y0v2Ravzd5AAXpcjndSf8TIBz168c4nU30KfpH5BkE2o9Z9np0lktKAxr8/zCnoY0dzZYhdyzO7WCoFJLEAdceja9oXaJX97HmvaBqWRmN7JVtFHt9xCgK/+Er7jFf/AEgUTNewnWtVJVSzSUVNl+bPJLUNMWFJmtJISWLHkI8nPHBPGM+H9pZbTM2FgABwy4a1ak0OyFhkecutgXlisW5+8V4zGtwdqE90fdOq+46HqOcFkkeHa7vtjKX3JwVIPDDm9wb/ANXxtc8ytaHUNdRSIUNLWzQ7lO1ktIw4Pl8Dwfa+NXIj0geFh44WfqtjtJ6+o/6eRvjs2OD2hw1rn3AtcWnUlE77YFmB3bmDWFgD+0PTyP44GFyHSJXULGtgjxgkAm7Ib9QDcj/FcYSiZEWSyRi7qzDaL8jmzEEjyPXEr7NdN0Op9UwUtbT1FZSQbGlo4vBLVzPIsUFHGw6NNKyJuHRS7W4wkkgiYXuyCGtLzQLc9knY5rXtg1RBpnRuUxVE7q1TPV1TmCko6UNZ6qok57uNSdpIuXbwqGbjHoV2WfQx7GdCZfTyaqyJO0HNVAb6znSNT5aG/wC4oEYBlv8AeqHdzx4V6YnPZz2YZd2XaZXS9FTRy5i8qT53W0YCJWVyLsO1V3WpoReGCJlsqKXN3kZsSG9ZBMpyyZYLk98hDRgrb1iLC4/vRgY4HSv7RTTvMVnN1o1jM+ez1XTWHRUTGh82J2alvstzbNNOA5dlOS5JkNJGgNP9Q07SRLGB91kVNygeTrcW6gY3FJrjVLzCWq1FJIUG2RaaTurDyZQfssPRgFPKnyOIl+mc8pJE7mGpnp+hkhUOt/8A1Z23+KKf3Y3rhu5T9JQwxMRcJXCKLg/3Ze6YD4EY50SSzu+5zj5lazmRRipa0DyUmptS6sWF5H1lXst9y1CTLtIPQlXBCH1VvCfIjphCTWOr0rmNVqnM2p2F2EElu74+0U294qn1BcD4dIVJm2V0tTIJc3yih2i6FM/pTGT/AIZJt8Z/wsynzGFKfVOnmI77U2moyDuBGfZegB/aAWcbT7ptPqrYs/3si60P/wAyhrZhiS3/ACqwqDUeoxCadNX5jJIo3RNLKGYr1FwQN49xzbzOH1ZmlfVxrS5rUvWQyKLxVccc1j5ja62cfK498QbK9Yacr5Ey3L9VZBXTymy0kGcUNVIz+gjSRi1/7qm/oMPhVSySmClpwBFdJYqaTZ3bDqGjO5R7golvXA6SaKl+8DvqkayKQ90A+FFFte/Rx7I+0ajqI6vR1HllRUgq9VkkCUblv78DBqSc/wByWJSegYE3xwb9IP6HepOzOpjr8lCV9BWytBltRRxOtPXSKCxpo1cl6WsCgn6nIzCQBvq8hI7rHpPBmtTSRd7aTay7SVUObehClrj5Eewxrc4gyPWeQ5lprUuXw5nk+bw/V8wpopLNLECCHAvujmjIDxyCzI6KRbG5o79oJLO4MnNW8x11voWrRbJQXRYO9V4llAzAxEOHSyX4VyTfafQN0I9eeDfCqAUqrPDNNIoRHijBW+1gdm4sLAdUPvuHHGLa+kJo2HSev8xyyjzqLMsypcwrcrzySGIxE5lS1Gw1HdkbSaiJqedlA29604Fr4pln+qykSyKLhiLOL7S19u3rw17eY5FjbHdse14vNyXNOBaaFbqknjqVkq4YZ0htscHaTBcgbXI8SHqLsu039+GrSOwZy8lpI0cXIKcgq+4f4gQWFwbruuPFhnE9RI8UddFFEgV5KaVSChH3gJQNx9gT4eRYXxsMnyWszzMqSjyxKmapqZo4IEQfrHlc7VRBwNze9hYFm8IOFJAFSmgVwCeZfltRmk3dQU4nawMm+ygJ+1I7nYqi3+8O3jhsXp2afQ87ZO0+nirdI6PrpMqdCxzJ41paBgeCUnqe7SRbWF41dSPM9cdC/Q4+jZkeZ0dN2nampIavJMuqC+TwiAz0+azxkq9SAVJakVwVWUrvqHVipjhVUbtb9L1FfK0lVUGqQcBZGDqo9AotYexTHOW7T7YH9nEKrWs2jTK2+5ecOZ/6M3tnhKTxaj0nJUJ0g/TcAlT+7coot088RHPPoA/SayKmJHZ9U5pTjpJlssVTYXudpgkY2J5tY28rY9RaikjqEHdxqwjGwEeIKg6I1gWAHNjYlRxZhxg+W0cFI4kpwYyRvDxXRrDzuhsQP2lJHqFxnR/tDO13eFR5fCuv0ZCRhgfP5XjHqrse7TdHVDDU2js3y6WM3VqykeK5HmHmjW3zuD54YU+Y6zpqR48tlzeiO0Mfqc86Rkjr/sXKqCPUW69OmPcSTUefRw9xJmMksDDhpkSUFf7wYG49+QfbEOzzs27NdVz91q7sj0XmjyAuKo5SlPOT/wCkis1/cYvR/tHC40eOXwSqztFGndPOvsF4rV+pczrwYc4my3MXTxJ+kqOln4/Ycsok+YYn9wbtk/Z3mUZlzPQsUM0ibrZdW1FEUPnfvBNGR7qOfQY9YtZfQe+jnq5P9VyXUOnppSSgy7NUrIw3p3Nasg+QIOKe1B/o3qSEGo0T2k5f3yk7Ys4yifK5VPoZacyQn4mLGnHpaxyZPA5etFSfYZm6qrhDT9ZU6MlWTQPax2jaHlSxSOOVqinZvIBqaRCR8Yjx1xLoe0Ttpz9vq1fmPZr2oGpSIxrmtLFT5lKsZJQpNamqiVJNtshsT54lvbV9F3tS7JYlr9SZRHHRhjFT17vHNldQWP8As1rIv1MbMeRHUJAWNrNewxRYpavK6uamzJamhrKeVGMVXHteGS+4MARZT0IAH95Ti4YYZheLQd+vjmq4fJGaVV1L9JTtI0jRrl2scm7U9O5UCA8DZimo8oNuDamzaKZbf3RNiT5b2m/RP19FBLmGVdn8+bFLSSGmrdDVzk+stG01Azf4kVcc+0+c5llirNk+c1dBJuJb6pWyRE35bdtPPW9yCGvceYwwqtQtWMRqigybOEI5eryyMy2PmZYO7lA+ZvY+mIjYyDWN7h519anmpBaKijmg+VPSi6T1B2J9htZlseeZbqnXWkqWcLTfWs0yuk1Zk6m52Ia7KJBMguW2s8VxfEGh+i3rTUU9RN2XZ7pDtMTeZSmj8+p6qpRTa4agqDDVx9BfwMAb9TipkyTSlFUrmWnZ9Q6YrotsnfZZX96iDqGAcpIB8JG6jnnExoe1ftb0lPJMuo9H65gqokilh1LlURq5Y+GRTLMiTjoCCs3kOcOP1UbcKOP+H5TR2LzjUc/haTV+htU6HzJcn1fprMchrnk2CjzSjmoJb2NyFnRAy+Xgc7ibdMRitjfL3ijrnWBwwASYvGtgLWta1vK69OLeuOgz9NzUlDQR5RqvItYaZoWXY9ElZHqPIqkWtzl+brKgHX/ZzLbyxvYM2+jL2wpHFl+Q6Wiq5og5/sznT6QzDvjwUGWZkajLpyev6uZL8jjA22Bo/fNLfHLiKjmgwV+wg+vArl3vWniJFSWCxCnNRYtI6Biy2v0Njt3HlrceuEvqwp5w0YkUHbZmeRC4tz9rqPe1/li+M87AdHpVtTaT7RYssraj9UuVa6yptLVUp8ljqiZctqDfoRLHfrfEL132HdpvZ5TQ5lrnRGa5PSOGKZnVZcDl0vI2OtbA0lPPuvx4lB98WGSMlFWGqic1zTRwVeVzS/VmninljnUCRGVyGRVN+8uLbbWutrEkccC+GmZZrqrOBJ/anUNZWpT19RVRCrl71DJKEEjlmO494FS9jza/XGxrIqmnWGpe2yQkrPIVdJW22vvN0b4344AGNbHDCjx3MR2KES4sABfhb3IHubE+1rYcWgkOIxCQOIFAUSPKo63K58xjr8vjjgkW9M1YrSMCedvl8FJDNza9jjTZjlETNIIbuoJN+lh63P8AXGNzPQQSVBZ1cRu1nFtveEj9W7KDwVa4/wCbCU8SSWVbnkKS5289PP8AebYACM0GmpQ+/cA09QN8V+COq+6/y88bXI9RZlpyuoqylzCqgejmWpoa6jlMc9JKpuskTixVgQDbCtdlzbXMsLqSb+JCv7+P698aRkmpLqVDxP1U9D7/AOeGPjDgnskLSvU76MP07dP9pMVFoftvzLLsq1nUBIMt1Q1oKDPyCNiVR+zTVd+N5sjk87Seeqq+kr6SoaMCemqKeS0ilirxMo3KwI5VgCbMD6HkY8DqSrko1cxqKijkt30MnI+fofRh/ljtz6KX088w7P6Sg0P2wVWYal0PSbYqPM1HfZtp5OgRgf8AzqkH7B8Sj7BFtp5DS37PtmPawYO5H4Pqt6w6SMYuPxHp8j0XXfbb2C6b7VYq/No4cqynVGaIiZi9chXJtRMv+zGZJH4qeoH+6zCGzoT4rg7T5xdr/YNqXs5zrM8vnyTMqSXJ173MMqzAK2YZXCTxKxQbKmlJ+zWQ/qzcd4Im4PrVlktDqfTFHqLSOp6LN8mzGMT0WY0BWognQHxbSw5B6PGwV0tYgEY0GtdFaX19k9Lp3WeW1ZTLZGkybMqCYQ5lkk7LZnoKk32i32qeTdE44IsbYqaN05JZH/T23Cms+/zn4qe16Ojnb2kHD4+PReLLp3aiw4I2ki3Q+n7/AD9sIsjDa+7buYlWLfeHmPM8+Y+eOvvpE/RGzTSbHUESZXT0tRN3NNqOkjFHkeZSk8Q1cRv+ha4k8q3+qSsfA0RPPKufafzzTecVWRZ7lVTluY0TKtRS1MRjlj8wCp6KQbgi6kG4JHOO3inZM0FpXNyROjNCtRtXapVQebncAwN+tz5/kcKgM1nvZeVC3vceoJ/C5/PBJEbdfdyTY36H0+OADNu7yR1YnggsSTfyPHTE1VEjsbeJWsNpuSQD7g26fvwW83dsqk7SAbWIXcPW/F7HGNG8SF3XkEC7kpf2A6n49BgHRY2KspAJtyoBHpcDg/H8sKhYdyx32MNl9yhSLX8+Ra546YCRuUDAgA7gLkbR+0P69cY91CiVHVvIMgsB8CcEdrMN6lefskk8+3mMKhYWQEbpDtuLbmB/G3Qe+MAG3el7gm+4j5hvn/lgCpa5JHJA9uR+WMdAGBVdoYA23W6dL38x69D8RgQgIDOfEu7jruO6w6/n54Fxfa1goAJ5PHw+PtjOrX27ubEAWuPK3v8Av5wW1gTvU3JueQD8PP8AAYRCI4WxO0cHnjy+P4YSCncSpvbptHN8LGzAEHm5Njx+Hr+/BWFrE87uQR1J+HoPXCoRLBSVBFyOCG4v69L/ACwRixY+fl/X88HcgE7jzbqBa39euE2UHpewufb/ACwISbMeOd3xxmBKsX5BHFrWtjMCERla+1Rza9vfAhQDewsRYEeeBYLySwHPJ9/6/LABSTZSCD+BwIRwBYGw2m458zjHQ7mZ7hzwbmxH8PlgELGwsOTbB1sTxe5PAIsQevn+IwIRLspFmFwel+Rb2xgW5IFrfZN/K/Q/DBzYi9yPO/S2AUmwNjwR/wBcIhF9ARyfKx64OoLW522PF26fl64zgoGKqLEgkH92Chgp5B4Iv7ccYVCPe78oTfk8cgevtbAqV4RiR57h09uP44Ke8a9jaw3Lb+H8sH3koEBtxdRfp7/vGBCBmYi4XcV56c+nzxhW4J+1bruHX5dcCF3IGKeE+g4BHXny4wc7mfft3FjyDze/9fHCISZBBBcXYgG56+3x6YA2282UA3uOgJ+P5Ym2h+yjVeuklzKhjpssyCllWGrzzM5TDl9M7dI+8sWllPlDEryt0C46e0Z2SaL7GjFmEkVaNRRqsyVmZUMS5yAfsyUtBMHhylDxtqa0SVRHijp064o2nSEcFQ3vEZ7B4nV4YnYCp4rO6UgbeutW9U92bfRfzzPHpMx7RZqzTtFUU/12nyqCFZM7rqYf79YJCqUdN61dW0UQv4RIbA3rFnOl9I6dGkNFafy2kyjcsjUkQM9JUyqVKz1UkqrJm8ykghp1SjQkd1TyC0mEMx1NVV9NPAqLS0M9SKuoponkl+tzAqDPVzylpqycB1dZZ2IXawRY7WxqDlFdmlNVZlvp6akoXR8wzCvqBTUdApSxaaZrhD0sgDSuQAiMMcxNb5bQ+jzXZsHgNu81OymS2orCyFt53XikszrqnO556/NK6pqJ6kxtPUTSmSRydyrvc8m23ao4C+EKFvbGi1LmmnezNRUdorVUubMglpdLUcnc18wtdWrJLf6hCQb8g1DggqkY/WYjOp/pA5No2F6TsgllkzAL3b6xr4CkyC1iuV0jX+ri1wKiTdNY8d10xzzmOa1WZzT1Eksp+sStLNLK5klnkY3Z3Y8sxJJJPzxdsmjHym9LgOf49fBQ2i3Boux9ddVUu7R+1jUvaNX00mbzU609AjRZZlNBF3OXZZGTcrDFc3Y9WkYl3PLMxxCy7bzIT39QfvHlV/mfywnFGz+BBtQ9T5n+vTDtEiiFh58Bugv6g46OOBsbbrRQLIfKXGpzSaUpJ3SEuzWJN+ScLqtkcJ9rbdQODwebg9OMCWBNpBcuBbbYXIPXjj5j8MK00FXV1MNHSRS1FRIwWCKNDI7t+yoA3N6WAOJslEkv9XKNaeRXDDajpyw6EC3BP8L4xEdixC/ZFyPOwvf5jr+OLDXskbIYlqO1PUFDpGCQKyUEy/Ws3mB6BKKI7kJ8jM0Y+PTGzjz/AEfowINP6FoqOQWcZxrT/X6xv78OXoO5jPmN6yf4sV3WhjcseuflVSNjc7Uojo7sy19riJ8wyPIZHyyI/rszrZkpKGG3XdVTFYwR6BifY4lmW6F7L8nnNHn+ssw1nXw/ay7RdLeFT5LJmNSoRV/vRxOPQ4jevO0ubWE9PUZlmmd6mqKUju5s4lWOkhA+7BSR+CNPa/yGIrmGe55mdOKaqzCQUt7LTU6iGAf8igD52wwOnl+0UG/qv+lOuxt+418OvlXFUdqdNoCoWPQ+S6O0Q6KG+s0Ea6gzu/mDWVG5Im4/3Xd2PliB5/2lVufZuc/r6ev1DmdiBmWqa1swkIPmI38A9gdwGIai7EtGqoLfaC8/if4YVhpaieoiiiid5JrBF6szHiyjq1/bB9MCP3rieQ68SUna0+wU59eS2Gcau1Vn8K0mb5/WS0q8pSRN3VMn+GJbIPkuNSkIUnbGo2+ZG79/GLw0t9EXthzvLqfPNRZHT6Oyae3dZhqmsTKYpA3I7uOUGolvzbu4Wvi0sg+jD2P5BC1bn2rs+1fLTFe+TJMsGWUMYJHjerrQ8xjF+XSmAtyOMRSW2x2MXQQNw/CkZZ7RaTUAneuR46Kpqou9SB5EVfE9vAD/AIj4R+OJ52d9gfa72onbobQObZvTKR3tZDDakhHmZKl9sKD/ABOMdhZflWR6RrGp9Edk+k8nqKazCvqqJ89rUBNw4qK4vGnqHjgVCOVbG4raqt1mZKzV2fZpmdQ05n3VU0tflwv0Bp0N6Xp1jVkPopxj2n9pGsH7ltTvWlDoV7j3zRUfp76D08bQxa57ZtDZNPKGtSU+YS1zJbqHlpYZYkPtvN/LDnPfoD57UwySdnnadozVdQo4ocv1DCKx/ZYKpIC5/uhtx6AHFw0slYW+qZZl/wBYooyQv6PqRL3JJuSjRguoJ5Mbxst7kBTc4d0mWVld3kVXFT10UQu8VSyOQPMMtyy8eqfhjKZp+1NdeeajZktD+yICLoXEmsOxbtA0SmYJnmSz3ygquZKYJIarLgTZRWU0gEkIPRZSDG3QOeMQGW0RDOGbawutrcX5F/LyPTHp3VSRvSUdLqGor6ihy+NloJ4ds+Z5LG48T5fI9+8itfvMvlL086bkUI9hjjbt87GV05nNZnOTU+X0wWOKqq6XLy36PqqSc/6vmmWluWoZrgGMktTSkxt4dtup0fpWC3tqzA7Fh2zR0tkOOIVHxEoWCgBh0seVYHrbz44+YwaolZ1ZAAsY8SIqhQl/T2v5X64F4pIVcOt9jKGZXP2Txytr/PpxbDdywvuJ4ut7X/639MaioJHLqyahqXq4GIkp5IqlDexujgg49a6bNBqKpqM7evQtmbrWxm8qL3cqLID44CpPj5KsRc8W6Y8kKZA1W8AH+1jdB8bXH5jHpV2QZ9T5j2XaKz/vghrNNUCzuRNJulp+8pXOxXC3vTrwBclviccT+1sQIjkpiCRxp/xK6TQL8Xs8OX6q3Uy+oRt6QyWHN12Ovx5RSD8LXw3znJZtVaaz/SUkQkOochzfKx+vBvJJQTND4G8R/Wxx9CbcY1okoKymcx1KpJ94rSwixBvtDN3lulri/wAwDiQ6QkqcvzrJq2WOlC0uZ0cr33kd33yLICxA+479AB18iMcfZpDFOyTY4Hmt6dnaQubtB9F5b62EVbnH6acWfNYYa4sBfxywo78eaklrjr5i/TEWZHlUywmRxEQilRuG49FBPkeRYgYsHtq082ltRSaaYEfoWorcqKHgt9TraimBt5+FENvcYreVI5Y2JNpIxa9jcL6lunXj2NrcY9esuELW7BThh7Lg5zWVztuPHFPEnnSmUAd2gjMQMaqJO6a52M1vFbxAHg2BHpizfohVC1X0j+z6mlVXpm1TTVJRh9p4I3aO/wAG5t64q16otStIxHKXN2tz5/HxC9vX44nv0SKxKb6Q3ZtPIbA6to4z/wCsJX+OItI//lkpsPoUtm/ijxXqdlmbJWUdLJX5dKjoFmMgi75TcXYF47tYknm3GNnJLHVlJMrqZKiIcXEyVKofIENZ1PwZTfyxHMmhMtHGe7lV0UDfCita3F7bkby98bmnneKWT6zmFE7gBSlfQuHYezEX/wDZ48iBoMV3pGOCqf6U3a1q3sV7NFzbSZXL9TajzBsmoMyBkvRRpAZqmpVJLkSrHtVL7gpcsDcLjy3znUsedV8+YZo2aZpUzuZJquur2kmmY9XYsCST15Jx6JfTvp2zHsw0om1dtLnufKipNLIqhskDi3eEleYz4QdvpjzkiysBLSABmF19z7j098eifs5AwWQFuvE6sSSPZcnpaR3bmvht1D5TdqzKz9jJyPjOT/DBfreX/wD2qH/w0/yw/GWK+0xR7lkXgXsR68njg3vhNssNgTCV6j7PmOv8Pxx0dwb+JWV2h3cAmX1ylDBoqMxEG4ZXuQfUcdcepf8Ao/8Atgz7tN7MMzyXtAqjnNfo7MKWhocyqgJaqWinhd442YkSSNEYmA2kvsa1iFGPLWoozHcbCNvniwOzzt47SOyikqqTs517XZBDmHcSVkUNJFKJZYlZUc71axAZgLW64zdK2I22zmFuZyrU0x8zkrditAgl7Q+w+F7YtmFLMd8GY0k0fRgx3kEeW4EOD7MCcL/X8u/RzxV9fQhmD92JpSYm9ArSgeL+7uDA9AceQR+m99JdlLt24ZizsLFmyqjLEfExXw3P03vpNiNqcdteYmJ12Mpyyk2lfQjuuccu39mbQQQXjg74WwdLsFO6eIU1+mBXJS/SN7QTv2qdV5qztusArGnHJsbD3seuKWBVnpywgVz3kPeMySM8LKbXZeLq3AbrY4c5xrfOe03OJtT6iz6bNc/r2qajNaqaARiSdyu2U7bLYhVBAAttPrhiZ2EQdB3RReQXVjE46WPTqPC3N/jjtbK0tiAdv9Vz0xBfhu9EYUsdJFGxdTE3i3EG6+X6wD7QHTeLOvmGXEn7O8urNU65yTRWTSTR1moK6kyNZUcF4TWyrFKysOCRBvG70kJxFamcLThGnaKQWBdjsAb9nw9W9vCB73xcP0FcqGd/Se0LPJDJJHl+Y12eOqLuLLRUcjKAOOd+0Acc4bbH3IXHrDFLCC54AXrhAmnaGSLJMgp6WnossjTLqCmX9U0FNAoiiRVbawAVBbbcG9/PC0lTWRCSL608UiyWRpzusB1BWQcg/G/ofLGroagyUcSZjE6jarNHIskNiQCQUnV1U3/vW98OaelSllqWhheCJpdyOpMcbCw5DRmSMj1JVfhjyYve+rjmTiu3DWsAaMhkt3VZlQUlO+Y1MSQxU0bTS1ErLHFDGouzGcNtVR5lith1GKqP0w/ouLmZy2m+kBo5Ku9272ql7jcPWYR92T6MGB98cvf6TrtOzLLck0z2X0eaVEGVZjRTZ/ndPE6o1daoEFJA7oADGGWWQi3PB6gW823zl5H5yygKDgIILAD43v8AnjrdHaEbaoRJKTjsoKfnXsWHa7e6J9xlPPFfQRpbXemdcUfe6Q1lpTPe8O9Vy3Oqapdz6BEfm/uoJ9cSaeSHLyBNTz0oPNnieNR6+GUbb/BsfOnFmuXxusq5W0Ei8iSmqXRh8L3xO9LfSD7XdGlBpHto11kyJ9mFM1maIf8AJv22/wCXFl/7NsA/duI8QD6KJuk3H7gD4Ej1XvqlfSVNlSWOoRjxseOQE/4JAGB+DHB6ippacmXc0ElrNujkjVwOl7bhcdAR8MeNOkP9I99KfTTJHV60yDVkCkXhzzJoJHdR5GRFjf57r49F/or/AEmMj+k/oSfUmV5e+S59kk0dJn2UJK0gpXcExTQvfe0Em1gN12RlKm/BOZbtE2iyQmUUcBnnUeStWa2RTSBhqCefmr3p4spzWOojtTzx1UTQ1UTRJKk8ZFmSaM+CVCL3DL8LY4r+l19EjINJafqu0bsyyaQacoU7zN8lpYzUy5LGTzWZcrcyUYJ/XUbGyA74ihBx2PVZlluWxJV10o297HCWJ3OjOwVSDYOF3EA9bXvgk2aU1RFIlVEKmkO6OcOBJG8bArIjsvNmUsp3KODziHR+lpNHmlat1j33KW02JtqFda8OM2pcxyytlpJ44nj3b45ISxikjJIDx8A2JDWuAwYMrgEEYYzxbzHHUvL4d8pUqHYFrFttvOw3FSbbjxbccX19LfRaaG7UczymkkWGKkzGqy5lcAK88SxMj3uLd/STUjtyAXikPU3xQE85G+OIRybCEWa6gCQG+wdFe37Ytb3x6BBK2eMSN1rmJGGN10pYVE0NUYYC4SFu7icJIjkG20kWNiRbgc8jr1xrzXNWVC99SrMkpFx98n0Yr1JHsCB8LYczSxmZgskjHbv/AFgKsxP2j9oki/2rNz7DCAqKiizCL9GyotTVSLTQvwXBkGxnI+6dha1ul+t74lJDRUpoFTRbKro890unczfX8nknUSRxSBkhmU8iyveORbeY3cG+EhS6bzEGHPdM0sjuVXvKZjRz7iL3vGDE3Qk7o/bk49GfoufRp0X2w/RV0/Xavgrstqszqsx/RVTlciXfLUqGjpxPTTh4KhrpJZiqvt2jdxfEK11/o19UZc0h0jmNBntMDuiWim/RlYP/AJzrC1O56cRTx9OMUmaQs73mO9QjUVObLKGh4FQuNdM5vrXQau/Zv2rZtk9NICr5bmQMlHKP2XUCSBx7PGPe2LH0T9K3XXZ/JDNqDQ9RlVCrsanMuz7MP0UlSrCxM9Momy6f4PAvnziP9o3Yxr7s5zs5Pn+T1dLVO3hhzCjfLamS/wCys57uS/8A3UsgOK3qqbNMozBhMajKKwBrI7PTzsfLk2LfGxUC9r4fJZIJTfu47RgeI901s8jBdrhsXSGU6++jt201pgzaj0hSV9bKLVqH+xObgsQLyNEs2UVbC/344b+18bLX/wBBbOcomap0Jr3J8wp6g3pEzgplJq1PQQVqtJl1SfL/AGsTH0GOXamOir7LnmSUtfIwA7wJ9Wnb1IkjADW9Sre/ON9orV+uuzGaWfsu7T8009HPcTZXmhD0FSPNJFs1PIPaSNfliN0NpixhfXc75HwU8SQv/iNpvHwVmu+zDX/ZvWrkmvtJZlkVTLdYo8zpmpVqr9DHK14ZhboY5Dcc4icsE8EjR1ySpKhXvFmp9j2B8yeWv87++OiaP6ZmtMj0q+ltZ6JTIqWqe5qslgirskqb9e9yqp7ykYHz7loj6AYlmTdnv0ZO3qgiqNL5tluUakZLzJpENDd/MtkOYODb1+p1HrZPLDf7QEQ/vTSzfmOIwHnRO+l7T+A69uyPD4quPFphLTGViIjuYKscIknlIP8Ae8KKOBc29gcMKqjgneUwpIQrA2mWzi6jg/O/PT0x0Trj6J3aRlslYuiYsv1/SZf+sq1040jZhTIBx9ZyqYJW0/F7kI62+9ijajLZiWA7lTSEwSrLIY5VIP2DGRuBU3+0Ab34ti+x7ZBeaahVXNLDRwoVD6illpZWlpmNhxwOvr8RglLVTQTrU0EhgnX7qng/D+RxIquhaZjtFuT4gL39evUD14GNHVZcygORZj0IHDH2+VvmcDmh2aVri1Xb9HP6VHaF2D549Roysgegr5FbNdNZgzfozMyOCyi96ee3AkSx6DkeE+mnYr27dm30hcrnzDs4zKeh1DSUqjN9L5iFavpSpPj2cCphAawljFwANwXHimzhjsqAQw4Eg6/P1/fiRaX1xnulc5y/PMuzevy7M8skWbL81oJzFVUrjoVcG5H90+XGMLSWhoraKkUOo6/yN2fotOy290BwPwfg717hU5mgpqqGU5fKtY/6PqKSSn7ylroXX/Yzwy3SQHkbDcH7pBxz72xfQ9yLWeVIeznLErKWBSINJVFaIqmiubt+gswmv3Yvz+j6rdCeQhTgiCdhX+kE09reki0f9Iv6lk+Y1RWOHWVJTn9HVzgjb+kIEF4JL2/XRgC/LKOuOxXikRaWqgC1tNXxrNFUU86Sw1URF+9ikF46hfM2IPnwQccqHW7QL6HvM5fg9Yrb/u2km7Hc/wAjrBeQnaD2I6m0YmaVlHDU1+X5O/dZmzUj02YZQx/3eZUMl5aQ/wB/xQt1WU3tisyroCp8O4An3Hl7W9v349stZ9n+jO1WGKfU0dQM1yoPSZdn+VVP1fOMt45SKdgRLCQeaacPEwNvPHEPbd9DWr07BXaih+qJl0JZn1LklBIctS/Q5plqBp8qc+c9OJKa/O1BzjrdH6ZgtraA97Zr/PlwCwbVo6SzmpGHL8ea4ncW8JuGsR4bC39f9cE5AvGo8PJC9AB5/wBX/diV6s0FqPRVZT0ufZYaf65H39DURSrNTVkP/FpaiMmOoT+8jG33gDxiNTQMSWNiVuQDa4/Dp8cbLXAioWc4EGhSA8HRuVO4g2Xn0seenrzhNybgHd0C/wCQwqQEXxH0v5Ei/kLfAnCbEkAhQR08V+ff0v8Anh1UiWa4N5ORb3AHrxYeftc4Jv4O0hAGt6Ei3NwScItKw8II5IIDW5I6cHm3ODoTf7I6Dg8Hn29MKhYbGzMVFuovzfzPvf1wBJszBGbaOSB5ep9MGIZDa20bT5eKx9v325xhW4DC+7oRci387+mBCSsDtG6wtY7haw97fwxhW5spVb2UeV/jc/ywYnkMGB69LfnfBSzD2FvQGw9fhgQky5IADGwuALdT/X5YLuIBVVtbrbrb05wchipDeZtY9QfT4nBSSRySwBvdhfj54EJM3JtYkDz9fljMHUFbEgm3Nj6HocZhEJAtZw263PFvLAkKARY/LBWAY8rax2/174OLlSSRcG5J4+GFQi7fDyBxf8vPB0K+W0qQOLcfMXxiqC32iDe9weQfW+DAWP2evN9t/wAcIhY1rg36gi9rg+eAIUqvJ9Txa3y9MCVGwEbT4iOgvyP3cYzxFvQkcLbxfC38cCEVQQD44zYcHd0F/wA/34FlKNY9QSLD264My7V4C9PQfhgCS1lBBJ6cf1/lgQjAgG4kIvzyRyfh64U2gqFYC9+fb3H9WwioJAsFFwBc8HnyJ9sbrS9dktFnVPU5/kjZxRQK7/UPrL04qHC3VGkTxqhPXbYm1gRe+GvfcaXUrROa284NT/RHZzrDX9VUU2lsmkq0o07ytqpHSGkoo/8AiVFRIVihT3dh7Xxfeifo+6S0xS02f61qqXNxVDfRz1sVRHlcxHU0tIuytzax+8BT0vHilYcHd9n/AGzZdrDSv1jK9G5HkL5FVQrSUUNOauiy0SITFPS0jjuFlZ0dTUTJPNvA8S3GDTZlV5jX1mdZhW1NdWVBEtVV1Mz1E05HIMrv+st0+14RbiwxzNt0jM4ljgW01D3I9qbDULWs1haQHE13/j5UyTWceWyQy6dSposxoVNJTZtWJF+kKaIjlaGOECkyuMjjbSL31vtTkm+I9CHr6+DKqCGaqrcyqGiiigjaaaeVlLbyou7neLOxuSrkscaiuzvTen8j/TnaBnMWX09YGmpoKcLUZnmZ5v3NLfaFv0nmKILWAe1jS+uPpB57mWXVWm9I0q6SySqBjqYaKYyZjmKelXWWDMv/AHSbIh+weuM+Gyz24gN+0cB4D48yFdknisgwz5+ZVta419obs6knoc/rU1Bnyt/8YclrFKwsCfDXV0d0Wx6xwb345kQk45/7RO2HWHaLJS02fVkH6Py4n9HZNQRdxltBfqY4h9pz5yOWdvvMcQaWpZ12ACGP9hOp+OCxozi32E/fjpbHoyOzY5u2/Gz12krGtFrfNgThsWTSyTyd5O5lk/IYWgp7uhmIubWueAMKRQIrBDZbnzH9c42WRZDnGpMzgyPIMmrcyr6q6wUlHA0sr/BFBJ9/IeuNLBoVMkuKZbVVSLkbT8xY8H+GHuUZHnOqcxhynT+UVeY5jKQyQUcLSO/97ao468twPW2J0dDaM0JIYe0XO3zbOkvbTWnaiOWSNv2autAaKH3SISOPMoejTM+0XMqqgk0vlaQ5HlE5s2QabUp9Y9PrVSSZJj67mf2C4rutH9Ar11u3qRsVfu6647koOz7S+mZlh7QdVrVZlf8A+MGmilbW7v2ZagXp4T62MrD08sOp+06bIKeTKNKUlLomklUxyw5I31jOKoHqs9c93W/mqFV/uYhdTPPSRvRSTRZdCRZqLLiC7e0stzz8SfhjWgxRginVYo7E8GzEeYJ63Hp8xiIRvnxecOvL18U8lseAz661eC2cmqK6ndxk1LHlBkJaSZLzVspPJLTN4gTyeNuNE6GonMkkjtJISWklbczH1JOBUbQSCVNwRbm3+fvhRWUSiSUAruDEr5jzt+eLbImx4gY9eahc9zs0FPTNMRGB42NgCRe/zP5YtzSP0Yu1HUeXxZ3mmUxaZySXxJmmoahctpnXzMZmtJN06QxyE4f/AEZO1rUulNVUmmsmzGloEz0T0cNYuXUr1dNWFCYTFUSRtJGGbahCEfaJ6jFw1NVWZ1mv9pMyqq+uzCoG6WtqpZKmdwfJmkJb/wALDp9nyxi23SslnkMIaAaVBz2+GzetCy6PE7b5OCimn+wnsX09VpT55nWfazrlXeIsuj/Q2XuRe6ioqVeqm/8AVwxexxauktVrpCAw6A03kOgmUlWqMkoQ1RMnAG/MZ+9q1bjm7Ri/IIxoYaKokhqEaGOriQEsUXv1MfkZUtuUjzJT54fREPVLFT0dT9bdVTZCDKj7RweSHjIHUm4t1xz89vnnqHuK2I7HDDiGrff2ohFTPVZrQpmT15Dz1Mla0lRJIDcP3sznc4NiGSYMDa2EZ9Q0M85n+vZjDUreS7wRioT1YqLM3ueL89ecV1qTWOlNMVtQ2qdaZPQAX2UtGPrdab/caKElOD/xGAI8sQbMvpJ5JllHLl+jNAHMI5B/53qWYLTg+TJRwkInw3sD5jEUVgmnAdG0nfkOJwPkpHWyOPBx68letBnVVNV/V8gp4c0MZuqZfGzBCTy3cC7xn12bB5lcLZzrzTGmKp4+0HUOmsqqozu7uqrlerB87xwCSZD/AIjz63xyBrDt47S9YQNR55r6sjojwMtyaNaGkA9CkQRW+LBjiuTWRRn9RSIPeQlj/AY0Yf2dfJjK6nh8mn+kqnJpi5hGOPwPlegeS6s7NddV7VuWaoy3OIqSFnmaNDUvTqP953biOoRR5v4gPMjridUFJlLBPqef0kjhLx97LIota4K99f43ST8ceaenM5zHLs4ps1ySrky3NqKT6xR1VMxRkkXkfu/geDjurTecy5/pzKtX5ZAlPSZ7RQ5gsEBj7uKRrrUxJGzAlUnSWyqOAwHAtjM0tor6FzQHmh206NcaeBV7R9v+pBvNxGzrqql9dkOsFKfVWq2jZz3c0UC1UNyeu+Ahl8vtKfLm4BxGO2+k1VWdkWTwZzTOkemM9qqJoiA0ZhzKAzxPG20HY0tPVwsvAvKtxcDGzhfMsmrJI5mDnaH7z6jJTd2x9G5LcenHvjcZjUZ3rTRmrtHT5pWVgzHIZ6rLYJqsVEaZll7rmFP3bXuGdYKiOxF/GBiPR8hhtLBUUOHHLmpbW0PhJ2Y/PJeeNfF9XzKejZjM0TtGVPBkjPQg+d12n1vhmEke4EZsU3NyRcL1b5fDEk13Qw0mbd4qs8EyHYR1shujD1JieI/PEZlYqqzXIs3iZuSQR0sfK2PR2OvtDhrXGPaWOLTqTBGWnzKFwSVEim59L847y+ifnf6T7DsrykVkazZLmuZ5UVceLYZIZ0AbyF5pePO+OC8xTu5vCLW9L29rXx1v9DHM4MwyPW2nq800UFFXZdnCSSu4b9cslO6IEUkks8Pp0POOa/aqIusXaD+Ug8Td/wBy19CPu2i7tBHv7Lp6my/MKS9e0F6eWVonlhAcAm4G5VN1Xyv0W4HniSuZ0oJ4YKnuJTTSbZFuGQlSVNwpI5HHTgfDETp1paOI1FI86bLyGSGiZt23kgEyi446AY29FmlBLJTzU1alRHOAwNP3agE2IAKsRc2PHDDzIx528ldeGhcf/TVoo6ftv1ZUwwqkeZZmc4iBQEKtfRUlaCAfVnlPzOOeZZIHYIIQhcXaMC6ggfaW/IHsenrYY6y+m3kvc5zp/Pm3g12mMpeSR9rHvqOeqy6QG9wbIlPe/rjlSUyhG7mUiOLx7I2I2kc7lvfafQXINjj13R0vbQB48f8AEA7/AHLgbZH2clOsDd9lrqmF+5jL+FXvtZuN1uPkP4+eJX9H+eSh7ZdC1I4NPq7KW59frCjEYqTCCy7oyG+9HEVUkj0vYfhbGw7NMxXKNZ5RXN1o86yyqHl9ipU4ltYvQuA2KODCQFevWU1iUUckCxrF43DoP1gVgxF9wRDx6lb+t8HkqKdVemmrpzG9niZypKt5qLNf/wAQHBxp6ipllzjMqdKRkb61VKgjkSVoz3r2UiNt1x6hflh3QZhWSxRh5GVlXxBZXujefVTbn1A+GPG2uqF6DdxwVQ/S1RJ+yChqwYmFFqlIuA9z9YyevToxIt+rH2TbHnVL3JWGWC5M0Yfu2Um3H7Q4Zb9Dwb2BHnj0f+liJqnsLqJRLE4h1jkbXRYbgPBmERuY7Ei7gDcoPPnjzao4nSCFpXkIWMA2G7aLdeDfj09L49J/Zj/8Y6/mcuQ0yKWg9ampVaeaUA7CTJbauwgP5sOg8gDb44VVNtkhWS7tssLbt/7It1PoRY2PthNQIL9VKdCL2DDkc3sL9b++AjnCFhDLERKo/Uuga/P2WU+h6H8DjpFjpvXoHlNOFjQRgozjfd2/vA3APl5DjGsGRSMO8B8Pnbnb6X9Mbjvu7YyywR7D17u9hx9pebj3U8HyxfP0c+wbJu3nK6yGl1LBk+bZRDFPVLmtf3EE8U8jJE0BSnlduUKurWsbEXBxFLMyEXpDQbSpI43ym6wVK5xiyW4/WSIhv99gBb1vgs+QSoGcE7V5JIt8/bHbVT/o7s0R52PaRpecgL3hGbTEqL3XrQeHDmD6BtQlKslZr/TCIOCzZ73Nr9PE9CAPnioNK2KtBK3iFY+gtVK9meBXHuUXyqiakRwSzXlXwncbeYYEceV/44UqqqpkVb1DAMCFLbWKg8eQAANiLW5t0xO+3Lsrg7G9WtoZc1OYV9FLLTZj3Esc1PSzpsYQo+xDIQsiFhYKrNZS3OK3Ina/eyiYBiWJG25b7QJ8r2BsfMXHne6yRsjbzDUKq5hYbrs0hU1TGmjvUozEWZRcMFHTn7y2PHmOfLHYf+jQ0wtZ2m53n1REjx5Xo2SJO8jVx39fWJGBtYgMTGr+HqfLHGmZmXurPKGVE2RhrBgvNgPUXP547/8A9HlkcD6O13W1tJBUR5lm2UZGsMlNBUCWOkppJ3URTEK4DOhNvEOCMZWnJeysjzu9aD3V7R0d+do3/n2XcQzCmoKmSkidY+7NgkSywC/mDEWuPkLe2EKjN6an72oCteJWlLr3TMFAuTt/VyHp5XJ8r4YCaY0oy6nrnpo4rARDcIwqm5Qw1IYKCOPBKLeWFFopaxhTIkjLUMsSpbcLuQotuVo2tfqHDcdMeZmpyXX0AGK82/8ASLZ1+n+3ary2OpMoynKckyxgQRZhSvVycHkc1KXB55xyx+iFO4Mm29uLfv8AMfHpi8fpR6li1Z2864zqKdWhrtQZqIpC+0dxFKlJD4rG11pjip37qSWNTcxP9oK+xgVHILC46WIIFiD6g49W0ay5ZmjrDD2XE2o3pSescfdRuTJw7bVPjvYW5BHmb+VsJSZOU2gFm3gFdo5+PwxI0eTc1MxhZ/Epfox56E9D8evkb4LJA1xE0hIF/DssC3U9CSL9bHg/HF5VlEJaV4vEfsjzx6I/6KWgr8q1Z2j50JSKePTmVUkpJspqJqsSR38rhEY44OqaSOciIDxS7Il8NrszAfx+OPTT/Rz6dko+ybVOrFSeE6h1g0EDxy92Xp6CmWNfusGXfM/hItxjF07aOwsbyNlOJA+Vo6Oj7Wdo3+lSuy82zGoq6CSCrW6zAKHK8XuCCA3mCARZje2GktNLI8+YJNsk2k32F1ufdQJAOfPfgEMqQFVZo4w6SSNCBA5APKuAGidT0YFVJHQg4f0eYQ01TG1TBEYknUs5UqBFfcxup4IVSbEWI88ecmryC4rqTRgNAvK36dmqoc4+kdrukoau0dPqCKBWFiBJS5VBBJweOWNjceXOOcfrdXGY3Sba17Ewkd2Vt0A5Cng+HlSOluhkXaXqh9e6+znWlRKS2dVtbm73IQ/65VSSIOeARCkfXEbY060+9GUKwLBhYggWvyBY/wCED3Jx6pYY+zs7W9bFxtodekSkLU0sjyTln8IYcgMSehF/K1ze3A5A5thzldE+YZq60FOXrUiKUaLcmSqnK09Ovrfe7t8hhvBUfV45KdU3EnbtFnAB5+yfDuv8LYvP6DuhF119IXR4nhM9LlFZPqytRU3E0+WranXaOu+pYC3nh1sk7OEk9DM8qpIG3ninWzmvV3Qum6Hs30TkPZ9Q0ki0Gl8rpMmiYxnY3cxKruWW4u0m9rnabnEqps4Vx9UaeSe63COVeTb+9x72J9zjRpPE8g+tI5qiCzmSO0l/O7AJL1Po3zw3qqnKXgd5opSKckyL3ay7bcllI2k2HJtscW6HHmMk75Hl4Oa66OJrWhpGS3OY01JnGXSZVmVBR5llc11lpZo0mga/k0EwaM/gpxQ2vvoWdi2sIJzk2W1GmJZbkx5QyPRsf7+XVe+nP/qmib0tiwW7U9Ax6ybs+XXmTy6khp4ag5bLWLFWNBMgeNl7wDvQVINhuNiNw88Szv2aQgxsrKAHR4z4QeneIeQPc3X3t0uWe22uyULHEDl8KGSCCf7gD1xXnh2jf6OntAyPdV6OjptQ0oBKJkjmKpCj1y6rcMfdYKhreS45X1R2a6x0/nU2TVOVVBrqQnfSdxJFWxepeknVZx8QrrfzNse1b5gZ4e4d4p6dW7vaV+yR9w7vMdAHsw8nIscMtT6S0x2g5SuUay05lOpKJF/V0+cUgqu5Pl3Zf9bCfeOQW8jjdsv7Tit2dvmPhZ82h8KxupuPyvDqk+t5bNIKGaSmcEpPGp2qABciRG69Dfct/KxvhhXQZFXSRvU5bLR1LHifLlEbBh0YxHwX/wAJU49Tu1X6BPZtraF6zR2dV2Q1wjutHmayZpSpbyjmutbCo92mUfsnHDvbP9E7tW7IKU5zmmTz1GRvMIRmtFMtVQbibKDUx2EZJsNk6RHm3OOis1vs1s/gvBOzXwWTNZZrP97aKF5L2rdp1I9BFV19Nr6nyi31Fq6WWPN8vUG/+rVKMlXDb0SRlHpi2pfpB9l/a1Srl/aXksGdZlAgiEeq5/qOdwkcFKXUNMitJbyjr4XHkXPXHNdXDUUVU8E8TRTwS7XV49hicdV2nofbi34YfVmdR5k6UWocoizMqNgkdmSVFA6rP9sAeYfco6WwpsbGm9Cbh3ZcMvQ70nbuIpJ3hvz45+yt/UH0f8h1HUJB2V6wZ8wqB3sWlNUmDKs4mHl9Wn3fUMyUeXcyxsf+HfFK6q0nm2lM3qdPalynMsqzWlJ77L66ilgqYvW8bgED3+yfInD/ACyuzPLqU5Tp3O4cwyqobcchz5UeCUn/AIZPg3ejKY39MWDk3bfmElPBoXXlDS5tk1KoWLTWuGlniph5/o/MxaqoenhG/ZfruGJGyyMwlHmPjMcxvTSxjj3D5H5/RULNlzOu4oI0PQk9f5/LGulgkgPNip5I/n6HjHS1b2N6H1rVRwdmWftk+f1Cb4dH6yq4qaomvyFy/M1tSVykmyhzDIeB4jim9X6Sz3Suc1On9V5DW5RmtIxSroa2lennh9C0bAG3oehHQnriYOa8VaaqMgtNColQ1tTROXpJOGFniYXVx6EeeL9+j59LPtM7EJP0XpLNYcx05VPeu0jnbtJltQT9poGJ3U0h5syFT67hxihKnL3jYsgK88A9cNWY7isynd6+f+eIJrOyYEOGfWPVVNHMWFezXY19JTs17eD9X0VmNRkWrikf1zSmcOgrjsQr/q7Gy1sVrfYtKAB4Tiy6muzOh7ivp0npKuGVoFqKebY0FhcgX8W08jYwIvcEHjHhnlupq6i7haktUxUzh4HEhSanYG4aKUeJCCAR5Y7E7Gv9IlqjJaaj072wU9XrjKaZO6gzRZEhz6jjIAKuzfqq5LAcSWfjhxjjLd+zckbu1shodnwdXn5EroLNpdpbcmxHPzGvrBdc687CNC60y+siyygoNO1GZSGprKRcsNXp7MZx1kq8tQhqWYm96uhZHF7sp6Y417ZPoeZ1paphqcnphlJr2K0dPXZnHU5ZXv8As5dnFlhkY+VPViCcdA0h69oaN7RNAdrOVSah0JrOgzfKJGpIpdk0lJVUVWWcLFJAf1sEjovG0uH7o2ZucWWlHHR5bXLmUM9RQVsO3MY58uV6asjA5FTFPaOov08Ue4ngG5xHZNP2mwv7G2tJp5Hz28jvT59GwWlnaQHPh+OsF4pai03nmk85qcg1HlVdk2Y0ZtUUtbE8FRGfVlYAj2IFj5E41DAvywuV5sfP1Pz/ACx199KbtC7LqbVmedmOmez2OXKdPTPRua3M5Kily+oEQMwy1SPrFH3ch2GITvASpHdjy5EcAW8Fj7m5B/gcdtZrQLQDhs5++0Lm5ouyTdt5DFCxUHkDkC/lY/uwFgpLKVb2AuOfj15wtID9k7z6cX8ubXOE2ZgpIJ2soPXqDyL+uLShRGkUDaxtyACByD6/wwYABWPi2HgcAg+t/TCbd6z3UE3Pi8PFz8PL8xhZVDG5Xwk/c8Pzvz88CEmxBuC5JI3bibkj1ueuCAG4Ejcn1sAfbj92Fdpsq8X6eHhb/D+hghZlBIY89bC24eh9v3YRCJtuwshJItYXJv8APBGuxN7gXubG4Pv74OxPUA2II8J558j6cYSN0G5TYDjpwMKhCitcbQRxa9rc4zBX27mKqQeOD5fHGYSiEk9ze4u3W9/LGJcMXfgcdPL2wJuPtHk8keZxmzxEAK3QXA/dhUIyMS20eK5sADxfy5xig7QebMeD0ufMYwWH2vLrYcE+vTj/ACwbcoSxUdeePby/rywiFjHkMPMWJtcXGDFQDcBdp4sTxf0/jgATe5uSOA1/L9x+GM32BsbXHiAsSfkemFQjFmYbyST1JYXPvzgnhKbxxfpexI9AcATufcFUEWJ46G3r6e2DkDao3Fub2PHJ8v8APCIQxgOSy3Bt5C9sKbzuVkCq3BUg8K3W48x8PlhNWAKkBjY3AY9fYj9+B7xm8RAIAAva4+Y/iMBxwQpl2Valp9M61gpq+pamyfPY2y+scf7hJWG2T/1UyxyfBT646Yheehpe4zWnhklhqHoqiGa5MMy3WTawPC7gQb2BVltji+Z+9DBrnYd9vUdG/LnHU2jdWtq7R2XZ9OsUtVLAMrzUhB3ktZSRgJIzdT3tLsNj96FiPPHMaZs9C2QeB9uXsFuaOmrVp8Ubt+0Ce0DS+Ra+ow7Z9lsZ0nm9Qg2vVSRQmfKaiaw8TTUay0xY9ZKMeZ55LEMlkKLzISLnyI6jHoX2Z6ffVsE+l5KoRrq2jOmGnkNxBmaSCpyWqv5bKtFiJ/YqiOl8cY9p2nZcj1dWIKA0seZRjNaaBhbunNxPAR5FJFlQj1jxf0Xa+2AaTqp5tz4ihHmq1us3Z1eNtfI/BqOCgsdP4v2vT3OHIQfZBsT8LfDACxvcmxHn+ODcgbuLkfZ5vzwf+uNlZaf5MculzijXPY6qSheoU1aUrrHM0Y5YIzgqrkAgEggX6HFj6o7U81/s1meV6C0/R6F0nHUrl1Xl+TSs9dXMVLKa6uf9dMrBW48MdwQIxiqHJhCyRsdyEMtx6evkcTLTTUua5s2UVMojpdV0X1ESHpHWLZqdz/6xUU+ztilacHgnL4x40ViIVaaZ9e6infpJDskYQwW4pqUW3f43P+fwwC1koH1eALSQNbcsVxcerH7Tf1xglUJlmZ5YTEzMd0e221wbMtvIhr4QL+IcC1h16db4siNuaiL3IzPucsihL+S9B6/LB4iwKFZAniHWwBPPH7x88JgDeFIXk8ndcc4DqCOFNrtuNvl/liRMRmNlts5uByfK3A+PvhKcfqyLMDbn8fTywdiCAD0Hkfz5wVuTZifS97m2BIlMjzOryqtirqCTu6qjmjrKduu2WJgyn8sdmmry/PaqOqoe/ijzakhzmkjCU1ninUSEIdik7GLoVJ3Db544lik+rVCSddpuR6jzx1P2N5/V5n2a0tN3scr6ZrXyxw8SyEQTXnpnBPKi/frx+yvrjnNPQDuTDVhx/IA81s6LloTGdft0eCsWiAp7mSqdo2U2WfLd9vXxRvcfLDjJp6LUGcx6R1LU0s2VZ8smnalJahj3cdYvdRTJvAdTFMYJB4j9lh0Jxp4swgEc0tTDOZCAUeIMoVg3JfjkFbi1wQbHkXGDSQVGZK9FT1ddvqlKQrHUoZASOCI5gCbGzDab8cY52JwieHnUtmQdows2rkzXGncw0ZXU+X9yYHkgPeholEkdRHI0M8ZNvuyxuPgRiKlJ5m3SMzE+bG+OovpY5EMynptcwURiTPYaXU4QLbYa5TBXJbyKZjSzXHl34xzY6lULGNlvyAw647yzEOZXrdyouVnqHU6380ySmJPW58vjg7QeDgAj88OjGBLYMCrgEEC3XkfPywDrwOPcnqcWVBVNKWY0VZDU9O6cN8QDzjr/AOjbqCszHszfIaeSSWXTOd1FOII6UVDGmq4+9jO3cpADwzXsbeLHH9SPvcEHzHni+vopZpD/AG1q8gqpVWHP8md4w1zeqo2EyHgE3KJKOAT4jbGD+0MAkshfrbjwz/yly1dESlk4btw4/mi6zpq/NqanjaaOaISExKkstZG/w2wtJYW8yAPLG00xn9LklTRagernkjyCoirZt8ySr3ETh3UsUVyhj7xbN6kc40sOYaajvEa+hsy3ZS1ufJtrKjdOouObHnASTZVmu7La+SGtgrIiHiMqKZAftArKlm46HdxjggCTWmK6s0IoSuTfpIaJXQuus60rSyl4NP5nU0UMl7h6eOXbC4Pnemlo2vin42LAxuoN1O08grY8j3Ht+GOsvpU6epaiPTOpFoqqJc2yaClrWqIwJGqaFmyydjbzMYy+QnzNj545OdrJ3aq5kUh/SxHDflj1CwyiWAOHVcfdcRa2FktHdUw9QmNcAY163Xgkm9x5W9MX79DzNvquus4oGkjVK/TTzAySrGokpJ4puWawXwRvyT54oeqAeM+FRax8PSx9MWd9FjNWyztq0oROIhWVFTlTEgEf6xTvGlwwIPjYcEEYraZi7WxSDcTwxHMKXRz7lpYd454LvTJMuzKEJHS01bD3cdhLDTiSFtvn3kFQ68WuLrh7l8EVVTiDM/0fJK7MSZKR6UTKTcNtkj2X9bEXONHlxpq6KGeko8uaoeC8l8vp1kQsvIJVUYG5te1ubgnC+VZyyiNpKSspWuBs74BlYcAeF1II/wCuPLS0ruQVWH0yNPUEvZxp6vykUqfVKrOsuaOlmSRIzLFSV8f2Cdt3iqDY+e7HF0xVnXvlSIHxBVbeoHmVYXNvZhce+O++3ipOo+yvOaaWTM5XyfMsjzLbVxPsWJ5KmhlMbsTe/wBch3AHiw9cefrMAQJo7uvEu0A8gW+yLHyPnj0f9nZL9jaNgHqR6ALjdLsu2g+PwfUlJ1Cm20MoVbqJGHPwP7Ptx8cNcmlMFe9UOBAYpefLbKhw7ldH3bTJtYcMQC1x52HTj3JHnjWUUyQ1skcpKx1KPCxbyDdCfnbG3Ji0rNjzC9fq3Np6rOaqCoMU8UtdJt3J30al/wBYoG5iejc7WUi4sPLDz63V0605XMIgrTd33c0xtZkO0L9YAZTu22Cvc3t6YpH6O3aXk3ab2aZdWzxUn9oNPx0+U5/C8IMgmQd3T1RIIOydFRdx471GUkF473OtRUz0b9yiLDIdksSG8Tn0ZXuAfQMPa9xjyCezus0ronjEdVXewytnjD2nNR3tB0ZXdpGks90XPFSUuZZpDST5bVT/AOrqK+jqRPBTzSMdqxTAywiRvsNIhJ2m+OFe0jsQzLRWfz5TQ1+WzxRsxSGuzGDLMxojfmnq6WpdJIpkPhNtyNYOjMrDHoJS0tNFWA/XzTxPE0UlFUq1PDJe1mVmBVSBcGxKkHkeeNjmOZZ9QS0cBmrlpYUESxVENPXRxxfd7tpEcqAfughbXAAIGNvQ+nBoyMxyNqNVFn6Q0Yba8PY6h3ry0TQGqJnvbKAqqVDLqLLrqOo/3xDDrwfXjAU/Z/qqQRwyy5OkaNYKc+oLBSbkgGXn2x6sZXnOaw7pBmEDo56fUqIqo8trCDg/EEH7y8E4fR5xnyWf9JUzd39pTQ0cIdT0v+qPdtf7MgLxk+E2HTY//l8Ff4buXys86AkH8w5/C8ms10BqKQR0URyjvKuXYvdZxT1LC54G2J2dj8FJ5IAx3Z9DrsV1b2M5DnWcavyuahq8/Sggo6arvT1UVHTGSXvZIzzC00sgKo9mCJuYC4GL6XOM8pq5ZUzipiaF9xWOnhpZSCLbW7uNWQ+6tY9QSMPKENIkklGsm53Lu8UzKWY8ksY3ZWPqzAE4ydKftF/aEZhiaQDnWnXNXbHon6Rwe9wPgnf6To5JUSaQQT1ULQiGcMgnUc2DA7WII3Aq1weg6jGtq6s0FLO1YrTU5hdZGILqVI6uF5t72I9duFnkziImOCvIiZbPCwR9zDzueD8G+TDES13rnKeznTOZ681bWJS5ZkMaz1KjvaapmkP+yp4RdkeWVhtUdLbmvtUnHONY6QhrBUnALVMgjBc40AXDX016lK36Q2uZYTvZNUZin7X2YqRefUeE4oaRP1D1CKkPdbYyWk3sL9AB0t5834xJdZ9oGd9pWe1msM+p6IZjmddXZvmBKnYZKyUOEABBsiou3kcWGIzUPEzRqjSXAIJZQCR7C/Hryceu2Bjo7O1jtVfUrhbSQ6So3egWtzCFTVU1KWDNPNGpO65sSPy9CDb4Y9LfoUQjJ/o/5RVBC6Z5n2dZxKrwJLCQssdNHvUjeOInsy2tc3x5lV00dLmVNLHIjLHMjllBC8HyBAOPRz6Dmpv7R9gw0/R1EMlbofPKykqaZ6YTMtLWN38Ew2kSKC4nQlTa4FxjG/agP+jJbu9f0Whoa79QL2/0XUi5vlwrFhEVRSyPAJY1jlDxG3BChiGBBI6NwDfCunM1ol1flhoa2P6ylUshjG1ZWKgkDa6pL1A6GQH88RyHM1YIlYkkcMTXWSFu+jU+4YB0+TH4YGtMCBa3L3dp6KRJ0elYboHUgpJtIOzkDyAPToTjzpryDeXWPjDmlq8kc9qp63NBWzhu+kiVnJtfvHlmd+oPO5z5eRwzjMfe371hISpR0VSUNyFJBsCDyu0dQb47y7Svoi5BriqzLPtA0WXQDMayWu+oTZomX1NBNKxaaKmqXVqWemaQs6wVAikiLMquy4pDVf0K+1fLN7x6R1VIACGlGSx5jFa45EmW1EpPTr3fwx61YbZZ54WmJwPrwXDWmzzRyG+0hc5kQSVCt0jDB9kYCm3lz5HnjyuRx54EiKOpR3jjbYNrA7vGfYgggD7pHPX3xPM47JNVZNMaSrkylqi5Bpaiu+oz3v8A8GtWBwfbnEbzjResMrhkq8z0vmlPT2utQ9K7wE+olj3J+DeV/jfqCqtCo/SbZc0pZJhxC7TvtsFCxqWsAPfbyefPHrx9D7Tkekvo39nGRzsizV2TNns0bToshlrZ5Jb7JQEa6939mQH1GPKfR2gNYa4zGLLchyiWprM9kGTZTEkdjVVUzAMI/NlRAzO/IVRckY9mMjajynT2XaWyDOIazLcjoKXKYBTzrKgWngSIsEILRklCeAVYG/njj/2otAuNjBzPp+q3tDwkuvbvX9FIqiuSgZI6tjTljZRKXpSR/dMgMbfASWOIR21albRPY12haugrCjZZpfMHp0kRopkqJozTwlLXSRS8wPhbwm1uDjaR1UlClUacpGsSCST6uWjupNrlIzY+d7x8YoX6bWt8vyb6O8+XRzR31PqKgy5+5aGzU1KHrZzeI2a/dRg3CnxC4xy+j2iW1Rtpr9MfZa9srHA413cV5q1Mb0maVkFPKuyCUUiMQCSIUWIAHyN1bnDUwzyxOuxmLg2Yi+63mw43W/aHiHnfCyzkwCQ2Ekg3zMBzve7ve/BuWPX26HCHfQ90VRGva77GKhreRU8MAPWx/LHq0bbjA06guMebzi7agnkanpJp6NLSiO6B3D2duAYyOGuSLE2I+V8d/wD+ja0TJkul9d9ozwPI9bVUujMvCRoztTUiietdQxG5WkaEGxv1x5/01QiV0UgQvDS3zCUbCgKRDcBZvVyo9LnHrz9HbSMXZz2A6D7P8xhiFfHlgzPMYZ9oLZhXH6zKWjkAEhRZI0ukisChFjxjnv2jtXY2YsGbsOOfIc1p6KivzB2zHh+SrVbM7xqp3JtsUjZnRb+wk3bPY3t8RhamqKXMDDl1T3ivVNss4WN0Q/acqx2yKouxMbm1r24xojQSAtTLHGzkXMMIJkI/9BLZyPZWI98Vh2964fs07FtbZ/l0pp6xMsOV5eqzNsFfXt9ViYI4DxuiPM/PI2XBtjg4I3SytiGs0XSzPEcZeNS8yPpS67p9fdteo+0DcaiDP66apy/a+1qehikMFJtI9YoVYezDG87Jvpu/SG7KEgy7T/aG+dZRANq5PqaP67Aq+ao7nfGP8Dr8MVDqHZnGaVM9MqmliK0tPtNwIYlEadPZb/PGjlyt1NkPyx6jBZGOs7Y5BXDI4+XlkuOkmuykt46/FemvZf8A6TLszzOFMt7WtIZ5o6sayrmNATmdCB6FTtnVPa8thxjqnR3afofXGTPm+gdcZTqWg4fv8uqVmaK/XvI1AkjP+ONbeZOPCGCfMsvBEMzBD1Q+JD8QeMbLT+pM0yjOqbONP1tXk2bU7b6esy6d4JFcdCCpBU/A/LGRaf2cs0g/d9zwxHA+xV+HSsrSLxvbjnx+V720WaifbMKhnhJuGADJf1DLcA/NcFraua07VMcFRFVRtBO7wKyVMLCxinQ3iqEI4IPUeQPOKD+i/wBq+cdsHZJpztC1NT/+WKgVWW5pUQU+z6zU0sip9aBiZHQyI6bgAy7wxsL4uSSuM9PK1FUO0kIu9izSKPVnjUSKPeSJ19TjiJIpbJO6KtHNOrdrXRMeyeIPpgQuA/pjfRmoOz+pj17oWgaLTmZzij+o7jIcorSrOlOrtzJSyosjU5bxRvG0JuNuOO6tOq1EZKhuHYk7r/tD7N/fi4Ivzj2A7Xsvj1V2Saz07V904qtPVuYUbhF8NXQxmtgkRkvFJaSntdSpG9roL48pNdJRRZ9MctiWKGr2TwMt7qkiiQJ727wL68fDHoWhNJG3w0f9wwPLHmuV0jZRZpTdyPXso9TinjhLM2/ht8bIQjC/NyDd+etrD3wtTZpIlIKCsEFbSqDekq0Z4x7xn7UZ8vARbzvhqZTGzjvQJYxYBQGMYHBFzx52C3wVW3ssdhtCru2dT53sTbpbjjG0s9OxPCKYUWT5gRT33tkubOJIGPX9TLwLnyvsb3OLp0L2p1mY6V09kXahpaLWGhKiSpjpqLUNSTWZXDTqDPLlmYp/rNMq3ChG3xFrDYecUbTQPM5lipVd7M/dqNzSuSFSO3ldyotbFkdqKJpWih0lBBSqMlpk03HJDYmVoSJa6Qm/iL1cjJf9lSuKs4uDuYEqeHvHv5BVrmy5bVZhWSZdFJS0bSu1PFUTCR44ixKIzgDewWwJAFyCbC+NLW5fJu5jKjaSNwsTb0xuO6p3iLSptaNPFvBTx+nIvyehHPPphsUBEoVSnd2C2N9txfd6c3+FsWgKABQE1NVopaaSE3Q8HGx0xkB1BnMFDLOKWn2PUVM/lDTxqWeT/wAINvU2wer+yEUeKU7S56gdSQBwOMXz9D7smg7Se0HKKbNqR5MrrKmTMc1RFuf0JlwWaoQf+mlEEA9dxxWtcvYx1BoThX1PkKnyU9mZ2j8cQMfx5mg813R9Fzsvg7F+yXJqGShkoc1zUDU+bBrmWGeqjApKdm+1eCjK39HqH9MTDt07VIezjsuznXsn1eSXJoEnoIzIZBV5hKxjoYbHjaZv1xHN0pz0GNnm2a1mbGqfvV+t1peqlsnCFnDOQl1JQcIdpuoUdLY40+nnr2ojzTIey8FFTI6cakzeCJyy/pGsQpRwknn9VTDdzyO9OPOLO/8AtXSRmcKtrXyGDR54DiurmZ9FZBEM8vM5n1XIuZVtTK7JXV71NTOzVFXPKSxllZizFj1JZyWPy9cauU3sTYKy23KvP49fa3XCpqN4/wBmS7HqOnx9/Yeg5whKb2LFOTcKvAY9L26g26npj0uCPs2UOZxPiuRlffdUZakQsxG9FCkEm4JIB9genz6eWCG5NrkHiy7b/l/D54M5O3vNpKMbFuVPT93xwmd3CuCBzcEA9f38+eJlGsCgybrAsDYW8z6YMpa3dqGHpx0b1t/XTBTJtQrdbm1wQSTbyt5YLvKsCqqfUEHn5X/LAhKoxMZ8YtcEs6+En5dMFZr9b+LgFzuVffBrMWK3LsfMC5562t18sFZiSLH7Vze3n7DywiESQ8klep4587/174ScAngjyJv/ABwoeSTvIPHNrX+F8Jy8te5AHT2/nhUJEE7gVCgjqCbYzAbASb8i9wDjMIhYCT158zxwev8ALrg6gnaDceG/A5AxnIcXYD1IHIB/r49cYosQAbEHy9fjhULASCLyegt5AHGLwxsBYg9B5Dn54ArchbDrawPX8+MDe92P2uoI8vc4EIQ1iwWwL9Be9z/P4YMP7vPFgLA3+Rwi11fkH0AuOnXCotbaT79LdfLCIQ2KRtc2sdtj1P8AdP7/AGOADLYEsDf1uAfYi2BctbbYhh046A+vvjPsqV2geI+IPcW9Px88CEUuCPGOTz9qxPv8f34Kb337lYsCCRfr8/PGMTuIB28ng9fXywhIxvY2uenXn54EIxn2OrG5APN+pB4OLa+j5nM36ar9CgszZ9EDQAMARmVNeSmsTwDIveQ+/eDFNuxJsMbXT2a1mT5pSZxl0pjq6KWOqp3B5E0LB1P5Yp26AWiFzBmfUYjgVZs0vZvB63rs/ReoZJZf0VluZ1UcGZ0yVFK0Ue1toN1kC+UsEyruA5FvQ2xqfpiach1BlVN2tZdSRxvX91qlkiWwjNVKabNoPhFmUJkA/ZrR5Y17tRDOUqckpXjoayaHOMr3sWiC1cfeqoYEMvJkjNiG8APJW2LsqMqpdf8AZfnWR1MdqWlmizikLvvb9FZ0goK9Gb1gzGKjmPub8Xxytjm7GS8NYveYzHiW1Hkuglj7Zl04/wAvHI+AdRedM6pG5jViEU+D/CeQQPgcJbW23UqQObWtf2A9bc4dVVNXZbJNl1chjqqGZ6OeNhyjoSPxBDj5YZO20cn7P9eeO2BqKhcsRQ0RauQbdoI9yPP3w8ymeSegmpI5Ck9Kwqqdh1Uqbm378aqYm9iOfS/TCuWVbUNbFUjkKw3D1XzGI5m3mYZhSROuuxyKsbXdMmcpFqqkQ9zqGn/TKn9mqB7uviv6iUCS37LA+eIAVudsd2PldbH3H9emLFyeq77S+caSSkSoOQ1Y1JQSBby/U2QR1kQPmpjMUhH/AHLHEDzGkfLa2ehZie4kZVPkR1VvgVsfe+ILK4UucPD9KedU+dpreTVVJUMQSCSBbz9cGAIYG53X4BH9fDAIwBPXdfoV4P8ALCha4UFFI87k/O5878H8LYuKukTcL18NwD74w8M1gOR6DphUsGPKE88jdyfPg249cEtcbn3WPBPHW2BImcwsbj54uH6O+oWg1DU6Yl8UWpsvejjUsVArac9/TG45BJUx8f8AEOKjnQlASPjYefvh/pjOKzJMyp80y9ttXl1RFX05v/vImDW49RfFLSFn+ps7o9Zy8dR8jQq1ZJeylDj1t5LsWKimlydM2ihnFI4UNUQy7xA/UBnWzROPLeu0+4wvR09ZPH3cr0tQjkbolFo5j+13R8UMnneJit+QPLDfLswhFd+kKBk+rVkUdbTOrGN/q86iSPxpyFAbaQwIBU8Yd1Qy1IKaekWq2VqzmaKSEKYJY3UMqshKyowZWDqAV6WvfHCl94YDPrr9F1QbTMrO1jL6jVnYqN9HUQ1+ms7lyyWKXxMcuzhO9hbcOCqZjRmx9Z/K+OLpQA24ghurD09jjvDTgg1bSV+gYuKvVeSVeV0p22/8oxAVtAy2O0/61SIoIt/tmBAPB4019Q0S6rrqrLUkSkryuY06Mo4iqEEoXjptLsv/AC46zQ85dEGO2en4I4Ln9JRXXlw6r+QeKjewFgCDGBYEsSdpHp529vLAS7rlGb7AJIHIHNyR6jz+eHdHQ1OYSCKlp555gNoWGMyNcdOBe3+WJnorsT7RdfzsmnNOVlV3P+1NLA8/depcp4Ixb/iOoxrvlZH9xWc2Nz/tCrWpS4ut/wABz+GLS+jYaym7U9L5kgkVKCpqqxmXgrFHTvuN/K5IX4kYsXK/or5TTU80+rdfZVSLTuIZ4qBXzeaOQ9EZKO8CMR0ElSMTbSmi9E6QzKtiyKhziRQiUi1uZRokjRrZtvdQHbCpezbVMhNl3McYeldIwugfE3EkEcRT3WpYLHIJA92AFDzVhz5/DVQwxVtP+uWMXWZUlJIAuRLGFcc/3WF78YTgFFAqzwb+4KkkPt8Bt0bwlPg3AP8AdPGEEGUV7cSyUc/HeGCrSpglA6FoZgkin4E4ez5XV9zEuSvHmFczrHTUiPLTTVDsbLFG6llDMTbnwjktYAnHGVaDQCi6bGlSof25V+V0vZPl8E2WJFJX6gzuCP8AWnc8f6KpnlmA3MLJPHTmymxI9TjjavD/AFqWUG3eSGQEADrzcHr1xcP0htcRam1FT6Y07mKV+TaUpZckgrKc3TMMwmkM2Y1Uf7SF7opHBSOM+eKZeQSjvUR+PtNt/r8+cd9oiJ8NnAfnQdcxwXI2+VsspLU1n3Mu5rAkHoOpv5e3n+ONp2eZ0NN6yybUTkgZRm1DmB+EU6k/lhhMm0FwysCL7lO4X8+f54Z05CzTRDnvYXUD3tcfmMaEzBIwtORVWJ114IXqFmdNWxZrWUUKrspKmSKLu5m2qoc7CnfRAgbStrORb1w4jgzOKjnTMIamjljZdkhDIHUi4YbSUPPH2ef31bF279ldNQ0Ga5r2xZRTNW0FFUPTRiprJoZGpYe+iZIYfCUkVwAznr1xrs4+l92KwU7UtNU6xz11lV4pKHK4KJeAfvTyMfmEB55Bx5ZHoq3OcA2MnypzOC7Z1vsoGLwrG1FTyZ1oTtAyhadHqZNHZlPEVjUMZaR6etVWClbG1M5Fo+bdccBary6Kg1HmdOFDxJWSvHYjlGbcOvThhyMdIH6bunMsn+sZR2HzVBVJ4hPmmpmVjHNBJBKvdwQogDRyuNput7G3GOadQZlR1lbDUUdZHNG1HSl33G0cgiVGVuhBBXkj483x22grPPZG9nM2mew7Nnmuc0nNHaHX4zs9/wALUksN6q1iRcqCLj0P/wCdjU18d370FebX2nz+Hl8MbnfJJF3JmbZuLKt7BW8/6HHnhvVRo8fJF7AC/B/d+eOiWQtt2b9puquzXUtPqvSOaiizOnjancSoJKetp2Fnp6iM+GSNgLFW4NgeCAR2XoD6ZfZvn8ZXUstVovOpFAeGsD1WXTG/IWcKzqPRZo3A85Da+OBpYWVvCMKQVVVCAgAdB9yRdw/PpjJt+iLPb8ZBjtGB68fJXrNbpbN9hwXrRlnaJpbMsokzHLtdaHdGhMkTpqSiSOUj7hSSXwsebG2248Si9wZteaOMaoe0XQ8RiNmp/wC09CY3H7SESloj7eJf8QPHk2cyU/ay6lJ9dp/ngn6RAJIoab5qT/HGJ/8AxOI5uPJaf9uSDIDmvWpu0Xs8jjkSbtL0bEziwc6goHMbftWEtmHqvF+oKsOXNH2q9m8dkqe1TQvd3udmpqMbT0JVi1728yORww88eRq5lbpQ0v8A4D/PB/0nIePqVN/8LP8APB//ABKDW53L4Sf27LsHA/K9d27V+y3upUPaxoMorsIVTUtMbpfhrFv1ZItdQePU9cN/+2TseS/6Q7W9AxBFLb6jO6SdVAHkY270ewU39MeSX6Tm2f8AmVNY/wDdn+eERWSu9xSwXHon+eGt/ZGEGt93L4Q7TkhGQ5/K9LdXfTh7C9IIw09qHONbZghIjoMpgkWkc+hnrE72Me6d5jjTt7+kFr7t0zWmn1Y9PQZXQSO2Vaay8t9WpC3WSS5JklYWDOxLHoAq8YqVqrNpgI0k7pT92JRH+Nv442GXUkEaEBleVWtYAk368i3FvO/542LBoOz2E32DHacT5YADyFd6oWnSMloFHHDl+VsMvpKhKNagku7AyNdrFhflh6gHrb7NhcW5wSVZe6deFjVgzEKCeegDenmBew5OFahW+9JPGyBSLuri46FWUA/xt5nBizqUvdt0d5JDtNlboQUF146lr3GNsCizq4rRZ3CUABUlrbiSm1vnbg4n3YF256k7EtaQ6uyOFK6nlgNBnWUyyMkWaUJIJQsvKOpAZJB4kdVYX5BiOYUu+IRLEvN9tgAXHl4r2cfDnEcmhlpZN63WxuCOLHEU0LZmljxUFSRyGM1C9hezTtU0n2k6eGuNA58+b5XZRVCpsK7LHP8Aua0IQ8TA8CXxQydVbkqJJPnFHVVSNLRLFIy2VpFured0miAdD8VYHrjx30V2hap0DnsGp9G6kzLT+cU/Eddl0zRvY9VYD7SnzU3B8wcdEad+n52oxmCLVekNFanMdi9R9VfLal7erUzohJ9TH1xw1u/ZaRpLrI4U2O+V0lm000gCcY7R8L0IFLS1VCZ2mlgJOxKiOoBZWHkJorEH2dR7g4JQQVMMhWsWkqpYXG2p2rTzofLdsWwb+9ZR7XxyXln+kJ0FUJI+f9kGpsnmmQRzS5VnEFbEyjp+rqIlLfAt04N8T7TH00Po0ZoYTUa/zzIiAE7jNtOP3UZ/aR6d5dnoVtt9Ap5xjzaDt8QxjJ8KH0WhHpSzu/np41C6hrM4kzTLjQ5vNVSxNwaWrcVMRt/cmWaM/EoMQnMeyfsrzYSVdR2Z6VWeQWepyyJ8mnb2MuXSRrf/ABR2xH8i7WuzLU6hdFdrWhs3mL7lpBnsVNO6cfYjqe7sevBUX6X5uJVSxZ9fvqfT2azD7QkplusqHnwEl4JCP2QY2/Zv5wC0W+xHFz28acMk8Nstp1NJ8q/K1GQdn/ZvoStap0LpikybM6qNoKqfMK+efMaiBusEdZVMwWI+aQlN3G64uMTWkpYHkhMtMaacJ3QhqUMRZOoCm+02+6Ua/lY412Q1WXVz1FL/AKvNJEx7+F6dqSoiv+2i3Hry0VvfCk1E0ANHkMkg3tc08Sgo/wAEG+CT5BT6gYpy2iS1O7SV14qwyNkIuMFAt88kEEwjmaQzL4VSSpj3f8qziMn/AJZPhjhv/SNayizHUeQ9n0KkTZBlElfmCsmx46vMpI0hRwWYhhTxbuSTZxzbF7dsP0iNMdgOTzLqStizLUky2yzSdNOS9VJ901UN3WmpwbFjcO4G1FFyw82dXax1DrjU+Zap1Jm5zHM8zr3zHM6ngLPVsCPB5KkSnagHhHlwMdT+zujnySttLxQDLftPhSo8T4rD0ramhpiaa7euaZTSd6JNxu62ABs3HoQ9jb4EjCLuxEpLkBQCbn7vkOhY/C5/I4CqaRfCkRVDbaCwcsvmLiw568AC/njKqdKaNpneElEJAsY2VR5bDzf4E35uTjv1zCmHYRoSLtL7WdL6J8Ipc9ziCmrXcGy0FP8A6zWuRfhRGnPwx61vqKGqFRmMU8sENXK86xPGFQo7EhbqXppLCwBPdPYc844R+gXpEQamz/W1RHQsckyinyKH67KsSNW5ixmqNhYgb1p4nWxZft2vjtLvoI6sRz0klJVkbghDozKfMWZWce6O49jjzr9p7WZrQ2MZAcz+KLrNDWe5GXnM+35Ugire+CUy3khdbrDEVdGb2gluAf8A0bj1F8arWOSaN7RMhm0Vr3TkeeZcsqVK0VQ0+6GZAQkqIzLUQuoZrPFI4sT4cHCU1VE1MJoy5ttSSNKheT0ZD3U45I/at5Y1GmdVaU1bWZjp3JdTZbmWY5RVzUlZQ0NYMx+rzxMVdTTS7aqKxUi6bl44Jxz0b5oz2kVajWNXwtZ7YndyTXqOtUXrz6FvZvqNpa7TuefUZ1vePOKd51X2FfRqtQnl/wCcU8trck4587RfoWdo2kqWoziLKax8tRS36SiRc1y/b5H65RBjD8Z4I7DqRj0Np6yKOrAJhnYblLJUiTYV68SWkjI8wSbeeBnqKIT02ZxwS0dTQzCamradnjMbWsw3xuDZhww3kH9nHSWL9q7TGKWgB/I/Hosq0aCgeaxEt5heOmoNH53lCCozGhBpW8K1tNKlRSuf/SxFlB9iQfYYZadyylhzaTNJlSakyulkrJtpupexWJLjzaRlFvj6Y9gNR9m3Zz2hVJn1VpLKq+unBP6Rpx9SzL1J+tUwSR/cTJKPXFe0v0XOyXSup6bUtLlOeZzPQTCqp1zbMaaeioqheY55KSmponmZT9lpropsWDC+Np/7T2GWA1JB2Ux8jl4LOboa0xyigBG2qmHYboKj7MuxrRPZ28iJmOUZUr5pFJThwMyqWNTUgScMGQyLEbEj9VyosTiY1VXU07hYVE7wKHCJudk/vhQRKn+OFzb36YYpQZoUYyz94alQoeZ1BeQm/LX2uSbm27km6nnjVZtV1WW1a0GZqtJI5URx1Y2xvJ6K7WG8/slkk9C3XHn89pfaZnTPFC41XTxQthjbE04AKM9u/ahDprsX1fmpjjfNM1pP0DQVHeIyvVV4anLSOm3eyQmof9ZEsgCXLMOceWVVUNmFZJOjuaeeUmHcCwMa2RLr5+BR0sb25x1n9P8A7Qa2au0n2bNPKJaShkzysDsS4qawmCmVifEe7po5HAYkr3x55xyNL3IiWGBVCN4QClxbooIPDcdOgHuceh/sxZTFZe2dm7HrgD5rktLSh81wauvlJ/Vipl8ZA2EEHgbSw4IN+B15JsbYIpUq67luzXsRsJ9ARY2Hn19MGaSKCmkiEcZtY3Sy+wsL7WF/NRe/njIor1C5eu15ZXEQZJdyhy1gRbhlBPJPpxxjplkqd9mGWVkecrqKnovrLZJBJnohZd3eSwsIqKMj7xerlj489mI1rCueXPYsmqKhZWyyMwyyqzSpNMrEyyE9TvmaVr9bbcWrl1K+leyOu1bQQx99XVqVcUhnQNDR07PSZd+rYhiJKo1E11vb6upNhzigKVxU1klQxba/hS7fdHAuT7DFNp7ac7G+qsGkcW8+i3U57yZDdmBuUJYki3UAnnpyL82vhlKtNvZnRCVuLlAwQXuQASB19fs/PC0imSZIS7eGPfv7tiHuQCQRY+Gw5684RlWGlQxg/ZueTyVHJNulvwv5i+Lir0TWSKasnShpowZZWSlhVVtd5CL8D2sPnj0f+iL2fZZp7s9zvP56MVCZxVRaboRuYf6hlbiaqlBU7j3le6LxyRTnqOMcBdmtNLUalGcCAzvk1NJmSRKL97VMQlPGB6mVoxj1dynSg7ONLZD2e0xknGksspsnkMC7nnq1BlrJRf8AaqpZiSfJRjk/2ltZhic1p/8AEeLsTwaKf/S3tEWe+4E+PkMuJP8AlWyXM6TMmXMs3rqf9HUrGrrqpikyx0cSmaqkDW4buo2s9g1yL8Y8s+2TXGY9omr811xmu+Ot1VXzZ5MhP+xjmYini+CU6qAPfHeX0lM8GmewjOoYaqogqNYVtPpOMSRlDHFKTVV7C/isKaFUN7270gG1hjzZ1DmzZvnE2YW7sTnvFQG2xDxGvsBGq/jil+zNjNO1Izx8hgOJrwCm0vPiWjVh5nPlREMiKpcWYngWYhR7kWF/YfywjIlws7AbTb7RsXN+bH09+mDwkxgBDY349v4fwwLxkgsUIU3YOzXTjrz1I9LeeO4XOJOQlXCvu68i23b/AA59R1wmVZSl0CeG4svX+9c8i/4eg88LGMlxGoYMCd3IAPyPT8efbGD9YVEW4szWBawF/c3/AMsKhNzdbWcLz4d3AuOo5vfAfZPIU7hyPT4+h9sLOhZQy7iUAswF7KOnHUAf0MJEi/Cm688jrfp0/D88CEIViqlgACOPQ/PzP4YBwLAEWta9xbafLnBkkDHeWVSxItfyA6f3v3nBSengCjddbckfzt6eXlhEIrWBbjcGvazXufXkcjCL82PekjoL82H8PhhU3JN1aw6XPX4/v45xjhQwJuxsOTxf+v4XwITYix4X4i/4H44zCrMNosQbfl7YzAhJkfeLAX5AA/rnGX8PXnrz0HsfPGMF9CfIi9v69vngGJPTxDzNrXwqEDvzYeQBt1ODKx4+14Rzzbz6m/T098EJBI8/LkXB98KFLKdxNr2JJvY+/nbAhARbadwF+LHw9PL44AgDiw4BuL3t88KnfyCDawN+SPjcXwmyneVLKb+a2Y/j0wiFnhAvc8/gD5jj8sCWF2Ba3Plzz6YJaQH7HIBv7YxXI+yAD0sTcEeh9/fAhFcqCejccAGwPthrK3ubfDCznqBz6eROG8htwb+nXCoSbeXFsLUsvdNuH3GDj5dfywhfB0ba4Pl5/DCEVFErTQrp7sszBM+7N6GhatC1mTVU+RhWfloZAaujYX/ZcVCfAgY6F7Gq/wDTGawaerD3NHn61Wn533eCODOIu7BN+ndZhHTyg+Vz545P+jhn9XSZrnWn4ZOc6ySYwrsVr1lCwqYrXB5KLIvrZji8NKZjVfpCqymgn2SVsNXBRSLwEmG2qpH+TAW+GOKtjXWe0lw1G95HP/cPBdLZSJYQDrFPj2K5p7eKOWh7TM0rpKXuJM9jizKaL/h1Tj/WF+K1CTr8sV/KUkAYH8T/ABx0p9NrK6Ov1bT6+yumMdNqBINQQoPuQ5nAtbs46bak16W8ttsc0nxR7d19hKj4eX78dXYj+5Df6cOGXEUKwrWP3pdtx4/nBNpl53G+EsOphdrnz546Ya+uLarKa6K1KMhzTKtSzIZosvnNLXxec1FMpSVD7FGkHzGDa5yefKc4q8rliBlygrSPKGv38HWnn+DRFBf2HriOaemQ1jUE1u6rUMBJ6KT9lvkbYm+oaibNNKZFquSJ5ZKWF9L5su7bdoRupyT7w2Av5wHFK72cuHXWPJWS68zrrZzUFXcSWAUBjYjra/Tr+/BiHZVKsLKdt72tf1wLtD3jtGJNo+yHAv8AAkce/wAsJkOFFgeOvpi5VVaI5cE3vtK8gD9w9PXAsQepW9r3Fvj+OCsGN7Fih4vt23PvgVJKhSTb0Hl74VCSkHeMLG5I9cJU8opayKc8qrAnjqPP8sOLXFzxxcemGc4APBvhCKihSg0NV072UapkqtAZdQVf1eQaerZcpJeOMSCCUmeAiU2YdZQAWt4cT9zRinRoJXJiq2YI0RiYpJELuA10I3KASHPXy4OKC7Ac4mOeVOn1lkU55QPHF3UayN9apf1sdlYFSWjDrYjndbF2ZXWU1ZVLSRLBeaIyJJDbL5WAtcGN7xMbG9vDjhbbB2Foe0ba8czxvLqrLL2sTSfDhlyothS1lblVTBnWSwTS1uW1UVfTNTPvaOohdZI98beMeJAL9bEg7hzhvr/s17Is613VZkdRJVZHUVNTU5TSZRlVRNWx000hqFpJu/MNOjwtLJHYGQhVHBFsKZllxoFiqHqoeoZUr4HpHt/cmXdH8w4HtjJcyqZaTfW0TZjQmYK3fuGZHHiAWVDYsPtKQQR5EdMJBaXxNozbnySywNkNXp3kcOgsroni0j2dZcksToYZc5qP0o08dyGtFaKljlHBUGKRWF1DBgLnmzd84rYY9QZ9WVaU5ZYMrrJttGgPlHCirDGVsLFY1ceat1xp5M4pI8xlgn72NpXMakAySsD03KovKCOtl3W58XTDjNGXI3NTq/MMuyWKNSqTZvWLTTkDoIgpaWUWsQDG69LW6YY+V7yWuOfPw2oayOPFoyW+y41NLvnhAYBhtammMUjAdAHgPdyWvxvRvzGFqjPJqDMmqp6KTa6eMNEFLAXsT4EBc3PJuPjiss77eOy7L4J4KeXOtWTVEe10oaf6jASDxeeYNIenlEL2B44xBs6+k5rRoRSaVy7ItJ06CymESV1T8S0zOoPwRfa2CPR1pnd/DoN+HL7v8qH26GMfdXwx/HNdHT0OR5hktRqapIyKip13yVteDSU6N6CZGVWPsFZieinFBdoHbpQ1ENXkmgq2dYp42pa7UcodJHhYWeGkRvGodSVaRrSMpK2jUsDTmptc59q2s+v6nz7Nc/ql+xJmFSzqnsqXNh7Cw9saaNaitlDSm4XhVHAHsB5Y3LDoPsjfmdUjIah8+dBuWXatJdoLrBQevW7itpJUGqmjkipWFJTr3UEAXdZT1vbqxPJP4dMFcOWMoDAgWN2KsPYn1+OMjUxqAmxh6GTZx6gn+h0wnIyCLwMqkEgg2sB6gjrjoWtDBQLIc4uNSjStIy2Vg9rEMRZjxyL9fxxrw8lLVLViHcI2+y4Ivx0NsbjJ8szTPJ/qmT5ZV5jMbfqaWCSdj8O7Ukex64sPLvow9u2Z0X16o7N6/KaRUaY1Od1EOWJsALFz9ZdDYKCeB0GGSyxxj944AbzRLG17j3BVVAaydhtjRI19FQfvPOChqlj4pnsOviOOl8o+hXqYRpUao1xkeXwPEk4OXUVZmngdQynvI40g5VgRaUjnrh/D2CfRy01WJDrHthzGYWPesKzLKDa3l+rSSqmI/wCQHFI6TsQwa8E7BieVVZ+ktJzaR44eq5beCVrck3642aokdIAndlFBF2KixPr5/I46mOY/QP0isU5yes1TNHcGPvcxrdx9Sz/VIufYHGwpvpgfR/0lHbQ30XtPKyOCs1VQ5dG5W37TxVDg3sftHEf9qOcaRQPPiA3/AFEeiX6Sn3vaPOvoCuVsuppMwZDR97UzP9uOKJpW3dOig3Hp7cYl+Udi/azqAA5F2YazzF3sF+rZDVlWBHXd3YW4P4gjzGL4q/8ASYdoVKncaa0PkeVxqLIiVdQAo9lgEK/liGah/wBIX9IbPgQNQUVGBewioe9Iv/eqHl/dgNp0g/7IAP8A2f8AAKOxswzkJ8B8kKO0/wBDr6TFaQydjOeU4ax/11qelI+IlkU/HG8pPoLfSInCit0tkmW35/1vUtCpHvZJGP5YiFV9MH6RFZG0cnafmkate4p4aan/APpcQOI/P9I7tzqVCT9rWsHjUWWMZ3UIoHsEZcIf7VJ/6Y/xH4RSyay48B8q3D9A7tNRiM01Zo2kZVuyxSZhWkX8v9XpXF/nh7D9AzPI0MuY9odKiKBdqfSmdTW/GnXHO2Z9peus5JfNdV5zWMeSajM6mT/20hxqH1BmshJkqXe/7bs37ziVjNIU7z21/wDU/wDJBNlGAB4j/iulq36G+SZXIYK3te2OQLqNH5gCR/zlcK0H0PtJVTePtwWmsbEz6b7tR83qxxjl1q+Z+TFASfWFT+/BWq5W47qD5QIP4YmuWqn3j/Cf+SjrBsPH8Lr+l+hDoirrZaCD6T+lTLDEkgLZfGqSbvuIxq+WW3iHQXHJxvKL/R4ZNXoJo/pI6XpiGt+vy9FB97rVtcY4hFTIpvsi+cSn+GFTmMxFu4pf/wAWj/liF0FvP2zD/D+U4Ps+th4/hd50v+jPzTMneHJfpEaBqZIYxKoFLOS4JIsuyRtx45UeRwtR/wCjD7USRWZf2paMG0WVZsqzNCfYgxHj09PLHBMWZVEZDJFTqw8xCoP5DDgamzpP9lWSx2/4cjp+44h+n0mMp2+bP/7BOv2bWw8fwuys9/0c/bZlkpjm1fol9t3LMc0hQg+e40hUG/PB+XOIHnf0Te0XSRWhzDV/ZXJPPGZIUk1dDSSFCSCUNSsQ23BHPTm1sUTlfanr/J1ByvWWe0ZXp9Xzaqj/APayDG8j+kd24pF3H/axrBov+E+d1Eif+F2YYljGkWu75Y4eBb/uckd9MRgCD4g+wU7/APiXu2GoBfLsi09mkZ6rlur8mqw/tsSqu34Xxqc6+jH265dSfWKvsZ1oI415mhyeSqj/APHBvBHviNxdvev3ZZK/M4a1wd2+ryygqCT7mWnYn8cKy9umePOahsnyVJWtulpcvFC7H1vSPEL+9sWRJPWjmjyKjMcdKhyi+d6E1Fp92XO9PZtlvF1NZl88AHse8QY0YoHY/qpI3uL2VwT+GLwy36VmtKOOBI8+1RRxwW8NHq3M1B/5Z3mT5bbe2Ngn0n4M1lkk1LRHNA3AjzjIMkzZT/zvTQy/+zw7tnDNhTez3hc+9zWQNde8Qjm4uMKDMcwjsHk3geUiBv3jF/J2kdh2e1EY1F2SaO2AXaooKbM8kle/3T9WnniBHr3ZHtjK3IvouZvTyVFMuscok3dcs1Dl+bIoPmIKuKknPwDE4TtonGjhxHynBkjRh6qgTmCOb1FBA/wBX9xtjd6e17n+l5O+0zqbUORSDo2XZlJF/wC1K4tat7Buzqtilk032yrAIyqqNSaWrqFGZugE9L9Zh+ZIGNRJ9GDtKr4pKrSNLkWsaeM2Mml9QUWZP/8ACUkEwPsY74dSOQUB5ptXtOIWxoPpmfSOy2Gnhh7ac5rVpl2xNmlLDWSxj0EsqM4/8WG+p/pifSM1XSSUOb9teoVp5RteLLdlAJB6MYQhb5k4rnVHZzrHRkhg1dpLPMicP3e3M8umpvF6XkUD88aaDLDKboAy+qtu/diIWGz1vXBXbQfCcbQ+lK4eJ+UaSvrK2okkjklMtQSZZpJC8sl+pZjzziQ5cr02XPToR4lBsEV+nW56r8RhpQ5akGyawJj54PPnf297Hy+GHcc3dR94iGVRNuR3AsrdSVK9fUjoR1GLQaG5KIuJRSKe9gyK202Tkkn9xHvhxFAlbUUWXyszRSSKZNxNhGl3e3P7KkYRasPhkZQ4YksqgBL/ALVlFj/hHWxB6YbS5nFTz1xlqGklNKaeAAi7mThiD0AC3Hzw2YkMN3NOjALxVekf0SdPf2X7D8krK/LwtVq+Wq1JWMX5IqJDHCCAdyqIIEI8LLaQ9OuLdeCiyxVjnneio3e/dyEPSAnpwQUjv62T445Y0T9OTsezOny/K9V6LzzRZoqSmy+B6VxmNFFDBCsUYBUJOgCoOCstiSRa+L50pr3Jda0v6T7ONb5dqKmjsXjoplmkiUj/AHkP+0S/PVI/cnHlmlLHamTvnmYQCc8xTUKjDJdnYbRAYmxsdiB5rdaq1JL2e5TmOs8wkjbLshy6rzlCW3JL9Wj3oi3LDmUwqNrff5HIGPIeuz+uGoJM6qqmqTNZJjVzVkMxSYVDsXZww6Hcx6Wx6I/TK1plmn+x6gyD6pTU9TrLN1paloUCl8tottXVW28eKU0qHz8FiTbHndV5a1UzVciDvZ2Mri/RmN7fnbHU/srZLkDpSPu9BgOdeSxdNz3pg3Z+p9leWhvpr9t2m4YMuzbUNHrjLqeRZIqbUcZariYcboqxWWdGtxfvDxwQRxjpzsy+ml2Ta+rqTKtRRZhoTUlTIkMMlXKxpJHY2C/XYQNoJsP18LLzy3njzcmoWjub4WoJpu7qYWkfwQmRAT9llINx6Y0rdoKyWsEubQ7Rgfg8FVs2kZ4TRhw2Zj8cV7PGrq4VePMqaOUU9SYpS22lljlRuQs8W6DvAem9I7+ZF8beety3O6eZMvqHkmoj+tpqqnMdVTnru2g7gDfiWCTab9De2KwyHPYtU5RlOtKakqqWqzvJqDM5KijkN5DJTpuZu6PeL4lYXeOROOo6Y2ZzevgRayn7mcod0crlFUG9+XjDRet7pHfm+PKxWlDmu1ArjqUyp9U0VI5yqXvI5JrQAVDLJDMx/wB2SwsWP7EiLu8rnEX15n2ntA6azDUPaLnn6J0xSRMkqVMfeGpVv/fOCCS/eytyqotgt9zBVUnHKP0nu1Ptq7GtcU1VpHtG1Tl2mNSZaub0tA88FU+XSmZ4qinSWRXLRLMl47NYJIg8scoaz7T9YdoWbjOdTagznP8AMxcLW5xXvVypf9jcdqc+gx1Vh/Zs2sMmc+rDjh6bvKqw7TpbsrzA2jsseseSkfa12g5p2sdpOda5zCiFJJmFQJIqJGLiigWNYqamB6nuoEUH1OIxObbTsI3HeBIpsVv9qx8ieLDg834wzolZI/q5bvJN+57BnZiet7Dnnyvz8sHJklHchNyo3gVASDuNuAfX4AY9BgibDGI25BctI8vcXFOmj76Qs7OpLbiSosDbkkLYgfAWA8sPdMZRV51nMdLk4EtdUTR0NApa2+qnbuoyAR0ALN7ADGuJ2Ew1Et1QhG2bSTboCx+0o/q4xYHZdk0ozWqzmNZY5cspUjpnUi8eYVp7mA835SIyy/8AIDh0jwxpcdSRjbzg1G7eM2paOKk01k1RG+XQrFBQOhuHoqRWpadwSL2kZauf374H0xVlDvp03mMiwspPT8R+7G57RtQ0+qdaVldl0ZTL4NlPQxC36uliRYoBbp/s0Qn3Y408UYjKEpsVz1BC8A2tfoLnoT52xDZWXI6nM4qSZ151BqT5opBFG5UkEnZwOeOfkQbE+4w0zBtsPcgjdKViAAsVW9+fle9+pJOH1O8UM5iimmZbFSzxqrHz2lGuAB++56Y1OYVBlnklLHZBGTGCAOW4HA498WVCuifoT6KpNU9q2m5MygV8uizZ9QZgr/YOX5TEakq3oHn7hPiRj0CmqK6pzOCmq5++kkEs9QHVF7ydm3uu82s13LKCQDc9TbHLP0G9GTUeT6s1K2XRyijoMs0tFvkEZMlVIa+sKFhtJCQ06EHjxAGwOOlqOKHMs4TIwDJBWv8AUgZVQoY5QiS7hcgJYuwXkWT4HHmn7RWjtbQ1mrF3+LL/AChq7DRUYjjc7y4Z8yVyb9PjW0suf6e7OKeVohpvJzUVcZ4ZMxzV97AjqGSjiiWx5Ab3xxutT30jyG/jYm3ltHA4+AGLS+kHr2LX2uc/1bBEqx5zmtdmcBBvtpi/1WijB9FghBHxxVFOpA2KPYCwJPtfrjtdEQdjZ2t2CnDPnj5rnbdJ2khO3HithGSwGwttXkgAkD+Hz98GSzWWJW7y912i9vU/G3UEdOb4TiQoymdCAOoIHA+J4HzwttQC0gJuOeb/ACHwPQ/EHGqqSGcxrdEkD7ebJyo9bE+/lyPTCcb94yxMoYnqEHLD2v1PtgXs0febrgWuGuD7Wt0t1/oYIwF+V4J8VlAP58W+FsCEdnLbklaRrDcbtbj59bcdbYRaJXBK7iOAtlIBPpb53wYsN43G7BvByD08ueo/rpghmXf1AJuQTxY36i3TAhAbcK0gIvsuzAAe2M8DHYBJc8heh55II8/yGB2h41axHBABW59Tbb+/BWUm6lQ29bkkixv5+pOFQsblbi1lJBa9x+Hr/LBX2sLg3IABt5f11GFGHIt9q3WxuPnbnBNoDne2255upI/6e3HthEJGSzkh2IPHW5JHl1xmDL4hwRcm568e3v8AmcZgQkZCLHgkfGx/6YEKGBvb/PBHNgbk3Hh68/15H5YGKxjIYEnyNxhUIFYXK8X9rH+jf5YWYneR9tQb7gCD8fUYQcchiQR1HTp8sGElwvg+zc2PNx5/LAhLNKTuO0BjaxA2kW+HU/vwQMCQ5PJ8jf8AG4OEyvkV6jz6H0wO42JtzwBe1wfwucIhCE4uoPHNvT4nBOFG0WPmvHPwOMZwb2tbyA8h6fDBQQQdxsCPXk4VCJIrcWB9ALW5w2f43w5e3Jtz5+uEHAPJ4ueMCEkeuMwJtgMCFOuyrUI01rXIc7eTZHRZnSVEpPTuS/dzD4FHOOoYaE6f1G1OJ1K5NmZpDUndsH1eVkWQgAtt2gAkcgc2NscY0Dn/AGY4MivF8yOPzx2BS5nPmGV0moiyE55l9HmykG4DyRLHMrehMkLtz6n1xy2m4yJGuGRqPj3W5o14LSDqx65J59IfJDmvY7ldQ0WyfIv0vp91RhIlqOrXMaWz8bg1HmVSFNvsxY4wja3hIubC5v6cHHd9fDS6h7NNV5fUSxuKCq09nEkfeeOnWV58nq969VPd1tP167AccLVtFPlWZVGW1HhmpJpKeT2ZTY/mMaOiZL0YacyAfMd0+g4qrpFlH3hlUj3HqUhJ9kAi9uPjhB+p5vzzhzIACNzc3GGxXy5v0PxxsLNQKxRgwNiDcYsvStQc7fNNJliYtV0InpUvZUzOnu8R+LWkT4TYreOCWS5RCQvU+Q+Jxu8lrKulp0rqGXbV5TUJX0zDyKML2/8AYH5HFa0NrQjPqnOimiOY66ombgFFAW1hu59D6/uwUXWZVkBsLAgi3lx/XpiZdpuVZbQ6kbNcs2rl+ewRZvQxqfsw1C7yv/JL3sZH90YhxCixUWAHA879cTRvEjQ4KJ7S11CjX23W5F7bj14+WAVmQ/a6XFj0IxhZuQejjeov6/D92CkqC0bXsputjcD+Yw9NWXIi2MCfO58v688IzINvKm9rG/l74WvYbRbw88jzvjHjJBZfECN1wbn349cCE/0ZqCq03nNLnNGSJ8sqYcwisbXMTAsPmtxjqrM0y+kzGo7uOKeh3CWjZxtvDIokjs68m6OtrhuMce0kv1asjlYeEN4gfNTwR+GLLk7btdZVk1JpzK84oMqGXUy0H1+gplavqIo7iO85uyWU7R3ZTgAG9sYOlLDJaJmvipXXq2bAcvdatitTYoyH+XXWSviqno8mo2zDUGbw6apJFDRpX1hiaU/tKm0SOPdYjxa3POIDm3brpnJlqqXStPmuf/Wdq1Erj9HUMu03XcBumkAPQlozfyxQdXm89VVSVszSVNTKdz1FXIZZHPqSevzvhrNNV1ZDTzM/mLngfAeWGQ6CbWszq7sh8/5h4J0uk3HBmHXDkrNznt97Q8wEkdJqOLT8EgCtDkNOKeRh/enH61vnIfhiuqzNBVVD1LxvNPIdzz1Mhlkc+pJ6n8cN1hA5Zhb1JxL8j7KdaZ/Q/pakyCWlyz72Y5i6UVGv/rpyqH4Ak+2NSKy2axirQG8v186qg6aWY4481D2lqJuHkbb6dB+AxiwX6m1vPyxaOU9mGk4rvm2sps5MV+8h03QtLEoHk1ZUd3CvuVD/ADxtKfUPZnpCqjkyfTWRRSU/++zC+f1bH1CNso1PxVhgdbGN7sYJO4deiBA92LjTxUB0X2a6519VGl0To3Os+lB2sMuoZJ1T/Eyjao9yRiyqb6Oea5LOkXaL2h6J0hK7bWoJcz/S2ZcDp9Ty4TOregcp8saHVHbhnOd0y5dPm+e5nRrGE+r5hmLJSX5uVo6bZCo5+yd3TEMl7QtRLTtR5fVjL6d7hoqFFpkI9D3YBPzJwgktMmTQPHr2CLkTc3VV9UfZl9HrIU7/AFDmuttSFGDbqqah0rRvxzZZWqaxx7rEDgD2n9gWiC65D2V6YNQHvG6ZZUZxIq+8+ZyLEWP7S09h5DHNQzGuaRmWdlZySxXwk/Ejk4CXvSNzXN+bnzwx9klmwlkNN2HFPE0ceMbcd+K6frPp36woaSSh0tkkNFA0axJFLUlYkAPVYKNKeIE9CDccYrbUH0qu2fPJJWi1WcsEoIYZZRwUrWPX9YqmQ/NzioLYUSnkfy64bHoewxOvdmCdpxPOqV9vtDhdvUG7D0W3zzWWp9RsJM+z7MMyYC16yrknt8A7ED5DGn7+UcK5W/7PH7sOaPLqqvnWkoqaaomc2EcMbSMfgFBOJdnHZHrzTmSQ5/qPQuocqy+ocRx1ldlc8EDsRwod1C3PkL8+WNANYzugKqXF2JUGu79STg3dvb7JxbvZF9HTUva1QZrnmT6l0vlGXZJKkVdNnFe0PdBk3B9qox2W+8SBfjrjO1fso0j2dZZlE2TdsuldZ5hWVE0NbRZIJD9RCqpRy7/bDEsOALFfO+G9qy/criluOu3qYKoWRhYkdcAFYttCm/piXaMo8jrtc6dyvUNO82WVmaU0FZGkpiZoWkVWAccqbHr5YvrPexzs70d9OOm7KY9NR1Wk5szpYIMvraiWZWhqKNXS8gZXYh2uOfK3TjA+VrDQ7K8ErYy7EbaLmOky95wSeLdb4CroTTgEsouLi7Dn4euO2NU9hGRdkP0vuz+qg0rQzaD1hmAkhy6pphUUkMwjK1FIyybgQrMkibr+CVOTtJwn2GVOQ6Z+mlr/ALJ8+pKGPItX1ObZN9XFNF3Ue8u8QRSto7RvIq7bWuPQYhNrZdvNxwveSf2Dq0O2i4poKKprpu4paWaZ7XKxRs7W+ABOJBp7QOqtTwNUaf0tnWZxo5jaShy6aoUPb7JMakA2I498dLfRQ0vmnZTqHt1zzMZZaOv0DpTMstmlRivdz948fl5kwXHxBxvvozaw7RtO/Q07SX7Ls1zml1Jl+o6WpoxlgaWch0pUlVYgGDFl3Enbfwk34wPtN0kAZUHFAhqASdp4LkrVWgtW6KNKNWaVzfJfrqO9OMxoZaUzKpszIJFBYAkAkdMbwdhXa+mjj2hy9l+qF0wsXfnNjlUophD/AMXdb/Z/37bfe2JJrbtE7Ws+7QdKS/SRrtQ1keU1EM8cGfUxSVaF6hXlIQorFG2E2tY2NsdG9vfad27dkfbBTdvWj82bUXZ9mtFBHQR980uVAPTbO4nSMjhm3OhNlkU7QeGRVfM9t0ACprr5DehsTSCScBuXJmk+xftP7Rcrnzjs/wCznUOoKOlmFNPPl1C88cUxUNsZl6HaQbehGE9Vdhna52f5XHneu+zTUeQ5dNOtLHV5hQPDE0zAsqBmFixCkgexxLewDXmu6Pte05kWT55mGSZRqHVFNPVZVQTvDRyB5QCvdg2KhfAAb+EAHEu+mFq/Wdb20650TVamzGfTeVahvR5ZJOTT07LEgXu06LYO1gOOThe1f2vZ0FKVSXG9nerroqT1NobU2h8xjyjW+l80yGueBahKbMqR6aVomJ2yBXAJU2PPTg42E/ZVrmHTI1nPofUEOn2UOM1kyudaPaejd8V27TcWa9uRzjrb6UGQ0Orfps9mWnM1pkq6OuoshiqIZOVlj72VmVgeoO2xHmLjDfPPpj6l7OvpWary3tDzzNcz7PqaSpyap0/BFFNE1MIQiqkbkL4mLbrmxBseAAI22p7mgtbUkV/ATjAGk1Oui4iqMuKsFXg3t088EqctqKN2hqoZIZUHijkQqw+RAOLz+j3oLKO0zt+0dpuGAnJqnPBWTowuVy6B2nIbnj9XGqnnqcWV9PnOYe0in7Mu2GOLuxqbKs3pXvJvINNmcwTxHk+CRbX6DgcDEzpw2Rse38/BTGxksL9i48EBYgKwG7pcgYcfo2Q3BBBHtjr+HO9E/Rf7EOzXNYOy/S+ps+7Q6eTNc2qc+oo6lhSgRkRR94rCNbSBQAALqzNuJG3Qao7IOzHtK+lNkOgex/UNBHpfVcNNmFSMpqFqY8rujvVQREk7dvdkqrE7O8A5AAwwWlpqSCAK4+GaUwkZHHDDxVXZv9E76QOTSUyVXZhmrLPST1byQPFNHSrDD30yVMiMUppUjszRSlXAZeORipmp5FXeVNrA3x3hTaR+j52xHtK7F+x/JtX6QznJKWprBmLagqamizt6RypNXC7lXDOCLlQw3bh0KYpTsR+jrl3a7oHUuuc+7Sct0ZlWnamlpnqsyoXnpnaaMsS8iODGFJRfstcuOmEbamhpMmFKc8tuaUxOqA3Gq53ViOl8KioqFFu8e3oTcfni0e1XsNTs4zLJKLIe0XSmvP7QmRaP+y9S9S4dZFQI6Fdwd2YbVFybHG11r9Ezt57P9IPrbVfZxV02U08YlqqiGqp6lqRSQN08cTs8QuQCWFgTYkHE3axupiMcky64V3KoKPOsxoHElHUyQODcNC7REH4oRjc0mv8AOadO7njpK1NwbbW0kVTY+zSKXX/lYY0jUgDEEEYTNK9/Dbjrz0wromOzCQPcNatzIPpMa3yFWTLK7M8rVmEmzK85qoYSw6FqaZ5oG44sUtbEvoe37RGqFZe0PROis5mKljPmmkUpahm96vKJKeT5tE3zxziIjfgXwVlYEjEX04H2OI86+tU8S/1AHrcuo4ss7BNZWXKsp1JkU8gUK2mdS0meRBupYUWYLT1gA67RK3xONXnnYTR1FRu072zaZq5pywjo9VU1Tpirbb91TVp9VJ9As9r45wDMpBsOPUXxvcq15qzJIvq+VZ7W00BFjAkzGFh6GM3Qj224Ls7ciD44c8fRFY3ZgjrrWrL1f2HdrGhKX9I6w7Ps5y7LWQFMzip/rlA56i1VTGSE3FrHdYD164qrMaKSeQzxOkq2tujIYfiMTjSXbrqrR07V2n66sySsmUrPUZBWy5S8oJ6utOwhkNuPFFb1vifR9s+ltfMG1/kOlc8q5WHeVWcZOMtrb26/pHK9hYj9qaFvfC9sWjvtI5+n4QIg7BpHoue0nqaXwK7L7dQflh1l2b1NBWR11JJNSVURDR1FJK0MsZ9QV6H4WxdNX2c9jmrS02Uaiz/SMwZgTXxJqDLAb8f63QhaiNfPx07cdTjQZt9HLtDp8ulzvS9Hl+tMohF3zDSlfHmkcYte8sUf+sQ/+tiW2Fa+OUd0oIkZgUwzztR1f2iZbRZdrnWOa5/PlQNLlzVz95JDTSv3k67/ALTEsq8sSbcdOMamqiVXMLLIrOpO1mUX/u7eov5En04HlrMpp6nLu+mkgkUgmOQkHwnzUkfYPsbHC7zDfeFE2MOoAsffjjn3vh0ULIQQzJMfIZDUpJoCwYPYAEi5HPTi/v8A0cMoYVSvWO3Ekbofe6nG2hjCd5aGNg5sm8bgrH9nnqfe4445xqqiRo8xp9ockSLwRz1th0mLSEMwcF6O/RzzddR9gPZ7JVbXWPLKmgKlO8Kmlq5F3bRZ1sjITJCwdBYlWS5WyqlqmGaEQVKyVIu6LM7NJInkQ8ZErL0s6mb3Ddcc9/RF15oWj7OV7L841vkGX6lybO60w5bmVR9X+uQy906GCZwIWkWRXGxmRwbFTzi9aiaOrY0ddlVZEe8YtDPTtckH7RiYFj6iRRJ5+LnHkukrN2Fqkwo0kkbwTX3Xd2ObtIW44gAHxGCon6beXjUPZzpvOFg7tstqM4y+QBo3B3xUtXGFZAAwvFUfdU9bqDfHEUFJFAxIIdQbEgXHtf0+Bt8cd/dvdAlf2PZwkcnfxUOocpliIk3939ZhraJxckkD9bH4TyD7WxwokdQBBUp4JXQMsiNaRePtAjyv5Hr6Y7r9mH1sQaNXyfwuY0w27aCesh8ohacyGNhMCp5S73X/AJDa3t1+GGrGQv3kjMFkVkWwv0+7ybgev7iMO60NA3dl0KMockXuG8yrHnaTzY+9rWw0ba20cvI7bW9x5D1Jv+GOjWQndJFDUOBUzxwxqbyFuuwAklVHJ4FvQE82GLcjqU0h2RNJJXU0eYZhE2c1FOGAlE9arQUoA67YqNJ34+yZ19sVppPJH1BnlNplHFP9fmSnqJIywMcAO+YvcmwWNGb5c4lmo4su1trGk0/mefyZFl9fE+YyVMdBJWClSRAtJG8UX6woIY4QxQMVDlgrdMU7X3gI60rn4DP481Ys/dJfTJVVTAzytMy+JiX6Xtf+A/hh/tUKpcqgFwFZLhweosOSD68WsMS3XHY/rbs6yukz3NKGnr9P1zbaDUWUVC1mVVjD7qVCcJJxzFJskHmoxEbF3eIKSwABCEksbc38zY8WxbGWCgOeKPAXCl5AgRrIjNNYlfNGNrm3ToPLrhHJaQ57n1BQMovX16K1hxsBA/C18BK4hSUSMxYb7La1r8nr0+GHOj98ObSZhH9rKsumqE/9KUIT/wBk4xDaXFkLi3OmHjq5qWztDpWg5VXpF9HjKY8v7F9OTNTNG2oq7N9QiVJTu/X1P1Wnulirp3FGDa6tY+Em9sbvU2cV2jOzXXmqaiopHzTIdN1hQwkiSKrqB9Uo7qbWLNUs4IBU24IIIxJMu0dPorLMi0Q9HFLDpzLMtyRe5nBffFTRLIO7tct37zNY35NxYm+K7+lVmdPlHYzFltO0Xe5/qXLqKZYybLBQQzZjOvPGy7QcDwk88G+PM3tFo0mS3EB3Jn4C68Ex2IVzI5u/VedWtu7p8zfL4H3R0si0kf8A6OnQRj8X3nGliJ4G3z4AP5c8HCmczPUZgDIx3CJSx6+N7u35tgkMd1IBAuRcDz/r8bY9KsrOzhaCuRncHSEhOUCKhBANiCNoB6+/kfY/ywraNrDfsA8z0HvYeXzvhNCuzaHjBPUkWv8AG/Hz/HAc72UP1awYi3n6+V+PyxYUKwOwjILkHyQg8X6+38Txg6k7NqDgeK9jz5C/qB0HxOCyhFYhnJC8Dfbj4W8uthg4iaPePGzABjaxsfO9/bnjkeeFQg2S23ADu5AOL2BsfO3v59MJuwIO07VN+nQfD+r4Vdrje7+fiYoT4fK9uny6YRZkBbw2J6EMSLX6/wAMCEKksoZmHhNm3Cxv7WuT6/vwBkN2UlbGxU23Et5kemA3KjKCSCRwDwQfl5E4LcMPEbeY9j738sIhHc3IPJDAc7No48gP43wSZTcMeAb8W/f6HGFgo3XWwa5axPXyI6H9+AAudxUXPUDgW9Ovl5YEJF1G4gX2gcE8flfjGYPIl0UuR6ix5+PtjMCEg17G625JFxgVJB5H4/y9DgzXLhiDY8Wvbj44ADkbRceRAwqEHkBYAk+lvzxhA4Ibwjrt8j7fDAXJBZeR526W/rz9cKFGZtq3uGtx19vn7YEInBuOlzbp64wgW6m9+hHr5j+WDuuwbWBJJuQbWB6AmxtjCS21VHW/B9f4cDCISJDEFhcgcbr/AC+WMUlSbNY/DqfT54UbobXPPkPLCduLkC1+fF5YVCIyk9Obc/DCD9SB/wBcLt948m3tax/hhGQC4NuDyLemBCRIF+uAHt1wZsFwIS9I5Ryw+6Q/4HHU/ZxNFW9mWlqmYAfVWzLJmlv9ruKgTRqRbxWSpPuB5EXGOVqaxl2/tAr+WOhOyjPqWDsfzKnrHU/VNUQyRhuCv1igkBZT5HdTofe3rzjD00wmNpGojmCPUrT0a4Nea7PyrcyGshqv7SacjLM2e6Fz7L4U6uZqSFcypkv1IWWhuoJO3cQDtIxxzrjMaXNtXZpnVBf6vX1TVacWt3niI/EnFyaX1/qg9pWRtoLIK3UNfBWNLFlVJA00lWjxSQzKAgLAPFI6kgbehNrYZZR2UaE0Y0g7U85/TWbUXMmmtO1cRFOR5Zhmh3U9KLixSETS348Bw6wMFmjD5TSgPM/gJlrk7Z5YwZ05KqtK6P1TrnN00/pLIK3Nq+UFhBSQl2VB1dz0RAOSzEKPM4mtR2eaO0Yrx6nzqLUmdoAXyzJqlf0fSH0qsw+w59Up9wP/ABAcONXdr0clBJpXTWX5fR5KGv8AofJ45KbK7jo0zMe/rnFv9pO5HooFhiu82zrM88mFXnNbJVSAKiR/ZijRAFVVUWAAFgAAAAMXg+Wc90UbtOfD9PNVS1kY7xqd3z+qeZ7nlBUll7mnlZDtggpI+5o4B/dX7Uh/vMefMnGoyvMDTZhHPMd0ZvHIvkY2FmFvgcNpVJYk26+WEfPE4hbdLTjVM7Q1BGpWbUU75t2ZGBkEtbobMe7JPV8tqzdSfVVnX5d/iG1c0UojaKUlpR41YWKsPXixBvwR6DE27NM1opc5y2izSVY8v1JST6YzJiOEMqgQzH/C/ct/6vEKrKOtyusqcuzCFo62kmaGRS1ijoxVx7+JSPliCzHEtOfVedT5hSzimIy6pyTdo1UmPcGuAeF6H0v/AEMYkRaN2Ia3RLLcM3pfysDfBGYs1txPnycGjVZJEDqLcBtvXnzPvi3iq6N5KXQqWFgbcG3HB/hgndXXcI+nUgm4/lhWmWV2Aj6KbsGUsgHqw9LdfbDmrSnikmRY0UJJZY+832JAuAw689D+/CVxoimFVqpYCzXPU+frgkdMWYXHyxMqHQ1cKZcy1JXUunKCRdyS5kW7+ZfWKnUGV7+R2qv97EqodP5JldGub5XpmOSmC3XOdXuIadyPvQUMZJk9gxlHrivLa44hUnrx9hjuUjIXvyVfZLozUWow75Fk1RVRRXMs4AWCH/HK1o0+bDEgyzR2n6SXZmmdT5vVJyaHT8Qm2+z1Lju0/wCVXwrn/aHT1m2HM62o1F3HEMdV+ooILf8ACpIiqgfEj/DiJ5xq7Os4j+rTVZjpR9mmhVYYQPaNAF/EHEV+0THui6OvP/SVJcijHeNT11rVk0mrNL6UjSoyNMjyKYKrL9QpP0vmi/4qmotDC3/owCMRvPO1aozCu/SCUUmYVqiy1+fVDZnUrz1USfqk+AQ4r4knzwZVLmw64c2wx1vPxPXHzqk+ocBRuHXWS2Oc6lzvPpTLmuYz1Fzfa7+EH2UWVfkBjWksepJw6ShJBv5C+JDonsy1n2h6nptH6TyV6vNapO+SGSRIB3XHjLSFQFswN79DxfFpoZE2gwAURLnnHEqK2LcYWpqfedziw6YtHsC7FD2rdo02kdQZhNlFJlcEtTmJVAZwsb7DGoIIDbjYkg7QCbMbAzLWuU/RSzns3zfNezvPc/07qTJXQU1Fm8pqBnaMbAooUmO4BbeCAtrOgDBgx0zQ+5QlOEZLbyadmPZd2L6jpNMU2Y5lr3UOpa++Y5jk+lslhrYYqdahkWjkcypJFNIiXMliq96vBIxS2eT0z5xXx0GVzZfSmqmNPSTSGSSmj3nbEzEAsVFlJIBJBxfP0O66sizXtAloHkFTT6RqKyAI5XdJTyrMoNvUpb538sZ9LzQ2VxZ/lHbPpaILkfaRRfpS8Y8MOZbEapTjpv7yOYf+kcfdxDHIRM6N3WvrwT3M/dh4TzT/ANGfstyJckybtm7X6rT2ptSQxTUmXUGWR1C0olNk715HXcb8EDatwQGNr4rLti7Ks37Gde5joDOqqmrJaHu5oK2mDCGsppUDwzoGAZQykXVhdWDKeVxbP0uRT59o3sk7VcqVg+aafFDUSoeRPEI5kufUGWQf8h9MKfS/LavyPsr7XIEDDUOnRQ1DXt+uh2zAE+R21LD/AJThIJXuLb5+6vkQnSxtAN3VTgVo/oJ6iXT/ANIShoKirmjp88oamhmVJGjL2VZgLqQesXkcXN2K9qev9W9uWufo+9rmpH1FlOYxZnkzxVcMYDGKbZwFUXHdkyC99rRowsRfHJXZZnq6M7VdKanMoSOjzamec36RNIEk5/wM2Op800RqvJvp8U+ptO6frp6Cvmpc4nnhp37lYqijKTsz22j9aHAF7luBc8YhtMbTI5xzu1B2EfqpIHm41o25eKhP0TcmkzKTte7Fc5qo4ZM4yOejYy/ZSeF3iDn/AAv3bEgXFr4qvtK7Gsp7MDlkVP2oaW1bWVr1EVXDkkxk+pGPaV3lrEh9xsdoF1I88Xn2V1un8s/0geoqTKpEq8pzLN86pWEXiDh0aSRFtwfGrqPK4xF9YfQ3fs7yrPtT5124dm1K9DT1VVQZSldI9VVqgLxwi6IqysLAA38WHskDJyXOpeANNur2TXMvRCgrSorzXOMkkmWZpTV6Fi1LKkwPPVWBHX4Y64+ldM+QfSd7Me0mlPd/pTLchzMOP24agxtz/hCY5NrpYp6UuqkM46n3GLj+kL24aP7S8j7KTkZzFs40jkyZfm5nphGjOrROBG2477Mri5A6jFiRpMjSBtChY6jCPArrjMc7o9Yds2v+xjMSseY6U1hSa70k8z/7k91JUU6nqEZJJVt0AmU2snHKX0mMxquzL6Yue6oy92DUOd0WcwleCY3iilt81JHzxp+0D6SZznt/y3t10ZkdXl8tFHSxy01ZOrmpESGKRXKi22SKykc259saLtp7Tl7btaVXaHUabiyWappKWlamjqWnBMEfdhy7AEllAB9x74q2ezva8FwwLcdxwqpZZWkd04g8l2h9IiipdH9mfb/r/LECQ9o8OmJ4WRbBzUrtnsfO7xysf8WKW+iVn+rso+jz23ZlojM63Ls8ymGjzCgqKInv45PErbLA8lAwPHQ4qjV/0me0rX/ZdlvY5nMOUNk2WLSBJoqeQVkgpVZYw7mQqeG5AUXIBxGey7t27VOxYZlB2cZ+Mq/SjxPU3o4ZmYxhghBkVitg7dLdecDbLJ2RY6hNQeFPhO7dgkDm5UPOvyn3aKO2PtBzjKc57U6jPZq+taLKaKuz6nemQrv8MfeuirtUyXPWwNzxjo7sG052y9i3bTQfR31flCav0ZqiBBW0qU8lTlpppk3SyxtIoCojgpKpABYEgb9jjnztG+kF2y9t2RUmU9pus584oMsqHq6SJ6eKNY5mTYzAooPKgCx4w9yn6Un0ict0+ukMu7VM5hy5IhAnKGaOO1rLOV7wWHAIa4tweMTSRSPjDQB8bCMFCyRrX3iSpTp7T+n9N/TJyfSenql5ssybXooKR2ff+ojqiEG7zIA2k+e2/rjWfS6LD6TfaCm0eLUrKCOttkI/Dpiq6TP9RaOz/LdU6czJ6fNMrqFrKaqCq7RTqbhrOCCb8+IHzw41prDVWtM7rNZaqzT69n+ZVIrayqaJFaWchfFtVQo4VeAAOOnOJhGRIH7qJt8Xbu+q7N7ZKlW/0g/ZeruNsVLlCg+957D8cL9hVPqSu+l12n6QzvRdFW5DmOZ1maV0mZZIs5i7lytMUklQqqSd90H2xtPlji7V/aV2m641RSdoGqdS1lbntIkK01eqpDLEImLRlTGqgFSSQRzi0k+nL9JepmyaWo1nDKuUVH1hY2oIhFVTd20YM6AAPYOxA4AY7rXAOKbrLIGBraHu0Vhs7L141zqp99D3LEy7/tQ7Ta3PMv09BQZZUaey7OK9bUtDWVsjXla1rBEWMWHP6xQMD9IrRVLkf0Q+zely3V+T6qg0nqDMct/S2UyF6d46lJJkUE8hgUsw8jigYu13WS9lNb2OU1PlUOQ1uaDN6uUUlqmadSCA8u7lAQvht90YVj7bc3HYdVdhFTkNDLQS50M6jre8kWeKUfaQD7BUi46A84kMEhl7TeOFKe5TBKy5c3c6q5/poZXV5hnvZD2bZNTPPWUuj6Gmp6aO12mnkCIgvYXYoBzg/wBCDQ2oNB/SZzei1tktVlGaaOyLMp66jq4tktPJsVbMPI7XJ+BxYXZT245frHsgftN1j2Uab1XrjslolqKGWWZ6atfLFKxGvsyd3JHHK6I4Qs6uxkRVNyKT7M/pFT6J7bc87Te0XLGzWn1nDV0uexUY2MiTkf7IEjhQoULuB2+d8RtEhgdFTIHzNdSeSztQ+uamH0Oq80NB249pMx8VBo2slWTz7yYzPYfguN/2HaY05mn0Kc6yDVnaFlWiqbV+pzCua5nE8kIanEO2KyEG7d03N7CxxCM+7Xuwbsy7I9XdnPYbPqbOK7Xxjp66tziFYVoKNCP1S2Ve8baCoNvvFieAuJDkuT9mna59GjQ3ZVQduejtL53lFZPnFZTZy80YknmMoMe4LYEBwb8g36i2Gygu75qAXDVkAM+KWMgd3MgHmfhVf2ejSfYX9JnIa3NtT5bqDItO54qy51lyFqWZDHYVMd/EURpA3r4DbnFodr+su1/6NfbHq3Uf+r6j0x2mUlXHl5rppZqCelmABRQjBQ8Ub7LfZMcisLq4OI12cfRp01q3W2vex3NNc0tfqjTeTtVacfI6yL9H5tV92H2GWVb7VLxhgoH+8JIC3xPs60T2g6N+hDn+nu3vKZcsqssz2lj0fTZi8bVSNvQ7UsxIQK1SAL8Rl/uhcPkLDK01rkCNtcQR16JGB3ZltKZke4K4+FE1PTmV28KKCbjkgY7J13mHYf2GdkXZNpTtE7Asj1jmWf6Y/SOYVqsKOupmeTeD30YDyEmUjxNwEsMcp5TQvqrUOR6PpCVOcZnTUXhBv+slVOn/ADc46q+lZ219l9RrvU3ZVqLsWyzPK7IsvGV5Ln/1ySKehneEMiEKQDHE7mwB5I5BucS2glz2sFdZNDT42qKIANc481yVryu0XmutM1zLs807U5HpuomDZdl9ZUmaWCLaBZpCTcltx6nrbFh9kf0Ve0ntmyKo1Tp1MmyzJKef6qMxzvMVo4Z5xa8UPhZpGFwCQNoJAvfjFaVdJHT0zSrYiEWN+Cbfzx0X9KsS6B7EOxzsXpW2tFkkmf5pCBa80p2qWHmQ7z/PnEsjnNLWMzO3co2NDgXOyCoPtO7L9W9leqarRetMnOX5nTKkgUSLLHNE4uksUikrJGwvZh6EGxBAh8dG8jHaVIAubMCRjrXsyzvLfpg9tOgMn1RpH6plGitPumaJJWtOtfDTkMu9tqsoZ2QML8gvY84nen9baX+kllfa3pfPOzTSeXaW0ZkdVXaezPLsuSnq8veJnEBEkYA8YiLFRZCt1KnhhEbUWG69uIpXdU0CeIQcWnwXBckBVrYIdyG6kgjzGOjuyn6L8faT2Hao7Xsz1DWZPJlX1yTL41pEmgqUpKfvJg92V1u7KispNubg2xV2h+x7tG7T8szXNNB6JzTPYclWNq76jGJHi7wMVAS+5zZWNkDEAXItidszHEiuWBTCxwAJGahVLmFZSTrUU9RLFKvSRHKuP+Yc43tB2gZ/R18GZvUtJWU7Bo6xXaGrQ+1REVkHzJxopaUoNosW9Aefwwh3TknjCmJjzUjFAe5uAK6CoO37I9XxQ0vadRU2fzIuw1WfU5kqbD9nM6MR1i/CZZ145GNwOx3QPaTHLmHZ1qSpyeWnTfJHmu3MMt56Kcxo4w8J9PrVMnu5xzGAQeeMbHKdRZ1kNYmY5JmlXQVcf2J6WZopF/5lINvbpiB8UrB+5d5HEfPqpWyRuP71vmMD8eisbXnZXrzs/p1rNR6YqI8vla0OY07JV5dOP+7q4S0JvboWDdLgYg9XE88Yq5EYbgGVy7Nf4E+nxxZGivpQa105VNPmYNS8vE1XQzGhq5B/3hQGGf3E0Tg+eJzmWsewvtR+rxT6ONLmNWrGetyNKbJ8wWX1eiZv0fW39Yvq0hPyxEy1yMNy0su7xi3jmPMBPdZ2OF6B9dxwPDI+RXN1bmlRVSf66BUEffceM/Fup+d8WV2Z/ST7WuzGGPLtM6xkqcpRgf0LnKCsoen3Ukv3fxjKH3wtm/YDm2YVEy9nWdU2rJUBd8sSF6HOoha/iy6ciR//AFDTD3xVNXlldQ1ctFWUs0E8DmOWKVCjxsOCrKQCp9iMTPs8FpjuOaC3ZSoUTZpIn3gSDzXXtT9K3RnaHojU+ndV6Ym05qPM8vpjRSxu1TR1NTT11PURoJG/WoCI5QocyAFrbwDbHOT065XWZglOsyCKpnjvG7KpVXI5sfIexxEqSSujdYIHJG77Lci/zxKMxqUnWaaXa1ZMxkk7u+1nP2tg8uerHgcgXPSCx2BlhcWwijTv1+fmpLRaTaW1eauTXdK3+sxK8rmT/anxKSObdOenN78eQw2Y7BYhXQE3UgWt16jkWvwcDGNuyQNeRBfvFuCfgR0t7fvwrGgeYU279XuJYkA7UAux6cAAE9ecaKqKbdmdFHRw5hmeZ0zNHUxLk8fdyd26pUjfUvusbFaVHHT/AHoxA86zhs81PXagpDLRmWqMlN3TWMEYNo1XofCgUCx8sWVqhBp/QNFSNSzQ1j0zSSuzWC1FaBIykebJSJEPbvPfFW0VO5S6xs1+WAW/y+AxUiIlkc/Zh8+ynfVjA3zV19mvb9n+i3qJ9SxVE1BmMYpa/M8tghnSsS1hHmVDMPq1evlaZUl81lBAOJ1m/ZR2H9q2T0+f9n+pcj0jnVZ3rNTpNPLp2Z0I3CQyBqrJJDcWWfvabkbZVHTnCmqqnLWSqhnnglZrEhRd4/NbNdWFvJ1t5Y29JmNPRZumd6SzDMNM5rAAUq6LciMSLkNGpLID57CyH9kA2AbO6LGA03HL8eWG5KJmvwlFd+v8+fFF7S+zftA7PM2GQa2yDMspq6pEalE7d5DVxEi0lPOpaOoiseHjZhiQ/Rv0tBqrtP09kVXsNPnWp8oy6Xdwv1cVIlnJ9AIomvib6C+kNPp/Lv7I9peQ5fU6fzKXfNA9CtfkVa/nK9ECv1ebz7+ieGUHkq3QylOzPRMmb5ZrD6P+qKDJK6sknNDkecZus2X1dRLA8JTL82YIFmCyHbS1ywzX6NJcEwTzGRnZOF12quRpjgcjl401KWKIMdfaajmK4Yjob11LL2p6Y7U/0pWUkiVEea1NQ81PudHRXqndFLKQQGHdWYHkG1xYjFI/S5z+aXTelqATbjPR6gzMWNwGqKqDLISPfZBP6eeK30hXar7NNV/oHUWV5lkeb5ZKi1GWZpStFNGtxcSRvbdGw8x4WB3K17E59JjVjZvFpmpeNImTR+SSGNLhYzOa3MGHPvMnXn1xykGjnWa1E9YmhW5Ja2ywNA6oFzFVyJPmNRKpsjyuV/wg2H5AYPH4rruVibG1+D7/ANcjCEMfgUkcgDn3w7XfwzEndwbi9h+PT5Y7trboDQuYJvGqUQb0IYMqluhS5VrWuSPy8sZbdchZBcdFA5Ppf+PxwdVtZLEG9mUX8S/1xwb4OqBgF2yDkki4vb2PUnjpbDkiRQSBt8ShWNhwTYAg8c+RwojwvYxRtsBAAJsQTz+P9eeCkkHwHcRySFJB9725+f54BnkAO4k2bgvc8bbefx5+QwIQmy28J4PiPAHJ4sfL1N8F7sWBP21PJZiQR5W44IwfyCSX8NrKx2gX8+PW17+fwwF0D8bgBxzbdt9eRbd74RCbMrx/q7hQ3NgDz6dRzgAzLY2ChgRc838jx6fH5YXlfYSpAC82UXJ/Enn39MJ9QWcn7I46n2uf69hhUJNWJBAb8ySCOPh+P44MbbABew4CsvUf1xhQxFRwoJHsCOnHyPPzthByiq1i1y3x/O/52wIRX3GwUXsLE36/10vjMJsFZuWs394WHyxmBCEnwmw6i7D0seL/AM/Q4I/K3ANj1PT5HCrCzKy8kcXA88J7RwQBdifP8vlgQhB3SXHyPTj0/DywaJQTZbAH1sQPS4wVbi/QdF5Nr3/6Xwa1xb7w5HTj+vbCIQCRiLkWFiLWtwfXBLKSQABYeIWtfn0wcAbb3HS9iPyv5YKxNugI9ugwIQXCmwa9r+IH0/r92CtYg8A+v/XBmFmIPW3HFsAUY2HNj58cjy+WFQiPfdc8W8N+hwlIDu5H44WbgA9PIAnp8PbCTBbcD4DzwISLDjyvgmFHA9B7+hwmRbjAhKU52zxn+8MWn2c5fLqHRerNLpmMFE8kmWVkUkwO3vI5JYwvHS/eWv09cVUp2sGB6EHFg6Drp6GPUwgXcXytpAD5GOpiYH5XxnaSDjF3M6jk4K3Y6X+9lj6FWXkuv81zQQ9mWSZpkHZ5keZFssqsvpJZKRKuqCkRnMa7meoikfaGLOI496nYqC+KZ1fDqnLc2qdM6py+bLJsqqZKeTK+67lKaVDZ12D7wPUm5NwbkEHGuqqmZq6aSqZpEne8nHIPQMB6geXpxi8chp6bt007R6IzSSL+32X06U2nK9nA/T0CLaLLpXY2+sqotSyMf1ig0zm4hIfExjCHPxJ661pshc4EMwCoZ2G1U2gW5so4/D1wTnnketx+GFaymmo6yWjqoJIpYnaOSN1KMrKbEEHlSCCLHkEWOEkUFXLOARztPn8Pf2xeVWiTcc8c8Xwgwtzhy45JQtby9bYRkW3vgQFtMiL1VPWZYkhWRo/rMBHUSx88e5Xd+WJl2pvHnGYZbreFAKbVdBHmDkDiOuT9VVr85UL+wkU4r7La18ur6euj+1BIr29QDyPmOMWHDRNmmjs/07BaR9N1C6jy4Hkmjl2R1Cj2AMDkf3GxSeOznDtv4B/281ZHfip1tHuoQI2ILHy/f0wp9aYwLTMoKoxYEgbgT1APkDxxzyMb+jqQkEk9Pl1Bk1DVxyRPVVifWZXR12ssCkeX3SoBB++MbWmyGHI4Y8zzCb9AQMBLFNVIJc0nXyaGAECEHydio9GbzWW1siH7zDYNZ8hXlVNbC5/2rXf2ZqMvpqeTVmZrkULp+rplQyZhVITuUCnBFhckhpSnXz4GN5SiDTMRmp4INLIwuKquAqs3kX1jjsBDf1Aj/wAZxG63XlPlrypo7LzQSSkmTMZ5O/r5iepMxHgv6RhfcnriIVFVUVUjS1EzyM53MWNyT6n1PviART2n7+63nwxHG94BS3o4ssT1r+KKX1et6HL6l6nTmXF6xmu2aZkwqatm/aG7wIfcAt/exGM0zrNc5qWq80zCeqmf7TyyF2PxJ5w0RC7BbgX8zgxgkjkaOVSrKbEHyxbis0cWIFTtOJ4qF8zn4akSxHUYGNDIxVRg7JYdbY6C+jBo7RGp8o1NmFRo6n1HqnIaeWvoqCuqG+r1EaxMyosYG0sSji77hcoABe+JJZBE28RVMjZ2hoFz9LDsUFWBB81IPPyxY+cdkFToTKNDa41fXw1WmdXFpHkyeQyT06RmMyxEuoTvRHKjgC46i9wbXfk1dpj6V3ZvqfJjoHJMl1xpilbMslmyqlSFp40Uv9XO0L3iOEdLNcrIYmUgM6nS6CdO1X6J2r9A2E2b6DnTUWWKT4mgQM0oHxgep4/7lMVzOajCmND55eymEQAOvCo91Ge2/sty/ss1Fp3Uuj55KzRuohDW5RWSSicq8ZjaaB32gMRuSRSVG6OZLi4YCf8Ab5qF9B/S30zrXv27mfLsnndmP2IGhale3oAqsQOgtjS9gubUfbL2V6g+jdqSriXMEiOZaUqJnt3NZHuMabj0Xc7xt/3dQxP+yGEvpnxyT5r2dZ3PFJFLX6TWnnWRbOkkVRKrqw8mUsVI8iDiEVMrYn50I8RhQp9AyMvZlUH1SvbHn+oPo4/Slr9e6Xp4jT53HFncdOSVjmiqRumQFfslahJNpHQqDY8g7nKKD6MH0kc5fKqTLcw7OdYZnveJqRQ1FLNa53U1yjDqSIGha1ysbHwlH6RDR66+j92R9s81LHV1FGrZLmKyEjvrWNjYhivfQzgkdO8HIJGCQ9qH0RqrOMi7UW0xnen8+yanjDabyinWOllqIhaKTdbaSv7W5d9lZ13btzcbge2t7EVG0bR14p2F4tNLueOw7Ef6M+is37N+2ntD7PtTQpFX5dp+vy6r7iQOm0soaSNrWZTG+9WI5BW48sTLTUPZl20djeofo+dmecZ1ndRp6lGb5BLm0Sif63GzFFiIVbo4eWEjbx3ydbDFAH6Ser4e1bVnatlWTZbHmGqaeWkENSrTR0sB7tVCgEb2CRKt24JJNumKwy+szrJZvrWU5pV0UskRid6adomZCQSpKkEi4HHTgYm7CR7u0JocOOtM7ZjW3AKjHgulOznOOzftc7BYexPXmuKbR+e6bzA1eWVmYj9UyBnspBK38MsqMoYMCI2AIuMazt+152ewdm+jOw3s61E+qItJyy1dZnZi2RPI0ZQRQ/tDxuzEXUARqCxDNjneoDH9aXMjP4mLG5c35PPU364VimjSMMXDueLLzYe+JW2cNferhWtN5UTpiW0puqiVDSOAb2ZSCLYtXUX0se3rUNA2UDWL5ZRugiePLIVpmPh2khxdwWHUhhe59cRfs77Mta9q2eHI9E5I1bPBH39TJJMkMFPHe26SVyFQEmwubk9AcSntN+j32idkmXUOY6xyWlGXZpI0VLmFBWx1lM8qrcws6fYk2+LawUkXK3ANnvET3hr6EhNaXtbVtaKqaP6zTytJFNIrnoysVb398OEywzMWaPf5k2v+JwStDUwDR87b7g3W3ljrqDJfo6diOW6E072hdm02rs11jSR1dfm1RVOkdGruI7xorqAquTwu19q7i9yAFklERGFSUjGF4XJklMJBsMiIo/aYKPxP7vwvhKeip5Y0VJIwfZgQx9uefh1x07pfTuiezb6Z+UacyCuy3PtM1Urx0yztBmAiSoppLQuxDI0kcg2hyLnwsQCTi29A9puZ9pn0htVfRx1/pXTWZaSZ8wy8QU+VRwukUTKEdiv2XswKyKFZXCkHqDC+1hv2iopXyUrLOXZmhrTzXD2T6SzjNsorc1y7IsxqsuyxQ1dVQUcssFKDzeWVVKIPMbiMS3IOyjtB1npDMdUaQ0hmGZ5Lk6zNmNVAsfd0phj7yQm7hiBGdxspNr9bYuD6MVEi6O+kJ2eitaenOn6kqytxKYBVIHsOOSqG/viU/QU1PSZZ2daqlzamSpy6DU2TrmEbkBRSVgFLKT6gK/I9CfTCS2ksDiBWhHOiI4A8gE5g8lzV2Pdims+2zP67INCQ5ZLU5fTCsqJK7MY6SFIi6oP1j8ElmHA56422hOwbU/aB2qZl2P0dfk2XagymSujqJa+qcUwekfZKqyRo5fm+0hbEc3GOl/o/9neY9gGddq0VUQJct1vp/RlNK/WWCetZ9492p2Q/Ej0xHOy1l07/AKRPVFMxYKc7z1wBySrq0th6m2GPtTqvu5AVHBObADdrrNCqp199GnM+zHSGa6nq+1zs5zZ8vMe/LMozZ6iqm3yCPwqY1BKlrnzsDim6eTZUM8qoV2jxfs268enni3u2L6N2adnn17VmYa60NmMEucGFaHJ84WprUE0jlS0YUbQoADcmxIxVGYRwwUcwO0lUZRwLgkHi9r4swvvtrWqhkbdNKUXQGQ/RZyup01p3UXaV2y6b0PVatpxLkmUV1JJUVFRC9tjSWdBHu3KwA3Fdy35O3EB1B2J6s092yZd2I6gno6fNqyvpaCnrEaSWieKoI7mpQhd5iYG9toYWYEAgjFmfTdIXPOysp9l9HUZt/wCsX+WLB7fBEfp1dmMV7fqcg59DvlP8sVmTyEBx1gnwop3RNrQaiBxUTb6C+pJays0zpftb7OdQ6iowS+R09bUUtaW48AE8QVSbrbftW7LdhcHFE6a7KNW641xT9n2mcnkmzuollianmIh7juiRK07NxEsZUhiehFuSQD2VWdj2Z6O+lFnf0lNYa105lOj6J5a+RxX97VlJaRadllhQExAEsbMdzHaqglgcV19GjVNLrftY7dO0uhD061mTZnmFIXFpIUnq3lHwayoT7jDGWt3ZufUGgHE6krrOA8NyqTwVa64+i1rrQmksw1rRag0vqnKsnkMOby6ezFqiTLmDBWMyPGh2hmAZlvtuCwA5xV2kezjPe0LVWXaP0ll4rM4zaRoqWnaVIQzKjO253IVQFRiWJtxiwuyPL/pC6J7I9W610n2dHMdDauyqpo8zzSeDvI0p4VkinkUrKpVk75rsVbkDjE1+iPDT5HmHaD20VQLU/Z1pCuqKdmttauqImhgX48N+Pvix2r2MdUgkZfnzUNxrnNpUKq9VS9tnZpSTdhGs8/zrLcopGhr5chGYiWjPeqs0bqI2ZGVtyvwSpPJF8RDVGQZ7Hk1NqGfI8zjyqrYxwV70ci0szrcMI5iNjkWNwCSCD6Yu/wCmu1RUav0Hq8SrJ/ajs9yOplnHG+eNDFI3HFzYfji1cw07TZp9B/LdEQVDtmUGk4dcU9J4jujp69u+kHlfbUSDpex9MM+oEbWOp9xxpzKeIi9zhXILhumyx2mMQe7bbjaQT+HzwpNk8tLCWkBIBtyOMdE9tVNHpz6MnYfkKRhZ6yjzXNy2xRIBKwsN1t1ryDi9sP8A6YGmNG6Czfs3yLLNJZbRPBpairc8aggEEte5dUJkK8FisT2bg3cknErJw8gUzryUboi2u6nNcxUddmGU19PX5ZVz0lRTsHhnp5DHJE4+8rLYqfwOJbm2rNba3MFTrHVecZ0aMFIPr9a8oiDdQu64W9ueOfPHTfZx2dfQ8+kPmWZ0ejOz/tH0gcspzW11S2cRPQ0cG778kwkA43EAkXCMb+E45raio6vVr6c0HLU5nTVmatl+USSKFmqY3mMcBYLYbmBUkW8+mFZKyRxwoRt1JHMc0DYUx0Zq+q7NO0PIde0mWUuYz5NVJXR0tVuETOt9obab8Hng+Qxces/pB9hnarkmeVeo/o/x5druvpZfqOd5dXyW+um22acBkDgeIklWJ4Bv1xEu3PsL1N2J53lmS6ozDJMxXN4JqijqcpqXljaOOYwtuDorKd46c8eeIpl2gM3OVvquhyavnyyCU0stalDM9NHNYMUaUKYw1iDYkGxGEuRzUlB8CClvPjBYVp1q6CmzKhkzONnokqoWqVC8mISKXW3uoOO4e036PVf28dtuU9ptbqXL/wDskqMlpjUZ5R5nArUNHDAxdQr3sxkJYGxSxO4qQQOI83y9JKfvECsjffjIK8denHy4xoxX5zT0r5KmYVSUMr75KdZ2ELn1KX2k8eYwskTnODmGhxHFEbwAWuFQusewauyXs+7J+3vtU0nPUTxUsMencjqJkVZhDMzbJXCiyli8THy4xqdD1T9mP0JNa6gtsr+0LOI8jpDyGemh8L2PmLLP7Xvipexnt/1r2J1eZrkVLluZ5VnKrFmeU5pTielqlW4G4eRszD0INiDiSa2+kNN2v6q0dHrnJ8v03oTTNZDtybT1HaKKHeplZULDe5Vdo5UKCbDk3gdC8vNRhUGvgMuKkEjbo20pxXV36KqdGfR+fseoRI2Y5d2U1mZVVLCfFNmmbVsMESEeZvJOFB8x7YVmhyn6LH0b9caYyv6tJm2nqCLLK7MYxc1eqszjtKiH9mlpnjUW5Ba/BviEaS+ljoSi1922dsE+dUtTAlFlVHpbLKgFJ8wWFpO7CRMLhO8Ikf8AZBJ9MRPt/h1PXdmfYz2LVdTLLq3XdXJrHPO+UrI9ZmU22FpFtfhZGt6LH7YpNideuPwBNTwqeGXnuVkvbSrcwKD0RMiqtBdh30XtHZnrvs5ynVs3aDnc+ZS5fmCL3sdAid2ssMlt6OEVNpBAJlN/LFd/Si7HNKdlXaRRZbpKOqpsrzzJafOYaKqkLy0LOzo8JZiWsClxuJYAkEta56ln7Q9PVP0lMg+jXR6B09n+S6eoqSioczqKCKsqsqrIaUSGaPeCpiVUhDpwdwJ3XNsUbpHSWe9u/wBJLVeoO2nPKWo07oeaefUtfDH3FIKKjkeOOmhW/gSQxmy3LFe8JJY3M0MrgTI/AUrnnXLBRyMaRdbnWnDNU9rjsD7QdC6ByDtO1FkkVNp/Uax/U5hVRtKjSKzxLNDfvIjJGhkW4IKEG4vbFXSxEEsLFR6Y7T+mrr/MdZ9hfZXmdfG0L6szHOdSw0lgBSUV0go4QB02w2+bNgMrzvQmovoyUHab9Ivs5yevpIczg0xp98hpVy3M5qaOMKZFlXw/qtj2DKVfaQRc3xYZaX9mHvbmSMPFROhbfLWnIVXFQBIwKySJwCbHqDyD8sX/AKV+jNVdtNZrHO+wqqafTuTZjDR5TDqKdKWtzEOhbarAd0JFAB2sVuGUAk4qbXPZ/q3s9ziXIdZabzDJa+Lk09bAY2K3tuU/Zdf7ykj3xZbIx5ug47NahLXNFSME3o9bZ1BBDRVkq19JAf1UFWDIsXoY2vviPujDF05D2waJ1zksWSdrNKNQyRL3cUucTGHMKcXO0UmcRqZAADYRVkcsXAsV64532n0xgJU3BIxG+ztOMZund8Ze+9SNmIweLw3/ADmuhK/6O8mcRJm/ZFnT6gWYNImRVyR0eera5IhiLGDMAB0alkZj1Ma9MVZXU1dTVdRRz0sqVFG5jngkiaKWBl4KyRN4kIIN/IG/TDDTmvM+05Tvl1POlTlsriSXLqte9pZGHR9hPgceToVceRxb2V9rWSdpCRZd2m5bFnLQoI6SrrcwNPm1NYWC02aFSXAH2Ya1ZE8g69cRiaWEkTjDaPcZjmN6eYo5cYTjsPsdfIqpo0I2OyK33tpNg3PnY8/Ec43Ok8oGdZzR5fJCxp6mW1WYwS4pYh31QR/yLb5++JTqvsorDFV53o2tmz2hoYu9rYPqnc5pl0Q+9VUQLXT/AOSIDJEepK9MB2d0FPJkua11XUN3Ve0WnaaSFrHum/1mvmBI8oY9t/8AvAOmJZJmiMyNNQoWRkvuEYpp2w6mbPKzK6KSoLvUxS5vULtK929W/eCPaeQUgjgjt5W44xBYGhVLtGjbhtO+5sT8CCMH1BmM+o9UV+eVMkshqp2cGV7tt6KGPsoUE+2DESSWju87g+AEl7D1CgWPxN8Fmj7OMDz44omdeeSkCq94dt+RZUF3sw4AJHNr/P2w5hkkpwzNII2uY9gciS4PK2HPX1tz+GG+1t777t3d7m3PzI6A+V+P3YIsjAq4ibcL7iRuUn4W6exvzidRp+lVIEljiYlJQO8SQAq3+Jeh/f6HFhdhfZ5qjVOs44ckrJ8nyphGc5KUy1UbQu22OAQSeCeWZwyRwycbgzErHG7rH+zPRWb9o+qabIMmggVwjVFXUTgrTUdIlu8mmKDiNbjhRuZ2VEDO6jFw9sfaDlnZ/kx7E+zFJI5ws36azJrCZSybJ1YqbfWJFAjlKkrDHaljJ/XvJQt0zQ3sGgFx1HIDadw5qzZmOvdpUgDZ6BSjM9Q5920ampPo/wDZZPR1Gn9OLLW1ueVhkzCHJyjMX+oTOTKlKilYBGh2VDguEUFbV39KmhhyHWWpNP0lXLV02QTUun6eeZQryR0mWUcAYqLhSSzGwJtcjFk/RK7F62gFRriSpYT1VM9HFHGSABICnNuvJFvLgjqDisvpV1q1WudZVDqd9XrDOrfBK9oh/wCxgGMKyyh1oETXXmtutxzJxJJPl4DUtSWMtiMjhQmp3agFQ0a8DgccAnj5X9MLjYrhxzY8EeE4S3A8grtHHh6DCoJUDeu3cL2J5Ppz/LHWLDSqbiCFFluQNvT/AD462wsX2r3IRRcXILCykn9keVvzwimxQr7F8N/tILWJ8/a/4c++FTJvt4yCNwuD4r+YHv7cXvfpgQiPTshO5bEsLEqtyfh5/hbGbCrg+o+6LA+3T8uuAVWUsqx3Zrkg2HuDccEefp1vjApW3eqARt433NrcE8/MH+WEQjItyFKX7siwuLC3W45HsefPywZvDZksBtJPIIX0sT+R+HngoBVu8icsV6hRyD/I+/TocZ4SHI2GwsGUeG1wTe3H9XwISckSvuIC8WY3G0E+45ufwOCbQhCX2gELZjtA+I6A++FfDsuHAJ8Rueg/rp+OCMwUgqxAIuotY2/rz/HCoQOytGY1jXZfeSTa38AD+fGEWBPiCr09SLj1GFPEH3MAjpYBbG35/mD8sGTaZB43IFyoCj8ubAf1bAhNDGADvkVF8gf65xmFtwd2VjYH2xmBCbsWsVuov1H5c/154Dp5Wvzx54MbIQsqk38+hA9R6/PCZ8IItYnyHN8CEoLEgA8ew45HW3ngALrtFjc+XIB9/wDLBFZgbk/ZFuB5fz98KKAwsXABHXqB8/8ALAhYFDgcFT0uT+/CRuzGy28woPQe3rhVbMt2NjxYkfZPtgpRg4utr3PIv/1wiEnY35Pn9oA4EEW3epvfzt/PpjCvG7jjj3/DGOPO9zexPlbyP44EIj+Lj3vYeYwk4Y8k3HS58vnhVitzcHdfk3GEpCQbkkHqPXCoSZB9PPzwm3nfrfCjsCbWNsJkc9b4EIMWR2ZrHPnua0cibhVZJmCc+RESyf8AuOK3xZ/ZDEk2snhkbaJMnzYbgLkWoGP7xihpH+C47jyVqyfxB4harVeRRUSzzwG4RmP2SPCb+R+H541P1ybJs1kyydzHGjKUJJGwkBgbjkDobjkGzDkYtrVGm5K6mqo4SJlFLIg4s25QT08xa/qR8CMQz6QmnYdNdpM1BCCFkyrJqzkec+V0sx/OQ4rWJ/1HcOWPllT3U1oBhNQp3qehi+kDklXq2lW/aZkVKajO6dVG/UtDEl2zCMD7VbCo/wBYUczR2qFuyy3ocqouD4h7H88bbSWrM30/m1BmeUZpPl2Z5ZPHUZfWwybJIZUa6EN5EHofK5BupIxZetNK5P2kaeru13QmXU+XVtAyHV2n6dAkdBNI21a+mT7tFNIdpXpTzMEP6t4yNKNxHcfmqcjQe83JU8wXqouP73JP4YRkXaLWN/TyOHLIQSm3lb3BH43+H5YScXFgp+Zvb+eJlEmZxNtN5lnDLl+Y5FV09LW0cM9DVS1AUwildSN0gYEFdrstrEk7bAm2IY4HVRbDuhzKSlpqihLuIKooZClg10JK/vPGK9ojMje7mP0PJTQvDTR2RUtzHV+X5JUmbJGkzLNz/tM4rkDSK3/cRG6wgeTNuf02dMQ2uzKuzKokqa6pkmllYu7OxZmY+ZJ5J9zzg4pKSU3p69L/ALMylD+PIw9oMoDygVdFLPEerUkqFx7gG4OI4o4YO8cXbTn+m4YbApHmSTADDYFpwt8KpGBybe9/LExg0VkGZyLFlWtqGjqmFxS55E9Ax9hL4ovmzKMGzrst11p2ijzLNNNVYy6Yfq6+ACejkH92ojLRH5sD7YsNlY7IqBzHBRGOGzEEG4xIpMvTO8lbMadL1dAm2pA6tGBw/wAhz8L/ALOGbQPSw9xVwPGJR4XkBUEjhWB6Na5HFxz7YeaQz+TTOfQZlJT99ArbKqAi/exH7Qt6jqPw88PTVH1jtdXXkcH2OLG+jzr6fs07Xciz1JwlNJUpSVVz4e7kYWJ9lcI3ywj2maHpdL5tFmGT3nyLNYFr8tmU3WSmY9L+qMdh9tp9cV7VXQ7lFrk9PLDXsEjS0605jixwcNS6yy3Ia/sD+mnS0Gm6WVspz6cVeXxohIOX1QMy8D7sEqFWPl3Bv54ZaUzPIew76Y+c6fpZof7M5nWvTNDLIEjjp6lEqY4XJNvD3nckkjgtfEZr/pla1OnMny/JMkoKbO6LKkyybPKmNaipKg3ZYwVG1GazlWLDeWNucURmdTXZ1XT5tmlZNVV1ZI89TPM5d5ZGYlmJPJJvzimyGSSvaYVFPMa1YdIxlLmNDXy2KUZ9KnZV2pV0+h8+pa+LIc2lGW19LMJoqiBXPdncvDBoyFYe7DEj7b+3fN+3evy6pr9OUOU02T/WvqqwyvLM3fsruJJGsGs6kiyrbcb364rEUybVDEGxt8L4cQJ3YKuLAG3HP5dDi32TS4POY1qtfddLRkUQSVUkSQz1EjRxAiNWc7UBNyAOg55+OCGiUgkrz5C/2vbDoBdzAOPXrwbefxHn7YMpABLW8Jsegth9E3FJJCoUEMPDyCR5YNOFU83T+6xuR8T5A+V+cLBVFrSAWO6/VAT7/wBc+uCEMnIJFvCwPN7+TeoPrhUJhPIY2UEAjcGIPQf0Mdt9u/Z7p7XXZtmlPpvT1Dl+d9msNLnAhoaSOI1GU1cAabhFG7uzGZBe5HcTftY4orodsIY7QefCOCPiMeiEmX1miu2HKe0+s1tpjLshrtMUmTVmVZjVlZ651LEAIRtCliguT4lZxbnFC2PMbmuGYqfHLDgrdnbfa4Hd7qiPowbM17K+2vS+VTyU+a1OQiqo5Y2tJeKOVwARzz3bpx/xffCnYpmVV2hfRW7VNAVkstXPpsU+osuSRi7xrGe9IQnpxFUD271/XGy7Osoy36PH0usw0Xmchh03n8T0+XNOdqy0dQyy0yMSbbvC8Bv0dWBxJuwrsg1H9HbOe0fUPaRU5bS6IXJKuh+ttWRv9eUue6KIDuBZGddrBW3OFAvezHva0udtukb8h14pzWOIDdlQVxZmzGSdVBPdseDzY+QN/PjHQ/0iC+pewXsa1uAC60dTlkzD1WOI2J+Mbn5nFASwL9WjJQgbB6/9CcXT2cdtnZZU9lCdjXbXpvNqrLMsr2zHKq7KpQk8EjbrryePtuOhBVrWBUNi1aAQWvaK0PtRQREEFpNKhRHSWk9Z9mmtOzjWupdOVmW5dnGY0tbls86KqVcIlQM6WJ4s462uCMdqZrrqsPbdrz6PtJJQaWzLO8g3ZLqHLoEhrKiqkp0fdUS23SN9obgQQoYjxAHHHfbd23ZFr+q0nkWg9My5JpfRFP3GVwzy755CXVmdzc2+woAuTwSSSeEe1Xt5zrtG7T8u7WcqyWLIMzyuCjiijjqWqAz05a0jMQt9waxW1rC3niCSF05DnCmB9QQpmSthaWg6x6YqzPoSU1ZlXbLq3QGf070k+YadzHKqqlk6pLHKiuh9x4/ww0+jKXPY/wBvOndzLKNNx1cdjYh4hOAfYgquKuqe3vtMi7TKntgyqtoMo1LWwvFNPQ0Eao25AjP3bhl3sFW7dSRfrc4itBrXXGTzZu+S6ir8u/T0bx5kKOUwpVRuWZkdVsCt2bw9OcPdA95ccq3eIOKY2VrQN1ea7s7VO1PLNVdlfZV2g5cZFre0HXGnK3MwFO01eXQ/V5x8d9r/ACxEY4zkX+k2qGlV49+cPOVZbH9Zle+1j7nHGkOZZ59Up8rlzStNLl8zTUsH1h+7gkYi7xpeyMSBcgAmwwOZz5jVuK6pq6iepvueeWV2dvi5Nz+OGix3QW1wII4/CX6ipBpkQeC6Y7WfoedpOU5vrTtMGa6NTJ6afMs9SMZyoqXpg7y7Fj7u/ebDYLcc8Y5kzaqqZ4x3jFgQAOAOPL4/vwgsFTK6mpaQ3Y/eLXt16nDycd6gXaQRbcD0PocWo2uaKONfJQPLSatC7C1TpvQn0qNJdnercl7XtL6azHTOURZNn2XZ5OYZIFjYMZY/2zwRZiqkFSG6gRvXXavpHtF+m9pHUOm81gmyLKMyyrLEzGRxHDOICd8oZ7AJuZgGNgQL9DjlhsunNxtW4kWPxAcEgn5dDg1Xl690BGR4TyxPX3xA2zUBF7UQN1VKZjnTYfGi7N05qHTOc/Ss7ZOzDOs4oH092iJVUS1DVKGn7+NFeJxJu2dGkIYHqotiD/Q6FPk+tu0fsnzvM6XLs71BkVXk1GZ5AI3rI2dWQML3PiLALckKbXNhjmc0LhEW20GwJIuCfUW6/LAS5fPDIJYCQVIIIbnjoQfLCGyi4WA5gcRr9Ev1BLg4jInmu4dDdkWruxT6MnbNkmv66gjzbMtPPWUWWU9etU8NDG6wNUlVJESSS1AAvYsVJPTGj0TmPZ12N/RSyuj7V8hz6uou1nNKisqKfK6gU1SaemEf1cliQdllV7cg96OCMckwV+aM1TUVdfVvLURd1MWqHJkj/ZYk+JbgcHjgYRzTOtT53R0WX5tneY11JlkbRUUNRVSSx06GwKxKxIQEKosAOg9MJ9M51bxzNTTDUgTgZDIUC6d+ls2ldT9i/Y7rPQceYLkkFNX5JRLXsGqkgj2NEkpHBYMkguOtsW5+nqTS/b72P9luZRgZZmXZo2l6uL0Nckot/wCONfnjhabtC19X6My7QVXqKqn07klUa2gy+UK0VPOSxLJxcX3tcXsbnjEs1J2xdouqNZ5P2pakzCnqM/yM0f1WVKVIYwtO/eRBkSwPJN+lwcN+lJaGE5XueSd24vFwGzlmrw+l5lr02tuxzstYQyvlmnaGicxAhWaWrWK4B5sViB5/aONT9NSsqNVfSkzDSGR0U9fWU1JlORUtJCNzSz9yCI0Hu0wHzv5YrbWX0hs61x2t6f7YtVaayySpydqP/wAm0byQ08yU7syi7F2XcWJPXpxiSaG+kxpqi+kln/b3rTSlW0uavUT0MFMyTtl8soRNwL7d5WIMgIsRuuOcNjjkiDXEVLWniSEr3skJANKkcAp122Zjlf0f+yei+jDoithlz/NY0zLXeY0zX3yyKCKRWH3bWW3/AAlHnMwxEPoi5DQ5NqrUXbdqCkBynszyqXNAW6S5g6stOnoW4Yj0JT2wTtFzH6MGqMp1NrnSnaZraDV1SZMwTLM6oI5466pklG5BOqgpwzEEngIBzid5Z2tdmXYf9HbSulKbKNI9pGY6wq6jMtT5XLV97DBdUMccpTxLJGBCq8WDLLa/XBlDdaCXOOOFPHPdgEg/iVJFAMNfhzWh+kXLV572VdguqMxlM1TX6bzJqiVgfFM1aJW/9lI2Jtknaxqr6OP0OdF5/obMYqDP9Y6orawtNCJlajG5W3IeGBSGDnrZja18R36WeZ0mb9hXYjqPTem/7P5bXQZl9Qy2J2kjpY3WFliR2F3Aa/Pvhj9KfTmdVWZ9j3YTp7LauolyzTlPSxRJExElZO4icA2sSvdEk+Vze2Gto9rGOGFSeFfwnOq1znDYOdE4+l3UUGf6F7J+1ip01l+Uaq1tlc9Tm0eXRfVxVKFjaORlXq4L2DkbmVwrFtoONxmn0Zfo+dmtJpbSXbT2l6io9e6sp0nVMupoHpKFnO1Q6shJUSXj3s67mVrBVAbG77dcnodY/S37KuxfLnWbLNJZdldJJGtioVb1Mt/jFFED7HEc1rJ/2/8A+kCgyWECWgy7PoMtUDoKXLl3y/jIkn/iGGxyPMYbepgXHwrgOtiV7GhxNK4gfKoHtn7IM57Gu0XNtAZxPDVTUBhmhqIFKpVU0yho5lQ8oSDZlN9puLm1zCJKUlDEGDKORt6X8xi5/pYa0TtD+kPrTOaNlmo6KsXKaYgggx0iCLcPUF1kPpzipYkPHiG0dDfgD4Y0Yi50bS/OgqqjwA4huS05pGiZil/8PocSHTfaXrbR2qsk1xludzSZxkS93QTVwFWtMgDAIqS7l2je5AtYE3FjhOpgMSdyRucdbi/cj9gH18z5Dgdb4YNRFlKlCQoueOAP4YdQPGOSSpaV0flP0+NWZbS1WZUXZZoqm1jPTfVV1LT0hjmVONpMYHjKkAqGcruVbqQAMbuDt37KexXsX0po/S1BkvaTXatapzbXUNYJFR5JLDuH3AMJEayoSrL4Gcf7QHHJ/wBTWzbY28RsLEHj0+OEZoWjhJRmLKbjqDYf1+WKxsURoAKD12cNSmFoeDVdZ/TPppM61v2W9lOk8nipGp9OU9NQZbESVpXq6g93CL3NlCqLnyBONJ9MrMsv0/mGivo+aZkL5b2f5LEs6R897X1CBmJA6uUCn4yn1xTmiO2vVekO0bIO0bO0TVVbkUUcFPDm80kgECqVRA4O5Sqs21udpN/IYlnZJqjTGufpKUOvO2rP6TLcvrc3kzyueoDGKSRW7yKnvYhVLiNbtZQqWJF8MbC6K7XENBPiT1zTjIH1pm48uvRdDdq/0eu1bSn0atKaG7OaGmlk09fUusBDmKU9XHVmIyrJZipKIN4BBBvCLA2vjj3WOu9c9olPllXrnVNfnD5RQihoWq5N7RQBi4W/mSzElj4j5k46H+lhoztpPaDqntxhzeKt0ZqV1o0zfIM17ykNA4EcNLOt1YbgoBVlZSxaxPXHNRjWWZYgAgRhuJ4UN5A+lhz+GHWVndvuoScfCuabO7G6Kj8LXRZVKIw8qFQxNmINmI6getunHrhKbLyoNlI8unT5YlUNMhRYokkZy21CHuHF/JbXPy69cbmm7PNV57VrSZbp+slqHICU0cLNMf8ADCgMp+SW98WslAqvmpniPiUj5YTSWSIko5F+D74uZOwPPIpzHq3P8g0yVNmpq+rM1dbztRUqy1HT9tUHqRiH6n0RlGUztFlubSVKqSO+rFSk3e4gDPIB/isfbEbp42m644p7WPdiAkdKdomc6eno3iraiN8vfvaKaKoaGejf9qnmXxQn25Q+a4snUOtq3PdMZvqqsnyvu2gagpEoqKOiY1tYV+uTzQx2UTGCDxMngbeGUC5GKWajyyE/6xmYYj7lPGWP4mwwvDWRy0/6OpEmFP3olYyvclgLDgcDFV8Ie4GLAVxwwVlsl0fvMwMNqc0gMYEk8YYNc2JKgE+dx/XrhXYxYQx7y6GwCOep8hbqfhxgh2CMlmbk7SCpAIPzsfhg0Yhke7II1K9C32bdDdRYn4jGhkqWaTcKoe6So4O0BSAp55BHl8B19sbbTGms+1nnuX6V01lbVmZZhL3VLDEArMxBJLOTYKqqzFmNkRWYkAHGvEd5Fihg7xpHCxooLu5PG0WsWBJHFrm4Ax0JkOf0v0WdO1VTU0lHU6+zenMcqyIJI8vXd/5rboyo6hqi3EkqR01ykdRuhnmMTKtFXahtPWakjj7R1DkiZ/2gZB9HrRk/Z5oZJZNSVLJJmuc7DExmC3SRFbxKVDH6ujW7tWNQ476SNYY/9GDQ3/aTVa2zPM0VlpMty6CPgkRibM6dLD08COPmb9Tii9QZ5mepc0qM1zSrnqJ6iSSoklncvJLI7Fnkdj9p2YksfMnHZf8Ao8Yvq2WdotRvmRmy/IgGifYw3ZhL0by+z19sY9vjNlscsx/iEVJV+zOEtoZGPtqusez/AEzRZPHlOTZXTRoBXUqvB34WSBGniBZk5uOBwdrAsftA485/pLSPNq7NHtfv9R51L/4s0rD/AAx6Y5P9XjzbKYaKojp5Fq6GRqdY9hb/AFuIMyj0uSH+8rML+B1OPMb6RIKatrI2YgJnWbLx1FsyrMYP7O1c68f6h6OWppajRQf0n1CqiPoC3p1scH6gjwgE8EfjxjERVberGxHi5sb+Xzw4FPwPA1yfMXFz1tzwfbz8sd2uZSUZ22sFAA6XI/Aj9/FsLmRiGC3Qm4I8NgPTgn8+RjI1AClQBY3uRtH5/wA74x2j3syrYAXUXv8AnYfjbCoSUbliAqyWJBKg9CPT2+OFBK1liXggcAWuGJ6eh9h0vf1wRo2JXhLHhSSAT8CfL28sZGlm2hSdpKbSOjeYuPLz+WBCUcux5YFQBt2cKB7eVv8APAborMSCSeLMd1h5kWtyOPhfGGQAWcEkHxMb2I/aK9Ljz498Ydpsi3Zt3AQdOOi263wiEm4jjCqoUbevjvYn58fxHvgoBsVd/DcttDAC3ra9hgzgqNlkNuCOv4cWv88EDEuii1gg2mwFh1/onAhYoeRv1Zdt58P94+w/njODbwkDhvtDhh7+mBYEgRgbiOl18RHW3uPz+WMZgDZwNx4te4Hx9MKhEIuxa22/mvA+Q8hjMG4+xZi49en/AF9sZhKoTY2BZLgkGx9T7fDBWBtcc7bDy49vXBmUKTYi46gi45/MHAXBB8S2HIN+h8+mFQsXkjbYEG6jrb+eDAdNvBHHrgoBW9x6gW8yPLCm4XsgYk2PS17f1bCISVwpHnbj0/PAMDYMQCBxyOo9/TCm5VsBu3Hp/K3ngQ+65UkhiCPUEcEW+HH4YEJPj7PPW3lf54KbbRcj5A9P44EhTclbW549h/XGCkgi5C2FySB+VsCER9vNvjyBx8cIseL2AuLm3Q4VYHm/PHPv/PCZC2B6c+uFQkz0ufy64IwwqfCPgfTzwkxv1FsCEXnFtdh0Im7QUUkgLlObMT5/+YtipTfF1dgFIZNaZrU98Y/qOnM2n/2aur3hWPYyt1Vt9jax5uCDjN0o67Z3ncfRXLEKygbwrZzHKZjQ5gaMRSNTNJcuSskaEGxDdCpsQN4sCCLjFTfS4j7vtdKAkhNO6ZjuevGRUP8APF356kaQZytJUhJY6GpljlRyGZViYPa/WxCbl5uGuQbcU/8ATFZajtlzWRVAMCZVQkDoGhyegjI+RU4zdCPoXV3e6uaSZgKb/ZUOeDibdnXaDnujM+pc+yOeEZlSo8JiqIxLT19NIuyWkqIzxJHIhKMp4ZTbghSIW6kFj7nBASOhscdC9l/xWUx107ld2vOzrTWdabXtP7Le+XJquYpNlsshknyqp2l3opG6uyqGeGT/AH8AJ/2sUq4qIKSwKgHad3HII6k/DFk6CzzUemtOZjnlDmMSU+b5XXw5hSywd5FVRQNF3LupNhIKiRGSRbMjpuB5INd1NRLNUyVE4QvMzNJtQKpJNz4RwBf06eWI4ZRKDuw9ksjLlN6YyobHgAX6DywgQRh86CxFrN6Hr/ngndL1Nh8OcTqKqZ2ODJJJGbo5X4HDgwMPiCRx0FsF7pfIdfL0ODNLVPabUmZwR9zLL38XTu5gHX8Gvjf6Z7R820vUGp0/mOYZLKTctllZJTBjaxugJRvmvOIkYOSOhBOEzEwPkfhiA2eM4gU8FIJnZE1V0UPbJlOcA02t9JZLnccp/WVMEIyqu58+9p1MEh/9LCb+eHn9geynVMJqdM6vmyCZz4KfPIu5jv6LVwd5TH/nWH5YokblNwSD6jDqkzXMKGUT0tVLE4+/G5RvxGGuikbjG7inB8bsHBdO0fZjq/L9GN2eaypYp8qqZGrdJZ6GV6OOuIO6jaoRmh7mpW6AhyFkKGwvxzpmeWmhqpKYiXu28Sd4tmK3Isw8mBBU+hBxK9D9uvaDoGWZsmzeqpUqh/rCwkKs6+YmhIMM6nzWWNr+oxte1PWOg+0Q0er9P5XHkWc1yk5xlMEBSkStUDfUU3JCQzrZzFf9XKrW8DixFJLW7K3zGSJI2UvMd8qtlp0AFgDby6H88GChgSL36mwBv7fG37sKstwQRZvLnr7e3tgQDcsoswPNha2LKgRe7Z73ZWIAAIFrjyt8RjFDBjGNhFgL7Txfz/Hi+FIxxsVW5BCBRYgnoLk2IJ8vwwceEKyrv3CxuLWHmD/lgQkYzZwrAgHqBa3Hx/nhSzEMysAyeK5NiPX3wJGwq0hb1BYi5HuLcj3weFjsZFHABZTuHFvL39rc4RCKGcFZd7LcFVJsSR58dMJspG1kFj0Yi5svuvngXeMWIVixF2KgEE+Text19fjg4V2YRqA5tuUL0YHzHr/QwJUwrS+wEtuUghSORi0PpGdqGm+2DU+T6h01llfSxZVkVNlExrdgaeSIue8UKTtFntYm/F8V1JEpiZ5OgNzz9ryv8cIhL/q0AKsCQAetuo/y9sNLAXBxzHunBxDS1b/tF7W9adqceTLq6opqlsipPqVLJFSrHJ3XBIdx4nJYF+T9p3Itc41NZn2ps9ggpM5z/MqyCnsYoqmrklSOwsNoYkDDdKcWLWI4vyQPgR64dRwgJHOq+GQkC33XX7S+3UEex9jgDWtoAEhc4rGdTHYEAC1r+pHI+N74ZNQCZmcmxBBNvQ/54fGKSadKenieWeQ2SNAXZ/YKOT8sHanqcveM5jBLSrOrFDMhj3LuKkrutfaykceakdcOJSBM3oY41UgABhdQOvHW+DQR7SvQqbnkfkPTDqqDCQx7Qu0iwDEg3HX0I9D6YJCywy97KoKIyl7PuBXobX9r+vTBXBFEaaERqCwAJva+CU9PCJoQwCo7rG3/ADeH+P5YPK31ZnjO5WjcqTfwtY8EdRzwebYIys4aOzXILL58+v7sBxCMkZIGYCMm9hY7uQD0sfTkEYUdVkiJXcbKeTyVNrWYf+5fj64WdR30k5UbZmWVeSp8ahjY/Engi2MW5Yqob7QsDawPuB/R+GEGIqilCitGj72MYJYBl28HlQR7YQZbkFihWRdoPI2leqkeR+OH1LE5jFgARBHJyvxB+AuB059MKyUNRJ3kvdtbi5YBSCOh5I5t+OAGoQc0ajgi/QM8vhJFfBHuIsOYZPX9+GtQjlwjKEEfN2sCPiT6dBf5YfUwpmyCqhmrKSJ/0pBKYvrEYbuhC6lgu7kAsOOuCrT01wseZ5ftA4b6/GR8lJ4+eGs1+Pwnu1eCa5ZSF8zpbhl2y7m3qAbWJ4t5cYLFEe5iBC8xra/UcD5efXy+HON5lFAqVQqBPSIqQ1DraZCofuXAVQD1JtYDzwhFllWY4e8o59yxndst0Ww6WJHwt6244wXu8UlMK9akxyyhSozOkgliLRNOpkB48CXd+PIWU410EJniEhXb3nLAeRPNv66gY3FPU/oqueVqd220s8J8X2DLG0Kt8AZL2weCmeOPu+6GzcTzHfb8+OPa+AVvE6sPdBpQDrUmlFlMdS8EEagyTSLCBfpuNr/16YNWmGqDGCH9W7uw54KXsLjyIFr42eWzmnqairYEGkpKifxDo5Tu0+PikUj4Y1NTH3MHdoNxuEFuTZeh+BJF/gMANXHcgijQtdLSpONptxYEkD5C/wABhpVZZFv3Je1+b+X8/jjeSR9zR0UIRjLLG9ZIygXUvxGOfRE3f8+EpCgVGEcaAqB4WPJ9fEbjn5DDmuvYpCLq0iZLUSNUAEpHTw97ISL25so+JJAGAoIHjLO5FwAAD5/16YldXCtHkFIspAlzQnMZL33CnQmOAHz8REr/APhxrHQ7iZFvt5e/JLHyPrYcfG+Gsdeqd/p+Urm3aBYc61EaClo1zqukhy2oFRQwPUO0VPKLeNEYlVPCjgeQxbuX/Tj+krlVD+i5tYQVoYue8q8thd9zfaJIABJvzcc+d8VCqre9gkajn9lR+8YfZpky5dMtPL4aiONHqF28IXG7u29wu2/94keWGyRxyENe0FKx72AlpopDoLtoz3Q/bTlHbFqKepz3MkrJKvMpZZAZqnvQVkIJ43bT4R9ngDgY6F0x2pfRS7IdT6m7eND6tz3OtVZxT1UmU5DU0ZX6lVVBLSAyEcjeR4mPhQMBvJBxx9WwCpZlkuGUk7j5j4fD8sJ0NHtlDO24ILgN0v5Ww2Sztk3YUw1jYnMmcwbflbQmWd3qapzJVTM00kjDlnY7mPre5J+eHEdoljqXYCWWxjJAOwXt3h9eeFv1Nz5YVpIaYqaiqiDxodqxhiBNIBfZfyUA3Y+ht1YWb1MzzvI8xV5G5cqRYD/COgA4AHQDElbxu6kyl0VOaKikPwqlUNwu8Nx6m/nzfkdeuCSUyIN29bKQAr8k3vzYdQLckedsL0FNHWVndNMAqRmV2v4Y41HLM/UKB6XJJAHJGCVJEl25COSUsm3gfd44Xi1x5e/XDq40SXcKlNzC0gbaoIZfEOoPIIt8/wCOG0lMACHJvuAFjdSbdL+3z9PLDxZY43YTIdjAFlvzz5c2HIsSOL8DjnASlGYImyVVS+9EYLbzJU88cAkentfCpFr2p0Yjf4yDe5tcH09sN62mmdxIhN7ck9TjYhQrAJtO0G7dLj0B9vU/zwjJVJCuxwNzWUFvsg+pHU2t0wJEtkVVqzMIF0plU2Z1cEk61EWVU7SSrJOoIEghW92ALcgX5OLD072WDLpN3aBn1NkrIN8lGzGorTfkkwQbmQn/ALxovjjR5F2kaxoMlj0NowzUVHWvtn+qotNNXyE/76VbSS+ioz7B0C+sVz2tz2imfK6xKmlkjYmSlmi7to2947AD8PniEuc7BoUoaBi5XxH2n9lWgIWg0no9MzqAtjXZvMWu3r9XpnVflJUv7jyxCdR/SQ1lmVM+VUuYNSZe1/8AUqG1HSn4wUojRv8AnMh9zioJDUTteR2c+pOFY6JmA3Gx688ceXOEEFcXmqO0p9oWzqtY57URvTxVslPTubtDAe6jJ9SqWB+d8adjPKDcmx5sOB+GNguXqgAYfl1/yw5NGvEY5e3l0Ht7n18h+WJWRsZ9oomOe52ZWkWJ2NgMbTL6eSMPcIt/Mi5+HXj8MKJTASFi0exSNzjxKL+XHX4D5cYdomwsrNfmwCc7vx6Dz8x6c4emohMxfwuzXXYA32iCLFb+vl62wrRU01TURU9NGTNI2xQWC8/E8AAXJLdACThRmp1lUTQkggFlluVX+41rFh0IYWYenXC2U55Lp/Nsvzf9VPDT1UElTE4SUTRJIrFGuOUO3kHr58cYQmgqlCufTOnMl7F8sTWmo3E+q56ZZ8qiPg/RkEibo6ohhdamZSWhBF4IP9YI7x4RiitY6qzLWWdyZlmFQZLhUWy7VVFG1FRfuoq8KvkOt2LE2f8ASUyrMKHNqGvkzrMa6PMqjM+/etl7x5KpKxi0m+wLb4np25v5dFCgVBBTbLORdb82P7j5HFSzET/3g68twU81Yv3XFId2C0gH3YGP547u+gJStDpDtDqkpzL/AOT9L3CuqEKausJYFuOLXsetscMMm2eoG1gPqz23DkjHc3+j/wA02ZDr/LkgaV5ch07OArW2CPMJ0Z7feA70cfDFDTv/AOR/gfRWtG/xm+I9V1FkWZ1WXZ9l6F4HV6mlcowKsy9/GCdpNiRxzyVtY3Fjjze+kpGV1xm6njutSZ3HyPTNawHHonl9RBS5hEaihpJYIJqWrpmkYs0UySLvFiNoBtuRlN/tA3FscEfS6yz9Hdo2qqfaL0+stRIfb/yrK4/KYY5v9mvuP/s30ctfTGIHgfVqoZLkr3a23DxEv18/McD88OtjttVXU3G4c2uR6A/u6emGzBbhgVIAC7ivT32/164egKVCBGPhFlvuB/w8cjHfLl0ksUo/XJGy7TcHjj2B8sKAd2yhXVwpspub3HoeCDgCgUbe5uXFjcEH4ehv6+WBLKWA7wEldoBPLAeVvI+3Q+RwIRSbboWdgHN3Pv7gev8AlhJ0WwkDMSAF2kgkjoAPf19PXCspkdVEiyWDAAAWA45AHlwDb8DgpW6yO0q3287lPIvxx5+XAwqEkpYC9rbefvcE+Xsfyt1wcSKDdzwPNvP8Oh9COR8MYjOrIm5uCdt1t8bA8H+WBNzHYKwBN7hW8ugB9P66YRCwMjgB3jtwCWaw59f8h74IxldjIzKNxsS/3iB15Fr2+HwxjEnwhiWYnrwfe3A5/hgm8pcL4SSB1B+Z8rfHAhKSgMGkDOys3nYvf1I/lhLdYllQMVseR4bX8x58+Q9MFeRjcBQCbfd6/C/I+XGChhuCg346Aj8uOfT3vgQhd2FrrYD7JvuLYzBFYKQ3dg+zchh+HljMCE3vuYeEnnoD1/o2wpGSb2F7eg64AkHoQLC4v5++DoAG+7fm1xxe35fEYVCLbz59Ta1zb8sYLgeG3XcFueSffBrAfYFgvPH+WA2kr0vfrzc/G/ywiEEhUEMQbMtw2645/r8cAGsbDbYdL3uPb44GzlRtNgRu487+eMCknkj06G5H8xhUIJCvS7EdeRb5YSCg83sR5+eFCCys6EixuLqenr7YBRcKSBZeDz5e3GBCKBuYEC9gLgAfjgrKBex463tzgT4iCQTYfgPLjGC5Xjj38xgQknBsSTz5nCDqQ1iMOGZQvS3PAHthBjYdfhxhEII03Son7TAficXp2EwxrUaxr3qO5MWTRwxta+5pa+IbR7lUP4HFIUCGWtgX1kX8ji/+wel/95zUuYNGrd/XZdSAsxWwSOonZgwBIIIUg2+PGMjTJpAfL/UFoaOFZR5+hVk6qjqpDWZY9NJIlRl7h3RQbGXcqMLnp4rgg3sTtLA2xU/0tRH/ANt2saVJVYwawzKm4/YhMcK/lF+WLm0fp+o1D2haTyeGEXzPNMuysqu1NqNXw3uOhUCSUWHTi3GOd+3bN/03r/M9QO13zXO83zNj6iSunYf+x24o6JbdpT+r/aVat5ND4e4VZOt0vbqSfzwgFG6xNuecOgp7lLeQ+GAoqOSvzCChiPjqZUhX4swA/fjpSQ0VKxW4mitDMUiy3sxhpAQkj0uX0RU/aJmklrZT7WVYB+GK4O0ruYm9h78fw+BxN+0TMqeeSKmopgaeWepqYvDtHdgrTQH2vHASP8WIM5I4C2IJv62+GKVgaRFeOZJPnrU9qP7yg1YIwBuCGIupHHHHSx9cAQpAMY+NwDb4euMUMW2Fhwet7j4+/wAcH2xgDYrAN0Fr8+fy9PPnF5V0V0Cjbe567h0YH93+WCLGet/tdf5fHC5jRAzNKFYKHQhTZvYH4dD6ixwndbfZN+nHp+/AhJOVJvxzyQMCwTaLgcizHp+eN3lmm8xzOIVtVLFRZe7/APntUCqORxaMAb5W/uoD72xOI49F6Co4qsOHr2QSCSZY5a8E9NsRvDRD0L95NzcKuIJbQyLDM7ApGROfiFF9Odm2aZ4EnzCVMppJBuRqiNnnnH/c06jvJP8AFZUHmwwvn9Lo7TCmgy0x1EqcPOZFnmY+d2W8UX+GPvGHm+NNn2u8xzQ1ENITT09S15kWRmM3oZXY75j/AIzb0UYjVpJmG5ifK58sMDZZcXmg2fKdVjBgKlOKutjlc9xCFB8zyfz/AI4PRLLv71ybrwtxe3HpgsdEQLsOf4YfxkhBHt4UWVr8r8PUex+VsWA0NyUZdVAtuAQoQmwsL7fcD09R8xhQWYkeByvRr3Hy9R8cFBK7WOxSFFyAbk+pPr+WBLoiDwjw8C4uLfzH7sKkWILjuyDwADaxuLf1+OHC+NWLFtyDc/HLD9r4ji/yPrhqU/aIW1zusBtI/rrhamCSsY2nhiazEMbi7W4U26E9PnzxgKEeTxXZfC3mAbX9x6H1HzHpgqRSVM4paaKSolkNkiijLsx9goJJvjdaY0/TZrLV1tdLakyuFaiop45Ns8wLbQiDkhb/AG3sdi82JIxMXzFtNPW5x2e5TS1GT5/U08Yy9pJUmyqqY+CEuWDSwt+sEbsxR+C1nS2IXzXTdGfJPbHexOS0lF2SaqlqI1zaKnyyEQfW5pHdaiSOINtdjDAWkDISN6EKVBDNYc43VLpvsryOilizej1JnM8s4vOrChdIrfagVO+Q8m7Ca3AXbzfG0ZzkOarXTQZpSVcVXMiw1EkVLOk8LbXErQsZldWax4Utfg2Jxrv0dmeaVc2bzwxCBou9qp6hFpqbahI3GRR3aXPhIFm3lSU5LYqiV8mLjQbsFPcazACqSgy7sbq2mpk0tqmcygJT7NRQtUMx6d0i0uyRrfdPFje+Ged9klNLmsVBoLN6jMHp4jU5umYpDTLki7lVPrNSkjQXYm1gdway23G2EaPUWRy55R5BlWpsyymizPdBnWZQ2hWqS94o0DC8ZvdGlJVW37mUAG8lqK+l09RRUeSaOyfJlSTaJFieSvjkjP6yGpeRrSN07yF49hBDR+uAukjfQE+f59vPYgNY5uPXW9R6DsU1xUSRw5F+hs8iM4gmly7M4mjo3a5DTGTYYYztbxsNl1IvfjG1o9C0mjHgzevr9P6ky6pkemlSFZZo6SrSxV3gkEZlQo10kB7t1bqDZS3rM2arzOap+rworFwlPIu9IkJFozvBJUWW5bk3U3ut8GgrKijikQOala/bLveBO87vxq5V+TGdwdHUcGwbpYhznyPbRxTQ1gNQpPp/UsmVzVFdpuKly+XMk+p1KZTGlLS1TRkldvdvvWUbj027xcFeCcLVHafnmnqCWatzt0o6ucySfrhPJUTmPZPtUq4kuNpYMwTdtckliMRjLZ5HyL6uI50hkrBmUToWjKTRRmIutxsZtp8VyrKygg887N9UtJVyyV2aVobMkmpquZEZZKczRiIzpCT3TvssGVRd43J4YA4gdGC7EVUoeaYFVxnGYrrDNo209pEU9T3b97DQRlzUAMWEphjXbG2zhu7AQldwA6YIdHamkyKTUj6czA5ZAds1T3XhivYXb7ygk2DMAp6A4nEWcZHk+WPQwCKljjQk09HMYo5ggtvic2M5ZgGBZt63PhsBiE6s1nqHVleavPMweUKAgu54UAKNznxSGwUbnJJAGLkTnuN1goBtUD2tGJOK0LkyqFLEMoAbyN14v+QwqCYw8twhTazXNgQfL0uCL46C7HvoW9pevUpNS64qqTs/0zWDvqeuzuB2ra2M2uaSgFpZQbcO4SP+9jsLs+7JPo1diapmOkdKZPqbUNMN36e1ik2cTxvb7UWX0aGnh55AaUN6nEVo0hZ7Kbrnd7Z1l5p8VklnxaMNq4h7JPovdu/bZHSy9nnZbneZURjaJ8wni+p0CjcdpNTNtQgK33d3TF8ac/0d2m9P1Mb9tP0mdJ5LV8d5k2mYGziuFvuEiyq1/PYQMX5q3XtX2iU8cOse07UeawyxkwZdFRx5fl0KAkD/AFOCYXAtb9bK/wAumNDQZemWUyRZHqqPL4gbLGmkaIpbqSRHURm1upLHqB1xgWjT8p7sIa3/ANqn/SPdakWi2DvSVPhQep9kpk/0Uvof5BBRzJoDtC1vJcwo+eZ6uW08rAbjaKHawW1za2LO0p2Q/R+p6SNMk+jn2R5bISbR5jRVWbSr6XZ1ufxxCoNW59QZaxm1vo9qGIWWXOdL11NTgnopemrJbXPojH2wfTOcavz/ADOSi1f9H7KYcq3bYtSUueU9LRT+d44M0po6iQci+1T7Xxnut+kZQXNkBAzoQ2nFteasdhZGEAsoTtBPoacl0bpTs/0plCxyJ2a6BhpjC6uuUaLjhZZDbYytKhG0C9wTfpjY5foeiSq3VWXaezOFjxBnGj6Gm4/uyIiqfmcUXN2cR1QH6JzPI9LTSh0WryyWuepi44dXp/q0IYdfEGU+YOJzoTQXaXkX1bu/pJanzTYQX7zT+WxmVR5OyG5v62v54jimmeKumI3XyTywPmU90cbcBH53RTnjwU2zzsL7La95arNexbs8qkkFyJ9Gwsg/9bBuPzIxFc5+hP8ARJ1rlbrR9kGj6TMSt1kycSRPE/8AgSWNiPawOJPS6b+kFD3lVk3a3pTNihJEefaR2SDzAaoy+pU+24x34uRif1GZ59R6YirNQ6dpc5zimoBLU0mVhZ1mqQPFFTSVJRuT9nvCtx1ONqNz2MDw8kHeSfwqBuuddLaEbsFw/rz/AEbWgy7/AKI09qWmYjbv09qNJpAt780OaoC3rtSoJuOMVTU/6NjV+YU1W/Zv2m5HnFRQC9Tk+oMoqckzSn9BJGveL8HAKG1wTj0Ny/t80NHTGn1nR6n0LMjBTFqfJKikgT4TkS0zL7iQDEwos8yfVFPDneQzZDqnK4UPdVmVTxzT0/qVMTsCvsjKw9Di8Jp424n3VctiecAvDztE+jh2zdmlPXVupuzzNEy1mijkzfLduZZegRyxDT027uyWCf7RUPhPGKjq4A0JaCSORWOxZImDKHPHNvsn2NsfQVqHSemtTwtn+W1tTRV6rZc0oKtoKlbeRqIxcgeaTo6+otjnbtY+j/2N6vpoZO1bssoM8rq2Zaeh1XpqNMqzKrkY8U8rU57k1drmO4MU5GwCOQhWki0iwGj+Sa+yEju8F4/VNmaaQxsE27EF+QiqFA487AfAk+uC5fRHOKymyyOXuzVSCIsGsEQ8s59lUMbn0x212mf6OfM8xy1dQfR517SaxoawPJQ5ZmoWhrqhUJDpBOLQTSIQVaJxFKrKQyg447zvSmqezrO800zrLTmaZFnscJpmosxpWp541c2kfaw5G0FQy3B3nnF1krHspEcQqxYQ4XxgtXntcM1zKavjVVgqAI4EI4jp0AWJB6WRVBHucISQdyUglkW+1ZGAB/VlubE9TZdpNul/jhMyIj75Y1ZR1A4v7A/l8DhIfWqyVYY0lqKmolCqictJIxsAq9SSSAPiMShoaA0ZAJhJdU6ytrpxFjaoz6qRZKPKCjCNxZKipY/qIbcXBKl2B+5E3rhiJaqeeSrlked2cvUzMTfcxJLMfUm5978Y2Oqj+iPq+kaV45IsldxWSowZZ8wfiZh+0qbRCp6WjY/exrqKjlzKsjgUqGlNtziyoALljboFW5PoAcNjIIMp1+n5z/RK8UPZjo9YI8FCJIpK+WmkKlikABAV5ALnrY2UEE+5UeeByzKhPLKv1gw00EZmq5SokWGO9rgftsSFVQfExHvZ/WJU6hzOjy7JKKSZWIo6CAkXKC53G/AJO53Y8AlieFwbOMxoY4Y9PZPIk9BSSCWoqlFvr1TaxlF/92ousY9CW6tw2844azyHXPcEtGjHUOZ65LX1NQtQwaKIRRIO7ii33ESDnaT5t5sfMknpbCEcdRPVRU0EMs0sxEMcaC7MzcAD3J+WMLlbCJCLkA7ep9Pnc/njZSsNO0bKWZsyq0ZD/wDI8R4ZR6FuQSOvKiw3EyE3BdbmmgXsXJCuEVDGcpy+ZZfEGq5oyCsrj7qk8d2h+yfvG7dNttZ3pR7ArZhbxDgD1Ivb4fj5YO9NNJCtVIAFkYrGCLBmHUAdLDzPlwBzxgoiKp4mBYA9WVTf3FiflhzG3RTNNcalDFLErt+smUkMFCixA89wa9/UkdfTGGRlKuJn75eUcWXZ7m3U28vTr6YyGGQPHNvhgUr3iCSRkLW+8Co3DnoeB6XwVWjjO5AfF5H99xw35H1GFSIZIUUIDE/d7LhWIG0/HzHoT5eXGNPnEUm+Mksdo2kHqD6WxvqaF6ol4YwFVS0jX8KeQJ9OeLG9zwL4RmSGWJViRlS2wsR4jbqv90f3euEvY0CWmFStBQ5pVUDXja6nhlIBDD0IPBHxxaOnO1mgraOHIta5NQZ5lkY2JS5mJHEA9YKlD9ZpT/haSP1S2K1qKAbSQhWx5J8sMJIXjbg38xbDHRB2IwO5PbIRgcQr1ruxPT+qKeLMuyzOpXnq2IgyLNZolqpm/Zo6pSKatP8AcBin/wC7JxWlbktbkeYVGW5xRVVJVUjFKiCWBkmif9h43AZW9mAxqMl1TnGR97HR1bCnqLCendQ8MwHTfG11b4kXHkRi48l7asj1ZQ02Q9qOVRZ7R06CCnkq6h46ujj8lpcwAaaEDyinE0PlZRiMSSRYSio2j3H6+ScY2SYxmh2H5/RVewgeImOOVTtupaW7EjiwsLL8/LphE7bNujEaE3sd21Ta1geT+N+eTi1dRdiMuZwvm3ZLmM2p6VIzUSZTLAsWd0sXm5p0LJVxAf76lZ182ROgqhroxVol8BItt5De563Hpx8MTse2QXmGoUTmlho4UKxUWQLGIIQ67nDsxDEccdbG3l8cHacMI22IoW5JiuvJ4Fr9B8OOfXAAvLGV3hVvvYvdgzWte5FgSOnr+GFFMYiWMOjvICJFeFjtN/sg3sQR58HyI88OTUk+1iWRgb8hS5Zgw97fngtURIssAsFdC3iKlgQP2hx68DjC26VpburyyBL8m52qLLx/dHqemG8iKEKm1t97dPLni1+vvgQr77Ramn199H/J8/jjaStyetyusnckEBKqkahqD/8AjGWwE+83vii40kWQRRK5baQAlt1va/X8yPLFvdiKSaq7O9RaDJV6iroswy+mTz79Y1zGlt799QzKP/S++KfDhlLg3ElmCk3Fuouvlbj0OKViowOhH8p5alYtFXESbQm8fjrZoWkR7Uc4CqDZPATa56nHWv0A8wA1BqijYygS6Clqf1Shm30ubU7AWJF+JDxfpfHJ9IrPm9OHLHvFliufO6EcHz646L+gVnK0XaxR5fM5WHNtN6ly2WwJuBSfWQOOesAPHPGKumGdpZnt3H/S5WLAbsodvHqF3NJCpoKuv3Ofqu+XbFKI2uqEFJY5ByCOfJgQCpJxxb9O/LzT9uPaIIlUQrq2qqYwL32z0lFPf5lmPzx2vlj0WYyVP66buHRqeaGVAWj6gOCv2twNyvQizJYhhjkb6eVPFmHafnGeU8cncZzkGnc4UspUsZMuNO7EHod9NY+4xyn7NPo54Oq6eZHutzS4Ja0/+w5A+y5D2btsiedwCPbyw4hJKsygtdDwG8I9SQOR/V8I90wC7EBUgE3SwJ9L+fxwoLJtkdJSxNwQV4J6LfqDby/LHoK5RAr7UUEkKLmwHC362HvwePjhQMViKFrLYjbbg/gP3nBA0bFrxXNuCebfE8fj64NcABHIsvIBuAR7gfZP94cHAhNvCAGS4a4F4yAePf1Btg/eWHAF/MLx09A3Tr098AWRT4HADX6kE/8ANb09vjzjCG5CuFQkAk8C3rf055t5YEIVYEABQFte5ubD0sD/AF5YFwO7O5VuvDeE+En49fYi+Ci+wgopvzYg7fxv1wVb2O4rtX7O48L68eV/XCoQlwFCi3XdcdPzHl+484TewNitwCLX4F+vkf3YNssrbl6nzUKBz634wBjBCgqQRxYjgjqLD8eD18rYRCTFzLbceeSSbkHr8/b3wZi3JF2XkWbz9j/XGD2eQBiwI55L3B4546+XTr0wQ+NC5HnfxD8v6+eFQkS25nQNuXi/v6H4YzB7EkMWuSTwBYDGYRCRZgzXKWUeQ8v6P78GBUAC+0EW5P8AMYDxMLA2t1v8MCpJuApBtew4/fhULOpUhSTa3Asbeot+7Bi1yWa58y3w8/lgLXssfJuBfz/yxilVK2b5dOfQfzwIWeDo1go6HYPw464y1+VQi4/DnyP8TgzE7r7iTflj+XTphI3LbuD67vX38xhEI11A3BTf1Hl/AD5/LBRawBPn7j8PIjA3sQCfAOqjyPqD0JwIJN1YWubgDoPf3wqEmQ1yxHN+CfO/XBSot4QxPkw8v54OAo6bbjzFhfCUgII3XJtcXHl/DAhJOL8mwHX2wibWwu3S5HJ/PCDefpfCITvKFvXBz/ukeT8FOOlexvLZYOy2imWindszz6uqO9jIsqxQQ08YIvezO8gBHQg39Mc35MngrZ/JINvzZgP3Xx1nprLWyPs60bRpJG0rZFFWPGGs0b1U885D+m5DGfgR6Y5/Tknda3aQORPrRa2jG94nd7gfKmXZVWUWU9rOm80zPMStNkFRXZw8jRGOMiloKmp7z4d5AlibXCg2F8caaylkY5WlRJvljymJpD57nFzf3uxx1bBUyLk+r6+J2dqXSWZxrvUqVeteny+IWYDk/XJOlxwbE45R7QZIG1TmkdP/ALGCZaWP/CpI/hg0UDVgp/UfQfKW3nBx8B6laOwVAHB5HXjg42eiQ0eoFzJUDnLYpa0D1ZEOwfNygxqJmAO3ocS7s7olMU9bIg2VFbS0hvyTGrGom/8AYQgH442LW4NhO/Djgs2zisgPmmur6Z6TOZMredJf0eEoC1+CYYwr8+neF+fbGkK92/iG0qOg6j0tb9+FaydqiqaolYl5CZGPu5Ln/wBthu7cgWHHHy/q+JYmljA05qN7rziQhsbcC4HmOl8DvKRnxfeFwDcg+/8Algkcc0s4gjhZpJGCoigsxJ6AW5PyxK6PSVDlBMurGlmq0Xd+iaaUK8foaqblYF/ui8h9F4OFe9rBUoa0nJaXKsrzPO2MOXRbxDcyyuwSKBSftSO1lRfifhziVUmUaTyGMVlXUQ5hNbctRNE31UEeUEBs9Sf777IvZsafONWLFBHRQJAUiO+Knii2UsDeqRHl2/7yXc3t54ilVXVeYTPNUzPI7m7FmJJ+JPXENZJssBzUlGszxKk2d67q6qqkqaaaVqh17v6zK4aYJ+yrABY1/uxhQPXEWZpqlryMTgyUx+188OUivawF/jbErImx5JjnlyJHRAkXPNxx/I4c/V1hblQbHkf1yMKx2Jte9+TcX+I98OYZaUENWRtKnTwNtkQeqseOPRrg+3XDjUYhNGOCa7Qv2lsARuI54Px/C2BVQpZWO1g3Tbf8cP5sob6vPW0hNXSw/wC1cJtkhv8A8WM+JPjyvvgmSZplGXZitTnuRvmVGE5iLugW5AWQ2tvA58JKhjbnCXwRVuKW6QaFJUuXZlmAqpcvoaqrjoIDU1Zp4WlFPCCAZJLfZUEi5Nuoxv8ARmW6SzGGtr9V5jVLSU4INNRVdPFVEAXMgE/hcDgBV5Zifs25mFDqfJtUd5RxSQzZPSRSU0OWPUrlMssbrd5dtMtii2AEd2LNYtcAjDLK+zzTFFqOGKsq2zOjrcsnqqfLpP8AV64ybxGFPIVWClpY5CQrhAStrjFV05IIdh4KYRYimKr+jqnpJ6euopds9PIJl3xqwRka6GxuG6C4ItxbkYlvaD2h5DrbLocyzahqZNQqUvMVVQtm8cbSrYzREElAwEkZ8O9ksMa6lXSGk9WZzluc1EOoMogE9JDJTRbjUgnwTRyBgIXWwbcNwuCtmUnGyzLRuT5jNlOi9GQZdmuoa/MGemqKXMhMXpyt44p2Noe8J5soG0AhrluHvey8CR57t/umNa6hAWsbUOntOaoOaaZpUraGro9ywVEzmWgkcC4SdQCJEIur2PhO1g3OJI9dDndItTqXIaikpaunY001RmDmaqkswp1RSAiU4kO4nZtNiAbnDjLMmrM5yrOzrPL8ograbMZYKeo/R7NJBWKyPL3hgun1dANvdkMGMngB2mzWVNS6nzCeDOaXS08G5pJszhrTFCOLAsYydi8BfEgVRYHaBxAXtd4jXX02qVrSPNJZblWZVsVZn+Z1ix5bTWSurqqQkAKgVUsT3kjFSAi+nh3W5DOfP8p7R9UZXkeY5hU5JlEUKwPOCHkrqiNWWKadWdYVmYbYu84CqF3E2JLTOXz/AEfNX6eqcoqslziVJKXMJjWtMKmhlVSkKggqYiu1g6swYbSLWxFlnGXzQ5hFDBOaeRZO7lQPGxUg7XU8EG3KnqMTMjLgXcExzwMOKt+ky3S2l6mir8h0ysgpqlhPU5oY6ypS19pCFfq6ncp+wJVYcbgbYY5xMKqeTNatqirqKyRp5ayoYmSRmO5mM4uG8RJBI2rwpsRjSJ2o5VKuaU9VNnlRDnIp6qRZVhk+q1CFi0UKt4RGdwsyhWAUC2N8cuyLPdN02e0MhqhPVik7x4FpquOZFV3SVIy0cwMbbllUKxNwQ9iBVDXMNZAfHPrWpyWuFGJkKyCTL/q60yRTJHErSRVbOrlNwJMZF0LMyX8W1SDbhiAnVSCeWPZVdyFdiSYDIE3oC7bUIkILAbgoNrlgLE4HOquohy2myiCnq50qZ4ZoYYkV5ZKlwyKQpF2LKSoKGzFLG5AOGtE8cdfDHU1sTxPKizGn70TQpt2klT9hkcrdOoZbAlcSjaoynk+YxUcH/k3M5VjdIqruXnPeKDGFJZE6vwTdbqyhN1rcGy6rkzFpYR9SaGnEjU81QiLLIZCP1bbSU2k7mRS3hO5d43WxrZ6CcV9PS0lPWNWGQxPTxhQ4ntcuGcCyMLsb/YsytwFJSC18FKmZ1FJJTrXRd5BN9bRjKhYr4woAKFlI/usouLHCgBJUp+aGHN6g5RW53R5bHOQ9RPVhglMA3+2KoC5Y2KiNFLNyCBtvi0+z/tf7AOwdpMw0Z2RZ3rrVaJ/q2p9SVlNRpRS/8SjoAkqxN6PKZHHUbTjmHNquWorDLC7BSNo2kg/PDrTGjNVa3zhMh0lp7M86zKX7FJl9JJUzN77UBIHueMSmFzm3b1AfLmCCm32g1piupar6cWb11RUV9V2U01ZWVb76irzDVNZUTzN6u1xu9haw8gMJT/Tfr5ojDXdlGWSXAC31FXts/wAIYsB+GKzpvomdpkdK1TnFAKWRLh6WJlmkRh92Rwe6Q/3dzMPMY1Uv0d9dUkjM2XyEqQO6SMu/z6fu+HrjCNm0TeIqK+Lv+S0xLbg2oB4D4Vpw/TG09Pmv1/Nuz7NKRiEVhl+eJJGdq7VOyaCwIHTnqcdAZHW12eZBRaiz+lzTROUVMIlgTPYqefNKmM8qYaZSFRD5S1BRTcFY5Mcn5HpDPOz+oo8/y7JKamzigDNFJmeWNVL3xPEigEiN0t4dyHabm5NrNc27YtfCplrNT5VkuasZC88komZ2Y9WciUPc+ZPOKcljs8rq2RlTvcTwBNOfkrDJ5Wik7uAHqu1NMdruhdPZnMMrzrKqfMoF7k1tZmYlzOO/mk0i/qrgcinji62AxL9O9ouWauqjVw5g02YRxkVDRkTvJEW8MyzKSJ1uNveBrXLLIsbdeDMo+kHk+UzNM2k6zK6oqvjy/OHMTDki8VQpJXm48fzxsc6+lnmNXUyVeTZTS0szG61FRurqgHYE3DfshV7C27axIsCWtinNom0y0Fw+ZFB5CtB4KeK2xR5OHDHmvQOo1ZWwUlRNmFPJkeV0saStnGaVdMlM7EkMiu7bInQgXEpMhuCi25w3yj6bnYJpQJlmddskOZyqNq/obLKzNLHzO/uo4ifdb48r9Y9qmrdeVcc2eVlXmcsI2wyZnUmp7kekcdhFGPZUHxxrZZ9TVuXSSVlar0UNg7zIpjQnoF4+16BecXINBCKjpzidVaela18lXk0iZMIxy+fyvazR305Pol53UQUcfbpT5fK4ZJY86yioy0gkWUxysgEZB63JB9Ri7tP6mpNWUCZhonVGn9ZUyrZmy+vhmaRfIrIhIDEfdkUC9+cfO1l1VqSKGNIi88FRuMUU3j7wA2JVTza/Fx5/DDrINWVGnc1XMMsqszyWtRripyuteGRD6gqVYf8AixpO0YwC6wnDYa8a/lVG2pxNTr2j4/C+iWHOKOmaI0s89Ga3eop2Vo90if7SN4GuBIvmnmPEu4cYrDUvZd2IZ3mYz2s0Fl+XZu7XXO9MySZRX7ixAcTUjIXNwRZgbnggHjHlhoH/AEh30nNDMmX/ANu6fXGVArfL9UwLVtZem2U7ZlI8iHNsdI6T/wBJD2Zayy/M8n7VtL5xoaszamAaZI5M0ys1QFjIVG2oSORLLKq7ibK6MrruNeeC1wgdiajnw/CljlgfXtBj1rXUOZVfaj2XSfprSfa/l2pcoXassOvstkEsK+SyZvQKJEU+UlTAy/38bHNe0rIYJKSs7Quz+q0LV5wqKuZQV0U+UZmWIKbK6K9JPuNiFnWKQmxU3Axyw/aZmmnqSh7QOyvXS6y0k7qKhaDM3rJsukP26eqUBZkNgSlQETvV+2BKGDyDTH0l/wCyn9p6PIlymfLKWJsyzjTmZqoyzM8skISaUxorLSVcEjKk5jQwSxyRzNEP1hVHTtkaWTNodo+OinNjuOvRmo39fhdA5pp9K7M8wzjTFRTQZjmMqy5nRVJkiy3PJVXh6hY7yUGYKoG2shvvAG8SAELptdU3Zh2laYk0Z9IvSb5zk1AUjfMswQLnGmnk4jeolh57l/8Ad11OTC/+8UG+KvyHtC0KslNX6CzOvo9H58sVFLk0+YFKnTFTOxECR1KMwfLZ5RshnRmFNUbFBCP3alrO1DP9KaipptQ1z53lFK6UbVFTHGJqaOqJRJVJsojlZWgqaVz3RlUOhQyREVnzvZ3q1cMj89V9VOImPFKYax8dU9Fzj9Jv/R8aq7FqOr1/oWon1t2fCJqhq2lCtW5bEylkaqjS4kitz9YiFiL7lXrjk7Ia6DTNRNnauP0hAhhy1hz3U7izVHxjQkr/AH3Q/dx6l6T7a897OsxzR9FaO1ZnmiIjI+WPQ5LU1FPNRsw7+npnKkSRqXaSK4K7kliJKSXxxl9MLsL0Jl9Y3bP2LNCNK19alJneTU1wcgzCUF4ykbeJaSoG4xhgO7cNEbEKBr2HScdtAilwceaz7RYn2Y9ozEBcyNsEqqkQYX7si9wQeAB++/qAcbeOaPKsqkV3U1daNjt5RUym5t6GRl/8C/38J09BFSoJ5Akg6KhupPqePsj1/LnGwoauDKIX1jmQSapjk7vKqeRQUlqFtundehih8Nl6M+1eitjVldQVz3bSqEbanrBOa930jlkuQLGq51mkKjMmI8dFSuAy0n92SQbWl8wuyPzcGLeKGodY1DI4sUZdwVva/nb8L84K81c8srVVRJJUTyGaeSXxyF2N2csObkm59cbzK6ehoMt/tBm8CSwLI0VHTPcfXZl6gjqII7guerEhOpayUETauxJ5nrgM9ZS/xHUbgByCXgEem6GPOKqBP0jVR76CMiwijN/17L5E/d9rsOoI0tO0M0rZjmksssZc7kB2vUOB9gN90WtdvujpyQMOJ5K7O/rWeZjUOybwJ5yBuklYXWNB0LED7I4VRzZQAWAYF1ZnCsq7Fjs21F9AfjzfqTycKxta1z17tw63pHGlNmr5S9ZmNRW1ZqJxECqhI40XbHEg+zGi+QF+B8Sbkk4GOSBUDyKJz1jj52sf2mtztHoOWPHABJwfVFimcteRFAiiYXMpPUsRwEXqR1YkDpfDR97Eb0AZftggE7j6+hPr5+1sSUwoMAmVxqc0aSdlYzSOzSyE723bfOw5t+AHA46YXy+gNYzVtTIKeijfbJMyl7tb7CLxvf8Au8epIGH2W5PRxUS5xqGRkoprGmpYjtnrbGxKf8OIdDKQb9EDNyG9bWVdXU968SRRp4KaGnG2GGO99qdRt8ySbk8sScMv3zdZlt+OqDknXbmLuCVrKuIwxwUkQp6eN/BCh3MWI5eU8bmI8/sgcKAL3YM0zq7Mq7Utew6elyf5/DCTzEt3oZG5IAZTa3wv+HP5YU37mBkR1IsVJjNhx7Dwg8EEe98SNaGigTSS41Kbyxhm3oPEwYX6rfr54RmpGIIZVI6qb/Z/kP8Arh8qK7BXY7Y+Q1+nPoebfmPLCjwC77nBsbEA2sPc+Xw6/DC1ASKOz0V1Lxcgefp7H3wxKPGfMEemJO0S27u4ZVJVQBbr+zbj5fmca6aj87XUHkgXF/S/9DAhLad1lnOnZIvqlXIIoZhPEhdgI5QbiRCpDRuP20Kt74vKTX3Zv2z00Q7Roaqn1DtCfp+iii/Sb8WHfL4IcyA4+33VTYfbk6Y55mpCObEWwgkksBupIv1HkcV32err8ZunkfEdHep2zG7ceKj08D0FcWv+xzPNG5Smqcrrsv1FpaWUQx6gyiR3pFlPIiqEcCWin9YZ1U9dpYYrwkqWXZcnw2tcL7j/AD4xJOzztm1TobMvrdBmMymSH6rPcJIJ6c9YJkkBjqIT5wzBl9CpscTqu0/2Zdqu2t0lJlejtQStxls1QyZFmMp6JDLIS2XSk9IJ2aAnhJV4XDmykG7IKHkmuiBF6M1HNU+sgU7lLG+4Md/G0i1vW/Xm/pgyRuCQyKONwVvv+Vhbz8+Px8sbLUWnNQadz6qyHUmU12X5rSNtqKWsjMcsZtwefIixBHhIsQSMa5Qe9EccitxwL9fVR5G3oPlzidQqwuwzViaQ1rW1kMZ2wfVM8hR3uSaKoSd1J4veAVCn1Bxou0fTNLo7XeotKDww5TmtVS0+3xBoFkJhb3HdtGb+hxpMgrv0dqzKaubckZm+qzBSVPdSAo4ueeVduuLF7bctleq0nqypWoQ6g0vQyVD90HJqqJny6p8xc7qRWP8AjF8UgOztR2OHMfgKc9+EblValIcxoZiST36FnuLEbh0t0FsXR9D2vpsk7f8AQUVbIEp01lBls5JteKrjkpmX4Hdb54oqukZlEySLYPcC3iv6n1P5Yl+ns9l0xrWfP6O4ly6ro88g/wAUM0c4t8r4bbWX2lu0e4Hun2c3euti9QMmWajlAbvGaMxXQELKroAu5b+FidvKEi5vtINwaH+m1RU1VFo/MaWQOtbo+uyeQ7GQpNlubFtrBgCCI6scemOjNVGMZ9X11BKslNV1s80RjUqWimYzR2B4N0ktxw1h0PGKX+k/k0+e9l2nNQTEytlWsMxy2Q2IKQ1+TpIikEAi81IeCBYi2OA0Eblocz/xPKjvZdTpHvRNfvHPD3XnxHfaGVQWAHAPPzF8CxcNJ9pWUNc3Pi9tvW/9cYWrIo4KmemZQe6lZVb0F7/xwkFtuCDcqt1uVBF+PcA889fwx6Y114BwXHEUNFkdl3bQCqAEcg2/5uoHy/ngS0ZNwQLD7o8vTk/xwdmaRgWuoVhsHPA9BYennghPfHagVjcm1rdPh5j0HxGFSIjFQdu9bKeDv+17gG2DRAEEkEG48XkvqT6D4YFi+4RqX6C7KD18vK5FuD+OAiuym+9QpFja4DHqCB148xyPPCoRxCCbkEWDXc/v9OP2ep+OCKCFvvAH2iQlhf0Pnb0HF+uA5RiTuuBbZYkcdLjqPM3/AHYUa2wsImYW+0DYAE+nnzzz0PvhEIiJuSyxhVvYAkg29LjkgflhOWN9y8bgWG0A23X8vXn1wa6h2aNF62JCgMD8b9f+uCysWZtxQ3Auxv4h53HpfzA5wqEFk3EnddOgAIPysMJScbWHNueBxY/M/O+FGkNiSEII8N+Bcew6fEYTk2kbwO7A62P4knzOBCRSQRsfCT5gE4zAgCORma+0k25Fz6HGYRCK1wpJIJ4v5j5+h/LGFCBuYXDc8nr/ADwLWVuqjw2v08/P8sZZBdhYW6Fev4HqMKhFBB2ktYgjkngfPBi+03RwhHTxWuP6/DBbujC9h58An/pgwdiSBci/UevpgQgYAAEbbHoCRf8ADGX2kEC/HX09r4zcwuOQOlr9T6enI/dgSpLB3C8+2BCK7E25APQi9rfxxnkL9CLegGAXkdeFINrdfh7/AB8sY17g3W/r0HXp7eWBCxvvHaQLW5HB+eE2QC5uB6c+fpgxU3Nl8/tW/fjCDboCT6DAhN5OnHmemEG5JPS+HEt/j6jDcgmwPGEQt3llFK+SSd1GzSV1bFSRW6ltpNrfFlx1/nlJT02cVlJBC0UVA8WURMzAhlpadKdWsOP9w5C+/vbHPfZJksWYax0PQVYP1ammqM/rObBYYA0pJ9Btp+vvi8Unq4RBLmEqd7VhzU7XBIMu1yzFb87pS/mwHNhfHI6WlL5g0aq8zQf6Oa3rBGAwk7vSvul8zqocp7Nc8rZKbecwznJMsSV5C3dpTJV5rMnuLxUqm3HItjkDMal6urknkYl553kY+p/6k46U7ds3iy3QumMkoqhmWsjznUcgsqWWeoTL6YlVAHMNBIQbdH8r45lcM0tz91ST8ScbWi4y1oJ1NA9SfVZ9ufeNBtPsPZARuBNh4unGJ5l88mn9IxER7XOX1NWG8xJWOKeM/ERRyEfHEGghmqqiGip03TVDrEg89zGwH4nE411UhoJKGlj2xPXikhUc7oaOIQJb1u7Ofck4sz9+VkerE+3ueCgi7rHO8vf2UVSaCJjKrNu+2jKL2YHwqb+VuvvhTLMlqcxj+sM8dJSIwR6qW4QH9lQOXb+6tz62xuaHT9Nlke/PIUmq4/H9Q7wosQPQ1Ug+wPMRr+sN+duG2b6l+rT95DKJapV2I6oIxCv7MSDiFfh4j5kHD3S43Y8SmtZhV2AW0WvyzTCNDQCekIXbJOzBa2YHqpYXFMh/YS7kdTiL5tqeorkWlpkSnp4yTHHGu1F9wPX+8xLe+NTLNPVyb5WuSeB0Av6YWipxxuNvc+WHMhxvPxKQyampFIWdiXJvfn3OHUcBvwvI5AHmP542GSZBmGdVDw0axRpEpkmnqJRFDAg6s7ngD25J8gTiQClyXIFWWiymPPponUzy1odYXBvdI4lKut7HxswbjgLhz5WtN0YlNawuFdSidkvZW4NhuvfB1aw+yODfn0B/PoMTwaa0XqPNRRZG+c0FbOWEFLIYpqepkFj3cVQ7RtGGFyGkUgEAFje+NfXaMyitaZ9P6noK1U8arPDLl87pcAyBHDIwBPO1uLE2sL4Rs7TuQYyFFUYKoYbbEDksRb8MOYNiyq1RF3qqwZkvtLj0J9Pf0w7bTVRlVRRjMXkkpKyVooZ6Fo6kSlSAwiZGKyEEjw3B5HrjcP2eZwYpqqjzDLfqEQkkjesqo6SaSJGCyMKeRu8DAn/ZjxNYlQw5w4yMyJSBrknlOXz651LRUkU60U0iO2Y1x4CQrzI9h1ASwC9SSFN73xYmZ9oWfVmdZ4zTS0eW1047ulmpvrvcpCLxQCJlK7gBv7oDYrXJ2gAmCU+ndbaWzl860/RfpKnoImqErqaAz0lXSl+5ZgjgM8bE7GUruUmzBSMJZHm+d5xrKiFVWVuXx0yVEax5U60JpEaNt4UniNTezluSpIJ6YqSxVxFLoH5+PBTsk25kqWZ1paHtFrIM+0NUZPQZ5DTUw/RVIhpp62qUHfNHsUU0c58JECPua1xdiRiM5d2eVtRnKxa2z+LLYayiTMvrQX9ISSrIrP8AZVhdl7uTvNzXTYQfFYY11TLTZaGpjJEXjPdyLHILG3lwSm0eoPx242ul56bO82mgkrJIoqjL6uCrqoYGnNLE6/7dol8RjX/eFN1kLEXtyoL424HBFGvdjmovpzIpdR5q1LEI5RTqJ2pxKI5qiNSNywg/afbztF2ABIBtbFh5RoTTU+sqqaPL5ky+myKlzqGjnzPZCzSBA7Gp8LinjLOxYeOyW5POIbJkmZ6Y1I9C9DT1GY0k/c9z3S1cUjMPCVUXEgIYMpHJBBHOJF/Z/Kcw0hW5tmjU1TmM+VPXxVwqwfqEsExRaNohYRXjAAUi5LRlSADhZnnAh2BoMN/489iIm1wIxCToMh09n+o62v05mOa6a0zRRK07xzPV1E8iLeRaZbR77jdIFkIKR3LMT1Qo8ozuvyhoajXUqV0UZq6bLkp2lEgj8eyWVfCJNqlgh3jw82uMP+zPUsR0pVZJJmgWopc0StoqR9hAEsDQ1Eqo1hKSgRGj3C6ngdSF65MzymoargjrMuMHihdmdEjAXwkSygqRfi5PiB28+cbnlryzZlX9OKUAFoKZ6xoKF6FtTQa/odTzVFZ9Wqt1NNS1IdlLrII5OGhIDAFCNtgNigjEDdNxkXbxIALHm1vMev7xixNM6uz3LaatybUDCfTmYUs0NdQxUkYpyzruinVEsN8cgSQSI29QDbglDnaP2ZDIqBNTaWkqsw0/KIe8lkIeWjkkHg3OoCywuwbup1AVrFWCSArizG+ncd5dYKJ7a94Kt1pokABA4NwfX2xMezfUNPlddnGSzRU8wzShCxRVAPdyTwuJUjZhYqWCuqupDKxUgjm8USaSLu2jcowHiXqGI6gjzH+eHukoM0bUci5K0MDvTTXlkhMzQRlbOUUAs0gHC2BPI6dQ6ZodGQUkRo8KSV71VZFRdxNVTwVEhrYP1zNtlclQysOj3A3WAIdG48QxmXzUMUlQKowhp93dzr4zcizWXkMr8mx5VvED9oY2WYpmlZlUNRmjCSaOiiSoargEcveLvB3bD4l8AAcgtuQ7irC50FbmL1FY1WJ3pp45VmWQuqOCoHiLAC8gIvfqeDyb3rsF5tFI40Kd1ss9bGab9JTlTEQQ9YbOpsSt3J3IbDob3sGFgDhxpnQ+vdfZ9BoHs8yHM9Q5lXOJqbL6GmMkm6w3yWHEa/tMSF8N2PAxYPYZ2Faj7cqmq1VndfBpPQ1HWGPMc/npwRUTG7Gno6cWFTVlbsVS0cYu0hVbA9eUmq9A9m2jX0F2V5VT6V0tVMq1NfVVX+uZ01uJKuoiHe1RbkiOMx0y9A7jnFS2W6OwNo4VdqaPfYOe4qzZ7I61GowG3rNVT2TfQH0ZkNQlT226i/tLnpIvpXTNX/qtPJ+xWZgisWI+9HTKxH7Yxf8ABn2kNCVFRoTSlTkOW5PToitkekqMzIXsd6zx0feSzeYLVUpuPIYgNPqmLVOV1mS5Tl9dn1JJB9WqYKWJkpljcXQinp+XuOVM0sguCCotjRamoNT1UEdJmFTl+n6CJwGy6qqqaggCftLHvUAIQGtsG4ErbjHK2rSE1tNyd9B/SDQees+dAtqCyx2YXo21O04ny1DmpOvbbp+iqZ8no9E6ikky2ZqW1c9DlMcNhdE3SNPIwCkWKQ2At54h+d9pDz1bifRGicsjmp2kimzOszDNGaS/kkklJTni5BMbKbWAPTFXamkhy6tppsp7QdFS7i4nWTUtDDNT2bwlV3sjhuSQeALC3oho7tWk0hqFdOVk2S6lyzUc8kVfLluTZTqbPZNyWVFeTc0kFwP1QEdl4Ug84tQQC6DA0DDZUnjVQyymvfcT509FZ1RncmbfU8vzyuSI5kFajposhpcsWqHl9Xhgp42qRyLBZJLn1xp67sZ0lqjLEzfVdBqnKZ6hWLULaXzWGpprMRtkc0bxMzAAgRsFFxc3BxLZczyvRFbluQ5XFnGQTZvHJLQUNfpfMtPONv2hHFM5hDkc7aeQdORxbEgyjU2bMk/cZnnJXZvL0lRVyBTfkkrcj3LBeD9rFSSRzHl7g6uonA8KUU0YD2gNIpx91WlT9Gb6NWepSx5flmvIisQQOM472YDpY08tMjr4riyi3tyMV3rX6I2h6Oinr9KdoOY0RpXPfR5zRr3cC3A/XNaOSEc27xozECRudAd2OhWqjXV0mcJNRTz5ksSV8WY0NJXrWSxpsRzBWhGeUJ+r3w1EchULcOQGxJ8r0ppQ1SRSqaaDbdIKOSrqYI5LEbqSOYtVQAbirRK9XFYkbBc4sWe3zt/63kfkgj0UUtljJoY/MdVXFWedi2l9EzSZfmlbqKizGnhFRPNm+TwpTpEykxyiGGZ5Xgc2tMrOhBvY9MQPtt0pXaH1jBpKuztcyp/0bl2ZU9TEAsDRVdLHODGq+HaBIVuOu25x1V2n6Vhi0dmdHRZVFBkuS1U1Or5XUipy6incizIhtJlUzG2+B1jpZ1J2qkiqMVJmmhq7ts7JaKDJqIz6/wCymklpKzLom7yXN9Mh2ljmpwOZZKNpJEdRdjA8bi4Rrbej3mWW/KcdVacR+MMVQtTLkd1g8aKxu0aFNEdg2pKnQultMQSUlLBls+YLlpfMYqZ6hYWkEz3Cjb3aAqOkpPUk4597aPpLa17eo8gg11lOl6eXJIXi+s5RkkNDNWO4Ud5UNGP1jBUUACyjkgXJx0R2Bds2jdd6Oi7PtW/o6XORQvk9Xl+YS91Bn9CyCMGOW4tOEWMMoIffFHLFubcmGNZ/o/8AKazOBmeV6j1rTZU7bzQPpxampRb/AGBUrIkTcffZE9SvliCw2uDRjXw2w3XVJqQcRqIOtOtMMlrc2Sz4inArl2k7NM6q9A1OvFoy+Xx18tEs3I8aRo5seh4ccdeDiKQ5lX0yJAlVIYrXMbHcg5PkeMde9vueaC7O9F0vY/paWEVVNCKSDK4qtKp6IuS0lTWSp4TUOWYiMeIsQSqLGitUeqOyb+wOU5Xp7UCmn1LnHd1+ZUTD9ZlVHs3U0Ew+5PICZ3jPKKYFazFgLtmtxmaXyt7pPdrnTaoJbOI6BhxGfiqxyzOMzySrGa5ZW1WX1ifZqaGqenmT4MuLN0x9JvW1BmdHWaoag1XFSCWF0zWmVKmSnmiaGaH6zHZirxuwO/cL2NuBiIQaKnr4e6hQmaJDNEv/ABoBy6j1ePlrdSobzTG+p+xeWqy6jWmeOWpziphhpZFlV+WJJUBTfwRh5JSR4QFXzvh8zrLIKTCtcOjq8kkbZm/YpX2C69NEmZaGzDMmTJMxpKqKllZwRRvJGdwJ6qGKQufu95Cj9ecdF/R9qKv6QWoa/WWsQ40VlcVPS12WyORT51mUiJK1HKB1pUdO+lA+1tgXgm+ObqvsCyHTFJmGYZtnRaOhpqioYqwUsI15CjzJJAt7N5C5vTs4zas7MtMaR0lTSCleiyqKtro7eJ6yutUyMbEeLYaeMX4KoACDyMLSEsb2ulgNSaUwpTafHIbs1qWVr2kMlFB1QL05yHW+m9U5JBpjWlDBJAEEcUndd0IwosOI7d2RwA0drW44HEf7Rvoq5NrfJa6llMGr8pzCn7mXLM9mIqHg3BgtPmcVp0IZVZe971AwB46449zvtjyfRGXZfqrU2sKLTmXVqOYpJ5XeSsdDb9TTRqZJwOhcKqA8b0N0xfH0c/px9mepopsuybVgziGhRZaykaklpqmCMsF75I5R4luQDtYjkXAJBNaxWguIltjTdH84rh4nI031T7RGGkts5Ff6dvgNXkuBfpPfRPzzsczGrzTTsOdZnpejdTXw1MCjN8jQnpUxp4JYjzsqorxNxu2NweZ8wq5s1zAzTU6woiLFBTpysES/YRPW3Un7xLN1OPZfMKLNvpp9uWr8u/tJNlPZl2X1bZDRR5bUNTzaizkRg1Tzzx2lNLCH2d0jBXLAte5A51+kb/o2avJTPnvZnRU+XSKjStRIXbLai3PhuWkpD7gvEPMRDnHSRW8RS3J8RkHDLwOw78uYWO+z9oysWB1t1+W0c/FcC5LkMFeZ6uqqPq+W0SCoraiMD9WhNlSPyMjnwovrcnhTgslLNqyunzOutluSZVGkbmMblpYee6pogftyt4rD7zF5G4ucb3UGQ5vSV/8AYnM6aoyKDJ3M+YrVxfrEkIAed1B8ZIssSqSCpXaTuZsamqzKXM5abLctgejyygZvqVKWBKlvtSyN0eZ7eJulgFWyqBi8HOkdUfoPk8hzgoGCh/U/A5+msfMJM0qIFiWOjoacGKmpw10gQnxEsftMerv1JHoAA7zrKxluwmrs7kN3ZQrJ3RH+0YH7G77qnxEeIgC13WfZImlqJznEb/pOrjV4YEfaaZWswmmsPtMvKxcHadzW8IMeinnzOripaSimra6rcIiXeV5ZXNgFVfE7E+5JxIxzXAGP7fXrbrTHAtNHjvenWxIVEsaOyiyBbNyxJt63PU/gMb2LJKXTTit1JSCbMXQSU2UzXAjBF1lqh1C2sVh4ZhYttX7W4+qZf2cgKr02YavQkOybZaXJWB5VDys1UD1bmOI8De/KROR6iqlmq5ZO9kdjNNO7kyXJuzEk/rCT8WvgBM+WDfX8b9erDElBHnn6flK1eYVtbUy1Ne0k88ou0jj7XHHAFhYcAAbbdLdMZAYmeMSRqWX7e4EgpbhivAJHqbA2F8NLukpew2E7/wBcockercgfh8OcELQL+sLMzsSWBHFvieTf8h68YnADRQKIkk4pcSuszSGoZ+84bgKXU+q+XT4YMveHvFRRGiguzKzKIgfUk9L/ADueL42eU6eqq+SWGenkMqwmoSmjmCSkDm7lzZVI6FvGTwiknB8209VZXBl1RqKaPKRXNBUwUjQSswoJUDJXLYWljIuAd3eEqRYYj7QE0anXCBUrRwSkkWVgT1PRj8PT9/uMLd6O7V4rLsJAKtax9L+XuLHCmY0D5bmVXl7xyg08rJaemeFyB0JibxJcWNmPF8NfESzMfFa97r08unQeVumJBQ4ppKWKFow4BQcq5vusR5X4t8OPjhMIG2u4C2A56Hr0Hr8Og+GFYC8jsIo7tfd4QQ5ufzHW/UYWjpZonkJVZ44rF1RgDEpa3iPIC7iAQLkFvIHCFwCUAlaySDeAtiWta4/P5Y19RTIxun2fTq35cY3tYBtu+wC9+7UH4fP4k+XlhiyE3CRbri5G4D5gHg/PChItDJAyc24w4y3Nq3Kp++ppSpKlGBAZXU9VZTwyn0PGHsiLKQwO6/HPl/LDOpoyORwcI5ocKOyStcWmozVtae7V8o1DklJo7tFyg51lFKhioQZxHWZXf/4AqmBMa+f1abfA3lsJ3DVav7OJckoZNUZBmKah0yXSL9IxxGKSjkY+GGtp7lqaXyFyY36o7DjFYHvIH9D+8Ymmhu0nPNI5glZl9Y8b901O4KLKssDfahljcFJ4W+9FICD5WNiIbroftxGxS1bLngVo6mnd6eQm+9bOCfb38uPLF09oOYrqrsS05mgqppKjT+dSB3dt7CnzSnE4U9OBU0lTYeW8C9zjTVmltKdpURr+zaODKNQyKe80y8x+q15tyctlkJIfz+qSneOkTycKM7P4MxzTRepdJ1EMiyS5DXxCFkKutXlk0dcgKt0YR/WFta9iwxXtTwSyZv8AKeRz5VTogaOjdrCq6shSWJ2ADMB1Isfjh9EyS1+USOSEr6D6pIfUgNH/AAXCCzb4yyqCp49R06/E/n8hgqP/AOQqSqU3bLswKX9FcBh+aNie0CoHDiPmibCaEr047L84qNZdnGg8+NTUJ9e03lwaTYjxfWadDRSCx5vuptpB4O4C4ONj250b1nYXr2mjor/oGLKdSQSRSu0aPRV6xzgRvdowIKuQ2uRa4FrYrb6Juf1td2OUuWxsrw6Z1LmlC0bw3VIKpIayJu9XmNt3f2DK0bjcDtIvi8smy+bVkeY6Cq4t9LqzKsz045DbWQ1dJKkIZT9q1QsQuDfkXHFz56ylk0sDqLuTvwV1TqzWEjWG+n6Lyn1VR/U9RVsPi27w32b2t4f/AHHDBgsT35sRcW8NwfO4Pn8/liU60pKhZcrr6gFZquhXvVIt+uUDeP8Axhx8sRIPtc+BW3C7WBF/kMeh2V1YW7sOGC5ScUkNPHjihY7yVW5/uqOp9refr584PtsdhSzg8jnePYji3rxzhME93dtrC4F/XjpwbH19cBuAQFtoFzz9np04vwfzxOokpI2zaZHIJJILmxI8iL+vr+HGC7fG0fhLDwlVfd8Ta/PvbzF8G37pCgZTc7mG8r8b34I+F/ztgASftOxjv0N7D5MP3HAEJNgFUEMgNywBAbj19be/8sGqNoBRUG9Ta4C3AtwoIH5/9cC/D7lG0yg7mIKkn1uT4h5evzwVVd143Hi5RTwPj5fPAhYql5CS4e9+WG1iQLgk9L9eD/ngZLmwCEFhcFY9pufMcfnhNjtN1YsCCAQCCRb0PTBL3BuOvJtZf4f15YEIyRBnumx2K7doa9x0tbrY9R/lgjIqqpG0XubryCOnX1vfCz7xuLBT1O0jrfrb0H/TzwkQVQAKvS67QF4+XGBCSZPCSBbceB0tx+736c4zB2VQxRARxfgC7e/HBxmBCbs5DXBufIED8sGVhs45X26f5HGFEjv473NrefTrgTxYrsb4c2/zwqFhFmBNvW7XHT36/LBWsUuRew+Nr+2MuqG1iATdbck+v/XB9ykL4xa5At19+PX288CETaSoJuQDwSb3Pt/XGBAXwnxBCTYdbDoSPe/nxgVsST63F/M/54E8WdgAWFxcdbcf0cIhFCjaLdfO/PPt/V8GslrEEG/PngwABChvs+foflwfjbBbngqLdbeotz/M4VCSNi1wBfp0A4/r88A3HB9P6ODNyL7hboLjCbnk7TYfu9xgQk5ADdTYWN7Yb2ZiFC8ngDDh/CLDiwuMO9P0oqs9pUfmOJ+/k9AkY3n8h+eGPfcaXHUnNF4gK7ux+hQag1JVKob9HZPT5BBx/vJ3RZOPP9XFUXHnf3xOs3nlpkq65opFmVHkCiAtvnAOxSB9k7rBXta3hPTEU7KKVafREGY1dKzVGcZnPnJlDEMiR3iiPHNr98SfLg9L4sjRkWX5lrnJzmUrnL6CsbPczSRjYUVEkmY1TMfJXMKpbkXfjjHITt7W0lngOGB51K6CJ3ZwXjvPXkqZ+krWCl19mmk4pGaDTEWX6VjO21/qFOiTn/mqDUN/zYpRDv7yQEDc3T2xJdZ57mOoMzqs7ztjJXZlJPmVQxPPf1EhlYn5ucRlQyRgHy8QHkTjqrK2jK7f09lhTGrqdbVItAwoNTLm8yLJDktNPmkm4cXhQsn4yd2PnjbZdm70uW5dWTRiklpqaSJKoPunmBld2aIMLRXLkNLz0G3njGs0pnVDkeUZvUVNBFUz1vdU8YmBMSor9425f95dlTwdDt8XHBj2Z5pW5tVPPUzO7SG7Fjy3pe3HHkBwOgGK9180z8KDAV8vyR5KWrY421zxNOvBO82z56p+5o1WGEElQl7AnqRfkk+bHxH2HGNUsJIucGRCWUEgX4ueg+ONnUZetDmEuXy1Ef6ptpl2sFPANwLbvP0vi21rYhdCgcXPxKZJCI7LMCm02Nxzf0Iw6SGcR988UgUfeMZt+NrY2uoKGhWmpc0p86o6iSqiJaGCmdG3qbMX3Hwkn8fIWw+0R2z9qXZ2n1PSmtc5y2mL7u4paoqpJ/um6n4EEYjMr3svxNqdhqPb2TuzDXXXnhj7qXdjtG2rMp1lpKVIg2Y6cqZ8sKIABWZdtrQwPPiZI5lJ891unGIfSGnly+Svp6pH3yqBGsoEqixY3Q82G1bMLi/BtfE1ou0rWdLrqLX1Xn5oc/ppxIM2hooUaKQptZmjRRGZAjNu8PitzziUdrHa5205VnGc6Fre0XO5aLLKg0kpWqWM1sZUESAxqqpG6OrKqi1mAJN8ZUU8r7QWtaO9Q5nClAad3H+XZ4K9JExkQcScMMvPb4qnaoxyqZaiPvqYIFMkVlYMfsGRSD+6xHQ4NF39EBDBX1NMRKsl4y10ZAbSbTYgXYjjqL+hxiwz1dRAlJDAdkTdzTEb2KWvbn7ZYXIXztwOmD07M5eKKCeolWnJ7tQZxEii/eLfxKqg3tfj3HGNbLBUM8U/izCZaioXJVbJ1rERaQU7Ax7toEiLKRvRntwbg/cJtbGrirqvLnkTvDE5e0oZeQ6ngkHo/l4ueovzgqSJ9VdzHuEsdtykFQfUr5njgHryMZUNFMWaLZdrt4Ax28c2vyR95eptuU9L4UNASFycR1T1dXDPJUNI5bkA/rY78HaOPxXk+YviRU+o/wBJ1VBDqafN0qKVzSR5zQTMaqGlIYGCRW/2yXPAYghSy+IEDGoBmr4qVJczho6ajo+5agolKF3U24L+CRpL94XubC4sLAYbPS1SV1hEZ/sMjU79+rIwuLG92BHPPNxawwwgPFCnVLTULb02T5rHIM1mzLR80KtHLPVUtIlTJTRvMsfemEIFWxYGzBfTrxjZZTTIkmaQS5rl8eo6Ovn2zZ0jfVnYOU7xJUA7iQW+xJeJwRypG0xV55I3MSqVVlZrBztKEccj7S8X9Dz5qcbKmNR9Yq6bLGqHeWJ4ZY43vLJEVHeCxN35W5HW6joecNdHhilDtYT/AC2PPNNZjV12dZxkwlqo2WWvpszjrKqAuwZpoDTuSJjYgEkAhiLi98PM3zJ/0rUamhp6FcxrqmGqnqY2G+OMqO6UtbYGcjvZWtZ2YC1t14rG1ZHO2x23NGPHGNytGwuCAOoNtwPkRfqCMbvIIoZ6qCnneGVYQ7CKaUxAxqCzBZ1F4zwTyrIDa6gXw1zP5inB2FFr58kjGpY9VR5NH+jEQV1TEdq07zq5BgVbGxeQLaLrZj93nC2vqmjra6lhoEljWpp0qZpYwIaasEh3RTJTrZImAJRwo27kJUC9sSHUeSQ09fWx5HDlU5y9aemp6bNRIJK6aRO+JWPcIknRJFQxE2JVtvXbiA5P/abW2s4srrJw+ZVjilUVNoVQopATaAAgAUgIB1sALnBGQ91+uXXXyMB/dF2ma3KZdUQVmXaQqlljrljDRSQq5anLbmMU8QBN1F23r4lVudw6SDT0ldm+n9QZPk2ZuMwzGOKmp6aec08FXTO4MyRh22Cdmji2o20MoJXxgDClLpyTR0GYRZlKmYS5oEggmhR41hRHvJIZnXcrH7BjsSQTvAGAq4ZNOZNNLW5S80aq1Pnsoa9ZSpMy9xKEDGCalIUAECzMxVmVu7bDHPFcOtf4/VOaMFX2Y0c1DUTU9fBJT1VM5SeGVCkiMDYhlPIYHqD8fPDaky7MsyzKko8jDtW1Mgjh7t9rbiDzfi3F7+wOJB2j6gyyvrKHKMtr0zeXK4vq8ucIW21S2GyJN4DtFHyqNJ47Er9lVA0mTVkVNm+Xzr3iiKqjYhWAbaWAZbnjkEi/T188XA4ujvUVctAdROs0qXyHLqehpKqqqYpA01LVSIYro+0sUS5K2kQi9+eQwHTF8/R1+jrRdo9Mnan2nUlZR6EpSTBQQ1CwT55LEAHVah7JS0isLS1TkW5SPc4JFMU40hlmfzVWr6Zq2jyhWMeWRPsGYTCSyxPKh8EfG52U3Kgqhu24L9pXb/2g9qstPHrDOXqcsoIkpsuyamUUmVZfAgtHFBSx2RVUcAm7eZJJJxVkEr2UgwO34GFTy9FMwMa6smWxdY9o30mexrJKqmy85wmewZRAKPLdO6Qp1hyqghBuIVqZgY9lwCe7ilLnxuztyKW1h9L3WmYGQaT0ZpvTkTuXSpqKT9KVpJ+8Z60sob1KRL+7HO7ZhVysRCVhB8olC/n1/PEroNBZhaNs0hcVUwV0p3B3IpFw0g6i45C9SOTYdaB0bZbMe0n7zj/VjXywCt/VTTC5HgBswTvUnbX2uawkddRdoOoM0STrFJmMyw/ARoVQD2AtiOxU2ZTOsrU0IeRgFAgDu7E8AcEkk9ByTi48t7IYchyqHOtRyihWti76Hv0LSzx/txxDxMn/AHh2x/3/ACxOOw7QOq6jUOU9o2k8xzLTGWPJVUGTVseVQZjmubTDwyyUFPIDHCsYBU1bOBGS21iwIDvqoIWEtaGtG6g68Am9i95oTVxSXZp9Dqolih1J245tUadpmCumnaLuxm8qsLg1JIMdApHQSB5z92HzxeFJm+S9neWNk3ZZp/LtG0Jfu2bJxItXVJ6zVhvVT3F7jeqE9I1uMaA63oMlrgra91dqqWOSRWEeeQQ0SODZh39NTWdr/a+rBlJtecHDxdR59UZUk9Tk+T5gWkZHmnpsw3ub3Cf+ffbAIt4bsFv54w7Za7RaTR7g1mzHnh608Fp2eCCIYCrtuHLrzUnyPOc81Tk1dJSa0r6HLaec0tXXPUSVSioABMa0cjEuRceOYxISfDusRjQHS2mpM3y3MNQar1PXVmXzmanny+DLNO1JbaVt9Yy+OSfZY3I3Wv8AjiK5ZBmuqdQtHpin+sZ1HMWH9naWb67Sxf8ACaeKRmEIBvaqYryRdRxiYU2htd5Tp7OK+pznTebZysE1bk2RfoGHNc1qmjALwS1NM8dJvCh3WNFmkIQgAsLYz2xmJ3ccATtGOOw0PnkFZe9r/uaTTYfUfqt1LlOlo3q6nL1rc1pq4LFmNHqbULSTRupvFU0mZzwj6vMrWvDVo0MikAspxsswzOvyqvqMmzfJmyyrhjgmnp6xadVZZEujDb3kLIQCpdXZd3hG22IjpjtdoI0lyqm1dTSrWU0uX1OWy1ojNVTyi0kE2XzohkB53Kq3JsV2kXxvMgy7SdNRt+gNE5RQGiTuI4aOPO4Kh93KrCsbzxyOD4+7eJl8DbrYY8kgCYUd4U4/pVSNFCTGajxr1xUmOntG61rQ1XmWapXSxSCOWgyikMndFbtEk0Jgnk8Iv3IZhewCkmxpXtK7Ne0TSuo8r7TuzTOnkyulgjlyPVJ1RTpFJJASBHHJMsM1JOlyrU0qm1mBLAkYtzTuqc9ipWyjUcGWVGY0bIn1nK6uWWlroWXdHUwx8PEGFw0bAGNlYC1gMb1dZauirY4NOxZs1fKHlqhHNKsMUINh38kcbMxcnwJy9gSSoxJDaDAaZkb8OW1I+MSCoXHlX2l/Re7Ycxml7bdA5ho/VM0hSu1BoiWGKCukBIaaegcNT7yeWaEoGPIAvfG8k0B9EuHLvqx+nL2hHLbf/GxdOncV/Z/887v52tjpKtyLJNfU9fT6x0hpCvzWZYzDN/Z8ZjLTkP8ArGlEst28P2fFGLm7XAthhk8XZbpGoK6c0dpZ5KaTazLp7L1cm/G4Rwbgb+Qk3Xta9741Dpk07ta7MDwJCoDR4r3gOY9yqH0pnXYtovPKPT/0a9FZ/X6nqIe8ptR6gy963NHHXdl9PFEYqU3BIlRJH8g68nFjV/0ftT9pRqtaa30jNlmf55XmqqhkNZRzV0m+xeWqyl5g0UkjXa9PKXvy8K3w/wC0X6QMeXpnNfR1CtV16Bao0xERqQqgLGzGzFFHCwi0QuLIBdjQ2pO02LP8xosr0bqHOphMY/rWYxq9GKqVrbYaaG+9YlJA3uQ8rAC0cYALI5JrQe0ofEmvoB1yc9kUQuYeAHuSrK1foTTVChTU2m8z0/WULxQiuyumeGsYs22DfRzIiTzsQADE0LuwNt+1mxIdMdnGV6Ovn1ZQwHO6hQlXNNTR2paUHd3cjwKAZWIXvJlGy/ClVUs2x0rTZnksa5nq/OM71DV0h7ykWor/AKzS0UpjEYkjN9xkKkJ3h2gKxMIs25lc1q8jfLqfMmzg5UlFK0n1pa/6vGkpS4Yux2xkhTYOVRiAFZWG3Gc+0OcOzFSNZ9tvW7G22NoN85qA9oulst1FNR6cYz04zqpiheSVQbUxbvZ5lkUWdEgSY3UsPF15tidxZNQau1NT1FPSwVgqKto4sozFTETKUZoaXvrholdliQoS8TBjsKm4DLINLZ5UZrJrDM6ioSpzDLu6yihqoBFI9EzK1RmElM1lSSp2pGikK8kMbu321vI66mpVoR9bcwliGVg7bFsQbo99y2sCATuFlKs3GIZ5bgEWzPxNK+niDVSRtv1ft9lyxnn0bPpU9r+e5jr3X2nDlsk6uTLm1VFT7tgOylpKVC0zgW2RxxRlRxe3JxPuyXsV1h2D5dUZnnmVmo1/qV48nynT8Lh3jmuJRTTODtD3VJaixtBDHZ2DyBR0dp3X3aAKyvhymvymE08bDMM2rMujR45doYwM9PGs9VUd2Q7Rqy7FKvNJGWAIUzUGX5pLXVFVNW5rUQCGpzOsijRjSltxpoIYrx0tIT4mijLPKfFLLIRbGpa9LiWAQ3LrcMNoGQ3DnTLNULPYCyS+HVdjjs37z75q1vo00eW9k+j8m0zTZ79bqIQZ6nMr7fr9bUOZZqgE/wDFkZtl+qhF62GOt8p1HlWqaNcvrVjMsihlAawe33kI5Vh6cEWxwtklZTxUz5dWCnkRi6FXG+OQHl0dR0BBuy8dSy/7tsWh2fasqctmGX1FVPJErAxPJLue1rgM/m62tu+8AG6lr5lj0jJA43sQ7Ouuqu2mxNlaLuBbko79Lz6Gundf5VU5nRU0dBmlMktRQZnDAT5F3jdEHI6s0aizjc8YWRSsnlNnuUZhoTP6rK86oBDmGWTbHgYh0EgAZXuOJIyCrqRw6lSODj6B8j1NFqbLocqqVilqTIgkYybGaLqJo/8AvFYKbevI648sf9I12JUeR6g/tTpuONINs00TRC8a0yyAVNOdoPhgeVJ4/SGeROkKjHV2G0RR3Yg6rXZbto6/JxLTHI+riKObnv39fgcSzLnesc1Wioop67Max2Z7vcyNyzyu7GwHVmZjYckkDGyizHLdFwPlmlaxZ6+WMxV+exkqdrCzQ0ZNikVuGm4eQXA2J9oI51y2CXK8qV3onA+uzSLskrwpvtb9iG/Kx+Zsz3Ngurr3+qVSmnlDldlVE7xoUK/aUsCOelmB44IxsXe0peHd1D5+OOwZ5NzLPWfj56LComVAFiSPu7dbjaAOLWH8ePjgP19TZ52LCIWEjN9keQHp8AP44kCZXlhNPn1LqOKFKhGrIZBRkCOp3Xam2qbgq9xdQy2aI2sTY1JlOQZNQ02earqVkFTEJqTJstqVNTMlyAaiXkUsZIPBvKfJVB3Yf2zU3sytPk+RV2ctOaaKIQ0w7ypqahhHBSqfvSSHhL+S8knopODy1WU5RIYMhL1FQps+ZTR7CD/3EZ/2Y/vtdz5BMK5xqeu1HHDRvT09HQUtxSZXRR93TQE/spcl3PnI5aRvNvTXQiGxaokcRqNq38QLfsrYbmt7cD1wl1z8ZMtg99vhl45pbwbgzj1+vggpszzPLpKtsqzOppjXxmmqTHIQJkJuUdvcjzPzweOkVqJg7hpFbYI2Bsq9b/C9wQOnXzw1qbzFk3+H7lwAoHwHA/rnE57LdFZVqGtlzDVkmY0eRUYYyfo/YslQ6gM6rLKDHCiId8krBtilQFZ5FUve9sbS92Ca1peboxUMZlknMks3L8s0rF2J8zcm5P7/AGweeHulUys8bbbhWH6xr9CF+6LebH4DyxZ+YdoHZ3kOn8x0X2cZU8qZ53kE1fm6RGpkhZ/C0rBbblAXu4owFQkyMXk2iOGQ6Nzk1cb17xUlJK+6SpqFZFsepHe93uPzxC20VxcLo1VzPlqTzFTAYqPRNI90SUICLbQxB+BPW3wsPbD2nq5aRCUCE7DHYgWClSGUqOtwSLfP0tLsh7OoNU6hpdK6d1XkMtZXyNDTrXZpDSRswVmAMrDukJ2kC78kgA8432V6JyvRUUc2p6STNMwzKgqKafKkggkgSOQgKyVm9lSZV2n9WHkjbg8NhHWuJowz2a+t+SUQvJVU1JK+NZ1LoQDc3J44+II4OCnL6vMAO5SR5Z5FQRLE7Fixt5CwAv0OLS1FlOWS5fTZplemKDT9XFDD+jaqi+sGog7m7IVIa07SFrtOwNjYXXgCS9lP0qvpDZbV1GmB2hT1WX0VLLNUR5jDFXUjwKLFHilUkh9wUbWU3YYhktzgwvY3LOp/BqpWWUOcGuOe5QXtw0vmOnc209XV0SRLPp7KqeSRAPFLHSoFZvLc0YQ38yDiupApCuUBjPG8A2v8/wB2Os4u2TtKp6DUeYdms+W6anp8qXMqymy6iiad6KkjWJY6Z6kSsI4VtdUsVBLEsORzVrTtT7S+1CaCfWmsM1zSKlYy00dVUmRYm82A4F7cXAHHGItGWt9oZdoMMzXHhT3UlugbC+oOepazMtN0dPl8OZVOdUImcN/qkbd5MvTZwvh5vyb2W3meMReaMoxKggY3kjiVdoVUtZyPO/p6EA/PnCVbS031dXE7GZ5Cpi7uwVABYlvMk34HQDr5Y02gtzNVROOQTTK84loZbyIk0TW7yOQEo4HS9uQR5MORi9NE9pMWbanyDPs6rEnio8xpVzWWZQaySlZWpZDPIP8AbKIJnXvbbiAocnaDigpaUr9kc9fljKStqaGZZaeVo2W9ip5F+D8iOCOhxWtVkFoYWg0qp4ZzEQSKqSZ9kNVpXUGa6ZrQRPlNXPQSsOCWicpcfG17Hgg41+XDfl+cZaRcmFKlPjE4v/7Fm/DG5zvUUusMxbUFXsFbUwRLVkDwyzIgTvOeAzKqk36m/rjVZW0cOf0scl9lQWppCRxaRSh/9tg/eOswdIO8ACfEY+oSd1sxDcj6Fdb/AEENWd4+tNKStda3JKHUCAmx73L6oQzlT5P9XqXI/wAOOsqCqrtNQ02cxxSyTUTCti7lO8cNAVcjaOSAU4IvwLW6HHnh9ELUx0l266PWtqGp6apzZ9P1zXHhp8wjalcm4I4dlYXBFwMeimXAuY6TNEjaWn3Uk6sgW7o53ArcC7WYcWN7FQQSBwP7SRmG1B7ciPT8UXUaHf2kJadXv+arhn6W+k6TSnatrLT9EpWjy7UtRVUdl4FDXAVcNv7oWqsP8OKEaJo3aK4YDjwjwhrdAOhPv8cdm/Tb0nJVZlpPVscMiDUempMkqS/X69lEpiQn3NLLTt/y44yeUOgcKrHgkFFB5+HJ+eO50dMJ477cjR3+IV9arm7ZGY3hp1VHA09KJvUEFQ6sX3AkgD7PqOeL/D9+Cq1izAugWx8JItzwbDr7+fphVnQtukdmAIv4uWHsbfmcE2shCurKRYi/htzwfUjGiqiUbvFBEgKBx94E7j638/W+BbvLl3jKi/2mUhj8/O/WxvgHjCC9zYk3HP42Pr8/PGEAMpYLGWUEts+75HrY/H0wIRJAEUq1lFgb/aAv0HHT5dDxhJ2UkgFdoPHPAPqPIH8/zGFWCKAAoXZxb7y39x0+FrYI0UjMOWDk7QjHk3HFrW/zwIRSqvtN734Hlc9fgMYe73jYw2qeu6wt/n1v64UKsE8RVrW5v0v19L9PywDIEBlmJBuCTcXN+ntz19BzgQiNZR/s2Ck2sWNz7X9cELMbD9ZtToOeLm56cX8/fB5AQwG5Qeh46n4Di/wtgjEEEXbbcHxCwHHPHW/Xn08sIhBtNiwF7G3QgA+/pjMCQgRQzAWAJ9Qbfu/njMIUJFgxJYr4eLm3Xy5/q3rjC2w2BKmxJHqMBuuxbd52ufL2+GBAuuwc35HF+f3/AIeuHIRQzBl7tWUt4Qb3v7EfywJHd3B8IJv1sPa98AOgPW/I/l64HbYGxRCPj/D+OBCMCAjKFBa9725+AwJ/ajDnceB94H1t5YAErfeLAi56H36H+GMV3BVegF/Dew+HzwiEBkYtuFiWHJHB/r4YJweAGIHRQOT8PTCm0iwZPEOQD0v5HBJCv2UZQQb23c+wv7YVCzaOm9eRydw/DCTGw5JsR69PlgWY7ipt/LASEngL159cCEk5Cjm3N7nofbG501HJFlub5hHGzSPElBAAOWklbkD32qR88aOQ28+vTFrdlOV0zag0hS1MReCmmm1JXqPOOBS0an492ot/fxTtsgjjqegMfanmp7Oy+7rXh7q5Y8vTTdEMjhnESZLBDl8bCIvuanjAlkuAQq9+81yQRx5YQzLNarIuzXXGbQ1LRHMqKj0tThFEa97mU/ez7VFx/wCaUbX28Wl98LCtqJ6RqetQmpgpe/mVmADO0gvJGwO7dvcjcOQSQwtziNdquYR5foPR2n42a1X+k9Y1QNvF30n1LLwR920NM7gdP1ptjmLG0ulMjtWPnq4rbtNGsDBrw8ta5+1BUCqrpZF6SOQB6C9h+S4YN4lHS3qAOfwweZ+9qmcHgdP3D8hgjgm6t1vY47CNtxgaNS557rzi4puA78Em1zYeWHlDRLPOqyMUiRWeVh1CD09zwB7kYTjjHNwRb8P8sbClFPAJBUo77o/AY5ApVr3F7jkeowOJAwQ3PFNp6XdKzRRhVflFHIHt8sOq/JM1yGKhfMqJqdcypErqXeR+sgcsFcAdASrWv6XwtkVJVZ7nNFklEqmetnSJP2UJPLH0AF2PsMXx9KTRtFSae7KNUZQrmhzTs6ycxORYP3E9TSsfidkTH/HiN0wa8R9bBxUjYi5hf1tXOhcSFrWvu5A9LcWxtNLU1DUZxBJXmQrTkypDEu6SeQfYjUdOvJJ4ABPPTGseL6tLdiGHB8JuMPtN5n+h81hrJKmqpogx7ySkVTKq2P2d3HXqLi4uMOlaXxlo1pkbg1wK3Ga1uez95URun1ZFYCiUExhL+Lr9tiSCWPiJ5HHGHuYatXVVHlMlRJMuZU2XxZbWMQA0op7pAwJ4YmApGQbcwj1w3eOKSBo5pxHuiTYGX9XKGIKeI22qV8+RZOQfJakosny2emhzjPKeOpkgaqjhlLNFHbmFZnUEguRusBwAt/tWFURRxgFooRs6xU/aPfgTgdqYzUlTTVP1CcQiWWRHSRXXYSRt/wBsG2bdwte9lYEG1sM590czyvOWnZiXdZDvDDz3A8m9wRzxzzh5mGp4J66qFZSzV9G8zyxu0xhm3uB3kgIBAEhG4ptsOLWPOF6WpyHUEsdZmmafUWAvVpsDyB1WwlQMyhw3BYBgwYEgG4xMC5oq4KOgOAK031iraY1ZlJeXxliPtX8uliD+GHsSSRQo0I5LN3bq4VlAG4kg9QvNmB45HOHGe5BU5NSUrxzzVdDVb5qepTcKWoUGxaK9iGU8OjAMjdca9TGrxkyKZBZgUXcFI9SbDjzC++HghwqEwgg0KcUKqlWq1NRDFErmzu+2NiRcfrCCLG32rci4HOHdO+X1VZTZNDUNXShFkheihI8UnjeNi4DXRifH6XAuMIyx0894oKaV5CfBGl22M33CD9297MvkbEYcSZ1mLGirqyaKaaGnOXhgNsppo2HhY+w/VgnnbweAMIanEJRTWi/WTOVmrEDyxliWXxRype4IHruHlbdY3seuuXe8pkbaxB3bgtzzzcjyPv5j3w4eeASVENMzLBJM7J3os6qWNiw6BipUNa4PB8hh3LIVyqWFkVhCTIN0YYq17MoYWZd1xxexNza4wYjJC2uSzCtpPqDUFPUASiWIMpjKtuDSIZVIKIyqxbcCF5ZdpvfT6s1Hl65mF0Q1RR0sEryLU96xmkZjcL3hAZo0HhQsAxF2YAtYbfIcunzfTOexacTvcyMSqlKapWmmpzzMadNgaVlRfEgNwrM1jZrRXLPqH1eSKrytarvCGSVahonj+FgVYHrYj54jYAXE7NXXWaeTQDfrW3yPVNSKGWgra6Oijijq6tmND9bbM6pwuyKoEh27eoD28PJ5bnCea5LnmqNS5vm+b6gyFEp5YoqjMUqL0jyMlolQoCzXEdt1rDaSxHmyzN8qhqkGVJOkDQREpUSBnSYL+sG4AAi9yLC1iBiS5Hl9C+i80q6uKfvq6vp4KTwNtlEQLSMhXm6d4A1/Ce8VR4sK4NZ+8AoTu2pAXON0moCzTkOp+zGWWu1NltbDk2Ys9IKimKTxfWlRZUdSrgbwrKQwYMFkLDcLqZOmqtOZXp1c0SqjqoxT1EVFSzRh2DyrtnoKmMMGFPIr96si+FZIyV27rLHdD0swrK/ToykZpllfTSGehhJaWodEYxfVxYEVCudyWG+yyAgrcGPZFo3NNS5NmOZx1MdM9HTySUkDxsZMwkiAeaOIAcFI9zknjjb9ogYidGx5Jfu8/j4809rnNFGrTPRVNJBDVCJlSZTsZluJLcG9+vlcdRcHzBxr3lYksSbg884lendSRS5ZPpnUaJLSTxhaeqKFzTuB+rYhfEQt+CviCkizKduNNn+WVGXSJDVooLIDFLGwaOWPyZGHDL7+XQ26Yth+N12ahLaCoR8zrafM6E1rtEaid0UoB4lIXxt7XNvxPpjQPCVYAEc9MGibupCjKzBvIdb+WOlewvsy/s1ldD2rZ/ldNPqDOd0mjctqo1eGmhjbZJnVRG1w0Ubgx0yMCJZlZiCkRDRSSssjC95wUjGOtDw1oxWu0B2L1ej0o6rPMqas1nXrE9Fk5g71srEljC08Vj3lW4KtHTkHYGV3BYrGJtPT5F2Tl83zr6tnWrJS0iJJsrKahk3eJ7NdKucN1kk3QIwIVahwSm4znVI7N6qtyKgzL6/qOp70ZhWuS01KzjdNDJJfc07lmaUA77NsZwO9Q1PqGRqqhzDMat5JZlU1Hek2L7Rba9rbQBYAqLKBYAAY5xsz7TJ20mvLw1LXdEyFvZsxopr2UaLzj6Q2u8wznXGbVUekcl25nqiumqyhmuSIqY1D8maYqV3sT3cayOAAijFy9s2r8qqKGJa3MMr0bpOWjjo6aKcvS/XaOMARxR0ygzNSqAO7pY1KkWeZnd2AoTMfpEVek9E5Z2RdhGUJTU2Xs9bmmpaymVqvMcylVRLUxxMClOiqBFEWDSLGoIKFjitsv03nup9StU6hzSXMMzqVeqq6qvneRo4kXdLPUSNdgiryb89ABdlGCWzOtDg6Q3WDED+Y7zs3VqQMwDVMjmEQIAq457PDerBzjtp0Np/Lky/QekM0zuSmLXr82n+qQbSxIVKeEmTYCxCh5RYNbaOBjoHsp7PNbU2UTap7dM/TI5c0p0jotD5XG1C7xP40XM5or1KIVIb6okgmZSDI0Sm5iugJuwP6PdPS667S8+jl1ehWoyTTdJQpmFdlS2uldVQ3EEda97xQyvtplIZkeU8N6j6TOf6lrjV9nHYbX1ZzN5TDmOps1mqZqzktIywU4iQre7OQWXqXY9cVbSxzI6WePHW5xHIuwr/6jBTRPDn1kdhsH491dWp8zq58ijyCnyL9B6Vyyn+szR5PlizmkcNZTHSRvFT08Fus7949772tYmA1OqRmGYpS5RPHIh27DJUmqndhyGWOnCIlrAjazkHkG+Kgh7dfpJ59qanyPSeYZdQ5rXO0FLR6dyWlWTp4wJnD7UCi7MX2BBuZgBixcqp9VZdRSz6v7UM+1NmUpmppVy7MZMsy2nlWwdENMIpqlkJ+2THGWtZXUhjnusb4gH2gtqdhJPp70VxloDjdhB4AD1Uvrtc6zoRTUHaNqeSPKM0UpDFn9M+Yd/x4Fiy+oWaoqN5sEsEUk8yKOcNqzMdK0eVUnd6NyGLM6avWbu5g0cyQhSGjqJKORFpRyLUlKZGUi80pI241NDkRq6ALqt6iaYoB4axcsNckIG3fKRJW3tyyoAFJNnCnDLIokyijOaU3ZXoHKYu8MYi/smM2qIyOSJKnMqnuw1rEksOT0OFY9j23AcePuKcEOa5pvU9vYraZr2uVVdQVuUZXDkMNZPS/VauHMc3bOsragLhpFqKGemafbcCz9+oDWJPlglFVHLIqan05o7s4yGuEnfVVZlulnpxNFa20QS1DwspuD3jgBbcA84Rk7Re0afM1pNHZuaCmCoktPp2PLIjKJG2qs1PQIKd0diE2zzlPW/TGwzXLInzyoidI8upIe5ibTunczkjy5KsD9YzzyhiG3cGOmVIla4DNha9kzsxQDZnXrfySNF93aHE7eqI2qM9zStyZJcwz+fNKWkAqViWJI6SLaf8AbHu1SkhA6B7SE3AXnFOZ9qvVPa1m7w5VT6v1m0UUscMWTw1VaFmK7Y1FhIEQMdzE+Ngu0bb7hI9T6s05pXtd0XlHaDli5lpOoyv9PPl2Z1DimrqwzSRpHI7G0ixCNlVWOzfcNYMcSnNaL6R+oZZ49L/Tj0TTaZlkL0VJTz1OW9zET4YzQU1KViYCwKruUW4ZhzixZWsidemIGFQXG6KEkYGhqRTHKldeqK0Pe7uxgnGhoKnhgq7yL6K/0gc8MVPn+m6fTNNUqWkqdWZ7S5a3d+aiOWRp9zXIN04BYDkk4sXIfoetQ11NmOedsGnaZ4Vdu60/ktfnLXKlSDKy08N7EgbXPXGib6Pva9VOamt+mdlMch5Z4VzlAf8AmFOt8avMuwXXiDbWfTSymdrfZkq87F//AKHi+61MmFGTRDzJ+PRVRA9hq6N58gPldDUf0duyeKjT/wB6PtDzGaKPdtpstoKJAgAvstDPILA/YQlgLmxscbSj0vojT08FZkulJaiuomLUldn9ZJmVTAeLtCJUSmhc2FpFhY3t4x1HJs/0b+0yqp2ky76SelK0DnxZtncfI6ctAVHzwxo+yT6TuUu0ek+2/Isyl5Ip6LtHSKQt6qlRLEb4q/RPtOEdpYTuoOYNVN27Y8Xwupvr+i67zDLoKzMYMzOaLM+bd7TiuecljM/ijeRmJYOJkUNu5BJvcWwwznL53qo4XrJcrWmWnqK2qpkVqiGeaMyxUNIrgotQY7zSTyBkpo3j2o8roF5mlyn/AEh+nadlk0/r7OaVvE1stgzyJ7DrdFmDceYONNN9Jbt40jmlce1js7iqBmFb9fnjzfKKjKpUnMMcLtG6bNm5IYrqVK3UEAG+GN0DaIauBa8jIV1+eHsnHSkUtGmrRrXT0T93FHS0gjoKCnR4qWnppCoortuvHIxLEs772kcl3kO9y25gHdBXOzU9VeKomp6do3Mf6pahVIMi7esMgKhwtjscecb4oOh+mB2Z5zGlPnOntSZEzOu54np8zpwLnd07mUCzMOC3U8HEyyHtB0TqN1l0f2g6ezOoZlUUk9acsrCAfDaOpCDev3drNY2tdTtxQlsNpjBdMwjnzFRzVyO1QOIbG4K1UrRU1dQII5g1KIlVYlCNJASSjLfwkq4YbG4uw2lQ18SXJtYww2jLM8TtHBvUFVWYndEjX5j71d/dP0LoiX3AjFNVmp87yecrm+S1VJUuTTtDLE0QqCwurwE2Vw6gsig/bR4rghCWba9iYJPS2lbmN4S5UTRsbvC5NvA9hZvuSokgselTsHOVjtABguvtM9ozpMYspro3nVQadr7SZVs0Zb08Wy/tJccXCx76QcumtYZLmVQgE1LUZfBrHLqbur7oGjYVsN+gDU0lbGw9UXzGKS0zq51zIVkVWZ1O2qjZhteenlLbHt5OHEikfdmSROjjEv1HqmngzHs1inkhmocwq8207IEbwtTzzI8a+1o6+QbT0t7Ykga+OQDz4Z8lHLRzb3l1wXnHrLK303nlfpuSSR5MqrZqJ2LXWQxOVRwPVkAb4thjQ6UzfVMBNOYqahy8H61X1cvdUlNGx3KJXI5a99qAF2vZVONp2pTin7Qcx+vLNITHTNMiMFMkggRHu5vtBZCSQCeePXGnrNSV2aUVPSVM6xQUbEU1DDHsp4AerKL8ufvM13bza1hj0WJz5ImlusZ9fp6LkpA1kjq8E4gz7KNLNJlujzUT1Csd+cVMfdygsLMaWI37gEC283lI806Y0kio8LGOEBlO47btdSeS3Nl+J5a/zw2rUcMkisp3cEK4DWv0Ppc9PyxZWWdkGaQ5NHm3aRnEGj6Ax/WIMvNHJVZlLGf96KNSpjQ+UtTJEGv4SRh5McGLjnxPWwJgvymgUDoqkQiSGnETGS3eM+11KjoArDbxc8m59LYxmnq6pI4ou+nltDHFHBu3HyVEA6+gXnFmZn2IVWV14zTIZNQ5jpUZUmZT5oNPTyNETy8KCPwSFRtJlVhENxBY7STB31dHRyPBo2lbK6eRe7kqhMJK6dT1DzLbu1P/AA4to9S2GtnbL/BFTwp4/HonGIs/iYBJ5tkD5CBTZu6RZkxVVyyCzzoT0EtiViJJHhJL88qMT7tVq4tLdmuR6Dmo5oqlGWgkny2dmopHp2MlYXksBPMZ5wpjHhj7pSd25CK8yrJ821G1ZR5NDRyNl1FUZlKstVFCRFCA0mzewDyAG4Rbu3NgbY3svaDrXVGia7Rc2SxZwjiPMJ6v6jJLVRLHbZOzJxuAITv2G7YxQsy7QGyRl72kmt01OrwQxwAIApVaug1PmscVNkmicspMokqXWnE9NDvrqh3IUbqhgXHJHEe0e2JbrXspjbOCmls6Wr+oZGlbmdTnmZRQvPVxztT1PcF+WUygFIzdwhucP6DsWfLc40zW0uf5tE0gpamtoGpUo84pJXjDqaeNpNk/VWCCRZ9rC8YLKcbY5TNSZzU5BU5rl5n+szZfJXR1TPAstVCskTgkd4P9Yp1uCN6s+xhuxVktLGuDoDtJw6Pgp2QOc2knl1kmK6YzHsxpJo6DLoKmepKxVGdQosxWCTYO4swdKZhMvhewMym6uYyRjWZhO1YKutp4kmlWplml+pRJDTJI7bnFxaFADwvdqCObcWxr5DQ08FTRUWfCWlqwwn+pCanUbW3rFJG4QuBIoO1gQrWZSDxhrDmtVAtJBV0Rru5DRhNx/UyK3hKqBYh736bm5ANxhgaXd52J4JxIb3Rkn+UzRVM7RVrCQUtOXVBUFYBD3l5SxFi4UtvZY9qtuJvxy3inqaenzuuyfK3qZGZBLK0Ygp4aeIlo0J4BZpCDsXk7UHUnBa2lmp8vRcxjglNXHFUrtO5Y0JJUrYkOw+9v+wDsIvyHJRYaNoq+cITUCIJJPuAkW4B23JIXg7wCVuu7eCCElY142j4NU9ji0rS5Pq/WWgM8oNRZo36TpDMlWxQ3CB1McsJ4GwPEzxMh8JB46A40Wb0+WUNZImWl5Mv75xQ1BQgTwE3S56bgpAYdeDib569NFTTMmafUz3L0syo/eSygqQsIhI2uXa7d8GKgXttIsYOtbVU1PBQVNR9agijCNT8iAL1C8feuSd48Q9eLYnsrWtJkaKHI0wrs63qCck0YTVMm8TbGn8rfa8vQL0t7YfUWnszrcqr86pqV3o8u7j65KpssXfOY4rgn7zKRx088FNIJXUUaSiOQ7dsiqXVvMEjgi3mLX9Ab46F+j12WT667G+2aOIr9ZqMipp8liJ8VVW0NQasRIOrM0FPVkAdQMXJZ2xAE6yBxUDIjISNgquZpYdzAKvJIUACwNza3xvhnVZfLHO0EsbRyRuY3VxYqwNiD6Wxsq2mXvFLB5I3AaNQbBlPQ/MfPByWlmY1Ew7xgCXN28Vuh8+gHPriapJUdKIlND9VpVQhrDqAftHz4vz8T5Y19fuUrLHJyp3Iy9cbFmiA8HO0+SgXHra9rX8hyODc4a1kYdSmxgebn7xI63+Hp5YELcJW1VDqWXMctkaKaqhizOkdeqzrtlUj3Dq2PU/L9SZRqotrSjq1p8szjL6bUbLURd7AIqqJJJNjjxKqzGRWQ7gDew4Ix5OtUucpyrMV+1l8z0rnz2k71H4Fxj0I+ipqCfUXYNk+WwCSafTmb1+QSojDesbMtbSPYkFgBJMFAN7rxfHFftJZ71mbJ/Sae3sOK6LRE1Jy3+ofn5U9+kTk1DnvYPm1QlPLFUaFzrLdWwCQqySZbOPqNa0Tr4WQCSJzbgAA480M9oXynNq/JS6SpTVUkautyGUNdSBcAixHN/hj1g03Q5bqUVHZvWLAMu1ZQ1+lp4lYRiI1sDKh7s2KHv1iPC2uARjy81pQVEMlBPXwlKtqZqKsUm1qqmYwyBv8Aw+Xpiz+zM5MLWnVVv+4ciQotMxUeSNx9j6BRMOVHW20dQCAT8P3g4ObrwzEMG5Dcgt16ev7/ACwohRFa5Qs3Fz0/5h59OPzwm57zwqhAK8Hkg/M8fPHWrCQneBtKyedtykG/Xoeovz7H2wCuLKQu1Twdpt08+RY+/wDLBCwsQTz1DL1J9/P58fDBQVUBA7Mv2gAOL9OAevH9cYEJYyCxUBwCLRi26/ysLm3T09Dgl9gMMajxW4B45HT3uPcDAd9sJCso9dlgW+HF/wB18YHiUCzpwbHbZlA8wQQbC/FuowiEG4lQ195Ym/F7DzNv4YFwrPZlChQAo629b+58zax6YO4W+1DGqoBexA23F7X+Hlzgr8EBgdqeIL0J+YB/ywqEnYswWQnjwi4Y29AByRgjF2BG8AIbi4Hn5+58vXCrcqqk2Y8NyfF5j06eQGE2AezWYAdb838iTfpf08sIhJBbsNxsOQNxuR/X8cZjPASwtcHnqQeOL3/LGYEJuSbkAcgXN/b2+HlhRbkWLsR1PHyvb+r4I1wxF+Cb8nnA3I4sSfO3r6Adfn5YVCEBi4UE+vB5N/O3XnGXRdouF6qOR+GBBYgC3itzzzf4euMFlNva5F7WPoePh88IhG5sOB03D1/fgBYeG4PpwCPjx5fHAKx2887eSAObeZHpgQ125BYi1wo9v4eWBCEAMQqi5II56X8v+uMLCw2twfTz+Xn8MEJtcE28unX4HGC7Gw2cnknj4/jhUIniQ/atb0+78MEKgXBPNuCPMWwsoXow2nyvzb3wjJsB4uD8MCEFNSyZlW0+X06ky1MqxIPVmIA/fi8+z7K6aSs1FmdPICkUtPklCDIEDRRESznceB4Y4hzx+s5xVWhKe2bVGcuNqZTTPU3PQOfAn/smv8sX9oWBcl0lllPA8cM706y1BkRtzT1UglIDA8WTuEIIINjjB0tOQCxu4ccT6DitTR8QJvO8fYep4LbmCrzJ1y+HLDLmtVKKSiTcq2qqhlijaxBI8cikhTze97WIrH6QeoqDNddZ6uRy7sqoaiLIMrsbg0GWxLSxMLeTtG8nxc4trSOb0eSzag1w07D+wmT1eewoIj3b1LMKfLUDjp/rdSG224EJsbGw5YzwzileTvQY6V46IFiS0km0s7fj1+OIdGxHN2sjl+cVJbXguoNQ9fwtNEpbe4HBNr+g6DCsUQmlSAWUsSB6BreG/wA7fjhJQURR6AA2N/fB95QEXsevHqOhx0ZywWMDjiiyStGA9gfRSL8/DFy9n2f/AEfc8ykZZrbs3rKLMkUD9KUeoKoU/u09OEkdV/vxbreaeeKZkVnBd2uW5JPmTzgMtmqqTMoJqFmWdHBj29SfIfPp88QWmIys7riDuNFNA8MdiKhXxl+g8nXUTVGlqGkp4M0zBdPZY9HmkmYR1CmSNampjmdEJUiVIgQo5L+YOOg+3DIqTM/oo9h2ss6oHzSg0xWVOn8xpYJxSOcnrp5ZMvHehW22akK95tPPW9ziq9F93kuoYKyZQ9HoLKa2tkI6NNR08k0r/F62VAPguL9oqY6w+idnXZ4HdnpuyLLMzWPuiCtZQgZhEyN947HlBt0ucYLZ3ANkf/MQPI1p7HxWm5gBLG/ygn59/Jc063y/6MeicsK1nZ1qqXO5Yyy0C65RxTki6me1GNl/JN28jkhRzigaypo8xzB2yjKnpIpWVYqVZXnKnpYMRuYk4V1VQyQ59WU4neSCKUtEWa91YBgSfMkMCT1OCZLn2a6cn+u5LMKSqUr3dYqDvojz/s2P2b3sTa5HGNiyQvjjvlxc4jWTThjT1VC0SMc+6BQDYFKqfIv7IZRJX6irGQ5rTvAlHDKGdSrg7WUggsCOvSJrXu3gxqc21FBmuVUeV0WXwQfV5jU94IVUxMwsYoybkx/ebcTuck2A66ABiASWYW5JufP18sKI0Y8INr88LcfzxYbFjVxqVCX6hkl5qiidRDV0ZjcdXpnC3+MbXH4EYzL6bLo5451zCOWI3SWKWIxybGFjbqpIBvwb8YdxajzSGlWjZ4qqnQgiKqp0mUW6AbhuAPsR+WEssWsTPI8wyHLTMsdQGjhnCyKAfuyWsNpBIJNhY9cIQW1rh5/OSdgaa/JSWNK2m0tUirytakQTxJFVLVqxonDNE/ewX8IlVUAYgXMancQbY0ctU8AK0iyxwThBVwyPuieVLkEHqOtwOq3NiRbDnJc4klmzeiryjNVlTJNGSyqiuN4Ur9pSFS3r3YtzhrVI8eYTIjgAvePad6kD7JFx4xaxuRz52OGMFCQeskOyBCewrTU0k5gkBDQ3UTqklrjxrY8bx5N5gG1icFnrI4+5mihubeGN/HArm43Bid1uvhPn5kC2GsUrd7CUjDztL9sncjs3Gwr0ufPpxxxa+F4Ejrnp6OmcrEJm20xcAxSNx4XIuykqAN32Seb3vh9KJic1ry1YgpUQ1ApoFijkWMiVk692/mTH4lDentayE09XXNVTh5pKqVu8mJdWWccmQuCbk3AN1JseffDc2ddtFSuyBGSQyITfz2kA2WxU25ve9utsKxtLFTs0cH+1gB2OVKDcpCnkci26x8j53GACgolOK2Wma1sizKn1DSrT15p0BiU7tpeVGDI72BRkVnN15DbSpN8SKLK9A6ho/qGSUB01mwkvSCqzZp6OrBAtA7yqpp5PNHP6tr7XKmzYi+WIKVXpaephdO8DjvmeJtxjHFvVeQbE+tiDglX9Zjekjdlp1ljdlacKojIYqym4O0Ajji3INh1xG5vevA0PWpOBwockfUujc+01JQtqTLKvKaavZu5nmiVkkVGAkZCrFXK3BsD6eoOJFRaY1ZlumK7JqvJayRqmthqcljnonapkis5lljg3bo45FERZjwSFAuecaDUMVTNojKDQPM1JQ11WlTT2FkqZ9rJJx03xpst+1CwGA1UaSp1DX1UEDw0/e92sTOXMCqqgREkkjaeBc28gQcNJdIACRXHVsOGtKAGGvWKb1NLm4ZpKqOqp2oqiMVKSM8UlNMeEaRTyh8lcjj1vbG/0/qDNKLUNJmuoaqomenKySSNAJJIYwSDIqqRyu4ydQWYEm+7CsLy6lgyal1DqasipYL0ffy0oqDSQM12uVbvZEVluEN9ouV4Fsa+KkM8ncUH1yWenJSZF2TKTcqD3bFJEFjxZja46c4HODhdd15pwaWmoUa1BlNfpDOqjJM0jlD072jkaJoxNGQCkqhgCVdSrD2YY17V860z07fraV33903Khv2lPVW9x1874mub53qY5FW6Yz2qqjksZieNKtmm7p4QRHHTqzMYSd21gGIVL38sabR2h8z7StYZRojQ0bz12dSJAi1LLHHCwBaV5H+yIo0VpGkNtqKxPTEzH3m1f1vUThdd3VL+wPsgo9f5pVaq1XBVLo7Tzxmv7pxHLmNS4LQ5dA54WSUI5eTpFCkkh6KD0DqbXGZZbmtdXUbRvqrNBDTgUVMVjyVFjCU8NLFyVeOECOnhAvCgEjgyuArXNa/TeitFU+TaOzeny3SOnDJS5fnVdEf8AW6uTa1RmRgHinqpyqmKnF+7hSn7woFYvz/mPa3mFOZMq0KMwy6GrLJLmM0gbNswLk7ryL/sVcm7LGbtc73fHPWh0ukZSWCrRgK4DxPwN1cMVrwiOxxi99xz2+AWy1JNQaFqXpc8zO1csiumVUTiWqicEkd6/Kwm5J2sWkG43UEnD7QOf12czV+p83y6kp8pycrFSUATvFqa9wTEJma5kSJQ0pX7N1QFbE4rg5dDSQSVDqDKF3O6m9r9Qh/HxdSenHW/dMaFrKTKKfRQySqo209ClRnVdVQgQ0lfVKJJAyHxPKqCKNIyLeDe/hABS09nFDR5q45nLDXQcBrdjgUkV98ndFB1Sp/QYLWaQ05Fq6sky7JaOaSvmnAAhXfI7sSLseu5iG8Xs3QAstvaV+j3pk9m1TqnVGpJMqyPNFNdX1cVTHSiSiikIo4nrJLpEksqtUsFV5HApgqtxiH6qhj7P8lrdBaUiGXVtQDFn1SQ0swVgN2XbgQzSSCwqnBB5FMm1UlvD9UZh2gdree01VrnVcVTLHUR0tFRsFp8tyWmMTs08FMlo1jiihk6XbwKCSWGKt1zxeY+6Ofx51w3qeobgW1PXW/cpvlMfZFFk9dnHZ/2b0uYGKo+oUebV9O7009UAGleOGYmWdY1KkvIUMjui7FXcMafMo82zOso8vyyhr84z3UMi5cqRTlqjMe8ICU4tZdlwLIQIltfaoUnFj5jnGgNH6LyTJ6ulNHR0tIGy7LzH3tfVwtdl/U3HLbjLLNIVQSSeEsE2nedhlNlceWN2vZlpowZnrE1FHpXJnqnklp8puYqiqeZVD7qqRXhDIu7uUmWMAS71qRyXQ6Z4IYDmTUnYK76VNMBipnsxEbaXtgGW3DonBZpXQuUaM0pX5Pl1dRmqrCsOfaje4pamUHctBTNtLNTI3KxIrSVLgyOpXu1VvUy0efZuMkpq7PXraqmd6jOKirNHTUyQrtMz0kQeoncDaqLJNGWLKrAfZxsO0jPdOZFqOkh152gZRldY9OyiMnvZMtp923uabLqfc6M5+xD4S1t0ztbxVZqf6SGn9LU01J2adnFQlNG4R851XUMhmkAuGNHTsC72JKxySsFB4jUXwz95bHdq1t4nLUOeobBUnYnXm2dtwmg17euCmk+T/o2jqNMaEyOshqq+WGWpr6ml+v12ZvG1+6r5FBlSJr/YhCxxnadr8nEfy3T1Lq05hW5fQ5nmc2VTHL6qSno0nSjmjPMLVMwWjVw1/F31m67bYi2hcs7Ue2PKqrtO7TdSV40FQySQUuWxzDJ6LNahLBkK04Gymj3ASuoaRmZYYryvdNxmWYZnqiupcsEM06UEJiy+ip6NVjpYV69xSqTFSRAfcQhgPFNMz7sWDZxBhO6rtdPckbN2CiEpf/CbRu/4B91sdL6lynM6mopaqjkpIqM/VN2ezhIUQsVlVY4I5paiJlBVljsp+6VNmw4zGSmyaZUodbU+arLLIkD0VNXUNTQQW3Ay/W41LRKpCJOjb2Ngy38WI9pSuzvVlLW/2MizHNIqGUUsk0NTHS0pqWQsIFmkJV5NiM5WNLJGjOzKo3YSq6fLm/QkGUa8zMZpWO1dm2ZQiCHK6VBGWWmpTUKZJX+8ayUqhAOyMjnDOwArUBuvEn0z5YqTtq0Fa+AHXNNNYvludafTQmaVNJJlsTvVU9LPOlWKadzeV4CWaWEtt8ex9rt1UnnFRVvZTS0ytLl+b5usO6NUMNK3d2Y9buykC3Iva/tycT/WnaXltTV/oTSWYV2p68pZY4HqqmONBbe7TS91GwC3HhgCAG+62I7mmrtWJJTwRHKcomqf/NYIY0qqlt5tfvZfAiergBRY2J8tGH6mFoDTSu34oT5qpL2MhJIqetahuY6BemjlkTNa2rWKSOEW2lBI7hVDSXMYFyfveR8hfEfzDSsUCM7ZnSuY2CkiQEdTzfy8rA8m4NgOcTXV9X2d5ZlM1BnOqs515qmWeF++oawxZRlyqSXjR5FLVUjfZ3qiRqLle8uCITV6iKoUy7TWU5YpBUM8Zqaggix8Upax+AXGrCZiAanhT88vZZ7xHU4ddb1J9Hdh/ajqmBc10tkldDl5P/x1nY0VEo9RPJtD/wDJuxPsv0DT6R3x6o7QDnFVKhi7r6/EtFATwXIlm7yUjyuiKDyd1rGjMx1dqvOY4Yc51LmlfHSxLTwR1NXJKsUaiyooZiFUDgADjGujrq6EmSCoeK/F1Nr4SSyWmeokeKbhjxPsAnMniioWg18fb8q/ZNQ6K09UxNDqGSnMchDNlNSRMwCkgb4zZATYFze1yQDjd5P9IPVGWuIMs7WdRw08w2PTy6gqKiIi1+Y53dSPLlT16Y5tjznPW5TM6vjqe+IAwumo8yit31c01vUhv/bKRiIaMc0UBr5/hT/Wtcan0/KvrMNZaa1VVk6s0VpLUAlBP1haFMurhIOqmegMQc+heNg3X2xHc30d2RV8W6lq9TaXdyAe9MWc0isegNhDOo/5XPscVUuoIWa81JCff6rHf8V2nDoahpGXaXkRWsdoklUfgWYYVtmniPdJ8svjkmmWKTMDrrarO0/nfaz2c0c0WhtaUWpsgW4qcvpJzWUjx3uRPl1QLqOL3EYKkAhgQDiYZDr3JO09u80/DJlWpVG6XI2maWGuIF2ajdzv7zi5p5CXIH6t3I2YogZ5TsySpVyLKhukgkXeh9msrD8cMa/MKhsxXOaeqJqlYO0iWjdmBuGup+0CAdwsbgHrhjrEJ3VeKO209dR5btqUTGMd3EbPjorqLTHaAtH3FTJL3kNLI87AfeppNq1aD4gRzj0eEnqxxcPadqF45uzSgSVEqv7UuXkjFjMI0oIxIfchhcj7VgTjl3LNTyalgpNayxqaqWXuc2CqAs1UF3d9YcAzx7iwHHeCUj7WLF1hnew6NqhNJK+ldK1OaVO+9xUu7ClU+5X6kL/DGPPAGyNdTEV40PvxqtCOYmIjw64KjO1ergzDW1bWU8i2ljSS/oCXI/K344iFIlRVTCkgheWSRwqoF3MX8gB1J8rDrhVilRmk1RmWYmKCEiH9WnezSCNQngXgfd+0xAHv0xJ9JZLq/VrVVPoPKTlOW07RQV+bzSAdyJWCIs1SQNm4niKIAt0Ctjq4j2ELW7BmclhPHayF20qWdl9HWadyvUWunohBW0JpsnyqpWATVNNXSSBp5IFN1WWOnVx3hBMZkQrZiCNgudT5bNUT5BSZ79ajCVdTVzxFaqnlWQyfWXX9YQVsB3kjXZu8vdAAH9K9XWUVJoPSFBmbUNDJ3NJTrEIppzHcSSvcgpI7GWSQMQY1eMMQIyMGyTLpTqKHN4c9VtPaSqop0kyqcRpmmbd3vmYTAWEKKCjSEG0KhI1LTC+VLJfe97+hhQcfU71dYwNa1rU0ybM891HrMZjHmvcZzXvJKtatXOk5mETyja6kyliqOqIrBWY7eFK4QzfN8g1tk9TFmGa5TqHM6qGR8pnoaZznP11QCkRjSFZZI3G4OJd6j7SOGBDKNLB30uc5RAlI1K/1uI5XEwEDIweOSNdzOqggEMS20RqCL8HTVWd6QoaaozPKsldNS1tatVNVRN9UGXyqG71qV0dWtIsm407KFjIJDMLDBFi6railKfnr1SPFMDRbrUtBpKny/Tq6sy7JanNtJZbDS1eUxACSrqiu5FnmgOySliQIzkHvhLJJCWHDK4TtO7QtTQSU66kraYI31h48tiNGNkUdhxSqrFIQoRVkOyNSPS+INJWV6RtLSST96FWqqGLRxrDIz7opFYNt2khW3gKCx22uMDBNFWUrywPT1Ewfv6haeWScBm4uzEch3NyPEpY2JvbExjDh38fZMD6YNwW+krK7UFPNLURy10MQeqqe9YzLGrW8dQ17WG4nvHPjboeAMM/03p3Oaipn1FnKrIfq8zzTmX6zMsTDasW0FXZwqhnlt3YVWubc6uSuq1mY0lTMKinqe/7hYSyVDKLkzxgC+0KdxYEEGwANzhGARTf6irRLJON+0ERtu37wq+UbNyVuLBfD1OHBgCQvW6zHMson1TV53Ln1dVpVM2bVc4y5T3VQXJ7tY53H1mIEqhLBNwvZTxhpTRQvU/WpGy6E1TSxnZNJT904BYpHydqDlBcFQwsDhhWVLpTBJEU/rTKsneWSpFwFKxEeAjk3BYkkGwAwgRVz/WNtNLFHTqPrCS9YEsNgkHFyPLjxMegOFu0bQFJXGq3oqqqtpHzCClaCKKM97JFPG+xiGKyMzBSGAtYWLFbg34toaKjesrClHRxq0CLI5lskcaIQdzjglVLL9pvW4IwWdp8xihy9KWeZluKOJIyT3rEGwPHjPUuLm3BsFw5oKiOijkoqbu3keZzG8Rdt3Gy8ZtYrcyIwv4xt81Q4cAWg0SVBKZ59FmNPXyZXV1rVK0jzOVhmDxpM8hDsgXwdVNyvhIN1xphmOVZfm0E+bUM2YUXeBpoop2pWceYEgDFD7gEY2VVGaeRkliLiRNiKpKlVFgW5s1hawBALfAYi+cqJJHVZQ4Ftr7bA/wBevnixEKihUMhxqupdNJ9EDOez7NM+yvSPaDFnGVUj1dfktXr+lgaWJRYS0rmiAqFBPiVSso/ZI5xYX0dsxl0ToPsrzaKnWnravX+SZ7KH5Q0tXUVGW08B9UMCVJPr3pOODsqjqat3owzMqI7pH1/WEbVsPW5GO9u1tE0X2eanky/iPQFRo6GEhbAJl06wFlt6yvKT7k4y7QHRWmOJzi68SRuAFKc67cNauR0fC57RSg44/hRjtD7MeyLsv7R840drDQGZ5nlOZRPnmmKij1JLQyz0k1QYhQrH3Lx74JhLEXJA2x3bFQ9r+b9geWfWtK9mHZ1W/XRFGGziq1VLXJBLuBdYo1hjSVRYqHbwtfcotYnr/wCmdpTK9V6NzCoR7Vujczqs+opE4dsnrZ0WtjFvKNqiiqR6DvD5nHm0IJ6SpkhnI72KRon3Hq4JBFz8MT6Ne+1ASuccBiN4wPpzSWxrYKsDdeB3Zj15JeQKVZxGgWwY7U6i9iPzwjMAsXjqfFJcsvJWw6XI8+vB9Plg0okA2k7T5gcE36gk+R/DAblSyk2awZTc3Hy6fjjaWYk6BjLSZplZWwkiFRGD+3Eb/wDtS+OpfoJ6ziizPVejq2PvhmeTR5zSRGTbvq8scs9j+0aWWY/BMctUsy5dnFJWSj9WkgEt/NTww/8ACTiyPo+6qj7L+2vTud1szx02SZ9B9akRgriimY005Ungfq5d1zxxzjH0nZxaIpISPuFfPL1DVoWSUxuZJsP5+V6RwVMseZS1kErS1MSRTUVRUSmYq6sHV4t32OVSxUn8Mcd/TZ0ZTae7XtWzZdTLHl2cTU2tsvVB4RS5lEHmC26BKgTg/DHVz0uYaZzCqyiURTTZbUS0ckrRiLvWhkKo7IllDHaGDoq9VIvcq1bfSvyE6l0RojXE1IFGU1lVomtcc7qKsjNZQsbdNsgqoh5WK2645LQM/Yyvj3VHi3PkSt/ScQlja/y8j+aLgGZWJ2kOxPW5Jv8Aj0/dggCg245Xiw2gnzI6n5+dsOqmlNJI9FUMwlpneFrm1mRrX68evS2G7q24AqikqGNxtH9W5/yx6IDUYLkcs0XcWQ7yDcqOvNv4jp1/LAAyOGRGYsRYAA39/iR6YDaxawO67ACw3XPoLfu9L4NGY49ziUA8hdq8m/HUdOPP+eBCHu2+w7OjNwRawBHoPP8AI4BzdriQgjp+sO5R7Cwt+OCqOBtTulY9WsPw8+P6thaRFiDI5jEjDcOSxAHw8IB+d+nGBCQAZArsou4uCVspHtfjr+fXAgArckqtwp8BuQf4j0PxHGFHFiFYWsLlWPhF+fCwv162IB+OCMYiB3aGxve7XBPl5AC344EIHFuCbq/Nl5Nr8ew/nhOWR0YM6MbG4HUkfH1t5nCiykgAoVS/JK77HzJvx8vMYRKK6N9kC58ufl/EdcCEmoswuwF/vHi3tbGYKqsW2LclhwLXxmFSJOVWQ3YWB5PHT/PGKWfxAe4XyP8AX4YPJcksWPWwP8B/LBQvk3jsfEfL8fb16YEqPGlztK8DjgfmPe+DhTbwtyRbg25/h88JllDAkkgDz4Ufz/dg6SseLXsT1F7jyvgQivtNwS1hwbn8/hgjby1yNxa3wP8AlhQ+MFiWsvS9r28vifL5YKFubEGwuwI5P5YEIAwCtd78eIAjp6/9MYxNiAwsfwP8xzgx3FA1uEN+vT44LcrYA2AA+weLfzGBCSLC9gNo6W64Tkbgs35c84UbwMVJsRwOOLYIElnkSCCMvJKwVFHJLE2AA+OEOSM1PdF5SKjIaTKjdZdS5iokYclaSAHe3w5lP/Ji7e7nqVpo46ORBUTpJHGSl3ZtxSMbTYEDaArWYkDjER0Tp2Bc8rUjG+DIaaDI4iOjyuD37D/lSoPH/EGJjmVR+jI65qSgimllWmmWm8L97Iy+AEKTdyZIl97DHI22btZKN8eNKcqLoLNH2bMfDh+arX9p2YDSHY8aB/BUaz1C8jqvX9G5KpjF+eklfUSm3T9Rx0xzpnSSR02XUDX3iI1k3/pJTfn/AJQuLd+kA6VXaZT9mtBOstHo+lo9JI6tcGSnUyV0l/71VJOSfhin83r/ANJZvWVqWWN5CIx0AjXwr8OAMbFhZi0bBX2HKvBZtofUOdtNPc+ybcBRdr+nyxuNKaWn1Xn9BkK1cNC+ZtJT09RWOIYGqNjGKMyNZV3uFTcTYFwSQMaUsSDz16++E4HWINJKDMym0aNyo9Sf3WxqPrTDNUW0ritlnGSZtkNfNlWe5dU0FbTsY5aeohZHRhwRY/Dr098P+zumhm1nl88ovFRzCqcML3EQMlj8SgHzxpaqpeeR5neR9x43sSVHkOfIY3WgpRDW1NUbf7CSPnp4yifuc4gtJd2DtpCmgAMo2K+vqrZN2Cdo2o6mokklqMpy7Lo3J6vW5mskoB9CtG/HXrjqrsLjigzei01V8xZjoKlysjcLBpdPMCrL1HXg9OMcl66nqovozZlSLTxq+f68y6mMiXvKtNl08nI6bg1SLkWBuCeecdH6Fzmop+28UsUneU+VVdVSMqoN0MNJlcqEEjkp+qPJuL+Y6Y5u1F3ZQj/yJ4ELYgAL5fCnIrz81LLvzCKUHlqChY+n/m0Yv+IwwV6KZR38NQpXklJFPzsR/HBs3cvVwc2KUVKnB/7lf54booC9G6+vFrdMdTG3uALEc7vlPO5y6YbVr5oiOR3tJcfAlGPHywY5YGVXgzTLZL+X1gxt+DqMMhtLC5Nr2NvTBy1k58XUdcLdcMneiS805hOxk2aMCEpHlF7jumWT5jaThKqoamkANVTVEKv4T30LRn4G4sR/VsNgqE7gqhgOSBa/v/Xxw7TM80plEcGZVkK/siofaR8CbflgpINnp8pas39cEfKM0rKCZoqCo2Gd47AEFC6nwbweGXkgg8WPtiTakhvl0lTBlz5X/wCUp6JqcylldERWZUJ8TIsu8fBkBvbEcyZKnUWfUGTTSIwramOFpe6QMqlhuYEAWsLm/tiR6qzIZo9ZmApGQ/WnZDvfdHTs3gWx4svh6WsZOfXFZ5/fAU8fQKVoHZk1WmSrrakNBNPANjIpqZqdS8QBO3dKBcLuspvfqOcKQtRCDvamkJLjwzU77EWS32WRwVNvMKR7YZUFectqPrCxIbjaVa5Ug8MrqftKwNiDz8xgFZ6UTJZxe8dySEfawtc/esALA8+eJy3FRApUsIz9Xp4lDCQkjZtDmwFj8OePI8g4f1TtHRR0kb09a0UpaJ0M6ukbcvEQyLuUtYjzVr24Y415g7uoeOomMcyFlkjkuXBub7unmTe5uDfyONi6PDsipy/ijVapSrArIftIR122Kklbbw3twx2aUI9BS0zd5V1FVHDBGQk31kd2EkJtsYkHcRa52Bja3A5wrSR5Dm0s2XZPJHUT0lVLJFULCVlrYZFUHaG8Td2ykhD4irHi9xh1Qah1Tl6F8uzmehWrUyveVUSQopBcmQEbyiryPtAKOuN9Sy1ddV1a1uVaYbOVeLLmzGuo2dq+q2hhCiIAqOLgPUDafsXILXMEjy2pPr+OvSVjA7AKOZO+b0VWIMoz05auYXgcLIVWVgdyKWsRywG1reF7XKm5wwo4KilElTFSd7a+9JCOSCdytze/2gb/AGgWBI4IdV9TXVeWQZtnFGYJKszAPv2iVVbYxa/IcNwTyX288rcr5bR0NXVtJm2YU1O8iEyCSa3fSEWG11BUXNi1yLc364kBoCT1RM1p3klBmf1CKuzBKOChoR9ZLzXSJhJfZJMjeJ+BtVFBeULb7N2LTUmc6ryDMoqHNM1o82mECTiWeltUwCQEiGU8OpCkXjJZQGA9sDmdBU5euXZfqWhy7LnMgWCSWMpUSIRYl5CSxiA8KubDkFeOcK64zvTr5kKzPMlaTNJIXjmp6Kt7oINu2GSVrN+tVQoKjhwqlrEm7BRzxUVBrs68U44NwwUYrqnMMzoKioqKjvZGO0gkKI4lF7KvAUEkcAeWLp0pLkPYL2f/AFfNoJKzWWt8vEmZ0MMhimpMlkAeKgMg5hNX4Jahx4hT91ELGZ7UZpvOKHI86o85zHLoM0ioXE4o6j/ZzuvKLIPOPdbco+0ARxe+EtU6wz/V+b12eZ9mM1XXZlO9TVzyHxzyuxZibcAXJso4AsAOBhZoXzO7JuDNfx85aqa0RvawXziesVtNY9oGdauzb9I5zUxVEsamKmggj7ujoYr8RQRDhVH5nkliScbLT+m3pcsTO8x3muzcH6oG6xUtyrTfFyDGnsHb0ONT2d6Ug1NnEkuZmSPKMqp2zDM5E4YQKQAin9uR2SJP70gPliyh9VrmzDWGpKdVyjLe7+sU9O3drNIylaeghP3dyoVuPsQxyP123q2hzYyLPEMs/Ye58tqsxVc3tZPL3PXstbJpmurf0VkmV0Yqc41G6U+V0w67ZH2LMfRSQQn91XfoAT1f9byvsh0pQZblsgqq2ORjllZK/ey1tbcfWs2nDcsQ9hEGPLdypuI2xVX0b6R80m1L2/a/mKLDvy6gkiiH6ksqpM1PH03JE0VLAg4DT+iMRvtV5rkoQ6/7TS1Ka5VOU5DQyBppqZCViiiLC0dMtypqHFndnaNZWay4NtPaSiDMNzprcdXkPU1WjZu6ztMq+n5Wjokoc1FdmiVKrR0LMaqWpmKJRhrne87DaVe5O4nvGJYbCTcvOy7Mc21DraqzbswzKry/KsgpHps3zz6pERN3/wDsqSkinRxG7mMsJ3G5ESR7Kq7WqbtIznUusXhSanioqOlcfo/IqBSsFNI1gLIbtJO5IBkkLSMSL7RZR0BkkEXY9oeh0zWyvltNlsrRzVscPfSZlnLgGr+rxXH1lkIWnAv3axxMWdAxuTMEUVRi52AGYG0nb+ddKpGOvyEZAZnrLrJQnWWQ5aub0tNqrMswhosyrUOe5vEklRVCjvuqZiTd3JU7E3GwaRL2FwI12j/Sb1XqnNa+i0FRyaUyExLSQ09CwSojoI0EcEElSvijjWJVURQ7QbEkuSSdvmtdqDtazObLZITkWQ0hFZU0yfrlpKZWC9/UtwamcsyxxRiyGWQJGoG5hPey7sy07W5vmGvs5yg5PpPs8iir6yKVBO0VRIT9U7w2tUTMUMv/AHrpHGoCMBh7bkTWm0NvOxoMx+TkK5DUccWkOe49kaDr9d6hvZxpXMuwmjp+1DMlp6fUZlaKKhqYlkWUsimagdDfjuZAaqU/7FZIolJmc7dDpjs1zft37S8n0nFV/V2r5ZGkkgpT3WU5dHeSokihW9hGn2VF2d2jBJL43epKjN+0TO3rzR/VliRo6SmM4ZMvo0LPsaRvCQC7yTzPYPLJJIx5UYZaYm7R8q/TFDobM4Mpy7UtJBST5xQq9PUzZcGZ+5hkchoYZmQyuwszJGhYqnBmbKGuMtRe2nLw8Ng18010ZcLoGGzrWrj7Y9ddl2h6uDRtJSZzm1fklPDlmTaSy+rRYMohjBWGKokTcsc53M8lu8naSWQt3V7Dm3ve0Dtf1VS6EpM4osjy/MXeSsiorx5fR0kQLTz1Ely86RIGZmdmBIAUEsMSDPcjy3R2Uw5VlMCpmFdSiWZwuw0lFKt0Ueay1CneSfElOUHDTm2i0x2p1eg9NZvl/Z9QuNV6nqEyyTOXgH+oUETBkgowRxLLIN7yW/VpFGo5uQ6yQdn+8bi/a7bt2Dbt1VKbPIXAMOA2DqpVydoGudEdmunafQFPHW5fltJSfU8u0lQqn6Xq4Hs7TZrMQUojUsqyywgPO4EcbLHFGFalKHUGsO1vU8ejdNxw0lRWM0s7U5cwUkCDdJNLMbyzbFBPFlvZVViVGNH/AGQrEpcyrqiV6g08Ynr6ssTuMj2VQx5ZpHNhfljdjwpxPezTIqfRnZJW60z2oOW0+r5istWg3THLqeW0VLTLcF5KipQmwIGym3MQt7ydnDDGZB3nE4V1k7ue05VUdZHODDgPZF1VTZLoSOXR2j3eszKRP9fqpQrmPabkzG5R5geq3MFP9kd5MGcU5nuoKxqOTK4qpapDL31S6C6zSdAZJD4piOig+Ec2BJJwrqPUOZajqRl2V0ZpqKSRUioon3l2vZO9ksDK97DyUE+EDA6gpqLT1ZluXrGlbLTyrI0SsR37jq7MOQrP4VA52IDxvxfgiuEF+JOPDf15qCV5dW7gAmaVM+S001LC0UdW3FZV7QXj4/2Mbfct94r4ieL2HOldw4ZY/Ci/aYixN/6+ONjNT1eYd9XVJUpG5VQo2xqxJHhHQC9+evBJwyNkqAkCpJFAxUMfsu1jd/fpcewGLbABU69agdXLUmzBIV5W7Dop6D3Pqfb8cIkPJeaYmw4/yGFG3OQbEBzcethzjKlSHaEDwwrY/Hz/ADxOMCoim7yM9lAsvko6YAp+zzhWKK43Hz6fD1wqlOrOACAPM+gxJUBMzTZYyQWbhV6/ywBuxvbjoPbD941YWVdqr0HW3+eAjiAO7YS1vIfngrrQmTRlRYg7jzb0wUcc42DohJAUn8ucNJomDXC+/GFCRWP2YZiZdM6vytmIMOX0+YRH0kiq4kv/AOCpkxZva/rGmoKGqeCFI581ipIEUdEpaVBHAv8AzSoXPtB6EYr3sTyCnnyfWOp81nWLK8uoqSmnG/Y9TJNWRstPEf8AiOIG56KodzwuI3qnOqrWOo5cxrTK2XxTrJVS08Z7uONmVCyi3hQLtRAeiqo6k4xHxMtNrLdTTj40aQPRaQkdFADrIw5pm1PT09J3rTwhzbbGGLSHy5AFl/5j8sWfpLKs71R2eZTl2TUMrw0+Z5lHUE2SlAKQyGold7RxlFLIXcg7OB5A6PtZySpizmlzRlkanzOhhaKZirtK0S92xaVBslcqI2LL1DqSAb413ZtkGVZrrLJMm1JNMcmrK6P69TxTshqI1BbuVsf9pIVEaE/edcXpHCWASjVjtyBqqrAWSXDrwVlZDIMu1TTQ1VS2pzWU3cx08Mk0gzChqlMEkcM0qqWBQyHvGF0aIbjtBI1miZsqo9TZ52YZXmecZxl1TmPe5bXZTRx1LSGn3qZnhkdEMbQklnLqqmJHJ2Ybz6gbQ2S1HZ/Lmf6ezAT5nktJp+Omm30yT900MjyixkO6WVRCnjDq6t4WIKPZdpaTSpzbO9bzZjpvM1qEyDKVqZJKKRK2RGkl72NlBZO7QRMGBjDTx70ZTih2dGOc7EGlN+OBp1nrVi/3mgZ6925S3NqOmytZaengyqnnqhURSUj1VPU5hEgjT6jUwVNOzO1S0hcybSsSgbduznGs/RVPHR/WHzGlqJSyTMkZVqqWO12nRbGSSMOrMVYb18Di6EgN0p8vzGKopXzBYJ6ljMI0LO0ZsSpaNGCU5VjYXkAAL2BHA3UEU1TSmjevNNDNUJV1VZUKsLfWIz4p6iVVLqYlJfb3gVAVKh94xXvFuHXXxip6AqGKcrq6uGsYxBjKWlZ4O9YJxv37fC6jw7vMcOPEDgab9JTRJfLZ6h55JYXYQhwHZCsx8FthG669Esobyw/TKs2q6Klzyr+tVeR108uX0uYUz06SxSJIzKJomIEUrAl2V7bg+4MbWxqc/wAskpnh7mUxvRqIHcjcY4wm4RSzLYu21zuDBRtAtZQAbQoTdqojgKrXiqoBPJUzCWWGWMwy01JVdwHC37uQuqNZBJdithuQgA3w7FOaLLoYa2IUk0lIkrQfV2jPdS+ISgsC0gZAX3hhYlAtrNhKvpMyjMkdTLWRVEjK699VHfIh5Xwsi2B8P27BhYAcglOrWOlEVtyCymPuyzR7TfdJFcbxFuDDi+yTcVuOk2BUSUJnofrUUlS6XiRqiJW7l5hvtcq1nfgrwu02O4qR1a1cE6RU0bUUJhgeRKdlp1D72NyGmWxla4sSxugsbDCSVhScvAvdxzqQ0hZo4ywvyS12e44udt/wwnKplCRoxMruHWX68BEgKEPEFB2AvcHcSDZdvOHAUSErFpxMiywTyImy1S8MxAkha53uzsNisfCqC3AYm56nqZYJYTDDDIyOrFQlLvRiACpI4U2HS3QcjnGvSjWnqF+uFoA4EsbGJrlvshgLHjb0PIBv5nDeWRhC0qAiNV2MyqyqzX+/Y/qyeo8r84fdqUytFuMm0vnmaUtRPlFB3oivK8FPAVqZYgNrSpF0mRGO1tjFkJ6WxoarKmhpe/qZEilsAsRJDgeTFCtwpHIJte/A88bFK2KGjgaaZ5TTPJUU8UdQ6sruo3FgptEpsC1vG9vLDDNszrqkyz5nVmqqnkLSSudwZuAAB027bADpa2Fj7S+a5ddakj7t0UzTzsnooa3tHyHLpEDLU5vl0Lr6qaqMt+S47G7SambUPYX2x1bN3gqMrizK+0gi+eQvzfg27w8jj2xx72KzpB2saXnkAAXPcvLAcWvUIv8A7ljtHKcpgzfsl7SMjjRY6iv0dnKsodtryUu2qUFb23XpnsQL8Wxk6Qf2dvicdo5lX7K29ZJBuPJbqLUkWr07M5c6rFkpNU6XyqmzIFGu1PXUByuqLPfaBvjDWtfcgPljgrUeT5hkuoa3LswgP1qncwVSsLjvoyYpbj07yJ8dUaR1AKv6PegcyQqJ6LKs0y9XJ5BpMwklS3uFnX5dMUt9JtKSLto1aYkW0+c101h0CyyLOv8A9PbD9FOMVpkh1VPqm26j4WSbgoborRuYa21BlmQ0CvSwVtXFTT5jLExpaNHcL3sr2AVEvcknp8MNtV6G1bojMnyfV2mMyymrVpESOtgeBX2tYuhYfrFv0ZfCeOTjTx1tdSU81NHVztTyRsgpzM3d3bz23sbC56YcVuazZrOK2rq6yolaNFInmaQpYWKqzE+Hi49OluMbg7XtMaXVmkMuYZrVV0J7vxXuq3seb/A42Zljnq8tzCVv1eaUpo6lj5OB3ZPx+w2GdZ9kkg2e9jbr/XTBaQPU5FWU6n9ZQSpVx+oU+B/z2H5YZaG1o7y44etD5J8JzHnw/FV6haG1VX9pfZppHXbS9/UZnk8CZixhEhTMKUfVKm+0h13NBG5IuLyA2xs9UaYfWnZVrXs8hDmvzLI5c0yiPaSRmOVOK2AKwG1tyLUILWPisVGOffocalps97PtTaZqY0kqshrabUdDcnctLW2pqtUNxtC1CU7+niOOjdJ6tr8lzWlzsRTVCZLV02ZbWHeFo0I7+Mn7QvE0o8YIIuNwKkY84mJ0dpO/qDq+RzHAkLroQLVYrmulPMZey81tdUlLHqSSrgjdqTNoYq6GRbcbwLnb58W8xYnzxGD6MDcc3VuT738j+/2OL0+k12cR9nWt9TaPp1LU+lc8mpaKRWsJMpq7VFE/nde6ljHwU4o6RAGCFbtwQPMn0v8AuHvcDHo1jJMQac24cMuVFydqAEhcMjjx/KSO0brsSz3uEYA+1xax8+OD8MFVBciwBtzc7QD5E3t+eF6iExqGZWRWJ2l+Lev8r4RLbE4AIvYWtwbeR9PUe+LSroqbY32xlfIbhzb5gcj9/wAcCXSNuEICncFPhJ/vAj7PPX88BvAcEEq3UsT6/uH5nAyWUhZFtYFh4hc/G/X4jywiEiSBfmyk2v8AZ56kenX0v8sGElvEZr2AU8knr1BPB97WwYqDtDMzBRbg7tv9eYGE3UKSG2keVxcX9j1/rywIRSSqkWUEWB6Ejnjz6+9sJszNeTcxubFib8/HCn6r72617cDn8Pj78YJZXILlma33RyPj6j88CETlgSwNt3NuPLi/54zApMGKoSLC9iP4e3XGYEIrgqbqht05Fx19vl8MAGa5a5G37wHT3t1/q+BcKzr3bbmvYEE3B6WH9WGC8i21+fzv5++FQh8TA25sCTzfj3wUqN5PkByT/LA2AIF+ebnra/TnAKLmxHXyt09/+uBCOCRa5N+h5tb8fxwHhYW2+fNhc+3xGMVbrvAJ6jwr/X44Agk2a91NwORf2v5H44EIzC1ht6gkEi34XwUk2BHPkD5k++MKsFYkFrnm/UDBXXcQShufbqMIhJt1JXwAeQ9cSDQECNqiPNahA8OURSZi+7oWjH6sH4yFMRxja9zext6cYnWg8hqs3yg5TRf+eanzSmyinPTagILt8Nzof+Q4r2twbERtw458sVNA29INytnQ9HUZdo3Le/VWqMwZ85mMq/aknP6u5PQ9zEjA/wB8ccnEk0TWU9Fq9dTV0RbLtMRVGq6uIDbEyUEZmjjKW431T0y8cX4wpPFRyzu9GWjpSe6gK8KsKWSBbN4bd3GoFxcEjmx50XaBO2meyrPZwu1tR5nS6djcoULU1Kor6wWJuLyNRIf8OOTjd9RNTaeuC3ZP3MXgqUgrKuRtQaozOsL1lPSyF5XN2lrqt/Gfc+Jz8sQ+NQI7Em4FxcXBHp7Y3mbtNS6WyuikjtPm1RJmUxLcsl9kf/uZ+eNNe1wByDa9+ox1Nkb90m004YetT5rDnP2s2D1x9KJSJG7yOGJS0j7Qu0bmLN5D35HTnnBq6hqcsqZcvr6WamqIHMcsUyFHjceTA8g3wOXFps2o4El7t2njAdbAqdwsfl1x0X9KDRGXajjyjty0tTLFl2uaR8zaKMeGnzKD9Xm1CfSSKcGoUdWhqCR9jE/aUkubqqMRksvjauaZj4CwtyOvpjfaL2hZ3eQotjuYc2tJE1/yxoZkP2A1v443Gj5UM01G7W73wf8AjUr/AO22Yitn8I0Ulm/iUXR1bQiq0L2S6alhR1zbtQqpJVQ3Ro4qfLYmI9rMxt5c431Fr56Ztda0iKmopcrz6po5RYMlRmTfU4gp9zVMbee0+mIvpaeeXSPZJmu3amUVutc5fd5CnhgI+d0XEF1dnMWnez+nyqmnL1WdSxV81+CKanVo6ZT/AI6iSeW3pGhxhOjvvjjOY+TX0Wg15a17tvwPlVbmMsUldU9047lZNkfQXWMBFN/gMNC0dyokB4B6jphzR1NTRrtpZ2TzNrEH3sQcSPR+U661zqSk01pOjTMMzqixRWp4AqIqlpJJHddscSKCzyMQqqCSRjoSTE3VQDWaeyywGvNTWp3flRqKmnqC4gp5JSiNK4SPeVRRdmNhcKB1PQYKw/Vg38LN5eR9/T+OOpu23ssyz6M/Y1p/IMl1BHmWre0GjkzXUmY0yskK5YskaU+XwA2YQySyd7JuAaQRx3CrZTzFHWwFh3mU0Bv1Kq8f/tWwyG0Cdt9gw62pz4uzNHGhTT7BuQWC+ht+eMYEEobK4tx0PxxsBWZWTsfKSB5FKyRf/bA4I0unwxZstr1YdSlap4/5o+mJO0P9J5fKbcH9Q5/C2vZvDHU61pMukFpK6KopYCFJKTSQOsZAHP2iB7bj6Yd1tFNEsjTTTUpMMUkdOzBhJTTqNpVvNdu2/wDeX1HEXmqaaKeKrymaugngYSROzIGjYG6kMtiCCAb+WLFzeQa8o4ZMqjQVdLHLUUcYIDVNM36yanXyMsMveusfVkdtvIANaUlsgeciOFPmvupmCrC0ZhRKrhy6jpqWpoamqjqtjpVd6FMe5twDR2H2Np2sGuQwuDY8Naeo7mnnikYg93tYAXJBFhdfPm1vxBGHc2X1k+WTZrS09RJQ0kscM8/cnZE0gJjDuPCGbaxA87HGrjDqGiuWQAgKWIVfI2HmOfMgcjE7aEKAp7S1M8mXyu0xaA2SosQSjX8LMOvJ+8OvIOF0o3jmAq32inMYiV3AJJAbaJAf1a7WJVzxbjjA5ZSZVV5FVrUrWJXxqJKCWFRKjuGG6GRLgpcXYSC9iNpXm4VeAZXVbJo/q4lVFPQoQwB+30te5HodynDXHEgdZJwQNVpHUxU+aqrR0ThRFI/d+HvA7qx5Ks3TcAQvAAtiVwV9JBXiqyRmm7if61RrONjoxl7095a4diwRT3Z5VFIFrgRCLM6nL6ukzDK6oK9OqP8AradZkDIzEJZgd0YUDjoQbHph93eXTzzV2UZtllJTVR75KSpmMb07t9qLuwrblU/ZYcbdpuDcYifGHZqRj7uS2lZmDxV1Ucg+uw0k4cyRVVarOzu/eMNyqVsHJ2WG4gsTY8Ya1tfUVIpdRa6zKWVatxTMYRFJWzxqD+uVbBZUUgLvc8m4Baxwvk2m8+q9QpllRStOZIvrEcg2zJKpDBCjA7XRpNieE9WAIBviLZ9nk+d5t+kqmMU0/dwxWDkbGjQLe/kSQSbWAJPTCMjDnXW7M9fXwhzqCpSOpc5Gpc5rs/eBYjVylxEn2VFgFUAcAWHQcC5AAFhhmYaOtpgzQimlQWLx3Kn3ZOo+K/hjd0tdBX97FnWUJXJGoJqomFPUpfp4gNr834dTfnnDJsppapWfI80iqlbnuJrQVI9tpO1/+Vj8MTtcGAMIpTh14phaXG8DWvXVFH6inlja72Zeiupup+BwmIz0PTD2pp5Mvd4Z1lim+9Eylbf4geuLD7DtMUldnVTrrPcuiq8n0ui1jUsptHW1W4LTU3uHmKAj9gSHywlotAs8RkONOZ1DzRDEZnhg/RTfLtC5hken8j7JsmyySo1PqCops0ziG+wipkS9DRMx+wsMMhnkJsFebn/Y41PaJU5dWRUWidHSHMcty2Q01DKiEHNa6YqstYV6/rXCJGvVYEiHUtebZhqPM9PaTrs+zjNzUap19HUBZxZTS5XM7CqquPsyVkgeFPSnScjiVMVXkWqMz0rqXK9RZRl9PWV+WGSro/rCkRwVAUiGcr5mJv1gU8blW/AtjDgEgaXOxeak7CfYahsFFpS3KhrcBl11mui+1zVOi/o+6UyfskgWmzXOtMIgloiRJTw12wl56hRxNKZJJWjgPgQNvluSkeOf6TMdS6vr313qSsqaiprpmkgaWQs7FfB3pJ6kW2Iei2baFCgY0WSaRzfV+oaHKpKmR6jNqn/WK+cl2VCTJNMSeSFUSSMx6ke+Lch0tUapzynyXTlK1F9cdKPLVaUKlJSohs8lxwI4UeVz6hjzcYibDHZGXSbzzUud51OGqpr5BPvPmNQKNGQ62INI1tL2eZvpvtKzbT5zhaOpqa7LspKk/XaimG2KWex8FKtUyBmvudonVehOJPpak1FryvHaZqLMa3Nc3zpzFQiZVuKVG2LBBTr4Ykln3hUUAbYwLksxwz7QqPN8wzNqHJKGOiWtFHkuT0MQJWBART0cMZPicDfuY/ed5XI5xbtVS0mgKTflsi97lSnIclqIrN3TQRrFLUAjr3MZuD/x6mPoVbFGeVsjBdwJwruzPqPFWYo7jjXV69VWizfIYcnoZMloJqeaPLap63PpoX8NXXIpDlZB/uaePvIoG6MwqZesi22pyOsy3s4yrK81kFBFU039rtQ1VZN3MST1hSSIzWvtipqRKONQQSJHPdqzkDEEzD9DPlPeZrXy5JlU9dTaemqZlYKqzzJ3sMarczbadXfaATbZezMLyjte1pmXaxqOpqaLTsuUZS9f31Bp+KPdIZRaKGSo699VhNqIn+zhBCRjfuY1yXSOqcAcPIahxz3Y6lNgzBufudZVQ6rzKv7Ra2j7OtGxz5blGc5hTZeHmj7ufMHklCiaoAP6tFG90p72XYWcvIbr0Rqil09p7T82oStHVZPSTQ5fDlDUw/1+RlEWX5esx5jibuN1Rt+1BBKLguBiuez3L6Cj+kDprRVFAK2i0bDm+c5zLS/rDX5pBSOsiQn7yQuyU0R+86yydJBiUdt2XHUGrF0DDPT0mU6GppTmk8ktqaCuKIa+eRh1jpkEVICOWMThfHKAZ5xSWNtKNaK+ZNKbzgN6ijNWOOs4eQ/VVXmWX5rXy1eb5lWyVkzyyVGYzyRhp5qiV/tbCLGSSQEIo44sbLExEZ0pkiZlX5/rDMnQplsK5bSiNt0cEkiGSYofvbIRYv1YyFumLAhD5xkFKi1EtNS5lTGpphWvaaOhcFEnnboJ50jZyL/q6YLGn+0dmU7MNP5DX5Nk1LWSUbZdm+ZZtnckR/2xy2CQComPmEFNRulz5y288TGQsBBzJAPnieQKaWh1CMs+uKgXaTS1WR5DQaKpIHFQtOub5yOpWuqId0EB9fq9I8YI8nqJvPEP17qHUurWyXS9XTCAady6kyGko4n8FDGieKMeXfOxZ5pPInaOhxceZzy5bkuZ9rmdywVOcV1T/qUKkSJNnlcTNGhA6rTRkzMPLZAp+1iuMk0Ez5fWxujGWPOqXK5aky3fvJkRXJXre00h3edzbpi9HO0986jQeleaqvhIN3j6qJ5VCnZ7m+R53X5NFmjwxjMVoWk2Bi6MYC/BIQDu3I6kSW4JviN0tFmX1qfUOZwSTTzb+6c8frXVlR+ONoIJAHFo7dMXh2k6MpKkaz1pIJx9SqTSZYIhtUVEs0qIp/upBC5sPKJL2GNJquPKU0fo7MMugjFPW5euZVKqBdWp6UUgT4iZKskerL6jEsdqDxeaM8D605pj4KGhOX6KvszpqVNOUmVUO76yss9RVxbCDFEkcaQ3PQ7gZW4/b9caOro6aCmYU0qykwLdkuQCxXd5dRyLeV8dBam7Oly7JNW1QRP1U2S5PC4S5IqHqqgsvqGjpYviHGK/l09TZZqCi0xmgEseZUE8oeJdzlSxY7bfaOwM6jqSqjqbYWC1B2Xjyr+ESQ0zVZ1lNHSNS95IBvSS9gTtYEcH8vxwFXQx/WK8xkujSuqMBxYjcp+dvzxYFf2eZmmS5uamDfVadljrpe55WSkOyOWaNujJ46aYN0KSg9L4ZaWyGgraz6jmdUKOCYHL62aYWFGSwEFW3pGkmxZP2Va/IIxbE4pUHL9VAYjXFV9ToShZDfzN+lsPIKSSVHaNNwjG97dQPW3p+7G4z/S1fozUdXpvUFHPStBO0VRAReSJ1NmX0JU3tbgjaRwwwzqqdsurTTOyShLMskTeF0I4Kkeo/DkHkYvNkDxVqqOYWmhTMwLu+0SbcgDkew9/3YAxizuLWHPHT53/AOuFbxg7TZQPf8vXCivsBMPHkZG+18vT9/uMPqdSaknh7so1QkiEjgFeWH8PmOnrhrVRqAxKgAfgP88P5qSpjpvrZpZlhIRxLsOwkk7Tu6c7WHJvwR5YLVZZVR02+rH1cSeJUkJ7yT0Ij629zYemC8BmUXSckhSZ9mC5KchgkaOlaYzyKD9tyNt//Dx8CR5nEiXU+UUXZ/Ppujyq+Y5hMr1VVuK+GOQunmd5tYAWCqLmzMQRGKaAILBPO3i5P8h+eH5d5EenQKiuAGIW5PpdjyfhwPbEZgYdWuvmn9q7bqorxzDM9K/2J1DoXTMmna4pFTRZZlVPXNI1RVSuoiqafaheorFud7kxja7Rsu0KA0yMZD2XZTT02YZfpXP86rJJmzHNYqRsxlyNTZEiiWRhTzWO8tJGDZjsV7gHCuQ5rlOf6ZqJezrTS6fkohT5ZmVBl2XpV5jMlSDEJ4ath30geQMrxLscBxtYi9olOYaCabKICk0FDJLElbTSPEzoi7SVRwQVX7I3WJAKkE4yrtQ6M1GNSNeWvVv113hXSRUP3YKdZr2k66rmkhl7SaNSIGp8vNHTRwxUv2ArUmyMPSkRAh3BDHcVLGQYa0+s9eZVlQlyzVecGmoo+7ki+tzzRwSPcEMJFeOGQNwsgKswVCrE3sy/Q+aZtlNNntFp/MHoIqeONJlhV4H7tWu6ggWjLOxKJvVSW3NfkKZxlWr0pZJxlNXlGWSzIlREspFNUJIf1LTkMBIBIO7XvF2gMm025xAGxijKDkpSXEXvlNaXU9PT1kVMWUo0DQEzu0dNAu5TAHdfEll3B3Ukr3o3H7WF6yt70PTGjh3xmOMSGOoMsLRsCF3O8iMpBA2/rFAUEbQ1wxyHL85rKuPL0apLzyikSCJEp/10jAJFYKNpHi3tJdVG5tpGHVbpfLlyhNQ0Gq6WfIRH3D1NDl9QZmaJwjIsFir3dlCs0iIwKnaDxhzmtB/VIHEhMRk8FLS1FZXU8cclSkdPAfq8hjeYsRIokZQAyp4ylyWDWBUDGnskUP6QeappfrZjkZKFlTum3d5DFscFZNqhXsxBJZbE2Ixt8/io6SnyicQZzDDU0CrNR5xO0xpXkJZVR2IQb4QkhjUIdr2v0wxoIaGtSXNqnMjS5fSIPr8stLtKKxJRdhZllZiPBEvWxN1Cs2JWHCp69eskx2wJnRR1+ZVppqeWtmlmqiyR1RVTJI3qSbd4xBHiPiNlBuBhlLUSxTs0s6Msbs26MjvHcHkHgNv5sA4G3m/XG1r2V6SmqqWmqWy6vhNREZInRJowdrAxksQFYWuGZLciwBxrqmqkpWiqZSKud3CoXZJ4wEIJUlSVlufDtYk7b7vu2maanJRlJ01dVZerVNLVTwT92O8mjZLuhNuSRYgn9ocEDpcHCU9UjxNWtUzs24EyFCVkvc8u9t3ToFtfz5w32ilKNMIkHNkDRm5ubgFWLG/lbpYDythbJpalM9pq6CSc1ccu+OUt3k24G6MCdwax5225FwVOJaDMJlTkkislJMwV2TbCHTuOA4NyOORySDceEbSOCMavad11Z/1ZDM0Z8QUHhrrzbzub/LGyPeorfWKql8TbijwhrsxuTGAQV58xtS1rjGkzSukgRo4QrGT/AGjq1yyE7iLj9ri/sBiVmxRuNE6CGtrAIzIwdrs1wS17lifzJvz1vbphvmaU/cidQ6yOBvQ35sLblI4IIHQ2I5646Ry7s5y7t40bPqnIaT6trvI6FnzikpoiTndDAqj9KwRL/tJo1ZEq4kG8jZUIGJdTz7qHK6vL6lqCqotsoQSoyyCSKVG+zJGV4kQjkEcfAgjCRSCShanSMMdQ5aXTuZzZRnUWZUzWmpWSqjI/bidZB/7THenZfmFJneq2y2MxtBnM2Z0EZtdhFmFDU91Y+hNQg9/ljz9VjQ1sUzDdsa7e48x+GOqPo6a8y7Jc90lVZpUMYKTNqDLKp1G4q0dQj0k3+Eo2w+wPpjK01BUslGr8fHEq3YJKMfHtC13ZfWyn6OeWI7qopc51BTqWvez0tC9lHruJ+V8QH6Qk8tR2u5+sjjdDWyQsQBa6Q08Z6+6EYtqPLIcg0vHoyOJIUk7S9R0puPsRRSUUbAegCKxJ9BjnzWWoJdUawzTUXDNmFXUVoLrfaJpnkXjp9gp14xLYxetr3tyx59FNtB/uzQVHSN54UHbxtQWA9+vnhWCinqmMNLTtMyqzsEViVUC7Nx90Dkk8AeeF9lXXuIwJqmWVgqggu7MTYAeZJNgAOvGOgsj7FMy0D9GXtB7as9p/1+Zmj0fkcgI2NLNN3mYTR/tJHFAabePC0jTAEhca8szYqA5n5oqEcbpMlzk4tGAwO+5Fx6e9/fAZC0cOcJTTN+prkekf4ONoPybaflhVgJB3r3YuLm4Bubev8ca+sBjKyxyWZCCp8x6fDD5Wdows2pI3XHB2xXx9DjVMGRdsGTadzmUxZfqI1Wj8xN7bYswjMcTE+WypWFwfIjHb+RS92i/pSRln2LHUPMCxEqXWQuyEMDu3B2HHQsD1x5iwV9VTaiGZZZI0M1dElZTuhsUqFIdSPcSocekxz/LdTRQ68yuWnNHn2X0Wpljs0bA1UW+ohuPC22dahRYhhyLm1scJ+0lnq5k414H19arptESULouHl+KKufpeacGcHSGqZIAp1Dp6s0fmLbtw/SGUvupnJ8y1LPGQfMIMcPeOKwlazAePd5N0I46c+Q5x6RdrmRQ6i7DNWQ00iy1+knpNa0qB1ZnFGRT1q2HIL0VSjEWFzF6DHnvrTLocv1NWQwu3cTgVcDKbqyOCSB7hg3x5xv6DtHbQMJ1jm3A8RQrM0nDckcBqPI4+tVoGlIJXfz736fH5dPP44LcMC3P43J9L+tvb+GCm2wsy8lbci4A8+fw8vgcHsEBLryOW528evTjG8spJ7ZNx2g8Drfj/AKfnhePcpZlO5SCEAHhPpdbfO3tghe9yy3sLN4Ldfe3T54EEs6oxYE3NyTc3HUg9BwOnNvU8YEISvBffu9NjXJPr7evPwwi7XYBjcAG1oygPqB5XOFNxcBQysenHBb35/f8APCDuFG2yjeQQWIIuPlwT6+nxwqEmSxsQpA6Hnp7f11wMVnI3lSL8gC1vnzY++MIufEeotcC45+6Pb+hg6bzuNix+Xht8f3fjhEJqCQ5KDbcnj/PGYdxhO9B23HO5L2/EeXljMCE0ZyXPhAsVHsPTr/XTGAi5Bva5FiL359PXCdyWttv52PkPXB16AgDoAbC/Hr+7CoRhxtaxKi/nwB54FB4Ps3t5E2uPPjzwAIU3L7XvcG9iD/XlgyxsL7UvtFyQPz9fxwIQlV4VwnlcsoFv6/DBWUqWDEqeR4jz/L54EFQvIAU/Nfz6/K2M2llG0eG9ib+fkPW1sIhEcsSQyG4JBBHF8FaxUXF+OLny+OMIspLc26E+mAYLYnbcjrx1/DCoSM0nJF7j4eWL07KsmqaLPVEUDvNpvIml2r1WurLRg/FRO59jED5YpzTmXpmuo6Kin/2LyiSYnyiQb3Pw2qcXv2eET5JV57Xbg+d5rJUsoYglIf1cY8JBH62ac8H7inyxkaVlux06x/F5aFgZV1esPzRSKuZacymvzGcQ0a7WE7l9lMI921WtyQAVsSSLAjjgQj6QNVXxtpPs8sy1OVZJBPWRHkrmOasKuYH3SJ6aL22WxZuUZCuqc6oNOyCBIc4zCCkqppmbfFSsDLUseikrBDMxfrYW63OKN1Hq+XXXaZnXaHXrY1E9bqF1NrRoSRTxgeQC90oHtjJsHdJkpiBz1ccQrtrxux7VCdWVKTagqIoGDwZeq0UPoFjG0/nc41xG3ji/mvrglOxlDyyMSzEsw9T6n0wdiQLECwHQi9x/X4Y6iGPso2sGoLFkffeXbU409AavUuWUq/7yrjTrfqwx2v8ARv1fpXX+lu1r6N+uqAnLf062ssqqKZC1ZlndMYaupoxexlhUwTGPpJElQh6jHF+jnEetMoY9ErYyP/EMWppDU2cdmHaTF2w5JTiqqNL5lHX1dK32auilmkp6mFv7rq4jPtJijaC42i400Jbgd9fwrUNBBeIqA7Hwotd299ieoOyHVFXleZUtOUjjjqUqKIlqOtpJf9jXUjHlqaX06xSbo2tYYqTL6o01YHDbVbwk+nofkQD8serupOzXSXbzoqbsqyfNaOWvp0Gddm9fV/ZaGshEy0UzeVLWRFVIv4KiNmHIx5f640dV6Rz2oy6opKmmCSyRmnqV2zU8sblJaeUeUkbgq3rYHoRh9jtbbZHR2DhgRsIz/CS0QGzvq3EZg7l0LlNSp7JY4ZpRClJlGplM3/wPFVZlRLLJ8olkUerOo88UHrXOzn+cNJTw93GpB7q/ESqoWOIf4EAB/vFjiUaf1xQUfZFqygqmM2Y5hJlNBAjkkJDFLLPKR6b3SC/rtN8QLKKOfMamOjpopJ6mokWNUVbs8jGwAHmSSLD1xXsMJa975P5TTkD7p9peC1rWa/miknZt2b592mahi0/kQpoFVHqa2uq5O6pKCljF5amok+5Gi8k9TcAAswB6s03l+mNB6O/RfZlkj1lHW5lSZLTTV8G2t1lnkpDU0VShv3NDGdtQ1IOAixd8WeUKsJ0/kFRp2kouybR9HHXZlmFXDFmTxtf9JZnvskJbp9Wp3uB91pBLKeES047NqrL9TdolbnuTVpn0r2b00mmtMVfQVmbVYZq/NvdigmkU/dU0i/dGM62Wk2oOmeaRNxpt2V8TlsGOeVmCAQlsbcXnl+iif074KbJdS6Y0plubvmdPl+lcoM2YSPvkzGrqJq2pq6x282mmYyX/AGdgHAGOW4isbiQG23zXjr16gj8eMdTf6QOagl7eMwyugplhpsqo8jymGKFbiNYcohbYqj9kz2t7Y5cmoqxCf9TkQE8Fl2W/EjGzo+QOgDnHE/AWfaGm/QavlN5HNyQgAYn7It8vTAbV3bmks17i6N/LCn1eW5WTu1v+1Mlr/j1wP1RC2019Co9dzMfhwMXb7dqhuOTN4y25Abi5F7cfztje6HeGHP6OhzCrWLLa+oipq0ufAsbMAZP7rJcMGHII9L4bx0eWFkWbO4F2i1xTTNYfJRhvmcWXwRkUebCqNuAtI8Y/Fjhjy2QFmOO4pQCw3vcKa6ryjKKOjqPqGYZnFJldc+XV9JmMCEwygMY2DQmxVtrgErwR1sb4jNBQtW0dZmEk9DGMtijlaCacRzVCGQKViW3611DG46hRfnbiXx5nklbl6Z41ZV1VM0NPlOfSPCFlCmNWgqttzdopFKX++I16FzjQZ5QNlNSaJ9gcNwYRdDcbgyE8MGUhl9QV874qwOcBcJx6r1sIU0jQTfGXXXFMc0yuPKc0qqA3cQyEK4APexWDI4PoUZTb5jzw4pp5K1lyqIJPJIo7q0oTxbbkKTxduBt6E9CGFylHPNPTItbW08kdIwgW4YSiK9xtB+0FuxA+7coeGGNjqeTO9L6irMtymhZaXK6hqSQvDHUCpZeCZbAqQfJOiggdQTiUuNQzX18pgbhe1JXJshrc5oJzHlsoekCxlaakvIzbj4NnG9rXNj4iARY4ZUiLLV0OU0lPO9ZJdYwYe6ZkN9p5NjtIYFr22dfsnGvqM+zjNY4Fq6x1ajkMsO39WyubeK45JFgATyAAAcbvSuYS6lzuNNW5jXVrpUqyM0ZkdkqHEUwaS90FpA3QgsPIthHBzGlxQ2jnBoSNdqZcok7vTJjmq4w6HM5Bu7tmFm+rKeIx/wB7be32ht4xqMsymizv/V4Kusiq9paRZaYzR8ck7o/EB5kleObnBY8ompwiUdfSVMpaWN4o5gkkZRto3B7AhuosTxe9sK1FJndFDIjUFfSU0wAcbWKyD+8yizfuw+gAow0KKkmrhgnCZXmlJSGkyhYMxjlVu/mpKhZyfIbUFnUAXFyt+TjTVFMsAtKwWQk3hdCrL7kEcYayxCOT6zC99pv4W5B+I5GHi55m7wdzU1rVMZ+7UgSj5brkfI4eGubvTC5rk1nqKvMXjpe8lqGuI4kYlzcmwC356+WOqcuyfJ+zbJY+znNcup6qlyCnWt1a6vsqZM0lUFaOB+gkRHWntY2eepbgxXxWn0edNJHm1f2u5lldPUUej2i/RdJKLR5jn0276jT2PBRCr1MvkI6dr2uMEzTNZqqoakfMzV0sU71ldmU7EfW6qQky1UnmdxZgi9drG3LnGNb3iaXsh9rc/E/A5naFo2P92wv1n0/PtvSWcZjm2q89q81zzuUqa1+8IjGynhjVQqRIOiRRRoiJ5BEUe+FKXKXqqSoo4oe9ihtPPL3R7yVAwC7r8iIMwsLAsWBew2jDrJcxjzOWu1LHSbcnyNkp6VKgC9fmbBjCJPLZGFedoxcKI1BuWxpM1rq+kyqslpZZhK8MneTvIVZY24d2PUliwUD7zP8AhAC+R10Chw8uhT9cpaMaKnFWL2dQx0eSZ1qxCXNY507QOYRMe7ISSvlVSPENvcQ3t0kkxYOl6GTLtOVOpZ6eEPqJ58sy8RRGMPQwteocAngzzoIeCLx001vtWxpMg05VxtpjsmyKnIr6WjpcvM8rWSKvqb1NVKwP3YjIST6U5vwLYlHaBmscmWVuU5TETp+igWjyyjqEVlSjiQJEfUO4tK3q8pPUXGVPLfcaZE0Hht88/Mq9Cy6PD169FH+zyQ572x5PU0SwvFp+CpzqIkWVp0Ap6QkdATUTKw/wDFh69zfKM4zSs1RV16UOltJ08OT0dXFBwKeFigaOIf7Woqqnv5An+8ZhuIRGYV59HGggddcapzHMoMspac0mXPmVWf1VDTxRyVE07+ZCFkIQcu2xRywxF9V6/h1xWpFlFDUZVpDIA8eS0M3MoUJtkram3DVMigA/sLtjTzJR0F+UtH2tAB8+9Txx8gPAJokuMB1mp9vZG0/qPP8AtO7UoNQzRfo3L9G5dNW5HlobemXu8qw07EgWedppRLJLa7On7KKosiimn0jl1VrWgrZo5aQihyQyC7R5hJGzfWOftfVqcvPu6NNJS8cnGq7MtJVWUwahqGMcVX9epaCtXqqNBTiaSH32tVKGtyGiuOmN12iVkkWZJpCplig/QsT0lQW5WOuqFWoqyx/uM9LT38vqnxw50gE+AwYAAOfqSfJK1n7oVzdWpWu+i3W5XknbPnGdSxPDT6Y0TXV8aW3EhJ6awHmegWx5LAjzxEtfVOcajzuh7BcqqCK/Nc3pY9VVm/cJMzmm8NKX+9HS95JJIfv1HeMbiNMN8j7Sqbsn1Vqvv5pKnUsujjR0EEiFyldNW0s8ccpt/uoY1kYHzDL1uMb/AOj5lOUtqXTufZj9cq5aSpzTO68RKJqqRqXL3Zgqk+OUzVClQT4nNvPEwaYXfUOGoU8QCa9a/BQg329kDrNeWCknbZC+Q5ZU6eyGngifNoDWKqEP3eWN+qy6AN0/WU0Cym3WPukFg7Y3XYdmsrdgNJlkyE0tJT5tTTxmFWWWSozqKMktbcW7oMipuAJkPBvfGg1RT0x+sZ9ma1NNRRUkj1A7xr/VaaJVWIX5AURxQheqO5Tpgn0Yddz5d2K64mzOaMroLMTqisgcgL3RUVNMqj73/lKniS37MpxWbWSzENGRHE1HqQpX0ZKC7YeWPos7S4qrW/brpbRV4JaHSlYmVzLCiKk2ZFu/zCQ7AAxVkWmDelPiA9mW6prdX0k8QFZVZNluraZL+KRqWSOSoVP7xhmdv+S/lixuy/SslPrDR9XmGZzVM+ZV1M0TOSRKa2nkYVAtxeSSckseSWte4Fq4oKqr7OM40l2mrkBnpNOU1PUZ1Q94GMuXTxxUk8V7C4kEpAB5uSPK+LDHhzeyZ4DxGI4kCvimEUN53ifPA8ldHarp+Oh/RlLVSQ1eXvT5prCUxqI4jTSyOaYgW8qKmK3P/Fb9rHLWgMwpsxykZHnYN8sqRmQW9jJRzCNa1F90ZIZgPILKfXHUeoZoavsi1PUTVpzamyHs1jyehqzJ4p6WSqSGgmF+m6iqVX2eGUHkY541H2a12XUNFqHKHeKamnkoK2QC5jmVA8cwU9Y5oHVyvqlSvlieyFga9pOFcPL0zUM1680jqv6K75aOpop6jsf1NS/XM7y7JcozVYl4kzOHKpqynmjT/vv0dOZUPmsK+am9QfSL0RmGVZjkedZFPsFBVChp6qmYqLzRrV0kqHqoe820dVZCv3cP+0jtSoc67WtFdomZS5vpwHTuWUtRmNGBJPlea0RaJqynH+9VJAsmw8vHIy9SDi+M703/ANr+i8/0kuU0VNqKryFK6jhyxu8opqumd6uhqaBx9ujqQKlIbWMTTS07AGJBhGB9neyY689lTgeHWaVxbI1zOtqpjsH7Q8m1tqLIcoz4UVPqKh7zLGy2umWmpM7o5QySUkcrWSGV0kkjVJCEIMexl2Khb6i7NtPaU11mPZTqKrNBVRyl9L53mI+rR5jRsSkdLW95xBOnigLv4A6NDNtHdyLqsx7GE1VpvL9RZZTCmrMyWB4WqHTuqhZ/DFdx4dryAxhj9iXwNxcJrM27RazW2h5NHdrK1UlZkxIynPqhC9RSTIO7+qVt/G0Tqndd5yyMkYbdsxM0Nc4ujqK4Eawdo2jPDZlqTSHNAD9WR29bVOtU9kuaa/0ocn1O6Zfn+nQuV5dnlXeAK6ALDlOdB+aaTbaOmrWPdsO7ikcrsdecvqc+U19Ro3W1JW5ZWZdO8AjqIzHJQzBvGkiEbgt+WUc/eXnrcHZN9KLMtF1VLkGtqCbPsnSm+pU1WJEjzSjpSNvcJLKGiqYANymlqVeIi6qUFrXFrTsY0L9JTTMGa9kmpson1PlcQSnpxvp5u4A8NJWU0zGekC9IpbzwJ/s+9WMoEvWeSSz/ALufLUeslVlY2Xvx8OuS4uzXLq/La40dZCoqNo7oRC6yKfsmMj7asDwRe/nzhOmo0nhWsnnigia+0sd8jW4O1F5PPmbD3xsM6h1TpXNqvSubUFXkuY5ZM9LU00qFamBwfEhJ8S+tlsDe4ve+NfT07d1tjUgAeQ5IHnbz9/l741heIx6681QNAU9/TFRTwpSZV30aL9l5H3OSTe6qPCnJvwCffDWmIleZXDOxO9pCCx3ee49bHyJ8/jhKnqkiYj6tFUXdWDSbgOAwIsCCQbjzH2RjYS5nUz0wiqa+KKn5KUsMfdoD7oth82/PAG3TgEE1GJTUQIkZmmqqaKNuQHfcx+CLdr/G2LB0VlGjJdBx6hznTE2b1FTm1VRyOM3longijihaPu1RWUljI25nVgDtHF8VxKj1UbClVnBHIRbn8umJvpLPcprdNUWnKzUsFFU5RFXVSR1/6mnbvJUvDE6gnvHW7EtxZAoIIBxBar4aLp1408NylgukmuxSTK6vI0oDkHZxpfURnzaqhkqFr6+KaoqZISe5giNOi7FQs0hdrHcEJ2qpuhrXXWR6ZzbM9PaWyLJ85MLLRjOa/bKJ3XgzCCO0BIZmAc7gSu+xJvgM7pst09pzMclyLOqPN9QapSmpKeiyOq+uVCwo/eStM8K7NrqFTuQWZiN7W22Osh7Ie2vNckpaGn7PZqelgnatWoqY6embcQAd80jglAFFlNgOeMU2Mjc7tZTQbzSuWOeNMtnKk7nPAuMGO7091IqfMKXWUkGSPnNHlOtNMyPDRZhljXjEkZPiiMYs9M/JdIxeJi0iKyPIobjLaqrp86qdQ6OpsuzWgSnlroVpovqdQs8myOeFgbCORiLpHuTxB49oBURav0TqEdopbtJ1PQ6ZrszAzZc2TbNTuXYlJIno9yKCytZlsqspBtg+Z9ptTVaynzTN4pqjJqqmfLhTsQZKenZgwmjt4UqBKoqCy2BkLeTG7uwJ/hGoz55A6wcfBN7Ufz4HL9esVPKIZ/QZZVZ/mGn2nhpaOWCgRou8qZX7po1jRwxnKRl97l1JVUC7gGtiJ6E0b2yabp4tVZdVRaWy8SEUdTn9QtHBPUtG0SPBFN9uVVdisoS0Zs25bA418Op8hbteGoqHNEaOvVjUV1Or0yxVk8DJJUJfa6qszd7zYDnqAMT45PDmFRLmVFqHLcxVDJT12YVecxVSUpjsCXnkcMQTyjAFGABQM4xC9zrPgQO8BmOX55KVrWzYg5dVUdynTvaHRxVOl9Z6VzGmp8tpi8ma9yojp4o7uhlkJ7mpj3HwkN3niAjY3ClrQ5NHrPI6/TMVAkOZVckVTlzFVjZK0JYU5K2UrPGSovYiULb7TY2ubZNq5IaynfTcslFk1WMuqayJ2qaaOYjfGzSXCLuQgiSy7k8W4HDDS1HldLWU2b5rnGTyZdl2YQGrWGeScK1iUEikh2XcosU8JUOocMVul8kF4wOeG3jr2JbgBunmtdX5pPq6m/SK1T0sqoq55Qohp1pu5UQxOp6CLb4VTgrIWWx3A4SzqeCtqvr7z0uYVFT3ULU8NJJRuIwgVV7sIE2bQFJQ3JO7rh8mmtQ5bLqSszvMKCvybNY2rznlDWK9NLVq7PEQo8Tszuy9wyq4LbrDZfEbjiDPRMHWKFz+s76J+4Ug+JpChJNzcMygEdRwMWGBle6cBs9OscBVREup3s0EwnzLMF72php6iZu6meqbuIyFW4ZyANhG2263JAJ5JwhI0bUykU7KXcMVmJdzYbvEBZAbHoBuBI9cOGZK6RiWlhpEvvWGUTlE3WNidveDyFyAR5k4SzF6Yu1PQRnuAzAfWWXvdvlv2eBze5B423AuQMTNzomHamsFJFFRNPG8caxIYpSliC97xOw9GBKE2sGUA2vjU5o/iLsWIvwW5Nz6nz/yxtlqqmGnnTvY2+twmGSVoVL7D1Ac8WNhyPgDjR5iq7dyJYFR1J49bDyF7++Jmg1xUTir37GtQay0/lOn+0DQFe9NqXSOoIKvL5L+FnkpAO5ceccv1d42Hnux2B2vdgvZv9KXsxo+3jsN05HQ1+oIZK6p09SWTvK5TasipgbLBXxyX3QG0VUpUgLIQx45+jbmENRkOqcmmfYy0lBXIx+4YqqSLcPcfWVPyx0p9HztMbsy7WZ+z+ojC6S7aKFM5yiBn2xUOpIy0NRAD93vJY5Yj0+3TNxYHGGZZIXTMZmzEDaKVI+N61C1klx7snYHdsXA+ssolySuNHLGOhMbhWUSLci9m5UggqVNmVlZTyMO9FZzVRNLl0dQEE4iUMxsI5Y3DwSf8sg2n2c47V+mf2KprDI6ztLymjUZvRSo2essYU1IkISDMSB0kZttPU8WLmnlNiznHBtIP0dXK0qmykrIhFrqeGBHw/PGjDaI9IWa+zxVSSF1kmuuy668F1n2+55S0OXPm4UxHONRapracDgxrVmhMzfFUMiD+82OYROsrvUzBEM7GRh5IPJR7AADE67TNYPqLSPZ3lk0lTLNT5PVzVksyEGeaXMJW3hj9q6xxgkfs28sbfsE0KM01ZBqHMmeny3LQ9S1WsIlanWIr3lQiNxJKheOOFDcNUzwA3CtaDRwNms5mmHeqRwcQOOadaf3z2xMyArxVrfRs+ixXavzukzLVkFfBTJOkNVBSnu6lpH27Mvhc8JVSB1M0nSlicA/rZAq2b9IzU2R9qGg+1PP9J0NHR6RyPP9P6D0rSUfFOKDK4K3vZIh07tqmpJU+YK3vbE67Q+0iLsX+jdnes6OBcqzZqYaa0zQxSl/qNVVLIXKOfFK0MP1iV5T4pJ3ErG7i1Q6104nZj9HSk7NpIGjqsqy/JJa4H/4OmrJp6kfFWqBF8IreWMwWyWZvauP3uo3wGZ44fnFXhAyN4YB9oqfE09uqLiiG5p4uAd1iNzC3T8ifX288N6lWljdSGG3w2Pl8f68sLU4ZqOM7b2QFQV4YnqPQ9OnxtzgGUlAQTwptYEgfG1x0t18gBjrlgIkUxGU0tYn+1yyr2n/AAN4l/NWHzx3Z9GbPmz/ALGaahhmaWbSmbVGSd2viZaSqBraQ26lb/XEA9ePPHCOXqHmqqBrgVUDbR/3i+JfzBHzx0j9CzUhl1BnWiZu7kXUuQu1PFIxVTmGWSCrgNwCbmEVCC3PixzumrOJbM9p1Yj1/wCS19HzGOVrhrw9vhdpaQ/Q/wClKWLOp3agqt2UVysQ6tltZCaWoS5XcFCShhdjbYOLAY85O0zS+ZaYrZsizuIjMtL5jVafrgb27yGRkv8AAtGxB/vjHoFOtLVU0lN9abupacxiRlEgsQQHHN+hJIB59Acc3fSyySaq19Wa0EMdtd6fodUlQt1NdGrUtcv/AOMUbNYc2mX1xifs3aSLzD/KQfI90+y09Lwg0cNYI4Yj3XKTAFiFIX1YEnj99vXGIwDWJZgLcM1h6245t8xheaQqhQeEHlSGvYe3H54ayAKCw4A6k8AC/FreXrf5cY71csgDKsnkouD1vbnoMHjS62AD7jypJsP7xPl8b84KrBbF1Ibzu3It8rW/oYEsCD41dL+AISF3eZF+vx88IhC7AhXa4UjwkHgke56kehxitLudbsGktuCjgj4fwxm8p0kYKWBIAsfS/ofjgiiMcyEIl7MQOL/16dOuBCE2FhKbpYmwNiD6H0v87YKFUg7iOR5dAPb2/H3wO8IBa43C3BG6/wAvL4f5YAtuAJZmP3rCxJ9/I/HCoROSbCMX6cEn8B5YzGBXY7VazHqVvyPQnGYRCRkAUBGubA+dhyb/ANA+uM+wLMObc+tvn/DAyPaU3A4a3IP4gg4LcsABwOgFzwPbAhYC1rXIuPI/d9b4PxaxI3EXBB5v5Ee2CKLuF+96WIJ9+cKqwALbmAUEjaLgevw+IwqEW5JZtl7eJiFHT5c398AFe+4e/IHG316fvwJCbAxA46Dfz+AGCBuDewK9Wt+BwiEYXaxUHcB5Ei1vh1wkzAXAN/gcGc3+1a9rgW/qw64Rme45N/QnqcCFvtIqIIc2zqS/6inFNHYclpL7rf8AIji3vjoOlyEZHlEeSSJtqKCljpVcSWvURhXlUg8EGWSYAjxC3mL4qHs7ydamXTdDLHuhqswbNKtR/wDA8F2IP/JBL/4sXjUfpGajR8ziiNTUs80g3CVXkdi7qDa4IZzxa/mpOOY0pLeeBXX6Yex4rbsMdG13fn4SlXXrp/R+s9TUs8ZmpsiOUUTouxlrM2l+qLe/IZaaKvawJHiuDY45nq3jjyDN61U5rq2KgpyOLRRDc1vbhBi4O1XNWybs4yXKFk4zrMa/Ukovy1PSj9H0Q+G9axx/ivimtQQyUNLkeRSBUelo/rkwv9+Y7+fcJsGJ7BFQNbtPIYnmBxUFpfVznbB64eleC1ESdwSQ1tp6+d7emDMw2BlW1uevHPHT+eBJLOu9z0te449PlhOQsAOg+6bDp6/H3x0Cy0fJZvq+fUU9+I5la/Tzx0B2d5bR6i1oNKZhNHHR6mkq9L1Ejmyw/XSRSzH2Sq+qt8Acc6LKaeqjmtzG4br1scW5pmtp2zSGOraHuq5UAeVdyJJ9lXPwZY2uLEXvjJ0gCyVko1dei0LJR8bmFdLfRR7Q8+y7RseRZpPPRZ92b5ocmqqeRQWjop53en3g3uIK1aiI/sioQAjBf9I3oLKNQPk/0h9K0S09FrhWgzmKP7NHqGljAlU26d/TgOP2jDf72IdrHP6fSHbxkPaxWqaHTvbHk7Qaj2gBaSvLilzIn+9DXQx1Y9mX1xcVXDJr3sg7R+yLNqWIZgcqmzyjp057rOcnLSvt/wDSQJVR8dQF63xnSE2S3CVv2yU49eqtspPZSw/cz069Fw1lFKh7EM7lEKGSbU+XR77DcFFNUNYH3v8AljOyVKqizqo1PBD+tyxQlET9yqkukb/FF7yQe6DDau2ZRonNcnhmMkQz6mmiIPDRmnk2N8drDEu7PKaBMloaQ0zzPOr1zqgvukbwRKfiBtFubti9aJbsL/8AydyoPbDzVaKMmRp/pHOpU8pM+k0H2Z5xrimMzZ3qGWXSGmUUXkj3RD9I1See9YJI6VGHO+qkI5XF06D0xQ6Ipsg7I4YxHU5KY6avnhKuJc5qnQ1Z23BshMcCsL3WlcW5xW1Vl0dF2zimtHU5F2F5XBT3Yb4qrUMstzx0YtmEzsR5x0Z9MO8m1hPoiPONQLVRZkdOZfNmMM87F5BXyN3VIVe/LPUSiUo1wNjEW5xkW29NE2CPXjxy5U4q3ZSGPdM7VgPdUv8AST1n/bTtg1Rn0b7oq/Psyq4je47kTCnht7d3TLb2OKvKxFiSOG55FyD7HDnMagVFa5EnepEq06Mx+0EFr/EncfnhobggC1xzx6+g9cdVZo+zia1YsrrzyUdo1KhklHgHKMAGA9QRww/P2wUtzYgNb1FxbBQQSVaxFrhhzYj+rHGM6FeHW5Nirf1yMTKNZ3Y8T8EAfZ3WNj6X64BoJkbZNCfDa9rHr06YXiindlMdM0wA8oyQvvfp+eMkpJFFpHgivyA8qLz62uThL7RrS3ScgnGTZ5Lkj1EYp1qKarheGeBjt3gqQObEixINvOw+OJBkk76k0XV0daLVOl445IqkqD3tHJMEMDX6skkm+Pz2mRelrRURJtBerVyot4Ed7fOw/fjY5Bqd9NyTRCjFZRVTxPURMe6kDxklHikW5jddzWPI55BGIJmXhfYO9h1wqFLG6huuOHXut3pPL4ZdXZNSrJ9ajkzCn3mnCyd54wdq7ztcnaFG7gk28sRnM6uvzLMa3NnmmSpqaiWeVrlX3M5Y7redyb4lVGKrUWYT5vp7Opq+sjdZRBMgWtQKQyssY8EmwgG8ZuLXK2vjfaqoNG1a/wBqsy0xm1PVZm8wq46HNaf9H/XuS6oVR3iLfb7pjxdgrFRcRNlaH4jE8evhPLHXcFE/0ZrSk0zDndbQ5ictktKKiaISRlGO1GO65CkggMQAx6E4W0xkObTS02c5bq7S+T1FYJIoY62vSNyrXUlk2sEBPQtbkAi1gcG1DnmZPlE09QiUlTqVUZzGNgOXwkJEgHXYzx3t0IgW3XEPhoQUDJMxueRbgj1BHX4dcSsZeachXcmOfQipqplmOks8yLLF0xqLLMmog9Y1XBmUqb3qCqbGiSpTcpjH2jGbMGNzbGjeOTIqqOGk1HFD3qCQzUlRI0adeG2AHdxyLHqMSGftCrqDJafIKuty/NaevpEWqaSF3kgFyoSRWIDTRABo5gdwBC7ivhDbUtH2Zy0kNHp/OZvrKzlUqiJ5e9hJ6zoyKI3tYgR7h1B/aw0F32vGewJTdzatPNXR5jxmGoIatiwYuaFi/APG6wNuenwwxhyqXM80pcnyNJ66srp0p6enjgO6aV2Coigm+4sQLe+FNQZNBp3Oq7KKXNIa+KlmMaVUSlFkFgbhTyOtiOeQeuJ92XZtk/Zpklf2o11Sp1DIsmWaXpkb9ZDM67ajMWP3BDGxSI8kzSBh/sycJI7sGXmY1yGGJ4cdgxQKyOoR4nH5Ut1vmFBonK6Dsxy6sQ5XpTv6OrnpiG/Sucy7f0hLEejKpRKZX6COnJvaQg1DneoKjNZY8vy9fEXAjhgJZVc8Cx6vIb23+V7LYY1me6hnzdxtUIiqIlCiwWMdI0H3UHp1J5JJOJH2V5QErJ9V1N0TLSI6Ntt71jglGsevdqGk+KoPPFMQCzRmeTF3ufk9UwVprzK8RMy9lPDQrQ09Fo+LdDlmnKeVZp1W6z1r2NXN/eG5VhUi/gpwOrYXrdOQVGs9J9m7xmSRszglz1Wa5Mq/rJICfSCEMh8u9aY+mFcqeTK6aq1pUi9LkbRR0EMp3RzZk4ZqcFTwyxLHJUOPPuo1P+051PZZNU5jrfMc+7yaV8jyerqkJJLyzylYRc9SzGY3PmSfXGeQ6KF0lcQDxOvzJqrXdfIGUwJHD8BdC6Pnr6ODP+0DM4YxU5rXT5JRIwt45Qs2Yy+vgieCmB8jPN6Y1eutVZRpvT2aaizvL2nhevfLMsoi1mzSpCKSEtyIlEgMrdVLlOWcWearpUymkjyfNc1ajyTQ9GlFmdYih2FS0hkrWjU8STzVkkscafeMak+BHIoafUeedqutHz6rIyrLsmg+q5TQCTctJAC3d00bH7Ttd3kkPLt3j9SoxnRWVs57Q/Y3nTV569nBW3zmKjG/c7kl4dT6jzrR9bpF6fuv/LD57nEq2C1dVMyRU0WxeBHEEdtvS6g/dxMtLZdlNNm8NXmNAz5bpejm1BXxuT+ujpgvcwMPMTVLU8fPJ7xyBzjX6EyiWDTsucNHGi5nmk3cWO5mSlVYVKqLkkPJMRYG5A8r4kNfQUWS9muYV8EEsMmpc2EW2RNjtS5eiu5cHnxVlXEeT/77j0xZDm3zhQDHDbhXnlwUJaboNc1YH0WQ+caVyrNdQ1yzPXZ9nmZZ00kW8VMUctNNUG9rX2Qykjg2Nx0OI3PnMH1fM+0DVOWQ1Jp2FcFll2/Xs1rZHkhonH2XRpWeR24KwwSg9Rht2FZlWJ2E5/DTVMkf1enznLVWJz4pa2ro4LgftmKSZQetmI88Nu0hYqnU66DogJoNM008lZ3S71mzOUrFUFB0bugFpUPS8NQ3nzV7P+8yF2V4k8cPVTh/7lt3OgHyq1oshMuW5zqGsqnzSsqs8rx+kgBvllj7vvJSx5KM0zsQPUeuLM+jfRR1ejc+WW8L0VLmOSGZL70lr66gRST6hQbHrfGr0nlcr9nFXRGOFJf7Y55RxwFxEIrQ0bgpIQEFhcbGsrA8EEDG2+jhO2XQdomQS2A7/JM1XrtMceZxwzkeo4Xn2xYtMrnRSAHEEetPQqKBjQ5pOw/Psp52zQSZlmGcIglEmodQRd7GEZ9lPTRRS1Ullvw9S0C2HJ7tupGKN/Q+f6eyjUGUZHVU81Hr2Cuy/MAqnwx5bmMLo0fr3khW4I+yTi2M81bT592z6r0bGCI8i02RG6MPHXQ1grKy3uGqpEv6wD0wwz3LZWp6bUEbw1EOW1+o0qNgdGkTZl87soIvwrlj7Ankc4gst+zDszsB5AjnTgnTXZu9vot32HZhBnXYbpjVS97NmuiM0hocxczfYhoa+KqplCW/4ElQS5Nj3KqBcHB/pK9nkulezPUlJlxEu3MaLJYik6tu35nOyK6EBo2206G5JV+GW3IxFOyXNaDRnblqnsknVYsj15I+WUrSMESLNliWSD2USGeWEegmU+WLU7aqir1J2IR5xJTwVdZVZ3p0zLPCR3WYUs89HVU8jDkpI8MMm08r379L4CDHag4faXB48z7HkkNHQ0GYFD5fIVVdhmoMt1B9GPtIynP6l48x0Vk8VDUxiPfI2SVOZQSpMFJ8f1WpWQMtwdk5AxcSaKp4IENQtMDWTJRtJM36lZr97RSqeQ0JeYxtJ0NPXsx+zYUjR6FzTLczzyq0Eab6jqvIpclmp6h9okyzMYkaGTfY3eGRUBv1enfob4lP0Xu1OXO6GfsP1xHF/aHI4ajLsugqhc5plyM/fZcT172G8pjH34mdBdo4xiaeO+XTR6jUjcRjzFfPcmROLCGP15Hwy63LQ9tXZNS1GgMyqstjmUZYBn9Mgp9rJS37qsjPoyxlS6dQ1HJ6EDR/Ri7d8y7IM1ptDa9zGSkyOKZ58ozfue+fI6hnV2lUdZaGV0Q1EA5FhKlnU7uxNHUOXzQS02YVstRLlspkrDNMHp8wpqlGjM08JBBFRGskczLYpUxyuP8AbqG5K1Z2MJlGb5z2f51SN3+m8wky6KqeW6ywgCWkcyLfb3lM8TLIL3tcBxcBYJ2mJ0LxVteRyI6zKdLETIJBg731q9KnT2TaK7S9U6MnyWCTRGf1VNqCgSFhVUkMOYx9+rU5vtcK6yoVBCVMKXUrPEt9H28dhVJHTxduek3jFGxeLUO6EVMUbIRHJVSoQBLH9iOrQgHlJ+CXIput1f2k9i+nMpoaeOozygyyeSkyJq2l72F6aYrLW5HXQhrr4hHVQhGtu3SwMA916b+jb9JDs/7QFmpMmplps3qoVhzzStee/qnjRSgqKTcAMxWONmjkgIFQ0BK2do1ZpHRSMPbDEZE8MT1gcDgoxIxw7M4Hrr0xXP8AL9G3TWs4KnJsugTKJ6pJKqCmUvUPQSqt3mpWtvraPj9dD/51AtpNsioWfmfU+X9oXZHql8izV6vLcyyiRWhlhqCQoZd0ctPPGeY3QhkeNtrKbjzx6H6qynR/YbqKLT+e5nFl2i8+qFbTea1czvR0s6BZFy+epTx00sKsr01WLboGS7I0bHDbtf7Ici7X9KSZRnyU9JnFDF3mX5xOEiVEdtwSrZLxmmldr/WoxsjlYS7e7kmxbs9sMTgycVByOrxGzeFDLZw8F8JxGY1riLU3bzqLtZySLLu0zLafUGe0ESQZXqWULFmcMakWgqJUAFXDa4AlG9D9l7eExVKjJKP/AOOGXZhVP5olYkC39LqrNh5qrQ+adnmeVundQZdW0lZRyyQPFNCEkV0ba6SKT4XU8MBcdCCVZSY/ehnlBqqyqiYC1kplewH/ADjG4GtDaNrTdVZhJvY5p7PUZPVZjBUU2njl1IIjeKSrlnWdw3LbzYg8gWFhx5YGszpItxynLcrpwgux+oAm3qGkZz6YXipdOy02855mSpGwViuWJ9pgbWvN6KefbBoKDRSG82bZ7Lf7seXQL+bSnDKsyoTTx90/v51A4JvLmmbIZKHNcwqA6WDw7wqi4uOE8JFiPXE30hpqh0pldJreeenkzrMIlqcqgkRGiooS7KKmQSAq0jbGCIQyoDve91XESyPRmR53QnOcwzeryXLKGqaCsrZqZZw5YgwxU0MZDST7A7MpIUBblgCL2JLm2mhJktHkup6ePJqaBaF2ngd/qUClikk0wiaQyyFmZlhS0N1Fz5V7RJdAjjHjQcvPrPGSJtTeefDFOci1jVd7VRQ5bp/JIq8GmzGTKcrShnnV1syB4WMg6q3dxd3G/IN/smKSUNHkoicTRoksxiLmJJCoJs9jsCtdSpTxAsA5G3dxKKivo6bLjkuTZ1pWCjy4vXvMtLPR19futGyJU1Kh6goh8MAMYNtwVjzjX189dUJWQ993AaaWEoZF4XcFaOVGHdSr0JWyvGT4LjgVGkhxNKA9dcFORhTMpzqLK9L66bJaXTGqKLJZMtofqFNl+cuI4ZD3ryF0rkHd75XkZmEwj2swQMVAOITR9nWfZvrim7PMwjjyPNah3jZM0jePZtRn+yqsz7gtk2A7yVseQcbXJoVy5Y8xdNlQrhopSqsbgso+2LbiwI3eNQbC1ybTuh7R9bxZLWZNLn9dDlxDUipJtcxq4tJHAzKZ4t6XDJAyDYSDYHEzZXwNuMxG/b7+eO9RGNspvHAqA1PZ52nR6djocliOp9MQyHMKepyULWRpvUbn2Kv1iE8AOrqtiBuHQ4RyTs87Wngihyzs6zGSDVyNllLUVeTFg4V0dnieRbxbSAWlW1l3C9r43uWVMWX59S6menzFcvy2pVJ5MsIpaiFZvCTC8bK28leAG2gLt4vfD76zLqSuz/OJBndVXzUDQZQtRnlS8tRGhE0tPPKWLSu1MGKICE3KQQxtcdaJAKUFPD8jcSdiBCyta9cFsNY5dQQ6niTM65o8ugWnpKKWohRpauOnjSGKWohRwHp5DEQe8IcKQF8ziMZ69DUZ1VVVBlEuTLQs4ipkmKiiB/2l3kBl2k/ZElyEIHIAONRqOCalpaOsyp1nySti72inigSJpIybFJinLSoQUYEWBAP2WBL7QeetS6hyWprcpp9Qrl9ZE8VDJSd/M0CEl40Ym1hztjfcAyqQAAcQsiLI71a0Hh+h2+e9SukDnUpRR5aWKorDLQQRxIGvI/G9ibWBPRR5BV55uTY2w/GVQ1KSmAd9V00LSSGFS0rU687+LXMYsGtYbdreRxJNUUFPmWdR02pNWiunl7yoyPUNTYR5lQOxtDKXt3FRFJuWz2VGLoxACEx/O9PZ/lUSrmEdL9SnnkoJJ0q4Z6eCdLEoZVdlikA5+1ZlJIB6Cw2QOoK0PXW7JQ3aY0WhQVQhNQhYRNIEUyWUOQeBt6m3JPJC8dSAMPaqTLIcphmp6moOaioRtixKIool+yyy3Jd2foAoAAJNybA1W2SvU0lPJnUfcySxwVFQkEjJTw7gNy38clhuNgBe3HJw0z/OaGsrZo8rQy0qTSFJ5ohG9QLkCaULxvZbeFbBR05LHE4q4jBRnAFJZnWy1la1TUxRBjL45YUEaSPa+7aPCpI54AB62HJxpcycEd0rHYpJXwbbk+ZHX8SbeXrh7DK6ymNp9sUpCOSNykDoQBflT0A5HwvhnWeONmW4Rja5HN/Q+h9v34naLuCidjipj2IZpPR59U0cUhVq/L6ulS3H60IJ4v8A2cGOoO0zSr6u7Bq3VGn6oJnGjayDWOXTQLtlhhbu4a4LbpYfUqkendSHHGuj83kyHPIcxjF2pJo6tQfvd224r803j547Z7Cdf5OsVJk2ZwrVUNJ39BWQ7earL2V0dG9d1NLIvTrt5+zjA0iXWa1ttDctfutayBs9nMRz1eyvHL+1DIu0vs8092j1sX1ik1Tk0qZ1lEYCwvMhNPmtKQBcliWkUsfD3kJUcA481e2vRlRoTtIzbTslSKlaOteJKgdKlOHin/8AWRPFJ8XOOn+yWgzDs5zbtM7As0qZZ6nReZ/2hygoyj63REJBUMpNwBJTyUdRfm4jbEI+k5p+jzDKsj1hGlpWpZctqGJuTNQTKFJNhe9LVxc+fdD0wWIDR9tdAPsdiOuskk9bVZmyfzDAqqe0CTMMxy7sxytLOYNJQw06285K+rbn5m+Onex/JaHLMuyHJI43UTRQZzVMEu31dS65dF8Tuqa5vVp4D9xbc8ZZllbnWo9EZdWZnSVm/TtJBRCnjCGmWWeaOONz96RWldi3uvpjq/sk1NQaczbUXafXtKuU5JT1WeSRxSFA9HRoFpKfj7rLHTRAHg7xiHScjnxiGPfxJUtia2NxleP0AWo7Z6ql1/8ASj0B2H1ELTac7NlOaahjRRskrAi1tchubXWOOlpLHoQw88ar6Sup5q7Is+aqkvUSzUAnYrY981RVzy/LdG1va2G30bUzHOcm1T2u6mkWrznUlbJStLI21qiRp1rK9r35UzNl8R9hIMVX28akepyWoiinlcZnm9VJG0kxkd4qdFplJY8kGWWqt/gNsNMQda47OzKMAeeaVri2B0z83mqo2lUdxFscElALLyb+hGDhgvi2lUUbdwU2B9OP6GA7zYkcaOjD7OwOOQfUeR/LBn3oRHKmxgDdZLoQPLhrAX/A465YCYPM1HmEFX4f1bq3ha4IBxPOyDV8nZl2oZVquFiE09nNLmZC/ept4SZfgYZGGIJmEO2Mg+HabFR6/Hz498O4ahTUZfUufBV05o5/zQ/kVPyxUtUYeKHIinv6VViB1Outy9NMwoa7KswqMtgP1oUFU1OsqOid8q3CsGsUV2j7thuG03AYWO4RT6QemjmfZTprOUXvJdK6plypx3TJIMvzeHvEVlJO3bW0jC1yAZQoJ4Jb9lmoo9Z6G0lqKsUx1uZ5JFTVdR3BYy1tE5o5gzL1YolM21wQ2/qDbE1z7LajU+i9aaCSocT57pavGXQ72KfpGgtmFMyE8k7qVwrAkndtazDnzvR1bLbexdtLT54etF1lpPb2YSjVR3D8VXmpV0n1OrqqFkYPSzPFYGzWBO3rx0t8MIPa/AUAAEWv1vze/n+XliSa4o6b+0kmYU0arT5pBBXREqDYOikgX6dV/HGgAQtsJEajrYG3sbevl6fvx6TBJ2sTX7QuQlZ2chbsKQdLbiAtl2gDr4b8dDYfutjD4LylDGpFmvcj8xz+/CjMIWAUkFeh5BBP7z636jywmrIXuxCeMkNc2Bt0/ui/mL4mUawA3KhCbjpa5/yPwOCnwOWBA4ALXvx7X6el/LAhVZCgUruNtrG9j58jj8fxwJJFnjYXUg+FgdhA8iOhwiER5I9uz0P4D91r8+owIvGjbowLcknnr056Drx+eANiFO5i1zb3P7+l/wDpgG2gHuRa1hfjz8+AB8PXAhA9rsxJAHFyAOfl/njMJs2/mQgkeXmMZhUIjCylbggG/Xm/9Hp88AFBJB8r8Hy9j/XBwLdd32z0selrdPzwK3JJJBta7Hrzxf8AH9+BCALu5KFgOviA/jg3eWABCjy44+f88EVL3AHPSxI/C+M9gTyB1HXnj4HAhH4t4SAQbKFG48c/0PPACwbcjtzzuHF/XGC4JYbdvT2/z+OCkk+Lm/Um/n/XrhEIrcAWX4G/nhB0dyFRSzObKB1JPTCshBNiOT68nD/TCo+o6F5RvSnl+sODzcRguf8A2uGyOuMLtic1t5wbtVr6EpDTZtmrdzK8OVUtNk6yKLojyuA27zAKQT8j9s+WJzmOZutFU1m0VTwxl4Y0iKlnPMcfeKf9p3hTqtjc+IG+NPoVGpNDsfqtQ8+c1tVmEs0UW4Dbtp4txv0G2pYeveYlWiocv/thkrVUqPRUU0mb5hCRcx0eXRPXSm37D/V41B45YqRcDHITEySnDLDhnzBW/GRHGDXPoeyqvt0T6/2jroTL5u8gyIUGkYHFvEaONUqJOP2qhp3PxOKy1BXDMtQZjmEbKYzOUisf90vhUj2soxJKTNqqpz6v1XmJMlTSUFVmkzE2vV1JNvnvmv8ALEMiXbGOTdfbr/XvjoLEzv4/yinmc+QafNZM7u74n0/UoTYAXt6kdL3xgu1rXFvDfd69P5YEDzF9oFrXvYe3qPbBWRTwLAEeX7/fGkqaa1C26DgYlGTZg1Tk6qrXmoiTb1W3P/sQp/5D64jU48J4HPl53wplFe+XVaybtqNwxte3obedvTzFx54r2mLtWYZjFT2eTs345FdJmpk7WewvOdLTXqM20851ZlR+9J3MKw5lAPd6VIZz5k0Lnzxbv0e+0KKur+z3V2Yzxu+ZVNFQZiWG4yTRzJQVe4+ZeCWGU/8ApWxzN2c65zDs91FQZ/lKJKaWoSrpo5DeNnXkwt5FGUlGH3o5T6nF26KyPLNJ6h1NpLIXkbIJpcl7RtHSMbscrmqooqiInqWjSREf+/Qt6Yw7S2/Zf/Qgjw/HstKIlk9D/MKFc667yN9O0eoMgckyZXmsFI4890Rnh/8AqeLP7Bo6bL9Vf2kzGJZaHR2VnUVUhtZo6Gn75EP+OoNMn/PiCdubyzdpfai0ZPc0+qatto6G9dUWv/4uMWb2HLkc2hNUVmohfL89zDKcrruORlNHG+aZiAevjWmpYuPNxiV0ZljF7I48Q1RmQNLhr/JTygnqNLaGy7T2bVix51nEzax1LUTmxatqoyaVH8/1NLI8pU/frPbFfdp2d0+Sabp9MUMpNRUSpmeY+RNSyEUsLe8cTvMw8nnUdVxsdS6rk1RmcupsxqYIazMZ5s1qhIbxLI7lyzD/AIMV1S33yiRrxuxTmd5w+dZo1SGlenR27szHxyljdpHP7bnkn4DoMMscRntBefPrrknWh4ihDR5JlAqWCbgoAtzcn8AMLiOmBVZZ53Y9AI1T4eJj/DG00vpDPNX5vFlOR0jTylXlYltqRxopd3djwiKoLMx6AetgbV+kppTLOyqryPsQy2KOSo0xSR1ee1zQ7JKzOayCOecEHxLHDE0ECIehWRiNzHG26VvaCIHH0WaGG7fIwVLpPSKt/wBHKxv1nldzf4DaPlhT9I1CACAwwKebwwqm4efNr3+eGrhkY7izXG07r8m3HXBdrBWcXOwXNvLyviS405qO+7UlZ5Zqg3eaSQjnxvuNvxIwVWUm8gAJ5DKtrjzBA8vfBB5mRRuuLHptPphQBT1HAvfm17+nph1KYBJXahdgoZlsTbpc3t+GMdELFX55tcKLX9ucHghkqWK0scs7D/hxk/j5DCwo5IuaqppKYjjbJJua3ptS5t7H8sIXtGFU4NJxWtMFiXjJBHTyN/l54d5Rm+Y5ItQIIIqijq1CVFNOpMMwBupIBHKnkMORzzYkFdYsjS8ktRW1hB+xAqwqPYs5LfPbhR83pQCtHkVDGw4D1G+pf8XO38Fwwuvil32/KcBdxJ91Y2kNY1tdSxUen9SSZaHliply2WXwUrTSbQadyrExB23BCQV6eIHdiL6vznJc/wBZZ3m6SxZfTVNfIY4VpWLgLZdzIoCBm2lyAbXY8DGlyvWeb5RmNVJStJDRVsK01RBEVBMYdXBU7bBlkRXU24ZR5XxNKmo7OtdZgghymSkznUc1RK9TBWMRSzItw31TaFAmcOTGsjkbgEtYKanZiB5fSgOzia9FTXu0aGqAVqacExLzZlUHyskcIP4ljhxQVGTrA7x5DTSsW8LzzySFBbpZSqn16YsDTLUMLwaZ01k+UZo2Y08cZfM8tV2zCr3pMsYLnfCjd33SBSrMGO7mSwxtSZFraaNe0aGeOqu8UFfliRQ1UcNyVj+r7e7mjQkhVPdsoG0OQABJfGuvHHgEwA6qKAtm5EbCOjo4Sp6xUaAAeXJBP54jtRLJO9mcnbwoJ4Avew9MTrUWn9MDIqmXTbajNfQyRvUJWiDu5oDcNII4zeMo23qXFnuSPOELEipvfxE4mjuZtCY+9rKRghnnmSnhieR5GCqiAksxNgAPMk8Yv/8AspNlX6O0LlCGeqykCmljij3mtzSV1E6oR1YSbIE9e6v54ifYnlFPST5p2l5isIg0qkQy5Zx+rnzeclaRSPMR7ZKhv7sHPXFi6Vnn0NktVr2WW+a1PfZXp2QuWYVJUCqzDngtAkgRG4vUVCsOYTjG0hKZ5xE3JvqfgH/NuWjY2COIvd/N6fk+ib63pKam7nSdHPDU0mnzMr1EHiSpzBiPrNSnUOgeNYU8jDTg/fONt9ErK4o881Fn2YtSRUeUT0dVUVFS9qaJaVJ6rdK/lEssdOWPmBtFywBryprCMsFLAHEKhoECeUY8aDny2uB8Fwhkuo86HZrW9neSK0X6Wzr9KZo6IR3qRoq0sZP3lB7yQIOC1ifsrirMx8sLommlSMdgrWvLjRTRuDJBIdVfSlOal3av2h02tJY8qyGSoXT2UtJLCZIyk9XMwPe11QP+PLztTpFEQg5ZiY/qT6rotRo4UsgzjLEWXM2t+rWqYB5lv1/UDu4bdNyMeDiTdlOnKOkmm1TVxLUZfpaBc5eGXkVFZ3gjoadvUy1W1mH/AA4JvIHFfa9jqI1qpa7MpqucxzNJMzHdJM775Gbz8bszkHzPpiWCKNjREMuuuaY97yTIrd0ZWUmVaP07FUJAkUOQx1MrTAshM001Q9wLXvxxcAkAHyGNx2p09U1Bl2Q1q3bINN0MdSDZQa2qR8yqjx576yBf/Vj0xG8uy/8ASOXZRkoPdtW5NlVFGVuzuskUakAdFG6QAsR7A3NsTrWkVPqXtBzzLMpijM+oc+zKjp43cgO31n6tSMD0G1IwjjyXaeRe2c3AOfrcfequuBIazUB7UUa7BdT5dpDsw1nFNWU/9pKfVNEcio5DcmoaN5FqWTzip2jE58i6RL97A5HlK5XWZ3mdRVzS0FFC4kL/AKyVaSmlWCpm/vNHM8dQw6kd8epxlDkGk9K51qLMpYGloKbLDmENaCC0tPTSIeT6VIClSPNlX1GHGhqgVWlcj1NVpCayroZqyr3myT/Wqmu76FrchJQ3dkWJuyEXIwkjw+9IMiRXxpq4c96GNuEMOYB65p92aTtmWntUUhAeGi144ZSDYCpoiL7gD50oI6g298Zkmo8q7JO1HROfagpYpclzfTFbDmEA6VEktbWS0gJHUCopaW/93jocRrsYVcozXtC0RXrJHNTfovMoYp27x0eCp+rndscfrBHVKCVa4YH3xO820/HneR6Cp4B39bJkcO2kkXc7PHW1tpoBb9Y8ZdlkjHLRSXtdQQOAikeHfbQDi3NI2rmtpnjyKgGnq2k03209n9VnU/1qLOKSLL8+ma9nqM2WZpdx9jMjfC2OjNG5HS1eUZjQZkZqpxnk01ahMcFVFUz5emyWBgFjDiSjkAjIVXUtG1ib45c19kdTm1LnlegqKd6bNDDTP0A+pQx0yOjA82ZCbj1+GOo9Faty/VeXrrSubvqfWWQZNmBWBUWSKvgq5KOuRl4BEc1QJGPXbIp6nENtxax+ulD4/d7u4JYcy3fUenxxXNXbHph8y1HrXMKV3WejzSbMIXS6PG1G0ME5TptISRXt1HdH0xe+l9a/9qv0c9WaiqRC2YUmc5Jmedwhwpp81Sup4qp7fdiqkEdSr9BItSvkMaA0+XTajzvMdQxRS0uWaxzZMwBgMne0M9XNS1a96D4CsUzSWINyic8cQIaJ1p2eHXuQadqKWojzzS2ZZJnFDO7HvpaZyz7NvWVZKUzRk8EMR0bFmrJGNY40LTQHhh5gV8fNREOBc5ozFeutaurKMjd6SrymKKnpKjKGZ4Y6hnhBoqmUh4+9UExGDMEkF2VkArjuAU3xzd2+6Gr8i7Wc2zbIxWZRWPPR6ko5mPdPAKtQ4cMh22SqV7SIdviVhxi9Ow7tlodR6byztFzY/wCvaerEy7WEauBI9NPGIpauNfSaJBNbkCppD/xBiU/SG7M4Ms1PpHPGyyn72ejzfIJqena8NcsEwqhCE6Xlp5ql1UAcooAuRhsUklll72YBB8qHmBmpHNbOygyOPH4KP9Gvt1PaMstdU0sado+SNLLn+RsBC2dUgVVqpqRQLByI0eSID9VLDHMg2bgst+k5BonTlbpTtjhzialyPOaWDIqjOqOAyQxQuXmy2pqKdb95T3+tU8iL+siKKUu0exuSO03s51LpnNqDtA7PMxrI84yYw1ArqOXZPIghM9DXxOp5aSCOVGI6y0kt/tkG6+xrteyP6VHZpqnsF1K9FlmqtSUVQ9NSooip5szus8VXRoOEEk8UbTU6/Yd2lhG2SRElbZ2OPax/Ycx/TXZu1jZRQvlc0XHZjI7fzqW/qcjodTZfPleaUFGyVFHD9ZT66Fp6qIsXhMdSgKqNxE1LWKCEdmuO7eeMVHrn6PNLm0bZ5RTS0ma0ztIuZwRfV5m7pgrPJCD4J4mZBLGCGRmQo7JLC7aLsu7RM97GqChotVQVeYaGqlYiQJ3lfpuSR/1uxbgvAJVYSQngspZTHLy3WNA0+cZQNX5ZTx11LLUUzQHLpUnTM6J1eOGajZrd5NEA4RWteN2pZgA692yr4XBodQY0d6g+481KA2QEkY6x6EfK5+oO3Sq1HkeZfRv+lsFWLMqKKLKdbqhdY2BZ6GqqiovJGrlgKkASqrzRyhgW2w/RXaT2hfRvmIpJTnmjKetNDX5Bm0neLldYwJaESqCacSLuaKeMNDPGbMjeNR0b2q9kia+0pBqXR6U1VnOnwczyJ8vZStfDKN0lPH3t+8SoVXZFdSVnjKOC7yriGae0npjU+TN3GnaAmuyvbUZdHuSjzfKHUMZacm7RxxtsLxG8lFKEdC0O4LYNobdoW4fzDftA5/oaw9jQk18D7E9eia9pWnOz36WGixnXZBUzyapyaiV1yaqRUzKOBBtFG4BIqAg/81qEJV470zkEQW4dpKAz1poqqtgoHjZleSoEmxSDa3hUsPPy+Njida6yfXf0ce0SkqdO53mdLDFIcw0/mkMhgqFQNY2Zf9nOjApLGPDuHIZWUnO0HtAyXtXro9fHL4Mr1dmEhGfwUtOI6bMKi1/r8Ma8RSSc99GPD3n6xLB2Vdyz3mtpWrTkVmzUc6tMRmFFJtMVEKlKTU2nZ0Zg21czVCSAQOJVU+Z498JDSWpGqKSBIKUCuqYqSKT6/A0XeyMFXcwc7VueSeAOuEKgi7RNKe8BvtZrXHwPnhnPAO6Zu7UttIuqj06Yshsmpw4fkKEuZ/Tz/CuPO8qTLGg0hBkGbS5XlaPQX+oveecP+vn4DeJpQWsescccZtYEaSu03UZLSUdVD9al753DxiiqBLTneVVX3qA29QHAQsrA2NnUDB9UrFBm801RW0s1VFS0sdU8Ti7VCU8ay/Z5Zu8BRiOrAkni51LZvPmNJOYKqoFljXulrJiJCbgDh9jc3LGwCgCwIOMyNryAa55q28tvEUTaupJ54UYIWgk/U94tzEjjkqXYBC4BsVYllBFz5Yf5fFmOaVtJkc+e5RSASIDWVWbRQwxiIXHeTC99q32HaX8l3cYM9Rq/KKV4KvMc7pqOrk7qakWplginiKDcGjYlTxtHiHi4554XrKBpT30a1DiAJHEUG10ZR4UG4fqnLXLBug4G5bHD3OwpgmgYrbZ3SDOpoqCLU9LWUGXomU5dXwVLmGNyS8SGOo2yrTly4DhFAZ95uL41kmdI9XWVsuS0eWFpWjqKSASQxUsi2DQ925IjUMrMFLBdxIB8sa+aSSaOoqM7pKGCqaRZNgpljRImj8JFxcRkAggHdcbiLm+C0ef0WdZLJlU2X00WYRVydxVfUJZswqad1Kukk6uAyptQiNkLNvtu4w1jDSmdE5zsahGXMqDM8warqAVbu1RXVLMihUXuxHcC1rjaCALMT1FgpqqpbO6Cdc9NAaecPRzVlWix07L+sVludq+JQSACDcrfnCGUGpoauKaGhpKhgwO2UOVcEtsYqWIYXHhHAbYL3GJNTZ6+QabhTL8y7rOauuLxy00QSWkpkG2UyyMhEiyuY+72N+rKSk2vYueS3Bgrq6z63pG44uSdXlBpK6akoqnKoMiz+kbPsqiraoRLS15bupKdOPF+sDRMn2Wi7tiQADiHVtC81WadqaeKpSVlqGlIjIkW4Ze7Twwqu0jqWNvLpib1Mdf2iZRVaOzTMqrMM6nlNVlE9XMZJDVRptNNvbkCeMbAgsO8ihHVjiO5dVUeocnfM6ykdMyhqEoq+wAjkLRkxzbT0kIiZXXgEqG6kgpE5zal2fVD55eIQ8A4Drr0S9PmlHX5HHQ6myp8xy5KhZonjqhSzx1OwLJaQq4dZFVDItiQVDXBPO6GejOqOvTQ9N/ZqsoqKWtkTKqTuqKeGBNz09VG7NuaIE7Klrli5VgCQcReuoazN5o6egkDzRRh5O+lCiCmAuzuQNkcY87Ec/dJNjpq7NqPLMkm07k1ctbPmEqS19XArrGY0uUgTcAzDcd7EgAkKALC+H9kH0Az5eer5Tb9BU5c1kmv9Uxv3GbTxVtBIWSpgFNDF36NbcrMqA7hYFSblWAIxts6yDMaTJKDNKKglqcuqbvSVOzZIqsrSbCBxcKjORYhbMQbOBiHsRPGI9puRZiet/QH0+OJDo1ag1dTlspkqjNldbSUUbyE93I6XIjB4DEK4AHUm3nizNH2YvxgCmahjff7r/Jaxw80a1TARIzd3d3uAQOeSOb9bAG3TDepVQpCsWY8EbSpIPTzN7+Qw+mi75XqGQrT0qeKSTgKhPhA9yTwByTyehtbHZTpZe23s4zPs5yijH9tMgqajUWmWp4g1RX0oQfpHLlA5kkCKtVCvUmKdV5cAuL7tKpA0urRUE8klLUh14kjbn2PocW/2S6tGUVVJmE8l4adkpqq56RniJ/hYbD7ouKxzzT9dkdWaOuhAJUSxOhuk0ZvtkRvNTY+4IIIBBGB05nT5ZVBHVXjcGMq5sro32kY+QPBB8mAPriG3Wft4iBmprJL2UmK7a1VV0sOsuy3trWNfq0U69nmq3Bvuo543jop3Po1NJJGT60YHpiAdujVEPZtmmTVrbqrJ9WTU0h9GfLpY5B83pgfwxr+yzUC6tpJ+xvMq7fl2uqJsly6rnO36vmKsJMvLn7rpVJHGf7szWNmwH0hs8lm02+cVdOYptVSZbn8sZFtlQ2XyQVCkeRFStQCPW+MJjXExPdmDd+PfgtIkDtGjIivyoNorM4H1bpeop1KHLNPwuSR1lijqJAf/GU/DFw9pU1bp3sBynRuSd9VZt2j5tT5NS0yLdnoaIxM6rbn9bVvSp5j9UwxSmRwwZdrPLKKWZIAdK0zSSyfZTfTiQsf+U46dqO5p+0qj1XXzD9F9iek8uyyjdgDHLqfMleqZhxYmAzTSkeRo0v1GJCQLQJP5WtvetOZUVSYrgzJotdrvMqPsr0VlHZRpelgra+ip0ymmnTdvqqtpG+sSjmxElVJKVPUIsXlbHMHannVFmeo6fJ8qmNRl+S08dBTyIL98kIYNL795O88vwdcWLrLXBqEzPXtXLtmlhky/I0vzGtu6mqgP+7Q9zGfOWQkX7o4oWEtLI9XMsZLkEJtLbV6Cw9hixomAveZn56/H8etVHb5Axojbl7flOLubRiK5+0iyKX3exXofh/LBou/ihM8LSBk4VyAVU+l7n8Lc2wn3c0bqsfha29dhsbeoHUfvtg8caQx3uoYqQy3tcX4BHl7X+XpjoFlJGaMGK4dbKPs2J+J9Bc416FmoZowfFBKsq+wPB/PbjaVCqbxyP3gjHha1mPqAf2fiPhjWRXFU8RtadClh09vzAwx4wrsT2Zrs/6JGepnHZvnmSz1BjOns9p81jBUMppsypzDKpFidoqKaA+HncRi+siz8ZfqLLNSO0c8eU5jTZoUMoKlUmO4Hm12g3xk9GFtxuuON/oc59LT65m04kQm/tLkeZZTHA0zQiSshUVtGO8XlG72ABWHQnzuRjq6ggjzwiSWmZ4qlGZSyDvInJF1YKBe4YWkUA3NmUgg4880xAbNbDI3Xjw/Iqur0c8S2e4eqrkr6RHZ+vZ3rLM9Ko7mPTecVuTRsbMe4imYQMfI3galb/nGKh3uyAxs0kackm9kv6g8L+4461+mRkclbU5bqdIGLahyOhqJyUKn69To+X1IPqxkoKRz/wClHrjkZn7wKyrbi9hbnzub/wAeMdvo+USxkjLPjj7rnLWwseAc8uGCLy7CNLvxYhibMPMW9B72NvwwmysWB6g8X5sR6C3FsKMWI2Bm2Ktyd/BBPW/x6X64CTakjB4ivgvdh4Sx8yB0B/f7YvKqkdyIDcrcngcfMbf4YHcoKxsUYWG2wFhfm3AFjf42+GDSPLwJEbn/AIlmuLcdfL+hhAjxiTwqP2QbX9hf+hgQlyAKcfe3G9txHzsOPL1wkxCgMDZiLk2UWb0HoP44xXZo+6JsbbTxYn25/djDcMoZXuwHkOT8/X0OFQiEgpynn9j0P7xjMGRbsFK+Zuo4APw8sZgqhJm7vsDAi3N2sAOpv6D+hgAfCLAn4A8extg7geF9tyRx6jk9PMdMF5kFzbnzBNzb8v6OBCKCpHiYIQdp8Nx8z5DAckHdYADr6+/v/LBiNu0Ou0W6gdB/HAsSz7y5Y3+0p+17j/PAhATbyQkHcL9b+o9cFI5uRcg3+PuR54ONo3AHdfgEixv/AF+eCkg33G9ib8j+hgQknIT7vHpf8hjaadtT0+Z5mbDuacRL8Xbn/wBijfjjUuwUcAAjiyi2JfoPJxmtfp7J5Eumc5wvfD1gjIDX9tveYrWp4ZHj1TE8gpoGlz8OtSuvL6WKgoqXLZQitlOX09GxfwhHSINIL+pmkkHqTYeeC1VVU5JofXWciTY82WQacpiQAyz5nU/rh72paGYf858jh9JNNXQPm7iCZ655a07T4ZO9YuUJHA+0y36ghDYFcaHtjzCHKey7TWWwVEsj5rm2aZxIZowjtT0UUWW0hIHWzrV8+ZufPHMWaskoGv3z54rZno2NU5POU0vmNcLhs2zSOnUD/hQIXa3zdPwxH1Y3DAgHaeo4N/IenrjaZ03c5dkWWBv9lRmrk2nkPM5b8dgTGuXaQQJFJuL+XB8/bnrjpbI26wu2k8sByAWNO6rgNg/J5lE37VZDssdvUXNx6H+r4MWgAJeOQ3IK7X5PPS5Hpfn1wBXe3gTqBYAnm3X5nrbDWqe4Fhb4YtqFWZ2qdmEOmFoM90zes09mmWQZlQ1Yv/rFKx2GYi52ukoaGZAbRzIfuumKwnjtyAb+eOr/AKMGcaP1xL/8T/2n109FleppDXaOzmBEebJNQSRBTEFchZKesUCKSFiFdxGfC3iFddtf0fdSdnOawsuXxvQ5j3hoaik3tR15jJWQUzOA6yIwKyUktqiFgVIYAHEEUzXAAnFSyRFpKrPTWeQxKcuzAAxvYIzNYXHQE/dIudreVyD4SbdE9iWsoK+lpMtzKXdU6Dqpq1N9w82na0iLNqcgcnuGkjrFXyBqCOMcr1ERifbYjqLHy9sSXRetsz0vm1LmlFV9xXUaulNUlA4CujRvFKp4eJ0d0YH7rHy4xWtFlDrxGThQhTxzkgNOYyKmfaFT5rDqztcyvMlP1uHM3qKni9yleQWv6HvQb++HGS6iGWdjq5e7oHzKunpRGSVvTgQvUuzeSt3VLFxzYOPPEk1HpmrzfUOY6uiR80yrXWhazNKKoI2tHU0dOn1mBrf72CalYEealG6ODiD9rbTZW2Tdny+FNN5bDSSIDx9YcCoqmPu00u0/+iA8sU2Mc9scORNCfJoB50Ut9oc+TV8mvooRmudzZlPIiSMYWYNI1tvekcLx91R0VfIe+NhpLSmd61zyj07p+kM9bWSCONeABfqxJ4AAuSTwACTwCRp46JmYLAN28hQL28Xp7Y6A7O8hTRumDRpSmXOs/p0krCq3kjo5LdzSoOoeo4d/Puu6T/eti5a5hYYaRDvavnrwUNnjNqlrJkOqdeKuTsN0RkNHVUGhdA19PUVGZZpRZTUZ60d1zKvZjMe4VhxQUcEM9WS3NRKlOzAIY1xUP04I5B9I/tAqWZiG1dmyKWJLbUWlVQb/AN0DF9diMUuT9o+ojSxmopuzqkGlkliYWn1Bmci/pOdW6WipoKiEDySKP1OKW+noixfSQ7RIrcLrHMiBf/iQUj+WM6wOLLSYyanAnbUg/A9BgFZtPfiv6sQPIhc3sxJ2Nfx8Gxvz5Ee4/nhZcurP9rPDHAh86hhED8L8/gMFjeeJj3M3d+XgAU/C/UfjbCRUMxY8k/aLDn8eoON43jkssXRmnC01JExY5i8pPhKUyEL06b3tf/wnGLUxQkCGihW335B3zD5Gy/lhuFYnwgsVuOnW/sPXA3U+AhV54Zbmx+B6j4YLlczVLe2BKT19bULaepeRB9zdZf8AwiwH4YRS5S4Tj9kC3+WBIk2CTZ4B4hcXBwKGMME+qwSFRuu5YXHX9oDDgA0YBITXNAqmwBS+31/n1GHUNLQRCKfMp22yeIQpwdn7Tt90GxtYEn2wNDFV1SGaHLIWjB5luYkT2Llgv44fLHku+Ns3rpaxI42jWnomNkBv/vmHQE3sqtfp54je/UOSexusrUr3JAkFOACb2LsQPbyxtNF5UZNYZfVVIlpqOhf9IVUoja0UEXiLewJAUH9pgPPBMwzCXJXMOWUdFTIeEqYX7+SRf2lkfoD/AHVW3Tgi2JFmWY1GS6YynJEidqXNqSPNKqdakqK6VnPLOt9wi292qtwjrIxG43xFI4ubdApewTmgNNa5LWx5hHLBGiuEZBuPd2BVutg1xY3vY8WNjcXsSZ/EMiqPqMVYlQjwxvvp1buzvjDlf1ihiwDbSLcG9iR4iWKnAkFTDBUiEr4jOn+zJ6EMtlZbHgkKORcdDh/JV070zQ5nGGgq6GmtKSSYJkQbJb8m1gVYDkofUDAcCKIGOBWlGYz5cI4kqmJZLqI5eBGy/ZDg3ANyCOoG4YY1BhrO7p4svjM7ARxfVkKsxvZRtF9xPT1OCZo0C5vWGCogqYRKe7lgR0jdRwCqv4gp/vc4srsLyiKkzCv7U80pBUU2lO7GWQvbZVZ1NcUcZJ42x7ZKh/7sFj1w20SizxGUjH1OocUsTDLIIwppT6QraGs0/wBj+UZlSUwyh5Zs8rJQr00eYPGHzCplvw0NLBGIr/8Acy25cX0Wu9W5dqTORU5JA1JkuWiLLsnpZBZ6fL4+ELjzlkMzTyHzlkk9Bbe5pM+ndE1OWfXJJs41LCk1ZLLGqyxZeT3sMLEfeqJAtRJzzGsCniRhiG1lbFoSVc9Yj9JOsdVlkcihvqw2gJWSKepBG6FD9pgHI2qN2DFUCjsXHmc/eu7Utd2GWDRyTfU1LVUVZlWgYpZI88zOaGjzKO1vqQcoiU//AKXZZpf2brH1DYkMctAZq6egjRYnkmmhUnagiXwR7jbhdirc+R5xXvZzmE1dr9NQ1JeQZbT1uZbpGLtujgkdWZjyzGQqSTySb4srsx0hTarzrK8nzrvlyuZzJmdSs4X6vltLC09c2y1zeCORQSbbnXg4kniuPEbjkAT4kmvCg8tqijfebeG3Dyp8lSjNcwh0xpLINFxQNHV5qn9ps2QiziaWErltM396Oj3S/wDpaxvTFO67nrMzy2aeWFIzCB3sinwsDYIF9m4IHop6YsDU+eV2qq2t1vUQpSVmYVdRVSQr9mnkilEsKL6KsDxIB6JbEc15Q00NG+n6dbQ0N6ub17+VN0URPpDCw/5pDh0T2tlB6w6AQ9pMZCtHsjgqM0192eUckcYp8w/s/IjXBciMKWIF72HdEXPHPrzjc6Rr4KyhbUNTIkqwCGIzlOFqavfdifVEkmc+YFj5DGm7Cs4pv0l2Z5r3G1qagq/rE56f+T4a6TafWy2Pzw705TZmcu03ozN8yn/1+hh+uyTPdoY54Y5XLORwaemWAqx+yUYdLg5kjDdIOFCRwJHsrjHUIO4c6FaTXWdRZZkultAVIVH1LmMs89NJYdzl0TvHTRk9Qj1ffyhel41YcEYc6Hq2yXs/0tmyOKBIslQS1TqHSQNUVcZjClTte4IDL4weV2/aFc69zCPWmY5trelX6vDOyJlaOp309HThY6RE/ZbYsdz6ufjizuyypjfsu0nUTTkN3GZQTGdmNMqrXSNykfj3kzC7A2CXBFjieeMRQA7xUb6OJ+PJRRPL5T4YeGAHyo9pbLqbT/adpatmSPL4NXZ7m+RTH7qQkUiQXA8lnkVvj5nrjoLTkwyzI8izuvpmlj0xk2ZZ6GRQWikoK/MioYE+E94IQrDkElTw3FB9ruQZhDonT5y+manrstrNQVsckR/2TpmcIUqR0AEfHsPbFyZ/qOkr+xGTVqSoq6whfLoYVO1FaqzZ66rX1sqwTKfMbh64htje0EdNeB8qkHgeSfASy9Xx44eo5qt9DxGTs1o/0qKSepmNXQVcFUAwGYU8hVnUXBZ2iniJF7ER3I8IONn9F6V5dQL2W5vUSL9ZSbOsmJ5cFpFhr4R6kvTRTWH/AAmPXGs7O8xFZmOttOyD9XNPS6sSBUQukc8fcVewSHa2x5ISQxAKIeRh5R6e1DkuaZL2g6VkjhzvQ2o5Mxog1wlVDVQrVx07c+FZJYK2Ln9u3Q4c6khkjeaXhUHZhUetPCqTFoY9orTDnT8qXaJ1Jl+rKfX2XVUMUeYZLqfMo62Ludxqsuq6ycpLu4sFZqiEn0lj9gZxTZfPmJyPPVpKbM60N3dfTt/q9TLLTGOCq+rv9ifvYPqtR9XmXce9co4JIxz/AKk1fP2H9qOmO0GmyyWOlzClqaXP8qDXM9BJPukjLHjeBKCj+TxRuMdE6VospqanM9EVEkuc6ezGiTNssrxVqJKmGJAoqoZDwkppJoZtzcl6OQNbyrSNIJlH2vxHiMx7+e5TxuBAjObcPI5dblydQJnXYrrqu1RkkEGbZfk1XU5NnmVyXjTMKGOcxPGR1s0YiZXHKOFYcqcdUay1ZkVX9HfJtd0U0+osh0dmOU18Tvb6zW5YrvREsR9ipjgnigkH3ZqMHpIhME7Qsvp6HXueVOZSPmVDnMVDmtXSmIJUKZ4Pq0lTECSFf6xTOkiElGZ0BtuDCu9KNqzs5fX/AGdUY/TGhtU6fqcxHcybhSmIxSR5jAjcnuXWLv47hu5V9wvGpxfqLUe/mKHxGseIBNFUo6D7f0O3kujtLxab1PRx0lBPHVyCmfNNN5kgKw1ELN9ZKlf2DLEZwh5RlzCPg8Y5r1r2MHKdf5zlmmZpstqMtqYM809X0e7vIqCqAqKZwycnuZC0RYcqV4NxzKdK5tUdlebx5JnNfV6f09PmH1/Kq6VDL/Y/PA6sTIo/21A8gUSqLjY6yKDuIe5e1/LMnmOnO0OWLLtLTU9Yun8xFTIzwZFWT7p6YySA2komleZdy+B6WtjkH2DaCIugxjJocvLbvzBHrUKVwEmDx0eqquKbS2os373NtQTHMcyzKqds0gqYE71cymjMsqFUASRZ0V6iCRABOoqI2AmiIMayfV+p/o15vU1uS0U+cdnctRTZhmGSyEBqDvn2fWaN2uF8S904PgcFA4syMtv5nCwr1op8sqqeohibL6rK5Km1QCsgeSj7+/21kCTUlRcCOaONjcSuWb1WT0OdZTBmObZOueUtZDJT1rQTNCma0dRdWaKAm1PJOsbHuz/sqymZFsYwC1srSO0Iq04OHuPY5g61I5jh3B9wyPt1gQrD0F2g6Ylnm1jkGf0lfo7PqqWrWqSnBOXvLtMs3dcvGyuF+uQEgraOqjvtZ3ivarT1HZ3rWmnkimpcu1HNPmlFNl7CObLs4gbdViBh4Y2PeLVIfsETTcNExXHKeSVOvvoxdomZ1OmHTO8loqtY8woqjctNmVIyiSnqSBzGXhdWDjlGLqbruU9a5VnGV9vHY/PkvZ/m7S1av3ulpJtq12V5rTxPLT0VQLFRNEBNAD9iopJRYHuWVZZoBGah1W7dY2V8NusKGOa8O8KHZ8ePqoprfQOTfSC0HXZCRQ0Of0e2pidl7iCKrP6uCpRee6pqg7aeUXtDIYw11SJjwpUZbU5DVz0NZSyUlXTyvBPDKCksMiMVdGHkysCCPUHHW/Y12u0erqymoUyqkyXV9KJYZMmqEZqPMVdCk8UCX3FZE3JLSXL2IaEsY0RIn9LfQdJVxUHbZpO9Rl2dSCizclg8iVieBZJiB/tzsMUzADdLEkv/AL8AY1tHzGJ5s8nXhuPqqVsiDwJWdeO8Ln3LM11E9QlLSZnLMWNlSbbKlvU7wwAHUnyGArNQyU1fS1tXleWzCkqY5XangEBkCOCRdLDm3W3vhjamYh4swamkHI76M2B9njvb47RhtWUFetO070/eQEEfWIiJY/8AxLcD52xqlkd7HDkqN593A+6sLPcyzXPaNqwZuuZ5bVToaWubK6TvAVcN3crrGJIpBu8am4Yc3Km4bVuoGzKulzeryjT0UzqY4o48nip6ckNbckcZWPdwb3uQdvhsL4h+lM5oKJKjKMzKx09dJC5nZTJGhj3WEiDqh38keJbAgHobIo4c4ozUUlPUx5Ws+xHqJp1WnqJCo2K3fl4qreLFWHIB6gcYpPb2Pdp4asMFMCZO8m0U2lYV77NdMvVzBt0iwVEtGZIHjFneUb1Uo/hCqhEgctcAY0ldPRmoWDOstzGtm3KskP1wwvKi9AfC9xtsotyoF1HIGJCtDLkS5nXZ/S02VxTy2jzOpjNFPDPYE/V4lDCZLfag7sx2+yyE3xrquDKM2ysZ3pmnbMo5JWgq/rEJNRTzAbhJ9ViNljYXKvd9pUg2IF2NIzzG2p668E41SeZ12kdQZwtZRZRnFDFNsZ6SprhWmnKrtUwTy7TMm0cRygMLeF8HmySppl+u6Yy2pr6HY7mso6KtUOv2GXfcSBR0IF1B456ltJkNbk9WseossnhmrIPrNLFLIpZt1rSSx33gEcqCEL8X4GEI8wWmdYprFlDKskly6GxsUdSrghueLqbX5th12oFw4cUl6n3BMo6rLZKpoVq5qZ9qRKqTtIrFBZEHeWcW/ZJtxYWIAO3qY22wWzejeqihQMzqO6ENrxlJPsc3JO9VYMG3eLjDpDVzUVTUzmfMauNzJORQJNIq28LzVDKzgG9rLuYeZFhjVwZi9ZG9VPKKqocsjTK/jZnF97kcFQ5U8n7VhyCcO+7HZ1sSZYLaZHV11EJc1nyfK6imoU2ua3mncMPAqsrBjJxdDGd4PIt1xuaTOpNRtUZhHoys1JWZg0MdTRUMjCqnlRmemqmeOMtJb9bDM+0M3hJIY7jCTV0IdO/iWMxqRtMhQ88F0JP2Tbnm4N1uQBZ7HpzOaBJM5zDNpdOZXVJ13vHPWISP1cFOGEkwJAsWtH5s2GvjYTV2B8/b2Tg51KBbXVmi+1efKK6TOsqy7TeW0VOKp8mSqgpZXjMiKP8AVg7TStudOZLm/nfEEzChOWwUgeel7ypjdjEk15qdkcoySpw0bgi9j1BBBxta6qpcs0XUw0Qkgl1DURIFkZd7UsDFmdioAKtMVA/9Eetr402U5xmWRd+kTU00dZt7+Kqpo6lJLcg2cGx5PKkH3xagvhpoBn4fOuqhku3gCff4TfbIWJlYuzC5Zj4vx/ng0eYz5bUQ1tK+2ankSaM+YZWBH5jGzkzrLZ1Zp9MZWrlSC9NJNAVB6sE3slx8LDrbGsrlyyW/dSVkO7kLKiSD/wASkH8sT3iRRzVFdp9pWx1fBHHHPJRsopa0wV1OFa47qQOyqfdSXS3kVOJb2CZjmWQZjLqXK82kyytyhqetoq2IkPSzrXUyJKvupe9vMXB4JxAaSr+vw02U1caGKiV+6MYKvPdt21mJ6LdyOL8keeLM7EqKkzLJ9Z5Y1pFn03WbT5ptrqM7x53AucZtteYoKHURwqrtlF6So2FdQfSS7Gsj7edLU3bx2a5J+jczzWjqa/UeQwJvSDNKZzHmk1Gq8kxyIXqKa25o3Spj3frBjgPOcunyuselqY1SVLXCkMpBFwysOGVgQQRwQb49JPoua9Si17V9lufyzw03aVlcOtNPTwnZJRanoFeCuWBvuySfV5mHqVjBBDkGi/pl9g9PlWaSdoOksvhjyusn7qrho4ysFJWyBpR3Kn7NNUqHmhXpHItRB91LrDbgLR2EmsVG8fhJLZ6xdq3Vget65k03qmoy9li+sNDIJI5I5gbbZEIMcl/uujAEN6DaeOlqdvHaDT9pmlMs1LHSx0dS+YVsdZSp0iqJGNTKFHkhllmdfQSW8sUZUUskTc259PTFs9mlJDrbso1poqeJDmGV1OX6ky6YgbgiyLR1cd+pBjqYnt/3V8La7O1rmTjCjgT54e/JNincWlhGo/K3+Taeoq/6Q+W5DnrFcvosuof0lx0pKbLYpagf/CopB88THti1nOlXTaSza9OIamo1DqdIH2tVZ7X7Zp6dT0C08H1ekDHhBFMerAFbUtNH2f8AaX2o9smaGnKUmd5hpbTVLMu9K2qQJHUTOl/FT00IG4dHkmhj6Frc4aq1Xmur8zqq+rqZpnq5nmnlksJJ3dy7M1uBdiW2ji5xRs8D7QGtGV1oJ3gY05eY3Kw6ZsffO0080bVGo6jU2Zsyd2lMhASOJdkSqo2oiL92NRwo+LHljhBIGljSNE3EXuFI/oYHIsoq6+thy/L6SWqqp22RQRKWdj7Afn5W62GOsvo//RMq86y6XtH1UcrTTuWRSVeY59nD3yDK4IuZZPCQ2azKOBDCRT7iqvK99h1XyxWCMMaMdQVMNfanl7j5qh6PswfLuybOe1XUdKoy2SogybIzI0kbVmYSnvGliA+3FDDHIWY+FmdQL4gqsIhujUN0cKTvt7EEC9+n9Xx0n9MjXlLn2k+z3SeR0M9FlcMFRqJY6lgayda5itLLV2ATvzSU0TlIwscYlWNFCoL80qGZ+7szEkAbeSLgW4vz6j0w6yOke0ulIqTqyG7f4+mSbO1jXAMyoseQbhtjtc3Tb4QL9OvX0v5WwwzBTFKkyuHMTAFgetvP4Y2CyEOxdQEszdNw49b3B5t8+MM62NpIWYkDYvi4AF/K1vP2xbUKlXZTrGTQWuMs1bDdhp/N6HOlQdXWKZe8X/mjZgfbHorXR0mSavzbI6GoheLL6uenhLt4u5RrwyW817tqdrj9u/kbeYeRbZczp6dzZayJ6V/iylR+e3Holo7Ov7R6R0jreWcmbOdKZbNUODZTU08b0MxPNw26hjb5k9Rzx37SxBt1/WP6Fb+h5CXFvXWK13b1Q1OddmFRKfHUafzeQKXmErdzXUi1kYvdrlp8pl6m95beeOEswhamqHgaUHuZHjC92OEUkA7vf8fPHorX5fT5ppnVlNKIC82UR14LRhXZ8tq4akmw6t9X+uLc82LDkWx5+60ozk+oqqgcENC2w7rEblvG179eUbFr9np70dzd6fiii0tFR17f6/mq0QcA2B8W4EgEEkeVwb3+Ywd7qm2+0BtxjAsAbfat7dPUYLuYkEsxt9l73IHsevyxhQh+pNhuAXkH5+/rjplipORLgr3YFzuPguD6H/MeWAWyuEdVRWUWB8IPtzfg+/HvjAqgLYWJG2yixPPxsfywBAUHYVI3bbgW3f15XwqEq0XdxgGMsSxU7lN7jrx8wL4RIVV2keXkSQePfn5YMFt42TwqbFibW9P+gwCCQ8oSFIPjtYH2v1v+eBCJwUCgBl4sR0B9v64xmDFVA7xmLWFrnkkX8/8APnGYRCRKbiCUY3vxbqB1xjMSAWPiPqevx/rnCSkFxcDcTwT6en9eWFwpdQtha9724A+Pp/XnhUJMSDcpXw8AjabH3vgFc3azC5uLLyD7XwcrFu7xrqOLknj48DjBCCQe8AJB5PmfnhEI1zYhW4IB+PxB/hxgHdjwzdByOTb0HtglyeD4jfrY8/PAbiGuWJ44seg9MCElPa24k36H5Ytvstyh5M/hmAsMi07LUAWvaeotEl/+arU/8uKmWF6ueKmjA3TOqA+pY2xePZ7NPQ0+pc3prBanMqLLo227vDCXn4HmLxwXHnwOL4y9KPIju9Y0HoSr1hbV9etvsFLanM4oZZ4XmjVYIvrNjKlxEl7k82NthB9POwIOID9IH6xT53QaWe4k09p3JcmYE9KiWH65UD/4dUvf3GJlT5bHqnUOWabCRRy5tVU9DCr2uy1EqQOFPRtquRewYqELAlbmBdrGcQ6r7Vc61BHMppc21ZmVbAxPH1WGTbH8tiWGMmwtDZK9bfQFXbWbwp1sVfaknSXUFaIbd3CVpY7eSxqEFv8Aw4YxOVdSrLu568/I39cAsrTu8zkbp3LncOLkk3v5H8sYGJA5Fxe1/T0x08TOzYGbAsV7rziUeQPtLWJJsGJvyfJr+RwyqTc8tc35Nrc4dsSrWFytgbX59jxxhtUKStyOnnh6RS/T/cV+T06VJcd24j3IbMoubMp9VO0j3GO4OzLtXpu2zQWaUuovqVdqzKYYotf5JmEXfUWpaOMKlPnhiuD3yqY0nmiKTIe7mDWLAcHaNqWKVFGv2yN8Y/vdR+agf82Lg0fmGoMopqHtU7PKiOl1boPfVhGj3pmOU2PfxSx9JFjR2Doft07yg/7LnEnZV7oCaVNQdh1LVjeWtbKBXDEbQp/25/RYy6TIa/XvZ09TU0NBEajM6OZvrFflMYNjLKygfXaIEgfWlUTQ8CdCPHjkLMsuq8rq5KOshMciEXF7ggi4II4IIsQRwQeMeoXZ12pZDm1PlHaXogzUuTZm36qiEqvLk9dGoFRRM7cMYw4KFuJYJY9279YBSf0vPoyZVS5AvbJ2Z0MEemKyYR1tFTA7MjrpSSsag8rRVDbu6v8A7CbdCeCuE0dpB7pDZbTg8c0lssbQwTwfaeSqL6Luvp31bpbs4zhY6jKcy1LTwv3ty0C1kTUU5j9O8jkRHB4PdxHquKt1zWVuZaxzbM8xN56qsqJZBe2099ICB8LY12i89qdJavyvPoARNlddT1qhuLNBMkov6cpibdvuRf2d7XtYZei7IFzutkprdO4kmaaO3xSZT88XmRBlrLtow8jj6hU3SXobo81HtA0FFm2qaZMzXdQUxNRVKR9qFAXdberAbB/jx0ZpLOKrT7Z12r5rFFKdK0UuoS+5SkmayOsNDEV6gLUSRttP3aaw4viguzemZIKyqUDdO8dOt/NbmRh89iD54ubW1BmOYdlmiOznJkH6V7TtUNKVFwWpaVhRUoPqpqKiqf8A5AfLGfayJrUGnIegxPW5Woj2VnNMz+iuv6ONENKdi2QpVVG3M85irtV10067+9kqpO5pjISb3MFK7hv+/wD72KM+nZUfWfpI9orA3P8AbCrXjn7NHTKfzxfFBnNBJqqXJchkDUD5jT5BlccikI9FG8VFB3br0vHGH2OLHdcHnHJ30jtVRa07WdW6jilVkzPUecV6m/DRtVNHGf8AwQripolzpbZJKR9xHocFYtwEdmjYNQPMhVgxH2GYC1yfO34YIzEsbm/Nzza+CNIFUAG1jcAixX15/hh3HSSFRNMUgTyeY7QfgPtN8AMdQSG5rFAJySI2gcqCwFjcH5c9QfLB41lqJDHBAWZhZtoLkD+GFUkooVHdwvUspsWkGxB8EHiIHuflgKuqnnURSTqIR9lY1Cxfgv8A7kL4becchx+P0TqAZlEWCnVmE86K1r7Yh3jg/EcD8cYKiGEXp6ONnPR6i0pPwX7I/A4SCi1iI2B4AvyD8fLBjGbKFALnwkbrc+R+B/eDhaVzKStMk5nSsrQJWnNWLDbeRdyn9nYT4T8Bb0wNNDSr/wCdZgkYPLJHG0sg9iOFB+LYSjiobf6zPLMQOUhjAUf8zeXwXDtJhUSR0dFkYqZSLIsskkzAeu0FV+ZFsMJLRhlw9U8AHPNa0/UkronnSV6Pv1MrMArNHuG7hSbG1/M4nVdLXZrn2b0lBRwPJlkLSJNFRExrRxsNn1eCNSsaBHDl23GxZi1zfEVzGnoEjjNcYu/jTa0FE4ILXPLMPAnFhZb9PLG8yipnzjTlPX09ROK/IZBQsFlKkUkm4xEEWNlbvIzc2syA8cYiebxDvLd18p7RSrUweKanJzGnkk7wEyvUKXvb1Zz1v6ng3tbDjNloostpY4M3ir5a6BJGEMDxrEguFBLgeL7Q2qAvANzwMGztKyOGGCrEm+MCOCPc14o2Bbu2HkrXJRTyDc2AONJvpaKSOomkNRC8au4T9WQ5Xlef2T4T5EDjDgL1D0U3KoSAoauvqqbK6Cmeoq6uVYIoolu00jEKiKB1YsQOOpOL/wAwo8p0HHRaQmjir8o0SZYq5Ue0Wb6hmUfWUDD7UaBBAWH2YYJTwZVvpexvJ5NOZJN235kI4a96p8j0VTmMt3mZsn66tCDlkpIn3CwO6d4lFyDjS6kz/LdLqsVWkdTX0CtBQ0MxE0WWgm7SVFvDPVuwDulyqmwcnaEGRa7R9RN2TBW7zOvyAwJ2kjMAG/ZohGztHHP0/Pxqqtlnuo0pFm1Vq6eOszOvY1cNLPHxMXN/rE6D7MN7bIeO82qOIlJNP55n1fqjNJKmsnlkaolMrtI255HPV3Pm3w4AsAAABhPMM2rs9q5ayuqJZS8hkdpHLvJIerM33m/yAsMZQUxvUVPhAgivc9AzHaPyucXLPZRZgZH4u9NwUMtoMxuM+31Up7Oo9g1RVqosmUilT0DTVMMf/td+Ll0RWjJ+zrUGbvAqz5tUppum2PuD0aulXmcynzBRaKI2+68g8yMVHoJ4qTR+os0njdkNfl0NlHLALUSkD/wL+WLhz/KavLs4ybQ+XUNRXy6fpqXKJqSlH62pzOrIkrlUH77TVIiW/lTqPLFCdze1kvbR/pb7nmrEdRGymz3KiMmayZZl2b1MlP3q5XVU7Rh1DJUVbmSNILeYYBXI/ZjPqMabMvrUtBUvmNQ9TUl556qci3ezklpH9PtHgHpxbDrXmYUUmvsk0FklfBWZdp7NIoamqhP6vMMy7xBU1I9UBQRRn/hxAj7RxsdP5X+ntUZDkE20/pPM4Y5luf8AZGYPJfj/AIaOfPjEQo1oeRnjvpq9CfOmpSE3iQPBS76O+WM00mn66Rg+n6rNzOpF+7grMjqRIPk8Vvi2Nt2iZzTQ6IjrER4s21lNUZbTxvw0NGpDV0y2+7sWKmU+jSi/hxFPo+6hq8x7Z80o6aVUfW2W5skBksEWd45ZIS1+AqqWv7Xwz1Xmr552hx5jSs0uX0GRvHk8Uq8Ll4vDA7DyM7PLUHz/AFqn0xWId9UQ/UL3ma+rgeCkqOxFNeHD8UWhqN1JpqpjMbPFL3ckSqbKssSttBtwbxyFhfj7JHRgLG7BzVVnZiMvhrZKaoynPKyAvEgMiQ1NNFMCjH/Z/wCxe7dQN1uTiHahp4/0HWLBHULaMAhArBYyQid4SR4bkEBQTc7vCCQdz2ZVlPSai7TNDzu5WkqaaviEZAv9SqPq8g5tcbZbkcA7eemJLVWSzuLdRB5gehSw0bKAdYPoT6hWRVUtG+hMlpzG8kcqZ1EgRiRIHzqRChe9wSoVlJPO1x1IxXNDmuoYOz1uzWWhMuXZVqJa7KZpJArPJX91CIQLWBC01Q179XYW64sRqf8ASGkdNVEbyM06ajWSnRgDMn6YuQpPHeAhXQngsAp8Lk4rTtqzybJNLaTqMt7qpr5s2n1FLMikJLDSFIKaXaeQHkNQSDzcEHxAkrE2/LcAxPqB8VHmkkIbHe2fK2SZrSaR1xpXXE0xhyrMRNldcjxlXOW1ituZhyjLGZGsQSLxC9umLqynJKnTeb6p03PleXU89ZQU0z0dKVYCpoZ+9dhybmWlqqk7mtu2FrBQDipdS6ay/OchzfKcsqWnhlgp82pYnUKG+tL3yLHybnY4Vum4q5H2CcSXQvaPT6myLRuvs2cv9SrqXINRoCQ0rxRGkeViOhly+dXv5tTS4qAula17P/U+BxHqVOAGktd4jxGa0XaLktJmFPRRzRxPSST5xlQEDF04pI5LoW8alXQho5LsjLwzIVONp9H3WWZ6Vq4Oy/NQuY5zpSq/T2lTC4YZvQoWauyuJgfE8kDzvCAeWMkfV1xvO0jJHocgpa+SoZ6/LtWJST1tOhWpWeOkqoXeWLkd6n1dW3J4ZU5te+Kq7RIf05Wad1Vl+Yx0FXlVVFT1VXl8hjTL3Mg7qriYAHulkKMG+6pC8Wth8JEwEL/t9HAmh88jtqmStLCZG5+ow/VWv9I+vrsm1VojWOljHnlFllBmlCVkZic+yd6iOpRGbzZqasVweocEjxKML5ZTUCR0Go9P1EFXlSwHOcsr5GSNXpmDRs0u4gBtrtTzp5lWUg3jw1rtTrr3QWU/X1/RmrNKaolyfUeURoN8U1RTTJ30CHgwyTQ7goHgclBYd2cQ3JdS0vZpql9F6mqIqDTOq5DmOXVkqnuslzeOUKznrtpptsYlAvtWSKUXMYvG9rnNEYHeaD5ivPDEUzFdyexwFXuOBpx/XA/qrByHLKHVejBp3M6mF4qGkSnp6iofcGy0MIqZ5Wbm9K8kdJMx6009JIf9gxEeyjPK2l09nn0bO0qoePL86yefKdOZtVC7Us0chloaKpJ47tKpdkUoJ7nv5Yz+rYbJdnrS5Pm1VmGV0UH1zJYC75bKjMJpWkEEtLUDnfHURVLwMVO3u2BW9lOHec6aoM+oaSsp5pK6KWjD0VbVQIKp6N90StPyA88TJJTzEi4mgfcQsiOCO0Vb2zdefjqckdFU9kfL3CqbsW7VK3M4Mu7M9fyCjz6KAUun6+tbatYiXjGXzsxG2aPxxwuxt9mJ+BGy3vo6eaorIFZnEObVDUUYjqO5eLNJSFkgfcLKtU0SqySC0dZHExH62S/NvatoSLPNQVC1AKR53SRZ9Adjd9FLIWhq9qWuxjq4Zg0dt1nBxLuyvtdXUtJV9jvaY4/tLIi09JmEUioc9G1TCGm4tWqERoJibSlUDkSKrNPPGHVnjH/sBsOse+w45KONxbSJ3kTt2dalaPbVo1lyui1XA/fT5ZNHk+YrXU31Vp8rqmJo5WJLK/c1LPCW3EhKlEcDaDim8lXVvZPmJ7R+y+eOmq8tMcuZ5Y0haGuoEnTY8gBvvpqju0aQWeMSQTA8OT0lQ5kvatpStybV9UtScwebI8+R4+7ENTJHY1aLYGMSeCrVTcpKlTGANmKB05Ln9PC9PmLZauosiqZaWrhkEndzVUfeU0sMykENFOgkikIIXZIDwUAwyzyljiHmuo7wcj1u1p80YcO6M8fAjMKufpFVuns91FmXat2f0VXR0dXmoXNsqrYe5noKmoT6wisFtYrIKmMSra5hVhbdbEy7H+2PIu1LK8z7Ou0zNRBNqmMZbW5jUyAR15cKkNRPeyx5hA6xOlRwtSkZimtJ3b4nNdpTKte6RmyyOtZKLP6SOgilr/8AbRxu5WkaWSx8dNWRrDIfUOekq44aqIcwyOvnp54Xp6mB3gnhdfssCVeNh58ggj2xr2WFtojLcnNP568tizp3mNwObSPx15qSao01nGjNS5rpLUdKabNcmrZaCthIK7Zo3Kta/kSLj1BGNbFJU0k/1mhqZaaYH7cTFDb3IPPwONrnusM611UUedZ+TLXU9FDl8tYx3SVCQrshaUnlmWMLHuNyyot7kElvTjLJAUqnNG56TRoZIj/jQeJfit/8ONgONzvjx1qhTvd0rJ6yCtpf/KuX0tRK97TRp3E49yyDa/zW/vhvBQpXwQwTZ8U+r7lhhrXZYkBNyEflVv5ggYXzDK8wo4Vq5o45KGQ2jqqeQSU7e28fZPqGs3tjVu0iNbob+YvfDWtYRWMpXOcDR4TjM8hzOlVJ62lmWC21Jwe8ht/ddSVt7XwjQ19dkkor8rzCrpJgrx99TStE+xgQw3KQbMCQR535wpQ1NfQs89BW1FNusG7lyoYn9oDj8Rg9VVpURstfQQvI6n9bD+pck+ZC+E/gDhe+MHCo61H5Sd04g062/hWNqsVOXDJ461KgZnSZXRCsr5VMTTTbC6lWIsWiR44hISSe75FhiNGaSVXr1npI46reEMDxoFCsAwKKxdbluB4Q3O0EcY3GS6mqtUxVmbCpmhzmiphPWww170gqoo1CtURkHarhVAkQjn7a83UnqKTIc3qs1rKfMnpqyrkSOiiz4WZiygsxq4yIlnVgyjvVVGQhiQb4oMJZ3XDEZ9bPZWngOxao4ry2EsIIkhcCORGeMp4bEbksVsR0Aub2vh5l01ZFJWVeamCq7vZDUJUU6z7yx8MaspVmkJBFg3Chixvh9PlmdrVU+SZrR1K1NMqQrSSAtIkIY2VEFmZfGWHdlhYnnnGvX6hl2YU9FmGn3rY3hdJKaWreHbussTRyKbqVAJBK/eZWGJKhwoPZMAIzW1zLUmZZVmRVZny7MxPM9ZTmjQ/VpDa3dpIh7lh/3e1woF7tyU1gpfqNZqTPd9cqt3aKkzM9ZOV3bDJ9qyrdpHvuCkKNrNcEzaaGuqklossy2mjkEVJBEiuQpU92JC58TsWN3dz4ielgBjX6m1BW0WoUy7L5qihTTsj0NIEcq6SKxEspI+/JJcn22r0GGsjvUDRQ6/j23JXOu1JxUfrMxrs3rnzGv2s8ihEEcYWONFFkRFHCqBYADp74ayRs17p77V8ONy+dyvNI2b5FRzOJD3sqA0cu4dQxjst7+qXvhtPU5TWoxgnqaXd0WpUSKD7OgB/FcXmuuilKdceSrltcarXltq7d99vVgBwfY9ScJhbAEhgbG3sf6v8AA4ex5dmbk1FLFDVRgXYwhZ1A91HiHzAw1PjZlVelyQgJ2fLqAPfDw4HJNLSM0znUbfO/v5Ys3sEqzT5nmyhj+syHMri9r7RDJa/l9jFazjeCxCKDxYdAf65xKuybNDQanhi3WFXFWUJ/9fTSIP8A2QXFPSLDJZngbFYsbrswJXSOt6zNtLaIyftS0tPMMz7MNcUmcUZ7t1ZaPMUDMpLcsv1miIJ5B788846R7X820bqChOW106pkWqqOCKOd6t5JRltaI56WdVZmt9XmeCUEWUGKRfvHFK6Eyen19oLUmixHGKrVumKmmogke0vXQRLXUt7dW72kZAfWQ+uB7KNRNrz6PORS1tRIVyenrtIV4ALMYY17+j4HU9xUOij1gB8scxM4vs8c4+6N1D4FbUYDZ3xHJ4r5hcd60yesyPO8xybMqVaesoKmWnniA4jlR2SVR/dEiNb2IxJPo7zVdV2n6e01l9VJTz59nNDlcjhQUNJPMqVCsD5FLH223xI/pKZWv9qct1TJtZtTZRS5hUsD1q9rU9T09Z6Vm+Mhxqvo1wjLu0WbVjAbNL5LnGeqW6B6eglMX/0WSL52x0U8olsLpBs57OKx2MMdou7+SP8ASB7RxrXUk2V5PMzZRQVNbFSE2BlieunnMrW43zSSNI3sIhc7AcR/st7PM41/nceUZbE8cXeRxy1Ip3m2s5tHHHGg3zTOfDHCnic/sqGZdDkWT1GdVsdBT2DbQZHYErEiqNztbmwHkOSbAckY7O7N9FwaZp6TRkUEmWVEcclPm017PBvA76iV1N1lcWFZUDkEikQrHHKXZPPHo2C43Prn0U+KJ9skvEYKbdkn0cNHUEMVKunoczDSLE9OXSsinmBA21MsZ210u7juIiKKI2X/AFtwWB+3vtgTtM1NU9i2VV4rOzrs5SPNNdVkcg7rOq+AkU2TwFLL9XSVdm1LKWSeQALEmF+3Ttmm7HtL5X2QdlkE03afrSnhp6aOjX9dk9FUDu4SijhauoDbYlt+qicyWBdLQDN+zzJOx/RGSdm8U8c4mqlfNq+GRGir6wgfWpF53GGKFJIoj5ojv/vsYTHvB+qn+532DYNbvGmS0HMaf3MeQ+47dg8Fz59JPP63Pu0d2zORZaymoqKCr2KEXv46OISBQOFVXd1A6DbbyxV6IXHCKQ4HDDcePTp+HTD/AFJqCbVWrMzz6otvzGomqiGYgKZZGlsD5WDAc8cWwzYFVEqxybWW6u428eo559vLzx1NiiMMDGHMBY9of2kpcERpUkb7fU2Y28ViLX29OB0AOEKm8oVAp2i4S53ED3HT8OmF3KO2wIoUqDuJew49Pjxx8sJyAyd6Qt1UA8nkL08VutvXyuL4tKBaqCVqdo5k+1BMGB/P+GO6OwLMmzzsbpaSLYTpXUWZZaz71Xu6SsSOspwbjkGVakdQBc/DHDDx7TMhIN1DC3scdSfRBzL68urMiKh2OVZfqCMt92Sjqfq0xHoe4rHuethjn/2ih7Syl2zr0JWpoqS5OB11kuqdEKajV+QZdJf6tmWYHJqqQ3aLua6KWje90uLCqVvEbAg/a8uCO1rJaqgz2nqaiPY9TTxvJuO39aY1Ml79P1yz/n547kyuaCGCXMaeeaKriYVtPLGQH3xESIbqbldyLdXH7RB4IxzJ9MXImoO0XPaym2yUjZzWTwNGbp3U9Q1XEB7CLMYgPbGP+z0l192uR9R+FpaVbVtdo9D+Vzyz3JZ/EdpFzcg+lv4DofbCbMWVroeRZroCfz/64PIRfgnjlrlQWP4/n5dcYUj33J2KeDYAW9hc2v8AP39sdwuaRVXaWAulzawH4gjBtwUhkbb1BsSQB6G/UfHGI7FVjUcKCNpG8LbngH38/jgrP4rBl3E7hbp78Hn+GBCxFR7Iig3ugswJv7cXPlxjDa5XuImZbgre9j58jzxkhbZujNhcg24B6G3PIPmPywVgzKXdPs2YWupB9v6tgQikkbk3qwuLkN9rjgn8x8cZgjMA+9o73Fh4vL+v64xmBCam4ZlspAHpb/r6c+2F0J22ub9RYn+j+7ASRuxudi3HAt1Hn/XXGC5v16244N/f18ucCEKsdwZSw3C3JvbBSNgC3a4HHlb8cCWLeI7bk8i1hf0AHwvx6nBSAg5AF+V29P54RCA3BU3uffn4YTNufF0PpbCnB3BRcE8cn54TkUcmxYjg24t/PAhP9LoJtR0bEcQs05v6RqW/hi8OzSlpv7JZclUsg+v1NbV94tm7sl0hUlDbcD3Eg4YMObX5GKR0zL3E9fWNcmKhkUH0L2T9zHF+ZOlJTaWySjalKz0uTUwLo1mLTNJO6MOhUmZTY+nFjjB0s6pA8PQ/IWrYBQV8fb4K3eiMzqKbWlNUvWr+jMohzHO3aOEFD9RoKioRwW8SHvFjHFr8Xxz1mANPFRwM6I1DkzSnfe++U8ge93xfNM0dLonXOaUczJLT6VkonSSYEmavrqSlTYm3kGLvQGDEEA8KQRihtZyIc4ztYiNkcsFEh8rKOR+K4j0Y0Om8vT5DkttdRteusFoI0siJb7trg9T8fI+2DqpEZYgncbdRzb8wfw+eDHdYjaPFawDdfQ/vGM3BQQOST4bXBHsb8k9Pzx0ayUUqCGYNfoDYdbi/4cflhCVRsBvz5nbbDkm25b34sSQR+NrkWPnhOQA3L2ZgbEdBb4jCIRMlq/qWZwytJsUsFZv2eeD8jY/LFv6D1RmOldRw1mXyxwywzh0Mib41IJsHX7ycsjr0ZC6/exScy7XNjfEtyXN5auiUq16mlAU82LDgKfnYC/7QX9o4zNIQF9JB1sV+xzUBYfH5XQcGd5d2B9oK1EUdXD2Vdo8YleFbzSZNURNztH356GWQgf8AGpZgDxNx1D2X6rp8nzbMdCa7o6TNNO59SNR5pRo+6mrqGoQHfFIT4kkjZJYZOPF3bfaLW5A0HqDIdb6XrOzXXWY/VsnzLZJDmDKXOUVsYIgrbdSiBmjmUctA7Wu0Udpz2K6hz3KKqu7Au0fK5f7VaJM02QAS7pJqRQZqnL45FJEg2FqykYXB/WqpImW2daoTaoe2ZhIz21q3Z5ewf2L/ALHKmvpKdjcvY32nV+n3qvrtLFJFLT1oFlr6Ccb6WsX2liNn/ZlSQHB/pGSjM10TrSLa39ptIZXVyuDe9TBG1DUX9+8owT/ix0X9IrLsm7Veyig1BTiN67SdQtH3oNy+U5izBOfNYa9Y2X0WrI4GObNSVC6i+jhpesdVFVo7UeY5LP4efq9dGlXCD7CaKrH/ADYvWK0i1MZKc/nVxHoqlqhMDnM1LW9nrRQ5TumRWjR5pZAxt4VjU3HuLHj+OL7Z58h7ZY6sIjDsc0FQUNIrG6jOquJViJB8xWZlJJ/6n2xRnZDlT6o1DpbSEI3NqHOKPK7f3ZaqNH/9gWxfOo62lrYtX6rFXTxT6+17m+bRbiQHy7LnanpQPK3fVMjW6HuR6Yz5zdkkJ11HE/F5TM74YzZQrKPUEekqVtQxXVdM0UmZRC+5ZKiJBDRKLgMC1TJD4T+y1iRjk3MJRJmElP3pcUyJTAqL7igszX923H54tntJzx8iyNNPrJ+ukkizOrW/O7awoYDfzCvJUMPLfH6YpmFNiFiCzNySfL3vi5oiGjTJt663FR6Qkq4N66+E5jWZXUU8R3uQq7V3OSTYAed/hgsgk5klHi6MSfFf0N/+uL7+iloWDMNS5j2n6jjIyrQ+TZrn1Orrf6xX0lDJPALfsxuYZG9zEPvYobMKCqoa2elrSXqIZGSYsee8B8XJ9TfrjTjmZJI5jdWvfs8sOKpPjc1gc7WkwfAfECepJPT0NrfxwRi3mQd9ibEG+ABNzwWANxxbj3t0woqs7Ike4m52ra556/HE+SiWXuSNwNvMMLEfPnC9JmFZTSf6qoaR43jIZBJdGFmFiLdPPyNiCMJAQxFhO7SMOe7jIuPXc3QfAX+WCvI8y7FAjjPIRBYH4+bfPDT3hSmCdS6hukZAZzI44tGf3v6/C5waSvrZKZ6YVPcQMPFFF4Ve37RHJ+JvhKMGV1ijhYsbALGN5v7Dr+ONtS5PR77ZjNLPMPEKGiIaT/1sn2Ix69T7DDXlrfuxStvHLBayjp6iukWmpoGlkteyDgD1Pko9+mNpkeYRabr5pfrCOk8DU7yonexRPcFHCmwlKsoNugNiLkYHMK8x0hpIFpoKUNc0tKSYifWSTrK3zI+HTDU5PJLFHmOdzNTQSC8KlbzTL/3cfkv942X49MRyEOb38AdWtPYLp7uJCkFYI2y+srK/PKQu8X1im3pNIK5i3iEbILB78nvAOQeQeoaI0ZJ2jaop8rrsziyrLYh9bzfNahi0dBRoR3kzdSxtZUTku7IouTiOpNN+i68U1LNFTU80Mipcuiltyku37TWHQWNunGEf0nqDLoGy4VEtGkzx1Eiodhc7f1Za3kAxIHluv1xE5spjLI3Udq3ZausdycCy8HOHdVpdpXbDSw16ZZozL5MuosppTlWUI5KzUOXgkiJObxPISZJpP9rI7tcotlxS9TUT1chnma/kABYKPQAcAYUeBt5L+K559T/nhbLaWKqr6WlnZliZ90rKt2CDliB5mwPGHWeyx2Rndz1nWiWZ07qaticw0gio1lkUhdt7287XP4A/u9cOcvUNlLMVbfWVfQfsIv8AN/yw2zqrkKLTd4WXm1hYKha4AHlfqfljcpTxLpzJoo4ik0sdRUSSbySVaXagA6L9hvfDLQ8hra6z7E+yfCBVxGoe4HurI7GstpJsjjqcyTflg1XHVVkYF91NSUclRIvzUbPi4xucx1bX6Qyeu13NVMuqtTy1dPlLKbNC0rN9fzAehUyPBCf+I8rD/ZjBOxSky09m2YVOf5x+icrGaVArK1F3SRU/d0/fd0v3pGSMxoPN5UB4JxXWu9RVOqc8n1DLRJQQPH3VBRRvvShpIwVhp1PntAsT1Zy7HlsYsTPqLTIDkHY7yNXzuFNa0JHdnCymdFpdKQhdV5LsUKsuZ0yIL9F71VGLR0Zups7zbOoJZh+hNO5xUhpVCuszRfVIyLcf7SqFj1xWeTyRUOpchlNl7mvpNxvzZJkuT874ujLVjpcv7RKg01NCJKmhytVhj7uNIpM1ed0VSSVBWjW9ycW7QR2oJ2fj3UMTe4RvWk0ro3KptRD9Il4aHLxJNNLHI0ZSmhQ98Sy9B3at8d23z5aaUzmu1hqTWufVuXRiprYKIR01wiwQtVRrHCvoERUUDpxja6vqG05pOsgeL/Wc8rGy9uft0FK6PPL7d5IaaJveKT1xD+z2oqqvNdU/UAVknpqVgsnoKuK4OKYY57HSH/xA8Lza8dfgrBc1rgwb/Q+nupc1K2Z5nkWXzQNukznLqUKtQBGqy1UYJ2/7wtzzyfPgY1Oma5Ju3+VfCabV9Zm+UybjZX+uy1EcZv7O8Z+IGJbojLpMw7QtE0Usb731NlkY7yIK0R+soxHHBYheWuSwANwBbFT6uoazKcvynPKGR4aqHuquKUE7t27eGHoQxH5EjpiaICRhiJ+6o4hRSEtcHjVjzXSOV5pI2k8jdlWI00mdMrFLgTNUUlQFI/8AST7LWtyQeDiH9s+X0k2f57liLAkeQUtPpwrGC6RvTxCSq2Em4JqpZ+t77Da5GLH07meTxS5RrFRG2RCpzzWkgZdwSlWjyzMTAR/8sxtT/PFK1jVVTpurrM6rlGcVVHUVtSxc3lkm7yR9/qS7sQeQTceFhc0Yg4vv5ZdeVKHxVhxF27nmt52a6hTN9AZNWVVYqnLe9yCdpFt3Ukfjga458cDshY/ZSGQDk86rSksehdf12lM8qko9Oa4jSmq2kIKUc+49xUm3RY5S6sf+HJJbjDHs9qocv1/qHs8niVabPYaeopFLbAK6mjWVbN93ejTpf1dcSfVuTZXV5rk9Lmfc/UkqostaZPEk1G7GIOGtz+rlDXNv9mDYdMSuuxzOj1Px8jiOBqB4IbV8Yfrb7YcwrM7b8+1fl/ZfS6uoqeGkzqjznIczqadorscyp6espswSTnxjvoGYqBYpKT0bEVo8uyDtAy0VlLlKDL82ginpI6cgSSU036o0xJuDIkpaLda9o1JBPJQfXFRqPswi0Jqh2OrtJ6vp6PNIS22Wrp4aaog+sqTwSDDGkl/vIjE+M40PZXJTZXqGv7JM0jYU2ayDN9OoWIvJf9fRA9fGI7pbrJAgH274rOjdGx2Heaa4f04V4Z+FU9rw9w/pIp59YcFPcpybLdQZZRazzegpajO4om07ndatpGFfQIoSZtobelRRiGTeniElJKwuCwMH7asqhzzRMm+NZWyWtp6uKeJQVejqwYJGRlZkZRLHTG6MV8RFlN1FmaeroMv1jX6HoKgRU2f01FRURiIh7vNId0+WSq1isTSfrKbz8NYC3HGNdm2m4tSCvpsiMeYUue5ZX0Mk9LTillhqZqczUy19EOKecTwQgTRgRSn0YjCtfSRto24+G3lknOb3DDs6HNQzsV7TqnOqaHsp1lK36epIIaXJaq/Oa0SOHjoXP3nTbeA3uy3h6iIC8tH5pGNS5zlE1I1XlNRK2YUSVAKuZBCor1XYQ0ZkjiFUoU33Uc3B7xgeQotLnVlLRzUtQtLWPGs9DKHKFJiA6xk/dVyRY/ccqejYvrsz7W6zXWTVFY031HX+QvDW5mzx2NQ1K90zJY/2gwC1SDyLyW2yybZLVG1pMrMjg4bK6xu99yjs7iR2b8x9p27utS2vbTkRodL5a61FUJckzjuYJKixkWlzBbWFQgAkUVUERvZXVpy1rMGNL6p0rlGr8ios9zSWVKqgbbVTwQASyUbG8jELw0kRIkFgNwEi9dt+h9b02X6i7O9T0OSPV0Qq9OVWa5VRd4zLTTUkgqVgPNu8p5KaohBH/CU9JBindC55HqTLKTVeXx0sGX1MhhzeA2CUVUF3MLm5RHUmSK1yfsC+zDLO98Ud/W00+OONE6VrJH3RkRXrwUz7Ee0TMZdZ1Gjtd5qYdWU0AyzMKuFt/wDaDLwu+nr4r8PW06sswI5qKdn4Mine17YtR02ku1iE5xVQZdT61y2DMRmmxXgpM1T/AFWsWQdDBLNTiXeLlDKsgupYGC6z0amb5DS5nlFSabNNKlIYamBmSZKAzkU0gPUNTVDCEnqI56f9nGp7Ytdr2l9leVVepJaej1xpHOXWupigRMxpKuFAKyEdOZIFEqLwGk3AAMQtpkEc0oLRg4UO7YRurq1VpkoHyPibQ5g9Dh1VX7oCveeszfSGYU1VDU5gJpYIXO6Rqwp+uhDcgmeFB3ZBsZqalIv3lzzN9JDTK0WvzqanaGak1RTjNBLEtonqdxjqivopmRpR/dnXEp7KO0OTUeUZfpl8zkpc9y8LFp2uMtpHdCHjoJGP+9RlD00h67e6J5XbIe3N4NfdlUGtsvggSWhzAZpJTwps+qPUMKbMaUKRdVjrEpZEU9IqxB9w4ksBkgnuyZ5Hw1H25aklruSxVb4j3HuuakjjiCd3u8XDC44PsfMfG3zwuIXqHFGZVp3DDd3oKgX824vYetvxGEC0bi4KlSODa3Hv/G+D/pGSnUUldB9ap04RSbSRj1jfqPhyp9MdG69TBY7aE4rFmzPTtbK2W1/dOfC7wSB4pl9/uyKfQg/AYdJXZRmdkzLL/qErf77L0BiJ9WpyQB/6sr7LhJcvWpiery6RqyGLxzIF2zwr5l054/vrdfW2G7pHK4SliLBhwBckj4f5/DEYa1+OvgevFPJczA5ddYJzWZfLBG1VGaapprBfrFIDst5Bhw0Z9mA+fXGtI3A71+15Aefr/XXC0c9TFL9YgmkjmQcuhINvc+/oeDgTV09U4+uRiAn/AH0MfF/Vox+9bH2OH1c3PHrZ14JlGuyw66+VsNEVESaj/RM8kS0udwPl0rMv2S9jGefs2lWO/tcdDjdvl9Y8K1yUtVGhLF3O17SrywZgdt127drjoeeuG+g6KmpKjNdUT91VTZMsC0QWNZ0WeZiFqCh+0saqzAEHxlLjBcznc18VeFjWa4qnkenjEvidbMSiAsQVLEshFr2vc4qPN6U3d3H9KKdouxiqWR610IWWcRXaVKeml/1VJD5xgE7Bz0W32l9xjdUs+XZoIoBk0FGBxSKsLiCMFSrI8p5PecEyMbiRQSApONNDmVXLNV5i+YU9TIZ55qiZp4rMXcszm9gyvc8bbHpYWwaGtp69EpqWkAJhDuq1DMhBPiPPCBRYncSBe1zcDDXMvBKHUTyTMstmjZKKiigUpLDGKiVGlUFNu9y7KgcWuu0WB/vYbZxHVTayrtRZbU0UYzeL69S01VNHKauKQWcElRE7hlbch2ncOORgczioqzJ4aekMKM8izvIwdGaoWMRuqM42urFRIACG3MRbpjQ6lOaUFflOUy0M31ynoljlpZ4gz95LI7lGjuSD4x4TZuhsDhY2Auw3jHZ0EOcQMUwrZpcxqF+t1Iu7FmkkNgt+WJsDyfhfCNUaJGWmol3LbeZC4djcdAwABHnYqCDfFk5BpHK8gppG1dk+V5xmFce7FAuYbhl9IiO9TUSNE36uZAqBAxYAlrg4rOCmEal2QsWtsBNmHn+Y6j8MWY5GyOIbkOChc0sALtayKJImEg6jzB5HuCPP/PDiIVtU4pYY56maYlQiRmV2IF7Cw3HgX49MIrM0ZKqi8jaRa4+HucbjQuf1GltaZRqymYs+n6uHMigJIkEUikx+4YXU+xxI8hrS5MaKkBR2cCRLgcEdQ1x+fOC5PVSUOYJVQkhoHWoT1ujBv3A4s7tz7KT2ea9zyjoADp+WuM2T1KglXo6iNamlvbpup5Y2UjhgrW+wwFVF2pZ0mQcq1xfz+OGVEsdW606hjfiuxOw7XtXp6Ciq6CRnlyGsjr6dN52OYpBKq26eJRsv1s+JF2Z0lL2a9o3bt2WZbMVoYY4dWZCxXw/UxJdHRW8zR5gnP/dn0xz12Xan/RNRSzx1DBFKQtze69YifiA0Z9Cgx0JLEM27Tey3U1Eb/wBpdP512dZjMR9uaCCUUjP/AHmp6ijt/g9sck6O520ByIJ824rec6vZyjUaeRVNfSKgtlGQSEk91JmEKEm52Crjf/20rH4k4jfZ3FTZX2R9pWq5pe7eXL6DT9Kb23S1tYHkHygpHJ+ONl2650uZ6e0zApBdqGWrYDqDPV+EH3tDjU5vTNlfYxo/TJUD+0ud1+oKsgc/VaVBSQn4eGrP/NjUsxrY42O1u9CXHiAqE4/vDi3Z64e6lvYNkbabp/7dSALVpNEaMvGGVKtgZIGIPBEMatVEHguaUEWx0Mmqcl7Gez6p7TM5jStgpJFo8myyd98eb5qy94kcgPiaGMfr6hhyy7UJ3TYr7sygr8x/s3pOCnokqWLzVJq/DBDNOBNUTSn7scFPHEpb7qQNgtPldF9JvtHaup6evXsg7PSmVZZTpIIKjN2di608RY2Wsr5EaWRzxBALk2iQGgWi2zumn+xuJ2bh1qwVwu+njEcX3FbrsTyTPc1zPOPpXa9y6pqc+1XUzzZWJKzu2poZLpV5kHcEqXYSU9MLWQd8w4iQmv8Atu7SIM1GYZmkfdU0MDZPRIHubyqDUSCwA8MACCwFjUL6WxPu2/tRqa8w6QyqnpY6hkSJoKfakFMqIFESbTtSGKNAgtwsUYJudxPJWs88Gb5gmXUE7y0dKCqErYyXbc0hHkXbxW8lEa/dxLZI3W+0du8Yahsb+flMne2yw9m3z8VqaaRyJKkHZI7FjbgIW689LW4tjFszFUeE8dQt149+gA9uMEUgJtWxAAN/Ic8/EfHoenXC9nIEXiAfhlufFzdbg+nUevnjqQKLEzQKIhdQGJUsBa48hY3sffgW4xglF+Gt0ZQAVs3930BHlzfAWVH2oHl3Du1Ckqdx6fZJvY+XngUDC8aXD7rHYbG3mLeZ/PrgSLXVe760oY+IqUK2+yLcDFz/AESMwkHallGSKwC59R5vkLA3szVFE5iBtzbvkjPHIIuOcUzXBEdGR1IV7Dnn3uOoxN+w3OW012n6Yz0SbBlGp8rq2P8A3f1gK/8A7FsZ+k4+0sj27jzFPdWrI67MD1tXeuWzwVtFls86kCpallmP1jvIwxF2uWVd3LLZmA5KhgDcmmfpU0Ej6XpZqqmVKn9AZS9ttiGpxUZbOWHkd+VwfiMXBWSNpnM6zJ3WGoiy2rqqYJVreOSBZZU7trWuhEZAPUA352i8G7ccsizTs9yBIopSrxZ/lbNM2+RY4czoKuFWf7xEeYS8+fJ88cVohwbKXDceBB9F0luF5gHlxC4mKlt6oN3nwLke9v4jBf1gLMngtb7K8X9LdD5nnCu0uvdzKLpdWFuCRxza5B+WCgA3NySTturchrcdehHv1HGPRAuSSdgBwoI428+X7vh0GCkg+DvEAY/ZZinzsQfwucHsT4gLgm1lUjp169Plg0sUsJAKta3gtu5HmBbn+ucCEhIqbLsb2PG6wuvwPnhMgIeGsOlyOP52/P0wc3a4DHa1jaw8XxHT88JgWub3HN13EAfMdb+uFQjggPzDyrEENz5cc/jzjMAF3DczGxPX+f8APzxmBCRZrsSDa7E2+H9fnggG03PHHsSPkMG2bRuUEhhcN6+vzB8+uDEXaw8hawHQW/f74EJPaSBusLiwBHIPoP34A+NWZjyepAsCfXBwCxEagG3Fgt+P3n54DxNezHgkEX5PvhEIrsSoBAH7JJ8sIy267bi3F7dfbCpsLA8Ei3Xg/wDXDaR7gjqcCFuMnTbp7OZyOZWp6dfXksxH/sRjoGsozR11bBTq++CTuFDNZf1EaQjjyUiO3zDeWKU0Tln6SXJcsIv+k9QU0G31G5F/9zOL0danMM1qJpKld8k08oMjhftSM9gegB3G1+ARzYNxzWkpP3uG/wD2j2WzZG/u+Hv8pPNHhh0JnfdK7PmeoNMZVd12kIgrqyVSPI3WO/lfnoRjnrNZ5KiaqndgGnzSaS49v+uL/wBWz1tLorTtPUoirVanzqsZhFsZ1y/L6WBCw8rPLKLfHHO5DPSUZdTukE0x97ta/wCWLWiG4uduHx/tVe3HADefb5RVUAePghg/Tr6gj9x88FZBe4uw6A88j1B8vh64M7Xv4OQLkjkW/h+7GIN52KbHkknp+H8cbizkYoVHeSEWcXv0v72PUfxwmx3GwI8Q29ev+eFAP1TN4rg+RuB63Hkf34TXba1gwYW58j54EJtNGCORzglHVzUFStRDYkXDKwurqeCpHmCOMbqbTmbx5RDnZpC9HUh2WRTuKhG2liPIX4J8uL2uL6aaLztYnDe68UzCUVYa61N8ozIqY80yiZ77gCl/Er9dp/veh+91HNxi3cjz1+0SjyOnyzMlyftA0o8UujM5VxGzvE++LLZnPAG/mmkPEbnuW/VumzmqkranL5TJA1tw2upF1dfQjzGJRRagjrejFJSLMp5Y/wDu49+H/wAWMt9nfZ39ozEe2wrRbKy0NuuwPWS7AfOMv7QdNrqHJqD9EUmtoMx03m2VLGUXJdQdyZmpNh5SNqmCCpgB6K0kfWI4oPs6pGzzTfav2fTRbJMyyFdUUEZPIqcvdakqB6mmmqh/y4sTse7Tchz6uqafV1dLT1WaxQU+ezQ8vXinbvKXNY1P/wBk6N1DkED61CJU/wBozCSMieo7EfpK5BFq9aZEp5qemzVqd99NUUFUrwyTRN0eGWlqO8Q+aFffEbIxE49kMDjTYc/ZD3GRtJDiOYWh+i/V0eX9pOmM8q2Bi0/mlVmJB8hFQTzof/FFjZ6v1NBkU2Qae7gVEWl9MZZRdwX4mrZkasnV/S81U28/sxHzxBqPI867Pe1rONBkyh8vrK3KKqURl9kCb4pprDyFOZGuOgN8R/Wmop82z/NKpWb/AMoVUs1yLbYmPgW3qUCX9hb1wj7MZbSRqIB5n5Ka2UMjDtYw65JrnmfVWocyepq6l57yPNLMeDPM5u8h9L8ADyUAYkfZxoes7RNTw5LR1FJTRhXqa2tqrino6aJS8s8tudiIpYgctYKOWAxCkVY0Fup8geT7YvLR9DWaJ0JBRUEW7N9VRR1lVZbt9U7z/VKf1/XTIJSPNYYR97m1apPpIRHFgTgPc9a8FDAzt3l8mQ6AXT3YHo+HV2Ta90XpbK/qmXV3ZbqWhyZKv/zwRvMiR1MgHH1quq6eoLjkJHSxRLwtzwvr2RZNVVdWqELmSwZhECObTwpIQPXksPkcehv0bdWQ6Y1n2kZtlXcyUmj8/wBJ6PpWkuRJTZWKiSs2+rSyRTOT0vLc44t+k/oWXs77U870zTnu6fJc4zXJISvH6iGqMsAv/wDK9RD8sVbDMGS9gMx/uF7HfgVNaGF0faaj7GnuqjVBE26QG4+6PtD4ny/f7YO0rvaJbRI3BC35/wAR6n93thEEIDcL7C38sHUb+AwB6+KwHra4/ljZu1xKzw6mSWijgDKRKVP3B3ZN/a1/PCscNNGe7LQRn1kJlf4BF4B+JwxZSQGQoQRuCswufljE2tZQLbhZR159reeAtrrQDRbE11OkYVBPOCPvMIYz/wAqeI/AthIS1tWGp4o4Vp0szKAI4I/dvIn43OESsUMhFQCzHkwq1uf7x8vgOfhjGnlkssoAjUnYi+FU+Hl8zc++GBo/lCcSdaeR1UNH+tgK1M4NlmmT9WhHQpGevsW/8Ixrp6itqKyWoqppaiWU3aRyWdj64zwuTGjBrtcWHX+P4YWCQtGEkD94wXgHg/E+XsOcLdDTXWi8XCmpSKkyNp9EyV8lEFkrsxEdPK6kM6xx2KqfNS8lj6lTY+E41OeQPDWVX1qYTSCWRWlUDbIVO2/PRbggewGJFnE2a0WSadENNUlctyeGonmjLbI1eqmeHePsi266nqSx4NjhLOszi08I8tokMs08cdR9ZkhCpHG68NBGxY72W/6yTxeQVb3xUY916oxqT7fhWHNF2hwpTr1UOqqOqo5Pq9XTSQyMiSBW4Oxhdbj0INxfnDzTuY0uU/pKplpBPNPRvR024eFHksGcn2TdYdbthu0cYkYoWKhmsT9tlvxf3ta58sNpFfaIi5CRkkC/Av5/HFotvtulV2m6ahI1T95Izs+9jclvI/DFgajymTJa45VIxvl0FPRMB1DrChkH/jlfEP01lv6Z1PlGUbL/AFytggIA6h5AOnwJxbK5RJ2mdobZTSzx0sups8khikYX2RzVDgyewSHdIT0Cp64zbbJdma3UASfMinoVcs7KxOO0jlWvqFq49S5bN2c5FoTLw6yQ5nWZxm8jIVUs7LHSxg+ahVaQnpdh6Yi+Yh0oSjKVkYrFyPvlgDx63U/hiW68zek1ZqrNM+y9DDleZSk0cQXYY6KImmgQ+m2COIkepxGa5VXP8iymWoWWSaqgnqLDoWdVW/uRdrf3hiCBrYq0GZLj45nnh4UUshL6V3D2HzxWpz5GoXtAwAp3WSIrJvHFjcMOtyL/AB9DjoevpEqsurYKCN56nU2p1mookHMjLQRNEhHoZcwtfy64o/UtCzwVUIiBYMx3belma9vj5k+mL1pc0TTXZno3tB2t+kJ8lqIsrLDwnN6h/qvffCnpaVpf8bQ4rWuS82MjM1HmQCPRTQC6XV3H1Huq37S8xpq7UtRRZbU/WMu0/CmQ0jK1+/ijuk8v/rZmlkB8wFGI52YVazZxqCGa7RT5NKxtblYpY3HUee0nDvUNJBl2WSSUVNJCYacQKHs1y5Cx28wftGx5BB6g4Z9mMLJr7LKGKyrmENVlzBvPdCwUfjbE77rbK8jUPTFRtqZ2jafXBXj2W09FD2r6HmqHjhqqPOYnhRYCySqiySMRKDww2g7WXn7p8sV9m+URzaQp4u6k3yUkMsYAvvBiVmjYe19yN5Ash4KWnXZnWUdf2i6LMLLHJ+k3ZEDbibUc7bgfMfZFvK1jiNySRSZPl1IyTbBl9M6IG8akQqXiBHUqQzIR9pCR1S2KbS5vXirJAdii5Nrejqvo7x6Rp543zymzuTKGpw/jOTyFK12t1AMsQi9LHb5jDKdaoaYzWSesjjhqaOUoJHRFkk2FgCHIG7aATbxFuACca36gtD2g1Om6bLqearotNpG0cvhD1ICVMoci3QO6nkcC3FsSHOVo8yglnky+kcQwNG8FNN3iIniBUSEcjcSxaxJIbnjDwGsdgMCb3+LFR4uHhhwwUa7SIq/IddzalytTDVZZWUdXTFx0ZaWncAj0Nxf2JxbWpoMvzzJu/wAnYDJtS0i1VAjkWiSdSQq3Ny8T94p29DAOgPMJ7RIEfMJkqNsgq8syaZSw8SmXK6ZgpPo1jtP7SkfeGEexbUlVW5LWaLmdJajT8k1dQs0avIKGYqKpE3AgbJBFN/dHenyxHaWO7NsmtlB5GnoaeRKfC4B5b/V6/keynmc5Cuo86btGgp0Wqz7Lstz6OaNgGE9VTiKpS9jZTVUtQOQRv8P3ua+7XMqmRtO6gySolpp6elq5Y5I/A8MsFZ9oEGwIJUgjrYHi9hb2WxU66NgrFklWo01m2ZZXFLGBKFpK2MZhFHKp4eAyRVyG4YA24tciCdoDNFTacaTL4lQfpYzRLUnYyCqpm3Ru12K2YEX3WB5LAXwxklJw5mrLwoU8spFdPWKkOV6joe0fSra1y+uXLK3NZxHnEECf+Z5zGgkWVLcokpjWojsOGSRBwgxPJ6uozauyPtMhpJFapkgzhu4Gyoy2YVNq5InXloEq4qnfTyXXZIhXabX5g0pqROzHXk1TWUdXNprNLUebUsA2SKhIkSSIN0mhO2SO/mpU8Owx0xpdnpdPZ7kUWcRvLQhNSZXmERKQT0VWIoKiWMnpExGX1IH3Ck4PKtiO1Q9kwtZkcW+44YeFM8U6zyX3AuzGB9iqmekost1NqbT8Kyw1mkM7q8vqkiPP1YVEgppkHmu090wPFzH+1xHdRU+Z6bly3tG0hmbZfqTIphURyKQWnpQ2wO3k7RteGQMLSROha9mw+7Y85rdJ9uGadoEUAqIszWiqsxpgwC1NPV0MDzIbcC7F7HydVPUY2VbRxVTRmGqWqy6aKCSjqxT/AO3o5riJyi++9ZOpEm9SCekzHFhZKMQ4CvmMRz57lGQHhzDmD+h691bfY9qvLNeJS1mTUxgonzKCWuyg3P6Kqp9tLmNESbk088TR1NOx+9Tyxm7C7c9aeqR2UZrT1VYkgyLMWfI8+giG5gYpCY6lQeBKhUSr7oy9GOJRpRtTdl2cZd2g6NjFXLltaMgzrLmk/V5nRv8Ar6VXPqyRyxq/VZKVGFiOR7YBlef6c1XmeRU06x/X6jM4FnhMcqKlZuAdfJhFVsG8robcYGtDJbg+x9B6jHwqCD7hGJbeP3Nr8/qFLMwr6TIK2+bh6ulRD9e+qi8VRl80e2Z43+8ssLLJH5hkivyOG+pOzyUUub5PP3FVUT5XmmWCoUIQ8qRCpjZfOztAkqED7NQOeuK67I9UnU2TwaTraiVqvJYpZKBBZmqqIXeaksSNzxHdLGDwUMq9QoxfWlMxpJMn03m1VGkgyurp8nrdwDFliASI7/79BUKt/M0Z8sMka6ztMZ+5pr5bet4Tm0m7wyK4XizCWhqTJELxuLSRkkBl6246EdQRyDyMdW9ieoqftkyrNNBV9SZc01BS1FLUb7b55Hp9iVyjzmV0plqQOXVIpxyspxy5qbJJco1HmeSvffl9XNSG45JjdkP/ALU4HS2p890XndDqLTtfLRZjltTFWUs8bENFNG25HHuCPmCQeCcdFabN9RHfjwdqPXXALJim7N112SWniljcvcxyn/aIem7zB9wbjjAxSx8CrgLI3CuhsQfY9D/hP5dcbPUOdUmos+zPP6ah+pRZjWS1f1YAbYmkYuyLboAzNtH7NhjXOslM4LqCJVD2ZbpKnkfce45B9Di20ktBOB2KuQLxok0eWnkWqpppI5IW3RzRkqyny5HQ/wBc4fNmtFVU8pzXLA1Q7AmanAj388kjor+44P3l88ILS/WkMuX7mZAWeA8yKPMj/iL+Y8x54YtOjWBfjyN/zwha2Q11jilvOZ4J/W0TRQ/W6KqXMKFOe9VCrw+0i/ajP4qfI4aSwyQMI5ozG7KrkN+ywuvIv1BB/lgYqiammWsoqhoJ04DobED0PqD6Hg42UIyrOE7qXu8rzA/ZZfDTVDehH+6Y/wDgP93CVdH92I29e3BFA77cCtt2e1NDLLW6Vr+4p3zt6Y0dZKo2wVkLlolc2No5C7Ix+6SjHhTh5WZrqSOihy6pRZKCkkZPq822VKdlfxRqb95ARazKjL52HOILO1VldaoqICslNKrFHv8AdYG35YsfVNLFFqXMmkoZYhLO0irLde9Rv1isSbDnf5Xt1tzitIxokvbceFB6UUrC67TYo9TR5bNnaV1dR1UdK8zt3UcneNGxFlKF1G8o1mKN9oDbcE3w/wA1y/MqSmkqJnjEFVKsnfxI/d1d/wDeGa2083IjbbYnlbjDpsnFRS0e/MMrhpjDLPNURh5aiEM+wRThPF3o42IbCzA7ueNJmExq80jocurjSrmFYsReCMxq6qFVHMEbFLXvxcm9/e6B144IpQYp4+o5dL5fRyZQ4jrM4ikqZK0rukp1ErRqkR5EbHaSzjxcgAgDnWaZpMwl1PlkmUxyS5jLXQmnC+NpJC48/Mm5uficbPU8AzHLKHIcmhpc0kyCSqmrswy2F+6cyMoAUEA7V7sksq7bsfW+GOh85g05qmir6ir7qjdZ6aeTnaqTQvExO3xAeMXK8gXIw+P+E5wGOPidnJI4G+AThh+Vuqihy/I8ozury+pqXhqaj9H0cT/7PuzMzd4D0cNFDbp0bm+NHW5hp7MIj3uVyZbUEWMlK7SQE/4Guy/IsMP8yzbKGo6XIpap5zDJJWSy0EqyqsxVI40DPYOAiEm1uXFuhxGcyqaeSULQRVMTEuZpZZF3S7iONqjaoFvLqThYWk51rtSyEaqUTaajq0jaaApUQr1eI3A+NunztiXdimSTah13lMTxh4pM0oopAVuCu/vGBHpsja/tiJ0qvETUQyPGy8b1JUj4nF5/Rly1pKypziJS1RBBnFZG9rkMlEKdG+UtanzGINKSuhsr6ZkYKWwta+YE6l03l+ksm+kz2AUelaahD600jkuWQZKYlCNmmWTU7zQUDseO+WWCrNK5++ncmwlN+ANWabqMlqWppxuuokikAKrJGb2YA8joQQeVZWU8rju/sWqpdB9qGkNIJUzik1d2e5vkUcamxauyiuqqijYEciQCJVBHPj98Vp9MrSFPqKOPtayXL1hfM6grncMK2SLMmjMvfoo4EdZEjTWHAqIKofeGK1mtl20CI/a4Aj063eCmmhvwl+tpoeusfFcpaXzk0M/1aUFla6hQbblJBKX8jcBlPkwHqcdI9j+sZM0o6agq5yZtJ6s05q6lmtbfT/W1oakj2ZJqckeRQjyxypLEVk3A9Te44xaXY7qVqfUtPBW18dPSZnTz5PWPK+1Iu/2mOYnyUTRxMfQgnzxLpGytIMo6/XJR2ackdmev0S3a1BJTaqlyOG80lBXzUMSAfaWnnlRFA92IFvfEy1hHltN2uz6GkZJcv0PkNHpViOQXi2NWt8WlaqF/fGk7La+i1l9IY6t1Am7Jsjrcw1XmAPK/V6RpKxlP+KRY099w9cI9iun63tX1/mua55BW1lLUVX6Vz0UgtNNE0hdoEY8K880iQqxIC72ckBGIrGJzY7jjS60cXH2APFPbIO0vbzyCtvNqDUE+m6DQmmVih1Z2oUsmY5hUysViyTS4cvLPM45RJ2QknqYIVUAmcA7LUHaPp3s60zR9nPZ9RI9Nlsb01LM6hpZXlsaipkXkGpndVZjf9WiRQg7Y3Jjuvu1VaSuzjJ8hraOtzTO5Y6jUec0a2gq3iAEFFRi24ZdSKiJCLASvGJmsqxAUHqTW0SvJDl7ionkG15DYqB6G3B/wjg/eL4rCzumIs7MWjE7zt8BkFYEgjrM84+gWy1lq56GOpggqRPX5im2eVTuHdk37tT5pcAlv94QPuAboLTxMUaaW53eJmZrXN+effCUYqauZ6uocyyu25mc3JPrh9HDvFjIqsOQd1uvr7+/n8cdDZrO2zsoM1lTTGV25HEZsYiJLMvLBhyvXgDqOB5+WBiCCmcvbcTa1rg+vmLX8wR8MbGu0znGU5Hl+e1lN3dLmM1RDAzxlC7wkd5a9twXcoJX7JIUndcDVkxlgpL7m5tvvf058/gRfE4IOSipRGIKqGWOPxBlYKeQPQjy9vbGOFfu498blyCQ1utrBdx6/uv54wxHZYoWuCb89L2ubX4v68H2wIS5SMi+7oqgeM+37Xtz8MCRM82U7Nm1wYyb7h5/17nC2UySR1eYPC21xSmoQ+jJtcH8Rglcrx0rgE7SAqrc2HNzYHp0598ZkFpM4ihJ4qaaSE/8ANGw/fiK0fwiTq9lLD/EC9FtTZ9TvrXM8yTMZK2KeqgrWjSYSSLFPBHUAdzKCrKO98IXwm5synrpu1DL0l7LZWy+pLpl2e968i8qUrMlmA2tYXDSZWhvbqbeWF9OVEOp9L6afNqtp6ePR+TVqoUDNBGtAqSKPsvctFcEMQC/UdMK6sWtfs61JHWuJVWp0/VzOE2ORHXy0xD2AVjszIjeoG7z5BJ84sP7qYRbiOS62ejob/gea4LzVDHmtbGpBEFTLcA2KjeTfnysev44aSuLlw1yOgAIAHne/2fxPth7nAvnFVfgu0b7bfeaNb/nfGvNiWsFuRbqOLdCB/XBx6REaxtO4LkXijiN6Fi3CEjcAW28Ei/Xjp8r3wG482hAvbaUP2f44BQB4Txtvc+R9wPLjy64OwUOvBIYfaIsRfpwfO3lfEqYskjBJLA3I3Djlubfu6Hz/ADwkwTw9NvXcAQCT6k/h6emFHGxiFQqG5Ata/uB8PQ8YRkYhrAEgnjmwPr62+HlhELOR9kElfXjj2/r8sZhMt3nhJX2HAII9cZgQikXPht4jfzHPlxfrxbAhBKCSRz4gfU9SPj+/AMwDAq1h7Dm/kfh5fjgoIVTsPlcj/P8AlhUIxsWsoChiAL+IL8PPANvS+0kg/C5wncglfM25HPxAwViByxJBsLny9PlgQgkAFltfqOhw3e/JJ+ZPXDg3KEX44ufLphvJYG4B/r3wIVmdjsfea00NDIoMaZo9YVA5Pdnfc/8AwsWxZ9OyxpSEvMveJEB9hm5C362Unm3kCbDzOK27IWWDVWR1LSKgospzOq3N0UinnI/eMWK6M8US1FNKi92EF13AqNvhtwDwx4B5BNjjkrfjaD4f7nLes2EQ8fZqY9pkstFpXSNI0jNt0/qXNA7Xue/zeeFSb83KwqOecULF4UgUmx+r3HNurE4u3tkE8em9NPUTSSvF2e5cQztdj9ZzOqmuSeSSHBN8UoAzmOMIpb6vGAqgncLE8388bWjRUPdvp7+6zLWftHWz2RXdnG5rEHqD5G3p19MZGCXUhGNm+z0JuLcE9cDchQCtylwW9B6EYFwUG0qAL3BHJt/EeeNNVFkQVWJKAi3N7g8dbHqD+OE2copIJAbg89fTBwzK7bXIdSbbTytvMHCchK7rWsbc9D/XscIhWL2M9qVPpBq3TmqssmzrSlY4nraKCRY6ulcDb9doJWBENTGD0IMcqbo5VZbWsrWn0V83zrJYO0HsyCas0tmaPNRZ3pymMoYJzJHWZYpNRRzx3HeLGJIxcMPCQTzPSVTUFbHVJ4ijXK+o6EfMXGLv7HO1TU3ZNXNX6boo88yiolQ5nkM08kS1gQboqinliIlpqtEa6TREOCD9obkNSQ9i/u61YY0SNqdSqTONKZrlSvJLFHNCjFWnppBLGp9Gtyh9mCn2xpCjLzb549LVp+zz6RGmx2g00FNrrJ1CRVNXX2o9TZJK3SnraylCyMbi0c8qSwygD7LXTFRa3+hnleexy1nZhnMldUkkHJ807mjzHf8AswVC2o6xv7j9xKfIE4gZpKPtDFL3XbD7KV1ifc7RmI3LjynzioiYd6O9A4DXs49LN14974uPIamh7d9J5R2d12YRw6009CaHTFZUssaZpRliy5TM7GyTIzMaWRjtO8wsV/V2rjVfZxqHSdbW0eZUNTE2Xz/V6tJqd4JqSX/hzwuA8LezCx8icaSjkmy2cShVJXhlceF181YehxdLRS8xV7xJuyLq/LsulnzOu7ZM5glgrsx7Mc6ocwjnRo5YM9pVjyqpV1YXVyk0MhB5vIb45dzd4qjNaupSwR532+gUEhfyAx0fobtCzbtW7O+1HJs+zGbMc4XT8GoMtqpVBqHWkkp4cxikYf7SRqVYXZz4nFKrNc3OObKiCWlkeOZRcFhz0JDEGx9bjFWyN/eOcRSmA8Pxl5J8+DQ3ejZHRjNc/oMrZwqT1CI7A3AUnk39hc/LHTXZ1XU+ddolBqOupGjpcsklz1oT9laPLqZqiOP2Cx0sCW6fjjnbQ0AbOmqtl+5hnZfLxd0QPzfF/ZDTjIuzvtK1JTyuwoNJNlUMcg8UcmY11NS8Hz/Vd/04te2KtvIfO1mvLif0VizgshJVgfRkjzSl7GKfMKyqj+s6p1FnGaS9/wD7/ZFT04YkkH7clR0N7k++Id9OWtgzLXNXm6xqr5rR6XzxuLkyVWSrHO1/O706/hifaAzmmyfsy7P9OGBS1JpaPMKkd0ZgprK2qqC7ot3ClDHdlDWHW3BxS/0q6hZs209ClUlQE0Xpde9jJKv+omdSpPltYWxnWN7pbfK47RTyNPRW5wGWSMDYeYqqGJ5v4SCCPD0IPH9XwAJFrgsD1AJH9H8sDEru3dqgZm4Fuow6X6rTMWlCVM1vsXvEn+Ij7R9hx6k9MdS51PFYgbVJ08Mndd+83cxcp3jfeB8lA5Y/Dgeoxj1CIpSkTuri29j+sf2v90ew+ZOEppZJn76ZjITxc9APIC32R8OMAtrXfoLX9P8Ar7YQNJxclrTJC8d9uxeLm4449sJsQo8ifLjgfywa5NvACelyL4A8KoCgAXuebNz5/uw9MR6arlpd5jAJcD7Q4B+Hn8OnxwcmaWneaaRth3EM3m5HIHqel/IYeJl0NBH3+bqTMVvHRAkMPQynqi/3ftH+6OcL6epabPtQp+mqhY6Ckhkq6raAo7iJS5iQDgbrBAB5tfEDpGgF4yGv4UwYahpz6zUn15mU2ZVD5fDWrHl6QU1VDBAhjiDNBECzAkkvGDYljYAHaFucQ2tpvq09TCwhJE+5DCQykgkNtYcMpAuD0ONpTrnmc5lW1RXu5HleWSzAAM112IfM2baB0sLngXwM6wCWeXuoJ1jBkjEEoiSQCwui2JtcXsvJ5PAxFE0RARjV+E+R3aEuKa11CscMlTSU9TJQykmmqJlVC8YNtzopNjfi17X9b41bKsNM8858DkhEPPeSD+A8z8sbfLp6iLJqnNIylPJRSKkdQigiYSsd0DA8OLAuAb2AN+oxpa1pp3FZXyFnZQI4wAvh8uBwq+gHX88SsJJIOpMcAACNaknZBF3esv0/KN0eQ0NXmzXHVoom7v8AGRkHzxYeh0GQ5Ln+qgqx1EdINNZZIw8ZrK2IrUyr6tFRrPz5NPH6jES7PEiy3Q2o88qCEbNKmlymMgdYUvVVFva0UKn/ABjEt1tDJk1NkekpGA/Q1LJNX++a1G2eqDe6IIYAfIwN64wrQ7tbU8+DfICp5uLStSJtyBo8TxNPQAprBk9DJNM+ZhocuoIjW1Dp9yn7sB1H94skaqPMyrir8uzKbM9c0eYSxqjz5jCwRT4Y/wBatlHsosB8MTftFzt8hyUaXVmWuqmSevU/ajYC8UDe6Bt7D/iOB1ixXulI2GqMm3cbq+mIHt3q4s2Rh7B8rtYoPCmfmVHO4CVkY249blP5KKqzXNBldOzvLV14o4geFvJUbOL8nlvIW64kmcZ/ndTT5VkGZdx+hNJ/X8symGNCdiSVbySyuPNmYBTbpGgt0wr2ewUX/azlcNUhFPRZ3PV1DHgLFS99Ox/+gjGvrZnqqOnlpoGNROkdU8DN4hMzbJlB6/7Sxv5B1PS+KDSHltRgKEedR6V4lTnAFarOjJPT5XQ1ERj72plmYDnwxqY4ufdzJ7G18R3K8yfS+pso1GoCjK81pqklRxtVgx/IHG7zivp585qoYJ1+r0DLl8EnHIgAG4Dz3SF2t740ufLS1mV5o9Iki91LRyMjL/smfvFZQfvLfkH0IB5GLkfeFxwwOB8/1oq7zd7wzHt+ivLQ9M+U9s+nqNkhlWl1BPT083dhWERiqBZbdAVdSbjng41+k8ji1bDkORVkimOVcvgncjaVhGxzN7FUSQE+RsfvnGdlFfHmfaB2aZxseR87qaEzuT4e/pw9NKD57iFjb0s/uMFyANQ6aq8wp6mRWGSJQQSIhUrPV93TAXPmsRqG/wDVH2xmlzhGA494YeYJB51V0AF5LcvbCnJQDSedHP8At0ObTLc6iqq1TGxvYVkUyonvZXQWxYFRRpS5fBKRAFnhCyS3GyOK+0qv7IvcMRzfwKOSTR+QZ2KDWFJqOJe7+qVsdbEB5LFIrqPhtW2Ogs8gFHmmZ5Ztp3jizNhSm9+5RZr7+P2l2KE8+XPA5mt4McrAMi0DgTX1CjshDmHxPP8ARM9fUT5pT0dbHLE7f2U07O0dtoeD6hHEb+ZCtEm7zUsrjo2KoyzOMx7O9d5dq+gp+/ahl76SnlXw1EJBWaGQekkbOjfE4tnMqcVuSaUMETxTS6RyyWInxK5iapjdPfhSCPvKXHW2Idq7IoZ8prJoEk2GnigiLc7XedQyE/tKoPHmDuHhbiftGiQxvxBwPngo7huB7cxiFdtBJlNEM4octnnlyTPMhhz/ACWeFysjJRS/WICrDpKtPLWREeUkLA+eKx7a80bI8q7O81EcdT9UrM7Mxgi7qOeD6xCrDYOE3AuCo4DE2Aw17BtdvS11FpXNZ3kbS2Z/pfL4iL/WKFm25lRgee6BpJkA80kA5fG77Y9NTZdk+ndOVki1Ay2LPYGmj5R40zJY0lHsyKjA+4xUhj+nnDZMaVHiCDQ8DxqNSlkk7WMlmFcfA1HvyWqzfTFBVL+jUIeKuWKSmqFN1YEXgmN+byGZCfcv6Y2nYNr85PqPLNFakiedstqJVyqKQ2+t0NUGjzDKueLTRSSSQ/s1CbRbvMRfQWYz5zpMUDSWzHS0n1Ka/JkoZWPcvbz7uUtGfRZY/TBNeafgzGtynMY2akgnq0ikkjku1KZXG7xefdyHcD5oB5g4dG2rnWeU+e/UfMU8jRDz3RI3rbw9lve3XJKyHUFNlsv+tNTaay+KWfoWEL1NOJbejJErG/mQOCRjSdjuaZnmdEnZyaySHMaR5M10/KG2yMpB+tUYJ6bwplUftxN5viT12d6l1nmUecavylaTNspy5MirKl+Iq6opZp0klYj/AGbOH2uCNrHcfvAYqjUbVek82yDPaGomo6hWmaCoUjdBLDUko6kcHadp44Iv64ljaZI/pjS9q14gHlhTwNFGTcf2ur2J6810fki5dJnEmQQLCabUUCZLTG/hNSrCbL5vnUxqgP7FUffGmr4Zc/lfISG7jM8qzCBC62lQS0UskY5+0N0S2vezK4B4Ixra3MYdV5FRag07T0+WjNVJVYeTR5nCwZ4QfuqkpjlQecUifsnE1iNFm+sNMZ3TEQU+c5tR5vTxTKWRFrjeeGOQH9W4mepQxm6N3fG1gb0GkiO+4YtJHh0QrZxddGRC460pntVpvPKTMaSselmhljmiqFFzBKpDJJbzseo81JHnjqnJ9QQ5pltcmTJ9Wo9S5ZNU0dNGdwo88y5WqDRA+QaNqlYz1aKpiHO045KziiOXZnUUn/BkaP8A8LFf4YsTsb11JleZLk1ZVJHFPPTzQyyt4aerha9NOT5BSzRP6xTN+yMdBb7OJKWhorTPeOusVl2WdzKwlB230cNP2xatmQXpqrMGq0cD7s6LMrW9CJL/ADxDZMskkpZa6GHvYqbb3zKQe7DGwJHXbfi/S/BtfFo/SCy1sr7SWSWlkpTU5LlFQIZB4kBoo12G/mNm3/lxV9HVVGVZglZQSjdHcWZd6spHiR1PDKRcEeYxasji6zMc3YPRVpWgTOa7aUUHaFYryoAJHJI9CPMfH8cPVq+6jENcpqaWQmRVDWPoXia3gYdCOnkw88LZjR0VbE+cZFEYoFG+qpN280l/NSeXhJ6MeV4VvJiwpqyODdBUQvPSSHdIimzK3/EQn7LD1PB6HjpNUPFQPLWmULDQ8dSNXU4y+ZKqhqDNATuhmsUbcPJxe6OPMA+4JGAqZ6OtAq6giOoJ8Um37R/vgdf8a/8AMPPC9XllRSRJPDKKiirbmCoAIWXbyVYdVkW/KnkdRcG+GAEcn6p7KHNi7kkJ+Av88K2jwHApHVabpCATMsqAtYRgi1r8H8iD+GCsFdQNwBvbgGxB9bdPl+GDSwmkKOVYI3KEMLH3Rhx/XTCib0H1hJd0cvh7wcgnzVx5H2PxGH1CZQpOR7RCCoAmRRZVJO5R6A9QPY8YnWWZqmoNPyZxNEn6RoZIqSukDlROhT/V5pACLNuQxs62udlwSxvBHLWCkBwfsljex9j/AAxttGZlltG2bZfmFatC2bU0dKtXIrNDGomSRw4UFhu7sAOASvpzcQTxgtvDMdFSxPxoVtIKOnd2dpYZbRszxq4V0BHBYMNqi5BIO4DzscO8veioS+pHooqePLoXpRDGndpPVSRsqLH1NwHaRmNygUc2ZcN4827P6Nlmps01FWTK8pAjpoYgpDGzBnYkb/teduhBONRqHO6vPKyGZqZKWhgUxUNNG5McadT4uCzk8ux5J8gAAIQ0yG6QQN+CeSGCutMUqHy4RS0U0kUkNtskTlJFPqGHTB5szOaSF83pw07cmphQJI3u68K/x4b3w2dCB9vbv6G9iPa/8fPGRLtVQJDcdebE+oP9HFwsDjXWoA4tFBklmoJI4mqadlngH25EBO32Zeq/MW98NCSwNrA222t932PTC+9oJBNBLIkg6MAUZfmOLflg/fU1RuNTEIZf+LCnDH+9H/FbfA4KubmigOSaS8RnkHjHQX0eqlqbT9UI3KN+i6eNmXqoqM7g3fMrTL8sc+1yvFESOVYcOp3K3wP8+fbFvdguYPFS5lESCpgysEMeCor26+guwxmaXF6z1Goj491dsGEtDrCu/VmrW0nP2LdqNSjL/ZbtHraauKtcNBUJl9S6hrDhlapFreZGLc1d2cVdRmefdnU1ZSEZlFVadpowrMTVwytLlsrMxIBFRBEnCgbahxzuOKi7WdJ1Opvo0ag/Rff1VRl2rMizCGlSIyTF6imraU2233klIvs9So8+MWV2g9oEhpafWdFPEKityLJtT03cjdN3xoKeo3t0CL30T8k3JFgMc5NK65Z5Wfy1B5fla8UdXysdrAK889RUMdLXz08ERSMPvjDdVRgGVfkGt8sE0jmcmVaioMyi295Qzx1S7hdSYmEgBB6jwdMWD9ILK6DLu2XVtLly7aM5tWPBYXAjaZpEA9tkqAfDFcZZSlsxJEZeySHavJPgPA9zcD4nHXtcJYLx1hYNLk1BtXT+p+zHN8nru1Gn0jl0cuY671bUaYyeNHWNYcsiYZjXys5IWOFEekV3NlVS9zYHFW6513pfR+nY+yrs5rJK3LYp1rM7zeFngbUFeFZVbyZKOJXdYYzZn3vK9i6qtg/SI7SdQaNb/srizGETUOVQ5bn3cID9azSV1q8wUubnuUmaGJkHEjUq7rhLY5sio2q3MzsWLtcseSST+8k4oWCJ00QfJlh54AA8Bx35STOEbu7mi1mdZjmAaJpO7hc3aKMbVJ9/NvnfCdNSOWW6Es3Ci3JPoB1J+GLn7Pfozas1XVMlZSz0zU4V56NDGk9MjDwtVSzFYKJSDcd83eMPsxNxjrrsb+iRkGXz08DNNU1FS6oIckken3g8APmEqfW5b+lPHTofIkc4ln0jZrG3Cnz7nyqnR2Wa0Ox6+PNcW6V7FtZ5xUUtPW0H6I+vMPqsNZE7VtV7U9DGrVMx/wAKBfVgOcdL5T9E3JuxnL6PWPbRQ12VNV009ZlmQSyxDUWbRQoXml2KWhyilVQe8ncyTKCFXbIyjHQ2t+1Ps9+i3mUvZp2GadyTO+1evjbv8syqNYqLKFVbvUZxXszTy92viaJ5rAC77RweNe2PtNzioybUGdZ1qys1HqDVRSizHOqm6vWwwOJJREht3FJ33cQxRAKNsUrEXJC1jbZpWggULsADgfEjH1x2BSMs0bTiagZn2r14qn+1DtEzPtM1m+ZVlJS5fQ5dSrlmVZXl8XdUmW0sd2WCCM3sqszEliWdizsS7E4iqIWDOJNuzlQ/3/8ADYG59fL1w2pklRSdod+pJN2uepA8/jzhzAxZniR9/G90JsFt9ony48yOcbUbBGwNCoPdfcXLAxs0rGIk3UXW1jx5Dwi3kD69DgLpJIe9YgSG7h7sSfUDru+Fh64x5QsneRKY1YbTaQsCPiR87c4xt/e7hHGouAVWQBR7A3Pxvzh6akKkI0Dse8bkqWbn4cjz9jhDIpBFnWVy+Xeqp/8AFb+OF5iWjkZnZrGwF91wPO444wwo32T0MnTZOP8A2wOI5ReYRuT4zdeDvXoJ2UVNPV9mPZ5DOsjGbSX1YyMC0bNT1+YQlbryjhLcg3K3tcqMKZ9WK+hdWZOm8D+zclbuWXw7qfNMsl8SdARsJ3IdrXbwqb30nZO4pOxvQtc7yqYpM+oSFJZJI48037GX9rdUAoy2IPqDY7irkep0vqhKuqkMr6G1LItPUON6haZZrxgeFkV4it06FbMobk+eNFLcT/5H1K6o/wD5fJcUatgkg1PmMS3tvNrD7qu6/wDuONGzKLcgEdFJP5C37+mJP2hwNDq2shIN0lnU8Xvaol/HriPO0i7k+wGIcBLqdwFuBe9reX8sd9ZTehady5icUkcEkJbMrKVO0WS46+3Pn7nnAElbN08XIuSR79eecYShACCxtfpcG/mMESR0axay3stuCF97cn+rYsKFHIj8SgJcG5G0i/vz1HX09cJOoJvcMAo63uPYkHm3H44cPsC7EdW5udp3bfn7+nPrxhGZk7v7RY9fEb8+3HA/H3wISLgs5KAG3pYWHpbGYwhgOjK1+bixHH5A3xmBCCTcrWuAvlfp8/6+GCNa5AsB155A9v8APB3FuNxB8tptx5g+2CMS32bkDpf0+WFQk7eOzbRa3mbAfvwYqHUi3mWtccHpf1OAtY2VV9LWsB8MGNrbObE+p6+uEQkr34624v8Axw3kIHHp5Hyw4lNrWP2r9ev44bSH/pgQrU7HS0eo2qFJHcZFXAH03U9v/c8WHXxjxiaMFIFYxlv92u4khf2QSxNh5lvW2K97INxzbMtrIuzJanxSGyr4IRcn54svOJIKekrJ46VphBFIxjC+JkKE/ha4a3RkuPtLjkraf7wfAdc1v2bGEea0X0hZVSjyeFI0RV0DouMKikAbqLvT1J5JJJ9SScUY4BkYE8COO4Pn4Bi6PpGTd5V0iKf9lpHRMPHQ2yZD/wC5YpqZQZn6JxGAAfRB643tHCkZ8R/pasm0nFvh7lFQWW7XuLj8Oh/A/lg3AeynkiwK8genXrY/1xjFUI43Fem8bRu6fA2/HABhvVlCEkf4jboQb8X88X1WRSABYqCSOF9/X44yQEkm+4CxLWuB+Pp0weYm5WIkKLAkEjd7+xP4XwntBfaC9ifM8+3thUiZ1C25t8/XEh0tmclhTLIUkSwUhiD1uhv5WY2v5bgfLGhlTeOnI4Pt/XTCdNUPSVCzKL2uCp+8DwR+GIZ4u1YW61LC/s3ArobszzbWFHqiDWfZDncmUa/iVomoFRTT6hhYfrKcxN+reZwDup3Gyex2Wk8LdJdlfa3pbtwaVMghfItaU0TfpHSjsW78J/tHoN43VEYtdqaQGaMAgd6q+HirLK/dSx5pTvuDFUm9Qx+y3HTdbr5OptYlcW2Mipu3GkXPsrzRco7TcueOWnzKSpFPHnkikd2s8twIK64AjqrhZmsspWUiR8OVkNqb2NowIwB2blqRmWzP7SHEaxtXXWcZT2edtOVQ6d1rSMMxhpxT5VntEUNdTqeAkMjnbPEehoqgtG/IieNvDjhP6RH0fdS9kWdJFLTRVFDWxyVNFWUSuaOugRrPLT7/ABoUPhlppP1sDcG6kNi+Oyvt0OtM4m0D2rU8Wmu0iCZqRqvMUFHS53UA7TBWq4C0leSLCcgRTHiUI5EpvHOsx0vqvJ6jsv7T8uzCojqJitQgj7uvy+siXatREX5irYlPAbiVLxSbkaNxVhtNp0TKIbTiw5FWZIYNIRmSHBwzC86uwrWiaC7Qsp1BWp3tBSz7q6DqKiidGiq4SPMPTSTC3sMOe2zRb9mmuMz0e9X9ZShqGWCoXlZ4SAY5QfMSRGGUH0lxve3vsaznsT1iqM0NXF+rraarpYitHmFHIT3NZBfpHJZkeI+KKVXjPlh12hIvaN2M6d17Ae9zHTHd6VzRurtFGjSZZO3n4qYTU5P7VGg88bzXgyCRuR668VkOYWsuHNVzoWqFPVzPYG0Tmxvzyl+mLi1PmNTTdgWuJofClXn+naA2cPdQmY1FrjqNyofkMUjoxttY4uLlXTr+0h/ioxd9FRpqLsrotPLZv072m5JRFR97/VJl5/8Ah354oWqjLWHEbPUKxGSbOQrE1bENMZvVUCq6vpfLqDK1MRs8UtJlkCkBhYo2/vLcgnm17EYqH6T2dDMe0euoe870ZatDk4a9j/qFBDTvz/6Uy/MHFl621flWfdpuc5vJK6UEmd1+Z1G5LB6KKpedvEPtKyRKpVuhII645p1VnFbqDUFTmeYFmqJmeepJNiZ5XaaX57pCPlinomEvmEjtQx68Qpra8NiuhamKGWWVKaCNneVxGsa/ack2A+Z+WN52haJrezzXGe6IrKlKmpyCtkoKiUJ3d5EsHAUm9g1x8r423ZPlFPm2r6errJvq+X5bIldXVLLf6tRwOstRLbzYIgRR5vIq+eLU+nD2fVGk/pDa3r3R3pc5zc57RyMtjJSZjGKmB/xMsZ94zjfE1Z+zGQHPCnuswxkRh51rniygg3DAcmx64xQxZhwzJcALx8wP5YM6+EgcX4v7f1+7DrLssqMyeSXfHT08JDVFTLcRwg9L25LHyUeI+Xriw5wYKuUTWlxoE2paWorp0pqOFpZZD4UQcnzJ9gOpPQetsPBUwZSSKFkqK4c/Wh4kh/8ARA9W/vn/AJR54GrzKCKB8uyVWjppOJ5ZNomqf8dj4U9EHHmSxxrxC7BnCMRH4mK+LaCbXNunJtz64jAMmL8Bs+fjjsTyQzBue34+eCMkjvuZzuLG5LN4ix6/HEj0blcLy1zZitTBDUUzUv1pYw6QB1ZizrfcQNgPF+A3tjRB1oV2Lbv2HiPXYPc+Z9hwPO54xuNKTMlPnOYT5WtZAMvkpzLKf9lLIyKrKT1cAtx6Enjrhloq+IhuCdDRrwSl83qadqLLElro2gGXFngo0s0DsW8Mpbhneyu9r2VlUdLYaUuaLkNBM6o8Oe96YAsqsslL9k96gNtrWuvPN2NuMKJSSPW/W6yWDuRSivaV5OGi3hOl77wQU2jxdbeRxrM9zqtz3No55qrvFgjFLTvIACkCX2Bj1JAPLNdvfphrG17urX8JXEjvIuY1OcZo8ZraqWrmUFo4mI2p5liBYD3/ADxraiSOwIkaZzzJIeBf2HU/E4fy1RggNDStxKLyt0Zx5A+gPW3pa/JOHOidMS6z1jk+loWCLmNWkUsl7COG95ZL+QVA7E+2JHPbFGZHYNAr5BNDTI4MGJKtXTmU0+R5NptMwphJTZHlzalr4XFhNPUupp4T/j2UiH+67+mNVHnslGlRq/OpxNJHUtLE8q7vrVdfe0hB4Kozd4w6FmjT7xtINUZ0dQV1Y0sn6Ny6tnXNq2UICaShRe6oKdV83ENikfnJIpPCkiodZapGo69UpqZaTLqNBBSUaPuEcakkLu+8blmd+rOzN5gDBs8L7VRjhSuLt1TUjxJJ8scMFqzSNhN8asBvpgD1r81qMxzKfNcwmzSsLyvKzOokbcSSSbsfPkkk+ZJPmcOtIrJJq7JtxJZ8ypRz1v3y4YxKzRySlL2sd3ko6dP3Y3XZtA1b2i6apwDZ84olt7d+mNq1ERWZ9Mg0+izoavmaTmSPVWblUE0OoNdZmnheCHMMvp2Y8LNW1Zpwbe0QqG/5cHmlTL6fNtVTwRww5OjVNGoYnfWTDuYYxfyurSn0ER9cLmWnpcsrMxmlEX6Yz6tqme1ysFOXRWI6kd5PKR7qcRDtNzoUOU0mlV8NQrGtrl/ZqZFASM+8UG1T/flk9MYETXTXYhrwPgPu9xuqtN5Ed55P66utyjul5HaiaYjeVqWjZrm43pdT78qeMSHLqdMxyrUanawjoaOduOoWujU/lIcRrQrrK1fl7SBC8SVKEoXBaJuV2jk3VmGJ/kVAlFLqKigVT3um64mMoQVaGSGYD8EYi3vi/aDdlczeD6FVoxWMO62LbdiFS9PmeV5RJ4ZtM6sy7MoSG8QpamVIJhf9kP3DH4nG61hM2nOzKJe/YSVkldW7PIJB/wCT6Yn3M09Ww9om9MVtludSaO1/S5zKzJRyTfV6rYP9yzWaw8ypAYe6jE/+kQwgzWl0hHURzGnqxliNF9iRaONnldT5q9ZV1TD/AA+2M58dbS0Oyd3uXeHv5q2HUhNMxh8fHkqIr4VpZY+7J7sHZf8Au22n9+L+XN1zTIsszSCWRautyujqWvGrRPIl6efdf72+A/EOeQRzR2bQCagQhQW2hhb0t5+4/HFj9nFa+YaDjkbc0uS5jJTMwk2GOKqTvI2vY+HvoJR8XHOLekG3omSawaeR+TRV7IbryzaK8PxVWbTTU76R0lLIqwrS5DURkm+6H6tm1YgmU/ejXvFSVeoV1kH2MQnWcrJk9c8ULxy1Ga5fT7ASDEYkqJHFhwbXW3ta3BGJpUSxf2E03KaowyUVdqGGKsuC1PIaqGVWa/DIRUFXB4IfnjnEO1RTQ9zpjI6mMU1TmmaZl3ccchZYmjp4oYdp81EgZVBvYEi5tinfHa3zvPDve3Wu1d/d3R1XD3Vc5nFXZRW0Gq8pmkpKylnEqTR2/UypZlPv5MPIqT74untC1pkuuafS2d5TAtPPFkRTM6GJmV6KvlqpHmRb8AECJ0BupRwp6DFfSUUVVlc1HUInd1CobgbiWCSFbW5+8bjrwL2GJTUUFMabK6+FI1lzrJMqr9jcCZ5KfuZAnS7d5Tt4bgkE7bkYnluyFryMW4eR1de5ULAW1Goqvcjz+m0XryPN6h3fKawNQ5iYkIP1aZbMwU8hlIDgeTRi2LGz7IqupZslmg70gNTTSIf1byuAXlQeaCJoglvtGQftG1X63glejzKolZmeN6YMHFnB3OLMPX1Pn14NxizOzvN4tRaGpq5pD+kcvSPJat7+Lu1jd6WQeYLxB4C3kIOOSMNtQuhs4/8AU+oPOm/BLAakxnxHv6eqm0u/Ossg1DCGjkznL6XNqyxA/wBaYtSVLg9FJqaRmJPhJnCtYEMKU7aMv+r5DkE6spjNbmiCylQCHhLDaeVO5jdTyDce+Ly0qsX9mY4+/mhFDX5hSU88R2tBFVQwVkauCCDGZIa0MrAi1zbjFSdvit/ZLT8sscSyPmub7u7QoCR9UF9pJ2k+YBK/s8WwWY/31tMia8WlLOK2Yk5/kKO9j2ro4Z59I5pWinos4MaCd2stLVrcU1QfQAs0Mh/4cpP3RjoTSMs1LlmXPPGqyaf1CmWV1NMviSCpnWpi+DRVUVaqn1a3Q44ugmemmWVQDbqD0YHqD7EcY6a7PNRyap0/XZlJVvLXVWXJRVNzd2rqN0q6GoI8zJFBUQsf21J6tixpSzBhLhk7A+Oo9bqKKxzlwAObfRUl2nUgpe0LU1Eq+GHNq2IfKpkHH4YicRennV1F7fgR6fhix+1OjjzLti1lBSSW77Pcx+qjyc/WHZV9twJA97euItSZUtcxp0W9QVLxIRYuwFzH/iIBsPUW88a0Dx2La7B6LPe09oabVJ9ca5m182Q5tmG+eupMipMprHe4aVqYuiPu827oxgkeY564i1VTPDsmSV2hc7openI6qfRl4uPgRwcK5fNSupy2qlVKeZg0crdKeXyc/wB09G9rHquF/rMuXvUZTmdPIkbN3dTFtBaN16Ov95fLyIJHQjCxMbA3smDAat34SPcZDfdr9UjSTVFM/wCksulaGen8UqJbwg8FwOhQ3syngX58J4JWNBMGnpqfuQfFJTKTsB8ynmB/dPTyJGCIs9DU95HIFliIZGQXVgRww8ipB6HyJBw6qqRJKU5vlybacOEniUkmlkPQevdtY7W+KnkC78GuFdfVPjqqYuHh1X56pmS5s+XpJA6LVZfV2+sUrsVWTb0ZT1SRequOR53UkEuZ5alMiV9BUGpy+V9sc20K8b9e7lUfZcfgw5W44DKCY0zG8ayI3DxsCAw9j5H0I5Hw4xsaaSbKoxmEIWWmrF2PSTqStRFfkOB0F+jDncLr0OBwLTeZn6/nooaQ4XXfosp6WiGWVVRV5oiMVAjo4D3klQx82H2Y1Hmx8XQAeiMqVcckdXHS01NDPGEjiRf1MyDyJP229STuB9MBLBSJuqqF5mo3NlDWMkLn/dv5X9G6MPe4wRK+ZIJKNWKxS27xNxKufUqeL+/B98AF7vJCad1ENKrpLLSd4OD3sN7sg9f76/3hyPP1w1aBHXcWHS4I559MOY0kRFqI3ZGWQqgUkOGA6i3pcfjh2ggzI903dU9ax43Wjjmb0PlG5+St/dOHXi3PJIBeyzWupqdUHSwBsD5g/D0OHUFS1LKRGqyxt4ZYpFukntt6/A8N6YTeGWnlemqlaKSIkNFIpDK3oB5fD54K7LK62jcMeviB+Y8/ex/HC4OFNSSpaU+MaVXeLlhtIQVallIMo9RGx4f4cN8euGBYKbkncPtKRaxHHx/djHCFABtYOLhtm0n0wstStSypmKu52gLUKt3AH7Q/3g/MeRPTCC83eOfXNLg7cUkBtI3GwIvuC8jy4v5422ldI1+r6vMcvytnM1DllbmaqI7mVaaFppFPPHgRjfk3Axr6ileIK3hMUnEcituRvYH19QbEeeOlvoTaPoq/tPoP0tIkdNm2S5vl7SSWCRSZjTSUFIZGPRSzSN8AD54jnnbEy9XPAJ8URe4jYuWZ1kSA+NrOOSDYMMWV2LV0NOlcHsd2WzGxNuaeohn/APalz8jiNa80zVaVzN8lrIHhkiQXR18UbAlJY29CkscqH3QYb9nebrlWapJMf1Mco74X6wSqYJv/AGEt/wDlxVtzfqLKS3x681Ysp7Ge6V1q+a1UnY/2j5bSNVQTxafp88pHjcoyy0GaUsm5WU3DKsk3iHxvzhxllR+lOwzQGd1EQqJa/TNTQOC0gVFo6qppwz2G0kqUCgG5tc8AYj3YlVNqLK6/RdTJukzXI88yQg9BJLl04BB8yZaeE29UB88JaUzyoi+jj2cwxvGGmiz2BAUDN4683A58P2uvvjmJWAwBmsPHofhbEby2epyLfhUr271Ak7Qq2RyHZ+7MnNufq1Nf92B+j5lVJLrg6xzmnE2VaRgkz+rjI8Mq0xVooj/6Spami9w7YjvaTnH6e1rmVZFufv6iQxhTe932oB8VRPxxM8//AP8AHPZDl+laSxzjWzLmdW6k7o8ugd46ZfX9bOaia3mscB54x0YBFkZBrcKfPJY7qOndIMgq+zB8819qp59k+Z5nmdUwURKXkqJ5HZnYDqSzs7fPngXx0t2O9juT6Fpcv1VqMtUZjVoamiNLULHLJGDt30shBWnh3XVq9lZ2IK0iNzMNJ2Ndn8Gg9P8A9s9T5Wk9VXd7R0lFMpIqSthLEwBB+rx3X6wQQZWK0qkAzkXRo/LqjNXzDX2ts6SjyigZazM80zFLoARsRmVBdnNgkMEa3NljjUKvhz7dpABvZQ5ZePyrdmsnevy4eytXQ2SwZlSU8U1HQ5fluXiSqp6eFBT0VFtF5pVEjEIQOZKid2kPJkkvZcVd2kfSo1RqLN63sq+inVLTmaKR861y8hgSCmUbZ5aaRxempgDZqtx3spIEaruVWgWude62+kvnM/ZF2Y040zoLLBFPnNVWvtRIFayVGaSR33uW/wBlRR3UMQqiSS7iS5xV6J7F9H0GQaOg7zL3Aq6qpkMclVmNWt9lVWOCVWRTfuqYXSmFz45CZBSijbZT9RaO/Kchs3nrw2qxI4zjsou6wZnb1+uxQyroso0Bpg6J0gjiXNWEma5pWoY6vOZgS/eVAa7Q0qENItOTuAUyzFpLKnN2t9Spn2dLBSFxRUwSOBXvu7pAdm7+8xZ5W/vSt6Yluv8AWlb3M8lTMRW5gmxobm8NO9m7s35DSeEsOojAB5kOK2pUJDTzMxklJJY2G4nr5c39sbmj4HvcZ5sSs21Sta0RR5JSylgC6AD7RBLqnoeOcCEEbJaNWUN7tb0J4sPlhQJURzBCCrKLMHTkLbxXHmLXuOuExsCKzb7sLodxW9vvev8APGuqCAM4I5J3ixO8AkHqOOQD6nBoyqBniZ0I5BUKSQTwDccW/A4KzPIu79YwPJBN9x9iRcH4E4NcqColsrC7BGJWQe3ne/kehwJUgzLta5Xm1+vPN/n8TjVglYY2/ZmP8MbdrrDtfcCTewb8bg+fuOv541D/APmp9pj+7AgbV3J2RwpXdl+k44Sz1VNnupUhhUtucM1A7bPK43XtcN5rdgBiToM0ly7P6RktTVWltRU7gxjljk9WwIP3OUuBYEq1jwgGI32Hr3vZBRKwqGE2f6gpyEH6kL3OVuWmYG4FwAosbk/dtfEopJ6yGWohnEZiqcmzxO9dgZHY5NmHmTeQcAFrdQDc7sebPJFoJ2E+pK64YwHwXHnayqrrvMY77LzykkAnrKx6fPEMjCi9owLE3JFrW/8Aan484m3bFH3WvKxXqIJmI3NJE4MbFlRjY/8APb44gzFpGbapY8dR0A8z/n+OO8sBrZ2HcuZtX8ZyPv5Edr7ibcdT7ehPmPPrhPaAwKhV3Gx3efxPl8sAXYAE7Qegtfk/H+BxlwwItweq+hHr/ni4q6PdlN2K+dueRx7j+vI4Ra5Ty4IsOOb/AMOPgMKX3bN4kdOgPPAHUC9wMAFUsUspNjcdLe/T+ItgQid2JLgA3U3JB59PwxmMKqxV5ALKLbQdov6/hjMCEnKfESq8DgHpf3P59PLBVAZdrKeTb+vM+WDbBtJUAWsLi/Hvz8MCVuGbaVsOQvIt5/8AX+GFQkmjsw3Gy+p/r8MZx964HnybfE+fxwcmzjjz444Hr54CQndyQLeW4cfDAhIzLwCOCfK97YaSX58hfDuQWUGx9PhhrJ1/PAhWt2Oukee5qW28ZFUNZjZeO4Nz8r4sjMO/NNVvSw1MpWmknPcR96/dBDuZQCAeCpvexAuORisuyCn7/P8ANAV3BtP1dh62jiOLGzAM9G9H3MbbO9dVjO0XKMOg4uLkji1nuAORjkbdQWg+AW9Zf4Sj/wBIlQK6lYkc6V0W1x0/+M0eKfrbxVcsdxcCPjyI7sYuD6RKSOuU1BBAl0XoqTkf/erZ/wC44qTOmMmb1MptaTum6WHijHp0/hje0fgwjrJqybTiR1rKaqw5CWcc2DA35HUj+XpgYnaORWRVLKCDu5VgR5j4X/oYBdo2q27x8WW24HyPxv8AiMCp2srBgNoBIW/ha/4fwxfVdCbnYqIQ5AHkQ34+vp088FOx7BlAu1mKkk/IHGMn2gepFyL9ebcD+GBvc33MGtba46cW4Pn+/wCOFSJKVRzZw58mA4PvfrhnOluefmMbDwqAQfLjcLj54aOrFbG1j5+mBKnuns8bK5jDMA9NN4XVgStj1vbnafO3PAI5UYnmVZtWabqRW0Uply+b9WXYhiu8f7OTqLkedtrryObgVdJGyk+EjGyyfPqjLj3LsWhIKEWDWUm5Ug8MpPO08X5FjzjOtljE3eb5jb+Vcs9ouUa7yK6pePQfbtkNPk2tK9svz+hgWDKtRrC008KKLCmrEF3q6YAWB5nhH2DKg7sSLQXaPmWmNTZd2U/SKFJQZvR0scWltbVUzT0U9JysEVZPHf6zQn7MVYl5Kf7L7owUXmHKs9koSlZl1TtAItH3h2k9Rtc8g+itZh5E4sPTnadlmd5ZJoftEyyXNtPzSvUCm70RVWX1D/aqqKZgRDKeN6kGKW36xb2kFBmDDDaBeZzCtuoXdrFg7kV132j6W/7Q9K1XZfqzLVpc+y5pZMqSskXdBXSoGNOZF8L0taoVRIpKGQ086G+/HGnZJXZZkusM07KtT1slNp3W1KcnerqE2mlZ5A9HVOvk9PVohf0HfjzxcmktVrkenKHQPaLq9s17Pbmg0p2hU0L/AFnS0zklaDMoATIlKSbtCxPdk97TOyggwn6UHZ/muQVMOqc5y6KmzOdlqa3uHWSnkqWUd9LBIvglpqpDHVxSL4TvqBwVICQQusncrVhyPpw9PJLJI20GpFHax69bVQGZZVnGi9XZhp3O6FqLM8vq5KKrp26w1EchV0+TKR8MX32AqmZppWKoa6UXaLS5mVb9mHLZZt3wHc4iHbaidoejtN9uFA3fZhKkendTsPtfpGCL/Vat/eppUW585aaY9Tja9iGo4I6WpIiu1DlWcZ0JB/u3iyqrhYH5zpbE2knfue1AUNmYbxjJyWv1DmZptFrXzsplzWFEI8zArK0l/wD0kwjT3EMnvipaiRrPKSzyOS7/AN5if4k42epdSVWZzU+XW2QUUcUKIPupGu1FPwBZj6tIxxqKVparMoIKdP1skqhFXpvJ8IH/ADEYt2CD6eIudmcVXtUnay3RkMFenZPpmCl01LRTxq0+dZXm2YVV/wD4EpKGqkjHweojMnuIY8dWfS10dR9sfYJoXtToVAzHTuRZJk2byAcvllfSxSUFW3tFVl4mPkJTjnjRqAZjqhqcXpss0TqOCm4/3cGWNTq3zLMx92OOpuxLOaTVXZn2f6W1BLE+Q6v0PRaTzdGA8NNMslKk3reGeKGQHy2nGM6Z0RZaT/M4g+BGHCg4LRDA8uhAybUeRx44rzWXLo6UPJnCyQLFI8ZhUjvZJFNmUX4UAixc8DyBOGdfmE9Zsh2JDTw3MVPHcJHfqeeWY+bG7H8sSztX09nGltZZhlee07U+aQzzUmYptttrqeRoKjj3dN//AKzEH2eIO/I6n3Hrjo4xfAe7P0WQ83SWDL1WLyQHNrixuP8ALCwqaiBGjpi0MZH2QbM3uff93l64wxIXjVL7yoDADgN1497dffCcjLay8Aiw6c/E4kzzUeSIxvCUubLyB7+2J5l5pazKtK5dU1dS9GYKnuqVisJFVuO7xAeJJW2qJDduqX4BxAyrRxbnA3HlV9vU+g9B5/DG4FZLUUGVkRQuYYnjbu1JdghP+0HsvmLCwF7WviKVt+nWoqSM3a9awt/lL5TQ1FdLm+XUlRllYpjrElBSTbfdeJ7lo5VYDZwSbWa43YiiLBBC/d37t5G2GQDc/PG63kotwOC3wxudS1EkeVZXlj1287qirlgalCSQyF9gvJ9qUFEBAPCEsPPnRo/16RY2MdPEoAvyVjQeZtyefTk3wkbQKuOv266wSvOIaEhsWNWkme6Em3Pic+dv4n+OLQ7IchnptL6j1u5jikqh/Z+ilkfYkMci766ct91Y6Yd2SPOoUC7EDFTyNJyDZj6/u+AxMtXa9kOnsp7PNPd1FlWSRt3s0Jua6rkYPLO7efisqL0VUB5Jvivb2SzBsEf8xxOwDHmaCmsEqWyOZGTK/Vl4nrNG15rSOvdspyYv9WR2dpHXa8kpG0yMPJtvhVf92nhHJY4iNBlhqSxkkEUKANLIRwg+HmT0A8z8zgcqy+Sqm2F1SNRulkf7Ma36n1N+AByTwMPaypBl+p08TwwU5O1GtuZuhd7fePtwBwPMmeGFsDeyj8yo5ZTK6+/gmuYTiULS0sZSBT4I+rO3S7W6t5ccDoPeUdkVJ3fajkCsQTQVLV0xB6dxG8zD5CPn3xomibLKFM1m/wDO6sH6kh6pHyDOR8bqnuGbyGJz9G/IKvNtaZlUQ0bVDUWT1WyMGxklnAgRL+rGQi/xOKek3htil2XSPM4epVixNJtMe2oPkMVvqNVpqOLMc0phJQaYpKdZUf7FTXSBpkpz7NK8kj+kcT+ovTuZzVufZpNWySPPNUSs5dvtSMxJZj7kkn2v7YnvaprCCWoGjcjqo5svoZ5nnqY+ldWSsPrFT/hYqqJ6QxoPvNeHThcsy8MTaprY/wBX5bIT1f4tyB/dBPmMJYYy1vauFC7IbtvnmeGpFpfU9mDgM+uSa6XrYss1JRyvJ+q7zuZGvxscFGP4G/yxb+g2bM9ZUOX1oUSZpFX5NsH3pJKOaJePeQJ+OKNqEERRbWYjcfa/T8sWfp3NWps009rKIKHgqqOsd1HiWaCZFlF/cBWt/fw3SEQLg8axT1p78E+yPN0t2Gvz7cUvm2XRZ1lsdQInZDFHI283aJmUFQ9vusTtDjgNa9rkY1Hak/6A1rQZCJ2kXTENNlzMTyZYxvqCfczSSj5YsbJ9Px0vaHFlFSrGlyrN6iKoYNYLTU00kklx6GGE+xuPPFI6vzabP9QVmbVbXmq5XqJTb/eSMZH/APZOcQ2MdrNQfaAeeA9DxT7QbkddZI5Ld1tN9XE8TBR3ZZWAPSxIU/Ei35Y3HZBVsM8rNLFrLqCgnpogTYGqhIngHxLRlB/6U4Z1ciVdLR1pcBKuljnew++AUc+5ulh6Xxo4aisyV4M+y87KvK8yimia17Pbcvy3R4Us+ogdGczydq4OSVEMofqHp+Qr8p6inXs9hqZp1KwaoqB9YCgpFHVZZTOCy/sXgIYehN8Vr2v18mS53pKOjHdTZRk1FWNErEiKeoker236myPGPgBixa0x5hkWe/ohGMGc5tlGcZaHAYFK6nqIFQr0IWUmIjp4cVV27VEc2vs5mpj/AKtHmk9FSgdBBSBKaMD2AjIGKtiAlnYSMCOdBhwJUtoJbGccQff5opzmkFOtRJugH1V+7qIQnJnilVZolUelm/8AYksbC2NpFQ97pXSlVJGYY2oM0yqZv9rDE1NmLOqyjm8WyrUblBKeFugOIxkNec60dktb4Hlo4nyaZm5Y90e8hAubAGKSx8z3VsTrJ67vdCxqWjpv0VqarBk+6iVmXwSAMRyqmSkJ3D7J56XxGz92xzDmPbDnnwUru84Hb74qt9fIq6RzdpyTUrW0EDlm3N0mNi33rW4bm4sbkWONF2MaiTLtQ/oWsqVho87jGWzO58MTs4anmPsk4S/9x3HniTdpMHd6KzCZoyjtnVLHtNrLtpJ2sLcWuwPHHNxwRimKeXuJg7A7ejAeh640bLCLTZXsOv4HoVUmkME7XDV8ldU6MzySnpdT5NUUZtDSU2ZSU81wUekrDBUJcc3ENdMtx6X8sQrt+oKai0lpijpNwRs0zqUISDsUyUq2uODyrcjj4dMSfRNZNqjMcsz1HDVGqMszfJa1rcNmAy9w5PvIVpJx/ekY+WIV2y5olbp/RJdEX61QVuZeHoVqK+babfCIYo2WptUZ14g+NHfBVieggcOtX4VOyU5W6kW8vniadj2sRpjWWWJW1Rgy+praeGrY9EiMybnt6rbcPgw88aCpplbLzWRjcIXEM1vLcCUb4GxHxX3xr46SSVDUwnxQsu+3kCbBvx4/DHQTMbPGY3ZHBZUbnRODgp12pNSJ2naw+pTI8Iz/ADARSI9wU+sPsKsPaxB+GNVmRkrKOPU9CWWVZlhrghsYanqkot0EgBYejqw8xhaip/7SZdWU5O/OKISV0TBfFVw/amj46unMi/3d48hhjkGaU2V14GYRvPltahp6+NOrwseWX++pAdT+0o9cRMaWRhoxcwAeI/Orf5qQ0LyT9rtez9Ne7yS2b0i5nRDU1FEEDuIcwgjXww1DXs4H7ElmI9GDr+zgYUfPstbunZsyyyHcV6tU0iDkj1eIfMx8/cweeCq0pnFXlksi1cMkfdyENaOspXAZHU+QYbXU/dYD0OGUbVOV5jBX5dWSh4nWemqAACeTY29eoK+oIPGHgXmi6d4PseuYTSaON4biPfr0TcyAbWO90Atw3Njzx8+fx9cLZbmM2X1S1MaxuroUkhcXjmib7SN6qfxBAIsQMPc7y+CohTUWU00cVJUyd3PAi3FHUkXMY/uMAWjPpdeqHGqUnbtZL7SSBa1/VfnyR739cStLZWYjxUbg6J2Hkthm+WU9PGmZ5aJJcuqHKxmRrvBJa5hkt94DkHo6+IeYDKKpmeXe0l2kIU94eCOlj5bfwt5Ye0dXJl8lSUjSty6XZFVKCwSVGJK8nlWBBKta6sD5EgoV9D9SmhqKWUT0VSxME7eEG3VHH3XXi4+YuCMIx103XeXW1K5t4Xmjr4SgZaFvrFNNG8h/Vyxou+FlP3WLEbwfYe4NxglRSRzQtX5af1CECaFiHemJPr95CejfJrHktJANrMkyuVXxFCePQ8gXAOD0/exWlpyENiLAX4I5UjzUjy8x8MOLTW83NNBwoUO4CRZFGxgR9glT8ASTz8fhgsqxgt4ncW6su0397k/xwpKokj76FbKgu6XuUHS/PVfQnp0PrhuW2Dcm0FRccgW+WHAgpDgtjBXQVNOtDnG5hENkNSgLS04H3T+2n93qv3SPslvPS1FFOhLA94u6KWM70kXpdT5j1HUeYBwl04Zeo3epN/Mc2w4pakQKYHUVNNId0kLErZvVT1R/cfO44wy6WGreHwnVDsHcUyYgEqSF8RsEHC+vHX5YN0Ygi7Dgheh9LH0thepiSEGWnd3gJtu+yy/3XA6H8j5emG4F1Ue3HGHggjBMIIOKf5BQV+fZ1S5JQyiI5g+yViP1aRAEvIwPkiBmv1G3HYGiNOxaV0ZU1c1O9MkUOS5hUAna6NWZrQQUsJPUGOiBJHk9S/pihOwDIMvzTPqipzZgKWoP1KdyeEoo0NTXNccgGCIQ3/7/ABfnanqHNf8AsN1NqauCpV57qLT8bvHIGXeZ6ipZVI8l7qMAGxAVRbjGDbZTJbGQjIEfPWwha1njDLK+Q50663qN/To0en9r67W1PCiTVub5tBmaxJtRMxgqylTYfdWaNqWrUftSVFvPHI1LOaGtEkguhusi/tIQQw/AnHoJ9Iuni1XF2m5c9MpemzKpz+EbgWvTTNFU8eW6mqpCSf8AhL6Y8/MxRopnil4ljYo4t0YGx/MHE+iJ/qbPddqwUNui7GW81X/9HnUD5d2o6LmqpSYKzUNBSTyA8CQzxxuT7PFIG+Z9MbTNBLo3s30dkj335LS6jO4njvEzZ4Yx/wDDFDf8hxTPZ9qh8oqqaqL2+p1dJOD5pJDMskMnysyH2YYsL6RupZ6PUI0h3lhRTV8pRev+s5jV1PP/ACzLYe4xnzQH6htnpiXV8gCCf8wVtjxc7bUBzqPhQ3s10ZJr7tAocoepNJTSSNPV1hHho6SJDJNO3oI4I5JPdlUeYxYuXwUHafr7M9bVCzZblNO4+pRxAF6DLaYJFEEB4MiJ3EES+c8hY8IcRqStfs07HahgFi1D2ilqRGb7VPk0Mo79h6d/URrDfzSllHRjicaByUUummmzrM1yXTuQxU9fqbOJYt/cSurGko4o7jv6kRtIyQg8yzySOVSMti3bS547hzwHvzoP1UFjo0ku1YqxacZbm1PmGp9XVaZDpnIKWGKeeJDLFRQi4pqGmQkGWUjcIo73lYyzOQpd8R2GbWP0kZIqxWqNFdmeR1rQwVDMah++2WMVMnh+vZm6cyS8RwpYXhjHif1ulV122V552t0tXpfQ2Up9Z01oGOrK19XHIAfrmYzgXgM4AZ5mXv5VskEaRhWXWdo3bur9zleSRUVPDQ0v1GipaGEQ09HTA37mCIErFFfxEXJZvHI7vyM9rGWU3WC9Jyb87+CtEutHed3WcypnqbVul+zTTUWldOZNQZfp6hJemoKaq78yTkWNVU1KgGeqZbhmkVFUeCIKguaA1nrozSNqHNkiWd4w2XUZQEAH7NRMD9pRbwIeZCAT4F5iOotX98zVuYVAqapzuRWPeIPTr/tPifAP75xC62vrc5qmqauRmBYtyxYknqSTySfX9w4xasejnSP7WXM5n4UE9rDG3GeQRpqqqzitesqZJJJJGZ7uSzMSblmPmxPJOHZKqg3OQV67xwfgenyNsIwxCIofDHuF1Msd1a3lzxhaBUcspmK8eEiMndyONo6C1z8rcXGOjaA0UCyCampSjd13OwKL3HKtx8rcfM4LIVaIN3YBW3i+1f2IJt+HT88CVPiBjkYbSFLsONvUkeluoHT1uDhPvuGjuVDeEllDXH+I+X4fHCpEVl2kP3PdKbgAkndbrYnzHywt30D2VQv2juLcgrYCxHn6k/hhOQM22xcGQ7rsOCRxxf7Xx9OMEsZLyBizbgCOAlv7tvL92BKhdkIYKrgA9Ge6gD4+LGvZb5Wznr9ZP/tMbKULErizncoZRt2lT5ix6i/n8D7YZSW/Qitddz1j8X5sIx/PDHmlKJzRWq7Z7Dlqk7JMvBnSOnOo8/O0ozSSuKfK9oiCi4k3chjZQL7uMSbLGovrL11PChYZXm8jOYWjIUZPmV1ZW5VlO4FTwp8PPBxpex6mnXshyB3kiio3z/UgleQFSHP6ORW7z7CqNpVlkIBuCDcDG3mpK+jp9RS1IqIUodNahqSC9rv+hapPGPTay7BzwWYkk8edPAdO4bz7rrGktgJ8Vx92sRd3rKZGO0Ku7jqfBH8vLEPd1Vi1gENzYkkgfHoR8vhiadsD79dVKbQO7GyxHpYeWIS46FTc/G9/gev447rR+NmZXYuZtX8ZyBrqxBPIG0gEHnzB8j/VsGBUWufDcWHkPh5f17YIqAWAsOAOeePXjpb35wqqXAIsGvY3API6ceXv6/DF1V0B2izcbgfEp4B9OfX+WCPtFurbrgXvx+Pn+WFpnDlQVA4sovwov0Hnb92E3XxfrFNyOQebeg/mPf2wISSPIGbYhcC21T5D1v19rD3xmCq8gkA8/UAi3+WMwqEDkBgfJehAtzjPCSQ5KgG/A5B9bYIftXF9t/Lz/ofuwYmwAUAAjm/AJ+f7xbCIQcM4N7k+QBP4YK5uAH44svHlf88CxNrXvc3tfof6/HCUm0s1hZTz9m1vfjCoWEl1vcHnj1+NsNZeT5/jfCzEkcjr5AW/DCL9LXvb+uuBCsrsjqauHUkqUjBZJ8mrY1YruA/1Rmvbz/2eLbiSiq6qBGppJSShaGmmCF7/AHY3N7P5jcD1AN/OpOyI21hlCE8VkEtNweu+mqEt+7Fk0U6yd042w+CIgrfYhKrbkfZ8VuR0utx6cnpFtbRhsHqVu2U0ix2/CZfSDEM2QaaqotxWXQOlyu77VojU0/Pv+rtimM8g7vMdqSrIr01LJuW4BBhU+fpfFx9tRM2itFTSElpNFLAzHzamz+ujN/cAgYp3PC27L6ncx3ZbTG4upAClOtzf7Nr9PYY2rEaCnWQ+FmWga+sytcFVgE8IVuLngfP2waF3Vi0bBSFIsTww8197+mCodkdrC+6wa/l6EfxwJIU/YNgOQRe3vxjRVVKOFudjuyA8Ejm388AXW20qxIPHiBNvTn8cCrBLlJLh1IIHW1+Qw/MH2wVpHUWNxwQb9CPP+eFSIrWDAMhF+p9R6j3wm6FWKlWDXA+Xv+WF7Bv1UbttcBgGsLP5j+R/HDWolZL7QVYMCb+Vv88CEaSIC4kBBXqCOR8jhrJFYkgiw88dA6Gbsc7UtN0WS6qkpNK6igX6slbW1Dx5VXvbjdUBXky6c/tMslK/VlhILY0/aN9GfX/Z+e/zHTme0NJKvewT1+Wu9NNGejxV1N3tLMh8nDgHEXbNAxUnZurQKloaiopX3QyMhIsfQj0I8xja0efKAEqYytuhQblH/KSCP+Uj4Ycf2M1HPCtTTZPPVRN0alZZ/wAkJI+YxqK7LKygbZV0s9O/7M0TIf8A2QGGOEM+FQTuTwZIsxxU109r/OdP1b1eR5kn66E01RA5WSGqgPWGaGUbZYz+y1xfkWNiLayLtc0fqXQp7I+0aLM8q0rVszZbUpA9W2kqxiT31K1y8uXyMT31GxLJcyRMzAhuaFiduQOMKx1VXSm0E8kf+FiAcM+mDft668E/tw7F2e7r3V56RyXMOyvVdZ2T9q3dQaS17Rig/StPKJ6GSJ3DUebUs6+GaOGcJJuXnY06MFLFcMuzXRub6a1B2oaQ1J9by/M8g0rm8UkEE20Gojnp45Y2I+3Gyb7joy2N8RzRfa/PlWTPonV2V0+odKVUpmmyisYrHHMwsaimkALUdRb/AHsfDWAkSReMXX2ein7YMwbO8jaqfU1JkOY6Tz2nls1RXUD0MqZZmR28MymOCmnYcXWnf77EV7ZXsnAimvh85cE+LBwINevZcx1UQFbUuW2s08hBP+I42GgYEn1rlffAbUq4nP8Aytu/9xxrppnqC0lrNIS5HxN7fHnDvQ8pi1XQSNcAVCA/M2/9yxatIIs7gNnsoLPjM0naukuzkK2TaqkkBDy6C1P0a6ljRI/APKng38j14xbnYbUnMOw/s/o4KuNKpsvzrLo4pV4Y0+YSFCjdQ4EymwPIB9MVJ2QrLmFRWZPtWQZjpLUVEvNnErZPUMgt95T3RseoIIPkcSHsTzlm7BdHVo3KMu1Dncby2ukZYUM43n7oO88nj1sOccxaml9heNhFOK2oSG2tp2gqPfT0paeftYq9WQwCNNV0mT6qUR2tvrqFUqLf/PFK5PuTjmHcCOlrHpe/Puf6GOofpdb8x05o7M5EVJYMprsvGw3QxQZwWhsbngLVEDk45eYCNe9IJ4JcDyHpfpjodHTdvAH9ZCvNZVrZ2cl3rNAi7jtNrehNvlzhVkWBC0qXlIBRD90H7zfwHn1PHXfa10Zm2g9SVWn83p5UnpxGyGWMLuR40kVrAkE7XXoSOQeQRiPlmJa7kNc3PJJ+fQ/PFwEPFRkqxBacc0kxsDuYknrfkk/HG1yHPJ8mpZqaMLtqTe6Rr3gUizAuedpH3L2LAE8Cx1LAsbAAcdPL4/D2w5hpSaTvWJV5ZNkV/vW5dj7AWHxPtgeGkUcht6tQpbT12nM0yrN6LPMxjajieCshlj7v6+JDeM90JLFxY/rI7i9lcE7eYjVLQ/W548taaSkEjCnNQqrKUv4S4UkA262JGAkrJpqZMvia9LFK0sY2DcXYBS17bjfaLC9hhKW6AqLD9pgOPgD6e/mfYYYxhYTvTnODgkKi7r3a32jkn9o+vw9MZSUjMwHAJ53HgBfUn0GFAt1LEkEc3PQD44cXMCmOSJRezbSbkex+NwSPhiQ7AmDelpqhEiSCnUrGPEtx4mJFu8YepHQeQ9ycEoqenu1RV80kFi4U2MrHpGp9T5nyFz6YRQNO4AazEkszdAOpJ+H8MDPKH2rEp7qK4QMeTfqx9zx8BYeWEIwuhOB/mKHMq2eumkr6plaSS1gosqgCyqo8lAAAHkBibaM1q3Z72b57Nlzlc71XKMvhlBsYKKNT3zqfJnaTYD5APbm1oK0DywS1JNo4QCxPqfsr7k2PyBOC01NLIsYqHdhGCEQ3O0E3sB7k9PfEFoszLSwRO+0EEjwxA40UkMzoXGQZkHHxwPKqHLKeOSq+tVytJTQWkmA4Li/CD3Y8fC/phxUSzZzmclXWi5cmacgWCxr5AeQAsoHwGHGaslFCmSx2vCxkqmX709rbL+iDw/4ixwnJC1FlUMKqfrFePrMv9yAH9WD6bjdvgFw6t7vDXgPDb78ElKYHVn49e609Uss07yutndiStunt/DEq0bVvU5Tm2nGG6R4Zaqlt171IzvQf4o7n4xriPoqxwy1TH/ZgKnPWRun4C5/DBtO5nNk2a01fTuFkhlWSMnpuU3F/byPsThlqj7WItbmMvEYp9nfckBdrz810RmVVQ79c6qjjePvMqijgcyXu2ZrCpPsRCKp+Mc11kjTzvUHgzMZCPS5vi5+0OoTKezqmkogyU+qamnmpd1+aampmjUA/3GqHQ+8eKaqU2x95z4uB8PP+X44o6IjAjLxroB5Ch51Vi3uq4NPjx/ClOnpXzDTiQXXfQVWwkmwEUviB+AkQ8/38OKfLzV0WoIXThaFK1CB9ruZ0DWHl4ZWI9hjXdnk6HOGyWodVjzaI0YLdFkYgxMfhKsfyJxNtJZbJmGpqTKUhK/peKsyllY895NTyKoPpaQJ+AxHaHdjI9v8A9D19QU+NvaMa7y9vQhT36P1dBneRSZdmLIz6InhzWYNb/wCNkE5rN/uEmRk9vrC++KM7QJah6nLlqpA870EVXMR5yVJadr+/6wYk+mq3OtNR5jVZQqhs9yOoyipV7+OnnVd22x+0CE9h5jEe7WFVNd5rRRABKKYUagfswRpF/wC4HDbFEW2xxrgakbjkeJJKW0vrANuA+PQLYdmmau9Dm2QE81MC1kAAue/pbvYD+9C0y/IYtPSZlj07qeni3zJBJk2YmIcMYhLPTOR6HbUpb5YoPSeb1GRZ9R5lSEd9TzJNGD0LqbhT7NypHo2L6yCSkoazUWXULy/Va/TVRV0jX5kijenroiPMkLCwv6qcNt0dycgZOx4UryA5p1nffiB/pw64lRntHWCn7MqbuZkkSrz2cxMo23SGkiXp903n+z5eXFsUzLEVlIt0NrYuPtblH9itMU9w3fz5rV7gLXD1aQr+K0uKpzCIGOCZTczQq/H7QurD8V/PF7RZPYVOsnkVVto/eYbArc+jdqaKDUdHk1V4mTNKDMaUHynp5bOo/wDSU0lQp9Si+2Nb2wx09Nn+Q5AUKR5HkGX0Eo6kXRpHP4zE/LFbZFm1XkOeUma0Epjmp5o54mB6MrBl/MW/HE27Wc3ps97RdQagy+KWPK8wzOoFGkhDMkMe1BGbceFQo+FjiMWa5b74ycCfMUHofVOdMXWfeCB5KMZVIuX5vLleanbSVQahq2t9gEjbJ/ysFcew98ZRKdO5zJSZxTMwgkelrYR1aO+1wD625B9QpwvmMUdXTUWa7QxmQ0s56/rogBc/4oyh/HDjOlbNcspc9JYz04SgrW67iq/qZD7tGCpP7Ufvi7ni7I4Hx69lDlgNWI8OvdFcV+m88Wagq9tTQTJNS1CgWccNHJbzDKQbehIwrqehpBPBneVQd1l2bBp4YkbimlBtNT/8jHw/3GQ4Ki/pXTn1kMTU5HtjlA6tRu1kf/kkO3/DIvpguVViy00+RVcgWGrcS07seIqtRZSfRXUlG+Kn7uCpwfrGB8OseIRQfZqOI65c06pQM/0/Jlkit+kcjiaekuPFNQ3LSxe5jJMij9lpB5DGiEzopTc6x8tfg7SRYn2B4v8AI+WFaSurMpr6fMKD9RVUcgkjNuQ6nkMPMdQR5i4w8zylo45480ymPZl2YKZ6eI89yb2kgPrsbgeqlD54c0dm+7qOI8dY9+Ka49oy9rGfhqPtwQZJmf6NqXSenapoqqPuaymB297Fe/hPk6kBlbyYDyJwGa5YcrnMP1tJoJ41mpKgA2nhY+FwADY8FSPusGHlhikfeU0uyddsRBaImxKHgEHq1jYEehB9bbnKpYsxy8acr5kjR5DJQVDuNtLO3BVj5RSWAb9lgr+TXV3cPaDz+fL08Akb3xcPl8da/Na5KmJKGemDPukdJSVFhZA3r/ivx5X9MHy6tp6eRqeuikqKKpsKiJCAwt9mSM9BIvl5HlTwcM6mnqaGeSnqo3gqIJDG8bLZo5FNiCPY8YQ3FrhU2+K4sxuBb7Pw/O2HljXjcUwOc0+C2OcZRLl86AzJUU00YlpamMeCeIm24XPBB4ZTyrAg+WGqbALEtY+E3Xj525HxHHww+yrMYRStkuZBzQzP3quF3NSy2t3qDzFgA6/eUftAHDarpaihmekqB+tTkMhDI6nlXU+akcg+nzwjXEG67P161pXAUvNyQRvNFIskLt3sbMAyndtI4PxHPSxuODhaqjiqITPRxCMou6emU/Yt1ePz2eo6p7ryGaqqSOoSwuPC5Cnp9oHpf1H/AFwpHJNC6uiPHIrEowazKV9AP8/Pyw5zcajNIDqOSbbrJ4XsFO4X6A/54PHIL7dm7fwo3CwPx/ryw6FP9aLVNEuydFLSwR8blHV4x5jzK+XUcdG+5XIZeebjgDcPlhWuvJpFEIV0G4NZhdSefwIPl8cIyuO6YxAi3DR36e49R+Y/PEs0xo+fU2WalzkTmKn01k0mbzvs3BwJoYUj9tzzrz7YieZ7THuBvb7PoR7YSocTTMJcQMVePYDTkZVmMzuEH1KChL2JKvXVQeQm3/yPSEG3kcWf9IOnlpexLIqayd5nHaAkZ2PuMgpsvjW59bNWdcQHsZliodNzxgLuOaREAg+LusvWw4IPWVuhvzxc8Gyu0yRdQ03YJpyEEjO9d5nMyllcE/XsuguCoAYWRrEhTwbgG+OWa8u0hf1Cp5V+VuOAbZLus0CsDXETZj2o6oyhHR6fMKzPssmjAS4E1LVxm/hD9VBHiI4tjz81LCsmb1M6m6z93UG/W8kaufzY47ardVxVPaJNm8ssu1s5zKvDd4CndhaxmPIBUbbcXI5vjiSvkL1hHUrBTI3sVgQG/wAxiX9nwWlwOwKPStKBaan75JJI4iRujcEeotf+GL4zjQs3aP8ASYz/ACyvz2ohy/LnavzjOKyzmhoaWBGqKmTaACUVSFW3icxpyWxVegMlTUeusnyQuqLmNfT0bMx4AllVGPwCsx+WOjNY6j0DoGTVGUaty2vrcx1nmEmotRUMEppm+pd8XyvJ6mYeOBGUrVzhAZW/1eMbGVmTTtUlLQA0Y3SOJHxxoqEQ/d13+n6quM1y6fte1LmHapmsUuluzvJZoMso5ZY+9MFNCoWly6lQkCpqjGoZlB2hmeSQqpJNi6u7bdMado8oo9AZDHVy5QrT5bNXwhqLKKiTxS1EMdQF+uV7mxkr6hLAqEghRER8UNrvte1VritpWlmioqHK4TTZZQ0UIgpcugJv3VNCLiFSeSRd3Pid3bnENmNZWSGaeSSVzyWkYsT8ziT6Z73A5ADDr8pBK1oxxKsPP+0fOc+nmqdR6wqKqSaRppbzyVUkrsbs7NwGY+bMxJ8ycRKu1IsimDLKVlF/9pMQzH4KBtH5n3xqo6VQf1zhSelz1xvco0hn+bkNlWn8zrUva9PRSSD/AMQW354cyxwwYu+B14pXWmSXBvXW5aOOCWolMkztI7G5JNyTh9HEisBu4HS3n626Yl9N2VaskqY6Stgy7L5JbhY62viSVj7QoWlY+wQn2xeegPoVa0znJBqrU9BXZdkSyRxPm2dD+z+VqWNlHfVa/Wp2J4CQUxZjYA84mFpi/lNfDH8KHsZNYp4rmIbeGudpsLSHr6eZ59LYwrtswKXv0Jt79OmLr7YdQdmGicrk7NuyeKDNpandDm+pHoO5SpgDA9xRQyFpIoTIlzNKxnlCD/ZxnY1NWDsFSUPuB4PhsPTnp/O2HxSdq28BgmyR9mbpSbgxyhGuClgQeqkfuPPPz64GIlZe83gd2b3TbYm/kTwQfw9sH2qkfeOWU/7MAqpBe3mDyvHJI6YSWERqHnR7g7OR9r4ny9z5joQeMSpiU3hSF2CNSSxHeOok+NuPmAMJ3XYSzgk+JtwIYe1+h/fgSjyMVWC6NdX2uVDHyuTew9AfLnGLZlUiUuzk3IN/bm3IJ+YI5wiElOCikOQSLA2IPUce38sEeBjlOVoV4nqZ2+PMa/zwo+2OJnFvslQPI+Vun5jC2wvFpml679z/APiqGH7lxBO67d8fYqaFt6o6zC7P7K6mT/sz07l6Rzzu5z3MY4hKEjdnzIoicn/aE0fhJBUeYvbD6rzGll0VngipJkp00BqiqRw22Cmd6FYhEqXvdmmPFgq7ABySMM+z+ijpuzPQySpIBV5GJ5HkivEHnzGvlRN/QMyc7Wtcci5AGFtTR1OV6K7QYTH4G0RmJVT1Bmq6CAN1v4u8HkLkX6knHBtb/fCP/L3K6Z7qWau5cs9sQUdoeZspFhLJY+XEjD+GIW6klbHhiRdha5t14v8AzOJR2k1TVOs62bcWYySk2Nus0nTEYIBJCNxbrbb8+tv547ewgts7Adi5u04yuoiDjliePW3H+Ifx5wsoKJ3nIJ4Vitgp8rEcEf8AXywiV5DMCCw5upBHz88CgVTYBbEEcjp7joOvyxcUCE37tQ1xc8DyJHqP4/jhFiBfY52m+0EXJ9bjphUv907bgjnofn6/v+OMlJMQ5OwXBubgn29PgfxwUQm9ruTvPi5+f9fwxmBZgFCrYHqwve/pjMCEkwu/htYi17AWwPAUPayj04sfQ++DSKxYspNhybfdBPQ+mCkWttJPmOORhUIlrvYXtf8Arn+hgvUFvO/UdMKMLWVgQ3IseLfD4+mCN18wL2Fxa/8AXvhEJNxbj8jxzhs3HX93TDggcrtI9jxhu17c9RhUKbdmtUKfVOm53J2R18Ktb0aXZ/8AVMWjR1VRltOl6cKI6bujxu2eEAOP7wswHlc+dsUnp2uegaCtQ+Klm71fihWQfmmL4miakzgPDQRt9YrZ4fqrm4kjaRvCeRfgqR/e2+pxzWkQGzY668j+VtWTGPDV17Iva5RqeyjR8plSV6UaqyslfSLMqerX8qy/zxSOaNvyfIKloAi/VZqYsrhi5jmY3K+XDgW+eL61dQT1/Y8wmcyT0GqZowWABK1+ThgxA6bpMtPpzfFAu6z6VoyWIalzCeM+E8LIiMDfp1VuL4vWFwLWk+PEFU7S0guHj6hMdhChgjDm249CPIfH+GBRypVoiVZSSrKbW+B8/ngUjG4xl+7Nxyx4uOhPp8f4YOyyd88b3RiSrBuOepBtjWVFFdY1WwcWChvgPcev9eeAZDvCi/HtYkEdT63H5YOrFo+6su5jwrED5i/IPztgN5UbSQe74B4PHp7jAhJll5Ug/MdD6+3wOG9QpKfw225+I64cum92VHvdeFAPX0BOE5PFdm4uOg8vTCoRMpzJsuqLtuMMllkUel+CPcHkf54vvsf+kZ2qdkFM+T6ZzqprtIZjKDVZBLmVRTU0cp5L080LrJSORysiMAejqwBGOepYypvY43ulc7Wkm+oVi95TzDYVLWuCb2BP2SDyp8j14Y4p2qNwHaMGOsbfyrMDwe47yXeWne1DQXavMcrTVOUx6ilC79M9qOS5XPVMx6ClzV4Y1qQeNveSQyN/e64Q1bo7I6KsOSan7DstoqxEJ7nJ6/MsiqLftiGSSqppB6ER7D645Wpc2oIqSHItV5LT59kjhvqwlcwTQLezGnnALQsD9qMh03faQ8E2rovUfbHpLJRR9kWpD2i6Tp1NRLo3UNIKuqo0HUpSli+0D/fUMgI6sqHjGS9pm70UlDsOI8q4haDXiPuysqNowPLAp3nfYP2S6hR1ynPa/S9bIPAmqcrBpgT0/wDKWWAhbnjdNSBR52xU3aT9GPtH7P6NM8zbI5FyOdtsGdQzxVuUzHyC19MWhU8jwy923qAeMdK6K+k52Ja9pDlGscufRObTQNRyLmc0lZlVQCQdgrUU1NKQwBXv45QhAvIADi4NOZZm2l6mkzXJsxlhhzqmvR5xllVFJBmsYHMMrxl6arYDybvFcdQrcYhdpC12A0tEeG0ZdcFILJZ7X/BfjsOfXFeV+a5JmOT1H1XMKSWnk27gsg+0vkynoy+4JGJj2Ndp2ouyvXmRa001WR0+b5HVrU0Mkp/VvzZ6ab1hlUsjA+TH3x3H2o9gPZjrimm2UFFo/M5NztUUFExySaU/enoY7vRMfOeiJQc74LDjirtZ7Ctb9l2cSZfnOTyxqYPrsLJItRDU0nlU008f6uqp/wDvY/s9HVDxjWgtkNuZget41em9UJbNLZHYj8qU9vuh9L0NfBr3s6o5Y9J6nSTM8tif7dJG0m2ail9JqSctTv6oaaTpIDimcvrGo8xWeM+NSGX/ABKQw/NbYuHsW19kOZ5XV9lfaHXJTZLnEizQ5jMu9cozAJ3UVcV6tCyEQVSDl4Csgu8C4rXtF0VnPZ5rDMNNZzQyUdZl9Q0MkDsHMTqeV3DhxyCrjh0ZHHDDFpneaY3ZqB4uuD2q/Po/aip6XtK0fWTyoKP9PwUVQW8qerY07X9tlT+WJH2StW6T7GK7J5pZIqvT3aTW5dMymxUnLVQg/FqY/MYobRGa1FLRPUUZtJFH3kY89yXeMj4Olv8AlGOku0CvosiTtizChkEUE+t9Oapodihlb9J5fVTBQvQgmULjmrSC1r4RtHqB7rUYe+yTcfRV/wDSMzqnm0dpKhijgjZ6HMJu6iXai97m7WIX7ob6uxsOOtsc+0pjqq2lo5E2xySqJPFe6DkgfIYsHtwzdZtSQafRtyZJSUuWMR5yQxF5iPbv6iX8MV5lUTSZvCDc7unzIH8cbNgZ2NkqNdTxVG0u7WcA9Yrr/W+hKftv+jNpntbyEvUah0dk1HlGs4jE26kjieWDL8yLW8cLwxCnnIvs7qJz9hscjVNLPRVU1NWQPDLCxjlifhlYdQbef9cg47X+ij2p1XZ12oUWnIJ4YaHVM2caK/1kBqT6y0yVmXxzoeGikkkmgYH7k7emIt9Lv6O2VadpMr7UezLJ6xNI6kkMNFTv45spqwC0mTVB/ajsxp5D9uNTEbsiXWy2thpGcKio8PwiezuxeNRoVyQsAf7T7LsFJtcWAufienHng9TOH2iNDHEi92kZNzt87n1JJJPqcYJXEbbWKhiCVUkA/LDqONKCFaiTmolAaBD/ALtT0lb3/ZH/ADHi177jTE+SpjKgTTa1MxQ3WTo39y/l8bdfTp1vjASVuLXHTxWP4dDgkpUA2vYdQTf5g/wOMQDowBs22xa3OHgJpKMm1CGvZg+5Ra4ta/5HywDkOftbrnr96/vgGBBNyASSD7fLDiBvqsYqQoEjH9SOtiDYyc+h4Hv8MITRKBXAoJaZqcGDnvBzNY32/wBz5efvx5YLT0zzFjHHvVRfk2BNr/uuT6AE4BRvspDbQeSOSP8AM4d1NUaeBqKMgNINkoXpa9+7B9AQCx+8w9F5biMBmlwOOpN6meMIlPC25IiWW4+056yEevQKPIAYc0EjUVM+btYNC/d01+d05F93uEHiPuUwzpaWaqmjpoYwZJW2pfgG/qfQckn0B9ML5hPBLKkFI96WmXu4Wbjfzdnt5Fm5+G0eWEdj+7Hn1v8AlOGHfPl1uSFFSLUThahz3CqZZnBue7HWx9TcAe7DBqypkq5pKiYbN7bmUewsFA9FAAA9sLs0VJRimHDz2mmPQhfuJ+e8+5X0wyVgp71SrFOelr+l/n+7Dhib3BIcBdShpjUyR0qnYEDPI3kvF3b5AAfLGrqAZpZJUTbGg4H7K9AMbiS9JlwjY3mrQJJCeqw3uo/5j4j7BfXCVRBFBRwUxH62U/WJifK4/Vr/AOG7H/EMMBNevMpxGHXktjPqXN9SZLp3TE5eSHJlnpqQM1wBPN3jWHkL8nDF6eGvrysI/wBXiVmuf+DGCST8bX+LYUp0akpJa5WswH1eE+juDuI+CX+bDBab/VcpqXIO+sZaVPaNbPJ+J7tfxwxrGxgtjFBXmTUnnXinFxeQX9AYBaaKV4qtZ9xRg2646j4fDFqrqH6jmuUa5jYBWqqfNJdi89/FMpqFB8gWUtb0lHlirZk2rJN7hB8Tyfy/fiX6KmbPcjzHSRG6qKNWZf6mVUIkiH+OIXA/ahT1xVt8YoJjk3A+Bz4YE7gVPZHk1i25eKtGhyWLLtfrpqqpT9Whz541k6r3UdRu+TLGp9ipN+gtSGe1cma53mGZSjxVMktQ1/IuS5/NsX7JO9RWZtq1HZe+0mmeX/alno46Qn499K9j6qcc9SMGllcEWYsov6Hp+QxBo0Fz3OOwcVJbO6xoWuBZG3qbFSCD74vrsfq11DPkbPMO9oJKjKZkPnSV0E0afJKh3X4TLijJI/1Uj+jKPxviVdluqo9M6oo5q2ZkojMgqCOoj3qWI91Kq490xZ0lAZ4SWfc3LhQ8iorHII30dkVI+0+om/QejctlJ3x5BTVD+u+eSpmP/wBMGIHLZ8vjHJaCZlsePC4DD81b8cWb285VJlvaHnWnJEVZNM5fl9KdpupCRxhivqD3oI9sVaCQpj8VmsTfi7D3+ZwujyH2dr268eOPumWqokIPhww9kY5Y0uXQVynwLUPTO3obB1/EE/hjbQxLUZDXUAUh6XZmEXkbA93KB8VZG/5MK5JafJM+y3ZZ1gizGIdbNA9nt/6uV/8Aw4bZRVRUeZQzVCgwEmKcdLxSKUe/rwxPyGJqkhw1g/B/CYKAtOoj8flGybdUUlblPVpYvrUA9JoQSR/zRmQfIYzLswam76FyVpq+L6vIW5U3IZHufNXCkfPDOGefKa+KdLNLQzA9eGKHkfA2t88OM1oo6TMJqenG6lcCanPHihkG5OvsbfEYfQFxacjj1ySVIAOsYdc0rkWZQ5Nmy1VVC0lPIHp62C/MkDjbKvxtyPdVwTOMsfJszqMrmkWQQkKsg+zNERdJB7MpVvnhkryys/hLEeJm9bdfj5Hj1xualmzbTlPU2vVZLtpZebl6R2Jhb/kcsnwZPTA7uPDtuB9vjzQ3vsLdYxHv8+Sb1tOlXly5tDdmRlhrUv8AZc/Yl/wuBY+jg/tDC2nI4syMumJ5VX6+weidmG2OsAsgJ8lkH6sn1KHywwoK6SkkdrBkmjaGaNzYSxnqp/AEHyIBwibMLtL9k2UAWYD9onywpYS0t4HrZ6JA8Ah3FFmjMTsJInWSNiGjkBDAg2ZSD0PUW9cCTtUqxDA8A+TDyPzH8sb7OZHzygj1SHDzySCmzO1v/ONt0mPtKqkk/to/rjQxtZiRIm0AljYHjpb3v/XTDo3321OevxTXtuGmpbSsqRnlAJpv/jlQxiORjyaqnUABj6yRiwPmyWPVTfUsCwCFjtHAJ8j5HjBiTCUZGKyR27og/Z/h/PB9m6FaqJQIWbayjpG9r7fgRyvtf0wjQI+7qQ438daBGaNd6uVfoVUHr58+npbn4Ww5hqUmphSVLAhLmmkvburnlG9EJ5/utyOCcNDu5JYgXABbzv5e+BV1e5jjPPQE3P8An+GHubeSA0R5N0RJDFWW4N7Ar6g+h/dhFG3EBeikWt6nyHpe3B9sOqeH64wpCGSoH+xuQDKB/u+fvfs36/Z9MNXVizMbjeftWtf8OmAGuGtIRRKd46yLUQuY5FYOjK7XQjoQSb/PGwocvk1BWrHSiOCofc9TuUrCiAXeckfYQC5YeX3eoGG2WUVZmVbBl1DSy1VXVSrBBDEm55pGayqo/aJIFvfHW3YX2TUOmP0c0rU0+eZrPCaOVgskAcSEGvcHhqWm2SvAp4nliecgxwx7qNutbbKAf5jl116kWrNAZidiNUdkT6I+jRq/LJcvraTUWeLpiaaKZQrUmS1VZUdwtSt7x1NTUwLUGM32QrAtwQwxxrmKKI+E2k9fjjubP+0iq152ba61TTF1h7QK3Os1oULFnGXZMlGmXB2PJZfq9Ufdnc9SccPZu4eprAOFE8hA9AWJH5HDLBIXOe12Yz8aY/HkltLQGtKuvsqnFRkTxKryN+k0OyIXciShUeEfeP6s2XztYc2xb2VzLmPah9Hr9bHLBleV55q+aVD4AqVtbUb+QCP/ADVeoBuORfHOPZRm/wBX75WLbY1pavg+cEpRiP8AklB+WOgcwzGbLs8r8+rmYJp3siocmgcEAipzea17+vd1U7H2U4xLRGWTyDcaeYu+pWiHX42eI9aqJ6mravL9I5lmlRVxyzvlIpVcIEbv620PiA4DbGqGuLdL45wm21Us9UQoMztJ1sACSR+VsWf2sahlTJ6XJUlb9bKamXceQxj2xJ/yROzEeTTgeWI32X6MqdeamjoRJBS0FHG9ZmFXUKTT0lLEA0s0tusaCxKjl2KRr4pFxp6LYIIDM/WqducZJQwKb9jumI9AZW3bpniRrUZLUAacpqhfBUZwEEkUjqR4oqVGWqlHmzUsfWUjFQ6hzPNtSZrUZhX1VVVS1c8lTLPUOXmqJnbc80jHlnYkkn3tixu1bXT66z2h0jpKgrBlGVp+j8rotu+cpvLtvC8NPLIWmmYcGRgo8ES2sLsz+j3RR5bR6r7Rq6lpaGsZhQQshqTXMv2xTU6OprdhuGffHSIQQ8zkGPEjJhCDaLR9xyGwJhiMlIoshmVSWj+y3VOrolrssy5Y8vWYQPmNW/c0gkP+7DnmST/u4w7ngBSSMX5pv6JdDlzL/bXNmopiocQZkJKSRlPQpl8Ky17gkcGVaVSD1GL80vklbL3cekaF8ljiQ0bViyiTMzCAdyGrUItMnB/U0aU8Q6EyW3GYZVpLs77K8hg1j2g5/lOSZYXcvW1sxpY6w2F+5VVNTWM3J/Uodrp1sScZNo02+V/ZQ4nUB85+eC0ItGBjb8uA2nqnqqv0j2H6Ky6Upp7TGd1tyoaVfqWSQhfM3WOsqyAfvd8rYt/Tv0fcmpaKbVOoNLaPy/LaEqZ8zzyKWqhgXzMtZmtRInS/SNSbcLitM8+nVTPWRdn30ZuyGo1TnlZIRS5pnOXNd255p8siJd1BJYNPIVXm6heBXeu9O6z1jLFrL6VXabX61r6clqXSlDmXd5ZSHzWSWD9StrWMVIrHizToRhghtbqPtcnZA6hQuPnjTmndpA3u2Zl87T9vDD2V9Zt9Ljs60EMzyX6NOQprOry87a7O4KGDINL5e3QF5aaGKerN+igpvt4d2OX+2L6QvaHnkVRrntB13mGfagzenem09E9qemyahe6TVlLSr4YJJxvhgY3kEYmlLXMdtRn2vKavysHN6aly/SOUyPHQ5NRR/VaaonABaCNE+wgBBnmu0m0hd7PItqA1PqHMdZ55Pm1bMZGke48IQWsALKOFUKFVVHCqqqOBi9ZYe3fdYCGDMnEncScfEZascaVJn9k284945AYAeQTeKeqrap6yXa0sh6W8IUCwUD0AAAHoOMKy98BaSMKD4wLHkDzueSPywEYVYwiBR5eJgAfYX6/IYOVYKwCcEjcF6EjkXvyD7dDjoQABQLKrU1KN3aSH/aOAq+JpG3eH148hwLDqbYTWnTekbgra9tpu1vPqbXv5DpyRfCiRhwiqrMx3BQFvYjpa/wBq4Jv0IwRlVR3KqWYkIw8JB5sBYXv6cnCoRwAFI/WRlRtckEhv7vkT/hP44JKGMVti9eCTz7rf8PlxzhJRI10HLi9hckgelrEj0/PC8bWjL7iiOFvZht2j2a++3tx74EJtWiGOlkEbMwv4WK2uLDi3qDcelrYeN3cGa5MAhUUeXpO9/M7WkJ/MYZ5mWEUikEuTsIIAIPoAPL0txyMbiagmr9YVWVQJeYwx5fGo58bLHCB/4mxUtJGR2H2HurEA17x8+y7To6Soy7TmRZCcnzaojTTOR0MwpIkdkH1BJTKq7wz2ec3UAG4BB450uoZh/wBmWppPruXuM5Gn8qZ6enSJ5u/zUSuzmwc7lojw4BUgjnqZPqCDKpNR5xHltZQ1kceZ1FKxSSNmVYT9WUXR1lVwsKqDbwDkXJtiG9pF6Ts9pkqpUmar1XSyvO6laiaOgy2qqm7z7rEfWIv1gVS17uN1yeHg/eSh2uteGPsuklF2OnkuT9VVX1vUFXMATvYGwPPN2/8Acsa0gRFgUYgci6lSPe3P8cK5k3e19TICnLhQG87KB8PLCKi3A23PBCi1/jf38sd5Z23YmjcFzExrI470KrwRtIJ8K+7ew6dOPY4G67iTOpUAC/Pl6W6Afz4wEYLHzIA9CL/jjJCVLbiCQOpAF7+RPmRiZRpMKbsu4ra4YLc3+F/33sOuMFiCw58g17i3p7/HpbGE7QRbbt5FiQbenx9MGc897uKlrEPfy9vb/PAhJG6cshY9CDYW+eMxlwpuDwOpt0/r3xmBCLIUYm20FQDYHkC45/DGID0DhifulevwI/y+eE08IAYm1uDa9/h74Mx2qSrA3HJsfPCoROLbFHDEC3mfQ4I11FyxCmwsRccYMzhfusGvYgHlfx64AFgSb7j5m38MIhItcKdxvbg3w3e3JubX4NsLt0AA8uMItyet8KhP8pO6OaLr0P4gqf34u3639YnhrJJlcTZfQ1diCbhqeO6k9FF1k+ZvijsnfbUuPWNiPl4v4YubI5Kh9HZLVW/UCklo3a19zwTuCD6EI8Z9x7Yw9KNo9rvLiK/7Vp2I1aR11ip0sor+z3XLySLIaIZBnp3C0jGlzF6WUsOu7u8wTd7G/njn1qP6vleocrkYg0NfBMoubWDyRMbdPvDHQ2naOXNq6tyAQLNJqDR2f5GVF9xrY6Y1MAPvvoYQPPxDFFVgFXn+bCKJXGb5Y9TEG4+1Cs9x7jY2G2IgMFOrpHsltAq816qPlRdQyoqhwXAsOeR6D29sYXiVm7ssyE28QsSPl0ODE7oVZiDcgKWP9cYBQ/eqwRbgbePEGt626k43Vloxj5MdyWta9789QPQ39RhN2BRV3EW6dbfG3lgfAEBUuLADpcA36fu5xgAcBCPEPEGYkE+g9Ph8cCFhjNl/Vmx4Pr8P5e2EyCigrY+ot5/x9jhUAbSFfjhvPxHy/DnnBZeWLJusTbmxNvlgQkJEDLdSGHW4P78NHUqfb2xsnYt4nkNwNpZjfp05+HmcNpo93TkW8ha+FQt9kGpEaI5ZmoaWGQj7wDFgLB1J4WQDgE8MPC3kRPch1DW6bmp5DOamlMgelqod0bB158J+1FKvmpsw6gkc4ptoypxuMn1JUUIamqCZIJAFYMNwIHQMp+0B5dGH3SMZdrsAf3o/Me439bKX4LWRQP49dc69Ry6j0N2q0p/7TtMtnlQV2/2hy2SOjz2D0Mktu6rQPSpQuegmGGeldGdtnZRLUZ79GzXk+rsncd9mOSw0t6p4xb/zzJZ94mAH+8g75RyQ6nFLZVniwOtXl1T3YPADTcA+iy8W/wAMlj7tiZ5P2kPA6wVfeQzxnfHIgKSxv5OBccg/eQqffFFs08HcIvN2FW3RwzY/a7aF0NoT6W2gO0KJsk7R8sm0JnkN45K6nhmrMpEgNryKN1VRc9QRMg/uDpY1RlKJpuPLNWZdS5/o3N5zVUM1FXJJRS1BFvrWWV0W5KarsftLbeLrPEwJtQVR2h6J7RmibtJyik1JWpGsK5s8rUOdRWFgUr4hvkI8hUpN6Xthtp/ItW6IrqrM+wfttpadMxJFXpzViQ0cWZD/AIcpffltWfLc7QvzcBTisIIJZL0B7N+w5HroKa/LGykvfbz+es1Ae3fsArdEZlFqvQ07ZnkdZM60dZDTiFjMq73ppoASKerVPE0IJjkS8kJZLqpslpqf6RuhaLRyCNe0LTdOKfTzk2Oc0aAkZS7HrOi7mpGPLoHpiSVhxdtH2mabhqZ9HdsGjKrsyznNIFp6rLs4+sPpnOVVgyGGqXfNRFX8cUymdInsyOi7lNSdufY9m3Y5qeLWGQztVZZmMK10M0Uiha6lJVu/jkhJjMsb7e8MLEK4SaPwOAmmyaWEgSjHVsO7z1bDuKomNklbhw59DXtG8KkNNZq2R14pqpdoDldsoKgNezI3mAbW9iAfI46I0/qvL9S6CzL6+ivJQ5VpyasVz4z+h62eI7h6mnkhGIXr3LaLt9yes7SNNKja5y+nas1LRRoqNnVOgHeZpFGvAqUFvrUSizf+cINrSBYr2NtmueVOptM5bLAtbmWks0iHeAk1KQRrU7Ft9/ZA4B9OMNt8DZIzOzDKvkapIJi0iJ+rJQzNKytzLNp8wrnDT1DtNJ/6SQmRvndvywXIpAc+pl3dZUFz/jXGVxUzzledzFuOtiAQbeYsca+jlaLMFkFgVNwB5Ecj92NMt/dXRsVJrv3ld66FzDKKzN8k1vBlk2zM8vp4tW5a0Zs6zZfOUqCp8mFPMZL+kGO1+y7tC052kaNEGt6N810R2oZOKvN8vgjAkoK8OUrHha91khrI3qI/vDvECnmx5L7KNS0eT55QZzmkLVuX5bWSGshIDbsvnVkq4x52amnkNul1xIux+szTsr1Fr7sFzisLz6JzWozfK5bE99RjalUyW5ZWhFJVgDnbG5Hnjk5A6Wy3mffEQR4ZfHNbrC1loo77XihVP/SB7Jsw7KdeZjkOdPDVVdHOsc00K7YMyjlXvaTMIrcdzUw2cgdJEkXi+Kjq5JJJXlmctI5LMWHJPy4P+WO+vpD6bynte7D4tdUstO+aaBMVBXTQyLIJcgrZwI5bjqKWtZWHok7jgdOBKqGopauaiqotk8MrRSJfo6mxH4j546SwWptriEgzp+vP2WRa4DZ3lmqvXL3SQ5FyBYWFr9OPP2OBRiw22F/IkdR74zg8FrXPHscKUsElTJHFBFvklIRV9z0HXF4kDFVKEpWKn75rzhlhhXdIwPNr8Kvu3QDp1PlgJqgzSMzKgJtZRwEA4UD0AHHvhSrkijtQUcgeGBjd1/3snRnt6eS+w98N5CxstgD0sfuj1PoMMbU94p7sO6EuTJSctxKwui3F0B+8bdGPl7c+mGfAO4XFvDz5YUuHI7lLDaAFA6geZ9T6/wCWD0sKzOzSq3cxL3kxB5CA2t8SSAPc4Wt0VKbS8aBOoWFBlzyAH6zXqUQk3MdP95vi5G0f3Q3qMM40JJZwWjjF2HkeeB8zx+ODz1T1MrTyhQz/AHV6IALBQelgAAPhgfrM8VO9Kk22MyCRkCXDECwuevFzx0FzhA0gY5lOc4E4ZBJOzuxLsC7Es3I5Jw4y6kgmnZ6sN9VpU7+pseWHRYx7sSF+ZPlhpuDEbyQB4j7AY2Fbakp0y3/elhUVZ8u9I8Kf8in/AMTH0wP1NGtDM7x1JNW/SOYtU1w/VEmeoCCwCDqq+g6KPiMNZ52nnkqZLBpGLsF8r/dt6Dp8sKlmjp+5AO6Uq735G0fZHzNz8hh1kUcS131+eMPT0CGrkUnhipGxP+Zyo/HCEhgLqZZdb0oq8hu3NBnANGYcpZRuoo/14Jved7M/4eFP+XCVcVHdUoHhp4hGfTefE5/8Rt8sJh3lqXqap2eQs08jNzvbqfxYj8cBS0s2Y1cNApJkqZVi3X82NiT+ZvgAuAV1eqCbxoNab5nT9xT0kJFneL6w49C58P8A7EKfnhDKMyqsmzKmzKklMU1PIssbr1VlNw3yIB+WH2bVS1+ZVNVELRO5WIW6RqNqD/wgY1VRGVYAEEBRyDfAG3mXX68/NBddfVupdC6pz2lm7Gl1PlVMkEFasGmpEXokkVVNX7F/ubZYrH0FvLFD9z4o0vt8zuHQAdcbGDUGbPpmDSEkkhohWnM403cCQx91e3S+0YZsoU3sRt8tpvc+2KlgspsrXA63E+WpT2uYTObTUOaK0JOXTyWHE0K/iH/ljWFGUki/A6+2N+iNLkNdNtJC1lKt/ikn8sNqGiFTJJEVu31eVgOhBVdw/ccW2upeOw+wULm/aNo9ypPT5zmOrs4znNM4kMtVX5LOCSxJJgp02cnk8QDEWe+wleeQwsL2+P8AnjfaLI/tTlsTW7uqkelLDoRLG6EH38WNGqsiRll8aABh5kC1x79MRxNbE8xsFAAKcx7J0ji9ge7Op9j7raaTmhh1FQrUqVgqJTRTm1l7qZTE35Pf5Y1xSqoaqSnn8M0DtBICLi6naQQeCLjpjJQTK6o/3isZ3dOfCR+WNxrFY5c8fNY1IjzaGDMQPeWMGQD4SCQfLDspRvHp+p4JM4/A+v6JrXR5bWQR19Cj08rLsqKcsWVZhwHRjzsYXNjyCpAJBGHFWordMUlZb9dlcpy+f1ML3khb5N3qf+HGtikX6pMljvM0RH+Ha/8AHGy000dTVyZJMypFnERo+eiS3DQv8pAo+DHDXtuNvD+U18tfIlKw3nXf6sPPVzWriA7kk8sGEqm9rr9lrfkePfD7J6+Ghrg9YN9LMjU9Wq2JaB+H6dSOGBt1UY17ozEmaO0l7EEWKkcWHoQbjGd6zW7w9eQ1uvx/r44mcwOBB1qJri0gjUnFfRy5bWTUU53NTuVYgXVx1DD2YEMD/ew3i8bclUsSVBbaAvpf2xt6hZcwyFK5C2+gZKOqAP24jcwOfWxDR39kxqY1Yjuw9ieiEWLfA9D8L4SN14Y5jA9c0r2hpwyK2eTV1JR1rw5jIzZdWw/Vq5kHKoTcSKOLtGwVweLlSPPDPMsuqsur6jLq5l76mYoxU3VgBcMp8wy2I9QRhvdL7EJLcHxErz5qf7w/PG9mjbONPmZdorskiRZlNi0tDusje5iZgp89jp+zhjv3b7+o4H2Ptw2Jw77busdH54rRxSEeFCUNrMbCx9Dbpz0PlfnDmgrYaSZ1qoTNBMO7nhXgul+o/ZZTyp9R6E4aoUW4dR6cEgn2t0wMsrMo3MWN7bQtlHyHX9+JXNDhQpjSWmoTivy5qSoMAkjqI2USQzqgXvoj0YHy9CPIgjyw3DFAUCkFltJfz9CPQ2/HDygq46uL9FVJXcWMlIzcbJT1T2V/yYA+Zw0bdE4vEDa4CP0Ug+nXr5dOMNa4/a7MJXAfcMiiqoXlwLDpx1+FvP8ArywvUSirQztfvx4pL/7wftj+9+16/a9cJI5CqgJ56cdfUW9vTA0tE2aZhBl/ebBK1nfyjQXLsfgoJ+WFfQC8dSG1OA1q7uwPR9DFHJq7PYZJFMBdI4yRItGzGGyMvKyVUt6dGHiSFauUcqhxafaTqLMtDdk+ptYTzx/pPUyDTuVNAgRS1RFeqkiUcKkVEghjA4VamNRa2NLoLL2ynJ8vSelmX60IqyrWJdzxM8QEEIHmYqYxqE6kyykcnG81Acv1d9IbS+jKwJU6Z7Jknr84jTlJ6yntV5gPgakUtEPXYotzjlu2E9sMr/tbU8Mh1nmtsx9lZ+zbmcPlbOqyg6NyGj0HV00appnSddk8jKLg1QymsmrOehIqZpFPmO7F8cQ5kl5piBfxD4/ZW+OvO0/OpKDKc4zGWeVq2uy7MKmpbeCs1RUr9XLi3Ul6mTxdSF9scg1MhmkmdejSMQfUXsPyGLmgyZA+Q6/VVtJNDLrBqTvSmeyZNPJIAGTupo3U/ejljKMPzU/EY6V7bs7hynI62PwPU5tnOV0Kp91qXJMogi8f9z6zUsSPPuyMcs5XTyz1vcQRmSSRSiIo5ZjwoHuSQPni0O2c6pk11L2ZTSU+YV2nj+iKiWhDus9T3m+baLbizTyFSBcsygDyGJ7XCH2poGsGvgKe9FFDJSLHVl5/hRvJ8s1H2tawodOZBRVWYV9dMKakhiXfLKzMSWt03MxLG9gL8kKpIsLVc1NpbLF7EuyuWmrZpZBU6iz+OUd1WzQ3a0cpsFoaY7mVzYSSBpzwIQpq2Wg+jnp6t0fRzrP2h53A1Jn88MgYZPTMLPlsbL/vmBtUup4H+rqeZjiWdi+h8whyCs13XzZVlVBHKgqM+z2VYMspqhTuDSEqxqDDYd1Sxo7ST+NkKQi7ZrRdAEY7uTd527x68w6OG8S5xx17k77LezfTfZ9Sirz3LI84zuoiVosvqo2WEI4DLLXLcP3RFmSiBV5VIeoZEKxPdNLpmGlhpe1XtO1NDRwVkql82zZ9v1qNVstPTxxrulCKAEgpYykYAACAHEDyKrz/ADJJKrsK7OKrU8TSM0vaBrqJKHJY5Wa7ywQVL91K5Zt/eVDzOxJYQg8YVm7JsmzrPJtWduXaxnPaJqKVAHpqOsfL6Bl4/VfXqhe/kjHQRU1PGhtZG6YzrTCXCtqfdrqzcev0orkTwDSztrv1Drqq2WafSgz6vzSTst+jR2dZzV5yZADmVblArsylcC2+ny5N8NNxY95OZXHB8BGNBT/RzzHN86m1p9JftUrZ83mJNTl2WVi5xmx8ttTWkvTUajzRO+dAP9mOMbzP+1TTunNKvpDSqZXkOSW8WW6fhaloprcfrrN31SwtYtUSSMCBwMVdm3bXnFdHJBldJE1PTxhe/jhVdluOZPCBby7y7e+ImTGNl2xsDRrJz68SVIYQXXrS+8dnXsArOzbWum9AZVU6Z0Fpyi0plsqhailon7+fMF6q1XWteSsB6gEiH9mLyFN6m1k2ZwPnef5jUxZaxKxKrhZ64r/u4L8Ii2AaYjanQBnsuIbqHXkKQ9/V7aqUg7Inv3PPmV4MnPUDah8y17YrjOM9zTUdc9VX1MkjPYeI/dH2V4sAoHRQAo8hizZbBJaDef5k5+Xzw2iGe1siF1nAe6f6t1dX6srlLiOGkgQQU9PCpWKGFSSsaKeiAknm7MxLMSxJwxpYoo0uw8XBAK3B+I4t+YwlFTsLIsbbj5bcOEDxm2xla1o26eL3vcEWv0x0kUbYWhjBQBYr3mR15yVldr7piGY22sWsAvS3Itt9D5YUkG9lLSAxoSrSKpYC/INupHFh+OChArRzI7KwN1sviAt6/Z58hb8MYu9THIzhgT4SQSDc28uSCfTnEiaEoJkdbiQJssARHwv/AC3vY888++E5R3h2IFJIuCilrAdRby9ha+CyqkkgYy2sLAhdvzUW9fXn1vgYlXcrmQjYN21eq26kedvUjkelsCVYC7oUWaRri52SF7r8B/HAM7NMXezs/LBup9Ctuntbj2wpJM/WV3IBvbdx8gPyPXBCsYkLJLJt5B3DbwfM2vt58r/jgSI1JT/W8zoaSSJ7S1MSFTywXcLgfK+LA7BcrXV/bvp4tf6vW6ro5ZS3lBHM1TJf4JCcQjT693mbViKpWhpKmrO03A2xkA/iw/LFufRLy1YM+zbVE0kUQ0/pfPs0Ekr7USV6daGAlvL9ZV8H2xlaSkuRvI2YeOJ9grtlZfLRv+B7roCaseuo48zqZKaOedhUsssix7RMrzMRuBKbS45tY3N+oYQTtrzNYNOafpJNwmFDqbM2vax76qpstiItcW2U8lrEjk2xL6qtOWSUKxwslLCiNtaWx7tTYvuQsrKFC7ZoyQCLOqhzanu3HO1ec0JjCHKtO5Hltr8K0kU2ZTW/9ZOl8cpYW1lwH6E091vWsgM62VVBl++mkmA8EjuSTyOW4GB5ZHLEkng+K9vc8cenOCxqFVWUEFFUXHBHxI/ccC0ahV2WPJK716eoPv8Avx34FBRcrWuKxVHkPETYeQH4cn5YVePvAUAuxU2AJax+fU4SjO5TwL2+7wP54OwjAVQpG3i977r+ZB6fLjCoSb94qAuCoABtYgE9CbdAT7YK8lj3ofYVso8hb9kW/d88HclBsLLYMeluD0Jv5n8jhJi5UFG3A8cWsf3H+OBCQMh3sBwp9+vx/fjMHKgdTdm549f68sZhEJK4DA2Xwjnp+/ywoW2qdpII44NuPc4IyFeChU33D4H+uD54xTZb32m/UDn8R1+frgQgvZ9yjk825Yj5nzwU2VeASAbKSbEYMzrcMR5C3kL+xwQ7+SSTfjxf1wcCESUMBc8Fj/Drhu2F2YXuAefzwiwIuCfPCoSuWuI8wgLdC4U/A8fxxb/Z9Ms+k6qmnkI+pZoNqk8MKiDaR7EtTi3v8cUwCQdwPI5vi1uzuWKpnzrK22layiWtiViQDJA6TW45/wBm1QOPIYytKMqy9sp648iVfsLu9d6y+Qrc0Tm509qvTOsplWWLKM5yypr9pAJiSZEdmB/ap3dT7ofXFSa9yBuz/tGTI68NbT+b1mQVXHJWmqpIre94mXFmZRlFJm5fI61ngFXC9GsgfdtDAr14NwGDc3vsuDfGi+k/BNnWbUuuwE3avybKdTsyDj61JB9Trre4rKSUn3bFCwuDqtO3kRRWbUCCHbvTFUhVUzUk81L4t1PK0XK9NpIufTphPaWIeNDZunrcdR8v5Y2WqbT5t+ko1ATMqeGtW3kXQbx8nDDGrUrvXcARfnk8+1/446GN19gJWQ8XXEI7HcrOXups3NwWJ/fgVW3PCspIuFHwN8AF2hFVftbgpZhYetvIHA2dkLBVIC3J27hfoD7Hy98OSLIixXaUN1BPXyHUG/U/DnBpja6B92w2S45C+XPmORgqKN6AC5v8f6/lbANwFVxyBbi9x7EevX92BCINy2N+DxuB5/r2OMdWkJB8Tg/E8ent7YOSxcEjcSBwRcEfLrhFX5IUjxDy56fHAhJPGG4A4PTCDREcg42AuLOyBxtuotxb1sOvw88Jd2AbHnzuP34VJVNoKmopH3wStGx4Nj1HoR5421HqRok7mphunkEsVH/IwKj/AJduNe0Ci/HHqOcENN5D8sRPhZJ9wUjJXM+0qRx6jp2I2S7fbcy/k24fnh7BqrMYXd4Mzb9Zy4LKwb4jdz+GIYYHtcDjAdyxPl+OIHWKMqZtqcFb+mO2bV+maJ8qps1pZ8qmP67Ka6ljrMvl/wAVLMWiHxVQffG+i+kHR0+m6nSsOisqjyqsqO/q8gWvlfI52PDSpSys0lFPYm01LNGebMrLxige6N7Yzumv06euFbZA0UrgkdaLxrTFXbk0fZg+a02qOzPX+d9nmoKWVKmCmzofXKOGdSSpizGmXegF7DvoBwSGZgTey8t0bXZRrLIu3bTOlaalzTJagZtqbT2VyRz0WZ5et1rMwyiSIsksDRPJ39MpL05bcB3dxHyREZonEkTMjL95TYjEr0t2h6h0rmNBm2VZrWUFbltUlbSVdFKYpqeoX7Mq28JYDg3HIJBuCRgmiLmFuYOB8Ekbm1rWhC2va9osaE15m2m4Klaqko5ytFUobrUUjAPTTKfMSU7wuD7n0xAnYRTq6G4U36Wx0xm1Rk/0pdP5YuQZXl+V9pWRUj0kmV5fD3UGo6JXeVJqGLotVDvkV6NftxEGEXTu8c65vlNVls7U9XEUdefZhe1wfMcfkQbEEYlhJLA12etMeLrqhWNorUP1eSk7+S8csf1eTwhwSosLqeDdCvXFtdpme1um807MPpF5dAamoyRotMZ8oAH1x6OJe4Z//lnLZBEfVqeTnHO2j6kVJ/RzyhJQwaJibAMPsknyHO0+m5T5Y6E7NazLO0TI857Ic1VaaTVNPHSUk8slkps3hdny6Yg/ZvM0lM58kq2/ZxiENs9puOHddgfPL4WmazQXm5jEeWa6B0TT5HR6ln0tK5n03quOfTNTVxxFY6nKs0hCQyyWFgwWaln3ftISPPHAWuctrMs1PV0WaIEr4yYqxSv/AL8xM0Mw+JkiY/PHQOitV5xUdlmX18YrqbOdHyzaWzBUdkqIBF3lRl0jrcdLVMJuDb6uB6DFb/Sfigm7ZNQ5zSgGnzjMJM1Tb0ArIoayw9LGobEOiGvstpfZ3Hb154FLbnNnhEo3deqqdgx8bMwuLE+fzw8ilahgZd9qmoTbcmxjiPX4Mw/BfjgYFhRGqahdyobKpPEknUKR6Dq3tb1wzeR5JHeVyzyEs7HqSfP0OOkIvmmpZA7orrWEeKwjvc+EAXIPtbAncN0RuvisyWvc/Ec4BtpBLyeIkLt55Hl04t5YBVIdk2EFTbaQbgjrx5f0cPTVg2m9wSOuNjV/6jAuVBQHDiaq2m36y3hS/ogJ/wCZm9BgMutTs2ZTIH+rkCBCL95OeVB9Qv2j8APPDIM7ybmLSMSSdx5Ynqb+vniP73bh6/jrJPHdbvPoinaWsCSPIEWv7YN4SDyGB4BPB+OMe43B15HkBYg+XGDDzBQ9eCD1xImJxRL3DNVyxgimIba335fuIfa43H2U4bfrZpDuLM7Es7sep6kn88O69VhVKBTdaa5mtyDM32vw4QfA+uGRWxC22lrhrXBK+nOI2Y9/b6J7sO7s9UddzjvCLk2DA+XoCPS1rY2FVahyWnpVsJsxf63ISekK3WFb+53v/wCHDXL6Nsxr4KQsE70hXkPGxBy7n0soJ+WBzSsXMK+asjj2RSHbCn7MQAVF+SgYR3eeG7Mfj58ko7rC7bh89b01J8W2wFhci3I544/P5Y2OT76cVmag3NLTMqNa362X9WnzAZ2/5ca4Ani/DAWF/IdLY2EsncZJTU/3qyZqtx6qoMcY/HvT88EgqLu3o8kRmhLtnQ5rVsuxWHRVHHqcEljAkkvYWJ8PS9vT3w4KIyszAbQdoF7XNr8n0H8sGpqfv6mGlfjvpUQnzO5gL/n1w8mmKYMcE5zGnFPWtThQRBFFAfQMI1vf5k4b94GBC3NmuPIlbAX/AC/PC2Y1Amr6uoXbaeola+4k7dxtx06YagMOm3kkEEHpbr+eGxijBVOeauJC2CHbputQHwvX0x/COT+eGy7oHRtxRgehuDY8EewsfPDiJh+hJo2N710RNxb/AHT+mGbKd3dspIuOPY4bGMXePsE5+TfD3Kd5TVnLs6oK1gUNJVwSMtugSRSfywvnVEKLPc0oO8CtTVk8Sh+FYCRh9ryNvXj3xqGL9225rlUK38yRcYkWtQsuqsxmDWFQ0dTe1/8AaxRyf+5YacJxvB5EfJSjGE7iOYPwtOzRhEii2HYguyi25jyRfzAvb5G2NvmDCs0rk9T1ky6apy1+L+BiJ4r/APjlHyxpRvjAIIseODfd7e+Ntk1qrKc7ypgu40yZhEo5tLTt4h8THJJx7YWUXbr9h9cD6pIzUlu0fn2WldRcFhfbzwxFxhQMQVZS0fRlfebqeoI/DBdl18JB3W68i5459MHk2/dYDaNoJuAQOnwP78TqJbbUTx1lZHmtOoCZrCKwoqg7ZuVmWx/7wMfgwxprtckE7vUE8n+ONrSqlbpmqiF+/wArqFrE9TBLaOT8HETfM412x0Ku1iW6hlJ+TX4N/TEMGDbn9OHxyopZcXXtuPzzWxyLMKamzE01airQV0bUdXtBuIntZ7nzRwrj/D8cM6mCWkqpaGrDd/BI0UibvtMpsdpP4gemEgihGLm6+f8ALn+Bw+zGX65l1Nmwcs9vqdXcX3Og/Vufdo7C/rGcKRcfXbh56vjgkreZTZ0evFa5TGbGRb82YDj88PcpzOfKK6LMIoopO6JV4WWySxsCrxm33WUlT8Qeow2IVAA63ktfYei3HRvM/Dy9cJiR0UBG2hvteI2t8uow9zQ8EOTGuLSCE+zihgoarbRTvJRTKKijd+S0LfZ3f3lsUb3U+2GheNrKVJVuHBO4E+XPX+IxsaORq+hOVsLzU5eejseTfmWK/wDeA3D+8p/axrWKPuIYlCLjzP4/xwyMml12Y6r1rT3gVvDIoj71S3IVuLDm/sL842BSLMKM13i+sU6gVQ/bW9lm/Gyt77W8zhjNMHbxRKHHDFRtBHrbyPw4ODUlc2X1C1MQ3MoKlGN1dSLMrexFwfj7Yc8E4jMJGmmBySMrBWJDEA9Rewb/ADxJuyzLKbPdWw0NawWlqXjpZmPG2F3DTH2/UpKPniO10UETiWk8UEoLRb+Tt81b+8p4PyPniT9kbmnzWeqWPe0dPOUW9rsyrEP/AKacVba+lmc5udFPZWDtgHLq7Smc5bkj512tztII9OUMuoZqcjbHPVlwtBTMvIYNVy04H9yN/TEH7IvrmgOzuo1dnUyVOa68qpKmWaoHeE5bTzMNxB+139b3jkHqtEPUYe6hhrNW6d0P2O5Sy0GadomYHO6lyWcRZdSGSCnlc24jDCvqiD92CP2OGnatq7Lq+eHTWmssfK6GlpoqKhSpKH6tQU8eyFmKkggRL3jk8lnkP3scuwmKHs6YvOP/AKjD5WwSJJL+pvqq97UtYx1OW1UkTjdmFSu2ygb46cElreV55QOP+CcUwyBIiAwa3QW6fPG91lWiqzOOniDLFCqBI2PijhUfq1b++btI396T2wnprSmfazzSnyLTWVVFdWVMyU8ENPE0kk8rnwRqg5aRvJR1sSbAEjprDC2zw12rHtUhlkUt+jxlML9oVNqjM6F6rLNLK2fVUSrcz/VirQwAebS1LU0QHUl7DnF0UegNY9leS5jrCWbK4O0nPKmaTM9QZrXw01DpSSQs0yRSOb1OaEswbuVf6rusP1xJjh76qyr6N2Vwae0lmVLmOuI3M+aZ1TSLPS5bVrdUgozytRNBucmo5iSZiyb2jRhQmY51X5hMZHkdRzbxlmsSSbseSSSST5kk+eIY2utUzpAO7ljrA9vXDzeSI4w0nH5Vo0dB2Q6JZ81ra+XtAz1LNDDMj5fkkcgv4pWkYVVWAbHYqwqT1ZgSDKa36SmbZwtBI2Q6YfMMthWGkrKylXMBQgEEiippbUlGpIBtFBu/vk8450ZD9rm59cAImJ4Bt8MWnWYuN69iomzNaKXahXrqDtz13qOoWrzzWtRXVUa7VnqCZGjHom+WyD2Ww9sRCfW9dJIz1WoZm3m5s8SX/DfiuxTSH7p/DAilkJAA5OK7dGRjM8gp3W5x1c1NK7WtA/iYtVSk3LSbprn18e1P/YHGkzPWGa5kVQMyqn2C7byn+EcKn/Koxq0o78/av6YVjgFwLCxNiRieOwwxmtK+KhfaZH66eCQ7qad+9ncsz8lmNyfc4eQwKgUtY39DyR/XywdIQPEFJIPIPha/9eYwYIAA1lPwNwfx5xcyVbNChiKMjqefDyALe9uePhgXKINkZXaW3NtHmOhB/G9sByjXuAGHhGwWKjj38/e49cKK7d5vIUkgHx9FPl162HS/B4wIQKXayFeehKgE89Lm/Pw498GkAd2RrWPBKvcMQBfnjn0B6eWE9jCM9APLxc28+QOR+7CsgkDhtjJa1iwsbWFhcjmw+fOBCOBMQZLjvAwXxRGzMb33Hpew6EXN/UYL4YmN1UGMruYuTsbyuLXBHzxiMGHj7wiQEhgN97fPn434wk7qD3igX423PS3x4+XrgQlCOO7Dr4nBG5Slz6kny/LnCZZ9pjRyUY2Kk7RcdD6H44zYiliX5t4bcgsem4+X7/LAldyqTcqTbcwHX4/wPOBKgiqPq+XZpJsBaWKOlV79Nz7mt63CW+eOnvo20RyPs71PmJLRSZpVZLkSMIDIe7jEuZVRuPsi6Uqk+rAdSMcyCEVlNlmWRJ+szCseQ2NyVLBF/c+Ot9E5e1DoLRtIrzJ9eGY6rmhU+FkqakU1NuW1z/q+XswsQQGuA3Q87pp1Ii06z6Up6Fa2jm98HZ+vuFJ1yarzyoOk6CvlC5jVU8PdUuWFik8xEJmDAgR2747mAAPG6+Oc+3bPqfO8+1Dm9Cy9xm2o8zkowv3aOKRKWmHyjhbHSGRVMeTVL6xqwXh0zR1eds4G4bqOjllRVcG3MwgFrnqvpjjjVrmlhynKiCr0tDAJebkysneyfPdOfwxnaHjvzA7CPevsrekXBrCN3XutJ3jlQrDjdwLX68XAwF7/AG9xXoVfj4i3lgUU33SXdSLkEfaHn64HlblLLbpc2a3ubXPy6Y7Jc8ikpvKIjX3eFR1Hw9/PywBLDiylWHU8m3qCeR7+eBeEsRuG1WW/jFrn28z8fPBSJFNguzzvcEA/LkYVCzmxs1h5WYEfK3XANzyqqD5Acn4XPJwdiCzneVSQC9wBcDmx4te+EXLXK2tex8uCP6+GEQiCwJdRe3HTi3v6YzAhd7ckrcdT5+3/AFxmBCIzvdY3LEDoL+V+fXCbEgnyANsGYm4PBPA9wR0wI4uDxtutvMfMfxwqEQncbEK3ht1t+Hp/HBCLqBtX+NvQ4M212ADKCeATwPjjCR8PK1/zvhEIkgsbA3873wgQObAYWcqRyQSw4OEev4/jhUJI+2Jz2dZvFl2oclrKlgIROtNOT/wnJie/t3cxPyxBzh/lh3xyQFrXI59A10P/ALYfhitaoxLEWnL5wU9neWPBC6QolqMvk+ryOEqadpKcna10qYgbqT0u22QDqObemHPaRlTag7KcqqadE25BnNdkpVTfZR5pEMwowB5KKqGvQf540seaVOc0dLnckarLmkEFd3yk7u/ACy8dOJo3b/mxONNwvnOns/0tToe/z/IKhqOnsxAzLLH/AElSGM9G3JFWxAHxAShT645uzvLJADmcPP8AGS2JmhzKjVj5LmGsBn03lNer+KkmnoHVV2sguJUO6/JO9vhbGphIRlZ9xjY7W8N/y9R1xvq5II6zPMphRzDUbMwotp4FvGL/ABjdhx6DGgtY2C89OCfFfHTQOvN58cViyijutWCWdmmu1lZxyDtsWt5H1P54IvUEBSDyLi9uOv8AXXGIVKKNxve9r8f4h6HAnxgKLsSLkDw2Prfz9b4mUSAKrBQABfgAnj53/fgu4kBlseeLjz9MG8LMQzbR6WsfkPj8vPBtg2nqAQAw68+h/ePngQkgBIo3FibAX2348vhgdhG+QhtrGxF/O3BPHl54O2/YHZUVeQzkfv8AK/78YjKwEbK1yN6spsQT+R6f10wISarYFSosPvcj5X8vjgjsEiLAMLcgE9DhV7rtkaPqPCTe3+YwWdfAUta3O3kAfLCoVnaAo+xjWlGMt1NmNTpPN1X/AM7ekkqsuksOr/Vz9YgHqwjmUegGJxB9DrtA1VG0/ZXTZTrimJOyXTGrMur7qD1MMhiqE+Dxg45sR5qWZZoZHjkQ3V1JUg+oPliQZJnOXSSlc0pisrtvFRTt3M6Sfto44Deqt4T6qecVJGyQ1cw1HHrirDbkuBwKtGu+il2n5FVyUerYsk0tPC1pIdQaly6hkU+6NIWH4YfUX0csiiiD5t2raPlkPJiyqqrc0e3xpqXZ+L4PpXtx7ZtD08SZP2wapTLZX2xM9QlfS7v2GhqgwRvWNiPa45xY0fbZ2j6nCT1ld2MaonK27jUmkafKJ2PtNEscZb/1t8Z0lqmfg14HEfKsMgjYauaeRUWg+j72UPSieo1lqBl7wQN9U0hJIVlIJVT39dGeQDbw2Njh/ln0eOyDMZVh/tfnNJ3lxHJVaJmljkIF7KaXMnJNuSoUta5ANjiYUOu9f5ZTSVOf/QsizGkqYGglq9J5tmL05Tghh3T1USlSoZSRwQLYTh+k52ERzTUWruzfXWQTS+CpEj0NUzMDe7KVppAwYbgwsyuAy2PWD/8AyDRVoDxud+QpwbITR1W+IHwVEcy+ihpmeFlyHtE0fUyuR3KS5pV5LLKD0smZU/c8+nf/ADxVPaV9G/tQ7OqOPNc60lm1PlcpslfJShqRj5WqoGkpmB8v1gPtjqfLu3HsDzqiqab/ALU3hhq1eJo9Qaaq6UyqwsDJJTGeLePNwq3PiIDXvNOz9dKpmm/sv7X8iWeriA/95/UlPTTEjjupoGaNZwRchmiN+QVw1ukLVZj+/idTdU+tfUJ5slnmH7p488Pj0XnLls+caWroqgxSwuGWZAWZLlTdZI3XkMCLq6G4PN8dBx57o/6TuX/UdWZvl2Q9olgIM9q2Smo89mPATMG+xS1jcAVgAhn4E4R7SHpTtL7DMu1LTTvrTsxqqfvxufUGm8mWm7w+T1NCAKGo95IHppfc9McmdpP0ZdQaNo59XaQzSlzbIYJO6kzOgkZqemduBFVJIBNQuemyoXY17LK3ncit8UzrwND1mMxz8VXksr4hdIqOsjl1kqh1VpXVfZvqus05qfJqvKc3yqUxVNJVwmOSM+jKfIg3BBKsDcEgg4kGn83qKmb9J5dOwqQhWSLcd0nHNj1LW+bAAjxA3m+Q9o1JqSgpuyv6QWS1uY0GWRilyzMFsmc5Gp+ytNLIQJoOb/VJj3ZBPdPESLxvX/Yxqnsxgj1lk9dTak0dWTGCk1FliuaV36inqY2tJR1AHWGYKwIuhYWY2pmMtTcPupxHwoInOgOP2q58mkp9aZ/lWrxJEtF2wU0uk9R+UdHqyDbJS1TAcAVDinnv59/VgfZOKr7fJ5H1Rl0NVEaeoGU5UtQj9YpFyylR1b3BQgj1FsIaF7RaCGlzPTecvNFlefxxLXmK/fUdXCxekzGMDrJBITcrYtFJMp5IOG/brqga11hTaySIhs7y+krKi0bIn1ruFiqDGWHiQzQylWFxYjzBGM+JkhtzXubTAg+OFOtgG1WJC1sDmg1GpQCaUOwVQVRRaNSeQOvwuTyf8sFAv4WYKFJNjx16/wDTBWP3j9kEXFufj74DaV8J4YefqPI/MY3gAMFlk1NSgFi5Km6gkAkWv6fDC6QyTOkaruY2VS3X4X62+PTCAG5wq7bNxYnj8Th3FJ3NOxv+snBVSBfbH0Y/832b+l/XCOywStxzSNTOoKxRG8MVwn965uze1zz8AMFuxYtIDdrtyLA4FidgFztNwBbg+vz98J2Kte545vbr8cKABgEhNTVKsb8W4W3l09r+Xww7orQiSse/6m2xeu6U/Z/CxY/4R64ZoL+EXJA5A5U+1vI4d1kiRhKOJiUprh2XzlP2z8rBR7L74a/Hupze73tia8vxuFxz4rc+vX92BBQ32oNp4BAuAPh6eeMBA5V7kkeJSfn16HGFizXCs5JAt5m/Tp5++HpqeQFaTL6moAtJUj6pEF/ZNmlb8Nq/8xwwY2uQoUni4Fvjh3XtGsy0sbFkpV7oFfvP1cj4sT8gMNST4ivIB49AD1viNmV7b0E9+d3YgVn6R/aJ2gdbnoBh7nDxmsanjIEVKq0qHyIjG0n3u248euE6ALHWJVuAyUqtUkAGx2C6j38W3Dbx3szeK3mL3Pz9cGb67PfrmjJnj7LIzfd52uQMOsufuq+BzbbG5m/8KluB5dPI4bIdg3DkghrHzFsKwuAJpyOBTyBTc8MbLb2+15YV+LSEjMHBIRqDGO95Nhf1PH78H2kAFwQSSQCLAi3l+GDBdyrtUjm45HPt7e3ljAGkkCxBnZ/shFLFvgMPKYl1LLlMrkGy10VyOOkbcXwiQO8ZG8duQCD4geQT7YmOS9meqc/0pNU5bF3tTJUtVxZcOKieKNSrtGp+067ge7HjK3YA2OIfPBVUc0lDVU80MsZKNHICskbeYseR7g4rwyseXBprip5GOaG1GpEmhVQVG4biWF16g+XsfyIxs83llkkoampkD/WMupmDbbbdqmMDj2jtfGui7ySQI7BF+yGY8Kff0W/4dfXGyzdCMnyGV7o3cVNM467THUMefk46Yc/B7T4jlX2TW4scPPnT3Wtc3BCjjcLkjzsb/lh/p6qjos8op5h+o74QzD/upAY3v7bXbDLxIhWw5t4T5+nH7jgjx3DxpfqV3i/9fjiR7b7S060xjrrg4akpWQz0FXLl07jfSStC4AtyjbSffp1wDgMrC6nyax6fzHv+7D/UMiV2Yrmi2/1+nhqnsL2cptk/9mrfjjXrbkEEjpdSBb35/dxhI3F7A45pzwGuIGSfZDUpT5rClQx+r1avRT3P3JRsv0HQlW+WGc8VRSyS0k25ZY2aKRbm25TY3HxGCsBKhD2W623LcW+XTD/O5Fq6iHNuAcwgSd+L/rh4JP8A2S3/AObDftkrt9ujwS/cymz365rX2XeA1nta1+Lj+F+nxw9y3unlky6RgsVcoRWfycG8T/JvCfZmwyBXnrf43/EH+GD7i6GMMbHlj1A9x6fHD3NvCia03TVE2srbXRUPIIFxbyP4c3xlrKVIXi4ve9h7W4/fheoR5glaLt31w9mtaVbbvxFm/wCY+mGxIBtsCgWIuSb/ABt6+2BrrwQ4URo5HikWSCVllRg6utrKwNwQep6A+WH2bRwyNFXU0IWGtBkKAf7OUG0kfwDcj+6wwyTu0JZlvySvi/iPL9+HVGe/V8tdlHfndCb8LMB4fgGHhPxX0w14um+NXolaagtPRTEgFwibSL26nafbnpgrxgG5TZYkG99wN/PAEsrkuOvUMbEHpY4WB7xmZzdjYhmHHS3I9PQ+WHpqbvMO7aFj4G5B/Zb1/gfb4YnPY1keZajzVcjycL9fzCSGjpyTwsktTCoY/wB1eWPspxAasWH2SD6gW/HyxZ/0fc0ynKMyzaszXNDlpbLJqaKu2EihWbwTVPu6QGcRr1aZ4h6kUNJYWdxG5WrGaSiqs/ONTwUr592jUDGM6hgGktIxgHfFpyiC08tQAOQaho1h45N6v3xSeqNWS0iPTQzd5PUczngg2a+zjgruALW4ZgB0Xl/rrtPh1NnMk+T0H1amjhjoMuo1cslDRQpsgp1bq21PtMPtu8jE+M4faO7EarMciXtI7SM2GnNKyM3c1k8e6bMSvWKiguGqWHQsCsKffkH2TQjgaw9rOKDAAa8MutZVp0jrtyM1JzK0XZh2bao7Uc7FBlNI0zyiSonqJZVijjiTmWeWV/BFEg5eZyEUerFVNoag7Q9D9juQVWjOy+tXMcxrYXpczz+nV4ZauJxZ6emLWkpaNvvvxU1I+0YoiIzoM319nGrMrn7P+ybJRpzRlO8cla0k4aSrZT+rlr6raO+YH/ZwqoiQn9VGWu5l2hexnSWlK2HMdeVMuY5y9pI8kggSfMCW5EksEoaOjQ34apWSU9Vp+hxJabQx1GymlP5R77OiFFDC7No8/jb6Kk8u0DrXWjxZrT5Y6U1S4hhqZgIKdj0EcN/t24GyMMRxxi4Mh+h/qwQwzarlOVpMODmdRT5NGR6g1rCof/lpjfyJx1BlFPqeCg/SNPkNXpGmlCpLXdy1JO8P/eV9c8VRIoHVIzDHb7MYtjUxZr2H6LqWra7tS0VTVq1CTu8dccwnZgrAqVpIpi6ksSd0lyDY8gHGe/S8zzchaTTZ80r6K8ywRt70jh5/r8qCZT9Ens6gpmev1bpmX6vEZXNNl+c5zcKLmzk0cN/YcfHG4f6KXZ3tjmWXMLSuYlCaLoghYDkLuzQsSOPPqR64lUP0r/oy6KkaufMtTagm+rtTj6hp94k6kq3eVk6BipI5MfNrm244aJ/pC9IVUMeTaC7B9UZ5Pfuou8zOCHeALIpjpKeRvCpKg7idpsScMD9NSYsYAN5+XEpXN0czNxJ3D8UUSb6Heic4njiotQiiaZWK/WtCVlkAYrukejr5QgLAqLrc2JAI5xp6z6AGqJV35HrPQlWxPhikzityqYj/ANHXU1h/4zizm+kx9KvNYmnyT6M2jdKUTeL69q2eojCcW3k1tRFfji4Q8cWtiP5j9I36TEVPKM1+klp/TkCKXkpNDaZikYW8jUNHHGP8XesB74njmt8R/eSRji7kCFC6OzSDuMeeA9ioLJ/o4fpJiI1+X9nsmbUai/e5VqPLKhbf4iVt88ayL6LNBo6eVe1nMNL6TWmQzVBzrW8M88aji31PK4pZmYnhVLIWPAPW2r1prLVerqOTOe1HtY11mOV1F3hOdZ1I71iDyp6NGVX95XKwj9pvsmlNUdodRm1DBpvT9DDk+RUbs9PR03nIeGmlksGmmYcGRug8KBF4xfs1qmtPdZjtOQ8sSSd2rXsNWSBkWL8NgzKkvaLqDstojLpvs5yM1+9gs+b1tI1M1geFp4GllkiBtyzyF7cbV5GILM2597Kym1zuB3H5jhvj19cM6KnVVZyCNpDFrXA+IHPzGHLMLlkVju4NhsF/L8OuNSNlwUrXxVNzrxrRCyqDcgg+K5Pkfh6fjgN0pk2RXRiCLLdQB5jk9PifhgUeMOHkZ0ZdrCwuR7geX5jAMYypPH2r3IAFvI7RwDz53tiRNRRuZTNuJBYC5ubEfHj097YOJNgNi6FrghbjcD5Gx4/ccA3N1VJvCAGUryT5E82+GAKFbOYRsYWVim0t7AevHPp14wiRCWeRy7EOzG7bhwbCwv06D+OB3vA5RZG5U8jcu8fMc/hgEmkCM5Vd4bcTYs3PmAeCffrgLGCUo20E2uX5Bv13Dnn8/nhUI+yRGVQwQsBtY9LHzB/iOcN6nakchjUg22gkeJvLp5f1zhxuDuoFiouQBHsDep4JI6XufLCb81USuqLHETMwUeSjzJuTzbrgQtnpmlqKjPXNHE0suX0hjp1A6zsBFEB7mWQfPHdWstPU+m6qfLaGKBaHLIItNUcyFmWSPL6aOjdyQDsYypUNtawO8MrXvjl/6NGXQw6qotU5jSmohymafUU0RPEq0Cb6eI+z1klOvyxfWd0+aVQy6lq6mYTRU5XvnjsUk7wCVmZGPBckl7lep4uccdpuXtJQ0HD4/NV0OjmXGVPVfxRaLtGkkoezHOysQSbVFTl+m4nQAP3Ekxq6kbx9oLFSRC5JIElj1xynn1bJmme1NWesshIW4+8d1vwIHyx0F225xDTQZBp+KaVkpctrs+kdpGIEle60tMLEmx+rQd5xx+tJxzdJN9ZeSW1u/LSW9btx/DGhoOIhl49V/ABVTST6uoOqflCQS+0ruN7G/ABHpfpb5DBgxY7iSAfew+F/U++CXtybXI5A5+Rwa/JZhdCbE7bg/PrjoVlLInay2W1up23/ACH7jxgCykPcBbHdYHkDytfrb+r4Da4AW1r9AT9r056H4YBmJjUXsADsBPhFzzYeQPOBCMSVXxd0G45W+4A8jocJO17EXK9OhN7YErcja1xbzHzt8b/LzwEgt4lUgADq1j+I4F8CEmZFuWYgegvc/L2+OMwTZwQL7b2v/DGYEIWN28LLzz52PsMCoFttvtenQfH/ACwDqCdijxfsg+duv9dPhgyHwks3BA8W3z/rjjCIRQNzAbT4j0B5v7+WCMwINifbp+GDgnkgji46efpfCbE2Fx1+Av8AwwIRH+z088JNwflhYg7ebWJNvcj88JEC9uevGFQkj8umFaNys229hICh+Y4/O2E2/qxwQEg3HXCEVFEoNDVXloPMxmmjPqrOA2UVbFlP2jBVgNx6hZo5Rb/vBifaV1dU6Smo88ihaabIa2nzmGIc73pn714/g8K1MXuCPXFOdkdWJdRfoUyrGM8gegjZuizv+spz8qiNV/5ziw8rqjFV0lQqtIryL3kITdIQLMUtwCwAYW8z/ixylpYIpnCmu8PX1BW/A/tIwfI9eFFEu2zS1N2d9plbRZd+uy3KsweGicdJ8ucCopG+D0k8Y+WK6zGiNDXy0gc92LtC3qnVTx04tjojtu01+l+zDTGpYE7ybKY5dHV7bbN32XD6xl0h/wDTZZUKoPmaU+mOd62Vquhoao2LRIaSQjzKfZv8VI/DG/ZXh47pw9jiPdZFoaW55+4wTXqvJO0m1r+fw/PAEKLhhuPv5fzv5fPAA24N/EbEm1rYEsA5LEFj1u1vzHn6YuKssjfa1+OB62A45/Hi+B5BYNcHj7PHPofUf5YKjKDZd97g3uBYdPhcHz98Cfs2sNoPrY29LH93l8MCELFD4SikjqTcW9iP6OCox3jcQVseHNl/ywdmDi3eghQOSOB8T19rn92AKMp2srCx5+P/AEtgQskcOqAgrYHoOvuT5nm2CMFv4FCj0BJ59vP3wchtoew4HkOAQf6+HwwVgLgEg3ubX6fyOBIm7xBhf5e98N2Rl+GH+0h9zp58lr8/E4TZN/3Sevl0GFRVOMl1DXZPMWjlLRSLskRgHWRP2WVvC6+x+VuuJnlmf5dURMKHMP0axHMbl5KY+wPLxfBg6/3rYruSMg24B8sEUvG25WKkeYNjinPY45jeyKsRWhzBQ4hXTkuqc30k36SjppaeNvs1tDUNGhPqs8BAB+JGLHyj6Ses5af6nXa7zKuoyPDFnEEGcRIfLw1SyFl9QGBHUHyxzHl2pM0yyQyUtTJGzcFonMbMPe3Df8wONhFqimnbfXUNLI/mxp+7Y/8ANEV/djOfYJGmvp8flXWWphFPX5XS76w0DqCL67qPso7L6/xfrpqfLpMvZ7/eBo5YiPcFLj3w1l099G/O2Zq3srzHLjEN7NkOtJWFvURVlPKePTd88c/RZ/k1txgZT/crW/8Ac4z+/AjUeXQMJKczqw6E1zC3/gjB/PDGxWpn2ud50+U4vgP3AdeSvjKf+ynSUslToHt07YdD1B+y0FNT1MVvdqSqhLD4p8sWfo36SVTkdQM4zj6R2lNbzxwtTtJqvs+r4qruGFmikq6UF3QjgpI0iG/KnHFk2p6VTugpaYN5N3Bla/xlJ/dhtWavzOsjELyu6L9lZXLKvwQWQfhi2IppRdkbUb6fCrl0TMWGnguke0nUf0a+0WpFFRpmGlJl3MkuRQyZ3k1Nc8rClQtPWUyE89yGlQfdA6Y0+j6KTQeZSVPZz9JzQ06VcH1eqy/PqKupKevp/wD4GqoaineCaI/syMdvVSp5xzfNVVNSbzSu4HkTx+HTBVeZDdXKn2NsWYrJ2QoDhsUL57+a6m1J9HnTPaBSx6i7P8x0zpDUM0ljlEepaWryOvqPJMvru9LUsjH7NNV2HNkmPC40FNonVGpNF512Ra007X5X2g9nEVVnWVZdWQNDU1OUse8rqREYAsYmH1yPbcMhqtt/Djn6LMqqAkbkcEWIkRWuPTkXxdXZ39J/VmmqXK8pz90zqjyKVJ8leuZjV5NKh3K9BWi89MAesR7yBhcNGQThXNe0VONMqZobdce6eKqORbXsQb8gjm/z9MEKCxN7EW+Y9PjjosdlfZz23hs+7MNQ0WXaozUGer0fIqUbpVMSZBlRd+4rKckkrT95HUJ9lVYWGKZ1X2f6p0bV1dHneV1ED0Mhiqt8LxvTvf7M0bhZIW9pFW/lfE8cgkAKgexzHUKjoVXcmRuOrEckj+rD4nGSmQv3wLIxsOOAo6AD28v+uAJaOyc+THnp6D87/PANvAO8Gx45H5YcMcUhwwQ3DHd9m4uQBzfzHtgjG9wb3IAv529D6+2DFmcBvK3Q8gHzH9euMVV2sdjEkgKFA5J/r8xhyRL09qVHrbhTHZYzYgd4eht/dF2+O3DYAoQtvcXHX5jC05UFYlAKxggEn7R+8fx4HsBhLbYAWsDewva3xGGtGtK46kIYEggWU3tf39SOmFaWRYnaqCWNONy7je7nhenlfn5YSPia4HX1tf8AEcHCrgJDHHxdj3r8geyj8Ln/AJsDhXu7UNNMUThRbdfyNun4/wAsFJsSGIBax9ATzjFWwuFJ28ceXpf2/lgwYABrgG1ueBbr+OHJqcIClHLb7U8iwgD9lfG3x52Ybk7XF15FuAD8bWwpWj9XTU62G2Pcw9Wc7j+W0fLCS+MHaCf2VCjn16eeGM27U9+obFhG1hwL8fHp+eDRMqQypw19iA24vuuf/ajBCNxt1IABvybfD+uuBIJKvydxI46kgD+eHHFNGGKWYeG4SyA3Xm1h8+p+GEZVBYbHBP2g44PxHvgxXat7Am24EC9/5f0MGBdmHN1AsABcE+tvM4EJalzjMosumy76zMUeVTt3naGHSS3TeBcA9ecJu7Wu0u7cSCxNufO5PQ/H8cHgVVo5CyeI1QFhz/uz08/44byFSCysCQRyp5+f+eGMY1taBPc4upVHVtgsT0PT3/d+GNlNNE2naKEypup66qG0HkI6RMD7cq34HGoEm5iTtJY7rWAUn1sMP4pQMqkUbSVq4ZACOCGjdT+YGCQfadh/HulZrG78+ybKVXcAoAIJFuh+Plb3GM3sVMhJa3h3Hrf0Pr8/xxhv9pFI5J4Nj0scECLvIuLA2uBe49R0NsSKNbKeJKjT9HWJdDT1c1I+7iysBKnyv3gwyZmAQhySpsD5gevvza1/W2HNG7fUa6k27gwhql8V/wDZsQSPbbIfwwzY+Ii/xAsbfL0xFHhUb/XH3T3mtDu/CC46NcgckkFep8r/AMsP4x3+TyoWN6GoEouORHL4W/B1U/8ANjXjzFubgiwA5Hth5QSBaowFl7usjamJU8EP0NuoIYKbe2FkGFRqxQw40OtMiCeTYAiwsCBf+vlhSPai7txuGsFAI4PU7unXythNVYkgqFJuG3cC46rjCwAsdqkDp5kfwOJExP8ALwal5MvBKvUWaBr9J1uV5H7QLL8xhlYBeJLk8gjgg+/ocYZJAwcXSRCvjtYqwsQQevoecOc0eOWYV8QRFq17wqAPA97Oo/5gfkRiP7X+Pr16J+bfBNwy8lha7chQLA/Pp68YMAoFjexBIte/ythuu02bwC/luAsfXCp8LLEzKSotxYe9vQ/HzxImJ3mI+swJmkZG+Q91VAeUoF91v76i/wAQ2GVwF5HPmOt/W2HFDUKkrw1BK09QBFIx52m91b/lNj8L43ujOzXW/aRnMWntD6azDN691LtDSxb+7QdZGNwscYtzI5VAOrYjb3O6cgnnvYhRoUc+Y1C00AeSRz1VSxt6gD93rYYuXN+zvVU1NR/R97PtOzZjndIwzjWEsO0R01ZsIjp55iRHFFSxsQ7Oyr38s3PhXGyodPdnvYcn6RrNXZZq3WcLBHpMuYvk+VWNyTmKMplqlYCxpldEINmY2Ig/aL286m1TRHS+VfUcm00snfDJcmpjR0Ekv/FlUs0tVJ595UO7XJtbGa97rVPdYO63bkTt301DWcQcFbDBDGC44u9PyrV072Iad7Lsng1JUTaM1fqB3K09VnedU1LpvL2B+2I5HE+ZuCPKNae//GHGI9rXT1BrSrqNTdqX0mcqzrMNqxiDJcprsw2ov2YkJjgp40UcKiEIo6AYoSozjNKzmWosT+woU/iBc/M4ZyNPJzK7t/iN/wB+LIs8rhV7he20+VEZGDBoNF132Q9ofYlo6kjlWCiOYUhZaOXVMda8VOxFjNBl+WpZpCCf1klQX9NuNpmna9kfdLkFP9KfPNM5VUPxlGiOzgZMhJ6kyvPE7k35aR2J8zjjSGtq6YBY5nAB+yTcfgeMbKm1VmVMoQSHZ5p9pD/yPuX8sVho5sVSwAk7dvmCpjae0IvEjrcum8u0Z9GzO6qXMdQau7TNQVJJIkzXMqOilc/tXCVD/g2JllFN9EPKKZY6vs0kzaojIO/MtV5hUoy/tbIhCDbzQgG1yL2tjkCHVdEzbqnKqMn1RJIDf1/VuB+WNiurMoPieKoHsMyl/wDcoyfzxTlgtjsLzvINAVqN9mGNB5krsGTtB7FNPVK1OkOzDs+ywqlg0WRQ1Ui+jLJUNKx9weo5HPGNPn/0mc2ooGiy3UNdSU/Raam200S3HQJCI0kU+vDA2HnjlZ9X5QsZRabef79RNJ+4IMM5Ne1FOCMqhSlJ+/DCsbf+LxP+DDFf+zJpj3y4/wDscPf0Un1cbPtDR4BXRnmvM51FPFm9fl9Hlkak7a2tYqZgfTcDNL7bQT79MQvUXaVlOVszZZF+ksxv4auuiBSAjoYaYllUjyeYuw6hAecVbWZ3mlfJJNPVSs8n233ksw92JLH8cN44C3JXF6HQzB/FOGwe+s+h2KrJpBx+zPb1+qfZvnub6irJavM62eeSZt8kkshd3PqzNyx+PA8gMJwU4ABJA6364PT043KCofi4XkbvYEef+eFxwhUR7ibm/Nx7WxtMY2Nt1ooFmueXmrkYrEoXxM523BUbQvzPP7sA72sDMhG0AFfTr09b9cG2CUBEQm4FgLXJt6D8gL/M4FLo7+JULEAngsPS3mPe1r4ckRd0so/2oFyXVF+yWNgQo556Gx684B5OBGrRMSbAqTYEHofI+x+XTCribeN7MWuVLXtY9OT+N78jCbCIEEAXXqpNwR7EdR+BwJUmrNI6BgHIawUrwelhb09umFCu0gkKF3EEgra/z4t68/jjJIluFRvtmwABN/Ycdb+344T3IWXyIJLX4HHQj0/gfjgQju4UL3R2lTuDX6npcegHS3njEuy7EEcdgxL322A6gbuB8PM8YMFL3lVEAjt1Kra/HCnr79fU4wdC4I5YeFTyPgP2fT3GBCGQosYO8pfgqhupX+8OoN+fQnCBLtDK5uTKy0636kDxN/AYVqZTBeRXG5BwYm4J9/5jDrJIFbN6ZJkLQ5ZE1XOLddvjI+Z2r88Mkf2bC7YnMbfcG7V0x2MaYhyzQOYV9QzJJWV1DksVkNmipP8AX60lhyimd6GItYgEc8XxNIqDMM5zKhoKKOGKsnkTL6KYmQSxzynuxtZFYK7M24obRsB1YYPl+XS6RyPL9HVEscdTp+ghizFpBtKZlPeqrSODvPfTxw8Am9N7YYR6hm0TpXOtbxmQVOnsvmaiZJJLNXTKKWkBQ7TdZZ2ksRYdyptcEngJi+WYjM1oPE4c11EdI4r26qoft31kmrdW6n1BTyK9JV5h+jcsKgBfqFDGKSm2heACqlrDi4OKoQCMAXW3pa/GN7qsRRVNLlEBumXQrT/4mQWZvnI0h+Qxp1cBLKbG/XdwR7f5nHa2CMRwimv0yHJc7aX3pKbOig5DhWboo8vL5c/15YG7MbsOrelyPYe+CEKXuVZWPQG5ufyIB+eFCh5RtwK+HceQLH+hcYuqusUK91Jvu6KSP8h+GBIsCxYXkWx5uT5EEeXl1+WCs7AgEqxY2JHPPuOh+XOA6kxqUZuSCBx05Av+OEQgDBWuy3A4IuQQfY3/AMsEJIVQbEXLcMDtv1uOo6YWNtgPIDsSNoA4sOfx8v3YRcci17i4t1IP9f8AXCoSezyIAYeYPkfLGYOhCi4Bv0O4C34euMwISTXsbXW/HB6n3wA4PC89TYdP5fuwo/2gWKg26nj5DBQQpDXHsAtiP4fjwcCEXgkGwuBzbpf+f5YJYE8H3BJ5/lhQgCxBCkHjmwwUqoazALa/B4sfjhEJM7txsePQH8vfCbLZuoJOFbE38K2PPoMJsCTcKeeb2wqElJzxzzzfCZwu/AIuo5HAGETYC2BC3Om6qeKtgNNP3M6SDuZP2JLho2+Tqv44vyWqXM639MZeCsFd3ObU4ZgFjeS0tvUbJhLGfYY5wo5Ak4BbaH8O79k+R+RscXr2f5lS5tpmZJNoqMtmDsL9IKhj19kqFdfYTrjA0vHdIl8uP5oFraPfeBZ1grN0lS1mqsh1B2exRieqzvLfruUQkc/pbKu8q6WMX5JlpGr6X3HdD0xzBVZatJmVdk0LbqepjWro3/aQrvjI+KEg+4x0LpzPszyDNMvzbIGkjzjLK6Ctyws3CV0DCeBWB+5IyPER5Fz5G2In9I7SWV5Lq6TVui6cLp6paHP8mW1iMpzC88MXxgm+s0zDyMRGHaOmugA+HHLngktkV4k7ceGfLFUdYshJ2gqbAHrb+Q6YyTaCRGT7A/u/h+GNlndPFTVYlpgWpqkd7Ew9P6t+eNdIwdFQJyPCPcenv7Y3FlIoAuCgtfp5kfPEki7NNfTRh00fnW2QBlP6OqDceRFkIOI1UtI4ZjfceWO21z/QxIJe0TMaoxzVdZnPfCKONzDmssaMVULu2+RIAv74gndM2nZCvXiFJGGH7zROZOy7tFjCn+wuoSCPFtyupIv6/Y6YyPsz14GVm0dnYF/LLKkH/wCl8YbwdodTFJvaq1A4sRt/Tky/mMKQ9pOYxDauYakF/wBnPZxfFftbX/R1/iU3Zwf1evwlansz7QIwinROfSAkm8eU1HHxHd8e2CL2b9ol1K9n+p3FtvGT1NwPbwYK/aVmbgD9I6jAA/8At9Ob89T+7CsPahm0RFs01PbzA1BOL/hhva20fyDr/wCkdnB/V1wRD2a9oK3/APeA1Na5A/8AI1SAff7HGE37P9fpw+htRi/FzlNQL/8AsOuHMfarnSXvnGqW5uP/AHoqgWHpxgx7V83YgvmWqDtO4X1HUcH1wCa2A4xjl8o7ODU/1+Fqn0HrRTZtH54pB5vlk/8A7phB9EavF76VzkfHLpv/AHXG8PapmLyb3r9Usb831HUc/PCi9rGYX/WVeqGHoNSVIw/t7Trj9PlJ2UP9fr8KLvo7VCmx05mq/Ghm/wDdcJnSepRycgzIf/OUv/uuJVJ2pTOf9tqgf/vJOf4Yz/tRlPWTUx4A41HOOnywdvaf+2eI+UvZQ/1jn8KJHTOoFNjkuYAj1pJP/dcGTS2opG2pkWYsfQUcp/8AccSj/tKTez21IC5BP/vRTc/E25wen7VaunnWaGo1IjL0I1HUgj5jC9vaKfwzxHyk7KL+sc/hRldG6pJ405mpt6UM3/uuFE0Zqn/7mc2P/wA4Tf8AuuJknbXnalrZ1qobv/wmrOPwODr2256ttmfapSwtcalrv/dsILRaNcZ5fKXsYv6xz+FEE0Zqu/Ol82B/+UJ/x+xgzaK1UFuNMZuD/wDKE3/umJvH27agDK76l1aWUAFv7TV9z/7PBn7ftTKwaPU+rbg9DqjMB/7nhRaZtcR5fKaYI/6xz+FAH0Pq2/Gl83F/Wgn/APdMJPovVkQO/TWbL8aGYf8AuGLKP0kdaFgE1Pq5FA+yurcxP73xlR9I/WUyCN9QapawsS+qcxJf3P6y18H1Mx/6R/y/KXsY/wDuDn8KtafK9UZbdkyqvRCbuj0jlGt6qVscXVpP6RWoKuhg012sae/tjlFLEIKf9Kd+tbRR9NtLmKf6zAB5RuZofWO2Isvbznci7Kiuz1x6NqKvIP4vhlU9r8bu8kuSvVtIdztNmdSzE2te7c9MQySTvNWxEHbUD3U0YjaKOeCNlCp/qLsy7KtaUE+f9mmsRSTIjTS5RnsS0VcpAuVjlQGjrT1tsaCU/wDDvxipq/Qeo6aXu4MuqKyMgMssFLPb4WaMEH1BHGJHlnbnmeTxSDIKCuyzvTuk+qZ7Vw7yOhYIQCffG8pvpUdo0MJjGcakB/bbVeZEj4DvLYc2e1sGMZPm35THRWdxwfTyPwq2OktTLcDT2Yg+1JNc/wDsMLx6S1QoMiZBmIbnZ/qcxtfz+x5Dpiy0+l92lxRLA2d6idVYOC2q8yuSOnIk6e2G0/0sO0OeQyvn+qASb+HV+Zgf/TcOFptBwMJ4t+U3sIhiJBwPwq9TRmq2Fhp/MD7fUp7W/wDhfl5HBTonVi8Lp7Mbe9FPz8f1eLJj+lbrsKb6h1bf21jmv/5XCMv0q9evIsi5/qslT56wzQ3/ABlw76mc4dieLflN7CMf9Qc/hV9HobVzg/8AvO5kVP2v9TmPH/g64K2ktYCVmbSucXY3uMumI+XgxZEH0r9fxEuupNYKT+zrLMx/9UwLfS07Re8JGp9Zc+ba0zM//VMJ9RaCf4R4j5S9jHl2g5/CrRtI6qT7emM55BuWy+f19dmCHSep3G3+z2ac+EkUE/Q9b+DFpN9LftIKi2rNY3HS+ss0/wDd8YPpcdpBFn1Zq4j/AObDM/8A3fC/U2j/ALR4j5SdhF/3Bz+FWD6V1VJO8h07mvL34y+e9h0+56YIukdUlGA03m1+lvqM4JH/AIMW9TfTG7RqcBRqbVBCrt51bmnPuf1vXBpPpodpqmwz3Ubr5/8AvW5qp/KXCC02jLsTxHylMEWZkHA/Cp5tJ6nADnTOZq3nehnX/wBw4OFJdNanYIX07mYABG00M9h4ugJT54vGj+nV2n0qon1zN5YkW3dS6ozVgfcnvb4Tpvpz9pdFHNDBV52Y5nLsJdW5q+2/kpMlwPbDfqrSf+if8QTuwhH/AFRwKo8ad1AGAOTV6sd1waSb2t9zCb5DnwuDk1dYcXFJMCPa+3nHQNJ9Pztao4kiizTNT3dyrSajzJm+Z7zn542+Wf6SbtzyiR5aHMFaRxa9ZWVlWB8FeSwOGPtdsA7sFf8A6CBZ4DnLyK5wbIMzjontS1RfvUkCPSzAsNlmF9vUHCDZFnMylny2uIPTdTyEj/2POOy2/wBLx299wIH0lohyo/2hoagE+/8AtMMz/pY+3mQTL+htMgzEEbIahe6AH3RvNsR/W20f/wCsf8TflAs8Ls5QPI/C4+On88dgv6GzFtx6mklIP4Jx8sOBpvPwk0f6GzAtIEI/1Ob7refg9CecdT1n+lC7b6uOWOemoAsylSYa2thZQfNWVxtPuMMYf9Jb2ygbO6VlAtd86zJj+PeXwv1ltLf/AM5/xBP+ns7ThNyK5lGmtQbeMkzNtp5tQTWtbrfZe9/lgjaZ1M1iNO5qRYgWoJ+T/wCDHQ9T/pDe1iSd5SKol+oGpM0A/DvbYZx/Tv7SJu8+sTZx+sUqdmq80Fri3H6w2PpiQWu167Of8TUz6eD/ALvIqjaXTWqYWZjpvNSGjdGX9Hz8hlIP3OvQ/LCf9k9XPHddK5ybeQy2fr8dmLul+nP2mIiQUtfnqoihQG1ZmjGw9byc41yfTO7TO9eY59qdGkBB2avzQCx9u8wotNprXsDxHylMMNKdryKqT+xOsrc6RzsehOXVA/8AcMJto/V6MbaVzncDcH9HT3Hn+xi21+lzrCMu7VmoZS6lSZNY5wbX8x+u6++Gq/Su1aXVpMy1OwQgqP7aZxwR/wCtxILTOc4Txb8phgirhIOB+FXM+ldVSyyVA0vnKhyWF8sqLDdyR9j1JwyfS2rEP/1t5uFBsD9Qn/8AdMWzV/S61XM7SiXU/eOxd2Gts3FyevHeYyl+l1rakqxmFLXajp5lFlK6wzZto6felIvhBaZwP4J4t/5IMERP8QcD8KpxozVzEW0rnJLLzty2c258/BgTo/WfdtGdJZ4VVt6n9Gzi1+G+77A/LFpv9L3Xwh7iLN9UhA26zayzci97/wDGweT6ZnaXJIjnMs/TaLbY9XZsoI9x3xwG0znKE8W/KUQxD/qDgfhVW2iNXqfFp3MV3ci9JML/APsMPsp7OdQVdbDBmNLJlkUrBBLU08oFyf8ABa3uSAOuLNpvppdpkAdfrucSB23lp9T5rKw+BM3A9hhzL9NvtJmi7iSuzrbe/g1RmgsfnLxiN9qteTYTxb8pzYIMzKOBSdFpDsM7OFFXrPN811vXBbrlmQRT0tLf0nzCWIMB6rTRXP8AxBjR647cdfawyWTRGjdMLprSF7jT+QZdJSUEhHRp1XdLWP8A36mST2Awhn30mdS6lOzPhm1dD/w6jUuYSqf/ABOcafK+2OjyuZqihyTMKaRurU+oq+L/ANq4xBGbQDWWNzvNvoCFI4Q3aMe0eTvgqGTab1jWTLJVZHm8pAsAaKbwj0A22A9hgf7E6pLcaYzgj1GXzn/3DFkn6SOeRACCbUikef8Aa7Mx/wDVMK0n0odXQlWbMdVXS9imss0Um58z3nl5fxxcFolaMIXf5f8AkoDCwnGQf5vhV9F2favaPeNNZiAfI0cykfIpgkug9VqN39ncxbm1lpJj+5MWen0pNRlT39fqqVzzdtaZuP3SjDQ/Sa1TyFzfVaA+Q1hmv8ZcILXMf+i7/L/yR9Oz/uN/zf8AFVpJofVW2x0xnA9P/J8//umETobVn/3LZwPW+Xz/APumLMP0k9UD7Gc6sF//AMMMz/i+Em+kbqtm3HPdVhut/wC1uZH/ANzwOtU4HdhPFv8AyQII9cg5/Crf+wmr7A/2Yzjn/wC90/8A7pgp0LrAddMZv/8Aw+f/AN0xY030idUShmOdaq7xhbedWZkSPgN9sJL9ILVJ4kz3VbLa1v7VZhz/AOzwC02gtqYjXy+UphiB/iDn8KvxoPWJFxpfOCPbLp//AHTC8XZ7rYkE6Pz1lP7OWT/+6Ynv/wAURqRAoGZasbb0/wDewzEf+54Tk+kVqpxt/SWrbXJB/tlmN/h9vphBabT/ANo8W/8AJHYw/wDcHA/Cho7PNccW0RqHrwBlVRz/AOwxsouyzXLou7TOaR7lDFXyysBX2P6nyxvW+kLqllsMx1aB/wDNlmP/ALvhWH6RWqEiWNa7VKlWLFxrHMgz8cKf1lrDrxb44X6i0U/hHi3/AJJOyi/7g5/C0p7J9bOFVMmrmU+uXVnX/wCEYWk7INbQkXo5txs3/mNde/l/77+WN7D9JnV0M4lOYanbaCFB1jmnHuLS4Rl+klrHaqU2c6pgiX7iawzQgfjLgFptJziPFvygwxD/AKg4H4Wq/wCx/W5Uj9HVIZgCSMvrvXr/ALD1wY9j+tV4GVVKlh1/R9bz6/8Avv8AnjcR/Sa1sIRTfpTUYjVzLxqrMQ5cqFuX7zd0HAvYXJtfCR+klrIRvH+ltVMHN+dY5px7D9bhfqLRriPFvygwxf8AcHA/C1MXZPr4zv3OR5gXgKlZFy6t8Xup7jm3S5+HOFG7He0Bjvk09X2vyTl1YLH/AOEdcbD/AOKT1qsYhGZ6jZA27a+rMza597y4CX6SWsZnVmrM9G1DGAuqcyAAPU/7W+7yvhv1FprhEaeLf+SOxhp/EHA/C1n/AGP6/LGRcmrlIPB+pVgP/wBJvhnmfZlrbKYPrtTp6ueCPxO60s6KgHqXjS34/DG4H0htXpIsi12oPBuGw6ozEqQfI/rL2+eNXnna9qfVOTVenairzH6rmEsUlWarNqqrDJG25VAlchRuseOfCPTDmzWlzgDHQbajzyJSGKICt/kfhRa5BjjkQhWN12sASfib8fHpgQyyqVRLHqCHuLDyAsPmcGlAD7UKgnlmbhefX1v6/hgE7uMK6yeHm+25PTyPmPj874vKsi7EknjV3ZoolMzgtewHl7E8DFu/Rn0nSZ3rnL831DTiTKaOWTUGbByAr0FARJ3RJ4/XVJghH+PFT01PPWdzR0yXnzOUIg9I1Nr/ADPPyx1P2UZTFp3s+SRJqKmbWcohikqWtsyTL5CqvcqVAqK8s3iIDLSWB5xl6UnMcd1vj8c8fJXbFFffU+Hzy9VKTXZnNV1UmcyulfMZaupZqgEfWHcvP+sW239Y7EEi3kSODiE9u2pP0bp7INNi8kmY1D6mq0L3ZqamBp6FSbkWedqhxyQRs6ixxMsupBmOc0+V5RUOJaiRaelkddi3JC977KoLMQbMoS/iU3xz32y63o9WaizvUmUosdBWVKUOToq7VXLKNRBTWWwtvYd4R6hsczo+EyzCoy9T0T4hbNqkuR0HVOqKtamYVNXPPI+83sGNzu28X+ZuefXCbPuuNp49x1/rjCKeBCigWtYm/It/HC0SkXO29/U8X9h5/DHcNAaAAubJqaoftdRYMLbQeB7cn+OChCvh2CxPRbixHt5H44OAtvFbYeGO4IT7+h+GANinKgm/IuePQcYckQMbCwAJPle54+PX5YKzN9pQLEAcci3mPXBwv2VYnjhfDcgdbWHQ3PXocA7bmG9t4AsR95QPIjywISZWQ8IVfebE8i/8j8cCeQLEMtwLg+f77emDStu8gB1K7iR+78j5YDaC3hBFvOw/ojAhFJDkNdieRuYfl7/HGYLvBvdit/Mn8reXyxmEKEDEXJseSeWNr+39fDBQbqbeXnfp64AsGLMOtuTe37+vwwKjcSFXnpb+v44VCA2AIYhTxfzHwwHiU9GBPFrWPyxipcjb7cWPI9cCVsTa3PBsevzwIQH7O6wPvf1/6YQNix6gr1I/l/LDjg2Ja1/Mj+r/ABwjxflQPLjywiEm97Dp539cJHgWws5spAvxYXB6YSPS/l8cKhExYHZdqClyrUMD5jJbL69Gocw56QS2R3+KN3Uo/wABxX5th5l0u2Qxkbr8hfXixHzUkYrWuAWiFzHa1PZ5DFIHBdFVW+jzBqJiY6qnlMcgmBAMsZIJDLx1BIPXm+Jlm1JTa87LKuicotVpdGzRDYHdkWYTqtWBb7tLmPdT+0dZIemK5yLMJtQ6foM7JmmmYHL6uSMbmFVCg7uRh1tJD3bXHO6N/fEi0JqebSOf0Oa1OX/pDL6Pv0r8sZbNVZdURtFmFKtuqSwPIQObSIhFsc1ZjcfSQ01H58jiNy2phVtWeI+OGBVEZjl9VRpXaYroStZlUrSRA9TGLhl+WNK224KgAADob/O5xbnbpoSr0FquWemrP0j+i544hWjlcxoZYlloa0eqz0zxk/3t48sVfmtLHT1Cmnt3EwE0B/7tubfI8Y6ezyGRnezGB8esVhzMDH93I4jwTFgviC2IHnYi/wCOETFGDaxI89ovbDkruu+0AKQpJPAJ8h+BwQqSFJB/Z63tb28vI4nUKRES3uFuAfMdfjgRCOV2j0+OHLnvdrbw7bfEdp5IPn6m1uRgiqSSLi173tzfCBKk+5Ufs2HiLDnjz/PywbuBby5F7jzGFQoX7IsTYnobfh7eWA2gAi1uenmOPTCpEl3JvYqLMBYgfa98HNKtr9LC/Pn88OKOaniraZqmm76ESbpo+m9ehA9Dbz9beWOkx2TdghzCXLxS9oqyIrSBXbKBuRTZiD0NiebfHpitNaorOQJDSqmjgkmFWBcwtTK1tosQfLGLBbyU/Dmx9MdKVfZH2Gs4Wnpu0Mkx97YzZSnguBu4jPHI/EYTg7IOxGpdKUZxq/KZmsVkzHK6avgA5+2aKRKlBx9pIpCOu09MRt0hZnmgcnmyTtFS1c4mmAP2bE+uANOu0kL0tfw9Pni0+1bsaq+zpqbMaKshrsqrYknhnhqFqYpIXYqk8M6gLPAzq6B9qOkitHKiOBeuNg2EA82t/Xri5VVjVMu4jPO03PQAcfjgDTxjxWtfyxb+geyLKqzL6fUPaDnVRlVDWxieioaOjFVmFZCSQJhGzpHDCbELLM432OxHAJxYbdm/YVHHEsejNc1Sut0mlzzLod9/QJRuov8A4j6Xvim/SEDHXaqyyyyuF6i5g+royXCgqBc+t/hgzUygWsLA2x0zB2edhFTI1HUaZ13QXH+0jrMprivXnupIadnsR0Eini18RntJ+j3DkmRJq3Quex57kzTCkWohhkhKVBQuKSpp5SZaKpKhmjR2eKYK3cylgVxLDaYpjRhxTHwSRCrhgqMMK26AcXH8cAIEPVfI2NvPC0oIjBZPgfU/D+OLi7PNKdmmb6Aoc41NkmfTV8lZU0hfLUodjrH3W0nv2DFz3oBtxxx54dLMyAAvOaayN0ho1Uo1PGo6fD1xgp1+8VBtwCbC3x9cWz21aP0dpSHTcujaKujpM5yuLM5DmUUSVSO81RC0R7ljGUBp9wNt125PlisQVDMASGHS3B9/hh7HiRt5qa5pYaFNkgQEi3I+zfzxsMp09mOpMwpsiyPK5q7MK2VYaeCBC0ksjGyqB/XFyeAcMpAtyDcjr6nEv7HcnyjUuslyjPpM0p6FoaiolqMrlWOsjEULvaNn8ADdGuDx0thk7+zjL9iWNt94btV5aY7KtB9n9BLl2v8Aswy3N4MqUzZxqTMqrMaaKJrDesSwuiuit+riABeZ7nhSLUr2jas0dqjMCmiezjK9J5TSs5gip5p56mZSQN08s0jkmwB2rZVuevXFr6p7K+yuDRep81oM17Q6itybK0r6WPM8wpJKSR3qIYFLd0N1h3wcEddtrjHOUe4Atc3RiNw8rcYpWCslZXSX9WFQOCs2sBlGBl3mT5oPqwkIS6jmxLGw/HywkKVb7bWJ6XHX2+OLK7I+y2r7R62pqKqvpMpyPKzGcyzSt3iCnDk7I7RgySyvYhIYlMj7WPhVWcXZB2LdgTERUtHrrOjEA0k1OuW5csik7d6QyCpl27vD4nuDYEAkYsS22GB115xUUdmklFWBcmCmRB4rL5c+XvjPqoYfZI+WOsazsc7BI4klXJ9cpFIbLIuZ5VLz04D0qAkeYLAi3NsMdS/R00NmWlMw1D2fZ+0j5VEamqiqaRqKspIQQplqaYPJFLTgkB6qlkPdbgZYQl2CQ2+Cc3WOxSyWSaIVcFy4sEZuuy1xw3of5YUSkTYWKkkeduvsMO6qkmo62Wnq43imhdo5EIB2sDYjjjr6cYufsR0H2eaw0dXZhrGnzIVNJmC0cLUFPSNvRoe8vIZ3UlrnaNvl15xNNOyztvyGgUUcT5TdbmqMNOnRdpI+V8FWkQm7EAXsWI4B9/b92Ll7b+z3RmjsoyCq0dS1vd5w1U07ZhDTrNE9PUCHbG1O7I0bBwxvzcW4xUixsHMSqS1yALefy4OHRStmbfYahI9jo3XXZpoKZL7Qo3fs+d/44OaO4JK9DwbflgZ/AlkBJPI98da5V2A9j1RlMUWbwasgziOKJa6Gkgy36vHUFFMioZJN4VXYoN3iupviG0WyGyU7Z1Kp8UEk/wDDFaLktKMfaKptHF2U2wVqVUaxUBr9MXj23dmmk9IaRy/O9E0uZrB+kZ8urmzSmpo6hJhCk0DI9M7I8UsZl+14g0LYpVjdBci/pbkj+HzxNFKydgew1BTJI3ROLXZpAwiwsQOtrjy+P88B3Ck3t5XP88KEFTcFRb7PNr38ucDtZvCbXv5W4+eJExAIB3fCOwvuICgjp19cJd0C9lAHHUYXIVrkXAFvEfIeV7f0cT7sMyLS+qtfLp3V1PVS5ZUUlVOWooYXqVeGB5VEYmZU8Wwg3PmMMkkbEwvdkE9jS8hrcyq7amfaUKA/BcYKW4/Vod3mLdMdq5P9H36O2dJGaebWqPKneKr0OUqxF7dDIeQeCPI8eYurmf0ZewOnlSJ63Wo7wLsIy7KeSwJUA7/Pa1vXabXIIxQdpextNC/HzVoaPtJxDfRcTGkDMFHW/mMG+qBVJXaHAJA639fyx29k30Y+wOqzbLcrjqdasMwqoqZZHy/KNoLuq7rb7mwYNYdR0xx3qh6aHMcwFLSQQJFPMkccW4IgViosCSR0v1PJPliezW6C1kiF1aKKazywU7QUqtCkCNyAePu+fy9sH+rc2K3PXgXt88dK9lHZR2Q6n7Psj1Fq+j1EmYZj9ZE5yuly400SRT9yJGNTIHFyVLNYqCw6DFhzfRs7B6WMh4NbiRiQqCkygsx9B+tFz7DnEMmlbJE6699DiOCkZYLQ8XmtwXEzUwtbZcWHlYg4KKUICy829v547Ak+jr2K1DMaX+2u4E2VqbKh06j/AGh59sQvtc7LOzPR3Z/JnelI87etGY0tFIc1pqNEEU8UsgkjanbcHBhtZrqQ3rhYtKWSZwYx9SfFI+xTxC89uC5zFMoNmsD0vb9/88OfqYUbWiuPPi2N3kMNHFn+WpW0UU9PNVRxTRsu8GNnVWNiRcgMSLnqBe+OsM/+j92BZXnlTkjQ65D0800cchoMpVJ1jcoWT9Yb9LlftAEEi3OJp7bBZiGyuoSmRWaWcExitFxa1ILkbbWNgOn/AF+OBFCTxt6c8C9vbHZZ7AOwMKjtDrTm53fVMosQou33z0HJ9BzhCr7A+wiuX6nl+a5nltSbqozbJEeMt5Az5bMZor+pp5BweDiJmk7K83WvUjrDaGipauOzSopF+PS+CNEpuFQEgcHof8x7Ysjtf7I867M8yRJ4TJl9XEKmkqI6iOoiqIGYqJoaiP8AV1ERYbd4Csr+CREbg15AVSaIuFa8iCzLccnzB4xeBBFQqlCDQpuKa4KlRzY8jCq0Q27hHceZA6Y6k7P+x/sa1DpLIc7z/L9WpW5nl8VTUfo2kylqVZXkkRY1M7K6BjHYF/DuIBYXGJJP2D9gzIhgo9cLe6kGhycFWBsVP6wWYHqDYg3v0xROkrK03S/EK0LFOQCGrjX6ugYB0Khrjpe/ng5owBu2lQcdgQ/R47EJg8iw65Z4huaM0uUIwUmwaxflb8bhxfg4r3t57OdBaJ0nkGY6Uoc0hlzKvrqapGaQUolT6ukDIY3pjtKv9YO4MLgoLHCxaSss0gjjfUnUkfY54233jBUCkCMxAtcfK49v5YOaWI/ZW5HBt0PwxMOy/Isr1F2hZFpzMYRPQ5tW09FU93EjSosjgEx954A46Atx646PpOw7sNrYlqEynVqrxujC5SJUJFwGHqRyLGx8rkEB89ugszrsrqFNis0swvMFQuPHpF+8LfL+ucYaVbFmQgdRx+X+eOu67sI7CW/WU8euHHdGayx5QPCPtGxHQefp52w1k7AOxLN6WWlynU2b5LWheP0/l8MVPuIuA1ZQySCH4y07IPvFRciNmk7I80DwnusNoYKlq5N+rKy8IAPxv/n/AEMAsEZAvz0tx1xMe0LQmcdm+pq3S+d0ksFTSsFdJQu6zKGW5UsjAqysrozI6sroxU4ib8KTYhriwv8AjfzxeBriFVRHpADslADD08v6/HAfV0DEbLhTbkHC9P3gqYVpqQTPvDLCw3CQ3FksOoY2HrzjrYdhvY1BPLR5rlGoBWUrGGc00eW9y86gCURB5N2wSb1UN4iFucVrRbIbKQJTSqmis8s9ezFVyMkER2gAXbgA+uMaCMruJW3v+63ri2+3vs50/omr0/mGkYKuDKs0p6hCK2GFJoqyCUCRGaBjG67HhkRhY7ZLHkXxVcgRNy2Yckjm9wfb+IOJYZWTsEjDUFMex0bix2YTNKdWbaAAT0BHX+RwpNSKtgdpb28vn542OQ1FLDnVJ9dpoaunZgssRS4KFgptyPEASRc2uBe+OpqPsL7GJtTR5Scq1g8bV31MSS0+VMou7IrsFkDFCVsSOcRT2yGymkrqVT4oJJv4YquRlpUI6AC3l5H92MEKEWboOBYXth9UWkqZmSnSJJJW2RxXsgLGyrck2HTnn3wkTdu6LKGJtuaQXX5jgfHFkKFNTTAWJswHAKr1+fngwpVYgAKbi4B4v8D5YX2E7iQoCgX8NreR4HPPnbGMtyzuyEX8Q3X+HHUe3GFQknoxGLun2uv8vf44yOLudxjK8fbF7ED15FvwwrGqk8KwKoSwAuCBz8v3YTZAy356Ajjr7g+QGESpeMEI8iIG2i7brdPUj+I6YIYvrciQQyWadhHa1iB5k+VgLm+AWwXiNCx8rXt8PU+/kMK5fFNUk/VFLVNW31OlUf3vtt+HF/c4CQBUoAqaBTLs40hmettQ0eVaeXbX5/XRafyctwsQc2knJ8lRNzMfLk46fzaTJJK5k07IDk1LDFlGUbSylctpkEMAa4CkyANMw5Bac8gm+IH2O5ZQ6d0pmmt0iJ8E2jdObeHbcivnFehPQrC6U6t0BqvUYl9ModUhYBIFO1LMQ4A52hbDkLzsF7jlSQGTHIaUnL3UJxOJ8NQ4Y+JK6CwxBuIyGHz8JpqHNX0voXOc5oJ1WvzBRpjKLABhVVaETSXuSRDSCSxubGeO1sco6sq4JcyShoiPqtGgihXy2L4VPz5b/nxcfbvqFoNSxaQpZgKXRcE0NUUIKtmtT+srGFuLxqqQj07rjrigmkeed55VALtuIv09h8Bx8saGhbMWtEjvHjlyx3VKp6QmvOLR1T8pZAu1vECApP2SbkdB6/PCm0AXKA+vIYfC4/dhJNykkEA+X9f11wpuH+7cLbp1/ePTHQLLRiWLkKfEQTa9unmPI/8AXGMSx70FfexszX8ueAMFIXbaVyQ1/Lr8Sff/AK4OVYm5J22AJ8/n8cCEBVghO0sB4iLXv/X5/HADmwY7VDcNt3AfLr+GB8IVfGDz0JIAP7vwxgUnbGUYkdN1lv789Pl+/CISXiCiyMDfg8Eeth7+uCyCQKWI8NrGwtfz/H+WD7hYXU/Z8gOl+LHz+fTGEozEsygj3/Dr8sKhIiSxJVA1h5njGYyUHfwxY2ubm9j8f5cYzCFCTcKWBF7Hkgfw/r1weNWVtpuOlyTx7298CRyeRY8iw/I/PAAre4CqALAsOl/M/uwqEZvFKXYFiSdwvYk4EkWJve9hcGx59/4HAX4+0AP2jc2Hv7YLdjYdeLWvxgQha6jawAsehvfCDWBPTpyCOMKGQizKennbCZ4NubX/AB9MCEQqqm1zYdB1OEtt1t1OFWDKl7cHz9R6YTYcdRxgQkyOcCjtG6yIbFSCPjjOefxwGBCtbsfzugOaz6VzV0TL9RRJTo8jELTVgYtSSkjoBKWiY/sSnE3jjniUzrFU0tXSzIRTzEErIrEOm4em1lN7G9uuOfsvqCh2m5C3Nh1K/eA97cj3GOiKTN49U6do9YKwaqlIoM2KEXSuRAVmA9J4gsn+NJRjmtIwdhN2gydn4j5GHkBrWzY5e0ZdOY9PwfUqX1dLlWvOzgUcxJqNIUjwXYeKbTM8/hv/AH8urZRf0pqsHomOcK7JaygkrtLVq2r8rkaSnv8AfS/iXF3aI1LmGmtQQ5zl2WJmUlA8k0uVTWU1tLJGYqulZPvxzQPJESt7Eo1gVxpu3Hs8TTVTBn+mcxfMqSGmpszyWuYc5nkk9/q0zf8AeLtenmHVZoJAftDFyyWihDjuB/2nzyPkorRDUFo8R7j3VEqbruEbAeotb4c9cFv4eeGU9CObH0ONjnEEDmLMqJB9UrAZIx/w3++h+BxrOrbbMQeOnXGyspKKxXzYFT0FwetiL+XGMBQXIIt6sOo9fbBO8Z2MjNuZj1PW/lz+GFFfbJboORtK7h73Hn/VsCVHZSGCOq8CxtzuB6HjqffzwVi6qocHkBrHkEeRF/L4YBVKkJtFzYdb9fMfH188GkAjUhlcSIQACOAPMeoIOEQkIXL1Q5vwbC5PljrPJ62ebU1BSw1/dpJntJFL4VMkLSTpGzKGHhJViCbbXHryMclQ7jVgNxwTwLY6lpe7i1nl9ZK8YSDO6SeZpTZERKiNmcN9ywW58iASRcXxgaZaHSMB2H2Wto4lrHHePdVTm3bPqLJ8/qqRcj08TSVE0cbtlxuAHZbgBwBcDmwA5xuNK9oR1lLNl1ZlVNQ11NC9XTyUKuscqJYyRtGzMA227KVIBK7SOQQXU30dtXZhXz5muoNGbppHYqurssIuXY3BM4NuR5YfaE7Ll0XJU1+eakymsr5oTSRU+VVS1kdOjkCSWWeP9XfaNqohY+IsbAAM60xWOOEuZS9u2psElodKA6t3epHqOqqM37ItTZbUyRyjKa2izWjIYkL9bc0lYovyveFaOQg/7ynVutyaN7PMtpdTa7yLI8wT/VKusjWqUm26IHc63HQFVI+eLf1/XLkegs4aWSPfm9XS0kTJe0q00jVEzAHyDNTofLc9h0IFWdjsBk7TsjRlveWQm3HHcyYfZpn/AEUj65A08h81TZ42m0sbTOleKviupJs1rqzO81WSFJklzCeWFdypDHHuYoAbFUjQKoBAAVV4GKqzT6QWdxKcqyTT+V02WwsxiirKf63Mb9Wd2YDcfMKqr7eeLp0/MctyTP4alJO7n0dnrIxAKOfqQAIZfA5H7QsfUA45CqT31VM4uQZGPHpc4r6NssVpDjKK0y2KW2zPhIEZpXNXDoTtfm1JquhyrUeVUNO9eyUMFbQwmExOzARiSPcVkj3EXsA4vuBJFjc+iqmOq1DLpWtSaOl1TQ1enswgPRlkjd4H3dC0NXFDMjdVZSRbcb8naECrrvTyk3U5rScjj/fL+GOudOwbdfZPUxyRPG+c94jx7gBcSH7LcqfmQfI4ZpBjbHKDDhhzTrK91oiIkxxXHeaSvI8kzfbk8b+zMAx/MnHQHY1Rw1vZdlULX76XN8xEYB2s1xTAqpIKEn9hxZrcEEc89V3ja/HKpz6eBcdD9kAgi7MMpjqFk2NmeYqWUbgLimsCv37nyBDC1xyLYuaaJEDaf1exVfR1DK7w9wtH9JNTT0+gYO9L20pTEllsQfr2Y8EXP7zimyqHxqQrCxKM3X3BP7jz6Xxc30lPHTaEJkDkaZp9zXJJ/wBezE9Tz+OKWCEdOLAE8Wt64vWE1gaVWtQ/elBIxKMo8S9bEefqPMfxxPPo8OYu0mncOykUlf4lvcf6rJyLc3xAZTYbtwueOGF/mOoxPPo/FU7RYmkdURaKvJZvsj/VJOT7euDSGNlk/wDU+iSx/wAdniPVXbnMEp7Pe0ORodqppuIh0Ud2xOaUfkOFY832+E8mym4xyzYLuIPLSEcGx+0cdVZvIp0B2iRywGOUaZp5Ab7gyHNKLkOPDIPRx1HUA3xydO53qo85P44paHJMT/H2Cs6Q+9vh7ldQaEU6c7OdMQ0rxx08tA+c1PF7zzyOryMtwTaKKFLi5CA8WLYj3aV2tZr2Y6kTT+RZbSSZnQwqayorYjMkMs0Ycwxx7gp2pIodmB3MSAAFBMlyCBans40/DJyG01CehuBacFhxY28yDddwJFucVf8ASqiii7fdWwQkFFqotvl1pYMZtgiZa7U9sor9x5/lW7VI6zwt7M0y9E7yT6TGfzVyHUml8gzCmaUSzCmpfqNQxsbsksbWD2JF3Rgb2IIOOkoWnyDOcqz/AEvQ1dbS97RZhTM1MQJaWdEYpIouoLQSSQyJcrZnXpYY4DYSQOHFxY3GHT1uYVbd60zMzEksxuSfU4v2rQzZHtdZzcpnhWuzWFWh0iWsIlF6qmXabk1NpTXmodNU27uMszSroog5uRHFO8ac+fgRR7jFu/RxphPoPMlshL53Eq3kKG5ozwGFwL/3lYf3fPHOk7y913lRK0j7QoLE3UDoOfLHRX0YyW0VWRMDtkz6Nd1tyo31M23AeIA2PiAO02uLXIdpcOZYsTiKJtgo607sUP0nu5j0roZqYyC8eak95GEN/r6dQGI9OQbHqLdBz2soYbgVUX5Vjyp9vUfn+/HQ/wBLCnemyHRim3AzcE7gRuFfHe1uMc3hyh3New6geWJ9Dm9Y2nx9SobcKTuVhdj2lE1f2hZbDVIZKLLt+bV9xe9NTDvnBHq7KkfxcY6SbVmRaT+saiz+JCmYV9FQzzRMYmWoqp3MkpNjdo1WR2VlIYAKbcHFcfRgyl6TTOc6kNOs9TnlSuU00QkVZTTwbZ5nVWN2BlamBt5KR541f0oMzFDl+ntMxu3eTmozqqv1Xe31eBeeRZIZTzz4xjHtn9/0gID9ow9z8LRs/wDdbJ2us/oPlWpr/KJNU6P1joSSNWzB6WTMqaMA3+v5Y0kxVR5loDXIPM7o+uOOCyogG4En+7+BufUfhjsHR+sJc3p9MdoskENXLHBR5jOm4q8zxXjqoyOkgcxyrfhh3pHK9Oau2DRUfZ52k59o+Bu9pcvrHFHLb/a0jWenf4NC8Tf82L2g5TcdA7Np/Xmq2k2C+JW5FRN9wLAgjyN+vvxgQbXsAOfTnBEZiVAB4I6Hk+nXjE87PuxTtE7Sklm01pmulpTT1NRHXvTyGld4I2doO9UFFdtpA3kC9hcXxsyzRwtvyuAG0rOZG6U3WCpUJ3WW/fKLAgBuQQeoA8wfT92LD+jfGKjtXokAuPqmY2Ujr/qUvGIfqzR+rdD1cWV6u07mGT1k8C1UcFdTmMzQt9mRfVTYgMPTEt+jXJs7V6N222FFmRO4XFhRS9R5jFS3uD7HI5hqLp9FYsoLbQ0O2j1XSWqNSVWhtEZ7qTLqXL6ypyaigqoo6qN3ilZqqKEh13K4IVyLhgeFBLAKFpWp+l3rDMDuk0jpNHC7CY8vmB27gxFzUG43AN/iAPXnF857leVa60xnelswzCroYM9y6OD67BTrWyUpSohmU928kbuh7vbZmZl3AhmFwKopfoo6EQySN2qZs4juXB0gPCLFr8V3SwJ+GOd0fHYHRk2ql6uvZQLVthtQeBBWlFHsv+mXr7J62nraXTGlmmpKiOriMtFOQJUbcpIE/PNzbpyfU4pHNq2fMHkrXQK1QzyuFPALEkgegucXzrX6OGjtMaG1Hqyi1fXZlU5PTwPHTS5N9SSQy1CQAiQVMtypkDbSoDDowxQijbEQwFipA/DHQWGKxtDn2QDfTresq0vnJDZzxXXHYtP9X7JdMtC5WWGbMxu2hlTdVWG4HqDyCDwQSp640fbD2+6l7Pc9ptLZXkGn6vL58ny+uC5hSSVDo08KyNHu7xbojXCXFwoAubXxvezKhpH7JNJzssyStJmbJNErEoBWWIJXxKD62YeowTV/ZX2ba1qabPtXz6smq6WgpaCZsqrKIRd3AndpLslj3jwgFgL8hiOOBzMLLM61yOtOIq7PbeW1KZxZ2CDOg4UVXUH0r9eU6wd1k+mjLTSd5FK+XO0gPNgzGa7gAkWa/HHTEd7QO3rVWv8AIzpavyDT9DSvWRVzNl1C8DGSJJEXrIyhQJX8IA5OLny76PXYtFW/U5ItfpLe6iSsykB1/aVu6IYWINx5EHFV9veh9KaG1ZSZFpanro6Z8spMw7yvaJqktUR7jG/c2iIQrZWUAkNzjassWjHTgQgXsxgs2d9sbH+8+3JVxkta8Oc0CtIW2zR8nyvImO8Mzrv05q+pyeSZk+t5tVRRuUDozqZXsCPsuApIDAEgHax5GOCMvjMmd0Y4FqiIcdP9ouO9dL0Lya5apCN3RzyuSSzJIscoiqTtexDI3AYAr8Dzij+0V0PYTnR3srWiibjhvHuuaan6WGt4oYIjpnS9PPAY545IqCYurgXVgWnPIufK1iQRY2xv+y7toTtCziDSWcZNSZfmVSzDL5aQsKWaUAsIXict3e6xClDt3WBUX3DnWsBkqWUi4KJ/7RcTbsBpO97YtIxA8tnVIo8usgxoWnRtlis73sbQgE4baKtDbJ3zNa41FaLprW9PFrLss1RkVbEVeiyybUmXu4J21VPsMzA8/wC3pjIklj4zFAzXZL441q0K1KIgP+1Av/zY67z36zl2RamVCyqdN50AVCmOQfVmUlXTwP8AEWPkwBxyfA8VTVQiSwc1Eak2/vj8cJoV7uxe0moBwSaRaBKCMyuuOzU/UOzPR2amCN2jyG19n6y3fVBZbnwujDhonFmBJUgixiPbz26aj7P+1HOtJ5Lk2QVVBl0tP9Vnq6FpahkaCORd8odS9t+27C+0AG/OJXpH6qnZnpGmFXUUs0+n0XvowCNvfVPkQVa3Xay2PkQcZqTsi7O+03U8mrNU5XrGbMc2SnXflGZ0SUUkiQpGBH3kTMNwjBCsTySLm2MSztsxtb5LV9tXZ7b36rTlM307WQZ4eiqGi+lnreJFUab0soR2kUDL5VCM32itpxsv5hbA+mIt2j9uGq+1Wgy7KM6y7K6anymapqYDR07xszzCMPvZ5HLC0KWHFrYtzM+wfsUytZaxMi7RKiKIMSGzbLVHh+0LmDrfjnzxUPbNpnT2iu0nPNKacSVcty2ZIaZ59veyIYkcNLtO3f47EpZTbgDG5YodGumD7OBeGPt7rLtL7W1l2bIpPsOqmHa7pRmZkCZzQn34lBx1DXV8uXZNm+e0FOfrGXZJmNfTyg74GkiiZ1VxazIW5KE2v5KTjmPsGpu/7YtKRsA27OqFenW8tsdOVuQtmuVZplkc1PSVOY5NX5bFUVEvdxCWamaNDK9vCpIA3NcC4u1hxm6ZDXW1gdsHqVe0dUWZxG9c/wAH0m9bK8cv6K00ktO/eRuMpsVf9ofrOp8/Xzvi0+yftDqe0fIsxrk0/SwZnllTEKqCk3RwzpNvKyxKWYxyBo2BQEq1xtAPBrWX6LuoYVZTrLRjybdzKNS0At/9Fv8AO2LS7M+zbLuznIKrLGz6mzPMc3kgqKuWnDfUohCGMUMcsir3rAuzs9hGSFVC3LYl0lBYIoC6IC9hSnPlVMsUtqfLR9aY5qP/AEi46Or0FpjNzEWlpKmqyOF3Ub46dFjqoYifNYzLURqPJXC8BQBzpsIRyykC3h55v5f54vP6SefSJlWn9HTH/XY5pc3q0P24jNFFFAG9GaKIykHm0iE8scURPJsjYqAPYeZ8vhjV0QXGyNL9/CuCoW+79Q66rH7AckizTtCo82rohJR5Ckmcyq32X+rW7lD676l4V9+cdDJmEGQ/or9NVizf2hzWLJaWZJth+sPG7vM4YEOgbugw4ZTNw3FjXfYvpNsu0BUZwyMJc4re5VxE0gFLR/aJVfFtaplPS5/UHjjEZ+kdqapy/VGQ6WpqtRNpihjqnaPcAtdUFagnxAG6ximXnnwc4xbUHaRt/ZNOAqD5D/kQFpQuFjst8jE9eitDthoINX9k2dUyowzDTzRaghiZLSBUYU9WvTn9VLExI4P1YHHKgWNgFsSP7rWI/Hg/v9MdpUuaZTqOsoNRVkJGS59DHLVpscBaOvhK1CAgkHas1QOgsYh6Xxx1qLIcw0hqHM9LZqAtbk9ZLQz8G3eROUJHxK3HscXtAzEwmJ2Y9/yqmlI6SCQZFMqFf/LEItY3Fxfp414x3LBTrl3aMgiqGkgbPwxUEHkySWNm6G4ILIQeAGRh4scL5c7DN4SBYXFv/GuO6JqhZNfPDTwSSD9OOyvGbbF7yRtzxP4lF7+OMshPULfFT9oSQW+DvZWNEgEHxHuuHi5KtdwgckEkDjn16rfpcflhOOUlbi4sCSAeRbztY8YKXZQAzW3i48RBP5Hj8MF3OAFFuBey3HN7c388dSMlia0Ytcd5u8QJYNuJ5PQC3U/y8sCbjxSOsl1sCH3bhfzPp8eb+WEACTxa546Dn5j+GFd5Mp23sbMOdvlY/jhUmSCQRbrqRx0FrH4nyJ9x19BjGIKFFKnZc77tyD5W8hfztfGLclRI17HYrbSQFHsObA8fzwDozysUH97gWCj1APQYEqFmeXbTx8yP4VbaBf1ufO2JboPTOc57mFHS6bomqc2zqrjyDT8HTvKiVgrSewAbk9Be/ljQZRltRm1VBQUiqKisYwRHoscY5llJPlbz9jjoTsby2nyXJ8x7TKUSoNk+kNGkLZwzoP0lmKeYKQyd0jD/AHlUtjdOKFtmDBdOWZ8NQ8zyqrNmjLjeHgPHb5fCn1ZlmS5U9HpXTlb9ayDStAMmyudfsz7HLT1thz/rNSZZQRc92IlIYC2Aj1Bkukcqrtd5hSvMdI06ZjTwyfYqMwZu7y+Cwurg1H607SRspmI23KjTSU3dtTeMRUoQxxxOEJhVUuG5G2RNqhd1gyFSpFirYgvbznn6Pny7szpGeF8o25vnm59wXMpYh3ULHz+rUxAPP+1klxykUZtlovSa8T4dYLckP08QYzwVM6mzCqEf1GarNRVVDvNVzM12lmZt0rk+d3sP+Q+uNHCvNm2G/N7Hj2N8JzzfW6l6jadh4QMfsqOAL/D+OFk3EAix3Dr1H/X2x2kEfZsprXPSPvuR2SxvwfO5Nx8PhjEYeMkFgTYdF/r4eeDs25Sdw3N4WsbEi3J6WI4GE/sqCTY3sbHkfxH8cTKNDx8d3h+B9OcHJuQLC542i9wPj5/DzwQjkBl3FiLW5I/6+nUYPuOxl3BlfyNyCPf/AC5wiEV3UIdwUXPX09x/I4ISwYgRgt8bc4Ua4IYM5IIJtzb4+v8AHBAvmNwvcAAEXHz6j8x74EIm4NZlIXb5k2P/AF+HXAq9tzcWvutew/yPtg4jA4G0i12sbWt6HBG+xYKp3G7G/Kn4f18sKhEfkEobhjyovdffGYA8rtPmLcckYzDUIDt3bxuJJ9h8vfBbbCqtwPtcjByovZLk26DyA/l+7BQAeVHQ3/6/154chDex3HkG/wBrm5/jgpa7BFW/sfX2wfaX/VqLljfw8En2v/RwQqCu4feP4/Hy64EIFvyRY+dyLDAKnNwOD1vzf/PAjcqBhcKPP+v3YLcfeAJ6C4v/AEMCEUxlr8EnqbDzwkR1AtyeLfww5kBZRchvIH0PvhF1BAJJv058sCEiQV6YKbeR4wo6EMQR04PsfTBCABgQsR2jdZF4Km4xYnZdqqlyXNZcszScx5NnUS0tW55EA3Xint6wy2b/AAM488Vz5Yd0M5U90Lbr3S/QnzU+xHGK1rgbaIix3W/yzU9nlMTw4LoKrjljlky0qfr2Wz7KiK5V0lX9hj04syuD4hb1xO9NLLrjR82gJKaWtzSjarznTESptllZl35plSr/APJEcf1qFR0qad1H+1xV+lM6bU+mErTWKMwyeGKkrO9BYz0W7ZBMSOQ0bEQufQxNiTZXV5hl+YZZPldVNl+YU861VNWUswZ6eoQhoZY/R1dFYA9fmRjmYZXQOLZBlg4dcRryqttzBKAW+I65FU/mOTpluYTZN3iy5fmgE9DOo8O8jcjL6BhyPfjyOIpNC9NM1PUcOhswtbp0P8sdQdsekKPV+TJrvJKCOiTPJ5lmoKdbJlGfKO/qqJP2YZtxraXoDHJPGOYiMc+ZxDLmtM9e8WzMKL9VWxAWP+MD08x8SPIY6WzTXxcJxHMaj1rWLaI7hvAYHkdYWgYbHdHsevIsQL88W8sCGbh2Y8G3LEH2OCNbcNi2+HII+WDqwsLOLXIYjk/H4YtKujKbANcNbwbDzx1tb0OCyEhvDIzA83vzb0PvgWZCQQgHqenPt6DGMG+6Q23noCP+mBCSp7LViwHRuPljqqggOZZ5Hl8axrLW5ilNE7sV2PLKsam4BBF2uVIvxcEcjHKlLb64lvfHTmmamOl7VMnoJVCGXUdCquu2xb63GbNbkH35v7YwNMMvvYNx9lqaPdda47x7qGaj7UdJLXVWVVtDnBlppXp5HVYD4kcqSLt6qfljVf8AbBprLgzZflGZ1bsVIWrqUhQFVAveIF7GwuAy39RiCa2h36qzR/26uZz8TK+NGtKOCxsByTi5Ho2C4M6eKrvtst4nBSLVmuc51rXLUZlKqxwoIoYYoxHFDGpJWONBwiAsxtcksxZiWN8bfsYlRe0/JnlBsHk46f7mTjEMVCgACkHp6EYdab1BUaY1JQZ9Txq8lDUJOEY2DhTypPkCLj4HFmazg2d0MYpUEcQoY5SJmyvOsFdLZHNU5pRZ7Q0CzWl0ln0ywK1w7JQlyQOhIVXPHNlOOa6vL6jLKiWkraeSCoja8kcilWW4uCR7ggj1BuMdEaeqqXJ6qg1tozPZo6OCpWoy/MYrCSjlHIjlUghJVBKujgpIpJG9HIwpqyp7Pu0KClOc6CzCGrgj7lFyDMKeOBVBNlhjqYXlgT0iErxp0QItlGRo+1xWe80q/a7PJMQ4KgtEQGbXOQGMcnNKS3/w1cdN6PzISa4yV0RjC2ZnxrJYW2ScMpFjb1BBt1B64hmncv7PtMRw5jkuj6iklnVwK7NM2FZURofD4FWKKGInkbrO3VQyXJw61Lm9BoOjm1NHUCOtqaWSLLaceCVmkjMZqXjPKoiu4Un7blQpZVZhFbJRbJg1g1UHipLPGbPCS466rnKudkYgDw7I+R67Fx0F2R9/N2aZY0FS8bx5hmNwpB3A/Vibxm4kW3Xg24PANxz/AFah0Z9li5v14UeQ/Cw+WLz7Pa6my3s2yWOV0UVFbmLBm+6ytBY+o8xccj8sXtMCsLB/5exVbR38Rx3e4TH6Qks8kOi2lRAx07ATsG1eK7MBwObD2xTskwUFEY3I9bW/ocY6hki0FqLLKA6x0ZmWfS5dTGjiqKHU5ogKcTSSKO5FLLu2mZ7lXJN77RY4050N2IPI0sPZzqFwvieP+2gJQftcUHiU/tDp52wWO3QxRBjjii0WaR8hc0LnNWuhIBsRYjriwOwYtHrtGjDFhQZgfCQG/wDNH6E8X+PGN720aZ0Rkml8gr9GaZqMrOZT16VQq8y+vybqZ41Uxy91EVVllO5Cp5UEHGh+j4+3tCjSSRIwaGvBZzZV/wBVk5J8hiW1TsnsUkjMqO5VCigidFaWsdnUK4qxpY+znX9K1YwD6dgk+rvGYzY5nR+Pbytj6oSDjmFkX6xFut/tBz88dR6on7rs/wBcQggwtp+J4GR1eNicyo9xRhcWNubHqLkAnHKVVIW6XBDnn5nEGhu9E+m32Cl0hhI3w9yurNPV0Z7OMipHhJh/QEAEgP2X/X2JtyvN7N6gi/NsVb9JmkrZO2vUWZyUsi09bUQtTyMpCyJ9Up2Fj5+B0PwYHEt7Ec2p9X6WpMipq80+oMhWWOGEIXeqoixkBjQAtKY2aQPGoZtjo4Vgr2n2b1ul8/yWm09rfJmzhaBI4aesy2shEwgjBEUUwkjkimWIMUjdlimRLRl3RVUULHIzR9qe2XA48yrVoY61wtMfWC4+q4dqmwIB5F/L1GLUyn6PedVFNS1U2tdG5aaiGGb6tmeeQ006CRFdVdG6NZhxfi4xZlL2e9jL5hDUU2kc9qpo3tDDmNXAKVpCLrvSnQGQWFwpkQNY9ehk2cZvluT0n9r8+qd1FSzrLOrN4qqRXDrTxfdZpCoHh4RGcnaFti5adMXXhkAr1hRQQ6P7pdKaLlXW+nJdKapzLSlRP31VlM70lWQq7VnjYrIqlWYOoYEBweethi7/AKO7U0XZ7mCShbtnsTguQBxRMLXNgDzxcj43xQufZnW6hzvMtR5hKJKnMKqWqncdHlkcu5HsWZrYuzsHzCag7P61ld0WTPoUZ0cggGib06jjpY/A4k0uHOsgDs8FHYSGz4ZYp/8ASlra6u0xoOatnmmcR5xGXluXstfGACepsOOebeuOeZNwS21iGHl5j0FuuL27fs0TMdIdnwSWCSNlzpkMG0KQcwT04B45HHN+B0xBOyXT9JnHaDla18Kvl+Wuc2r1PAaCmHeMn/MwRP8AnxJYnts1hD3ZNBPMploaZrTdGshX9lmmjpnRmU6QniiZ8mo0SqUcmOsl/XTFja6kSSBAw4/VAHnGo7Qex3TnaHm1RqSv7Tf0dLJFBFDRPpytmelgiiWONGkjPdv9knepIYsSOtsP9TZ3JR5Ln2sqipL1NDSSVk0LNcPWTzBY9rKQwDSyEkXBARgbjnFRQ/Sf1FHMJH0XphrbhfbWCwb7QFp+AfMdPbGFYWWtxM8ABONa7TiVq2p1nAEUpoOgre0VpmHRGnF07BnAz6KCpeop6laWSkZEktvi2SXIs6hgb2O91NiRiC/STyE1tDo3XNKrWajl09WOCSe9oyGpyW9TSTwKPXuT6YX7Ou2mbXOs8u05WaeyrKoq8yUqy0b1BPfOh7kN3sjjYXAU2APi4PAxPNa0mX6y7MtR5HTTK08dKuo6CNtveLLSbhOhtYbvqs9Te3UUwOJ7I+WzW+kwoXZ+f5CitDY5rLWM1DcvJckuljdQoNvTr8R0xstJ6z1LoavkrsgrpofrEZgnRZmVZYm+2jAGxDLdSGBFmONZve5WXaHVijBjYbgeecAw7w2B2i42nbax/fz746eSNkrSyQVB1LEY90bg5poQnGY5hV5pVTZhXVUs0spJDzSbio+6vpYCwAAAFuBidfR3jLdplKYnCOaLMgG62JopRivpk7s2IIHTqNwPwB4+BxP/AKODhu1eghk2iOSmzBXD8gD6nLe/I/ePjirb2htkeBkGn0U9lN6dpO0eq6KzzM8y0TorONV0eX5ZUy5FRRVAgrIHMVQHqIYdpCOhUjvSwdGF7WK83xTdT9LTUUc6z0midLQSKRzHBVeRuBzUG4BFx6Hpi4u0Gry+s7INc0tNULI36Go3Fn3XH6SpQff1HPPS5PBPFf1bqT15NjjD0RYILXG58zakGnILSt1rlgeGxmgp8q2c7+kfn+otLZppCp03k8FDm8UMMv1WKUSRiKYTII2eVgo3qoIsfDwLYqmeR1h56/l8sGggCvdQeht54CqQshe4JI8ub46OCzRWYFsQoDisiWZ85q81XYnYe4zHsw0rl1RLDAU+v90Z7pHIz1zXUyHwoeltxCnkXBtiLfSB7YO07sx1TSaUyDUskFC2RZVXCKWCCoeOaenDy2kkjZ7b72BY26DjG77LM/iyzso0lSOQx3Zi4QWLH/XW+6T4ri42+Y6c2B22aaX7HO0Gpy/NdX6bzGvqEoKbLFr6DVX1alMVOvdxXhNHI0RC2VgWNiLkKMcdEIGW2SS1AFlXYEVxvLoZO1dZ2NgONBrpqXNr/SD7WJYEiGpwqRHwBaKn8JuTx+r46np640Oo9e6m13mKZvqvMmrq2KnipVnaNUbuY02RpZQBZRwOL+t8dQZj2L9hVHRz1Ddmep9lM1plOsxuj8QUmwoPIsPxHriJdp/Zn2Zae7LazUWjtN1NHMc4osvEtXnS5kphmgnkLJemhaGQNABe54LAqMbtmtejRIOwADjgKNpn5LMms9suEy1oMc6rnnKnJzyitx/rEPn/AN4uO9cqilPaStRTVjRyDOayGpjRiBOnd1O3ePPaeVPlcjHCOTQD+0FGrG4FRERYf96uO3M1z+HKtdVGcARz/U85qJJadZjG8iCWVXVWCtzZzYgEgj7JGKH7Qd6ZjRsd/tVnRjf3TzvHuuFbbo+8J6ol/P7oxNOwlp07VtOV6giKkzGOpkfyRIlaR2PsFQk4vt9IdiNHl60L9m+bmOCNR3n9oYHYoej7zl92BHNwfh6Yf5Hk2lcmWbJ9B6ViylcxiC1NSKtsxr6iO4Pc7mjG2JrK2yKFVewEjMBYXbVpWzyQOYK4ghV4LDMyQOOopfOsxJ0bqyvhp9q5jprMxUeMqrb4RZ2QcF1Y23cHxc3FscfVMhp6qORFI/XBgRYjhr4vzto1zS6V07W6IynNYKrMc22w1sVMQYsvp1YO0BYEgzSOqF1ViI0QITudlXnxU76pgTnc0igk/HD9DRGOFz3ZE4eCTSEgfIGjMZrrTQNJmWadn2lqiCdW+q5EmxQXR1IlqCCHW4Fm5BK9R15xF+2rtf7R+zHtFzbTGnc8ho6KCOjlSFMupeGlpopWJJiJJLuTc83xNuz0UEHZ5o9ZKlaeZ8jUCUSbGX/WKgWLdAD5bha/3lPVxrPTvZJrrMZtWax0hmdfm00cEVfUUWrlpYnaGJIQ6wmkk2eFELDewBuRwcYMDoPrXm0gFnezFcbw+CtOVsv07exPew10woqBk+kr2uVvfSTaihdp0aObvMvpWEikWIYd1Y8cX64hGdamz/VWdVmqNRVz1WY5g5eoqXVVLnaF6ABQLKBwB0x09/2WfR8opSj6A1KsiNteKXWqK6H3U0F/wv1HrfEJ7e9J9nmS6GyDNNDaYky01uZVlFVPVZquYyMIY4HRkk7iIxj9eQRyDYdLY37JaNHtkDLO0BzsMBTf7LLnhtbmF0xqBvqq57Dq76r2vaUkUg2zqjItze0l/njrSoz+Wi0vmGoqOhpK16HLqvNLOCscwhj3vskjsV3EHcCLBgCVVi27kHshp3TtT0uWGzbnFLwR08Rx0cmYN/ZPVLCojmWq0vmrd7G4IlvSnxFl8MnThvtDocZum4RJamncPUq7o2Usgd5+ii2h+2Wg7RtXw5FU6cGS1VZFIlNLBXySpNVAbkjZXA2lgGUEfe2i3NsT3P5tQwaPzSp0TV0/6cy2jfMKP6xRpO7wxHfUJEHuFkWO8yXBt3coHUHHHlJXTZPmaZhTTPBLHIJI5UNmjdWDK6+6sAR8MdU5F2gy6hbJNZ5LLSU9XUWqShFhBXREd9GwvZo99m22v3UxHI4wlvsbbJOx7W1j1g44+dcx6b0WW0utETmk9/h6LlqrzbMs1zKqrc8qp6qsmlZ5pZ5C8kjsbszMbkknqTgtPSVGY5hBTU0TSyMwKRjks1wFX3uxUfPE+7e9DUmkNWw53pyDbp3UkH6QyxCdwphvKTUjH9unmV4T5lVib7wwt9G+noZO0MajzhUag0xTvnc2+1neAj6vGL9S9S9OLexx0M07YrMZWZALJiiL5gx21dCwZbTaVgy3S0tCJKbJIKfLiUY3m7lr1LD7rF5nqDwQeQPMYrbVPYNkOsdS5rqrMO2COafM66aolePTNeyozkvs23DLZSLD9kXFxziTa41LlWmNAZhqigzCKvqIKqmgjgnU7amplZmfvV4JAjSVyykNuK2awxUFP9JXPoJXlGlNPAvGIWBWqIKA3UH9d903KnqtzYgEg81o4WujprO0Ek0NePutm1mzktjlNKK8dMUtDprTVDpSDN4s/jymGSjNUkE1JuieRpY0McgEikd5Kt72YWsSRY1H9IvJlfVOWaziEjx6ky9ZKli24mtpj9XqLnzLBIZTfk98cbHQPa/U9oer4MmzLI8roJKmlnihnpDMrNKqGVUfvJGBVu7ZbcWLAixxI+1ajpdR9lGYfUZInq9O1EOexKrhj9Wk201VYA9LtROfTa+JLGZLLbqTChdn/wDRPuFHabk1mHZmobl5fhc10jEZtEb7juUX/wCZemO3cskqk7QIopIdynO5rGRDe15fHGensTcGwF16HHEWVkSZrC5Xqym3vvXHZmcZjHkOtMwq6KaHvqbOalzDHtHekSNdW4IN1ZhuF2UkEggDEmnwHuazaHeybos3WudvHuuNkV3VCVLeEC4F7Dy+OCS+mwMOl7cEe39fzx1jSaE7DhBBRp2a568bkxxGXVMUjq9r9018vDB7A2v1KkDyvrZtGdidD3847LM4lFMqvY6ojTcrMACCaC9jewPrwcazdK2Z2APJUDYJxmOa5ftdrLELnwgEWPsb3scHjEy2jVtrg7iFcG/W4B9fUA89cS7tkyXKdNdq+q9M5LAsWWZVms+X0oIW/dxnapcqApYjksAATc2F7YiqxeAb1sh4AZgoYevtz54vxvEjQ9uRxVRzS0lp1IhjaNQ0gK3+1uFieDz72NsYsLTusThUUKZGax8KeZHlz7eeA3JCu5ZLyKQFKC+6/l7H8QcSDT2n4sxqZIszlaLLcvQV2d1IP2IgfDCp/bY+ED1PscJLI2Fhe7IJ0bDI4NbmpX2daGzrVNXlelsiENPnuuSYKeWY2jyvJogWqKuU/dTYjux/Yjc+Yxf9fJkjTUWUZFl81NpjKaSPJ8jSRNrrRKxImccK0s0jSVEouDuksCDHbGk0NlU+kdKVef5wkFJqbtGoYy1MSN2VacuGpqNb8I1SEjlYNx9XijUkfWDh/DC2y0BqHqi6dykKEzPMxAQKDcbixUAEm+4csrHbyekZ3F3Zk45nx2eQw5resUQAvjLIfPml6zUGWaOyPMNb5i0FTBkDi9A8LIZ8xcn6rSsp8LBnBkLJcGKJz94Y5Q1Pm1dLFPLmFW9VmecTyVdZUObtIzuWkcn1d7/JffFp9uer6TNNQ0+isvqKdcp0mZnr5qUHuanNnAFVLHc8xxBUgi9owR9o4o2sqZcyrJKspt3kBEB+wg4VR8AAMaGirHcAc7M4n2HvyKqW60XyaeA9z7IsCEqSRbyv/PCygAE+ZNrEeeCQooJfaSBySePkcLr3i2lSQgi6hgfLzAP8OuN9ZSMxfdsIcH7oZf6tjFChRdwvHHr+Hx9eMAWsNoA8m8J4HHsTz/RwMbWVm3EE9No8z6f5/LCoQXsLHqLEBf389PbzwcbibDd4rn4k83I8z1wi/gsDe3l639vf2OFCxKlQb8gi9+Tb8cIhYCha2/dce3HuD/lbAAoSA24W8ww6/h5/5YIW9/kep/hbCiKzMCAAT0uw4Pr8PY4EIpH+0dHHHhswBF/S/wDHBC0QALA+psOnthTaSu4bB4evt8T0/DCMjAlTuJHQFhtv7m3TCoRQ1nDofF1uDbGYTYWJA5seDfj5YzAhCeliARxa46nBkN18v2SAef8ApgljckHkdb4FDdT4ufTytgQjBhuDFh8fIH1GCgW4Nr3vwB6+ft7YFdxJCC/PA81wBNz0uenQH8OMCED7T6KQb39b+3tgqizAEcX8uLfPBxcXCsL+/wDXX44y5BD3YFbefI9ucCFhPAF7m1iPU+mCNbkkLbp62wdtyP3ewe48/W38cJsd12BH7vngQiOPsjz8sIuPGfL2wubX5BKjy9MJleliPO2BCStbGA7SCOvrgTa3vgLHywIUr0Tquq0zndPmtOiSqS0c9O/2JlcbZIm/uyJdfY2PkMXVVxUhliny6rafL8xiWpop2HMtObhd399CGjdfJ0b1GObYZO7azC6kWYe2LX7LNSitH9isylQmebv8tkdtojq2ABTcfspOAqk9BIsbeuMHSlkuu+pYMs/Db5elczRa1hnvDsneXj+fXcrU0TqeDLaysyrVv1mbTWexLl2erSkvPAkT74a2nB5NVRykVEfm6GWO5EhxEu1zs/zfS+fVdYi0k2a5aq/WkozupsypHjEiTwH70MkRWeFvOJ7faiYYePFUwl6xYJ0see7A3qyk+F0NmDKQRxyLEeWJtpWSTtFyCk0BGm/UWWJIdITIdr1S7mmlyQ34Dlt9RRX4WbvoPszACKyzGoAzGXuFLNGKEnI5+xXLOa0ENOIq2hcvQVoLwN5oR9qJvQr+6xwyRbMNrruPhFvPjj+WJ5qrT9PlLnMVjI0/m7K0/dRm1JM19s0ankKSG8PUEPGeVGIVW5dPl1XLQ1W3fGQQym6Op5VlPmpHIPocb8UolaHBY8kZiddKSch9tgVIWzLbjcPP2NvzwBZQjXCk8W4Pn5j0t54WpqSpr6laOjppqiedtsccSF5HbyACglifQY2z9n2uv/uLz8ccg5VUf+6YcXtbgSmAE5BR1JfqtSlQY94U8qbgH24xKm7Ydf8A16DNU1BItZTTx1UM4ggDJKjBkcHu+oIB+WGR0FrRiV/sfnu70/RlRf5+DBRoDWZK/wDvHZ7/APwyf/3TEEjLPKQ6QAkbaKVj5WC60kJlXZjUZtVS5lVBEmns8gHQsTdmt7sSbdBewtxhBVDOxC8AbgL9B7E43o0Brcpu/sVn9/8A9V1HP/sMHh7P9dSHami8/YXuP/JVSQPwjJxIJI2igIUZa4mpCjxHPO4gXAHQ2wznSzE4mcnZpriFd0mlc6RWI5fKqscnoP8AY+eGs3Z7rQgINKZ1L7pldT4fjeMYXtWHWEga4alpsj1Vn2nJzUZLmdTSSMuxmhkKll/Za3DL7MCPbEhp+1/WMBuaijkJNyXy+mYk+v8As8a5uz7WoY30dnvXgfoyo/8AdMB/2fay6tpHPBb/AO9k/wD7piCSGyymrwCfJTslmYKNJC2tV2ya7mMhgzZaYyszu1NSQQsWbqQyIGF/YjnEbmzGvzeoaszSrkmd23u8jlmdv2mJN2PuSTjZr2f6z6ro7PWv6ZXUf+6YcRaG1pETs0jntwP/ALV1HH/sOMOjjs8X8MAeFEx8krxRxJUeqmuhVRfZ0b+uvxxvMp7T9YZDlkWT5Rmhp6KF2kSEwxSKsjAB2HeIxBbaL2NuMC2hdaSNuOj8+byJ/RlRz/7DrhOTQWrAP/rUzq//AOraj/3TD5GQyikgB8UjHvj+0kJ+e2XtBYFP09YNybUlMP8A6lgIu1/tAhfvEz91e+7cKeC9/W/d8H3wxj7PtZyWMejs9a3B25XUH/3DCh7PdbdDovPh7HKqjj/2GIvp7LqY3gE/t5v6jxQ5vr3U2r6WCkz+t+sxUTyvTDYkYjaVg0tgigEsVU3PTbx1wwybU+b6QzIZrkFU9JXqCqVCW3IrKVYWIINwbG4w9/sNrGnBL6VzdB5b6GZP/bKDhtLo3VlQfDpjNz/+z5if/a4kEUIZ2YAu7NXBM7SS9fqa7VtG7XNd5jTVeU1+bvVUWZosFZTGOJFqIw6uFYogYWZFa4PVcRiWAbdvXqSSOvwxIMu7PtWTuI49M5qHt9+jkj/NgBhWfQWsoCySadrh6ju+D7+x98LFHFDhGAPDBI+R8mLySorT1tVl0qzUsrRujB1ZSQVYG4YEcgjyIsRibR9u/aYoX6zqN65lFhJX00FXIfjJLGzn5scac6J1U976YzdiOLrl8zA/NVIOEm0PqpTb+y2c/PLp/wD3TCSQwzfxGg+KcyWSP7SQpEe3jX7FZBV5eJEQRqwyijBCg3A/2Xlc29L4juf6y1RrOvFfqTOaqtkVe7DTylhGnmqDoi+ygD2xg0ZqiM+LS2cA+X/k6b/3XB6fSGpZXKLkNfGV85YGiA+bgDDY7LZ4nXmMAPglfPLIKOcSFq5zsQxDbe3Uc/njdab7TtY6Ty2TKdPZp9TpZZVqJI/q8MoeZVKB/wBYjEHbxxbBJtDarj2sclqZEc2Ap179h8kuQPjgBonUm8AabzNQT9+kkX8yuJJI45hdkAI3pjHvjNWmhSmoNd6r1vQ01LqHMY6hMu7z6mzRpD3QkffKqiNVU7msxuL3HBwhpjtA1Ho4zvp6qFHPVRGnqJRGkhkh3K/dkSKwtuUHgX4xtKrs81hSUySS6br40cXRjDww/HnGmbSGpSxZdN5qOfKhmP57cNMEJZ2ZaLuzVwSiWRrr4JrtW4zPtU1xqbLJ8hzfNjU0dRJHO8fcRRjvIwyo940UkqHewNx4ibX5xEnpluLAWPS2N5TaP1G4cnI62PYNx72B4r/DcBf4DCp0jqS25cjriL3t3DEYdFFHEKRgAbkj5HyGrzVabLMyrchrYsxoJXiqKd0lidDZo5FYMrD3DKD8sS2Pt27Toq79Iw6i7upvIS8dDSoT3isrg7Yhwyu4I8wxHnjTS6Xz2wjOR5iZCOAtLI/42Bw3/spn6PvfIMyW3XdRSgD/ANjhr7PDI6+9oJ2kJWzSNF1pICSndaiVqgwJAzopaKMkgELyRfpe3S/F8EhbwkC5NvD6X87+3XG6p9IagqEXZl7kDlQZEH5Fr4GbSGfQ7i2WyWHNlKuR7cE3xMCBgmUJxWoY7j4CNrAiwIIPt8f6GFch1PnOjM1XOtPVDUmYxbu6q0AZo1ZCjgBgVO4MQbg4e1Ok9TUyLUT5JWbJheNmjIFvnz+OGoyXNAf12W1F/RYmb918Ne1sjbjxUFK1zmG8MFuqvtp7Rs4o6nJ8zz+Sqoq+JYKmB4YUWWNXVwrFEDAB0VuD1UYje1QpCncbcM3H9H44W/s3m88m2DKKoi3V4WQfMnjGwptI56wBNHtIHnMhA+d/yw2KKKEUjaB4CiV73yYuNVp7srgENdTYA3uB54LUqdrMObA/1fzxJG7PNUKFdaamVTzdqtB/0wWfRmoo/wBX9QeYAEqIJFmsPgl7H2xJUJlKI2T9tPaDp7L4MoyPPGpKGkZ2p4Pq8EoiLnc+1pI2YXa5PPU4cS9vXajMSZNSFixuSaKluT6n9VzjQSaL1LvKppvNLHz+ozW/9rhVdB6oCbm05mYB9aKW/wAvDit9JZjiWN4BT/UzDJx4qSw/SG7XkglpxrWt7qf/AGsZjhKvwByO7seAB8AMa7NO1TXGp8plyHOs5aoy+Wojq5IjDGgM8aMkchKICzKruBckWY+2NfFojVDWtpbOX+GWzn/3DC8Og9Wtc/2bzGNQee9pZIv/AG6jCMsVmjN9kbQdwCHWqZ4uueSFHhXyUNYtTELsnKkn7LAgg39QQDib1H0h+1isleaq1U8skrmR2aipLsxNySe56388R6u0XqaJzHLkVazdR3UBksP+S9sJxaNz/bdshzGw9KOQ8/hh8lmgmIdK0EjaKprJpIxRhI8FIf8At87Ul5j1VOvhKkCCnAsevHdW588MMy7Ye0fN6RsuqtVZgKZgVaGObuY2B6grEFBHsQRjWNpHPr2XTma/EUM3/uuHVNorUcz7DkNfAx5HfUskQ/FlAw0WOzNNQwcAnfUzEfceK0qLJUP3sxLvawFhwPS3QD8sA8hp5Y6hFH6tgwF7jg3xL4OzLVkvC0MYU/t1CKPzP5YbVnZzq6JHY5ckm0XtFUI7H/lBufkL4sd2lFD3q1S9P239otDllFklBnzQ0GXRmGkg+rU7iGPczbQzRliLs3UnrhWl7fO1KjYtT6nkhv17ukpV/dFjRw6L1K6jbpnNzzfw5dOT/wC0wZtCarZjbSGde9stqP3bMVjYrKc428AphaZv6jxW/n+kP2t1MK08+sKmaKMjYstPTPst0AvFwBfoMa7Oe0zWOuMuhynUVf8AXYKOWSop0EEMQikfb3jARooYsEQG9zZeMMk0Bq1j/wDWjnnHn+jKj/3TDqPs41kAu3TlYFaxIeB0PzDKLYG2OzRkOawAjWAAh1pmcC0uJB3rUZdn9bpvNI8yyyR4aumkSannjYb4ZUa6upsRcH5Y3cvbH2hVSVNLJnjtDXRSU1Sgp4E76KQWkRmSNWswuDY3wnVdn2r1kMbaUzdih6xUE0qH4MqkHCEOgtWs5C6WzlWH7WXTj/3DEj7NDI6+9oJ2kJjZ5Gi6HEBR+tVmQMbXub2PT2xutLdoerNJUcmX5JmXcU0sy1DxtBFKDIF2hgJEaxtxcWuOuFm0Drl32nR2e2PHGV1Bv+CYP/2da0vY6J1DfzP6Kqbf+0w6SOOZt2QAjekY90ZvMNCnGbdp2rtaZemntTZsKvLoamSuijkiiQQTuio8ibFWxZUQMOjbVvyAcM8k7QdSaNp6ij0tWnL0q1RKwpGj/WgkhePeHVh4Sbi1ug9MPIezTWqqJ4tN1y3B4eFo2H/K4BHzGNcdDav3sj6UzklTYlKCVwfmFscMNngLOzui7s1cE4TSXr4Jrt1pTPu0bWGsqKPK9QZs9XTQSNNEndRxhJGUKWtGq3JVQvN7C9uuNA9KtgQot0uMSKDQOrCDs01miW697RyRf+3UYWPZ3rIDf+gqgfLD4444m3YwANgTHPdIauNSo5lWZ5jpzM6fN8qqHgq6SVJ4JkPiikRgysOo4I8+MSeftm7QapJ4Zs6TZV08tJOEoaaMyQSrtkjJSMHaw4IB9PQYbz6F1VCt5NPV7W6BIGcn/wAIOAotAasrZhFHprMY/PdPSyRLb/E4Aw18EMjr72gnaQE5ssjRdaTRaUVJy+RaiKBGYxkL4riNiQQb+ZFvPzxLqjt97VKypmrKrVDyzVDF5XajpSXY9Sf1XJwzzHs11nRqHkyCskv/APA8TzfkinGubQ+rQo/95POgPO+Wz/j9jCSWeCU3pGgneKpWzSMF1pIW6Pbp2mSR90+o2ZeOtJTH8+69h+GMqe3LtLnQpLqWR9ylW3UtMbgm5H+y9ecaZtCasRA7aczNFYffpHXj5jCtNoLVNSwSPTmZOT0ApWuT7YZ9FZf+23gE/wCpn/qPFNc0zXMtRV1Xn+fVhrsyzGV6ieZ2DM8rG7O9gBf2w3bxG6KFUngBrDpe3Pp+WHdZR1OWTy0VbAYqiFjFKjqA6HoQfIMPO/zscIx0ZqHMJdo4o17yaVhxHGOp6m5PQeptiy0BoAbkoCScSj5fBUzyxVEMLTzO4p6OEC5lnPAsPa/42xevZVoDJ8wrnp9RU/1zRuipoq7UpEoQZ7nD7jTZYrkjwsysGa9kijnkJuFvDOzzSmfZpmmU0Wmsu73VGpWNHpyjdtq0dNY97XSseEUKHYubBVV3PAxf0uU5XlORZXonQ9Qs+m8iR/q9btO/NK6UqKnM5EJBXviiJGpvsgjiRgGdr4mkLaG98ZDLe7b4N1b/AAC1LJZie6cznuGzxPolK7Mcwziur9TZ5ItRWTzSVdXNJAYW71iA6MnWJVFlVVPgRI9t1UgxfW+qp+zrTMmoYe7pc6zbvaDTaQ1Ec6xkj/WcwEkYClIVfu47AfrpHNgYiMSOlinkpWbNZHynKcviNXmuaGxGX0cJuzoo5JuyokTDmSSPZZXZcc4a616+r9QT6wly1aLK6KKPLshyotuSkpI7mCn/AL1rmWRvvOxJ+3jIsUBnkMjhUDmdnzu2K9aJRG3s24ewUOz6aOjposngDBgBJPfqB1VD787j7kDyxqoVNt224HNyvT3wWR5amZ552aSWZi7u5N2Ym5J+eF40VTv+yVI5Dcj+eOvhj7NtNetYEj77qjJHCv0sxJ5Btf8Ar4jCoAY3BPJsx8x6A/zwUbUPLAq3WwNiPwsMHVWbb4WNweLXuPb19xiVMQsG27beIchLC/v/AP8AOAGw36EMPtbuh/j/AFbA2uAuw9CxW3NvI/15YOhZWJMjhlG4lW5A9cCElxvuyBmJHW5v7H1/fgAzx3sWW/B6gkfvA/fgxBYtYAuTuIU9P4/MYTKrZfCpuTcgXv8AnhELL2AAA9eePwt0PvgEClSBYAX6kHp6jrbBrEn7a8DqVuL/ANefTBd3NmY/EC/Pl0wqEaQi1ltc8+I8k/DpY8fhgryHYEW/BNzfgt7fzxivcjlCejC/A+Xp7YyQ3tuIJ9bXFvK/r8cCEh0JZAOTxcdMZgCL3JJJHX0+WMwIQyA7iQLAdQOSP5e3ljFXcLFQD16XOAa4J5vfArtII9vxwIWCxsrX8Xla4vfBiu7i3J9+SPTGEbrnk8cnz48zgCevTlvjgQhBLKVJsD1B4HHr64AeIXAYEeXUjAqd12v158Itb3+H7sAbjawa1ub++EQgPIIIuB18gB/DAAEi3Fzz7/5YPcKLkWIINvTnGFQtjYjnpbr/AF784VCRZeN23i/B5PPpfCbDbdd3PU28vY4X2qbiyc+YB/nhIqx683HHGBCRYetucEwq4uL+eCYEIDh1RVDo6hRd1vtv0Yea/PDXy5xgJHIvcHjDXNDhQpzXFpqF0LlOo4u0DTgzh5ic7yxEjzVb81kHCpWf+kU7Y5j5/q5T1kwqply6VXmnqVjmZbui93JBIpDJJHKrXWRXUFTYWZQcUxpLU+YaczimzrLpxHLC/jUruVrgghlPDKwJVlPBDEHg4ulamirKCDNspu+VVZKRozF2o5rbmpXPVltdo2P20HPijfHL2mzmxyXR9py3bvcbsNQrvQyi0MvHPX8/KsbWclJ2n6SzHXf1WiOYxlRrOjWPu4zNMwRM6iUfYpqqQKlUg4p6zZKLLLjmzNcilpHXTtaHDqzjKKiYbWuG8dHL5BgTx6MRbwuMXLo3Uub6NzylzzIJaMz0glDU1bF3tNU00id3PBPGf9rTyxkxzR9dpVxyqkOO1Ps909meVU2faTp6yTSudyNBSRSP31XlFdGm58rnf71TCh3QSHiqpSpB3oQNGy2jG+PP5+VTtENRd4fHwuYjWVuXVsdRTPJTVFLJvUpeN0dT7cgjEpXtX1GrFhnupwGN9i6gqAo9hze3zwlnWSV1ZUSUtWivm9PCJt8fK5jTgXE8f7TBRyOpAPmpGIsEjIJWxI6i3Ue2NN8Mc1HOxWe2R8eAUvXta1OkgkXUOqCAb7G1BUFT7Hm9vngT2t6mPLZ9qdj/APNBVf8Au2IgIgBylvPp5YU+rxK1ibi1wVH4dfzwz6OHZzKd9RJ0ApbH2warisY9QaoUg8Eahqx8uGwp/wBseqBYrm+pLf8AzR1n/u+IZ3QPIUC/lgVjUA8cetr2whsUJ1cz8o+pk6AU3HbVqpo+7kzXUkighrNqWusCOht3nlhGPtj1VHJ3n6Tz4ny/96GuH7pMQ0Bd+0pweLdD8cLLTGThAbgXItzhBY4Rq5n5QbQ/oD4U0/7b9XuQz5zqMbeFC6krrD8ZMLxdu+rohb9L6ka3k2pa+35S4g/1F1jLNGbdRx19sN3RQBYG4PmtvlgNigOrmflAtMnQHwp8vbvq+JmZM41CpdixH9o6/qf/AFuFv/igdcHZtz3Pk235XUGYXN/U9/iu0h7w2EZ3eXHXBxSEoZAt1Bsx9D7/AMMH0UGzmUv1Mg/QKw0+kP2gBNv6fz9T1J/tFmRv/wDR8YfpB6+PTU+oU5v4c+r/AOM+K57lQPsqSeCPP5YM8CbrKjBWPAI5t/PB9FDs9Un1MisMfSD7QEHGq9TN6A6izAfumwpF9InXpNm1DqAf/vHmR/8Aq+K3lhEahmSwPt+eJr2e9jmca9plro85yfJY62pbLcoOaztCuaZgE3/VIWAIDEFRvcrGGkjUsC4w19ms8Yq/DzTm2iV2XotrUfSL7QuO61TqLjyOocxt/wDT8DQfSP7Q6di757ncjHzbUGZfwnGK9lyypp6meiraWSCqppGimglQq6OpIZSDyCCCCD6YSFKF/wB3cEeQ5+IwosUNKU5lN+pkr+ArSn+lJ2mOvdpnecJY3DLqHM7j8ai2DQ/Sd7R5JBJLqTUUrAfe1DmBA+XfYqwUJkXwLuI9MFMLU5u0ZHkb4DYYaUA5n5Si1SbeQVtP9KXtMp5RNTakz5GHpqHMrfh9YwU/Sq7Tnj7p9R503uc+zIn/APqcVCWjclGUqeLcXwqaUpYPGVv0JXg/A4QWCDZzPylNrl6AVrRfSk7R4Y+6XO83te5LZ9mbE+x/1rCp+lh2pFIohm+YFIb7b51mIPJv1+sHFSNAgXhDxwSPP3wURqDa6AKLBgPtf54U2GA5jmflAtcoyPIK5oPpe9qtMP1eY1xNurZ9mf8ACoGAT6Yfa+s5kGd16ggDYM9zTaPf/wA564p4wggBrdLAny+OE+4QLuIA5w0aOs41cz8o+sl28h8K7X+mb2yiNYo87qFUG/OaZgTf4movgIfpkdsMR3tn9ax9Dm2ZEfh9aGKSWMHkAHnqcLimBUEpcEeE9P8Ar8MING2YZN5n5QbbMczyCueT6afbU690M8mWMsGsMxzC5t7/AFk8YdL9NztsWPYM5dr/ALWYZgf/AO5xSYoFKcx2J6E/yw3eAIxjKte/TofwwHRllP8AJ6pwt04/mV3RfTW7Zkm75s6l3g8Wra63/wDU429P9PbtzgTbHmqN676itb99TigDR7EZwht6FfL3w2WNLldtwDYgDp739MRu0RY3ihZ6pW6QtDcnLpKH/SG/SJhI7rN6O3uanp/8Ow7T/SNfSUgkWojz+hQp0AgkYH4hpDfHNJpjs7zu7ni/UXxkcSM3d2ZfQemGHQdgOBiCDpC0/wBZXWE/+lP+lBPCI5KjS5a1u8OT7WPzWTGkf/STfSSmD99mOSneb3SCdCPgVlGOcZcsl2iSVHZQbAtc2H9eWGU1LsFhC173BCnB/Yej/wDtjn8obb7Q0UDvRdOx/wCkg+kSUlT65lUrSJs3SLUnaPUfruvvhin+kL+kTE6uM1ofD0BapI/DvbY52paRnchY2a4sRtINvXCz5e1twj3XBIsv5/10wDQdgBqIhz+Uv9pWn+tX5V/T27dquIp+lEpnZtxkgra5WJ+c5FvljQS/TS7cZKhqhtTVXeMCL/pCtAW/mAJwL+hxSiJyUaNwff19RhU0UjMrLEWB89txiZui7IzJnr8prrdO7Au9FddN9MntlMSpPqLMJSONz5nXH/8AuMLr9M/tjiV4kz+rCv8A/fPMLj4H6zxil4srY8bDz0HPGGlbCKORQ8TEeYIt+eFOjbKf5eZ+UgtswyKul/podswJCZ3MeLXfM8yJP/8ANYLSfTK7ao6lZl1DOjKLArmOYcfjUnED0L2Q59r/ACmozTKc105SyLOKSko8yzqGkqcwnCb2SnSQgMQCv2ioJYKCTxjWZ/o6XSuoK3T9TmOXV81C6RyT5dUiop+8KKzIsgsGKFijEXG5WsSOSz6GxvJbdBKd9XaGi9U0VvZh9NPtyqxF9Z1ZWyCK+0DMK5OvW+2cX+d7YSj+mn2xU8UkYzyqkMnVpM0zEkfC1SLYpaoomjjLGFwtvtEG38sNVpw5BN/Qm1yP54cNGWYYXfVN+tmOtXV/8WZ2xs+855U+tv0nmP8A/tYeJ9N/t0pqb6rT6iqggbfuOZV5a/xNQePbFGLTCxOwXA3WPP8A1wBprk2Vel7WthXaMsrs28z8obbZm5FXg303u3uQujasq40kbe4StrCWb1v318Av0zO2ua/1nVuYOrEEhq+t5sbjpUDFIiCIEEgjb1tyflhTuk2N4QDaxt5/H+Ywh0XZP6Ev10/9Su+u+mn2t1dzNm024m/6vMMyj/C1XxjUt9LDtOnLl8/zUh2LWOeZpZL/AHVH1rpioZaYWDBCxv0Ck/nhaGnuBeP7XHKdPh6YezR8DPtB4lI61yuzPJWmPpWdqMLDu9QZqArhiP0zmR3W+6b1XQ+fnhyv0tu1lmEqapziOx4Rc5zDavsAak8YqOfL5LB44H8Ju3gJ/oYBaR2FhAdx55XaBhTYIHZjmUgtcoyPIK4P/ivu1iNPFqPOZnN7ySagzO5+QqQBjXxfSr7UYplnTPM03r5/p3M7n4/6ziqZqcoQrx2F+bDz9sZHRkhbwNYfe8j6YPoYKUpzPyj6qXbyCt9fpX9qa3f9O5y+++4NqXNQOf8A55wQ/Sx7UWjSI51nNkFr/wBqM23N7sfrPJxWLZYgAIVyG+zdbE/LDGamWJ+7aM7xwbjqfLCCwQbD/id8o+rkOscB8K1JPpSdpcws2e58PYanzT+M5w1b6SvaaJu+j1NqJbrtC/2mzLaD6/7a98Vg0TK5/UlehA2EYWigE3hSMk/DkWwosUIyB4n5SG0yHOnAfCseP6S/amqFDqrUDXHJbUeYk/8A07Dmn+kz2kU0HcnUGoZWJuXbU+Ygn5CW2K4iy9X8O0k+lsFqaIQj/YSD1LC1z7YV1jidmDxPykbaXty9B8Kx2+k/2m7t39otQD46lzE//VsYv0ne0wIypqfUSliSbajzE/8A1fFWtGlwAtybC1uL/wAfhhRICW2yIwPutjhPoYRqPE/KX6uTdwHwrKi+kz2qQyiWPV2odw/az7MCLfAz4NF9JrtSg4XWuqRfr/70Fdf/AOnYrVoFU2Cj5/vwXuEsLhOlxc8/uw76SKlKcyk+oef0CsWb6SnavJZP7danKAlrvn9aWJPr+tAsPLAQ/SQ7V4ZEmg1nqAyo26MtnVabN5G3feR5+WK67lL2IBa4ABHX4YEQqjAstj9k2HI/nhfpITgRzKT6h4NQU6E9XWy7pZXqKqZ2Luz3ZnYklmJ6kkk3PuTiT6Zy3LJKSozrPTfT2VOGqLHacyqwPBTJ/dHUnyW5PJGGWl9N1WeT1NMkwpaSih7/ADavYXSip722D1kc2UKOSSFHQ4vrsx01RZPT5X2rZ9lkdPSUiyL2d5FVxiSOaRHtJnFWh4khjkXwqeJ51C3EUMpxXtlpDQYmmm07BsG86tgx2Vms0JcQ8iuwbfwOeS3eT6fzDQGSZhNn8I/t3q+jVs+jCMpyPKCqvFlKKvKyyp3ctQoG5IRFDYtJKuF2WponphDT1lTJVbPq5g3SPUM42psZRZy28BStj4ttrEouRS1FU5r66verrqmZ5Z5KiQmaWd3Ls0hazmRnLMSeSWcgsBfEZ7QtZT9nOmqJKOrC6ozenmOTlHN8sy2RmvmDW4EsilkhIF9peYWLxnHMvJtcoY0U1AbBT0C2RSzMvE11neVG+2bVcFOj9kuQVETwUlUtVqivjkDJW5hECEpQ68NT0oLgkcPK0rjju7UTmNd9fnCRFzTw3WEHqbm5Y+5PJ+Q8sPM5q0pKZcmphtew7+33QORH8b+Jvew+7jUwJxySPUY6ew2ZsTARkMvnz9NyxbTKXuIOev4S67jYv6fePphdQdt+TY9LkcHCaBlYACx6r6/0cOFG0N5HaSbqb/8AX+rY0FVQSC6ghy1upsePnfp8sJnweGRUB4PKgAfEfuwobm5t90E+lj6EHp5YArEQuwEkceY+A55/h5YEISwFmLDpuG43AHqeOb9Dg7BiFRrAA3AYlbX8xfkX464ISQriRmsepb1+f9HGMbtdSoU8GzcXtgQjMCXKFBe1rEW/LyP/AFxjoBZm8Rtcjzt7n/qcJhrfdsD+XwGDM4sSCLdbjgW+HlgQssFC8EXJ225vbyv5fDCYXaCQfwH9WwYkAcAX6t/D+uowT7Z3SO22xA9CR5fP8jgQs6GzeIIenJH9e+Be/wBrdck+ZDYxd5Xwte/A2kk4C+2MXYgHng8cYRCR6eQJP9X/AK98Zgb3kI9PU9PxxmAoRStvh1BwpZl4ZfD9pefI/wBdMF6MwFnINz5nGAW4KhQ3HI6/L1wqELFb2I9PLofI8+uCta4YbgehsbWPp8MAButZevFhfA3YhfG1lFgfQenwwIQ9NvJDXPTpfGKSvRRwbkEGxH8vbBftL5DngE2GDcORdueDcta3z+WEQjrdU8IsFBPSx+frgpJBKqL+vucYWuLm4sePQfLA7Va1rWI6Am3p5YVCAksWuPF5/wCL+B/o4TKqBfcfXof6vg55NwCbnoPX2OCuL3BYkfDz9/PAhIsPD1Fh0464TKgdbeh88LMLeIWNxfrx8sJsObD4cnAhJE4DBrN6XwHTAhGjkMTBhz5EeRHpid9n2tjp6rejq1abLq8CKaIvt3DdcWP3XVrMrfdYA9CwMBwpDKEJVgSjdR/HFe02dlpjLHhTwTugdeauj6maOmp0ljkSb6wolpJu7I7xQSO8Fvs7SGVh1Vgyn7t99oLV0mQV9XTZvlH6U09nECUue5T9aEX1ymV90Zhl/wBzVwv+tpZ+Cr+E3R3GKq7MNcUkUf8AZbU0w+oyuZqepKFzSTFdomUdWQqAssY5ZQCPHGt5zUUM9JNWUlSHhkjjjuFcOsitZl56SI1lZWHDCxBF8cx+8sUtw5+o6z/RbouWll4ZLYdsXZe2XrT5rledJmOV1cT5xkOo6eIwiohDhXqjGOYXSS0dbT/agmAlA7t74pHPMolrxU5lBRClzahB/S+XhQLf/JUQHBjN7sBwpO4eFuOk9Ca3y3KRPpjVLVeY6RzecVM/1FA1Xl1Yqd2uZUSv0qo1OyWFvDUxFopLnYwival2TVumMxp5sqzKhSWGmTMsjzigctQ1VDI22KeB25ahkYlNr+KllLQyjaQcbVktIu1H2+h+FmWmzkmmv1/K5xTgK4s1xblfTy64MxCnwXFuR5H+vfG8zvKxOKrMaHLjQT0b7M0y3aQaOS9i6qee6J/8BO08FTjQu5KhCW4PS/FvTGpms3JHurb1EakBDY3JPHnfCTyBFaxYMePb8cDuZSCjEfDr8MIz229fkOmCiFburaLKB2taIii0/ktLT1lDlb1NLDQrHTSszNuZ4r2a/F79bYkNXo7J6Ltq0rmtJkGWHT2p7V0dA0G+iWSPclTT7CfsCRSwW/CyJ7Y02tJqtO2bQ09TIZZYaTKDd+eEY2B9gABiZaWzunzvP6vTdYEFTpbV0stCzf8AAmnaKSMel0YW94FHpjnZbRKyMPDjQs5mtDxotiOGN0l0j+bkKVHBQbJNe5Tler830vL2f6YzBK/UM6w1Fblsc5pITIUWKIODtjW1wo9cNu1LVlEM7rOz9dAaToZMuziNRmGX5WlNO6pdWRtvVG3glT+yMRKbfH2l1AU+MZ3NYj1E7c423alSyL2t6gnkcu4zkm58zuTn88abW/vxiftrmc+KpEjsshnTIKd627Qcg0Lqhslk7ItB16RxRTm+TRRiRZU3beLkWBtcG/FxiK6RpqCt7Htb1tRlGXzVVPU0Zp6ySG9TThpY1ZY5L3CkMQQQb3vwcSPta7Q8zyDMcx0XW6ayDNIZsrjhSrqstiNVEk0YZSkwXfuQnwtuvxbpiJaKldOx3XA8mlovPzE0WK0LnmzsfU1JZiTWuI34V8lNI1vbPbQUAdkMs9y1PZ13EPaLkKVOWUWZ05qwz0lbTiaGfwnwvGSAy38r4lWmMgyCu7QtZV+ZZPBNlunDW1cOWgFYSFlfahANyiqreG/NlBNr3ivZq5qu0fTyi4Y1Q56fdOJvpQmLVHaypHP1HNLD4Gb+WJLZK9j3hp/lbzdRNssbXNaXDWeTaptqXUHZ7r3syrc0nyHTemNTZTUoaJMtjFO9dASoeJ4kAR7B9yyWDAxupJDCzqmz3KdF9k2l80fs90rnUmYSVMLzZjl/eShlO/dv3c8MB0+6MaiOg0HnHY3mmZ02jpKbPMmSCGXMjmUzCaZpVu4h/wBmFKPtt5WviQUusqzRPZNo2um0zkeeUlTLUKtPmtBHUKkiMSWUupKllIUgdQBiJ/eusF40eRStD9pNK1y14lOa2l5xoO6DWmGYGVPZarsrz/JtU9sNRnddoLTbUVRSyStkslIXodyqgO1GbwliGN7+EuSMSDJND5Npztc1RkM2WUOa5QMoqq7KhXU4mQ0s8YeFwCeJERtu4cq6HEU7HJos87VazMBllPSRV8FfJ9TgW0MSyIx7tR5KAbAemLH0FmY1hkWRahZ4zW5ZkWY5FU3+26hHlQ/ANu97Sj0wy1zvge+6SBdA8DifaidBEyRjajGp4YKA9nVTp/TGgc07RM30vQagroamKio4MyTvaWEsFvI0R8MjEkjx3VQDZSWuM7Xq3s61HpnTOr9J0OUZLn1eGjzrJ8sa1OpveKdYekLGzK6Kdh8DBVJIxr8sV27AM1VVJ/8ALVNe3lcC37sROp0hqbLdOZPq6tyqRMnzeoaCkqw6MkkkbWdLKxZSLHhgLjkXGLkV58jnl5FHEU1EUypzwxVd9GtDQ0YgHwxz9tiln0g8ry7Ldc065RkmXZXC+UUUzQUFN3ERYo132gnxG3J8zid6105klN2NR5fDkOV0mfaZXKJ8xraem2VNTFVpIrrM+4iTbK0VuARtI88a7tjy05z2xZDlLR7jVZflSED9izlvyBxM6nTWaZ3qHXkIqsuqaDVuT/U8vjgqd0ySwRiSANGQCp3x2Pu3F74ofVvEUDi6mAcd+IHpU+St/TtL5Rd1kDdgT60XNMi7QviBU35H9cH2/fgAgIAuCSLG46f54RWUugdtpLDpyCOPzwuZthJRm6WBHBtbocdAsmiTkNmC7rG1uWufx/njfa10VqHs+1JmWk9U5eaPNMnm+q1sJIOyTaG6+Y2sCD6HGy7LtVPleaSRVkOWmly9Jq+OSenhEyzBQECzMpZRu2m3PTgYuX6RH0gs27WtHZXqSXKtG1Wc1NRHlecZ5Bk9O2ZySU0UbQSCp2hzvTcHuvWMgeE2xRfbSycQFpx16vNWW2a9EZa5alzWtyF3bT7gWJHvbr7YySQxI/ia4vxa1jb19cGkkDOztYsxLEjzPXp5YaVT+Cw9Dc+ZxeVVXz2gRZRluruyxaDTeSU8Vdl9A1fT09AkcVcWnQMZ0BtIzA2J4vh92h6GySk7YtGagyPT+URZDqCuWlqMugpwaJaumcJNGYrkBXUxOVv1dxxbGj7RKmoqtU9k4d+UyvLUS3kPrK4s5sxgzzVustIVU/eZjpzWkWoqFnPjkQ8ToD6srOfdgg8xjmXWmaNjH3j9pr5mgPkSPJbTYY3vc26PuHIVI8xVQzTOUaXz/tO7Uuzx9M5JGa+CtGTFaIXy+oi3FBTG/wCqBZlFhfgWxEeyvLMsoOy3tG1tnWQ5VmDU9JBl9Ca+kExgqZiRviuRskXepvz9nB4c/bS/0o63M43ARNR1EbWPDKzkW+ZtiT9p2WQ6J7Ia/TFMNpzvX2ZshtYmlpQvdkD0JdMWDLKHsiBPfuHh93Ic1B2bLrpKfbeHx68kXK9U6f0J2H6X1FL2Z6Mz+prqqopJpczymKScsNzhmlYFm44A4tYYjGjWyHtm7ZaIvovKMioZoI9+VZYDBSzyRKq9FIKCRiC20gnmxF7iTZPriv7OexTSmZJkWUZxDW1tRCafM6KKojjdNxDr3iNtaxI46jriqzmmq9Z9o+Yaq0Npx6PMJZ6jNlo8mprxUcYBd9qKLLEqg3BAW1xbE9nvvZJQkGrgCThmdVcKeCZLca5lQKUFQBjkNdNfirN1Z2r9mWp9C5/kGedmuV6Q1XlM6w5HLkuWClmYq5EkVWq2Vk2gglwZEcKQxG4HbdlOush1LpPUk9T2M9nryaJ06K9pZ8lWWWveOyEyOTfc3LFuecafMc7yLt07Pc81Jn2Vw0WtNK0qVj5lEthmECjmKb9s7AxRz41KbCXRhs1vYcRFpPteCt9rR0o9L3kwyR393NC4ODmgipwqWg46wQa47dSc1v74VAIIJrQY0B1aslsuzrUmTa+7Xq6tq+zzTGX5fJp2sX9E09CPqaSxweGdY7gLLex3Dz5xr+yCq05kWgNT9puodNUOpKrJ2gpKGgzKPvaNZZAoM0kdwJGuwADXUAMbEkENvo9kP2j1zngrkWZXt/6MC+GekXv9HTXEKC7DM6Fj7AmMA/jh8znMe6JpNKsGZyJNcc0yNrXNDyBWjtW5KdqGf9n+tNHZDrLJch0/pzU71U1JmWV5NGIYZoALpOYF8MTBlK3WwdXUlbqWM67BtN6dzbRUmXagyPJq6q1dmFXlGVz1lLvqKORKNyskD7h3bd8VYmxvsA87YjGusr7Osy7Fcl1lpbQMen61s1SiqJfr89Q86pTy94SHYoA0kZYWUEdMTTJdK59p6j7NanLKrKYV0sP0jXU1TmCQVDzTlZJDGjD9Z4DYWPUWxDa7Qfp7sbyDUipzwqacaBSWeEGWr2gigy30x4VKqzsuyzL9Ra9yvI88gaWjO6ergVzGziNbtGWXxKC3BI5ABsQTfG+zXWfZV2maQ1JS5ro/Tmj83yiAT6fqcsg+rSVTBuaaVFuJVZAfE13Vtp3EFhjZadyj9DfSfznLIEtFLJW1FOoH+6nQSpb5SDFT6U0XnOo4M6zjK6Ramn09RnMa89+iNHAH2llViC9uu1bmwJtYHFkSfUSl14tFGEY7SfLHL0UNzsowLtTVw4U9M1Nfo/ZTp2gzXONY6uyLLc3yvJoaeI02Y0wnhd6idEuUJF2Cb7eha45AxrdR6XTTPahX6EjcwRw52+Wq0dmIhM9kK36/q2Ui/na/nib6Q0Dm2fdjNVTZDV5bS12e5ylXur6wU6mlpvCu1iLEmS/HHAvhLteyOem7etMaglhCpqEZdXEhrgyo/cygEdbNF+44gZaq2mSjtRAGy6B73lO6z/uY7zdYx21r7UUko9T9nuZdos3YdL2P6ZbKQ8tCuYrG65ssyISZhW7t+8WJF7oSLFNpsOe6+lShzGsohJv+rTyQMyrYEo5Xdbyva/tfFvaORJPpduZVO39OV3l0tE+Kr1GBFqHNSoBvX1Q5v0E79COQcWbG9wfcJJ7rTtxNa+igtDWll8CneI8hRaz7TWJJ282Hl6fL3xuNKaS1DrjPqLSumaF67MsxZxT00b8yMkbOeL+Sox+WNLKYyb92qjyANrD28zi+fo9dtGednOXnUtG2nqCoyeogoqOpkoYKaWcMGd1mqgvekAKtipueAbg4s2q0fTR9oBXEc8FDBF2z7taKiBGSilDv32ttPisegwmrFH2kcqbBSOh8wR/DFn9veu6zWurI5pqLTsVFJTpmVPNlOUUtE1Q1QitI8zU6qJH3hgb8Ag2A5xWXCncrHpbkDm/Xj+vXEkMgmYJBrTJGGNxYdSs/6OtLlua63r4M1ybKsygp8pnmSnzOjWqhLKVO4oSPFa4DDkXxuKjOdGdqehdS5lSdmmR6WznTlOtdBUZNG8Cyre5ikTcUdSobmwZWCkEglcaT6NlaaLtCrphCkpTKqh9j/ZfaQ20+xtY+xONxq/Usmd9jX9otDaXyLTtDX5g1DndJllKUMW7mMbmZmKsdy3PNiALBjfLmkf8AVFoJzbTHAZ1qN4GxXmMb9OHGn81cMdVOC3OmNaZDo/sa03qDN+zLRuo6qvzCtoXnzLKkkqGEbFlZpCbtYeEC3AAxHtIS6d7We2eBptH5VlOWTQCY5JlqmmppXijVFjGw3USPZn2kEjdYgm4e6f12+guxbS9XPpbJdQQVuY10Zps1o46iON1bdvXerbWINja1x1xXFJVax1h2g1mp9Baclp8xLz5otFk9Nujo4VG6SyAWESre9xttcHCwiR7ZSCQauAJOGZphXCm2ibJcY5lQKYVAGOQ1091NdW9ovZ5rDS2dZdm3Z7k2lNRZVUKMkfJcqFJLLZ7PDVolkZdqtcv+sR9tiwLjCerKagXsN0LmaZLlcNdU11clRXQ0apVVIUHaJZQbuBfgW8hjYZ3mmTdtehM11bW5PDl+sNOU61VRWxCy18IHKy/ttsDMjm7gpsLOrKUa6tfd9HrQMZsduZV/T/Df+P54VsxvxNBI79CCa07pOesZFI+MFsjqA92oIH/kOGxONfLlmWVHZLJl+R5XSnMMmopq9YaUIta5qApecA/rGKghjxcYi3bNU5fp/tpz+bJ8jyuGho81SVMvipttLsCo3dCO5tGTcFb8g43vaZUB37I9oF4sko1F+RxU9cRvt8JPa/qw7eDWAiy8f7JMFikke9gcT9rv9QS2ljWhxaNY9FYcnaNlEnZ2O0H/ALHuzu8mcnL/AKomQR2C7dwIYk+YItbpjR9hkGSVWrdT611LpfJZcnoaczPl9VRCelhaoqFVFWMkW2DcFsbgYZZfEG+jtFIwNhqpenX/AGT4kGhtL5hmfZDmkGVV+WUFXqHNY3EuY1PcRGmpWXwh7HkuDwbCwOI3uIheLxFXXa1OArvOyqe1oMre6DRtaUGOCiee5LNpDtfn0s1PAY6TUX1JopIw8ckPfgqpU9UZGXjzFsWJmOqtBZ12h1vZNnPZJpSGnlqXo4sxyqi+p1ULFdwlR0bwsvUBtyG1iLG4jvbRFKe2PTuoi0cqZ3+jqmQxNdGnidYJQG8+Y1N/Qg418ySzfSXkkiQ3OayG1vIQXP5YDM+WJslTURk56wgRNY8sI/np5Jj2cZTkeRT6xz7UeQU2oBpKEiOjqGdIJ5TK0YaQIQzINt9gYAlhuJC2O51VVaY7Seyuu7QqHQuTaVzjIc0p6KdclR4qStgmFvHCzMEkUlCHS24bgwJ2nCGmY45cg7aJC1/1F19OKtsNNMTTD6PWt4hGoQZzlrFm8/ElwPfpf2w90z7xfXJzBuobtcPMpojYGhlM2uPmCfhVzJsYBUJsbXJ4ufh/M4SIsLj7N7XsbX9L4H7KW2lmFrdLAfvPw8sAzFVbaSFPXn8j6/HGws9Y1lW5JAIsbHj5n8+MbbT+QZjnOZw5Vl6IK2RDKzynbFRQAXeeVuigDn248yBjMjyXM81r6bLMtpDVZhVDdTwE2SJALmaUnhUUXPPAAueMXdoXs1ynNstq8kObT0Wk8vmhl1dqSKMGqzWpa7w5dRK9g0j2JRG8KqGqJbIgBo2u19j+6Z9x4NH9R9hrPmRas1n7T94/7RzOwdYJLs50VkmeZdHXVtHVS9m+Q1xUU4cwVGr83VNzRlx4khRCGlcX7mEgC80qA2Dm+dZvneZVOc59Gs1VVGONTTRdxHSxxpsighi+zFDEgVIkWwVAOhLs5c5rhmc0NNHlFJleXZVAKDLcro5HlpcupFYssKSEbi+4mV5XAkklLyNcbVXO8oYMjqdRalzyuyjIcoCSV1ZHZpWvfZDCjGzVEpBVEPA2tIT3asBzU8xkIiiy5knX59bFtRx9mL8mfIDYtLmGYaf0fp2fWWrKA1lFTVBo8ty0vzm1cAGFMLAMIVBR5mFiEKxizSDbzxqnVGc5rmVbqrUuYNXZ9nE5qJZWAtu6AgDhUSwVFHh8IAG1Bjea/wC0Gs15nDatzmlTL8vpIzSZNlEMhMdHTXLCJWPLOxYvJKfEzO7nxOAK4qaiavqWqag3d/ICwUDgADyAHAHljYsFiDPuz1n2G7btPks202kux4fPwiIrM+9nuTyWJ/PDmOIH7AJuSPQk4TRbMBYn29cOVjAsZLIL2JJuB/HG0s5GSwAO23nwP3/PB9yqNvTaSQPO3rf8cF3MWY3563Nrg/E9MG3PYR/rNl+AOQG9x5YVCNyAoRjcMXA+78R5G/nbjBdrcgHxC5Pt7g+fw8sCWIW1gtj1tzfyb4+lsFDs6E3RwwIY7bH4EeXrhEIUAQ92o2lBzZgCPb4exwRn8QuSAeGJFz8+P4YMOFtvUE8Dp+70PGEyxLDaSfu3HXjy9R8MKhAAQTuuoQm9h0BNuf54MhcNxy+087evw8uPXCRI8LbRwL3Jt54XKEixY2I9re3xH8cIhA23d1BJ6lXuD8OOffCbXPJF78EkfgB/lgZOn3bDobdPXnr/AEcFvY9LC3HmLevw98CEN7BjcC/vb92MN1It6dbYEnr/AHubX8/e/lghbyBAt0so59vfCoRN1ul79eeT74zAFrNuPQ9LYzAhGNiSb3N+eensMYGAB4vYC/HW+AbbcqzAEeflfBrkMQoYEn5i+BCKdvQFRfkXwPBJ8PPuefy64Abr7b7STY28/wDLGAeHy454Fz/0wIQ7bC1/K3/L7YEE3Pl7/L0xijqSSVJBN+OTjAehsb34uP4+eEQgFyu77IH2j6nyxhPiBuR1Hi/r/rgSlgCF8rdT/wBDgLXULyevHHH+ZwqEW3UkN048Xl+F8AVJ5HW1zbi2FHFm4JJ8wep/lgGB4JaxW20lv6tgQkrdCOOb8j+OCWubAAgAnjywq1ug29LcdL+18FN2HU2H4j5YEJBh04+WCnrY9OmFZCDe5IHkuCMpHTocCEn54H5YEj8sB088CEtBOyEDeV2m6sOqn1xb/Z5r2nzali0nqWqhp3itHQV8pISDkt3UhHSFiSb890xLgFC64pk4dUdS0MisH2sv2W9PY/z8sU7ZZGWplDmMjsVqzWl0Dty6EeGbKq2so5oys8V1lSTgpIBYbrdCOlwSCCCCRtOJ/oXW+VQZL/YjWv1v+zE08tZTVUNMJ6zT1dIu2Wrp4jxNE6+Gqo/s1EXjUCRQTTegdVZdqJFyPPKtaPNQiQUNdPLsgcLwtPUE/ZFuEl+5ezXT7ErkiqcsqZqKuy+ZKiCQRT08pMbRsvkSCdjrwVYXtcG5VuedDn2aW677hwI+PRbNGzsqMvT8pXtQ7M8zyHOkraSSkpa+Ckjq6KvpKj6zQVlBINsU0czf+cUEg8KSt4ojeGcAqDim880+lVHV1+U0D0VVQX/SuTvffRkGxkjvy0N7e6XAN1s2OldFayy2CBNC6virm05PUSVlHPRxJJXZDVyCz11AjeF1e1qmhb9XUKDYCQKTGO1jsszfIq2lzOnqKKGpipBmGT51lUhkoaygB2iop3IvLREnaysO9pGJilXZYjas1pFKj7eY/CzbRZyTQ/dyP5XN0khdFZpi9h9m5utsITbmFyOevxHriV53kjV8lVPRZX9QzSkBfMcqVbBAOTPAPOI9SovtvcXXpGFAKgXFr9Op/wCmNMEEVCziKZqXZ72kTZ3q/KNYyZMkbZXBTwCn+sMwkEQNiWIuL38hxbCGW9pNVlOva/W0WTxuuY1MlS9G0z7AWlEqjeOfC6gg4jbqLA2uDyB64zu1Y+EiwHFyL/PFcWSENuXcKU8lL9RJevVxrXzT1M6ePUo1NJRKxetasMG8gHcxbbu6+fXG21Xqs6h1JW6skokglrasVjU6uXVGup27jyR4evviOHbdullsftD8sJu6FgSRwLHnriXsmXr1MaU8lGXuIp5qz3+kFLVVSz5p2baRzR1AjL1+WR1Erxi4CNKy7yADYG9wALHjEd0R2myaLoK7Kv7NZZmkFdKHlirYxLEQALKUYEGxUMD1BGIgIk+0u0m3Nj0wASzX23GIhZImsuAYYazqy1qQzvLr2vwGvyU3qu1OCbVmRaspdD5JlsmS3Dw5fCKWOrO5iC4jAAIDbbgXIAvhtlXaZnWm9a1utMtgjjlzGeWeanbxJaRy5XxA3tcjkEEEgggkYiTFRYL9nqP88A7iV/Eb8AEnCmzROrUVqKZnLNAnkFKGlDXIKcZ52rnOcqzXI6LSWV5ZQ5jTrH3NEvcpHKJUkMoVRtJIj22974eZL21tlOl8u0pVaHyHOKbLlPdnM6cVNnJN3UMLISCAbfsjFfhVC2AU/wB0+YwOxFPhWx8r84b9JFQNpka5mtcq1rVH1D6k7cMhlwU3o+1ODKdfSa5y3ROV0Cz0gp/0bRM0FMjd0sbSKAPDuILkDi7Hywy7P+0iv0E1fAmWpXRVyWCPMyCN9jpvFgb+FyCPYYijMCblgT7nzwccAs6kk+Y4wr7JDIC14rWlc9WSRs72EFpy91KNDdpub6FgrcqSipcwyzME7upo6qJZI5VNrgqwKsCVU2IuCqkEEYR1Tris1Y9In6OpcvoMuH+rUdMu1FPFybAc2AA44HxJMdtEbHaPhe/z5xjlVFkNwevH9fhh300d/tKY9aknbPu3a4KfZr2wTZjr6h7Qm03CslHRJRxUpq3ZboCokLEXPDHi1saPQ/aFmWiNaR6zFElbKryO8DzMivuO7qLkWNj8sRxrfeI+ZwDogKt4PEegIP8AXzw0WOFrOzDcKXdeQ1c0/wCokLr1ca180/rainrKqprqWjSjimneZKZHLCCN2JVBf7SgEC/XgYaSHxWIJJ6gG1sG73gMxBCnoLcj0wQFCLBgLAXuwuff3/hixSigSD07M5Y25F+OmF6aEJuB5Djn2PUH+vfGMyDkH369f8sGTZceNTfkWIPX+OBFUeQMNrMGUMAyg+nnb1w1qR4Wt58i48vb/LCzSDaULlA3XaeAfUj+XOE7BSAWWxF+HBX44VAUyzjtMfPs30pmzZGsA0xT09MIhVM31jupA99xF0uRawvbCtd2rV6dptd2kQZXHE2aSyyTUKzvsCugG3f9rgqrA+oGIbeFIyd8fw8JJwBDVb2UhwvQdfmL/uxV+jhpS7hQjXkcVN9RJWtddfNP9QahmzzVddrUUqU0tdmD5h3CuWEbNJv2gnkgHi+Jl2jdr1f2rfo6WvyqCgTKkqEjijnZy5lcMWJYDoqIg9lGIC0TEBALj8sYYzGO5O0AH7y9P8sP+njvNfTFuA3VwTe2fdcyuBzVgad7eZtP6Uo9IVGiMjziloS7RnM6dKle8Yt4wjqQrWYi48saRe1bOKLXcGvtPZNleSVEcMcLU1BTiGnkCqFJZFsAWsCSLeIXxFPq36yxQlr9Bzf4Wwu0AsFMViBwbc/54aLJE2tBnWuJpjnhWicbQ803bhXDfmrF1D20Vmo8iqdOZJpTJdO0uYsXzA5fAEM5JBa9gAu7aNx9BtG0Egx3S+uZ9H0Wp8op8tWrXU+WnLXkaYoYFJuXAAsx9jxiP01HWmllzGCkmNLBIsclSIyYo2a+0M9rAmxsD1scY8LxncylT9oEjrhW2SFrDGBgTXMnEUxrnqCQzyFweTqp5eC3Gg9e1GgNSVGeR5YteZ6KeiMbTGMASqBuuAenphXs87Tc27Nqmup1y6jzTLM0hEFdl9YgeGdB0uCCPxH5gER1ljJvtt8OR/PGS0l7tKlredrW+WFfZ431vDOlfLJI2V7aAHL3U91D2xT6vyhdOSaTy+hymkzGmrqekpGMSxRxrKrQAAWs/ekkgXBHnc41vaZrao19qI6rNCuWssEUMNNHK0ixLGOLMQC3Jv0xGo4lWzxgBW/f6YGpmADI7HpYFm4Htb098NjskMRDmDKu3XSvGiV88jwQTnTlkpLm/a9nGY68yvX+V5VHl2aUNNT07iOV3WoaJdgY35F02CwPVQRziwu2jVmaUOlKJMnyDJ8tg1xE9TmE9HT/AFeeSaFwJ4pYCqlGDkAyDwSWJQC7jFFSIrRjuwAV8w4xsKjM80zSYVud5rPWzRosSvVVRdxGOiguSbDngYidYY+0Y5owb46suBxxUgtLgxwOZ6PFSfUGt/7Vaa07kEeULl9PpumenVkqGczltpLtcCxupPn9rAas7T8w1BprSGSplaU1XoxZFgzFJ2aSZWZWG5CLCzLfgnqb4iAmUgosqW5+8Of88ClmABKFbjjeAT/LEzbLC0ABuRJHia141KjdPISSTnQcMlZn/wAUTnLwmtotJ5TBqWSl+ptnaRDv+7sB1tc2AFrk2AANwAMVyzO/+1ZmcmzXJuzdTc+Z87+98F2RQvuQhb9LML/D44NYbiQyv68i1vgf3YdFBHBW4OhlnqSPldJ9xSMgLrtsAx4JA5J98NXglZLByEDXAJ4v58euHz7RtHU/aIvYc9Py5wQgjjdc9BexxMowUlTL3SWsbg34BP5jocOI2j6uSA124AJ/DjBRHtF3AuGupuB5eVsFNlBk7xLk8cr+7y+NrYELdaJ1w2gc8q82jyla56ijelRXmaNU32u3A54HTjD3QfaINL5BqLSea5CuaZZqKmELoagxNTyr9iZLAgsrBCAePD74ihjErBy28HrzyD74MVC9CLdOOMQOs8TyS4YmnLJSiZ7QADt55qeZD21Sac0vS6Tl0Rp/OIKSSSVWzSlWrUSOxJdY5AVQ2NuPLGtXtbzeh7QIO0LTuRZVkdRHFFF9Ty6mEFM21ApYRpYKWKhjb7wviIOsW6x2kgm9iP34PHHdgSycHpfkf5e+GiyQtrQfdWuJ154VSmeQ0qcssBqVkal7Z6vU+nqrT2T6QyTT1PXHfXNl1OsRn5BNwqgDcVXcfQADaCb6/SHbJWaY0++j840tlGosp776xDTZlAJBDLz4lv0PLdLcMwN1NsQsrHACyShTfgX59rev9XwmiJICwAuPy9/69cILHCGXKYVrma121rXmlNpkLr1dVMhSnhkpDqftAzfVOp6LVGZUkKx5eYlpaKI7Y44Y33bB5i5uSbdTwAAAJdmPb3SZ6temY9k2kKqbMIJoXq6ihSSqUvGUWTvtu8uhIYNe91GKwYxqNjMtxfgWJ+eB2wb3MBAF/CGKk29/LC/SRYEClBQUJGHkUnbvxqc8dSllFrhz2df9nIylLJmZzQVxqGDEhSNnd2225PN74Zap19PqHS2nNLU+UR0NNp6KRA6zM7VDuQWdrgbedxsOPFjQb7Ag7Bci5uBhMqOb2N+QemHCzRNNQNZOvM4EpDM85nUB5BSao7SpKvINJZRUZJG8+k6p54ar6w9542dXMbL0ABRbEe/ribQ/SBzaWCrlyvRuQZdmNYjJLmiUqtVuCbgtJtBciwtuNrhSQbYqQUqggMFANiLsBa+FEkSM7FY2BIuPPDDY4HZjbrOs1OvbqSi0yDI8hqwUg0n2gZ32f5tU12VJT1MdZG0FXTVcYmhqY2vuV1bhr3N79bnD3U3ajmerMpi09R5LluQ5PHL3xocti7qJ5AbgkefJv5km3PAAiFu8QblPn7g/ywpt7mMBRYMB7hhh5s0Rf2hGPWrJNEzw24DgsZy9txFkWwH7I9vbG2ybIczzbMafKMto/rWZVPiigNgsKAXMsxPCqBzzwByfTCmntO5lmmaLlmXU6SZgVMzd422GhiHLTTueFCjnnp53Nhi4+z7s/ptQ0GYUeS5nUZXo2hljj1Pqs03eVWZzt4koqSIkd5K9iY6cG1h3sxVBiC1Wrsj2cf3cgNp9hmeJE1ns/ad9/wBvMnYPnVyR+zPRMMktZT5TnQpciy94jqjVrUvftUSvcxUdHCbGeRyD3UHG8gyy7Yk4m9VmcOa1FNleU5V+isgyeOSPLsrE3fGmR2Bmmlk476pkbY08xHiJjCfqQFDnMJzVZbTUGW6fGR6fyrfFk+Vx1CzinjcfrXmk472slIBmmYAP4VQpGgQoZX+tkrq+KWmpo6GheSrrqsslLQxB7GpnuN21Dwsf2pJSsagkNbmbROH1ZHjXM6yfLkOC244y0Bz8KZDUB1mUlNlFPXVFWma1tPRUOTUkc+Z5nWQ7oaCnA8RcDxM28rGkYO+SQhEsNxSke0ftATXNRTxUVPPlmjciLjKqGdg80rMbPV1BHElRKVubeEWCLaNOXnaf2oDXcaZJkkU2W6Ny6b6wEnNqjNalRtNZV26tYlUjF1iUlEuxdzU+a5tLmcixoNlPEf1aevluPvYfACwGNSwWFzKF33a93hv9FQtVpDsstW/f4ImY5hLmdR3jjZGnhjjBuEW9/mT1J8zgkcXnwAfPBI0tZrW/hhdObePoevn/ANcbzWhgutyWW5xcalGC2Atzc+V/lhRB3dzcdLfZ6+3t8cYE3AnaB1cW9B5D0wYC7AlubkmwwqRB0bb4h6c3+XwPpgyWsx48RFuD5eRH8DjBF9gKCLm1rXv8z/DBxGVNgxVj0sSCSB0thUIrc9UuxNhcXsfY+ftjAbclFZmAPU889Lj8LYMIzYX3WY3uWuD7cf8AXGHg2JUki+4EE/AkYRCId/i2uWDE3AYG/rf0PtgOWYE7CbW5XqPIH1wY7zZVIN+AjD+X7x0wW5f7IUj2HFv69cKhCHbdw7HcegNrn9wxl22Eel7C3LH04/fgnNypUdeLj+h/0wpztYi3lfzt7g+R/fhEIhADgbmJU9eL4TDEePgkNf0/64UkG4lACx/ZGCG97M17ftDafw6YEIDbzJ46+XOMJFyNot+Fh7YO/kVBBF7C5HHr/DBLEi/Swvf0wIQdfCB8j6e+MwUIrSbdpPU26jGYELH3ITu8JB+Y9cYvuzHysD19r4BQQ20g2JvweuBUqAL7rDqL8f5YVCFFN9oTcfTp8RjPLhrWG4G3ODdetmAHPNt1vTAC5+y4VvLjkf5/vwiEZOSWPmOigAH0tjGNj7nzBuSPY/8ATBRcLciy3t7g+ntjCTv8Q288i2638zgQgsQVXaFNh5dD8/xxm7i/IHPNr2/jgSLACwHna3GCncNzAjwi3Hl8Lf0cCEIIupG5egu3Fj68Yx+bk8X+1c8nnrgt+eQBceQH78GVhwQxB9QbfIYVCKSeABwB529fzwQjk2Hv8PfByRweF69RuwDKDwAALeQwIRSONwPUfO+EmXj4YUYttH2bW6kWuMAQeBsHwwISTC3l0/q2CHnnCrADkLYdBghHPS/ngQi8dMZ5Ywg2wGBCdUtW0bL49rLba17fI/wPli4dKdokWosri03qSYR19IgioMwka1o1/wDfecjkxgfZblovK8ZKrSlsOaWqeJ1O8qyEFHBsRbpz/HyxRtliZam45jI7OtY1q3ZrS6B245rpN4a2lmWizCOWGSBl72IWEgPBJDC4BK2IZeD4WBIN8TXR2s6LJ6CfTOex12baWrK16s0sciRVuX1n2fr+XyNdYKzYRvjN4KlCUlHRhTmge0Sjzemg01quo7iSFRFQ14jLNCL3EbqOXivc7BypJaPzjaUSU9bS1T5fOqEuu5djB0lRrlXRxwynkqw9TxfcuOcaZbNLdfg4cCN2704E7REc7KjEcx+VuO0nsjelOV5jk9dHV0FYJJdN6hyqFkiqe75kSGI+OGaP/f5dJ+tiO5ot8ZANGZzkMtdWPTmlhpM72979Xgt9WzKM/wC+piONx5JQcE322IK46I0Hr2r0jFVZVnGXU2d6czsxrmmU18jxU9aYzaOTvF8dLVRkfqqyPxxMNrbozYP+0zsl09qTT0mr9O5hNnGn56pYv0hVLHS1mV17i60mZqngpKw8d3VL/qtYLElHN8bVltVB3cRrGsfj0WbaLNXPPUdv5XHzloELhQpRuQR5g9CP688SejrNPV0AqSuS0RdiDBPPU7kI87qpFj5fhh9q/IZkqny/U7fV6+OQ08WavEYxM6/7msTrHMOPERe1r7hZsQiuyusy2oairYXgnj5KnkMD0YEcEHyI4ONFze2aHMdTwVBp7JxDgpgf7OhWP1vTTEkWDS1Zt6n7H/TGINMlt0lVpcADhQ9WLn1J2X+WIKY367j0wpHC0hEe+x6KSbA+2IjZ3jG+eXwniVuV0KcX01tKtPpQ34uJKzgef3fTjCkB0VE+6Sn026X5UV+YKxHoDbj44gqwlX/WhjYkMOhH+eMFPvBCScjgejYOwf8A1nl8I7Vv9IU5NRobvS36L00UJJAatzE2Hpfzw8Wv7PVlLf2c0myg8XzTMwCP34rx6Pa4VZrAqrXb7txex9/8sJGA8frAp6WY2wogeR95Sdo3+kK0Xzbs4YIw0hpGQ2swOc5qoXnoOefXCMufdm8LqDoLS8oIuTHneace3JxWgh3XPeAelz1wb6tdh47jzIN/zwNgcDi8oMjafaFZ9Dqfs1pxum7OdKzMQRZ88zQWHpwbYcjXHZGIVMfY9p6EsvKjUubk39/FbFUfV7i4BNjY2PTBBTt79T1w8RH+opt8bFbg1z2MxqrTdjGnpjbxW1NnAJP44Vh7Q+xEbhP2D6fYEeE/2ozgkHy6N8sVIlH3lo0vvP2QT9o+nxwApAxKjdcD0wdn/wCRSXhXJXIe0HsNQBYew/T1nIbw6pzgbTbzDA+uGM+vuxcxuF7EcmLA+HbqjNeT6+n44qcUwJC3IYXvzx7DCgo1dfCHuv2j1BuePhgEW8oLxsVpjtB7HpOG7CMslJ6N/arMgf34eUWt+wxpmWu7FsqpUETsD/anNGLPbwrYXtc+fS18U4aZVPUced8YKccncPxwpiJFASEgeAakK5Y9bdhcjhJex7I0H7barzZgPkEucZBrrsLMak9i+QK7N4t2qc349+FxTi04VuQT1BH9eeM+rBVP3jbk+Q+GG9if6inX27FbH/aN2KM5Udg+TgoSAf7TZoQ3v8Pzxh7Q+xAoVfsFyq/qNT5qP4YqmOmS25h4r/Ij+H8cLR0Xei0EMrsOTtUtb5AcYd2W8pL+5WPJ2k9iwui/R8y57cBv7UZl/EjEbfV2lHqppaHsuyiCF5C0cT1lZL3a34XcZLm3qcaOHJaqapWBKQmaQ2Ee4b2Psn2vyxZOnvo3dsOf0iZjTaBz2Ohbn63UZe9LTgf+nqjFH/7LB3GCpPNJi7ABV/mWcisqpJaTJcvy+JzdYYYywjFugZ2LH15J6418TySzbRGJpGsAoQk/IL/LHRGnfoyZLMw/tf2jZFRG5X6tlzSZ1UswF9oWkC0ytYdHqLYmGW9n/Y7o3LYM3odE1+eQyoJBWarr1oKMX8zSUxRePNZKlj7HFObStnhwrU7lPHYppBUDBcz6f0RqvVOaQ5Np/Iaqrr5jaOkpIZJp2v6RR7n+ZAA8zi58g+jH+hoxJ2n56uUVPdvKmTUUK5jmkqpbd+pjb6vT2uLmeUsL3MeJJqHt809lWX1OVwajU5eZgYcm0pRx5ZloTZykkiKiTWcXDMJiQbEHriptSfSQ1dVUy5Zpaf8AQdLGHVDQOxnAYWYCocBkBHBESx3HBuMVDarbaxSBl0bT17U3qw2CCDGV1fBdKaU0D2d5bHU6Zp+zKWRainWtd489aPU8aLde9FwYWiuSGpnpe6Nr7vvYh+sPov6f1KXqOzHUeX5hO9//ACbJ3eSZpu81+rzOaGpb/wBBPGx/YHTHJEec5rFXnNjUS/Wt/ed9vbvA/wC13l9+73vfFpaY+knqmktR6xgg1NS8KWzBitUFHkKlBuceglWQYkEekLKKtcJBsyPkceaQusk2BBadufFMNW9i+tNH5oclzPJ6qLMUHioKulkoK/jz+rzhTJ/iiLgjEGq463LKhqWop3gnVjugqImjcX6+E2I/DHXOmvpO6ZzXIYtOPqRoctLgHINY0UWZ5SF8xGzLJHH14IWA+d74kc2meyfVeVvW5to3MMqoWP8A59pbMY82ysAnr9UrjKqC/lHUofYYkbpaMYTtLDvHodaabA938Ih3h8LiiDP5KKMLLkWXzFXDb5YixI/ZPNiPPpcHG4yLWWmI80E+o9C5TmFIIpAYGaopgW2+E74HDjn5c846Z1J9FDQWaQP/AGR1rpionJ8EM9XUacqifIGKsE1G59kqF+WKt1T9EDtf03E9bNovUiUI6Vhydqyl2+pqaBp4re9xi6ySOYVjdXwKrOY+M0e2nktBF2ldjkkQab6PuRFv/mjzUf8AuRwtB2odjKDxdguRr/cGpM2BHzucQ2o0BnlHKYI4KOrkQ2MUFXGZPnE5WQfNcR+syuegcwVtFNBL5pNGVPx5/eMP7MbUy9uVqHtO7E3Jv2B5GL+bamzY4Mnaf2KIFH/YDp9rDm+ps1NzioVpIrfaY3F7ed/T/PA/Vk6bDcDyN8Fzei8NiuAdq/Ynfcfo8aZNv/whzW5wf/ta7GYpGSb6OWnAAQQHz/NBcEeVjinVoWku6I1gOQoNvxweGgp5ZdlRM0Y23DC3B9Dci2ELABmlDq4UVtTdrfYt3xCfRx04VHQpqDNAP/ZHB4+1XsXmp5mP0dtORMliN+pMyBYH9kedvPpipHoaKOQRQu8gNgzPtHJ9CCR8ThOakUMERgVbglefF6XHX5euFDRTApCccVb79p/Yk0cd/o/aeUsSG26lzTw2HBPHQ9PP3wmvaR2MTyBB2Gadjvx4tS5oFHxNsVN9UkpZDDJEyyDqHANsKNAkbb0cHcu4MAL89R8b/wBc4S4TkU4OAzCten7RexGFnLdh2m3fcAp/tDm1reeDt2mdicjsj9hOmFFjZjqDN/TpxinrFVYcDcu08DClPl1JIf1tYYyVvtMTMb+VrcH8sJcLRUuPXki8DqHXmrZh7R+xLahPYZpwWtw+pc3JGFf+0nsWJ79exLSoMlx3Rz3OB3dvMm9jfFQrlZYXVTwLm7AAD44yPJq2YgLSuRa9+7cC3xtbCFv/AJHl8JQ7/wAfX5VtzdqXYxTRbh2B6WrNzWCx5/mwZR6m5FxgydpnYrGspXsI0xHKHUIRn+bFXUg7je/HkLWv1xUSZTUykiCklkCcttU8fy+eF5MirUH6ylkQBttytiSfLk4QgUpeKN931Vpv2k9jPBTsO0uPZdRZwP44Ie0fsXtc9h+mmYAEW1HnFjc8jk+WKtqcilpYBLMkiF2YLvWwO3rYg2Y88jyw0koUikaNpA4U23RkMDx5HzGFa28MHFITTUrfi7Q+xiRZL9i+lYGt4d2fZw1z8j+/DlO0TsTjhXf2M6RldbAk59nAJ9T0xSq0RZgFJdjYBV8z/Xl1wpVZcaRjBM368fbQG/dex/veo8vjgMZr9x5fCUPFPtHP5VzTdonY0q93B2RaPa56/pnNrD8RjS5vrzsn+qVcdL2PabkrKinkjppabNs0P1eVhZZCsjBX23JA8yBfi+KrNMwNtxGHVJQtJLZALqNzMzWVF8yx8h/QvhRCRm48vhIZW6mjn8oYFcskcaGR38IUckn4Yk+k9LV2eVMy0M9NTR0SiTMM1qG/1XLY/wBot96Q9ABck8KPPBMh04Mwpmrpah6HKAxSbMGT9bVEfaip0PX3PQdWI6YvHSHZfHX5Hluf66oJso0SifXci0xTTmGtz832iplktuhp2YbTVMu5z+rp1J8S1rVa7oLYjSmZ1N+TsGrXsM0FnvEOeM8hrPwN/BazReh8szzId++vyHs0hrO5rcy2KM21XXR8mnpg3huLgkk9zThg8pZyqNNs5zysrfq9FlmU0dBkWUwNR5Xk9C7Cly+nexkC7vHLM5s007/rJWF+EVUDTOs+q9SZilTm9FDFBSRLR0NDRRd1RZfTKD3cFPECe7hHJAuWe7GQtId+BNNlUVPm+bZjnZy7I8oqWjrM0mjJWMsSyRRrcGpqHG4LCvJHjYom5sc5NN2n7qKtCcdZcdp2nd5DUFssjEffkz5AbAnIlrKySaors0GX0FDQJJmNZUSNHS0sF7CeQAXvJwAi3dpAUUG5IpvtH7SRrSH9B5Kaug0VQz/WFjqHtU5tUqNv1uqK+YHhSMeGJTsS7F3LbtH7R59dCKmSOfKtJZcwegy6SbfNUsBtFTVSC3eylRa4AVR4Iwqgk1fmWaS5hIFHghW2xBwOOht/DoMaVg0eW9533en59OapWq1XsNXr+PVHzbNnzGQRRjZTofAgAF/IEge3QdAPndpGnqMERLsB68YcRqAv2gL9L+Y/hjeYxsYutWW5xealDH4VvfgnofPCwJvZmJ9fPn1/ywRVdbX63sR5/P2wuoYMrLfcPEoJ4t7/AM8PTUKbxddvHB8XP7unxHOFUFzcgLbkEgNz8PX+jgqg3MKK24i4Rxcm3PztzhQnYLyIBuAIO8Ee9/b0/DCIWNdmJa5BP2WN7+/HT3H4YyxNrMSSTYACx9uefzwBYcnab9bAWH4+3vxgQVKnfbk25Fh/iHw6HAhYSHJsLL97bYG/uDbBXAvYsBxxZftW6cD73v0PXASXI5sStiQ12IHqW6W8ucZ4QQSiCx4UcXHnfz59fwwIRQ/hZACu7kgNcNby9b/Drgjd2Rcjq3F+bnAbG3btqHgmxFgfW/n/ABwBIJLNwD72+AwIWXPO1QBYnra1+nJ64U23AG7oOOOnrf1wkSTZh79R5+Zwa4PW1rXAVuD5Hk4VCEg2F2LW4Bv/AFb4YItrEbSSL8efw/rnBiQD9mx6ng8+nv0wRla1gy89ObWHsP4nCIRrnarDhrgqQev9fzwVuBu3cjjzFsCSSNpVSb2A6G/wwRrFCd1+bED9/pgQilgWFmNj0vx8emMwZQruLdPcfljMCEUk83Y8m1v2hgfwuCeL2/PAEtuv9m46dPzwKlQLbTa9uDa+FQha9iGRrgdD5D1wUFgWA8uD7+2DKALWLXBtwbW+BwQLtNuLHi1jhEJQ7SAWYm/NwL2Hv/LADbexPFrdPfrbGACwPJsOuBKso8JIYdD0+Y9RbAhDtFy24+x9D7jCYFxe17c39B/DBydvAINwLrcjp+/4YDgjdsUkm/IwqEBIBAPIPXjn2JxgO0+EC34g/H1wK7y21X/w2a39WxjHcxN+vAt0+AwIRfFtF7Dnqeb4z7JHJ5HHlYjqBbBhwQbDrbpcn2PtjAiF+SPU3/LkYEIjKPtcgHq3vgpHAABv6A9Ph/LCtjc7iCx5seOfS/lgrqbC9zbi3rgQkrLc2PXr8PXCbg3v1HXpbCxBvbjjz8vlgjA9QfhgQkiBe18FsPXCljfp06fDBSCeffzwIRb4DA9fbGYEJanqGjIBJsOh9P8ALFmaL7SEgp1yTUZaSk57qoU3kgLdWUnjnzB8Lfes1nFWdMKQzNERybfuxUtVkjtTbrlZs9pfA6oXTLI+XUC1NSBWUGY7no6unb9XJIoAcXblWtYOjWdSFbngnZaH1rqPSWYvnmmZxFPFTtSVlM9OKunraF/twVFM3hqqZhffC32ftxlSMUv2f9pdbpwtldZDHXZPV7VqKObcY3t0It4lZeSrr40+7uW6G16Oly40Dag01Wy12WNJEzq0oSooGJupkKWHJ4SZDsci3ha6Y5iVkthko/yOo+Ow8jq2LdjfHamVb5jrUp3qLs80n2p5RPqHs0y6VK2koTNnGkWlauqaKjTrUZe3280yxf2P/O6T++otjnvUelnyiijjnhfM8haMz07QSrNPRITzLTyjieC/UeXRxG3OLUocxqMnrIa7Kauto6ujrP0hQVdHOYKqlnJv3kMosYpQb88XHDAg4nK1On+1xZP01FlGndZ1EweStYigyLUtSeA0+zjJs0boKhQKaZjZwtzfTsluvHYdmo+Gw9bVRtFlAG0bdY8doXHOZZJU5dFDXxyJV5dUMVgrIeUc2+ywPKPbqjc+lxzhiFF/JgR8Di9dbdl+eaRzLOVo8skoKihZafPMnzWlEHdsfspWwg7YWa946mM9xJcFWjJ2msK/S61LTnIYahKuAFqrJ6kf63T25Jj/AONH53HiA6gjnGzFM2UYLLkhdEcVoV+rtCwmdkkUXR7XDD9lv4H5HjBWVo7q67DYGxHIHkf+mEyd12Q3BJIwtDUKV+q1CGSEE7bcPHf9kny/ung+x5w8gtyTM0R3gaVhArleq7xdrW5Jtx6/LGd8QpFlIYFR4QLHyI464GemaHa8MvexKeJUNgD5Ag8ofY/K+EpTdGJAvyetufhhwIISGoT1MuzILuSlqEDi1yoW/wCNsY9HmCSbno2LDg3tz8bHnEg0V2ya50QBBlmawtDt2CKvy6mr4QPQxVEbqflY4ufIu2zNdWCGKko9I0uZyoFNGdIZKVqWHV6dzSeIn/hNZx93d0xVktBgFZW4bQp2Q9qaRniueDSVt9r05b1UlR/HCv6DzVYw/wBRfaPvF0/92x0i2vNaXeKJdLiUE7I20bkw3EfaT/zThx1A87EdeMFg7RdWVBKLHpZd9grf2NyVuvnb6p09wfzxCdIxAVCkFil1rmwZfW79q01j0szoQPz/AH4cx6az4QT165XI8FKFaV0kQrECbAkhuOeMdIv2ga7hKzodMHu7LLD/AGOyQMnoVb6pZlb7rEDnwttOFKftN1vJWPTmXTOx0Vo3OisksDzuDA0d/IkefDA8jlP7RipVL9FLkuZ1os4qxdIXkAP/ABUPPza+DnIc2ZQy5ZKL/a/WRkfLnHUR1/2i0lUacT6YCvaSG2isjG9eLqT9T4a/F+nKnoca2btW17DOaeurtKK5O5GbQeSDvEvwyt9TtfyIJBDBgfLCt0jC77UhsUrc1zU+S5j1NE5I62dOnv4sGGSZntJWia7Cwu6C3v8Aax1HS9pWuJIAWn0o7WIYronIgGN7g/8AmfAIK/DxeQw4n7T9efVXagbTDTUsm5kOisiHewPwpH+p+F1cFG6i5U8A4Q6UhBolGj5SK4cVyquTZq7BY8vcsvlvSw9r7umFBpnPpDdsrlYeRWSM2P8A4umOpaftD7Uqinqaqhl0tKYIROVTRuR7gt7XYGjuATcBuQG8LWNrlre0zX7QLVxVmmJorByG0Nkfebbc+EUfiItyoN+GtcgXBpOEmlUfQSjUuXKrTmeZfQHMqvKp4qVZBCZjtKCQgkKSCbEgE/LB9BZ/X5Jq6izDLzCxWohieKeFZoJo2kVWjkjfwyIQeVbjpiy+3TWmcap0zpWrztcpNRF+k6cSZdlNLQxyIlVHsLLTRxo52tYMRextir9EUq1Wo6RSBtarpRYe86DjEvbiWBzyMMRwNFGISyUM1rrGt7dNUdm4jih1M2mMtqZqgUa6byKly5qhYpNjgmkjjbct1B3TDqCMVfrD6Q+UZtV/pCbKs21BXrFJF9b1Fm7TkhyCf1YLm4IFjvuOfXBvpIUTUukdKiSpWoP1/UG2QC11+t0oHzFyMUAsSE/wxm2Gxx2yJs8pJrXXhmRuVu0Wh1neY2ADyVj1/wBIftCnDRZdmUWVxtcWoIAjDm/Ekm9x8iMQnM9R57n1Qa3Mq6oq5yb99UytPJf/ABOTb5WwjFlVRMoaKjmcdQViZr/MC2FxllaqlVy+qZuCR3D9PTpjWiskEP8ADaAqL7RJJ9xKYtDNM3e1EjuxHLE7vxPlhSGCx8IHxtzjYQZPmkzfqcrrnNui0shNvkOmE5aSSmZ4aiJ4ZEFzFIhR7eu02OLAUNSt0ulzNomq1S1RSxiGrVI43mUNKlirhR1LBrHb1K3PQYifcd6bkKBz0OF5DJKoiMkhiBuFJ4U2tcDp04xuabTOdPTLUx5NXujAMrLSSMpU9DcLYg4awOFbxqnOLTS6FHhTzQ2khlIYcGxIIxscl1ZqLTNUK7J8yqqGcG/fUs7wSH4tGQT874evkWdSnauS5iQPIUknH/scIz6czpuWyWv+JpZB/wC44HMa8UcKoa8tNQVP8n+k12gU8wOc1VLnIKhGbMaVXkZR5GaExyH5lsWjpb6WuRZcEnlyHMcnrRb/AFvIM4anfgWA2sI3AHkBIcc0rprNAG7zKK5F67jSycD/AMPTDVqeKNTba3HPnjNm0NZJTeDaHcaeiuR6RnYKVqN+Pqu7YPpSVHaDkuYw5fqGs1LNRUv1iel1Tk1JmKpHcKGY1UcpcbiBZZbnk3xzR9JSpeo7UczSKkpMuoaaplgoqKhgWCkpItsTlIoV8KJukZrDzYnCf0ellV9W92rEvkypZIy7c1Ua8KvLHnoObXw5+kKI37R8zeRlsa6pBKj9mOnW4/DFayNdBbTBeJaNprqqprQ4S2cSXQCdgprVa0GSZpXQGsgoJWgD90ZF27N9r7bkjmxvb3w4fTedoLrQTDi3Lx83/wCbFpfR81Ln+l59WV+mKymp5f0VCD9ay2kro9v1yJSTHUxSJcBjYgAi/UjjFgSdsPafXV5oMsr8gE6AyTF9FZG0cSWJBcij+0fJRckcmwGL77exkjmHV8A+6qtsjntDhrXNUmns+IMb5fIWNmG6WPr68tgqacze5aWhlDA3P62Mc+5LY6nTtQ7VRFFUrmGQyLGpkljOh8hDMANzKlqQjcq3PPDcdMOpu1XtkpoqCvEuTQ5dmpK0VTJojIkWoAW+9FNFuKnmzEBSQQCbE4jGlIDkU/6GULk58gzpbD6gxtcqRJGTz7hucGOn8/i3WoGHh8X6yIkj/wAX/XHVEPar2sF5Wqcz0u5uwgK6KyNS3JsCDRegDMegB4vjcya17aKigTMVrNMQ00sRlikqdE5EBLa20Rr9S7yTcTYFV2+d7C+Gv0rZ25lObo+YrkJclzyfkZdvVBZd5juo9PtC/wCeMOR5wrXGV3aQWYtJHY/DxcfHHVB7T+1GVYkSp0wkikiZU0PkLsfAbhb0fNmtz5qR54SHaN2pRgTVeo9LQop3Af2FyLcsQvdn/wBU8JJsqjzszGwAuDSlnOv1QbBMNS5ak07nZ+1k6gc8idP4sfjg8eQajbn6iy7jyY54rk/+Ln92OqKTta19N3iNm+n5iIx3Mn9hciF3LAAFTR+m42vyB5YNX9q3aUio1HnGk4UMwVzLofIjHHFtJMm76kCbgAhbdW29QcL/AGpBWh9Cg2CWlcOP4XKz5JqmOUd1l8sTK26wkQXNrH73n/nhOXTmpWJZ6KpkAa1mmST5Ebjjqig7V9dtVrG+a5EXkcKltBZEGcnoAgoybnyXk/PDLWP0jNT6U+v09VmemczzFCUgoho3IlSiBH2quVKMjeObQJdum4ryMK3SUTnBrRUnckNika2pIHmubqXQ2sqwH6lkFbJ4djFGT7J+6x3dPjjRAGJzGwVWUlGDAcEGxHyxL+0Htn132kU0WW55mFCtFA28QZfk9Fl0Jb17umiQfiTiGiNgqKAG4sOOfyxfZeNS4UVR1BkU5ENLIxRZKgut2OynW20dTywI+eFIYaWYtEHqS7DdEGRLMR13c+EW872HnxgkdLNAsoqZmjjIXeL8F+oQjzI67R04vbCc1REYe6gLqrt47m7PbzY+l+i9Pj1wmJOBRgBUpbvoqCYNR1BaXkNOl7L6iM9f+b8LDnDZ43ViSoUsAwAIIAPTp7YRZzfxcew6fEY2WX5VWVdXT0EVFNU1tU22noYVJlkJ9QOg/P4DDsGC8SmirjQJpHTo0ffzOIqdDZpCL3P7Kj7x/IeeJnlukkpaekqdTUc0aVZRsvyKEk1lezGyPNtG5UYkAADe17IBe+JFo/QWbVWfw5LprK4NS6xETT7YzGctyWFPtyySORF+r6tM5EMfqzYtnIIMp7NHkrtG502d63qAXrtaOGtTBx40yoOA8Y2k7q51EzLu7hIlG9su0Wy+2oddZt1u3N2DfwpgVoRWW66hFXbNQ8fjjsTfLOz6LQFXBnXaZllJmmrqdL0WlZoQ+XafVRuU5jGDtknUeNcvU2UDfUtwYi5zHPcwzuvq82z7O6uprq0/WamoqX7x3dVKkubeS2UAAIFGzaqhcMqPMIaOIwGQRQqx8UvPdHcSWN+u1/GT1J3E33PdhqXN8k0NT0uca0jnepnQyZVpiCcpU1m77E0zKN1PT8CzH9bIOIwFPeDElkdaXCNgoNQ69dS0mhtnBe446yt3VVWR6byebVmsMyNDlUbmONo1DVVdOoBaCkVuHkFxvkb9XECGYliEaiu0LX2aa0ngr8+pocsyfLt6ZPp+nJMNIjG7M1/FJK/DPI/jc2LFVCrjXax1vm+os2/tJqyqiqMw7sQ0VJAgSmy+nUnZDBGPDGi+VvO58TEtiC1lbUV8pmqHufIeSj2xq2HR9w3jnt9h88FQtNqv4cvn4SuZZnUZnL3kzWUHwrfgf5/10w1Vb34wKrfocLRpzfpY9fLG21oYLrclmucXGpQxofQ+vHUDCyKAtrGxYWH7+mMRPFe/vusbg+uFVHF0IIN7E+XuP88OSIF+6Ra/oG/q2FNqhfsHaCBtPPXzPzwRAPJRtItcNYfC/kRhSFrnbuCc+vI9QAOuBCEoVTaQFB8VmBIsOvHW2FFBAsgUMCDdRe3w/wA8EVlYKw3Kym5NyCpHH7v34NwF8CsLcWsOB73/AHWwiFhZUIuNrX3c9D72PH44NEq7Cb7uN7EWJt+1Y/n8jgNwt4Ra4KsNpHHWxFzx8MKOQ1rhV22UXYtYW4txcD8cCElIyrwrvYi/i8z7gYRDeHk9ABYC5+WFZV43A33c8G4NsIsAGJ6+hZuLeWFQsFtr/ZN7W23F/fnoRjOVCtuNjfqt/wB3vfAkSLYksPLnkG/PGAsxC7Sp8vMi/rfywIRArs20Wa/IF/65t+IwbdzuLhgo4t0+XH54K3Atdr2tyPzwPjKgoSTfdcL4if6+XrgQguSQQDx1JN7/AB9/fAE2U3AFzYm/UeQ/rrg+wsxc7D7sRb4XwW+43sxZjzz5e3rgQgALWLX8QtweSP6498BwBYdBwDfpg/R7rc34uOCfLBCCbFrcdPEDwfQYEIoZFk3dRboPX1xmCNck29bkf54zCIRr2O0OGsOnofP+vXAqSPGCV2+n9c4Idu7ixF+eOR74EX5IANj188KhGA8Q6KOnS+MO03LAkXFrcf0bYxja/pbrYjj54wSEccH19/lhEI6/ZAAAJPJH4/j74BgtriM89LcEYzfwPEQRz16fxGCktfZY7h6dcKhA1r7txsQRYDi/vgbceIgX/f58YLuutiefcdMCpJG5bX67R5YEJQAm9lva3HmRgLm7KHYXI68fC+A56G/hvweb36j2xl1BNiF4sAOTb0t54RCyxI9Bfza/54wAtdgLkn0HPr0/djAAPEAvpext+eB2oOSpI6BR1Pz9MCFhFrA7rHk3Xn9+AICr0Bt6cAj0/wA8B4owAR05O4YFebCyrx0sbYEIhUcixuOOeowV7EkkjyvY2v7+mFG322kcdOeowXYDwCSL8fzthUJEi9xz5jnBdt+fL0wswa3Cj0wVhe/n7364EJEgeeAK+gOFD1IB/LBbc+fPXAhEIwGDW44F8BYjr54EIySMhuOnpiV6P15m2k66OuoKh9q3Dx8G6twwIYFWDDgggq3RgeCIicCpINwbWxFLCydpY8VBUkcronXmldOaVzjJNcGCoyJo46tpQJspJsZb9fqpY3L+fcOS/B7tpR4QWGCoaCVKY2kpZpKORgL7kLEBXVuGU9Crcc24Ixzrlma1OXTiall7uToR91he9iDwR7HjzFji69I9p+T6uP1DVdR9RzmQBBmRVpBPboKhRdpBxbfzKv8A3o4HM2uwS2Q3mVcziR8jn46tyzWtlowdg7kfhXPpXXtNUZfQaT7UIMzqKTKY2pslzjLyhzrIomuGipmm8FdRG/jy6p4Av3TL0xFe1HscWCioM7pny+oyitnMeRagyV2TL6idTfu6aSSz0NSPv5dVFbG/ctawxrK2CaFfqOYSH9fEHFpu9jnj6LLE4urr0swNrixsbrjd6Q7RM80RUVEWXS0lVRZ1TLTZzlOZUv1rLc2ijIVoaynPEqi6skgKyxhro4tbBZ7ZWhefAjP8hOmgGIZwOX4KonUmSiSpeHVXd5fX94YUzqOBkgnlH+7rIrboZfVgt/Mqw8eIhmeTZjkdT9TzSmMMroHja4aOVD0dHHhdT5MpIx1/m/ZZprtEies7I4KySu7gmq0JX1AqsyiiXlv0VUyELm9Ko5FNJtqoxbbyN2KSrtG1FFl7Q5GsWc5Kah45MmqZHUxVA+2tNI4DwTjzhkCycWtKOcb8Vra6gcfPUVjSWYtqW8NYVSwyTQT74nKki3QWPswPBHxweQ09Qto41hkb7v3G/wAJP2fgePQ43GYaa3rPV6elmrIqe/1iknj2VtHbqJI/vAftrx6hemI8zCRDbkYtEVNQqwNMEjNTuGIIIZeoIsR8cZTVc1ISBZkY3ZG5U/5+/XDynmDKKepjMsY8Km9pE/wnzH908fDCNTS7GKowdQbbgLfiD0P9c4TPuuTh3cWlWbknaxDWZc1LqxZa9oEVYasVCxV0dvsqzsCtQg8i9pF4s1uMPx2p6PhcyT0ObMWsGKTUrFvU+xPr62PXFN/V2DWv1/LBu6iVkMpYoCN23gke18Z7tFwkkioGwK223yYK8Ju1zQVckjJRZ4hBJCyS0qgX6/G/mMNYe1PRy+JqTNxIpuHFXTBr8c329eFN/VQfW8W0FoTs81dNURZ32qx6SK7BTfX8nqatKhibFA9MG2t0+2qg+uJ1N9HXs1gmkp5fpB0yTwtskjbRWdBkb0IMPGKw0fZRheI8SVMbXPXKqSk7YNJVMccMkOaxmGUSRss1KQoIIZLG3hN728ug4AGFR2u6Qp3vT/pkRsbsn1qmG73NjyeBz149zhNfo+dmi3Mn0gqYLdQG/sXnBF2vYcRedjb4YXP0dOzQC5+kPSDi/wD9ZWddOOf9j7j8cJ/Zlk/r5pfrZ/6eSby9q2hVMhpqTOIw7l0VaqkHd3JO0ccgXa1/IkHg4LF2s6Ig7t46TOS6MysrVNIUaNhZo7Wva1v/AAr6YVX6PfZdIf1P0iqKS3W2i8749f8Ac4UX6OnZo4j2fSGoy0jbEB0bnHia17A911tzg/syyDN/NILZPqbyS0PbVomhWGoo4dQRVcDXEoqKNgwIsykW8QYcMDww634ONRWdrmh6mOWnp6DPIFZ9yWmpzsB5K3vewNiDe4IHphxUdgHZnT7BJ2/0qmRgibtGZyu5j0H+x6nAwfR97Oy/Pbvlx2tZg2ks7BHtb6v1wrdG2Nnevc0Otlodhd5KK9oOrcg1bkOT0GVGoSbKpKtpWqVjVqg1EqyblEfhG3Yb3te4I88R7QZMeo6MLYH67RWt5H6ymJn2m9nOi9FaYyfM9M6ypdTNm71DLU01LWUaQLA6xtG8VUisXLPcEcAKepPEF0M7LqfL19a+jHw/Xpi2I2R2VzIzUY+uPNVw9z5w54ofwrz+kiUrtFaYmCx74q/UO9kULvJrqPkgcX8ViRa9r4537sWFgQRfcLce39e2Oi+3iDfoHLZFCBY6rPbFHDA/+UqEXBBI59Mc9SsHDyKiru52oLAfAYZoc/3Ro2VS28fvyVc/0cc2oaWn1PHmel9L54kdPSywxZ9lH15InNT3bFLSIyXVjexN7C44xblXmWRlgrdjHZOjEBlvpZ1DA8ghhV2IPkw4xQ/YKzF8/wBrqo+r0ty17f8AnVxcgEjnz8sWX2jZxqLSnZfLnGQ5lU5ZV/2gpoVmp5ACYpKKd3QEXG1niVjbi4v5nGbap5xa3RRupUinAK7BHF2Ae9tcPdThazIaKlpc0j7KOzQrLF3siRZBOjot7XBFYCBfgnyPBtcYLWUvZ12iwLkWcaXy/J5ahhHTSpV1ByvvSdqpN9YeWbLyxIVauGVokYr30Jj3MOVD219qBqIp5Nb5u7Qm695P3gFxY+FgQQRwRaxHBx0N2d5pTar0vluoqiiipVr1lpauGND3PfRtsk2j7qOrKdvQFmHSwCzSW3R1JJnBwJp77AiNlmtlWRtoVQuvtHpoLVWY6fmNaDRyvF3dZB3NRE6kq8MyXIWWNwyOFJUlbqSrA4ursD1bDRdnddDWaR0dn7Q52kMb59kornggekLlY3MiMqB1B23IFzYC5xEvpaVK1motPZ0Tvq8z07QTVkv3pZ41kpXkPqzCljJPmbnD/wCjnC82jcwtEZT+mYgEVwrsTSH7APDEfs9SOnTFq3WtzrCJ2YE0Vey2cC09k/GlVdEWZ5XVwvPJ2TdlyNESrK2kFUhlAJX/AM4texB9wVI4YHD79N6YjRYZ+yXspXvRsjdtKKP1hHAJWqAseRu6Btt+GuKz7ctQ51o/Q+RZppHNa/KZ6jM6+nqHgk7p5I4oaZo4mtfwo0shUdV3kdOBz3L2z9qDSbjrjObqeP8AWen5YzrLDbrbH2sclAdp8tiuzy2WzvuOZiurO0MZbB2eaurcs0XovLJY8jknpqzKMi+qVdNMKqlQMkwmdluksike5GOLgpVHUkDaSAPYH88SVu13tHzKmnyqv1bmtXR1sfcVNLLUsYp49ytsdRa43Ih+KjGglRghVlN2G7nz56/jjb0fZ7RZ2ObaHXjXArLtcsUrgYhRW39HWFXi1azKjFMlSTawuDtr4TY/hhp9Jlvq/adnMSHcyZhU3A8rpAf54V7ApTFR6wkbcFGQMSygEr/rcZBIJFxccgG/mL2thl9JmoSo7VtRMp8P6Tqdv/ghxSgB/tN56+1WZD/dGjd7haDsw7Qcs0cM5jzamrJos3o1o70kkayxFZ45Q36wFSDst6i+JxSdtGgYpiBp7Oe6Ykugnp499+oJUX58/M+tuMQ7sf7OdH9oD5uurNdSaVjyumSrFU2TVOYxSK0qxbCtNeRG3OtmsVN7cYs6P6N/Y8QGi+kTQuW5UHRmei4/+FYsWmxWSWRz3uoTnjTUPaijhtFoY0NaKgblp37b9Dx0j036Oz1nYG016MsCb2uAACBfj2ABuOMGTtr7PJpllnoNSM6iO8jtSGSRkXarSOxJaw4C8Ko4UADCr/R87Jpgs0P0gaQI4JF9F57ew8/9if34AfR67KBtVvpC5erNcgHSGeXIBsTbuPXjEP8AZ1i/q5qQ2u1H+XklJe2Ts1eB4zTakJIIU7qIkX+fOGtR2zaHrJJGLaqAkREk3SUe5iEClixNyTtFubAcCw4wq30fOycRq6/SDy873Mak6Pz0XYC5A/UHoBfBT9HPs1a4j+kDRuwAOxdF56TYi4PMA4I5wg0ZYxm7mg2u0n+Xkn9P229nsVTHVin1BuDM8oZqFu8JAv52FyAbDz5GN9Q9vnZPHDKgyzURmlLEyl8tDAnqVLX5xGW+jf2XwrtqPpEUMUjIHUPozO+QenSHDSP6PvZm2YQZXB9ISikrqksIKddFZ4XlI67R3HNvP088QP0NYpB954p4t9pb/LyW0re2fs0LQqabVshhDqztLlpMu7zawA3ADbf098Intl7M5IpIVy3UyKyBVKy5eu0h94PQ3sb/ACY+uIjrTs/7KdNU8bUfa/S6nI3CRcnyKspyHHGzvKsIrHrfYrAWOKqmhiaZ/qyusW47A5BYL7kcYmZoizSGgJw8aJrrfM3EgYq2dU9s7xxSUWjlmy6CZDHJUmVGzCdT9pTLGAsCHoViszfeY9MVa889Ww7whY0+xGosq/Af0cFjpwNoBDX9OuH9NRExCpqJDBTkmzkXMhH7C/e+PAHmcadnssNkbRgVKWeS0HFJ09M00iRRRMSfIEfM36Ae56Y2B+rUADxkSS7fCynwj/ATyf8AGRb9m/XDeorImhSCki7mFeWUm7uQeGdvvfDhR6eeGsrliJJOSb72LEkn1/DjE9C7PJQ4NyWVM005HeMfCDtUdFF+bfz/ABwFPDLUt3UKd49rs17KF9WJ4Hxw4joD+rlqWkiWXiNFXdLMT5Knv6nj0viZZVoxUqaPL86oquetrZUjotM5arS11XKxsiylAWUsT9gAufILhsszIRQ56gMz1w2p0cTpcstupabT+ma3MxLV5eUSnpmC1Wazqfq9KT91Ba7yHyABY+QA5xb/AGe9l9RnuStmeX1smktFTyNTV+qq6DvcxzqUfbpaGEEGZ/WNGEScGeVRiZ0nZ/lGk2ppe1amy/N85y8mOh0Hl85XKcmNrn9KTwm8so6tRwOZD/vpU5QhqTUGeZ3mRzXO6/61MaYUsSokcVPTU68R08MSARQU46LEgCA9b33YxLTbTX97QnU0ZD/2Os7shzWnBZaDuYD+rWfDYN/6Jw+dZVl2TDQ2h8ijyLSveK8sDyrNVZnIo8NRX1AAFRKrAlY1Ap4rEIhI706mKizCasigpaapqquokenp0gUvLNMCCECi53XIJHNiVblHBU2Sq1bR1uZZgKfKctywL9fr62XZS0hborMAWZzbwxKrSsRwp5bFZaz7WZMxOZZL2etVZdlVf4cxzmoAjrsxULtKKASKaErx3aEkjiWR+FGcI57XIXE+eodbB+VbMkcDA1oUo1j2h5PoKseg039SzzV8I/2g2z5bkj9NwPKVVSPXmGIgD9c43Cjs3z+qasqswzGvnzPOK12kqqupkMkju3Uszck/mfPjjGurM0jij+qZaNiD7Ug6t8P5/hYY1gBOOgsthbEMfyfHYN3HFZM1oLjXodbUaSWSeRpZWLOxuScYse7pzfAqt+Dx74V2DobgfDGiNyqIFSw4S49cLKAvWw9+v9DGKCD1XkcAn+vLCyJwPQkmw5Pxt54VCAqLHaAPIgnj5ev8MCBYluSOAxHl5C/9e+DiMKRtIsPMXI/HAFVLDbzxYjjjn8RgQgAZtyXN2AJA8x8PPGRM17WN3Bt7euB2jYVsLXuBbi2MkF2D+Ihj0YXJt+/CIQgmw+0dw3C7EfO+BiYMLeTG739B6fDrgDGu4HmwNvs3/I4MjMp2qWIb3ub4EJYo0ZDFDc/ZGzYT7kG9/lfGBSCEA4J44AJ9yenTz4wQsVJjIC83NnA597cfDAXWw62IJsBbr5WPH8PTAhYxHLqxPncWv8fc4I1rXUWAHHi9Txfy/q2DLdnPgBsL9bW464B2vYgm3lfob+XoR7YEIoCrex5XoLfj8P44C4K3IJA8r8nn8PlgWZiNpNgG4Ba4+P8A1+WACjaVckMR04vbzI/zwIQHcCdtyoBBt+dh5f8AXBSD9ki9iRze3wwY2AtYWtxcXwnvAkG4/ZFr2sT6D3wISlhGAWW4tfn/AC6fvwXcNoCFtpNx74KL2BJ9ienPxPXGWuSwN7jn2+GFQjlV878+ZAH7sFuvALAc29wD5/5YDcDc3tuHlwB/X88A7GwDP1JFiemBCKD0c9bWFvMYzBN/hIuTboMZhEI3BBAN7HywZVYkbULDoOQMEJuTYEi/GMubEjqB6YEIxJYAbhzzbnGAgkbibed/5YKWNwRxx062ODKfL7IB4N+ntgQhJK8NwR1uOb/D1wBKk7SeB908cYAA/skXPQeRwF+Bf4fG3lhULFBHIY26deuDfZtdQOPMXsPn1wUkG/hHHPwxgU3FvO9wT+frhEIb3JHN+oA/jgbjhW55I6c/HBR16Ej0vbrgdzWt6cDjj8MKhKjbclrHoLWuPkMCzKLHeCQLXta/tbBNwUBr+V7emALGwvtNuL36j0wiEAUiwACg8hR1Bwf1Y7gevvf2wK7ibqDYni54b2/ywTqDwAOvTkfPAhC1vthLE9LHqfe+AUo1haw9P4YAncSxsN3J/wAsDuKtY8kHm/r74EICo46Ajjrz/kMFZR0IwqCACOdtvOxv8bYIQCDcGy/l/nhUJG3v1/LAFfa98KsFsetjza+ClT9oXtwD8cCEmVwQrYjgDC9vL+r4Kym/A+Y/dgQkGHJwGFGHzHXBCDe3ngQg4wrHNYgOTx0YdRhLi+B4tfCEVSgkKydG9qFfk9KMmz62YZS8hkCvcmNyLGRGHiRrdWXk/eDjjFk0U9LmVJHmGU1f16lQ94WBAmgB4/WqONhNv1i3W/XYfDjm+OZ4idh4PVT0ON1kOo81yOtir8jrpqeeJ+8CrIQQ3qpHn+duDcYxrXosOJfDgeR+Dv4rTs9upRsmI5/nwXQ6FJYpJ4Z3Lw1ZtsbhSpvG6EEWsbgFSGBAKnnE4j1zlGv+9HanFVJnPdrT/wBsMtp0qMxaIiyLmdK1os4p7C247KpbcPIRbFP6a7QNP6sX6rVtBkeasLNIg7qlnPnvRQREb87kGz9pE+1iUGmqqerFNWxt33cNHLFJb9dA/R1KmzjdZldSVJvY4x2Sy2VxY4UOw6/kb/CtFpOjjnaHDEbQl9f9k9TRS5bm2ayU31fMHEeQaqyWt72gr5B/u6arezJKvQ0VZsmXoHHANTat0jLl87/2sphQSmQxJndJTMKaWX9irgsGgl9bAHz2sPEbt0/rrPNHmafKp6KojzeBY89yjMadarLM5hUgf65SMbSe0i7ZYyNyuMSSn0/prtFl/RvZgWpM5nhWCTQmdZikskqcnu8nzGcbKxPSgrAJRa0bMeRs2e2l2DOBz8j8rNmsgH3cR7j4XHebZPmOSOi10K91OCYKiFhJBOvqjjhvh1HmBhijeMMCenl6YurPOzrMckqcxp9NwpEIZjFm2m82geFIZhxslglPe0knpuNv2JSOMV9XaPiq6yWjySGooM1iP63IsxO2oBPP6iQ2Ey+YBs9um7rjUjlbKMOCznxujOKjPdxm7KQp9DwP/wA393uMJTRE70dCpBBseowrNFNA8lPUxPFLC2x45FKuh9CDyPhgVmBRI5FLqDaxPT/CbXH5j2w/EJiYRSz0cm6JrX4IIuGHoR5jFm6G7UDTRw5TnlO1fRxqIooy4FTTr5CCR7h1H/Bk4/ZKnnFfTU6uSYSZB6W8Q+Xn8Rhm0RHKnjEFos0dqbdcpoZ3QOqF0tFPTZjQfWstrIKyhqt6pMrLCVII3xyo5DRSA7TY3syqylhyddFHmDxDvI4SyMUIE8XX1tv+ybnp5Fh1tigDmFcxDSSByoC7nRWaw6C5F8BCtbmM4ggiWSVr2UIoJ/LGWNFPYDV4p4fkeyvG3h1KNPFdFNl0zSU1Q9NTqHYq5aqjul/vHa4JAbkkclXc9RbDmij728M0ccQdbOtRPAwbn7JYNtfnm9lPqL+I0ZDoHX1RRpWQaJzuSB7lZUyeV0a3owjIPyOEP7G63jbnSeag++Uyf+6YQ6Le8YSDh+UC3hp+zn+Ff0uU0qqyfWKQQ2Bs9XEwBJ6r473UgFl6kEOh3BlwqtR3Own6tUsTyFrYNw4t1L2J4t6Mu0gqyg4oRdIa6lSy6TzY+XGUy/8A5PBBoXtAvYaQzof/ALIm/wDyeEGiZCMXjh+Up0g2uDef4Vh9uJpqbT+mKWGrhlu+aytsZbgvVow3KCdpPXk89cVzoSF5dSUAjUlvr9Ha3W/frgmbaZ1nkmVNW5xk9dQ0FROkN6mnMAklCswCq4BYgXuQLC4ueRh/2YsBq/KieCMyoCPj9ZTFoRGy2NzCakVx8TX3UDXie0hwFK/CvDty7hOzuiWmEXdNPnUg7sgqS2Z0JJ497/PHOErERsq3seote+Oke3kI2hInSKOMiozfeI+FLfpOjuQPK/Bt0vjmpiRGeOv44j0Ga2Su8+yXSYpP5e5Vq9gYIh1IyXuKakNlNiR9aHAPri09Q6P1Tr/s6rdKacpxmeYDUVFXCONlErUhoqqPvzH9ogO6q20GzdeuKx+j5IkcepXkjDJ9Tpd3F7D64nNuP69MWJqOLIKTI2zjPqyBMvFUtCXeN5bzSRtIoVkRiCVViGIBBUg36nNtT+ztznAVxHoFdhbfswbWmHuq2m+if20QTBqjS31elBHeVdZURUdPGPVpJ2QAfj8Di4skyvL9F6Ry7R+WVpzRaJH72op42C1FbI+52iVgGKX2RqGAZ1G6wJAxDI9edltNTLJDnbrVRWAdcnYkADki6gXJtcE7TzZV4xH9SdueXU0Dw6SpJ0rJFMZrpWCugIIvDGC3dtYkb3dmUE7ApswLRJa9IuEXZkAGuRHr7IhZBYwX3qkjrJaXt8z+DPNXQ5ZS1KVEeQ0EGVmVDdXliDmZgfMd9LKAfMLfzxOOwBYf7A5rDNGJFfN4Rtte/wDqbE8efAPHXjjFASSvIGne257mwFgPS3sOmL77CopZOzzNJ4ULGPOYGazFStqM2YMPskEjnyPXF7ScIhsAirlT1CrWKXtLUX+Pop9qfINKa6yOk05qnWOZ5PFlOY1Vakn6GlzQTx1ENOmwukyMpRoDy17qy88HEbrPo0dlcYhmPaTX93UEiN00ZWFWNr2v9btexva/TC2tdX1HZ9ktDqnM9PPWrmtdVUIWOtWBkaCOJi1hGwIYTDjgBlIta1oxH9KHJhRSUM+gpZKeVAjRfpMLa3KlSIQVKnlSOnI6EjFCyzW+OECzR1bjjUbcdY1q1NHZHSEzOofP4UW7Y+zHTHZzLkP9m89qM3TOaSerMsmXyUJjEc7QGLupJJCW3RltwaxDAWvc4rmTlG2Hds8Rt931uMTftO7UaLtKpcihpsmmoGyOlnpi0tT37T99UvMXLBEC23bbAc9cQRwoTqAR1HQjHR2R0z4QZxR2sefwsacRtkIiNWq1ewVo2odYJL3hVsiAKxx72saxLkLcXte9r3sDbDH6RNMYe1nUg3K4GaVK+o4EYuMPOwON5KTWKxMEk/QV1JYLa1VGTYnobXGEfpHMx7UdQsqkM2Z1hPH/AKLGXCaaTeOvtCvSD+6NO73ROwVPrE2p6INEGqMrjCq8iIW21sDHbvYAkAE2BvYYtaSmAuk01MiC5KrWQOVAB5IEq7mI8r7RcL6seZ8myPUueQSrkeUV1fFAwaUU9K84iJ4Bbap238ietjh43Z/rQ+JtP1gJ8jQyD8tmH2jRjp5nSB4AOqldQG3ckhtzYowy7Wm/fXYuhsuoTWKlQyrRtJyPrNbTkr/zJJYf+EC5Ptg9clQsizw0mUVLsViMctVDsjRV4siyDco+yqE7erMGJ559h0DrnYe703XH3/R8n/5PDdtCa2L7RpnMHb0XLpCfyTEI0M+9XtBw/Kk/tFt2l3n+FfUFFUoJYVNO0buX7s1cVgbWvff5iwv7A+oxtaanirhDJUiKmk2OkkgroCZGDgAgFwUFhe1z4jxYcY52PZr2j9zJUDQ+d91Eu95Dk86oq+pYxgAe5ONDURZhQTPS1MKxyR23qUQkX9eMK7RDyaCQV8Pykbb2jNvP8LpurrsuyOnizLP66ClpqdtrSRyJO0kwPhWKJHLSORyV4UXIZgOcVXr3tcqczjqcqyKA5bSVCmKVI3U1FRH+zPKoFk/7iLbGPvbzzis4pqkbhGwTeLMVUKSPS4GFUhCjgAnyxYs+imRuvyG96cOvJQzW5zxRgoiRtNI/eTMWIFufIfwGHKU7ysEjUuxF7D0/ljI4TA6vKWVgb7V+2fx4Hz/DCks5ZTHEirH0IXpcevm3xPyGNWupqo4VqUrakptgLCZ7jcQu6NPWwP2z+C/HCU081TK0ksjO5+8/J2jp7Ae3QYRjO505ZtzWsOSfgPXGzpMpnqamKgWGoqayZtsVDSrvmc+htfb+Z+GCgaLzijF2ATGCCaqJjp03EDcxJsEHqx8hjbZLkVVmDNJQCFkha02Y1Phpac+i3+03oLE+gxMNN9ntZWZvRZDUZZJnOaVU3dUmncpV5zLN5Rv3V2lf1RCbfedRfFzQdnel9NTw5h2rSUWpM8pFMOX6LyuYJlOXuL2SvqYCFkYWYmlpGLEqRLOlipz5reC2sRw/qOXlt9PFXI7Ia0kz2a/PYq47POy/Ps/jn1Lp+tXJsgo5TBmWts5V0hEliTBSqoMk05ANoIA8p43FFuRauR1mnez3Lfq3ZJDVUD18J+vaoq3QZ7mdO4IcK6kjL4W6GOAlmF1lnJ4DLUOrc+1FW0VXnU0M0lFEtNQ08EaU1LQU9xaGkp4gscEfI8CAEnaxZjuu0pjJm9E1HRU1KVonNVU1EzJDFRR7iWkkmawhU3axJ5DWAbhcYc1tcaiOormdZ+PALTiszW0MmrVqCCKKnooBTokMaUxK7FbaAt7jg22/zPqG3I6i1JpXQ9J3mvknetlBNJkNFJsr6xG5UztYijibjxFTK622Rr9sQfU/a9S5bX1MHZnIamvLsH1DPGVSDk8UcTC4YDgTON/7Cx2BxUNXmqQTTTrM9ZXTszzVMzl2Z2N2JY3JJPXnnzJxJZdHOlNZBTdr89g5+CZPbA0UYfP42+ilOuu0DO9aVMNVqWSno6CkuMuyTL0MVHRqeoSO5JY/ekZmkbq7k8YhOY5tNX/q1URQC1o19ul/5dBhnLNLO5lmcux8zgoBPGOihszIgMMuA8FkPlLlgFzbBwnHBHrjFTm/rhVQOvB+eLCiWRob3N/b3wqqgDiw4PBOMVQOSPy88KhTexBPNiPf29MCFgQg3HUjiw6/EfxGFEAHJvYW6n8RgoUkW23uLix6/wA8CoAa+0GxAHNgP428uMCEob2YjxCx5BuB74ANwQDdh4bEeXne1vL88FupIFtxA8+L+1/I+nlgRc2G4kdLtxx1Hw4wiEbgm/X145PzwQqbfaubWsDfjy/ywIJNrgDz62sfP/pjASRYX2qebDoT5k+uBCxdx2kbRYfePBHpbp74EAsCwC/tC7Hn2uMBe1gCSb/s8Bv69MCHJHIIsebfx9cKhHNjyL26el/UA9L4TYqrbhe9jYA8geVz6/DBS7XAQkX4PUfAE9D7YBz93c3Bv7D2GEQjrIVjHUgm+3yv5keYPtgjOPEeBfng+Xt74JuaNWC38RW5+A8/x/LAi4BG61xwF6HCoWA8EDbY89QPL0/jgwNuQ3Qg8C9v5HBONtyAbCwB5sP5YC55UdOgB9/b88CEZpLsFZja9ut73+PXGDp4nO1vLd5+nPTBGPN2A9PhgGueoI4HX0wIRjcm5sOLGx8vTGXZh9kdOvPT9wOALNcgtwehN/6GMFgQeh45PX2wIQMWIBsfmOCcDwL8kC/Xi34YKeeBe58r/wBfhgpPI8PHPngQit04sLm9hjMD1AAHXnnnGYRCAgCwHNuhv5YwEXI8vcYy5uTa1hYYEGym/N+L4VCwgXC+lrEYE2A8rHAbibKQPmBg77WuqWAPi9vlgQsY7WIIC26+Rt/PBVPNrEEjoPP3wbcLLx74LtHyPtxfAhCfKxt5WuSPz6YxrC2zqPU/hxjFAY2Y344AHAt8P34y/A54ubc269cIhCbi1uRewuLc+YtjGUGzAW46XPAwI238Rt6i1/kLfjjAvhJ3A8dF/f8A1+7CoRfu9bi1uvXGAk2IIuf2ugwLAWutiP3/ABwAO3kjqbH1GEQjACzci4FuvXB7gkM3PNyB/XX2wUsT4SSfMD0+GBF78MBYX8P9dcCEFmFxZetrgcWwmxt4QFJHrzg7Ak2I6Di+AKkjkXA69b/5jAhYoIAvcHke3XyOMueSOAelgOuB8K3NzuPn6n1N8Z0AY2I+Pn8vLCoWc7TxYeQ6gDBT4bG4BtwMYTdgBb0uBa2D7vDxwCOT159fhgQk2BuCB9r8MEIAJHHX4jChFixNrnr54KwB8QNgPXm388CERhc7SLW9MJH7ViLYVKkXJAF/ywU2J5vz6DAhJkcnBT8MKlTax8umCkEjpfAhE4JxnwxlvMYy2BCdQVZVgZCwYWKyKfECOh/rnE/0j2p5hlFOmTaihGa5OX3IGYh6dz1eFwCYnPnYbW++j4rb5XweKV4muhHPUHkH4jFW0WSO0tuvFfbwOpWIbQ+F1WnrftXUVFV5JmsUOd6eipc+pKeImohljtUwJe5aREO5QDyJoiyA33BPs4LEmV1JlR4l+pVy/V5e8AcRkm8L3PDbH2jdx4W8rY55yDPa7Jq6HMMmzCahrIZBJGUcrZx0ZWBBB97g++LVyLtTyXOpDTazgXK61yQ+YwQEwSbuD9ZgQAi/nLEAfNo364520WCaz4tq4f5h5a/EY5YLahtccudAeXHV58VdH/aTT6io6HJO2XK67UkWWqKTL9QUNUINTZMlvCkNXJ4a6nt0pqvd5qsikY1GvOymGryCTP2qcu1zoukssmosop3p6nJiebZjRWNRlb36sFkpWPIBvuxHPqzk0kFSYpocyWWmhqYpRNBUlR3kbrIvhcHxA2swubhTxja6bz3MtO51l2e5BqHMspzLL4mgpszoJzFW0o6d0z2tJH1UxyBkYWHnh0NuIp2mO8Z9biklsgOLMNxy68FWuqtHZhR0kU+aRSarydot9NmdIV/SdNCOhJUslTEPUF1/vR9MQCq0xIlK+bZHmEWb5YgvJPAhEkA/7+E+KP8AxeJPRsdY0meaC1fWSvqelptD59JJvk1FkOWO2S5g5/3mZZPGQ1Ox86qhIt9ox9cRTtK7GJ8lnotSZlCuSS5k3/krVeQ1yVeU5o3/AHNYm2KZj5xS9zUDoRIeuzDbKipxG0e41LLkstDhgdh9iuXjZSrABtpv7HAPdyXY3ueSBz+GLE1XpGbLmk/tjlYpGD7RnuVQk07N6VNNYGNvcBG/ut1xEs00xmmVUq5hIsVVl8h2x5hSP3tO59Nw+w391wre2L7XB4vNNVTc0tNCtRFTSz7gijwi5JNgObdf54RqEeMhdmx069Qb4XZfESVAv5AWGAY71Ae5CAqAT0HoP34VIpr2ddt/aB2fn6hk+rs8pctka81BBm9VTQSH9ody6mKQeTr08wRxi6qLtd17qSiOY6f7WtcVBgTfU0cupK1aulX9p1WXbJH/AN7GNv7Sx9ccryxDy4byGHGVZ1mOT1UNZQ1MsMtO4kikjkKPEw+8jDlT8PnjPtVh7TvxG6eR66wwVyC0hndkFQuq6btN7TJxEg7T9bsNpVQNS11yGNxY97yb8LfqDa5IUloNcdp8tQo/7TNaOQyi41JXncD5gd7c8Am3XhvNSMU4O2jMWKvNluVySFf1kn1VozK3mzIjhLnz2qAetsFqO2jNaibvZsrylyEKXNLICVJuQbSc8i9zzfnrzjNFmtg28VcM0G5S7tnzrUupdG6bm1TqPN84lpJs1jgbMq+WreFfrcQKq0jMQOnQ+QxW/Z4RHqzLXP3cxoL/AP4ymNnqvtOn1tk8VHmsUEdTRse4angKCQO4aQyEsbvdVsep5vzjVdnmw6qy1f2szoB//Mpi3FHIyyObKManmaqu57XWhpjyp7K4+2useTRSRiSGVGqc5IeF9ysP0nSWI9OnQ2Ixz842wlmA54GL37bKgU+lqOBgWNfXZyqHyG2upWN//Diipk2xG49ufXDtDNu2XDKpTNIG9L5fKtb6P7mKm1O6zGI/U6Paw5sfrg/liW9qtR9Y7Ia/vSu5tTUDC1ubUFYb+/FufbFc9kOs8h0jHm4zqqngNfTwJA8NN9YAaOo3sHUMpAK3scb3X+vtGag0XUZBkFbUtVz5pHmUiy0xgiCx0k8bbSzHxM0q+EdeT1NsUpIpDby66aVGNMMhrVpr2iygVFaKlzTv164dU1OL8DxeQPH4YWRFBCjzFva/vg6IzN3ZB4PKnrf4euOjWOkXC7CliL3txi9uwuolptAV8iK5AzuDlRcj/Um5A8+L8efTzxRk5AiI44J4HXFpdlnaFprS2lKrJ83zOWmqJ8whrF20Jqo5IhTtEyttddrbiDYj34OMzSrHSWe60VNR6q7YXBktSaKadvtXT5j2V6adY1BiznOVIVty/wCyoLFSebWPnz5c9cc4tBGWstjf0xcvaj2gaR1ZonL8p09XPJU0OYV1ZULJTNANs6UqJs3uxc3ha9ugscVGE5B3cXtc+RwmiYzFZ7rhTE5+KS3OD5ag1RYUEcRI4v1sP4jpjGXbETY7bce+BIYNcDxD5WwUmwJvYFR1PFsaipqyOxOURZfrBXp4p0kyVUZJASpBq0vwOfK497YH6QVVHUdqGoBHNdlzKp3c8gERW587i+FexmmhmybWAkERX9BqGEhsh/1tBZvbnGs7eYXg7VtTKp4GZVC2PsEGMWG6dIv2/gLRkqLK3rWnvYhnmpdLDVGbaW1Jm2TVMeVRlpctzCakd1+uQgqzRMpK8njF5Rdq3bTQUU5ftW15GJYwGibUtbdeQQQTJdW4/wCZSbcjHMOitcSaMWuMNHTVT18Qp5o6qN2RUWRJFKlGUht6D5DEhl7bKmQtI2mcnkdm3EsKnlutz+u64itUFpfM4sbgThwA9QVLBLCyNt4409yrrm7Vu1WqV2PbFrpOCQj6prgfj/tuQPM9ABc9RdHNe1ntA0rTrXah7YddqZYxLT0cWp60VNQh6P4pSIYz/wAVxz9xX6ilqftpradu8ptPZJBKu9opRSySmFmIIKpJIyMVt4d6kL6GwxCs1znMs/rJqyvqZpXnkMsryyF3lc9Xdjy7e5+VhghsNoe6jyQOfXWGaH2uJoq0VKnOu+3XtO1yJsuzLX+parLZTb6lPnVXPTKvkAssjFz/AHnuT5AYgEcBsBbg9Pf29zg9PTkjjyvckgAfEngYVRxAxMcodiLeElQR5+hP5Y2mMbELrFmuc6Q3nIDTGAWlXY4O4A8sPYr1HztjI3WPeQWVwBtIAv7+fHHmOcEaVnY+IWsbWFgD8BhWnop6lO9TZHCLh55TtjHsT5n2FziSlc0yuxJOQEuCeDc9LfEYVhoWMQqKqVaWmPSSQcv/AIV6t+73xvcnyCbMommyqkSWGI2lzSu/V0kJ/uqftn8T7YsXs67Mc71LHU53pOippKHL5AmZa21DIKXKsvc9FjeQFd/7KIskzfdQYryWlrCWtxPp4nV67lKyBzgHOwHWW1QOh0nLT0i5nmcxyHL5F3LNUDdXVKescfG1T+0dq+7YuHSnZTWw6cp81zyZuzfSOZRGSGpqIGqc+z+LzNLTArJJGeneOYaYH779MSjLv7BdnVRJX6QpZNYaqjXvn1XqCj3JC4YAtl2XS7grAkBZ6rvJB1WKM9NdX11dmVfX5vm2fVuY5rUu4rK6uleplqCrmzuSS99oF18ara6gAkYw7RbA/wC83js/lHz5+VFqQWa7l3Rt1n46qtzHqOg05p6syDsxyM6UyeeAU9dM1R3ucZpDYMRWVYUHu2W7inp1jp/AwIlIviNSVjwMqPHGqqgUgEhQB0UDpt8KgW8gvvh7mNRL9X/TOYT0GR5VSSLfMauoMlITbdtg2gtUbgTaOIMdrbTsF7V9n/bCtIslD2V0s2WRjch1FXj/AF4pz4aVLkUq8kbgWmI6yLcrimO2tbqnH0Hn7Cp9VZJjs4o38lSnU+f5JpGEVGtqisp65m72nyKkkAzCc3uDMCCtIhvwZFMpudsfR8VJrLXed6op0ps4liyvJIpTNS5JQ7lgV/25LktNL6u5Z+Tyo4xE5c2p6V5Hpw1RUyMzSVEp3MzE8knqSf8AqTjVTTS1Ehlmcsx8z/XGNqy6PbGQ457fgavHNZk9rL8OXzt9E5rczkqf1cSiKECwVeLj3/lhkMCoB64OE88arWhgoFRLi41KKF9cKom42Ixix3uADb87YWSO3I4HX44ckQBRyennyf3YOE4IIAPB4Hlg6rcEk3sPPywfbfheb2HHQ4EInh8JLbiePgPT+WDhGKkgMQTa4BP4/wAjgQreIleF9eRbz/o4MNoAvZrHgtzx7H+eEQgsSCOGANzaxA+eMVupuCgA6gcA+RwL3BVyRfnafn/XGChytiQbjyC9L8/52wIRjfaQTu9gb2HqQOMZ0FmUciy3APHz6YGRXIBI6ngt963Nvh74LdVHQH14t+P+WFQhIuFBa9r9Bc24PXqf5YBkZh4uTewtz16Ee2MUixJLcc3X2PX2wJCPv4I546EDCIQre1t4UKB0PTyI6flhNwhIWO5F7+K1x/n/AFzg24keE24t6AWwO4AWFxxb0+XN8KhFC2HAJLHyt4h/X9WwU3cFib25JHT8uB+74YMVFyPI+R6++MJYnddS5uCW5I9sIhFY7iJCDcf3ju/HBTbpxYdDbjBn6AAcH4E29/ngCdpILWNrXN+npxhUIrk2Jv5WBLW49LYAi/PUet7/AC+ODXBUXJsCSRbgfxwVrWBuCfO1ubYELOAFG31IFvLBCAvUKCehtg9hzx15BHmf5fvxhBLXBFuvPIOBCBD1PmLXA5wAIDAE3F+b+v8ADGbeSbA+/S/+WMtYg8WHFvMe2BCwW5uR6263GAsLG/wvbBzcgg2t8Ov8cEJVbDm4Fyb3tgQgsxJAY8/n8sZgU23FmtbrfGYRCLzcj288CCFBsPz/AKvghN79b/DywYfYt69OMKhFYqWsQABzb0wYH9odeth0wQ7gT4T74MD4LHgflhEI3RQelzwRgbWJ6AXtY+Xz8vjjORwRz94eY9sB0uBxxc36j+eFQsI6sOQB5H8sChBBJvYdfL5Yzi4BYEdbXvxjCST14PHPX4/HAhCCb9FI+HQ/DzwP2vtAE9b35J+OAUkkE3Hl4fTyvg3AtwL9RbCIQehvyBY35vjLADjafPgcfD3wPFrC3UG9yb/LGA8W48+mFQgIBJWxv7nnArawueetwefhxjLKUva6i1/Ue/PlgWbabEkEHp6YRCyw3FT5+tgPmbc/11wFulr26/HGbgCLC1uQb/ywFwB1sCeD7jAhGAv6gedj198A3TggWPP+WB3ErckC9twF+MG4YqviYjgXUfZ9OOpwqEmQdx3beep3Xt78YG6+pFxwP6/fgD5Hbbnlb8Yx9wbaQtwbWtaxwiEKoAAWNhe17cWwDbi1pBdvO/XA3JF2vboSSP4YJdiLG9vTz/PAhEdbAjg9CCD+WMC+W3oOeLWwbbuUm9x6jn4DGG4O4AG/IsT+N8CEXbztA8+LdTgrqLcj8PPB7km54HS5wHlf3tbAhIstvDfpgCLX9MKEXNgL36D0wU3twuFQk7Wxg5wJUX6nAEeVsCEBw7gzCRFWKcGVF+zc2ZPgf4dMNemAw1zQ7ApzXFuIUz0hrrUOjZpJtO10c1LMQ1Vl9VGJaee3QyRHi48nQhh5EYtDTmtdL6rqA1NWtk+ZzPuNDmFX+rlYnkU9Y1gfaOexHAEjdMc+glSGQkEHgg9MPY65JPDVqQT/ALxALn/EvRvyOM606OjlJeMHbR7jX67CrsFsczunLZ8bPTaukZ4KyiappZmmhlpGE80EkRjeIEi5CnmIm97jwHgjzON5k3aNneg8wqafS2ZmClry1PqDLKyiSqyrNEHBWqo5bxVHW+6yyAE2cEYovT/adn2QUMeX1ixZ5k6RmGOCpdj3CHqIZh+sg4+7yh81OLA05qCg1IEGlMzepqlAP6Krii1y8W/VHhKgW48G1yOqHGK+KWyG+7AbRl57PPDeVpNkjtAuDh8bfLHcrS7js21g6x5RVQ6AzmId0uW5nXyy6dn3dEpswYNUZaG8oaoTU/Ng6jEB1R2X51pLUj5S+U5hozUVTF3i0rwRmHMoD0dYVLU9bCfKSmZhY37vA0kkDSfVxUCmrYI9qwzXR2iJsYmDWN1PG02O3aR0IxJsn1xnmmMpOkpqGizvSc7GWfTGfUxq8tEoFy8KgrLRyFbkS0rRtcG4NsWYbbddR+B2jI+I+FDJZbzatxGw/KojVOn6KnmAzvLYtPVEjFFraHdPlVQ46jaPHA1+qjp5ouInm+n81yZEmrabdSzH9TVwsJaeX/DIvhPw4I9Mdjy6Y7OO0SnQ6Szb9E5vVARDTuqsxj21JtxHQ5y6iGpHksGYIknksuKkz7smzrSOoK/IssjzHTWcRrvrcgzTL3TenrJSSbi0ZH+8hM0duQVHONiO1tIq/LaMlmvsxBo3PYc/yuf3HtceeA7hW/n5n3xPc70vlYlVc4pDpapmJEdQu6oyqob+663aL4eIDzC4jecaazjIO6fMaYGmm4gq4HWWnmH9yRbqfhe/qBi2HBwqFWILTQrSfVrYwwDoDhzu68E/Efn8cC20BSynjzB5+IPmPbCoSAj2LYD7Q/HGy0jmUOTZ9SZhOFf6pVU9UsbuUWXu5FfYXsdtwLXsbYbEwBQpMjNt5YMLAnyAtz+OGEy8lrn4+uI3sEjS0605jrjg5diaT1loHPMtqVzPKf0plSTz1EsywxDO8hefh3Mb/qKqmcbQ8bgwyWBDU8hBxWHbL2M0eR5YdaaPqqWpyOokMSVFIZGoXktfYhk/W0s1uTR1NnH+7kmSxxTmRamzPIa2nraGtqKeamN4ZoX2yReoB81PQqbgjg4vbsz7XacVDvHUZfleZV0f1WqgqIkbJs6hJ5gqIZLxxk/8OQd0TyjQtZsZDDNooXSL8fMfI9PAVWg5sduN5puv5H468FzuIH3lWBFjY+uHNNThW3kXt7Xtjq7Mex/s/wC06RqHS1A+ndRJ+q/s3U1AEveHlVy6pnIEwPVaOrYORxBPJwMc/wCtOzvUWh6yePMKaSSnp6hqVqlYZIjFMOsM0cgElPMPOKUK3puFidaGaO0MD4zUKhJG+J114oVF9ylgPCVA8+lvgMY8paR24Ukm4Jvx5Dn0/HB+ZmHfyGxuCQo/o/vwaIK8itJKiXKgsyXC28yB1A/MYkOCYMU3maRbqT9obWIYEH5+eGggLAm9vjjbVmU1MRll3wSU6MVFRCxMEjABtqNaxaxB28G3kMNFVftN5eVvPACDkg4JvDEEbcbXHrhyjNc7WKhvK9/h8cGCL3Y8PivYkdCP5/15Y3mQ6LzrUNVBSZbSvI88wp07uNpTJK3SJFQFpZD/AMNAT62HOEc9rBVxQ0F2S1aI1SyQwU+6RyqKqqSWPoAOrE+mLG0R2K5jqClmz3Oainy7J6KQR12Y1W76nSOf90zIC9RUG/FLAGk83KLcidac7J8q0FKqaqpP0lnl+7GTQTEsreaVk8JOy1uaanYydRLLHyuGmvu2Q00YjqKqizCppITR0dLTxImWZXD5wwxRWjPQXSLwk8ySObjGVNpC8ezgFT111RXo7L/NJgFs6uh0PoPIWgyilnyXLJ0KfpOqjSXN88cMGISnBMMcQIsIlvGnWSWRgBilO0TVsettaZrqZKVaU5jVzVZgjfvVg3kWjDG2/aqi7cc3tjRZ9qnN9S10tbmFZNNLMArySNd3UcBeOFUDoigKB0GGSjYgCnk/n/liSx2J0T+3lNXnrr2yTbRaGvb2UY7oSRpwzE9fywQwc2w+Qx7mAXrcLc83t7e/ywFhtUkD1AI5+PsMaapJokFvF6dMOYdo2gQKxvawLcn4DB1jABNjwLkj+uPnhekoqmtZvqkDsEHjkJCIg9WY8D8cIhN5J3kJMgsFNlUDaq/AD+vfB4aeozBiaaBVjj+27PaNPdmbp/XGN7p7S1RndW1FkOV1GfVUYvJ3IMdJAPWSQ249yVX44sLT/ZRLnecUunY4anV2oJgWp8jyOnd4IwOrMUtuUeb3SMebkYqy2uKF3Z5u2DE8NXiaKdkD5BeyG05deCrXKslarZpMvo/0kyMFeeUGOkib3JtvPt5+hxaOjuxvMdV5oaDJsoq9a5/TRCaaAAU2WZXF/wASqkcqkMQ9ZWjU+h6Gyf7EaS0usMevM+pcyq6SwXTel6gGCm5tsq8zjBijFxZo6JZGJNjMp5w5zbWOZZrlNJpqBctyzTlNK8kOnsopfq2VISLLKYgS88ymx72d5JCbgkXxmWm3E4PNNwPq7V5K/BZBm0V3n2HytZSaU7NtKtBU6tzCLtPzyN+7hoqRpKTSmXsOt3TZPmAUixEQhhNvtyC1x1fqfUOrammTUUjTRZeqwUVDDTxwUNBE1v1VNSxBYqdbMo/VqCeSWY84aUArXqIcshR6qapl2rHCveSM20eNVUbjewBAB8O1uNtsa7UusNKaEnkjzuuXNc3LsVyLKJkd0Yn/AN+apQyQcCzJDvkIABMflmPnmtB7OMYbB1zOCuCKKAX3HHaU6g+sZtmI+qg1Mdu+qmEiqkCKTHJLKzERxeFVJZyAetyDiN592o6cyCrrIcjSDWGaSzyStNJvTKqYmRiOu16iwsLfq4uBw/nXGrdeZ9qiEUuczwZdlCSNLDk+XgxUysTfcwuWlf8AvuWb+8MQ6pzaRk7ikXuYh0twf8v3++NCz6Mv/wATHdq8zr8BhvKqTW2n24eqkOqdV5pqDMRm+rc4kzWtC7IoyAsUCDokcYASNB5KFA/u+eIzW5lU1x/WNtTyQdP88NuTzycCPhjaigZEBTVy8As58rnooFzYYOFN7fngwTnpb3wcABhcW49ehxOokULbi3zwYAew4+WBRTe1x1/HCoUAWJv5/DAhAqgi4Ab4jrhUD7oW9/K+BVQoIawYHz4wcCM/aLEnnn7Pt7/PAhYgZiG3NcHgg/xwfawcMUtYkW6kH4n1wU+Dqtjax6g/hjFY9QSLAMb+nrhEJR7CwDE+dmvf53/zwRrbSCwFr2Fzz/L4+WDbySVJY2Fxe3X1/rrgN7AgKCQGHQ35+f7jgQsQmM+drc+Rb2J/dgAAOQwJ63JFvj8fY9MYSL7NoJJJsz8cdeD0OMNx4mYPbncLW+YNuR7YEICDxdCOv2T4vj/XGAK3XaFBuvFxwb+gwJNlFgbA+vH44VCrt3eEWFixsRz5W6/PywISA6LdgOb3v/H+vLAbjbqdo+1YWsP3fI4VsFBW5ANtyk9PQ8dR7jCTKtyeHAFr3456dPPAhCpLXXkc+nX8D16YBSCNoNiSOC3mL+2AU2bxWIHFyOD/ADwdfs3Aa17G3r8L4EIyghSQoJPFib29+PPAMNo8SkBuRx6e+MHiuLKVHNgwsP8AL92Cbjc3bysx8/bAhYSAoJA8R4Pn8v44L1F73Uc26D/rjAeQbc9T74BgARYEcX68W9b4VCzqu3nn3Nj8sAfGAeeOg6kDGO17ng8eXF8ACQp3A9bX9PUYRCMAByGHPmeb/H+eMIGzk9eAfK/pf1wAfgbmW1uP6H8cCCFbjxX87AEe4wqFhI3DqLdT1A+fXAAm522HFr9OP69cFJ5BsBY+9r4zxBgwPIN7jAhG6WNl+BW/GEpEtwQTbocKE2FrjgmwJwUgEjrc8++BCID4hbgeuMxhPy9uuMwIQEki5HH8cYBfxenUWxhvu+zb0vjF6jng884EITY82Nx7YEc2svHmt+uC26br7fMe2B8rX6evl88CEcWU9Txf5YLxe3U/uP8AHGA2t19/L8MFsAfECQfTpgQjkkn19rWxm1gLqwuOfS+DEEWG29wTc9D8PXAjwkkMACbHy4Pp7fDAhFG5jcE/j0wLbeTe46+v54wXClfIc2/dgG9CDa3n1PGEQhNyOD5+Z6fLGdPF0J56/u9RgBcgcAnj4n2vg/iNyV48+Ov8sCEBIHRj7A+ZwBPQ7iCosDbp/XrjLXupO4W62/hgbEcnm3Ug9fQH+eBCLYi4vwDyBx+GFLXuSbgcm/Ujpz8MFWx8rWN+D/P8sHvYbfM9benmOf4YVCIy8AbLAjg36W88YljYixN+B0+eAYAk+EE9OMHHCm5O1vO3BI9cCEF7EXud3NwbG39euAINgbcDg28x/O2BaxJI6nyt+Pw/jgoXkEA+wsef5YRCHaTbhetwQbn/ADwJFib2I6jgm4H5Wxi+FmJA49r3PyxjCwDDqwLci3HqPXnAhFJDOGNx8Cb4w2K3uDc8+/zwYi9jxe/UemAII5AF7cc3/fhUIGCsxJAJ6k8j/rgLWte1h19sCWsAVJNulwOnscYbAgDkci9vL4YEJN9xJ3k36HzN8E2+pv59euFNt2G2zfDz+WA49r9CP5YEIjJc39fPBNvt+OFWXp0J8vhgoA5sPz64EJLbfkA4C3X2wsUB8vewwSwPC29sCETpzjDgxHn5+fOCkWNsCEaKaWBt8MhU+3nhytTDLbeohe99yjwk/Dy+WGeBw0sDsU4PIwVpaf7WMzphS0musuGpKKBO5pqwyha+lj/ZiqbHeg/4UwdfQL1xZGU5jlepKVptKZuMyjhAd6ZYzHWQqOfHTEk2HmYy6+d1xzVBVT0xJhewb7SnkN8R0ONhR5lGk8dRDK9FUxMHjkjdl2sOhVh4kP44x7Totru9F3Tuy4avLxNVpQW4jB2Pjnx+fJdAU0rhJaaeqWaJ3aB90e5JQOlzYiTi3Xn1F8TjJO0zPMsyym0VqLLss1lpRLSUeRag7yaKkcfbFBVoRU0EguGUwvsII8HNsUVlvazXj9RrKmfM0nUK2YUzLFWkDoWaxiqLerjd/eGJ3ki0WpsteoyDM4c3emPeOsCslVHFY3MtKbyAqTfdH3i28xYYzD2tkN53d35g9b6HYro7O0C6Md2vrwVpx6Q0Nr2MwaI1CI62u/VnS2s62GmrZW/4dHm5X6nmA8hHVpFL0G65xWef9jec6Vz2t0/lEOZ6bzhAWrdPZpl7jvFHN5aGXczJ/wB7CZ47cgqMDI09bl7QCSnk74loHYh45ZF52N1DBgNvTi4uBYHE2yDtdzuPTVHp7UNLQav0rCytDkWpRJNDRE/fo6pSKrLnvcboX2A9UGLcVsu/d3Tuy8woJLNey73jn5Fc857pjKo3C5xl7aZqHbalZTs1VlM7ezC7w/C7W81GIznmmc4yARz19ODSzG8FVDIJaaYeZSRfCfhwfUDHYFVp7RWtKzZorUjUOZ1g7v8As3rethgqJ2/4dJnQX6pW+gjrUjkPTcTis9RdmGaaXzyp09SUGZ6Sz9xuqchzCgMf1lf2mopSUmQ24kpmkUjkKL41Y7XUVfltGI/Czn2aho3PYc1zq20eErx6gkX98JMpa/Qm/PxxYWo9L0FMO8zbLf0DKX2Ctot1Tlcj+YZf9pTtf7p5H7IxD8309muTRJV1FOk1FK36urppBLTv7Bx0PsbH2xba4OFQqxBBoVpXSxvgYZ5adi0bWuLMCLhh6EeeFD4gbqF9f54I8TX9MKRXAoBIxCsLR3arV5dTR5FnlOMzypV7uOCVgJqVT1FNMwO0efdOGjPoDYjoDL+0bKtYZSH1dUVWp8rip0of05BHGM5y2DolPVxzEpVU4PSnqiyf8CojNhjjkgjjzxt9P6pzfTlfHmOWV01NURqUWaI+IKeqEHh0PQowIIxmS2B0b+2spuu1jUfH59VeZag9vZzio5jw64K6e0bsBmy6jXVei6qkzDI6uUU8NXSs4onnI4gJl/WUVSf/AIFqbE/7qSUWxTFXSVVBUyUdbTywVELmOWGZCjow6hlPIPscXj2cdtEVNUd/ldfR6czSph+qVUUiLJk+ZwHrBPDKGREbzilDwX5UxGxxOM70n2b9qrrlk+VR6W1JGixJltRWdzTOT9gUFZOT3Aa/hpatngbpDUpwuJLPbb57OYXXb9fgo5bKWi/GatXMC6hzSPJJNPCo30UsvetC4BQMPsso+64JPiHNjY3HGGVNTzVLLFTqZJZG2pGoJdmPkFA5xYWZdhetaHXEGhY6KapzKo37KSKmcV6bTYpLSH9ZE46+KybfGHKeLFh6e0PpHs+k7iaiotV55tO+iiqA+W09vtGrqkZRVbfvQwslODxJNJ9nEk1oisoJ81HHDJOQFA+zzsZzLUcC6h1BVU2Taeim7mXMqwOacyDkxRLH+sq5v+5gvb/ePGOcWy+rtO9n2SkaREuR5fURPTPm9TtbN8zj+9DF3R208J84KUhf+NO/TEH7Q+1qiNV9ebNV1FmccX1eGaSJVy2iiHSGmgUKjop6IipALC4kPJpXPdS5xqOtkrsyrZ6iaXh5JX3Ow8hfoFHkqgKPIYoNbPbzU91vPy2dYK3+6suGZ5KZ6x7W67M6f9D5NCKKgCd00cbfrJl/ZlkW11/7pLJ67jzivZ5aisk76eQu1rDyAHkAOgHsMAsJAB+XwPvhwse309OOv5404LNFZm3Ywqc075jVxRIo9rAXt72P8MLFnC7dzBSvK8i4+f78YyKFILHmxttIN/XDikoqirBlhQCNftyyMEjX4sevwGLChTZQqcEIx6ck/uwvT0tVWFngjZwv25CQI1H95jwMbfJsm+vuYstoJM3kj8LSyHuKOL/E7EXHsSPhif6e7MqnO83o8lrKTMdU5xUi9FkGS07hWHqERe8ZR5vtRfV8VprXHEbpxdsGJ68aKaOB8gvZDaclXuSZDU5vVtR5NQS5zVAXcRXSniH7Ujm3h9yVGLFyHs0grsxpMnrGq9W5zUc0uQ5FDI0QP/q13ygeZRVX1kxZq6FyTTUbZfr/AFCjNQsGbReiZoZJYW9K2vAampiPMJ9Zm56KcPM51jmn1Go0tpSiynSWRTALUZbkpkX62f8A5NqpCamsb1719noig8ZVptjyKPddGxv3eZ1eXEq/BZm5tFd5y8hr80lR6RyfIaf9F6/zxYEpn8OkdIyQvKrj7lVWDdS0pvwQn1me/HhOF67V1ZJlo0zkVDQaZ0tVS3qMvydZBDUbR/tKuZyZ65xe5MzlRxtjXpjQVD1eTwuJokpopheHutpj3rwdrDwg2sCBzuVCQPNvmua5XpukSs1lnUeURSIZIaTuu8rpb/ejpwQwBAHilKLYLbcBY5nbPd+7ibQHUMz4nMq6WNb35DUjbkPAZBKzVdOaKkjVN7xEwq6ptMpJspRDyNymxjtY8cDcba7P82yHRVRDPrCtNM2ws2Sw2kr3foAy3tAp4IMpDgcFGIBMI1B2uZrWK9NojLv7M0e0o2YSy97mU6m9/wBdYCIEfdhCD1ZsVpNmNHSMWo1M07El5pDclj1N/wCX4nFuDR734PwGwZ8ch5V8lXltgGLeJ+MyrA1J2t6kzihnynJUj0pkcy7JoKRyaqrQm9qioNncH9gbU9ExXcmaQ06mLLoQo6F2HX+vf8MMZ6iapffNIWPkPIYSt6jG3BY44W3QMOs9Z81myTueak9eyNJJJMxklcux8ycABgQhPQXwZU87e1sWssFAgVfbB0HN7/zwYIeu38OowYKB59PbCoQheL2sevW+BZd1rEnji/p5YOoO087QD5njAgHpxyfK44+GBCKEUEc3B9OOfngwt53454F/3+WBK3sWJAta/XjAFSSFG0H0GBCUjNuV5Bva37iPTBwByQOLetvy98J8cEqALDzPT1wqouN27aSTck9SP3H+eEQikLwVsCR0b8ucGVbbXV/mF88FvtsLg29Ba/rhRWPUD7R6DofjfnCoWMG2kW4v0H2T7gYDaSNpuzDkkjn4etvc4y4Ac7AFY82N2A9ATgQVYcEHjkngEf11GEQgJDKLi3ntB4+Qt/PBXsrWDrbqDcG/pyMGYixW5INuF829ecJ+5YA38QYXF/T3wIR0Yqwa5W3WwNxYdP8Arg3iuw3XI6H09cEYhidoAA45F9vt6/PGEBedpIHy+Yv/ABwIQMfEApFyfLnjrx/XlgSSQSzg8X59enz+OCsLDa446m/pgVDKQ/KkC46Gx9eDwfiMCERjYgtaxUXbcL/PAoL38I5HmQAPmcFubgAdTxwOfS/vjAxU2DW3DysR6G3v74EI/KsV2jwAkjbb5/EeeCX9SQTx1vz/AF54FTcgke4HP44xtq+Y568n5cWtgQg3LyAxHNrnr74KV8SgXJtz63/hgxG1tpULbgi/9fhgo6WI8gCfO+BCFgxvZgTf14vgG2jdwPMC/FvngW2gkEEEjgf1+WA3AcC5v+ftbCoRTcbQOp8vM4zgADqL8kcD/PBrgMR0FrDp19xjD4SL2Hre+BCDkWZbW5HJP7sAFN7spuvy+ftjOg4B56H1+OALE29ATx5X/r54EITxcbh1/wAP/XBebXAB5va/F8Dyw639xbpjLG/AYcWt14wiEU8dCbnzxmClmW428k/PGYVCMQGPB5OAFrHny/rnGXt0YXv54Dg2BYD+GBCHy5Hzv5++C2JHw63PGD+Ig2BB4wUbRe1+vB8sCFliBY+RPnbAg9bkX6ceuMNreE/C+AFr2BFut8IhHB+9+HOBBXpuBH44J1Hlz52xgB234OFQjA2FgOfbAMb39b/hgAb2t1Hp1OB8uGNjzYYRCzkCz8X9OLYMCLhlIB8vUfzwG1ifDtNh/XXA8gWYk3N/I/hgQhJUAXuPQG3P8sCpYPY3+P8AXTAFbKW/hfjAWBN7/wCY/lgQjht3PB9iebegwDEgGxPHPItbAKCGsOTxyDf3tg1mBFhc9QB14P8AnhUILi+61weLE2vjIwWuOF5I+BwPNieOPK3J/HGIAthcEDz63wIRrXABPA6bhyMAwWxuFsx9LX98CDYW37tvF7ci/t6Ywi7BUAI6WBuSPa+EQgAuTt/IW/IYxjz3hViN3Jvzf2PkcAOBtLC1+p/K3njCBc8AW4te5t/PAhYD7keXwGMsD4SCD8+f88Dt4Fwbn4cj+eAuD4fbz88CFhuwsovf1A5/l74zwXBsCV/P3+WAYgEA3Funn164FiCeSqnqPL8sKhFKi3w4JuDx8cFt5FQoPBuP4YVZSGAutmF7Ngh/vX563FycCERr7b83v0OCx3N1/LB9txY+R+f44MoJJ5tfjp1thEJMLxcoQB1PXrgpAXm1vbCoALC4PPn6H0+GCsNtx1INyAOB88KhJWsLnn2wTabj3wqwAHB/r298FCDr8xgQkyD69cBbCrKCQOCfW9sF48rc4EInlc4DBiPTAWsMCEpDUzQcRt4T1U8qfiMbGizMRzxVMFRLRVULBopY3ZSjDoVceJDjU+WMtiN8bX5p7ZC1Wtk3avmEE0f9r6AZmrOrfpGnZYa0n1L2Mc/wkUk/tDE+oK7LdRpK2lM4jzKOQM/1VU7mtp93LA07ElluSfAZF58vLnCCqnpie6fwnqpF1b4g8Ydw1sTSJIjNSyoQyspO0H1B+0vyxlTaLbnF3fThq8iN60Irccn4+OfH5XQsE5mhaanYyxSRmOop57SFWjO2SNlIuVF1YEgkBhcYl2VdpuoMl02ul8wo6DU+lInDDTWpqdq6hp2vY/V33Coy9/2Xp5FQnqt7DFG5R2o1qiCLWWWDO0iP6qvSXua+IWA8NQARKLAeGUN0tcYmuR51l2ooHjyPO5MwkRvDRyKI64Rn+4Sd9uPsFgbDjm2M4smsjr2Q2jLrxAVy/FaG3Tidhz68FcUmTdn+tUjn07ns2ls1nHcjI9a1walnPlDR5+qbbc2WHMY/O2/zxWOr+xys0rn75NW5Xmuic9qV3fUKulVI6xD0ZISxgq4yPv0zuPMJgtNUNPFVUcwdIx+pniZTYenhPIB/ZPTkW4tiYaU7QtV6TyJdFiSgzzTMxDDT2fUa5nk9QD5pTub08h8paV4yfTdxixFbQ3B/dO0ZcFFJZi77e8Nhz4qhtTaFOUs75/lDZWobb+kstRqigLf95FxJAfbg/wB3ERzPTeZZbB9cCx1VAx8NZSv30B+Y5Q+zAHHYVMeyvUTiHJ82k7PM0eIMMr1LVTZhp6ZCdtqbN0X61RrfgR1aSRqbAtivO0TsSzXReYwVtdlVbouqzMXo6uOSKXLczv0+r1ULGjqlP7KvG39y+NWO1VFXYjaPcalQfZsaNwOwrmVkG0Hy9cJsthfFlak0TVUPe/2hyGSkaPl8xyuEmJfeemNmj/xAKPS+IhmGmcxo6T9IU/d11B/8F0jd5GP8X3kPswGLbXteKtNVVc0sNHBaWKWWB+8jdlI/qxxO9JdptZlkMWT5vCMxytAypTSylDCG+13EvJiv+wQ0Z81xBGW3zwQgjEU9njtDbsgUkUz4TVhXUX/axFPo2HLajtFrqzJqaF6OLLDYZkICQ/1NlIstOCA3MjU4PKpfwCnNadp1XnkZy2hghpKBbWpYGLRtboZXNmnYeVwEX7qjEHFRVNT/AFYTP3V77b8YKI8U4dGsY+/IS46q6lPLbC8UYKbVjNNUyGSWRndvNjzhVIbdeL4xVA6i3Hnh7HSyPGJyBDCf97K21fl5t8saWWCpE1TdWHJZiSB1txxh3TUc9VH3ylY4R1mlbai/M9flfG+yXRma5pElXR0WylY8ZhmAMcH/AKtOWkPwDfLE+0T2atqHMZMv0dp7MdaZvSJvqJjCv1WiUffkLEQUyCx8c78fsjFWW2Rxm63vHYPc6lMyzucLzsBv6xUGyXSFTmkSTZXlr1aO20Vtb+ppSfPYn25T8L/DEx0z2ZSakz6LTeVZbmutNQ9VyzLqUskC+bOikLGg82lZQByRix4ct0Bp9u91lqSo1nmFirZNpWrakyuMDqlVnDqZJlHG6Ojj29bOBzhTOtY5/nWTnR1BS0GQ6ak2/wDvO6YgNBQMRyGn5M1YfV6h32kXta4xnTWx1f3jqbhn5n4V6KzDNo8z7BOItEaL0aFj1xrCDMMzg4/s9o2eKcUxHVKrN2U08BHIKUkcrj9oHnBc317mb0DaW0tBTaWyCsmAq8ryFHiWvBBH+t1TsaisckBf1shUk8IvAxH62fLZ/qwgjSKNQIkiphckW4MajqvovmCV6hThPPJMk0oYpdaZ2MnlCMxy+NRUZk4PK7YVIERuFIaUpwPsnjGcZ3PN1gpXUMz7lXBG1nedjvPVAtmr0WVKiSRRFZGVUjgAjCLbgC3RgWUhepuQQBbGtziupNMU7VOuc8XKWkG6Ok2GStlU82jp77o04sDIUFgOvOILnHbBnBElPoig/s9B4g+Yzus+ZSAkknvrBYbkk2iC9TdjitqnNadJ5KhWkq6qVizzzMXZmPUljyT/AFfEtn0c9/34eGfHIc/JRS2xrft69zyVkZz2xZqI3ptDZSmnqdmLGvqH7/MJDa24OQEhNv8Ahqp/vHFa1Oaxiokq2llrKuVi8lRO5dmY9WLHkn36++NfU1lTVG80lx6DgDCHXG3BY44RgKdazmVmyWh0hz68EtUVdRVG80ha/NvLCQGBUDB1K8euLgAAoFASTiUXYT5YNs6A+eFNoKkgdOmDBAAOSL+owqRJhLEXNh64MB+1gWFrXFvfB0T1HHNgeOmBCAL5m/vbr88GA4NuT5+pwcbfCtiLnzPX0wKgAFgpIJ5vyCf4YEICLKTwLjhieuD7DfbYg9SADwfK1+uBKG17i3TgWGBH3QLcCw55v7eX9XwiEFmB3KApHU34/wAvgcBtUkKLAdRf+unphQ+EkWIK88W49T8fUYFvQgX4uq3G4+p/q2BCILeIABST6cH+WFEPC7jdSQSpJA9LH3t5j1wXaN1iGIHkTb5ceeMVrAW5INiAL9PT2/dgQsIAIJNxfm3H/TASMVF2G1Tfpex+Z88YXNmBFiB5Lzb09sFUgBlsASOBz08wb4EICCu5XupC3sT5eg/HoMCrbPUeXzwUOFG1T7g/1+/A+hDBSObk9PT4fDAhDd2YAkkgXB6BQeo+BwP+zF/s87T5ED+vLBhGyoga9nG4L1uL8G344I2zbYsGJNiCSCPxHHscCEXdezG9r28K8n35woD1uWHAufn1J+Nh6DCdgqk7gBfdz0H8j8OMDuC2N16cC/S/w/PAhG3EA2VR52Fh878/54C4twG8PiuB+PHp+eM2kjcAdv2dtvP4+fr64JyWv5AcgH8wcCEZUYtZSFJ5BBHXAlbIRdT5kL4QT/HAB1UXKg3vfqbj5YAurGzWJvc3P7z5/HrgQhBC8EWDC9vM/Lr/AF6YKz9eLi3F+B+HpjGHADAknzPVvf8AhfBWJve9r9fX5YEIORwRa3BHnf0tg17OpJJA49OPQXwBXzsp6C5b+B64L4d3BAYn8PjgQjWI4Y38rKP6/wA8ArkGzc8ce2BNgBbpz14/L+eCnoRc29On54ELAeQBfgeoxhNzcv4m4J/dfGAkg2Y2A5H7I/iMYTY2tb+vLCoQE3UdAPbi2CkX9eLdPK+BawNjxYf188AL8eIc+d+uBCDdYm5vzjGZ7efpgwDhgouCbgD1vgjC5G0W8/c++BCKbk2IIIxmMI6j93njMCEJPJC/HArY2vex9OuCgm9xz54FDcedvM+mBCENe1up5GMYnkX68HGFrWBA6/HANuJNjut68kYEIQBweoPAsfTGC/Nx055wVRc8demBHWw4HlgQhJv5f54wX6evr5YAEjj19r3wI8/Y2NsCEK9bcC/vx+ODlfvdb+Z/ngg63uPXrxgxPHHN+tvTCIQ2LWQ/IWwIK9TY/MDn4YAelhY/j/lgUbqAb+XHOBCNcD0Fhzbr88APIG/I8ja38MAwbcLKt1FwD5Yzkk8eV+t/zwUQhU3FwCLfv+WBIudtunAFvL+NsFBDC/4+XXpgUIFrW44N+n9fxwIWFuLNyDYsB69P5YFLK37JPHHQYC3lzzx6A4EMVY3Ww6WPHHp8cKhCd24ADp5bri/x/q2MO/aSBcDy5JOAHNxe3lYi39DAcqSDuW3mD19j54RCMVKcNcNcHiw4I4wFrkA3tz8B74EBbhTewHxAB8vh+7BQG6AAXv0a/Hv8cCEJIHA6EWAHB+GA8Vr3Um3AIvx7HywO0Hm38fy64zcCWa/LX6MT198KhAF8JKk28gf3e2BG6wA3Hm9vL8MAXO02HJX0vx8fLGECw8ulr9RhEI9j+0eTybXI/n8cE2mwP77/AIYxtx5C9eeD0+eBPndjf0PS+BCEWBvbgHi/n88FC+QPJGDG7AnkngX6gfPGKQfOx+PJ+XphUINvFwD0v58fh1wmOPvG63IPp8f44UaxKgMD1vfgn5eeClSDybefXnCIRAu424v5AfnjFWym5P77/L+ODgdVI2i3Nhyf5DGNZb34J9TwMKhEdOvBPA5B/hghQC5t/XrhYlWbm5vzfrgrL6EE/wBdcCEhsAI3KR+/AMvvfCqqDYixJ448jgSo9CPIi2BCQK/uvgpFsLbbiw+HIwVlJuLdOTgQkvPGfDBivrYYD54EJSGpmpzeKQgHqOoPxHnh1FWQSFS47hwQQy3K39fVT7jDDGDDHMDk9ry3BWblfarnsMEVJqSKPUFKlljkqJSlXGB5R1S+I/4XDD2xP8mzbT2pICmmM8L1jN3hynMFWCs3eewg91OT6xlHPmhxzvHNLESUYi/UdQflh2lfHKAk67ffkr/Mfn8MZc+i2Oxjw8MuHxQ71eitzhg/Hx+fldEzZjWtVRb6plqaViTFOhUm/DxyAgEBwdrXA8m6jG/0hr/UOg4J8q0bmTQZXmBZKzJayOOqyyvX9ipoZwYJHH2SyhX4uGF8UXkXaNn+XCKGvljzqij8KQ1zszIvpHODvj+F7f3TiwMg1PpvOiseXZlHRVcjc5fm2xQ548KTWEcl+ni2NwPPGY+Kaym8MN49+qb1ebJHOKHn1+VclFmXZZqGiiEgk7OK0oTsliqc10rfcQehOYZRc3B2NNCpvxbED152G5rkE8GfTUsenhmZ/wDJefZXmENTlGa36CCuitTTk3/2cncy+RVjjUzTZpT5pJTXq8urY2MyI6lHUnqbH7SsAD5hrHm+NvpLtD1Ho6rmk03XJlkWZ/qM2o0hSfLMxU8A1NFKGgnXdYMHTcAbq4xNDbMbzhQ7R7jIqOSz4UaajYfY5hU1q3SGYZRWS0+rNPyU0sf26/L4CnwaWnNretxt9r4idXpyqipzW0Tx19GOTUU12Cj++v2kPxHzx2TQ1XZ9rHLIaaUwaArtpIgMdRmOl2k+8oS7V2UHrzE00A67QOkT1t2AUNFNRZhVB9K1maC+VZvTVcM2WZqT50tfERS1QP7JMM3PKE41Y7Z3auxG0e4zCoPsuNG4HYfY5FcnjaByQRbDlIJCgkd1hj/bk8/gOpxYmf8AZhrak1O2T55l2W5dPSU/ez5k8RjWWItZZChtd7gjhRf88THQvZFU5xDPnWlMnirqOhO2u1Xn1THSZTRt53qJf1KsPJE76U+Sg4JNIRNN2PvHd7nIeu5MbZH5v7o3+w1qq8l0lmWYCOojpkpKeT7NZXqfH/6KIAs5+APuRiydC9m1Rnz1FRpbTzZ7NQC9dnmayRQ5fl49ZJpSKWmA9HZ39Evxie/obs6yBFqFWp7SM7mTcamsjnoMhjI5ulPcVleo63laCIjnYR11Odai1HrKKBM+zOSspMm716OgjRIMupF3IkXcUsSrDECzjlV3ccscZ01qdJ/EPkMvM5nl4K3HAG/YPM+wW+y7TvZnl1QjZ/mdZ2mZ1IjMlLSzz5bp8FRcqaggVleB12RLBGQDYkYbZ/rDVeraKnyeveCm09RF2h03ldGuX5TSMtutKllLAdZJjI/U7rjmO5ctXW5hFSxUk9bULIssccCFnNupK/C/Jt0FyBhnn+p9MabMkOpNQtm+YA2NDlkwnkuPKWckxR9B4R3hFhwCMUjLJIezj4DrmaK0I42C+7PaUvCslT+ropZZ6v7UcdiZJCBeMqtrtuXwEAeZB4IOFM2znK9NLu1ln4o3CB4clpEWozA8eESgHbERfnvG/wCTnFb552oair45osoWLTtBMCrpRuwmlX0kqG/WPfzClV/u4gkmYxRXFMm9jyXIsCfX1Pz/AAxbh0e+TF+G4Z8ch5V3FQSWtrft68lY2YdrGeJTNQ6NoYNL0ezu2qIn7yvdOfC1SwBQcnwxhABxzivZswpo3aSK88zMWaRySCx6kk8sffjGumqJp23SyFj5eg+A8sE6+WNeGyMiyHW85lUJLQ56VqKypqzeaQkDoo4UfAYRwYAkWwZAPO2LQAaKBQEk4lEAwYKPTCoQgXI5uevlgQoIBIH/AFwqRJgC9jcfwwcR355AGDIgA3Ei4NrdcK7BbdYA+nPPzwIRdthyfy64ML7b2I9+oI+HrgbCxFlYDzvfnBgBtFiBuPl1/DywISZXkEAjyFxf8sCAdm2/HU/54Ej7IuoPra9h7+mDMACQCwF7g2BPxv0wIQ2twbcDyN+ff1+GDbV+0Sbgc2BPy+HxwC3vba1yOBwLkfkfhg5ZCvFrMLcdT6fn5YRCIp2/ZNrHqDzg9vCCqm3p1v6fD4YFbi788Cx4I/D+fljFC7doFz6C5J/6YEIvO6ygEdSOg+Xrgbg8Hk9QCLH8xzgXC8FbNa4upuD7j+XpgrHwklh04YNcfngQsJHO0A38hbm3PPyxl2ZiCdwtyriwt6j0+Ixh3Jxvvbz8PHwt54BblhyBfytYfG/9WwIQck8c28iLi3rY4y5A6evJY/n7++B3EgWAuTu+0ODa2BUXC7d3FrWFzf8ArywIRFDMNo+8OCep87YOSFsycm/krWv/AM3X4YMwVnuyjk8lTe/z/hggvc3uLjxEcAj44EIH+0CLdetrYy5Fgr/ZvYgm2DKFVrlBx1BNhb3wnvYbtoAuCCTxwPbAhC202HXkEC/PzPT4YIxLWub8cDr+R5wYAkkm5JHkb/Hp+7BHuGsR8Ph7e2BCFCtrEE2PW17/AMcHZrk+EEk3IJ/lhO4sRu8+bHqfYjrgWsLck+d26keeBCAsST7G4waNtvmQAfLz/wA8AQCSenFgDySMALgjcPL7N7jAhKHw+HhT1Fj1H88Ec9FHlzax4+IxnSxIFj5Dg4wceIbiR0NrD+vbAhAzX/Zbzv16/uwCm5AL+XUnn4DAsQAOb2Ppz8L4LyCfe4Pv+OFQhuLkbQP66+2APoovf88YzDgfnbr8sFDen7sIhCLm3mfLGbiOB62+GMBHNrC3T/PAdOBbjn3wqEF7dRfjn0+GBBAFvPrgGY/ZIHPFvIe+MPPQe3S3zwiFhAB6CwF/+oxh+z1HPX1xh9eBf1PN8ZxYXB6c8/1xgQiseTwLmw9uMZgDtI9vjjMKhFIFze+Dr5X+Nxgv2iQDyemDC9rqbYELL+KxNr88YwWBJ6fuwF+T5X9PLBrnk9Rb+jgQi2NyALe2DDxdLDAdOPPyuOuMAuevv74ELLAC5JwbaDz6H8MAQTY39PfGA2uL3I98CEI3C9wRbqcDc9St/L54AAXuDbnr74wi4B5A9cCEfffg88elrjBSSGuSD/iHBGA8zYC3XnnALcEWB4N7+mEQlD9k8eG/BA4tgTcgB/hycF4BNhckcEG+BUBbWPi5488CENzcg824F+RjNwIvcXv/AF8cByenF+nnf3xnO7zt5X/fhUIxXcN1hyt734tjDbmzWHUehwU7bDaB0HQc4Nzf+6evFv8AqcCEAUHjzPW3TB7Wa3FgR16/A4Kth9tgCelxwT74MDaykkDkDjgfLywiFjoVXxra4uDfpfp8RhPwqS1rg+vX/PBy3QkEXP8AXOCAkC63v19ecAQjjxCxU9LceXPT4YE8/wC85b7Q9vbAKt7KouR4uVBsPP8ADGFgU5O4Mb8/w98CEQjrcAefP8sH2yeYNvU+YwBAJPIPHAH9cYN/eUKASQBa1vw64EILbhcPa5sT5D3OM2hrsLAcXPv6fHBSxUC1jbg/9MHAuN/kpsCB09jhULD5nrx0vYf5fHAJdVuxIHXn88CzWsvNib2A5I/r5YBVAuSb/Bfz+GBCN3fhPiBC82vb9+AZLg2XoOb/ALzgRdWKm9gOAep9wP4emM5PIQttPn0YemBCIPElhe/mD/XTAg8kj06gcn8fL4YELe6lblefF1A/rzGAK7nARbluLMQf6PvhEIoAa7c3PvYH/PGArcDnj2/hgxW6gXBF/Q/hjLWCgngeXXj39MCEG0W6csCPX5YIRtAJNh5k9PnhReVBLDrt9FwVrnwW4uT9nobYVCTte4JA8rk9B6YAhW8rH064MEO5vDawvY+WBt0NhYeVr4EJIr52wRhbgg3wt187+H1wkQb/ACwiEQg+YwFvPChW/QYAqfO98KhE6cYD3wYofTGW88CEMU0sLbonKnzt54ew5pwElAA6Gwup+Xl8vwww87HAYY5jXZpzXluSsbIO0XPMpoo8uNTHmGVxG6UdbeaGL/0bXDwH/AVHqDiW5XqfTmcNYVH6OlmuGpq2QGJr8EJUAAcjykCn3OKOWR4m3xuVPqDbDqDMXjPi4v1K+fxHT92M6bRzXm83A9dbVdjthbgV0ZR1FVlNdFDmgqI4pLq55UlT9moQjqVNiSp/at5Y3mT9oOrOz1Kr+zOdPTUmYyCPNMukp46vLMxJ6fW6GYGCYN037Ve5BDDdigtO9oed5HAKOlqI56Etuaiqo+/pSfXuz4o2/vIVPvid5FrXS+bRz085/RktVCYu5r5TLS7r3Ux1NtyWPRZQV5+2OMZbrPaLK++OI+P1G1XRLFM26efz+itL9MaD1LW5fnlL2PZe2a08TUkGXy55U1umqdwS4aHLZP16OSTamknNOCSRuHgxpdSa6zLVlRBLqHNJs0alRYoDIyRQ0kZ57qmgjUQ0q7QR+qRQLi+4g4biirWpaSoOWsJJRKtQSQsSKig72mvsEZJDKxaw5scQ7NdYaUyWeWWKdtQ1rk94tNKYaMG97PPbfKbgEiMAf3+TiHtJLU6gx8PfV6BP7NkArl4qXUkNZnlS9JQQy91CX/Vo3ECg9ZWNgCBtBZiB1NrY0WotWaIyF+7bNJM9qob2pcpk2xb7i/eVRBULwPDErnj7Q64rnVvaBqDUsZps2rYqWg3b0y6jhFPSg+vcry7f35CW98RGXM2+zTpb+8wufw6DF2z6Ne43nnDYPc/FPFVprYBg3ryU31D2iahzqleglrIcmyeT/wCx2WgxRyD/ALx7mSc+7lvliGTZrFEO7oadVA6Mw/h/O+Ne8jysXkcsx8ybnBbY14rJHEKAde/mqD53PNUeWaWZ980hc+pOCdcCF9TxgwUk9PY4sgUwCgJqgC+mDKh4v08uPPBlQtwBfy9sKBbE8n3F+h+GFQiFL9BfB0UEbbgX64Ut5X+yQbgcj+jgQCAeSPMAm35YRCLt8RIIHxPT2xhFuSevmfjgxB8QK25ta/l8fXGKAOABz52tcYELAAF6dTxg5JBC2I2m3I5HtgpAsPW/r/DAgdCbKL9SMCEJIt1J4t/L4jALIRxfjzv6H1Hpgq2uCbXHW/PXyIwoI1sRcmw6C37j1/HCoWE7jboFPHNyPfA7QBYr+B/Dp5fzxlgpO1ebDp5YBmU+LcGW/wBocE/Lr+OEQjPst1ufPxE/1/XOAYC54JsAT05979beuMSwHHPobW2/10t88YOqlDbi9h1tgQjbtv2iT7jg/DjBwfDtBtzybi/wJHngm1SBYgWN+ABb3NsYTsBuPK1unH8v6vgQjmQk8vuY2ueCSfM8YTkkAPFgb8eWA3MSxJ5uBzwSPYeeALG53A2YdLkXwIWNtubXvzwbXv8A9cZuu24FRc9R09fljBcNYtZSRe91AHv7YBQL7lBB+0OLEe/88CEJswJJFuOLdT/l/HBlYFLEliLte/Qjp7/jjByvPkL8ngfzwYiwUMTu8vb/AC/oYVCBWux2n7IuR5EfvA/PGKd1iV5vcj09ueRgpYFb7zxyD9ofhgCVIuSB7kfuwiFh8QJ3XHXnqebfC/ngpVvYWtYk8j0wc2C7lsOBz7+3kfb0wVrKlr3v4r9AB/Xn5YELBbxFVN+vXp8MENgALdeevUeWMVuBcrfy9D/Xrgy26hh5keZHr8sKhFKlbgcbetz+XHTBr7RuuQPXp+OBlvcBhtt4T6L8MFB2i9+nAsfL3+eEQsHT22gX45xjhdqi3h87cAfyxilA3j5J6sWA/M4xnJUHdax4senwwqEUgXIK9Bz5WHvjDzYsV683+Hn/ADwDKLg2sbcC3B9wemBIPUEAdR53H8cIhZcE3Yk+ZN+bfPBSbXF+lrD2wN245Fr3sPI+vxwUtckG3h4tfnCoQeZ/HnAqp9eetv8ArgLiwtf4njBiegNgQOmBCDkCxPQ3Ht7Yw8Hbbz88ARY2Ppfr1+GAsb8Lx58cDAhAxF+gF7/M4MB4T16Xv6YKOSDzccc84zgcCxHvzbAhGIIXjjjr+/BG3XtfnywbcTwcFIK2B4/LCIWAXYqAAcZgqgHgevn5YzAhBcX56HAg+/yGMPHB4wIHAFsKhCBc9eBjASB6+eMUgHocYxvcXuPMDzwIWNe1yOvp64E24LcXAwU2sORx1tjAbefJwIRuLfZ58/hjL88c+dhjOSCT5efngOD7euBCFWIFvywIuPPmxv14wUWJuQbefGBv6EW6cDAhCLi3T4e+BQG918xx/lgoHQ26e+DLbdz8+euBCwehHTrfBtxC23E29enP54C5UAEC4/LAAjggEXHn54EIRvJIa/A+HOM44FwB6YMCoFl4FrXPJwBAHhvwTcjzwIWC1txv1t/XkcCHU8kkH1tyMBuuQWG7j1/djOpF2ubcdefhhAhZdhb7pPXBr8EXscFUC9ub3FrEfhzgQpFhb168YEIbE8bT6dPXpgBe4QE8n35OM4HnxbgdcYxBAHlYG9sKhDwTusCLXO4dcYLg9PInk9MFJPmPPGcX4/D0wIRwRsIItexJ9v6/lgNwvY82HNj/ABwVSOQTze1jg3NgQLeXB88IhZax6rbkXP8AXrg3huDew8gOMAoIBJLeGxNuLYEMeSoI9ABz/l+7CoQdLGwIvewJ/oYE3YgBi1/P+umMU9bGxA446j0wUAsp4FgL2HT8MIhYCbkXNzyfX3woSCvPlxyL3/r38sJ3N7qxAv15+QwO4g8/A2FrDCoR2ANt9iB8yMAGuPCTz6cXP78YG3i4I3XvyQL/AC9cYPtsSLnqSbm/8cCEdSGNkNx0HF/y8x+eCkAKLAW5/H3wUk7QLXN7+v5WwJ56cXAF7dMIhAb7b+Z63/d7YAqSbeYF7YPYh7q9yo8jyPXpx8cA37IQDzAsB87YVCIFAUlT4h1Nug98Y3UlbgNxYHr/ADwNg242vYceEm3zBwa5uVBIvcnn+vywISJFvO9vPBdniIte/phYKo5FjYXPN7D39/PCbr5efr1ucCEVRxb0PnjCoItcDy+GD2W/IANrWA6e+DWsDf1/q2BCQZfIj4WwUrcWvhcrfzsPO/F8E28nkH0NsCEgU9ucBb0wqUU8+mAsOthz64EJIixxmFCvHngpX08sCECkqdymxHmMOoa9kbdJuv8AtIbH5jocNMZ8sNc0OzShxbkpEmfz/of9FTZrM1AsvfrSd6/c95a27ub7d3vjV1GaSyNeEbLdGPLfyHywxwIUnEbYGNNU90rnICSxJYkk9ScCL34GB28XHODgc8+mJlGiAe2MA9fLC3d+RHw5wIjNibk34HPQ4EJOw6C/qOOuDrH0IBIPmMHVR5H0thQICRfkHz8sCERABcEc26YEjm4a/H4/HA8W2jd5gep+Prgykfl5XuR/DAhZtNgSBYdT5H2xgAtfpbnn0/jgxsOePj/P1xnB4tYg3uT5YEJN1Ci1wQBx/LCkYBvYXPTm9j6f16jAkEE3NrfePAt/LAhSSUJYgcHzJ+XngQsC8Bj0a497+YxnO3aR0FiT1t/L4YNYbVN9xIte174LwAqpfk3sCRz0+OBCBhe9yCwFyCbWH8cCCoU+K/HFhx+Pl8cYoFybCx9vPBWstl629+vxB88CFknNiCtr+fp74B1blW4cep5t74HcH6Jx5C/8TgbA9G2i/pbn1tgQhTaUUHhiegFhf4np+7GFxc2dSevHF7emCgEnxXA9h0+AwZSSB47WPm1vxwIQlm232ngkeXJPw8/34Lew6lhbyNhfGEgN0t5G3mPl1GANidxB5bz5v8cIhAXK3HIA8v8AL+WBKKgI4A8/b39BjLLfhQLX6dAfLjywZVIsBZfTyJ9hgKFnKDxXv0IPXGFgVsGvze1zz8v44AnnaLW8+h/6YwsOSCtxfm4v8sKhG22UFgBzuHw+P8PL54K7G1/lceftY/wxhdRax8J5Xn8/Y4TI38WBJNuvrhEIF3tdrdDxb93xxl2CdTweQfL5YM3iPiIPFr+QwIUiyWNx7cH+GBCAEKLBug8ufnbz98Zc9bjnyHl8P5YwA2UC6qL29r9cF9hwPfrgQsJAO5hfzt6j5YHeRyepFhfocBxcNx624vjLEcEbbDyNr/5YVCNuY3J9Ovt+/CZIb7TXAPPHl64OTYXv1NltwcE4I8xfjj+GBCze5J8Ra/Wx64LvIPPn+WBCG1+nPkODgQthfcRxfg9cIhYSAoBsevF8GBuSTf1NvI/zwVSLeIMR7eftfA7iQC1yQLDk4VCA3tfrxzbj+hgrWtwLm9ib9cHPPFxwep5OEyVDXJv87YELL2s3Ngf6GBAHUEWvx7YxhcAFR7YywuAOD7nAhCSx6nqb29/XAMtwW6/uwPCtyCFHQeeAYg38VuR8cCEUdQep9DjLsLg9OnIwJPUkgk8/54wWJuDex48sCEBuAefY4KbgncbH3wfnrfz6e+CtxyAb9TfAhEVrMRa59BjMYLeXHrjMJVCw/aN/XAgDjrbGYzCoWdbi/wDngRYe1/yxmMwIWetug9r4DnyItjMZhEIRyTZumBB9r268YzGYVCA2NweffBiAOfX8sZjMCFlvInr88ZfkC44GMxmBCxbA3BFuowbggEg/xxmMwIQnaLc8/wBeeDXNyGNgDe3W3w9cZjMCEB8xcX8h/n64DeTwWsBc8fnjMZgQsA6Gwt15F7YNcAXPHPn/ABxmMwIWEcGwHToRzfGEEckdcZjMCEDABSdw+IwRX9R1/LGYzAhDbm37XrhRLAXYbieovjMZgQjM21dw2En0PT5DBNxAAPQ+Xkf54zGYRCEuApAA59unw+OMVgbX8vMdcZjMAQsuerkHi1uTf4YAWF13Hn15sP44zGYVCMpINgSCfQ8YwkbAbg39+f8AL44zGYEiLuN9u4E/CwP9e+DLuPiNx5fA4zGYRKjg+rWPrbp7/wDTGAWsQDzfn197g4zGYVCTYi4A54tzjFcr4dwtf5E/154zGYRCNe3Llxt5Fhc/D0wS92BBCg8i37v68sZjMKhKdPEFFv2RwL/LBCTzY/O+MxmBCAnd0uCevS3ywF1PAUj482/njMZgQsJAHlfqPbBFU24NvnbGYzCIWML9LH4HphPbcEg/jjMZgQilRu4F8AFvjMZgCEbbby6YOqC174zGYVCMEtfzPt6YxEsPF1/fjMZgQlLeHr8Rbr8/TAhGJv7cW6YzGYELL2IFrH1wcX4v0vcjcTf5AYzGYEJMNc3HuTf2wYNYkeH4df6+OMxmBCV6CxHl5m5t/X+eAuxADEhb36m1/X8OL4zGYRCFVbaWNhttfnjnpxg3gXbYEC3nzYfLy9sZjMCEWQyKQTYnztYj8elsASNhPAuefX/pjMZgQiA9SLXHQ2xh9SByOL+nr7fHGYzAhYD0NlK/gT7H0+OBBNtwtb098ZjMAQjA7mKsV3E+nn+4YwgKSN/PW5BA+eMxmFQiWYHoQCOp8v69cCfCeAeRa/Uf18sZjMJkhY8jBRe4t0G4k4FWBvbcfQk+frjMZgQigG91YA+Vx19f+mAcE3BsPPg4zGYM0IisQ3hIv5eXP9euDFrWFrXPp1xmMwIQqbEqQUPQknp/lgzEsLea+QPHx98ZjMCFgPhD2Hz8/b4YAHgqt/F5Xt06AnzxmMwqFm8gEKxU9SOgwXqSQAAOBf0/hjMZgQsY2FgSPj54Abrm7Dk3vbGYzAhYSDySF8j6jBeR53HX3xmMwIWWJANxx+fywJItcWHPHXjGYzCIQbiCBwPf+vLAKSSG5uOnrjMZhUIDwLXt8PXGBiSxPn5YzGYELOR1+AI64wsSSCDf2xmMwIRXI9eR6YwcW59wMZjMCEYHgncLnrxhM8k9evQnGYzAhFuB164zGYzAhf/Z"""


def escrever_asset_padrao(nome_arquivo: str, conteudo_b64: str) -> Optional[Path]:
    try:
        caminho = DOSSIE_ASSETS_DIR / nome_arquivo
        if not caminho.exists():
            caminho.write_bytes(base64.b64decode(conteudo_b64))
        return caminho if caminho.exists() else None
    except Exception:
        return None


def caminho_arquivo_configurado(valor: str, nomes_padrao: Optional[List[str]] = None) -> Optional[Path]:
    candidatos: List[Path] = []
    valor = str(valor or "").strip()
    if valor:
        p = Path(valor)
        candidatos.append(p if p.is_absolute() else BASE_DIR / p)
    for nome in nomes_padrao or []:
        candidatos.append(BASE_DIR / nome)
        candidatos.append(DATA_DIR / nome)
        candidatos.append(DOSSIE_ASSETS_DIR / nome)
    for caminho in candidatos:
        try:
            if caminho.exists():
                return caminho
        except Exception:
            pass
    return None

DICOR_AZUL = "002B5B"
DICOR_DOURADO = "D4AF37"
DICOR_CINZA = "F4F6F8"


def caminho_brasao_pf() -> Optional[Path]:
    externo = caminho_arquivo_configurado("", ["brasao_pf.png", "brasao_pf.jpg", "pf.png", "pf.jpg"])
    caminho = externo or escrever_asset_padrao("brasao_pf_padrao.jpg", DOSSIE_BRASAO_PF_PADRAO_B64)
    return limpar_imagem_brasao_dossie(caminho)


def caminho_brasao_dicor() -> Optional[Path]:
    externo = caminho_arquivo_configurado("", ["brasao_dicor.png", "brasao_dicor.jpg", "dicor.png", "dicor.jpg"])
    caminho = externo or escrever_asset_padrao("brasao_dicor_padrao.jpg", DOSSIE_BRASAO_DICOR_PADRAO_B64)
    return limpar_imagem_brasao_dossie(caminho)


def obter_assinaturas_dossie(dados: Dict[str, Any]) -> List[Dict[str, Any]]:
    registros = carregar_assinaturas_dossie()
    delegado_geral = registros.get("delegado_geral", {})
    delegado_dicor = registros.get("delegado_dicor", {})
    agente_id = str(dados.get("agente_encerramento_id") or "")
    agente_reg = registros.get(f"agente_{agente_id}", {}) if agente_id else {}

    return [
        {
            "titulo": "DELEGADO GERAL",
            "nome": delegado_geral.get("nome") or ASSINATURA_DELEGADO_GERAL_NOME,
            "texto": delegado_geral.get("texto") or ASSINATURA_DELEGADO_GERAL_TEXTO,
            "imagem": caminho_assinatura_registrada(delegado_geral) or caminho_arquivo_configurado(ASSINATURA_DELEGADO_GERAL_IMAGEM, [
                "assinatura_delegado_geral.png", "assinatura_delegado_geral.jpg",
            ]),
        },
        {
            "titulo": "DELEGADO DICOR",
            "nome": delegado_dicor.get("nome") or ASSINATURA_DELEGADO_DICOR_NOME,
            "texto": delegado_dicor.get("texto") or ASSINATURA_DELEGADO_DICOR_TEXTO,
            "imagem": caminho_assinatura_registrada(delegado_dicor) or caminho_arquivo_configurado(ASSINATURA_DELEGADO_DICOR_IMAGEM, [
                "assinatura_delegado_dicor.png", "assinatura_delegado_dicor.jpg",
            ]),
        },
        {
            "titulo": "AGENTE RESPONSÁVEL",
            "nome": agente_reg.get("nome") or dados.get("agente_encerramento") or "Agente responsável",
            "texto": agente_reg.get("texto") or ASSINATURA_AGENTE_RESPONSAVEL_TEXTO,
            "imagem": caminho_assinatura_registrada(agente_reg) or caminho_arquivo_configurado(ASSINATURA_AGENTE_RESPONSAVEL_IMAGEM, [
                "assinatura_agente_responsavel.png", "assinatura_agente_responsavel.jpg",
            ]),
        },
    ]


def linhas_tabela_conteudo() -> List[List[str]]:
    return [
        ["Página", "Seção"],
        ["1", "Capa"],
        ["2", "Resumo Executivo e Índice"],
        ["3", "Lideranças Identificadas"],
        ["4", "Integrantes Identificados"],
        ["5", "Painel da Organização"],
        ["6", "Localização"],
        ["7", "Produção e Fabricação"],
        ["8", "Baús e Armazenamento"],
        ["9", "Informantes"],
        ["10", "Materiais Apreendidos"],
        ["11", "Resultado Operacional"],
        ["12", "Conclusão"],
    ]


def montar_conclusao_dossie(dados: Dict[str, Any]) -> str:
    est = dados.get("estatisticas", {}) or {}
    res = dados.get("resultados", {}) or {}
    return (
        f"A mesa de investigação referente ao processo {dados.get('processo')} foi encerrada na circunscrição operacional de "
        f"{dados.get('cidade_operacional', DOSSIE_CIDADE_OPERACIONAL)}, após a consolidação dos registros vinculados "
        f"à operação {dados.get('nome_operacao')}. Durante a análise, foram "
        f"verificadas {est.get('mensagens_analisadas', 0)} mensagens, {est.get('evidencias', 0)} evidências anexadas, "
        f"{est.get('imagens', 0)} imagens, {est.get('videos', 0)} vídeos e {est.get('links', 0)} links registrados. "
        f"Os elementos reunidos indicam a necessidade de preservação integral do material para consulta posterior, "
        f"subsidiando ações de inteligência, diligências futuras, cruzamento de dados e eventual reabertura da investigação. "
        f"Resultados operacionais informados: prisões ({res.get('prisoes', 'Não informado')}), drogas ({res.get('drogas', 'Não informado')}), "
        f"armas ({res.get('armas', 'Não informado')}), veículos ({res.get('veiculos', 'Não informado')}) e materiais diversos "
        f"({res.get('materiais', 'Não informado')}). O dossiê é arquivado com autenticação digital e QR Code de reabertura."
    )


def resumo_contexto_operacional(dados: Dict[str, Any]) -> str:
    return (
        f"A investigação foi conduzida pela Polícia Federal de {dados.get('cidade_operacional', DOSSIE_CIDADE_OPERACIONAL)}, "
        f"tendo como alvo a comunidade/organização {dados.get('comunidade', 'Não informado')}, "
        f"com possível vínculo à facção ou estrutura {dados.get('faccao', 'Não informado')}. O objetivo operacional foi: "
        f"{dados.get('objetivo', 'Não informado')} O material abaixo foi extraído automaticamente da mesa de investigação, "
        f"incluindo tópicos, anexos, links, registros textuais e dados enviados pelos integrantes autorizados."
    )


def pdf_paragrafo(texto: Any, estilo) -> Any:
    seguro = escape(texto).replace("\n", "<br/>")
    return Paragraph(seguro or "Não informado", estilo)


def pdf_img_fit(caminho: str, max_w: float, max_h: float) -> Optional[Any]:
    if not caminho or not Path(caminho).exists() or RLImage is None:
        return None
    try:
        w, h = max_w, max_h
        if PILImage is not None:
            with PILImage.open(caminho) as img:
                iw, ih = img.size
                if iw > 0 and ih > 0:
                    escala = min(max_w / iw, max_h / ih)
                    w, h = iw * escala, ih * escala
        return RLImage(caminho, width=w, height=h)
    except Exception:
        return None


def pdf_tabela_metadados(dados: List[List[Any]], col_widths: List[float], estilo_body) -> Any:
    linhas = []
    for row in dados:
        linhas.append([pdf_paragrafo(c, estilo_body) if not hasattr(c, "wrap") else c for c in row])
    tabela = Table(linhas, colWidths=col_widths, hAlign="LEFT")
    tabela.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#C7CCD4")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{DICOR_AZUL}")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("PADDING", (0, 0), (-1, -1), 6),
    ]))
    return tabela


def pdf_add_header_footer(c, doc, dados: Dict[str, Any]):
    c.saveState()
    largura, altura = A4
    pagina = c.getPageNumber()

    # Marca d'água dupla e discreta
    c.setFont("Helvetica-Bold", 58)
    c.setFillColor(colors.Color(0.82, 0.86, 0.91, alpha=0.16))
    c.translate(largura / 2, altura / 2)
    c.rotate(35)
    c.drawCentredString(0, 0, "DICOR")
    c.setFont("Helvetica-Bold", 18)
    c.drawCentredString(0, -1.05 * cm, "INTELIGÊNCIA OPERACIONAL")
    c.rotate(-35)
    c.translate(-largura / 2, -altura / 2)

    # Faixa superior azul + linha dourada
    c.setFillColor(colors.HexColor(f"#{DICOR_AZUL}"))
    c.rect(0, altura - 0.62 * cm, largura, 0.18 * cm, stroke=0, fill=1)
    c.setStrokeColor(colors.HexColor(f"#{DICOR_DOURADO}"))
    c.setLineWidth(1)
    c.line(1.45 * cm, altura - 1.35 * cm, largura - 1.45 * cm, altura - 1.35 * cm)

    c.setFillColor(colors.HexColor(f"#{DICOR_AZUL}"))
    c.setFont("Helvetica-Bold", 8.7)
    c.drawString(1.5 * cm, altura - 1.1 * cm, f"POLÍCIA FEDERAL • DICOR • {dados.get('cidade_operacional', DOSSIE_CIDADE_OPERACIONAL).upper()}")
    c.setFont("Helvetica", 8)
    c.drawRightString(largura - 1.5 * cm, altura - 1.1 * cm, f"Processo: {dados.get('processo', 'N/I')}")

    # Selo lateral reservado
    c.setFillColor(colors.HexColor(f"#{DICOR_DOURADO}"))
    c.roundRect(largura - 4.0 * cm, altura - 1.75 * cm, 2.45 * cm, 0.36 * cm, 4, stroke=0, fill=1)
    c.setFillColor(colors.HexColor(f"#{DICOR_AZUL}"))
    c.setFont("Helvetica-Bold", 6.8)
    c.drawCentredString(largura - 2.78 * cm, altura - 1.62 * cm, "DOCUMENTO RESERVADO")

    # Rodapé
    c.setStrokeColor(colors.HexColor("#C7CCD4"))
    c.line(1.45 * cm, 1.25 * cm, largura - 1.45 * cm, 1.25 * cm)
    c.setFillColor(colors.HexColor(f"#{DICOR_AZUL}"))
    c.rect(0, 0, largura, 0.20 * cm, stroke=0, fill=1)
    c.setFont("Helvetica", 8)
    c.setFillColor(colors.HexColor("#555555"))
    c.drawString(1.5 * cm, 0.85 * cm, f"Polícia Federal de {dados.get('cidade_operacional', DOSSIE_CIDADE_OPERACIONAL)} • Sistema DICOR • Uso interno GTA RP")
    c.drawRightString(largura - 1.5 * cm, 0.85 * cm, f"Página {pagina}")
    c.restoreState()


def pdf_add_secao_titulo(story: List[Any], titulo: str, estilo_h1) -> None:
    faixa = Table([[Paragraph(titulo, estilo_h1)]], colWidths=[16.5 * cm], hAlign="LEFT")
    faixa.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F7F9FC")),
        ("LINEBEFORE", (0, 0), (0, 0), 4, colors.HexColor(f"#{DICOR_DOURADO}")),
        ("LINEBELOW", (0, 0), (-1, -1), 0.8, colors.HexColor(f"#{DICOR_DOURADO}")),
        ("PADDING", (0, 0), (-1, -1), 7),
    ]))
    story.append(faixa)
    story.append(Spacer(1, 10))


def pdf_add_imagens_evidencias(story: List[Any], evidencias: List[Dict[str, Any]], estilo_body, limite: int = 6) -> None:
    imagens = [ev for ev in evidencias if ev.get("tipo") == "imagem" and ev.get("local")][:limite]
    if not imagens:
        story.append(pdf_paragrafo("Nenhuma imagem anexada nesta seção.", estilo_body))
        return
    for idx, ev in enumerate(imagens, start=1):
        img = pdf_img_fit(ev.get("local", ""), 8.5 * cm, 5.5 * cm)
        legenda = pdf_paragrafo(
            f"Imagem {idx} • {ev.get('arquivo')} • {ev.get('data')} • Agente: {ev.get('autor')} • Origem: {ev.get('origem')}",
            estilo_body,
        )
        if img:
            t = Table([[img, legenda]], colWidths=[9 * cm, 7 * cm], hAlign="LEFT")
            t.setStyle(TableStyle([
                ("BOX", (0, 0), (-1, -1), 0.45, colors.HexColor("#C7CCD4")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("PADDING", (0, 0), (-1, -1), 6),
            ]))
            story.append(t)
            story.append(Spacer(1, 6))


def pdf_add_pessoas(story: List[Any], pessoas: List[Dict[str, str]], titulo_vazio: str, estilo_body) -> None:
    if not pessoas:
        story.append(pdf_paragrafo(titulo_vazio, estilo_body))
        return
    cab = ["Foto", "Nome", "RG", "Função/Cargo", "Periculosidade/Obs."]
    rows = [cab]
    for p in pessoas[:30]:
        foto = pdf_img_fit(p.get("foto", ""), 2.2 * cm, 2.2 * cm) or pdf_paragrafo("Sem foto", estilo_body)
        rows.append([
            foto,
            pdf_paragrafo(p.get("nome", "Não informado"), estilo_body),
            pdf_paragrafo(p.get("rg", "Não informado"), estilo_body),
            pdf_paragrafo(p.get("funcao") or p.get("cargo") or "Não informado", estilo_body),
            pdf_paragrafo(f"{p.get('periculosidade', 'Não informado')}\n{p.get('observacoes', '')}", estilo_body),
        ])
    tabela = Table(rows, colWidths=[2.7 * cm, 4.1 * cm, 2.7 * cm, 3.4 * cm, 4 * cm], repeatRows=1, hAlign="LEFT")
    tabela.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{DICOR_AZUL}")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#C7CCD4")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("PADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(tabela)


def pdf_add_assinaturas_dossie(story: List[Any], dados: Dict[str, Any], style_center) -> None:
    """Adiciona assinaturas usando arquivo/imagem ou texto."""
    assinaturas = obter_assinaturas_dossie(dados)
    linha_imagens = []
    linha_vazia = []
    linha_titulos = []

    for ass in assinaturas:
        imagem_assinatura = limpar_imagem_assinatura_dossie(ass.get("imagem")) or ass.get("imagem")
        img = pdf_img_fit(
            str(imagem_assinatura or ""),
            5.25 * cm,
            2.35 * cm,
        )
        texto_assinatura = str(ass.get("texto") or "").strip()
        if img:
            linha_imagens.append(img)
        elif texto_assinatura:
            linha_imagens.append(pdf_paragrafo(texto_assinatura[:400], style_center))
        else:
            linha_imagens.append(Spacer(1, 1.15 * cm))
        linha_vazia.append("")
        linha_titulos.append(pdf_paragrafo(ass.get("titulo"), style_center))

    tabela = Table(
        [linha_imagens, linha_vazia, linha_titulos],
        colWidths=[5.55 * cm, 5.55 * cm, 5.55 * cm],
        hAlign="CENTER",
    )
    tabela.setStyle(TableStyle([
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LINEABOVE", (0, 1), (0, 1), 0.8, colors.HexColor("#333333")),
        ("LINEABOVE", (1, 1), (1, 1), 0.8, colors.HexColor("#333333")),
        ("LINEABOVE", (2, 1), (2, 1), 0.8, colors.HexColor("#333333")),
        ("PADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 2),
        ("TOPPADDING", (0, 2), (-1, 2), 8),
    ]))
    story.append(tabela)
    story.append(Spacer(1, 0.45 * cm))
    story.append(pdf_paragrafo(
        f"Assinaturas digitais vinculadas ao processo {dados.get('processo')} • Gerado em {dados.get('data_encerramento')}",
        style_center,
    ))

def gerar_pdf_dossie(dados: Dict[str, Any], caminho_pdf: Path) -> None:
    if SimpleDocTemplate is None:
        raise RuntimeError("Dependência ausente: instale reportlab, qrcode e pillow.")

    caminho_pdf = Path(caminho_pdf)
    caminho_pdf.parent.mkdir(parents=True, exist_ok=True)

    doc = SimpleDocTemplate(
        str(caminho_pdf),
        pagesize=A4,
        rightMargin=1.6 * cm,
        leftMargin=1.6 * cm,
        topMargin=1.8 * cm,
        bottomMargin=1.8 * cm,
        title=f"Dossiê Operacional {dados.get('processo')}",
        author="Polícia Federal - DICOR",
    )

    styles = getSampleStyleSheet()
    style_title = ParagraphStyle("DICORTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=24, leading=29, alignment=TA_CENTER, textColor=colors.HexColor(f"#{DICOR_AZUL}"), spaceAfter=8)
    style_subtitle = ParagraphStyle("DICORSub", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=11.2, leading=15, alignment=TA_CENTER, textColor=colors.HexColor(f"#{DICOR_DOURADO}"), spaceAfter=10)
    style_h1 = ParagraphStyle("DICORH1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=16.5, leading=20, textColor=colors.HexColor(f"#{DICOR_AZUL}"), spaceBefore=8, spaceAfter=6)
    style_h2 = ParagraphStyle("DICORH2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=12.2, leading=15.5, textColor=colors.HexColor("#333333"), spaceBefore=6, spaceAfter=4)
    style_body = ParagraphStyle("DICORBody", parent=styles["Normal"], fontName="Helvetica", fontSize=10.2, leading=14.2, alignment=TA_JUSTIFY, textColor=colors.HexColor("#222222"), spaceAfter=6)
    style_center = ParagraphStyle("DICORCenter", parent=style_body, alignment=TA_CENTER)

    story: List[Any] = []

    # Página 1 — Capa
    brasao_pf = caminho_brasao_pf()
    brasao_dicor = caminho_brasao_dicor()
    img_pf = pdf_img_fit(str(brasao_pf), 3.45 * cm, 3.45 * cm) if brasao_pf else pdf_paragrafo("PF", style_center)
    img_dicor = pdf_img_fit(str(brasao_dicor), 3.45 * cm, 3.45 * cm) if brasao_dicor else pdf_paragrafo("DICOR", style_center)
    cab = Table([[img_pf, pdf_paragrafo(dados.get("cabecalho_institucional") or DOSSIE_INSTITUICAO_CABECALHO, style_center), img_dicor]], colWidths=[4.05 * cm, 9.3 * cm, 4.05 * cm])
    cab.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("ALIGN", (0, 0), (-1, -1), "CENTER")]))
    story.append(cab)
    story.append(Spacer(1, 0.75 * cm))
    story.append(Paragraph("POLÍCIA FEDERAL", style_title))
    story.append(Paragraph(f"CAPITAL MORADA DO VALLEY • DIRETORIA DE INVESTIGAÇÃO E COMBATE AO CRIME ORGANIZADO - DICOR", style_subtitle))
    faixa = Table([[pdf_paragrafo("DOCUMENTO RESERVADO • INTELIGÊNCIA • INVESTIGAÇÃO • RESULTADO", style_center)]], colWidths=[16 * cm])
    faixa.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(f"#{DICOR_AZUL}")),
        ("TEXTCOLOR", (0, 0), (-1, -1), colors.white),
        ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor(f"#{DICOR_DOURADO}")),
        ("PADDING", (0, 0), (-1, -1), 8),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
    ]))
    story.append(faixa)
    story.append(Spacer(1, 0.25 * cm))
    selo_operacional = Table([[
        pdf_paragrafo("UNIDADE OPERACIONAL", style_center),
        pdf_paragrafo(dados.get("cidade_operacional", DOSSIE_CIDADE_OPERACIONAL).upper(), style_center),
        pdf_paragrafo("ARQUIVO DICOR", style_center),
    ]], colWidths=[5.33 * cm, 5.33 * cm, 5.33 * cm], hAlign="CENTER")
    selo_operacional.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F7F9FC")),
        ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor(f"#{DICOR_DOURADO}")),
        ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#C7CCD4")),
        ("PADDING", (0, 0), (-1, -1), 7),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
    ]))
    story.append(selo_operacional)
    story.append(Spacer(1, 0.55 * cm))
    capa_rows = [
        ["DOSSIÊ DE INTELIGÊNCIA OPERACIONAL", ""],
        ["Processo Nº", dados.get("processo")],
        ["Investigação Nº", dados.get("numero_investigacao")],
        ["Nome da Operação", dados.get("nome_operacao")],
        ["Comunidade Investigada", dados.get("comunidade")],
        ["Cidade Operacional", dados.get("cidade_operacional", DOSSIE_CIDADE_OPERACIONAL)],
        ["Data de Abertura", dados.get("data_abertura")],
        ["Data de Encerramento", dados.get("data_encerramento")],
        ["Delegado Responsável", dados.get("delegado_responsavel")],
        ["Integrantes da Investigação", ", ".join(dados.get("integrantes_investigacao", []))],
    ]
    capa = []
    for i, row in enumerate(capa_rows):
        if i == 0:
            capa.append([Paragraph(f"<b>{escape(row[0])}</b>", style_center), ""])
        else:
            capa.append([pdf_paragrafo(row[0], style_body), pdf_paragrafo(row[1], style_body)])
    t_capa = Table(capa, colWidths=[5.2 * cm, 10.8 * cm])
    t_capa.setStyle(TableStyle([
        ("SPAN", (0, 0), (1, 0)),
        ("BACKGROUND", (0, 0), (1, 0), colors.HexColor(f"#{DICOR_AZUL}")),
        ("TEXTCOLOR", (0, 0), (1, 0), colors.white),
        ("BACKGROUND", (0, 1), (0, -1), colors.HexColor(f"#{DICOR_CINZA}")),
        ("BOX", (0, 0), (-1, -1), 1.2, colors.HexColor(f"#{DICOR_DOURADO}")),
        ("GRID", (0, 1), (-1, -1), 0.35, colors.HexColor("#C7CCD4")),
        ("PADDING", (0, 0), (-1, -1), 8),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.append(t_capa)
    story.append(Spacer(1, 2.0 * cm))
    story.append(pdf_paragrafo("Documento gerado automaticamente a partir do encerramento da mesa de investigação.", style_center))
    story.append(PageBreak())

    # Página 2 — Resumo Executivo + índice
    pdf_add_secao_titulo(story, "2. RESUMO EXECUTIVO", style_h1)
    story.append(pdf_paragrafo(dados.get("objetivo"), style_body))
    story.append(pdf_paragrafo(resumo_contexto_operacional(dados), style_body))
    meta = [
        ["Campo", "Informação"],
        ["Status da Investigação", dados.get("status")],
        ["Prioridade", dados.get("prioridade")],
        ["Cidade Operacional", dados.get("cidade_operacional", DOSSIE_CIDADE_OPERACIONAL)],
        ["Facção Investigada", dados.get("faccao")],
        ["Agente Responsável pelo Encerramento", dados.get("agente_encerramento")],
        ["Integrantes da Investigação", ", ".join(dados.get("integrantes_investigacao", []))],
    ]
    story.append(pdf_tabela_metadados(meta, [5.3 * cm, 11.2 * cm], style_body))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Tabela de Conteúdo", style_h2))
    story.append(pdf_tabela_metadados(linhas_tabela_conteudo(), [2.8 * cm, 13.7 * cm], style_body))
    story.append(PageBreak())

    # Página 3 — Lideranças
    pdf_add_secao_titulo(story, "3. LIDERANÇAS IDENTIFICADAS", style_h1)
    pdf_add_pessoas(story, dados.get("liderancas", []), "Nenhuma liderança foi identificada automaticamente nos tópicos da mesa.", style_body)
    story.append(PageBreak())

    # Página 4 — Integrantes
    pdf_add_secao_titulo(story, "4. INTEGRANTES IDENTIFICADOS", style_h1)
    pdf_add_pessoas(story, dados.get("integrantes", []), "Nenhum integrante foi identificado automaticamente nos tópicos da mesa.", style_body)
    story.append(PageBreak())

    # Página 5 — Painel
    pdf_add_secao_titulo(story, "5. PAINEL DA ORGANIZAÇÃO", style_h1)
    story.append(pdf_paragrafo(dados.get("resumos", {}).get("painel"), style_body))
    pdf_add_imagens_evidencias(story, filtrar_evidencias_por_topico(dados.get("evidencias", []), ["painel"]), style_body, 6)
    story.append(PageBreak())

    # Página 6 — Localização
    pdf_add_secao_titulo(story, "6. LOCALIZAÇÃO", style_h1)
    story.append(pdf_paragrafo(dados.get("resumos", {}).get("localizacao"), style_body))
    story.append(pdf_paragrafo(f"Comunidade/Base: {dados.get('comunidade')}\nCanal de reabertura: {dados.get('reabrir_url')}", style_body))
    pdf_add_imagens_evidencias(story, filtrar_evidencias_por_topico(dados.get("evidencias", []), ["localizacao"]), style_body, 6)
    story.append(PageBreak())

    # Página 7 — Produção
    pdf_add_secao_titulo(story, "7. PRODUÇÃO E FABRICAÇÃO", style_h1)
    story.append(pdf_paragrafo(dados.get("resumos", {}).get("producao"), style_body))
    pdf_add_imagens_evidencias(story, filtrar_evidencias_por_topico(dados.get("evidencias", []), ["producao"]), style_body, 6)
    story.append(PageBreak())

    # Página 8 — Baús
    pdf_add_secao_titulo(story, "8. BAÚS E ARMAZENAMENTO", style_h1)
    story.append(pdf_paragrafo(dados.get("resumos", {}).get("baus"), style_body))
    pdf_add_imagens_evidencias(story, filtrar_evidencias_por_topico(dados.get("evidencias", []), ["baus"]), style_body, 6)
    story.append(PageBreak())

    # Página 9 — Informantes
    pdf_add_secao_titulo(story, "9. INFORMANTES", style_h1)
    pdf_add_pessoas(story, dados.get("informantes", []), "Nenhum informante foi identificado automaticamente nos tópicos da mesa.", style_body)
    story.append(Spacer(1, 10))
    story.append(pdf_paragrafo(dados.get("resumos", {}).get("informantes"), style_body))
    story.append(PageBreak())

    # Página 10 — Materiais apreendidos
    pdf_add_secao_titulo(story, "10. MATERIAIS APREENDIDOS", style_h1)
    res = dados.get("resultados", {}) or {}
    materiais = [
        ["Categoria", "Registro"],
        ["Drogas", res.get("drogas", "Não informado")],
        ["Armas", res.get("armas", "Não informado")],
        ["Munições", res.get("municoes", "Não informado")],
        ["Dinheiro", res.get("dinheiro", "Não informado")],
        ["Veículos", res.get("veiculos", "Não informado")],
        ["Outros Itens", res.get("outros", "Não informado")],
    ]
    story.append(pdf_tabela_metadados(materiais, [4.2 * cm, 12.3 * cm], style_body))
    story.append(Spacer(1, 10))
    anexos = [["Tipo", "Arquivo", "Data", "Origem"]]
    for ev in dados.get("evidencias", [])[:30]:
        anexos.append([ev.get("tipo"), ev.get("arquivo"), ev.get("data"), ev.get("origem")])
    story.append(Paragraph("Provas e anexos registrados", style_h2))
    story.append(pdf_tabela_metadados(anexos or [["Sem anexos", "", "", ""]], [2.2 * cm, 6.2 * cm, 3.2 * cm, 4.9 * cm], style_body))
    story.append(PageBreak())

    # Página 11 — Resultado
    pdf_add_secao_titulo(story, "11. RESULTADO OPERACIONAL", style_h1)
    est = dados.get("estatisticas", {}) or {}
    tabela_resultado = [
        ["Indicador", "Quantidade/Status"],
        ["Quantidade de presos", res.get("prisoes", "Não informado")],
        ["Procurados capturados", res.get("procurados_capturados", "Não informado")],
        ["Mensagens analisadas", est.get("mensagens_analisadas", 0)],
        ["Evidências coletadas", est.get("evidencias", 0)],
        ["Imagens anexadas", est.get("imagens", 0)],
        ["Vídeos anexados", est.get("videos", 0)],
        ["Links registrados", est.get("links", 0)],
        ["Tópicos analisados", est.get("threads", 0)],
    ]
    story.append(pdf_tabela_metadados(tabela_resultado, [5.5 * cm, 11 * cm], style_body))
    story.append(PageBreak())

    # Página 12 — Conclusão
    pdf_add_secao_titulo(story, "12. CONCLUSÃO", style_h1)
    story.append(pdf_paragrafo(montar_conclusao_dossie(dados), style_body))
    story.append(Spacer(1, 1.2 * cm))
    pdf_add_assinaturas_dossie(story, dados, style_center)
    story.append(Spacer(1, 0.6 * cm))
    qr_path = dados.get("qr_reabertura")
    qr_img = pdf_img_fit(qr_path, 2.5 * cm, 2.5 * cm) if qr_path else None
    if qr_img:
        story.append(Paragraph("QR Code para reabrir a investigação arquivada", style_h2))
        story.append(qr_img)

    doc.build(story, onFirstPage=lambda c, d: pdf_add_header_footer(c, d, dados), onLaterPages=lambda c, d: pdf_add_header_footer(c, d, dados))


def docx_set_cell_shading(cell, fill: str) -> None:
    if OxmlElement is None:
        return
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def docx_set_cell_text(cell, texto: Any, bold: bool = False, color: Optional[str] = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(texto or "Não informado"))
    run.bold = bold
    if color and RGBColor is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if Pt is not None:
        run.font.size = Pt(9)


def docx_add_heading(doc, texto: str, level: int = 1):
    p = doc.add_heading(texto, level=level)
    try:
        p.runs[0].font.color.rgb = RGBColor.from_string(DICOR_AZUL)
        p.runs[0].font.name = "Arial"
        p.runs[0].font.bold = True
    except Exception:
        pass
    try:
        barra = doc.add_paragraph("▔" * 48)
        barra.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = barra.runs[0]
        run.font.color.rgb = RGBColor.from_string(DICOR_DOURADO)
        run.font.size = Pt(6)
    except Exception:
        pass
    return p


def docx_add_paragraph(doc, texto: Any, bold: bool = False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(str(texto or "Não informado"))
    run.bold = bold
    try:
        run.font.name = "Arial"
        run.font.size = Pt(10)
    except Exception:
        pass
    if align is not None:
        p.alignment = align
    return p


def docx_add_info_table(doc, rows: List[List[Any]], widths: Optional[List[float]] = None):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    try:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
    except Exception:
        pass
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            if r == 0:
                docx_set_cell_shading(cell, DICOR_AZUL)
                docx_set_cell_text(cell, value, bold=True, color="FFFFFF")
            else:
                if c == 0:
                    docx_set_cell_shading(cell, DICOR_CINZA)
                    docx_set_cell_text(cell, value, bold=True)
                else:
                    docx_set_cell_text(cell, value)
            try:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            except Exception:
                pass
    doc.add_paragraph()
    return table


def docx_add_picture_safe(paragraph, caminho: str, width_inches: float = 2.2) -> bool:
    try:
        if caminho and Path(caminho).exists():
            paragraph.add_run().add_picture(str(caminho), width=Inches(width_inches))
            return True
    except Exception:
        return False
    return False


def docx_add_pessoas(doc, pessoas: List[Dict[str, str]], vazio: str):
    if not pessoas:
        docx_add_paragraph(doc, vazio)
        return
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Foto", "Nome", "RG", "Função/Cargo", "Periculosidade/Obs."]
    for idx, h in enumerate(headers):
        cell = table.cell(0, idx)
        docx_set_cell_shading(cell, DICOR_AZUL)
        docx_set_cell_text(cell, h, bold=True, color="FFFFFF")
    for p in pessoas[:40]:
        row = table.add_row().cells
        row[0].text = ""
        par = row[0].paragraphs[0]
        if not docx_add_picture_safe(par, p.get("foto", ""), 0.9):
            docx_set_cell_text(row[0], "Sem foto")
        docx_set_cell_text(row[1], p.get("nome", "Não informado"))
        docx_set_cell_text(row[2], p.get("rg", "Não informado"))
        docx_set_cell_text(row[3], p.get("funcao") or p.get("cargo") or "Não informado")
        docx_set_cell_text(row[4], f"{p.get('periculosidade', 'Não informado')}\n{p.get('observacoes', '')}")
    doc.add_paragraph()


def docx_add_imagens_evidencias(doc, evidencias: List[Dict[str, Any]], limite: int = 8):
    imagens = [ev for ev in evidencias if ev.get("tipo") == "imagem" and ev.get("local")][:limite]
    if not imagens:
        docx_add_paragraph(doc, "Nenhuma imagem anexada nesta seção.")
        return
    for idx, ev in enumerate(imagens, start=1):
        p = doc.add_paragraph()
        if docx_add_picture_safe(p, ev.get("local", ""), 4.8):
            docx_add_paragraph(doc, f"Imagem {idx} • {ev.get('arquivo')} • {ev.get('data')} • Agente: {ev.get('autor')} • Origem: {ev.get('origem')}")


def configurar_docx(doc: Any, dados: Dict[str, Any]) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    header = section.header.paragraphs[0]
    header.text = f"POLÍCIA FEDERAL • DICOR • {dados.get('cidade_operacional', DOSSIE_CIDADE_OPERACIONAL)} • Processo {dados.get('processo', 'N/I')}"
    try:
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header.runs[0].font.bold = True
        header.runs[0].font.color.rgb = RGBColor.from_string(DICOR_AZUL)
        header.runs[0].font.size = Pt(9)
    except Exception:
        pass
    footer = section.footer.paragraphs[0]
    footer.text = f"Polícia Federal de {dados.get('cidade_operacional', DOSSIE_CIDADE_OPERACIONAL)} • Sistema DICOR • Documento reservado • Uso interno GTA RP"
    try:
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.runs[0].font.size = Pt(8)
    except Exception:
        pass


def docx_add_assinaturas_dossie(doc, dados: Dict[str, Any]) -> None:
    """Adiciona assinaturas no DOCX usando arquivo/imagem ou texto."""
    assinaturas = obter_assinaturas_dossie(dados)
    tabela = doc.add_table(rows=3, cols=3)
    try:
        tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    except Exception:
        pass

    for col, ass in enumerate(assinaturas):
        cell_img = tabela.cell(0, col)
        p_img = cell_img.paragraphs[0]
        try:
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass
        imagem = ass.get("imagem")
        imagem_limpa = limpar_imagem_assinatura_dossie(imagem) if imagem else None
        texto_assinatura = str(ass.get("texto") or "").strip()
        if imagem_limpa and Path(imagem_limpa).exists():
            try:
                p_img.add_run().add_picture(str(imagem_limpa), width=Inches(2.15))
            except Exception:
                p_img.add_run("\n")
        elif texto_assinatura:
            run_ass = p_img.add_run(texto_assinatura[:400])
            try:
                run_ass.italic = True
                run_ass.font.size = Pt(13)
                run_ass.font.name = "Segoe Script"
            except Exception:
                pass
        else:
            p_img.add_run("\n")

        tabela.cell(1, col).text = "________________________________"
        tabela.cell(2, col).text = str(ass.get("titulo") or "")
        for row in range(1, 3):
            p = tabela.cell(row, col).paragraphs[0]
            try:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9.5 if row == 2 else 8.5)
                    if row == 2:
                        run.bold = True
            except Exception:
                pass

    docx_add_paragraph(
        doc,
        f"Assinaturas digitais vinculadas ao processo {dados.get('processo')} • Gerado em {dados.get('data_encerramento')}",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

def gerar_docx_dossie(dados: Dict[str, Any], caminho_docx: Path) -> None:
    if Document is None:
        raise RuntimeError("Dependência ausente: instale python-docx.")

    caminho_docx = Path(caminho_docx)
    caminho_docx.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configurar_docx(doc, dados)

    # Página 1 — Capa
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for caminho in [caminho_brasao_pf(), caminho_brasao_dicor()]:
        if caminho:
            try:
                p.add_run().add_picture(str(caminho), width=Inches(1.45))
                p.add_run("   ")
            except Exception:
                pass
    docx_add_paragraph(doc, dados.get("cabecalho_institucional") or DOSSIE_INSTITUICAO_CABECALHO, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    docx_add_paragraph(doc, "POLÍCIA FEDERAL", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    docx_add_paragraph(doc, "CAPITAL MORADA DO VALLEY • DIRETORIA DE INVESTIGAÇÃO E COMBATE AO CRIME ORGANIZADO - DICOR", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    docx_add_paragraph(doc, "DOCUMENTO RESERVADO • INTELIGÊNCIA • INVESTIGAÇÃO • RESULTADO", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    docx_add_info_table(doc, [
        ["Campo", "Informação"],
        ["Processo Nº", dados.get("processo")],
        ["Investigação Nº", dados.get("numero_investigacao")],
        ["Nome da Operação", dados.get("nome_operacao")],
        ["Comunidade Investigada", dados.get("comunidade")],
        ["Cidade Operacional", dados.get("cidade_operacional", DOSSIE_CIDADE_OPERACIONAL)],
        ["Data de Abertura", dados.get("data_abertura")],
        ["Data de Encerramento", dados.get("data_encerramento")],
        ["Delegado Responsável", dados.get("delegado_responsavel")],
        ["Integrantes da Investigação", ", ".join(dados.get("integrantes_investigacao", []))],
    ])
    docx_add_paragraph(doc, "Documento gerado automaticamente a partir do encerramento da mesa de investigação.", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # Página 2 — Resumo
    docx_add_heading(doc, "2. RESUMO EXECUTIVO", 1)
    docx_add_paragraph(doc, dados.get("objetivo"))
    docx_add_paragraph(doc, resumo_contexto_operacional(dados))
    docx_add_info_table(doc, [
        ["Campo", "Informação"],
        ["Status da Investigação", dados.get("status")],
        ["Prioridade", dados.get("prioridade")],
        ["Cidade Operacional", dados.get("cidade_operacional", DOSSIE_CIDADE_OPERACIONAL)],
        ["Facção Investigada", dados.get("faccao")],
        ["Agente Responsável pelo Encerramento", dados.get("agente_encerramento")],
        ["Integrantes da Investigação", ", ".join(dados.get("integrantes_investigacao", []))],
    ])
    docx_add_heading(doc, "Tabela de Conteúdo", 2)
    docx_add_info_table(doc, linhas_tabela_conteudo())
    doc.add_page_break()

    docx_add_heading(doc, "3. LIDERANÇAS IDENTIFICADAS", 1)
    docx_add_pessoas(doc, dados.get("liderancas", []), "Nenhuma liderança foi identificada automaticamente nos tópicos da mesa.")
    doc.add_page_break()

    docx_add_heading(doc, "4. INTEGRANTES IDENTIFICADOS", 1)
    docx_add_pessoas(doc, dados.get("integrantes", []), "Nenhum integrante foi identificado automaticamente nos tópicos da mesa.")
    doc.add_page_break()

    docx_add_heading(doc, "5. PAINEL DA ORGANIZAÇÃO", 1)
    docx_add_paragraph(doc, dados.get("resumos", {}).get("painel"))
    docx_add_imagens_evidencias(doc, filtrar_evidencias_por_topico(dados.get("evidencias", []), ["painel"]), 8)
    doc.add_page_break()

    docx_add_heading(doc, "6. LOCALIZAÇÃO", 1)
    docx_add_paragraph(doc, dados.get("resumos", {}).get("localizacao"))
    docx_add_paragraph(doc, f"Comunidade/Base: {dados.get('comunidade')}\nCanal de reabertura: {dados.get('reabrir_url')}")
    docx_add_imagens_evidencias(doc, filtrar_evidencias_por_topico(dados.get("evidencias", []), ["localizacao"]), 8)
    doc.add_page_break()

    docx_add_heading(doc, "7. PRODUÇÃO E FABRICAÇÃO", 1)
    docx_add_paragraph(doc, dados.get("resumos", {}).get("producao"))
    docx_add_imagens_evidencias(doc, filtrar_evidencias_por_topico(dados.get("evidencias", []), ["producao"]), 8)
    doc.add_page_break()

    docx_add_heading(doc, "8. BAÚS E ARMAZENAMENTO", 1)
    docx_add_paragraph(doc, dados.get("resumos", {}).get("baus"))
    docx_add_imagens_evidencias(doc, filtrar_evidencias_por_topico(dados.get("evidencias", []), ["baus"]), 8)
    doc.add_page_break()

    docx_add_heading(doc, "9. INFORMANTES", 1)
    docx_add_pessoas(doc, dados.get("informantes", []), "Nenhum informante foi identificado automaticamente nos tópicos da mesa.")
    docx_add_paragraph(doc, dados.get("resumos", {}).get("informantes"))
    doc.add_page_break()

    res = dados.get("resultados", {}) or {}
    docx_add_heading(doc, "10. MATERIAIS APREENDIDOS", 1)
    docx_add_info_table(doc, [
        ["Categoria", "Registro"],
        ["Drogas", res.get("drogas", "Não informado")],
        ["Armas", res.get("armas", "Não informado")],
        ["Munições", res.get("municoes", "Não informado")],
        ["Dinheiro", res.get("dinheiro", "Não informado")],
        ["Veículos", res.get("veiculos", "Não informado")],
        ["Outros Itens", res.get("outros", "Não informado")],
    ])
    anexos = [["Tipo", "Arquivo", "Data", "Origem"]]
    for ev in dados.get("evidencias", [])[:40]:
        anexos.append([ev.get("tipo"), ev.get("arquivo"), ev.get("data"), ev.get("origem")])
    docx_add_heading(doc, "Provas e anexos registrados", 2)
    docx_add_info_table(doc, anexos)
    doc.add_page_break()

    est = dados.get("estatisticas", {}) or {}
    docx_add_heading(doc, "11. RESULTADO OPERACIONAL", 1)
    docx_add_info_table(doc, [
        ["Indicador", "Quantidade/Status"],
        ["Quantidade de presos", res.get("prisoes", "Não informado")],
        ["Procurados capturados", res.get("procurados_capturados", "Não informado")],
        ["Mensagens analisadas", est.get("mensagens_analisadas", 0)],
        ["Evidências coletadas", est.get("evidencias", 0)],
        ["Imagens anexadas", est.get("imagens", 0)],
        ["Vídeos anexados", est.get("videos", 0)],
        ["Links registrados", est.get("links", 0)],
        ["Tópicos analisados", est.get("threads", 0)],
    ])
    doc.add_page_break()

    docx_add_heading(doc, "12. CONCLUSÃO", 1)
    docx_add_paragraph(doc, montar_conclusao_dossie(dados))
    doc.add_paragraph("\n")
    docx_add_assinaturas_dossie(doc, dados)
    if dados.get("qr_reabertura") and Path(dados.get("qr_reabertura")).exists():
        docx_add_heading(doc, "QR Code para reabrir a investigação arquivada", 2)
        pqr = doc.add_paragraph()
        pqr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        docx_add_picture_safe(pqr, dados.get("qr_reabertura"), 1.35)

    doc.save(str(caminho_docx))



# =====================================================
# AUDITORIA GERAL DE AÇÕES DO BOT
# =====================================================


def _nome_interacao_log(interaction: discord.Interaction) -> str:
    try:
        if interaction.type == discord.InteractionType.application_command:
            return f"/{getattr(interaction.command, 'qualified_name', None) or interaction.data.get('name', 'comando')}"
        if interaction.type == discord.InteractionType.component:
            custom_id = (interaction.data or {}).get("custom_id", "componente")
            return f"Botão/Menu `{custom_id}`"
        if interaction.type == discord.InteractionType.modal_submit:
            custom_id = (interaction.data or {}).get("custom_id", "modal")
            return f"Modal `{custom_id}`"
    except Exception:
        pass
    return str(getattr(interaction, "type", "interação"))


@bot.listen("on_interaction")
async def auditoria_todas_interacoes(interaction: discord.Interaction):
    """Registra todo comando, botão, menu e modal usado no bot."""
    try:
        canal_desc = getattr(interaction.channel, "mention", None) or f"ID {getattr(interaction, 'channel_id', 'desconhecido')}"
        guild_desc = interaction.guild.name if interaction.guild else "DM/sem servidor"
        acao = _nome_interacao_log(interaction)
        await enviar_log(
            "🧾 **Ação executada no bot**\n"
            f"Ação: {acao}\n"
            f"Usuário: {interaction.user.mention} (`{interaction.user.id}`)\n"
            f"Canal: {canal_desc}\n"
            f"Servidor: {guild_desc}\n"
            f"Data: {agora_br()}"
        )
    except Exception:
        pass


@bot.listen("on_guild_channel_create")
async def auditoria_canal_criado(canal):
    """Registra criação de canal quando for possível identificar como ação do bot."""
    try:
        await asyncio.sleep(1)
        autor = None
        try:
            async for entry in canal.guild.audit_logs(limit=3, action=discord.AuditLogAction.channel_create):
                if entry.target and getattr(entry.target, "id", None) == canal.id:
                    autor = entry.user
                    break
        except Exception:
            autor = None

        if autor and bot.user and autor.id != bot.user.id:
            return

        await enviar_log(
            "📁 **Canal criado pelo bot**\n"
            f"Canal: {getattr(canal, 'mention', canal.name)} (`{canal.id}`)\n"
            f"Categoria: {getattr(getattr(canal, 'category', None), 'name', 'Sem categoria')}\n"
            f"Data: {agora_br()}"
        )
    except Exception:
        pass


@bot.listen("on_guild_channel_delete")
async def auditoria_canal_apagado(canal):
    """Registra exclusão de canal quando for possível identificar como ação do bot."""
    try:
        await asyncio.sleep(1)
        autor = None
        try:
            async for entry in canal.guild.audit_logs(limit=3, action=discord.AuditLogAction.channel_delete):
                if entry.target and getattr(entry.target, "id", None) == canal.id:
                    autor = entry.user
                    break
        except Exception:
            autor = None

        if autor and bot.user and autor.id != bot.user.id:
            return

        await enviar_log(
            "🗑️ **Canal apagado pelo bot**\n"
            f"Canal: #{canal.name} (`{canal.id}`)\n"
            f"Categoria: {getattr(getattr(canal, 'category', None), 'name', 'Sem categoria')}\n"
            f"Data: {agora_br()}"
        )
    except Exception:
        pass


@bot.listen("on_message")
async def auditoria_mensagem_do_bot(message: discord.Message):
    """Registra mensagens enviadas pelo bot, sem espelhar o próprio canal de logs para evitar loop."""
    try:
        if not bot.user or message.author.id != bot.user.id:
            return
        if message.channel and getattr(message.channel, "id", 0) == LOGS_CHANNEL_ID:
            return

        conteudo = (message.content or "").strip().replace("\n", " ")
        if len(conteudo) > 450:
            conteudo = conteudo[:450].rstrip() + "..."
        if not conteudo:
            conteudo = f"Mensagem com {len(message.attachments)} anexo(s)."

        await enviar_log(
            "📨 **Mensagem enviada pelo bot**\n"
            f"Canal: {getattr(message.channel, 'mention', 'Sem canal')} (`{getattr(message.channel, 'id', '0')}`)\n"
            f"Conteúdo: {conteudo}\n"
            f"Anexos: {len(message.attachments)}\n"
            f"Link: {getattr(message, 'jump_url', 'Sem link')}\n"
            f"Data: {agora_br()}"
        )
    except Exception:
        pass


# =====================================================
# ATENDIMENTO AUTOMÁTICO DE BOLETINS - DICOR
# =====================================================

boletins_processando: set[int] = set()

def carregar_atendimentos_boletins() -> List[Dict[str, Any]]:
    dados = carregar_json(BOLETIM_ATENDIMENTOS_JSON, [])
    return dados if isinstance(dados, list) else []

def salvar_atendimentos_boletins(lista: List[Dict[str, Any]]) -> None:
    salvar_json(BOLETIM_ATENDIMENTOS_JSON, lista[-5000:])

def buscar_atendimento_por_mensagem(message_id: int) -> Optional[Dict[str, Any]]:
    for item in carregar_atendimentos_boletins():
        if str(item.get("mensagem_original_id")) == str(message_id):
            return item
    return None


def buscar_atendimento_por_numero(numero: str) -> Optional[Dict[str, Any]]:
    alvo = str(numero or "").strip().upper()
    if not alvo:
        return None
    for item in carregar_atendimentos_boletins():
        atual = str(item.get("numero") or "").strip().upper()
        if atual == alvo:
            return item
        try:
            if numero_curto_boletim(atual) == numero_curto_boletim(alvo):
                return item
        except Exception:
            pass
    return None

def buscar_atendimento_por_area(channel_id: int) -> Optional[Dict[str, Any]]:
    for item in carregar_atendimentos_boletins():
        ids = {str(item.get("area_id")), str(item.get("thread_id")), str(item.get("forum_thread_id"))}
        if str(channel_id) in ids:
            return item
    return None

def atualizar_atendimento_boletim(chave: str, valor: Any, atualizacoes: Dict[str, Any]) -> Dict[str, Any]:
    lista = carregar_atendimentos_boletins()
    alvo = None
    for item in lista:
        if str(item.get(chave)) == str(valor):
            item.update(atualizacoes)
            alvo = item
            break
    if alvo is None:
        alvo = {chave: valor, **atualizacoes}
        lista.append(alvo)
    salvar_atendimentos_boletins(lista)
    return alvo

def numero_boletim_de_texto(texto: str) -> str:
    texto = str(texto or "")
    for p in [r"BO[- ]?DICOR[- ]?(\d{1,6})", r"BOLETIM\s+DE\s+OCORR[ÊE]NCIA\s*[—\-–]?\s*(?:N[º°O]\.?|Nº|N)?\s*(\d{1,6})", r"N[º°O]\.?\s*(\d{1,6})"]:
        m = re.search(p, texto, flags=re.I)
        if m:
            return f"BO-DICOR-{int(m.group(1)):03d}"
    return gerar_numero_boletim()

def eh_boletim_valido_para_atendimento(message: discord.Message) -> bool:
    if not message or not getattr(message, "channel", None):
        return False
    if getattr(message.channel, "id", 0) != BOLETINS_CHANNEL_ID:
        return False
    if getattr(message, "type", None) != discord.MessageType.default:
        return False
    texto = coletar_texto_embed(message) if 'coletar_texto_embed' in globals() else (message.content or "")
    if not texto.strip() and not message.attachments:
        return False

    n = normalizar_busca(texto)
    # Não processar mensagens soltas de provas/anexos como se fossem outro boletim.
    # Ex.: "Provas — BO-DICOR-002" ou "Anexos do boletim" não abre novo atendimento.
    termos_provas = ["provas", "anexos do boletim", "anexo do boletim", "lote final"]
    if any(t in n for t in termos_provas) and "boletim de ocorrencia" not in n:
        return False

    # A mensagem válida precisa ter o título/modelo de boletim.
    # Apenas aparecer BO-DICOR-000 não é suficiente, porque anexos também usam esse número.
    return "boletim de ocorrencia" in n

def carregar_rodizio_boletim() -> Dict[str, Any]:
    dados = carregar_json(BOLETIM_RODIZIO_JSON, {})
    return dados if isinstance(dados, dict) else {}

def salvar_rodizio_boletim(dados: Dict[str, Any]) -> None:
    salvar_json(BOLETIM_RODIZIO_JSON, dados)

def membros_elegiveis_rodizio(guild: discord.Guild) -> List[discord.Member]:
    if guild is None:
        return []
    cargos_alvo = set(BOLETIM_RODIZIO_CARGOS_IDS) & set(BOLETIM_RODIZIO_CARGOS_FIXOS_IDS)
    cargos_excluir = set(BOLETIM_RODIZIO_EXCLUIR_CARGOS_IDS)
    membros = {}
    for member in guild.members:
        if member.bot:
            continue
        roles_ids = {role.id for role in member.roles}
        # Precisa ter Estagiário ou Investigador.
        if not roles_ids.intersection(cargos_alvo):
            continue
        # Diretoria/Inspetor/Delegado não entra no rodízio de atendimento comum.
        if roles_ids.intersection(cargos_excluir):
            continue
        membros[member.id] = member
    return sorted(membros.values(), key=lambda m: (m.display_name.lower(), m.id))

async def escolher_agente_rodizio(guild: discord.Guild, numero: str) -> Optional[discord.Member]:
    elegiveis = membros_elegiveis_rodizio(guild)
    if not elegiveis:
        await enviar_log(f"⚠️ Nenhum agente elegível para rodízio do boletim `{numero}`.")
        return None
    ids_atuais = [m.id for m in elegiveis]
    estado = carregar_rodizio_boletim()
    usados = [int(x) for x in estado.get("ciclo_usados", []) if int(x) in ids_atuais]
    ultimo_id = int(estado.get("ultimo_id", 0) or 0)
    disponiveis = [m for m in elegiveis if m.id not in usados]
    if not disponiveis:
        usados = []
        disponiveis = elegiveis[:]
    if len(disponiveis) > 1:
        disponiveis = [m for m in disponiveis if m.id != ultimo_id] or disponiveis
    escolhido = disponiveis[0]
    usados.append(escolhido.id)
    historico = estado.get("historico", []) if isinstance(estado.get("historico", []), list) else []
    historico.append({"numero": numero, "agente_id": escolhido.id, "agente_nome": str(escolhido), "data": agora_br()})
    estado.update({"ultimo_id": escolhido.id, "ciclo_usados": usados, "ultima_atribuicao": agora_br(), "historico": historico[-1000:]})
    salvar_rodizio_boletim(estado)
    await enviar_log(f"🔁 **Rodízio de boletim**\nBoletim: `{numero}`\nAgente escolhido: {escolhido.mention} (`{escolhido.id}`)")
    return escolhido

async def baixar_anexo_persistente(anexo: discord.Attachment, pasta: Path, prefixo: str) -> Optional[Path]:
    try:
        pasta.mkdir(parents=True, exist_ok=True)
        caminho = pasta / f"{data_caso()}-{slugify(prefixo)}-{secrets.token_hex(4)}-{nome_arquivo_seguro(anexo.filename)}"
        async with ClientSession(timeout=ClientTimeout(total=45)) as session:
            async with session.get(anexo.url) as resp:
                if resp.status != 200:
                    await enviar_log(f"⚠️ Erro ao baixar anexo `{anexo.filename}`: HTTP {resp.status}")
                    return None
                with open(caminho, "wb") as f:
                    async for chunk in resp.content.iter_chunked(1024 * 128):
                        f.write(chunk)
        return caminho
    except Exception as erro:
        await enviar_log(f"⚠️ Falha ao baixar anexo `{getattr(anexo, 'filename', 'arquivo')}`: {erro}")
        return None

async def arquivos_para_reenvio_de_mensagens(mensagens: List[discord.Message], pasta: Path, prefixo: str) -> List[Path]:
    caminhos, vistos = [], set()
    for msg in mensagens:
        for anexo in msg.attachments:
            chave = (anexo.id, anexo.filename, anexo.size)
            if chave in vistos:
                continue
            vistos.add(chave)
            caminho = await baixar_anexo_persistente(anexo, pasta, prefixo)
            if caminho:
                caminhos.append(caminho)
    return caminhos

async def enviar_arquivos_em_lotes(destino, caminhos: List[Path], legenda: str = "Arquivos") -> int:
    enviados, lote, tamanho_lote = 0, [], 0
    limite = 24 * 1024 * 1024
    for caminho in caminhos:
        try:
            tamanho = caminho.stat().st_size
            if tamanho > limite:
                await destino.send(f"⚠️ `{caminho.name}` ultrapassa o limite de envio. Arquivo preservado no backup interno.")
                continue
            if len(lote) >= 10 or tamanho_lote + tamanho > limite:
                await destino.send(content=f"{legenda} — lote", files=lote)
                enviados += len(lote)
                lote, tamanho_lote = [], 0
            lote.append(discord.File(str(caminho), filename=caminho.name))
            tamanho_lote += tamanho
        except Exception as erro:
            await enviar_log(f"⚠️ Falha preparando arquivo `{caminho}`: {erro}")
    if lote:
        await destino.send(content=f"{legenda} — lote final", files=lote)
        enviados += len(lote)
    return enviados


async def enviar_texto_com_anexos_path(destino, conteudo: str, caminhos: List[Path], legenda_restante: str) -> tuple[Optional[discord.Message], int]:
    """Envia texto com o primeiro lote de anexos junto e o restante em lotes."""
    limite = 24 * 1024 * 1024
    selecionados: List[Path] = []
    restantes: List[Path] = []
    tamanho = 0

    for caminho in caminhos:
        try:
            tam = caminho.stat().st_size
            if len(selecionados) < 10 and tam <= limite and tamanho + tam <= limite:
                selecionados.append(caminho)
                tamanho += tam
            else:
                restantes.append(caminho)
        except Exception:
            restantes.append(caminho)

    mensagem = None
    try:
        if selecionados:
            arquivos = [discord.File(str(p), filename=p.name) for p in selecionados]
            mensagem = await destino.send(content=conteudo[:1900], files=arquivos)
        else:
            mensagem = await destino.send(content=conteudo[:1900])
    except discord.HTTPException as erro:
        await enviar_log(f"⚠️ Não consegui enviar texto com anexos juntos. Enviando separado: {erro}")
        mensagem = await destino.send(content=conteudo[:1900])
        restantes = caminhos
        selecionados = []

    enviados = len(selecionados)
    if restantes:
        enviados += await enviar_arquivos_em_lotes(destino, restantes, legenda_restante)
    return mensagem, enviados

async def criar_area_atendimento_boletim(message: discord.Message) -> Optional[Dict[str, Any]]:
    if buscar_atendimento_por_mensagem(message.id):
        return None

    numero = numero_boletim_de_texto(coletar_texto_embed(message) or message.content or "")

    # Impede dois atendimentos/tópicos para o mesmo número de boletim.
    existente = buscar_atendimento_por_numero(numero)
    if existente:
        await enviar_log(
            f"⚠️ Atendimento duplicado ignorado.\n"
            f"Boletim: `{numero}`\n"
            f"Mensagem ignorada: {message.jump_url}\n"
            f"Área existente: <#{existente.get('area_id') or existente.get('thread_id')}>"
        )
        return None

    agente = await escolher_agente_rodizio(message.guild, numero)
    canal_atendimento = message.guild.get_channel(BOLETIM_ATENDIMENTO_CHANNEL_ID) if message.guild else None
    if canal_atendimento is None:
        await enviar_log(f"❌ Canal de atendimento de boletins `{BOLETIM_ATENDIMENTO_CHANNEL_ID}` não encontrado.")
        return None

    texto_original = coletar_texto_embed(message) or message.content or "Sem texto."
    pasta = BOLETIM_ARQUIVOS_DIR / numero.replace("/", "-")
    anexos = await arquivos_para_reenvio_de_mensagens([message], pasta, numero)
    titulo = f"📋 BOLETIM DE OCORRÊNCIA — Nº {numero_curto_boletim(numero)}"

    # A menção do agente NÃO vai aqui. Ela será enviada por último dentro do atendimento,
    # para ficar mais organizado e evitar parecer que existem dois tópicos/mensagens principais.
    conteudo_abertura = (
        "📋 **NOVO BOLETIM RECEBIDO PARA ANÁLISE**\n"
        "━━━━━━━━━━━━━━━━━━━━━━━\n\n"
        f"**Boletim:** Nº `{numero_curto_boletim(numero)}`\n"
        f"**Status:** 🟢 Em atendimento\n"
        f"**Responsável pelo atendimento:** {agente.mention if agente else 'Não definido'}\n"
        f"**Autor do boletim:** {message.author.mention if message.author else 'Não identificado'}\n"
        f"**Data de recebimento:** {agora_br()}\n"
        f"**Mensagem original:** {message.jump_url}\n"
        f"**Anexos localizados:** `{len(message.attachments)}`\n\n"
        "Toda a análise, provas e decisões deste boletim deverão permanecer neste tópico."
    )

    try:
        if isinstance(canal_atendimento, discord.ForumChannel):
            criado = await canal_atendimento.create_thread(
                name=titulo[:100],
                content="📋 Atendimento criado. As informações do boletim serão organizadas abaixo.",
                allowed_mentions=discord.AllowedMentions.none(),
                reason="Atendimento automático de boletim DICOR",
            )
            area = getattr(criado, "thread", criado)
        elif isinstance(canal_atendimento, discord.TextChannel):
            msg_principal = await canal_atendimento.send(
                f"📋 Atendimento criado para **Boletim Nº {numero_curto_boletim(numero)}**. Use o tópico vinculado abaixo.",
                allowed_mentions=discord.AllowedMentions.none(),
            )
            area = await msg_principal.create_thread(name=titulo[:100], auto_archive_duration=10080)
        else:
            await enviar_log(f"❌ Canal de atendimento `{BOLETIM_ATENDIMENTO_CHANNEL_ID}` não é texto/fórum compatível.")
            return None

        # Tudo fica dentro do mesmo tópico/thread/postagem: dados, texto, anexos, painel e, por último, a menção.
        await area.send(conteudo_abertura, allowed_mentions=discord.AllowedMentions(users=False, roles=False, everyone=False))

        texto_boletim_completo = (
            "📄 **BOLETIM ORIGINAL COMPLETO**\n"
            "━━━━━━━━━━━━━━━━━━━━━━━\n"
            f"{cortar_discord(texto_original, 1650)}\n\n"
            f"📎 **Anexos do boletim:** `{len(anexos)}` arquivo(s) reenviado(s) neste atendimento."
        )
        if anexos:
            await enviar_texto_com_anexos_path(area, texto_boletim_completo, anexos, f"📎 Anexos restantes do boletim {numero}")
        else:
            await area.send(texto_boletim_completo)

        painel_msg = await area.send(
            "🛠️ **Mini painel de gerenciamento**\n"
            "Use os botões abaixo para finalizar, solicitar comparecimento ou iniciar cadastro de procurado.",
            view=BoletimAtendimentoView(),
        )

        # Última mensagem: somente aqui o agente é mencionado.
        if agente:
            await area.send(
                f"{agente.mention}, este boletim foi encaminhado para sua análise.",
                allowed_mentions=discord.AllowedMentions(users=True, roles=False, everyone=False),
            )

        atendimento = {
            "id": f"ATD-{message.id}",
            "numero": numero,
            "mensagem_original_id": message.id,
            "mensagem_original_url": message.jump_url,
            "canal_origem_id": message.channel.id,
            "area_id": getattr(area, "id", None),
            "thread_id": getattr(area, "id", None),
            "painel_msg_id": painel_msg.id,
            "agente_id": agente.id if agente else None,
            "agente_nome": str(agente) if agente else "Não definido",
            "autor_id": message.author.id if message.author else None,
            "autor_nome": str(message.author) if message.author else "Não identificado",
            "status": "EM ATENDIMENTO",
            "data_criacao": agora_br(),
            "anexos_salvos": [str(p) for p in anexos],
            "historico": [{"acao": "Atendimento criado", "usuario": "Sistema", "data": agora_br()}],
            "comparecimento_status": "não solicitado",
            "procurado_status": "não solicitado",
        }
        lista = carregar_atendimentos_boletins()
        lista.append(atendimento)
        salvar_atendimentos_boletins(lista)
        await enviar_log(
            f"📋 **Atendimento de boletim criado**\n"
            f"Boletim: `{numero}`\n"
            f"Área: <#{area.id}>\n"
            f"Agente: {agente.mention if agente else 'Não definido'}\n"
            f"Original: {message.jump_url}\n"
            f"Anexos salvos: `{len(anexos)}`"
        )
        return atendimento
    except Exception as erro:
        await enviar_log(f"❌ Erro ao criar atendimento do boletim `{numero}`: {erro}")
        return None

class FinalizarBoletimAtendimentoModal(Modal, title="Finalizar Boletim"):
    resultado = TextInput(label="Resultado final do boletim", style=discord.TextStyle.paragraph, max_length=1600)
    async def on_submit(self, interaction: discord.Interaction):
        await finalizar_boletim_atendimento(interaction, str(self.resultado.value))

class ComparecimentoBoletimModal(Modal, title="Solicitar Comparecimento"):
    nome = TextInput(label="Nome do convocado", max_length=120)
    rg = TextInput(label="RG", max_length=50)
    motivo = TextInput(label="Motivo do comparecimento", style=discord.TextStyle.paragraph, max_length=900)
    data_hora = TextInput(label="Data e horário", placeholder="Ex: 12/07 às 20h", max_length=120)
    local = TextInput(label="Local", placeholder="Ex: Sede da Polícia Federal", max_length=160)
    async def on_submit(self, interaction: discord.Interaction):
        await interaction.response.defer(ephemeral=True, thinking=True)
        atendimento = buscar_atendimento_por_area(interaction.channel.id)
        if not atendimento:
            return await interaction.followup.send("❌ Atendimento não encontrado.", ephemeral=True)
        registro = {"nome": str(self.nome.value), "rg": str(self.rg.value), "motivo": str(self.motivo.value), "data_hora": str(self.data_hora.value), "local": str(self.local.value), "solicitado_por_id": interaction.user.id, "solicitado_por_nome": str(interaction.user), "data": agora_br()}
        historico = atendimento.get("historico", []) or []; historico.append({"acao": "Solicitação de comparecimento", "usuario": str(interaction.user), "data": agora_br(), "dados": registro})
        atualizar_atendimento_boletim("id", atendimento.get("id"), {"comparecimento_status": "solicitado", "comparecimento": registro, "historico": historico})
        await interaction.channel.send(f"📩 **SOLICITAÇÃO DE COMPARECIMENTO REGISTRADA**\n**Boletim:** Nº {numero_curto_boletim(atendimento.get('numero'))}\n**Convocado:** {registro['nome']}\n**RG:** {registro['rg']}\n**Motivo:** {registro['motivo']}\n**Data/Horário:** {registro['data_hora']}\n**Local:** {registro['local']}\n**Solicitado por:** {interaction.user.mention}\n\n⚠️ Documento oficial de comparecimento ainda será configurado.")
        await enviar_log(f"📩 Comparecimento solicitado no boletim `{atendimento.get('numero')}` por {interaction.user.mention} (`{interaction.user.id}`).")
        await interaction.followup.send("✅ Solicitação registrada.", ephemeral=True)

class ProcuradoBoletimModal(Modal, title="Cadastrar como Procurado"):
    nome = TextInput(label="Nome", max_length=120, required=True)
    rg = TextInput(label="RG", max_length=50, required=True)
    caracteristicas = TextInput(label="Características", style=discord.TextStyle.paragraph, max_length=600, required=False)
    crimes = TextInput(label="Crimes cometidos", style=discord.TextStyle.paragraph, max_length=1200, required=True)
    outras = TextInput(label="Outras informações", style=discord.TextStyle.paragraph, max_length=900, required=False)
    def __init__(self, dados: Optional[Dict[str, str]] = None):
        super().__init__(); dados = dados or {}
        for campo, chave in [(self.nome, "nome"), (self.rg, "rg"), (self.caracteristicas, "caracteristicas"), (self.crimes, "crimes"), (self.outras, "outras")]:
            try: campo.default = str(dados.get(chave, ""))[:campo.max_length]
            except Exception: pass
    async def on_submit(self, interaction: discord.Interaction):
        await solicitar_autorizacao_procurado_boletim(interaction, {"nome": str(self.nome.value), "rg": str(self.rg.value), "caracteristicas": str(self.caracteristicas.value or "Não informado"), "crimes": str(self.crimes.value), "outras": str(self.outras.value or "Não informado")})

def extrair_dados_procurado_de_texto(texto: str) -> Dict[str, str]:
    return {"nome": extrair_valor_por_rotulos([texto], ["Nome", "Indivíduo", "Individuo"], 120).replace("Não informado", ""), "rg": extrair_valor_por_rotulos([texto], ["RG", "Registro"], 50).replace("Não informado", ""), "caracteristicas": extrair_valor_por_rotulos([texto], ["Características", "Caracteristicas"], 600).replace("Não informado", ""), "crimes": extrair_valor_por_rotulos([texto], ["Crimes cometidos", "Crimes", "Crime"], 1200).replace("Não informado", ""), "outras": extrair_valor_por_rotulos([texto], ["Outras informações", "Outras informacoes", "Informações", "Informacoes"], 900).replace("Não informado", "")}

async def solicitar_autorizacao_procurado_boletim(interaction: discord.Interaction, dados: Dict[str, str]) -> None:
    await interaction.response.defer(ephemeral=True, thinking=True)
    atendimento = buscar_atendimento_por_area(interaction.channel.id)
    if not atendimento: return await interaction.followup.send("❌ Atendimento não encontrado.", ephemeral=True)
    if not limpar_rg(dados.get("rg")): return await interaction.followup.send("❌ RG obrigatório.", ephemeral=True)
    if procurar_por_rg(dados.get("rg")): return await interaction.followup.send("❌ Já existe procurado cadastrado com esse RG.", ephemeral=True)
    atendimento.update({"procurado_solicitado": dados, "procurado_status": "aguardando_autorizacao", "procurado_solicitado_por_id": interaction.user.id, "procurado_solicitado_por_nome": str(interaction.user)})
    hist = atendimento.get("historico", []) or []; hist.append({"acao": "Solicitação de cadastro como procurado", "usuario": str(interaction.user), "data": agora_br(), "rg": dados.get("rg")}); atendimento["historico"] = hist
    atualizar_atendimento_boletim("id", atendimento.get("id"), atendimento)
    embed = discord.Embed(title="🚨 Solicitação de cadastro como procurado", description=(f"**Boletim:** Nº {numero_curto_boletim(atendimento.get('numero'))}\n**Nome:** {dados.get('nome')}\n**RG:** {dados.get('rg')}\n**Características:** {dados.get('caracteristicas')}\n**Crimes:**\n{dados.get('crimes')}\n\n**Outras informações:** {dados.get('outras')}\n**Solicitado por:** {interaction.user.mention}\n**Agente do boletim:** <@{atendimento.get('agente_id')}>"), color=discord.Color.red())
    await interaction.channel.send(embed=embed, view=AutorizacaoProcuradoBoletimView())
    await enviar_log(f"🚨 Solicitação de procurado no boletim `{atendimento.get('numero')}` por {interaction.user.mention} (`{interaction.user.id}`).")
    await interaction.followup.send("✅ Solicitação enviada para autorização de Inspetor ou superior.", ephemeral=True)

class AutorizacaoProcuradoBoletimView(View):
    def __init__(self): super().__init__(timeout=None)
    @discord.ui.button(label="Autorizar Cadastro", emoji="✅", style=discord.ButtonStyle.green, custom_id="dic_bo_autorizar_procurado")
    async def autorizar(self, interaction: discord.Interaction, button: Button):
        await interaction.response.defer(ephemeral=True, thinking=True)
        if not isinstance(interaction.user, discord.Member) or not usuario_pode_fechar_mesa(interaction.user):
            await enviar_log(f"🚫 Tentativa sem permissão de autorizar procurado no boletim por {interaction.user.mention} (`{interaction.user.id}`).")
            return await interaction.followup.send("❌ Somente Inspetor ou superior pode autorizar.", ephemeral=True)
        atendimento = buscar_atendimento_por_area(interaction.channel.id)
        if not atendimento: return await interaction.followup.send("❌ Atendimento não encontrado.", ephemeral=True)
        if atendimento.get("procurado_status") != "aguardando_autorizacao": return await interaction.followup.send("⚠️ Solicitação já foi analisada.", ephemeral=True)
        atendimento.update({"procurado_status": "autorizado", "procurado_autorizado_por_id": interaction.user.id, "procurado_autorizado_por_nome": str(interaction.user), "procurado_autorizado_em": agora_br()})
        hist = atendimento.get("historico", []) or []; hist.append({"acao": "Cadastro como procurado autorizado", "usuario": str(interaction.user), "data": agora_br()}); atendimento["historico"] = hist
        atualizar_atendimento_boletim("id", atendimento.get("id"), atendimento)
        for child in self.children: child.disabled = True
        try: await interaction.message.edit(view=self)
        except Exception: pass
        await interaction.channel.send("✅ **Cadastro autorizado.**\nEnvie agora a **foto do indivíduo** e a **foto do RG**. Depois clique em **📸 Confirmar Fotos e Publicar**.", view=FotoProcuradoBoletimView())
        await enviar_log(f"✅ Procurado autorizado no boletim `{atendimento.get('numero')}` por {interaction.user.mention} (`{interaction.user.id}`).")
        await interaction.followup.send("✅ Autorizado. Aguardando fotos obrigatórias.", ephemeral=True)
    @discord.ui.button(label="Recusar Cadastro", emoji="❌", style=discord.ButtonStyle.red, custom_id="dic_bo_recusar_procurado")
    async def recusar(self, interaction: discord.Interaction, button: Button):
        await interaction.response.defer(ephemeral=True, thinking=True)
        if not isinstance(interaction.user, discord.Member) or not usuario_pode_fechar_mesa(interaction.user):
            return await interaction.followup.send("❌ Somente Inspetor ou superior pode recusar.", ephemeral=True)
        atendimento = buscar_atendimento_por_area(interaction.channel.id)
        if not atendimento: return await interaction.followup.send("❌ Atendimento não encontrado.", ephemeral=True)
        if atendimento.get("procurado_status") != "aguardando_autorizacao": return await interaction.followup.send("⚠️ Solicitação já foi analisada.", ephemeral=True)
        atendimento.update({"procurado_status": "recusado", "procurado_recusado_por_id": interaction.user.id, "procurado_recusado_por_nome": str(interaction.user), "procurado_recusado_em": agora_br()})
        atualizar_atendimento_boletim("id", atendimento.get("id"), atendimento)
        for child in self.children: child.disabled = True
        try: await interaction.message.edit(view=self)
        except Exception: pass
        await interaction.channel.send(f"❌ **Cadastro como procurado recusado por {interaction.user.mention}.**")
        await enviar_log(f"❌ Procurado recusado no boletim `{atendimento.get('numero')}` por {interaction.user.mention} (`{interaction.user.id}`).")
        await interaction.followup.send("❌ Solicitação recusada.", ephemeral=True)

class FotoProcuradoBoletimView(View):
    def __init__(self): super().__init__(timeout=None)
    @discord.ui.button(label="Confirmar Fotos e Publicar", emoji="📸", style=discord.ButtonStyle.green, custom_id="dic_bo_confirmar_fotos_procurado")
    async def confirmar(self, interaction: discord.Interaction, button: Button):
        await publicar_procurado_autorizado_boletim(interaction)

async def publicar_procurado_autorizado_boletim(interaction: discord.Interaction) -> None:
    await interaction.response.defer(ephemeral=True, thinking=True)
    atendimento = buscar_atendimento_por_area(interaction.channel.id)
    if not atendimento: return await interaction.followup.send("❌ Atendimento não encontrado.", ephemeral=True)
    if atendimento.get("procurado_status") != "autorizado": return await interaction.followup.send("❌ O cadastro ainda não foi autorizado ou já foi processado.", ephemeral=True)
    dados = atendimento.get("procurado_solicitado") or {}
    if procurar_por_rg(dados.get("rg")):
        atendimento["procurado_status"] = "duplicado"; atualizar_atendimento_boletim("id", atendimento.get("id"), atendimento)
        return await interaction.followup.send("❌ Já existe procurado com esse RG.", ephemeral=True)
    imagens = []
    async for msg in interaction.channel.history(limit=120, oldest_first=True):
        for anexo in msg.attachments:
            if tipo_anexo_dossie(anexo) == "imagem": imagens.append(anexo)
    if len(imagens) < 2: return await interaction.followup.send("❌ Envie a foto do indivíduo e a foto do RG antes de confirmar.", ephemeral=True)
    foto_ind = await salvar_anexo_publico(imagens[-2], f"boletim-individuo-{dados.get('rg')}"); foto_rg = await salvar_anexo_publico(imagens[-1], f"boletim-rg-{dados.get('rg')}")
    registro = {"id": data_caso(), "caso": f"DICOR-{data_caso()}", "data": agora_br(), "status": "A PROCURAR", "nome": dados.get("nome", "Não informado"), "rg": dados.get("rg", "Não informado"), "caracteristicas": dados.get("caracteristicas", "Não informado"), "crimes": dados.get("crimes", "Não informado"), "numero_boletim": atendimento.get("numero"), "boletim": atendimento.get("numero"), "ultimo_avistamento": dados.get("outras") or "Vinculado ao boletim de ocorrência.", "informacoes": dados.get("outras", "Não informado"), "foto_individuo": foto_ind, "foto_rg": foto_rg, "autor_id": atendimento.get("procurado_solicitado_por_id"), "autor_nome": atendimento.get("procurado_solicitado_por_nome"), "autorizado_por_id": atendimento.get("procurado_autorizado_por_id"), "autorizado_por_nome": atendimento.get("procurado_autorizado_por_nome"), "agente_boletim_id": atendimento.get("agente_id"), "agente_boletim_nome": atendimento.get("agente_nome"), "mensagem_id": None, "mensagem_url": None}
    msg = await postar_procurado_oficial(registro)
    if msg: registro["mensagem_id"] = msg.id; registro["mensagem_url"] = msg.jump_url
    await salvar_procurado_catalogo(registro)
    atendimento.update({"procurado_status": "publicado", "procurado_publicacao_id": registro.get("mensagem_id"), "procurado_publicacao_url": registro.get("mensagem_url")})
    atualizar_atendimento_boletim("id", atendimento.get("id"), atendimento)
    await interaction.channel.send(f"✅ **Procurado cadastrado com sucesso.**\nRG: `{dados.get('rg')}`\nPublicação: {registro.get('mensagem_url')}")
    await enviar_log(f"🚨 Procurado cadastrado automaticamente via boletim `{atendimento.get('numero')}` | RG `{dados.get('rg')}` | Por {interaction.user.mention} (`{interaction.user.id}`).")
    await interaction.followup.send("✅ Procurado publicado e catálogo atualizado.", ephemeral=True)

class BoletimAtendimentoView(View):
    def __init__(self): super().__init__(timeout=None)
    @discord.ui.button(label="Finalizar Boletim", emoji="✅", style=discord.ButtonStyle.green, custom_id="dic_bo_finalizar")
    async def finalizar(self, interaction: discord.Interaction, button: Button):
        atendimento = buscar_atendimento_por_area(interaction.channel.id)
        if not atendimento: return await interaction.response.send_message("❌ Atendimento não encontrado.", ephemeral=True)
        if atendimento.get("status") == "FINALIZADO": return await interaction.response.send_message("⚠️ Este boletim já foi finalizado.", ephemeral=True)
        await interaction.response.send_modal(FinalizarBoletimAtendimentoModal())
    @discord.ui.button(label="Solicitar Comparecimento", emoji="📩", style=discord.ButtonStyle.secondary, custom_id="dic_bo_comparecimento")
    async def comparecimento(self, interaction: discord.Interaction, button: Button):
        await interaction.response.send_modal(ComparecimentoBoletimModal())
    @discord.ui.button(label="Cadastrar como Procurado", emoji="🚨", style=discord.ButtonStyle.danger, custom_id="dic_bo_cadastrar_procurado")
    async def cadastrar_procurado(self, interaction: discord.Interaction, button: Button):
        atendimento = buscar_atendimento_por_area(interaction.channel.id); texto = ""
        if atendimento:
            try:
                canal_origem = interaction.guild.get_channel(int(atendimento.get("canal_origem_id") or 0))
                if canal_origem:
                    msg = await canal_origem.fetch_message(int(atendimento.get("mensagem_original_id")))
                    texto = coletar_texto_embed(msg)
            except Exception: texto = ""
        await interaction.response.send_modal(ProcuradoBoletimModal(extrair_dados_procurado_de_texto(texto)))

async def finalizar_boletim_atendimento(interaction: discord.Interaction, resultado: str) -> None:
    if not interaction.response.is_done(): await interaction.response.defer(ephemeral=True, thinking=True)
    atendimento = buscar_atendimento_por_area(interaction.channel.id)
    if not atendimento: return await interaction.followup.send("❌ Atendimento não encontrado.", ephemeral=True)
    if atendimento.get("status") == "FINALIZADO": return await interaction.followup.send("⚠️ Este boletim já foi finalizado.", ephemeral=True)
    if not isinstance(interaction.user, discord.Member) or not usuario_tem_equipe(interaction.user):
        await enviar_log(f"🚫 Tentativa sem permissão de finalizar boletim `{atendimento.get('numero')}` por {interaction.user.mention} (`{interaction.user.id}`).")
        return await interaction.followup.send("❌ Você não possui permissão para finalizar este boletim.", ephemeral=True)
    await interaction.channel.send(f"✅ **BOLETIM FINALIZADO**\n\n**Boletim:** Nº {numero_curto_boletim(atendimento.get('numero'))}\n**Agente responsável:** <@{atendimento.get('agente_id')}>\n**Finalizado por:** {interaction.user.mention}\n**Data e horário:** {agora_br()}\n**Resultado final:**\n{resultado}")
    mensagens = []
    try:
        async for msg in interaction.channel.history(limit=None, oldest_first=True): mensagens.append(msg)
    except Exception as erro: await enviar_log(f"⚠️ Erro coletando mensagens do atendimento `{atendimento.get('numero')}`: {erro}")
    pasta = BOLETIM_ARQUIVOS_DIR / str(atendimento.get("numero", "boletim")).replace("/", "-") / "arquivamento"
    anexos = await arquivos_para_reenvio_de_mensagens(mensagens, pasta, str(atendimento.get("numero", "boletim")))
    canal_arq = interaction.guild.get_channel(BOLETINS_ARQUIVADOS_CHANNEL_ID) if interaction.guild else None
    if canal_arq is None:
        try:
            canal_arq = await bot.fetch_channel(BOLETINS_ARQUIVADOS_CHANNEL_ID)
        except Exception:
            canal_arq = None
    if canal_arq is None: return await interaction.followup.send("❌ Canal de arquivamento de boletins não encontrado. Nada foi apagado.", ephemeral=True)
    try:
        texto_msgs = []
        for msg in mensagens:
            conteudo = coletar_texto_embed(msg) if 'coletar_texto_embed' in globals() else (msg.content or "")
            if conteudo.strip(): texto_msgs.append(f"[{formatar_data_discord(msg.created_at)}] {msg.author}: {conteudo[:900]}")
        texto_arquivo = f"📁 BOLETIM ARQUIVADO — Nº {numero_curto_boletim(atendimento.get('numero'))}\n\nNúmero: {atendimento.get('numero')}\nAutor do boletim: {atendimento.get('autor_nome')} (`{atendimento.get('autor_id')}`)\nAgente responsável: {atendimento.get('agente_nome')} (`{atendimento.get('agente_id')}`)\nFinalizado por: {interaction.user} (`{interaction.user.id}`)\nData de recebimento: {atendimento.get('data_criacao')}\nData de finalização: {agora_br()}\nMensagem original: {atendimento.get('mensagem_original_url')}\n\nRESULTADO FINAL:\n{resultado}\n\nHISTÓRICO DE MENSAGENS DO ATENDIMENTO:\n" + "\n\n".join(texto_msgs)
        partes = [texto_arquivo[i:i+1900] for i in range(0, len(texto_arquivo), 1900)] or [texto_arquivo]
        for parte in partes: await canal_arq.send(parte)
        enviados = await enviar_arquivos_em_lotes(canal_arq, anexos, f"📎 Arquivos do boletim {atendimento.get('numero')}")
    except Exception as erro:
        await enviar_log(f"❌ Falha no arquivamento do boletim `{atendimento.get('numero')}`: {erro}")
        return await interaction.followup.send("❌ Falha no arquivamento. A área de atendimento foi mantida.", ephemeral=True)
    atendimento.update({"status": "FINALIZADO", "resultado_final": resultado, "finalizado_por_id": interaction.user.id, "finalizado_por_nome": str(interaction.user), "data_finalizacao": agora_br(), "arquivamento_status": "concluido", "anexos_arquivados": enviados})
    atualizar_atendimento_boletim("id", atendimento.get("id"), atendimento)
    await enviar_log(f"✅ **Boletim finalizado e arquivado**\nBoletim: `{atendimento.get('numero')}`\nFinalizado por: {interaction.user.mention} (`{interaction.user.id}`)\nMensagens coletadas: `{len(mensagens)}`\nAnexos arquivados: `{enviados}`")
    try: await interaction.followup.send("✅ Boletim finalizado, arquivado e área temporária removida.", ephemeral=True)
    except Exception: pass
    try: await interaction.channel.delete(reason="Boletim finalizado e arquivado com sucesso")
    except Exception as erro: await enviar_log(f"⚠️ Boletim arquivado, mas não consegui apagar a área `{getattr(interaction.channel, 'id', 0)}`: {erro}")

@bot.listen("on_message")
async def atendimento_boletim_automatico(message: discord.Message):
    numero_lock = None
    try:
        if message.id in boletins_processando:
            return
        if not eh_boletim_valido_para_atendimento(message):
            return
        if buscar_atendimento_por_mensagem(message.id):
            return
        numero_lock = numero_boletim_de_texto(coletar_texto_embed(message) or message.content or "")
        if buscar_atendimento_por_numero(numero_lock):
            return
        if numero_lock in boletins_processando:
            return
        boletins_processando.add(message.id)
        boletins_processando.add(numero_lock)
        await criar_area_atendimento_boletim(message)
    except Exception as erro:
        await enviar_log(f"❌ Erro no processamento automático de boletim `{getattr(message, 'id', 0)}`: {erro}")
    finally:
        boletins_processando.discard(getattr(message, 'id', 0))
        if numero_lock is not None:
            boletins_processando.discard(numero_lock)

@bot.listen("on_ready")
async def persistencia_boletim_atendimento_views():
    try:
        bot.add_view(BoletimAtendimentoView())
        bot.add_view(AutorizacaoProcuradoBoletimView())
        bot.add_view(FotoProcuradoBoletimView())
        await enviar_log("✅ Views persistentes do atendimento de boletins carregadas.")
    except Exception as erro:
        await enviar_log(f"⚠️ Falha ao carregar views persistentes de boletins: {erro}")


# =====================================================
# MAIN
# =====================================================

async def main():
    carregar_boletins_pendentes_memoria()
    garantir_56_organizacoes()
    garantir_senha_catalogo()
    if not DISCORD_TOKEN or DISCORD_TOKEN == "COLE_O_TOKEN_DO_BOT_AQUI":
        print("ERRO: Coloque o token correto nas variáveis do Railway ou no arquivo .env.")
        return
    await start_web_server()
    await bot.start(DISCORD_TOKEN)

if __name__ == "__main__":
    asyncio.run(main())
